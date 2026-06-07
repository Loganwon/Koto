#!/usr/bin/env python
# -*- coding: utf-8 -*-
# Copyright (C) 2024-2026 Koto AI. All rights reserved.
# SPDX-License-Identifier: LicenseRef-Koto-Proprietary
"""
Word Track Changes 修订模式实现（改进版）
在文档中插入可以 accept/decline 的修改建议

改进点：
1. 保留原有格式（粗体、斜体、颜色等）
2. 支持同一段落内多处修改
3. 更精确的文本定位
4. 详细的成功/失败统计
"""

import logging
import os
import re
import shutil
import zipfile
from datetime import datetime
from typing import Any, Dict, List

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from lxml import etree

logger = logging.getLogger(__name__)


def _normalize_text_for_match(text: str) -> str:
    """Normalize AI-returned text so it matches the plain text stored in Word para.text.

    format_for_ai() decorates paragraphs with Markdown symbols (##, -, >, **...)
    that are NOT present in the Word document's para.text.  When the AI echoes those
    symbols back in the '原文片段' field, exact matching fails silently.

    Normalization steps (safe to apply to both original and modified text):
      1. Unicode NFC  – prevent composed/decomposed CJK mismatch
      2. Strip leading block-level Markdown (headings, lists, blockquotes)
      3. Remove inline bold / italic markers anywhere in the string
      4. Strip the image-caption suffix added by format_for_ai
      5. Normalize all Unicode whitespace variants to a regular space
      6. Collapse runs of whitespace to a single space
    """
    import unicodedata

    t = unicodedata.normalize("NFC", text)
    # Block-level prefixes (only at string start)
    t = re.sub(r"^#{1,6}[ \t]+", "", t)  # ## Heading
    t = re.sub(r"^[-*+][ \t]+", "", t)  # - list item
    t = re.sub(r"^\d+\.[ \t]+", "", t)  # 1. list item
    t = re.sub(r"^>[ \t]+", "", t)  # > blockquote
    # Inline bold / italic (anywhere in string)
    t = re.sub(r"\*\*([^*]+?)\*\*", r"\1", t)  # **bold**
    t = re.sub(r"\*([^*]+?)\*", r"\1", t)  # *italic*
    t = re.sub(r"`([^`]+?)`", r"\1", t)  # `code`
    # Image caption suffix from format_for_ai: "  *(图片，无法直接编辑文本)*"
    t = re.sub(r"\s*\*\([^)]*图片[^)]*\)\*\s*$", "", t)
    # Unicode spaces → regular space
    t = re.sub(r"[\u00a0\u3000\u2009\u2003\u200b]", " ", t)
    # Collapse whitespace runs
    t = re.sub(r"[ \t]{2,}", " ", t)
    return t.strip()


# Keep old name as alias so any external call still works
_strip_md_formatting = _normalize_text_for_match


def _find_para_with_text(
    all_paragraphs,
    original_text: str,
) -> tuple:
    """Return (para, effective_text) for the first paragraph containing original_text.

    Uses a two-level search:
      Level 1 – exact match against para.text (fastest, no transformation)
      Level 2 – normalized original vs raw para.text
               (handles ## / - / ** markers echoed by AI)

    Returns (None, '') if no paragraph found at any level.
    """
    if not original_text:
        return None, ""

    norm = _normalize_text_for_match(original_text)
    candidates = [original_text]
    if norm and norm != original_text:
        candidates.append(norm)

    for candidate in candidates:
        for para in all_paragraphs:
            if candidate in para.text:
                return para, candidate

    return None, ""


class TrackChangesEditor:
    """Word Track Changes 修订编辑器（改进版）"""

    def __init__(self, author: str = "Koto AI"):
        self.author = author
        self.change_id = 0

    def apply_comment_changes(
        self, file_path: str, annotations: List[Dict[str, str]], progress_callback=None
    ) -> Dict[str, Any]:
        """
        以右侧批注气泡的方式标注修改建议

        原文保持不变，修改建议以 Word 批注(Comment)显示在右侧边栏。
        用户在 Word 中可在「审阅」里逐条查看、接受或忽略。

        Args:
            file_path: Word 文档路径
            annotations: 标注列表，同 apply_tracked_changes
            progress_callback: 可选的进度回调函数 callback(current, total, status, detail)

        Returns:
            修改统计
        """
        try:
            doc = Document(file_path)

            applied_count = 0
            failed_count = 0

            logger.info(f"[Comments] 💬 开始添加批注...")
            logger.info(f"[Comments] 📊 共 {len(annotations)} 条批注")

            # 预处理标注
            normalized = []
            for anno in annotations:
                original = anno.get("原文片段", anno.get("原文", "")).strip()
                comment_text = anno.get(
                    "批注内容", anno.get("批注", anno.get("comment", ""))
                ).strip()
                modified = anno.get(
                    "修改后文本", anno.get("修改建议", anno.get("改为", ""))
                ).strip()
                reason = anno.get("修改原因", anno.get("原因", "")).strip()

                if original and comment_text:
                    normalized.append(
                        {
                            "original": original,
                            "modified": comment_text,
                            "reason": reason,
                            "label": str(anno.get("批注标签") or "差异说明：").strip()
                            or "差异说明：",
                        }
                    )
                elif original and modified and original != modified:
                    normalized.append(
                        {
                            "original": original,
                            "modified": modified,
                            "reason": reason,
                            "label": "建议改为：",
                        }
                    )

            logger.info(f"[Comments] ✅ 有效批注: {len(normalized)} 条")

            if not normalized:
                doc.save(file_path)
                return {"success": True, "applied": 0, "failed": 0, "total": 0}

            # 获取或创建 comments part
            comments_el, comments_part_ref = self._get_or_create_comments_part(doc)
            self.change_id = max(self.change_id, self._max_comment_id(comments_el))
            initial_comment_count = self._comment_count(comments_el)

            # 通知开始应用
            if progress_callback:
                progress_callback(
                    0, len(normalized), "start", f"开始添加 {len(normalized)} 条批注"
                )

            for idx, anno in enumerate(normalized, 1):
                original = anno["original"]
                modified = anno["modified"]
                reason = anno["reason"]
                label = anno.get("label") or "建议改为："

                # 通知当前进度
                if progress_callback:
                    progress_callback(
                        idx,
                        len(normalized),
                        "processing",
                        f"正在处理: {original[:30]}...",
                    )

                found = False

                # 先在正文段落中查找
                for para in doc.paragraphs:
                    if original in para.text:
                        self.change_id += 1
                        cid = self.change_id

                        # 先在正文里插入批注范围；成功后再追加 comments.xml，
                        # 避免出现没有锚点的孤儿批注。
                        success = self._add_comment_markers_to_paragraph(
                            para, original, cid
                        )
                        if success:
                            self._add_comment_element(
                                comments_el, cid, modified, reason, label=label
                            )
                            applied_count += 1
                            found = True
                            detail_msg = (
                                f"✅ #{idx}/{len(normalized)}: '{original[:25]}...'"
                            )
                            logger.info(f"  💬 {detail_msg}")
                            if progress_callback:
                                progress_callback(
                                    idx, len(normalized), "success", detail_msg
                                )
                            break

                # 再在表格中查找
                if not found:
                    for table in doc.tables:
                        if found:
                            break
                        for row in table.rows:
                            if found:
                                break
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if original in para.text:
                                        self.change_id += 1
                                        cid = self.change_id
                                        success = (
                                            self._add_comment_markers_to_paragraph(
                                                para, original, cid
                                            )
                                        )
                                        if success:
                                            self._add_comment_element(
                                                comments_el,
                                                cid,
                                                modified,
                                                reason,
                                                label=label,
                                            )
                                            applied_count += 1
                                            found = True
                                            detail_msg = f"✅ (表格) #{idx}/{len(normalized)}: '{original[:20]}...'"
                                            logger.info(f"  💬 {detail_msg}")
                                            if progress_callback:
                                                progress_callback(
                                                    idx,
                                                    len(normalized),
                                                    "success",
                                                    detail_msg,
                                                )
                                            break

                if not found:
                    failed_count += 1
                    detail_msg = f"⚠️ #{idx} 未找到: '{original[:30]}...'"
                    logger.info(f"  {detail_msg}")
                    if progress_callback:
                        progress_callback(idx, len(normalized), "failed", detail_msg)

            # 将 comments XML 写回 part
            comments_bytes = etree.tostring(
                comments_el, xml_declaration=True, encoding="UTF-8", standalone=True
            )

            if comments_part_ref is not None:
                # 已有 comments part，更新内容
                comments_part_ref._blob = comments_bytes
            else:
                # 新建 comments part
                from docx.opc.packuri import PackURI
                from docx.opc.part import Part

                COMMENTS_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
                COMMENTS_RT = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"

                new_part = Part(
                    PackURI("/word/comments.xml"),
                    COMMENTS_CT,
                    comments_bytes,
                    doc.part.package,
                )
                doc.part.relate_to(new_part, COMMENTS_RT)

            # 保存文档
            doc.save(file_path)
            # python-docx 对已有 comments part 的持久化并不总是可靠；
            # 直接注入 zip 包，确保 Word 与前端预览都能读取到新批注正文。
            if applied_count > 0:
                self._inject_comments_to_docx(file_path, comments_el)

            persisted_comment_count = self._docx_comment_count(file_path)
            expected_comment_count = initial_comment_count + applied_count
            if persisted_comment_count < expected_comment_count:
                return {
                    "success": False,
                    "error": (
                        "DOCX 批注正文未可靠写入："
                        f"预期至少 {expected_comment_count} 条，实际 {persisted_comment_count} 条。"
                    ),
                    "applied": 0,
                    "failed": len(normalized),
                    "total": len(normalized),
                }

            success_rate = (applied_count / len(normalized) * 100) if normalized else 0
            logger.info(f"\n[Comments] 💾 文档已保存")
            logger.info(
                f"[Comments] 📊 成功: {applied_count}, 失败: {failed_count}, 成功率: {success_rate:.1f}%"
            )

            return {
                "success": True,
                "applied": applied_count,
                "failed": failed_count,
                "total": len(normalized),
            }

        except Exception as e:
            logger.error(f"[Comments] ❌ 错误: {str(e)}")
            import traceback

            traceback.print_exc()
            return {"success": False, "error": str(e)}

    def _get_or_create_comments_part(self, doc):
        """获取或创建文档的 comments 部分"""
        for rel in doc.part.rels.values():
            if "comments" in rel.reltype:
                part = rel.target_part
                el = etree.fromstring(part.blob)
                return el, part

        # 新建空的 comments 元素
        el = etree.fromstring(
            b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            b' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
        )
        return el, None

    def _comment_count(self, comments_el) -> int:
        try:
            return len(
                comments_el.xpath(
                    ".//w:comment",
                    namespaces={
                        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                    },
                )
            )
        except Exception:
            return 0

    def _max_comment_id(self, comments_el) -> int:
        max_id = 0
        try:
            for comment in comments_el.xpath(
                ".//w:comment",
                namespaces={
                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                },
            ):
                raw = comment.get(qn("w:id")) or comment.get("w:id") or ""
                try:
                    max_id = max(max_id, int(str(raw)))
                except Exception:
                    continue
        except Exception:
            return 0
        return max_id

    def _docx_comment_count(self, file_path: str) -> int:
        try:
            import zipfile

            with zipfile.ZipFile(file_path, "r") as archive:
                if "word/comments.xml" not in archive.namelist():
                    return 0
                root = etree.fromstring(archive.read("word/comments.xml"))
            return self._comment_count(root)
        except Exception:
            return 0

    def _sync_comments_part(self, doc, comments_el, comments_part_ref):
        """将 comments XML 同步到文档 part，并返回当前 part 引用。"""
        comments_bytes = etree.tostring(
            comments_el, xml_declaration=True, encoding="UTF-8", standalone=True
        )

        if comments_part_ref is not None:
            comments_part_ref._blob = comments_bytes
            return comments_part_ref

        from docx.opc.packuri import PackURI
        from docx.opc.part import Part

        comments_content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
        comments_rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"

        new_part = Part(
            PackURI("/word/comments.xml"),
            comments_content_type,
            comments_bytes,
            doc.part.package,
        )
        doc.part.relate_to(new_part, comments_rel_type)
        return new_part

    def _add_comment_element(
        self, comments_el, comment_id, modified, reason="", label="建议改为："
    ):
        """在 comments XML 里添加一条批注"""
        WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        comment = etree.SubElement(comments_el, qn("w:comment"))
        comment.set(qn("w:id"), str(comment_id))
        comment.set(qn("w:author"), self.author)
        comment.set(qn("w:date"), datetime.now().isoformat() + "Z")
        comment.set(qn("w:initials"), "K")

        label_text = str(label or "建议改为：").strip() or "建议改为："

        # 第1段：批注主内容
        p1 = etree.SubElement(comment, qn("w:p"))
        r1 = etree.SubElement(p1, qn("w:r"))
        # 加粗标签
        rpr1 = etree.SubElement(r1, qn("w:rPr"))
        etree.SubElement(rpr1, qn("w:b"))
        t1 = etree.SubElement(r1, qn("w:t"))
        t1.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t1.text = label_text

        # 批注内容（不加粗）
        r1b = etree.SubElement(p1, qn("w:r"))
        t1b = etree.SubElement(r1b, qn("w:t"))
        t1b.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t1b.text = modified

        # 第2段：原因
        if reason:
            p2 = etree.SubElement(comment, qn("w:p"))
            r2 = etree.SubElement(p2, qn("w:r"))
            rpr2 = etree.SubElement(r2, qn("w:rPr"))
            # 灰色小字
            color = etree.SubElement(rpr2, qn("w:color"))
            color.set(qn("w:val"), "888888")
            sz = etree.SubElement(rpr2, qn("w:sz"))
            sz.set(qn("w:val"), "18")  # 9pt
            t2 = etree.SubElement(r2, qn("w:t"))
            t2.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t2.text = f"原因：{reason}"

    def _add_comment_markers_to_paragraph(self, para, original, comment_id):
        """
        在段落中为原文片段添加批注标记

        添加 commentRangeStart / commentRangeEnd / commentReference
        """
        try:
            pos = para.text.find(original)
            if pos == -1:
                return False
            return self._insert_comment_range_markers(
                para,
                pos,
                para,
                pos + len(original),
                comment_id,
            )

        except Exception as e:
            logger.warning(f"[Comments] ⚠️ 添加批注标记失败: {str(e)}")
            return False

    def apply_tracked_changes(
        self, file_path: str, annotations: List[Dict[str, str]], progress_callback=None
    ) -> Dict[str, Any]:
        """
        应用 Track Changes 修订到文档

        Args:
            file_path: Word 文档路径
            annotations: 标注列表，每个包含：
                - 原文片段: 要修改的文本
                - 修改后文本/修改建议/改为: 修改后的文本
            progress_callback: 可选的进度回调函数

        Returns:
            修改统计
        """
        try:
            doc = Document(file_path)

            applied_count = 0
            failed_count = 0

            logger.info(f"[TrackChanges] 📝 开始应用修订...")
            logger.info(f"[TrackChanges] 📊 共 {len(annotations)} 条修改建议")

            # 预处理标注：标准化字段名
            normalized = []
            for anno in annotations:
                original = anno.get("原文片段", anno.get("原文", "")).strip()
                # 支持多种字段名
                modified = anno.get(
                    "修改后文本", anno.get("修改建议", anno.get("改为", ""))
                ).strip()

                if original and modified and original != modified:
                    normalized.append({"original": original, "modified": modified})

            logger.info(f"[TrackChanges] ✅ 有效修改: {len(normalized)} 条")

            if progress_callback:
                progress_callback(
                    0, len(normalized), "start", f"开始应用 {len(normalized)} 条修订"
                )

            def _emit_saved_progress(current_index: int, original_text: str) -> None:
                if not progress_callback:
                    return
                try:
                    doc.save(file_path)
                except Exception as save_error:
                    logger.warning(f"[TrackChanges] ⚠️ 增量保存失败: {save_error}")
                    progress_callback(
                        current_index,
                        len(normalized),
                        "failed",
                        f"⚠️ 增量写回失败: {str(save_error)[:80]}",
                    )
                    return
                progress_callback(
                    current_index,
                    len(normalized),
                    "saved",
                    f"已写回原文 {applied_count}/{len(normalized)}：{original_text[:20]}...",
                    file_updated=True,
                    applied=applied_count,
                    file_path=file_path,
                    updated_in_place=True,
                )

            for idx, anno in enumerate(normalized, 1):
                original = anno["original"]
                modified = anno["modified"]

                if progress_callback:
                    progress_callback(
                        idx,
                        len(normalized),
                        "processing",
                        f"正在处理: {original[:30]}...",
                    )

                # 查找文本位置并应用修订
                found = False

                # 先在正文段落中查找
                for para in doc.paragraphs:
                    if original in para.text:
                        success = self._apply_change_to_paragraph(
                            para, original, modified
                        )
                        if success:
                            applied_count += 1
                            found = True
                            detail_msg = f"✅ #{idx}/{len(normalized)}: '{original[:25]}...' → '{modified[:25]}...'"
                            logger.info(f"  {detail_msg}")
                            if progress_callback:
                                progress_callback(
                                    idx, len(normalized), "success", detail_msg
                                )
                            _emit_saved_progress(idx, original)
                            break

                # 再在表格中查找
                if not found:
                    for table in doc.tables:
                        if found:
                            break
                        for row in table.rows:
                            if found:
                                break
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if original in para.text:
                                        success = self._apply_change_to_paragraph(
                                            para, original, modified
                                        )
                                        if success:
                                            applied_count += 1
                                            found = True
                                            detail_msg = f"✅ (表格) #{idx}/{len(normalized)}: '{original[:20]}...'"
                                            logger.info(f"  {detail_msg}")
                                            if progress_callback:
                                                progress_callback(
                                                    idx,
                                                    len(normalized),
                                                    "success",
                                                    detail_msg,
                                                )
                                            _emit_saved_progress(idx, original)
                                            break

                if not found:
                    failed_count += 1
                    detail_msg = f"⚠️ #{idx} 未找到: '{original[:30]}...'"
                    logger.info(f"  {detail_msg}")
                    if progress_callback:
                        progress_callback(idx, len(normalized), "failed", detail_msg)

            # 保存文档
            doc.save(file_path)

            success_rate = (applied_count / len(normalized) * 100) if normalized else 0
            logger.info(f"\n[TrackChanges] 💾 文档已保存")
            logger.info(
                f"[TrackChanges] 📊 成功: {applied_count}, 失败: {failed_count}, 成功率: {success_rate:.1f}%"
            )

            return {
                "success": True,
                "applied": applied_count,
                "failed": failed_count,
                "total": len(normalized),
                "updated_in_place": True,
            }

        except Exception as e:
            logger.error(f"[TrackChanges] ❌ 错误: {str(e)}")
            import traceback

            traceback.print_exc()
            return {"success": False, "error": str(e)}

    def _apply_change_to_paragraph(self, para, original: str, modified: str) -> bool:
        """
        在段落中应用 Track Changes 修订

        策略：尽可能保留格式，精确定位文本
        """
        try:
            p = para._element
            runs = list(p.findall(qn("w:r")))

            if not runs:
                return False

            # 构建完整文本和 run 映射
            text_parts = []
            run_map = []  # [(start_pos, end_pos, run_element)]

            for run in runs:
                run_text = self._get_run_text(run)
                start = len("".join(text_parts))
                text_parts.append(run_text)
                end = len("".join(text_parts))
                run_map.append((start, end, run))

            full_text = "".join(text_parts)

            # 查找目标文本
            pos = full_text.find(original)
            if pos == -1:
                return False

            target_end = pos + len(original)

            # 找到涉及的 run
            start_run = None
            end_run = None
            start_offset = 0
            end_offset = 0

            for i, (s, e, run) in enumerate(run_map):
                if start_run is None and s <= pos < e:
                    start_run = i
                    start_offset = pos - s
                if s < target_end <= e:
                    end_run = i
                    end_offset = target_end - s
                    break

            if start_run is None:
                return False
            if end_run is None:
                end_run = len(run_map) - 1
                end_offset = len(self._get_run_text(run_map[end_run][2]))

            # 生成修订 ID
            self.change_id += 1
            del_id = str(self.change_id)
            self.change_id += 1
            ins_id = str(self.change_id)
            date_str = datetime.now().isoformat()

            # 单 run 内的修改（最常见情况）
            if start_run == end_run:
                run = run_map[start_run][2]
                run_text = self._get_run_text(run)

                before = run_text[:start_offset]
                target = run_text[start_offset:end_offset]
                after = run_text[end_offset:]

                # 获取格式
                rPr = run.find(qn("w:rPr"))
                rPr_xml = self._clone_rPr(rPr)

                # 构建新元素
                new_elements = []

                if before:
                    new_elements.append(self._make_run(before, rPr_xml))

                # 删除标记
                new_elements.append(
                    parse_xml(
                        f"""<w:del w:id="{del_id}" w:author="{self._esc(self.author)}" w:date="{date_str}"
                        xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:r>{rPr_xml}<w:delText xml:space="preserve">{self._esc(target)}</w:delText></w:r>
                    </w:del>"""
                    )
                )

                # 插入标记
                new_elements.append(
                    parse_xml(
                        f"""<w:ins w:id="{ins_id}" w:author="{self._esc(self.author)}" w:date="{date_str}"
                        xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:r>{rPr_xml}<w:t xml:space="preserve">{self._esc(modified)}</w:t></w:r>
                    </w:ins>"""
                    )
                )

                if after:
                    new_elements.append(self._make_run(after, rPr_xml))

                # 替换原 run
                idx = list(p).index(run)
                p.remove(run)
                for i, elem in enumerate(new_elements):
                    p.insert(idx + i, elem)

                return True

            else:
                # 跨多个 run：简化处理
                para_text = para.text
                parts = para_text.split(original, 1)

                # 清空段落
                for run in list(p.findall(qn("w:r"))):
                    p.remove(run)

                if parts[0]:
                    p.append(self._make_run(parts[0], ""))

                p.append(
                    parse_xml(
                        f"""<w:del w:id="{del_id}" w:author="{self._esc(self.author)}" w:date="{date_str}"
                        xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:r><w:delText xml:space="preserve">{self._esc(original)}</w:delText></w:r>
                    </w:del>"""
                    )
                )

                p.append(
                    parse_xml(
                        f"""<w:ins w:id="{ins_id}" w:author="{self._esc(self.author)}" w:date="{date_str}"
                        xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:r><w:t xml:space="preserve">{self._esc(modified)}</w:t></w:r>
                    </w:ins>"""
                    )
                )

                if len(parts) > 1 and parts[1]:
                    p.append(self._make_run(parts[1], ""))

                return True

        except Exception as e:
            logger.warning(f"[TrackChanges] ⚠️ 段落修订失败: {str(e)}")
            return False

    def _get_run_text(self, run) -> str:
        """获取 run 中的文本"""
        parts = []
        for t in run.findall(qn("w:t")):
            if t.text:
                parts.append(t.text)
        return "".join(parts)

    def _iter_commentable_paragraphs(self, doc: Document) -> list[Any]:
        paragraphs: list[Any] = []
        try:
            body = doc._element.body
        except Exception:
            body = None

        body_paragraphs = list(doc.paragraphs)
        body_tables = list(doc.tables)
        para_index = 0
        table_index = 0

        if body is not None:
            for child in body.iterchildren():
                if child.tag == qn("w:p"):
                    if para_index < len(body_paragraphs):
                        paragraphs.append(body_paragraphs[para_index])
                        para_index += 1
                elif child.tag == qn("w:tbl"):
                    if table_index >= len(body_tables):
                        continue
                    table = body_tables[table_index]
                    table_index += 1
                    for row in table.rows:
                        for cell in row.cells:
                            paragraphs.extend(cell.paragraphs)
            if paragraphs:
                return paragraphs

        paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        return paragraphs

    def _build_comment_anchor_index(self, doc: Document) -> tuple[str, list[dict[str, Any]]]:
        paragraphs = self._iter_commentable_paragraphs(doc)
        spans: list[dict[str, Any]] = []
        parts: list[str] = []
        cursor = 0

        for index, para in enumerate(paragraphs):
            para_text = para.text or ""
            start_offset = cursor
            parts.append(para_text)
            cursor += len(para_text)
            spans.append({
                "para": para,
                "start": start_offset,
                "end": cursor,
            })
            if index + 1 < len(paragraphs):
                parts.append("\n")
                cursor += 1

        return "".join(parts), spans

    @staticmethod
    def _coerce_anchor_int(value: Any) -> int | None:
        if value in (None, ""):
            return None
        try:
            return int(value)
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _matching_prefix_len(expected: str, actual: str) -> int:
        limit = min(len(expected), len(actual))
        matched = 0
        while matched < limit and expected[matched] == actual[matched]:
            matched += 1
        return matched

    @staticmethod
    def _matching_suffix_len(expected: str, actual: str) -> int:
        limit = min(len(expected), len(actual))
        matched = 0
        while matched < limit and expected[-(matched + 1)] == actual[-(matched + 1)]:
            matched += 1
        return matched

    def _score_comment_anchor_candidate(
        self,
        full_text: str,
        start_offset: int,
        end_offset: int,
        occurrence_index: int,
        anchor_meta: dict[str, Any],
    ) -> int:
        before_context = str(anchor_meta.get("anchor_context_before") or "")
        after_context = str(anchor_meta.get("anchor_context_after") or "")
        score = 0

        if before_context:
            actual_before = full_text[max(0, start_offset - len(before_context)):start_offset]
            score += self._matching_suffix_len(before_context, actual_before) * 4
        if after_context:
            actual_after = full_text[end_offset:end_offset + len(after_context)]
            score += self._matching_prefix_len(after_context, actual_after) * 4

        occurrence_hint = self._coerce_anchor_int(anchor_meta.get("anchor_occurrence"))
        if occurrence_hint is not None:
            score += max(0, 96 - min(96, abs(occurrence_index - occurrence_hint) * 24))

        start_hint = self._coerce_anchor_int(anchor_meta.get("anchor_start_offset"))
        if start_hint is not None:
            score += max(0, 96 - min(96, abs(start_offset - start_hint)))

        end_hint = self._coerce_anchor_int(anchor_meta.get("anchor_end_offset"))
        if end_hint is not None:
            score += max(0, 48 - min(48, abs(end_offset - end_hint)))

        return score

    def _offsets_to_paragraph_range(
        self,
        spans: list[dict[str, Any]],
        start_offset: int,
        end_offset: int,
    ) -> dict[str, Any] | None:
        if not spans or end_offset <= start_offset:
            return None

        def _locate(offset: int, *, is_end: bool) -> tuple[int, int] | None:
            for index, span in enumerate(spans):
                span_start = int(span["start"])
                span_end = int(span["end"])
                para_length = max(0, span_end - span_start)
                if span_start <= offset <= span_end:
                    return index, max(0, min(offset - span_start, para_length))
                if index + 1 < len(spans):
                    next_start = int(spans[index + 1]["start"])
                    if span_end < offset < next_start:
                        return (index, para_length) if is_end else (index + 1, 0)
            last_index = len(spans) - 1
            last_span = spans[last_index]
            last_length = max(0, int(last_span["end"]) - int(last_span["start"]))
            if offset >= int(last_span["end"]):
                return last_index, last_length
            return None

        start_loc = _locate(start_offset, is_end=False)
        end_loc = _locate(end_offset, is_end=True)
        if start_loc is None or end_loc is None:
            return None

        start_index, start_in_para = start_loc
        end_index, end_in_para = end_loc
        if start_index > end_index:
            return None
        if start_index == end_index and end_in_para <= start_in_para:
            return None

        return {
            "start_para": spans[start_index]["para"],
            "start_offset": start_in_para,
            "end_para": spans[end_index]["para"],
            "end_offset": end_in_para,
        }

    def _resolve_comment_anchor_range(
        self,
        doc: Document,
        original_text: str,
        anchor_meta: dict[str, Any] | None = None,
    ) -> dict[str, Any] | None:
        full_text, spans = self._build_comment_anchor_index(doc)
        if not spans:
            return None

        anchor_meta = anchor_meta if isinstance(anchor_meta, dict) else {}
        search_texts: list[str] = []
        stripped_text = _strip_md_formatting(original_text)
        for candidate in (original_text, stripped_text):
            if candidate and candidate not in search_texts:
                search_texts.append(candidate)
        if not search_texts:
            return None

        start_hint = self._coerce_anchor_int(anchor_meta.get("anchor_start_offset"))
        end_hint = self._coerce_anchor_int(anchor_meta.get("anchor_end_offset"))
        before_context = str(anchor_meta.get("anchor_context_before") or "")
        after_context = str(anchor_meta.get("anchor_context_after") or "")

        if start_hint is not None and end_hint is not None and end_hint > start_hint:
            range_info = self._offsets_to_paragraph_range(spans, start_hint, end_hint)
            if range_info is not None:
                candidate_slice = full_text[start_hint:end_hint]
                normalized_slice = _normalize_text_for_match(candidate_slice)
                text_matches = any(
                    candidate_slice == candidate
                    or normalized_slice == _normalize_text_for_match(candidate)
                    for candidate in search_texts
                )
                before_matches = (
                    not before_context
                    or full_text[max(0, start_hint - len(before_context)):start_hint] == before_context
                )
                after_matches = (
                    not after_context
                    or full_text[end_hint:end_hint + len(after_context)] == after_context
                )
                if text_matches or (before_context or after_context) and before_matches and after_matches:
                    return range_info

        best_match: dict[str, Any] | None = None
        for candidate in search_texts:
            occurrence_index = 0
            search_from = 0
            while True:
                hit = full_text.find(candidate, search_from)
                if hit == -1:
                    break
                match_end = hit + len(candidate)
                range_info = self._offsets_to_paragraph_range(spans, hit, match_end)
                if range_info is not None:
                    score = self._score_comment_anchor_candidate(
                        full_text,
                        hit,
                        match_end,
                        occurrence_index,
                        anchor_meta,
                    )
                    if best_match is None or score > int(best_match["score"]):
                        best_match = {
                            **range_info,
                            "score": score,
                        }
                occurrence_index += 1
                search_from = hit + max(len(candidate), 1)

        if best_match is not None:
            best_match.pop("score", None)
            return best_match

        all_paragraphs = [span["para"] for span in spans]
        para, matched_text = _find_para_with_text(all_paragraphs, original_text)
        if para is None or not matched_text:
            return None
        pos = para.text.index(matched_text)
        return {
            "start_para": para,
            "start_offset": pos,
            "end_para": para,
            "end_offset": pos + len(matched_text),
        }

    def _split_run_at_paragraph_offset(self, para, split_offset: int) -> None:
        if split_offset <= 0 or split_offset >= len(para.text or ""):
            return

        p = para._element
        accumulated = 0
        for run_el in list(p.findall(qn("w:r"))):
            run_text = self._get_run_text(run_el)
            run_length = len(run_text)
            if run_length <= 0:
                continue

            run_start = accumulated
            run_end = run_start + run_length
            if not (run_start < split_offset < run_end):
                accumulated = run_end
                continue

            split_index = split_offset - run_start
            left_text = run_text[:split_index]
            right_text = run_text[split_index:]
            rpr_xml = self._clone_rPr(run_el.find(qn("w:rPr")))
            insert_at = list(p).index(run_el)
            p.remove(run_el)
            if left_text:
                p.insert(insert_at, self._make_run(left_text, rpr_xml))
                insert_at += 1
            if right_text:
                p.insert(insert_at, self._make_run(right_text, rpr_xml))
            return

    def _find_run_at_or_after_offset(self, para, offset: int):
        accumulated = 0
        first_text_run = None
        for run_el in para._element.findall(qn("w:r")):
            run_text = self._get_run_text(run_el)
            run_length = len(run_text)
            if run_length <= 0:
                continue
            if first_text_run is None:
                first_text_run = run_el
            run_start = accumulated
            run_end = run_start + run_length
            if offset <= run_start or run_start <= offset < run_end:
                return run_el
            accumulated = run_end
        return first_text_run if offset <= 0 else None

    def _find_run_at_or_before_offset(self, para, offset: int):
        accumulated = 0
        last_text_run = None
        for run_el in para._element.findall(qn("w:r")):
            run_text = self._get_run_text(run_el)
            run_length = len(run_text)
            if run_length <= 0:
                continue
            run_start = accumulated
            run_end = run_start + run_length
            if run_start < offset <= run_end or run_end == offset:
                return run_el
            accumulated = run_end
            last_text_run = run_el
        return last_text_run

    def _insert_comment_range_markers(
        self,
        start_para,
        start_offset: int,
        end_para,
        end_offset: int,
        comment_id: str | int,
    ) -> bool:
        if start_para is None or end_para is None:
            return False
        if start_para is end_para and end_offset <= start_offset:
            return False

        self._split_run_at_paragraph_offset(end_para, end_offset)
        self._split_run_at_paragraph_offset(start_para, start_offset)

        start_run_el = self._find_run_at_or_after_offset(start_para, start_offset)
        end_run_el = self._find_run_at_or_before_offset(end_para, end_offset)
        if start_run_el is None or end_run_el is None:
            return False

        cid = str(comment_id)
        wns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        start_marker = parse_xml(
            f'<w:commentRangeStart w:id="{cid}" xmlns:w="{wns}"/>'
        )
        end_marker = parse_xml(
            f'<w:commentRangeEnd w:id="{cid}" xmlns:w="{wns}"/>'
        )
        ref_run = parse_xml(
            f'<w:r xmlns:w="{wns}">'
            f'  <w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
            f'  <w:commentReference w:id="{cid}"/>'
            f"</w:r>"
        )

        start_run_el.addprevious(start_marker)
        end_run_el.addnext(end_marker)
        end_marker.addnext(ref_run)
        return True

    def _clone_rPr(self, rPr) -> str:
        """克隆格式属性"""
        if rPr is None:
            return ""
        from lxml import etree

        return etree.tostring(rPr, encoding="unicode")

    def _make_run(self, text: str, rPr_xml: str):
        """创建 run 元素"""
        return parse_xml(
            f"""<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                {rPr_xml}<w:t xml:space="preserve">{self._esc(text)}</w:t>
            </w:r>"""
        )

    @staticmethod
    def _esc(text: str) -> str:
        """转义 XML"""
        if not text:
            return ""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&apos;")
        )

    def apply_hybrid_changes(
        self, file_path: str, annotations: List[Dict[str, str]], progress_callback=None
    ) -> Dict[str, Any]:
        """
        混合应用两种标注方式：
        1. 精确的短文本修改 → Track Changes（修订标记）
        2. 大段落方向建议 → Comments（批注气泡）

        自动判断规则：
        - 原文 <= 30字且有精确替换文本 → Track Changes
        - 原文 > 30字或只有方向建议 → Comment

        Args:
            file_path: Word文档路径
            annotations: 标注列表
            progress_callback: 进度回调

        Returns:
            {"success": True, "tracked": 15, "commented": 8, "failed": 2}
        """
        try:
            doc = Document(file_path)

            # 分类标注
            track_changes_items = []  # 精确修改
            comment_items = []  # 方向建议

            for anno in annotations:
                original = anno.get("原文片段", anno.get("原文", "")).strip()
                modified = anno.get(
                    "修改后文本", anno.get("修改建议", anno.get("改为", ""))
                ).strip()
                reason = anno.get("修改原因", anno.get("原因", "")).strip()

                if not original:
                    continue

                # 判断标注类型
                # 策略调整：只要提供了具体修改文本且未显式标记为"建议"，均视为修订(Track Changes)
                # 放宽长度限制，允许整句重写
                is_suggestion = (
                    modified.startswith("建议")
                    or modified.startswith("批注")
                    or "建议：" in modified
                    or "原因：" in modified  # 某些情况下AI可能会把原因混入
                )

                is_precise = (
                    modified  # 有替换文本
                    and modified != original  # 不是重复
                    and not is_suggestion  # 不是建议
                    and len(original) < 500  # 长度安全限制，防止整页替换
                )

                if is_precise:
                    # 精确修改 → Track Changes
                    track_changes_items.append(
                        {"original": original, "modified": modified, "reason": reason}
                    )
                else:
                    # 方向建议 → Comment
                    comment_items.append(
                        {"original": original, "modified": modified, "reason": reason}
                    )

            total_items = len(track_changes_items) + len(comment_items)

            logger.info(f"\n[Hybrid] 🎯 混合标注模式")
            logger.info(
                f"[Hybrid] ✏️  精确修改: {len(track_changes_items)} 条（Track Changes）"
            )
            logger.info(f"[Hybrid] 💬 方向建议: {len(comment_items)} 条（Comments）")

            # 先应用 Track Changes
            tracked_success = 0
            tracked_failed = 0

            if track_changes_items:
                logger.info(f"\n[Hybrid] 📝 第1步：应用精确修改...")

                for idx, item in enumerate(track_changes_items, 1):
                    if progress_callback:
                        progress_callback(
                            idx,
                            len(track_changes_items) + len(comment_items),
                            "tracking",
                            f"修订标记: {item['original'][:20]}...",
                        )

                    success = self._apply_single_track_change(
                        doc, item["original"], item["modified"], item["reason"]
                    )

                    if success:
                        tracked_success += 1
                        doc.save(file_path)
                        if progress_callback:
                            progress_callback(
                                idx,
                                total_items,
                                "saved",
                                f"已写回原文 {tracked_success + commented_success}/{total_items}：{item['original'][:20]}...",
                                file_updated=True,
                                applied=tracked_success + commented_success,
                                file_path=file_path,
                                updated_in_place=True,
                            )
                    else:
                        tracked_failed += 1
                        logger.warning(
                            f"[Hybrid] ⚠️  修订失败: {item['original'][:30]}..."
                        )

            # 再应用 Comments
            commented_success = 0
            commented_failed = 0
            comments_el = None
            comments_part_ref = None

            if comment_items:
                logger.info(f"\n[Hybrid] 💬 第2步：添加批注建议...")

                # 获取 comments part
                comments_el, comments_part_ref = self._get_or_create_comments_part(doc)

                for idx, item in enumerate(comment_items, 1):
                    if progress_callback:
                        progress_callback(
                            len(track_changes_items) + idx,
                            len(track_changes_items) + len(comment_items),
                            "commenting",
                            f"批注建议: {item['original'][:20]}...",
                        )

                    success = self._apply_single_comment(
                        doc,
                        comments_el,
                        item["original"],
                        item["modified"],
                        item["reason"],
                    )

                    if success:
                        commented_success += 1
                        comments_part_ref = self._sync_comments_part(
                            doc, comments_el, comments_part_ref
                        )
                        doc.save(file_path)
                        if progress_callback:
                            progress_callback(
                                len(track_changes_items) + idx,
                                total_items,
                                "saved",
                                f"已写回原文 {tracked_success + commented_success}/{total_items}：{item['original'][:20]}...",
                                file_updated=True,
                                applied=tracked_success + commented_success,
                                file_path=file_path,
                                updated_in_place=True,
                            )
                    else:
                        commented_failed += 1
                        logger.warning(
                            f"[Hybrid] ⚠️  批注失败: {item['original'][:30]}..."
                        )

            # 将 comments XML 写入文档 Part（python-docx OPC 方式）
            if commented_success > 0 and comments_el is not None:
                comments_part_ref = self._sync_comments_part(
                    doc, comments_el, comments_part_ref
                )

            # 保存文档
            doc.save(file_path)

            total_success = tracked_success + commented_success
            total_failed = tracked_failed + commented_failed

            logger.info(f"\n[Hybrid] ✅ 完成！")
            logger.info(
                f"[Hybrid] 📊 修订标记: {tracked_success}成功 / {tracked_failed}失败"
            )
            logger.info(
                f"[Hybrid] 📊 批注建议: {commented_success}成功 / {commented_failed}失败"
            )
            logger.info(f"[Hybrid] 📊 总计: {total_success}成功 / {total_failed}失败\n")

            return {
                "success": True,
                "tracked": tracked_success,
                "commented": commented_success,
                "failed": total_failed,
                "total": len(annotations),
                "applied": total_success,
            }

        except Exception as e:
            logger.error(f"[Hybrid] ❌ 混合标注失败: {e}")
            import traceback

            traceback.print_exc()
            return {
                "success": False,
                "error": str(e),
                "tracked": 0,
                "commented": 0,
                "failed": len(annotations),
                "total": len(annotations),
                "applied": 0,
            }

    def _apply_single_track_change(
        self, doc: Document, original_text: str, modified_text: str, reason: str = ""
    ) -> bool:
        """Apply a single Track-Changes revision (del + ins) preserving run formatting."""
        try:
            from copy import deepcopy

            # Build paragraph list: body paragraphs + all table cell paragraphs
            all_paragraphs = list(doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_paragraphs.extend(cell.paragraphs)

            # Multi-level normalized search
            found_para, effective_original = _find_para_with_text(
                all_paragraphs, original_text
            )
            if not found_para:
                return False

            para = found_para
            # Strip markdown from modified_text so Track Change text looks clean in Word
            effective_modified = (
                _normalize_text_for_match(modified_text) or modified_text
            )

            # Map runs to character positions
            run_map = []
            current_pos = 0
            full_text_parts = []
            for i, run in enumerate(para.runs):
                text = run.text
                run_map.append(
                    {
                        "start": current_pos,
                        "end": current_pos + len(text),
                        "run": run,
                        "index": i,
                    }
                )
                full_text_parts.append(text)
                current_pos += len(text)

            full_text = "".join(full_text_parts)
            start_idx = full_text.find(effective_original)
            if start_idx == -1:
                return False
            end_idx = start_idx + len(effective_original)

            # Find runs overlapping [start_idx, end_idx]
            target_runs = []
            for info in run_map:
                if max(start_idx, info["start"]) < min(end_idx, info["end"]):
                    target_runs.append(info)

            if not target_runs:
                return False

            runs_to_move = []
            parent = para._element

            s_info = target_runs[0]
            s_run = s_info["run"]
            s_offset = start_idx - s_info["start"]

            e_info = target_runs[-1]
            e_run = e_info["run"]
            e_offset = end_idx - e_info["start"]

            current_time = datetime.now().isoformat()
            self.change_id += 1
            cid = str(self.change_id)

            # ── Case A: change within a single run ───────────────────────────
            if s_info["index"] == e_info["index"]:
                original_run_text = s_run.text
                prefix = original_run_text[:s_offset]
                middle = original_run_text[s_offset:e_offset]
                suffix = original_run_text[e_offset:]

                s_run.text = prefix
                insert_point = s_run._element

                middle_run_elem = deepcopy(s_run._element)
                t = middle_run_elem.find(qn("w:t"))
                if t is None:
                    t = etree.SubElement(middle_run_elem, qn("w:t"))
                t.text = middle

                if suffix:
                    suffix_run_elem = deepcopy(s_run._element)
                    t2 = suffix_run_elem.find(qn("w:t"))
                    if t2 is None:
                        t2 = etree.SubElement(suffix_run_elem, qn("w:t"))
                    t2.text = suffix
                    parent.insert(parent.index(insert_point) + 1, suffix_run_elem)

                runs_to_move.append(middle_run_elem)

            # ── Case B: change spans multiple runs ───────────────────────────
            else:
                s_text = s_run.text
                s_run.text = s_text[:s_offset]
                insert_point = s_run._element

                s_del_elem = deepcopy(s_run._element)
                t = s_del_elem.find(qn("w:t"))
                if t is None:
                    t = etree.SubElement(s_del_elem, qn("w:t"))
                t.text = s_text[s_offset:]
                runs_to_move.append(s_del_elem)

                for info in target_runs[1:-1]:
                    r_elem = info["run"]._element
                    parent.remove(r_elem)
                    runs_to_move.append(r_elem)

                e_text = e_run.text
                e_del_elem = deepcopy(e_run._element)
                t = e_del_elem.find(qn("w:t"))
                if t is None:
                    t = etree.SubElement(e_del_elem, qn("w:t"))
                t.text = e_text[:e_offset]
                runs_to_move.append(e_del_elem)

                e_run.text = e_text[e_offset:]

            # ── Build <w:del> and <w:ins> ─────────────────────────────────────
            base_idx = parent.index(insert_point)

            del_el = parse_xml(
                f'<w:del w:id="{cid}" w:author="{self.author}" w:date="{current_time}"'
                f' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
            for r in runs_to_move:
                for t_el in r.findall(qn("w:t")):
                    t_el.tag = qn("w:delText")
                del_el.append(r)
            parent.insert(base_idx + 1, del_el)

            self.change_id += 1
            ins_el = parse_xml(
                f'<w:ins w:id="{str(self.change_id)}" w:author="{self.author}"'
                f' w:date="{current_time}"'
                f' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                f'<w:r><w:t xml:space="preserve">{self._esc(effective_modified)}</w:t></w:r>'
                f"</w:ins>"
            )
            parent.insert(base_idx + 2, ins_el)

            return True

        except Exception as e:
            logger.error(f"[TrackChange] Error: {e}")
            import traceback

            traceback.print_exc()
            return False

    def _inject_comments_to_docx(self, file_path: str, comments_el) -> bool:
        """
        将 comments_el (lxml Element) 注入到 docx zip 包中。
        python-docx 默认不会保存手动创建的 comments part，
        所以需要在 doc.save() 之后，操作 zip 文件来注入:
        1. word/comments.xml — 批注内容
        2. [Content_Types].xml — 添加 comments 的 content type
        3. word/_rels/document.xml.rels — 添加批注的关系
        """
        try:
            # 序列化 comments XML
            comments_xml = etree.tostring(
                comments_el, xml_declaration=True, encoding="UTF-8", standalone=True
            )

            # 使用临时文件来安全修改 zip
            tmp_path = file_path + ".tmp"

            with zipfile.ZipFile(file_path, "r") as zin:
                with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                    for item in zin.infolist():
                        if item.filename == "word/comments.xml":
                            continue
                        data = zin.read(item.filename)

                        if item.filename == "[Content_Types].xml":
                            data = self._add_comments_content_type(data)
                        elif item.filename == "word/_rels/document.xml.rels":
                            data = self._add_comments_relationship(data)

                        zout.writestr(item, data)

                    # 添加 word/comments.xml
                    zout.writestr("word/comments.xml", comments_xml)

            # 替换原文件
            shutil.move(tmp_path, file_path)
            logger.info(f"[Hybrid] 💾 comments.xml 已注入 ({len(comments_xml)} bytes)")
            return True

        except Exception as e:
            logger.error(f"[Hybrid] ❌ 注入 comments.xml 失败: {e}")
            tmp_path = file_path + ".tmp"
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
            return False

    def _add_comments_content_type(self, content_types_data: bytes) -> bytes:
        """在 [Content_Types].xml 中添加 comments 的 Override"""
        try:
            root = etree.fromstring(content_types_data)
            ns = "http://schemas.openxmlformats.org/package/2006/content-types"

            for override in root.findall(f"{{{ns}}}Override"):
                if override.get("PartName") == "/word/comments.xml":
                    return content_types_data

            override = etree.SubElement(root, f"{{{ns}}}Override")
            override.set("PartName", "/word/comments.xml")
            override.set(
                "ContentType",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
            )

            return etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone=True
            )
        except Exception as e:
            logger.warning(f"[Hybrid] ⚠️ Content_Types 修改失败: {e}")
            return content_types_data

    def _add_comments_relationship(self, rels_data: bytes) -> bytes:
        """在 document.xml.rels 中添加 comments 关系"""
        try:
            root = etree.fromstring(rels_data)
            ns = "http://schemas.openxmlformats.org/package/2006/relationships"

            for rel in root.findall(f"{{{ns}}}Relationship"):
                if "comments" in rel.get("Type", "").lower():
                    return rels_data

            existing_ids = [
                rel.get("Id", "") for rel in root.findall(f"{{{ns}}}Relationship")
            ]
            max_id = 0
            for rid in existing_ids:
                if rid.startswith("rId"):
                    try:
                        max_id = max(max_id, int(rid[3:]))
                    except ValueError:
                        pass
            new_id = f"rId{max_id + 1}"

            rel = etree.SubElement(root, f"{{{ns}}}Relationship")
            rel.set("Id", new_id)
            rel.set(
                "Type",
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
            )
            rel.set("Target", "comments.xml")

            return etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone=True
            )
        except Exception as e:
            logger.warning(f"[Hybrid] ⚠️ Rels 修改失败: {e}")
            return rels_data

    def _apply_single_comment(
        self,
        doc: Document,
        comments_el,
        original_text: str,
        suggestion_text: str,
        reason: str = "",
        anchor_meta: dict[str, Any] | None = None,
    ) -> bool:
        """应用单个批注（内部方法）"""
        try:
            # Strip Markdown formatting from original_text (same as track-change path)
            stripped = _strip_md_formatting(original_text)
            effective_original = (
                stripped if (stripped and stripped != original_text) else original_text
            )
            anchor_range = self._resolve_comment_anchor_range(
                doc,
                effective_original,
                anchor_meta=anchor_meta,
            )
            if not anchor_range:
                return False

            self.change_id += 1
            comment_id = str(self.change_id)

            comment_content = suggestion_text
            if reason:
                comment_content = f"{suggestion_text}\n\n原因：{reason}"

            comment_xml = f"""
                    <w:comment w:id="{comment_id}" w:author="{self.author}" w:date="{datetime.now().isoformat()}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:p>
                            <w:pPr>
                                <w:pStyle w:val="CommentText"/>
                            </w:pPr>
                            <w:r>
                                <w:t>{self._esc(comment_content)}</w:t>
                            </w:r>
                        </w:p>
                    </w:comment>
                    """
            comments_el.append(parse_xml(comment_xml))

            return self._insert_comment_range_markers(
                anchor_range["start_para"],
                int(anchor_range["start_offset"]),
                anchor_range["end_para"],
                int(anchor_range["end_offset"]),
                comment_id,
            )

        except Exception as e:
            logger.info(f"[Comment] 单条批注失败: {e}")
            return False


if __name__ == "__main__":
    logger.info("Track Changes Editor 已准备就绪")
