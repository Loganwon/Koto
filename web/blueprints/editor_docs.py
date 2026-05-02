# Copyright (C) 2024-2026 Koto AI. All rights reserved.
# SPDX-License-Identifier: LicenseRef-Koto-Proprietary
"""
Editor Document Store API — /api/editor/docs

Lightweight REST API for the Univer Canvas file assistant.
Documents are stored as JSON files under workspace/editor-docs/.

Routes:
  GET    /api/editor/docs              — List all documents
  POST   /api/editor/docs              — Create a new blank document
  GET    /api/editor/docs/<id>         — Get document content
  PUT    /api/editor/docs/<id>         — Save/update document
  PATCH  /api/editor/docs/<id>         — Rename document
  DELETE /api/editor/docs/<id>         — Delete document
  POST   /api/editor/docs/import       — Import a file (txt/md/docx/pdf)
"""

from __future__ import annotations

import json
import logging
import os
import re
import tempfile
import time
import uuid

from flask import Blueprint, Response, jsonify, request

_logger = logging.getLogger("koto.routes.editor_docs")

editor_docs_bp = Blueprint("editor_docs", __name__)

# ── Storage directory ──
_DOCS_DIR: str | None = None


def _get_docs_dir() -> str:
    global _DOCS_DIR
    if _DOCS_DIR is None:
        from web.shared import WORKSPACE_DIR
        _DOCS_DIR = os.path.join(WORKSPACE_DIR, "editor-docs")
    os.makedirs(_DOCS_DIR, exist_ok=True)
    return _DOCS_DIR


def _doc_path(doc_id: str) -> str:
    # Sanitize id to prevent path traversal
    safe_id = re.sub(r"[^a-zA-Z0-9_-]", "", doc_id)
    if not safe_id:
        raise ValueError("Invalid document ID")
    return os.path.join(_get_docs_dir(), f"{safe_id}.json")


def _source_path(doc_id: str, ext: str = ".docx") -> str:
    """Return the path where the original source file is stored. `ext` defaults to .docx for backward compat."""
    safe_id = re.sub(r"[^a-zA-Z0-9_-]", "", doc_id)
    if not safe_id:
        raise ValueError("Invalid document ID")
    return os.path.join(_get_docs_dir(), f"{safe_id}.source{ext}")


def _read_doc(doc_id: str) -> dict | None:
    path = _doc_path(doc_id)
    if not os.path.exists(path):
        return None
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _write_doc(doc: dict) -> None:
    path = _doc_path(doc["id"])
    with open(path, "w", encoding="utf-8") as f:
        json.dump(doc, f, ensure_ascii=False, indent=2)


def _new_id() -> str:
    return uuid.uuid4().hex[:12]


# ── List all docs ──
@editor_docs_bp.route("/api/editor/docs", methods=["GET"])
def list_docs() -> Response:
    docs_dir = _get_docs_dir()
    docs = []
    for fname in sorted(os.listdir(docs_dir)):
        if not fname.endswith(".json"):
            continue
        try:
            with open(os.path.join(docs_dir, fname), "r", encoding="utf-8") as f:
                d = json.load(f)
            docs.append({
                "id": d["id"],
                "name": d.get("name", "未命名"),
                "updatedAt": d.get("updatedAt", ""),
                "createdAt": d.get("createdAt", ""),
                "size": len(d.get("content", "")),
                "sourceExt": d.get("sourceExt", ""),
            })
        except Exception as e:
            _logger.warning("Bad doc file %s: %s", fname, e)
    # Sort by updatedAt descending (most recent first)
    docs.sort(key=lambda x: x.get("updatedAt", ""), reverse=True)
    return jsonify({"docs": docs})


# ── Create blank doc ──
@editor_docs_bp.route("/api/editor/docs", methods=["POST"])
def create_doc() -> Response:
    data = request.get_json(silent=True) or {}
    name = (data.get("name") or "未命名文档").strip()[:200]
    now = time.strftime("%Y-%m-%dT%H:%M:%S")
    doc = {
        "id": _new_id(),
        "name": name,
        "content": "",
        "snapshot": None,
        "createdAt": now,
        "updatedAt": now,
    }
    _write_doc(doc)
    _logger.info("Created doc %s: %s", doc["id"], name)
    return jsonify({"id": doc["id"], "name": doc["name"]}), 201


# ── Get doc ──
@editor_docs_bp.route("/api/editor/docs/<doc_id>", methods=["GET"])
def get_doc(doc_id: str) -> Response:
    doc = _read_doc(doc_id)
    if not doc:
        return jsonify({"error": "Not found"}), 404
    return jsonify(doc)


# ── Save/update doc ──
@editor_docs_bp.route("/api/editor/docs/<doc_id>", methods=["PUT"])
def update_doc(doc_id: str) -> Response:
    doc = _read_doc(doc_id)
    if not doc:
        return jsonify({"error": "Not found"}), 404
    data = request.get_json(silent=True) or {}
    if "content" in data:
        doc["content"] = data["content"]
    if "snapshot" in data:
        doc["snapshot"] = data["snapshot"]
    if "name" in data:
        doc["name"] = str(data["name"]).strip()[:200]
    if "pptxData" in data:
        doc["pptxData"] = data["pptxData"]
    doc["updatedAt"] = time.strftime("%Y-%m-%dT%H:%M:%S")
    _write_doc(doc)
    return jsonify({"ok": True})


# ── Rename doc ──
@editor_docs_bp.route("/api/editor/docs/<doc_id>", methods=["PATCH"])
def rename_doc(doc_id: str) -> Response:
    doc = _read_doc(doc_id)
    if not doc:
        return jsonify({"error": "Not found"}), 404
    data = request.get_json(silent=True) or {}
    name = (data.get("name") or "").strip()[:200]
    if not name:
        return jsonify({"error": "Name required"}), 400
    doc["name"] = name
    doc["updatedAt"] = time.strftime("%Y-%m-%dT%H:%M:%S")
    _write_doc(doc)
    return jsonify({"ok": True})


# ── Delete doc ──
@editor_docs_bp.route("/api/editor/docs/<doc_id>", methods=["DELETE"])
def delete_doc(doc_id: str) -> Response:
    doc = _read_doc(doc_id)
    if not doc:
        return jsonify({"error": "Not found"}), 404
    source_ext = doc.get("sourceExt", ".docx") or ".docx"
    os.remove(_doc_path(doc_id))
    # Also delete stored source file if present (use correct extension)
    source_path = _source_path(doc_id, source_ext)
    if os.path.exists(source_path):
        os.remove(source_path)
    _logger.info("Deleted doc %s", doc_id)
    return jsonify({"ok": True})


# ── Unicode-safe Content-Disposition helper ──
def _content_disposition(filename: str) -> str:
    """Return a Content-Disposition header value that is safe for non-ASCII filenames.

    Uses RFC 5987 `filename*=UTF-8''<percent-encoded>` for full Unicode support
    and falls back to an ASCII-only `filename=` parameter for older clients.
    """
    import urllib.parse
    ascii_name = re.sub(r"[^\x20-\x7E]", "_", filename)
    encoded_name = urllib.parse.quote(filename, safe=" !#$&+-.0-9A-Z_a-z~")
    return f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{encoded_name}"


# ── Get original source file (e.g. raw .docx binary) ──
@editor_docs_bp.route("/api/editor/docs/<doc_id>/source", methods=["GET"])
def get_doc_source(doc_id: str) -> Response:
    doc = _read_doc(doc_id)
    if not doc:
        return jsonify({"error": "Not found"}), 404
    source_ext = doc.get("sourceExt", "")
    if not source_ext:
        return jsonify({"error": "No source file"}), 404
    spath = _source_path(doc_id, source_ext)
    if not os.path.exists(spath):
        return jsonify({"error": "Source file missing"}), 404
    with open(spath, "rb") as fh:
        raw_source = fh.read()
    _MIME_MAP = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }
    mime = _MIME_MAP.get(source_ext, "application/octet-stream")
    safe_name = doc.get("name", "document").replace('"', "'")
    return Response(raw_source, mimetype=mime, headers={
        "Content-Disposition": _content_disposition(safe_name + source_ext),
    })


# ── Get document metadata (page size, margins, default font) for docx-preview layout ──
@editor_docs_bp.route("/api/editor/docs/<doc_id>/meta", methods=["GET"])
def get_doc_meta(doc_id: str) -> Response:
    """Return page-layout metadata extracted from the .docx source via python-docx."""
    doc = _read_doc(doc_id)
    if not doc:
        return jsonify({"error": "Not found"}), 404
    if doc.get("sourceExt") != ".docx":
        return jsonify({"error": "Not a docx"}), 404
    spath = _source_path(doc_id)
    if not os.path.exists(spath):
        return jsonify({"error": "Source file missing"}), 404

    try:
        meta = _extract_docx_meta(spath)
    except Exception as e:
        _logger.warning("meta extraction failed for %s: %s", doc_id, e)
        # Return sensible A4 defaults so the viewer still renders well
        meta = _default_docx_meta()

    return jsonify(meta)


def _emu_to_pt(emu: int) -> float:
    """Convert EMU (English Metric Units) to points. 1 pt = 12700 EMU."""
    return round(emu / 12700, 2)


def _default_docx_meta() -> dict:
    return {
        "pageWidth": 595.28,    # A4 pt
        "pageHeight": 841.89,   # A4 pt
        "marginTop": 72.0,
        "marginBottom": 72.0,
        "marginLeft": 90.0,
        "marginRight": 90.0,
        "defaultFont": "Calibri",
        "defaultFontSize": 11.0,
        "orientation": "portrait",
    }


def _extract_docx_meta(source_path: str) -> dict:
    """Extract page layout and default font info from a .docx source file."""
    try:
        import docx
        from docx.oxml.ns import qn
        doc = docx.Document(source_path)
        section = doc.sections[0]

        page_width  = _emu_to_pt(section.page_width)
        page_height = _emu_to_pt(section.page_height)
        margin_top    = _emu_to_pt(section.top_margin)
        margin_bottom = _emu_to_pt(section.bottom_margin)
        margin_left   = _emu_to_pt(section.left_margin)
        margin_right  = _emu_to_pt(section.right_margin)

        orientation = "landscape" if page_width > page_height else "portrait"

        # Default font — from Normal style or document defaults
        default_font = "Calibri"
        default_font_size = 11.0
        try:
            normal_style = doc.styles["Normal"]
            if normal_style.font.name:
                default_font = normal_style.font.name
            if normal_style.font.size:
                default_font_size = round(normal_style.font.size.pt, 1)
        except Exception:
            pass

        # Try document-level defaults if Normal style has no explicit values
        if default_font == "Calibri":
            try:
                docDefaults = doc.element.body.getparent().find(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}docDefaults"
                )
                if docDefaults is not None:
                    rFonts = docDefaults.find(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
                    )
                    if rFonts is not None:
                        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                        font_name = (
                            rFonts.get(f"{{{ns}}}ascii")
                            or rFonts.get(f"{{{ns}}}hAnsi")
                            or rFonts.get(f"{{{ns}}}cs")
                        )
                        if font_name:
                            default_font = font_name
            except Exception:
                pass

        return {
            "pageWidth": page_width,
            "pageHeight": page_height,
            "marginTop": margin_top,
            "marginBottom": margin_bottom,
            "marginLeft": margin_left,
            "marginRight": margin_right,
            "defaultFont": default_font,
            "defaultFontSize": default_font_size,
            "orientation": orientation,
        }
    except ImportError:
        _logger.warning("python-docx not available; returning default meta")
        return _default_docx_meta()


# ── Import file ──
@editor_docs_bp.route("/api/editor/docs/import", methods=["POST"])
def import_doc() -> Response:
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "Empty filename"}), 400

    original_name = f.filename
    ext = os.path.splitext(original_name)[1].lower()
    raw = f.read()

    # Extract text/data based on file type
    text = ""
    extra: dict = {}
    try:
        if ext in (".txt", ".md", ".csv", ".json", ".html", ".rtf"):
            text = raw.decode("utf-8", errors="replace")
        elif ext == ".docx":
            text = _extract_docx(raw)
            extra["sourceExt"] = ".docx"
        elif ext == ".pdf":
            text = _extract_pdf(raw)
        elif ext == ".pptx":
            from web.blueprints.pptx_editor import _parse_slides as _pptx_parse
            parsed = _pptx_parse(raw)
            extra["sourceExt"] = ".pptx"
            extra["pptxData"] = {
                "slideWidthEmu": parsed["slide_width_emu"],
                "slideHeightEmu": parsed["slide_height_emu"],
                "slides": parsed["slides"],
            }
            text = f"[PowerPoint: {len(parsed['slides'])} 张幻灯片]"
        elif ext == ".xlsx":
            from app.core.file.file_parser import parse_xlsx as _xlsx_parse
            with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tf:
                tf.write(raw)
                tmp_path = tf.name
            try:
                sheets_data = _xlsx_parse(tmp_path, original_name)
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass
            extra["sourceExt"] = ".xlsx"
            extra["sheetsData"] = sheets_data
            n_sheets = len(sheets_data.get("sheetOrder", []))
            warnings = sheets_data.get("_warnings", [])
            text = f"[Excel: {n_sheets} 个工作表{'（含公式，已静态化）' if warnings else ''}]"
        else:
            text = raw.decode("utf-8", errors="replace")
    except Exception as e:
        _logger.error("Import parse error for %s: %s", original_name, e)
        return jsonify({"error": f"无法解析文件: {e}"}), 400

    now = time.strftime("%Y-%m-%dT%H:%M:%S")
    doc_name = os.path.splitext(original_name)[0][:200]
    new_id = _new_id()
    doc: dict = {
        "id": new_id,
        "name": doc_name,
        "content": text,
        "snapshot": None,
        "createdAt": now,
        "updatedAt": now,
        "importedFrom": original_name,
        **extra,
    }
    # Persist binary source for formats that need it
    if ext in (".docx", ".pptx", ".xlsx"):
        spath = _source_path(new_id, ext)
        with open(spath, "wb") as sfh:
            sfh.write(raw)
    _write_doc(doc)
    _logger.info("Imported %s → doc %s (%d chars)", original_name, doc["id"], len(text))
    return jsonify({"id": doc["id"], "name": doc_name, "size": len(text), "sourceExt": doc.get("sourceExt", "")}), 201


@editor_docs_bp.route("/api/editor/docs/import_path", methods=["POST"])
def import_doc_from_path() -> Response:
    """Import a document from a server-side file path (JSON body: {"path": "..."})."""
    data = request.get_json(silent=True) or {}
    file_path = data.get("path", "")
    if not file_path:
        return jsonify({"error": "Missing 'path' in request body"}), 400

    try:
        with open(file_path, "rb") as fh:
            raw = fh.read()
    except FileNotFoundError:
        return jsonify({"error": f"File not found: {file_path}"}), 404
    except OSError as e:
        return jsonify({"error": str(e)}), 400

    original_name = os.path.basename(file_path)
    ext = os.path.splitext(original_name)[1].lower()

    text = ""
    extra2: dict = {}
    try:
        if ext in (".txt", ".md", ".csv", ".json", ".html", ".rtf"):
            text = raw.decode("utf-8", errors="replace")
        elif ext == ".docx":
            text = _extract_docx(raw)
            extra2["sourceExt"] = ".docx"
        elif ext == ".pdf":
            text = _extract_pdf(raw)
        elif ext == ".pptx":
            from web.blueprints.pptx_editor import _parse_slides as _pptx_parse
            parsed = _pptx_parse(raw)
            extra2["sourceExt"] = ".pptx"
            extra2["pptxData"] = {
                "slideWidthEmu": parsed["slide_width_emu"],
                "slideHeightEmu": parsed["slide_height_emu"],
                "slides": parsed["slides"],
            }
            text = f"[PowerPoint: {len(parsed['slides'])} 张幻灯片]"
        elif ext == ".xlsx":
            from app.core.file.file_parser import parse_xlsx as _xlsx_parse
            with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tf:
                tf.write(raw)
                tmp_path = tf.name
            try:
                sheets_data = _xlsx_parse(tmp_path, original_name)
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass
            extra2["sourceExt"] = ".xlsx"
            extra2["sheetsData"] = sheets_data
            n_sheets = len(sheets_data.get("sheetOrder", []))
            warnings = sheets_data.get("_warnings", [])
            text = f"[Excel: {n_sheets} 个工作表{'（含公式，已静态化）' if warnings else ''}]"
        else:
            text = raw.decode("utf-8", errors="replace")
    except Exception as e:
        _logger.error("Import_path parse error for %s: %s", file_path, e)
        return jsonify({"error": f"无法解析文件: {e}"}), 400

    now = time.strftime("%Y-%m-%dT%H:%M:%S")
    doc_name = os.path.splitext(original_name)[0][:200]
    new_id = _new_id()
    doc: dict = {
        "id": new_id,
        "name": doc_name,
        "content": text,
        "snapshot": None,
        "createdAt": now,
        "updatedAt": now,
        "importedFrom": original_name,
        **extra2,
    }
    # Persist binary source for formats that need it
    if ext in (".docx", ".pptx", ".xlsx"):
        spath = _source_path(new_id, ext)
        with open(spath, "wb") as sfh:
            sfh.write(raw)
    _write_doc(doc)
    _logger.info("Imported path %s → doc %s (%d chars)", file_path, doc["id"], len(text))
    return jsonify({"id": doc["id"], "name": doc_name, "size": len(text), "sourceExt": doc.get("sourceExt", "")}), 201


# ── Text extraction helpers ──

def _extract_docx(raw_bytes: bytes) -> str:
    """Extract text from .docx — includes body paragraphs AND table cell text, in document order."""
    try:
        import docx
        import io
        doc = docx.Document(io.BytesIO(raw_bytes))
        parts = []
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for block in doc.element.body:
            tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag
            if tag == "p":
                # Top-level body paragraph
                texts = [r.text for r in block.iter(f"{{{W}}}t") if r.text]
                line = "".join(texts)
                parts.append(line)
            elif tag == "tbl":
                # Table — iterate rows/cells, produce TSV rows
                from docx.table import Table
                tbl = Table(block, doc)
                for row in tbl.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = " ".join(
                            p.text.strip() for p in cell.paragraphs if p.text.strip()
                        )
                        row_texts.append(cell_text)
                    parts.append("\t".join(row_texts))
        return "\n".join(parts)
    except ImportError:
        pass
    # Fallback: extract from XML inside zip
    import io
    import zipfile
    import xml.etree.ElementTree as ET
    try:
        z = zipfile.ZipFile(io.BytesIO(raw_bytes))
        xml_content = z.read("word/document.xml")
        tree = ET.fromstring(xml_content)
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        W_p = f"{{{W}}}p"
        W_tbl = f"{{{W}}}tbl"
        W_tr = f"{{{W}}}tr"
        W_tc = f"{{{W}}}tc"
        W_t = f"{{{W}}}t"
        parts = []
        for block in tree.find(f"{{{W}}}body") or []:
            if block.tag == W_p:
                line = "".join(r.text or "" for r in block.iter(W_t))
                parts.append(line)
            elif block.tag == W_tbl:
                for tr in block.iter(W_tr):
                    row_texts = []
                    for tc in tr.findall(W_tc):
                        cell_text = " ".join(
                            "".join(r.text or "" for r in p.iter(W_t)).strip()
                            for p in tc.iter(W_p)
                            if "".join(r.text or "" for r in p.iter(W_t)).strip()
                        )
                        row_texts.append(cell_text)
                    parts.append("\t".join(row_texts))
        return "\n".join(parts)
    except Exception as e:
        raise ValueError(f"Failed to parse docx: {e}") from e


def _extract_pdf(raw_bytes: bytes) -> str:
    """Extract text from PDF using PyPDF2/pypdf if available."""
    try:
        import pypdf
        import io
        reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
        texts = []
        for page in reader.pages:
            texts.append(page.extract_text() or "")
        return "\n".join(texts)
    except ImportError:
        pass
    try:
        import PyPDF2
        import io
        reader = PyPDF2.PdfReader(io.BytesIO(raw_bytes))
        texts = []
        for page in reader.pages:
            texts.append(page.extract_text() or "")
        return "\n".join(texts)
    except ImportError:
        raise ValueError("PDF 解析需要 pypdf 或 PyPDF2 库，请安装后重试")


# ── Download PPTX with applied text edits ──
@editor_docs_bp.route("/api/editor/docs/<doc_id>/pptx_download", methods=["GET"])
def pptx_download(doc_id: str) -> Response:
    """Re-render and download the .pptx with in-place text edits applied."""
    doc = _read_doc(doc_id)
    if not doc or doc.get("sourceExt") != ".pptx":
        return jsonify({"error": "Not a PPTX document"}), 404
    spath = _source_path(doc_id, ".pptx")
    if not os.path.exists(spath):
        return jsonify({"error": "Source file missing"}), 404
    pptx_data = doc.get("pptxData") or {}
    slides_edits = pptx_data.get("slides", [])
    with open(spath, "rb") as fh:
        orig_bytes = fh.read()
    try:
        from web.blueprints.pptx_editor import _apply_edits
        result = _apply_edits(orig_bytes, slides_edits)
    except Exception as e:
        _logger.error("pptx_download apply_edits failed for %s: %s", doc_id, e)
        return jsonify({"error": f"导出失败: {e}"}), 500
    safe_name = doc.get("name", "presentation").replace('"', "'")
    return Response(
        result,
        mimetype="application/vnd.openxmlformats-officedocument.presentationml.presentation",
        headers={"Content-Disposition": _content_disposition(safe_name + ".pptx")},
    )
