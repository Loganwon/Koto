from docx import Document
doc_path = r'C:\Users\12524\Desktop\Koto\workspace\雷鸟访谈问题.docx'
doc = Document(doc_path)
print(f'原始表格数量: {len(doc.tables)}')
paragraphs = doc.paragraphs
print(f'段落数量: {len(paragraphs)}')
new_doc = Document()
for index, para in enumerate(paragraphs, start=1):
    print(f'copying paragraph {index}')
    new_para = new_doc.add_paragraph(para.text)
    if para.style.name:
        new_para.style = para.style
new_doc.save(doc_path)
print('表格已删除，文档已保存')
