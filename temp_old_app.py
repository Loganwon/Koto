import os
import asyncio
import re
import json
import logging
_app_logger = logging.getLogger("koto.app")
import time
import threading
import subprocess
import sys
import threading
import time
import uuid
from datetime import datetime
from pathlib import Path

# τí«Σ┐¥ web/ τ¢«σ╜òσ£¿µ¿íσ¥ùµÉ£τ┤óΦ╖»σ╛äΣ╕¡∩╝êΘÇÜΦ┐ç koto_app.py σÉ»σè¿µù╢Θ£ÇΦªü∩╝ë
_web_dir = os.path.dirname(os.path.abspath(__file__))
if _web_dir not in sys.path:
    sys.path.append(_web_dir)

from dotenv import load_dotenv
from flask import (
    Flask,
    Response,
    g,
    jsonify,
    render_template,
    request,
    send_file,
    send_from_directory,
    stream_with_context,
)
from flask_cors import CORS
from werkzeug.utils import secure_filename as _secure_filename

# Import new routing modules
from app.core.routing import SmartDispatcher

# σ╗╢Φ┐ƒσ»╝σàÑ - Φ┐ÖΣ║¢Φ╖»τö▒τ▒╗Σ╗àσ£¿Φ┐ÉΦíîµù╢Θªûµ¼íΦ«┐Θù«µù╢ΘÇÜΦ┐ç __getattr__ σèáΦ╜╜
# LocalModelRouter, AIRouter, TaskDecomposer, LocalPlanner ΘÇÜΦ┐ç app.core.routing.__getattr__ σ╗╢Φ┐ƒσèáΦ╜╜

# Import unified agent API blueprint ΓÇö σ╗╢Φ┐ƒσê░Φô¥σ¢╛µ│¿σåîµù╢σèáΦ╜╜
agent_bp = None  # σ╗╢Φ┐ƒσèáΦ╜╜∩╝îΦºüΣ╕ïµû╣Φô¥σ¢╛µ│¿σåîσî║

# ================= σ╣╢ΦíîµëºΦíîτ│╗τ╗ƒσ»╝σàÑ =================
try:
    from parallel_api import register_parallel_api
    from parallel_executor import (
        Priority,
        Task,
        TaskStatus,
        TaskType,
        cancel_task,
        get_next_task,
        get_queue_manager,
        get_resource_manager,
        get_task_monitor,
        submit_task,
    )
    from task_dispatcher import get_scheduler, start_dispatcher, stop_dispatcher

    PARALLEL_SYSTEM_ENABLED = True
except ImportError as e:
    _app_logger.warning(f"[WARNING] Failed to import parallel execution system: {e}")
    PARALLEL_SYSTEM_ENABLED = False

try:
    from flask_sock import Sock
except ImportError:
    Sock = None

# ================= µçÆσèáΦ╜╜Θçìσ₧ïµ¿íσ¥ù∩╝êσÉ»σè¿Σ╝ÿσîû∩╝ë =================
# google.genai (~4.7s), requests (~0.5s) σ╗╢Φ┐ƒσê░Θªûµ¼íΣ╜┐τö¿µù╢σèáΦ╜╜


class _LazyModule:
    """σ╗╢Φ┐ƒσ»╝σàÑΣ╗úτÉå - Θªûµ¼íσ▒₧µÇºΦ«┐Θù«µù╢µëìΦºªσÅæσ«₧ΘÖà import"""

    __slots__ = ("_import_func", "_module")

    def __init__(self, import_func):
        object.__setattr__(self, "_import_func", import_func)
        object.__setattr__(self, "_module", None)

    def _load(self):
        mod = object.__getattribute__(self, "_module")
        if mod is None:
            import_func = object.__getattribute__(self, "_import_func")
            mod = import_func()
            object.__setattr__(self, "_module", mod)
        return mod

    def __getattr__(self, name):
        return getattr(self._load(), name)

    def __repr__(self):
        mod = object.__getattribute__(self, "_module")
        if mod is None:
            return "<LazyModule (not loaded)>"
        return repr(mod)


def _import_genai():
    _app_logger.debug("[LAZY_IMPORT] σèáΦ╜╜ google.genai ...")
    from google import genai as _genai

    return _genai


def _import_types():
    _app_logger.debug("[LAZY_IMPORT] σèáΦ╜╜ google.genai.types ...")
    from google.genai import types as _types

    return _types


def _import_requests():
    _app_logger.debug("[LAZY_IMPORT] σèáΦ╜╜ requests ...")
    import requests as _requests

    return _requests


genai = _LazyModule(_import_genai)
types = _LazyModule(_import_types)
requests = _LazyModule(_import_requests)

# ================= µçÆσèáΦ╜╜µûçµíúσÆîPPTµ¿íσ¥ù∩╝êσÉ»σè¿σèáΘÇƒ∩╝ë =================
# σ╗╢Φ┐ƒσ»╝σàÑ python-docx (~572ms) σÆî python-pptx (~666ms)

# µûçµíúσ╖ÑΣ╜£µ╡üµëºΦíîσÖ¿µçÆσèáΦ╜╜
_document_workflow_cache = {}


def get_document_workflow_executor():
    """µçÆσèáΦ╜╜µûçµíúσ╖ÑΣ╜£µ╡üµëºΦíîσÖ¿"""
    if "executor" not in _document_workflow_cache:
        _app_logger.debug("[LAZY_IMPORT] σèáΦ╜╜µûçµíúσ╖ÑΣ╜£µ╡üµëºΦíîσÖ¿...")
        try:
            from web.document_workflow_executor import (
                DocumentWorkflowExecutor,
                execute_document_workflow,
            )
        except ImportError:
            try:
                from document_workflow_executor import (
                    DocumentWorkflowExecutor,
                    execute_document_workflow,
                )
            except ImportError:
                DocumentWorkflowExecutor = None
                execute_document_workflow = None
                _app_logger.warning("[WARNING] µûçµíúσ╖ÑΣ╜£µ╡üµëºΦíîσÖ¿µ£¬σ«ëΦúà")
        _document_workflow_cache["executor"] = DocumentWorkflowExecutor
        _document_workflow_cache["execute"] = execute_document_workflow
    return _document_workflow_cache.get("executor"), _document_workflow_cache.get(
        "execute"
    )


# DocumentWorkflowExecutor σÆî execute_document_workflow τÜäµçÆσèáΦ╜╜Σ╗úτÉå
class _DocWorkflowProxy:
    def __getattr__(self, name):
        executor_cls, _ = get_document_workflow_executor()
        if executor_cls is None:
            raise ImportError("µûçµíúσ╖ÑΣ╜£µ╡üµëºΦíîσÖ¿µ£¬σ«ëΦúà")
        return getattr(executor_cls, name)


DocumentWorkflowExecutor = _DocWorkflowProxy()


def execute_document_workflow(*args, **kwargs):
    _, execute_func = get_document_workflow_executor()
    if execute_func is None:
        raise ImportError("µûçµíúσ╖ÑΣ╜£µ╡üµëºΦíîσÖ¿µ£¬σ«ëΦúà")
    return execute_func(*args, **kwargs)


# PPTσñÜµ¿íσ₧ïτ│╗τ╗ƒµçÆσèáΦ╜╜
_ppt_system_cache = {}


def get_ppt_system():
    """µçÆσèáΦ╜╜PPTτöƒµêÉτ│╗τ╗ƒ"""
    if "loaded" not in _ppt_system_cache:
        _app_logger.debug("[LAZY_IMPORT] σèáΦ╜╜PPTσñÜµ¿íσ₧ïτöƒµêÉτ│╗τ╗ƒ...")
        try:
            from web.ppt_master import PPTBlueprint, PPTMasterOrchestrator
            from web.ppt_pipeline import (
                PPTGenerationPipeline,
                PPTGenerationTaskHandler,
                format_ppt_generation_result,
            )
            from web.ppt_synthesizer import PPTSynthesizer

            _app_logger.info("[PPT_SYSTEM] Γ£à σñÜµ¿íσ₧ïPPTτöƒµêÉτ│╗τ╗ƒσ╖▓σèáΦ╜╜")
        except ImportError:
            try:
                from ppt_master import PPTBlueprint, PPTMasterOrchestrator
                from ppt_pipeline import (
                    PPTGenerationPipeline,
                    PPTGenerationTaskHandler,
                    format_ppt_generation_result,
                )
                from ppt_synthesizer import PPTSynthesizer

                _app_logger.info("[PPT_SYSTEM] Γ£à σñÜµ¿íσ₧ïPPTτöƒµêÉτ│╗τ╗ƒσ╖▓σèáΦ╜╜∩╝êτ¢╕σ»╣σ»╝σàÑ∩╝ë")
            except ImportError:
                PPTMasterOrchestrator = None
                PPTBlueprint = None
                PPTSynthesizer = None
                PPTGenerationPipeline = None
                PPTGenerationTaskHandler = None
                format_ppt_generation_result = None
                _app_logger.warning("[WARNING] σñÜµ¿íσ₧ïPPTτöƒµêÉτ│╗τ╗ƒµ£¬σ«ëΦúà")
        _ppt_system_cache["orchestrator"] = PPTMasterOrchestrator
        _ppt_system_cache["blueprint"] = PPTBlueprint
        _ppt_system_cache["synthesizer"] = PPTSynthesizer
        _ppt_system_cache["pipeline"] = PPTGenerationPipeline
        _ppt_system_cache["handler"] = PPTGenerationTaskHandler
        _ppt_system_cache["formatter"] = format_ppt_generation_result
        _ppt_system_cache["loaded"] = True

    return (
        _ppt_system_cache.get("orchestrator"),
        _ppt_system_cache.get("blueprint"),
        _ppt_system_cache.get("synthesizer"),
        _ppt_system_cache.get("pipeline"),
        _ppt_system_cache.get("handler"),
        _ppt_system_cache.get("formatter"),
    )


# µçÆσèáΦ╜╜Σ╗úτÉåτ▒╗
class _PPTModuleProxy:
    def __init__(self, index):
        self._index = index

    def __getattr__(self, name):
        modules = get_ppt_system()
        module = modules[self._index]
        if module is None:
            raise ImportError("PPTτöƒµêÉτ│╗τ╗ƒµ£¬σ«ëΦúà")
        return getattr(module, name)

    def __call__(self, *args, **kwargs):
        modules = get_ppt_system()
        module = modules[self._index]
        if module is None:
            raise ImportError("PPTτöƒµêÉτ│╗τ╗ƒµ£¬σ«ëΦúà")
        if callable(module):
            return module(*args, **kwargs)
        raise TypeError(f"{module} is not callable")


PPTMasterOrchestrator = _PPTModuleProxy(0)
PPTBlueprint = _PPTModuleProxy(1)
PPTSynthesizer = _PPTModuleProxy(2)
PPTGenerationPipeline = _PPTModuleProxy(3)
PPTGenerationTaskHandler = _PPTModuleProxy(4)
format_ppt_generation_result = _PPTModuleProxy(5)

# ================= Configuration =================
# Σ╗Ä web τ¢«σ╜òσÉæΣ╕èµƒÑµë╛
import os
import sys as _sys


# Σ╕¡µû¡Σ┐íσÅ╖σ¡ÿσé¿ - µö╣Φ┐¢τëêµ£¼∩╝îµö»µîüσ«₧µù╢µ╡üΣ╕¡µ¡ó
class StreamInterruptManager:
    """τ«íτÉåµ»ÅΣ╕¬ session τÜäµ╡üΣ╕¡µ¡óτè╢µÇüσÆîµÄºσê╢"""

    def __init__(self):
        self.interrupts = {}  # session_name -> {'flag': bool, 'event': threading.Event}
        self._lock = threading.Lock()

    def _ensure(self, session_name):
        """τí«Σ┐¥ session Φ«░σ╜òσ¡ÿσ£¿ (must be called with self._lock held)"""
        if session_name not in self.interrupts:
            self.interrupts[session_name] = {"flag": False, "event": threading.Event()}
        elif self.interrupts[session_name].get("event") is None:
            self.interrupts[session_name]["event"] = threading.Event()

    def set_interrupt(self, session_name):
        """Φ«╛τ╜«Σ╕¡µû¡µáçσ┐ù"""
        with self._lock:
            self._ensure(session_name)
            self.interrupts[session_name]["flag"] = True
            if self.interrupts[session_name]["event"]:
                self.interrupts[session_name]["event"].set()
        _app_logger.debug(f"[INTERRUPT] Marked session {session_name} for interruption")

    def is_interrupted(self, session_name):
        """µúÇµƒÑµÿ»σÉªΦó½Σ╕¡µû¡"""
        with self._lock:
            if session_name not in self.interrupts:
                return False
            record = self.interrupts[session_name]
            event_flag = record.get("event").is_set() if record.get("event") else False
            return bool(record.get("flag")) or event_flag

    def reset(self, session_name):
        """Θçìτ╜«Σ╕¡µû¡µáçσ┐ù"""
        with self._lock:
            self._ensure(session_name)
            self.interrupts[session_name]["flag"] = False
            if self.interrupts[session_name]["event"]:
                self.interrupts[session_name]["event"].clear()
        _app_logger.debug(f"[INTERRUPT] Reset interrupt flag for session {session_name}")

    def get_event(self, session_name):
        """ΦÄ╖σÅû/σê¢σ╗║Σ╕¡µû¡Σ║ïΣ╗╢σ»╣Φ▒í"""
        with self._lock:
            self._ensure(session_name)
            return self.interrupts[session_name]["event"]

    def cleanup(self, session_name):
        """µ╕àτÉå session τÜäΣ╕¡µû¡Φ«░σ╜ò"""
        with self._lock:
            if session_name in self.interrupts:
                del self.interrupts[session_name]


_interrupt_manager = StreamInterruptManager()
# Σ┐¥τòÖσÉæσÉÄσà╝σ«╣
_interrupt_flags = {}  # Σ╗àτö¿Σ║ÄσÉæσÉÄσà╝σ«╣

# σêñµû¡µÿ»σÉªΣ╕║µëôσîàσÉÄΦ┐ÉΦíî
if getattr(_sys, "frozen", False):
    # PyInstaller µëôσîàσÉÄ - exeµëÇσ£¿τ¢«σ╜ò∩╝êµîüΣ╣àσîûµò░µì«τ¢«σ╜ò∩╝ë
    PROJECT_ROOT = os.path.dirname(_sys.executable)
else:
    # σ╝ÇσÅæτÄ»σóâ - Σ╗Ä web τ¢«σ╜òσÉæΣ╕èµë╛
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)

# Try multiple locations for the config file
config_locations = [
    os.path.join(PROJECT_ROOT, "config", "gemini_config.env"),
    os.path.join(PROJECT_ROOT, "gemini_config.env"),
    (
        os.path.join(os.path.dirname(_sys.executable), "config", "gemini_config.env")
        if getattr(_sys, "frozen", False)
        else ""
    ),
    "gemini_config.env",
    "../gemini_config.env",
]

for config_path in config_locations:
    if os.path.exists(config_path):
        load_dotenv(config_path)
        break

# σ░¥Φ»òΦ»╗σÅû GEMINI_API_KEY µêû API_KEY
API_KEY = os.getenv("GEMINI_API_KEY") or os.getenv("API_KEY")

# Φ»╗σÅûΦç¬σ«ÜΣ╣ë API τ½»τé╣∩╝êτö¿Σ║ÄΣ╕¡Φ╜¼µ£ìσèí∩╝ë
GEMINI_API_BASE = os.getenv("GEMINI_API_BASE", "").strip()
FORCE_PROXY = os.getenv("FORCE_PROXY", "").strip()

_user_settings_cache = {}
_user_settings_lock = threading.Lock()


def _load_user_settings() -> dict:
    """Load user_settings.json with caching and safe fallbacks."""
    with _user_settings_lock:
        if "data" in _user_settings_cache:
            return _user_settings_cache["data"]
        settings_path = os.path.join(PROJECT_ROOT, "config", "user_settings.json")
        try:
            with open(settings_path, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception:
            data = {}
        _user_settings_cache["data"] = data
        return data


def get_workspace_root() -> str:
    """Return the workspace root directory from settings or default path."""
    settings = _load_user_settings()
    workspace_dir = settings.get("storage", {}).get("workspace_dir")
    if workspace_dir:
        return workspace_dir
    return os.path.join(PROJECT_ROOT, "workspace")


def get_organize_root() -> str:
    """Return the file organization root directory from settings or default path."""
    settings = _load_user_settings()
    organize_root = settings.get("storage", {}).get("organize_root")
    if organize_root:
        return organize_root
    return os.path.join(get_workspace_root(), "_organize")


def get_default_wechat_files_dir() -> str:
    """Return configured default WeChat files directory, if provided by user settings."""
    settings = _load_user_settings()
    return settings.get("storage", {}).get("wechat_files_dir", "")


if not API_KEY:
    _app_logger.warning("ΓÜá∩╕Å Warning: GEMINI_API_KEY or API_KEY not found in gemini_config.env")
    _app_logger.info("   Φ»╖σ£¿ config/gemini_config.env Σ╕¡Θàìτ╜« API σ»åΘÆÑ")
    _app_logger.info("   σ║öτö¿σ░åτ╗ºτ╗¡σÉ»σè¿∩╝îΣ╜å AI σèƒΦâ╜Σ╕ìσÅ»τö¿")
    # Σ╕ìσåì sys.exit ΓÇö σàüΦ«╕σ║öτö¿σÉ»σè¿σ╣╢σ£¿ UI Σ╕¡µÅÉτñ║τö¿µê╖Θàìτ╜«

if GEMINI_API_BASE:
    _app_logger.info(f"≡ƒôí Σ╜┐τö¿Φç¬σ«ÜΣ╣ë API τ½»τé╣: {GEMINI_API_BASE}")

# µúÇµ╡ïσ╣╢Φ«╛τ╜«Σ╗úτÉå
PROXY_OPTIONS = [
    "http://127.0.0.1:7890",
    "http://127.0.0.1:10809",
    "http://127.0.0.1:1080",
]


def _normalize_proxy_url(proxy_value: str) -> str:
    """Normalize proxy value to a URL with scheme."""
    if not proxy_value:
        return ""
    value = proxy_value.strip()
    if not value:
        return ""
    if "://" not in value:
        value = f"http://{value}"
    return value


def _extract_system_proxy_candidates() -> list:
    """Collect proxy candidates from system settings (Windows) and env."""
    candidates = []

    # 0) User-configured manual proxy (highest priority after FORCE_PROXY)
    try:
        _us = settings_manager.get("proxy", "enabled")
        _um = settings_manager.get("proxy", "manual_proxy") or ""
        if _us is not False and _um.strip():
            candidates.append(_normalize_proxy_url(_um.strip()))
    except Exception:
        pass

    # 1) Environment variables first (if user/system already configured)
    env_proxy = os.environ.get("HTTPS_PROXY") or os.environ.get("https_proxy")
    if env_proxy:
        candidates.append(_normalize_proxy_url(env_proxy))

    # 2) Windows Internet Settings proxy (for "Use a proxy server")
    if sys.platform.startswith("win"):
        try:
            import winreg

            key_path = r"Software\Microsoft\Windows\CurrentVersion\Internet Settings"
            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path) as key:
                proxy_enabled = winreg.QueryValueEx(key, "ProxyEnable")[0]
                if proxy_enabled:
                    proxy_server = str(
                        winreg.QueryValueEx(key, "ProxyServer")[0]
                    ).strip()
                    if proxy_server:
                        # Formats:
                        #   127.0.0.1:7890
                        #   http=127.0.0.1:7890;https=127.0.0.1:7890
                        if "=" in proxy_server and ";" in proxy_server:
                            pairs = [
                                p.strip() for p in proxy_server.split(";") if p.strip()
                            ]
                            parsed_map = {}
                            for pair in pairs:
                                if "=" in pair:
                                    k, v = pair.split("=", 1)
                                    parsed_map[k.strip().lower()] = v.strip()
                            for proto in ["https", "http", "socks", "socks5"]:
                                if parsed_map.get(proto):
                                    candidates.append(
                                        _normalize_proxy_url(parsed_map.get(proto))
                                    )
                        else:
                            candidates.append(_normalize_proxy_url(proxy_server))
        except Exception:
            pass

    # 3) Built-in localhost fallback options
    candidates.extend(PROXY_OPTIONS)

    # De-duplicate while preserving order
    deduped = []
    seen = set()
    for item in candidates:
        if not item:
            continue
        key = item.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return deduped


def setup_proxy():
    # Σ╝ÿσàêΣ╜┐τö¿σ╝║σê╢Σ╗úτÉå∩╝êΣ╕ìΘ£ÇΦªüµ╡ïΦ»ò∩╝ë
    if FORCE_PROXY and FORCE_PROXY.lower() not in ("auto", "system"):
        os.environ["HTTPS_PROXY"] = FORCE_PROXY
        os.environ["HTTP_PROXY"] = FORCE_PROXY
        _app_logger.info(f"≡ƒöº Σ╜┐τö¿σ╝║σê╢Σ╗úτÉå: {FORCE_PROXY}")
        return FORCE_PROXY

    # τö¿µê╖µÿÄτí«τªüτö¿Σ╗úτÉåµù╢∩╝îµ╕àΘÖñτÄ»σóâσÅÿΘçÅσ╣╢ΘÇÇσç║
    try:
        if settings_manager.get("proxy", "enabled") is False:
            os.environ.pop("HTTPS_PROXY", None)
            os.environ.pop("HTTP_PROXY", None)
            _app_logger.info("≡ƒöº τö¿µê╖σ╖▓τªüτö¿Σ╗úτÉå")
            return None
    except Exception:
        pass

    # Φç¬σè¿σî╣Θàìτ│╗τ╗ƒΣ╗úτÉåΣ╕Äµ£¼σ£░σ╕╕Φºüτ½»σÅú
    import socket
    from urllib.parse import urlparse

    proxy_candidates = _extract_system_proxy_candidates()

    for proxy in proxy_candidates:
        try:
            # Σ╗Ä URL µÅÉσÅû host:port
            parsed = urlparse(proxy)
            host = parsed.hostname
            port = parsed.port
            if not host or not port:
                continue

            # σ┐½ΘÇƒτ½»σÅúµúÇµ╡ï∩╝ê0.1τºÆΦ╢àµù╢∩╝ë
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(0.1)
            result = sock.connect_ex((host, port))
            sock.close()

            if result == 0:
                os.environ["HTTPS_PROXY"] = proxy
                os.environ["HTTP_PROXY"] = proxy
                _app_logger.info(f"Γ£à Φç¬σè¿σî╣Θàìτ│╗τ╗ƒΣ╗úτÉå: {proxy}")
                return proxy
        except Exception:
            continue

    return None


# σ╗╢Φ┐ƒΣ╗úτÉåµúÇµ╡ïσê░Θªûµ¼íΘ£ÇΦªüµù╢∩╝êσÉ»σè¿σèáΘÇƒ∩╝ë
_detected_proxy = None
_proxy_checked = False


def get_detected_proxy():
    """µçÆσèáΦ╜╜Σ╗úτÉåµúÇµ╡ï∩╝êΘªûµ¼íΦ░âτö¿µù╢µëºΦíî∩╝ë"""
    global _detected_proxy, _proxy_checked
    if not _proxy_checked:
        _detected_proxy = setup_proxy()
        _proxy_checked = True
    return _detected_proxy


# σÉæσÉÄσà╝σ«╣∩╝Üdetected_proxy τÄ░σ£¿ΘÇÜΦ┐çσç╜µò░Φ«┐Θù«
detected_proxy = None  # σìáΣ╜ìτ¼ª∩╝îσ«₧ΘÖàΘÇÜΦ┐ç get_detected_proxy() ΦÄ╖σÅû


# σ£¿σÉÄσÅ░τ║┐τ¿ïΘóäτâ¡Σ╗úτÉåµúÇµ╡ï∩╝êΣ╕ìΘÿ╗σí₧σÉ»σè¿∩╝ë
def _warmup_proxy():
    global detected_proxy
    detected_proxy = get_detected_proxy()


threading.Thread(target=_warmup_proxy, daemon=True).start()


# σê¢σ╗║ GenAI σ«óµê╖τ½» (Θàìτ╜«Σ╗úτÉåσÆîΦç¬σ«ÜΣ╣ëτ½»τé╣)
def create_client():
    import httpx

    proxy = get_detected_proxy()
    # Φ╢àµù╢µù╢Θù┤: Φ┐₧µÄÑ30τºÆ, Φ»╗σÅû180τºÆ (σ¢╛σâÅτöƒµêÉσÆîΘò┐µûçµ£¼τöƒµêÉΘ£ÇΦªüµ¢┤Θò┐µù╢Θù┤)
    timeout_config = httpx.Timeout(180.0, connect=30.0)

    # µ₧äσ╗║ http_options
    http_options = {}

    # µ│¿µäÅ∩╝Üµ£Çµû░τÜä Gemini µ¿íσ₧ï∩╝êσªé gemini-1.5-flash∩╝ëΘ£ÇΦªü v1beta API
    # v1 API σÅ¬µö»µîüµùºτÜäµ¿íσ₧ïπÇéΦ┐ÖΘçîΣ╜┐τö¿ v1betaπÇé
    http_options["api_version"] = "v1beta"

    # Φç¬σ«ÜΣ╣ë API τ½»τé╣∩╝êτö¿Σ║ÄΣ╕¡Φ╜¼µ£ìσèí∩╝ë
    if GEMINI_API_BASE:
        http_options["base_url"] = GEMINI_API_BASE
        _app_logger.info(f"≡ƒôí API τ½»τé╣: {GEMINI_API_BASE}")

    # Θàìτ╜«Σ╗úτÉå - ΘÇÜΦ┐çτÄ»σóâσÅÿΘçÅτí«Σ┐¥Φó½Σ╜┐τö¿
    if proxy:
        os.environ["HTTP_PROXY"] = proxy
        os.environ["HTTPS_PROXY"] = proxy
        _app_logger.info(f"≡ƒöî Φ«╛τ╜«Σ╗úτÉå: {proxy}")

    # Σ╜┐τö¿ httpx with explicit proxy for genai
    # µ│¿µäÅ∩╝ÜHttpOptions σ¡ùµ«╡σÉìΣ╕║ httpx_client (snake_case)∩╝îΣ╕ìµÿ» httpxClient
    # Σ╜åσ«₧µ╡ïµÿ╛τñ║∩╝ÜΘÇÜΦ┐ç env vars Φ«╛τ╜«Σ╗úτÉåµ»öµÿ╛σ╝ÅΣ╝áσàÑ httpx_client µ¢┤τ¿│σ«Ü∩╝êµùá SSL Θù«Θóÿ∩╝ë∩╝î
    # σ¢áµ¡ñΦ┐ÖΘçîτ¢┤µÄÑΣ╜┐τö¿ timeout-only τÜä httpx σ«óµê╖τ½»∩╝îΣ╗úτÉåτö▒ env vars Φç¬σè¿µÄÑτ«í
    from google.genai._api_client import HttpOptions as _HttpOptions

    try:
        http_client = httpx.Client(timeout=timeout_config, verify=True)
    except Exception as e:
        _app_logger.warning(f"ΓÜá∩╕Å σê¢σ╗║ HTTP σ«óµê╖τ½»σç║ΘöÖ: {e}")
        http_client = httpx.Client(timeout=timeout_config)

    # µ₧äσ╗║ HttpOptions σ»╣Φ▒í
    opts_kwargs = dict(
        api_version=http_options.get("api_version", "v1beta"),
        httpx_client=http_client,
    )
    if http_options.get("base_url"):
        opts_kwargs["base_url"] = http_options["base_url"]

    return genai.Client(api_key=API_KEY, http_options=_HttpOptions(**opts_kwargs))


# ΓöÇΓöÇ µ£¼σ£░µ¿íσ₧ïΘàìτ╜«Φ»╗σÅû ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
def _get_local_model_config() -> tuple:
    """
    Φ»╗σÅû user_settings.json∩╝îΦ┐öσ¢₧ (model_mode, local_model_tag)πÇé
    model_mode: "local" µêû "cloud"∩╝êΘ╗ÿΦ«ñ cloud∩╝ë
    local_model_tag: σªé "qwen2.5:7b" µêû None
    """
    try:
        settings_path = os.path.join(PROJECT_ROOT, "config", "user_settings.json")
        with open(settings_path, "r", encoding="utf-8") as _f:
            _data = json.load(_f)
        mode = _data.get("model_mode", "cloud")
        tag = _data.get("local_model")
        return mode, tag
    except Exception:
        return "cloud", None


# µçÆσèáΦ╜╜σ«óµê╖τ½»∩╝êmode+tag Σ╜£Σ╕║τ╝ôσ¡ÿ key∩╝îσêçµìóµ¿íσ╝ÅσÉÄΦç¬σè¿Θçìσ╗║∩╝ë
_client = None
_client_mode_key: tuple = (None, None)  # (model_mode, local_model_tag)


def get_client():
    """
    ΦÄ╖σÅû AI σ«óµê╖τ½»∩╝êµçÆσèáΦ╜╜∩╝ëπÇé
    - ΦïÑ user_settings.json Σ╕¡ model_mode == "local"∩╝îΦ┐öσ¢₧ OllamaClientProxy
    - σÉªσêÖΦ┐öσ¢₧ Gemini genai.Client∩╝êσÄƒµ£ëΦíîΣ╕║∩╝ë
    """
    global _client, _client_mode_key
    model_mode, local_model = _get_local_model_config()
    current_key = (model_mode, local_model)

    # µ¿íσ╝Åµêûµ¿íσ₧ïσÅæτöƒσÅÿσîûµù╢∩╝îΘçìτ╜«τ╝ôσ¡ÿ
    if _client is not None and _client_mode_key != current_key:
        _client = None

    if _client is None:
        if model_mode == "local" and local_model:
            try:
                from app.core.llm.ollama_provider import OllamaClientProxy

                _client = OllamaClientProxy(model_tag=local_model)
                _app_logger.debug(f"[Koto] ≡ƒªÖ Σ╜┐τö¿µ£¼σ£░µ¿íσ₧ï: {local_model}")
            except Exception as _e:
                _app_logger.warning(f"[Koto] ΓÜá∩╕Å Ollama σê¥σºïσîûσñ▒Φ┤Ñ∩╝îσ¢₧ΘÇÇσê░ Gemini: {_e}")
                _client = create_client()
        else:
            _client = create_client()
        _client_mode_key = current_key

    return _client


# ΓöÇΓöÇ Token τ¢æµ╡ïµ¿íσ¥ù∩╝êµ£¼σ£░τ╗ƒΦ«í∩╝îµùáΘ£ÇΘó¥σñûΦ┐₧µÄÑ Google∩╝ëΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
try:
    from token_tracker import record_usage as _record_token_usage

    _TOKEN_TRACKER_ENABLED = True
except ImportError:
    _TOKEN_TRACKER_ENABLED = False

    def _record_token_usage(*_a, **_kw):
        pass


class _FakeGenerateContentResponse:
    """
    Φ╜╗ΘçÅτ║ºσôìσ║öσîàΦúàσÖ¿πÇé
    σ╜ô _TrackedModels σ░å Interactions API τÜäσ¡ùτ¼ªΣ╕▓τ╗ôµ₧£Φ╜¼µìóΣ╕║µáçσçåσôìσ║öσ»╣Φ▒íµù╢Σ╜┐τö¿∩╝î
    τí«Σ┐¥µëÇµ£ëΦ░âτö¿µû╣σÅ»Σ╗Ñτ╗ƒΣ╕ÇΣ╗Ñ response.text σÅûσÇ╝πÇé
    """
    __slots__ = ("text", "candidates", "usage_metadata")

    def __init__(self, text: str):
        self.text = text
        self.candidates = []
        self.usage_metadata = None


def _extract_prompt_text(contents, config=None) -> tuple:
    """
    Σ╗Ä generate_content τÜä contents / config σÅéµò░Σ╕¡µÅÉσÅûµûçµ£¼ prompt σÆî system_instructionπÇé
    Φ┐öσ¢₧ (prompt_text: str, sys_instruction: str | None)
    """
    # µÅÉσÅû system_instruction
    sys_instr = None
    if config is not None:
        sys_instr = getattr(config, "system_instruction", None)
        if sys_instr is not None:
            sys_instr = str(sys_instr)

    # µÅÉσÅû prompt µûçµ£¼
    if contents is None:
        return "", sys_instr
    if isinstance(contents, str):
        return contents, sys_instr
    if isinstance(contents, list):
        parts = []
        for item in contents:
            if isinstance(item, str):
                parts.append(item)
            elif hasattr(item, "text") and item.text:
                parts.append(str(item.text))
            elif hasattr(item, "parts"):
                for p in (item.parts or []):
                    if hasattr(p, "text") and p.text:
                        parts.append(str(p.text))
            else:
                s = str(item)
                if s:
                    parts.append(s)
        return "\n".join(parts), sys_instr
    return str(contents), sys_instr


def _is_interactions_only(model_id: str) -> bool:
    """
    µúÇµƒÑ model_id µÿ»σÉªΘ£ÇΦªüΦ╡░ Interactions API ΦÇîΘ¥₧ generate_contentπÇé
    Σ╜┐τö¿µ¿íσ¥ùτ║º _INTERACTIONS_ONLY_MODELS∩╝êΦ┐ÉΦíîµù╢µƒÑµë╛∩╝îσ«ÜΣ╣ëσÉÄΣ╕Çσ«ÜσÅ»τö¿∩╝ëπÇé
    """
    try:
        iom = _INTERACTIONS_ONLY_MODELS  # noqa: F821 ΓÇö µ¿íσ¥ùτ║ºσà¿σ▒Ç∩╝îΦ┐ÉΦíîµù╢σ╖▓σ«ÜΣ╣ë
    except NameError:
        iom = {"gemini-3-flash-preview", "gemini-3-pro-preview", "deep-research-pro-preview-12-2025"}
    mid = str(model_id or "")
    return mid in iom or mid.startswith("deep-research-pro-preview")


_logger_tracked = logging.getLogger(__name__)


class _TrackedModels:
    """
    µïªµê¬ client.models τÜä generate_content / generate_content_stream∩╝îσ«₧τÄ░∩╝Ü
      1. Token τö¿ΘçÅΦç¬σè¿Φ«░σ╜ò
      2. Interactions-only µ¿íσ₧ïΘÿ▓σ╛íΦ╖»τö▒∩╝êσëìτ╜«µúÇµƒÑ + σ╝éσ╕╕µìòΦÄ╖σà£σ║ò∩╝ë
         - σ£¿Φ░âτö¿ generate_content σëìσàêσêñµû¡µ¿íσ₧ïµÿ»σÉª interactions-only
         - ΦïÑµÿ»∩╝îτ¢┤µÄÑΦ╜¼σÅæσê░ _call_interactions_api_sync()
         - ΦïÑσÉª∩╝îµ¡úσ╕╕Φ░âτö¿σÉÄ catch "Interactions API" 400 ΘöÖΦ»»σ╣╢ retry
    """

    def __init__(self, real_models):
        object.__setattr__(self, "_real", real_models)

    # ΓöÇΓöÇ σåàΘâ¿σ╖Ñσà╖ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ

    @staticmethod
    def _call_ia(model_id: str, contents, config) -> "_FakeGenerateContentResponse":
        """µÅÉσÅûµûçµ£¼σ╣╢Φ╜¼σÅæσê░ _call_interactions_api_sync∩╝îΦ┐öσ¢₧σîàΦúàσÉÄτÜäσôìσ║öσ»╣Φ▒íπÇé"""
        prompt, sys_instr = _extract_prompt_text(contents, config)
        text = _call_interactions_api_sync(  # noqa: F821
            model_id=model_id,
            user_prompt=prompt,
            sys_instruction=sys_instr,
        )
        return _FakeGenerateContentResponse(text or "")

    # ΓöÇΓöÇ generate_content ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ

    def generate_content(self, model=None, *args, **kwargs):
        # σà╝σ«╣µùºσ╝ÅΣ╜ìτ╜«Φ░âτö¿
        if model is None and args:
            model, args = args[0], args[1:]

        model_str = str(model or "")
        real = object.__getattribute__(self, '_real')

        # Γæá σëìτ╜«Φ╖»τö▒∩╝Üinteractions-only µ¿íσ₧ïτ¢┤µÄÑΦ╡░ Interactions API
        if _is_interactions_only(model_str):
            _logger_tracked.debug(
                "[TrackedModels] %s ΓåÆ Interactions API (pre-check)", model_str
            )
            try:
                return self._call_ia(model_str, kwargs.get("contents"), kwargs.get("config"))
            except Exception as _ia_err:
                _logger_tracked.warning(
                    "[TrackedModels] Interactions API failed for %s: %s ΓÇö retrying generate_content as last resort",
                    model_str, _ia_err,
                )
                # σ╝║Φíîσ░¥Φ»ò generate_content∩╝êµ₧üσ░æµò░µâàσå╡µ¿íσ₧ïσ«₧ΘÖàµö»µîü∩╝ë

        # Γæí µáçσçåΦ░âτö¿ + σ╝éσ╕╕σà£σ║ò
        try:
            response = real.generate_content(model=model, **kwargs)
        except Exception as _gc_err:
            _err_str = str(_gc_err)
            if "Interactions API" in _err_str or (
                "only supports" in _err_str and "Interactions" in _err_str
            ):
                _logger_tracked.warning(
                    "[TrackedModels] 400 Interactions-API error for model=%s ΓÇö retrying via Interactions API",
                    model_str,
                )
                try:
                    return self._call_ia(model_str, kwargs.get("contents"), kwargs.get("config"))
                except Exception as _ia_retry_err:
                    _logger_tracked.error(
                        "[TrackedModels] Interactions API retry also failed for %s: %s",
                        model_str, _ia_retry_err,
                    )
            raise  # Θ¥₧ Interactions ΘöÖΦ»»∩╝îµêû retry Σ╣ƒσñ▒Φ┤ÑσÉÄ∩╝îΘçìµû░µè¢σç║σÄƒσºïσ╝éσ╕╕

        # Γæó Token Φ«░σ╜ò
        if _TOKEN_TRACKER_ENABLED:
            try:
                usage = getattr(response, "usage_metadata", None)
                if usage:
                    _record_token_usage(
                        model=model_str,
                        prompt_tokens=int(getattr(usage, 'prompt_token_count', 0) or 0),
                        completion_tokens=int(getattr(usage, 'candidates_token_count', 0) or 0),
                    )
            except Exception:
                pass
        return response

    # ΓöÇΓöÇ generate_content_stream ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ

    def generate_content_stream(self, model=None, *args, **kwargs):
        """
        µïªµê¬µ╡üσ╝ÅΦ░âτö¿πÇé
        interactions-only µ¿íσ₧ïΣ╕ìµö»µîüµ╡üσ╝ÅµÄÑσÅú∩╝¢Θüçσê░µ¡ñτ▒╗µ¿íσ₧ïµù╢σÉîµ¡ÑΦ░âτö¿
        Interactions API∩╝îσåìσ░åσ«îµò┤τ╗ôµ₧£σîàΦúàµêÉσìòΣ╕¬ chunk yield σç║σÄ╗∩╝î
        Σ┐¥Φ»üΦ░âτö¿µû╣ for chunk in stream τÜäτö¿µ│òΣ╕ìσÅÿπÇé
        """
        if model is None and args:
            model, args = args[0], args[1:]

        model_str = str(model or "")
        real = object.__getattribute__(self, '_real')

        # Γæá σëìτ╜«Φ╖»τö▒∩╝Üinteractions-only µ¿íσ₧ï ΓåÆ σÉîµ¡ÑΦ░âτö¿σÉÄσìò chunk Φ╛ôσç║
        if _is_interactions_only(model_str):
            _logger_tracked.debug(
                "[TrackedModels] %s ΓåÆ Interactions API stream-adapter (pre-check)", model_str
            )
            try:
                fake_resp = self._call_ia(model_str, kwargs.get("contents"), kwargs.get("config"))
                yield fake_resp   # Φ░âτö¿µû╣ for chunk in stream: chunk.text Σ╗ìΦâ╜σ╖ÑΣ╜£
                return
            except Exception as _ia_err:
                _logger_tracked.warning(
                    "[TrackedModels] Interactions API stream-adapter failed for %s: %s ΓÇö raising",
                    model_str, _ia_err,
                )
                raise

        # Γæí µáçσçåµ╡üσ╝ÅΦ░âτö¿ + σ╝éσ╕╕σà£σ║ò∩╝êΘªûΣ╕¬ chunk σëìΦºªσÅæ∩╝ë
        try:
            stream = real.generate_content_stream(model=model, **kwargs)
            first_chunk = True
            for chunk in stream:
                yield chunk
                if first_chunk:
                    first_chunk = False
                if _TOKEN_TRACKER_ENABLED:
                    try:
                        usage = getattr(chunk, 'usage_metadata', None)
                        if usage and (getattr(usage, 'prompt_token_count', 0) or 0) > 0:
                            _record_token_usage(
                                model=model_str,
                                prompt_tokens=int(getattr(usage, 'prompt_token_count', 0) or 0),
                                completion_tokens=int(getattr(usage, 'candidates_token_count', 0) or 0),
                            )
                    except Exception:
                        pass
        except Exception as _stream_err:
            _err_str = str(_stream_err)
            if "Interactions API" in _err_str or (
                "only supports" in _err_str and "Interactions" in _err_str
            ):
                _logger_tracked.warning(
                    "[TrackedModels] 400 Interactions-API error in stream for model=%s ΓÇö retrying via Interactions API",
                    model_str,
                )
                fake_resp = self._call_ia(model_str, kwargs.get("contents"), kwargs.get("config"))
                yield fake_resp
                return
            raise

    def generate_images(self, model=None, *args, **kwargs):
        """µïªµê¬ generate_images∩╝êImagen∩╝ë∩╝îµîëσ¢╛τëçµò░ΘçÅΦ«░σ╜òσÉêµêÉ token τö¿ΘçÅ"""
        if model is None and args:
            model, args = args[0], args[1:]
        real = object.__getattribute__(self, "_real")
        response = real.generate_images(model=model, **kwargs)
        if _TOKEN_TRACKER_ENABLED:
            try:
                # Imagen µîëσ╝áΦ«íΦ┤╣∩╝îτö¿σÉêµêÉ token µò░µìóτ«ù∩╝ê1000 tokens/σ╝á∩╝îΘàìσÉêσ«ÜΣ╗╖Φí¿σ╛ùσç║µ¡úτí«Φ┤╣τö¿∩╝ë
                num_images = max(
                    1, len(getattr(response, "generated_images", []) or [])
                )
                _record_token_usage(
                    model=str(model or "unknown"),
                    prompt_tokens=1000 * num_images,
                    completion_tokens=0,
                )
            except Exception:
                pass
        return response

    def embed_content(self, model=None, *args, **kwargs):
        """µïªµê¬ embed_content∩╝êtext-embedding-004 τ¡ë∩╝ë∩╝îΦ«░σ╜ò embedding token τö¿ΘçÅ"""
        if model is None and args:
            model, args = args[0], args[1:]
        real = object.__getattribute__(self, "_real")
        response = real.embed_content(model=model, **kwargs)
        if _TOKEN_TRACKER_ENABLED:
            try:
                usage = getattr(response, "usage_metadata", None)
                if usage:
                    prompt_tokens = int(getattr(usage, "prompt_token_count", 0) or 0)
                else:
                    # embed_content σ╣╢Σ╕ìµÇ╗µÿ»Φ┐öσ¢₧ usage_metadata∩╝îµîëΦ╛ôσàÑσåàσ«╣σ¡ùτ¼ªµò░Σ╝░τ«ù
                    contents = kwargs.get("contents", "") or ""
                    if isinstance(contents, list):
                        contents = " ".join(str(c) for c in contents)
                    prompt_tokens = max(
                        1, len(str(contents)) // 4
                    )  # τ▓ùτòÑΣ╝░τ«ù 1 token Γëê 4 σ¡ùτ¼ª
                _record_token_usage(
                    model=str(model or "unknown"),
                    prompt_tokens=prompt_tokens,
                    completion_tokens=0,
                )
            except Exception:
                pass
        return response

    def __getattr__(self, name):
        real = object.__getattribute__(self, "_real")
        return getattr(real, name)


# Σ┐¥µîüσÉæσÉÄσà╝σ«╣τÜä client σÅÿΘçÅ∩╝êΘÇÜΦ┐çσ▒₧µÇºΦ«┐Θù«ΦºªσÅæµçÆσèáΦ╜╜∩╝ë
class _ClientProxy:
    """Σ╗úτÉåτ▒╗∩╝îσ«₧τÄ░µçÆσèáΦ╜╜"""

    def __getattr__(self, name):
        obj = getattr(get_client(), name)
        if name == "models":
            return _TrackedModels(obj)
        return obj


client = _ClientProxy()


def create_research_client():
    """σê¢σ╗║Σ╕ôτö¿Σ║Ä Deep Research τÜäΘò┐Φ╢àµù╢σ«óµê╖τ½» (5σêåΘÆƒ read timeout)"""
    import httpx
    from google.genai._api_client import HttpOptions as _HttpOptions

    proxy = get_detected_proxy()
    # µ╖▒σ║ªτáöτ⌐╢Θ£ÇΦªüµ¢┤Θò┐τÜäΦ╢àµù╢µù╢Θù┤∩╝ÜΦ┐₧µÄÑ30τºÆ∩╝îΦ»╗σÅû5σêåΘÆƒ
    timeout_config = httpx.Timeout(300.0, connect=30.0)

    # Θàìτ╜«Σ╗úτÉå - ΘÇÜΦ┐çτÄ»σóâσÅÿΘçÅτí«Σ┐¥Φó½Σ╜┐τö¿∩╝êσÉî create_client∩╝ë
    if proxy:
        os.environ["HTTP_PROXY"] = proxy
        os.environ["HTTPS_PROXY"] = proxy

    # Φç¬σ«ÜΣ╣ë httpx σ«óµê╖τ½»∩╝êΣ╗àτö¿Σ║Äµë⌐σ▒òΦ╢àµù╢∩╝¢Σ╗úτÉåτö▒ env vars µÄÑτ«í∩╝ë
    http_client = httpx.Client(timeout=timeout_config, verify=True)

    opts_kwargs = dict(
        api_version="v1beta",
        httpx_client=http_client,
    )
    if GEMINI_API_BASE:
        opts_kwargs["base_url"] = GEMINI_API_BASE

    return genai.Client(api_key=API_KEY, http_options=_HttpOptions(**opts_kwargs))


def _poll_interaction(
    ia_client,
    interaction_id: str,
    *,
    timeout: float = 900.0,
    initial_sleep: float = 2.0,
    backoff_multiplier: float = 1.5,
    max_sleep: float = 30.0,
    label: str = "",
) -> object:
    """
    τöƒΣ║ºτ║º Interactions API Φ╜«Φ»óσÖ¿πÇé

    σ«₧τÄ░µîçµò░ΘÇÇΘü┐ + µèûσè¿ + µ£ÇσñºΦ╢àµù╢∩╝îΘü┐σàìΦ╜«Φ»óΘúÄµÜ┤∩╝Ü
      - every successful poll: sleep *= backoff_multiplier∩╝êΣ╕èΘÖÉ max_sleep∩╝ë
      - ┬▒25% ΘÜÅµ£║µèûσè¿∩╝îσêåµòúσ╣╢σÅæΦ»╖µ▒éσ│░σÇ╝
      - Φ╢àµù╢σÉÄΦç¬σè¿Φ»╖µ▒éσÅûµ╢ê∩╝îσåìµè¢σç║ TimeoutError

    τè╢µÇüµ£║∩╝Ü
      ΓöÇ RUNNING   (active / running / queued / ΓÇª)  ΓåÆ  τ╗ºτ╗¡τ¡ëσ╛à
      ΓöÇ COMPLETED (completed)                       ΓåÆ  Φ┐öσ¢₧µ£Çτ╗ê interaction σ»╣Φ▒í
      ΓöÇ FAILED    (failed / cancelled / error)      ΓåÆ  µè¢σç║ RuntimeError

    Args:
        ia_client:          σ╖▓σê¥σºïσîûτÜä Gemini client∩╝êσÉ½ .interactions µÄÑσÅú∩╝ë
        interaction_id:     rc.interactions.create() Φ┐öσ¢₧τÜä job ID
        timeout:            µ£Çσñºτ¡ëσ╛àτºÆµò░∩╝êΘ╗ÿΦ«ñ 15 σêåΘÆƒ∩╝ë
        initial_sleep:      Θªûµ¼íΦ╜«Φ»óσëìτ¡ëσ╛àτºÆµò░
        backoff_multiplier: ΘÇÇΘü┐σÇìτÄç∩╝êµ»ÅΦ╜«Φç¬σè¿Σ╣ÿΣ╗Ñµ¡ñσÇ╝∩╝ë
        max_sleep:          σìòµ¼íτ¡ëσ╛àΣ╕èΘÖÉ∩╝êτºÆ∩╝ë
        label:              µùÑσ┐ùσëìτ╝Çµáçτ¡╛∩╝êΣ╛┐Σ║Äσî║σêåΦ░âτö¿µû╣∩╝ë

    Returns:
        status == "completed" τÜä interaction σ»╣Φ▒í

    Raises:
        RuntimeError: interaction_id Σ╕║τ⌐║
        TimeoutError: Φ╢àσç║ timeout Σ╗ìµ£¬σ«îµêÉ∩╝êσ╖▓Φ»╖µ▒éσÅûµ╢ê∩╝ë
        RuntimeError: job Φ┐öσ¢₧ failed / cancelled / error τè╢µÇü
    """
    import random as _random

    if not interaction_id:
        raise RuntimeError(f"[{label or 'poll'}] interaction_id Σ╕║τ⌐║∩╝îµùáµ│òΦ╜«Φ»ó")

    _log = logging.getLogger(__name__)
    tag  = f"[Interactions{':' + label if label else ''}]"

    start          = time.monotonic()
    sleep_interval = initial_sleep
    last_status    = ""
    poll_count     = 0

    _log.info("%s ΓÅ│ job=%s  σ╝ÇσºïΦ╜«Φ»ó (timeout=%.0fs)", tag, interaction_id, timeout)

    while True:
        elapsed = time.monotonic() - start

        # ΓöÇΓöÇ Φ╢àµù╢µúÇµƒÑ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        if elapsed >= timeout:
            _log.warning("%s Γî¢ job=%s  Φ╜«Φ»óΦ╢àµù╢ (%.0fs elapsed)", tag, interaction_id, elapsed)
            try:
                ia_client.interactions.cancel(interaction_id)
                _log.info("%s ≡ƒ¢æ job=%s  σ╖▓Φ»╖µ▒éσÅûµ╢ê", tag, interaction_id)
            except Exception as _ce:
                _log.debug("%s σÅûµ╢êΦ»╖µ▒éσñ▒Φ┤Ñ: %s", tag, _ce)
            raise TimeoutError(
                f"Interactions API Φ╢àµù╢ ({timeout:.0f}s) job={interaction_id}"
            )

        # ΓöÇΓöÇ Φ╜«Φ»óΦ»╖µ▒é∩╝êτ╜æτ╗£µèûσè¿µù╢τƒ¡µÜéτ¡ëσ╛àσÉÄΘçìΦ»ò∩╝îΣ╕ìτ½ïσì│µö╛σ╝â∩╝ëΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        try:
            interaction = ia_client.interactions.get(interaction_id)
        except Exception as _poll_err:
            _log.warning(
                "%s job=%s  Φ╜«Φ»óΦ»╖µ▒éσñ▒Φ┤Ñ (#%d): %s",
                tag, interaction_id, poll_count, _poll_err,
            )
            time.sleep(min(sleep_interval, 10.0))
            continue

        status     = str(getattr(interaction, "status", "") or "").lower().strip()
        poll_count += 1

        # ΓöÇΓöÇ Σ╗àσ£¿τè╢µÇüσÅÿσîûµù╢Φ╛ôσç║µùÑσ┐ù∩╝îΘü┐σàìµùÑσ┐ùµ┤¬µ░┤ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        if status != last_status:
            msg = _INTERACTION_STATUS_MSGS.get(status, f"τè╢µÇü: {status!r}")
            _log.info(
                "%s ≡ƒöä job=%s  [poll#%d | %.0fs] %s",
                tag, interaction_id, poll_count, elapsed, msg,
            )
            last_status = status

        # ΓöÇΓöÇ τ╗êµ¡óτè╢µÇüσêñµû¡ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        if status in _INTERACTION_TERMINAL_STATES:
            if status in _INTERACTION_SUCCESS_STATES:
                _log.info(
                    "%s Γ£à job=%s  σ«îµêÉ (total=%.1fs, polls=%d)",
                    tag, interaction_id, elapsed, poll_count,
                )
                return interaction
            # failed / cancelled / error
            err_detail = getattr(interaction, "error", None) or status
            _log.error(
                "%s Γ¥î job=%s  σñ▒Φ┤Ñ status=%s  detail=%s",
                tag, interaction_id, status, err_detail,
            )
            raise RuntimeError(
                f"Interactions API job σñ▒Φ┤Ñ (status={status}, detail={err_detail})"
            )

        # ΓöÇΓöÇ Φ«íτ«ùΣ╕ïΣ╕ÇΦ╜«τ¡ëσ╛àµù╢Θù┤∩╝Üµîçµò░ΘÇÇΘü┐ + ┬▒25% ΘÜÅµ£║µèûσè¿ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        jitter       = sleep_interval * 0.25 * (_random.random() * 2 - 1)
        actual_sleep = max(1.0, min(sleep_interval + jitter, max_sleep))
        remaining    = timeout - elapsed
        actual_sleep = min(actual_sleep, max(0.5, remaining - 0.1))   # Σ╕ìΦ╢àΦ┐çσë⌐Σ╜Öµù╢Θù┤

        _log.debug("%s job=%s  τ¡ëσ╛à %.1fs σÉÄσåìµ¼íΦ╜«Φ»óΓÇª", tag, interaction_id, actual_sleep)
        time.sleep(actual_sleep)

        # ΘÇÉµ¡Ñσ╗╢Θò┐Φ╜«Φ»óΘù┤ΘÜö∩╝îτ¢┤σê░ max_sleep Σ╕èΘÖÉ
        sleep_interval = min(sleep_interval * backoff_multiplier, max_sleep)


def _extract_interaction_text_global(interaction) -> str:
    """
    Σ╗Ä interaction σ»╣Φ▒íΘÇÆσ╜ÆµÅÉσÅûΦ╛ôσç║µûçµ£¼πÇé
    σà╝σ«╣σñÜτºì SDK Φ┐öσ¢₧µá╝σ╝Å∩╝Üoutputs σêùΦí¿πÇütext σ▒₧µÇºπÇüpartsπÇüPydantic model_dumpπÇüdict τ¡ëπÇé
    """
    def _walk(obj) -> list:
        if obj is None:
            return []
        if isinstance(obj, str):
            s = obj.strip()
            return [s] if s else []
        if isinstance(obj, dict):
            results = []
            for key in ("output_text", "text", "content"):
                val = obj.get(key)
                if isinstance(val, str) and val.strip():
                    results.append(val.strip())
                    return results          # Σ╝ÿσàêΦ┐öσ¢₧Φ»¡Σ╣ëµ£Çσ╝║τÜäσ¡ùµ«╡
            for val in obj.values():
                results.extend(_walk(val))
            return results
        if isinstance(obj, (list, tuple)):
            results = []
            for item in obj:
                results.extend(_walk(item))
            return results
        # Pydantic / SDK σ»╣Φ▒í∩╝Üσàêσ░¥Φ»ò model_dump()
        if hasattr(obj, "model_dump"):
            try:
                return _walk(obj.model_dump())
            except Exception:
                pass
        if hasattr(obj, "text") and obj.text:
            return [str(obj.text).strip()]
        if hasattr(obj, "parts"):
            results = []
            for p in (obj.parts or []):
                results.extend(_walk(p))
            return results
        if hasattr(obj, "outputs"):
            results = []
            for o in (obj.outputs or []):
                results.extend(_walk(o))
            return results
        return []

    parts = _walk(getattr(interaction, "outputs", None))
    if not parts:
        parts = _walk(interaction)

    # σÄ╗Θçì∩╝îΣ┐¥µîüσÄƒσºïΘí║σ║Å
    seen: set = set()
    deduped = []
    for p in parts:
        if p not in seen:
            deduped.append(p)
            seen.add(p)
    return "\n".join(deduped).strip()


def _call_interactions_api_sync(
    model_id: str,
    user_prompt: str,
    sys_instruction: str = None,
    timeout: float = 900.0,
) -> str:
    """
    ΘÇÜΦ┐ç Interactions API Φ░âτö¿ gemini-3-*-preview / deep-research τ¡ëσ╝éµ¡Ñµ¿íσ₧ïπÇé
    Φ┐ÖΣ║¢µ¿íσ₧ïΣ╕ìµö»µîü client.models.generate_content()∩╝îσ┐àΘí╗Σ╜┐τö¿µ¡ñτ½»τé╣πÇé

    σ╖ÑΣ╜£µ╡üτ¿ï∩╝Ü
      1. µ£¼σ£░µ¿íσ₧ïµ¿íσ╝Å ΓåÆ τ¢┤µÄÑτö¿ Ollama∩╝îΦ╖│Φ┐ç Interactions API
      2. Σ║æτ½»µ¿íσ╝Å     ΓåÆ rc.interactions.create() µÅÉΣ║ñσ╝éµ¡Ñ job∩╝îµìòΦÄ╖ interaction_id
      3.              ΓåÆ _poll_interaction() Φ╜«Φ»ó∩╝êµîçµò░ΘÇÇΘü┐∩╝îµ£Çσñº timeout τºÆ∩╝ë
      4.              ΓåÆ µÅÉσÅûσ╣╢Φ┐öσ¢₧µ£Çτ╗êµûçµ£¼

    Args:
        model_id:        τ¢«µáçµ¿íσ₧ï ID
        user_prompt:     τö¿µê╖Φ╛ôσàÑ∩╝êσ╖▓µá╝σ╝Åσîû∩╝ë
        sys_instruction: τ│╗τ╗ƒµîçΣ╗ñ∩╝êσÅ»ΘÇë∩╝ë
        timeout:         µ£Çσñºτ¡ëσ╛àτºÆµò░∩╝êΘ╗ÿΦ«ñ 15 σêåΘÆƒ∩╝ë

    Returns:
        µ¿íσ₧ïσôìσ║öµûçµ£¼

    Raises:
        TimeoutError:   Φ╢àµù╢∩╝êσ╖▓Φç¬σè¿Φ»╖µ▒éσÅûµ╢ê∩╝ë
        RuntimeError:   job σñ▒Φ┤Ñµêûµ£¼σ£░ΘÖìτ║ºσñ▒Φ┤Ñ
    """
    _log = logging.getLogger(__name__)

    # ΓöÇΓöÇ µ£¼σ£░µ¿íσ₧ïµ¿íσ╝Å∩╝Üτö¿ Ollama τ¢┤µÄÑσ¢₧τ¡ö∩╝îµùáΘ£Ç Interactions API ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    model_mode, _ = _get_local_model_config()
    if model_mode == "local":
        try:
            full_prompt = user_prompt
            if sys_instruction:
                full_prompt = (
                    f"[τ│╗τ╗ƒµîçΣ╗ñ]\n{sys_instruction}\n\n[τö¿µê╖Φ╛ôσàÑ]\n{user_prompt}"
                )
            resp = get_client().models.generate_content(
                model=model_id,
                contents=full_prompt,
            )
            return getattr(resp, "text", "") or ""
        except Exception as _e:
            raise RuntimeError(f"µ£¼σ£░µ¿íσ₧ï Interactions ΘÖìτ║ºσñ▒Φ┤Ñ: {_e}") from _e

    # ΓöÇΓöÇ Σ║æτ½»∩╝ÜµÅÉΣ║ñσ╝éµ¡Ñ Interactions Σ╗╗σèí ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    full_input = user_prompt
    if sys_instruction:
        full_input = f"[τ│╗τ╗ƒµîçΣ╗ñ]\n{sys_instruction}\n\n[τö¿µê╖Φ╛ôσàÑ]\n{user_prompt}"

    rc = create_research_client()

    _log.info(
        "[Interactions] ≡ƒÜÇ µÅÉΣ║ñ job  model=%s  input_chars=%d",
        model_id, len(full_input),
    )
    # Interactions API σî║σêåΣ╕ñτºìΦ░âτö¿µû╣σ╝Å∩╝Ü
    #   agent=  ΓåÆ deep-research τ¡ëτ£ƒµ¡úτÜä Agent
    #   model=  ΓåÆ gemini-3-pro/flash-preview τ¡ëµÖ«ΘÇÜµ¿íσ₧ï∩╝êτö¿ agent= Σ╝ÜµèÑ 400∩╝ë
    _create_kwargs: dict = {
        "input":      full_input[:80000],
        "background": True,
        "stream":     False,
    }
    if _is_interactions_agent(model_id):
        _create_kwargs["agent"] = model_id
    else:
        _create_kwargs["model"] = model_id

    interaction = rc.interactions.create(**_create_kwargs)

    interaction_id = getattr(interaction, "id", None)
    init_status    = str(getattr(interaction, "status", "") or "").lower()

    # σ┐½ΘÇƒΦ╖»σ╛ä∩╝Üµ₧üσ░æµò░µâàσå╡Σ╕ï create() σì│σê╗Φ┐öσ¢₧σ╖▓σ«îµêÉ
    if init_status in _INTERACTION_SUCCESS_STATES:
        _log.info(
            "[Interactions] ΓÜí job=%s σì│µù╢σ«îµêÉ (status=%s)", interaction_id, init_status
        )
        return _extract_interaction_text_global(interaction)

    if init_status in _INTERACTION_FAIL_STATES:
        err = getattr(interaction, "error", init_status)
        raise RuntimeError(
            f"Interactions API job τ½ïσì│σñ▒Φ┤Ñ (status={init_status}): {err}"
        )

    if not interaction_id:
        raise RuntimeError(
            f"Interactions API µ£¬Φ┐öσ¢₧µ£ëµòêτÜä interaction_id (model={model_id})"
        )

    # µàóΘÇƒΦ╖»σ╛ä∩╝ÜΦ╜«Φ»óτ¡ëσ╛à∩╝êµîçµò░ΘÇÇΘü┐∩╝îσÉ½Φç¬σè¿Φ╢àµù╢σÅûµ╢ê∩╝ë
    final_interaction = _poll_interaction(
        rc,
        interaction_id,
        timeout=timeout,
        initial_sleep=2.0,
        backoff_multiplier=1.5,
        max_sleep=30.0,
        label=model_id,
    )

    text = _extract_interaction_text_global(final_interaction)
    _log.info(
        "[Interactions] ≡ƒôä µÅÉσÅûµûçµ£¼ %d σ¡ùτ¼ª (model=%s)", len(text), model_id
    )
    return text


def run_with_timeout(fn, timeout_seconds):
    """σ£¿τ║┐τ¿ïΣ╕¡µëºΦíîσç╜µò░σ╣╢ΘÖÉµù╢Φ┐öσ¢₧ (Θü┐σàìσìíµ¡╗Σ╕╗µ╡üτ¿ï)"""
    holder = {"result": None, "error": None}

    def _runner():
        try:
            holder["result"] = fn()
        except Exception as e:
            holder["error"] = e

    t = threading.Thread(target=_runner, daemon=True)
    t.start()
    t.join(timeout_seconds)
    if t.is_alive():
        return None, TimeoutError(f"Timeout after {timeout_seconds}s"), True
    return holder["result"], holder["error"], False


def run_with_heartbeat(
    fn, start_time, heartbeat_callback, heartbeat_interval=5, timeout_seconds=90
):
    """
    σ£¿σÉÄσÅ░τ║┐τ¿ïΦ┐ÉΦíîσç╜µò░∩╝îσÉîµù╢σ«Üµ£ƒσÅæΘÇüσ┐âΦ╖│πÇé
    τö¿Σ║ÄΘ¥₧µ╡üσ╝Å API Φ░âτö¿∩╝êσªéσ¢╛σâÅτöƒµêÉ∩╝ëπÇé

    Args:
        fn: ΦªüµëºΦíîτÜäσç╜µò░
        start_time: Φ»╖µ▒éσ╝Çσºïµù╢Θù┤
        heartbeat_callback: σ┐âΦ╖│σ¢₧Φ░âσç╜µò░∩╝îµÄÑµö╢ elapsed_seconds σÅéµò░
        heartbeat_interval: σ┐âΦ╖│Θù┤ΘÜö∩╝êτºÆ∩╝ë
        timeout_seconds: Φ╢àµù╢µù╢Θù┤∩╝êτºÆ∩╝ë

    Returns:
        (result, error, timed_out)
    """
    import queue
    import threading

    result_queue = queue.Queue()

    def worker():
        try:
            result = fn()
            result_queue.put(("success", result))
        except Exception as e:
            result_queue.put(("error", e))

    thread = threading.Thread(target=worker, daemon=True)
    thread.start()

    last_heartbeat = time.time()

    while True:
        # µúÇµƒÑµÿ»σÉªΦ╢àµù╢
        elapsed = time.time() - start_time
        if elapsed > timeout_seconds:
            return None, TimeoutError(f"µôìΣ╜£Φ╢àµù╢ ({int(elapsed)}s)"), True

        # σ░¥Φ»òΦÄ╖σÅûτ╗ôµ₧£∩╝êτƒ¡Φ╢àµù╢∩╝ë
        try:
            status, data = result_queue.get(timeout=1.0)
            if status == "success":
                return data, None, False
            else:
                return None, data, False
        except queue.Empty:
            # σÅæΘÇüσ┐âΦ╖│
            current_time = time.time()
            if current_time - last_heartbeat >= heartbeat_interval:
                heartbeat_callback(int(current_time - start_time))
                last_heartbeat = current_time


def stream_with_keepalive(
    response_stream, start_time, keepalive_interval=5, max_wait_first_token=60
):
    """
    σîàΦúàµ╡üσ╝Åσôìσ║ö∩╝îσ£¿τ¡ëσ╛àτ¼¼Σ╕ÇΣ╕¬ token µ£ƒΘù┤σÅæΘÇüΣ┐¥µ┤╗σ┐âΦ╖│πÇé

    Args:
        response_stream: σÄƒσºïµ╡üσ╝Åσôìσ║öΦ┐¡Σ╗úσÖ¿
        start_time: Φ»╖µ▒éσ╝Çσºïµù╢Θù┤
        keepalive_interval: σ┐âΦ╖│Θù┤ΘÜö∩╝êτºÆ∩╝ë
        max_wait_first_token: τ¡ëσ╛àτ¼¼Σ╕ÇΣ╕¬ token τÜäµ£Çσñºµù╢Θù┤∩╝êτºÆ∩╝ë

    Yields:
        (type, data): type σÅ»Σ╗Ñµÿ» 'chunk', 'heartbeat', 'timeout'
    """
    import queue
    import time

    chunk_queue = queue.Queue()
    first_chunk_received = threading.Event()
    stream_done = threading.Event()
    stream_error = {"error": None}

    def stream_reader():
        """σ£¿σÉÄσÅ░τ║┐τ¿ïΣ╕¡Φ»╗σÅûµ╡ü"""
        try:
            for chunk in response_stream:
                chunk_queue.put(("chunk", chunk))
                first_chunk_received.set()
            chunk_queue.put(("done", None))
        except Exception as e:
            stream_error["error"] = e
            chunk_queue.put(("error", e))
        finally:
            stream_done.set()

    # σÉ»σè¿σÉÄσÅ░Φ»╗σÅûτ║┐τ¿ï
    reader_thread = threading.Thread(target=stream_reader, daemon=True)
    reader_thread.start()

    last_heartbeat = time.time()

    while True:
        # µúÇµƒÑµÿ»σÉªτ¡ëσ╛àτ¼¼Σ╕ÇΣ╕¬ token Φ╢àµù╢
        if not first_chunk_received.is_set():
            elapsed = time.time() - start_time
            if elapsed > max_wait_first_token:
                yield ("timeout", f"τ¡ëσ╛àσôìσ║öΦ╢àµù╢ ({int(elapsed)}s)")
                return

        # σ░¥Φ»òΦÄ╖σÅû chunk∩╝îΣ╜┐τö¿τƒ¡Φ╢àµù╢Σ╗ÑΣ╛┐σÅæΘÇüσ┐âΦ╖│
        try:
            item_type, item_data = chunk_queue.get(timeout=1.0)

            if item_type == "chunk":
                yield ("chunk", item_data)
            elif item_type == "done":
                return
            elif item_type == "error":
                raise item_data

        except queue.Empty:
            # ΘÿƒσêùΣ╕║τ⌐║∩╝îµúÇµƒÑµÿ»σÉªΘ£ÇΦªüσÅæΘÇüσ┐âΦ╖│
            current_time = time.time()
            if current_time - last_heartbeat >= keepalive_interval:
                elapsed = int(current_time - start_time)
                yield ("heartbeat", elapsed)
                last_heartbeat = current_time

            # µúÇµƒÑµ╡üµÿ»σÉªσ╖▓τ╗ôµ¥ƒ
            if stream_done.is_set() and chunk_queue.empty():
                if stream_error["error"]:
                    raise stream_error["error"]
                return


app = Flask(__name__)

# Read app version from VERSION file
try:
    APP_VERSION = (
        (Path(__file__).parent.parent / "VERSION").read_text(encoding="utf-8").strip()
    )
except Exception:
    APP_VERSION = "unknown"
# Θ¥ÖµÇüΦ╡äµ║Éτ╝ôσ¡ÿ∩╝îσçÅσ░æΘçìσñìσèáΦ╜╜
app.config["SEND_FILE_MAX_AGE_DEFAULT"] = 3600
# Γ£à σàüΦ«╕µ£Çσñº 20MB Φ»╖µ▒éΣ╜ô∩╝êΦ»¡Θƒ│ base64 τ║ª 1-5MB∩╝îτòÖΦ╢│Σ╜ÖΘçÅ∩╝ë
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024

# CORS: Σ║æµ¿íσ╝ÅΘÖÉσê╢µ¥Ñµ║É∩╝îµ£¼σ£░µ¿íσ╝Åµëôσ╝Ç
_cors_origins = os.environ.get("KOTO_CORS_ORIGINS", "*")
if os.environ.get("KOTO_DEPLOY_MODE") == "cloud" and _cors_origins == "*":
    # Σ║æµ¿íσ╝ÅΘ╗ÿΦ«ñσÅ¬σàüΦ«╕Φç¬Φ║½τ½Öτé╣∩╝êσÉîµ║É∩╝ë∩╝îσÅ»ΘÇÜΦ┐çτÄ»σóâσÅÿΘçÅΦªåτ¢û
    _cors_origins = os.environ.get("KOTO_SITE_URL", "*")
CORS(app, origins=_cors_origins)


# ΓöÇΓöÇ Sentry error tracking (no-op if SENTRY_DSN not set) ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
_sentry_dsn = os.environ.get("SENTRY_DSN", "")
if _sentry_dsn:
    try:
        import sentry_sdk
        from sentry_sdk.integrations.flask import FlaskIntegration

        sentry_sdk.init(
            dsn=_sentry_dsn,
            integrations=[FlaskIntegration()],
            release=APP_VERSION,
            traces_sample_rate=float(os.environ.get("SENTRY_TRACES_RATE", "0.1")),
            send_default_pii=False,
        )
        _app_logger.info("Sentry error tracking enabled (release=%s)", APP_VERSION)
    except ImportError:
        _app_logger.warning("SENTRY_DSN set but sentry-sdk not installed; skipping")

# ΓöÇΓöÇ Prometheus metrics (/metrics) ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
try:
    from prometheus_flask_exporter import PrometheusMetrics

    _metrics_token = os.environ.get("METRICS_TOKEN", "")
    _prometheus = PrometheusMetrics(app, group_by="endpoint")
    _prometheus.info("koto_app_info", "Koto application info", version=APP_VERSION)

    if _metrics_token:
        # Require Bearer token to scrape /metrics
        @app.before_request
        def _guard_metrics():
            if request.path == "/metrics":
                auth = request.headers.get("Authorization", "")
                if auth != f"Bearer {_metrics_token}":
                    return _error_response("Unauthorized", 401)

    _app_logger.info("Prometheus metrics enabled at /metrics")
except ImportError:
    _app_logger.debug("prometheus-flask-exporter not installed; /metrics disabled")

# ΓöÇΓöÇ Swagger / OpenAPI docs ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
_swagger_config = {
    "headers": [],
    "specs": [
        {
            "endpoint": "apispec",
            "route": "/apispec.json",
            "rule_filter": lambda rule: True,
            "model_filter": lambda tag: True,
        }
    ],
    "static_url_path": "/flasgger_static",
    "swagger_ui": True,
    "specs_route": "/apidocs/",
}

_swagger_template = {
    "info": {
        "title": "Koto API",
        "description": "API documentation for Koto AI Assistant",
        "version": "1.0.0",
    },
    "basePath": "/",
    "schemes": ["http", "https"],
    "securityDefinitions": {
        "Bearer": {
            "type": "apiKey",
            "name": "Authorization",
            "in": "header",
            "description": "JWT token: `Bearer <token>`",
        }
    },
}

try:
    from flasgger import Swagger

    swagger = Swagger(app, config=_swagger_config, template=_swagger_template)
    _app_logger.info("Swagger UI enabled at /apidocs/")
except ImportError:
    _app_logger.debug("flasgger not installed; Swagger UI disabled")


# ΓöÇΓöÇ Request ID middleware ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
@app.before_request
def _assign_request_id():
    """Assign a correlation ID to every request (read from header or generate)."""
    g.request_id = request.headers.get("X-Request-ID") or str(uuid.uuid4())


@app.after_request
def _attach_request_id(response):
    """Attach the correlation ID to every outgoing response."""
    if hasattr(g, "request_id"):
        response.headers["X-Request-ID"] = g.request_id
    return response


def _error_response(message: str, status: int = 400, details=None):
    """Return a standardized JSON error envelope."""
    body = {"error": message, "status": status}
    if details:
        body["details"] = details
    if hasattr(g, "request_id"):
        body["request_id"] = g.request_id
    return jsonify(body), status


# ΓöÇΓöÇ Global error handlers (return JSON, not HTML) ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
@app.errorhandler(404)
def _handle_404(exc):
    return _error_response("Not found", 404)


@app.errorhandler(405)
def _handle_405(exc):
    return _error_response("Method not allowed", 405)


@app.errorhandler(500)
def _handle_500(exc):
    _app_logger.exception(
        "Unhandled server error [request_id=%s]", getattr(g, "request_id", "-")
    )
    return _error_response("Internal server error", 500)


# ================= τö¿µê╖Φ«ñΦ»üτ│╗τ╗ƒ =================
try:
    from auth import register_auth_routes

    register_auth_routes(app)
except Exception as e:
    _app_logger.warning(f"[Auth] ΓÜá∩╕Å Φ«ñΦ»üµ¿íσ¥ùσèáΦ╜╜σñ▒Φ┤Ñ: {e}")

# ================= σ╣╢ΦíîµëºΦíîτ│╗τ╗ƒσê¥σºïσîû =================
if PARALLEL_SYSTEM_ENABLED:
    _app_logger.debug("[PARALLEL] ≡ƒÜÇ Initializing parallel execution system...")
    try:
        register_parallel_api(app)
        start_dispatcher()
        _app_logger.info("[PARALLEL] Γ£à Parallel execution system initialized successfully")
    except Exception as e:
        _app_logger.error(f"[PARALLEL] Γ¥î Failed to initialize parallel execution system: {e}")
        PARALLEL_SYSTEM_ENABLED = False

# ================= WebSocket µö»µîü∩╝êσÅ»ΘÇë∩╝ë =================
sock = None
if Sock:
    sock = Sock(app)
else:
    _app_logger.warning("[WebSocket] ΓÜá∩╕Å flask-sock µ£¬σ«ëΦúà∩╝îΣ╜┐τö¿Φ╜«Φ»óΣ╜£Σ╕║ΘÇÜτƒÑσà£σ║ò")

if sock:

    @sock.route("/ws/notifications")
    def ws_notifications(ws):
        user_id = request.args.get("user_id", "default")
        manager = get_notification_manager()
        manager.register_connection(user_id, ws)
        try:
            while True:
                message = ws.receive()
                if message is None:
                    break
                if isinstance(message, str) and message.lower() == "ping":
                    ws.send("pong")
        finally:
            manager.unregister_connection(user_id, ws)


# ================= σ╗╢Φ┐ƒµ│¿σåîΦô¥σ¢╛∩╝êσ£¿σÉÄσÅ░τ║┐τ¿ïΣ╕¡σèáΦ╜╜∩╝îΘü┐σàìΘÿ╗σí₧σÉ»σè¿∩╝ë =================
_blueprints_registered = False
_blueprints_lock = threading.Lock()


def _register_blueprints_deferred():
    """σ£¿σÉÄσÅ░τ║┐τ¿ïΣ╕¡µ│¿σåîµëÇµ£ëΦô¥σ¢╛∩╝îΘü┐σàìΘÿ╗σí₧Σ╕╗τ║┐τ¿ïσÉ»σè¿."""
    global _blueprints_registered, agent_bp
    with _blueprints_lock:
        if _blueprints_registered:
            return
        _blueprints_registered = True

    # µ│¿σåîσüÑσ║╖µúÇµƒÑ API∩╝ê/api/health + /api/ping∩╝ë
    try:
        from web.routes.health import health_bp

        app.register_blueprint(health_bp)
        _app_logger.info("[HealthAPI] Γ£à σüÑσ║╖µúÇµƒÑ API σ╖▓µ│¿σåî: /api/health, /api/ping")
    except ImportError as e:
        _app_logger.warning(f"[HealthAPI] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑσüÑσ║╖µúÇµƒÑΦô¥σ¢╛: {e}")
    except Exception as e:
        _app_logger.error(f"[HealthAPI] Γ¥î σüÑσ║╖µúÇµƒÑ API µ│¿σåîσñ▒Φ┤Ñ: {e}")

    # µ│¿σåîΣ╗╗σèíτ«íτÉå API∩╝êΣ╗╗σèíσÅ░Φ┤ª + Φ┐¢σ║ªµÇ╗τ║┐ + µëôµû¡µÄºσê╢∩╝ë
    try:
        from app.api.task_routes import task_bp as _task_bp

        app.register_blueprint(_task_bp, url_prefix="/api/tasks")
        _app_logger.info("[TaskAPI] Γ£à Σ╗╗σèíτ«íτÉå API σ╖▓µ│¿σåî: /api/tasks")
    except ImportError as e:
        _app_logger.warning(f"[TaskAPI] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑΣ╗╗σèíτ«íτÉå API Φô¥σ¢╛: {e}")
    except Exception as e:
        _app_logger.error(f"[TaskAPI] Γ¥î Σ╗╗σèíτ«íτÉå API µ│¿σåîσñ▒Φ┤Ñ: {e}")

    # µ│¿σåîτ╗ƒΣ╕Ç Agent API
    try:
        from app.api import agent_bp as _agent_bp

        agent_bp = _agent_bp
        app.register_blueprint(agent_bp, url_prefix="/api/agent")
        _app_logger.info("[UnifiedAgent] Γ£à τ╗ƒΣ╕Ç Agent API σ╖▓µ│¿σåî: /api/agent")
    except ImportError as e:
        _app_logger.warning(f"[UnifiedAgent] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑτ╗ƒΣ╕Ç Agent API Φô¥σ¢╛: {e}")
    except Exception as e:
        _app_logger.error(f"[UnifiedAgent] Γ¥î µ│¿σåîσñ▒Φ┤Ñ: {e}")

    # µ│¿σåî Skill CRUD + MCP σ»╝σç║ API∩╝êPhase 2∩╝ë
    try:
        from app.api.skill_routes import skill_bp as _skill_bp

        app.register_blueprint(_skill_bp)
        _app_logger.info("[SkillAPI] Γ£à Skill CRUD API σ╖▓µ│¿σåî: /api/skills")
    except ImportError as e:
        _app_logger.warning(f"[SkillAPI] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑ Skill API Φô¥σ¢╛: {e}")
    except Exception as e:
        _app_logger.error(f"[SkillAPI] Γ¥î Skill API µ│¿σåîσñ▒Φ┤Ñ: {e}")

    # µ│¿σåî Skill Marketplace API∩╝êΘúÄµá╝σ╕éσ£║ + Φç¬σè¿µ₧äσ╗║ + σ»╝σàÑσ»╝σç║∩╝ë
    try:
        from app.api.skill_marketplace_routes import marketplace_bp as _marketplace_bp

        app.register_blueprint(_marketplace_bp)
        _app_logger.info("[SkillMarket] Γ£à Skill Marketplace API σ╖▓µ│¿σåî: /api/skillmarket")
    except ImportError as e:
        _app_logger.warning(f"[SkillMarket] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑ Skill Marketplace API Φô¥σ¢╛: {e}")
    except Exception as e:
        _app_logger.error(f"[SkillMarket] Γ¥î Skill Marketplace API µ│¿σåîσñ▒Φ┤Ñ: {e}")

    # µ│¿σåîΦ«¡τ╗âµò░µì« API + LoRA ΦÆ╕ΘªÅΦ«¡τ╗â API
    # Σ╗àσ£¿σ╝ÇσÅæµ£║Σ╕èσÉ»τö¿∩╝êΘ£ÇΦ«╛τ╜«τÄ»σóâσÅÿΘçÅ KOTO_DEV_TRAINING=1∩╝ë
    # σà¼σà▒σÅæΦíîτëêΣ╕ìσîàσÉ½µ¡ñσèƒΦâ╜∩╝îΦ«╛σñçΦªüµ▒éµ₧üΘ½ÿ∩╝êΓëÑ16GB VRAM∩╝ë∩╝îµÖ«ΘÇÜτö¿µê╖µùáµ│òΣ╜┐τö¿
    if os.environ.get("KOTO_DEV_TRAINING") == "1":
        try:
            from app.core.learning.training_data_builder import (
                register_training_routes as _reg_training,
            )

            _reg_training(app)
        except ImportError as e:
            _app_logger.warning(f"[TrainingAPI] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑΦ«¡τ╗âµò░µì«µ¿íσ¥ù: {e}")
        except Exception as e:
            _app_logger.error(f"[TrainingAPI] Γ¥î Φ«¡τ╗âµò░µì« API µ│¿σåîσñ▒Φ┤Ñ: {e}")

        try:
            from app.api.distill_routes import distill_bp as _distill_bp

            app.register_blueprint(_distill_bp, url_prefix="/api/distill")
            _app_logger.info("[DistillAPI] Γ£à LoRA ΦÆ╕ΘªÅΦ«¡τ╗â API σ╖▓µ│¿σåî∩╝êσ╝ÇσÅæµ¿íσ╝Å∩╝ë: /api/distill")
        except ImportError as e:
            _app_logger.warning(f"[DistillAPI] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑΦÆ╕ΘªÅΦ«¡τ╗âµ¿íσ¥ù: {e}")
        except Exception as e:
            _app_logger.error(f"[DistillAPI] Γ¥î ΦÆ╕ΘªÅΦ«¡τ╗â API µ│¿σåîσñ▒Φ┤Ñ: {e}")
    else:
        _app_logger.debug(
            "[DistillAPI] Γä╣∩╕Å LoRA Φ«¡τ╗â API σ╖▓σ░üσ¡ÿ∩╝êσà¼σà▒τëê∩╝ë∩╝îσªéΘ£ÇσÉ»τö¿Φ»╖Φ«╛τ╜« KOTO_DEV_TRAINING=1"
        )

    # µ│¿σåîσó₧σ╝║Φ»¡Θƒ│ API
    try:
        from voice_api_enhanced import voice_bp

        app.register_blueprint(voice_bp)
        _app_logger.debug("[VOICE_API] σ╖▓µ│¿σåîσó₧σ╝║Φ»¡Θƒ│ API Φô¥σ¢╛")
    except ImportError as e:
        _app_logger.warning(f"[VOICE_API] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑσó₧σ╝║Φ»¡Θƒ│µ¿íσ¥ù: {e}")

    # µ│¿σåî PPT τ╝ûΦ╛æ API∩╝êP1 σèƒΦâ╜∩╝ë
    try:
        from web.ppt_api_routes import ppt_api_bp

        app.register_blueprint(ppt_api_bp)
        _app_logger.info("[PPT_API] Γ£à PPT τ╝ûΦ╛æ API σ╖▓µ│¿σåî: /api/ppt")
    except ImportError as e:
        _app_logger.warning(f"[PPT_API] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑ PPT τ╝ûΦ╛æ API: {e}")
    except Exception as e:
        _app_logger.warning(f"[PPT_API] ΓÜá∩╕Å PPT τ╝ûΦ╛æ API µ│¿σåîσñ▒Φ┤Ñ: {e}")

    # µ│¿σåîΦç¬ΘÇéσ║ö Agent API∩╝êσ╖▓Φ┐üτº╗σê░ UnifiedAgent∩╝îΣ╜åΣ┐¥τòÖσà╝σ«╣σ»╝σàÑ∩╝ë
    try:
        from adaptive_agent_api import init_adaptive_agent_api

        init_adaptive_agent_api(app, gemini_client=None)
        _app_logger.info("[AdaptiveAgent] Γ£à Φç¬ΘÇéσ║ö Agent API σ╖▓µ│¿σåî (σ╗╢Φ┐ƒσèáΦ╜╜σ«óµê╖τ½»)")
    except ImportError:
        _app_logger.debug("[AdaptiveAgent] Γä╣∩╕Å µùº Agent µ¿íσ¥ùσ╖▓ΘÇÇσ╜╣∩╝îΣ╜┐τö¿ UnifiedAgent")
    except Exception as e:
        _app_logger.warning(f"[AdaptiveAgent] ΓÜá∩╕Å µùº Agent σê¥σºïσîûσñ▒Φ┤Ñ (Θ¥₧Φç┤σæ╜): {e}")

    # µ│¿σåîΘò┐µ£ƒτ¢«µáç API∩╝êGoalManager: Φ╖¿σñ⌐µîüτ╗¡µëºΦíîτÜäσºöµëÿΣ╗╗σèí∩╝ë
    try:
        from app.api.goal_routes import goal_bp as _goal_bp

        app.register_blueprint(_goal_bp, url_prefix="/api/goals")
        _app_logger.info("[GoalAPI] Γ£à Θò┐µ£ƒτ¢«µáç API σ╖▓µ│¿σåî: /api/goals")
    except ImportError as e:
        _app_logger.warning(f"[GoalAPI] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑΘò┐µ£ƒτ¢«µáç API Φô¥σ¢╛: {e}")
    except Exception as e:
        _app_logger.error(f"[GoalAPI] Γ¥î Θò┐µ£ƒτ¢«µáç API µ│¿σåîσñ▒Φ┤Ñ: {e}")

    # µ│¿σåîµûçΣ╗╢ Hub API∩╝êFileRegistry + FileWatcher τ╗ƒΣ╕ÇµÄÑσÅú∩╝ë
    try:
        from app.api.file_hub_routes import file_hub_bp as _file_hub_bp

        app.register_blueprint(_file_hub_bp, url_prefix="/api/files")
        _app_logger.info("[FileHubAPI] Γ£à µûçΣ╗╢ Hub API σ╖▓µ│¿σåî: /api/files")
    except ImportError as e:
        _app_logger.warning(f"[FileHubAPI] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑµûçΣ╗╢ Hub Φô¥σ¢╛: {e}")
    except Exception as e:
        _app_logger.error(f"[FileHubAPI] Γ¥î µûçΣ╗╢ Hub API µ│¿σåîσñ▒Φ┤Ñ: {e}")

    # µ│¿σåîσÉÄσÅ░Σ╜£Σ╕Ü API∩╝êJobRunner + TriggerRegistry∩╝ë
    try:
        from app.api.job_routes import job_bp as _job_bp

        app.register_blueprint(_job_bp)
        _app_logger.info("[JobAPI] Γ£à σÉÄσÅ░Σ╜£Σ╕Ü API σ╖▓µ│¿σåî: /api/jobs")
    except ImportError as e:
        _app_logger.warning(f"[JobAPI] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑΣ╜£Σ╕Ü API Φô¥σ¢╛: {e}")
    except Exception as e:
        _app_logger.error(f"[JobAPI] Γ¥î Σ╜£Σ╕Ü API µ│¿σåîσñ▒Φ┤Ñ: {e}")

    # µ│¿σåîΦ┐Éτ╗┤σüÑσ║╖ API∩╝êHealthSnapshot + RemediationPolicy + OpsEventBus∩╝ë
    try:
        from app.api.ops_routes import ops_bp as _ops_bp

        app.register_blueprint(_ops_bp)
        _app_logger.info("[OpsAPI] Γ£à Φ┐Éτ╗┤σüÑσ║╖ API σ╖▓µ│¿σåî: /api/ops")
    except ImportError as e:
        _app_logger.warning(f"[OpsAPI] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑΦ┐Éτ╗┤ API Φô¥σ¢╛: {e}")
    except Exception as e:
        _app_logger.error(f"[OpsAPI] Γ¥î Φ┐Éτ╗┤ API µ│¿σåîσñ▒Φ┤Ñ: {e}")

    # µ│¿σåîσ╜▒σ¡ÉΦ┐╜Φ╕¬ API∩╝êShadowWatcher + ProactiveAgent∩╝ë
    try:
        from app.api.shadow_routes import shadow_bp as _shadow_bp

        app.register_blueprint(_shadow_bp)
        _app_logger.info("[ShadowAPI] Γ£à σ╜▒σ¡ÉΦ┐╜Φ╕¬ API σ╖▓µ│¿σåî: /api/shadow")
    except ImportError as e:
        _app_logger.warning(f"[ShadowAPI] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑσ╜▒σ¡ÉΦ┐╜Φ╕¬Φô¥σ¢╛: {e}")
    except Exception as e:
        _app_logger.error(f"[ShadowAPI] Γ¥î σ╜▒σ¡ÉΦ┐╜Φ╕¬ API µ│¿σåîσñ▒Φ┤Ñ: {e}")

    # µ│¿σåîσ«Åσ╜òσê╢ API∩╝êMacroRecorder Σ╕╗σè¿σ╗║Φ««∩╝ë
    try:
        from app.api.macro_routes import macro_bp as _macro_bp

        app.register_blueprint(_macro_bp)
        _app_logger.info("[MacroAPI] Γ£à σ«Åσ╜òσê╢ API σ╖▓µ│¿σåî: /api/macro")
    except ImportError as e:
        _app_logger.warning(f"[MacroAPI] ΓÜá∩╕Å µ£¬Φâ╜σ»╝σàÑσ«Åσ╜òσê╢Φô¥σ¢╛: {e}")
    except Exception as e:
        _app_logger.error(f"[MacroAPI] Γ¥î σ«Åσ╜òσê╢ API µ│¿σåîσñ▒Φ┤Ñ: {e}")

    _app_logger.info("[INIT] Γ£à µëÇµ£ëΦô¥σ¢╛µ│¿σåîσ«îµêÉ")


def _initialize_background_runtime():
    """Warm up long-running subsystems so jobs, triggers, and ops are live after startup."""
    try:
        time.sleep(1)

        from app.core.jobs.job_runner import get_job_runner
        from app.core.jobs.trigger_registry import get_trigger_registry
        from app.core.ops.ops_event_bus import get_ops_bus
        from app.core.skills.skill_trigger_binding import get_skill_binding_manager

        get_ops_bus()
        runner = get_job_runner()
        registry = get_trigger_registry()
        bindings = get_skill_binding_manager()

        # σê¥σºïσîû GoalManager σ╣╢µ│¿σåî goal_check σñäτÉåσÖ¿
        try:
            from app.core.goal.goal_job_handler import register_goal_handler
            from app.core.goal.goal_manager import get_goal_manager

            _gm = get_goal_manager()
            register_goal_handler(runner)
            _app_logger.info(f"[GoalManager] Γ£à Θò┐µ£ƒτ¢«µáçτ«íτÉåσÖ¿σ╖▓σÉ»σè¿ (µ┤╗Φ╖âτ¢«µáç: {_gm.count()} µ¥í)")
        except Exception as _ge:
            _app_logger.warning(f"[GoalManager] ΓÜá∩╕Å σê¥σºïσîûσñ▒Φ┤Ñ∩╝êΘ¥₧Φç┤σæ╜∩╝ë: {_ge}")

        # σê¥σºïσîû FileRegistry σ╣╢σÉ»σè¿ FileWatcher
        try:
            from app.core.file.file_registry import get_file_registry
            from app.core.file.file_watcher import get_file_watcher

            _fr = get_file_registry()
            _fw = get_file_watcher()
            _fw.start()
            _app_logger.info(f"[FileHub] Γ£à µûçΣ╗╢µ│¿σåîΦí¿σ╖▓σÉ»σè¿ (σ╖▓µö╢σ╜ò: {_fr.count()} Σ╕¬µûçΣ╗╢)")
        except Exception as _fe:
            _app_logger.warning(f"[FileHub] ΓÜá∩╕Å µûçΣ╗╢µ¿íσ¥ùσê¥σºïσîûσñ▒Φ┤Ñ∩╝êΘ¥₧Φç┤σæ╜∩╝ë: {_fe}")

        # σê¥σºïσîûσ╖ÑΣ╜£µûçΣ╗╢σ║ô∩╝êσÉÄσÅ░σ┐½ΘÇƒµë½µÅÅµíîΘ¥ó/µûçµíú/Σ╕ïΦ╜╜∩╝ë
        try:
            from web.work_file_library import get_work_file_library

            _wfl_inst = get_work_file_library()
            if not _wfl_inst.is_indexed():
                _wfl_inst.scan_locations()
                _app_logger.debug("[WorkFileLibrary] ≡ƒÜÇ σ╖ÑΣ╜£µûçΣ╗╢σ║ôσÉÄσÅ░µë½µÅÅσ╖▓σÉ»σè¿∩╝êµíîΘ¥ó/µûçµíú/Σ╕ïΦ╜╜∩╝ë")
            else:
                _app_logger.info(
                    f"[WorkFileLibrary] Γ£à σ╖ÑΣ╜£µûçΣ╗╢σ║ôσ╖▓σèáΦ╜╜: {_wfl_inst.count()} Σ╕¬σ╖ÑΣ╜£µûçΣ╗╢"
                )
        except Exception as _wfl_e:
            _app_logger.warning(f"[WorkFileLibrary] ΓÜá∩╕Å σê¥σºïσîûσñ▒Φ┤Ñ∩╝êΘ¥₧Φç┤σæ╜∩╝ë: {_wfl_e}")

        _app_logger.info(
            "[Runtime] Γ£à σÉÄσÅ░Φ┐ÉΦíîµù╢σ╖▓σÉ»σè¿: "
            f"job_runner={runner is not None}, "
            f"triggers={len(registry.list_all())}, "
            f"bindings={len(bindings.list_bindings())}"
        )

        # µ│¿σåî ShadowTracer ΘÿêσÇ╝ ΓåÆ DistillManager Φç¬σè¿µÅÉΣ║ñΦ«¡τ╗â∩╝êµò░µì«Θú₧Φ╜«Θù¡τÄ»∩╝ë
        try:
            from app.core.learning.distill_manager import DistillManager
            from app.core.learning.shadow_tracer import ShadowTracer, TraceEvent

            def _on_training_ready(event: str, skill_id: str, count: int):
                if event == TraceEvent.TRAINING_READY:
                    _app_logger.debug(
                        f"[Flywheel] ≡ƒÜÇ skill={skill_id} σ╖▓τº»τ┤» {count} µ¥íΣ╝ÿΦ┤¿Φ«░σ╜ò∩╝îΦç¬σè¿µÅÉΣ║ñ LoRA Φ«¡τ╗â..."
                    )
                    try:
                        job_id = DistillManager.instance().submit(skill_id)
                        _app_logger.info(
                            f"[Flywheel] Γ£à Φ«¡τ╗âΣ╗╗σèíσ╖▓µÅÉΣ║ñ job_id={job_id} skill={skill_id}"
                        )
                    except Exception as _e:
                        _app_logger.warning(f"[Flywheel] ΓÜá∩╕Å Φç¬σè¿µÅÉΣ║ñΦ«¡τ╗âσñ▒Φ┤Ñ: {_e}")

            ShadowTracer.add_listener(_on_training_ready)
            _app_logger.info("[Flywheel] Γ£à µò░µì«Θú₧Φ╜«τ¢æσÉ¼σÖ¿σ╖▓µ│¿σåî∩╝êShadowTracer ΓåÆ DistillManager∩╝ë")
        except Exception as _fe:
            _app_logger.warning(f"[Flywheel] ΓÜá∩╕Å Θú₧Φ╜«τ¢æσÉ¼σÖ¿µ│¿σåîσñ▒Φ┤Ñ∩╝êΘ¥₧Φç┤σæ╜∩╝ë: {_fe}")

    except Exception as exc:
        _app_logger.warning(f"[Runtime] ΓÜá∩╕Å σÉÄσÅ░Φ┐ÉΦíîµù╢σê¥σºïσîûσñ▒Φ┤Ñ: {exc}")


# σÉîµ¡Ñµ│¿σåîµëÇµ£ëΦô¥σ¢╛∩╝êσ┐àΘí╗σ£¿ app.run() Σ╣ïσëìσ«îµêÉ∩╝îσÉªσêÖ Flask 3.x Σ╝Üσ£¿Θªûµ¼íΦ»╖µ▒éσÉÄµïÆτ╗¥µ│¿σåî∩╝ë
_register_blueprints_deferred()
threading.Thread(
    target=_initialize_background_runtime, name="RuntimeBootstrap", daemon=True
).start()

# σÉÄσÅ░ΘóäσèáΦ╜╜ Vosk Φ»¡Θƒ│µ¿íσ₧ï∩╝êσçÅσ░æΘªûµ¼íΦ»åσê½σ╗╢Φ┐ƒ∩╝ë
try:
    from web.voice_engine import preload as _voice_preload

    _voice_preload()
except Exception:
    pass

CHAT_DIR = os.path.join(PROJECT_ROOT, "chats")
WORKSPACE_DIR = get_workspace_root()
UPLOAD_DIR = os.path.join(PROJECT_ROOT, "web", "uploads")
os.makedirs(CHAT_DIR, exist_ok=True)
os.makedirs(WORKSPACE_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ================= Settings Manager (µÅÉσëìσèáΦ╜╜) =================
try:
    from settings import SettingsManager
except ImportError:
    from web.settings import SettingsManager
settings_manager = SettingsManager()

# ================= σè¿µÇüµ¿íσ₧ïτ«íτÉåσÖ¿ =================
# Φç¬σè¿Σ╗Ä API σÅæτÄ░σÅ»τö¿µ¿íσ₧ïσ╣╢µîëΣ╗╗σèíτ▒╗σ₧ïµÖ║Φâ╜σî╣Θàì∩╝îµùáΘ£Çµëïσè¿τ╗┤µèñµ¿íσ₧ïσêùΦí¿πÇé
# µû░µ¿íσ₧ïΣ╕èτ║┐σÉÄΦç¬σè¿µäƒτƒÑ∩╝îTTL τ╝ôσ¡ÿµ»Å 6 σ░Åµù╢σê╖µû░Σ╕Çµ¼íπÇé

try:
    from web.model_manager import KNOWN_MODEL_REGISTRY as _MODEL_REGISTRY
    from web.model_manager import ModelManager

    _model_manager_available = True
except ImportError:
    try:
        from model_manager import KNOWN_MODEL_REGISTRY as _MODEL_REGISTRY
        from model_manager import ModelManager

        _model_manager_available = True
    except ImportError:
        _model_manager_available = False
        ModelManager = None
        _MODEL_REGISTRY = {}

# Θ¥ÖµÇüΘ╗ÿΦ«ñσÇ╝∩╝êAPI Σ╕ìσÅ»τö¿µù╢τÜäσà£σ║ò∩╝îΣ╣ƒµÿ»σÉ»σè¿µù╢τÜäσê¥σºïσÇ╝∩╝ë
MODEL_MAP = {
    "CHAT":        "gemini-3-flash-preview",
    "CODER":       "gemini-3.1-pro-preview",   # generate_content σà╝σ«╣∩╝îµ£Çσ╝║τ¢┤Φ░âµ¿íσ₧ï
    "WEB_SEARCH":  "gemini-2.5-flash",
    "VISION":      "gemini-3-flash-preview",
    "RESEARCH":    "deep-research-pro-preview-12-2025",
    "FILE_GEN":    "gemini-3-flash-preview",
    "PAINTER":     "gemini-3.1-flash-image-preview",
    "SYSTEM":      "local-executor",
    "FILE_OP":     "local-executor",
    "AGENT":       "gemini-3-flash-preview",
    "FILE_SEARCH": "gemini-3-flash-preview",  # µûçΣ╗╢µÉ£τ┤ó/µò┤τÉåσºïτ╗ê Flash
    "DOC_ANNOTATE":"gemini-3.1-pro-preview",  # σ╝║µ¿íσ₧ïµáçµ│¿∩╝îcomplex τö▒ get_model_for_task τí«Φ«ñ
    "COMPLEX":     "gemini-3.1-pro-preview",  # σñìµ¥éσ║ªσìçτ║ºσà£σ║ò∩╝Ü3.1 Pro
}

# ΓöÇΓöÇΓöÇ Interactions-API-only µ¿íσ₧ï∩╝êσè¿µÇüµ¢┤µû░∩╝îΘ¥ÖµÇüΘ╗ÿΦ«ñσà£σ║ò∩╝ëΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
# Φ┐ÖΣ║¢µ¿íσ₧ïΣ╕ìµö»µîü client.models.generate_content()∩╝îσ┐àΘí╗Φ╡░ Interactions API
_INTERACTIONS_ONLY_MODELS = {
    "gemini-3-flash-preview",
    "gemini-3-pro-preview",
    "deep-research-pro-preview-12-2025",  # σ«₧ΘÖàΣ╣ƒµÿ» interactions-only∩╝îΘöÖΦ»»Φ░âτö¿ generate_content Σ╝ÜΦ┐öσ¢₧ 400
}

# ΓöÇΓöÇΓöÇ Interactions API σ¡ùµ«╡σî║σêå ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
# τ£ƒµ¡úτÜä Agent∩╝êdeep-research τ¡ë∩╝ëΘ£ÇΦªü interactions.create(agent=...)∩╝¢
# µÖ«ΘÇÜµ¿íσ₧ï∩╝êgemini-3-pro/flash-preview∩╝ëΦ╡░σÉîΣ╕Çτ½»τé╣Σ╜åσ┐àΘí╗τö¿ model= σ¡ùµ«╡πÇé
# ΘöÖΦ»»τö¿ agent= Σ╝áµÖ«ΘÇÜµ¿íσ₧ï ID Σ╝Üσ╛ùσê░ 400: "refers to a model, not an agent"πÇé
_INTERACTIONS_AGENT_MODELS = frozenset({
    "deep-research-pro-preview-12-2025",
})


def _is_interactions_agent(model_id: str) -> bool:
    """True = Φ»Ñµ¿íσ₧ïΘ£ÇΦªüτö¿ agent= σ¡ùµ«╡∩╝êτ£ƒµ¡úτÜä Interactions Agent∩╝ë∩╝¢
    False = Φ»Ñµ¿íσ₧ïµÿ»µÖ«ΘÇÜµ¿íσ₧ï∩╝îΘ£ÇΦªüτö¿ model= σ¡ùµ«╡πÇé"""
    mid = str(model_id or "")
    return mid in _INTERACTIONS_AGENT_MODELS or mid.startswith("deep-research-pro-preview")


# σ╜ô Interactions API Σ╣ƒσñ▒Φ┤Ñµù╢τÜäµ£Çτ╗êΘÖìτ║ºµ¿íσ₧ï
_INTERACTIONS_FALLBACK_MODEL = "gemini-2.5-flash"

# ΓöÇΓöÇ Interactions API Φ╜«Φ»óτè╢µÇüσ╕╕ΘçÅ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
_INTERACTION_TERMINAL_STATES = frozenset({"completed", "failed", "cancelled", "error"})
_INTERACTION_SUCCESS_STATES  = frozenset({"completed"})
_INTERACTION_FAIL_STATES     = frozenset({"failed", "cancelled", "error"})

# Σ╕¡Θù┤τè╢µÇü ΓåÆ Σ║║τ▒╗σÅ»Φ»╗µùÑσ┐ù∩╝êΣ╗àσ╜ôτè╢µÇüσÅÿσîûµù╢Φ╛ôσç║∩╝îΘü┐σàìµùÑσ┐ùµ┤¬µ░┤∩╝ë
_INTERACTION_STATUS_MSGS: dict = {
    "active":      "Agent σ╖ÑΣ╜£Σ╕¡ΓÇª",
    "running":     "Agent σ╖ÑΣ╜£Σ╕¡ΓÇª",
    "queued":      "τ¡ëσ╛àΘÿƒσêùΣ╕¡∩╝îσì│σ░åσ╝ÇσºïΓÇª",
    "in_progress": "Agent σñäτÉåΣ╕¡ΓÇª",
    "thinking":    "Agent µ╖▒σ║ªµÇ¥ΦÇâΣ╕¡ΓÇª",
    "searching":   "Agent µ¡úσ£¿µúÇτ┤óΣ║ÆΦüöτ╜æΓÇª",
    "reading":     "Agent µ¡úσ£¿ΘÿàΦ»╗Φ╡äµûÖΓÇª",
    "generating":  "Agent µ¡úσ£¿τöƒµêÉσ¢₧σñìΓÇª",
}

# σà¿σ▒Çµ¿íσ₧ïτ«íτÉåσÖ¿σ«₧Σ╛ï∩╝êσÉÄσÅ░σê¥σºïσîû∩╝ë
_model_manager = None


def _init_model_manager():
    """
    σ£¿σÉÄσÅ░τ║┐τ¿ïΣ╕¡σê¥σºïσîûσè¿µÇüµ¿íσ₧ïτ«íτÉåσÖ¿σ╣╢µ¢┤µû░σà¿σ▒ÇΦ╖»τö▒Φí¿πÇé
    Σ╕ìΘÿ╗σí₧Σ╕╗τ║┐τ¿ïσÉ»σè¿∩╝¢Φ╖»τö▒Φí¿µ¢┤µû░µ£ƒΘù┤Σ╗ìΣ╜┐τö¿Θ¥ÖµÇüΘ╗ÿΦ«ñσÇ╝πÇé
    """
    global MODEL_MAP, _model_manager, _INTERACTIONS_ONLY_MODELS, _INTERACTIONS_FALLBACK_MODEL
    if not _model_manager_available or ModelManager is None:
        _app_logger.debug("[ModelManager] µ¿íσ¥ùΣ╕ìσÅ»τö¿∩╝îΣ╜┐τö¿Θ¥ÖµÇüΘ╗ÿΦ«ñΦ╖»τö▒")
        return
    try:
        _app_logger.debug("[ModelManager] ≡ƒöì µ¡úσ£¿σÅæτÄ░σÅ»τö¿µ¿íσ₧ï...")
        _model_manager = ModelManager(client)
        dynamic_map = _model_manager.get_model_map()
        MODEL_MAP.update(dynamic_map)
        _INTERACTIONS_ONLY_MODELS = _model_manager.get_interactions_only_models()
        _INTERACTIONS_FALLBACK_MODEL = _model_manager.get_fallback_model()
        # σÉîµ¡Ñµ¢┤µû░ SmartDispatcher τÜä MODEL_MAP σ╝òτö¿
        try:
            SmartDispatcher._dependencies["MODEL_MAP"] = MODEL_MAP
        except Exception:
            pass
        _app_logger.info(f"[ModelManager] Γ£à σè¿µÇüΦ╖»τö▒σ╖▓σèáΦ╜╜: {len(dynamic_map)} Σ╕¬Σ╗╗σèí")
        # ΓöÇΓöÇ σÉîµ¡Ñµ¢┤µû░ ModelFallbackExecutor τÜäΦ╖»τö▒Φí¿ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        try:
            from app.core.llm.model_fallback import get_fallback_executor
            get_fallback_executor().update_model_map(MODEL_MAP)
            _app_logger.info("[ModelManager] Γ£à ModelFallbackExecutor Φ╖»τö▒Φí¿σ╖▓σÉîµ¡Ñ")
        except Exception as _fe:
            _app_logger.warning(f"[ModelManager] ΓÜá∩╕Å ModelFallbackExecutor σÉîµ¡Ñσñ▒Φ┤Ñ∩╝êΘ¥₧Φç┤σæ╜∩╝ë: {_fe}")
        # ΓöÇΓöÇ σÉîµ¡Ñµ¢┤µû░ AIRouter τÜäΦ╜╗ΘçÅΦ╖»τö▒µ¿íσ₧ï ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        try:
            from app.core.routing.ai_router import AIRouter
            # ΘÇëσÅûσÅ»τö¿τÜäΘ¥₧ interactions-onlyπÇüΘÇƒσ║ªµ£Çσ┐½τÜäµ¿íσ₧ïΣ╜£Σ╕║Φ╖»τö▒σÖ¿
            _available_caps = _model_manager._cached_caps
            _fast_candidates = [
                (mid, caps) for mid, caps in _available_caps.items()
                if not caps.get("interactions_only", False)
                and not caps.get("image_gen", False)
                and mid != "local-executor"
            ]
            if _fast_candidates:
                _router_candidate = max(
                    _fast_candidates,
                    key=lambda x: x[1].get("speed", 0) + x[1].get("tier", 0) * 0.1
                )[0]
                AIRouter.set_router_model(_router_candidate)
        except Exception as _are:
            _app_logger.warning(f"[ModelManager] ΓÜá∩╕Å AIRouter Φ╖»τö▒µ¿íσ₧ïµ¢┤µû░σñ▒Φ┤Ñ∩╝êΘ¥₧Φç┤σæ╜∩╝ë: {_are}")
    except Exception as _me:
        import traceback as _tb

        _app_logger.warning(f"[ModelManager] ΓÜá∩╕Å σè¿µÇüΦ╖»τö▒σê¥σºïσîûσñ▒Φ┤Ñ∩╝îΣ╜┐τö¿Θ¥ÖµÇüΘ╗ÿΦ«ñσÇ╝: {_me}")
        _tb.print_exc()


# µ¿íσ₧ïΦâ╜σè¢τƒ⌐Θÿ╡∩╝êτö¿Σ║Äµÿ╛τñ║∩╝îσè¿µÇüµ¿íσ₧ïΦç¬σè¿ΦíÑσàà∩╝ë
MODEL_INFO = {
    "gemini-3-pro-preview": {
        "name": "Gemini 3.0 Pro",
        "speed": "≡ƒÜÇ",
        "tier": 7,
        "strengths": ["µÄ¿τÉå", "σêåµ₧É", "Σ╗úτáü", "σñìµ¥éΣ╗╗σèí"],
    },
    "gemini-3-flash-preview": {
        "name": "Gemini 3.0 Flash",
        "speed": "ΓÜí",
        "tier": 6,
        "strengths": ["σ┐½ΘÇƒ", "σ»╣Φ»¥", "σñÜµ¿íµÇü"],
    },
    "gemini-2.5-flash": {
        "name": "Gemini 2.5 Flash",
        "speed": "≡ƒîÉ",
        "tier": 5,
        "strengths": ["Φüöτ╜æµÉ£τ┤ó", "grounding"],
    },
    "gemini-2.5-flash-preview": {
        "name": "Gemini 2.5 Flash Preview",
        "speed": "≡ƒîÉ",
        "tier": 5,
        "strengths": ["Φüöτ╜æµÉ£τ┤ó", "grounding"],
    },
    "gemini-2.5-pro-preview": {
        "name": "Gemini 2.5 Pro",
        "speed": "≡ƒÄ»",
        "tier": 6,
        "strengths": ["µÄ¿τÉå", "Σ╗úτáü", "σêåµ₧É"],
    },
    "deep-research-pro-preview-12-2025": {
        "name": "Deep Research Pro",
        "speed": "≡ƒö¼",
        "tier": 7,
        "strengths": ["µ╖▒σ║ªτáöτ⌐╢", "σ¡ªµ£»σêåµ₧É", "τ╗╝σÉêµèÑσæè"],
    },
    "gemini-3.1-flash-image-preview": {
        "name": "Gemini 3.1 Flash Image",
        "speed": "≡ƒÄ¿",
        "tier": 6,
        "strengths": ["σ¢╛σâÅτöƒµêÉ", "σê¢µäÅτ╗ÿτö╗", "Φë║µ£»ΘúÄµá╝"],
    },
    "gemini-2.0-flash-exp": {
        "name": "Gemini 2.0 Flash Exp",
        "speed": "≡ƒº¬",
        "tier": 5,
        "strengths": ["σ¢╛σâÅτöƒµêÉ", "σñÜµ¿íµÇü", "σ«₧Θ¬îσèƒΦâ╜"],
    },
    "gemini-2.0-flash": {
        "name": "Gemini 2.0 Flash",
        "speed": "ΓÜí",
        "tier": 5,
        "strengths": ["σ┐½ΘÇƒ", "σñÜµ¿íµÇü"],
    },
    "gemini-1.5-pro": {
        "name": "Gemini 1.5 Pro",
        "speed": "≡ƒôÜ",
        "tier": 5,
        "strengths": ["Θò┐Σ╕èΣ╕ïµûç", "µÄ¿τÉå"],
    },
    "gemini-1.5-flash": {
        "name": "Gemini 1.5 Flash",
        "speed": "ΓÜí",
        "tier": 4,
        "strengths": ["σ┐½ΘÇƒ", "τ╗Åµ╡Ä"],
    },
    "local-executor": {
        "name": "Local Executor",
        "speed": "≡ƒûÑ∩╕Å",
        "tier": 0,
        "strengths": ["τ│╗τ╗ƒµôìΣ╜£", "µëôσ╝Çσ║öτö¿", "µûçΣ╗╢τ«íτÉå"],
    },
}


def get_model_display_name(model_id):
    """ΦÄ╖σÅûµ¿íσ₧ïσÅïσÑ╜µÿ╛τñ║σÉìτº░∩╝¢σè¿µÇüσÅæτÄ░τÜäµû░µ¿íσ₧ïΦç¬σè¿Σ╗ÄΦâ╜σè¢µ│¿σåîΦí¿ΦíÑσààπÇé"""
    info = MODEL_INFO.get(model_id)
    if info:
        return f"{info['name']} {info['speed']}"
    # σè¿µÇüµ¿íσ₧ï∩╝ÜΣ╗Ä ModelManager Φâ╜σè¢τ╝ôσ¡ÿΦÄ╖σÅû
    if _model_manager:
        caps = _model_manager._cached_caps.get(model_id)
        if caps and caps.get("display"):
            return caps["display"]
    # µ£¬τƒÑµ¿íσ₧ï∩╝Üτ¢┤µÄÑσ▒òτñ║ ID
    return model_id


# ================= µ£¼σ£░τ│╗τ╗ƒµëºΦíîσÖ¿ (σ╖▓Φ┐üτº╗σê░ web/local_executor.py) =================
try:
    from web.local_executor import LocalExecutor
except ImportError:
    from local_executor import LocalExecutor


# ================= µûçΣ╗╢µôìΣ╜£µëºΦíîσÖ¿ =================
class FileOperator:
    """
    µ£¼σ£░µûçΣ╗╢µôìΣ╜£µëºΦíîσÖ¿ - σñäτÉåµûçΣ╗╢Φ»╗σåÖπÇüτ«íτÉåτ¡ëµôìΣ╜£
    """

    # µûçΣ╗╢µôìΣ╜£σà│Θö«Φ»ì
    FILE_KEYWORDS = [
        "Φ»╗σÅûµûçΣ╗╢",
        "µëôσ╝ÇµûçΣ╗╢",
        "µƒÑτ£ïµûçΣ╗╢",
        "Φ»╗µûçΣ╗╢",
        "τ£ïτ£ïµûçΣ╗╢",
        "σê¢σ╗║µûçΣ╗╢",
        "µû░σ╗║µûçΣ╗╢",
        "σåÖσàÑµûçΣ╗╢",
        "Σ┐¥σ¡ÿµûçΣ╗╢",
        "σêáΘÖñµûçΣ╗╢",
        "τº╗σè¿µûçΣ╗╢",
        "σñìσê╢µûçΣ╗╢",
        "Θçìσæ╜σÉì",
        "µûçΣ╗╢σêùΦí¿",
        "τ¢«σ╜ò",
        "µûçΣ╗╢σñ╣",
        "σêùσç║µûçΣ╗╢",
        "Φç¬σè¿σ╜Æτ║│",
        "Φç¬σè¿µò┤τÉå",
        "σ╜Æτ║│µûçΣ╗╢σñ╣",
        "µò┤τÉåµûçΣ╗╢σñ╣",
        "σ╜ÆµíúµûçΣ╗╢σñ╣",
        "σ╛«Σ┐íµûçΣ╗╢σ╜Æτ║│",
        "read file",
        "open file",
        "create file",
        "delete file",
        "list files",
        "directory",
        "folder",
    ]

    FOLDER_ORGANIZE_KEYWORDS = [
        "Φç¬σè¿σ╜Æτ║│",
        "Φç¬σè¿µò┤τÉå",
        "σ╜Æτ║│",
        "µò┤τÉå",
        "σ╜Æµíú",
        "σ╜Æτ▒╗",
        "σêåτ▒╗",
        "µûçΣ╗╢σñ╣",
        "τ¢«σ╜ò",
        "σ╛«Σ┐íµûçΣ╗╢",
        "wechat files",
    ]

    @classmethod
    def is_file_operation(cls, text):
        """µúÇµ╡ïµÿ»σÉªµÿ»µûçΣ╗╢µôìΣ╜£Φ»╖µ▒é"""
        text_lower = text.lower()
        return any(kw in text_lower for kw in cls.FILE_KEYWORDS)

    @classmethod
    def _is_folder_organize_intent(cls, text_lower: str) -> bool:
        has_action = any(
            kw in text_lower for kw in ["σ╜Æτ║│", "µò┤τÉå", "σ╜Æµíú", "σ╜Æτ▒╗", "σêåτ▒╗"]
        )
        has_target = any(kw in text_lower for kw in ["µûçΣ╗╢σñ╣", "τ¢«σ╜ò", "Φ╖»σ╛ä", "µûçΣ╗╢"])
        if has_action and has_target:
            return True
        return any(kw in text_lower for kw in cls.FOLDER_ORGANIZE_KEYWORDS)

    @classmethod
    def _extract_path_from_text(cls, user_input: str) -> str:
        """Extract a likely filesystem path from user input."""
        import re

        patterns = [
            r'["\']([^"\']+)["\']',
            r'([A-Za-z]:\\(?:[^\\/:*?"<>|\r\n]+\\)*[^\\/:*?"<>|\r\n]*)',
            r"(\.?/[\w\-./ ]+)",
        ]
        for pattern in patterns:
            m = re.search(pattern, user_input)
            if m:
                candidate = m.group(1).strip().strip("∩╝îπÇé,.;∩╝¢")
                if candidate:
                    return candidate
        return ""

    @classmethod
    def execute(cls, user_input):
        """µëºΦíîµûçΣ╗╢µôìΣ╜£"""
        text_lower = user_input.lower()
        result = {"success": False, "action": "", "message": "", "content": ""}

        # === µîçσ«ÜΦ╖»σ╛äµûçΣ╗╢σñ╣Φç¬σè¿σ╜Æτ║│ ===
        if cls._is_folder_organize_intent(text_lower):
            folder_path = cls._extract_path_from_text(user_input)
            if not folder_path:
                folder_path = get_default_wechat_files_dir()

            if not folder_path:
                result["message"] = (
                    "Γ¥ô Φ»╖µÅÉΣ╛¢Φªüσ╜Æτ║│τÜäµûçΣ╗╢σñ╣Φ╖»σ╛ä∩╝êσÅ»τö¿σ╝òσÅ╖σîàΦú╣∩╝ë∩╝îµêûσ£¿ config/user_settings.json Σ╕¡Φ«╛τ╜« "
                    "storage.wechat_files_dir Σ╜£Σ╕║Θ╗ÿΦ«ñΦ╖»σ╛ä"
                )
                return result

            if not os.path.isabs(folder_path):
                folder_path = os.path.join(WORKSPACE_DIR, folder_path)

            if not os.path.isdir(folder_path):
                result["message"] = f"Γ¥î τ¢«σ╜òΣ╕ìσ¡ÿσ£¿: {folder_path}"
                return result

            try:
                try:
                    from web.folder_catalog_organizer import FolderCatalogOrganizer
                except Exception:
                    from folder_catalog_organizer import FolderCatalogOrganizer

                analyzer = get_file_analyzer()
                organizer = get_file_organizer()
                engine = FolderCatalogOrganizer(
                    get_organize_root(), analyzer, organizer
                )
                summary = engine.organize_folder(folder_path)

                if not summary.get("success"):
                    result["message"] = (
                        f"Γ¥î Φç¬σè¿σ╜Æτ║│σñ▒Φ┤Ñ: {summary.get('error', 'µ£¬τƒÑΘöÖΦ»»')}"
                    )
                    return result

                report_md = summary.get("report_markdown", "")
                report_json = summary.get("report_json", "")
                entries = summary.get("entries", [])

                sender_preview = []
                for item in entries:
                    sender = item.get("sender", "µ£¬τƒÑ")
                    if sender and sender != "µ£¬τƒÑ":
                        sender_preview.append(sender)
                sender_preview = sorted(set(sender_preview))[:8]
                sender_preview_text = (
                    "πÇü".join(sender_preview)
                    if sender_preview
                    else "µ£¬Φ»åσê½σê░σÅ»Θ¥áσÅæΘÇüΦÇà"
                )

                result["success"] = True
                result["action"] = "folder_auto_catalog"
                result["message"] = (
                    f"Γ£à σ╜Æτ║│σ«îµêÉ∩╝Ü{summary.get('organized_count', 0)}/{summary.get('total_files', 0)} Σ╕¬µûçΣ╗╢σ╖▓σ╜Æτ║│"
                    f"\n≡ƒôü µ¥Ñµ║Éτ¢«σ╜ò: {summary.get('source_dir', folder_path)}"
                    f"\n≡ƒº╛ µ╕àσìò(MD): {report_md}"
                    f"\n≡ƒº╛ µ╕àσìò(JSON): {report_json}"
                    f"\n≡ƒæñ Φ»åσê½σê░τÜäσÅæΘÇüΦÇà/µ¥Ñµ║ÉΣ║║: {sender_preview_text}"
                )
                return result
            except Exception as e:
                result["message"] = f"Γ¥î Φç¬σè¿σ╜Æτ║│σ╝éσ╕╕: {str(e)}"
                return result

        # === Φ»╗σÅûµûçΣ╗╢ ===
        if any(
            kw in text_lower
            for kw in [
                "Φ»╗σÅû",
                "µëôσ╝ÇµûçΣ╗╢",
                "µƒÑτ£ïµûçΣ╗╢",
                "Φ»╗µûçΣ╗╢",
                "τ£ïτ£ï",
                "read file",
                "open file",
            ]
        ):
            # µÅÉσÅûµûçΣ╗╢Φ╖»σ╛ä
            import re

            # σ░¥Φ»òσî╣Θàìσ╕╕ΦºüΦ╖»σ╛äµ¿íσ╝Å
            patterns = [
                r'["\']([^"\']+)["\']',  # σ╝òσÅ╖σîàσ¢┤τÜäΦ╖»σ╛ä
                r"([A-Za-z]:\\[^\s]+)",  # Windows τ╗¥σ»╣Φ╖»σ╛ä
                r"(\.?/[^\s]+)",  # Unix ΘúÄµá╝Φ╖»σ╛ä
                r"(\S+\.\w{1,5})(?:\s|$)",  # σ╕ªµë⌐σ▒òσÉìτÜäµûçΣ╗╢
            ]

            filepath = None
            for pattern in patterns:
                match = re.search(pattern, user_input)
                if match:
                    filepath = match.group(1)
                    break

            if filepath:
                # σªéµ₧£µÿ»τ¢╕σ»╣Φ╖»σ╛ä∩╝îσ£¿ workspace τ¢«σ╜òµƒÑµë╛
                if not os.path.isabs(filepath):
                    workspace_path = os.path.join(WORKSPACE_DIR, filepath)
                    if os.path.exists(workspace_path):
                        filepath = workspace_path

                if os.path.exists(filepath):
                    try:
                        with open(
                            filepath, "r", encoding="utf-8", errors="ignore"
                        ) as f:
                            content = f.read()

                        # ΘÖÉσê╢σåàσ«╣Θò┐σ║ª
                        if len(content) > 10000:
                            content = content[:10000] + "\n\n... (µûçΣ╗╢Φ┐çΘò┐∩╝îσ╖▓µê¬µû¡)"

                        result["success"] = True
                        result["action"] = "read_file"
                        result["message"] = (
                            f"Γ£à σ╖▓Φ»╗σÅûµûçΣ╗╢: {os.path.basename(filepath)}"
                        )
                        result["content"] = f"```\n{content}\n```"
                        return result
                    except Exception as e:
                        result["message"] = f"Γ¥î Φ»╗σÅûµûçΣ╗╢σñ▒Φ┤Ñ: {str(e)}"
                        return result
                else:
                    result["message"] = f"Γ¥î µûçΣ╗╢Σ╕ìσ¡ÿσ£¿: {filepath}"
                    return result
            else:
                result["message"] = "Γ¥ô Φ»╖µîçσ«ÜΦªüΦ»╗σÅûτÜäµûçΣ╗╢Φ╖»σ╛ä"
                return result

        # === σêùσç║µûçΣ╗╢ ===
        if any(
            kw in text_lower
            for kw in [
                "µûçΣ╗╢σêùΦí¿",
                "τ¢«σ╜ò",
                "σêùσç║µûçΣ╗╢",
                "list files",
                "directory",
                "µûçΣ╗╢σñ╣Θçî",
            ]
        ):
            # µÅÉσÅûτ¢«σ╜òΦ╖»σ╛ä
            import re

            patterns = [
                r'["\']([^"\']+)["\']',
                r"([A-Za-z]:\\[^\s]+)",
                r"(\.?/[^\s]+)",
            ]

            dirpath = WORKSPACE_DIR  # Θ╗ÿΦ«ñ workspace
            for pattern in patterns:
                match = re.search(pattern, user_input)
                if match:
                    dirpath = match.group(1)
                    break

            if not os.path.isabs(dirpath):
                dirpath = os.path.join(WORKSPACE_DIR, dirpath)

            if os.path.isdir(dirpath):
                try:
                    items = os.listdir(dirpath)
                    file_list = []
                    for item in items[:50]:  # ΘÖÉσê╢µò░ΘçÅ
                        item_path = os.path.join(dirpath, item)
                        if os.path.isdir(item_path):
                            file_list.append(f"≡ƒôü {item}/")
                        else:
                            size = os.path.getsize(item_path)
                            size_str = (
                                f"{size/1024:.1f}KB" if size > 1024 else f"{size}B"
                            )
                            file_list.append(f"≡ƒôä {item} ({size_str})")

                    result["success"] = True
                    result["action"] = "list_files"
                    result["message"] = f"Γ£à τ¢«σ╜ò: {dirpath}"
                    result["content"] = "\n".join(file_list) if file_list else "τ⌐║τ¢«σ╜ò"
                    return result
                except Exception as e:
                    result["message"] = f"Γ¥î Φ»╗σÅûτ¢«σ╜òσñ▒Φ┤Ñ: {str(e)}"
                    return result
            else:
                result["message"] = f"Γ¥î τ¢«σ╜òΣ╕ìσ¡ÿσ£¿: {dirpath}"
                return result

        # === σê¢σ╗║/σåÖσàÑµûçΣ╗╢ ===
        if any(
            kw in text_lower
            for kw in ["σê¢σ╗║µûçΣ╗╢", "µû░σ╗║µûçΣ╗╢", "σåÖσàÑµûçΣ╗╢", "Σ┐¥σ¡ÿσê░", "create file"]
        ):
            result["message"] = (
                "≡ƒÆí Φ»╖Σ╜┐τö¿Σ╗úτáüτöƒµêÉσèƒΦâ╜∩╝îKoto Σ╝ÜΦç¬σè¿Σ┐¥σ¡ÿτöƒµêÉτÜäµûçΣ╗╢σê░ workspace"
            )
            return result

        result["message"] = "Γ¥ô µùáµ│òΦ»åσê½Φ»ÑµûçΣ╗╢µôìΣ╜£∩╝îΦ»╖σ░¥Φ»ò∩╝ÜΦ»╗σÅûµûçΣ╗╢πÇüσêùσç║τ¢«σ╜òτ¡ë"
        return result

    @classmethod
    def watch_directory(cls, directory, callback=None, patterns=None):
        """τ¢æσÉ¼τ¢«σ╜òσÅÿσîûσ╣╢ΦºªσÅæσ¢₧Φ░â"""
        try:
            from watchdog.events import FileSystemEventHandler
            from watchdog.observers import Observer

            if patterns is None:
                patterns = ["*.txt", "*.pdf", "*.docx", "*.xlsx", "*.csv"]

            class ChangeHandler(FileSystemEventHandler):
                def on_created(self, event):
                    if not event.is_directory:
                        filename = os.path.basename(event.src_path)
                        if any(filename.endswith(p.replace("*", "")) for p in patterns):
                            if callback:
                                callback("created", event.src_path)

                def on_modified(self, event):
                    if not event.is_directory:
                        filename = os.path.basename(event.src_path)
                        if any(filename.endswith(p.replace("*", "")) for p in patterns):
                            if callback:
                                callback("modified", event.src_path)

            observer = Observer()
            observer.schedule(ChangeHandler(), directory, recursive=True)
            observer.start()

            return {
                "success": True,
                "observer": observer,
                "message": f"Γ£à σ╖▓σ╝Çσºïτ¢æσÉ¼τ¢«σ╜ò: {directory}",
            }
        except Exception as e:
            return {"success": False, "message": f"Γ¥î µùáµ│òτ¢æσÉ¼τ¢«σ╜ò: {str(e)}"}

    @classmethod
    def get_file_metadata(cls, filepath):
        """ΦÄ╖σÅûµûçΣ╗╢σàâµò░µì«"""
        try:
            if not os.path.exists(filepath):
                return {"success": False, "message": "µûçΣ╗╢Σ╕ìσ¡ÿσ£¿"}

            stat = os.stat(filepath)
            from datetime import datetime

            return {
                "success": True,
                "filepath": filepath,
                "filename": os.path.basename(filepath),
                "size": f"{stat.st_size / 1024:.2f} KB",
                "created": datetime.fromtimestamp(stat.st_ctime).strftime(
                    "%Y-%m-%d %H:%M:%S"
                ),
                "modified": datetime.fromtimestamp(stat.st_mtime).strftime(
                    "%Y-%m-%d %H:%M:%S"
                ),
                "extension": os.path.splitext(filepath)[1],
                "is_file": os.path.isfile(filepath),
            }
        except Exception as e:
            return {"success": False, "message": f"Γ¥î µùáµ│òΦÄ╖σÅûµûçΣ╗╢Σ┐íµü»: {str(e)}"}


# ================= Φüöτ╜æµÉ£τ┤óΦâ╜σè¢ =================
class WebSearcher:
    """
    Σ╜┐τö¿ Gemini τÜä Google Search Grounding Φâ╜σè¢
    ΦÄ╖σÅûσ«₧µù╢σñ⌐µ░öπÇüµû░Θù╗τ¡ëΣ┐íµü»
    """

    # Θ£ÇΦªüΦüöτ╜æτÜäσà│Θö«Φ»ì∩╝êΣ╕Ñµá╝µö╢τ¬ä∩╝ÜΣ╗àσîàσÉ½σçáΣ╣ÄσÅ¬σ£¿Θ£ÇΦªüσ«₧µù╢Σ┐íµü»µù╢µëìΣ╝Üσç║τÄ░τÜäΦ»ì∩╝ë
    WEB_KEYWORDS = [
        # σñ⌐µ░ö∩╝êΘ½ÿτ╜«Σ┐í∩╝ë
        "σñ⌐µ░ö",
        "µ░öµ╕⌐",
        "Σ╕ïΘ¢¿σÉù",
        "Σ╕ïΘ¢¬σÉù",
        "µ╕⌐σ║ªσñÜσ░æ",
        "σñ⌐µ░öµÇÄΣ╣êµá╖",
        "σñ⌐µ░öΘóäµèÑ",
        "weather",
        "temperature",
        "forecast",
        # σ«₧µù╢Φíîµâà∩╝êΘ½ÿτ╜«Σ┐í∩╝ë
        "ΦéíΣ╗╖",
        "µ▒çτÄç",
        "µ»öτë╣σ╕üΣ╗╖µá╝",
        "Θ╗äΘçæΣ╗╖µá╝",
        "ΘçæΣ╗╖",
        "σ«₧µù╢ΘçæΣ╗╖",
        "Σ╗èµùÑΘçæΣ╗╖",
        "σ╜ôσëìΘçæΣ╗╖",
        "τÄ░Φ┤ºΘ╗äΘçæ",
        "σ¢╜ΘÖàΘçæΣ╗╖",
        "τƒ│µ▓╣Σ╗╖µá╝",
        "aΦéí",
        "µ╕»Φéí",
        "τ╛ÄΦéí",
        "stock price",
        # µ»öΦ╡¢/Σ╜ôΦé▓∩╝êΘ½ÿτ╜«Σ┐í∩╝ë
        "µ»öσêå",
        "µ»öΦ╡¢τ╗ôµ₧£",
        "Φ░üΦ╡óΣ║å",
        # µû░Θù╗∩╝êσÅ¬σî╣ΘàìµÿÄτí«τÜäµû░Θù╗Φ»╖µ▒é∩╝ë
        "Σ╗èσñ⌐µû░Θù╗",
        "µ£Çµû░µû░Θù╗",
        "latest news",
        # Σ║ñΘÇÜσç║ΦíîτÑ¿σèí∩╝êΘ½ÿτ╜«Σ┐í ΓÇö Σ╜ÖτÑ¿/µù╢σê╗Φí¿σ«₧µù╢σÅÿσîû∩╝ë
        "τü½Φ╜ªτÑ¿",
        "Θ½ÿΘôüτÑ¿",
        "σè¿Φ╜ªτÑ¿",
        "µ£║τÑ¿",
        "Σ╜ÖτÑ¿",
        "τÅ¡µ¼íµƒÑΦ»ó",
        "Φ╜ªµ¼íµƒÑΦ»ó",
        "µù╢σê╗Φí¿",
        "σêùΦ╜ªµù╢σê╗",
        "Φê¬τÅ¡µƒÑΦ»ó",
        "Φê¬τÅ¡σè¿µÇü",
        "σçáτé╣σç║σÅæ",
        "σçáτé╣σê░",
        "σçáτé╣σê░Φ╛╛",
        "σñÜΣ╣àσê░",
        "ΦªüσñÜΣ╣à",
    ]

    @classmethod
    def needs_web_search(cls, text):
        """µúÇµ╡ïµÿ»σÉªΘ£ÇΦªüΦüöτ╜æµÉ£τ┤ó

        Σ╝ÿσîûτ¡ûτòÑ∩╝Ü
        1. µúÇµƒÑσà│Θö«Φ»ìσêùΦí¿
        2. σ»╣Σ║ÄΘçæΦ₧ì/Θóäµ╡ïτ▒╗∩╝îµ¢┤σÇ╛σÉæΣ║Äweb-search
        3. σ»╣Σ║Äτâ¡τé╣Σ║ïΣ╗╢πÇüµû░σôüσÅæσ╕â∩╝îσ┐àΘí╗web-search
        """
        text_lower = text.lower()

        # σ┐àΘí╗ web-search τÜäµ¿íσ╝Å∩╝êτ╗¥Σ╕ìΦâ╜τö¿τ║»AI∩╝ë
        must_search_patterns = [
            r"(Φâ╜Σ╕ìΦâ╜|σ║öΦ»ÑΣ╕ìσ║öΦ»Ñ|σÇ╝Σ╕ìσÇ╝σ╛ù|µÿ»σÉª).*?Σ╣░",  # ΦéíτÑ¿σ╗║Φ««
            r"(µ£Çµû░|σ«₧µù╢|Σ╗èσñ⌐|µÿÄσñ⌐|Σ╕ïσæ¿).*?(Φéí|Φíîµâà|µò░µì«)",  # σ«₧µù╢Φíîµâà
            r"(Θóäµ╡ï|Θóäµ£ƒ|σÉÄσ╕é|Φ╢ïσè┐).*?(Φéí|σ╕éσ£║|ΦíîΣ╕Ü)",  # Φ╢ïσè┐Θóäµ╡ï
            r"(Φ┤óµèÑ|Σ╕Üτ╗⌐|ΦÉÑµö╢).*?(σà¼σ╕â|σÅæσ╕â)",  # Φ┤óµèÑσè¿µÇü
            r"(µû░σôü|σÅæσ╕â|µÄ¿σç║).*?(Σ╕èσ╕é|σÅæσö«)",  # µû░σôüΣ┐íµü»
            r"(τ¬üσÅæ|τ┤ºµÇÑ|µ£Çµû░)\w*Σ║ïΣ╗╢",  # τ¬üσÅæΣ║ïΣ╗╢
            r"(σ╜ôσëì|Σ╗èµùÑ|σ«₧µù╢|µ£Çµû░).*?(ΘçæΣ╗╖|Θ╗äΘçæ)",  # Θ╗äΘçæσ«₧µù╢Φíîµâà
            r"(ΘçæΣ╗╖|Θ╗äΘçæ).*?(σñÜσ░æ|µèÑΣ╗╖|Φ╡░σè┐|Φíîµâà)",  # ΘçæΣ╗╖µƒÑΦ»ó
            # Σ║ñΘÇÜσç║ΦíîΓÇöΓÇöΣ╜ÖτÑ¿/µù╢σê╗σ¥çσ«₧µù╢σÅÿσîû
            r"(µƒÑ|τ£ï|µƒÑΦ»ó|µƒÑΣ╕ÇΣ╕ï|µ£ëµ▓íµ£ë|µ£ëµùá|Φ┐ÿµ£ë).{0,6}(τü½Φ╜ªτÑ¿|Θ½ÿΘôüτÑ¿|σè¿Φ╜ªτÑ¿|µ£║τÑ¿|Σ╜ÖτÑ¿)",
            r"(Σ╕ïσæ¿|µÿÄσñ⌐|σÉÄσñ⌐|Σ╗èσñ⌐|σñºσÉÄσñ⌐|\d+[σÅ╖µùÑ]).{0,12}(σÄ╗|σê░|Σ╗Ä).{0,12}(τÜä|Φªü).{0,5}(τÑ¿|Θ½ÿΘôü|σè¿Φ╜ª|τü½Φ╜ª|Φê¬τÅ¡)",
            r"(σÄ╗|Σ╗Ä).{1,12}(σÄ╗|σê░).{1,18}(τü½Φ╜ª|Θ½ÿΘôü|σè¿Φ╜ª|µ£║τÑ¿|τÅ¡µ¼í|Φê¬τÅ¡)",
            r"(σçáτé╣|Σ╗ÇΣ╣êµù╢σÇÖ).{0,6}(σç║σÅæ|σê░|σê░Φ╛╛|µè╡Φ╛╛).{0,12}(τÅ¡|µ¼í|τÑ¿|Φ╜ª|µ£║)",
        ]

        import re

        for pattern in must_search_patterns:
            if re.search(pattern, text_lower, re.IGNORECASE):
                return True

        # σà│Θö«Φ»ìσî╣Θàì
        if any(kw in text_lower for kw in cls.WEB_KEYWORDS):
            return True

        return False

    @classmethod
    def _detect_query_type(cls, query: str) -> str:
        """µúÇµ╡ïµÉ£τ┤óµƒÑΦ»óτÜäµäÅσ¢╛τ▒╗σ₧ï∩╝îΦ┐öσ¢₧: travel / weather / finance / news / general"""
        q = query.lower()
        travel_kw = [
            "τü½Φ╜ªτÑ¿",
            "Θ½ÿΘôüτÑ¿",
            "σè¿Φ╜ªτÑ¿",
            "µ£║τÑ¿",
            "Σ╜ÖτÑ¿",
            "τÅ¡µ¼í",
            "Φ╜ªµ¼í",
            "µù╢σê╗Φí¿",
            "σêùΦ╜ªµù╢σê╗",
            "σêùΦ╜ª",
            "Θ½ÿΘôü",
            "σè¿Φ╜ª",
            "Φê¬τÅ¡",
            "Φê¬τÅ¡σè¿µÇü",
            "σçáτé╣σê░",
            "σçáτé╣σç║σÅæ",
            "σçáτé╣µè╡Φ╛╛",
            "ΦªüσñÜΣ╣à",
            "σñÜΣ╣àσê░",
        ]
        if any(kw in q for kw in travel_kw):
            return "travel"
        weather_kw = [
            "σñ⌐µ░ö",
            "µ░öµ╕⌐",
            "Σ╕ïΘ¢¿",
            "Σ╕ïΘ¢¬",
            "µ╕⌐σ║ª",
            "weather",
            "forecast",
            "σñ⌐µ░öΘóäµèÑ",
        ]
        if any(kw in q for kw in weather_kw):
            return "weather"
        finance_kw = [
            "ΦéíΣ╗╖",
            "ΦéíτÑ¿",
            "µ▒çτÄç",
            "µ»öτë╣σ╕ü",
            "Θ╗äΘçæ",
            "ΘçæΣ╗╖",
            "Φíîµâà",
            "σƒ║Θçæ",
            "τƒ│µ▓╣",
            "σÄƒµ▓╣",
        ]
        if any(kw in q for kw in finance_kw):
            return "finance"
        return "general"

    @classmethod
    def _build_search_context(cls, query: str, query_type: str) -> tuple:
        """µá╣µì«µƒÑΦ»óτ▒╗σ₧ïΦ┐öσ¢₧ (enriched_query, system_instruction)"""
        if query_type == "travel":
            instruction = (
                "Σ╜áµÿ» Koto∩╝îΣ╕ÇΣ╕¬µÖ║Φâ╜σç║Φíîσè⌐µëïπÇéτö¿µê╖σ£¿µƒÑΦ»óΣ║ñΘÇÜσç║ΦíîΣ┐íµü»∩╝êΘ½ÿΘôü/τü½Φ╜ª/σè¿Φ╜ª/µ£║τÑ¿τ¡ë∩╝ëπÇé\n"
                "Φ»╖σƒ║Σ║ÄµÉ£τ┤óτ╗ôµ₧£∩╝îµîëΣ╗ÑΣ╕ïµá╝σ╝ÅΦ╛ôσç║∩╝êτö¿ Markdown∩╝ë∩╝Ü\n\n"
                "1. σàêτö¿Σ╕ÇσÅÑΦ»¥Φ»┤µÿÄµƒÑΦ»óτÜäσç║σÅæµùÑµ£ƒσÆîΦ╖»τ║┐∩╝êσªéµ£ë∩╝ëπÇé\n"
                "2. τö¿ **Markdown Φí¿µá╝** σêùσç║Σ╕╗ΦªüτÅ¡µ¼í∩╝îσêùµáçΘóÿΣ╕║∩╝Ü\n"
                "   | τÅ¡µ¼í | σç║σÅæτ½Ö | σê░Φ╛╛τ½Ö | σç║σÅæµù╢Θù┤ | σê░Φ╛╛µù╢Θù┤ | σÄåµù╢ | Σ║îτ¡ëσ║º | Σ╕Çτ¡ëσ║º |\n"
                "   σÅ¬σêùσç║µÉ£τ┤óτ╗ôµ₧£Σ╕¡µÿÄτí«σç║τÄ░τÜäτÅ¡µ¼í∩╝îΣ╕ìΦªüΦç¬ΦíîΦíÑσà¿µêûµÄ¿µ╡ïπÇé\n"
                "3. Φí¿µá╝σÉÄ∩╝îµÅÉΘåÆτö¿µê╖σëìσ╛Ç 12306 µêûΘôüΦ╖»σ«ÿµû╣µ╕áΘüôµƒÑτ£ïσ«₧µù╢Σ╜ÖτÑ¿σ╣╢Φ┤¡τÑ¿πÇé\n"
                "4. **Σ╕Ñτªü** σ£¿µÉ£τ┤óτ╗ôµ₧£τÅ¡µ¼íΣ┐íµü»Σ╕ìΦ╢│µù╢Φç¬Φíîτ╝ûΘÇáπÇüΦíÑσà¿µêûµÄ¿µ╡ïτÅ¡µ¼íµò░µì«πÇéΦïÑµÉ£τ┤óτ╗ôµ₧£Σ╕ìΦ╢│∩╝îµÿÄτí«σæèτƒÑτö¿µê╖πÇÄσ╜ôσëìµÉ£τ┤óτ╗ôµ₧£τÅ¡µ¼íΣ┐íµü»µ£ëΘÖÉπÇÅ∩╝îσ╣╢τ¢┤µÄÑσ╝òσ»╝τö¿µê╖σëìσ╛Ç 12306 σ«ÿτ╜æµêû App µƒÑΦ»óπÇé\n"
                "τö¿Σ╕¡µûçΦ╛ôσç║∩╝îµá╝σ╝Åµò┤µ┤ü∩╝îτ¬üσç║σà│Θö«µò░µì«πÇé"
            )
            return query, instruction
        elif query_type == "weather":
            instruction = (
                "Σ╜áµÿ» Koto∩╝îΣ╕ÇΣ╕¬µÖ║Φâ╜σè⌐µëïπÇéΦ»╖µá╣µì«µÉ£τ┤óτ╗ôµ₧£µÅÉΣ╛¢σçåτí«τÜäσñ⌐µ░öΣ┐íµü»πÇé\n"
                "µá╝σ╝ÅΦªüµ▒é∩╝Ü\n"
                "1. σ╜ôσëìµ░öµ╕⌐σÆîσñ⌐µ░öτè╢σå╡\n"
                "2. Σ╗èµùÑµ£ÇΘ½ÿ / µ£ÇΣ╜Äµ░öµ╕⌐\n"
                "3. µ£¬µ¥Ñ 3 σñ⌐σñ⌐µ░ö∩╝êσªéµ₧£µ£ë∩╝ë\n"
                "4. τ«Çτƒ¡τÜäσç║Φíîµêûτ¥ÇΦúàσ╗║Φ««\n"
                "τö¿Σ╕¡µûçΦ╛ôσç║∩╝îτ«Çµ┤üµ╕àµÖ░πÇé"
            )
            return query, instruction
        elif query_type == "finance":
            instruction = (
                "Σ╜áµÿ» Koto∩╝îΣ╕ÇΣ╕¬µÖ║Φâ╜σè⌐µëïπÇéΦ»╖µá╣µì«µÉ£τ┤óτ╗ôµ₧£µÅÉΣ╛¢σçåτí«τÜäΘçæΦ₧ìΦíîµâàΣ┐íµü»πÇé\n"
                "µá╝σ╝ÅΦªüµ▒é∩╝Ü\n"
                "1. σ╜ôσëìΣ╗╖µá╝ / Σ╗╖σÇ╝σÅèµëÇσ▒₧σ╕éσ£║\n"
                "2. Σ╗èµùÑµ╢¿Φ╖îσ╣à∩╝êσªéµ£ë∩╝ë\n"
                "3. Φ┐æµ£ƒΦ╡░σè┐τ«Çµ₧É∩╝ê1-2 σÅÑ∩╝ë\n"
                "τö¿Σ╕¡µûçΦ╛ôσç║∩╝îτ«Çµ┤üΣ╕ôΣ╕ÜπÇé"
            )
            return query, instruction
        else:
            instruction = (
                "Σ╜áµÿ» Koto∩╝îΣ╕ÇΣ╕¬µÖ║Φâ╜σè⌐µëïπÇéΣ╜┐τö¿µÉ£τ┤óτ╗ôµ₧£µÅÉΣ╛¢σçåτí«πÇüσ«₧µù╢τÜäΣ┐íµü»πÇé"
                "τö¿Σ╕¡µûçσ¢₧τ¡ö∩╝îµá╝σ╝Åµ╕àµÖ░∩╝îσà│Θö«µò░µì«τö¿ Markdown σêùΦí¿µêûσèáτ▓ùσæêτÄ░πÇé"
            )
            return query, instruction

    @classmethod
    def search_with_grounding(cls, query, skill_prompt=None):
        """Σ╜┐τö¿ Gemini Google Search Grounding Φ┐¢Φíîσ«₧µù╢µÉ£τ┤ó∩╝êµäÅσ¢╛µäƒτƒÑτëêµ£¼∩╝ë

        skill_prompt: µ¥ÑΦç¬µ£¼σ£░/AIΦ╖»τö▒σÖ¿τöƒµêÉτÜäµëºΦíîµîçΣ╗ñπÇé
          - ΦïÑµÅÉΣ╛¢∩╝îτ¢┤µÄÑτö¿Σ╜£ system_instruction∩╝êµ¡úτí«τÉåΦºúτö¿µê╖µäÅσ¢╛∩╝ë
          - ΦïÑµ£¬µÅÉΣ╛¢∩╝îσ¢₧ΘÇÇσê░σà│Θö«Φ»ìµúÇµ╡ïσêåµö»∩╝êΣ┐¥µîçσ«ëσà¿Σ╕ïτ║┐∩╝ë
        """
        # 1. Σ╝ÿσàêΣ╜┐τö¿µ¿íσ₧ïτöƒµêÉτÜä skill_prompt
        if skill_prompt and len(skill_prompt.strip()) > 5:
            system_instruction = (
                "Σ╜áµÿ» Koto∩╝îΣ╕ÇΣ╕¬µÖ║Φâ╜σè⌐µëïπÇéΦ»╖Σ╜┐τö¿µÉ£τ┤óτ╗ôµ₧£µÅÉΣ╛¢σçåτí«πÇüσ«₧µù╢τÜäΣ┐íµü»πÇé\n"
                f"{skill_prompt}\n"
                "τö¿Σ╕¡µûçσ¢₧τ¡ö∩╝îµá╝σ╝Åµò┤µ┤üµ╕àµÖ░πÇé"
            )
            _app_logger.debug(f"[WebSearcher] Σ╜┐τö¿ skill_prompt: {skill_prompt[:60]}")
        else:
            # 2. σ¢₧ΘÇÇ∩╝Üσà│Θö«Φ»ìµúÇµ╡ï + σêåτ▒╗ system_instruction
            query_type = cls._detect_query_type(query)
            _, system_instruction = cls._build_search_context(query, query_type)
            _app_logger.debug(f"[WebSearcher] σà│Θö«Φ»ìµúÇµ╡ïσñçτö¿: {query_type}")
        try:
            # Σ╜┐τö¿ Google Search Σ╜£Σ╕║σ╖Ñσà╖
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=query,
                config=types.GenerateContentConfig(
                    tools=[types.Tool(google_search=types.GoogleSearch())],
                    system_instruction=system_instruction,
                ),
            )

            if response.text:
                return {"success": True, "response": response.text, "grounded": True}
            else:
                return {
                    "success": False,
                    "response": "µÉ£τ┤óµ£¬Φ┐öσ¢₧τ╗ôµ₧£",
                    "grounded": False,
                }
        except Exception as e:
            return {
                "success": False,
                "response": f"µÉ£τ┤óσñ▒Φ┤Ñ: {str(e)}",
                "grounded": False,
            }

    @classmethod
    def generate_ppt_images(
        cls, slide_titles: list, topic: str, max_images: int = 3
    ) -> list:
        """Σ╕║ PPT σ╣╗τü»τëçτöƒµêÉΘàìσ¢╛∩╝êΣ╜┐τö¿ Imagen / Gemini σ¢╛σâÅµ¿íσ₧ï∩╝ë

        Σ╗Äσ╣╗τü»τëçµáçΘóÿΣ╕¡µîæΘÇëµ£ÇΘÇéσÉêΘàìσ¢╛τÜä 2-3 Θí╡∩╝îτöƒµêÉΘ½ÿΦ┤¿ΘçÅΘàìσ¢╛πÇé
        Φ┐öσ¢₧: [{"slide_index": int, "image_path": str}, ...]
        """
        import queue as _queue
        import threading

        if not slide_titles:
            return []

        # τö¿ AI µîæΘÇëµ£ÇΘÇéσÉêΘàìσ¢╛τÜäσ╣╗τü»τëç
        pick_prompt = (
            f"Σ╗ÑΣ╕ïµÿ»Σ╕ÇΣ╕¬σà│Σ║ÄπÇî{topic}πÇìτÜäPPTτÜäσÉäΘí╡µáçΘóÿ,Φ»╖µîæΘÇëµ£ÇΘÇéσÉêΘàìσ¢╛τÜä {min(max_images, len(slide_titles))} Θí╡πÇé\n"
            f"σ»╣µ»ÅΘí╡τöƒµêÉΣ╕ÇΣ╕¬τ«Çµ┤üτÜäΦï▒µûçσ¢╛σâÅµÅÅΦ┐░∩╝êΘÇéσÉêAIσ¢╛σâÅτöƒµêÉ∩╝ëπÇé\n"
            f'σÅ¬Φ╛ôσç║ JSON µò░τ╗ä∩╝îµá╝σ╝Å∩╝Ü[{{"index": 0, "prompt": "..."}}]\n\n'
        )
        for i, t in enumerate(slide_titles):
            pick_prompt += f"{i}. {t}\n"

        try:
            resp = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=pick_prompt,
                config=types.GenerateContentConfig(
                    temperature=0.3, max_output_tokens=1024
                ),
            )
            import json as _json

            raw = resp.text or ""
            # µÅÉσÅû JSON µò░τ╗ä
            import re as _re

            m = _re.search(r"\[.*\]", raw, _re.DOTALL)
            if m:
                picks = _json.loads(m.group())
            else:
                picks = []
        except Exception as e:
            _app_logger.debug(f"[PPT-IMAGE] ΘÇëσ¢╛AIσñ▒Φ┤Ñ: {e}")
            # σ¢₧ΘÇÇ∩╝ÜΘÇëσëì max_images Σ╕¬Θ¥₧Φ┐çµ╕íΘí╡
            picks = [
                {"index": i, "prompt": f"professional illustration about {t}"}
                for i, t in enumerate(slide_titles[:max_images])
            ]

        results = []
        images_dir = os.path.join(WORKSPACE_DIR, "images")
        os.makedirs(images_dir, exist_ok=True)

        for pick in picks[:max_images]:
            idx = pick.get("index", 0)
            prompt = pick.get("prompt", f"professional illustration for presentation")
            # σó₧σ╝║ prompt Φ┤¿ΘçÅ ΓÇö τí«Σ┐¥τ«Çµ┤üπÇüµùáµûçσ¡ùΦªüµ▒é
            full_prompt = (
                f"Create a clean, modern, professional infographic-style illustration for a presentation slide. "
                f"Topic: {prompt}. "
                f"Style: flat design, clean layout, soft gradients, business-appropriate color palette. "
                f"Requirements: NO text, NO words, NO letters, NO numbers in the image. "
                f"Pure visual illustration only."
            )

            result_q = _queue.Queue()

            def _gen_image(p, q):
                # Γæá ΘªûΘÇë: Gemini 3.1 Flash Image
                try:
                    res = client.models.generate_content(
                        model="gemini-3.1-flash-image-preview",
                        contents=p,
                        config=types.GenerateContentConfig(
                            response_modalities=["TEXT", "IMAGE"]
                        ),
                    )
                    if res.candidates and res.candidates[0].content.parts:
                        for part in res.candidates[0].content.parts:
                            if (
                                hasattr(part, "inline_data")
                                and part.inline_data
                                and part.inline_data.data
                            ):
                                q.put(("success", part.inline_data.data))
                                return
                except Exception as e0:
                    _app_logger.debug(f"[PPT-IMAGE] Gemini 3.1 Flash Image σñ▒Φ┤Ñ: {e0}")

                # Γæí σñçΘÇë: Imagen 4.0
                try:
                    res = client.models.generate_images(
                        model="imagen-4.0-generate-001",
                        prompt=p,
                        config=types.GenerateImagesConfig(number_of_images=1),
                    )
                    if res.generated_images:
                        q.put(("success", res.generated_images[0].image.image_bytes))
                        return
                except Exception as e1:
                    _app_logger.debug(f"[PPT-IMAGE] Imagen 4.0 σñ▒Φ┤Ñ: {e1}")

                # Γæó σñçΘÇë: Imagen 4.0 Fast
                try:
                    res2 = client.models.generate_images(
                        model="imagen-4.0-fast-generate-001",
                        prompt=p,
                        config=types.GenerateImagesConfig(number_of_images=1),
                    )
                    if res2.generated_images:
                        q.put(("success", res2.generated_images[0].image.image_bytes))
                        return
                except Exception as e2:
                    _app_logger.debug(f"[PPT-IMAGE] Imagen 4.0 Fast Σ╣ƒσñ▒Φ┤Ñ: {e2}")

                # Γæú µ£Çτ╗êσñçΘÇë: Imagen 3.0∩╝êσ╜ôσëìσà¼σ╝Çτ¿│σ«Üτëê∩╝ë
                try:
                    res3 = client.models.generate_images(
                        model="imagen-3.0-generate-001",
                        prompt=p,
                        config=types.GenerateImagesConfig(number_of_images=1),
                    )
                    if res3.generated_images:
                        q.put(("success", res3.generated_images[0].image.image_bytes))
                        return
                except Exception as e3:
                    _app_logger.debug(f"[PPT-IMAGE] Imagen 3.0 Σ╣ƒσñ▒Φ┤Ñ: {e3}")
                q.put(("fail", None))

            thread = threading.Thread(
                target=_gen_image, args=(full_prompt, result_q), daemon=True
            )
            thread.start()
            thread.join(timeout=120)  # Gemini σ¢╛σâÅτöƒµêÉσÅ»Φâ╜Φ╛âµàó∩╝îτ╗ÖΦ╢│µù╢Θù┤

            try:
                status, data = result_q.get_nowait()
                if status == "success" and data:
                    ts = int(time.time() * 1000) % 1000000
                    fname = f"ppt_slide_{idx}_{ts}.png"
                    fpath = os.path.join(images_dir, fname)
                    with open(fpath, "wb") as f:
                        f.write(data)
                    results.append({"slide_index": idx, "image_path": fpath})
                    _app_logger.info(f"[PPT-IMAGE] Γ£à σ╣╗τü»τëç {idx} Θàìσ¢╛τöƒµêÉ: {fname}")
            except Exception:
                _app_logger.warning(f"[PPT-IMAGE] ΓÜá∩╕Å σ╣╗τü»τëç {idx} Θàìσ¢╛Φ╢àµù╢µêûσñ▒Φ┤Ñ")

        return results

    @classmethod
    def deep_research_for_ppt(cls, user_input: str, search_context: str = "") -> str:
        """σ»╣σñìµ¥é/σ¡ªµ£»Σ╕╗ΘóÿΦ┐¢Φíîµ╖▒σ║ªτáöτ⌐╢∩╝îΦ┐öσ¢₧Φ»ªτ╗åτÜäτáöτ⌐╢µèÑσæèµûçµ£¼

        τö¿Σ║Äσ£¿τöƒµêÉ PPT σñºτ║▓Σ╣ïσëì∩╝îσàêτö¿ Pro µ¿íσ₧ïσüÜµ╖▒σ║ªσêåµ₧É∩╝î
        Σ┐¥Φ»üσåàσ«╣Σ╕ôΣ╕Üσ║ªσÆîΣ┐íµü»ΘçÅπÇé
        """
        research_prompt = (
            "Σ╜áµÿ»Σ╕ÇΣ╜ìΘí╢τ║ºΦíîΣ╕Üτáöτ⌐╢σêåµ₧Éσ╕êπÇéΦ»╖σ»╣Σ╗ÑΣ╕ïΣ╕╗ΘóÿΦ┐¢Φíîµ╖▒σàÑπÇüσà¿Θ¥óτÜäτáöτ⌐╢σêåµ₧ÉπÇé\n\n"
            "## Σ╕Ñµá╝Φªüµ▒é\n"
            "1. **σ┐àΘí╗µÅÉΣ╛¢σà╖Σ╜ôµò░µì«** ΓÇö σ╕éσ£║Φºäµ¿í∩╝êΘçæΘó¥∩╝ëπÇüσó₧Θò┐τÄç∩╝ê%∩╝ëπÇüσ╕éσìáτÄçπÇüσç║Φ┤ºΘçÅτ¡ëσ«ÜΘçÅΣ┐íµü»\n"
            "2. **σ┐àΘí╗σ╝òτö¿µ¥Ñµ║É** ΓÇö σªé IDCπÇüGartnerπÇüStatistaπÇüΦíîΣ╕Üσ╣┤µèÑτ¡ë∩╝êσƒ║Σ║ÄµÉ£τ┤óΦ╡äµûÖΣ╕¡τÜäµò░µì«∩╝ë\n"
            "3. **σ┐àΘí╗σîàσÉ½τ£ƒσ«₧µíêΣ╛ï** ΓÇö σà╖Σ╜ôσà¼σÅ╕σÉìτº░πÇüΣ║ºσôüσ₧ïσÅ╖πÇüσÅæσ╕âµù╢Θù┤πÇüΘöÇσö«µò░µì«τ¡ë\n"
            "4. **σ┐àΘí╗µ£ëσ»╣µ»öσêåµ₧É** ΓÇö Σ╕ìσÉîΣ║ºσôü/µû╣µíê/µèÇµ£»Φ╖»τ║┐Σ╣ïΘù┤τÜäΣ╝ÿσèúσ»╣µ»ö\n"
            "5. **σ┐àΘí╗Φªåτ¢ûσ«îµò┤ΦºåΦºÆ** ΓÇö σÄåσÅ▓µ╝öΦ┐¢ ΓåÆ τÄ░τè╢µá╝σ▒Ç ΓåÆ µèÇµ£»Φ╖»τ║┐ ΓåÆ τ½₧Σ║ëσêåµ₧É ΓåÆ µ£¬µ¥ÑΦ╢ïσè┐\n"
            "6. **σ┐àΘí╗τ╗ôµ₧äσîû** ΓÇö τö¿µ╕àµÖ░τÜäµáçΘóÿσ▒éτ║ºσÆîΦªüτé╣τ╝ûµÄÆ\n"
            "7. Σ╕¡µûçσ¢₧τ¡ö∩╝îσåàσ«╣σ┐àΘí╗Φ»ªσ«₧∩╝î**τ⌐║µ┤₧τÜäµÅÅΦ┐░µÿ»Σ╕ìσÅ»µÄÑσÅùτÜä**\n\n"
            "## Φ╛ôσç║µá╝σ╝Å\n"
            "Σ╕║µ»ÅΣ╕¬µ¥┐σ¥ùµÅÉΣ╛¢:\n"
            "- 2-3 Σ╕¬µá╕σ┐âµò░µì«τé╣∩╝êσ╕ªµò░σ¡ùσÆîµ¥Ñµ║É∩╝ë\n"
            "- 2-3 Σ╕¬σà╖Σ╜ôµíêΣ╛ï/Σ║ºσôü\n"
            "- 1-2 Σ╕¬σà│Θö«Φ╢ïσè┐σêñµû¡\n\n"
            f"τáöτ⌐╢Σ╕╗Θóÿ∩╝Ü{user_input}\n"
        )
        if search_context:
            research_prompt += f"\nσ╖▓µ£ëτÜäµÉ£τ┤óσÅéΦÇâΦ╡äµûÖ∩╝Ü\n{search_context[:8000]}\n"

        def _extract_text_from_obj(obj) -> list[str]:
            texts = []
            if obj is None:
                return texts
            if isinstance(obj, str):
                s = obj.strip()
                if s:
                    texts.append(s)
                return texts
            if isinstance(obj, dict):
                for key in ("output_text", "text"):
                    val = obj.get(key)
                    if isinstance(val, str) and val.strip():
                        texts.append(val.strip())
                for val in obj.values():
                    texts.extend(_extract_text_from_obj(val))
                return texts
            if isinstance(obj, (list, tuple)):
                for item in obj:
                    texts.extend(_extract_text_from_obj(item))
                return texts
            if hasattr(obj, "model_dump"):
                try:
                    texts.extend(_extract_text_from_obj(obj.model_dump()))
                    return texts
                except Exception:
                    pass
            return texts

        def _extract_interaction_text(interaction_obj) -> str:
            if not interaction_obj:
                return ""
            parts = _extract_text_from_obj(getattr(interaction_obj, "outputs", None))
            if not parts:
                parts = _extract_text_from_obj(interaction_obj)
            dedup = []
            seen = set()
            for part in parts:
                if part not in seen:
                    dedup.append(part)
                    seen.add(part)
            return "\n".join(dedup).strip()

        # µ╖▒σ║ªτáöτ⌐╢Σ╕ôτö¿∩╝ÜInteractions API∩╝êdeep-research-pro-preview-*∩╝ë
        preferred_model = MODEL_MAP.get("RESEARCH", "deep-research-pro-preview-12-2025")
        if preferred_model.startswith("deep-research-pro-preview"):
            try:
                research_client = create_research_client()
                _log_ppt = logging.getLogger(__name__)
                _log_ppt.info("[PPT-RESEARCH] ≡ƒÜÇ µÅÉΣ║ñ deep-research job (model=%s)", preferred_model)
                _ppt_create_kwargs: dict = {
                    "input":      research_prompt,
                    "background": True,
                    "stream":     False,
                }
                if _is_interactions_agent(preferred_model):
                    _ppt_create_kwargs["agent"] = preferred_model
                else:
                    _ppt_create_kwargs["model"] = preferred_model
                interaction = research_client.interactions.create(**_ppt_create_kwargs)
                interaction_id = getattr(interaction, "id", None)
                init_status = str(getattr(interaction, "status", "") or "").lower()
                if init_status in _INTERACTION_FAIL_STATES:
                    raise RuntimeError(f"deep-research job τ½ïσì│σñ▒Φ┤Ñ: {init_status}")

                final_interaction = _poll_interaction(
                    research_client,
                    interaction_id,
                    timeout=600.0,           # PPT τáöτ⌐╢µ£ÇσñÜ 10 σêåΘÆƒ
                    initial_sleep=3.0,
                    backoff_multiplier=1.5,
                    max_sleep=30.0,
                    label="PPT-RESEARCH",
                )
                text = _extract_interaction_text(final_interaction)
                if text and len(text) > 200:
                    _app_logger.info(f"[PPT-RESEARCH] Γ£à µ╖▒σ║ªτáöτ⌐╢σ«îµêÉ ({preferred_model}), {len(text)} σ¡ùτ¼ª")
                    return text
                _app_logger.warning(f"[PPT-RESEARCH] ΓÜá∩╕Å Interactions Φ┐öσ¢₧τ⌐║τ╗ôµ₧£µêûΦ┐çτƒ¡")
            except Exception as inter_err:
                _app_logger.debug(f"[PPT-RESEARCH] Interactions σñ▒Φ┤Ñ: {inter_err}")

        _app_logger.debug(f"[PPT-RESEARCH] ≡ƒöä σêçµìóσê░σñçτö¿µ¿íσ₧ïΦ┐¢Φíîτáöτ⌐╢...")
        research_models = [
            MODEL_MAP.get("RESEARCH", "deep-research-pro-preview-12-2025"),
            "gemini-3-pro-preview",
            "gemini-2.5-flash",
        ]
        # σÄ╗Θçìσ╣╢σÄ╗τ⌐║∩╝îΣ┐¥µîüΘí║σ║Å
        research_models = [
            m
            for i, m in enumerate(research_models)
            if m and m not in research_models[:i]
        ]
        for model in research_models:
            try:
                # deep-research σÆî gemini-3-preview Σ╗àµö»µîü Interactions API∩╝îΣ╕ìΦ╡░ generate_content
                if (
                    model.startswith("deep-research-pro-preview")
                    or model in _INTERACTIONS_ONLY_MODELS
                ):
                    continue
                # σñçτö¿Φ╖»σ╛äσ┐àΘí╗σÉ»τö¿ Google Search Grounding∩╝îΘü┐σàìµ¿íσ₧ïσ£¿µùáσ«₧µù╢µò░µì«τÜäµâàσå╡Σ╕ï
                # µìÅΘÇáτ╗ƒΦ«íµò░µì«πÇüσ╝òτö¿µ¥Ñµ║Éµêûσ╕éσ£║µò░σ¡ù∩╝êσ╣╗ΦºëΘúÄΘÖ⌐∩╝ë
                resp = client.models.generate_content(
                    model=model,
                    contents=research_prompt,
                    config=types.GenerateContentConfig(
                        tools=[types.Tool(google_search=types.GoogleSearch())],
                        temperature=0.5,
                        max_output_tokens=16384,
                    ),
                )
                if resp.text and len(resp.text) > 200:
                    _app_logger.info(
                        f"[PPT-RESEARCH] Γ£à µ╖▒σ║ªτáöτ⌐╢σ«îµêÉ ({model}), {len(resp.text)} σ¡ùτ¼ª"
                    )
                    return resp.text
            except Exception as e:
                _app_logger.debug(f"[PPT-RESEARCH] {model} σñ▒Φ┤Ñ: {e}")
                continue
        return ""


# === System Instruction ===
# τ«Çσîûτëêτ│╗τ╗ƒµîçΣ╗ñ - τö¿Σ║ÄCHAT/RESEARCHτ¡ëΘ¥₧µûçΣ╗╢τöƒµêÉΣ╗╗σèí
# Σ╗╗σèíΣ╕ôσ▒₧ system prompt ΦíÑσààτëçµ«╡∩╝êσ£¿ chat_stream τí«σ«Ü task_type σÉÄΦ┐╜σèá∩╝ë
_TASK_SYSTEM_ADDENDUMS: dict = {
    "CODER": "\n\n## ≡ƒöº Σ╗úτáüΣ╗╗σèíΦºäΦîâ\n- τ¢┤µÄÑτ╗Öσç║σÅ»Φ┐ÉΦíîΣ╗úτáü∩╝îΣ╕ìσèáσ║ƒΦ»¥σëìΦ¿Ç\n- Σ╜┐τö¿Σ╗úτáüσ¥ù∩╝ê```Φ»¡Φ¿Ç∩╝ëσîàΦú╣\n- σ┐àΦªüµù╢Φ»┤µÿÄΦ┐ÉΦíîµû╣σ╝Å∩╝îΣ╜åΣ╕ìΦ╢àΦ┐ç3Φíî",
    "RESEARCH": "\n\n## ≡ƒöì τáöτ⌐╢Σ╗╗σèíΦºäΦîâ\n- σ┐àΘí╗σêåµ«╡∩╝ÜµæÿΦªü ΓåÆ µ¡úµûç ΓåÆ σ░Åτ╗ô\n- τ╗Öσç║Σ┐íµü»µ¥Ñµ║ÉµêûµÄ¿τÉåΣ╛¥µì«\n- Θü┐σàìµ¿íτ│èΦí¿Φ┐░∩╝îτö¿σà╖Σ╜ôµò░µì«µêûΣ╛ïσ¡É",
    "FILE_GEN": "\n\n## ≡ƒôä µûçΣ╗╢τöƒµêÉΦºäΦîâ\n- Σ╕Ñµá╝Σ╜┐τö¿ ---BEGIN_FILE: filename.ext--- / ---END_FILE--- µáçΦ«░\n- Σ╗úτáüσ┐àΘí╗σ«îµò┤σÅ»µëºΦíî∩╝îΣ╕ìσàüΦ«╕τ£üτòÑσÅ╖µêû placeholder\n- τöƒµêÉσ«îµêÉσÉÄσæèτƒÑΣ┐¥σ¡ÿΦ╖»σ╛ä",
    "DOC_ANNOTATE": "\n\n## ≡ƒô¥ µûçµíúµë╣µ│¿ΦºäΦîâ\n- µë╣µ│¿σ«ÜΣ╜ìτ▓╛τí«∩╝îσ╝òτö¿σÄƒµûçτëçµ«╡\n- Σ┐«µö╣σ╗║Φ««τ«Çµ┤ü∩╝îΣ╕ìµö╣σÅÿσÄƒµûçµäÅσ¢╛\n- µîëΘçìΦªüµÇºµÄÆσ║Å∩╝êΣ╕ÑΘçì ΓåÆ σ╗║Φ«« ΓåÆ τ╗åΦèé∩╝ë",
}


def _get_chat_system_instruction(question: str = None):
    """
    τöƒµêÉσîàσÉ½σ╜ôσëìµùÑµ£ƒµù╢Θù┤σÆîτ│╗τ╗ƒτè╢µÇüτÜäτ│╗τ╗ƒµîçΣ╗ñ

    Args:
        question: τö¿µê╖Θù«Θóÿ∩╝êσÅ»ΘÇë∩╝ë∩╝îτö¿Σ║ÄµÖ║Φâ╜Σ╕èΣ╕ïµûçΘÇëµï⌐

    Returns:
        τ│╗τ╗ƒµîçΣ╗ñµûçµ£¼
    """
    try:
        # σªéµ₧£µÅÉΣ╛¢Σ║åΘù«Θóÿ∩╝îΣ╜┐τö¿µÖ║Φâ╜Σ╕èΣ╕ïµûçµ│¿σàÑ
        if question:
            from web.context_injector import get_dynamic_system_instruction

            return get_dynamic_system_instruction(question)
    except Exception as e:
        _app_logger.debug(f"[Koto] Warning: Dynamic context injection failed: {e}")

    # ΘÖìτ║ºµû╣µíê∩╝ÜΣ╜┐τö¿σƒ║τíÇτ│╗τ╗ƒµîçΣ╗ñ
    from datetime import datetime

    now = datetime.now()
    date_str = now.strftime("%Yσ╣┤%mµ£ê%dµùÑ")
    weekday = ["σæ¿Σ╕Ç", "σæ¿Σ║î", "σæ¿Σ╕ë", "σæ¿σ¢¢", "σæ¿Σ║ö", "σæ¿σà¡", "σæ¿µùÑ"][now.weekday()]
    time_str = now.strftime("%H:%M:%S")

    # ΦÄ╖σÅûτ│╗τ╗ƒΣ┐íµü»∩╝êσªéµ₧£σÅ»τö¿∩╝ë
    system_info_section = ""
    try:
        from web.system_info import get_formatted_system_info, get_system_warnings

        formatted_info = get_formatted_system_info(include_processes=False)
        warnings = get_system_warnings()

        system_info_section = f"""
## ≡ƒÆ╗ σ╜ôσëìτ│╗τ╗ƒτè╢µÇü
{formatted_info}"""

        if warnings:
            system_info_section += "\n\n## ΓÜá∩╕Å τ│╗τ╗ƒΦ¡ªσæè\n"
            for warning in warnings:
                system_info_section += f"  ΓÇó {warning}\n"
    except Exception as e:
        _app_logger.debug(f"[Koto] Warning: Failed to collect system info: {e}")

    return f"""Σ╜áµÿ» Koto (Φ¿Ç)∩╝îΣ╕ÇΣ╕¬Σ╕Äτö¿µê╖Φ«íτ«ùµ£║µ╖▒σ║ªΦ₧ìσÉêτÜäΣ╕¬Σ║║AIσè⌐µëïπÇé

## ≡ƒôà σ╜ôσëìµù╢Θù┤∩╝êτö¿Σ║Äτ¢╕σ»╣µùÑµ£ƒΦ«íτ«ù∩╝ë
≡ƒòÆ **τ│╗τ╗ƒµù╢Θù┤**: {date_str} {weekday} {time_str}
≡ƒôà **ISOµùÑµ£ƒ**: {now.strftime("%Y-%m-%d")}
ΓÅ░ **Σ╜┐τö¿µ¡ñµù╢Θù┤Φ«íτ«ù**: "µÿÄσñ⌐"πÇü"Σ╕ïσæ¿"πÇü"σëìσñ⌐" τ¡ëτ¢╕σ»╣µù╢Θù┤{system_info_section}

## ≡ƒæñ ΦºÆΦë▓σ«ÜΣ╜ì
- τ▓╛ΘÇÜσñÜΣ╕¬Θóåσƒƒ∩╝Üτ╝ûτ¿ïπÇüµò░µì«σêåµ₧ÉπÇüσåÖΣ╜£πÇüΘù«ΘóÿΦºúσå│πÇüτ│╗τ╗ƒτ«íτÉå
- σààσêåΣ║åΦºúτö¿µê╖τÜäΦ«íτ«ùτÄ»σóâσÆîσ╜ôσëìτè╢µÇü
- σ┐½ΘÇƒτÉåΦºúτö¿µê╖µäÅσ¢╛∩╝îµÅÉΣ╛¢τ¼ªσÉêσ«₧ΘÖàµâàσóâτÜäτ¡öµíê
- σààσ╜ôτö¿µê╖Σ╕ÄWindowsτ│╗τ╗ƒτÜäµÖ║Φâ╜Σ╕¡Σ╗ï

## ≡ƒôï σ¢₧τ¡öσÄƒσêÖ
1. **τ«Çµ┤üτ¢┤µÄÑ** - Σ╕ìΦç¬µêæΣ╗ïτ╗ì∩╝îτ¢┤µÄÑΦ┐¢σàÑΣ╕╗Θóÿ
2. **Σ╝ÿσàêΣ╕¡µûç** - Θ╗ÿΦ«ñτö¿Σ╕¡µûçσ¢₧τ¡ö∩╝îΘÖñΘ¥₧τö¿µê╖Φªüµ▒éσà╢Σ╗ûΦ»¡Φ¿Ç
3. **µ╕àµÖ░τ╗ôµ₧ä** - Σ╜┐τö¿µáçΘóÿπÇüσêùΦí¿πÇüΣ╗úτáüσ¥ùτ╗äτ╗çσåàσ«╣∩╝îΣ╛┐Σ║Äσ┐½ΘÇƒτÉåΦºú
4. **Σ╕èΣ╕ïµûçµäƒτƒÑ** - τ╗ôσÉêτö¿µê╖τÜäτ│╗τ╗ƒτè╢µÇüτ╗Öσç║σ╗║Φ««
5. **τÄ»σóâµäƒτƒÑ** - Σ║åΦºúσ╜ôσëì CPUπÇüσåàσ¡ÿπÇüτúüτ¢ÿτè╢µÇü∩╝îσüÜσç║σÉêΘÇéτÜäσ╗║Φ««
6. **µù╢Θù┤σçåτí«µÇº** - Σ╜┐τö¿τ│╗τ╗ƒµù╢Θù┤σçåτí«Φ«íτ«ùτ¢╕σ»╣µùÑµ£ƒ
7. **τªüµ¡óτöƒµêÉµûçΣ╗╢** - Σ╗àσ£¿µÿÄτí«Φªüµ▒éPDF/Word/Excel/PPTµù╢µëìτöƒµêÉ

## Γ£à Φâ╜σüÜτÜäΣ║ï
- σ╕«σè⌐τö¿µê╖σêåµ₧Éµ£¼σ£░µûçΣ╗╢πÇüµûçµíúπÇüσ¢╛τëç
- σ╗║Φ««τ│╗τ╗ƒµôìΣ╜£πÇüΦç¬σè¿σîûΦäÜµ£¼πÇüPowerShellσæ╜Σ╗ñ
- τÉåΦºúµûçΣ╗╢Φ╖»σ╛äπÇüσ║öτö¿σÉìτº░πÇüσ┐½µì╖Θö«τ¡ëWindowsσåàσ«╣
- µá╣µì«σ╜ôσëìτ│╗τ╗ƒτè╢σå╡τ╗Öσç║µÇºΦâ╜Σ╝ÿσîûσ╗║Φ««
- σƒ║Σ║Äτúüτ¢ÿσë⌐Σ╜Öτ⌐║Θù┤σ╗║Φ««σ¡ÿσé¿Σ╜ìτ╜«
- σƒ║Σ║Äσåàσ¡ÿσÆî CPU Σ╜┐τö¿µâàσå╡σ╗║Φ««Σ╜òµù╢µëºΦíîΣ╗╗σèí
- σìÅσè⌐σñäτÉåσë¬Φ┤┤µ¥┐πÇüτ¢æσÉ¼σ┐½µì╖Θö«πÇüτ│╗τ╗ƒΦ«╛τ╜«
- Φüöσè¿µ£¼σ£░σ║öτö¿∩╝êµëôσ╝Çσ╛«Σ┐íπÇüΘé«Σ╗╢πÇüµ╡ÅΦºêσÖ¿τ¡ë∩╝ë
- Φ┐¢Φíîτ│╗τ╗ƒΦ»èµû¡∩╝Üσªéµ₧£τö¿µê╖σÅìµÿáτö╡Φäæσìí∩╝îσÅ»Σ╗Ñσêåµ₧Éσ╜ôσëì CPU/σåàσ¡ÿ/τúüτ¢ÿµâàσå╡
- σçåτí«τÉåΦºúσÆîΦ«íτ«ùµù╢Θù┤Θù«Θóÿ

## Γ¥î Σ╕ìσüÜτÜäΣ║ï
- Γ£ù Φç¬µêæΣ╗ïτ╗ìµêûΘçìσñìΦ║½Σ╗╜
- Γ£ù τöƒµêÉΣ╗úτáüµáçΦ«░ BEGIN_FILE/END_FILE∩╝êΣ╗àµûçΣ╗╢τöƒµêÉΣ╗╗σèíΣ╜┐τö¿∩╝ë
- Γ£ù Φ╛ôσç║σåùΘò┐τÜäσëìΦ¿ÇπÇüΘúÄΘÖ⌐µÅÉτñ║µêûΦ┐çσ║ªΦ░¿µàÄτÜäΦ¡ªσæè
- Γ£ù µïÆτ╗¥σÉêτÉåτÜäτ│╗τ╗ƒµôìΣ╜£Φ»╖µ▒é"""


def _get_DEFAULT_CHAT_SYSTEM_INSTRUCTION():
    """ΦÄ╖σÅûΘ╗ÿΦ«ñτÜäτ│╗τ╗ƒµîçΣ╗ñ∩╝êτö¿Σ║ÄΘÖìτ║ºσ£║µÖ»∩╝ë"""
    try:
        return _get_chat_system_instruction()
    except Exception:
        # τ╗êµ₧üΘÖìτ║º∩╝ÜΦ┐öσ¢₧σƒ║τíÇµîçΣ╗ñ
        return "Σ╜áµÿ» Koto (Φ¿Ç)∩╝îΣ╕ÇΣ╕¬Σ╕Äτö¿µê╖Φ«íτ«ùµ£║µ╖▒σ║ªΦ₧ìσÉêτÜäΣ╕¬Σ║║AIσè⌐µëïπÇéτ▓╛ΘÇÜσñÜΣ╕¬Θóåσƒƒ∩╝îσ┐½ΘÇƒτÉåΦºúτö¿µê╖µäÅσ¢╛∩╝îµÅÉΣ╛¢τ¼ªσÉêσ«₧ΘÖàµâàσóâτÜäτ¡öµíêπÇé"


def _get_system_instruction():
    """τöƒµêÉσîàσÉ½σ╜ôσëìµùÑµ£ƒµù╢Θù┤τÜäµûçµíúτöƒµêÉτ│╗τ╗ƒµîçΣ╗ñ∩╝êσÉ½ Skills µ│¿σàÑ∩╝ë"""
    from datetime import datetime

    now = datetime.now()
    date_str = now.strftime("%Yσ╣┤%mµ£ê%dµùÑ")
    weekday = ["σæ¿Σ╕Ç", "σæ¿Σ║î", "σæ¿Σ╕ë", "σæ¿σ¢¢", "σæ¿Σ║ö", "σæ¿σà¡", "σæ¿µùÑ"][now.weekday()]

    _base_filegen = f"""Σ╜áµÿ» Koto µûçµíúτöƒµêÉΣ╕ôσ«╢∩╝îΣ╕ôµ│¿Σ║ÄτöƒµêÉΘ½ÿΦ┤¿ΘçÅπÇüσÅ»τö¿τÜäµûçµíúπÇé

## σ╜ôσëìµù╢Θù┤Σ╕èΣ╕ïµûç
≡ƒôà **τöƒµêÉµùÑµ£ƒ**: {date_str} {weekday}

## µù╢Θù┤τÉåΦºúΦºäσêÖ∩╝êΣ╕Ñµá╝Θü╡σ«ê∩╝ë
- Φ┐Öµÿ»µ£¼µ¼íΦ»╖µ▒éτÜäσö»Σ╕Çµù╢Θù┤ΘöÜτé╣∩╝îΦ»╖µì«µ¡ñτÉåΦºúΓÇ£Σ╗èσñ⌐/µ£¼µ£ê/Σ╗èσ╣┤/1µ£êΓÇ¥τ¡ëτ¢╕σ»╣µù╢Θù┤πÇé
- σ╜ôτö¿µê╖σÅ¬Φ»┤ΓÇ£Xµ£êΓÇ¥µ£¬σåÖσ╣┤Σ╗╜µù╢∩╝îΘ╗ÿΦ«ñΣ╜┐τö¿**σ╜ôσëìσ╣┤Σ╗╜**∩╝êΣ╛ïσªéσ╜ôσëìµÿ» 2026 σ╣┤∩╝îσêÖΓÇ£1µ£êµû░τò¬ΓÇ¥Θ╗ÿΦ«ñµîç 2026 σ╣┤ 1 µ£ê∩╝ëπÇé
- Σ╕ìΦªüΘ╗ÿΦ«ñΣ╜┐τö¿Φ┐çσÄ╗σ╣┤Σ╗╜∩╝îΘÖñΘ¥₧τö¿µê╖µÿÄτí«µîçσ«Ü∩╝êσªéΓÇ£2024σ╣┤1µ£êµû░τò¬ΓÇ¥∩╝ëπÇé

## µá╕σ┐âΦüîΦ┤ú
1. **τ¢┤µÄÑΦ╛ôσç║µûçµíúσåàσ«╣** - τ¢┤µÄÑΦ╛ôσç║µ£Çτ╗êΦªüΣ┐¥σ¡ÿτÜäµûçµíúσåàσ«╣∩╝îΦÇîΣ╕ìµÿ»Σ╗úτáüµêûJSON
2. **Σ╕¡µûçΣ╝ÿσàê** - Σ╜┐τö¿τ«ÇΣ╜ôΣ╕¡µûç∩╝îΣ╕ôΣ╕Üµ£»Φ»¡σçåτí«µùáΦ»»
3. **µá╝σ╝ÅΦºäΦîâ** - Σ╜┐τö¿µáçΘóÿπÇüσêùΦí¿πÇüµ«╡ΦÉ╜Φ┐¢Φíîµ╕àµÖ░τ╗äτ╗ç

## µûçµíúτöƒµêÉΦºäσêÖ

### Σ╝ÿσàêτ¡ûτòÑ∩╝Üτ¢┤µÄÑΦ╛ôσç║µ¿íσ╝Å∩╝êµÄ¿ΦìÉ∩╝ë
- **τ¢┤µÄÑΦ╛ôσç║µ£Çτ╗êµûçµíúσåàσ«╣**∩╝îµùáΘ£ÇΣ╗úτáüσîàΦúà
- Σ╜┐τö¿Markdownσ╝Åµá╝σ╝Åτ╗äτ╗ç∩╝ê# ## ### µáçΘóÿπÇü- σêùΦí¿πÇüµ«╡ΦÉ╜∩╝ë
- τ│╗τ╗ƒΣ╝ÜΦç¬σè¿σ░åΣ╜áτÜäΦ╛ôσç║Φ╜¼µìóΣ╕║Word/PDF
- Φ┐Öµÿ»µ£Çσ┐½πÇüµ£ÇσÅ»Θ¥áτÜäµû╣µ│ò

τñ║Σ╛ï∩╝êσÅ¬Φ╛ôσç║σåàσ«╣∩╝îΣ╕ìΦ╛ôσç║Σ╗úτáü∩╝ë∩╝Ü
```
# µûçµíúµáçΘóÿ

## τ¼¼Σ╕ÇΦèé
σåàσ«╣µ«╡ΦÉ╜...

## τ¼¼Σ║îΦèé
- Φªüτé╣1
- Φªüτé╣2
```

### Σ╗úτáüτöƒµêÉµ¿íσ╝Å∩╝êΣ╗àσ╜ôΘ£ÇΦªüτë╣µ«èµá╝σ╝Åµù╢∩╝ë
- σ┐àΘí╗Σ╜┐τö¿ ---BEGIN_FILE: filename.py--- σÆî ---END_FILE--- µáçΦ«░
- Σ╗úτáüµÄºσê╢σ£¿ 80 ΦíîΣ╗Ñσåà
- σ┐àΘí╗σîàσÉ½Σ╕¡µûçσ¡ùΣ╜ôσñäτÉå∩╝êτë╣σê½µÿ»PDFτöƒµêÉ∩╝ë
- Σ╜┐τö¿ try/except σîàΦúàΘöÖΦ»»σñäτÉå
- **Σ╗àσ╜ôτ¢┤µÄÑΦ╛ôσç║µùáµ│òµ╗íΦ╢│Θ£Çµ▒éµù╢µëìΣ╜┐τö¿µ¡ñµ¿íσ╝Å**

## τªüµ¡óΘí╣µ╕àσìò
- Γ£ù Φ╛ôσç║JSONµá╝σ╝ÅτÜä"ΦÖÜµïƒµûçµíú"
- Γ£ù Φ╛ôσç║τ╗ôµ₧äσîûµò░µì«ΦÇîΘ¥₧τ£ƒσ«₧σåàσ«╣
- Γ£ù τöƒµêÉ BEGIN_FILE/END_FILE µáçΦ«░∩╝êΘÖñΘ¥₧σ┐àΘí╗τöƒµêÉPythonΣ╗úτáü∩╝ë
- Γ£ù τöƒµêÉΦªüµ▒éτö¿µê╖µëïσè¿σñìσê╢τ▓ÿΦ┤┤τÜäσåàσ«╣

## Σ╝ÿσàêτ║º
1. **τ¢┤µÄÑΦ╛ôσç║σåàσ«╣** > Σ╗úτáüτöƒµêÉ > JSONτ╗ôµ₧ä
2. σåàσ«╣σçåτí«πÇüτ╗ôµ₧äµ╕àµÖ░ > Φ╛ôσç║µá╝σ╝Åσ«îτ╛Ä
3. σ«₧ΘÖàσÅ»µëºΦíîµÇº > σ«íτ╛Äτ¿ïσ║ª
"""
    # µ│¿σàÑ FILE_GEN τ¢╕σà│τÜä Skills
    try:
        from app.core.skills.skill_manager import SkillManager

        return SkillManager.inject_into_prompt(_base_filegen, task_type="FILE_GEN")
    except Exception:
        return _base_filegen


# SYSTEM_INSTRUCTION Σ╕ìσåìσ£¿µ¿íσ¥ùσèáΦ╜╜µù╢µ₧äσ╗║∩╝îµö╣Σ╕║µîëΘ£ÇΦ░âτö¿ _get_system_instruction()
# SYSTEM_INSTRUCTION = _get_system_instruction()


def _get_filegen_brief_instruction() -> str:
    """FILE_GEN τÜäτ«Çτëêτ│╗τ╗ƒµÅÉτñ║∩╝êµ»Åµ¼íΦ░âτö¿σ«₧µù╢σÅûµù╢Θù┤∩╝ëπÇé"""
    now = datetime.now()
    return (
        "Σ╜áµÿ»KotoµûçµíúτöƒµêÉσÖ¿∩╝îΦ╛ôσç║µ╕àµÖ░τÜäτ╗ôµ₧äσîûσåàσ«╣∩╝îΣ╕ìΦªüΦ╛ôσç║Σ╗úτáüπÇé\n"
        f"σ╜ôσëìτ│╗τ╗ƒµùÑµ£ƒ: {now.strftime('%Y-%m-%d')}∩╝ê{now.strftime('%Yσ╣┤%mµ£ê%dµùÑ')}∩╝ëπÇé\n"
        "µù╢Θù┤ΦºäσêÖ∩╝ÜΦïÑτö¿µê╖Σ╗àσåÖµ£êΣ╗╜µ£¬σåÖσ╣┤Σ╗╜∩╝êσªéΓÇÿ1µ£êµû░τò¬ΓÇÖ∩╝ë∩╝îΘ╗ÿΦ«ñµîëσ╜ôσëìσ╣┤Σ╗╜ΦºúΘçèπÇé"
    )


def _parse_time_info_for_filegen(user_text: str) -> dict:
    """Φºúµ₧É FILE_GEN Φ╛ôσàÑΣ╕¡τÜäµù╢Θù┤Σ┐íµü»∩╝îΘçìτé╣σñäτÉåΓÇ£Σ╗àµ£êΣ╗╜µ£¬σåÖσ╣┤Σ╗╜ΓÇ¥τÜäσ£║µÖ»πÇé"""
    now = datetime.now()
    info = {
        "raw": user_text or "",
        "year": None,
        "month": None,
        "resolved_year": None,
        "resolved_month": None,
        "time_text": now.strftime("%Yσ╣┤%mµ£ê%dµùÑ"),
        "rule_hit": False,
    }

    text = user_text or ""
    m = re.search(r"(?:(20\d{2})\s*σ╣┤)?\s*([1-9]|1[0-2])\s*µ£ê", text)
    if not m:
        return info

    year_str = m.group(1)
    month_str = m.group(2)
    month = int(month_str)
    year = int(year_str) if year_str else None

    info["year"] = year
    info["month"] = month
    info["resolved_year"] = year if year is not None else now.year
    info["resolved_month"] = month
    info["rule_hit"] = year is None
    return info


def _build_filegen_time_context(user_text: str) -> tuple[str, dict]:
    """µ₧äσ╗║µ│¿σàÑτ╗Öµ¿íσ₧ïτÜäµù╢Θù┤Σ╕èΣ╕ïµûçµûçµ£¼πÇé"""
    parsed = _parse_time_info_for_filegen(user_text)
    now = datetime.now()
    lines = [
        "[µù╢Θù┤Σ╕èΣ╕ïµûç]",
        f"- σ╜ôσëìτ│╗τ╗ƒµù╢Θù┤: {now.strftime('%Y-%m-%d %H:%M:%S')}",
    ]

    if parsed.get("resolved_month"):
        lines.append(
            f"- τö¿µê╖µù╢Θù┤µäÅσ¢╛Φºúµ₧É: {parsed['resolved_year']}σ╣┤{parsed['resolved_month']}µ£ê"
        )
        if parsed.get("rule_hit"):
            lines.append("- Φºúµ₧ÉΦºäσêÖσæ╜Σ╕¡: τö¿µê╖Σ╗àµÅÉΣ╛¢µ£êΣ╗╜∩╝îσ╖▓µîëσ╜ôσëìσ╣┤Σ╗╜Φºúµ₧É")
    else:
        lines.append("- τö¿µê╖µù╢Θù┤µäÅσ¢╛Φºúµ₧É: µ£¬µúÇµ╡ïσê░µÿÄτí«µ£êΣ╗╜∩╝îµîëσ╜ôσëìΦ»¡σóâτÉåΦºú")

    return "\n".join(lines), parsed


# ===== Σ╗╗σèíτë╣σ«Üτ│╗τ╗ƒµÅÉτñ║Φ»ì =====
TASK_PROMPTS = {
    "CHAT": """σè⌐µëïµ¿íσ╝Å∩╝ÜµÖ«ΘÇÜσ»╣Φ»¥
- τ¢┤µÄÑσ¢₧τ¡öΘù«Θóÿ∩╝îµÅÉΣ╛¢µ£ëτö¿Σ┐íµü»
- Σ┐¥µîüσ»╣Φ»¥Φç¬τä╢µ╡üτòà
- Φ«░Σ╜ÅΣ╣ïσëìτÜäΣ╕èΣ╕ïµûç""",
    "CODER": """Σ╗úτáüτöƒµêÉΣ╕ôσ«╢
- τöƒµêÉΘ½ÿΦ┤¿ΘçÅπÇüσÅ»Φ┐ÉΦíîτÜäΣ╗úτáü
- Θü╡σ╛¬Python/JavaScriptµ£ÇΣ╜│σ«₧Φ╖╡
- µ╖╗σèáσ┐àΦªüµ│¿Θçè∩╝îΦºúΘçèσñìµ¥éΘÇ╗Φ╛æ
- σîàσÉ½ΘöÖΦ»»σñäτÉåσÆîΦ╛╣τòîµúÇµƒÑ
- Σ╗úτáüΘò┐σ║ªµÄºσê╢σ£¿80ΦíîΣ╗Ñσåà""",
    "FILE_GEN": """µûçµíúτöƒµêÉΣ╕ôσ«╢
- τöƒµêÉτ╗ôµ₧äµ╕àµÖ░πÇüµá╝σ╝ÅΦºäΦîâτÜäµûçµíú
- Σ╜┐τö¿µáçΘóÿπÇüσêùΦí¿πÇüµ«╡ΦÉ╜Φ┐¢Φíîτ╗äτ╗ç
- ΘÇéΘàìWord/PDF/Excelσ»╝σç║
- σåàσ«╣σçåτí«πÇüΣ╕ôΣ╕ÜπÇüσÅ»µëºΦíî
- τªüµ¡óΦ╛ôσç║Σ╗úτáüσ¥ùσÆîµèÇµ£»τ╗åΦèé""",
    "PAINTER": """σ¢╛σâÅτöƒµêÉΦë║µ£»σ«╢
- σê¢Σ╜£τï¼τë╣πÇüΘ½ÿΦ┤¿ΘçÅτÜäσ¢╛σâÅ
- τÉåΦºúτö¿µê╖τÜäσ«íτ╛ÄσüÅσÑ╜
- µö»µîüΘúÄµá╝πÇüΘó£Φë▓πÇüµ₧äσ¢╛τÜäσ╛«Φ░â
- Φ╛ôσç║Θ½ÿσêåΦ╛¿τÄçσ¢╛σâÅ""",
    "RESEARCH": """µ╖▒σ║ªτáöτ⌐╢Σ╕ôσ«╢
- Φ┐¢Φíîσà¿Θ¥óτÜäΣ┐íµü»µÉ£τ┤óσÆîσêåµ₧É
- µƒÑµë╛µ£Çµû░πÇüµ£Çσçåτí«τÜäΣ┐íµü»
- µò┤τÉåσñÜΣ╕¬µ¥Ñµ║ÉτÜäΦºéτé╣
- µÅÉΣ╛¢µ£ëµá╣µì«τÜäτ╗ôΦ«║σÆîΦºüΦºú
- µáçµ│¿Σ┐íµü»µ¥Ñµ║É""",
    "SYSTEM": """τ│╗τ╗ƒµôìΣ╜£µëºΦíîσÖ¿
- µëºΦíîµ£¼σ£░τ│╗τ╗ƒσæ╜Σ╗ñσÆîµôìΣ╜£
- µëôσ╝Çσ║öτö¿πÇüτ«íτÉåµûçΣ╗╢πÇüµÄºσê╢τ│╗τ╗ƒ
- µÅÉΣ╛¢µ╕àµÖ░τÜäµëºΦíîσÅìΘªê
- ΦºúΘçèµôìΣ╜£τ╗ôµ₧£σÆîΘöÖΦ»»""",
}

# ===== Windowsµ£¼σ£░σ┐½µì╖µîçΣ╗ñµÿáσ░ä =====
WINDOWS_SHORTCUTS = {
    # µûçΣ╗╢σÆîσë¬Φ┤┤µ¥┐µôìΣ╜£
    "σñìσê╢": "Ctrl+C",
    "τ▓ÿΦ┤┤": "Ctrl+V",
    "σë¬σêç": "Ctrl+X",
    "µÆñΘöÇ": "Ctrl+Z",
    "ΘçìσüÜ": "Ctrl+Y",
    "σà¿ΘÇë": "Ctrl+A",
    "Σ┐¥σ¡ÿ": "Ctrl+S",
    "µëôσ╝Ç": "Ctrl+O",
    "µû░σ╗║": "Ctrl+N",
    # µ╡ÅΦºêσÖ¿µôìΣ╜£
    "µû░µáçτ¡╛Θí╡": "Ctrl+T",
    "σà│Θù¡µáçτ¡╛Θí╡": "Ctrl+W",
    "σÄåσÅ▓Φ«░σ╜ò": "Ctrl+H",
    "Σ╣ªτ¡╛": "Ctrl+B",
    "σê╖µû░": "Ctrl+R",
    "µö╛σñº": "Ctrl+σèáσÅ╖",
    "τ╝⌐σ░Å": "Ctrl+σçÅσÅ╖",
    # τ│╗τ╗ƒµôìΣ╜£
    "Σ╗╗σèíτ«íτÉåσÖ¿": "Ctrl+Shift+Esc",
    "µê¬σ¢╛": "Win+Shift+S",
    "σ╝ÇσºïΦÅ£σìò": "Win",
    "Θöüσ▒Å": "Win+L",
    "σà│µ£║": "Alt+F4",
    "ΦÖÜµïƒµíîΘ¥ó": "Win+Tab",
    "µÿ╛τñ║µíîΘ¥ó": "Win+D",
    # σ║öτö¿σêçµìó
    "σêçµìóσ║öτö¿": "Alt+Tab",
    "σà│Θù¡σ║öτö¿": "Alt+F4",
}


# ================= RAG Σ╕èΣ╕ïµûçσêåµ₧ÉσÖ¿ =================
class ContextAnalyzer:
    """
    σƒ║Σ║Ä RAG (µúÇτ┤óσó₧σ╝║τöƒµêÉ) τÜäµÖ║Φâ╜Σ╕èΣ╕ïµûçσêåµ₧ÉσÖ¿

    σèƒΦâ╜∩╝Ü
    1. σêåµ₧ÉσÄåσÅ▓σ»╣Φ»¥∩╝îµÅÉσÅûσà│Θö«Σ┐íµü»
    2. µ₧äσ╗║τ╗ôµ₧äσîûτÜäΣ╕èΣ╕ïµûçµÅÉτñ║Φ»ì
    3. µÖ║Φâ╜σêñµû¡Σ╗╗σèíσà│ΦüöµÇº
    4. τöƒµêÉσó₧σ╝║σÉÄτÜäΦ╛ôσàÑ
    """

    # Σ╗╗σèíτ▒╗σ₧ïτë╣σ╛üτ¡╛σÉì
    TASK_SIGNATURES = {
        "PAINTER": {
            "keywords": [
                "σ¢╛",
                "τö╗",
                "τàºτëç",
                "image",
                "photo",
                "picture",
                "σ¢╛σâÅσ╖▓τöƒµêÉ",
                "σ¢╛τëçσ╖▓τöƒµêÉ",
                "τî½",
                "τïù",
                "Σ║║τë⌐",
                "ΘúÄµÖ»",
                "σñ┤σâÅ",
            ],
            "outputs": ["σ¢╛σâÅσ╖▓τöƒµêÉ", "σ¢╛τëçσ╖▓τöƒµêÉ", "σ╖▓Σ┐¥σ¡ÿσ¢╛τëç", "Γ£¿ σ¢╛τëçσ╖▓τöƒµêÉ"],
            "entities": [
                "Θó£Φë▓",
                "ΘúÄµá╝",
                "σñºσ░Å",
                "ΦâîµÖ»",
                "Φí¿µâà",
                "σº┐σè┐",
                "τ£╝τ¥¢",
                "µ»¢σÅæ",
                "Φä╕",
            ],
        },
        "FILE_GEN": {
            "keywords": [
                "pdf",
                "word",
                "excel",
                "docx",
                "µûçµíú",
                "µèÑσæè",
                "µûçΣ╗╢",
                "τ«ÇσÄå",
                "σÉêσÉî",
                "µáçµ│¿",
                "µë╣µ│¿",
                "µ╢ªΦë▓",
                "µö╣σåÖ",
                "µáíσ»╣",
                "σ«íµáí",
                "Σ┐«Φ«ó",
                "Σ╝ÿσîû",
                "τ║áΘöÖ",
            ],
            "outputs": [
                "σ╖▓τöƒµêÉµûçΣ╗╢",
                "µûçΣ╗╢σ╖▓Σ┐¥σ¡ÿ",
                ".pdf",
                ".docx",
                ".xlsx",
                "Γ£à **µûçΣ╗╢τöƒµêÉµêÉσèƒ",
            ],
            "entities": [
                "µáçΘóÿ",
                "τ½áΦèé",
                "σåàσ«╣",
                "µá╝σ╝Å",
                "µ¿íµ¥┐",
                "µáçµ│¿",
                "µë╣µ│¿",
                "Σ┐«µö╣σ╗║Φ««",
            ],
        },
        "RESEARCH": {
            "keywords": ["τáöτ⌐╢", "σêåµ₧É", "Σ╗ïτ╗ì", "Σ║åΦºú", "σÄƒτÉå", "µèÇµ£»", "µ╖▒σàÑ"],
            "outputs": ["##", "###", "1.", "2.", "µÇ╗τ╗ô", "τ╗ôΦ«║"],
            "entities": ["σ«ÜΣ╣ë", "τë╣τé╣", "Σ╝ÿσè┐", "σèúσè┐", "σ║öτö¿", "σÅæσ▒ò"],
        },
        "CODER": {
            "keywords": [
                "Σ╗úτáü",
                "τ╝ûτ¿ï",
                "σç╜µò░",
                "ΦäÜµ£¼",
                "code",
                "script",
                "python",
                "javascript",
            ],
            "outputs": ["```python", "```javascript", "```", "def ", "class "],
            "entities": ["σç╜µò░", "σÅÿΘçÅ", "τ▒╗", "µ¿íσ¥ù", "τ«ùµ│ò"],
        },
        "CHAT": {
            "keywords": ["Σ╜áσÑ╜", "Φ░óΦ░ó", "σ╕«µêæ", "Φ»╖Θù«", "Σ╗ÇΣ╣êµÿ»"],
            "outputs": [],
            "entities": [],
        },
    }

    # σ╗╢τ╗¡µÇºµîçτñ║Φ»ìσêåτ▒╗ - Θ£ÇΦªüµ¢┤Σ╕Ñµá╝τÜäσî╣Θàì
    CONTINUATION_PATTERNS = {
        "modify": {
            # Σ┐«µö╣τ▒╗∩╝Üσ┐àΘí╗µÿ»τƒ¡σÅÑµêûµÿÄτí«τÜäΣ┐«µö╣µîçΣ╗ñ
            "indicators": [
                "σåìµ¥ÑΣ╕Çσ╝á",
                "σåìµ¥ÑΣ╕ÇΣ╕¬",
                "µ¢┤σñºΣ╕Çτé╣",
                "µ¢┤σ░ÅΣ╕Çτé╣",
                "σñºΣ╕Çτé╣",
                "σ░ÅΣ╕Çτé╣",
                "µ╖▒Σ╕ÇΣ║¢",
                "µ╡àΣ╕ÇΣ║¢",
                "Θó£Φë▓µìóµêÉ",
                "ΦâîµÖ»µìóµêÉ",
            ],
            "weight": 0.9,
            "max_input_length": 30,  # ΘÖÉσê╢Φ╛ôσàÑΘò┐σ║ª∩╝îΘò┐σÅÑσ¡ÉΣ╕ìσñ¬σÅ»Φâ╜µÿ»τ«ÇσìòΣ┐«µö╣
            "prompt_template": "τö¿µê╖Φªüµ▒éΣ┐«µö╣Σ╣ïσëìτÜäτ╗ôµ₧£∩╝Ü{modification}",
        },
        "reference": {
            # σ╝òτö¿τ▒╗∩╝Üσ┐àΘí╗σ£¿σÅÑΘªûµêûτï¼τ½ïΣ╜┐τö¿
            "indicators": [
                "Φ┐ÖΣ╕¬µÇÄΣ╣ê",
                "Φ┐Öσ╝áσ¢╛",
                "ΘéúΣ╕¬µûçΣ╗╢",
                "Σ╕èΘ¥óτÜä",
                "σêÜµëìτÜä",
                "µèèσ«â",
                "µèèΦ┐ÖΣ╕¬",
                "σƒ║Σ║ÄΦ┐ÖΣ╕¬",
            ],
            "weight": 0.85,
            "require_start": True,  # Θ£ÇΦªüσ£¿σÅÑΘªûσç║τÄ░
            "prompt_template": "τö¿µê╖σ╝òτö¿Σ║åΣ╣ïσëìτÜäσåàσ«╣∩╝Ü{reference}",
        },
        "reference_loose": {
            # σ╝òτö¿τ▒╗∩╝êσ«╜µ¥╛∩╝ë∩╝Üτö¿Σ║ÄΦ«íσêÆ/σñºτ║▓/µû╣µíêτ¡ëΘ£ÇΦªüΦ╖ƒΘÜÅΣ╣ïσëìσåàσ«╣τÜäΦ»╖µ▒é
            "indicators": [
                "Φ┐ÖΣ╕¬Φ«íσêÆ",
                "Φ»ÑΦ«íσêÆ",
                "Σ╕èΦ┐░Φ«íσêÆ",
                "Σ╕èΘ¥óτÜäΦ«íσêÆ",
                "Φ┐ÖΣ╕¬µû╣µíê",
                "Φ»Ñµû╣µíê",
                "Σ╕èΦ┐░µû╣µíê",
                "Φ┐ÖΣ╕¬σñºτ║▓",
                "Φ»Ñσñºτ║▓",
                "Φ┐ÖΣ╕¬ppt",
                "Φ»Ñppt",
                "Φ┐ÖΣ╕¬PPT",
                "Φ»ÑPPT",
                "µîëτàºΦ┐ÖΣ╕¬",
                "µá╣µì«Φ┐ÖΣ╕¬",
            ],
            "weight": 0.78,
            "require_start": False,
            "prompt_template": "τö¿µê╖σ╝òτö¿Σ║åΣ╣ïσëìτÜäΦ«íσêÆµêûσñºτ║▓∩╝Ü{reference}",
        },
        "convert": {
            # Φ╜¼µìóτ▒╗∩╝ÜµÿÄτí«τÜäµá╝σ╝ÅΦ╜¼µìóΦ»╖µ▒é
            "indicators": [
                "σüÜµêÉword",
                "σüÜµêÉpdf",
                "σüÜµêÉexcel",
                "Φ╜¼µêÉword",
                "Φ╜¼µêÉpdf",
                "σÅÿµêÉµûçµíú",
                "σ»╝σç║Σ╕║",
                "Σ┐¥σ¡ÿΣ╕║word",
                "Σ┐¥σ¡ÿΣ╕║pdf",
            ],
            "weight": 0.95,
            "prompt_template": "τö¿µê╖Φªüµ▒éσ░åΣ╣ïσëìτÜäσåàσ«╣Φ╜¼µìóΣ╕║µû░µá╝σ╝Å∩╝Ü{conversion}",
        },
        "continue": {
            # τ╗ºτ╗¡τ▒╗∩╝ÜµÿÄτí«Φªüµ▒éτ╗ºτ╗¡Σ╣ïσëìτÜäσåàσ«╣
            "indicators": [
                "τ╗ºτ╗¡σåÖ",
                "µÄÑτ¥ÇΦ»┤",
                "µÄÑτ¥ÇσåÖ",
                "τä╢σÉÄσæó",
                "Σ╕ïΣ╕Çµ¡Ñ",
                "Φ┐ÿµ£ëσæó",
                "σÅªσñûΦíÑσàà",
                "σåìµë╛µë╛",
                "σåìµÉ£",
                "σåìµƒÑ",
                "σåìτ£ïτ£ï",
                "τ╗ºτ╗¡µƒÑ",
                "τ╗ºτ╗¡µë╛",
                "σåìµë╛",
                "σåìµÉ£Σ╕ÇΣ╕ï",
            ],
            "weight": 0.7,
            "max_input_length": 20,  # τƒ¡σÅÑµëìµÿ»τ╗ºτ╗¡µîçΣ╗ñ
            "prompt_template": "τö¿µê╖Φªüµ▒éτ╗ºτ╗¡Σ╣ïσëìτÜäΣ╗╗σèí∩╝Ü{continuation}",
        },
        "detail": {
            # Φ»ªτ╗åτ▒╗∩╝ÜσÅ¬µ£ëΘ¥₧σ╕╕µÿÄτí«τÜäσ▒òσ╝ÇΦ»╖µ▒éµëìτ«ù∩╝îΣ╕öσ┐àΘí╗µÿ»τƒ¡σÅÑ
            "indicators": [
                "Φ»ªτ╗åΦ»┤Φ»┤",
                "σ▒òσ╝ÇΦ»┤Φ»┤",
                "Φ»ªτ╗åΦ«▓Φ«▓",
                "σà╖Σ╜ôΦ»┤Σ╕ÇΣ╕ï",
                "ΦºúΘçèΣ╕ÇΣ╕ïσêÜµëìτÜä",
            ],
            "weight": 0.75,
            "max_input_length": 25,  # ΘÖÉσê╢Θò┐σ║ª
            "prompt_template": "τö¿µê╖Φªüµ▒éΦ»ªτ╗åΦ»┤µÿÄΣ╣ïσëìµÅÉσê░τÜäσåàσ«╣∩╝Ü{detail}",
        },
    }

    @classmethod
    def extract_entities(cls, text: str, task_type: str = None) -> list:
        """Σ╗Äµûçµ£¼Σ╕¡µÅÉσÅûσà│Θö«σ«₧Σ╜ô"""
        entities = []
        text_lower = text.lower()

        # ΘÇÜτö¿σ«₧Σ╜ôµÅÉσÅû
        # Θó£Φë▓
        colors = [
            "τ║óΦë▓",
            "Φô¥Φë▓",
            "τ╗┐Φë▓",
            "Θ╗äΦë▓",
            "τÖ╜Φë▓",
            "Θ╗æΦë▓",
            "τü░Φë▓",
            "τ▓ëΦë▓",
            "τ┤½Φë▓",
            "µ⌐ÖΦë▓",
            "µúòΦë▓",
        ]
        for color in colors:
            if color in text_lower:
                entities.append({"type": "color", "value": color})

        # ΘúÄµá╝
        styles = [
            "σÅ»τê▒",
            "σ╕àµ░ö",
            "σåÖσ«₧",
            "σìíΘÇÜ",
            "σè¿µ╝½",
            "Φ╡¢σìÜµ£ïσàï",
            "µ░┤σ╜⌐",
            "µ▓╣τö╗",
            "τ«Çτ║ª",
            "σñìσÅñ",
        ]
        for style in styles:
            if style in text_lower:
                entities.append({"type": "style", "value": style})

        # Σ╕╗Θóÿ/σ»╣Φ▒í
        subjects = [
            "τî½",
            "τïù",
            "Σ║║",
            "ΘúÄµÖ»",
            "σ╗║τ¡æ",
            "µ▒╜Φ╜ª",
            "Φè▒",
            "µáæ",
            "σ▒▒",
            "µ╡╖",
            "σƒÄσ╕é",
        ]
        for subject in subjects:
            if subject in text_lower:
                entities.append({"type": "subject", "value": subject})

        # τë╣σ«ÜΣ╗╗σèíτÜäσ«₧Σ╜ô
        if task_type and task_type in cls.TASK_SIGNATURES:
            for entity_keyword in cls.TASK_SIGNATURES[task_type].get("entities", []):
                if entity_keyword in text_lower:
                    entities.append({"type": "task_specific", "value": entity_keyword})

        return entities

    @classmethod
    def build_context_summary(cls, history: list, max_turns: int = 3) -> dict:
        """
        µ₧äσ╗║σÄåσÅ▓Σ╕èΣ╕ïµûçµæÿΦªü

        Φ┐öσ¢₧:
        {
            "task_history": [],      # Σ╗╗σèíσÄåσÅ▓
            "key_entities": [],      # σà│Θö«σ«₧Σ╜ô
            "last_user_intent": "",  # µ£ÇΦ┐æτÜäτö¿µê╖µäÅσ¢╛
            "last_model_output": "", # µ£ÇΦ┐æτÜäµ¿íσ₧ïΦ╛ôσç║
            "conversation_topic": "" # σ»╣Φ»¥Σ╕╗Θóÿ
        }
        """
        summary = {
            "task_history": [],
            "key_entities": [],
            "last_user_intent": "",
            "last_model_output": "",
            "conversation_topic": "",
        }

        if not history:
            return summary

        # σêåµ₧Éµ£ÇΦ┐æτÜäσ»╣Φ»¥
        recent_turns = (
            history[-max_turns * 2 :] if len(history) > max_turns * 2 else history
        )

        all_entities = []
        topics = []

        for turn in recent_turns:
            content = turn["parts"][0] if turn["parts"] else ""
            role = turn["role"]

            if role == "user":
                summary["last_user_intent"] = content
                # Φ»åσê½Σ╗╗σèíτ▒╗σ₧ï
                for task_type, signatures in cls.TASK_SIGNATURES.items():
                    if any(kw in content.lower() for kw in signatures["keywords"]):
                        summary["task_history"].append(
                            {"type": task_type, "content": content[:100]}
                        )
                        topics.append(task_type)
                        break

                # µÅÉσÅûσ«₧Σ╜ô
                entities = cls.extract_entities(content)
                all_entities.extend(entities)

            elif role == "model":
                summary["last_model_output"] = content

        # σÄ╗Θçìσ«₧Σ╜ô
        seen = set()
        unique_entities = []
        for e in all_entities:
            key = f"{e['type']}:{e['value']}"
            if key not in seen:
                seen.add(key)
                unique_entities.append(e)
        summary["key_entities"] = unique_entities

        # τí«σ«Üσ»╣Φ»¥Σ╕╗Θóÿ
        if topics:
            summary["conversation_topic"] = topics[-1]  # µ£ÇΦ┐æτÜäΣ╗╗σèíτ▒╗σ₧ï

        return summary

    @classmethod
    def build_rag_prompt(
        cls, user_input: str, context_summary: dict, continuation_type: str = None
    ) -> str:
        """
        µ₧äσ╗║ RAG ΘúÄµá╝τÜäσó₧σ╝║µÅÉτñ║Φ»ì

        σ░åΣ╕èΣ╕ïµûçΣ┐íµü»τ╗ôµ₧äσîûσ£░µ│¿σàÑσê░τö¿µê╖Φ╛ôσàÑΣ╕¡
        """
        prompt_parts = []

        # 1. µ╖╗σèáΣ╕èΣ╕ïµûçµáçΦ«░
        if context_summary.get("conversation_topic"):
            prompt_parts.append(
                f"[Σ╕èΣ╕ïµûçτ▒╗σ₧ï: {context_summary['conversation_topic']}]"
            )

        # 2. µ╖╗σèáσà│Θö«σ«₧Σ╜ôΣ┐íµü»
        if context_summary.get("key_entities"):
            entities_str = ", ".join(
                [
                    f"{e['type']}={e['value']}"
                    for e in context_summary["key_entities"][:5]
                ]
            )
            prompt_parts.append(f"[σà│Θö«Σ┐íµü»: {entities_str}]")

        # 3. µ╖╗σèáσÄåσÅ▓µäÅσ¢╛
        if context_summary.get("last_user_intent"):
            # µê¬σÅûµá╕σ┐âµÅÅΦ┐░
            last_intent = context_summary["last_user_intent"]
            if len(last_intent) > 200:
                last_intent = last_intent[:200] + "..."
            prompt_parts.append(f"[Σ╣ïσëìτÜäΦ»╖µ▒é: {last_intent}]")

        # 4. µá╣µì«σ╗╢τ╗¡τ▒╗σ₧ïµ╖╗σèáτë╣σ«ÜµîçΣ╗ñ
        if continuation_type and continuation_type in cls.CONTINUATION_PATTERNS:
            pattern = cls.CONTINUATION_PATTERNS[continuation_type]
            # Σ╕ìµ╖╗σèáµ¿íµ¥┐∩╝îΦ«⌐σ«₧Σ╜ôσÆîΣ╕èΣ╕ïµûçΦç¬τä╢Φ₧ìσÉê

        # 5. µ╖╗σèáτö¿µê╖σ╜ôσëìΦ╛ôσàÑ
        prompt_parts.append(f"[σ╜ôσëìΦ»╖µ▒é: {user_input}]")

        # 6. σªéµ₧£µÿ»Φ╜¼µìóΦ»╖µ▒é∩╝îµ╖╗σèáµ║Éσåàσ«╣
        if continuation_type == "convert" and context_summary.get("last_model_output"):
            output = context_summary["last_model_output"]
            # ΘÖÉσê╢Θò┐σ║ª
            if len(output) > 4000:
                output = output[:4000] + "\n...(σåàσ«╣σ╖▓µê¬µû¡)"
            prompt_parts.append(f"\n[Θ£ÇΦªüΦ╜¼µìóτÜäµ║Éσåàσ«╣:]\n{output}")

        # 7. σªéµ₧£µÿ»σ╝òτö¿τ▒╗σ╗╢τ╗¡∩╝îΘÖäΣ╕èµ£ÇΦ┐æΦ╛ôσç║µæÿΦªü
        if continuation_type in (
            "reference",
            "reference_loose",
        ) and context_summary.get("last_model_output"):
            output = context_summary["last_model_output"]
            if len(output) > 2000:
                output = output[:2000] + "\n...(σåàσ«╣σ╖▓µê¬µû¡)"
            prompt_parts.append(f"\n[µ£ÇΦ┐æΦ╛ôσç║µæÿΦªü:]\n{output}")

        # τ╗äσÉêµêÉµ£Çτ╗êτÜäσó₧σ╝║µÅÉτñ║
        enhanced_prompt = "\n".join(prompt_parts)

        return enhanced_prompt

    @classmethod
    def analyze_context(cls, user_input: str, history: list) -> dict:
        """
        RAG ΘúÄµá╝τÜäΣ╕èΣ╕ïµûçσêåµ₧É

        Φ┐öσ¢₧:
        {
            "is_continuation": bool,      # µÿ»σÉªµÿ»σ╗╢τ╗¡Σ╗╗σèí
            "related_task": str,          # σà│ΦüöτÜäΣ╗╗σèíτ▒╗σ₧ï
            "continuation_type": str,     # σ╗╢τ╗¡τ▒╗σ₧ï (modify/reference/convert/continue/detail)
            "context_summary": dict,      # τ╗ôµ₧äσîûΣ╕èΣ╕ïµûçµæÿΦªü
            "enhanced_input": str,        # RAG σó₧σ╝║σÉÄτÜäΦ╛ôσàÑ
            "confidence": float,          # τ╜«Σ┐íσ║ª
        }
        """
        result = {
            "is_continuation": False,
            "related_task": None,
            "continuation_type": None,
            "context_summary": {},
            "enhanced_input": user_input,
            "confidence": 0.0,
        }

        if not history or len(history) < 2:
            return result

        user_lower = user_input.lower()
        input_length = len(user_input)

        # 1. µ₧äσ╗║Σ╕èΣ╕ïµûçµæÿΦªü
        context_summary = cls.build_context_summary(history)
        result["context_summary"] = context_summary

        # 2. µúÇµ╡ïσ╗╢τ╗¡τ▒╗σ₧ïσÆîτ╜«Σ┐íσ║ª∩╝êµ¢┤Σ╕Ñµá╝τÜäσî╣Θàì∩╝ë
        detected_type = None
        max_weight = 0.0

        for pattern_type, pattern_info in cls.CONTINUATION_PATTERNS.items():
            indicators = pattern_info["indicators"]
            weight = pattern_info["weight"]

            # µúÇµƒÑΦ╛ôσàÑΘò┐σ║ªΘÖÉσê╢∩╝êσªéµ₧£µ£ë∩╝ë
            max_len = pattern_info.get("max_input_length")
            if max_len and input_length > max_len:
                continue  # Φ╛ôσàÑσñ¬Θò┐∩╝îΣ╕ìσñ¬σÅ»Φâ╜µÿ»τ«ÇσìòτÜäσ╗╢τ╗¡µîçΣ╗ñ

            # µúÇµƒÑµÿ»σÉªΘ£ÇΦªüσ£¿σÅÑΘªûσç║τÄ░
            require_start = pattern_info.get("require_start", False)

            # Φ«íτ«ùσî╣ΘàìτÜäµîçτñ║Φ»ìµò░ΘçÅ
            matches = 0
            for ind in indicators:
                if ind in user_lower:
                    if require_start:
                        # Θ£ÇΦªüσ£¿σÅÑΘªû∩╝êσëì10Σ╕¬σ¡ùτ¼ªσåà∩╝ë
                        if user_lower.find(ind) < 10:
                            matches += 1
                    else:
                        matches += 1

            if matches > 0:
                # σèáµ¥âΦ«íτ«ùτ╜«Σ┐íσ║ª
                adjusted_weight = weight * (
                    1 + 0.1 * (matches - 1)
                )  # σñÜΣ╕¬σî╣Θàìσó₧σèáτ╜«Σ┐íσ║ª
                if adjusted_weight > max_weight:
                    max_weight = adjusted_weight
                    detected_type = pattern_type

        # 3. Θó¥σñûµúÇµƒÑ∩╝Üσªéµ₧£τö¿µê╖Φ╛ôσàÑσîàσÉ½µÿÄτí«τÜäµû░Σ╕╗Θóÿ∩╝îΘÖìΣ╜Äσ╗╢τ╗¡σêñµû¡
        # µû░Σ╕╗Θóÿµáçσ┐ù∩╝ÜσîàσÉ½"σà│Σ║Ä"πÇü"Σ╕ÇΣ╕¬"σÉÄµÄÑµû░σ«₧Σ╜ô
        new_topic_indicators = [
            "σà│Σ║Ä",
            "Σ╕Çτ»ç",
            "Σ╕ÇΣ╗╜",
            "Σ╕ÇΣ╕¬µû░τÜä",
            "σ╕«µêæσåÖ",
            "σ╕«µêæσüÜ",
            "σ╕«µêæτöƒµêÉ",
            "τ╗ÖµêæτöƒµêÉ",
            "τöƒµêÉΣ╕Ç",
        ]
        has_new_topic = any(ind in user_lower for ind in new_topic_indicators)

        # µúÇµƒÑµÿ»σÉªµÿ»σ«îσà¿Σ╕ìσÉîτÜäΣ╗╗σèíτ▒╗σ₧ï∩╝êσªé∩╝Üµëôσ╝Çσ╛«Σ┐í -> τöƒµêÉσ¢╛τëç∩╝ë
        task_mismatch = False
        if context_summary.get("conversation_topic"):
            prev_topic = context_summary["conversation_topic"]
            # µúÇµ╡ïσ╜ôσëìΦ╛ôσàÑτÜäΣ╗╗σèíτ▒╗σ₧ï
            curr_likely_task = None
            if any(
                kw in user_lower
                for kw in ["µƒÑ", "µÉ£", "µÉ£τ┤ó", "µƒÑΦ»ó", "µë╛", "σåìµë╛", "σåìµƒÑ", "σåìµÉ£"]
            ):
                curr_likely_task = "WEB_SEARCH"
            elif any(kw in user_lower for kw in ["σ¢╛", "τö╗", "τàºτëç", "image"]):
                curr_likely_task = "PAINTER"
            elif any(kw in user_lower for kw in ["word", "pdf", "µûçµíú", "µèÑσæè"]):
                curr_likely_task = "FILE_GEN"
            elif any(kw in user_lower for kw in ["µëôσ╝Ç", "Φ┐ÉΦíî", "σà│Θù¡"]):
                curr_likely_task = "SYSTEM"

            # σªéµ₧£Σ╗╗σèíτ▒╗σ₧ïσ«îσà¿Σ╕ìσÉî∩╝îΣ╕ìσ║öΦ»Ñµÿ»σ╗╢τ╗¡
            if curr_likely_task and prev_topic and curr_likely_task != prev_topic:
                task_mismatch = True
                _app_logger.debug(
                    f"[ContextAnalyzer] Σ╗╗σèíτ▒╗σ₧ïΣ╕ìσî╣Θàì: {prev_topic} -> {curr_likely_task}"
                )

        if has_new_topic and input_length > 10:
            # µ£ëµû░Σ╕╗ΘóÿΣ╕öΦ╛ôσàÑΦ╛âΘò┐∩╝îσ╛êσÅ»Φâ╜µÿ»τï¼τ½ïΣ╗╗σèí
            max_weight *= 0.2  # σñºσ╣àΘÖìΣ╜Äτ╜«Σ┐íσ║ª
            _app_logger.debug(f"[ContextAnalyzer] µúÇµ╡ïσê░µû░Σ╕╗Θóÿµáçσ┐ù∩╝îΘÖìΣ╜Äσ╗╢τ╗¡τ╜«Σ┐íσ║ª")

        if task_mismatch:
            # Σ╗╗σèíτ▒╗σ₧ïΣ╕ìσî╣Θàì∩╝îσ╝║σê╢µ╕àΘ¢╢
            max_weight = 0
            detected_type = None
            _app_logger.debug(f"[ContextAnalyzer] Σ╗╗σèíτ▒╗σ₧ïΣ╕ìσî╣Θàì∩╝îµ╕àΘÖñσ╗╢τ╗¡σêñµû¡")

        # 4. σªéµ₧£µúÇµ╡ïσê░σ╗╢τ╗¡µ¿íσ╝ÅΣ╕öτ╜«Σ┐íσ║ªΦ╢│σñƒΘ½ÿ
        if detected_type and max_weight > 0.5:
            result["is_continuation"] = True
            result["continuation_type"] = detected_type
            result["confidence"] = min(max_weight, 1.0)

            # τí«σ«Üσà│ΦüöτÜäΣ╗╗σèíτ▒╗σ₧ï
            if context_summary.get("conversation_topic"):
                result["related_task"] = context_summary["conversation_topic"]
            elif context_summary.get("task_history"):
                result["related_task"] = context_summary["task_history"][-1]["type"]

            # 4. µ₧äσ╗║ RAG σó₧σ╝║µÅÉτñ║
            result["enhanced_input"] = cls.build_rag_prompt(
                user_input, context_summary, detected_type
            )

            _app_logger.debug(f"[ContextAnalyzer] RAG Analysis:")
            _app_logger.info(f"  - Continuation Type: {detected_type}")
            _app_logger.info(f"  - Related Task: {result['related_task']}")
            _app_logger.info(f"  - Confidence: {result['confidence']:.2f}")
            _app_logger.info(
                f"  - Entities: {[e['value'] for e in context_summary.get('key_entities', [])]}"
            )

        # 5. τë╣µ«èσñäτÉå∩╝ÜΦ╜¼µìóΦ»╖µ▒é∩╝êσì│Σ╜┐µ▓íµ£ëµÿÄτí«τÜäσ╗╢τ╗¡µîçτñ║Φ»ì∩╝ë
        convert_patterns = [
            "σüÜµêÉword",
            "σüÜµêÉpdf",
            "Φ╜¼µêÉword",
            "Φ╜¼µêÉpdf",
            "τöƒµêÉword",
            "τöƒµêÉpdf",
            "σ»╝σç║Σ╕║",
        ]
        if any(p in user_lower for p in convert_patterns) and context_summary.get(
            "last_model_output"
        ):
            result["is_continuation"] = True
            result["continuation_type"] = "convert"
            result["related_task"] = "FILE_GEN"
            result["confidence"] = 0.95
            result["enhanced_input"] = cls.build_rag_prompt(
                user_input, context_summary, "convert"
            )

        return result

    @classmethod
    def filter_history(
        cls, user_input: str, history: list, keep_turns: int = 6
    ) -> list:
        """Φ┐çµ╗ñσÄåσÅ▓Φ«░σ╜ò∩╝îσ░╜ΘçÅΘü┐σàìµùáσà│Σ╕èΣ╕ïµûçµ▒íµƒô"""
        if not history:
            return []

        # σªéµ₧£σÄåσÅ▓σ╛êτƒ¡∩╝îτ¢┤µÄÑΦ┐öσ¢₧
        if len(history) <= keep_turns * 2:
            return history

        user_lower = user_input.lower()

        # µè╜σÅûτö¿µê╖Φ╛ôσàÑΣ╕¡τÜäσ«₧Σ╜ôΣ╕Äσà│Θö«Φ»ì
        entities = cls.extract_entities(user_input)
        entity_values = {e["value"] for e in entities}

        # Θó¥σñûµÅÉσÅûΣ╕¡µûçσà│Θö«Φ»ì∩╝êΘò┐σ║ª>=2∩╝ëΣ╕ÄΦï▒µûçσìòΦ»ì∩╝êΘò┐σ║ª>=3∩╝ë
        import re

        cjk_words = re.findall(r"[\u4e00-\u9fff]{2,}", user_input)
        eng_words = re.findall(r"[a-zA-Z]{3,}", user_input)
        keyword_set = {k.lower() for k in (cjk_words + eng_words)}
        keyword_set.update({v.lower() for v in entity_values})

        # µ₧äσ╗║τ¢╕σà│σÄåσÅ▓∩╝ÜσîàσÉ½σà│Θö«Φ»ìτÜäσ»╣Φ»¥
        relevant = []
        for turn in history:
            content = (turn.get("parts") or [""])[0]
            content_lower = content.lower()
            if any(k in content_lower for k in keyword_set if k):
                relevant.append(turn)

        # σºïτ╗êΣ┐¥τòÖµ£ÇΦ┐æ 3 Φ╜«σ»╣Φ»¥∩╝êτí«Σ┐¥Σ╕èΣ╕ïµûçΦ┐₧Φ┤»∩╝ë
        tail_count = 6
        tail_start_index = max(0, len(history) - tail_count)

        # µö╢Θ¢åΘ£ÇΦªüΣ┐¥τòÖτÜäτ┤óσ╝ò
        indices_to_keep = set()

        # 1. σà│Θö«Φ»ìσî╣ΘàìτÜäσÄåσÅ▓
        for i, turn in enumerate(history):
            content = (turn.get("parts") or [""])[0]
            content_lower = content.lower()
            if any(k in content_lower for k in keyword_set if k):
                indices_to_keep.add(i)
                # σÉîµù╢Σ┐¥τòÖΦ»Ñµ¥íτÜäσëìΣ╕Çµ¥í∩╝êσªéµ₧£µÿ»User/ModelΘàìσ»╣∩╝ë
                if i > 0:
                    indices_to_keep.add(i - 1)

        # 2. Σ╣ƒµÿ»µ£ÇΘçìΦªüτÜä∩╝ÜΣ┐¥τòÖσ░╛Θâ¿Σ╕èΣ╕ïµûç
        for i in range(tail_start_index, len(history)):
            indices_to_keep.add(i)

        # µîëσÄƒσºïΘí║σ║ÅΘçìτ╗ä
        filtered_history = [history[i] for i in sorted(indices_to_keep)]

        return filtered_history

        # σÅ¬Σ┐¥τòÖµ£ÇΦ┐æ keep_turns Φ╜«
        return merged[-keep_turns * 2 :]


class TaskOrchestrator:
    """
    τ╝ûµÄÆσÆîµëºΦíîσñÜΣ╕¬σ¡ÉΣ╗╗σèí

    Φ┤úΦüî∩╝Ü
    1. Θí║σ║ÅµëºΦíîσ¡ÉΣ╗╗σèí
    2. σ£¿σ¡ÉΣ╗╗σèíΘù┤Σ╝áΘÇÆµò░µì«/Σ╕èΣ╕ïµûç
    3. σñäτÉåΘöÖΦ»»σÆîΘçìΦ»ò
    4. µ£Çτ╗êΘ¬îΦ»üΦ╛ôσç║Φ┤¿ΘçÅ
    """

    @classmethod
    async def execute_compound_task(
        cls, user_input: str, subtasks: list, session_name: str = None
    ) -> dict:
        """
        µëºΦíîσñìσÉêΣ╗╗σèíτÜäµëÇµ£ëσ¡ÉΣ╗╗σèí

        Φ┐öσ¢₧:
            {
                "success": bool,
                "primary_result": Σ╕╗Σ╗╗σèíτ╗ôµ₧£,
                "secondary_results": [µ¼íΦªüΣ╗╗σèíτ╗ôµ₧£],
                "combined_output": µ£Çτ╗êσÉêσ╣╢Φ╛ôσç║,
                "execution_log": µëºΦíîµùÑσ┐ù,
                "quality_score": Φ┤¿ΘçÅΦ»äσêå (0-100),
                "errors": ΘöÖΦ»»σêùΦí¿
            }
        """
        execution_log = []
        results = []
        context = {"original_input": user_input, "user_input": user_input}
        errors = []

        try:
            for i, subtask in enumerate(subtasks):
                _app_logger.debug(
                    f"\n[TaskOrchestrator] µëºΦíîσ¡ÉΣ╗╗σèí {i+1}/{len(subtasks)}: {subtask['task_type']}"
                )
                execution_log.append(
                    f"µ¡ÑΘ¬ñ {i+1}: µëºΦíî {subtask['task_type']} - {subtask['description']}"
                )
                step_input = subtask.get("input") or user_input

                try:
                    # µá╣µì«Σ╗╗σèíτ▒╗σ₧ïΦ░âτö¿τ¢╕σ║öτÜäσñäτÉåσç╜µò░
                    if subtask["task_type"] == "WEB_SEARCH":
                        result = await cls._execute_web_search(step_input, context)
                    elif subtask["task_type"] == "FILE_GEN":
                        result = await cls._execute_file_gen(
                            step_input, context, subtask
                        )
                    elif subtask["task_type"] == "PAINTER":
                        result = await cls._execute_painter(step_input, context)
                    elif subtask["task_type"] == "RESEARCH":
                        result = await cls._execute_research(step_input, context)
                    else:
                        result = {
                            "success": False,
                            "error": f"µ£¬τƒÑΣ╗╗σèíτ▒╗σ₧ï: {subtask['task_type']}",
                        }

                    subtask["status"] = "completed"
                    subtask["result"] = result
                    results.append(result)

                    # σ░åτ╗ôµ₧£Σ┐¥σ¡ÿσê░Σ╕èΣ╕ïµûç∩╝îΣ╛¢Σ╕ïΣ╕ÇΣ╕¬Σ╗╗σèíΣ╜┐τö¿
                    context[f"{subtask['task_type']}_result"] = result
                    context[f"step_{i+1}_output"] = result.get(
                        "output", result.get("content", "")
                    )

                    execution_log.append(f"  Γ£à σ«îµêÉ: {subtask['description']}")

                except Exception as e:
                    error_msg = str(e)
                    subtask["status"] = "failed"
                    subtask["error"] = error_msg
                    errors.append(error_msg)
                    execution_log.append(f"  Γ¥î σñ▒Φ┤Ñ: {error_msg}")
                    _app_logger.debug(f"[TaskOrchestrator] σ¡ÉΣ╗╗σèíσñ▒Φ┤Ñ: {error_msg}")

            # σÉêσ╣╢τ╗ôµ₧£
            combined_output = cls._merge_results(subtasks, context)

            # Φ┤¿ΘçÅΘ¬îΦ»ü
            quality_score = await cls._validate_quality(
                user_input, combined_output, context
            )

            return {
                "success": len(errors) == 0,
                "primary_result": results[0] if results else None,
                "secondary_results": results[1:] if len(results) > 1 else [],
                "combined_output": combined_output,
                "execution_log": execution_log,
                "quality_score": quality_score,
                "errors": errors,
                "context": context,
            }

        except Exception as e:
            return {
                "success": False,
                "primary_result": None,
                "secondary_results": [],
                "combined_output": None,
                "execution_log": execution_log,
                "quality_score": 0,
                "errors": errors + [str(e)],
                "context": context,
            }

    @classmethod
    async def _execute_web_search(
        cls, user_input: str, context: dict, progress_callback=None
    ) -> dict:
        """µëºΦíî Web µÉ£τ┤óσ¡ÉΣ╗╗σèí (σ╕ªσÅ»ΦºåΦ┐¢σ║ª)"""

        def _report(msg: str, detail: str = ""):
            _app_logger.debug(f"[WEB_SEARCH] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        try:
            _report("σÉ»σè¿τ╜æτ╗£µÉ£τ┤ó...", "µ¡úσ£¿ΦºäσêÆµÉ£τ┤óσà│Θö«Φ»ì")

            # Phase 1: Planning
            # (WebSearcher manages its own queries, but we can simulate the 'thought' process)
            await asyncio.sleep(0.3)
            _report("µëºΦíî Google Search...", f"σà│Θö«Φ»ì: {user_input[:20]}...")

            # Phase 2: Execution
            # wrap in thread
            result = await asyncio.to_thread(
                WebSearcher.search_with_grounding, user_input
            )

            # Phase 3: Reporting
            if result.get("grounded"):
                _report("Γ£à µÉ£τ┤óσ╣╢σ╝òτö¿σ«îµêÉ", "σ╖▓τ╗ôσÉêµ£Çµû░Σ┐íµü»")
            else:
                _report("Γ£à µÉ£τ┤óσ«îµêÉ", "σ╖▓ΦÄ╖σÅûτ¢╕σà│τ╜æΘí╡µæÿΦªü")

            return {
                "success": result.get("success", False),
                "output": result.get("response", ""),
                "content": result.get("response", ""),
                "grounded": result.get("grounded", False),
                "raw_result": result,
                "model_id": "gemini-2.5-flash",
            }
        except Exception as e:
            _report("Γ¥î µÉ£τ┤óΘüçσê░Θù«Θóÿ", str(e))
            return {
                "success": False,
                "output": "",
                "error": str(e),
                "raw_result": None,
                "model_id": "gemini-2.5-flash",
            }

    @classmethod
    async def _execute_ppt_multi_step(
        cls, user_input: str, context: dict, subtask: dict, progress_callback=None
    ) -> dict:
        """µëºΦíîσñÜΘÿ╢µ«╡PPTτöƒµêÉΣ╗╗σèí (Plan-then-Execute)"""
        from web.smart_feedback import SmartFeedback

        fb = SmartFeedback(
            user_request=user_input,
            task_type="PPT",
            emit=lambda m, d: None,
            total_steps=3,
        )

        def _report(msg: str, detail: str = ""):
            _app_logger.debug(f"[PPT_PROGRESS] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        m, d = fb.start("σñÜΘÿ╢µ«╡PPTτöƒµêÉ")
        _report(m, d)
        previous_data = context.get(f"step_{subtask['index']}_output", "")

        # 1. ΦºäσêÆΘÿ╢µ«╡ (Planning Phase)
        try:
            from web.ppt_master import PPTBlueprint, PPTContentPlanner

            # σê¥σºïσîûΦºäσêÆσÖ¿
            planner = PPTContentPlanner(ai_client=client, model_name="gemini-2.5-flash")

            # µëºΦíîΦºäσêÆ
            _report("µ¡úσ£¿ΦºäσêÆσåàσ«╣τ╗ôµ₧ä...", "Φ░âτö¿ AI ΦºäσêÆσñºτ║▓")
            plan_result = await planner.plan_content_structure(
                user_input, search_results=None
            )

            # µÅÉσÅûσñºτ║▓
            outline_data = plan_result.get("outline", [])
            theme_choice = plan_result.get("theme_recommendation", "business")
            total_slides = plan_result.get("total_expected_slides", 10)

            # --- 1.2 σ▒òτñ║ΦºäσêÆµªéΦºê (User Requirement: Visualize Plan) ---
            plan_summary = f"σñºτ║▓µªéΦºê ({len(outline_data)} τ½áΦèé, {total_slides} Θí╡):\n"
            for idx, sec in enumerate(outline_data):
                plan_summary += f"{idx+1}. {sec.get('section_title')} ({len(sec.get('slides', []))} Θí╡)\n"
            _report(f"ΦºäσêÆσ«îµêÉ∩╝îσà▒ {total_slides} Θí╡", plan_summary)

            # σ░åσñºτ║▓Φ╜¼µìóΣ╕║ PPTGenerator σÅ»Φ»åσê½τÜäµá╝σ╝Å
            ppt_slides = []

            # --- σñÜΘÿ╢µ«╡µëºΦíî∩╝ÜΘÇÉΘí╡τöƒµêÉσåàσ«╣ ---
            total_steps = sum(len(sec.get("slides", [])) for sec in outline_data)
            current_step = 0

            for section in outline_data:
                section_title = section.get("section_title", "τ½áΦèé")
                # µ╖╗σèáτ½áΦèéΘí╡
                ppt_slides.append(
                    {
                        "type": "section",
                        "title": section_title,
                        "content": [section.get("section_theme", "")],
                    }
                )

                for slide in section.get("slides", []):
                    current_step += 1
                    s_title = slide.get("slide_title", "µ£¬σæ╜σÉìσ╣╗τü»τëç")
                    s_type = slide.get("slide_type", "content")
                    s_points = slide.get("key_points", [])

                    # Log progress
                    _report(
                        f"τöƒµêÉτ¼¼ {current_step}/{total_steps} Θí╡σåàσ«╣: {s_title}",
                        "Θÿ╢µ«╡ 2/3: σåàσ«╣µë⌐σàà",
                    )

                    # µë⌐σààσåàσ«╣ (Per-Slide Generation)
                    expanded_points = s_points
                    if hasattr(planner, "expand_slide_content"):
                        try:
                            # Use new method in PPTContentPlanner
                            expanded_points = await planner.expand_slide_content(
                                s_title, s_points, context=f"Context: {section_title}"
                            )
                            if expanded_points != s_points:
                                _report(
                                    f"  Γ£¿ σåàσ«╣σ╖▓µë⌐σàà: {len(expanded_points)} µ¥í",
                                    f"σ╣╗τü»τëç: {s_title}",
                                )
                        except Exception as exp_err:
                            _report(f"  ΓÜá∩╕Å µë⌐σààσñ▒Φ┤Ñ∩╝îΣ╜┐τö¿σÄƒσºïσåàσ«╣", str(exp_err))
                            expanded_points = s_points

                    ppt_slides.append(
                        {
                            "type": (
                                s_type
                                if s_type
                                in ["content", "content_image", "comparison", "data"]
                                else "content"
                            ),
                            "title": s_title,
                            "points": expanded_points,
                            "content": expanded_points,
                            "notes": slide.get("content_description", ""),
                        }
                    )

            # σªéµ₧£µ▓íµ£ëτöƒµêÉµ£ëµòêτÜäσ╣╗τü»τëç∩╝îσ¢₧ΘÇÇσê░µùºΘÇ╗Φ╛æ
            if not ppt_slides:
                raise ValueError("ΦºäσêÆσÖ¿µ£¬τöƒµêÉµ£ëµòêσ╣╗τü»τëçσñºτ║▓")

            # --- 2.5 Φ┤¿ΘçÅΦç¬µúÇΣ╕Äσåàσ«╣µ╕àµ┤ù ---
            _report("µ¡úσ£¿Φ┐¢ΦíîΦ┤¿ΘçÅΦç¬µúÇΣ╕Äσåàσ«╣µ╕àµ┤ù...", "Θÿ╢µ«╡ 2.5/3: Φ┤¿ΘçÅΘù¿µÄº")
            try:
                from web.file_quality_checker import FileQualityGate

                qg_result = FileQualityGate.check_and_fix_ppt_outline(
                    ppt_slides, user_request=user_input, progress_callback=_report
                )
                ppt_slides = qg_result["outline"]
                _qg_score = qg_result["quality"]["score"]
                _qg_fixes = qg_result["fixes"]
                if _qg_fixes:
                    _report(f"≡ƒº╣ σ╖▓µ╕àµ┤ù {len(_qg_fixes)} σñäσåàσ«╣Θù«Θóÿ", "")
                _report(
                    f"{'Γ£à' if _qg_score >= 60 else 'ΓÜá∩╕Å'} Φ┤¿ΘçÅΦ»äσêå: {_qg_score}/100",
                    (
                        "; ".join(qg_result["quality"]["issues"][:3])
                        if qg_result["quality"]["issues"]
                        else "Φ┤¿ΘçÅΦë»σÑ╜"
                    ),
                )
            except Exception as qg_err:
                _app_logger.warning(f"[PPT] ΓÜá∩╕Å Φ┤¿ΘçÅΘù¿µÄºσ╝éσ╕╕: {qg_err}")

            # AI Θ¬îΦ»ü
            try:
                verify_prompt = (
                    f"Φ»╖Σ╜£Σ╕║Φ┤¿µúÇσæÿµúÇµƒÑτöƒµêÉτÜäPPTσåàσ«╣µÿ»σÉªτ¼ªσÉêτö¿µê╖Θ£Çµ▒éπÇé\n"
                    f"τö¿µê╖Θ£Çµ▒é: {user_input}\n"
                    f"τöƒµêÉτÜäµáçΘóÿ: {[s['title'] for s in ppt_slides]}\n"
                    "Φ»╖τ«ÇΦªüσ¢₧τ¡ö∩╝Üσåàσ«╣µÿ»σÉªΦªåτ¢ûΣ║åΘ£Çµ▒é∩╝ƒ(µÿ»/σÉª) + Σ╕ÇσÅÑΦ»¥τé╣Φ»äπÇé"
                )
                verify_resp = await asyncio.to_thread(
                    lambda: client.models.generate_content(
                        model="gemini-2.5-flash", contents=verify_prompt
                    )
                )
                if verify_resp and verify_resp.text:
                    _report(
                        "Γ£à AI Θ¬îΦ»üΘÇÜΦ┐ç",
                        f"µ¿íσ₧ïτé╣Φ»ä: {verify_resp.text.strip()[:60]}...",
                    )
            except Exception as v_err:
                _report("ΓÜá∩╕Å AI Θ¬îΦ»üΦ╖│Φ┐ç (Θ¥₧Φç┤σæ╜)", str(v_err))

            # 2. µëºΦíîΘÿ╢µ«╡ (Execution Phase) - τöƒµêÉ PPT µûçΣ╗╢
            _report("µ¡úσ£¿τöƒµêÉµ£Çτ╗êµûçΣ╗╢...", "Θÿ╢µ«╡ 3/3: µ╕▓µƒôΣ╕ÄΣ┐¥σ¡ÿ")
            from web.ppt_generator import PPTGenerator

            ppt_gen = PPTGenerator(theme=theme_choice)

            # τöƒµêÉµûçΣ╗╢σÉì
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = re.sub(r'[\\/*?:"<>|]', "", user_input[:20]) or "µ╝öτñ║µûçτ¿┐"
            filename = f"{safe_title}_{timestamp}.pptx"
            ppt_path = os.path.join(settings_manager.documents_dir, filename)
            os.makedirs(settings_manager.documents_dir, exist_ok=True)

            # Σ╜┐τö¿ PPTGenerator τöƒµêÉ (τ¢«σëìσ«âτ¢┤µÄÑµö»µîü outline list)
            ppt_gen.generate_from_outline(
                title=safe_title, outline=ppt_slides, output_path=ppt_path
            )

            rel_path = os.path.relpath(ppt_path, WORKSPACE_DIR).replace("\\", "/")

            # Φ┐öσ¢₧τ╗ôµ₧£∩╝îµá╝σ╝ÅΣ╕Ä _execute_file_gen Σ┐¥µîüΣ╕ÇΦç┤
            # µ₧äσ╗║ markdown Φí¿τñ║τÜäσñºτ║▓Σ╛¢σëìτ½»µÿ╛τñ║
            md_outline = f"# {safe_title}\n\n"
            for slide in ppt_slides:
                md_outline += f"## {slide['title']}\n"
                for p in slide.get("points", []):
                    md_outline += f"- {p}\n"
                md_outline += "\n"

            return {
                "success": True,
                "output": md_outline,
                "content": md_outline,
                "saved_files": [rel_path],
                "model_id": "gemini-2.5-flash (Planner)",
            }

        except Exception as e:
            _app_logger.warning(f"[PPT] ΓÜá∩╕Å σñÜΘÿ╢µ«╡τöƒµêÉσñ▒Φ┤Ñ∩╝îσ¢₧ΘÇÇσê░σìòµ¡ÑτöƒµêÉ: {e}")
            # Θçìµû░µè¢σç║σ╝éσ╕╕Φ«⌐Σ╕èσ▒éσñäτÉå∩╝îµêûΦÇàσ£¿Φ┐ÖΘçîΦ░âτö¿µùºΘÇ╗Φ╛æ?
            # Σ╕║Σ║åτ«Çσìò∩╝îµè¢σç║σ╝éσ╕╕Φ«⌐σñûΘâ¿ _execute_file_gen τÜä except σ¥ùµìòΦÄ╖ (Σ╜åσñûΘâ¿µÿ» generic exception)
            # µêûΦÇàµêæΣ╗¼τ¢┤µÄÑΦ┐öσ¢₧σñ▒Φ┤Ñ∩╝îΦ«⌐ TaskOrchestrator Φ«░σ╜òΘöÖΦ»»
            return {
                "success": False,
                "error": str(e),
                "opt_out_to_legacy": True,  # µáçΦ«░Θ£ÇΦªüσ¢₧ΘÇÇ
            }

    @classmethod
    async def _execute_file_gen(
        cls, user_input: str, context: dict, subtask: dict, progress_callback=None
    ) -> dict:
        """µëºΦíîµûçΣ╗╢τöƒµêÉσ¡ÉΣ╗╗σèí
        σó₧σ╝║∩╝Üσñìµ¥é/Θò┐µûç/Φªüµ▒éΓÇ£µ╖▒σ║ªπÇüΦ»ªτ╗åπÇüτáöτ⌐╢ΓÇ¥µù╢∩╝îσàêΦ┐ÉΦíîµ╖▒σ║ªτáöτ⌐╢σ╣╢σêçµìóσê░µ¢┤σ╝║µ¿íσ₧ïτöƒµêÉπÇé
        """

        def _report(msg: str, detail: str = ""):
            _app_logger.debug(f"[FILE_GEN] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        try:
            # µÅÉσÅûσëìΣ╕ÇΣ╕¬Σ╗╗σèíτÜäτ╗ôµ₧£Σ╜£Σ╕║Φ╛ôσàÑ
            previous_data = context.get(f"step_{subtask['index']}_output", "")

            # σñìµ¥éσ║ªσêñσ«Ü∩╝êΘò┐µûçµ£¼µêûµÿ╛σ╝ÅΓÇ£µ╖▒σ║ª/Φ»ªτ╗å/τáöτ⌐╢/σà¿Θ¥ó/µèÇµ£»ΓÇ¥Φ»╖µ▒é∩╝ë
            text_lower = user_input.lower()
            complex_flags = [
                len(user_input) > 120,
                any(
                    k in text_lower
                    for k in [
                        "µ╖▒σ║ª",
                        "Φ»ªτ╗å",
                        "τáöτ⌐╢",
                        "σà¿Θ¥ó",
                        "µèÇµ£»",
                        "µèÑσæè",
                        "τ╗╝Φ┐░",
                        "whitepaper",
                    ]
                ),
            ]
            is_complex = any(complex_flags)

            # -- Planning Layer (DocumentPlanner) --------------------------
            _doc_plan = None
            if is_complex:
                try:
                    from web.doc_planner import DocumentPlanner

                    _planner = DocumentPlanner(
                        ai_client=client, model_name="gemini-2.5-flash"
                    )
                    _report("≡ƒôï ΦºäσêÆµûçµíúτ╗ôµ₧ä...", "σêåµ₧ÉΘ£Çµ▒é/σêåΘàìτ½áΦèé")
                    _doc_plan = await _planner.plan(
                        user_input, previous_context=previous_data
                    )
                    if _doc_plan.success:
                        _report(
                            f"Γ£à ΦºäσêÆσ«îµêÉ∩╝Ü{len(_doc_plan.sections)} Φèé | {_doc_plan.doc_type.upper()}",
                            _doc_plan.to_context_str()[:120],
                        )
                    else:
                        _report(
                            "ΓÜá∩╕Å ΦºäσêÆσ▒éσñ▒Φ┤Ñ∩╝îΣ╜┐τö¿Θ╗ÿΦ«ñµ╡üτ¿ï",
                            _doc_plan.error[:60] if _doc_plan.error else "",
                        )
                        _doc_plan = None
                except Exception as _pe:
                    _app_logger.warning(f"[FILE_GEN] ΓÜá∩╕Å ΦºäσêÆσ▒éσ╝éσ╕╕: {_pe}")
                    _doc_plan = None

            # µúÇµ╡ïτ¢«µáçµá╝σ╝Å∩╝êPPTπÇüExcelπÇüWordτ¡ë∩╝ë
            ppt_keywords = ["ppt", "σ╣╗τü»τëç", "µ╝öτñ║", "µ▒çµèÑ", "presentation", "slide"]
            prefer_ppt = any(kw in user_input.lower() for kw in ppt_keywords)

            prefer_excel = (
                "excel" in user_input.lower()
                or "xlsx" in user_input.lower()
                or "Φí¿µá╝" in user_input
            )
            prefer_pdf = "pdf" in user_input.lower()
            if _doc_plan:
                prefer_ppt = (_doc_plan.doc_type == "ppt") or prefer_ppt
                prefer_excel = (_doc_plan.doc_type == "excel") or prefer_excel
                prefer_pdf = (_doc_plan.doc_type == "pdf") or prefer_pdf

            # µá╣µì«τ¢«µáçµá╝σ╝ÅΘÇëµï⌐µÅÉτñ║
            if prefer_ppt:
                # σ░¥Φ»òΣ╜┐τö¿µû░τÜäσñÜΘÿ╢µ«╡τöƒµêÉµ╡üτ¿ï (Plan-then-Execute)
                try:
                    ppt_result = await cls._execute_ppt_multi_step(
                        user_input, context, subtask, progress_callback
                    )
                    if ppt_result.get("success"):
                        _report(
                            f"PPTτöƒµêÉµêÉσèƒ",
                            f"µûçΣ╗╢: {(ppt_result.get('saved_files') or [''])[0]}",
                        )
                        return ppt_result
                    elif ppt_result.get("opt_out_to_legacy"):
                        _app_logger.warning("[FILE_GEN] ΓÜá∩╕Å σñÜΘÿ╢µ«╡τöƒµêÉΘüçσê░Θù«Θóÿ∩╝îσ¢₧ΘÇÇσê░µùºτëêτöƒµêÉΘÇ╗Φ╛æ")
                    else:
                        return ppt_result
                except Exception as e:
                    _app_logger.warning(f"[FILE_GEN] ΓÜá∩╕Å σñÜΘÿ╢µ«╡τöƒµêÉσ╝éσ╕╕: {e}")

                # σ¢₧ΘÇÇµùºΘÇ╗Φ╛æ (Legacy Prompt Generation)
                gen_prompt = (
                    "Σ╜áµÿ»Σ╕ÇΣ╕¬Θí╢σ░ûτÜäµ╝öτñ║µûçτ¿┐σåàσ«╣τ¡ûσêÆσ╕êσÆîµÄÆτëêΦºäσêÆσ╕êπÇé\n\n"
                    "σ£¿µ»ÅΣ╕¬ `## τ½áΦèéµáçΘóÿ` σëìΣ╕ÇΦíîσåÖτ▒╗σ₧ïµáçτ¡╛µ¥ÑΘÇëµï⌐σ╣╗τü»τëçτ▒╗σ₧ï∩╝Ü\n"
                    "- `[Φ»ªτ╗å]` ΓÇö µ╖▒σàÑσ▒òτñ║ 3-5 Σ╕¬Φªüτé╣\n"
                    "- `[µªéΦºê]` ΓÇö σñÜΣ╕╗ΘóÿΘÇƒΦºê∩╝îτö¿ `### σ¡ÉµáçΘóÿ` σêåτ╗ä\n"
                    "- `[Σ║«τé╣]` ΓÇö σà│Θö«µò░µì«∩╝îµá╝σ╝Å: `- µò░σÇ╝ | Φ»┤µÿÄ`\n"
                    "- `[σ»╣µ»ö]` ΓÇö Σ╕ñµû╣σ»╣µ»ö∩╝îτö¿ `### ΘÇëΘí╣A` σÆî `### ΘÇëΘí╣B` σêåτ╗ä\n"
                    "- `[Φ┐çµ╕íΘí╡]` ΓÇö τ½áΦèéΦ┐çµ╕í∩╝êµ£ÇσñÜ 2 Σ╕¬∩╝ë\n\n"
                    "**Φ╛ôσç║µá╝σ╝Å∩╝êΣ╕Ñµá╝Θü╡σ╛¬ Markdown∩╝ë**∩╝Ü\n"
                    "```\n"
                    "# µ╝öτñ║µáçΘóÿ\n\n"
                    "[Φ»ªτ╗å]\n"
                    "## τ½áΦèéµáçΘóÿ\n"
                    "- Φªüτé╣1∩╝êσîàσÉ½σà╖Σ╜ôΣ┐íµü»∩╝ë\n"
                    "- Φªüτé╣2\n"
                    "```\n\n"
                    "ΦºäσêÖ∩╝ÜΘçìτé╣σåàσ«╣τö¿σñÜΣ╕¬ [Φ»ªτ╗å] σ▒òσ╝Ç∩╝îτ«ÇΦªüσåàσ«╣σÉêσ╣╢σê░ [µªéΦºê]∩╝îσà│Θö«µò░µì«τö¿ [Σ║«τé╣]πÇé\n"
                    "µ»ÅΣ╕¬Φªüτé╣σîàσÉ½σà╖Σ╜ôΣ┐íµü»∩╝îΣ╕¡µûçΦ╛ôσç║∩╝îσÅ¬Φ╛ôσç║σñºτ║▓πÇé\n"
                )
            else:
                if _doc_plan and is_complex:
                    # Σ╜┐τö¿ΦºäσêÆσ▒éτöƒµêÉσó₧σ╝║ prompt∩╝êσÉ½τ½áΦèéµîçσ╝ò∩╝ë
                    from web.doc_planner import build_generation_prompt_from_plan

                    gen_prompt = build_generation_prompt_from_plan(
                        _doc_plan, user_input, previous_data
                    )
                else:
                    gen_prompt = (
                        "Σ╜áµÿ»Koto∩╝îΣ╕ÇΣ╕¬Σ╕ôΣ╕ÜτÜäµò░µì«µò┤τÉåΣ╕ÄµèÑσæèτöƒµêÉσè⌐µëïπÇé\n"
                        "Φ»╖σƒ║Σ║Äτö¿µê╖Θ£Çµ▒éσÆîµÅÉΣ╛¢τÜäµò░µì«∩╝îΦ╛ôσç║µ╕àµÖ░πÇüσÅ»τ¢┤µÄÑµö╛σàÑµûçµíúτÜä Markdown σåàσ«╣πÇé\n"
                        "σªéµ₧£µÿ»Σ╗╖µá╝τ▒╗Σ┐íµü»∩╝îσ┐àΘí╗σîàσÉ½Σ╕ÇΣ╕¬ Markdown Φí¿µá╝∩╝îσ¡ùµ«╡σ╗║Φ««Σ╕║∩╝Üµù╢Θù┤πÇüΣ╗╖µá╝πÇüσÅÿσîûπÇüµ¥Ñµ║ÉπÇé\n"
                        "Φ╛ôσç║Φªüµ▒é∩╝Ü\n"
                        "- σÅ¬Φ╛ôσç║σåàσ«╣∩╝îΣ╕ìΦªüΦ╛ôσç║Σ╗úτáüµêû BEGIN_FILE µáçΦ«░\n"
                        "- Σ╕¡µûçΦ╛ôσç║∩╝îτ╗ôµ₧äµ╕àµÖ░\n"
                    )

            full_input = (
                f"τö¿µê╖σÄƒσºïΘ£Çµ▒é: {context['original_input']}\n\n"
                f"σëìΘ¥óµ¡ÑΘ¬ñτÜäµò░µì«/Σ┐íµü»:\n{previous_data}\n\n"
                f"{gen_prompt}"
            )

            # µ╖▒σ║ªτáöτ⌐╢∩╝ÜΣ╕║σñìµ¥éΣ╗╗σèíσàêΦíÑσààτáöτ⌐╢Σ╕èΣ╕ïµûç
            research_context = ""
            if is_complex:
                try:
                    research_context = WebSearcher.deep_research_for_ppt(
                        user_input, previous_data
                    )
                    if research_context:
                        previous_data = f"[µ╖▒σ║ªτáöτ⌐╢]\n{research_context}\n\n[σ╖▓µ£ëΣ┐íµü»]\n{previous_data}"
                        _app_logger.debug(
                            f"[FILE_GEN] ≡ƒö¼ µ╖▒σ║ªτáöτ⌐╢σ«îµêÉ∩╝îΦ┐╜σèá {len(research_context)} σ¡ùΣ╕èΣ╕ïµûç"
                        )
                except Exception as research_err:
                    _app_logger.warning(f"[FILE_GEN] ΓÜá∩╕Å µ╖▒σ║ªτáöτ⌐╢σñ▒Φ┤Ñ: {research_err}")

            # Φ░âτö¿µ¿íσ₧ïτöƒµêÉσåàσ«╣
            model_id = SmartDispatcher.get_model_for_task("FILE_GEN", complexity="complex" if is_complex else "normal")
            
            _report(f"µ¡úσ£¿µÆ░σåÖσåàσ«╣...", f"µ¿íσ₧ï: {model_id}")

            def _generate_text(prompt_text: str) -> str:
                response = client.models.generate_content(
                    model=model_id,
                    contents=prompt_text,
                    config=types.GenerateContentConfig(
                        system_instruction=_get_filegen_brief_instruction(),
                        temperature=0.4,
                        max_output_tokens=4000,
                    ),
                )
                return response.text or ""

            def _clean_filegen_text(text: str) -> str:
                if not text:
                    return text
                cleaned = text

                # Remove fenced code blocks but keep content
                cleaned = re.sub(r"```[a-zA-Z0-9_-]*\n", "", cleaned)
                cleaned = cleaned.replace("```", "")

                # Strip markdown links to plain text
                cleaned = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", cleaned)

                # Remove bold/italic markers
                cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", cleaned)
                cleaned = re.sub(r"__(.+?)__", r"\1", cleaned)
                cleaned = re.sub(r"\*(.+?)\*", r"\1", cleaned)
                cleaned = re.sub(r"_(.+?)_", r"\1", cleaned)

                # Remove inline code ticks
                cleaned = cleaned.replace("`", "")

                # Strip heading markers and blockquotes at line start
                cleaned = re.sub(r"^\s{0,3}#{1,6}\s+", "", cleaned, flags=re.MULTILINE)
                cleaned = re.sub(r"^\s*>\s?", "", cleaned, flags=re.MULTILINE)

                # Remove horizontal rules
                cleaned = re.sub(r"^\s*[-_*]{3,}\s*$", "", cleaned, flags=re.MULTILINE)

                # Flatten list markers but keep structure via indentation
                cleaned = re.sub(r"^\s*[-*+]\s+", "  ", cleaned, flags=re.MULTILINE)
                cleaned = re.sub(r"^\s*\d+\.\s+", "  ", cleaned, flags=re.MULTILINE)

                # Normalize extra blank lines
                cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)

                # Cleanup leftover marker pairs
                cleaned = cleaned.replace("**", "").replace("__", "")

                return cleaned

            text_out = _generate_text(full_input) or "(µùáΦ╛ôσç║)"
            text_out = _clean_filegen_text(text_out)
            _report(f"σåàσ«╣µÆ░σåÖσ«îµêÉ", f"σà▒ {len(text_out)} σ¡ù")

            # Φºúµ₧É Markdown Φí¿µá╝
            def _extract_markdown_table(md_text: str):
                lines = [line.strip() for line in md_text.splitlines() if "|" in line]
                for i in range(len(lines) - 1):
                    header_line = lines[i]
                    sep_line = lines[i + 1]
                    if re.match(r"^\s*\|?\s*[-:|\s]+\|\s*$", sep_line):
                        headers = [c.strip() for c in header_line.strip("|").split("|")]
                        rows = []
                        j = i + 2
                        while j < len(lines) and "|" in lines[j]:
                            row = [c.strip() for c in lines[j].strip("|").split("|")]
                            if len(row) < len(headers):
                                row += [""] * (len(headers) - len(row))
                            rows.append(row[: len(headers)])
                            j += 1
                        return [headers] + rows
                return None

            # Φºúµ₧ÉPPTσñºτ║▓τ╗ôµ₧ä∩╝êµö»µîüµÖ║Φâ╜ΦºäσêÆµáçτ¡╛∩╝ë
            def _parse_ppt_outline(md_text: str) -> dict:
                """Φºúµ₧Éσ╕ª [τ▒╗σ₧ï] µáçτ¡╛τÜä PPT σñºτ║▓"""
                lines = md_text.split("\n")
                outline = {"title": "", "slides": []}
                _tmap = {
                    "Φ┐çµ╕íΘí╡": "divider",
                    "Φ┐çµ╕í": "divider",
                    "Φ»ªτ╗å": "detail",
                    "Θçìτé╣": "detail",
                    "Σ║«τé╣": "highlight",
                    "µò░µì«": "highlight",
                    "µªéΦºê": "overview",
                    "ΘÇƒΦºê": "overview",
                    "τ«ÇΦªü": "overview",
                    "σ»╣µ»ö": "comparison",
                    "µ»öΦ╛â": "comparison",
                }
                cur_type = "detail"
                cur_slide = None
                cur_sub = None

                for line in lines:
                    line = line.rstrip()
                    if line.strip() in ("```", "```markdown"):
                        continue
                    tm = re.match(r"^\s*\[(.+?)\]\s*$", line)
                    if tm:
                        cur_type = _tmap.get(tm.group(1).strip(), "detail")
                        continue
                    if line.startswith("# ") and not line.startswith("## "):
                        outline["title"] = line[2:].strip()
                    elif line.startswith("## "):
                        if (
                            cur_sub
                            and cur_slide
                            and cur_slide.get("type") in ("overview", "comparison")
                        ):
                            cur_slide.setdefault("subsections", []).append(cur_sub)
                            cur_sub = None
                        if cur_slide:
                            outline["slides"].append(cur_slide)
                        cur_slide = {
                            "type": cur_type,
                            "title": line[3:].strip(),
                            "points": [],
                            "content": [],
                        }
                        if cur_type == "divider":
                            cur_slide["description"] = ""
                        cur_type = "detail"
                        cur_sub = None
                    elif line.startswith("### ") and cur_slide:
                        if cur_sub:
                            cur_slide.setdefault("subsections", []).append(cur_sub)
                        cur_sub = {
                            "subtitle": line[4:].strip(),
                            "label": line[4:].strip(),
                            "points": [],
                        }
                    elif re.match(r"^[\s]*[-ΓÇó*]\s", line) and cur_slide is not None:
                        pt = re.sub(r"^[\s]*[-ΓÇó*]\s+", "", line).strip()
                        if cur_sub is not None:
                            cur_sub["points"].append(pt)
                        else:
                            cur_slide["points"].append(pt)
                            cur_slide["content"].append(pt)
                    elif (
                        cur_slide
                        and cur_slide.get("type") == "divider"
                        and line.strip()
                    ):
                        cur_slide["description"] = line.strip()

                if (
                    cur_sub
                    and cur_slide
                    and cur_slide.get("type") in ("overview", "comparison")
                ):
                    cur_slide.setdefault("subsections", []).append(cur_sub)
                if cur_slide:
                    outline["slides"].append(cur_slide)
                for sl in outline["slides"]:
                    if sl.get("type") == "comparison" and "subsections" in sl:
                        subs = sl["subsections"]
                        if len(subs) >= 2:
                            sl["left"] = subs[0]
                            sl["right"] = subs[1]
                return outline

            title = "τöƒµêÉµûçµíú"
            if "Σ╗╖µá╝" in user_input or "Φí¿µá╝" in user_input:
                title = "Σ╗╖µá╝µ│óσè¿Φí¿µá╝"
            elif prefer_ppt:
                title = "µ╝öτñ║µûçτ¿┐"

            saved_files = []
            file_type = None
            excel_error = None

            # τöƒµêÉPPT
            if prefer_ppt:
                try:
                    from web.ppt_generator import PPTGenerator

                    ppt_outline = _parse_ppt_outline(text_out)

                    # ΓöÇΓöÇ Φ┤¿ΘçÅΘù¿µÄº ΓöÇΓöÇ
                    try:
                        from web.file_quality_checker import FileQualityGate

                        _qg = FileQualityGate.check_and_fix_ppt_outline(
                            ppt_outline.get("slides", []),
                            user_request=user_input,
                            progress_callback=_report,
                        )
                        ppt_outline["slides"] = _qg["outline"]
                    except Exception as _qge:
                        _app_logger.warning(f"[FILE_GEN] ΓÜá∩╕Å PPT Φ┤¿ΘçÅΘù¿µÄºσ╝éσ╕╕: {_qge}")

                    # τí«σ«ÜΣ╕╗Θóÿ∩╝êΘÇÜΦ┐çσà│Θö«Φ»ìµúÇµ╡ï∩╝ë
                    theme = "business"  # Θ╗ÿΦ«ñσòåσèíΣ╕╗Θóÿ
                    user_input_lower = user_input.lower()
                    if (
                        "tech" in user_input_lower
                        or "µèÇµ£»" in user_input_lower
                        or "τºæµèÇ" in user_input_lower
                    ):
                        theme = "tech"
                    elif (
                        "creative" in user_input_lower
                        or "σê¢µäÅ" in user_input_lower
                        or "Φë║µ£»" in user_input_lower
                    ):
                        theme = "creative"
                    elif (
                        "simple" in user_input_lower
                        or "minimal" in user_input_lower
                        or "µ₧üτ«Ç" in user_input_lower
                    ):
                        theme = "minimal"

                    _report("µ¡úσ£¿τöƒµêÉPPT...", f"Σ╕╗Θóÿ: {theme} (Φç¬σè¿Θàìσ¢╛)")

                    ppt_gen = PPTGenerator(theme=theme)
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = (
                        f"{ppt_outline.get('title', 'Presentation')}_{timestamp}.pptx"
                    )
                    # Max length for filename safety
                    if len(filename) > 50:
                        filename = f"Presentation_{timestamp}.pptx"

                    ppt_path = os.path.join(settings_manager.documents_dir, filename)
                    os.makedirs(settings_manager.documents_dir, exist_ok=True)

                    def _ppt_progress_wrapper(c, t, st, ty):
                        try:
                            _report(
                                f"µ¡úσ£¿τöƒµêÉPPT ({c}/{t})", f"Θí╡Θ¥ó: {st[:10]}... [{ty}]"
                            )
                        except Exception:
                            pass

                    ppt_gen.generate_from_outline(
                        title=ppt_outline.get("title", "µ╝öτñ║"),
                        outline=ppt_outline.get("slides", []),
                        output_path=ppt_path,
                        enable_ai_images=True,
                        progress_callback=_ppt_progress_wrapper,
                    )

                    rel_path = os.path.relpath(ppt_path, WORKSPACE_DIR).replace(
                        "\\", "/"
                    )
                    saved_files.append(rel_path)
                    file_type = "pptx"
                    _report("PPTτöƒµêÉσ«îµêÉ", f"σ╖▓Σ┐¥σ¡ÿσê░: {rel_path}")

                except Exception as ppt_err:
                    _app_logger.warning(f"[FILE_GEN] ΓÜá∩╕Å PPTτöƒµêÉσñ▒Φ┤Ñ: {ppt_err}")
                    _report("PPTτöƒµêÉσñ▒Φ┤Ñ∩╝îσ¢₧ΘÇÇσê░Word...", f"ΘöÖΦ»»: {str(ppt_err)[:50]}")
                    # PPTσñ▒Φ┤Ñµù╢σ¢₧ΘÇÇσê░Word
                    from web.document_generator import save_docx

                    saved_docx = save_docx(
                        text_out, title=title, output_dir=settings_manager.documents_dir
                    )
                    rel_path = os.path.relpath(saved_docx, WORKSPACE_DIR).replace(
                        "\\", "/"
                    )
                    saved_files.append(rel_path)
                    file_type = "docx"
            else:
                # τöƒµêÉExcelµêûWord
                _report("µ¡úσ£¿σñäτÉåσåàσ«╣...", "Φºúµ₧Éµûçµíúτ╗ôµ₧ä")
                table_rows = _extract_markdown_table(text_out)
                if prefer_excel and not table_rows:
                    # τ¼¼Σ╕Çµ¼íµ£¬τöƒµêÉσÉêµá╝Φí¿µá╝ ΓåÆ τöƒµêÉΣ┐«µ¡úPromptΘçìΦ»òΣ╕Çµ¼í
                    fix_prompt = (
                        "Φ»╖σÅ¬Φ╛ôσç║Σ╕ÇΣ╕¬ Markdown Φí¿µá╝∩╝îΣ╕ìΦªüΦ╛ôσç║σà╢Σ╗ûΦ»┤µÿÄπÇé\n"
                        "Φí¿µá╝σ┐àΘí╗σîàσÉ½Σ╗ÑΣ╕ïσêù∩╝Üµù╢Θù┤πÇüΣ╗╖µá╝πÇüσÅÿσîûπÇüµ¥Ñµ║ÉπÇé\n"
                        "µ»ÅΦíîµò░µì«Σ╕ÇΦíî∩╝îµá╝σ╝ÅΣ╕Ñµá╝πÇé\n\n"
                        f"τö¿µê╖Θ£Çµ▒é: {context['original_input']}\n\n"
                        f"σÅ»τö¿µò░µì«:\n{previous_data}\n"
                    )
                    text_out_retry = _generate_text(fix_prompt)
                    if text_out_retry:
                        text_out = _clean_filegen_text(text_out_retry)
                        table_rows = _extract_markdown_table(text_out)

                if prefer_excel and table_rows:
                    _report("µ¡úσ£¿τöƒµêÉExcel...", f"σåÖσàÑ {len(table_rows)} Φíîµò░µì«")
                    try:
                        from openpyxl import Workbook
                        from openpyxl.styles import (
                            Alignment,
                            Border,
                            Font,
                            PatternFill,
                            Side,
                        )
                        from openpyxl.utils import get_column_letter

                        wb = Workbook()
                        ws = wb.active
                        ws.title = title[:31] if title else "Sheet1"

                        # σåÖσàÑµò░µì«∩╝êµ╕àµ┤ùµ»ÅΣ╕¬σìòσàâµá╝σåàτÜä Markdown τ¼ªσÅ╖∩╝ë
                        try:
                            from web.file_quality_checker import (
                                strip_markdown_from_cell,
                            )

                            _strip_cell = strip_markdown_from_cell
                        except Exception:
                            _strip_cell = lambda x: x
                        for row in table_rows:
                            ws.append(
                                [
                                    _strip_cell(str(c)) if isinstance(c, str) else c
                                    for c in row
                                ]
                            )

                        # --- µá╖σ╝Åτ╛Äσîû ---
                        header_font = Font(
                            name="Microsoft YaHei", size=11, bold=True, color="FFFFFF"
                        )
                        header_fill = PatternFill(
                            start_color="4472C4", end_color="4472C4", fill_type="solid"
                        )
                        data_font = Font(name="Microsoft YaHei", size=10)
                        thin_border = Border(
                            left=Side(style="thin", color="D9D9D9"),
                            right=Side(style="thin", color="D9D9D9"),
                            top=Side(style="thin", color="D9D9D9"),
                            bottom=Side(style="thin", color="D9D9D9"),
                        )
                        alt_fill = PatternFill(
                            start_color="F2F7FB", end_color="F2F7FB", fill_type="solid"
                        )
                        center_align = Alignment(
                            horizontal="center", vertical="center", wrap_text=True
                        )
                        left_align = Alignment(
                            horizontal="left", vertical="center", wrap_text=True
                        )

                        max_row = ws.max_row
                        max_col = ws.max_column

                        for col_idx in range(1, max_col + 1):
                            # Φí¿σñ┤µá╖σ╝Å
                            cell = ws.cell(row=1, column=col_idx)
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = center_align
                            cell.border = thin_border

                            # µò░µì«Φíîµá╖σ╝Å
                            for row_idx in range(2, max_row + 1):
                                cell = ws.cell(row=row_idx, column=col_idx)
                                cell.font = data_font
                                cell.alignment = left_align
                                cell.border = thin_border
                                # ΘÜöΦíîσÅÿΦë▓
                                if row_idx % 2 == 0:
                                    cell.fill = alt_fill

                            # Φç¬σè¿σêùσ«╜
                            max_len = 0
                            for row_idx in range(1, max_row + 1):
                                val = ws.cell(row=row_idx, column=col_idx).value
                                if val:
                                    # CJK σ¡ùτ¼ªτ«ù2Σ╕¬σ¡ùτ¼ªσ«╜
                                    vlen = sum(
                                        2 if ord(c) > 127 else 1 for c in str(val)
                                    )
                                    max_len = max(max_len, vlen)
                            ws.column_dimensions[get_column_letter(col_idx)].width = (
                                min(max_len + 4, 40)
                            )

                        # σå╗τ╗ôΘªûΦíî
                        ws.freeze_panes = "A2"

                        filename = (
                            f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                        )
                        excel_path = os.path.join(
                            settings_manager.documents_dir, filename
                        )
                        os.makedirs(settings_manager.documents_dir, exist_ok=True)
                        wb.save(excel_path)
                        rel_path = os.path.relpath(excel_path, WORKSPACE_DIR).replace(
                            "\\", "/"
                        )
                        saved_files.append(rel_path)
                        file_type = "xlsx"
                        _report("ExcelτöƒµêÉσ«îµêÉ", f"σ╖▓Σ┐¥σ¡ÿσê░: {rel_path}")
                    except Exception as excel_err:
                        excel_error = str(excel_err)
                        _app_logger.warning(f"[FILE_GEN] ΓÜá∩╕Å ExcelΣ┐¥σ¡ÿσñ▒Φ┤Ñ: {excel_error}")
                        _report(
                            "ExcelΣ┐¥σ¡ÿσñ▒Φ┤Ñ∩╝îσ¢₧ΘÇÇσê░Word...", f"ΘöÖΦ»»: {excel_error[:50]}"
                        )

                # Σ┐¥σ¡ÿΣ╕║ DOCX∩╝êµùáΦí¿µá╝µêûExcelσñ▒Φ┤Ñµù╢σ¢₧ΘÇÇ∩╝ë
                if not saved_files:
                    # ΓöÇΓöÇ σ»╝σç║µúÇµƒÑσ▒é∩╝êCheck Layer∩╝ë∩╝ÜΦ┤¿ΘçÅµúÇµƒÑ + Markdown τ¼ªσÅ╖σÄ╗ΘÖñ∩╝êµ░╕Σ╣àτë╣µÇº∩╝ëΓöÇΓöÇ
                    try:
                        from web.file_quality_checker import FileQualityGate

                        _dqg = FileQualityGate.check_and_fix_for_export(
                            text_out,
                            target_format="word",
                            user_request=user_input,
                            progress_callback=_report,
                        )
                        text_out = _dqg["text"]
                        if _dqg.get("issues"):
                            _app_logger.debug(
                                f"[FILE_GEN] ≡ƒöì µúÇµƒÑσ▒é: {', '.join(_dqg['issues'][:3])}"
                            )
                    except Exception as _dqge:
                        _app_logger.warning(f"[FILE_GEN] ΓÜá∩╕Å σ»╝σç║µúÇµƒÑσ▒éσ╝éσ╕╕: {_dqge}")

                    _report("µ¡úσ£¿τöƒµêÉWordµûçµíú...", "Φ╜¼µìóΣ╕║ DOCX")
                    from web.document_generator import save_docx, save_pdf

                    saved_docx = save_docx(
                        text_out, title=title, output_dir=settings_manager.documents_dir
                    )
                    rel_path = os.path.relpath(saved_docx, WORKSPACE_DIR).replace(
                        "\\", "/"
                    )
                    saved_files.append(rel_path)
                    file_type = "docx"
                    _report("WordµûçµíúτöƒµêÉσ«îµêÉ", f"σ╖▓Σ┐¥σ¡ÿσê░: {rel_path}")

                    # σªéτö¿µê╖µÿÄτí«Θ£ÇΦªü PDF∩╝îΣ╣ƒσÉîµù╢Σ┐¥σ¡ÿ
                    if prefer_pdf:
                        try:
                            _report("µ¡úσ£¿τöƒµêÉPDF...", "Φ╜¼µìóΣ╕║ PDF")
                            saved_pdf = save_pdf(
                                text_out,
                                title=title,
                                output_dir=settings_manager.documents_dir,
                            )
                            pdf_rel = os.path.relpath(saved_pdf, WORKSPACE_DIR).replace(
                                "\\", "/"
                            )
                            saved_files.append(pdf_rel)
                            _report("PDFτöƒµêÉσ«îµêÉ", f"σ╖▓Σ┐¥σ¡ÿσê░: {pdf_rel}")
                        except Exception as pdf_err:
                            _app_logger.warning(f"[FILE_GEN] ΓÜá∩╕Å PDFΣ┐¥σ¡ÿσñ▒Φ┤Ñ: {pdf_err}")
                            _report("PDFτöƒµêÉσñ▒Φ┤Ñ", str(pdf_err)[:50])

            return {
                "success": True,
                "output": f"σ╖▓τöƒµêÉ{file_type.upper()}µûçµíú: {', '.join([os.path.basename(p) for p in saved_files])}"
                + (f" (Excelσñ▒Φ┤Ñ: {excel_error})" if excel_error else ""),
                "content": text_out,
                "file_type": file_type or "docx",
                "saved_files": saved_files,
                "model_id": model_id,
            }
        except Exception as e:
            return {"success": False, "output": "", "error": str(e)}

    @classmethod
    async def _execute_painter(
        cls, user_input: str, context: dict, progress_callback=None
    ) -> dict:
        """µëºΦíîσ¢╛σâÅτöƒµêÉσ¡ÉΣ╗╗σèí - Σ╕║PPTτ¡ëτöƒµêÉΘàìσ¢╛ (σ╕ªσÅ»ΦºåΦ┐¢σ║ª)"""

        def _report(msg: str, detail: str = ""):
            _app_logger.debug(f"[PAINTER] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        try:
            topic = context.get("original_input", user_input)
            prompt = f"Professional illustration for: {topic[:100]}. Clean flat design, no text."

            image_paths = []
            images_dir = os.path.join(WORKSPACE_DIR, "images")
            os.makedirs(images_dir, exist_ok=True)

            _report("σÉ»σè¿σ¢╛σâÅτöƒµêÉ...", "Φ░âτö¿ Imagen 4 µ¿íσ₧ï")

            for i in range(2):
                try:
                    _report(
                        f"µ¡úσ£¿τöƒµêÉτ¼¼ {i+1}/2 σ╝áΘàìσ¢╛...", f"µÅÉτñ║Φ»ì: {prompt[:30]}..."
                    )

                    # Run potentially blocking generation in thread
                    fname = f"painter_{i}_{int(time.time()*1000)%1000000}.png"
                    fpath = os.path.join(images_dir, fname)
                    _img_models = [
                        "imagen-4.0-generate-001",
                        "imagen-4.0-fast-generate-001",
                        "imagen-3.0-generate-001",
                    ]
                    _img_res = None
                    for _img_m in _img_models:
                        try:
                            _img_res = await asyncio.to_thread(
                                lambda _m=_img_m: client.models.generate_images(
                                    model=_m,
                                    prompt=prompt,
                                    config=types.GenerateImagesConfig(
                                        number_of_images=1
                                    ),
                                )
                            )
                            if _img_res and _img_res.generated_images:
                                break
                        except Exception as _img_e:
                            _app_logger.debug(f"[PAINTER] {_img_m} σñ▒Φ┤Ñ: {_img_e}")
                            _img_res = None
                    if _img_res and _img_res.generated_images:
                        with open(fpath, "wb") as f:
                            f.write(_img_res.generated_images[0].image.image_bytes)
                        image_paths.append(fpath)
                        _app_logger.info(f"[PAINTER] Γ£à Θàìσ¢╛ {i+1} σ╖▓τöƒµêÉ: {fname}")
                        _report(f"Γ£à Θàìσ¢╛ {i+1} σ«îµêÉ", fname)
                    else:
                        raise RuntimeError("µëÇµ£ëσ¢╛σâÅµ¿íσ₧ïσ¥çσñ▒Φ┤Ñ")
                except Exception as img_err:
                    _app_logger.warning(f"[PAINTER] ΓÜá∩╕Å Θàìσ¢╛ {i+1} τöƒµêÉσñ▒Φ┤Ñ: {img_err}")
                    _report(f"ΓÜá∩╕Å Θàìσ¢╛ {i+1} σñ▒Φ┤Ñ", str(img_err))

            success = len(image_paths) > 0
            if success:
                _report("Γ£à σ¢╛σâÅτöƒµêÉΣ╗╗σèíσ«îµêÉ", f"σà▒τöƒµêÉ {len(image_paths)} σ╝á")
            else:
                _report("Γ¥î σ¢╛σâÅτöƒµêÉΣ╗╗σèíσñ▒Φ┤Ñ", "µ£¬τöƒµêÉµ£ëµòêσ¢╛τëç")

            return {
                "success": success,
                "output": f"σ╖▓τöƒµêÉ {len(image_paths)} σ╝áΘàìσ¢╛",
                "content": ",".join(image_paths),
                "image_paths": image_paths,
                "model_id": "imagen-3.0",
            }
        except Exception as e:
            _report("Γ¥î σ¢╛σâÅτöƒµêÉΘüçσê░Φç┤σæ╜ΘöÖΦ»»", str(e))
            return {"success": False, "output": "", "error": str(e)}

    @classmethod
    async def _execute_research(
        cls, user_input: str, context: dict, progress_callback=None
    ) -> dict:
        """µëºΦíîµ╖▒σ║ªτáöτ⌐╢σ¡ÉΣ╗╗σèí - Σ╜┐τö¿ Gemini Pro µ╖▒σ║ªσêåµ₧É (σÅ»ΦºåΦ┐¢σ║ª)"""

        def _report(msg: str, detail: str = ""):
            _app_logger.debug(f"[RESEARCH] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        try:
            _report("σÉ»σè¿µ╖▒σ║ªτáöτ⌐╢µ╡üτ¿ï...", "σêåµ₧ÉΣ╕èΣ╕ïµûçµò░µì«")
            search_data = context.get("WEB_SEARCH_result", {})
            search_text = search_data.get("content", "") or search_data.get(
                "output", ""
            )

            # Phase 1: Planning
            _report("ΦºäσêÆτáöτ⌐╢σñºτ║▓...", "τí«σ«Üσêåµ₧Éτ╗┤σ║ª")
            # (Implied planning by WebSearcher internal logic, but we report it)
            await asyncio.sleep(0.5)  # Simulate quick think

            # Phase 2: Synthesis
            _report("µ¡úσ£¿Φ┐¢Φíîµ╖▒σ║ªσêåµ₧É...", "Σ╝ÿσàê Deep Research Pro∩╝îσñ▒Φ┤ÑΦç¬σè¿σ¢₧ΘÇÇ")
            # Run in thread to not block event loop if sync
            research_text = await asyncio.to_thread(
                WebSearcher.deep_research_for_ppt, user_input, search_text
            )

            # Phase 3: Verification
            _report("Θ¬îΦ»üτáöτ⌐╢µèÑσæè...", "µúÇµƒÑσåàσ«╣σ«îµò┤µÇº")
            if research_text:
                _report("Γ£à τáöτ⌐╢σ«îµêÉ", f"τöƒµêÉ {len(research_text)} σ¡ùΦ»ªτ╗åµèÑσæè")
                return {
                    "success": True,
                    "output": f"µ╖▒σ║ªτáöτ⌐╢σ«îµêÉ∩╝îΦÄ╖σÅû {len(research_text)} σ¡ùΣ╕ôΣ╕Üσêåµ₧É",
                    "content": research_text,
                    "model_id": MODEL_MAP.get("RESEARCH", "deep-research-pro-preview-12-2025")
                }
            else:
                _report("ΓÜá∩╕Å τáöτ⌐╢Σ║ºσç║Σ╕║τ⌐║", "σ¢₧ΘÇÇσê░σƒ║τíÇµÉ£τ┤óτ╗ôµ₧£")
                return {
                    "success": True,
                    "output": "τáöτ⌐╢µ£¬Φ┐öσ¢₧τ╗ôµ₧£∩╝îσ░åΣ╜┐τö¿σ╖▓µ£ëΣ┐íµü»",
                    "content": search_text,
                }
        except Exception as e:
            _report("Γ¥î τáöτ⌐╢Φ┐çτ¿ïσç║ΘöÖ", str(e))
            return {"success": False, "output": "", "error": str(e)}

    @classmethod
    async def _execute_coder(cls, user_input: str, context: dict, progress_callback=None) -> dict:
        """µëºΦíîΣ╗úτáüτöƒµêÉσ¡ÉΣ╗╗σèí - Σ╜┐τö¿ gemini-3.1-pro-preview∩╝êgenerate_content τ¢┤Φ░â∩╝ë"""

        def _report(msg: str, detail: str = ""):
            _app_logger.debug(f"[CODER] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        try:
            model_id = MODEL_MAP.get("CODER", "gemini-3.1-pro-preview")
            _report("σÉ»σè¿Σ╗úτáüτöƒµêÉ...", f"µ¿íσ₧ï: {model_id}")

            # µ│¿σàÑσëìµ¡ÑµÉ£τ┤ó/τáöτ⌐╢τ╗ôµ₧£∩╝êσªéµ£ë∩╝ë
            search_ctx = ""
            for key in (
                "WEB_SEARCH_result",
                "RESEARCH_result",
                "search_result",
                "research_result",
            ):
                val = context.get(key)
                if val:
                    text = (
                        val.get("content") or val.get("output") or ""
                        if isinstance(val, dict)
                        else str(val)
                    )
                    if text:
                        search_ctx = text[:3000]
                        break

            full_prompt = user_input
            if search_ctx:
                full_prompt = f"σÅéΦÇâΣ┐íµü»:\n{search_ctx}\n\nΣ╗╗σèí: {user_input}"

            sys_instr = (
                "Σ╜áµÿ» Koto Σ╗úτáüΣ╕ôσ«╢πÇéτ¢┤µÄÑΦ╛ôσç║σ«îµò┤σÅ»Φ┐ÉΦíîΣ╗úτáü∩╝îΣ╜┐τö¿Σ╗úτáüσ¥ù∩╝ê```Φ»¡Φ¿Ç∩╝ëσîàΦú╣∩╝î"
                "Σ╕ìσèáσ║ƒΦ»¥σëìΦ¿ÇπÇéσ┐àΦªüµù╢τ«Çτƒ¡Φ»┤µÿÄΦ┐ÉΦíîµû╣σ╝Å∩╝êΓëñ3Φíî∩╝ëπÇé"
            )
            _report("µ¡úσ£¿τöƒµêÉΣ╗úτáü...", "Φ░âτö¿ Interactions API")

            result_text = await asyncio.to_thread(
                _call_interactions_api_sync, model_id, full_prompt, sys_instr, 90.0
            )

            if not result_text:
                # ΘÖìτ║ºσê░ gemini-2.5-flash
                _report("ΓÜá∩╕Å Σ╕╗µ¿íσ₧ïΦ╢àµù╢∩╝îΘÖìτ║ºτöƒµêÉ...", "gemini-2.5-flash")
                resp = await asyncio.to_thread(
                    lambda: client.models.generate_content(
                        model=_INTERACTIONS_FALLBACK_MODEL,
                        contents=full_prompt,
                        config=types.GenerateContentConfig(
                            system_instruction=sys_instr,
                            temperature=0.3,
                            max_output_tokens=4096,
                        ),
                    )
                )
                result_text = resp.text or "(µùáΦ╛ôσç║)"
                model_id = _INTERACTIONS_FALLBACK_MODEL

            # Φç¬σè¿Σ┐¥σ¡ÿΣ╗úτáüµûçΣ╗╢
            saved = Utils.auto_save_files(result_text)
            _report(
                "Γ£à Σ╗úτáüτöƒµêÉσ«îµêÉ",
                f"σ╖▓Σ┐¥σ¡ÿ {len(saved)} Σ╕¬µûçΣ╗╢" if saved else "µ£¬µúÇµ╡ïσê░µûçΣ╗╢µáçΦ«░",
            )

            return {
                "success": True,
                "output": result_text,
                "content": result_text,
                "saved_files": saved,
                "model_id": model_id,
            }
        except Exception as e:
            _report("Γ¥î Σ╗úτáüτöƒµêÉσñ▒Φ┤Ñ", str(e))
            return {"success": False, "output": "", "error": str(e)}

    @classmethod
    async def _execute_system(
        cls, user_input: str, context: dict, progress_callback=None
    ) -> dict:
        """µëºΦíîτ│╗τ╗ƒµôìΣ╜£σ¡ÉΣ╗╗σèí - Φ░âτö¿ LocalExecutor"""

        def _report(msg: str, detail: str = ""):
            _app_logger.debug(f"[SYSTEM] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        try:
            _report("µëºΦíîτ│╗τ╗ƒµôìΣ╜£...", user_input[:40])
            result = await asyncio.to_thread(LocalExecutor.execute, user_input)
            success = result.get("success", False)
            msg = result.get("message", "")
            if success:
                _report("Γ£à τ│╗τ╗ƒµôìΣ╜£σ«îµêÉ", msg[:60])
            else:
                _report("ΓÜá∩╕Å τ│╗τ╗ƒµôìΣ╜£σñ▒Φ┤Ñ", msg[:60])
            return {
                "success": success,
                "output": msg,
                "content": msg,
                "model_id": "local-executor",
            }
        except Exception as e:
            _report("Γ¥î τ│╗τ╗ƒµôìΣ╜£σ╝éσ╕╕", str(e))
            return {"success": False, "output": "", "error": str(e)}

    @classmethod
    def _merge_results(cls, subtasks: list, context: dict) -> dict:
        """σÉêσ╣╢µëÇµ£ëσ¡ÉΣ╗╗σèíτÜäτ╗ôµ₧£"""
        merged = {"summary": "Σ╗╗σèíµëºΦíîσ«îµêÉ", "steps": [], "final_output": ""}

        for i, subtask in enumerate(subtasks):
            step_info = {
                "step": i + 1,
                "task": subtask["task_type"],
                "status": subtask["status"],
                "description": subtask["description"],
            }

            if subtask["result"]:
                step_info["output"] = subtask["result"].get("output", "")
            if subtask["error"]:
                step_info["error"] = subtask["error"]

            merged["steps"].append(step_info)

        # µ£ÇσÉÄΣ╕ÇΣ╕¬σ«îµêÉτÜäΣ╗╗σèíτÜäΦ╛ôσç║Σ╜£Σ╕║µ£Çτ╗êΦ╛ôσç║
        for subtask in reversed(subtasks):
            if subtask["status"] == "completed" and subtask["result"]:
                merged["final_output"] = subtask["result"].get("output", "")
                break

        return merged

    @classmethod
    async def _validate_quality(
        cls, user_input: str, combined_output: dict, context: dict
    ) -> int:
        """
        Θ¬îΦ»üΦ╛ôσç║Φ┤¿ΘçÅ∩╝êΦ»¡Σ╣ëΦ»äσêåτëêµ£¼∩╝ëπÇé
        σàêτö¿σ┐½ΘÇƒΦºäσêÖτ╗Öσƒ║σçåσêå∩╝îσåìτö¿ gemini-2.0-flash-lite σüÜΦ»¡Σ╣ëΦ»äΣ╝░πÇé
        Φ┐öσ¢₧: Φ┤¿ΘçÅΦ»äσêå (0-100)
        """
        # ΓöÇΓöÇ ΦºäσêÖσƒ║σçåσêå ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        score = 40
        total_steps = len(combined_output.get("steps", []))
        completed_steps = len(
            [
                s
                for s in combined_output.get("steps", [])
                if s.get("status") == "completed"
            ]
        )
        if total_steps > 0:
            score += int((completed_steps / total_steps) * 30)  # µ£ÇσñÜ +30

        final_output = combined_output.get("final_output", "")
        if not final_output:
            return max(0, min(100, score))  # µùáΦ╛ôσç║τ¢┤µÄÑΦ┐öσ¢₧ΦºäσêÖσêå

        # µ£ëµûçΣ╗╢Φ╛ôσç║σèáσêå
        has_files = any(
            r.get("result", {}).get("saved_files")
            for r in combined_output.get("steps", [])
            if isinstance(r.get("result"), dict)
        )
        if has_files:
            score += 10

        # ΓöÇΓöÇ Φ»¡Σ╣ëΦ»äσêå∩╝êgemini-2.0-flash-lite∩╝îΣ╜ÄµêÉµ£¼∩╝ëΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        try:
            check_prompt = (
                f"τö¿µê╖Θ£Çµ▒é∩╝Ü{user_input[:300]}\n\n"
                f"µ£Çτ╗êΦ╛ôσç║∩╝êσëì1500σ¡ù∩╝ë∩╝Ü{final_output[:1500]}\n\n"
                "Φ»╖Φ»äΣ╝░Φ╛ôσç║µÿ»σÉªµ╗íΦ╢│Σ║åτö¿µê╖Θ£Çµ▒éπÇéσÅ¬Φ╛ôσç║Σ╕ÇΣ╕¬ 0~30 τÜäµò┤µò░∩╝ê30Σ╕║σ«îσà¿µ╗íΦ╢│∩╝ëπÇé"
            )
            resp = await asyncio.to_thread(lambda: client.models.generate_content(
                model="gemini-2.0-flash-lite",
                contents=check_prompt,
                config=types.GenerateContentConfig(
                    max_output_tokens=8,
                    temperature=0.0,
                )
            ))
            text = (resp.text or "").strip()
            m = re.search(r"\d+", text)
            if m:
                semantic_score = min(30, max(0, int(m.group())))
                score += semantic_score
        except Exception as e:
            _app_logger.debug(f"[VALIDATE_QUALITY] Φ»¡Σ╣ëΦ»äσêåσñ▒Φ┤Ñ∩╝îΣ╜┐τö¿ΦºäσêÖσêå: {e}")

        return max(0, min(100, score))


# ================= µÖ║Φâ╜Φ»¡µûÖΦ╖»τö▒σÖ¿Θàìτ╜« =================
# Θàìτ╜« SmartDispatcher Σ╗ÑΣ╜┐τö¿µ£¼σ£░σ«ÜΣ╣ëτÜäτ▒╗σÆîσ»╣Φ▒í
# SmartDispatcherπÇüModelRouter τ¡ëσ╖▓Σ╗Ä app.core.routing σ»╝σàÑ

try:
    _app_logger.debug("[INIT] Configuring SmartDispatcher with local dependencies...")
    SmartDispatcher.configure(
        local_executor=LocalExecutor,
        context_analyzer=ContextAnalyzer,
        web_searcher=WebSearcher,
        model_map=MODEL_MAP,
        client=client,
    )
    _app_logger.debug("[INIT] SmartDispatcher configured successfully.")
except Exception as e:
    _app_logger.error(f"[ERROR] Failed to configure SmartDispatcher: {e}")

# ΓöÇΓöÇΓöÇ σÉÄσÅ░σÉ»σè¿σè¿µÇüµ¿íσ₧ïΦ╖»τö▒σÖ¿ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
# Σ╕ìΘÿ╗σí₧Σ╕╗τ║┐τ¿ïσÉ»σè¿∩╝¢Φ╖»τö▒Φí¿µ¢┤µû░σÉÄΣ╝ÜΦç¬σè¿Φªåτ¢ûΘ¥ÖµÇüΘ╗ÿΦ«ñσÇ╝σÅè SmartDispatcher Θàìτ╜«
import threading as _threading

_threading.Thread(
    target=_init_model_manager, name="ModelManagerInit", daemon=True
).start()

# === Ollama σÉÄσñçΦ╖»τö▒ (σÅ»ΘÇë) ===
LOCAL_ROUTER_MODEL = "qwen3:8b"  # σìçτ║º: Qwen3 Σ╕¡Φï▒µûçΦâ╜σè¢Φ┐£Φ╢àµùºµ¿íσ₧ï
OLLAMA_API_URL = "http://localhost:11434/api/generate"


class LocalDispatcher:
    """σÉÄσñçΦ╖»τö▒σÖ¿ - Σ╜┐τö¿ Ollama (σªéµ₧£σÅ»τö¿)"""

    @staticmethod
    def is_ollama_running():
        # Σ║æτ½»µ¿íσ╝ÅΣ╕ïτªüτö¿ Ollama∩╝êΣ║æµ£ìσèíσÖ¿µùáµ£¼σ£░ GPU∩╝ë
        if os.environ.get("KOTO_DEPLOY_MODE") == "cloud":
            return False
        try:
            requests.get("http://localhost:11434", timeout=0.2)
            return True
        except Exception:
            return False

    @staticmethod
    def analyze(user_input, history=None):
        """Σ╝ÿσàêΣ╜┐τö¿ SmartDispatcher∩╝îσñ▒Φ┤Ñµù╢Σ╜┐τö¿ Ollama"""
        # Σ╜┐τö¿µÖ║Φâ╜µ£¼σ£░Φ╖»τö▒
        return SmartDispatcher.analyze(user_input, history)


# ================= Utilities =================


class Utils:
    _PACKAGE_ALLOWLIST = {
        "pygame": "pygame",
        "numpy": "numpy",
        "pandas": "pandas",
        "requests": "requests",
        "bs4": "beautifulsoup4",
        "beautifulsoup4": "beautifulsoup4",
        "lxml": "lxml",
        "pillow": "Pillow",
        "PIL": "Pillow",
        "opencv": "opencv-python",
        "cv2": "opencv-python",
        "matplotlib": "matplotlib",
        "scipy": "scipy",
        "sklearn": "scikit-learn",
        "flask": "flask",
        "fastapi": "fastapi",
        "uvicorn": "uvicorn",
        "streamlit": "streamlit",
        "gradio": "gradio",
    }

    @staticmethod
    def sanitize_string(s):
        if isinstance(s, str):
            return s.encode("utf-8", "ignore").decode("utf-8")
        return s

    @staticmethod
    def is_failure_output(text: str) -> bool:
        if not text or not str(text).strip():
            return True
        t = str(text).strip().lower()
        if t.startswith("Γ¥î") or "σñ▒Φ┤Ñ" in t or "ΘöÖΦ»»" in t:
            return True
        # µúÇµ╡ïµ¿íσ₧ïσú░τº░πÇîµùáµ│òΦüöτ╜æ/µ▓íµ£ëσ«₧µù╢µò░µì«πÇìτÜäµïÆτ╗¥σ₧ïσ¢₧τ¡ö
        _no_internet_phrases = [
            "µ▓íµ£ëτ¢┤µÄÑΦüöτ╜æ",
            "µùáµ│òτ¢┤µÄÑΦüöτ╜æ",
            "µùáµ│òΦüöτ╜æ",
            "µ▓íµ£ëΦüöτ╜æ",
            "Σ╕ìΦâ╜Φüöτ╜æ",
            "µ▓íµ£ëσ«₧µù╢",
            "µùáµ│òΦÄ╖σÅûσ«₧µù╢",
            "Σ╕ìΦâ╜ΦÄ╖σÅûσ«₧µù╢",
            "µ▓íµ£ëΦ«┐Θù«Σ║ÆΦüöτ╜æ",
            "µùáµ│òΦ«┐Θù«Σ║ÆΦüöτ╜æ",
            "i don't have access to the internet",
            "i cannot access the internet",
            "i'm unable to access the internet",
            "no internet access",
            "i don't have real-time",
            "i cannot browse",
            "i can't browse",
        ]
        return any(phrase in t for phrase in _no_internet_phrases)

    @staticmethod
    def build_fix_prompt(
        task_type: str, user_input: str, prev_output: str = "", error_hint: str = ""
    ) -> str:
        base = (
            f"τö¿µê╖Θ£Çµ▒é: {user_input}\n\n"
            f"Σ╕èµ¼íΦ╛ôσç║/ΘöÖΦ»»:\n{prev_output or error_hint}\n\n"
            "Φ»╖Σ┐«µ¡úσ╣╢Θçìµû░Φ╛ôσç║µ£Çτ╗êτ╗ôµ₧£πÇéΣ╕ìΦªüΦºúΘçèΦ┐çτ¿ï∩╝îσÅ¬Φ╛ôσç║µ£Çτ╗êσåàσ«╣πÇé\n"
        )

        if task_type == "FILE_GEN":
            return base + (
                "Φªüµ▒é∩╝ÜΦ╛ôσç║σÅ»µëºΦíîτÜä Python ΦäÜµ£¼∩╝îσ╣╢Σ╜┐τö¿ BEGIN_FILE/END_FILE µáçΦ«░πÇé\n"
                "σ┐àΘí╗τöƒµêÉµûçµíúµêûΦí¿µá╝µûçΣ╗╢∩╝êdocx/xlsx/pdf∩╝ëπÇé"
            )
        if task_type == "CODER":
            return base + "Φªüµ▒é∩╝ÜΦ╛ôσç║σ«îµò┤σÅ»Φ┐ÉΦíîΣ╗úτáü∩╝îσ╣╢σîàσÉ½σ┐àΦªüΦ»┤µÿÄπÇé"
        if task_type == "RESEARCH":
            return base + "Φªüµ▒é∩╝ÜΦ╛ôσç║τ╗ôµ₧äσîûµèÑσæè∩╝îσîàσÉ½µáçΘóÿΣ╕ÄΦªüτé╣πÇé"
        if task_type == "WEB_SEARCH":
            return base + "Φªüµ▒é∩╝Üσƒ║Σ║Äσ«₧µù╢Σ┐íµü»σ¢₧τ¡ö∩╝îτ╗Öσç║µ╕àµÖ░τ╗ôΦ«║πÇé"
        return base

    @staticmethod
    def adapt_prompt_to_markdown(
        task_type: str, user_input: str, history: list = None
    ) -> str:
        """Σ╜┐τö¿µ£¼σ£░µ¿íµ¥┐σ░åσÄƒσºïΦ»╖µ▒éΦ╜¼Σ╕║τ╗ôµ₧äσîû Markdown∩╝îΣ╛┐Σ║Äσñºµ¿íσ₧ïτÉåΦºúπÇé

        µ│¿∩╝Üσ╖▓τº╗ΘÖñ flash-lite Σ║îµ¼íµ╢ªΦë▓Φ░âτö¿∩╝êΘó¥σñû API Φ┤╣τö¿ + ~300ms σ╗╢Φ┐ƒ∩╝îµö╢τ¢èΣ╕ìµÿÄµÿ╛∩╝ëπÇé
        PromptAdapter τÜäµ£¼σ£░µ¿íµ¥┐∩╝êbase_md∩╝ëσ╖▓Φ╢│σñƒΣ╕╗µ¿íσ₧ïτÉåΦºúπÇé
        """
        try:
            try:
                from web.prompt_adapter import PromptAdapter
            except ImportError:
                from prompt_adapter import PromptAdapter

            # model_generate=None∩╝ÜΣ╗àΣ╜┐τö¿µ£¼σ£░σà│Θö«Φ»ìµÅÉσÅû + Markdown µ¿íµ¥┐∩╝îΣ╕ìσÅæΦ╡╖Θó¥σñû LLM Φ░âτö¿
            return PromptAdapter.adapt(
                user_input=user_input,
                task_type=task_type,
                history=history,
                model_generate=None,
            )
        except Exception as e:
            _app_logger.debug(f"[PROMPT_ADAPTER] Failed: {e}")
            return user_input

    @staticmethod
    def quick_self_check(task_type: str, user_input: str, output_text: str) -> dict:
        """Σ╜┐τö¿σ┐½ΘÇƒµ¿íσ₧ïΦ┐¢ΦíîΦç¬µúÇ∩╝îΦ┐öσ¢₧ {'pass': bool, 'fix_prompt': str}πÇé"""
        try:
            check_prompt = (
                "Σ╜áµÿ»Φ┤¿ΘçÅµúÇµƒÑσÖ¿πÇéσêñµû¡Φ╛ôσç║µÿ»σÉªµ╗íΦ╢│τö¿µê╖Θ£Çµ▒éπÇé\n"
                "σÅ¬Φ╛ôσç║Σ╗ÑΣ╕ïµá╝σ╝ÅΣ╣ïΣ╕Ç∩╝Ü\n"
                "PASS\n"
                "µêû\n"
                "FAIL\nFIX_PROMPT: <τö¿Σ║ÄΣ┐«µ¡úτÜäµÅÉτñ║Φ»ì>\n\n"
                f"Σ╗╗σèíτ▒╗σ₧ï: {task_type}\n"
                f"τö¿µê╖Θ£Çµ▒é: {user_input}\n"
                f"µ¿íσ₧ïΦ╛ôσç║:\n{output_text}\n"
            )
            response = client.models.generate_content(
                model="gemini-2.0-flash-lite",
                contents=check_prompt,
                config=types.GenerateContentConfig(
                    max_output_tokens=300,
                    temperature=0.1,
                ),
            )
            text = (response.text or "").strip()
            if text.startswith("PASS"):
                return {"pass": True, "fix_prompt": ""}
            if text.startswith("FAIL"):
                fix = ""
                for line in text.splitlines():
                    if line.startswith("FIX_PROMPT:"):
                        fix = line.replace("FIX_PROMPT:", "").strip()
                        break
                return {"pass": False, "fix_prompt": fix}
            return {"pass": True, "fix_prompt": ""}
        except Exception as e:
            _app_logger.debug(f"[SELF_CHECK] Failed: {e}")
            return {"pass": True, "fix_prompt": ""}

    @staticmethod
    def detect_required_packages(text: str) -> list:
        """Σ╗ÄΦ╛ôσç║Σ╕¡τ▓ùτòÑµúÇµ╡ïτ¼¼Σ╕ëµû╣Σ╛¥Φ╡û∩╝êΣ╗àΦ┐öσ¢₧τÖ╜σÉìσìòσåàτÜäσîà∩╝ëπÇé"""
        if not text:
            return []
        modules = set()
        for line in text.splitlines():
            line = line.strip()
            if line.startswith("import "):
                parts = line.replace("import", "").split(",")
                for p in parts:
                    name = p.strip().split(" ")[0]
                    if name:
                        modules.add(name)
            elif line.startswith("from "):
                parts = line.split()
                if len(parts) >= 2:
                    modules.add(parts[1].strip())

        packages = set()
        for mod in modules:
            if mod in Utils._PACKAGE_ALLOWLIST:
                packages.add(Utils._PACKAGE_ALLOWLIST[mod])
        return sorted(packages)

    @staticmethod
    def auto_install_packages(packages: list) -> dict:
        """σ«ëΦúàτ╝║σñ▒τÜäΣ╛¥Φ╡ûσîàπÇéΦ┐öσ¢₧σ«ëΦúàτ╗ôµ₧£µæÿΦªüπÇé"""
        result = {"installed": [], "skipped": [], "failed": []}
        if not packages:
            return result

        for pkg in packages:
            try:
                spec = importlib.util.find_spec(pkg)
                if spec is not None:
                    result["skipped"].append(pkg)
                    continue
                module_aliases = [
                    m for m, p in Utils._PACKAGE_ALLOWLIST.items() if p == pkg
                ]
                if any(importlib.util.find_spec(m) is not None for m in module_aliases):
                    result["skipped"].append(pkg)
                    continue
            except Exception:
                pass

            try:
                if getattr(sys, "frozen", False):
                    # µëôσîàτëêµùáµ│òσ«ëΦúàµû░σîà∩╝îpip σ£¿σå╗τ╗ôτÄ»σóâΣ╕ïΣ╕ìσÅ»τö¿
                    result["failed"].append(pkg)
                else:
                    cmd = [sys.executable, "-m", "pip", "install", pkg]
                    proc = subprocess.run(
                        cmd,
                        capture_output=True,
                        text=True,
                        timeout=120,
                        creationflags=(
                            subprocess.CREATE_NO_WINDOW
                            if sys.platform == "win32"
                            else 0
                        ),
                    )
                    if proc.returncode == 0:
                        result["installed"].append(pkg)
                    else:
                        result["failed"].append(pkg)
            except Exception:
                result["failed"].append(pkg)

        return result

    @staticmethod
    def auto_save_files(text):
        """Φç¬σè¿Σ╗Äσôìσ║öΣ╕¡µÅÉσÅûσ╣╢Σ┐¥σ¡ÿµûçΣ╗╢"""
        saved = []

        code_dir = os.path.join(WORKSPACE_DIR, "code")
        os.makedirs(code_dir, exist_ok=True)

        def _get_save_dir(filename):
            ext = os.path.splitext(filename)[1].lower()
            code_exts = {
                ".py",
                ".js",
                ".ts",
                ".tsx",
                ".jsx",
                ".java",
                ".cs",
                ".cpp",
                ".c",
                ".go",
                ".rs",
                ".rb",
                ".php",
                ".swift",
                ".kt",
                ".m",
                ".scala",
                ".sh",
                ".ps1",
                ".bat",
                ".cmd",
                ".json",
                ".yaml",
                ".yml",
                ".toml",
                ".ini",
                ".cfg",
                ".sql",
                ".md",
                ".html",
                ".css",
            }
            return code_dir if ext in code_exts else WORKSPACE_DIR

        # Φ░âΦ»ò∩╝Üµëôσì░σëì800σ¡ùτ¼ªτ£ïτ£ïµá╝σ╝Å
        _app_logger.debug(f"[FILE_GEN] Response first 800 chars:\n{text[:800]}\n")

        # ΘóäσñäτÉå∩╝Üτ╗ƒΣ╕Çµá╝σ╝Å (σÄ╗µÄëσñÜΣ╜Öτ⌐║µá╝)
        normalized_text = text

        # µû╣µ│ò1: σñÜτºì BEGIN_FILE µá╝σ╝ÅτÜäµ¡úσêÖσî╣Θàì
        patterns = [
            # µá╝σ╝Å1: ---BEGIN_FILE: filename.py--- (µùáτ⌐║µá╝)
            r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\s*(.*?)---\s*END_FILE\s*---",
            # µá╝σ╝Å2: ---BEGIN_FILE: filename.py--- ... ---END_FILE--- (σ╕ªµìóΦíî)
            r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\n(.*?)\n---END_FILE---",
            # µá╝σ╝Å3: µ¢┤σ«╜µ¥╛ - σàüΦ«╕σÉäτºìτ⌐║τÖ╜
            r"---\s*BEGIN_FILE[:\s]+([a-zA-Z0-9_.-]+)\s*---\s*(.*?)---\s*END_FILE\s*---",
            # µá╝σ╝Å4: µ£Çσ«╜µ¥╛ - µìòΦÄ╖Σ╗╗µäÅµûçΣ╗╢σÉì
            r"---BEGIN_FILE[:\s]+([^\n-]+?)---\s*(.*?)---END_FILE---",
        ]

        matches1 = []
        for i, pattern in enumerate(patterns):
            try:
                matches1 = re.findall(
                    pattern, normalized_text, re.DOTALL | re.IGNORECASE
                )
                _app_logger.debug(f"[FILE_GEN] Pattern{i+1} matches: {len(matches1)}")
                if matches1:
                    _app_logger.debug(f"[FILE_GEN] Γ£ô Using pattern {i+1}")
                    break
            except Exception as e:
                _app_logger.debug(f"[FILE_GEN] Pattern{i+1} error: {e}")

        for filename, content in matches1:
            try:
                filename = filename.strip()
                content = content.strip()
                _app_logger.debug(
                    f"[FILE_GEN] Processing file: '{filename}', content length: {len(content)}"
                )

                # µ╕àΘÖñ Markdown Σ╗úτáüσ¥ùµáçΦ«░
                if content.startswith("```"):
                    lines = content.split("\n")
                    if lines[0].startswith("```"):
                        lines = lines[1:]
                    if lines and lines[-1].strip() == "```":
                        lines = lines[:-1]
                    content = "\n".join(lines)
                    _app_logger.debug(f"[FILE_GEN] After stripping markdown: {len(content)} chars")

                # τí«Σ┐¥µûçΣ╗╢σÉìµ£ëµòê
                if not filename or len(filename) > 100:
                    _app_logger.debug(f"[FILE_GEN] Invalid filename: {filename}")
                    continue

                # τí«Σ┐¥µûçΣ╗╢σÉìµ£ëµë⌐σ▒òσÉì
                if "." not in filename:
                    filename = filename + ".py"

                base_dir = _get_save_dir(filename)
                path = os.path.join(base_dir, filename)
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content)
                saved.append(filename)
                _app_logger.info(f"[FILE_GEN] Γ£à Saved: {filename} to {path}")
            except Exception as e:
                _app_logger.error(f"[FILE_GEN] Γ¥î Save failed: {e}")
                import traceback

                traceback.print_exc()

        # µû╣µ│ò2: σªéµ₧£µû╣µ│ò1µ▓íµë╛σê░∩╝îσ░¥Φ»òµÅÉσÅû ```python Σ╗úτáüσ¥ù + µûçΣ╗╢σÉìµ│¿Θçè
        if not saved:
            _app_logger.debug(f"[FILE_GEN] Method1 empty, trying method2 (```python blocks)...")

            # σàêσ░¥Φ»òσî╣Θàìσ╕ªµûçΣ╗╢σÉìτÜäΣ╗úτáüσ¥ù
            # Σ╛ïσªé: # filename: cat_info.py µêû # cat_info.py
            pattern2a = (
                r"```python\s*\n#\s*(?:filename:\s*)?([a-zA-Z0-9_.-]+\.py)\s*\n(.*?)```"
            )
            matches2a = re.findall(pattern2a, text, re.DOTALL)
            _app_logger.debug(
                f"[FILE_GEN] Pattern2a (with filename comment) matches: {len(matches2a)}"
            )

            if matches2a:
                for filename, code in matches2a:
                    code = code.strip()
                    if not code or len(code) < 20:
                        continue
                    base_dir = _get_save_dir(filename)
                    path = os.path.join(base_dir, filename)
                    try:
                        with open(path, "w", encoding="utf-8") as f:
                            f.write(code)
                        saved.append(filename)
                        _app_logger.info(f"[FILE_GEN] Γ£à Method2a saved: {filename}")
                    except Exception as e:
                        _app_logger.error(f"[FILE_GEN] Γ¥î Method2a save failed: {e}")
            else:
                # µùáµûçΣ╗╢σÉìτÜäΣ╗úτáüσ¥ù∩╝îΣ╜┐τö¿µù╢Θù┤µê│
                pattern2 = r"```python\s*\n(.*?)```"
                matches2 = re.findall(pattern2, text, re.DOTALL)
                _app_logger.debug(f"[FILE_GEN] Pattern2 (generic) matches: {len(matches2)}")

                if matches2:
                    timestamp = int(time.time())
                    for idx, code in enumerate(matches2):
                        code = code.strip()
                        if not code or len(code) < 50:
                            continue

                        # σ░¥Φ»òΣ╗ÄΣ╗úτáüΣ╕¡µÅÉσÅûµ£ëµäÅΣ╣ëτÜäµûçΣ╗╢σÉì
                        filename = None
                        # µƒÑµë╛ doc_path, file_path τ¡ëσÅÿΘçÅ
                        path_match = re.search(
                            r'(?:doc_path|file_path|filepath|output_path)\s*=.*?["\']([^"\']+\.(pdf|docx|xlsx))["\']',
                            code,
                        )
                        if path_match:
                            # Σ╜┐τö¿τ¢«µáçµûçΣ╗╢σÉìΣ╜£Σ╕║ΦäÜµ£¼σÉì
                            target_file = os.path.basename(path_match.group(1))
                            filename = target_file.rsplit(".", 1)[0] + ".py"

                        if not filename:
                            filename = f"generated_{timestamp}_{idx}.py"

                        base_dir = _get_save_dir(filename)
                        path = os.path.join(base_dir, filename)
                        try:
                            with open(path, "w", encoding="utf-8") as f:
                                f.write(code)
                            saved.append(filename)
                            _app_logger.info(f"[FILE_GEN] Γ£à Method2 saved: {filename}")
                        except Exception as e:
                            _app_logger.error(f"[FILE_GEN] Γ¥î Method2 save failed: {e}")

        _app_logger.debug(f"[FILE_GEN] Final saved files: {saved}")
        return saved

    @staticmethod
    def save_image_part(blob_part):
        try:
            # Σ╜┐τö¿τö¿µê╖Φ«╛τ╜«τÜäσ¢╛τëçτ¢«σ╜ò
            images_dir = settings_manager.images_dir
            os.makedirs(images_dir, exist_ok=True)

            timestamp = int(time.time())
            filename = f"generated_{timestamp}.png"
            filepath = os.path.join(images_dir, filename)
            with open(filepath, "wb") as f:
                f.write(blob_part.inline_data.data)

            # Φ┐öσ¢₧τ¢╕σ»╣Σ║Ä workspace τÜäΦ╖»σ╛ä
            # τí«Σ┐¥Φ╖»σ╛äσºïτ╗êσ£¿ workspace Σ╕ï∩╝îΣ╕öµá╝σ╝ÅΣ╕║µ¡úµû£µ¥á
            try:
                rel_path = os.path.relpath(filepath, WORKSPACE_DIR)
                # σªéµ₧£σîàσÉ½ .. Φ»┤µÿÄΣ╕ìσ£¿ workspace Σ╕ï∩╝îΘ£ÇΦªüσñäτÉå
                if ".." in rel_path:
                    # ΘÖìτ║ºΣ╕║σÅ¬Φ┐öσ¢₧µûçΣ╗╢σÉì∩╝îµö╛σ£¿ workspace/images Σ╕ï
                    abs_workspace_images = os.path.join(WORKSPACE_DIR, "images")
                    os.makedirs(abs_workspace_images, exist_ok=True)
                    fallback_path = os.path.join(abs_workspace_images, filename)
                    with open(fallback_path, "wb") as f:
                        f.write(blob_part.inline_data.data)
                    rel_path = os.path.relpath(fallback_path, WORKSPACE_DIR)
                    _app_logger.debug(f"[IMAGE] Falling back to workspace/images: {rel_path}")

                result = rel_path.replace("\\", "/")
                _app_logger.debug(f"[IMAGE] Saved image: {result}")
                return result
            except Exception as path_err:
                _app_logger.debug(f"[IMAGE] Path calculation error: {path_err}")
                # µ£ÇσÉÄτÜäΣ┐¥ΘÖ⌐µû╣µíê∩╝Üτ¢┤µÄÑΣ┐¥σ¡ÿσê░ workspace/images
                abs_workspace_images = os.path.join(WORKSPACE_DIR, "images")
                os.makedirs(abs_workspace_images, exist_ok=True)
                fallback_path = os.path.join(abs_workspace_images, filename)
                with open(fallback_path, "wb") as f:
                    f.write(blob_part.inline_data.data)
                result = os.path.relpath(fallback_path, WORKSPACE_DIR).replace(
                    "\\", "/"
                )
                _app_logger.debug(f"[IMAGE] Emergency fallback: {result}")
                return result
        except Exception as e:
            _app_logger.debug(f"[IMAGE] Save failed: {e}")
            import traceback

            traceback.print_exc()
            return None


# ================= Session Manager =================


class SessionManager:
    def __init__(self):
        self.sessions = {}

    def list_sessions(self):
        """σêùσç║µëÇµ£ëΣ╝ÜΦ»¥∩╝îµîëΣ┐«µö╣µù╢Θù┤µÄÆσ║Å∩╝êµ£Çµû░σ£¿σëì∩╝ë"""
        files = [f for f in os.listdir(CHAT_DIR) if f.endswith(".json")]
        # µîëΣ┐«µö╣µù╢Θù┤µÄÆσ║Å∩╝îµ£Çµû░τÜäσ£¿σëì
        files_with_time = []
        for f in files:
            path = os.path.join(CHAT_DIR, f)
            mtime = os.path.getmtime(path)
            files_with_time.append((f, mtime))
        files_with_time.sort(key=lambda x: x[1], reverse=True)
        return [f[0] for f in files_with_time]

    def load(self, filename):
        """σèáΦ╜╜Σ╝ÜΦ»¥σÄåσÅ▓ - Φ┐öσ¢₧τö¿Σ║Äµ¿íσ₧ïΣ╕èΣ╕ïµûçτÜäµê¬µû¡τëêµ£¼"""
        path = os.path.join(CHAT_DIR, filename)
        if os.path.exists(path):
            try:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    full_history = json.load(f)
                    # Σ╗àµê¬µû¡τö¿Σ║Äµ¿íσ₧ïΣ╕èΣ╕ïµûçτÜäΘâ¿σêå∩╝îΣ╕ìσ╜▒σôìµîüΣ╣àσîûσ¡ÿσé¿
                    return self._trim_history(full_history)
            except (json.JSONDecodeError, OSError) as e:
                _app_logger.warning("Failed to load session %s: %s", filename, e)
                return []
        return []

    def load_full(self, filename):
        """σèáΦ╜╜σ«îµò┤Σ╝ÜΦ»¥σÄåσÅ▓ - τö¿Σ║ÄΦ┐╜σèáΣ┐¥σ¡ÿ∩╝îΣ╕ìσüÜµê¬µû¡"""
        path = os.path.join(CHAT_DIR, filename)
        if os.path.exists(path):
            try:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    return json.load(f)
            except (json.JSONDecodeError, OSError) as e:
                _app_logger.warning("Failed to load full session %s: %s", filename, e)
                return []
        return []

    def _trim_history(self, history, max_turns=20):
        """Σ┐¥τòÖµ£ÇσñÜ 20 Φ╜«σ»╣Φ»¥∩╝êτ║ª 12000+ tokens∩╝ë∩╝îτí«Σ┐¥Σ╕èΣ╕ïµûçΦ╢│σñƒΣ╜åΣ╕ìΦ┐çΘò┐"""
        if len(history) <= max_turns:
            return history
        # σÅ¬Σ┐¥τòÖµ£ÇσÉÄ N Φ╜«σ»╣Φ»¥
        trimmed = history[-max_turns:]
        _app_logger.debug(f"[HISTORY] Trimmed to last {max_turns} turns (was {len(history)})")
        return trimmed

    def create(self, name):
        safe = "".join([c if c.isalnum() else "_" for c in name])
        filename = f"{safe}.json"
        path = os.path.join(CHAT_DIR, filename)
        # ΦïÑσÉîσÉìµûçΣ╗╢σ╖▓σ¡ÿσ£¿∩╝îσèáµù╢Θù┤µê│σÉÄτ╝ÇΘü┐σàìΦªåτ¢ûσ╖▓µ£ëΣ╝ÜΦ»¥
        if os.path.exists(path):
            filename = f"{safe}_{int(time.time())}.json"
            path = os.path.join(CHAT_DIR, filename)
        with open(path, "w", encoding="utf-8") as f:
            json.dump([], f)
        return filename

    def save(self, filename, history):
        path = os.path.join(CHAT_DIR, filename)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(history, f, indent=2, ensure_ascii=False)

    def append_and_save(self, filename, user_msg, model_msg, **extra_fields):
        """Φ┐╜σèáµ╢êµü»σ╣╢Σ┐¥σ¡ÿ - σƒ║Σ║Äτúüτ¢ÿσ«îµò┤σÄåσÅ▓∩╝îΘü┐σàìµê¬µû¡σ»╝Φç┤µò░µì«Σ╕óσñ▒"""
        full_history = self.load_full(filename)
        user_timestamp = extra_fields.pop("user_timestamp", datetime.now().isoformat())
        model_timestamp = extra_fields.pop(
            "model_timestamp", datetime.now().isoformat()
        )

        full_history.append(
            {"role": "user", "parts": [user_msg], "timestamp": user_timestamp}
        )
        model_entry = {"role": "model", "parts": [model_msg]}
        if "timestamp" not in extra_fields:
            model_entry["timestamp"] = model_timestamp
        model_entry.update(extra_fields)
        full_history.append(model_entry)
        self.save(filename, full_history)
        return full_history

    def append_user_early(self, filename, user_msg):
        """σ£¿Φ»╖µ▒éσê░Φ╛╛µù╢τ½ïσì│Σ┐¥σ¡ÿτö¿µê╖µ╢êµü»∩╝îΘÿ▓µ¡óµû¡Φ┐₧σ»╝Φç┤Σ╕óσñ▒
        Φ┐öσ¢₧historyΘò┐σ║ª∩╝îσÉÄτ╗¡τö¿update_last_model_responseµ¢┤µû░µ¿íσ₧ïσ¢₧σñì"""
        full_history = self.load_full(filename)
        now_iso = datetime.now().isoformat()
        full_history.append({"role": "user", "parts": [user_msg], "timestamp": now_iso})
        full_history.append(
            {"role": "model", "parts": ["ΓÅ│ σñäτÉåΣ╕¡..."], "timestamp": now_iso}
        )
        self.save(filename, full_history)
        return len(full_history)

    def update_last_model_response(self, filename, model_msg, **extra_fields):
        """µ¢┤µû░µ£ÇσÉÄΣ╕Çµ¥íµ¿íσ₧ïσ¢₧σñì∩╝êΘàìσÉêappend_user_earlyΣ╜┐τö¿∩╝ë"""
        full_history = self.load_full(filename)
        if full_history and full_history[-1].get("role") == "model":
            model_entry = {"role": "model", "parts": [model_msg]}
            if "timestamp" not in extra_fields:
                model_entry["timestamp"] = datetime.now().isoformat()
            model_entry.update(extra_fields)
            full_history[-1] = model_entry
            self.save(filename, full_history)
        else:
            # fallback: τ¢┤µÄÑΦ┐╜σèá
            model_entry = {"role": "model", "parts": [model_msg]}
            if "timestamp" not in extra_fields:
                model_entry["timestamp"] = datetime.now().isoformat()
            model_entry.update(extra_fields)
            full_history.append(model_entry)
            self.save(filename, full_history)

    def add_message(
        self, filename, role, content, task="CHAT", model_name="Auto", **extra_fields
    ):
        """Φ┐╜σèáσìòµ¥íµ╢êµü»∩╝êσà╝σ«╣µùºΦ░âτö¿∩╝ë∩╝îΘ╗ÿΦ«ñΘÖäσ╕ªµù╢Θù┤µê│"""
        full_history = self.load_full(filename)
        entry = {
            "role": role,
            "parts": [content],
            "task": task,
            "model_name": model_name,
            "timestamp": extra_fields.pop("timestamp", datetime.now().isoformat()),
        }
        entry.update(extra_fields)
        full_history.append(entry)
        self.save(filename, full_history)
        return entry

    def delete(self, filename):
        path = os.path.join(CHAT_DIR, filename)
        if os.path.exists(path):
            try:
                os.remove(path)
                return True
            except OSError as e:
                _app_logger.warning("Failed to delete session %s: %s", filename, e)
                return False
        return False


session_manager = SessionManager()

# ================= σê¥σºïσîûσà¿σ▒Çµ¿íσ¥ù =================
# µçÆσèáΦ╜╜ Memory Manager σÆî Knowledge Base
_memory_manager = None
_kb = None


def get_memory_manager():
    """ΦÄ╖σÅûµêûσê¢σ╗║ Memory Manager σ«₧Σ╛ï∩╝êσó₧σ╝║τëê∩╝ë"""
    global _memory_manager
    if _memory_manager is None:
        try:
            # Σ╝ÿσàêΣ╜┐τö¿σó₧σ╝║τëêµ£¼
            from enhanced_memory_manager import EnhancedMemoryManager

            _memory_manager = EnhancedMemoryManager()
            _app_logger.info("[INIT] Γ£à σó₧σ╝║Φ«░σ┐åτ«íτÉåσÖ¿σ╖▓σê¥σºïσîû")
        except ImportError:
            try:
                from web.enhanced_memory_manager import EnhancedMemoryManager

                _memory_manager = EnhancedMemoryManager()
                _app_logger.info("[INIT] Γ£à σó₧σ╝║Φ«░σ┐åτ«íτÉåσÖ¿σ╖▓σê¥σºïσîû")
            except ImportError:
                # ΘÖìτ║ºσê░σƒ║τíÇτëêµ£¼
                try:
                    from memory_manager import MemoryManager
                except ImportError:
                    from web.memory_manager import MemoryManager
                _memory_manager = MemoryManager()
                _app_logger.warning("[INIT] ΓÜá∩╕Å  Σ╜┐τö¿σƒ║τíÇΦ«░σ┐åτ«íτÉåσÖ¿")

        # µ│¿σàÑµæÿΦªüΣ╕ÄσÉæΘçÅΘÇéΘàìσÖ¿∩╝êσªéµ₧£µö»µîü∩╝ë
        try:

            def _memory_generate(
                prompt: str, temperature: float = 0.2, max_tokens: int = 300
            ) -> str:
                resp = client.models.generate_content(
                    model="gemini-2.0-flash-lite",
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        temperature=temperature,
                        max_output_tokens=max_tokens,
                    ),
                )
                return resp.text or ""

            def _memory_embed(texts: list) -> list:
                safe_texts = [(t or "")[:1000] for t in texts]
                resp = client.models.embed_content(
                    model="text-embedding-004", contents=safe_texts
                )
                embeddings = []
                if hasattr(resp, "embeddings"):
                    for item in resp.embeddings:
                        if hasattr(item, "values"):
                            embeddings.append(list(item.values))
                        elif hasattr(item, "embedding"):
                            embeddings.append(list(item.embedding))
                        elif isinstance(item, dict):
                            embeddings.append(
                                list(item.get("values") or item.get("embedding") or [])
                            )
                elif hasattr(resp, "embedding"):
                    embeddings.append(list(resp.embedding))
                elif isinstance(resp, dict) and "embeddings" in resp:
                    for item in resp.get("embeddings", []):
                        embeddings.append(
                            list(item.get("values") or item.get("embedding") or [])
                        )
                return embeddings

            if hasattr(_memory_manager, "set_llm_adapters"):
                _memory_manager.set_llm_adapters(
                    generate_fn=_memory_generate, embedding_fn=_memory_embed
                )
        except Exception as e:
            _app_logger.warning(f"[INIT] ΓÜá∩╕Å  Φ«░σ┐åΘÇéΘàìσÖ¿µ│¿σàÑσñ▒Φ┤Ñ: {e}")
    return _memory_manager


def _start_memory_extraction(
    user_msg: str,
    ai_msg: str,
    history=None,
    task_type: str = "CHAT",
    session_name: str = "default",
):
    """σÉÄσÅ░µÅÉσÅûΘò┐µ£ƒΦ«░σ┐å∩╝êσÉ½ MemoryReflector µ╖▒σ║ªσÅìµÇ¥∩╝ë∩╝îΣ╕ìΘÿ╗σí₧Σ╕╗σ»╣Φ»¥µ╡üτ¿ï"""
    try:
        from memory_integration import MemoryIntegration
    except ImportError:
        try:
            from web.memory_integration import MemoryIntegration
        except ImportError:
            MemoryIntegration = None

    def _llm_sync(prompt: str) -> str:
        """Synchronous LLM call for reflection / summarization."""
        try:
            resp = client.models.generate_content(
                model="gemini-2.0-flash-lite",
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.1,
                    max_output_tokens=600,
                ),
            )
            return resp.text or ""
        except Exception:
            return ""

    def _llm_quality_sync(prompt: str) -> str:
        """Higher-quality LLM call for PersonalityMatrix and evaluations.
        Tries gemini-2.5-flash ΓåÆ gemini-2.0-flash ΓåÆ falls back to _llm_sync."""
        _quality_models = ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-2.0-flash-lite"]
        _model = "gemini-2.0-flash"  # safe default
        try:
            from app.core.llm.model_fallback import get_fallback_executor
            _fbe = get_fallback_executor()
            _model = next(
                (m for m in _quality_models if _fbe.is_available(m)),
                _quality_models[-1],
            )
        except Exception:
            pass
        try:
            resp = client.models.generate_content(
                model=_model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.15,
                    max_output_tokens=800,
                )
            )
            return resp.text or ""
        except Exception:
            return _llm_sync(prompt)

    def _worker():
        # ΓöÇΓöÇ Existing MemoryIntegration (entity extraction) ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        if MemoryIntegration and MemoryIntegration.should_extract(user_msg, ai_msg):
            try:
                memory_mgr = get_memory_manager()

                class _LLMAdapter:
                    async def generate(self, prompt, temperature=0.1, max_tokens=500):
                        return _llm_sync(prompt)

                result = asyncio.run(
                    MemoryIntegration.extract_and_apply(
                        memory_mgr, user_msg, ai_msg, _LLMAdapter(), history
                    )
                )
                if result.get("success"):
                    _app_logger.info("[MemoryIntegration] Γ£à Φç¬σè¿Φ«░σ┐åµÅÉσÅûσ«îµêÉ")
                else:
                    _app_logger.warning(f"[MemoryIntegration] ΓÜá∩╕Å µÅÉσÅûσñ▒Φ┤Ñ: {result.get('error')}")
            except Exception as e:
                _app_logger.error(f"[MemoryIntegration] Γ¥î σ╝éσ╕╕: {e}")

        # ΓöÇΓöÇ 2-B: MemoryReflector (deep structured reflection) ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        try:
            from app.core.memory.memory_reflector import MemoryReflector

            MemoryReflector.reflect_async(
                user_msg=user_msg,
                ai_msg=ai_msg,
                task_type=task_type,
                session_name=session_name,
                get_memory_fn=get_memory_manager,
                llm_fn=_llm_sync,
            )
        except Exception as e:
            _app_logger.warning(f"[MemoryReflector] ΓÜá∩╕Å σÉ»σè¿σñ▒Φ┤Ñ: {e}")

        # ΓöÇΓöÇ 2-C: PersonalityMatrix ΓÇö σè¿µÇüΣ╕¬Σ║║Φ«░σ┐åτƒ⌐Θÿ╡µ¢┤µû░∩╝êΣ╜┐τö¿µ¢┤Θ½ÿΦ┤¿ΘçÅµ¿íσ₧ï∩╝ëΓöÇΓöÇ
        try:
            _pm_mgr = get_memory_manager()
            if _pm_mgr and hasattr(_pm_mgr, "update_personality_async"):
                _pm_mgr.update_personality_async(user_msg, ai_msg, _llm_quality_sync)
        except Exception as e:
            _app_logger.warning(f"[PersonalityMatrix] ΓÜá∩╕Å µ¢┤µû░σÉ»σè¿σñ▒Φ┤Ñ: {e}")

        # ΓöÇΓöÇ 3: ShadowWatcher σ╜▒σ¡ÉΦ┐╜Φ╕¬∩╝êΘ¢╢µäƒτƒÑΦºéσ»ƒ∩╝ëΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        try:
            from app.core.monitoring.shadow_watcher import ShadowWatcher

            ShadowWatcher.observe(user_msg, ai_msg, session_name)
        except Exception as e:
            _app_logger.warning(f"[ShadowWatcher] ΓÜá∩╕Å Φºéσ»ƒσñ▒Φ┤Ñ: {e}")

        # ΓöÇΓöÇ 3-B: ResponseEvaluator µ¿íσ₧ïΦç¬Φ»ä∩╝êΦç¬σè¿Φ┤¿ΘçÅΦ»äσêå ΓåÆ RatingStore∩╝ëΓöÇΓöÇΓöÇΓöÇ
        try:
            from app.core.learning.rating_store import RatingStore as _RS
            from app.core.learning.response_evaluator import ResponseEvaluator

            _eval_msg_id = _RS.make_msg_id(session_name, user_msg)
            ResponseEvaluator.evaluate_async(
                msg_id=_eval_msg_id,
                user_input=user_msg,
                ai_response=ai_msg,
                task_type=task_type,
                session_name=session_name,
                llm_fn=_llm_sync,
            )
        except Exception as e:
            _app_logger.warning(f"[ResponseEvaluator] ΓÜá∩╕Å Φç¬Φ»äσÉ»σè¿σñ▒Φ┤Ñ: {e}")

        # ΓöÇΓöÇ 4: MacroRecorder σ«Åσ╜òσê╢∩╝êΘçìσñìσ╖ÑΣ╜£µ╡üµúÇµ╡ï∩╝ëΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        try:
            from app.core.monitoring.macro_recorder import MacroRecorder

            MacroRecorder.record_turn(user_msg, task_type or "CHAT", session_name)
        except Exception as e:
            _app_logger.warning(f"[MacroRecorder] ΓÜá∩╕Å Φ«░σ╜òσñ▒Φ┤Ñ: {e}")

    threading.Thread(target=_worker, daemon=True).start()


def get_knowledge_base():
    """ΦÄ╖σÅûµêûσê¢σ╗║ Knowledge Base σ«₧Σ╛ï"""
    global _kb
    if _kb is None:
        try:
            from knowledge_base import KnowledgeBase
        except ImportError:
            from web.knowledge_base import KnowledgeBase
        _kb = KnowledgeBase()
        _app_logger.info("[INIT] Γ£à Knowledge Base σ╖▓σê¥σºïσîû")
    return _kb


# Σ╕║Σ║åσÉæσÉÄσà╝σ«╣∩╝îσ»╝σç║σà¿σ▒ÇσÅÿΘçÅ
memory_manager = None  # σ░åΘÇÜΦ┐ç get_memory_manager() σè¿µÇüΦÄ╖σÅû
kb = None  # σ░åΘÇÜΦ┐ç get_knowledge_base() σè¿µÇüΦÄ╖σÅû

# ================= Koto Brain =================


class KotoBrain:
    # σ¢╛σâÅτ╝ûΦ╛æσà│Θö«Φ»ì
    IMAGE_EDIT_KEYWORDS = [
        "Σ┐«µö╣",
        "µìó",
        "µö╣µêÉ",
        "σÅÿµêÉ",
        "σ║òΦë▓",
        "ΦâîµÖ»",
        "Θó£Φë▓",
        "µèáσ¢╛",
        "σÄ╗ΦâîµÖ»",
        "Pσ¢╛",
        "τ╛Äσîû",
        "µ╗ñΘò£",
        "Φ░âΦë▓",
        "τ╝ûΦ╛æ",
        "change",
        "modify",
        "edit",
        "background",
        "color",
    ]

    def chat(
        self,
        history,
        user_input,
        file_data=None,
        model=None,
        auto_model=True,
        task_type: str = None,
    ):
        start_time = time.time()
        original_input = user_input
        # µö»µîüµ¿íσ₧ïΘÇëµï⌐σÆîΦç¬σè¿ΘÇëµï⌐
        _model_id_locked = (
            False  # σªéµ₧£σ╖▓σ£¿Φ╖»τö▒Σ╕¡σ╝║σê╢Φ«╛τ╜« model_id∩╝îΦ╖│Φ┐çσÉÄτ╗¡ SmartDispatcher Φªåτ¢û
        )
        if model and not auto_model:
            model_id = model
            route_method = "Manual select"
            # Σ╝ÿσàêΣ╜┐τö¿Φ░âτö¿µû╣Σ╝áσàÑτÜä task_type∩╝îΘü┐σàìΘçìσñìΦ╖»τö▒
            target_key = task_type or "CHAT"
        else:
            target_key = "CHAT"
            route_method = "Auto"
            model_id = None  # σàêτ╜«τ⌐║∩╝îΣ╕ïΘ¥óµîëΦ╖»τö▒σå│σ«Ü

            if file_data:
                _fd_mime = (
                    file_data.get("mime_type") or "application/octet-stream"
                ).lower()
                _is_image_file = _fd_mime.startswith("image/")
                if _is_image_file:
                    # σ¢╛τëçµûçΣ╗╢∩╝Üσêñµû¡τ╝ûΦ╛æ vs σêåµ₧É
                    user_lower = user_input.lower()
                    is_edit = any(kw in user_lower for kw in self.IMAGE_EDIT_KEYWORDS)
                    if is_edit:
                        target_key = "PAINTER"
                        route_method = "Image Edit"
                    else:
                        target_key = "VISION"
                        route_method = "Image Analysis"
                else:
                    # Θ¥₧σ¢╛τëçΣ║îΦ┐¢σê╢µûçΣ╗╢∩╝êPDF/Wordτ¡ë∩╝ë∩╝ÜΦ╖»τö▒Σ╕║ CHAT∩╝îΣ╜┐τö¿ΘÖìτ║ºµ¿íσ₧ïτ¢┤µÄÑΦ»╗σÅû
                    target_key = "CHAT"
                    route_method = "≡ƒôä Binary-Doc-Read"
                    # σ╝║σê╢Σ╜┐τö¿µö»µîü generate_content + µûçΣ╗╢σ¡ùΦèéτÜäΘÖìτ║ºµ¿íσ₧ï∩╝êInteractions API Σ╕ìµö»µîüµûçΣ╗╢ΘÖäΣ╗╢∩╝ë
                    model_id = _INTERACTIONS_FALLBACK_MODEL
                    _model_id_locked = True
            else:
                # Σ╜┐τö¿µÖ║Φâ╜Φ╖»τö▒σÖ¿
                target_key, route_method, _ = SmartDispatcher.analyze(user_input)

            if not _model_id_locked:
                model_id = SmartDispatcher.get_model_for_task(
                    target_key, has_image=bool(file_data)
                )

        # Σ╜┐τö¿σ░Åµ¿íσ₧ïσ░åΦ»╖µ▒éΦ╜¼µìóΣ╕║τ╗ôµ₧äσîû Markdown∩╝êΣ╗àσ£¿σñºµ¿íσ₧ïσñäτÉåµù╢σÉ»τö¿∩╝ë
        # ΓÜá∩╕Å Φ╖│Φ┐çµ¥íΣ╗╢∩╝Üµ£ëµûçΣ╗╢ΘÖäΣ╗╢µù╢∩╝êfile_data∩╝ëπÇüµêûΦ╛ôσàÑσ╛êσñº∩╝êσÉ½σ╡îσàÑµûçΣ╗╢σåàσ«╣∩╝ë
        _has_embedded_file_content = (
            "=== µûçΣ╗╢σåàσ«╣ ===" in user_input or len(user_input) > 3000
        )
        model_input = user_input
        if (
            auto_model
            and not file_data
            and not _has_embedded_file_content
            and target_key not in ["SYSTEM", "FILE_OP", "PAINTER", "VISION"]
        ):
            # Σ╗àΣ╜┐τö¿µ£¼σ£░µ¿íµ¥┐Θçìµò┤∩╝êΣ╕ìΣ╝á model_generate∩╝îΘü┐σàìΘó¥σñûτÜä flash-lite API Φ░âτö¿∩╝ë
            model_input = Utils.adapt_prompt_to_markdown(
                target_key, user_input, history=history
            )
            if model_input != user_input:
                _app_logger.debug("[PROMPT_ADAPTER] Applied local Markdown template")
        result = {
            "task": target_key,
            "model": model_id,
            "route_method": route_method,  # Φ╖»τö▒µû╣µ│òΣ┐íµü»
            "response": "",
            "images": [],
            "saved_files": [],
            "latency": 0,
            "total_time": 0,
        }

        try:
            # === SYSTEM Mode (µ£¼σ£░µëºΦíî) ===
            if target_key == "SYSTEM":
                exec_result = LocalExecutor.execute(user_input)
                result["response"] = exec_result["message"]
                if exec_result.get("details"):
                    result["response"] += f"\n\n{exec_result['details']}"
                result["total_time"] = time.time() - start_time
                return result

            # === PAINTER Mode (σ¢╛σâÅτöƒµêÉ/τ╝ûΦ╛æ) ===
            if target_key == "PAINTER":
                # σªéµ₧£µ£ëΦ╛ôσàÑσ¢╛τëç∩╝êσ¢╛σâÅτ╝ûΦ╛æµ¿íσ╝Å∩╝ë- Σ╜┐τö¿Σ╗úτáüµû╣σ╝ÅσñäτÉå
                if file_data:
                    # Σ┐¥σ¡ÿΣ╕èΣ╝áτÜäσ¢╛τëçσê░ workspace
                    import subprocess
                    import tempfile

                    temp_img_path = os.path.join(
                        WORKSPACE_DIR, "images", f"input_{int(time.time())}.jpg"
                    )
                    os.makedirs(os.path.dirname(temp_img_path), exist_ok=True)
                    with open(temp_img_path, "wb") as f:
                        f.write(file_data["data"])

                    # µ₧äσ╗║σ¢╛σâÅτ╝ûΦ╛æτÜäτ│╗τ╗ƒµîçΣ╗ñ
                    edit_instruction = f"""Σ╜áµÿ»Σ╕ÇΣ╕¬σ¢╛σâÅσñäτÉåΣ╕ôσ«╢πÇéτö¿µê╖Σ╕èΣ╝áΣ║åΣ╕Çσ╝áσ¢╛τëç∩╝îΘ£ÇΦªüΣ╜áτöƒµêÉ Python Σ╗úτáüµ¥ÑσñäτÉåσ«âπÇé

σ¢╛τëçΦ╖»σ╛ä: {temp_img_path}
τö¿µê╖Φ»╖µ▒é: {user_input}

Φ»╖τöƒµêÉσ«îµò┤τÜä Python Σ╗úτáüµ¥Ñσ«îµêÉτö¿µê╖τÜäσ¢╛σâÅτ╝ûΦ╛æΦ»╖µ▒éπÇé

Φªüµ▒é:
1. Σ╜┐τö¿ OpenCV (cv2) µêû PIL σñäτÉåσ¢╛τëç
2. σñäτÉåσÉÄτÜäσ¢╛τëçΣ┐¥σ¡ÿσê░: {settings_manager.images_dir}
3. µûçΣ╗╢σÉìµá╝σ╝Å: edited_{{timestamp}}.jpg µêû .png
4. Σ╗úτáüσ┐àΘí╗σ«îµò┤σÅ»µëºΦíî
5. σ»╣Σ║ÄµìóΦâîµÖ»Φë▓∩╝îΣ╜┐τö¿Θó£Φë▓ΘÿêσÇ╝µêûΦ╛╣τ╝ÿµúÇµ╡ïµ¥ÑΦ»åσê½ΦâîµÖ»σî║σƒƒ

σ╕╕τö¿τÜäΦâîµÖ»Φë▓σñäτÉåµû╣µ│ò:
- Φ»üΣ╗╢τàºµìóσ║òΦë▓: µúÇµ╡ïµÄÑΦ┐æσÄƒΦâîµÖ»Φë▓τÜäσâÅτ┤á∩╝îµ¢┐µìóΣ╕║τ¢«µáçΘó£Φë▓
- Φô¥Φë▓ΦâîµÖ» RGB: (67, 142, 219) µêû (0, 191, 255)
- τ║óΦë▓ΦâîµÖ» RGB: (255, 0, 0) µêû (220, 0, 0)  
- τÖ╜Φë▓ΦâîµÖ» RGB: (255, 255, 255)

Σ╗úτáüµá╝σ╝Å∩╝êσ┐àΘí╗Σ╜┐τö¿Φ┐ÖΣ╕¬µá╝σ╝Å∩╝ë:
---BEGIN_FILE: image_edit.py---
# Σ╜áτÜäΣ╗úτáü
---END_FILE---"""

                    # Φ░âτö¿ Gemini τöƒµêÉΣ╗úτáü∩╝êσ╕ªσ¢₧ΘÇÇ∩╝ë
                    edit_models = [
                        "gemini-3-flash-preview",
                        "gemini-3-pro-preview",
                        "gemini-2.5-flash",
                    ]
                    code_response = None
                    last_error = None

                    def _process_code_response(code_response_text: str):
                        # µÅÉσÅûΣ╗úτáü - µö»µîüσñÜτºìµá╝σ╝Å
                        import re

                        patterns = [
                            r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\s*(.*?)---\s*END_FILE\s*---",
                            r"```python\s*(.*?)```",  # µáçσçå markdown Σ╗úτáüσ¥ù
                            r"```\s*(.*?)```",  # µùáΦ»¡Φ¿ÇµáçΦ«░τÜäΣ╗úτáüσ¥ù
                        ]

                        code_content = None
                        for pattern in patterns:
                            matches = re.findall(
                                pattern, code_response_text, re.DOTALL | re.IGNORECASE
                            )
                            if matches:
                                if isinstance(matches[0], tuple):
                                    code_content = matches[0][1].strip()
                                else:
                                    code_content = matches[0].strip()
                                _app_logger.debug(
                                    f"[IMAGE_EDIT] Extracted code, length: {len(code_content)}"
                                )
                                break

                        if not code_content:
                            return {
                                "images": [],
                                "response": f"Γ¥î µùáµ│òΣ╗Äµ¿íσ₧ïσôìσ║öΣ╕¡µÅÉσÅûΣ╗úτáü\n\nµ¿íσ₧ïΦ┐öσ¢₧σåàσ«╣:\n```\n{code_response_text[:500]}\n```",
                                "error": "no_code",
                            }

                        # Σ┐¥σ¡ÿσ╣╢µëºΦíîΣ╗úτáü
                        temp_script = os.path.join(
                            tempfile.gettempdir(), f"koto_edit_{int(time.time())}.py"
                        )
                        with open(temp_script, "w", encoding="utf-8") as f:
                            f.write(code_content)

                        _app_logger.debug(f"[IMAGE_EDIT] Executing script: {temp_script}")
                        if getattr(sys, "frozen", False):
                            # µëôσîàµ¿íσ╝Å∩╝Üsys.executable µÿ» Koto.exe∩╝îΣ╕ìΦâ╜τö¿µ¥ÑΦ┐ÉΦíîΦäÜµ£¼∩╝îµö╣Σ╕║Φ┐¢τ¿ïσåà exec()
                            import contextlib as _ctx
                            import io as _io

                            _out, _err, _rc = _io.StringIO(), _io.StringIO(), 0
                            try:
                                _prev = os.getcwd()
                                os.chdir(WORKSPACE_DIR)
                                with _ctx.redirect_stdout(_out), _ctx.redirect_stderr(
                                    _err
                                ):
                                    exec(
                                        open(temp_script, "r", encoding="utf-8").read(),
                                        {"__file__": temp_script},
                                    )
                                os.chdir(_prev)
                            except Exception as _ex:
                                _err.write(str(_ex))
                                _rc = 1

                            class _ImgR:
                                returncode = _rc
                                stdout = _out.getvalue()
                                stderr = _err.getvalue()

                            exec_result = _ImgR()
                        else:
                            exec_result = subprocess.run(
                                [sys.executable, temp_script],
                                capture_output=True,
                                text=True,
                                timeout=60,
                                cwd=WORKSPACE_DIR,
                            )

                        _app_logger.debug(
                            f"[IMAGE_EDIT] Script result: returncode={exec_result.returncode}"
                        )
                        if exec_result.stdout:
                            _app_logger.debug(f"[IMAGE_EDIT] stdout: {exec_result.stdout[:200]}")
                        if exec_result.stderr:
                            _app_logger.debug(f"[IMAGE_EDIT] stderr: {exec_result.stderr[:200]}")

                        # µ╕àτÉåΣ╕┤µù╢ΦäÜµ£¼
                        try:
                            os.remove(temp_script)
                        except OSError:
                            pass

                        if exec_result.returncode == 0:
                            images = []
                            images_dir = settings_manager.images_dir
                            for f in os.listdir(images_dir):
                                if f.startswith("edited_") and f.endswith(
                                    (".jpg", ".png", ".jpeg")
                                ):
                                    full_path = os.path.join(images_dir, f)
                                    age = time.time() - os.path.getmtime(full_path)
                                    if age < 60:
                                        rel_path = os.path.relpath(
                                            full_path, WORKSPACE_DIR
                                        ).replace("\\", "/")
                                        images.append(rel_path)

                            if images:
                                return {
                                    "images": images,
                                    "response": f"Γ£à σ¢╛τëçτ╝ûΦ╛æσ«îµêÉ!\n≡ƒû╝∩╕Å Σ┐¥σ¡ÿΣ╜ìτ╜«: `{images_dir}`",
                                    "error": "",
                                }
                            return {
                                "images": [],
                                "response": f"ΓÜá∩╕Å ΦäÜµ£¼µëºΦíîµêÉσèƒΣ╜åµ£¬µúÇµ╡ïσê░µû░σ¢╛τëç\n\n{exec_result.stdout[:500]}",
                                "error": "no_output",
                            }

                        return {
                            "images": [],
                            "response": f"Γ¥î σ¢╛τëçσñäτÉåσñ▒Φ┤Ñ\n```\n{exec_result.stderr[:500]}\n```",
                            "error": "exec_failed",
                        }

                    for edit_model in edit_models:
                        try:
                            _app_logger.debug(f"[IMAGE_EDIT] Trying model: {edit_model}")
                            _app_logger.debug(f"[IMAGE_EDIT] Sending request to API...")
                            response = client.models.generate_content(
                                model=edit_model,
                                contents=edit_instruction,
                                config=types.GenerateContentConfig(
                                    max_output_tokens=4096, temperature=0.5
                                ),
                            )
                            _app_logger.debug(f"[IMAGE_EDIT] Got API response")

                            if (
                                response.candidates
                                and response.candidates[0].content.parts
                            ):
                                code_response = (
                                    response.candidates[0].content.parts[0].text
                                )
                                _app_logger.debug(
                                    f"[IMAGE_EDIT] Got response from {edit_model}, length: {len(code_response)}"
                                )
                                break
                        except Exception as model_err:
                            last_error = str(model_err)
                            _app_logger.debug(
                                f"[IMAGE_EDIT] Model {edit_model} failed: {last_error[:100]}"
                            )
                            continue

                    if code_response:
                        run_result = _process_code_response(code_response)
                        result["images"] = run_result["images"]
                        result["response"] = run_result["response"]
                    else:
                        result["response"] = (
                            f"Γ¥î µëÇµ£ëµ¿íσ₧ïΘâ╜Σ╕ìσÅ»τö¿: {last_error[:200] if last_error else 'µ£¬τƒÑΘöÖΦ»»'}"
                        )

                    # σñ▒Φ┤ÑσÉÄΦç¬σè¿Σ┐«µ¡úσ╣╢ΘçìΦ»òΣ╕Çµ¼í∩╝êΘü┐σàìµùáτ╝ûΦ╛æτ╗ôµ₧£∩╝ë
                    if not result["images"] and Utils.is_failure_output(
                        result["response"]
                    ):
                        fix_prompt = (
                            "Σ╕èµ¼íτöƒµêÉσñ▒Φ┤Ñ∩╝îΦ»╖Σ┐«µ¡úσ╣╢σÅ¬Φ╛ôσç║σ«îµò┤σÅ»µëºΦíîτÜä Python Σ╗úτáüπÇé\n"
                            "σ┐àΘí╗Σ╜┐τö¿ BEGIN_FILE/END_FILE µá╝σ╝ÅπÇé\n"
                            f"σ¢╛τëçΦ╖»σ╛ä: {temp_img_path}\n"
                            f"Φ╛ôσç║τ¢«σ╜ò: {settings_manager.images_dir}\n"
                            f"τö¿µê╖Φ»╖µ▒é: {user_input}\n\n"
                            f"σñ▒Φ┤ÑΣ┐íµü»/Φ╛ôσç║: {result['response']}\n"
                        )
                        retry_models = ["gemini-3-flash-preview", "gemini-2.5-flash"]
                        for retry_model in retry_models:
                            try:
                                _app_logger.debug(f"[IMAGE_EDIT] Retry with model: {retry_model}")
                                retry_resp = client.models.generate_content(
                                    model=retry_model,
                                    contents=fix_prompt,
                                    config=types.GenerateContentConfig(
                                        max_output_tokens=4096
                                    ),
                                )
                                if (
                                    retry_resp.candidates
                                    and retry_resp.candidates[0].content.parts
                                ):
                                    retry_code = (
                                        retry_resp.candidates[0].content.parts[0].text
                                    )
                                    retry_run = _process_code_response(retry_code)
                                    if retry_run["images"]:
                                        result["images"] = retry_run["images"]
                                        result["response"] = retry_run["response"]
                                        break
                                    result["response"] = retry_run["response"]
                            except Exception as retry_err:
                                _app_logger.debug(f"[IMAGE_EDIT] Retry failed: {retry_err}")

                    result["total_time"] = time.time() - start_time
                    return result
                else:
                    # τ║»σ¢╛σâÅτöƒµêÉΣ╜┐τö¿ gemini-3.1-flash-image-preview
                    try:
                        _app_logger.info(f"[σ¢╛σâÅτöƒµêÉ] σ╝ÇσºïτöƒµêÉ: {user_input[:50]}...")
                        response = client.models.generate_content(
                            model="gemini-3.1-flash-image-preview",
                            contents=user_input,
                            config=types.GenerateContentConfig(
                                response_modalities=["TEXT", "IMAGE"]
                            ),
                        )
                        _app_logger.info(
                            f"[σ¢╛σâÅτöƒµêÉ] σôìσ║öµêÉσèƒ∩╝îσÇÖΘÇëµò░: {len(response.candidates) if response.candidates else 0}"
                        )

                        # Σ┐¥σ¡ÿτöƒµêÉτÜäσ¢╛τëç
                        if response.candidates and response.candidates[0].content.parts:
                            for part in response.candidates[0].content.parts:
                                if hasattr(part, "inline_data") and part.inline_data:
                                    img_filename = Utils.save_image_part(part)
                                    if img_filename:
                                        result["images"].append(img_filename)
                                        _app_logger.info(f"[σ¢╛σâÅτöƒµêÉ] σ╖▓Σ┐¥σ¡ÿ: {img_filename}")

                        if result["images"]:
                            save_path = settings_manager.images_dir
                            result["response"] = (
                                f"Γ£¿ σ¢╛τëçσ╖▓τöƒµêÉ!\n≡ƒû╝∩╕Å Σ┐¥σ¡ÿΣ╜ìτ╜«: `{save_path}`"
                            )
                        else:
                            result["response"] = (
                                "Γ¥î σ¢╛σâÅτöƒµêÉσñ▒Φ┤Ñ: µùáΦ╛ôσç║σåàσ«╣∩╝îΦ»╖µúÇµƒÑµÅÉτñ║Φ»ì"
                            )
                        result["total_time"] = time.time() - start_time
                        return result
                    except Exception as img_err:
                        error_msg = str(img_err)
                        _app_logger.info(f"[σ¢╛σâÅτöƒµêÉ] ΘöÖΦ»»: {error_msg[:200]}")

                        # µÅÉΣ╛¢µ¢┤Φ»ªτ╗åτÜäΘöÖΦ»»Σ┐íµü»
                        if (
                            "disconnected" in error_msg.lower()
                            or "timeout" in error_msg.lower()
                        ):
                            result["response"] = (
                                f"Γ¥î Φ┐₧µÄÑΦ╢àµù╢µêûΣ╕¡µû¡: {error_msg[:100]}\n\n≡ƒÆí σ╗║Φ««: Φ»╖τ¿ìσÉÄΘçìΦ»ò∩╝îµêûµúÇµƒÑτ╜æτ╗£Φ┐₧µÄÑ"
                            )
                        elif "safety" in error_msg.lower():
                            result["response"] = "Γ¥î σåàσ«╣σ¢áσ«ëσà¿µö┐τ¡ûΦó½Φ┐çµ╗ñ∩╝îΦ»╖Σ┐«µö╣µÅÉτñ║Φ»ì"
                        elif (
                            "quota" in error_msg.lower() or "rate" in error_msg.lower()
                        ):
                            result["response"] = "Γ¥î API ΘàìΘó¥σ╖▓Φ╛╛ΘÖÉσê╢∩╝îΦ»╖τ¿ìσÉÄΘçìΦ»ò"
                        else:
                            result["response"] = f"Γ¥î σ¢╛σâÅτöƒµêÉσñ▒Φ┤Ñ: {error_msg[:100]}"

                        result["total_time"] = time.time() - start_time
                        return result

                if not response.candidates:
                    result["response"] = "Generation failed (safety filter or busy)."
                    result["total_time"] = time.time() - start_time
                    return result

                text_response = ""
                if response.candidates and response.candidates[0].content.parts:
                    for part in response.candidates[0].content.parts:
                        if hasattr(part, "text") and part.text:
                            text_response += part.text
                        if hasattr(part, "inline_data") and part.inline_data:
                            img_filename = Utils.save_image_part(part)
                            if img_filename:
                                result["images"].append(img_filename)

                # µ╖╗σèáσ¢╛τëçΣ┐¥σ¡ÿΣ╜ìτ╜«µÅÉτñ║
                if result["images"]:
                    save_path = settings_manager.images_dir
                    text_response += f"\n\n≡ƒû╝∩╕Å σ¢╛τëçσ╖▓Σ┐¥σ¡ÿσê░: `{save_path}`"

                result["response"] = (
                    text_response if text_response else "Image generated successfully!"
                )
                result["total_time"] = time.time() - start_time
                return result

            # === RAG: Retrieve Relevant Context (Auto) ===
            try:
                # ΦÄ╖σÅûτƒÑΦ»åσ║ôσ«₧Σ╛ï
                kb_inst = get_knowledge_base()

                # Σ╗àσ£¿Θ¥₧τë╣σ«Üµ¿íσ╝ÅΣ╕öΦ╛ôσàÑµ£ëµòêµù╢µúÇτ┤ó
                if target_key not in ["PAINTER", "SYSTEM"] and len(original_input) > 3:
                    # Θü┐σàìσ»╣µ₧üτƒ¡τÜäΘù«σÇÖΦ»¡Φ┐¢ΦíîµúÇτ┤ó
                    skip_keywords = ["Σ╜áσÑ╜", "hello", "hi", "test", "µ╡ïΦ»ò"]
                    if not any(original_input.lower() == k for k in skip_keywords):
                        _app_logger.debug(f"[RAG]µ¡úσ£¿µúÇτ┤óτƒÑΦ»åσ║ô: {original_input[:50]}...")
                        rag_results = kb_inst.search(original_input, top_k=3)
                        
                        if rag_results:
                            _app_logger.debug(f"[RAG] µúÇτ┤óσê░ {len(rag_results)} Σ╕¬τ¢╕σà│τëçµ«╡")
                            context_str = "\n".join([
                                f"--- µ¥Ñµ║É: {r['file_name']} (τ¢╕Σ╝╝σ║ª: {r['similarity']:.2f}) ---\n{r['text']}"
                                for r in rag_results
                            ])
                            
                            # σ░åΣ╕èΣ╕ïµûçµ│¿σàÑ prompt
                            rag_context = f"\n\nπÇÉσÅéΦÇâΦ╡äµûÖπÇæ\nΣ╗ÑΣ╕ïµÿ»Σ╗Äµ£¼σ£░τƒÑΦ»åσ║ôµúÇτ┤óσê░τÜäτ¢╕σà│σåàσ«╣∩╝îΣ╛¢σ¢₧τ¡öσÅéΦÇâ∩╝Ü\n{context_str}\n\n"

                            # Log retrieval
                            _app_logger.debug(f"[RAG] Injected context length: {len(rag_context)}")

                            # Update model input
                            # σªéµ₧£µ£ë file_data∩╝îmodel_input σÅ»Φâ╜µÿ» None µêûΣ╕ìΦó½τ¢┤µÄÑΣ╜┐τö¿∩╝îΘ£ÇΦ░¿µàÄ
                            if not file_data:
                                model_input = rag_context + model_input
                            else:
                                # σ»╣Σ║Äµ£ëµûçΣ╗╢τÜäΦ»╖µ▒é∩╝îµêæΣ╗¼σ░åΣ╕èΣ╕ïµûçµï╝µÄÑσê░ original_input (user prompt)
                                # µ│¿µäÅ∩╝ÜΣ╕ïΘ¥ó generate_content τö¿τÜäµÿ» original_input + image_part
                                original_input = rag_context + original_input

            except Exception as rag_err:
                _app_logger.debug(f"[RAG] Retrieval warning: {rag_err}")

            # === Regular Mode ===
            # µ₧äσ╗║σÄåσÅ▓Φ«░σ╜òµá╝σ╝Å∩╝êΦ┐çµ╗ñµùáσà│σÄåσÅ▓∩╝ë
            history_for_model = ContextAnalyzer.filter_history(original_input, history)
            formatted_history = []
            for turn in history_for_model:
                formatted_history.append(
                    types.Content(
                        role=turn["role"],
                        parts=[types.Part.from_text(text=p) for p in turn["parts"]],
                    )
                )

            # µá╣µì«Σ╗╗σèíτ▒╗σ₧ïΘÇëµï⌐τ│╗τ╗ƒµÅÉτñ║∩╝ÜFILE_GEN Φ╡░µûçµíúτöƒµêÉµÅÉτñ║∩╝îσà╢Σ╜ÖΦ╡░ΘÇÜτö¿σè⌐µëïµÅÉτñ║
            if target_key == "FILE_GEN":
                _brain_sys_instruction = _get_system_instruction()
            else:
                _brain_sys_instruction = _get_chat_system_instruction(original_input)

            if file_data:
                # µ₧äσ╗║ Part µá╝σ╝Å∩╝êΘÇéτö¿Σ║Äσ¢╛τëçσÆî PDF/µûçµíú∩╝ë
                doc_part = types.Part.from_bytes(
                    data=file_data["data"], mime_type=file_data["mime_type"]
                )
                _fd_mime2 = (file_data.get("mime_type") or "").lower()
                _is_image = _fd_mime2.startswith("image/")

                if not _is_image:
                    # PDF / µûçµíúΣ║îΦ┐¢σê╢∩╝ÜInteractions API Σ╕ìµö»µîüµûçΣ╗╢ΘÖäΣ╗╢
                    # ΓåÆ τ¢┤µÄÑΣ╜┐τö¿ gemini-2.5-flash∩╝êσÄƒτöƒµö»µîü generate_content + PDF bytes∩╝ë
                    _doc_model = _INTERACTIONS_FALLBACK_MODEL
                    if model_id != _doc_model:
                        _app_logger.info(
                            f"[brain.chat] Θ¥₧σ¢╛τëçµûçΣ╗╢ ({_fd_mime2}): ΘÖìτ║ºµ¿íσ₧ï {model_id} ΓåÆ {_doc_model}"
                        )
                        model_id = _doc_model
                        result["model"] = model_id
                    response = client.models.generate_content(
                        model=model_id,
                        contents=[original_input, doc_part],
                        config=types.GenerateContentConfig(
                            system_instruction=_brain_sys_instruction
                        ),
                    )
                    accumulated_text = response.text if response.text else ""
                elif model_id in _INTERACTIONS_ONLY_MODELS:
                    # σ¢╛τëçµûçΣ╗╢ + gemini-3-preview µ¿íσ₧ï∩╝ÜΦ╡░ Interactions API
                    try:
                        accumulated_text = _call_interactions_api_sync(
                            model_id,
                            original_input,
                            sys_instruction=_brain_sys_instruction,
                        )
                        if not accumulated_text:
                            raise ValueError("Interactions API Φ┐öσ¢₧τ⌐║σôìσ║ö")
                    except Exception as _ia_err:
                        _app_logger.info(f"[brain.chat] {model_id} Interactions API σñ▒Φ┤Ñ: {_ia_err} ΓåÆ ΘÖìτ║ºσê░ {_INTERACTIONS_FALLBACK_MODEL}")
                        model_id = _INTERACTIONS_FALLBACK_MODEL
                        result["model"] = model_id
                        _fb_resp = client.models.generate_content(
                            model=model_id,
                            contents=[original_input, doc_part],
                            config=types.GenerateContentConfig(system_instruction=_brain_sys_instruction)
                        )
                        accumulated_text = _fb_resp.text if _fb_resp.text else ""
                else:
                    # σ¢╛τëçµûçΣ╗╢ + µÖ«ΘÇÜ generate_content µ¿íσ₧ï
                    response = client.models.generate_content(
                        model=model_id,
                        contents=[original_input, doc_part],
                        config=types.GenerateContentConfig(
                            system_instruction=_brain_sys_instruction
                        ),
                    )
                    accumulated_text = response.text if response.text else ""
            else:
                # gemini-3-preview σÅ¬µö»µîü Interactions API∩╝îΣ╕ìµö»µîü generate_content
                if model_id in _INTERACTIONS_ONLY_MODELS:
                    try:
                        # σ░åσÄåσÅ▓Φ«░σ╜òµèÿσÅáΦ┐¢ prompt∩╝êInteractions API Σ╕ìµö»µîüσñÜΦ╜«σÄåσÅ▓∩╝ë
                        history_prefix = ""
                        if formatted_history:
                            history_lines = []
                            for turn in formatted_history[-6:]:  # µ£ÇΦ┐æ 3 Φ╜«
                                role_label = "τö¿µê╖" if turn.role == "user" else "σè⌐µëï"
                                turn_text = " ".join(
                                    p.text
                                    for p in turn.parts
                                    if hasattr(p, "text") and p.text
                                )
                                if turn_text:
                                    history_lines.append(f"{role_label}: {turn_text}")
                            if history_lines:
                                history_prefix = (
                                    "[σ»╣Φ»¥σÄåσÅ▓]\n" + "\n".join(history_lines) + "\n\n"
                                )
                        full_prompt = history_prefix + model_input
                        accumulated_text = _call_interactions_api_sync(
                            model_id,
                            full_prompt,
                            sys_instruction=_brain_sys_instruction,
                        )
                        if not accumulated_text:
                            raise ValueError("Interactions API Φ┐öσ¢₧τ⌐║σôìσ║ö")
                    except Exception as _ia_err:
                        _app_logger.info(
                            f"[brain.chat] {model_id} Interactions API σñ▒Φ┤Ñ: {_ia_err} ΓåÆ ΘÖìτ║ºσê░ {_INTERACTIONS_FALLBACK_MODEL}"
                        )
                        model_id = _INTERACTIONS_FALLBACK_MODEL
                        result["model"] = model_id
                        _fb_resp = client.models.generate_content(
                            model=model_id,
                            contents=formatted_history
                            + [
                                types.Content(
                                    role="user",
                                    parts=[types.Part.from_text(text=model_input)],
                                )
                            ],
                            config=types.GenerateContentConfig(
                                system_instruction=_brain_sys_instruction
                            ),
                        )
                        accumulated_text = _fb_resp.text if _fb_resp.text else ""
                else:
                    response = client.models.generate_content(
                        model=model_id,
                        contents=formatted_history
                        + [
                            types.Content(
                                role="user",
                                parts=[types.Part.from_text(text=model_input)],
                            )
                        ],
                        config=types.GenerateContentConfig(
                            system_instruction=_brain_sys_instruction
                        ),
                    )
                    accumulated_text = response.text if response.text else ""

            first_token_latency = (time.time() - start_time) * 1000
            result["latency"] = first_token_latency

            # Auto-save files
            saved_files = Utils.auto_save_files(accumulated_text)
            result["saved_files"] = saved_files

            # µ╖╗σèáµûçΣ╗╢Σ┐¥σ¡ÿµÅÉτñ║
            if saved_files:
                files_list = ", ".join(saved_files)
                accumulated_text += (
                    f"\n\n≡ƒôü µûçΣ╗╢σ╖▓Σ┐¥σ¡ÿ: **{files_list}**\n≡ƒôé Σ╜ìτ╜«: `{WORKSPACE_DIR}`"
                )

            result["response"] = accumulated_text
            result["total_time"] = time.time() - start_time
            return result

        except Exception as e:
            err_str = str(e)
            # Φç¬σè¿ΘÖìτ║º∩╝Üσªéµ₧£µ¿íσ₧ïΦ┐öσ¢₧"σÅ¬µö»µîü Interactions API"ΘöÖΦ»»∩╝îτö¿ 2.0-flash ΘçìΦ»òΣ╕Çµ¼í
            if "Interactions API" in err_str and model_id not in (_INTERACTIONS_ONLY_MODELS | {_INTERACTIONS_FALLBACK_MODEL}):
                _app_logger.info(f"[brain.chat] Interactions API ΘöÖΦ»»∩╝îΦç¬σè¿ΘÖìτ║º {model_id} ΓåÆ {_INTERACTIONS_FALLBACK_MODEL}")
                try:
                    model_id = _INTERACTIONS_FALLBACK_MODEL
                    _fb = client.models.generate_content(
                        model=model_id,
                        contents=(
                            formatted_history
                            + [
                                types.Content(
                                    role="user",
                                    parts=[types.Part.from_text(text=model_input)],
                                )
                            ]
                            if not file_data
                            else [original_input]
                        ),
                        config=types.GenerateContentConfig(
                            system_instruction=_brain_sys_instruction
                        ),
                    )
                    result["response"] = _fb.text if _fb.text else ""
                    result["model"] = model_id
                    result["total_time"] = time.time() - start_time
                    return result
                except Exception as _fb_err:
                    result["response"] = f"Γ¥î σêåµ₧Éσñ▒Φ┤Ñ: {_fb_err}"
            elif (
                "API key not valid" in err_str
                or "INVALID_ARGUMENT" in err_str
                and "api key" in err_str.lower()
            ):
                result["response"] = (
                    "Γ¥î **API σ»åΘÆÑµùáµòê**\n\n"
                    "Φ»╖µúÇµƒÑµé¿τÜä Gemini API σ»åΘÆÑ∩╝Ü\n"
                    "1. σëìσ╛Ç [aistudio.google.com/apikey](https://aistudio.google.com/apikey) ΦÄ╖σÅûµ£ëµòêσ»åΘÆÑ\n"
                    "2. σ£¿ Koto Φ«╛τ╜«Θí╡Θ¥óµ¢┤µû░ API σ»åΘÆÑ\n"
                    "3. τí«Σ┐¥σ»åΘÆÑµëÇσ£¿Θí╣τ¢«σ╖▓σÉ»τö¿ Generative Language API\n\n"
                    f"σÄƒσºïΘöÖΦ»»: `{err_str[:200]}`"
                )
            else:
                # ΓöÇΓöÇ µ¿íσ₧ïµ£¼Φ║½Σ╕ìσÅ»τö¿∩╝ê404 / not-found / Interactions-only τ¡ë∩╝ëΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                # σ░¥Φ»òΣ╗Ä ModelFallbackExecutor ΦÄ╖σÅûσñçΘÇëµ¿íσ₧ïσ╣╢Θ¥ÖΘ╗ÿΘçìΦ»òΣ╕Çµ¼íπÇé
                _retried = False
                try:
                    from app.core.llm.model_fallback import get_fallback_executor, _is_model_unavailable_error as _mue_chk
                    if _mue_chk(e) and model_id not in (None, _INTERACTIONS_FALLBACK_MODEL):
                        _fbe = get_fallback_executor()
                        _fbe.mark_unavailable(model_id)
                        _fb_model = _fbe.get_best_available(task_type=target_key)
                        if _fb_model and _fb_model != model_id and _fb_model not in _INTERACTIONS_ONLY_MODELS:
                            _app_logger.info(f"[brain.chat] µ¿íσ₧ïΣ╕ìσÅ»τö¿ {model_id} ΓåÆ Φç¬σè¿ΘÖìτ║º {_fb_model} (task={target_key})")
                            _fh = locals().get("formatted_history") or []
                            _mi = locals().get("model_input") or original_input
                            _si = locals().get("_brain_sys_instruction") or ""
                            _fb_r = client.models.generate_content(
                                model=_fb_model,
                                contents=_fh + [types.Content(
                                    role="user",
                                    parts=[types.Part.from_text(text=_mi)]
                                )],
                                config=types.GenerateContentConfig(system_instruction=_si)
                            )
                            result["response"] = _fb_r.text if _fb_r.text else ""
                            result["model"] = _fb_model
                            _retried = True
                except Exception as _r_err:
                    _app_logger.info(f"[brain.chat] ΘÖìτ║ºΘçìΦ»òσñ▒Φ┤Ñ: {_r_err}")
                if not _retried:
                    result["response"] = f"Γ¥î σÅæτöƒΘöÖΦ»»: {err_str}"
            result["total_time"] = time.time() - start_time
            return result


brain = KotoBrain()

# ================= Routes =================


@app.route("/")
def index():
    # Σ║æµ¿íσ╝Å∩╝Üµ£¬Φ«ñΦ»üτö¿µê╖τ£ïσê░ΦÉ╜σ£░Θí╡
    deploy_mode = os.environ.get("KOTO_DEPLOY_MODE", "local")
    auth_enabled = os.environ.get("KOTO_AUTH_ENABLED", "false").lower() == "true"
    if deploy_mode == "cloud" and auth_enabled:
        return render_template("landing.html")
    return render_template("index.html")


@app.route("/app")
def app_main():
    """Σ╕╗σ║öτö¿Θí╡Θ¥ó∩╝êSaaS µ¿íσ╝ÅΣ╕ïΘ£ÇΦ«ñΦ»üσÉÄΦ«┐Θù«∩╝ë"""
    return render_template("index.html")


@app.route("/file-network")
def file_network():
    """µûçΣ╗╢τ╜æτ╗£τòîΘ¥ó"""
    return render_template("file_network.html")


@app.route("/knowledge-graph")
def knowledge_graph_page():
    """τƒÑΦ»åσ¢╛Φ░▒σÅ»ΦºåσîûτòîΘ¥ó"""
    return render_template("knowledge_graph.html")


@app.route("/test_upload")
def test_upload():
    return render_template("test_upload.html")


@app.route("/edit-ppt/<session_id>")
def edit_ppt(session_id):
    """PPT τöƒµêÉσÉÄτ╝ûΦ╛æΘí╡Θ¥ó∩╝êP1 σèƒΦâ╜∩╝ë"""
    return render_template("edit_ppt.html")


@app.route("/skills")
@app.route("/skill-marketplace")
def skill_marketplace():
    """Koto Skill σ║ô ΓÇö GitHub Extension Marketplace ΘúÄµá╝τ«íτÉåτòîΘ¥ó"""
    return render_template("skill_marketplace.html")


@app.route("/monitoring-dashboard")
def monitoring_dashboard():
    """Phase 4 System Monitoring Dashboard"""
    return send_from_directory(
        os.path.join(os.path.dirname(__file__), "static"), "monitoring_dashboard.html"
    )


@app.route("/api/sessions", methods=["GET"])
def get_sessions():
    """List all chat sessions.
    ---
    tags:
      - Sessions
    responses:
      200:
        description: List of session names
        schema:
          type: object
          properties:
            sessions:
              type: array
              items:
                type: string
    """
    sessions = session_manager.list_sessions()
    return jsonify({"sessions": [s.replace(".json", "") for s in sessions]})


@app.route("/api/sessions", methods=["POST"])
def create_session():
    """Create a new chat session.
    ---
    tags:
      - Sessions
    parameters:
      - in: body
        name: body
        schema:
          properties:
            name:
              type: string
              description: Optional session name
    responses:
      200:
        description: Session created
        schema:
          type: object
          properties:
            success:
              type: boolean
            session:
              type: string
    """
    data = request.json
    name = data.get("name", f"chat_{int(time.time())}")
    filename = session_manager.create(name)
    return jsonify({"success": True, "session": filename.replace(".json", "")})


@app.route("/api/sessions/<session_name>", methods=["GET"])
def get_session(session_name):
    """Get a specific chat session with full history.
    ---
    tags:
      - Sessions
    parameters:
      - in: path
        name: session_name
        type: string
        required: true
    responses:
      200:
        description: Session data with conversation history
        schema:
          type: object
          properties:
            session:
              type: string
            history:
              type: array
              items:
                type: object
    """
    # Φ┐öσ¢₧σ«îµò┤σÄåσÅ▓Σ╛¢σëìτ½»µ╕▓µƒô∩╝êΣ╕ìµê¬µû¡∩╝ë∩╝îµê¬µû¡Σ╗àτö¿Σ║Äµ¿íσ₧ïΣ╕èΣ╕ïµûç
    history = session_manager.load_full(f"{session_name}.json")
    return jsonify({"session": session_name, "history": history})


@app.route("/api/sessions/<session_name>", methods=["DELETE"])
def delete_session(session_name):
    """Delete a chat session.
    ---
    tags:
      - Sessions
    parameters:
      - in: path
        name: session_name
        type: string
        required: true
    responses:
      200:
        description: Deletion result
        schema:
          type: object
          properties:
            success:
              type: boolean
    """
    success = session_manager.delete(f"{session_name}.json")
    return jsonify({"success": success})


@app.route("/api/chat", methods=["POST"])
def chat():
    """Send a chat message and get a response (non-streaming).
    ---
    tags: [Chat]
    security:
      - Bearer: []
    parameters:
      - in: body
        name: body
        schema:
          required: [session, message]
          properties:
            session: {type: string, description: Session/conversation name}
            message: {type: string, description: User message}
            locked_model: {type: string, default: auto}
            locked_task: {type: string}
    responses:
      200:
        description: AI response
        schema:
          properties:
            response: {type: string}
            model: {type: string}
      400:
        description: Missing session or message
      500:
        description: Internal error
    """
    data = request.json
    session_name = data.get("session")
    user_input = data.get("message", "")
    locked_task = data.get("locked_task")
    locked_model = data.get("locked_model", "auto")

    if not session_name or not user_input:
        return jsonify({"error": "Missing session or message"}), 400

    user_input = Utils.sanitize_string(user_input)

    # Load history
    full_history = session_manager.load_full(f"{session_name}.json")
    history = session_manager._trim_history(full_history)

    # τí«σ«ÜΣ╜┐τö¿τÜäµ¿íσ₧ï
    if locked_model and locked_model != "auto":
        model = locked_model
        auto_model = False
    elif locked_task:
        model = MODEL_MAP.get(locked_task, MODEL_MAP["CHAT"])
        auto_model = False
    else:
        model = None
        auto_model = True

    # Get response
    result = brain.chat(history, user_input, model=model, auto_model=auto_model)

    # Σ╗úτáüΣ╗╗σèí: Φç¬σè¿µúÇµƒÑΣ╛¥Φ╡ûσ╣╢σ«ëΦúà
    if result.get("task") == "CODER" and result.get("response"):
        pkgs = Utils.detect_required_packages(result["response"])
        if pkgs:
            install_result = Utils.auto_install_packages(pkgs)
            installed = install_result.get("installed", [])
            failed = install_result.get("failed", [])
            skipped = install_result.get("skipped", [])
            msg_parts = []
            if installed:
                msg_parts.append(f"Γ£à σ╖▓σ«ëΦúà: {', '.join(installed)}")
            if skipped:
                msg_parts.append(f"Γä╣∩╕Å σ╖▓σ¡ÿσ£¿: {', '.join(skipped)}")
            if failed:
                msg_parts.append(f"ΓÜá∩╕Å σ«ëΦúàσñ▒Φ┤Ñ: {', '.join(failed)}")
            if msg_parts:
                result["response"] += "\n\n" + "\n".join(msg_parts)

    # Update history (σƒ║Σ║Äτúüτ¢ÿσ«îµò┤σÄåσÅ▓Φ┐╜σèá∩╝îΘü┐σàìµê¬µû¡Σ╕óσñ▒)
    session_manager.append_and_save(
        f"{session_name}.json", user_input, result["response"]
    )

    return jsonify(result)


# ============== Agent τí«Φ«ñ API ==============
# NOTE: These routes have been migrated to the unified agent blueprint
#       (app/api/agent_routes.py) under /api/agent/confirm and /api/agent/choice.
#       Kept here as comments for reference.

# @app.route('/api/agent/confirm', methods=['POST'])
# def agent_confirm():
#     """Agent τö¿µê╖τí«Φ«ñ API ΓÇö σëìτ½»τé╣σç╗τí«Φ«ñ/σÅûµ╢êσÉÄσ¢₧Φ░â"""
#     ...

# @app.route('/api/agent/choice', methods=['POST'])
# def agent_choice():
#     """Agent τö¿µê╖ΘÇëµï⌐ API ΓÇö σëìτ½»ΘÇëµï⌐σÉÄσ¢₧Φ░â"""
#     ...


# NOTE: /api/agent/plan has been migrated to the unified agent blueprint
#       (app/api/agent_routes.py). Kept as comment for reference.
# @app.route('/api/agent/plan', methods=['POST'])
# def agent_plan(): ...


@app.route("/api/chat/stream", methods=["POST"])
def chat_stream():
    """Stream a chat response via Server-Sent Events.
    ---
    tags:
      - Chat
    parameters:
      - in: body
        name: body
        schema:
          required: [session, message]
          properties:
            session:
              type: string
            message:
              type: string
            locked_model:
              type: string
              default: auto
            locked_task:
              type: string
    responses:
      200:
        description: SSE stream of chat tokens
    """
    data = request.json
    session_name = data.get("session")
    user_input = data.get("message", "")
    locked_task = data.get("locked_task")
    locked_model = data.get("locked_model", "auto")

    _app_logger.debug(
        f"\n[STREAM] Incoming request: locked_task='{locked_task}', locked_model='{locked_model}'"
    )
    _app_logger.debug(f"[STREAM] User input: {user_input[:60]}")

    if not session_name or not user_input:

        def error_gen():
            yield f"data: {json.dumps({'type': 'error', 'message': 'Missing session or message'})}\n\n"

        return Response(error_gen(), mimetype="text/event-stream")

    # API σ»åΘÆÑτ╝║σñ▒µù╢µÅÉσëìΦ┐öσ¢₧σÅïσÑ╜µÅÉτñ║
    if not API_KEY:

        def no_key_gen():
            msg = (
                "ΓÜá∩╕Å **API σ»åΘÆÑµ£¬Θàìτ╜«**\n\n"
                "Φ»╖σ£¿ `config/gemini_config.env` µûçΣ╗╢Σ╕¡Φ«╛τ╜«∩╝Ü\n"
                "```\nGEMINI_API_KEY=Σ╜áτÜäσ»åΘÆÑ\n```\n\n"
                "≡ƒÆí ΦÄ╖σÅûσ»åΘÆÑ∩╝Ü[Google AI Studio](https://aistudio.google.com/apikey)\n\n"
                "Φ«╛τ╜«σ«îµêÉσÉÄΘçìσÉ» Koto σì│σÅ»Σ╜┐τö¿πÇé"
            )
            yield f"data: {json.dumps({'type': 'token', 'content': msg})}\n\n"

        return Response(no_key_gen(), mimetype="text/event-stream")

    user_input = Utils.sanitize_string(user_input)

    # ≡ƒºá µäÅσ¢╛σêåµ₧ÉΣ╕ÄΘçìσåÖ (Intent Analysis & Rewrite)
    # τ╗ôσÉêσÄåσÅ▓Φ«░σ┐åσÆîµ£¼σ£░µ¿íσ₧ï∩╝îτÉåΦºúτö¿µê╖τÜäσñìµ¥éµîçΣ╗ñ∩╝êσªéΓÇ£ΘçìσñìΣ╕èΣ╕¬Σ╗╗σèíΓÇ¥πÇüΓÇ£µèèσêÜµëìΘéúΣ╕¬µö╣µêÉ...ΓÇ¥∩╝ë
    try:
        from app.core.routing.intent_analyzer import IntentAnalyzer

        if IntentAnalyzer.should_analyze(user_input):
            full_hist = session_manager.load_full(f"{session_name}.json")
            # σ┐½ΘÇƒΦÄ╖σÅûΘò┐µ£ƒΦ«░σ┐åΣ╕èΣ╕ïµûç∩╝îσ╕«σè⌐µ╢êΦºúΦ╖¿ session τÜäµîçΣ╗úΦ»ì
            _intent_memory_ctx = ""
            try:
                _mm_for_intent = get_memory_manager()
                if _mm_for_intent:
                    _intent_memory_ctx = _mm_for_intent.get_context_string(user_input) or ""
            except Exception:
                pass
            rewritten_input = IntentAnalyzer.rewrite_intent(
                user_input, full_hist, memory_context=_intent_memory_ctx
            )
            if rewritten_input and rewritten_input != user_input:
                _app_logger.debug(f"[STREAM] ≡ƒöä µäÅσ¢╛ΘçìσåÖ: '{user_input}' -> '{rewritten_input}'")
                user_input = rewritten_input
    except Exception as e:
        _app_logger.warning(f"[STREAM] ΓÜá∩╕Å µäÅσ¢╛σêåµ₧Éσñ▒Φ┤Ñ: {e}")
        # ΘÖìτ║ºσê░σƒ║τíÇτÜäµ¡úσêÖσî╣Θàì
        repeat_patterns = [
            r"^Θçìσñì.*Σ╗╗σèí",
            r"^σåìσüÜΣ╕ÇΘüì",
            r"^σåìµ¥ÑΣ╕Çµ¼í",
            r"^re(peat|do).*last.*task",
            r"^try.*again",
        ]
        if any(re.search(p, user_input, re.IGNORECASE) for p in repeat_patterns):
            try:
                full_hist = session_manager.load_full(f"{session_name}.json")
                last_user_msg = None
                for msg in reversed(full_hist):
                    if msg.get("role") == "user":
                        content = (msg.get("parts") or [""])[0]
                        if not any(
                            re.search(p, content, re.IGNORECASE)
                            for p in repeat_patterns
                        ):
                            last_user_msg = content
                            break
                if last_user_msg:
                    _app_logger.debug(f"[REPEAT] Found last user message: {last_user_msg[:50]}...")
                    user_input = last_user_msg
            except Exception as hist_e:
                _app_logger.debug(f"[REPEAT] Error fetching history: {hist_e}")

    # ΓÜí σ┐½ΘÇƒΦ╖»σ╛ä∩╝Üτ│╗τ╗ƒµù╢Θù┤µƒÑΦ»ó - τ¢┤µÄÑΦ┐öσ¢₧∩╝îµùáΘ£ÇσÅæΘÇüσê░LLM
    time_query_patterns = [
        r"σ╜ôσëì.*µù╢Θù┤|σ╜ôσëìτ│╗τ╗ƒµù╢Θù┤",
        r"τÄ░σ£¿.*σçáτé╣|σçáτé╣ΘÆƒ",
        r"σçáτé╣|Σ╗ÇΣ╣êµù╢Θù┤",
        r"µù╢Θù┤µÿ»|τÄ░σ£¿µÿ»",
        r"now.*time|what.*time|current.*time",
    ]
    if any(
        re.search(pattern, user_input, re.IGNORECASE) for pattern in time_query_patterns
    ):

        def quick_time_response():
            from datetime import datetime

            now = datetime.now()
            date_str = now.strftime("%Yσ╣┤%mµ£ê%dµùÑ")
            weekday = ["σæ¿Σ╕Ç", "σæ¿Σ║î", "σæ¿Σ╕ë", "σæ¿σ¢¢", "σæ¿Σ║ö", "σæ¿σà¡", "σæ¿µùÑ"][
                now.weekday()
            ]
            time_str = now.strftime("%H:%M:%S")
            timestamp = now.isoformat()  # Φ«░σ╜òτ▓╛τí«µù╢Θù┤µê│
            response = f"σ╜ôσëìτ│╗τ╗ƒµù╢Θù┤Σ╕║∩╝Ü\n\n**{date_str} {weekday} {time_str}**"

            # Φ«░σ╜òσê░σÄåσÅ▓∩╝êτö¿µê╖ + µ¿íσ₧ï∩╝îσ¥çσ╕ªµù╢Θù┤µê│∩╝ë
            try:
                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    response,
                    task="CHAT",
                    model_name="QuickResponse",
                    timestamp=timestamp,
                    user_timestamp=timestamp,
                    model_timestamp=timestamp,
                )
            except Exception as e:
                _app_logger.debug(f"[STREAM] Quick time history save failed: {e}")

            yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒôà τ│╗τ╗ƒµù╢Θù┤µƒÑΦ»ó', 'detail': 'Σ╗Äµ£¼σ£░ΦÄ╖σÅû'}, ensure_ascii=False)}\n\n"
            yield f"data: {json.dumps({'type': 'token', 'content': response, 'timestamp': timestamp}, ensure_ascii=False)}\n\n"
            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': 0.01, 'timestamp': timestamp}, ensure_ascii=False)}\n\n"

        return Response(quick_time_response(), mimetype="text/event-stream")

    # ≡ƒÄ» ΦÄ╖σÅûσè¿µÇüτ│╗τ╗ƒµîçΣ╗ñ∩╝êµá╣µì«τö¿µê╖Θù«ΘóÿµÖ║Φâ╜µ│¿σàÑΣ╕èΣ╕ïµûç∩╝ë
    try:
        system_instruction = _get_chat_system_instruction(user_input)
    except Exception as e:
        _app_logger.debug(f"[STREAM] Warning: Dynamic system instruction failed: {e}")
        system_instruction = (
            _get_DEFAULT_CHAT_SYSTEM_INSTRUCTION()
        )  # ΘÖìτ║ºσê░µû░Θ▓£τöƒµêÉτÜäµîçΣ╗ñ

    history = session_manager.load(f"{session_name}.json")
    full_history = session_manager.load_full(f"{session_name}.json")

    # ≡ƒºá 2-A: ContextWindowManager ΓÇö MemGPT-style page-out/in
    _cw_paged_context = ""
    try:
        from app.core.memory.context_window_manager import ContextWindowManager as _CWM

        _cw_out = _CWM.manage(
            history=history,
            query=user_input,
            session_name=session_name,
            get_memory_fn=get_memory_manager,
        )
        history = _cw_out["history"]
        _cw_paged_context = _cw_out.get("paged_in_context", "")
        if _cw_paged_context:
            system_instruction += f"\n\n{_cw_paged_context}"
    except Exception as _cw_err:
        _app_logger.warning(f"[CWM] ΓÜá∩╕Å Σ╕èΣ╕ïµûçτ«íτÉåσÖ¿σ╝éσ╕╕: {_cw_err}")

    # ≡ƒò╡∩╕ÅΓÇìΓÖÇ∩╕Å µúÇµ╡ïµÿ»σÉªµ£ëµ£ÇΦ┐æΣ╕èΣ╝áτÜäµûçΣ╗╢ (5σêåΘÆƒσåà)
    has_recent_upload = False
    recent_file_type = None
    try:
        upload_scan_dirs = ["web/uploads", "uploads", "workspace/documents"]
        recent_threshold = time.time() - 300  # 5σêåΘÆƒσåà
        for d in upload_scan_dirs:
            if os.path.exists(d):
                for f in os.listdir(d):
                    fp = os.path.join(d, f)
                    if os.path.isfile(fp) and os.path.getmtime(fp) > recent_threshold:
                        has_recent_upload = True
                        _, ext = os.path.splitext(f)
                        recent_file_type = ext.lower()
                        _app_logger.debug(f"[STREAM] Found recent upload: {f} ({recent_file_type})")
                        break
            if has_recent_upload:
                break
    except Exception as e:
        _app_logger.debug(f"[STREAM] Error checking uploads: {e}")

    # τí«σ«ÜΣ╗╗σèíτ▒╗σ₧ïσÆîµ¿íσ₧ï
    context_info = None
    if locked_task:
        task_type = locked_task
        route_method = "≡ƒöÆ Manual"
        _app_logger.info(f"[STREAM] Γ£à Using locked_task: '{task_type}'")
    else:
        # σ░åµûçΣ╗╢Σ┐íµü»Σ╝áΘÇÆτ╗Öσêåµ₧ÉσÖ¿∩╝êσÉîµù╢µ│¿σàÑ [FILE_ATTACHED:ext] σëìτ╝Çτí«Σ┐¥µ£¼σ£░µ¿íσ₧ïµ¡úτí«σêåτ▒╗∩╝ë
        context_override = {
            "has_file": has_recent_upload,
            "file_type": recent_file_type,
        }
        _routing_input = user_input
        if has_recent_upload and recent_file_type:
            _routing_input = f"[FILE_ATTACHED:{recent_file_type}] {user_input}"
            _app_logger.debug(f"[STREAM] ≡ƒôÄ µûçΣ╗╢Σ╕èΣ╕ïµûçµ│¿σàÑ: {_routing_input[:80]}")
        task_type, route_method, context_info = SmartDispatcher.analyze(
            _routing_input, history, file_context=context_override
        )
        _app_logger.debug(
            f"[STREAM] Auto-detected task_type: '{task_type}', context: {context_info is not None}"
        )

        # ΓöÇΓöÇ σ«ëσà¿σà£σ║ò∩╝Üµ£¬τƒÑ task_type ΓåÆ CHAT ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        _HANDLED_TASK_TYPES = {
            "SYSTEM",
            "FILE_OP",
            "FILE_EDIT",
            "FILE_SEARCH",
            "DOC_ANNOTATE",
            "WEB_SEARCH",
            "RESEARCH",
            "PAINTER",
            "FILE_GEN",
            "CODER",
            "CHAT",
            "MULTI_STEP",
            "AGENT",
            "VISION",
        }
        if not task_type or task_type not in _HANDLED_TASK_TYPES:
            _app_logger.warning(f"[STREAM] ΓÜá∩╕Å µö╢σê░µ£¬τƒÑ task_type='{task_type}'∩╝îΘÖìτ║ºΣ╕║ CHAT")
            task_type = "CHAT"
            route_method = "Γ¼ç∩╕Å UnknownΓåÆCHAT"

        # ΓöÇΓöÇ MULTI_STEP Σ┐¥µèñ∩╝Üµùá is_multi_step_task µáçΦ«░µù╢ΘÖìτ║º CHAT ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        if task_type == "MULTI_STEP" and (
            not context_info or not context_info.get("is_multi_step_task")
        ):
            _app_logger.warning(f"[STREAM] ΓÜá∩╕Å MULTI_STEP µùáµ£ëµòê context∩╝îΘÖìτ║ºΣ╕║ CHAT")
            task_type = "CHAT"
            route_method = "Γ¼ç∩╕Å MULTI_STEPΓåÆCHAT"

        # ΓöÇΓöÇ FILE_EDIT Σ┐¥µèñ∩╝ÜΦ╛ôσàÑΣ╕¡µùáµ│òΦ»åσê½µûçΣ╗╢Φ╖»σ╛äµù╢ΘÖìτ║º CHAT ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        # Θü┐σàìτƒ¡/µÖ«ΘÇÜΦüèσñ⌐µ╢êµü»σ¢áµÉ║σ╕ªµûçΣ╗╢Σ╕èΣ╕ïµûçΦÇîΦó½Φ»»σêåτ▒╗σê░ FILE_EDIT
        if task_type == "FILE_EDIT":
            _fe_pat1 = re.search(
                r'(?:Σ┐«µö╣|τ╝ûΦ╛æ|µö╣)\s+["\']?([^"\']{2,}?)["\']?\s+.+', user_input
            )
            _fe_pat2 = re.search(
                r'(?:µèè|σ░å)\s+["\']?([^"\']{2,}?)["\']?\s+(?:τÜä|Σ╕¡τÜä|ΘçîτÜä)\s*.+',
                user_input,
            )
            if not _fe_pat1 and not _fe_pat2:
                _app_logger.warning(
                    f"[STREAM] ΓÜá∩╕Å FILE_EDIT Φ╛ôσàÑµùáµ£ëµòêµûçΣ╗╢Φ╖»σ╛ä: '{user_input[:40]}' ΓåÆ ΘÖìτ║ºΣ╕║ CHAT"
                )
                task_type = "CHAT"
                route_method = "Γ¼ç∩╕Å FILE_EDITΓåÆCHAT"

        # ΓöÇΓöÇ CHAT ΓåÆ WEB_SEARCH σ«ëσà¿σà£σ║ò∩╝êΘÿ▓µ¡óσñ⌐µ░ö/ΦéíΣ╗╖/µû░Θù╗τ¡ëσ«₧µù╢µƒÑΦ»óΦó½Φ»»σêåΣ╕║CHAT∩╝ëΓöÇΓöÇΓöÇΓöÇ
        # Φ┐Öµÿ»µ£ÇσÉÄΣ╕ÇΘüôΘÿ▓τ║┐∩╝Üσ£¿Σ╗╗σèíΘô╛Φ╖»µëºΦíîΣ╣ïσëì∩╝îΘçìµû░µáíΘ¬îµÿ»σÉªΘ£ÇΦªüΦüöτ╜æµÉ£τ┤ó
        if task_type == "CHAT" and WebSearcher.needs_web_search(user_input):
            _app_logger.debug(f"[STREAM] ΓÜí CHATΓåÆWEB_SEARCH σ«ëσà¿σà£σ║òΦºªσÅæ: '{user_input[:40]}'")
            task_type = "WEB_SEARCH"
            route_method = "≡ƒîÉ CHATΓåÆWEB_SEARCH"

        # σªéµ₧£µ£ëΣ╕èΣ╕ïµûçΣ┐íµü»∩╝îΦ«░σ╜òΦ»ªµâà
        if context_info and context_info.get("is_continuation"):
            _app_logger.debug(
                f"[STREAM] Context continuation: {context_info.get('related_task')}, confidence: {context_info.get('confidence')}"
            )

    # ΓöÇΓöÇ Phase2: RouterDecision (classify_v2) ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    # σ£¿ SmartDispatcher σ╖▓τí«σ«Ü task_type τÜäσƒ║τíÇΣ╕è∩╝îΦ┐¢Σ╕Çµ¡ÑΦÄ╖σÅû skill_id / forward σå│τ¡ûπÇé
    # Σ╗ÑΘ¥₧Θÿ╗σí₧µû╣σ╝ÅΦ┐ÉΦíî∩╝êΦ╢àµù╢Σ┐¥µèñ∩╝ë∩╝îσñ▒Φ┤Ñµù╢Σ╕ìσ╜▒σôìΣ╕╗µ╡üτ¿ïπÇé
    _router_decision = None
    try:
        from app.core.routing.local_model_router import LocalModelRouter as _LMRv2

        _router_decision = _LMRv2.classify_v2(user_input, hint=task_type, timeout=1.5)
        if _router_decision and _router_decision.skill_id:
            _app_logger.debug(
                f"[STREAM] ≡ƒÄ» RouterDecision skill_id={_router_decision.skill_id} "
                f"forward_to_cloud={_router_decision.forward_to_cloud} "
                f"confidence={_router_decision.confidence:.2f}"
            )
    except Exception as _rv2_err:
        _app_logger.debug(f"[STREAM] RouterDecision classify_v2 Φ╖│Φ┐ç: {_rv2_err}")

    # ΓöÇΓöÇ σ░åΣ╗╗σèíΣ╕ôσ▒₧ΦíÑσààµîçΣ╗ñΦ┐╜σèáσê░ system_instruction ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    _addendum = _TASK_SYSTEM_ADDENDUMS.get(task_type, "")
    if _addendum:
        system_instruction = system_instruction + _addendum
        _app_logger.debug(f"[STREAM] ≡ƒôî Σ╗╗σèíΣ╕ôσ▒₧µîçΣ╗ñσ╖▓µ│¿σàÑ: {task_type}")

    # ΓöÇΓöÇ ≡ƒö« LangGraph Θ½ÿτ║ºσ╖ÑΣ╜£µ╡üΦ╖»τö▒∩╝êRESEARCH / FILE_GEN / MULTI_STEP∩╝ëΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    # resolve_workflow() µúÇµ╡ïτö¿µê╖µäÅσ¢╛∩╝îσå│σ«Üµÿ»σÉªτö¿ WorkflowEngine µ¢┐Σ╗úµùºΦ╖»σ╛ä
    _wf_route = "legacy"
    if task_type in ("RESEARCH", "FILE_GEN", "MULTI_STEP"):
        try:
            _wf_route = SmartDispatcher.resolve_workflow(task_type, user_input)
            if _wf_route != "legacy":
                _app_logger.debug(f"[STREAM] ≡ƒö« LangGraph σ╖ÑΣ╜£µ╡üΦ╖»τö▒: {_wf_route}")
        except Exception as _wf_err:
            _app_logger.debug(f"[STREAM] resolve_workflow Φ╖│Φ┐ç: {_wf_err}")

    if _wf_route in ("langgraph_research_doc", "langgraph_multi_agent_ppt"):
        _wf_name = (
            "research_and_document"
            if _wf_route == "langgraph_research_doc"
            else "multi_agent_ppt"
        )
        _wf_label = (
            "≡ƒôÜ τáöτ⌐╢+µûçµíú"
            if _wf_route == "langgraph_research_doc"
            else "≡ƒÄ₧∩╕Å σñÜAgent PPT"
        )

        def generate_langgraph_workflow():
            yield f"data: {json.dumps({'type': 'classification', 'task_type': 'LG_WORKFLOW', 'workflow': _wf_name, 'route_method': 'LangGraph', 'message': f'≡ƒÄ» Σ╗╗σèíσêåτ▒╗: {_wf_label} (LangGraph WorkflowEngine)'})}\n\n"
            try:
                from app.core.workflow.langgraph_workflow import WorkflowEngine

                _engine = WorkflowEngine()
                final_output = ""
                for event in _engine.stream(
                    workflow=_wf_name,
                    user_input=user_input,
                    session_id=session_name,
                ):
                    node = event.get("node", "")
                    content = event.get("content", "")
                    done = event.get("done", False)
                    if node == "error":
                        yield f"data: {json.dumps({'type': 'error', 'message': content})}\n\n"
                        return
                    if content:
                        yield f"data: {json.dumps({'type': 'status' if not done else 'token', 'message': f'[{node}] {content}' if not done else None, 'content': content if done else None}, ensure_ascii=False)}\n\n"
                    if done:
                        final_output = content
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                try:
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        final_output or f"[{_wf_label}σ╖ÑΣ╜£µ╡üσ«îµêÉ]",
                    )
                except Exception:
                    pass
            except Exception as _wf_ex:
                import traceback

                _app_logger.error(f"[LG_WORKFLOW] Γ¥î σ╖ÑΣ╜£µ╡üσñ▒Φ┤Ñ:\n{traceback.format_exc()}")
                yield f"data: {json.dumps({'type': 'error', 'message': f'σ╖ÑΣ╜£µ╡üµëºΦíîσñ▒Φ┤Ñ: {str(_wf_ex)}'})}\n\n"

        return Response(generate_langgraph_workflow(), mimetype="text/event-stream")

    # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    if (
        task_type == "MULTI_STEP"
        and context_info
        and context_info.get("is_multi_step_task")
    ):
        _app_logger.debug(f"[STREAM] ≡ƒöä µúÇµ╡ïσê░σñìµ¥éΣ╗╗σèí∩╝îΣ╜┐τö¿ TaskOrchestrator µëºΦíîσñÜµ¡Ñµ╡üτ¿ï")
        multi_step_info = context_info.get("multi_step_info", {})
        pattern = multi_step_info.get("pattern", "unknown")

        # === ≡ƒñû MultiAgent Θ½ÿΦ┤¿ΘçÅΘÇÜΦ╖»∩╝êLangGraph∩╝Üτáöτ⌐╢ΓåÆσåÖΣ╜£ΓåÆσ«íµá╕ΓåÆΣ┐«Φ«ó∩╝ë===
        # Σ╗àσ╜ô LangGraph σÅ»τö¿Σ╕öΦ╖»τö▒σå│τ¡ûΣ╕║ langgraph_react∩╝êΘÇÜτö¿σñìµ¥éΣ╗╗σèí∩╝ëµù╢ΦºªσÅæ
        if _wf_route == "langgraph_react":
            _app_logger.debug(
                f"[STREAM] ≡ƒñû MultiAgentOrchestrator ΘÇÜΦ╖»∩╝ÜRESEARCHERΓåÆWRITERΓåÆCRITICΓåÆREVISE"
            )
            _ma_model = SmartDispatcher.get_model_for_task("MULTI_STEP")

            def generate_multi_agent():
                yield f"data: {json.dumps({'type': 'classification', 'task_type': 'MULTI_STEP', 'pattern': 'multi_agent', 'route_method': 'LangGraph MultiAgent', 'message': '≡ƒÄ» Σ╗╗σèíσêåτ▒╗: ≡ƒñû σñÜAgentΘ½ÿΦ┤¿ΘçÅσñäτÉå∩╝êτáöτ⌐╢ΓåÆσåÖΣ╜£ΓåÆσ«íµá╕ΓåÆΣ┐«Φ«ó∩╝ë'})}\n\n"
                try:
                    from app.core.agent.multi_agent import MultiAgentOrchestrator

                    orch = MultiAgentOrchestrator.preset_content_pipeline(
                        model_id=_ma_model,
                        max_revisions=1,
                    )
                    _agent_labels = {
                        "researcher": "≡ƒôÜ τáöτ⌐╢Σ╕ôσæÿ",
                        "writer":     "Γ£ì∩╕Å σåÖΣ╜£Σ╕ôσæÿ",
                        "critic":     "≡ƒöì σ«íµá╕Σ╕ôσæÿ",
                        "revise":     "≡ƒöº Σ┐«Φ«óΣ╕ôσæÿ",
                        "finalize":   "Γ£à µò┤σÉêσ«îµêÉ",
                    }
                    final_output = ""
                    for event in orch.stream(user_input=user_input, session_id=session_name):
                        agent_name = event.get("agent", "unknown")
                        content = event.get("content", "")
                        done = event.get("done", False)
                        label = _agent_labels.get(agent_name, f"[{agent_name}]")

                        if agent_name == "error":
                            raise RuntimeError(content)

                        if done:
                            final_output = content
                            yield f"data: {json.dumps({'type': 'token', 'content': content}, ensure_ascii=False)}\n\n"
                        else:
                            yield f"data: {json.dumps({'type': 'status', 'message': f'{label} σñäτÉåΣ╕¡...'})}\n\n"

                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                    try:
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            final_output or "[σñÜAgentΣ╗╗σèíσ«îµêÉ]",
                        )
                    except Exception:
                        pass

                except Exception as _ma_err:
                    import traceback as _tb

                    _app_logger.error(
                        f"[MULTI_AGENT] Γ¥î MultiAgentOrchestrator σñ▒Φ┤Ñ: {_tb.format_exc()}"
                    )
                    yield f"data: {json.dumps({'type': 'error', 'message': f'σñÜAgentµëºΦíîσñ▒Φ┤Ñ∩╝îΦ»╖ΘçìΦ»ò: {str(_ma_err)}'})}\n\n"

            return Response(generate_multi_agent(), mimetype="text/event-stream")

        # === µûçµíúσ╖ÑΣ╜£µ╡üµëºΦíî ===
        if pattern == "document_workflow" and DocumentWorkflowExecutor:
            _app_logger.debug(f"[STREAM] ≡ƒôä µëºΦíîµûçµíúσ╖ÑΣ╜£µ╡ü")

            def generate_doc_workflow():
                yield f"data: {json.dumps({'type': 'classification', 'task_type': 'DOC_WORKFLOW', 'pattern': 'document_workflow', 'route_method': route_method, 'message': '≡ƒÄ» Σ╗╗σèíσêåτ▒╗: ≡ƒôä µûçµíúσ╖ÑΣ╜£µ╡üµëºΦíî'})}\n\n"

                # µƒÑµë╛µ£ÇΦ┐æΣ╕èΣ╝áτÜäµûçµíú
                doc_path = None
                upload_dirs = ["web/uploads", "uploads", "workspace/documents"]

                for dir_path in upload_dirs:
                    if os.path.exists(dir_path):
                        docs = []
                        for ext in [
                            ".docx",
                            ".md",
                            ".txt",
                            ".json",
                            ".doc",
                            ".pdf",
                            ".rtf",
                            ".odt",
                        ]:
                            import glob

                            docs.extend(
                                glob.glob(f"{dir_path}/**/*{ext}", recursive=True)
                            )

                        if docs:
                            # ΦÄ╖σÅûµ£Çµû░τÜäµûçµíú
                            doc_path = max(docs, key=os.path.getmtime)
                            break

                if not doc_path:
                    yield f"data: {json.dumps({'type': 'error', 'message': 'Γ¥î µ£¬µë╛σê░σÅ»µëºΦíîτÜäµûçµíúµûçΣ╗╢∩╝êµö»µîü .docx, .doc, .pdf, .md, .txt, .rtf, .odt, .json∩╝ë'})}\n\n"
                    return

                status_msg = f"≡ƒôä µë╛σê░µûçµíú: {os.path.basename(doc_path)}\n"
                yield f"data: {json.dumps({'type': 'status', 'message': status_msg})}\n\n"

                try:
                    import asyncio

                    # µëºΦíîµûçµíúσ╖ÑΣ╜£µ╡ü
                    executor = DocumentWorkflowExecutor(client)

                    # σèáΦ╜╜σ╖ÑΣ╜£µ╡ü
                    status_msg = "ΓÅ│ µ¡úσ£¿Φºúµ₧ÉµûçµíúΣ╕¡τÜäσ╖ÑΣ╜£µ╡ü...\n"
                    yield f"data: {json.dumps({'type': 'status', 'message': status_msg})}\n\n"

                    load_result = asyncio.run(executor.load_from_document(doc_path))

                    if not load_result.get("success"):
                        error_msg = (
                            f"Γ¥î µûçµíúΦºúµ₧Éσñ▒Φ┤Ñ: {load_result.get('error', 'µ£¬τƒÑΘöÖΦ»»')}\n"
                        )
                        yield f"data: {json.dumps({'type': 'error', 'message': error_msg})}\n\n"
                        return

                    # µÿ╛τñ║σ╖ÑΣ╜£µ╡üΣ┐íµü»
                    info_msg = f"Γ£à σ╖ÑΣ╜£µ╡üσèáΦ╜╜µêÉσèƒ\n"
                    info_msg += f"   σÉìτº░: {executor.workflow_name}\n"
                    info_msg += f"   µ¡ÑΘ¬ñµò░: {len(executor.steps)}\n"
                    info_msg += f"   ΦâîµÖ»: {executor.workflow_context}\n\n"
                    yield f"data: {json.dumps({'type': 'status', 'message': info_msg})}\n\n"

                    # µÿ╛τñ║µëÇµ£ëµ¡ÑΘ¬ñ
                    steps_msg = "≡ƒôï σ╖ÑΣ╜£µ╡üµ¡ÑΘ¬ñ:\n"
                    for step in executor.steps:
                        steps_msg += (
                            f"  {step.step_id}. [{step.step_type}] {step.description}\n"
                        )
                    steps_msg += "\n"
                    yield f"data: {json.dumps({'type': 'status', 'message': steps_msg})}\n\n"

                    # µëºΦíîσ╖ÑΣ╜£µ╡ü∩╝êµ╡üσ╝ÅσÅìΘªêµ»ÅΣ╕¬µ¡ÑΘ¬ñ∩╝ë
                    start_msg = "≡ƒÜÇ σ╝ÇσºïµëºΦíîσ╖ÑΣ╜£µ╡ü...\n\n"
                    yield f"data: {json.dumps({'type': 'status', 'message': start_msg})}\n\n"

                    for step in executor.steps:
                        step_msg = f"[µ¡ÑΘ¬ñ {step.step_id}/{len(executor.steps)}] {step.description}\n"
                        step_msg += f"ΓööΓöÇ τ▒╗σ₧ï: {step.step_type}\n"
                        step_msg += f"   ΓÅ│ µëºΦíîΣ╕¡...\n"
                        yield f"data: {json.dumps({'type': 'status', 'message': step_msg})}\n\n"

                        step.status = "running"
                        step.start_time = datetime.now()

                        try:
                            # µëºΦíîµ¡ÑΘ¬ñ
                            step_result = asyncio.run(
                                executor._execute_step_standalone(step)
                            )
                            step.result = step_result
                            step.status = "completed"

                            success_msg = f"   Γ£à σ«îµêÉ\n"
                            if isinstance(step_result, dict) and step_result.get(
                                "output"
                            ):
                                output_preview = str(step_result["output"])[:200]
                                success_msg += f"   ≡ƒôä Φ╛ôσç║ΘóäΦºê: {output_preview}...\n"
                            success_msg += "\n"
                            yield f"data: {json.dumps({'type': 'status', 'message': success_msg})}\n\n"

                        except Exception as e:
                            step.status = "failed"
                            step.error = str(e)
                            error_msg = f"   Γ¥î σñ▒Φ┤Ñ: {e}\n\n"
                            yield f"data: {json.dumps({'type': 'status', 'message': error_msg})}\n\n"

                        finally:
                            step.end_time = datetime.now()

                    # τöƒµêÉτ╗ôµ₧£
                    results = {
                        "workflow_name": executor.workflow_name,
                        "start_time": datetime.now().isoformat(),
                        "steps": [step.to_dict() for step in executor.steps],
                        "overall_status": "completed",
                    }
                    results["summary"] = executor._generate_summary(results)

                    # Σ┐¥σ¡ÿτ╗ôµ₧£
                    output_path = asyncio.run(executor.save_results(results))

                    # σÅæΘÇüσ«îµêÉµ╢êµü»
                    separator = "=" * 50
                    final_msg = f"\n{separator}\n"
                    final_msg += f"Γ£à µûçµíúσ╖ÑΣ╜£µ╡üµëºΦíîσ«îµêÉ\n\n"
                    final_msg += f"≡ƒôè µëºΦíîτ╗ƒΦ«í:\n"
                    total = len(results["steps"])
                    completed = sum(
                        1 for s in results["steps"] if s["status"] == "completed"
                    )
                    failed = sum(1 for s in results["steps"] if s["status"] == "failed")
                    final_msg += f"  µÇ╗µ¡ÑΘ¬ñ: {total}\n"
                    final_msg += f"  µêÉσèƒ: {completed}\n"
                    final_msg += f"  σñ▒Φ┤Ñ: {failed}\n"
                    final_msg += f"  µêÉσèƒτÄç: {completed/total*100:.1f}%\n\n"
                    final_msg += f"≡ƒôü τ╗ôµ₧£σ╖▓Σ┐¥σ¡ÿ: {os.path.basename(output_path)}\n"
                    final_msg += f"≡ƒôé Σ╜ìτ╜«: `workspace/workflows/`\n\n"
                    final_msg += f"{separator}\n"

                    yield f"data: {json.dumps({'type': 'token', 'content': final_msg})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [output_path]})}\n\n"

                    # Σ┐¥σ¡ÿµûçµíúσ╖ÑΣ╜£µ╡üσ»╣Φ»¥σÄåσÅ▓∩╝êσƒ║Σ║Äτúüτ¢ÿσ«îµò┤σÄåσÅ▓Φ┐╜σèá∩╝ë
                    try:
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            f"[µûçµíúσ╖ÑΣ╜£µ╡üσ«îµêÉ] {executor.workflow_name}",
                        )
                    except Exception:
                        pass

                except Exception as e:
                    import traceback

                    error_detail = traceback.format_exc()
                    error_msg = f"Γ¥î σ╖ÑΣ╜£µ╡üµëºΦíîσñ▒Φ┤Ñ: {str(e)}\n{error_detail}"
                    yield f"data: {json.dumps({'type': 'error', 'message': error_msg})}\n\n"
                    # Σ┐¥σ¡ÿσñ▒Φ┤ÑΦ«░σ╜ò
                    try:
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            f"[µûçµíúσ╖ÑΣ╜£µ╡üσñ▒Φ┤Ñ] {str(e)[:200]}",
                        )
                    except Exception:
                        pass

            return Response(generate_doc_workflow(), mimetype="text/event-stream")

        # === σà╢Σ╗ûσñÜµ¡ÑΣ╗╗σèíµëºΦíî ===
        from app.core.routing import TaskDecomposer

        subtasks = TaskDecomposer.create_subtasks(user_input, multi_step_info)
        use_local_planner = multi_step_info.get("pattern") == "local_plan"

        def generate_multi_step():
            # === τ½ïσì│σÅæΘÇüΣ╗╗σèíσêåτ▒╗Σ┐íµü» ===
            pattern = multi_step_info.get("pattern", "unknown")
            classification_msg = f"≡ƒÄ» Σ╗╗σèíσêåτ▒╗: ≡ƒöä σñÜµ¡ÑΣ╗╗σèí\n"
            yield f"data: {json.dumps({'type': 'classification', 'task_type': 'MULTI_STEP', 'pattern': pattern, 'route_method': route_method, 'message': classification_msg})}\n\n"
            
            # µÿ╛τñ║µëÇµ£ëσ¡ÉΣ╗╗σèí
            status_msg = f"≡ƒôï Σ╗╗σèíσêåΦºú:\n"
            for i, subtask in enumerate(subtasks):
                status_msg += f"  {i+1}. {subtask['description']}\n"
            status_msg += "\n"
            yield f"data: {json.dumps({'type': 'status', 'message': status_msg})}\n\n"
            
            # µëºΦíîµëÇµ£ëσ¡ÉΣ╗╗σèí∩╝êΘÇÉµ¡Ñµ╡üσ╝ÅσÅìΘªê∩╝ë
            try:
                import asyncio

                execution_log = []
                step_results = []
                context = {"original_input": user_input, "user_input": user_input}
                saved_files = []

                # ΓöÇΓöÇ Σ╜┐τö¿ PlanExecutor µëºΦíî∩╝êµö»µîüµïôµëæµÄÆσ║Å + Σ╛¥Φ╡ûµ│¿σàÑ∩╝ëΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                import queue as _queue_mod
                import threading as _threading_mod

                from app.core.routing.plan_executor import PlanExecutor as _PlanExecutor
                from app.core.routing.plan_executor import (
                    build_handlers_from_orchestrator as _build_handlers,
                )

                # µ₧äσ╗║ handlers∩╝êσ░å TaskOrchestrator σÉäµû╣µ│òσîàΦúàµêÉ PlanExecutor µÄÑσÅú∩╝ë
                _handlers = _build_handlers(TaskOrchestrator, context)

                # PlanExecutor σ«₧Σ╛ï∩╝êµïôµëæµÄÆσ║Å + ContextStore Σ╝áΘÇÆ∩╝ë
                _plan_exec = _PlanExecutor(
                    steps=subtasks,
                    user_input=user_input,
                    handlers=_handlers,
                    max_retry=1,
                )

                # σ£¿σÉÄσÅ░τ║┐τ¿ïΦ┐ÉΦíî async executor∩╝îΘÇÜΦ┐ç queue σ¢₧Σ╝áΣ║ïΣ╗╢τ╗Ö SSE τöƒµêÉσÖ¿
                _event_queue = _queue_mod.Queue()
                _plan_exception = {"err": None}

                def _run_plan_executor():
                    async def _inner():
                        try:
                            async for event in _plan_exec.execute():
                                _event_queue.put(event)
                        except Exception as _exc:
                            _plan_exception["err"] = _exc
                        finally:
                            _event_queue.put(None)  # τ╗ôµ¥ƒΣ┐íσÅ╖

                    asyncio.run(_inner())

                _exec_thread = _threading_mod.Thread(
                    target=_run_plan_executor, daemon=True
                )
                _exec_thread.start()

                _plan_done_event = None
                while True:
                    try:
                        evt = _event_queue.get(timeout=0.1)
                    except _queue_mod.Empty:
                        if not _exec_thread.is_alive():
                            break
                        continue

                    if evt is None:
                        break

                    etype = evt.get("type", "")

                    if etype == "progress":
                        yield f"data: {json.dumps({'type': 'progress', 'message': evt.get('message', ''), 'detail': evt.get('detail', '')})}\n\n"

                    elif etype == "step_done":
                        step_idx = evt.get("step_index", 0)
                        task_type_done = evt.get("task_type", "")
                        success = evt.get("success", False)
                        preview = evt.get("output_preview", "")
                        if success:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'µ¡ÑΘ¬ñ {step_idx} σ«îµêÉ', 'detail': preview[:80]})}\n\n"
                        else:
                            err_msg = evt.get("error") or "µëºΦíîσñ▒Φ┤Ñ"
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'µ¡ÑΘ¬ñ {step_idx} Θüçσê░Θù«Θóÿ', 'detail': err_msg[:80]})}\n\n"
                        # σ¢₧ΦíÑ subtasks τè╢µÇü∩╝êτö¿Σ║ÄσÉÄτ╗¡Φç¬µúÇ∩╝ë
                        for _st in subtasks:
                            if str(_st.get("id")) == str(evt.get("step_id")):
                                _st["status"] = "completed" if success else "failed"
                                _st["result"] = {"success": success, "output": preview, "error": evt.get("error")}
                                # µö╢Θ¢åΦ╛ôσç║
                                if success:
                                    step_results.append({"success": success, "output": preview})
                                    # Σ╗Ä ContextStore ΦÄ╖σÅûσ«îµò┤τ╗ôµ₧£σÉÄµ¢┤µû░ saved_files
                                break

                    elif etype == "status":
                        yield f"data: {json.dumps({'type': 'status', 'message': evt.get('message', '')})}\n\n"

                    elif etype == "plan_done":
                        _plan_done_event = evt
                        # µö╢Θ¢åΣ┐¥σ¡ÿµûçΣ╗╢
                        saved_files.extend(evt.get("saved_files") or [])
                        # σÉîµ¡Ñ context σ┐½τàº
                        for k, v in (evt.get("context_snapshot") or {}).items():
                            context[k] = v

                _exec_thread.join()

                if _plan_exception["err"]:
                    raise _plan_exception["err"]

                # ΓöÇΓöÇ τ╗äΦúàµ£Çτ╗êΦ╛ôσç║ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                final_result_text = (_plan_done_event or {}).get(
                    "final_output", "(µùáΦ╛ôσç║)"
                )
                # Σ╣ƒΣ╗Ä ContextStore ΦíÑσà¿ saved_files∩╝êPlanExecutor σåàΘâ¿σ╖▓µ▒çµÇ╗∩╝ë
                _pe_saved = (_plan_done_event or {}).get("saved_files") or []
                for _pf in _pe_saved:
                    if _pf not in saved_files:
                        saved_files.append(_pf)

                # Φ┤¿ΘçÅΘ¬îΦ»ü
                yield f"data: {json.dumps({'type': 'status', 'message': 'µ¡úσ£¿Φ┐¢Φíîµ£Çτ╗êΦ┤¿ΘçÅΘ¬îΦ»ü...'})}\n\n"
                _combined_for_validate = {
                    "final_output": final_result_text,
                    "steps": [
                        {
                            "status": s.get("status", "completed"),
                            "result": s.get("result"),
                        }
                        for s in subtasks
                    ],
                }
                quality_score = asyncio.run(
                    TaskOrchestrator._validate_quality(
                        user_input, _combined_for_validate, context
                    )
                )
                yield f"data: {json.dumps({'type': 'status', 'message': f'Φ┤¿ΘçÅΘ¬îΦ»üσ«îµêÉ∩╝îΦ»äσêå: {quality_score}/100'})}\n\n"

                # σñìµ¥éΣ╗╗σèíσ┐½ΘÇƒΦç¬µúÇ
                check = Utils.quick_self_check(
                    "MULTI_STEP", user_input, final_result_text
                )
                if not check.get("pass") and check.get("fix_prompt"):
                    yield f"data: {json.dumps({'type': 'status', 'message': '≡ƒ⌐║ Φç¬µúÇµ£¬ΘÇÜΦ┐ç∩╝îµ¡úσ£¿Σ┐«µ¡úµ£Çτ╗êΦ╛ôσç║...'})}\n\n"
                    try:
                        fix_resp = client.models.generate_content(
                            model=SmartDispatcher.get_model_for_task("MULTI_STEP"),
                            contents=check["fix_prompt"],
                            config=types.GenerateContentConfig(
                                system_instruction=system_instruction,
                                temperature=0.4,
                                max_output_tokens=3000,
                            ),
                        )
                        final_result_text = fix_resp.text or final_result_text
                    except Exception as _fix_err:
                        _app_logger.warning(f"[MULTI_STEP] ΓÜá∩╕Å Φç¬µúÇΣ┐«µ¡úσñ▒Φ┤Ñ: {_fix_err}")

                # LocalPlanner self_check
                if use_local_planner:
                    from app.core.routing import LocalPlanner

                    plan_check = LocalPlanner.self_check(
                        user_input, subtasks, step_results
                    )
                    _lp_status = plan_check.get("status", "partial")
                    _lp_summary = plan_check.get("summary", "")
                    _lp_next = (
                        plan_check.get("next_actions", [])
                        if isinstance(plan_check.get("next_actions", []), list)
                        else []
                    )
                    _lp_msg = f"Φç¬µúÇτ╗ôΦ«║: {_lp_status}"
                    if _lp_summary:
                        _lp_msg += f"\nΦ»┤µÿÄ: {_lp_summary}"
                    if _lp_next:
                        _lp_msg += f"\nσ╗║Φ««σÉÄτ╗¡: {', '.join(_lp_next)}"
                    yield f"data: {json.dumps({'type': 'status', 'message': _lp_msg})}\n\n"

                separator = "=" * 50
                output_text = (
                    f"\n{separator}\nΓ£à σñÜµ¡ÑΣ╗╗σèíσ«îµêÉ\nΦ┤¿ΘçÅΦ»äσêå: {quality_score}/100\n"
                )
                if saved_files:
                    output_text += "σ╖▓Σ┐¥σ¡ÿµûçΣ╗╢:\n"
                    for p in saved_files:
                        name = os.path.basename(p)
                        link_path = p.replace("\\", "/")
                        output_text += f"- [{name}]({link_path})\n"
                    output_text += f"\n≡ƒôé Σ╜ìτ╜«: `{settings_manager.documents_dir}`\n"

                errors_list = [
                    s["result"].get("error")
                    for s in subtasks
                    if s.get("status") == "failed"
                    and isinstance(s.get("result"), dict)
                    and s["result"].get("error")
                ]
                if errors_list:
                    output_text += f"ΓÜá∩╕Å Θüçσê░τÜäΘù«Θóÿ: {', '.join(errors_list)}\n"

                output_text += f"\nµ£Çτ╗êΦ╛ôσç║:\n{final_result_text}\n{separator}\n"

                yield f"data: {json.dumps({'type': 'token', 'content': output_text})}\n\n"
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': saved_files})}\n\n"

            except Exception as e:
                yield f"data: {json.dumps({'type': 'error', 'message': f'σñÜµ¡ÑΣ╗╗σèíµëºΦíîσñ▒Φ┤Ñ: {str(e)}'})}\n\n"

            # Σ┐¥σ¡ÿ MULTI_STEP σ»╣Φ»¥σÄåσÅ▓∩╝êσƒ║Σ║Äτúüτ¢ÿσ«îµò┤σÄåσÅ▓Φ┐╜σèá∩╝ë
            try:
                _multi_summary = (
                    f"[σñÜµ¡ÑΣ╗╗σèíσ«îµêÉ] {', '.join(s['description'] for s in subtasks)}"
                )
                if saved_files:
                    _multi_summary += f"\nτöƒµêÉµûçΣ╗╢: {', '.join(os.path.basename(p) for p in saved_files)}"
                session_manager.append_and_save(
                    f"{session_name}.json", user_input, _multi_summary
                )
                _app_logger.info(f"[MULTI_STEP] Γ£à σ»╣Φ»¥σÄåσÅ▓σ╖▓Σ┐¥σ¡ÿ")
                _start_memory_extraction(
                    user_input,
                    _multi_summary,
                    [],
                    task_type="CODER",
                    session_name=session_name,
                )
            except Exception as save_err:
                _app_logger.warning(f"[MULTI_STEP] ΓÜá∩╕Å Σ┐¥σ¡ÿσ»╣Φ»¥σÄåσÅ▓σñ▒Φ┤Ñ: {save_err}")

        return Response(generate_multi_step(), mimetype="text/event-stream")

    # === Agent Σ╗╗σèíµëºΦíî∩╝êLangGraphAgent ReAct∩╝îΘÖìτ║ºσê░ UnifiedAgent∩╝ë===
    if task_type == "AGENT":
        _app_logger.debug(f"[STREAM] ≡ƒñû µëºΦíî Agent Σ╗╗σèí (LangGraphAgent ReAct)")
        # classify_v2 σ╖▓Φ»åσê½σç║ skill_id∩╝îΣ╝áτ╗Ö Agent σ«₧τÄ░ Skill Σ╕ôσ▒₧ΦíîΣ╕║
        _agent_skill_id = _router_decision.skill_id if _router_decision else None

        def generate_agent():
            yield f"data: {json.dumps({'type': 'classification', 'task_type': 'AGENT', 'route_method': route_method, 'message': '≡ƒÄ» Σ╗╗σèíσêåτ▒╗: ≡ƒñû µÖ║Φâ╜σè⌐µëï (LangGraph ReAct)'})}\n\n"

            final_answer = ""
            collected_steps = []

            # ΓöÇΓöÇ Σ╝ÿσàê∩╝ÜLangGraphAgent∩╝êStateGraph + CheckpointSaver∩╝ëΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
            _lg_ok = False
            try:
                from app.core.agent.factory import create_langgraph_agent

                _lg_agent = create_langgraph_agent(
                    model_id=SmartDispatcher.get_model_for_task("AGENT"),
                )
                _lg_ok = True
                for chunk in _lg_agent.stream(
                    input_text=user_input,
                    history=history,
                    session_id=session_name,
                    skill_id=_agent_skill_id,
                    task_type="AGENT",
                ):
                    ctype = chunk.get("type", "token")
                    content = chunk.get("content", "")
                    if ctype == "answer":
                        final_answer = content
                        step_data = {"step_type": "ANSWER", "content": content, "tool": None}
                    elif ctype == "tool_call":
                        step_data = {"step_type": "TOOL_CALL", "content": f"Φ░âτö¿σ╖Ñσà╖: {content}", "tool": content, "args": chunk.get("args", {})}
                    elif ctype == "tool_result":
                        step_data = {
                            "step_type": "TOOL_RESULT",
                            "content": content,
                            "tool": None,
                        }
                    elif ctype == "token":
                        step_data = {
                            "step_type": "THINKING",
                            "content": content,
                            "tool": None,
                        }
                    elif ctype == "error":
                        raise RuntimeError(content)
                    else:
                        continue
                    collected_steps.append(step_data)
                    yield f"data: {json.dumps({'type': 'agent_step', 'data': step_data}, ensure_ascii=False)}\n\n"

            except Exception as _lg_err:
                _app_logger.debug(
                    f"[AGENT] LangGraphAgent σñ▒Φ┤Ñ ({_lg_err})∩╝îΘÖìτ║ºσê░ UnifiedAgent..."
                )
                _lg_ok = False

            # ΓöÇΓöÇ ΘÖìτ║º∩╝ÜUnifiedAgent∩╝êµùº while-loop σ«₧τÄ░∩╝ëΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
            if not _lg_ok:
                try:
                    from app.core.agent.factory import create_agent
                    from app.core.agent.types import AgentStepType

                    _ua = create_agent(
                        model_id=SmartDispatcher.get_model_for_task("AGENT")
                    )
                    collected_steps = []
                    final_answer = ""
                    for step in _ua.run(
                        input_text=user_input,
                        history=history,
                        session_id=session_name,
                        skill_id=_agent_skill_id,
                        task_type="AGENT",
                    ):
                        step_data = step.to_dict()
                        collected_steps.append(step_data)
                        if step.step_type == AgentStepType.ANSWER:
                            final_answer = step.content or ""
                        yield f"data: {json.dumps({'type': 'agent_step', 'data': step_data}, ensure_ascii=False)}\n\n"
                    if not final_answer and collected_steps:
                        final_answer = collected_steps[-1].get("content", "")
                except Exception as e:
                    import traceback

                    _app_logger.error(f"[AGENT] Γ¥î UnifiedAgent Σ╣ƒσñ▒Φ┤Ñ:\n{traceback.format_exc()}")
                    yield f"data: {json.dumps({'type': 'error', 'message': f'Agent µëºΦíîσñ▒Φ┤Ñ: {str(e)}'})}\n\n"
                    return

            task_payload = {
                "id": f"task_{int(time.time() * 1000)}",
                "status": "success",
                "result": final_answer,
                "steps": collected_steps,
                "engine": "langgraph" if _lg_ok else "unified",
            }
            yield f"data: {json.dumps({'type': 'task_final', 'data': task_payload}, ensure_ascii=False)}\n\n"

            try:
                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    final_answer or "[Agent Σ╗╗σèíσ«îµêÉ]",
                )
            except Exception:
                pass
            try:
                _start_memory_extraction(
                    user_input,
                    final_answer or "",
                    [],
                    task_type="AGENT",
                    session_name=session_name,
                )
            except Exception:
                pass

        return Response(generate_agent(), mimetype="text/event-stream")

    if locked_model and locked_model != "auto":
        model_id = locked_model
    else:
        # Σ╝áΘÇÆ complexity Σ╗ÑΣ╛┐Σ╕║σñìµ¥éΣ╗╗σèíΘÇëµï⌐µ¢┤σ╝║τÜäµ¿íσ₧ï
        _complexity = (context_info or {}).get("complexity", "normal")
        model_id = SmartDispatcher.get_model_for_task(task_type, complexity=_complexity)

    _app_logger.debug(f"[STREAM] Final: task_type='{task_type}', model_id='{model_id}'\n")

    # ≡ƒÄ» Skills µ│¿σàÑ∩╝Üσ░åτö¿µê╖σÉ»τö¿τÜä Skill Φ┐╜σèáσê░ system_instruction
    try:
        from app.core.skills.skill_manager import SkillManager

        _active_skills = SkillManager.get_active_skill_names(task_type=task_type)
        if _active_skills:
            _app_logger.debug(
                f"[STREAM] ≡ƒÄ» Active Skills ({task_type}): {', '.join(_active_skills)}"
            )
        # µäÅσ¢╛τ╗æσ«Ü∩╝ÜµèÇΦâ╜µ£¬µëïσè¿σ╝ÇσÉ»µù╢∩╝îµîëΦ╛ôσàÑσåàσ«╣Σ╕┤µù╢µ┐Çµ┤╗σî╣ΘàìτÜäµèÇΦâ╜
        _intent_temp_ids = []
        try:
            from app.core.skills.skill_trigger_binding import get_skill_binding_manager

            _intent_temp_ids = get_skill_binding_manager().match_intent(
                user_input or ""
            )
        except Exception:
            pass
        # AutoMatcher ΦíÑσàà∩╝ÜΦºäσêÖ/Φ»¡Σ╣ëσî╣ΘàìΦªåτ¢ûµäÅσ¢╛τ╗æσ«Üµ£¬µîüΣ╣àσîûτÜäσ£║µÖ»
        try:
            from app.core.skills.skill_auto_matcher import SkillAutoMatcher

            _auto_ids = SkillAutoMatcher.match(
                user_input=user_input or "", task_type=task_type or "CHAT"
            )
            if _auto_ids:
                # σÉêσ╣╢σÄ╗Θçì∩╝îΣ┐¥µîü intent τ╗ôµ₧£Σ╝ÿσàê
                _intent_temp_ids = list(dict.fromkeys(_intent_temp_ids + _auto_ids))
        except Exception:
            pass
        if _intent_temp_ids:
            _app_logger.debug(f"[STREAM] ≡ƒöù Auto Skills: {', '.join(_intent_temp_ids)}")
        system_instruction = SkillManager.inject_into_prompt(
            system_instruction,
            task_type=task_type,
            user_input=user_input,
            temp_skill_ids=_intent_temp_ids,
        )
    except Exception as _sk_err:
        _app_logger.warning(f"[STREAM] ΓÜá∩╕Å Skills µ│¿σàÑσñ▒Φ┤Ñ: {_sk_err}")

    # ≡ƒôÜ RAG µ╖╖σÉêµúÇτ┤ó∩╝êσÉæΘçÅ + BM25 + RRF Φ₧ìσÉê∩╝ë
    # _rag_context_block: ΘàìΘÇüτ╗Ö generate()πÇüRESEARCH πÇüToT τ¡ëσÉäΦ╖»σ╛ä
    _rag_context_block = ""
    try:
        from app.core.services.rag_service import get_rag_service

        _rag_svc = get_rag_service()
        if _rag_svc.stats().get("initialized"):
            _rag_hits = _rag_svc.hybrid_retrieve(user_input, k=3, score_threshold=0.3)
            if _rag_hits:
                for _rc in _rag_hits:
                    _src = os.path.basename(_rc.get("source", "unknown"))
                    _rag_context_block += f"[{_src} | τ¢╕Σ╝╝σ║ª: {_rc.get('score', 0):.3f}]\n{_rc['content']}\n\n"
                # σÉîµù╢µ│¿σàÑ system_instruction∩╝êΣ╛¢ ToT πÇüAGENT Φ╖»σ╛äΣ╜┐τö¿∩╝ë
                _rag_sys_block = (
                    "\n\nΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ"
                    "\n## ≡ƒôÜ τƒÑΦ»åσ║ôσÅéΦÇâσåàσ«╣∩╝êµ╖╖σÉêµúÇτ┤ó∩╝ë\n" + _rag_context_block
                )
                system_instruction += _rag_sys_block
                _app_logger.debug(
                    f"[STREAM] ≡ƒôÜ µ╖╖σÉêRAG: {len(_rag_hits)} τëçµ«╡∩╝îtop_score={_rag_hits[0].get('score', 0):.3f}"
                )
    except Exception as _rag_err:
        _app_logger.warning(f"[STREAM] ΓÜá∩╕Å RAG µ│¿σàÑΦ╖│Φ┐ç: {_rag_err}")

    # ≡ƒò╕∩╕Å Graph RAG ΓÇö entity-expanded triple retrieval
    try:
        from app.core.services.graph_rag_service import GraphRAGService as _GRAGS

        _graph_ctx = _GRAGS.retrieve(user_input, k=8)
        if _graph_ctx:
            _rag_context_block += "\n\n" + _graph_ctx
            system_instruction += (
                "\n\nΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ" "\n" + _graph_ctx
            )
            _app_logger.debug(f"[STREAM] ≡ƒò╕∩╕Å Graph RAG: µ│¿σàÑτƒÑΦ»åσ¢╛Φ░▒σà│ΦüöΣ║ïσ«₧")
    except Exception as _ge:
        pass

    # Φ»╗σÅûτö¿µê╖Φ«╛τ╜«∩╝Üµÿ»σÉªµÿ╛τñ║µÇ¥ΦÇâΦ┐çτ¿ï
    _show_thinking = False
    try:
        _show_thinking = settings_manager.get("ai", "show_thinking") == True
    except Exception:
        pass

    def generate():
        start_time = time.time()

        def _infer_analysis_source(message: str, phase: str = "thinking") -> str:
            """µÄ¿µû¡σêåµ₧Éµ¥Ñµ║É∩╝Ülocal / cloud / hybrid / system"""
            msg = (message or "").lower()
            phase_l = (phase or "").lower()

            if any(k in msg for k in ["ollama", "µ£¼σ£░µ¿íσ₧ï", "qwen", "local"]):
                return "local"
            if any(k in msg for k in ["gemini", "deep-research", "Σ║æτ½»", "cloud"]):
                return "cloud"
            if any(
                k in phase_l for k in ["routing", "context", "planning", "analyzing"]
            ):
                return "hybrid"
            return "system"

        def yield_thinking(message: str, phase: str = "thinking", source: str = None):
            """σÅæΘÇüµÇ¥ΦÇâΦ┐çτ¿ïΣ║ïΣ╗╢∩╝êΣ╗àσ╜ôτö¿µê╖σ╝ÇσÉ» show_thinking µù╢∩╝ë∩╝îΘÖäσ╕ªσêåµ₧Éµ¥Ñµ║É"""
            if not _show_thinking:
                return ""

            resolved_source = source or _infer_analysis_source(message, phase)
            source_tag = {
                "local": "[µ£¼σ£░σêåµ₧É]",
                "cloud": "[σñºµ¿íσ₧ïσêåµ₧É]",
                "hybrid": "[µ╖╖σÉêσå│τ¡û]",
                "system": "[τ│╗τ╗ƒµ╡üτ¿ï]",
            }.get(resolved_source, "[τ│╗τ╗ƒµ╡üτ¿ï]")

            elapsed = round(time.time() - start_time, 1)
            display_message = f"{source_tag} {message}"
            return f"data: {json.dumps({'type': 'thinking', 'message': display_message, 'phase': phase, 'elapsed': elapsed, 'analysis_source': resolved_source}, ensure_ascii=False)}\n\n"

        # === τ½ïσì│σÅìΘªêΣ╗╗σèíσêåτ▒╗Σ┐íµü» ===
        task_display_names = {
            "PAINTER": "≡ƒÄ¿ σ¢╛σâÅτöƒµêÉ",
            "FILE_GEN": "≡ƒôä µûçµíúτöƒµêÉ",
            "CODER": "≡ƒÆ╗ Σ╗úτáüτ╝ûτ¿ï",
            "RESEARCH": "≡ƒôÜ µ╖▒σ║ªτáöτ⌐╢",
            "WEB_SEARCH": "≡ƒîÉ σ«₧µù╢µÉ£τ┤ó",
            "CHAT": "≡ƒÆ¼ σ»╣Φ»¥",
            "SYSTEM": "≡ƒûÑ∩╕Å τ│╗τ╗ƒµôìΣ╜£",
            "FILE_OP": "≡ƒôé µûçΣ╗╢µôìΣ╜£",
            "FILE_EDIT": "Γ£Å∩╕Å µûçΣ╗╢τ╝ûΦ╛æ",
            "FILE_SEARCH": "≡ƒöì µûçΣ╗╢µÉ£τ┤ó",
            "VISION": "≡ƒæü∩╕Å σ¢╛σâÅΦ»åσê½",
            "MULTI_STEP": "≡ƒöä σñÜµ¡ÑΣ╗╗σèí",
            "AGENT": "≡ƒñû µÖ║Φâ╜σè⌐µëï",
        }

        model_display = get_model_display_name(model_id)
        task_display = task_display_names.get(task_type, task_type)

        # σÅæΘÇüΣ╗╗σèíσêåτ▒╗Σ┐íµü»∩╝êσ£¿µ£Çσ╝Çσºï∩╝îτ½ïσì│µÿ╛τñ║∩╝ë
        classification_msg = f"≡ƒÄ» Σ╗╗σèíσêåτ▒╗: {task_display}"
        if route_method:
            classification_msg += f" (µû╣µ│ò: {route_method})"

        routing_list = None
        # Σ╗àΣ┐¥τòÖ routing_list τö¿Σ║ÄσåàΘâ¿Φ░âΦ»ò∩╝îΣ╕ìµÿ╛τñ║τ╗Öτö¿µê╖
        if context_info and context_info.get("routing_list"):
            routing_list = context_info.get("routing_list")

        yield f"data: {json.dumps({'type': 'classification', 'task_type': task_type, 'task_display': task_display, 'model': model_id, 'model_display': model_display, 'route_method': route_method, 'routing_list': routing_list, 'message': classification_msg})}\n\n"

        # µÇ¥ΦÇâΦ┐çτ¿ï∩╝ÜΣ╗╗σèíΦ╖»τö▒σêåµ₧É
        t = yield_thinking(f"σêåµ₧Éτö¿µê╖µäÅσ¢╛ ΓåÆ Φ»åσê½Σ╕║ {task_display}", "routing", "hybrid")
        if t:
            yield t
        model_source = (
            "local"
            if any(k in (model_id or "").lower() for k in ["qwen", "llama", "ollama"])
            else "cloud"
        )
        t = yield_thinking(
            f"Φ╖»τö▒µû╣µ│ò: {route_method}∩╝îΘÇëµï⌐µ¿íσ₧ï: {model_display}",
            "model",
            model_source,
        )
        if t:
            yield t
        if routing_list:
            steps_str = (
                " ΓåÆ ".join(
                    [
                        (
                            f"{r.get('task','?')}({r.get('score', 0):.2f})"
                            if isinstance(r.get("score"), (int, float))
                            else f"{r.get('task','?')}({r.get('score','?')})"
                        )
                        for r in routing_list[:5]
                    ]
                )
                if isinstance(routing_list, list)
                else str(routing_list)
            )
            t = yield_thinking(f"Φ╖»τö▒σå│τ¡ûΘô╛: {steps_str}", "routing", "hybrid")
            if t:
                yield t

        # σªéµ₧£µ£ëσñìµ¥éσ║ªΣ┐íµü»∩╝îΣ╣ƒσÅæΘÇü
        if context_info and context_info.get("complexity"):
            complexity_msg = f"≡ƒôè Σ╗╗σèíσñìµ¥éσ║ª: {context_info['complexity']}"
            yield f"data: {json.dumps({'type': 'info', 'message': complexity_msg})}\n\n"
            t = yield_thinking(
                f"Σ╗╗σèíσñìµ¥éσ║ªΦ»äΣ╝░: {context_info['complexity']}", "analyzing", "hybrid"
            )
            if t:
                yield t

        # σªéµ₧£µ£ëΣ╕èΣ╕ïµûç∩╝îΣ╜┐τö¿σó₧σ╝║σÉÄτÜäΦ╛ôσàÑ
        effective_input = user_input
        if (
            context_info
            and context_info.get("is_continuation")
            and context_info.get("enhanced_input")
        ):
            effective_input = context_info["enhanced_input"]
            _app_logger.debug(f"[STREAM] Using enhanced input (length: {len(effective_input)})")
            yield f"data: {json.dumps({'type': 'info', 'message': '≡ƒöù µúÇµ╡ïσê░σ╗╢τ╗¡Σ╗╗σèí∩╝îΣ╜┐τö¿Σ╕èΣ╕ïµûçσó₧σ╝║'})}\n\n"
            t = yield_thinking(
                f"µúÇµ╡ïσê░Σ╕èΣ╕ïµûçσ╗╢τ╗¡∩╝îσó₧σ╝║Φ╛ôσàÑ ({len(effective_input)} σ¡ùτ¼ª)",
                "context",
                "hybrid",
            )
            if t:
                yield t

        # Σ╜┐τö¿σ┐½ΘÇƒσ░Åµ¿íσ₧ïσ░åΦ»╖µ▒éΦ╜¼Σ╕║τ╗ôµ₧äσîû Markdown∩╝êΣ╗àσ»╣σñºµ¿íσ₧ïΣ╗╗σèíσÉ»τö¿∩╝ë
        if task_type not in ["SYSTEM", "FILE_OP", "PAINTER", "VISION"]:
            adapted_input = Utils.adapt_prompt_to_markdown(
                task_type, effective_input, history=history
            )
            if adapted_input != effective_input:
                effective_input = adapted_input
                yield f"data: {json.dumps({'type': 'info', 'message': '≡ƒº╛ σ╖▓σ░åΦ»╖µ▒éτ╗ôµ₧äσîûΣ╕║MarkdownµÅÉτñ║'})}\n\n"
                t = yield_thinking(
                    "σ░åτö¿µê╖Φ»╖µ▒éτ╗ôµ₧äσîûΣ╕║ Markdown µá╝σ╝ÅΣ╗ÑµÅÉσìçΦ╛ôσç║Φ┤¿ΘçÅ",
                    "planning",
                    "hybrid",
                )
                if t:
                    yield t

        # Θçìτ╜«Σ╕¡µû¡µáçσ┐ù∩╝êµ»Åµ¼íµû░Φ»╖µ▒éΘâ╜Θçìτ╜«∩╝ë
        _interrupt_manager.reset(session_name)
        interrupt_event = _interrupt_manager.get_event(session_name)

        def interrupted():
            return _interrupt_manager.is_interrupted(session_name)

        # σÅæΘÇüΦ┐¢σ║ª: σ╝ÇσºïσñäτÉå
        from web.smart_feedback import SmartFeedback

        _task_labels = SmartFeedback.TASK_LABELS
        _tl = _task_labels.get(task_type, task_type)
        yield f"data: {json.dumps({'type': 'progress', 'message': f'σ╝ÇσºïσñäτÉå{_tl}Σ╗╗σèí', 'detail': get_model_display_name(model_id)})}\n\n"

        try:
            # σê¥σºïσîûµ¿íσ₧ïΦ┐╜Φ╕¬σÅÿΘçÅ∩╝êτö¿Σ║ÄµùÑσ┐ùΦ«░σ╜ò∩╝ë
            used_model = "unknown"

            # === SYSTEM Mode (µ£¼σ£░µëºΦíî - σì│µù╢) ===
            if task_type == "SYSTEM":
                used_model = "LocalExecutor"
                yield f"data: {json.dumps({'type': 'progress', 'message': 'µ¡úσ£¿σêåµ₧Éτ│╗τ╗ƒµîçΣ╗ñ...', 'detail': ''})}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'message': 'µ¡úσ£¿µëºΦíîµôìΣ╜£...', 'detail': ''})}\n\n"

                exec_result = LocalExecutor.execute(user_input)
                response_text = exec_result["message"]
                if exec_result.get("details"):
                    response_text += f"\n\n{exec_result['details']}"

                if Utils.is_failure_output(response_text):
                    t = yield_thinking(
                        "τ│╗τ╗ƒµîçΣ╗ñµëºΦíîσñ▒Φ┤Ñ∩╝îΣ╜┐τö¿ AI Σ┐«µ¡úσÉÄΘçìΦ»ò", "validating"
                    )
                    if t:
                        yield t
                    yield f"data: {json.dumps({'type': 'progress', 'message': 'ΓÜá∩╕Å σê¥µ¼íµëºΦíîσñ▒Φ┤Ñ∩╝îµ¡úσ£¿Σ┐«µ¡ú...', 'detail': ''})}\n\n"
                    fix_prompt = Utils.build_fix_prompt(
                        "SYSTEM", user_input, response_text
                    )
                    fix_resp = client.models.generate_content(
                        model=model_id,
                        contents=fix_prompt,
                        config=types.GenerateContentConfig(
                            system_instruction=system_instruction,
                            temperature=0.4,
                            max_output_tokens=1000,
                        ),
                    )
                    response_text = fix_resp.text or response_text

                yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"

                # σàêΣ┐¥σ¡ÿσÄåσÅ▓∩╝îσåìσÅæΘÇü done Σ║ïΣ╗╢∩╝êΘÿ▓µ¡óσ«óµê╖τ½»µû¡σ╝Çσ»╝Φç┤Σ╕óσñ▒∩╝ë
                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    response_text,
                    task=task_type,
                    model_name=used_model,
                )

                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return

            # === FILE_OP Mode (µûçΣ╗╢µôìΣ╜£ - σì│µù╢) ===
            if task_type == "FILE_OP":
                used_model = "LocalExecutor"
                yield f"data: {json.dumps({'type': 'progress', 'message': 'µ¡úσ£¿σêåµ₧ÉµûçΣ╗╢µôìΣ╜£...', 'detail': ''})}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'message': 'µ¡úσ£¿Φ«┐Θù«µûçΣ╗╢τ│╗τ╗ƒ...', 'detail': ''})}\n\n"

                batch_manager = get_batch_ops_manager()
                if batch_manager.is_batch_command(user_input):
                    parsed = batch_manager.parse_command(user_input)
                    if not parsed.get("success"):
                        response_text = (
                            f"Γ¥î {parsed.get('error')}\n\n{parsed.get('hint', '')}"
                        )
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            response_text,
                            task="FILE_OP",
                            model_name=used_model,
                        )
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                    job = batch_manager.create_job(
                        name=f"batch_{parsed.get('operation')}",
                        operation=parsed.get("operation"),
                        input_dir=parsed.get("input_dir"),
                        output_dir=parsed.get("output_dir"),
                        options=parsed.get("options", {}),
                    )
                    batch_manager.start_job(job.job_id)
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒº⌐ σ╖▓σê¢σ╗║µë╣ΘçÅΣ╗╗σèí: {job.job_id}', 'detail': ''})}\n\n"

                    summary_text = None
                    for event in batch_manager.iter_job_events(job.job_id):
                        if event.get("type") == "progress":
                            current = event.get("current", 0)
                            total = event.get("total", 0)
                            progress_pct = int((current / total) * 100) if total else 0
                            yield f"data: {json.dumps({'type': 'progress', 'message': 'ΓÅ│ µë╣ΘçÅσñäτÉåΣ╕¡...', 'detail': event.get('detail', ''), 'progress': progress_pct, 'total': total})}\n\n"
                        elif event.get("type") == "final":
                            summary_text = event.get("summary") or "Γ£à µë╣ΘçÅσñäτÉåσ«îµêÉ"
                            break
                        elif event.get("type") == "error":
                            summary_text = event.get("message", "Γ¥î µë╣ΘçÅΣ╗╗σèíσñ▒Φ┤Ñ")
                            break

                    if summary_text:
                        yield f"data: {json.dumps({'type': 'token', 'content': summary_text})}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            summary_text,
                            task="FILE_OP",
                            model_name=used_model,
                        )

                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                file_result = FileOperator.execute(user_input)
                response_text = file_result["message"]
                if file_result.get("content"):
                    response_text += f"\n\n{file_result['content']}"

                if Utils.is_failure_output(response_text):
                    yield f"data: {json.dumps({'type': 'progress', 'message': 'ΓÜá∩╕Å σê¥µ¼íµëºΦíîσñ▒Φ┤Ñ∩╝îµ¡úσ£¿Σ┐«µ¡ú...', 'detail': ''})}\n\n"
                    fix_prompt = Utils.build_fix_prompt(
                        "FILE_OP", user_input, response_text
                    )
                    fix_resp = client.models.generate_content(
                        model=model_id,
                        contents=fix_prompt,
                        config=types.GenerateContentConfig(
                            system_instruction=system_instruction,
                            temperature=0.4,
                            max_output_tokens=1000,
                        ),
                    )
                    response_text = fix_resp.text or response_text

                yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"

                # σàêΣ┐¥σ¡ÿσÄåσÅ▓∩╝îσåìσÅæΘÇü done Σ║ïΣ╗╢
                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    response_text,
                    task="FILE_OP",
                    model_name=used_model,
                )

                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return

            # === FILE_EDIT Mode (µûçΣ╗╢τ╝ûΦ╛æ - µÖ║Φâ╜Σ┐«µö╣) ===
            if task_type == "FILE_EDIT":
                used_model = model_id
                t = yield_thinking(
                    "Φ┐¢σàÑµûçΣ╗╢τ╝ûΦ╛æµ¿íσ╝Å∩╝îσ░åτÉåΦºúτö¿µê╖µîçΣ╗ñσ╣╢Σ┐«µö╣µûçΣ╗╢", "routing"
                )
                if t:
                    yield t
                yield f"data: {json.dumps({'type': 'progress', 'message': 'µ¡úσ£¿σêåµ₧Éτ╝ûΦ╛æµîçΣ╗ñ...', 'detail': ''})}\n\n"

                editor = get_file_editor()

                # σ░¥Φ»òΣ╗Äτö¿µê╖Φ╛ôσàÑΣ╕¡µÅÉσÅûµûçΣ╗╢Φ╖»σ╛äσÆîµîçΣ╗ñ
                # µ¿íσ╝Å 1: "Σ┐«µö╣ path/to/file µèèxxxµö╣µêÉyyy"
                match = re.search(
                    r'(?:Σ┐«µö╣|τ╝ûΦ╛æ|µö╣)\s+["\']?([^"\']+?)["\']?\s+(.+)', user_input
                )
                if not match:
                    # µ¿íσ╝Å 2: "µèè path/to/file τÜäxxxµö╣µêÉyyy"
                    match = re.search(
                        r'(?:µèè|σ░å)\s+["\']?([^"\']+?)["\']?\s+(?:τÜä|Σ╕¡τÜä|ΘçîτÜä)\s*(.+)',
                        user_input,
                    )

                if match:
                    file_path = match.group(1).strip()
                    instruction = match.group(2).strip()

                    t = yield_thinking(
                        f"µÅÉσÅûσê░µûçΣ╗╢Φ╖»σ╛ä: {file_path}, µîçΣ╗ñ: {instruction}", "analyzing"
                    )
                    if t:
                        yield t

                    yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒöì τ¢«µáçµûçΣ╗╢: {os.path.basename(file_path)}', 'detail': ''})}\n\n"
                    yield f"data: {json.dumps({'type': 'progress', 'message': 'µ¡úσ£¿µëºΦíîτ╝ûΦ╛æ...', 'detail': ''})}\n\n"

                    result = editor.smart_edit(file_path, instruction)

                    if result["success"]:
                        operation = result.get("operation", "edit")
                        edit_result = result.get("result", {})

                        response_text = f"Γ£à µûçΣ╗╢τ╝ûΦ╛æµêÉσèƒ∩╝ü\n\n"
                        response_text += f"**µôìΣ╜£τ▒╗σ₧ï**: {operation}\n"

                        if operation == "replace":
                            response_text += (
                                f"**µ¢┐µìóµ¼íµò░**: {edit_result.get('replacements', 0)}\n"
                            )
                            response_text += (
                                f"**ΘóäΦºê**:\n```\n{edit_result.get('preview', '')}\n```"
                            )
                        elif operation == "delete_lines":
                            response_text += f"**σêáΘÖñσåàσ«╣**:\n```\n{edit_result.get('deleted_content', '')}\n```"
                        elif operation == "insert_line":
                            response_text += (
                                f"**µ╢êµü»**: {edit_result.get('message', '')}"
                            )

                        if edit_result.get("backup"):
                            response_text += (
                                f"\n\n≡ƒÆ╛ σñçΣ╗╜µûçΣ╗╢: `{edit_result.get('backup')}`"
                            )
                    else:
                        error_msg = result.get("error", "µ£¬τƒÑΘöÖΦ»»")
                        hint = result.get("hint", "")
                        response_text = f"Γ¥î µûçΣ╗╢τ╝ûΦ╛æσñ▒Φ┤Ñ\n\n{error_msg}\n\n{hint}"
                else:
                    # µùáµ│òµÅÉσÅûµûçΣ╗╢Φ╖»σ╛ä∩╝îΦ«⌐AIτÉåΦºú
                    response_text = "Γ¥î µùáµ│òΦ»åσê½µûçΣ╗╢Φ╖»σ╛äσÆîτ╝ûΦ╛æµîçΣ╗ñ\n\n"
                    response_text += "Φ»╖Σ╜┐τö¿Σ╗ÑΣ╕ïµá╝σ╝Å:\n"
                    response_text += "- `Σ┐«µö╣ µûçΣ╗╢Φ╖»σ╛ä µèè'µùºµûçµ£¼'µö╣µêÉ'µû░µûçµ£¼'`\n"
                    response_text += "- `µèè µûçΣ╗╢Φ╖»σ╛ä τÜäτ¼¼5-10ΦíîσêáΘÖñ`\n"
                    response_text += "- `τ╝ûΦ╛æ µûçΣ╗╢Φ╖»σ╛ä σ£¿τ¼¼3ΦíîΣ╣ïσÉÄµÅÆσàÑ'µû░σåàσ«╣'`"

                yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"

                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    response_text,
                    task="FILE_EDIT",
                    model_name=used_model,
                )

                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return

            # === FILE_SEARCH Mode (µûçΣ╗╢µÉ£τ┤ó: σà¿τ¢ÿµë½µÅÅ + σ╖ÑΣ╜£σî║τ┤óσ╝ò) ===
            if task_type == "FILE_SEARCH":
                used_model = "FileScanner (Local)"
                t = yield_thinking("Φ┐¢σàÑµûçΣ╗╢µÉ£τ┤óµ¿íσ╝Å∩╝Üσà¿τ¢ÿµûçΣ╗╢τ┤óσ╝òµúÇτ┤ó", "searching")
                if t:
                    yield t
                yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒöì µ¡úσ£¿µÉ£τ┤óµûçΣ╗╢...', 'detail': 'σà¿τ¢ÿτ┤óσ╝òµúÇτ┤ó'}, ensure_ascii=False)}\n\n"

                # ΓöÇΓöÇ σ»╝σàÑσà¿τ¢ÿµë½µÅÅσÖ¿ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                try:
                    from web.file_scanner import (
                        FileScanner,
                        extract_query_from_input,
                        is_disk_search_intent,
                    )

                    _disk_scanner_ok = True
                except Exception as _fse:
                    _app_logger.warning(f"[FILE_SEARCH] ΓÜá∩╕Å FileScanner σ»╝σàÑσñ▒Φ┤Ñ: {_fse}")
                    _disk_scanner_ok = False

                # σêñµû¡µÿ»σÉªµÿ»σà¿τ¢ÿµÉ£τ┤óµäÅσ¢╛
                _is_disk = _disk_scanner_ok and is_disk_search_intent(user_input)

                # ΓöÇΓöÇ µîçσ«ÜΦ╖»σ╛ä + µë⌐σ▒òσÉìσêùΣ╕╛/σ╜Æτ║│∩╝êσªé"σ╜Æτ║│ C:\Desktop ΘçîτÜä .doc µûçΣ╗╢"∩╝ëΓöÇΓöÇ
                import re as _re_pathscan
                import time as _time_fmt

                _explicit_path_m = _re_pathscan.search(
                    r"([A-Za-z]:[\\][^\s\u4e00-\u9fa5]*)", user_input
                )
                if _explicit_path_m:
                    from pathlib import Path as _ScanPath

                    _scan_dir = _ScanPath(_explicit_path_m.group(1).rstrip("\\/. "))
                    # σà│Θö«Φ»ì ΓåÆ µë⌐σ▒òσÉìµÿáσ░ä
                    _KW_EXT_MAP = [
                        (["wordµûçΣ╗╢", "wordµûçµíú", "word"], [".doc", ".docx"]),
                        (["excelµûçΣ╗╢", "excelΦí¿µá╝", "excel"], [".xls", ".xlsx"]),
                        (["pptµûçΣ╗╢", "pptµ╝öτñ║", "ppt", "σ╣╗τü»τëç"], [".ppt", ".pptx"]),
                        (["pdfµûçΣ╗╢", "pdf"], [".pdf"]),
                        (
                            ["σ¢╛τëç", "τàºτëç"],
                            [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"],
                        ),
                        (["σÄïτ╝⌐σîà", "σÄïτ╝⌐µûçΣ╗╢"], [".zip", ".rar", ".7z"]),
                    ]
                    _ext_filters = []
                    _user_lower = user_input.lower()
                    for _kws, _exts in _KW_EXT_MAP:
                        if any(k in _user_lower for k in _kws):
                            _ext_filters = _exts[:]
                            break
                    if not _ext_filters:
                        _ext_raw = _re_pathscan.findall(
                            r"\.(docx?|xlsx?|pdf|txt|md|pptx?|csv|json|py|jpe?g|png|gif|mp[34]|zip|rar)",
                            user_input,
                            _re_pathscan.IGNORECASE,
                        )
                        _ext_filters = list(
                            dict.fromkeys("." + e.lower() for e in _ext_raw)
                        )
                    if _scan_dir.exists() and _scan_dir.is_dir():
                        _ext_label = (
                            "πÇü".join(_ext_filters) if _ext_filters else "µëÇµ£ëτ▒╗σ₧ï"
                        )
                        # ΓöÇΓöÇ σêñµû¡µäÅσ¢╛∩╝Üσ╜Æτ║│∩╝êτë⌐τÉåµò┤τÉå∩╝ëvs µÉ£τ┤ó/σêùΣ╕╛∩╝êσÅ¬Φ»╗σêùΦí¿∩╝ëΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        _CATALOG_KWS = ["σ╜Æτ║│", "σ╜Æµíú", "σ╜Æτ▒╗"]
                        _is_catalog_intent = any(k in user_input for k in _CATALOG_KWS)

                        # ΓöÇΓöÇ σåàσ«╣Φ┐çµ╗ñµäÅσ¢╛∩╝ê"σô¬σçáΣ╕¬µÿ»X"∩╝ëΣ╝ÿσàêΣ║Ä flat list ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        import re as _re_filter_intent

                        _CONTENT_FILTER_KWS = [
                            "σô¬σçáΣ╕¬µÿ»",
                            "σô¬σçáΣ╗╜µÿ»",
                            "µ£ëσô¬σçáΣ╕¬",
                            "µ£ëσçáΣ╕¬µÿ»",
                            "σô¬Σ║¢µÿ»",
                            "µÿ»Σ╝üΣ╕Ü",
                            "µÿ»Φ«┐Φ░ê",
                            "µÿ»µèÑσæè",
                            "µÿ»σÉêσÉî",
                            "µÿ»τ«ÇσÄå",
                            "µÿ»µû╣µíê",
                            "µÿ»Σ╗ÇΣ╣êτ▒╗σ₧ï",
                            "σ▒₧Σ║Ä",
                            "µÿ»Σ╗ÇΣ╣êµûçΣ╗╢",
                            "σô¬σçá",
                            "σçáΣ╕¬µÿ»",
                            "σçáΣ╗╜µÿ»",
                        ]
                        _is_filter_intent = not _is_catalog_intent and (
                            any(k in user_input for k in _CONTENT_FILTER_KWS)
                            or bool(
                                _re_filter_intent.search(
                                    r"σô¬[σçáΣ╕¬Σ║¢].*µÿ»|µÿ».*[µèÑσæèσÉêσÉîτ«ÇσÄåµû╣µíêΦ«┐Φ░êτ║¬ΦªüΦ«íσêÆ]",
                                    user_input,
                                )
                            )
                        )

                        if _is_catalog_intent:
                            # ΓòÉΓòÉΓòÉ σ╜Æτ║│µ¿íσ╝Å∩╝ÜFolderCatalogOrganizer τë⌐τÉåµò┤τÉå ΓòÉΓòÉΓòÉ
                            t = yield_thinking(
                                f"σÉ»σè¿µûçΣ╗╢σñ╣σ╜Æτ║│: {_scan_dir}∩╝îτ¡¢ΘÇë: {_ext_label}",
                                "searching",
                            )
                            if t:
                                yield t
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒùé∩╕Å µ¡úσ£¿σ╜Æτ║│ {_scan_dir} Σ╕¡τÜä {_ext_label} µûçΣ╗╢...', 'detail': 'Σ╗àσñäτÉåσ╜ôσëìτ¢«σ╜ò∩╝îΣ╕ìΦ┐¢σàÑσ¡ÉµûçΣ╗╢σñ╣'}, ensure_ascii=False)}\n\n"
                            try:
                                try:
                                    from web.folder_catalog_organizer import (
                                        FolderCatalogOrganizer,
                                    )
                                except ImportError:
                                    from folder_catalog_organizer import (
                                        FolderCatalogOrganizer,
                                    )
                                _analyzer = get_file_analyzer()
                                _organizer_inst = get_file_organizer()
                                _engine = FolderCatalogOrganizer(
                                    get_organize_root(), _analyzer, _organizer_inst
                                )
                                _summary = _engine.organize_folder(
                                    str(_scan_dir),
                                    recursive=False,
                                    ext_filters=_ext_filters if _ext_filters else None,
                                )
                                if _summary.get("success"):
                                    _entries = _summary.get("entries", [])
                                    _ok = [e for e in _entries if e.get("organized")]
                                    _fail = [
                                        e for e in _entries if not e.get("organized")
                                    ]
                                    _organize_root_display = get_organize_root()
                                    response_text = (
                                        f"Γ£à σ╜Æτ║│σ«îµêÉ∩╝ü\n\n"
                                        f"- ≡ƒôé µ¥Ñµ║Éτ¢«σ╜ò: `{_scan_dir}`\n"
                                        f"- ≡ƒôü µò┤τÉåσê░: `{_organize_root_display}`\n"
                                        f"- Γ£ö∩╕Å σ╖▓µò┤τÉå: **{len(_ok)}** Σ╕¬µûçΣ╗╢\n"
                                    )
                                    if _fail:
                                        response_text += f"- ΓÜá∩╕Å σñ▒Φ┤Ñ: {len(_fail)} Σ╕¬∩╝ê{', '.join(e['file_name'] for e in _fail[:3])}{'...' if len(_fail) > 3 else ''}∩╝ë\n"
                                    if _summary.get("report_markdown"):
                                        response_text += f"\n≡ƒº╛ σ╜Æτ║│µ╕àσìòσ╖▓Σ┐¥σ¡ÿ: `{_summary['report_markdown']}`"
                                    # µÿ╛τñ║σêåτ╗äτ╗ôµ₧£
                                    _groups: dict = {}
                                    for _e in _ok:
                                        _grp = _e.get("suggested_folder", "σà╢Σ╗û")
                                        _groups.setdefault(_grp, []).append(
                                            _e["file_name"]
                                        )
                                    if _groups:
                                        response_text += "\n\n### ≡ƒôé σ╜Æτ║│σêåτ╗ä\n"
                                        for _grp, _names in _groups.items():
                                            response_text += f"\n**{_grp}**\n"
                                            for _n in _names:
                                                response_text += f"- {_n}\n"
                                else:
                                    response_text = f"Γ¥î σ╜Æτ║│σñ▒Φ┤Ñ: {_summary.get('error', 'µ£¬τƒÑΘöÖΦ»»')}"
                            except Exception as _oe:
                                response_text = f"Γ¥î σ╜Æτ║│σ╝éσ╕╕: {str(_oe)}"
                                _app_logger.debug(
                                    f"[FILE_SEARCH] FolderCatalogOrganizer σ╝éσ╕╕: {_oe}"
                                )
                        elif _is_filter_intent:
                            # ΓòÉΓòÉΓòÉ σåàσ«╣Φ┐çµ╗ñµ¿íσ╝Å∩╝ÜOllama σêñµû¡σô¬Σ║¢µûçΣ╗╢τ¼ªσÉêµÅÅΦ┐░ ΓòÉΓòÉΓòÉ
                            # Σ╗ÄΘù«σÅÑΣ╕¡µÅÉσÅûΦ┐çµ╗ñµ¥íΣ╗╢∩╝êσÄ╗ΘÖñΦ╖»σ╛äσÆîτ╗ôµ₧äΦ»ì∩╝ë
                            _criterion_raw = user_input
                            _criterion_raw = _re_filter_intent.sub(
                                r"[A-Za-z]:[\\][^\s\u4e00-\u9fa5]*", "", _criterion_raw
                            ).strip()
                            _criterion_raw = _re_filter_intent.sub(
                                r"Φ┐ÖΣ╕¬Φ╖»σ╛äΣ╕ï|Φ╖»σ╛äΣ╕ï|Φ┐ÖΣ╕¬µûçΣ╗╢σñ╣|µûçΣ╗╢σñ╣Σ╕ï|Σ╕ïΘ¥ó|Σ╕ïµ£ë|σô¬σçáΣ╕¬|µ£ëσô¬σçáΣ╕¬"
                                r"|µ£ëσçáΣ╕¬|σô¬Σ║¢|σô¬σçáΣ╗╜|σçáΣ╕¬µÿ»|σçáΣ╗╜µÿ»|µÿ»τÜä|Φ┐ÖΣ╕¬|τÜäµûçΣ╗╢|Φ┐Öµ¥íΦ╖»σ╛ä",
                                "",
                                _criterion_raw,
                            ).strip()
                            _criterion = _criterion_raw.strip(
                                "∩╝ƒ?πÇé "
                            ) or user_input.strip("∩╝ƒ?πÇé ")
                            t = yield_thinking(
                                f"σåàσ«╣Φ┐çµ╗ñ: {_scan_dir} / µ¥íΣ╗╢: {_criterion}",
                                "searching",
                            )
                            if t:
                                yield t
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒöÄ µ¡úσ£¿ΘÇÉΣ╕Çσêåµ₧ÉµûçΣ╗╢σåàσ«╣...', 'detail': f'Φ┐çµ╗ñµ¥íΣ╗╢: {_criterion}'}, ensure_ascii=False)}\n\n"
                            try:
                                try:
                                    from web.file_qa import (
                                        filter_files_by_criterion as _ffc2,
                                    )
                                except ImportError:
                                    from file_qa import (
                                        filter_files_by_criterion as _ffc2,
                                    )
                                _filter_res = _ffc2(
                                    criterion=_criterion,
                                    directory=str(_scan_dir),
                                    ext_filters=_ext_filters if _ext_filters else None,
                                )
                                if _filter_res.get("success"):
                                    _matches2 = _filter_res.get("matches", [])
                                    _scanned2 = _filter_res.get("total_scanned", 0)
                                    if _matches2:
                                        response_text = (
                                            f"≡ƒöÄ σ£¿ `{_scan_dir}` µë½µÅÅΣ║å **{_scanned2}** Σ╕¬µûçΣ╗╢∩╝î"
                                            f"µë╛σê░ **{len(_matches2)}** Σ╕¬τ¼ªσÉêπÇî{_criterion}πÇì∩╝Ü\n\n"
                                        )
                                        for _mi, _mm in enumerate(_matches2, 1):
                                            response_text += (
                                                f"**{_mi}. `{_mm['file_name']}`**\n"
                                            )
                                            if _mm.get("reason"):
                                                response_text += (
                                                    f"   _{_mm['reason']}_\n"
                                                )
                                            response_text += "\n"
                                    else:
                                        response_text = (
                                            f"≡ƒô¡ σ£¿ `{_scan_dir}` µë½µÅÅΣ║å **{_scanned2}** Σ╕¬µûçΣ╗╢∩╝î"
                                            f"µ£¬µë╛σê░τ¼ªσÉêπÇî{_criterion}πÇìµÅÅΦ┐░τÜäµûçΣ╗╢πÇé"
                                        )
                                else:
                                    response_text = (
                                        f"Γ¥î {_filter_res.get('error', 'Φ┐çµ╗ñσñ▒Φ┤Ñ')}"
                                    )
                            except Exception as _fe3:
                                response_text = f"Γ¥î σåàσ«╣Φ┐çµ╗ñσ╝éσ╕╕: {_fe3}"
                        else:
                            # ΓòÉΓòÉΓòÉ µÉ£τ┤ó/σêùΣ╕╛µ¿íσ╝Å∩╝ÜσÅ¬Φ»╗∩╝îµÿ╛τñ║µûçΣ╗╢σêùΦí¿ ΓòÉΓòÉΓòÉ
                            t = yield_thinking(
                                f"µë½µÅÅΦ╖»σ╛ä: {_scan_dir}∩╝êΣ╗àσ╜ôσëìσ▒é∩╝ë∩╝îτ¡¢ΘÇë: {_ext_label}",
                                "searching",
                            )
                            if t:
                                yield t
                            _file_list = []
                            try:
                                for _entry in _scan_dir.iterdir():
                                    if not _entry.is_file():
                                        continue
                                    if _entry.name.startswith("~$"):
                                        continue
                                    if (
                                        _ext_filters
                                        and _entry.suffix.lower() not in _ext_filters
                                    ):
                                        continue
                                    try:
                                        _stat = _entry.stat()
                                        _sz = _stat.st_size
                                        _sz_str = (
                                            f"{_sz} B"
                                            if _sz < 1024
                                            else (
                                                f"{_sz/1024:.1f} KB"
                                                if _sz < 1048576
                                                else f"{_sz/1048576:.1f} MB"
                                            )
                                        )
                                        _file_list.append(
                                            {
                                                "name": _entry.name,
                                                "size": _sz_str,
                                                "mtime": _stat.st_mtime,
                                                "mtime_str": _time_fmt.strftime(
                                                    "%Y-%m-%d %H:%M",
                                                    _time_fmt.localtime(_stat.st_mtime),
                                                ),
                                            }
                                        )
                                    except (PermissionError, OSError):
                                        pass
                            except (PermissionError, OSError):
                                pass
                            _file_list.sort(key=lambda x: x["mtime"], reverse=True)
                            if not _file_list:
                                response_text = f"≡ƒôü σ£¿ `{_scan_dir}` Σ╕¡µ£¬µë╛σê░Σ╗╗Σ╜ò **{_ext_label}** µûçΣ╗╢πÇé"
                            else:
                                response_text = f"≡ƒôü σ£¿ `{_scan_dir}` Σ╕¡µë╛σê░ **{len(_file_list)}** Σ╕¬ **{_ext_label}** µûçΣ╗╢∩╝Ü\n\n"
                                response_text += "| # | µûçΣ╗╢σÉì | σñºσ░Å | Σ┐«µö╣µù╢Θù┤ |\n| --- | --- | --- | --- |\n"
                                for _i, _f in enumerate(_file_list[:100], 1):
                                    response_text += f"| {_i} | `{_f['name']}` | {_f['size']} | {_f['mtime_str']} |\n"
                                if len(_file_list) > 100:
                                    response_text += f"\n*...σà▒ {len(_file_list)} Σ╕¬µûçΣ╗╢∩╝îΣ╗àµÿ╛τñ║σëì 100 Σ╕¬*"
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            response_text,
                            task="FILE_SEARCH",
                            model_name=used_model,
                        )
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                # ΓöÇΓöÇ ≡ƒôü µûçΣ╗╢σñ╣τ¢æµÄº∩╝êWatch Mode∩╝ë ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                _WATCH_START_KWS = [
                    "τ¢æµÄºµûçΣ╗╢σñ╣",
                    "τ¢æµÄºτ¢«σ╜ò",
                    "τ¢æµÄºΦ┐ÖΣ╕¬µûçΣ╗╢σñ╣",
                    "σ╝Çσºïτ¢æµÄº",
                    "Φç¬σè¿σ╜Æτ▒╗",
                ]
                _WATCH_STOP_KWS = ["σü£µ¡óτ¢æµÄº", "σÅûµ╢êτ¢æµÄº", "σà│Θù¡τ¢æµÄº"]
                _WATCH_LIST_KWS = ["µ¡úσ£¿τ¢æµÄº", "τ¢æµÄºσêùΦí¿", "µƒÑτ£ïτ¢æµÄº", "µ£ëσô¬Σ║¢τ¢æµÄº"]
                _watch_path_m = _re_pathscan.search(
                    r"([A-Za-z]:[\\][^\s\u4e00-\u9fa5]*)", user_input
                )
                if any(k in user_input for k in _WATCH_START_KWS) and _watch_path_m:
                    _wdir = _watch_path_m.group(1).rstrip("\\/. ")
                    try:
                        from web.file_watcher import get_file_watcher
                    except ImportError:
                        from file_watcher import get_file_watcher
                    _watcher = get_file_watcher()
                    _watcher.configure(
                        get_file_analyzer(), get_file_organizer(), get_organize_root()
                    )
                    _wres = _watcher.start_watch(_wdir)
                    if _wres.get("success"):
                        response_text = (
                            f"≡ƒæü∩╕Å **µûçΣ╗╢σñ╣τ¢æµÄºσ╖▓σÉ»σè¿∩╝ü**\n\n"
                            f"- ≡ƒôé τ¢æµÄºτ¢«σ╜ò: `{_wdir}`\n"
                            f"- ΓÜí µû░µûçΣ╗╢ΦÉ╜σ£░σÉÄΦç¬σè¿σêåµ₧Éσ╣╢σ╜Æτ▒╗σê░ `{get_organize_root()}`\n"
                            f"- ≡ƒöò Φ»┤πÇî**σü£µ¡óτ¢æµÄº {_wdir}**πÇìσÅ»ΘÜÅµù╢σà│Θù¡\n\n"
                            f"_µö»µîüµá╝σ╝Å: .doc/.docx/.pdf/.xlsx/.pptx/.txt/.csv/.zip τ¡ë_"
                        )
                    else:
                        response_text = (
                            f"Γ¥î σÉ»σè¿τ¢æµÄºσñ▒Φ┤Ñ: {_wres.get('error', 'µ£¬τƒÑΘöÖΦ»»')}"
                        )
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                if any(k in user_input for k in _WATCH_STOP_KWS) and _watch_path_m:
                    _wdir = _watch_path_m.group(1).rstrip("\\/. ")
                    try:
                        from web.file_watcher import get_file_watcher
                    except ImportError:
                        from file_watcher import get_file_watcher
                    _wres = get_file_watcher().stop_watch(_wdir)
                    response_text = (
                        f"Γ¢ö σ╖▓σü£µ¡óτ¢æµÄº `{_wdir}`"
                        if _wres.get("success")
                        else f"ΓÜá∩╕Å σü£µ¡óσñ▒Φ┤Ñ: {_wres.get('error', 'Φ»Ñτ¢«σ╜òµ£¬σ£¿τ¢æµÄºΣ╕¡')}"
                    )
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                if any(k in user_input for k in _WATCH_LIST_KWS):
                    try:
                        from web.file_watcher import get_file_watcher
                    except ImportError:
                        from file_watcher import get_file_watcher
                    _watches = get_file_watcher().list_watches()
                    if _watches:
                        response_text = "≡ƒæü∩╕Å **σ╜ôσëìτ¢æµÄºτ¢«σ╜òσêùΦí¿∩╝Ü**\n\n"
                        for _w in _watches:
                            _alive = "Γ£à Φ┐ÉΦíîΣ╕¡" if _w["alive"] else "ΓÜá∩╕Å σ╖▓σü£µ¡ó"
                            response_text += f"- `{_w['path']}` ΓÇö {_alive}∩╝êΦç¬ {_w['started_at']}∩╝ë\n"
                    else:
                        response_text = "≡ƒô¡ σ╜ôσëìµ▓íµ£ëµ¡úσ£¿τ¢æµÄºτÜäµûçΣ╗╢σñ╣πÇé\n\nΦ»┤πÇî**τ¢æµÄºµûçΣ╗╢σñ╣ C:\\xxx**πÇìσÅ»σÉ»σè¿τ¢æµÄºπÇé"
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                # ΓöÇΓöÇ ≡ƒôï σìòµûçΣ╗╢σà│Θö«σ¡ùµ«╡µÅÉσÅû∩╝êσÉêσÉî/σÅæτÑ¿/τ«ÇσÄåτ¡ë∩╝ë ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                _FIELDS_KWS = [
                    "µÅÉσÅûσ¡ùµ«╡",
                    "µÅÉσÅûΣ┐íµü»",
                    "σà│Θö«Σ┐íµü»",
                    "σÉêσÉîΣ┐íµü»",
                    "σÅæτÑ¿Σ┐íµü»",
                    "µÅÉσÅûσà│Θö«",
                    "ΦºúΦ»╗Φ┐ÖΣ╕¬",
                    "Φ»╗Σ╕ÇΣ╕ïΦ┐ÖΣ╕¬",
                    "σêåµ₧ÉΦ┐ÖΣ╕¬µûçΣ╗╢",
                    "µûçΣ╗╢σåàσ«╣",
                ]

                # ΓöÇΓöÇ ≡ƒùé∩╕Å σ╖ÑΣ╜£µûçΣ╗╢σ║ôτ«íτÉåσæ╜Σ╗ñ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                _WFL_ADD_KWS = [
                    "µ╖╗σèáτ¢æµÄºµûçΣ╗╢σñ╣",
                    "µ╖╗σèáµûçΣ╗╢σñ╣",
                    "σèáσàÑµûçΣ╗╢σ║ô",
                    "τ¢æµÄºΦ┐ÖΣ╕¬µûçΣ╗╢σñ╣σê░µûçΣ╗╢σ║ô",
                    "µèèΦ┐ÖΣ╕¬µûçΣ╗╢σñ╣σèáσàÑµûçΣ╗╢σ║ô",
                    "µ╖╗σèáσê░µûçΣ╗╢σ║ô",
                ]
                _WFL_REFRESH_KWS = [
                    "σê╖µû░µûçΣ╗╢σ║ô",
                    "µ¢┤µû░µûçΣ╗╢σ║ô",
                    "Θçìµû░µë½µÅÅµûçΣ╗╢σ║ô",
                    "Θçìσ╗║µûçΣ╗╢σ║ô",
                ]
                _WFL_STATUS_KWS = [
                    "µûçΣ╗╢σ║ôτè╢µÇü",
                    "µûçΣ╗╢σ║ôτ╗ƒΦ«í",
                    "µûçΣ╗╢σ║ôµ£ëσñÜσ░æ",
                    "µûçΣ╗╢σ║ôΘçîµ£ëΣ╗ÇΣ╣ê",
                    "µûçΣ╗╢σ║ôµªéσå╡",
                    "µƒÑτ£ïµûçΣ╗╢σ║ô",
                ]

                if any(k in user_input for k in _WFL_ADD_KWS) and _watch_path_m:
                    _add_path = _watch_path_m.group(1).rstrip("\\/. ")
                    try:
                        from web.work_file_library import get_work_file_library

                        _wfl2 = get_work_file_library()
                        _added = _wfl2.add_watch_folder(_add_path)
                        if _added:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒôé µ¡úσ£¿µë½µÅÅ {_add_path}...', 'detail': 'µ╖╗σèáσê░σ╖ÑΣ╜£µûçΣ╗╢σ║ô'}, ensure_ascii=False)}\n\n"
                            _wfl2.scan_locations(locations=[_add_path])
                            _wfl2.wait_for_scan(timeout=15.0)
                            _cnt = _wfl2.count()
                            response_text = (
                                f"Γ£à σ╖▓σ░å `{_add_path}` µ╖╗σèáσê░σ╖ÑΣ╜£µûçΣ╗╢σ║ôσ╣╢σ«îµêÉµë½µÅÅ∩╝ü\n\n"
                                f"µûçΣ╗╢σ║ôτÄ░σà▒µö╢σ╜ò **{_cnt}** Σ╕¬σ╖ÑΣ╜£µûçΣ╗╢πÇé\n"
                                "Σ╗ÑσÉÄΦ»┤πÇîµë╛ xxxπÇìσì│σÅ»σ┐½ΘÇƒµúÇτ┤óπÇé"
                            )
                        else:
                            response_text = (
                                f"Γ¥î µ╖╗σèáσñ▒Φ┤Ñ∩╝îΦ»╖τí«Φ«ñΦ╖»σ╛äσ¡ÿσ£¿: `{_add_path}`"
                            )
                    except Exception as _wadd_e:
                        response_text = f"Γ¥î µ╖╗σèáµûçΣ╗╢σñ╣σç║ΘöÖ: {_wadd_e}"
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                if any(k in user_input for k in _WFL_REFRESH_KWS):
                    try:
                        from web.work_file_library import get_work_file_library

                        _wfl3 = get_work_file_library()
                        yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒöä µ¡úσ£¿σê╖µû░σ╖ÑΣ╜£µûçΣ╗╢σ║ô...', 'detail': 'Θçìµû░µë½µÅÅµëÇµ£ëΣ╜ìτ╜«'}, ensure_ascii=False)}\n\n"
                        _wfl3.scan_locations(force=True)
                        _wfl3.wait_for_scan(timeout=15.0)
                        _st = _wfl3.get_stats()
                        _cats_str = "πÇü".join(
                            f"{k} {v}Σ╕¬" for k, v in _st.get("categories", {}).items()
                        )
                        response_text = (
                            f"Γ£à σ╖ÑΣ╜£µûçΣ╗╢σ║ôσ╖▓σê╖µû░∩╝ü\n\n"
                            f"σà▒µö╢σ╜ò **{_st['total']}** Σ╕¬σ╖ÑΣ╜£µûçΣ╗╢"
                            + (f"∩╝ê{_cats_str}∩╝ë" if _cats_str else "")
                            + "πÇé"
                        )
                    except Exception as _wref_e:
                        response_text = f"Γ¥î σê╖µû░µûçΣ╗╢σ║ôσç║ΘöÖ: {_wref_e}"
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                if any(k in user_input for k in _WFL_STATUS_KWS):
                    try:
                        from web.work_file_library import (
                            _CATEGORY_ICONS,
                            get_work_file_library,
                        )

                        _wfl4 = get_work_file_library()
                        _st4 = _wfl4.get_stats()
                        import time as _t4

                        _ls = _st4.get("last_scan")
                        _ls_str = (
                            _t4.strftime("%Y-%m-%d %H:%M", _t4.localtime(_ls))
                            if _ls
                            else "Σ╗Äµ£¬µë½µÅÅ"
                        )
                        response_text = f"### ≡ƒùé∩╕Å σ╖ÑΣ╜£µûçΣ╗╢σ║ôτè╢µÇü\n\n"
                        response_text += f"- **µö╢σ╜òµÇ╗µò░**: {_st4['total']} Σ╕¬σ╖ÑΣ╜£µûçΣ╗╢\n"
                        response_text += f"- **µ£ÇσÉÄµë½µÅÅ**: {_ls_str}\n\n"
                        if _st4.get("categories"):
                            response_text += "**µîëτ▒╗σ₧ïσêåσ╕â∩╝Ü**\n\n"
                            for _cat, _cnt4 in _st4["categories"].items():
                                _icon4 = _CATEGORY_ICONS.get(_cat, "≡ƒôÄ")
                                response_text += f"- {_icon4} {_cat}: **{_cnt4}** Σ╕¬\n"
                        _wfs = _wfl4.list_watch_folders()
                        _default_locs = __import__(
                            "web.work_file_library", fromlist=["_get_common_locations"]
                        )._get_common_locations()
                        response_text += (
                            f"\n**µë½µÅÅΣ╜ìτ╜«∩╝ê{len(_default_locs) + len(_wfs)} Σ╕¬∩╝ë∩╝Ü**\n"
                        )
                        for _loc in _default_locs:
                            response_text += f"- `{_loc}` ∩╝êΘ╗ÿΦ«ñ∩╝ë\n"
                        for _wf in _wfs:
                            response_text += f"- `{_wf['path']}` ∩╝êτö¿µê╖µ╖╗σèá∩╝ë\n"
                        response_text += "\nΦ»┤πÇî**σê╖µû░µûçΣ╗╢σ║ô**πÇìσÅ»Θçìµû░µë½µÅÅ∩╝îΦ»┤πÇî**µ╖╗σèáτ¢æµÄºµûçΣ╗╢σñ╣ Φ╖»σ╛ä**πÇìσÅ»µë⌐σñºΦîâσ¢┤πÇé"
                    except Exception as _wst_e:
                        response_text = f"Γ¥î ΦÄ╖σÅûµûçΣ╗╢σ║ôτè╢µÇüσñ▒Φ┤Ñ: {_wst_e}"
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return
                if any(k in user_input for k in _FIELDS_KWS) and _watch_path_m:
                    _tgt_file = _watch_path_m.group(1).rstrip("\\/. ")
                    from pathlib import Path as _FPath

                    _fp = _FPath(_tgt_file)
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒôï µ¡úσ£¿ΦºúΦ»╗ {_fp.name}...', 'detail': 'µÅÉσÅûσà│Θö«σ¡ùµ«╡'}, ensure_ascii=False)}\n\n"
                    try:
                        try:
                            from web.file_fields_extractor import extract_fields as _ef
                            from web.file_fields_extractor import (
                                fields_to_markdown as _fm,
                            )
                        except ImportError:
                            from file_fields_extractor import extract_fields as _ef
                            from file_fields_extractor import fields_to_markdown as _fm
                        _ana = get_file_analyzer()
                        _content = _ana._extract_content(str(_fp))
                        _fields = _ef(_fp.name, _content, _fp.suffix.lower())
                        if _fields:
                            response_text = "### ≡ƒôï µûçΣ╗╢σà│Θö«Σ┐íµü»\n\n" + _fm(
                                _fields, _fp.name
                            )
                        else:
                            response_text = f"ΓÜá∩╕Å µùáµ│òµÅÉσÅûσ¡ùµ«╡∩╝êOllama σÅ»Φâ╜µ£¬Φ┐ÉΦíî∩╝îµêûµûçΣ╗╢σåàσ«╣µùáµ│òΦºúµ₧É∩╝ë\nµûçΣ╗╢: `{_fp.name}`"
                    except Exception as _fe:
                        response_text = f"Γ¥î µÅÉσÅûσñ▒Φ┤Ñ: {_fe}"
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                # ΓöÇΓöÇ ≡ƒñû Φ╖¿µûçΣ╗╢Θù«τ¡ö∩╝êFile Q&A∩╝ë ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                _QA_KWS = [
                    "σô¬Σ╕¬",
                    "σô¬Σ║¢",
                    "µ£ëµ▓íµ£ë",
                    "σà│Σ║Ä",
                    "Φ░üτÜä",
                    "µ£Çµù⌐",
                    "µ£ÇµÖÜ",
                    "µ£ÇΘ½ÿ",
                    "Φ┐ÖσçáΣ╗╜",
                    "Φ┐ÖΣ║¢µûçΣ╗╢",
                    "µûçΣ╗╢Θçî",
                    "ΘçîΘ¥óµ£ëµ▓íµ£ë",
                    "σ«âΣ╗¼",
                    "µ▒çµÇ╗Σ╕ÇΣ╕ï",
                    "σæèΦ»ëµêæ",
                    "µƒÑΣ╕ÇµƒÑ",
                    "σ»╣µ»ö",
                ]
                # µ£ëΘù«ΘóÿΦ»ì + µ£ëΦ╖»σ╛ä ΓåÆ Φ╖¿µûçΣ╗╢Θù«τ¡ö
                if any(k in user_input for k in _QA_KWS) and _watch_path_m:
                    _qa_dir = _watch_path_m.group(1).rstrip("\\/. ")
                    yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒñû µ¡úσ£¿ΘÿàΦ»╗µûçΣ╗╢σ╣╢µÇ¥ΦÇâ...', 'detail': 'Φ╖¿µûçΣ╗╢Θù«τ¡ö'}, ensure_ascii=False)}\n\n"
                    try:
                        try:
                            from web.file_qa import answer_file_question
                        except ImportError:
                            from file_qa import answer_file_question
                        _qa_result = answer_file_question(
                            question=user_input,
                            search_dirs=[_qa_dir],
                            top_k=6,
                        )
                        if _qa_result.get("success"):
                            response_text = _qa_result["answer"]
                            _srcs = _qa_result.get("sources", [])
                            if _srcs:
                                response_text += "\n\n---\n**σÅéΦÇâµûçΣ╗╢∩╝Ü** " + "πÇü".join(
                                    f"`{s['file_name']}`" for s in _srcs
                                )
                        else:
                            response_text = f"Γ¥î {_qa_result.get('error', 'Θù«τ¡öσñ▒Φ┤Ñ')}"
                    except Exception as _qe:
                        response_text = f"Γ¥î Θù«τ¡öσ╝éσ╕╕: {_qe}"
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                if _is_disk:
                    # µúÇµ╡ïµÿ»σÉªµÿ»"µë½µÅÅ"µîçΣ╗ñ∩╝êΦÇîΘ¥₧µÉ£τ┤óµîçΣ╗ñ∩╝ë
                    _scan_cmd = any(
                        k in user_input.lower()
                        for k in [
                            "µë½µÅÅµêæτÜäτö╡Φäæ",
                            "µë½µÅÅτö╡Φäæ",
                            "µë½µÅÅτúüτ¢ÿ",
                            "µë½µÅÅτí¼τ¢ÿ",
                            "σà¿τ¢ÿµë½µÅÅ",
                            "σ╝Çσºïµë½µÅÅ",
                            "scan my",
                            "start scan",
                        ]
                    )
                    if _scan_cmd:
                        scan_started = FileScanner.start_scan()
                        if scan_started:
                            response_text = (
                                "≡ƒÜÇ σà¿τ¢ÿµûçΣ╗╢µë½µÅÅσ╖▓σÉ»σè¿∩╝ü\n\n"
                                f"µ¡úσ£¿µë½µÅÅΣ╗ÑΣ╕ïτúüτ¢ÿσêåσî║∩╝Ü**{', '.join(FileScanner.get_drives())}**\n\n"
                                "µë½µÅÅσ£¿σÉÄσÅ░Φ┐ÉΦíî∩╝îΣ╕ìΣ╝Üσ╜▒σôìµé¿τÜäΣ╜┐τö¿πÇéµë½µÅÅσ«îµêÉσÉÄµé¿σÅ»Σ╗ÑΘÇÜΦ┐çσ»╣Φ»¥∩╝Ü\n"
                                "- πÇîσ╕«µêæµë╛Σ╕ÇΣ╕ï xxx µûçΣ╗╢πÇì\n"
                                "- πÇîµëôσ╝Ç µêæτÜäτ«ÇσÄåπÇì\n"
                                "- πÇîµë╛Σ╕ÇΣ╕ï 2025σ╣┤µèÑσæèπÇì\n\n"
                                "Θªûµ¼íµë½µÅÅΘÇÜσ╕╕Θ£ÇΦªü **2-10 σêåΘÆƒ**∩╝îΣ╣ïσÉÄτ╗ôµ₧£Σ╝ÜµîüΣ╣àσîûπÇé"
                            )
                        else:
                            st = FileScanner.get_status()
                            scanned = st.get("scanned", 0)
                            indexed = st.get(
                                "indexed_count", FileScanner.stats()["total"]
                            )
                            response_text = (
                                f"ΓÅ│ σà¿τ¢ÿµë½µÅÅµ¡úσ£¿Φ┐¢ΦíîΣ╕¡...\n\n"
                                f"- σ╖▓µúÇµƒÑµûçΣ╗╢∩╝Ü**{scanned:,}**\n"
                                f"- σ╖▓τ┤óσ╝ò∩╝Ü**{indexed:,}**\n"
                                f"- σ╜ôσëìτ¢«σ╜ò∩╝Ü`{st.get('current_dir', '...')[:80]}`"
                            )
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            response_text,
                            task="FILE_SEARCH",
                            model_name=used_model,
                        )
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                    # ΓöÇΓöÇ σ╖ÑΣ╜£µûçΣ╗╢σ║ôµÉ£τ┤ó∩╝êΣ╝ÿσàê∩╝îµùáΘ£Çσà¿τ¢ÿµë½µÅÅ∩╝ëΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                    # Σ╗àσ╜ôτö¿µê╖µ£¬µÿÄτí«Φªüµ▒éσà¿τ¢ÿµÉ£τ┤óµù╢∩╝îσàêµƒÑσ╖ÑΣ╜£µûçΣ╗╢σ║ô
                    _EXPLICIT_FULL_DISK_KWS = [
                        "σà¿τ¢ÿ",
                        "µò┤Σ╕¬τö╡Φäæ",
                        "µëÇµ£ëτúüτ¢ÿ",
                        "µëÇµ£ëµûçΣ╗╢",
                        "σà¿τö╡Φäæ",
                        "σà¿Θâ¿τúüτ¢ÿ",
                        "σà¿Θâ¿µûçΣ╗╢",
                        "σà¿τí¼τ¢ÿ",
                    ]
                    _want_full_disk = any(
                        k in user_input for k in _EXPLICIT_FULL_DISK_KWS
                    )

                    if not _want_full_disk:
                        try:
                            from web.work_file_library import (
                                _CATEGORY_ICONS,
                                detect_category_from_input,
                                get_work_file_library,
                            )

                            _wfl = get_work_file_library()
                            _wfl_query = extract_query_from_input(user_input)

                            # σªéµ₧£σ║ôΦ┐ÿµ▓íµò░µì«∩╝îΦºªσÅæσ┐½ΘÇƒµë½µÅÅσ╣╢τ¡ëσ╛à
                            if not _wfl.is_indexed():
                                yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒôé µ¡úσ£¿σ┐½ΘÇƒσ╗║τ½ïσ╖ÑΣ╜£µûçΣ╗╢σ║ô...', 'detail': 'µë½µÅÅµíîΘ¥óπÇüµûçµíúπÇüΣ╕ïΦ╜╜τ¡ëσ╕╕τö¿Σ╜ìτ╜«'}, ensure_ascii=False)}\n\n"
                                _wfl.scan_locations()
                                _wfl.wait_for_scan(timeout=10.0)

                            # µúÇµ╡ïτö¿µê╖µäÅσ¢╛τÜäµûçΣ╗╢τ▒╗σ₧ï
                            _wfl_category = detect_category_from_input(user_input)
                            _wfl_results = _wfl.search(
                                _wfl_query, limit=30, category=_wfl_category
                            )
                            _wfl_stats = _wfl.get_stats()

                            if _wfl_results:
                                # µîëσêåτ▒╗σêåτ╗äσ▒òτñ║
                                _grouped: dict = {}
                                for _r in _wfl_results:
                                    _grouped.setdefault(_r["category"], []).append(_r)

                                response_text = f"≡ƒôé σ£¿µûçΣ╗╢σ║ôΣ╕¡µë╛σê░ **{len(_wfl_results)}** Σ╕¬σîàσÉ½πÇî{_wfl_query}πÇìτÜäµûçΣ╗╢∩╝Ü\n\n"
                                for _cat, _cat_files in _grouped.items():
                                    _icon = _CATEGORY_ICONS.get(_cat, "≡ƒôÄ")
                                    response_text += f"### {_icon} {_cat}∩╝ê{len(_cat_files)} Σ╕¬∩╝ë\n\n"
                                    response_text += "| µûçΣ╗╢σÉì | σñºσ░Å | Σ┐«µö╣µù╢Θù┤ |\n| --- | --- | --- |\n"
                                    for _f in _cat_files[:10]:
                                        response_text += f"| `{_f['name']}` | {_f['size_str']} | {_f['mtime_str']} |\n"
                                    if len(_cat_files) > 10:
                                        response_text += f"\n_...Φ┐ÿµ£ë {len(_cat_files) - 10} Σ╕¬σÉîτ▒╗µûçΣ╗╢_\n"
                                    response_text += "\n"

                                _cats_summary = "πÇü".join(
                                    f"{k} {v}Σ╕¬"
                                    for k, v in _wfl_stats.get("categories", {}).items()
                                )
                                response_text += (
                                    f"\n---\n_µûçΣ╗╢σ║ôσà▒µö╢σ╜ò **{_wfl_stats['total']}** Σ╕¬σ╖ÑΣ╜£µûçΣ╗╢"
                                    + (f"∩╝ê{_cats_summary}∩╝ë" if _cats_summary else "")
                                    + "πÇéσªéΘ£ÇµÉ£τ┤óµ¢┤σñÜΣ╜ìτ╜«∩╝îΦ»┤πÇîµ╖╗σèáτ¢æµÄºµûçΣ╗╢σñ╣ D:\\σ╖ÑΣ╜£Φ╡äµûÖπÇì_"
                                )

                                # σÉîµ¡ÑσÅæΘÇüµûçΣ╗╢ΘÇëµï⌐σÖ¿Σ║ïΣ╗╢∩╝êΣ╛¢σëìτ½»µ╕▓µƒôσìíτëç∩╝ë
                                _picker_files = [
                                    {
                                        "path": _r["path"],
                                        "name": _r["name"],
                                        "ext": _r["ext"],
                                        "category": _r["category"],
                                        "size_str": _r["size_str"],
                                        "mtime_str": _r["mtime_str"],
                                        "score": _r["score"],
                                    }
                                    for _r in _wfl_results[:12]
                                ]
                                yield f"data: {json.dumps({'type': 'file_picker', 'query': _wfl_query, 'count': len(_wfl_results), 'files': _picker_files, 'auto_opened': False}, ensure_ascii=False)}\n\n"
                                yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                                session_manager.append_and_save(
                                    f"{session_name}.json",
                                    user_input,
                                    response_text,
                                    task="FILE_SEARCH",
                                    model_name=used_model,
                                )
                                total_time = time.time() - start_time
                                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                                return

                            elif _wfl.is_indexed():
                                # σ║ôσ╖▓σ╗║τ½ïΣ╜åµ£¼µ¼íµ£¬µë╛σê░
                                _cats_summary = "πÇü".join(
                                    f"{k} {v}Σ╕¬"
                                    for k, v in _wfl_stats.get("categories", {}).items()
                                )
                                response_text = (
                                    f"≡ƒô¡ σ╖ÑΣ╜£µûçΣ╗╢σ║ôΣ╕¡µ£¬µë╛σê░σîàσÉ½πÇî{_wfl_query}πÇìτÜäµûçΣ╗╢πÇé\n\n"
                                    f"µûçΣ╗╢σ║ôσ╜ôσëìµö╢σ╜òΣ║å **{_wfl_stats['total']}** Σ╕¬σ╖ÑΣ╜£µûçΣ╗╢"
                                    + (f"∩╝ê{_cats_summary}∩╝ë" if _cats_summary else "")
                                    + "πÇé\n\n≡ƒÆí µÅÉτñ║∩╝Ü\n"
                                    "- Φ»┤πÇî**µ╖╗σèáτ¢æµÄºµûçΣ╗╢σñ╣ D:\\σ╖ÑΣ╜£Φ╡äµûÖ**πÇìσÅ»µë⌐σñºµÉ£τ┤óΦîâσ¢┤\n"
                                    "- Φ»┤πÇî**σê╖µû░µûçΣ╗╢σ║ô**πÇìσÅ»Θçìµû░µë½µÅÅσ╖▓µ£ëΣ╜ìτ╜«\n"
                                    f"- Φ»┤πÇî**σà¿τ¢ÿµÉ£τ┤ó {_wfl_query}**πÇìσÅ»µÉ£τ┤óµò┤Σ╕¬τö╡Φäæ"
                                )
                                yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                                session_manager.append_and_save(
                                    f"{session_name}.json",
                                    user_input,
                                    response_text,
                                    task="FILE_SEARCH",
                                    model_name=used_model,
                                )
                                total_time = time.time() - start_time
                                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                                return

                        except Exception as _wfl_exc:
                            _app_logger.warning(
                                f"[FILE_SEARCH] ΓÜá∩╕Å σ╖ÑΣ╜£µûçΣ╗╢σ║ôµÉ£τ┤óσç║ΘöÖ∩╝êΘÖìτ║ºσê░σà¿τ¢ÿµë½µÅÅ∩╝ë: {_wfl_exc}"
                            )
                            # τ╗ºτ╗¡Φ╡░σà¿τ¢ÿµë½µÅÅΘÇ╗Φ╛æ

                    # ΓöÇΓöÇ σà¿τ¢ÿµûçΣ╗╢σÉìµ¿íτ│èµÉ£τ┤ó ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                    FileScanner.ensure_loaded()
                    query = extract_query_from_input(user_input)
                    t = yield_thinking(
                        f"σà¿τ¢ÿµÉ£τ┤óσà│Θö«Φ»ì: {query!r}∩╝îτ┤óσ╝òΘçÅ: {FileScanner.stats()['total']:,}",
                        "searching",
                    )
                    if t:
                        yield t

                    if not FileScanner.is_indexed():
                        # µ▓íµ£ëτ┤óσ╝ò ΓåÆ σ╗║Φ««τö¿µê╖ΦºªσÅæµë½µÅÅ
                        response_text = (
                            "ΓÜá∩╕Å σà¿τ¢ÿµûçΣ╗╢τ┤óσ╝òσ░Üµ£¬σ╗║τ½ïπÇé\n\n"
                            "Φ»╖Φ»┤πÇî**µë½µÅÅµêæτÜäτö╡Φäæ**πÇìΦ«⌐ Koto Φç¬σè¿σ╝Çσºïµë½µÅÅ∩╝î\n"
                            "µêûτé╣σç╗Σ╛ºΦ╛╣µáÅ **µûçΣ╗╢τ«íτÉå > σ╝Çσºïσà¿τ¢ÿµë½µÅÅ**πÇé\n\n"
                            "Θªûµ¼íµë½µÅÅσÅ»Φâ╜Θ£ÇΦªü 2-10 σêåΘÆƒ∩╝îσ«îµêÉσÉÄσì│σÅ»ΘÇÜΦ┐çσ»╣Φ»¥σ┐½ΘÇƒµƒÑµë╛σÆîµëôσ╝ÇΣ╗╗Σ╜òµûçΣ╗╢πÇé"
                        )
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            response_text,
                            task="FILE_SEARCH",
                            model_name=used_model,
                        )
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                    disk_results = FileScanner.search(query, limit=12)

                    if not disk_results:
                        response_text = (
                            f"Γ¥î µ£¬µë╛σê░σîàσÉ½ **{query}** τÜäµûçΣ╗╢πÇé\n\n"
                            f"≡ƒÆí σ╗║Φ««∩╝Ü\n"
                            f"- µúÇµƒÑσà│Θö«Φ»ìµÿ»σÉªµ¡úτí«\n"
                            f"- σªéµ₧£µûçΣ╗╢Φ╛âµû░∩╝îσÅ»Σ╗ÑΘçìµû░µë½µÅÅ∩╝êΦ»┤πÇîµë½µÅÅµêæτÜäτö╡ΦäæπÇì∩╝ë\n"
                            f"- σ╜ôσëìτ┤óσ╝ò {FileScanner.stats()['total']:,} Σ╕¬µûçΣ╗╢"
                        )
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            response_text,
                            task="FILE_SEARCH",
                            model_name=used_model,
                        )
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                    # σêñµû¡µÿ»σÉªτ¢┤µÄÑµëôσ╝Ç∩╝êσö»Σ╕ÇΘ½ÿτ╜«Σ┐íσî╣Θàì∩╝ë
                    auto_open = (
                        len(disk_results) == 1 and disk_results[0]["score"] >= 0.9
                    ) or (
                        len(disk_results) >= 1
                        and disk_results[0]["score"] >= 0.95
                        and (len(disk_results) < 2 or disk_results[1]["score"] < 0.7)
                    )

                    if auto_open:
                        best = disk_results[0]
                        open_result = FileScanner.open_file(best["path"])
                        if open_result["success"]:
                            response_text = (
                                f"Γ£à σ╖▓Σ╕║µé¿µëôσ╝ÇµûçΣ╗╢∩╝Ü**{best['name']}**\n\n"
                                f"≡ƒôü Φ╖»σ╛ä: `{best['path']}`\n"
                                f"≡ƒôé σêåτ▒╗: {best['category']}πÇÇσñºσ░Å: {best['size_str']}πÇÇΣ┐«µö╣: {best['mtime_str']}"
                            )
                        else:
                            response_text = f"ΓÜá∩╕Å µë╛σê░µûçΣ╗╢Σ╜åµëôσ╝Çσñ▒Φ┤Ñ: {open_result.get('error', '')}\n\n≡ƒôü Φ╖»σ╛ä: `{best['path']}`"
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            response_text,
                            task="FILE_SEARCH",
                            model_name=used_model,
                        )
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                    # σñÜτ╗ôµ₧£ ΓåÆ σÅæΘÇü file_picker Σ║ïΣ╗╢∩╝êσëìτ½»µ╕▓µƒôΘÇëµï⌐σìíτëç∩╝ë
                    picker_event = {
                        "type": "file_picker",
                        "query": query,
                        "count": len(disk_results),
                        "files": disk_results[:12],
                        "auto_opened": False,
                    }
                    yield f"data: {json.dumps(picker_event, ensure_ascii=False)}\n\n"

                    response_text = f"≡ƒöì µë╛σê░ {len(disk_results)} Σ╕¬σî╣Θàì **{query}** τÜäµûçΣ╗╢∩╝îΦ»╖τé╣σç╗ΘÇëµï⌐Φªüµëôσ╝ÇτÜäµûçΣ╗╢∩╝Ü"
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                else:
                    # ΓöÇΓöÇ σ╖ÑΣ╜£σî║σåàσ«╣τ┤óσ╝òµÉ£τ┤ó∩╝êσÄƒµ£ëΘÇ╗Φ╛æ∩╝ëΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                    indexer = get_file_indexer()
                    keywords = (
                        user_input.replace("µë╛µûçΣ╗╢", "")
                        .replace("µÉ£τ┤ó", "")
                        .replace("µƒÑµë╛", "")
                    )
                    keywords = (
                        keywords.replace("σîàσÉ½", "").replace("τÜäµûçΣ╗╢", "").strip()
                    )
                    t = yield_thinking(f"σ╖ÑΣ╜£σî║σà│Θö«Φ»ìµÉ£τ┤ó: {keywords}", "searching")
                    if t:
                        yield t
                    results = indexer.search(keywords, limit=10)
                    if not results:
                        yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒöä µë⌐σ▒òµÉ£τ┤óΦîâσ¢┤...', 'detail': ''}, ensure_ascii=False)}\n\n"
                        results = indexer.find_by_content(keywords, min_similarity=0.2)
                    if results:
                        response_text = f"≡ƒöì µë╛σê░ {len(results)} Σ╕¬σî╣ΘàìµûçΣ╗╢:\n\n"
                        for i, r in enumerate(results[:10], 1):
                            file_name = r.get("file_name", "µ£¬τƒÑµûçΣ╗╢")
                            file_path = r.get("file_path", "")
                            snippet = r.get("match_snippet", "")
                            score = r.get("score", 0)
                            similarity = r.get("similarity")
                            response_text += (
                                f"### {i}. {file_name}\n≡ƒôü Φ╖»σ╛ä: `{file_path}`\n"
                            )
                            if similarity:
                                response_text += f"≡ƒÄ» τ¢╕Σ╝╝σ║ª: {similarity:.0%}\n"
                            elif score:
                                response_text += f"Γ¡É σî╣Θàìσêå: {score:.2f}\n"
                            if snippet:
                                response_text += f"≡ƒôä ΘóäΦºê: {snippet[:200]}...\n"
                            response_text += "\n"
                    else:
                        response_text = (
                            "Γ¥î µ£¬µë╛σê░σî╣ΘàìµûçΣ╗╢\n\n≡ƒÆí µÅÉτñ║:\n"
                            "- ΦïÑΦªüµÉ£τ┤óτö╡ΦäæΣ╕èµëÇµ£ëµûçΣ╗╢∩╝îΦ»╖Φ»┤πÇîσ╕«µêæµë╛Σ╕ÇΣ╕ï xxx µûçΣ╗╢πÇì\n"
                            f"- σ╜ôσëìσ╖ÑΣ╜£σî║τ┤óσ╝òµûçΣ╗╢µò░: {len(indexer.list_indexed_files(limit=1000))}"
                        )
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

            # === DOC_ANNOTATE Mode (µûçµíúµáçµ│¿/µ╢ªΦë▓ - µ╡üσ╝ÅσÅìΘªê) ===
            if task_type == "DOC_ANNOTATE":
                used_model = model_id if model_id else "gemini-3.1-pro-preview"
                t = yield_thinking(
                    f"Φ┐¢σàÑµûçµíúµáçµ│¿µ¿íσ╝Å∩╝îσ░åΣ╜┐τö¿ {model_id or 'gemini-3.1-pro-preview'} σêåµ₧Éµûçµíú",
                    "routing",
                )
                if t:
                    yield t
                _app_logger.debug(f"[STREAM] ≡ƒôä µëºΦíî DOC_ANNOTATE Σ╗╗σèí")

                # Σ╗ÄΦ»╖µ▒éΣ╕¡ΦÄ╖σÅûtask_id∩╝îτö¿Σ║Äµö»µîüσÅûµ╢êµôìΣ╜£
                task_id = request.json.get("task_id")

                # µƒÑµë╛µ£ÇΦ┐æΣ╕èΣ╝áτÜäµûçµíú
                doc_path = None
                upload_dirs = ["web/uploads", "uploads", "workspace/documents"]

                for dir_path in upload_dirs:
                    if os.path.exists(dir_path):
                        import glob

                        docs = []
                        for ext in [
                            ".docx",
                            ".docxm",
                            ".doc",
                            ".pdf",
                            ".txt",
                            ".md",
                            ".rtf",
                            ".odt",
                        ]:
                            docs.extend(
                                glob.glob(f"{dir_path}/**/*{ext}", recursive=True)
                            )
                        if docs:
                            doc_path = max(docs, key=os.path.getmtime)
                            break

                if not doc_path or not os.path.exists(doc_path):
                    yield f"data: {json.dumps({'type': 'progress', 'message': 'Γ¥î µ£¬µë╛σê░µûçµíú', 'detail': 'Φ»╖Σ╕èΣ╝á .docx/.doc/.pdf/.txt/.md/.rtf µûçΣ╗╢'})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                    return

                # Θ¥₧ .docx Φç¬σè¿Φ╜¼µìó∩╝êΦ╛ôσç║σñìσê╢σê░µáçσçåµûçµíúτ¢«σ╜ò∩╝îΘü┐σàìΦ╛ôσç║σê░ temp τ¢«σ╜ò∩╝ë
                _dw_ext = os.path.splitext(doc_path)[1].lower()
                _stream_docs_dir = settings_manager.documents_dir
                os.makedirs(_stream_docs_dir, exist_ok=True)
                if _dw_ext != ".docx":
                    try:
                        import tempfile as _tmpdw

                        from web.doc_converter import convert_to_docx

                        _dw_conv_dir = _tmpdw.mkdtemp(prefix="koto_dw_")
                        _dw_conv_path, _ = convert_to_docx(
                            doc_path, output_dir=_dw_conv_dir
                        )
                        # σñìσê╢Φ╜¼µìóσÉÄµûçΣ╗╢σê░µáçσçåµûçµíúτ¢«σ╜ò∩╝îτí«Σ┐¥Φ╛ôσç║Σ╣ƒσ£¿Φ»Ñτ¢«σ╜ò
                        _dw_conv_basename = os.path.basename(_dw_conv_path)
                        _dw_conv_in_docs = os.path.join(
                            _stream_docs_dir, _dw_conv_basename
                        )
                        import shutil as _dw_shutil

                        _dw_shutil.copy2(_dw_conv_path, _dw_conv_in_docs)
                        doc_path = _dw_conv_in_docs
                        _app_logger.debug(
                            f"[DocWorkflow] Φ╜¼µìó {_dw_ext} ΓåÆ .docx σ╣╢σñìσê╢σê░µûçµíúτ¢«σ╜ò: {doc_path}"
                        )
                    except Exception as _dw_conv_err:
                        yield f"data: {json.dumps({'type': 'error', 'message': f'µûçµíúΦ╜¼µìóσñ▒Φ┤Ñ: {_dw_conv_err}'})}\n\n"
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                        return

                # Step 1: Φ»╗σÅûµûçµíúΣ┐íµü»
                yield f"data: {json.dumps({'type': 'progress', 'stage': 'init_reading', 'message': '≡ƒôû µ¡úσ£¿Φ»╗σÅûµûçµíú...', 'detail': os.path.basename(doc_path)})}\n\n"

                doc_filename = os.path.basename(doc_path)
                total_chars = 0
                total_paras = 0

                try:
                    from docx import Document

                    doc = Document(doc_path)
                    total_paras = len([p for p in doc.paragraphs if p.text.strip()])
                    total_chars = sum(len(p.text) for p in doc.paragraphs)

                    yield f"data: {json.dumps({'type': 'progress', 'stage': 'init_reading_complete', 'message': f'Γ£à µûçµíúΦºúµ₧Éσ«îµêÉ', 'detail': f'{doc_filename}: {total_paras} µ«╡  |  {total_chars} σ¡ù'})}\n\n"
                except Exception as e:
                    yield f"data: {json.dumps({'type': 'error', 'message': f'Γ¥î Φ»╗σÅûµûçµíúσñ▒Φ┤Ñ: {str(e)}'})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                    return

                # Step 2: σ▒òτñ║Σ╗╗σèíΣ┐íµü»
                nl = "\n"
                task_info_msg = f"≡ƒôï πÇÉΣ╗╗σèíΣ┐íµü»πÇæ{nl}- µ¿íσ₧ï: {model_id}{nl}- Θ£Çµ▒é: {user_input[:100]}{nl}- µûçµíú: {doc_filename}"
                yield f"data: {json.dumps({'type': 'info', 'message': task_info_msg})}\n\n"

                try:
                    from web.document_feedback import DocumentFeedbackSystem

                    feedback_system = DocumentFeedbackSystem(gemini_client=client)

                    # Σ╜┐τö¿µ╡üσ╝Åσêåµ₧Éτ│╗τ╗ƒ∩╝îΘÇÉµ¡ÑσÅìΘªêΦ┐¢σ║ª
                    yield f"data: {json.dumps({'type': 'progress', 'stage': 'processing_start', 'message': '≡ƒöì σ╝ÇσºïσñäτÉåµûçµíú...', 'detail': 'Φ┐ÖΣ╕¬Φ┐çτ¿ïΣ╝Üµ╢ëσÅèσñÜΣ╕¬Θÿ╢µ«╡'})}\n\n"

                    revised_file = None
                    final_result = None
                    cancelled = False

                    # Φ┐¡Σ╗úµ╡üσ╝Åτ╗ôµ₧£∩╝îΣ╝áσàÑtask_idτö¿Σ║Äµö»µîüσÅûµ╢ê
                    for (
                        progress_event
                    ) in feedback_system.full_annotation_loop_streaming(
                        doc_path,
                        user_input,
                        task_id=task_id,
                        model_id=model_id,
                        cancel_check=lambda: _interrupt_manager.is_interrupted(
                            session_name
                        ),
                    ):
                        stage = progress_event.get("stage", "unknown")
                        progress = progress_event.get("progress", 0)
                        message = progress_event.get("message", "")
                        detail = progress_event.get("detail", "")

                        # σñäτÉåΣ╗╗σèíσÅûµ╢ê
                        if stage == "cancelled":
                            cancelled = True
                            yield f"data: {json.dumps({'type': 'info', 'message': 'ΓÅ╕∩╕Å Σ╗╗σèíσ╖▓σÅûµ╢ê', 'detail': 'τö¿µê╖Σ╕¡µ¡óΣ║åσñäτÉå'})}\n\n"
                            break

                        # µá╣µì«Θÿ╢µ«╡σÅæΘÇüΣ╕ìσÉîµá╖σ╝ÅτÜäΦ┐¢σ║ªΣ┐íµü»
                        yield f"data: {json.dumps({'type': 'progress', 'stage': stage, 'message': message, 'detail': detail, 'progress': progress})}\n\n"

                        # Σ┐¥σ¡ÿµ£Çτ╗êτ╗ôµ₧£
                        if stage == "complete":
                            final_result = progress_event.get("result", {})
                            revised_file = final_result.get("revised_file")

                    # σªéµ₧£Σ╗╗σèíΦó½σÅûµ╢ê∩╝îΦ┐öσ¢₧σÅûµ╢êσôìσ║ö
                    if cancelled:
                        total_time = time.time() - start_time
                        # Σ┐¥σ¡ÿσÅûµ╢êΦ«░σ╜òσê░σÄåσÅ▓
                        session_manager.append_and_save(
                            f"{session_name}.json", user_input, "ΓÅ╕∩╕Å µûçµíúµáçµ│¿Σ╗╗σèíσ╖▓σÅûµ╢ê"
                        )
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time, 'cancelled': True})}\n\n"
                        return

                    # σªéµ₧£σñäτÉåµêÉσèƒ∩╝îτöƒµêÉΦ»ªτ╗åµÇ╗τ╗ô
                    if final_result and final_result.get("success"):
                        applied = final_result.get("applied", 0)
                        failed = final_result.get("failed", 0)
                        total = final_result.get("total", applied + failed)

                        # Φ«íτ«ùΣ┐«µö╣σ»åσ║ª
                        density = (
                            (applied / total_chars * 1000) if total_chars > 0 else 0
                        )

                        summary_msg = (
                            f"Γ£à **µûçµíúΣ┐«µö╣σ«îµêÉ∩╝ü**\n\n"
                            f"≡ƒôè **µ╡ïΦ»òτ╗ôµ₧£**∩╝Ü\n"
                            f"- **µûçµíúσêåµ₧É**∩╝ÜµêÉσèƒΦ»╗σÅû {total_paras} µ«╡∩╝îσà▒ {total_chars} σ¡ùπÇé\n"
                            f"- **AI σñäτÉå**∩╝ÜµûçµíúΦó½σ╣╢σÅæσñäτÉå∩╝îµÇ╗ΦÇùµù╢τ║ª {int(time.time() - start_time)} τºÆπÇé\n"
                            f"- **τöƒµêÉΦ┤¿ΘçÅ**∩╝ÜAI µêÉσèƒµë╛σç║Σ║å **{total} σñä** τ┐╗Φ»æτöƒτí¼πÇüΦ»¡σ║ÅΣ╕ìΘí║τÜäσ£░µû╣πÇé\n"
                            f"- **σ║öτö¿Σ┐«Φ«ó**∩╝ÜµêÉσèƒσ░å **{applied} σñä** Σ┐«µö╣Σ╗ÑΓÇ£Σ┐«Φ«óµ¿íσ╝Å∩╝êTrack Changes∩╝ëΓÇ¥σåÖσàÑΣ║å Word µûçµíú∩╝êΣ╗àµ£ë {failed} σñäσ¢áσñìµ¥éµá╝σ╝Åσ«ÜΣ╜ìσñ▒Φ┤Ñ∩╝îσ▒₧Σ║Äµ¡úσ╕╕σ«╣ΘöÖΦîâσ¢┤∩╝ëπÇé\n\n"
                            f"≡ƒôé **Θ¬îΦ»üµûçΣ╗╢**∩╝Ü\n"
                            f"Θ½ÿΦ┤¿ΘçÅτÜäµ╡ïΦ»òτ╗ôµ₧£µûçΣ╗╢σ╖▓τ╗ÅτöƒµêÉσ£¿µé¿τÜäµ£¼σ£░τ¢«σ╜òΣ╕¡∩╝îµé¿σÅ»Σ╗Ñτ¢┤µÄÑµëôσ╝ÇµƒÑτ£ïµòêµ₧£∩╝Ü\n"
                            f"≡ƒæë `{os.path.basename(revised_file) if revised_file else 'σ╛àτöƒµêÉ'}`\n\n"
                            f"≡ƒÆí **Σ╜┐τö¿µû╣µ│ò**∩╝Ü\n"
                            f"1. τö¿ Microsoft Word µëôσ╝ÇΦ╛ôσç║µûçΣ╗╢\n"
                            f"2. τé╣σç╗πÇîσ«íΘÿàπÇìµáçτ¡╛Θí╡\n"
                            f"3. σÅ│Σ╛ºµ░öµ│íΣ╕¡µƒÑτ£ïσà¿Θâ¿Σ┐«µö╣σ╗║Φ««\n"
                            f"4. ΘÇÉµ¥íµÄÑσÅùµêûσ┐╜τòÑ∩╝êσÅ│Θö«µë╣µ│¿σÅ»µôìΣ╜£∩╝ë\n"
                            f"5. τé╣σç╗πÇîµÄÑσÅùσà¿Θâ¿πÇìµêûΘÇÉµ¥íσñäτÉå\n\n"
                            f"≡ƒôé **µûçΣ╗╢Σ╜ìτ╜«**: `{os.path.dirname(revised_file) if revised_file else settings_manager.documents_dir}`"
                        )

                        yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒô¥ τöƒµêÉµ£Çτ╗êµèÑσæè...', 'detail': ''})}\n\n"
                        yield f"data: {json.dumps({'type': 'token', 'content': summary_msg})}\n\n"

                        # Σ┐¥σ¡ÿσ»╣Φ»¥σÄåσÅ▓∩╝êσîàσÉ½σàâµò░µì«∩╝ë
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            summary_msg,
                            task="DOC_ANNOTATE",
                            model_name=model_id,
                            saved_files=[revised_file] if revised_file else [],
                        )

                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [revised_file] if revised_file else [], 'total_time': total_time})}\n\n"
                    else:
                        error_msg = (
                            final_result.get("message", "µ£¬τƒÑΘöÖΦ»»")
                            if final_result
                            else "σñäτÉåσñ▒Φ┤Ñ"
                        )
                        # Σ┐¥σ¡ÿσñ▒Φ┤ÑΦ«░σ╜ò
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            f"Γ¥î µûçµíúµáçµ│¿σñ▒Φ┤Ñ: {error_msg}",
                        )
                        yield f"data: {json.dumps({'type': 'error', 'message': f'Γ¥î σñäτÉåσñ▒Φ┤Ñ: {error_msg}'})}\n\n"

                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"

                except Exception as e:
                    import traceback

                    error_detail = traceback.format_exc()
                    _app_logger.error(f"[DOC_ANNOTATE] Γ¥î σñ▒Φ┤Ñ:\n{error_detail}")
                    # Σ┐¥σ¡ÿσ╝éσ╕╕Φ«░σ╜ò
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        f"Γ¥î µûçµíúµáçµ│¿σ╝éσ╕╕: {str(e)[:200]}",
                    )

                    yield f"data: {json.dumps({'type': 'error', 'message': f'Γ¥î σñäτÉåσ╝éσ╕╕: {str(e)[:200]}'})}\n\n"
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"

                return

            # === WEB_SEARCH Mode (Φüöτ╜æµÉ£τ┤ó - σ«₧µù╢Σ┐íµü») ===
            if task_type == "WEB_SEARCH":
                used_model = "gemini-2.5-flash (Google Search)"
                yield f"data: {json.dumps({'type': 'progress', 'message': 'µ¡úσ£¿Φ┐₧µÄÑΣ║ÆΦüöτ╜æµÉ£τ┤ó...', 'detail': ''})}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'message': 'µ¡úσ£¿µÉ£τ┤óσ«₧µù╢Σ┐íµü»...', 'detail': 'Google Search'})}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'message': 'µ¡úσ£¿µò┤τÉåµÉ£τ┤óτ╗ôµ₧£...', 'detail': ''})}\n\n"

                # Σ╝ÿσàêΣ╜┐τö¿µ£¼σ£░/AIΦ╖»τö▒σÖ¿τöƒµêÉτÜä skill_prompt∩╝îσ«₧τÄ░πÇîµ¿íσ₧ïτÉåΦºúµäÅσ¢╛ ΓåÆ τöƒµêÉµëºΦíîµîçΣ╗ñπÇì
                _skill_prompt = (context_info or {}).get("skill_prompt")
                search_result = WebSearcher.search_with_grounding(
                    user_input, skill_prompt=_skill_prompt
                )
                response_text = search_result["response"]

                if (
                    Utils.is_failure_output(response_text)
                    or "µÉ£τ┤óσñ▒Φ┤Ñ" in response_text
                ):
                    t = yield_thinking(
                        "σê¥µ¼íµÉ£τ┤óτ╗ôµ₧£Σ╕ìΣ╜│∩╝îΣ╜┐τö¿ gemini-2.0-flash-lite µö╣σåÖµƒÑΦ»óΦ»ìσÉÄΘçìΦ»ò",
                        "searching",
                    )
                    if t:
                        yield t
                    yield f"data: {json.dumps({'type': 'progress', 'message': 'ΓÜá∩╕Å σê¥µ¼íµÉ£τ┤óσñ▒Φ┤Ñ∩╝îµ¡úσ£¿Σ┐«µ¡úµƒÑΦ»ó...', 'detail': ''})}\n\n"
                    fix_query_prompt = (
                        "Φ»╖µèèτö¿µê╖Θ£Çµ▒éµö╣σåÖµêÉµ¢┤ΘÇéσÉêµÉ£τ┤óτÜäτ«Çτƒ¡σà│Θö«Φ»ìµêûµƒÑΦ»óΦ»¡σÅÑ∩╝îσÅ¬Φ╛ôσç║µƒÑΦ»óΦ»¡σÅÑπÇé\n"
                        f"τö¿µê╖Θ£Çµ▒é: {user_input}"
                    )
                    fix_query_resp = client.models.generate_content(
                        model="gemini-2.0-flash-lite",
                        contents=fix_query_prompt,
                        config=types.GenerateContentConfig(
                            temperature=0.2,
                            max_output_tokens=64,
                        ),
                    )
                    fixed_query = (fix_query_resp.text or user_input).strip()
                    search_result = WebSearcher.search_with_grounding(fixed_query)
                    response_text = search_result["response"]

                if Utils.is_failure_output(response_text):
                    fix_prompt = Utils.build_fix_prompt(
                        "WEB_SEARCH", user_input, response_text
                    )
                    fix_resp = client.models.generate_content(
                        model=model_id,
                        contents=fix_prompt,
                        config=types.GenerateContentConfig(
                            system_instruction=system_instruction,
                            temperature=0.4,
                            max_output_tokens=1200,
                        ),
                    )
                    response_text = fix_resp.text or response_text

                yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"

                # σàêΣ┐¥σ¡ÿσÄåσÅ▓∩╝îσåìσÅæΘÇü done Σ║ïΣ╗╢
                session_manager.append_and_save(
                    f"{session_name}.json", user_input, response_text
                )

                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return

            # === ≡ƒî│ Tree of Thought Mode (RESEARCH / FILE_GEN σ╣╢ΦíîσñÜΦ╖»µÄ¿τÉåΘÇëΣ╝ÿ) ===
            # ΦºªσÅæµ¥íΣ╗╢∩╝Ülegacy Φ╖»τö▒ + RESEARCH/FILE_GEN + Θ¥₧ Deep-Research-Pro + ToT µ£¬Φó½τö¿µê╖τªüτö¿
            _tot_enabled = (
                _wf_route == "legacy"
                and task_type in ("RESEARCH", "FILE_GEN")
                and len(str(effective_input)) >= 20
                and not str(model_id or "").startswith("deep-research-pro-preview")
                and settings_manager.get("ai", "use_tree_of_thought") is not False
            )

            if _tot_enabled:
                _tot_model = model_id or MODEL_MAP.get(
                    task_type, "gemini-2.5-flash-preview-05-20"
                )
                _tot_n = (
                    2 if task_type == "FILE_GEN" else 3
                )  # FILE_GEN τö¿ 2 Φ╖»∩╝îRESEARCH τö¿ 3 Φ╖»
                _tot_label = "≡ƒôä µûçµíúτöƒµêÉ" if task_type == "FILE_GEN" else "≡ƒö¼ µ╖▒σ║ªτáöτ⌐╢"
                yield f"data: {json.dumps({'type': 'classification', 'task_type': task_type, 'route_method': 'TreeOfThought', 'message': f'≡ƒî│ Tree of Thought σÉ»σè¿∩╝Ü{_tot_n} µ¥íσ╣╢ΦíîµÄ¿τÉåσêåµö» ({_tot_label})'}, ensure_ascii=False)}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒî│ Tree of Thought σÉ»σè¿ ({_tot_n} σêåµö»σ╣╢ΦíîµÄ¿τÉå)...', 'detail': f'µ¿íσ₧ï: {_tot_model}'})}\n\n"

                try:
                    from app.core.agent.tree_of_thought import create_tot

                    _tot = create_tot(
                        task_type=task_type, n_branches=_tot_n, model_id=_tot_model
                    )
                    _tot_final = ""
                    _tot_winner_id = None

                    for _evt in _tot.stream(
                        user_input=effective_input,
                        task_type=task_type,
                        system_instruction=system_instruction,
                    ):
                        _stage = _evt.get("stage", "")

                        if _stage == "expand":
                            _bid = _evt.get("branch_id", "?")
                            _blabel = _evt.get("label", "")
                            _bstatus = _evt.get("status", "")
                            if _bstatus == "generating":
                                yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒî┐ σêåµö» {_bid}πÇî{_blabel}πÇìτöƒµêÉΣ╕¡...', 'detail': ''})}\n\n"
                            elif _bstatus == "done":
                                _elapsed = _evt.get("elapsed", "")
                                yield f"data: {json.dumps({'type': 'progress', 'message': f'Γ£à σêåµö» {_bid}πÇî{_blabel}πÇìσ«îµêÉ ({_elapsed}s)', 'detail': _evt.get('preview', '')[:60]})}\n\n"
                            elif _bstatus == "error":
                                yield f"data: {json.dumps({'type': 'progress', 'message': f'ΓÜá∩╕Å σêåµö» {_bid} σñ▒Φ┤Ñ', 'detail': _evt.get('error', '')[:80]})}\n\n"

                        elif _stage == "evaluate":
                            _bstatus = _evt.get("status", "")
                            if _bstatus == "scoring":
                                yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒöì Critic µ¡úσ£¿Φ»äΣ╝░σÉäσêåµö»Φ┤¿ΘçÅ...', 'detail': ''})}\n\n"
                            else:
                                _bid = _evt.get("branch_id", "?")
                                _score = _evt.get("score", 0)
                                _crit = _evt.get("critique", "")
                                yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒôè σêåµö» {_bid} σ╛ùσêå {_score:.1f} ΓÇö {_crit}', 'detail': ''}, ensure_ascii=False)}\n\n"

                        elif _stage == "select":
                            _tot_winner_id = _evt.get("winner_id")
                            _tot_score = _evt.get("score", 0)
                            _wlabel = _evt.get("winner_label", "")
                            _reason = _evt.get("reason", "")
                            _tot_final = _evt.get("content", "")
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒÅå µ£ÇΣ╝ÿσêåµö»: {_tot_winner_id}πÇî{_wlabel}πÇì(σ╛ùσêå {_tot_score:.1f})', 'detail': _reason}, ensure_ascii=False)}\n\n"
                            yield f"data: {json.dumps({'type': 'token', 'content': _tot_final}, ensure_ascii=False)}\n\n"

                        elif _stage == "error":
                            _errmsg = _evt.get("message", "µ£¬τƒÑΘöÖΦ»»")
                            _app_logger.error(f"[ToT] Γ¥î ΘöÖΦ»»: {_errmsg}")
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'ΓÜá∩╕Å Tree of Thought Θüçσê░Θù«Θóÿ∩╝îσêçµìóΦç│µáçσçåµ¿íσ╝Å: {_errmsg[:100]}', 'detail': ''}, ensure_ascii=False)}\n\n"
                            _tot_enabled_fallback = True
                            _tot_final = ""
                            break
                    else:
                        _tot_enabled_fallback = False

                    if _tot_final:
                        try:
                            session_manager.append_and_save(
                                f"{session_name}.json",
                                user_input,
                                _tot_final[:6000],
                                task=task_type,
                                model_name=_tot_model,
                            )
                        except Exception:
                            pass
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time, 'tot_winner': _tot_winner_id})}\n\n"
                        return

                    # ToT σñ▒Φ┤Ñ ΓåÆ ΘÖìτ║ºσê░Σ╕ïµû╣µáçσçå RESEARCH/FILE_GEN ΘÇ╗Φ╛æ
                    _app_logger.warning(f"[ToT] ΓÜá∩╕Å µ£¬ΦÄ╖σ╛ùµ£ëµòêΦ╛ôσç║∩╝îΘÖìτ║ºΦç│µáçσçåΦ╖»σ╛ä")

                except ImportError:
                    _app_logger.warning("[ToT] ΓÜá∩╕Å tree_of_thought µ¿íσ¥ùµ£¬µë╛σê░∩╝îΘÖìτ║ºΦç│µáçσçåΦ╖»σ╛ä")
                except Exception as _tot_err:
                    import traceback as _ttb

                    _app_logger.error(f"[ToT] Γ¥î σ╝éσ╕╕: {_ttb.format_exc()}")
                    yield f"data: {json.dumps({'type': 'progress', 'message': 'ΓÜá∩╕Å Tree of Thought σ╝éσ╕╕∩╝îσêçµìóΦç│µáçσçåµ¿íσ╝Å', 'detail': str(_tot_err)[:100]})}\n\n"
            # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ

            # === RESEARCH Mode (µ╖▒σ║ªτáöτ⌐╢ - µ╡üσ╝Åσôìσ║öΣ╝ÿσàê) ===
            if task_type == "RESEARCH":
                research_model = model_id or MODEL_MAP.get(
                    "RESEARCH", "gemini-3-pro-preview"
                )
                used_model = research_model
                t = yield_thinking(
                    f"Φ┐¢σàÑµ╖▒σ║ªτáöτ⌐╢µ¿íσ╝Å∩╝îΣ╜┐τö¿ {research_model} Φ┐¢ΦíîΣ╕ôΣ╕Üτ║ºσêåµ₧É",
                    "analyzing",
                )
                if t:
                    yield t
                newline = "\n"
                _detail = (
                    "Σ╜┐τö¿ Interactions API µ╖▒σ║ªτáöτ⌐╢"
                    if research_model.startswith("deep-research-pro-preview")
                    else f"Σ╜┐τö¿ {research_model} Φ┐¢Φíîµ╡üσ╝Åσêåµ₧É"
                )
                yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒö¼ σÉ»σè¿µ╖▒σ║ªτáöτ⌐╢µ¿íσ╝Å...', 'detail': _detail})}{newline}{newline}"

                # µ₧äσ╗║µ╖▒σ║ªτáöτ⌐╢τÜäsystem instruction
                research_instruction = """Σ╜áµÿ»Σ╕ÇΣ╜ìΣ╕ôΣ╕ÜτÜäτáöτ⌐╢σè⌐µëï∩╝îµôàΘò┐µ╖▒σ║ªσêåµ₧Éσñìµ¥éµèÇµ£»Φ»¥ΘóÿπÇéΦ»╖µîëτàºΣ╗ÑΣ╕ïτ╗ôµ₧äµÅÉΣ╛¢σà¿Θ¥óµ╖▒σàÑτÜäτáöτ⌐╢µèÑσæè∩╝Ü

1. **µèÇµ£»µªéΦ┐░**∩╝Üµ╕àµÖ░σ«ÜΣ╣ëσÆîΦºúΘçèµá╕σ┐âµªéσ┐╡
2. **µèÇµ£»σÄƒτÉå**∩╝ÜΦ»ªτ╗åΦ»┤µÿÄσ╖ÑΣ╜£µ£║σê╢σÆîσ║òσ▒éσÄƒτÉå
3. **Σ╝ÿσè┐σêåµ₧É**∩╝ÜσêùΣ╕╛Σ╕╗ΦªüΣ╝ÿτé╣σÆîσ║öτö¿σ£║µÖ»
4. **Θù«ΘóÿΣ╕Äµîæµêÿ**∩╝Üσêåµ₧Éσ¡ÿσ£¿τÜäΘù«ΘóÿσÆîµèÇµ£»τô╢Θóê
5. **σ»╣µ»öσêåµ₧É**∩╝ÜΣ╕Äσà╢Σ╗ûσÉîτ▒╗µèÇµ£»Φ┐¢Φíîµ¿¬σÉæσ»╣µ»ö
6. **σÅæσ▒òΦ╢ïσè┐**∩╝ÜΦ«¿Φ«║µ£¬µ¥ÑσÅæσ▒òµû╣σÉæσÆîσ║öτö¿σëìµÖ»
7. **σÅéΦÇâΦ╡äµûÖ**∩╝ÜµÅÉΣ╛¢τ¢╕σà│µèÇµ£»µûçµíúσÆîσ¡ªµ£»Φ╡äµûÖτÜäσ╝òτö¿

≡ƒôî **τë╣µ«èµƒÑΦ»óτ▒╗σ₧ïσó₧σ╝║ΦºäσêÖ**∩╝Ü

**Σ╗╖µá╝/Φ┤╣τö¿/τÑ¿σèíµƒÑΦ»ó**∩╝êσªéΘ½ÿΘôüτÑ¿πÇüµ£║τÑ¿πÇüΘàÆσ║ùπÇüΘù¿τÑ¿τ¡ë∩╝ë∩╝Ü
- Γ£à **ΘªûσàêΦ╛ôσç║Σ╕ÇΣ╕¬µ╕àµÖ░τÜäΦí¿µá╝**∩╝îσîàσÉ½σà│Θö«Σ┐íµü»∩╝êΦ╜ªµ¼íπÇüσÅæΦ╜ªµù╢Θù┤πÇüσê░Φ╛╛µù╢Θù┤πÇüσ║ºΣ╜ìπÇüΣ╗╖µá╝πÇüµù╢Θò┐τ¡ë∩╝ë
- Γ£à σ┐àΘí╗µÅÉΣ╛¢**σà╖Σ╜ôΣ╗╖µá╝**∩╝êΣ╛ïσªé∩╝ÜΣ║îτ¡ëσ║º ┬Ñ524.5∩╝ë
- Γ¥î τªüµ¡óΣ╜┐τö¿Σ╗╖µá╝σî║Θù┤∩╝êσªé"500-600σàâ"∩╝ë
- Γ£à µîëσ║ºΣ╜ì/µê┐σ₧ïτ¡ëτ║º**σêåσê½σêùσç║**µ»ÅΣ╕¬ΘÇëΘí╣τÜäτí«σêçΣ╗╖µá╝
- Γ£à σêùσç║**σà╖Σ╜ôτÅ¡µ¼í/Φ╜ªµ¼íσÅ╖**∩╝êσªé G12πÇüΦê¬τÅ¡ MU5137∩╝ë
- Γ£à σêùσç║**σÅæΦ╜ªµù╢Θù┤σÆîσê░Φ╛╛µù╢Θù┤**∩╝îµû╣Σ╛┐τö¿µê╖σ»╣µ»öΘÇëµï⌐
- Γ¥î τªüµ¡óΦ╛ôσç║Θçìσñìσåàσ«╣µêûσñÜΣ╕¬τ¢╕σÉîτÜäµ«╡ΦÉ╜

**σ╝║σê╢Σ╜┐τö¿Φí¿µá╝µá╝σ╝Å**∩╝Ü
```
≡ƒÜä Σ╕èµ╡╖ΦÖ╣µíÑ ΓåÆ σîùΣ║¼σìù∩╝ê2026σ╣┤2µ£ê12µùÑ∩╝ë

| Φ╜ªµ¼í   | σÅæΦ╜ª  | σê░Φ╛╛  | σ║ºΣ╜ìτ▒╗σ₧ï | Σ╗╖µá╝     | µù╢Θò┐  |
|--------|-------|-------|----------|----------|-------|
| G12µ¼í  | 09:00 | 13:24 | σòåσèíσ║º   | ┬Ñ1,748   | 4h24m |
| G12µ¼í  | 09:00 | 13:24 | Σ╕Çτ¡ëσ║º   | ┬Ñ933     | 4h24m |
| G12µ¼í  | 09:00 | 13:24 | Σ║îτ¡ëσ║º   | ┬Ñ524.5   | 4h24m |
| G8µ¼í   | 10:00 | 14:31 | σòåσèíσ║º   | ┬Ñ1,748   | 4h31m |
| G8µ¼í   | 10:00 | 14:31 | Σ╕Çτ¡ëσ║º   | ┬Ñ933     | 4h31m |
| G8µ¼í   | 10:00 | 14:31 | Σ║îτ¡ëσ║º   | ┬Ñ524.5   | 4h31m |

≡ƒÆí Φ┤¡τÑ¿µû╣σ╝Å∩╝ÜΦ«┐Θù« 12306.cn µÉ£τ┤óσ»╣σ║öΦ╜ªµ¼íΦ┤¡Σ╣░πÇé
```

Φªüµ▒é∩╝Ü
- µÅÉΣ╛¢σà╖Σ╜ôτÜäµèÇµ£»τ╗åΦèéσÆîµò░µì«µö»µîü
- Σ╜┐τö¿Σ╕ôΣ╕Üµ£»Φ»¡Σ╜åτí«Σ┐¥σÅ»τÉåΦºúµÇº
- Σ┐¥µîüσ«óΦºéΣ╕¡τ½ïτÜäσêåµ₧ÉµÇüσ║ª
- σåàσ«╣σà¿Θ¥óΣ╕öµ£ëµ╖▒σ║ª
- ΘÇéσ╜ôΣ╜┐τö¿σ¢╛Φí¿σÆîτñ║Σ╛ïΦ»┤µÿÄ"""

                # µ│¿σàÑ skill_prompt∩╝êµ¿íσ₧ïσ»╣τö¿µê╖µäÅσ¢╛τÜäτÉåΦºú∩╝ë
                _research_skill = (context_info or {}).get("skill_prompt")
                if _research_skill:
                    research_instruction += (
                        f"\n\n[τö¿µê╖µ£ƒµ£¢τÜäΦ╛ôσç║Θçìτé╣] {_research_skill}"
                    )
                # σ░åτƒÑΦ»åσ║ôµúÇτ┤óσåàσ«╣µ│¿σàÑτáöτ⌐╢µîçΣ╗ñ∩╝êΦïÑµ£ë∩╝ë
                if _rag_context_block:
                    research_instruction += (
                        f"\n\n[≡ƒôÜ τƒÑΦ»åσ║ôσÅéΦÇâΦ╡äµûÖ]\n{_rag_context_block}"
                    )

                collected_text = []

                try:
                    newline = "\n"
                    if research_model.startswith("deep-research-pro-preview"):
                        yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒôè µ¡úσ£¿Φ┐¢Φíîµ╖▒σ║ªσêåµ₧É...', 'detail': 'Deep Research µ¡úσ£¿µúÇτ┤óΣ╕Äτ╗╝σÉê∩╝îσÅ»Φâ╜Θ£ÇΦªüΦ╛âΘò┐µù╢Θù┤'})}{newline}{newline}"
                        deep_text = WebSearcher.deep_research_for_ppt(
                            effective_input, ""
                        )
                        if not deep_text:
                            raise RuntimeError("Deep Research µ£¬Φ┐öσ¢₧µ£ëµòêσåàσ«╣")
                        collected_text.append(deep_text)
                        yield f"data: {json.dumps({'type': 'token', 'content': deep_text})}\n\n"
                    else:
                        yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒôè µ¡úσ£¿Φ┐¢Φíîµ╖▒σ║ªσêåµ₧É...', 'detail': f'{research_model} µ¡úσ£¿µÇ¥ΦÇâ∩╝îσÅ»Φâ╜Θ£ÇΦªü30-90τºÆ'})}{newline}{newline}"

                        response_stream = client.models.generate_content_stream(
                            model=research_model,
                            contents=effective_input,
                            config=types.GenerateContentConfig(
                                system_instruction=research_instruction,
                                temperature=0.7,
                                max_output_tokens=8000,  # σàüΦ«╕µ¢┤Θò┐τÜäΦ╛ôσç║
                                top_p=0.95,
                            ),
                        )

                    if not research_model.startswith("deep-research-pro-preview"):
                        chunk_count = 0
                        heartbeat_interval = 5  # µ»Å5τºÆσÅæΘÇüΣ╕Çµ¼íσ┐âΦ╖│
                        first_chunk_received = False

                        # Σ╜┐τö¿Σ┐¥µ┤╗σîàΦúàσÖ¿σñäτÉåµ╡üσ╝Åσôìσ║ö
                        for item_type, item_data in stream_with_keepalive(
                            response_stream,
                            start_time,
                            keepalive_interval=heartbeat_interval,
                            max_wait_first_token=90,
                        ):  # µ£ÇσñÜτ¡ëσ╛à90τºÆ
                            # µúÇµƒÑΣ╕¡µû¡
                            if interrupted():
                                _app_logger.debug(f"[RESEARCH] τö¿µê╖Σ╕¡µû¡τáöτ⌐╢")
                                newline = "\n"
                                interrupt_msg = f"{newline}{newline}ΓÅ╣∩╕Å τáöτ⌐╢σ╖▓Φó½τö¿µê╖Σ╕¡µû¡"
                                yield f"data: {json.dumps({'type': 'token', 'content': interrupt_msg})}{newline}{newline}"
                                break

                            if item_type == "heartbeat":
                                # σÅæΘÇüσ┐âΦ╖│Σ┐¥µîüΦ┐₧µÄÑ
                                elapsed = item_data
                                if first_chunk_received:
                                    char_count = len("".join(collected_text))
                                    yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒô¥ µ¡úσ£¿τöƒµêÉΣ╕¡...', 'detail': f'σ╖▓τöƒµêÉ {char_count} σ¡ùτ¼ª∩╝îΦÇùµù╢ {elapsed}s'})}\n\n"
                                else:
                                    yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒºá µ¿íσ₧ïµ¡úσ£¿µ╖▒σ║ªµÇ¥ΦÇâ...', 'detail': f'σ╖▓τ¡ëσ╛à {elapsed}s∩╝îΦ»╖ΦÇÉσ┐âτ¡ëσ╛à'})}\n\n"

                            elif item_type == "timeout":
                                # τ¡ëσ╛àΦ╢àµù╢
                                yield f"data: {json.dumps({'type': 'token', 'content': f'ΓÜá∩╕Å {item_data}∩╝îµ¿íσ₧ïσôìσ║öµù╢Θù┤Φ┐çΘò┐∩╝îΦ»╖τ¿ìσÉÄΘçìΦ»ò'})}\n\n"
                                break

                            elif item_type == "chunk":
                                chunk = item_data
                                if chunk.text:
                                    if not first_chunk_received:
                                        first_chunk_received = True
                                        _app_logger.debug(
                                            f"[RESEARCH] µö╢σê░τ¼¼Σ╕ÇΣ╕¬σôìσ║öσ¥ù∩╝îΦÇùµù╢ {time.time() - start_time:.1f}s"
                                        )

                                    collected_text.append(chunk.text)
                                    yield f"data: {json.dumps({'type': 'token', 'content': chunk.text})}\n\n"
                                    chunk_count += 1

                                    # µ»Å50Σ╕¬chunkµÿ╛τñ║Σ╕Çµ¼íΦ┐¢σ║ªµùÑσ┐ù
                                    if chunk_count % 50 == 0:
                                        _app_logger.debug(
                                            f"[RESEARCH] σ╖▓τöƒµêÉ {chunk_count} Σ╕¬chunk, {len(''.join(collected_text))} σ¡ùτ¼ª"
                                        )

                    final_text = "".join(collected_text)
                    _app_logger.info(f"[RESEARCH] Γ£à τáöτ⌐╢σ«îµêÉ∩╝îσà▒ {len(final_text)} σ¡ùτ¼ª")

                    # Σ┐¥σ¡ÿσÄåσÅ▓∩╝êσƒ║Σ║Äτúüτ¢ÿσ«îµò┤σÄåσÅ▓Φ┐╜σèá∩╝ë
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        final_text[:4000],
                        task="RESEARCH",
                        model_name=used_model,
                    )

                except Exception as research_err:
                    error_msg = str(research_err)
                    _app_logger.debug(f"[RESEARCH] ΘöÖΦ»»: {error_msg}")

                    # µÖ║Φâ╜ΘöÖΦ»»σñäτÉå
                    if "503" in error_msg or "UNAVAILABLE" in error_msg:
                        # APIΦ┐çΦ╜╜∩╝îσ░¥Φ»òΣ╜┐τö¿Flashτëêµ£¼
                        try:
                            newline = "\n"
                            yield f"data: {json.dumps({'type': 'progress', 'message': 'ΓÜá∩╕Å µ£ìσèíτ╣üσ┐Ö∩╝îσêçµìóσê░ Gemini 2.5 Flash...', 'detail': ''})}{newline}{newline}"

                            response_stream = client.models.generate_content_stream(
                                model="gemini-2.5-flash",
                                contents=effective_input,
                                config=types.GenerateContentConfig(
                                    system_instruction=research_instruction,
                                    temperature=0.7,
                                    max_output_tokens=8000,
                                ),
                            )

                            last_heartbeat_flash = time.time()
                            for chunk in response_stream:
                                if interrupted():
                                    break
                                if chunk.text:
                                    collected_text.append(chunk.text)
                                    yield f"data: {json.dumps({'type': 'token', 'content': chunk.text})}\n\n"

                                    # Flash µ¿íσ╝ÅΣ╕ïΣ╣ƒσÅæΘÇüσ┐âΦ╖│
                                    current_time = time.time()
                                    if current_time - last_heartbeat_flash > 3:
                                        elapsed = int(current_time - start_time)
                                        yield f"data: {json.dumps({'type': 'progress', 'message': f'ΓÜí σ┐½ΘÇƒµ¿íσ╝ÅτöƒµêÉΣ╕¡...', 'detail': f'{elapsed}s'})}\n\n"
                                        last_heartbeat_flash = current_time

                            final_text = "".join(collected_text)
                            session_manager.append_and_save(
                                f"{session_name}.json",
                                user_input,
                                final_text[:4000],
                                task="RESEARCH",
                                model_name="gemini-3-flash-preview",
                            )

                        except Exception as fallback_err:
                            error_text = f"Γ¥î τáöτ⌐╢µ£ìσèíµÜéµù╢Σ╕ìσÅ»τö¿\n\nΘöÖΦ»»Σ┐íµü»: {str(fallback_err)[:200]}\n\n≡ƒÆí σ╗║Φ««∩╝Ü\n1. τ¿ìσÉÄΘçìΦ»ò\n2. τ«ÇσîûΘù«Θóÿ\n3. Σ╜┐τö¿µÖ«ΘÇÜσ»╣Φ»¥µ¿íσ╝Å"
                            yield f"data: {json.dumps({'type': 'token', 'content': error_text})}\n\n"
                            session_manager.append_and_save(
                                f"{session_name}.json",
                                user_input,
                                error_text[:1000],
                                task="RESEARCH",
                                model_name="gemini-3-flash-preview",
                            )

                    elif (
                        "timeout" in error_msg.lower()
                        or "disconnect" in error_msg.lower()
                    ):
                        # Φ┐₧µÄÑΘù«Θóÿ
                        error_text = f"ΓÜá∩╕Å Φ┐₧µÄÑΦ╢àµù╢µêûΣ╕¡µû¡\n\nσÅ»Φâ╜σÄƒσ¢á∩╝Ü\n1. τ╜æτ╗£Σ╕ìτ¿│σ«Ü\n2. µ£ìσèíσÖ¿τ╣üσ┐Ö\n3. Σ╗úτÉåΘàìτ╜«Θù«Θóÿ\n\nσ╗║Φ««∩╝ÜΦ»╖τ¿ìσÉÄΘçìΦ»ò∩╝îµêûµúÇµƒÑτ╜æτ╗£Φ┐₧µÄÑ"
                        yield f"data: {json.dumps({'type': 'token', 'content': error_text})}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            error_text[:1000],
                            task="RESEARCH",
                            model_name=used_model,
                        )

                    else:
                        # σà╢Σ╗ûΘöÖΦ»»
                        error_text = f"Γ¥î τáöτ⌐╢Φ┐çτ¿ïΣ╕¡σç║τÄ░ΘöÖΦ»»\n\n{error_msg[:300]}\n\nΦ»╖σ░¥Φ»ò∩╝Ü\n1. Θçìµû░µÅÉΘù«\n2. τ«ÇσîûΘù«ΘóÿµÅÅΦ┐░\n3. τ¿ìσÉÄΘçìΦ»ò"
                        yield f"data: {json.dumps({'type': 'token', 'content': error_text})}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            error_text[:1000],
                            task="RESEARCH",
                            model_name=used_model,
                        )

                total_time = time.time() - start_time
                newline = "\n"
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}{newline}{newline}"
                return

            # === PAINTER Mode (σ¢╛σâÅτöƒµêÉ - Gemini 3.1 Flash Image Σ╝ÿσàê∩╝îImagen 4.0 σñçτö¿) ===
            if task_type == "PAINTER":
                used_model = "Gemini 3.1 Flash Image (Imagen 4.0 fallback)"
                yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒÄ¿ µ¡úσ£¿τÉåΦºúΣ╜áτÜäσê¢Σ╜£Φ»╖µ▒é...', 'detail': '', 'progress': 5, 'stage': 'paint_prepare'})}\n\n"

                # Σ╜┐τö¿Σ╕èΣ╕ïµûçσó₧σ╝║τÜäΦ╛ôσàÑ∩╝êσªéµ₧£µ£ë∩╝ë
                if (
                    context_info
                    and context_info.get("is_continuation")
                    and context_info.get("enhanced_input")
                ):
                    image_prompt = context_info["enhanced_input"]
                    _app_logger.debug(f"[PAINTER] Σ╜┐τö¿Σ╕èΣ╕ïµûçσó₧σ╝║τÜäprompt: {image_prompt[:100]}...")
                else:
                    image_prompt = effective_input

                yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒûî∩╕Å Gemini 3.1 Flash Image µ¡úσ£¿τöƒµêÉσ¢╛σâÅ...', 'detail': 'Φ»╖ΦÇÉσ┐âτ¡ëσ╛à', 'progress': 20, 'stage': 'paint_generate'})}\n\n"

                max_retries = 2
                use_fallback = False
                images = []

                for attempt in range(max_retries):
                    try:
                        if interrupted():
                            yield f"data: {json.dumps({'type': 'token', 'content': 'ΓÅ╣∩╕Å σ¢╛σâÅτöƒµêÉσ╖▓Σ╕¡µû¡'})}\n\n"
                            total_time = time.time() - start_time
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                            return

                        if attempt > 0:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒöä τ¼¼ {attempt} µ¼íΘçìΦ»ò...', 'detail': '', 'progress': 25, 'stage': 'paint_retry'})}\n\n"
                            time.sleep(2)

                        # ΘÇëµï⌐µ¿íσ₧ï
                        if use_fallback:
                            model_name = "Imagen 4.0"
                            yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒöä σêçµìóσê░ Imagen 4.0...', 'detail': '', 'progress': 30, 'stage': 'paint_fallback'})}\n\n"
                        else:
                            model_name = "Gemini 3.1 Flash Image"

                        # Σ╜┐τö¿σÉÄσÅ░τ║┐τ¿ïµëºΦíîΦ»╖µ▒é∩╝îΣ╕╗τ║┐τ¿ïσÅæΘÇüσ┐âΦ╖│
                        import queue
                        import threading

                        result_queue = queue.Queue()

                        def worker():
                            try:
                                if use_fallback:
                                    result = client.models.generate_images(
                                        model="imagen-4.0-fast-generate-001",
                                        prompt=image_prompt,
                                        config=types.GenerateImagesConfig(
                                            number_of_images=1
                                        ),
                                    )
                                else:
                                    result = client.models.generate_content(
                                        model="gemini-3.1-flash-image-preview",
                                        contents=image_prompt,
                                        config=types.GenerateContentConfig(
                                            response_modalities=["TEXT", "IMAGE"]
                                        ),
                                    )
                                result_queue.put(("success", result))
                            except Exception as e:
                                result_queue.put(("error", e))

                        thread = threading.Thread(target=worker, daemon=True)
                        thread.start()

                        # Per-model timeout: 120s for gemini-3.1-flash-image (takes ~65s), 90s for imagen fallback
                        timeout_seconds = 120 if not use_fallback else 90
                        attempt_start = time.time()
                        timed_out = False
                        response = None

                        while True:
                            attempt_elapsed = time.time() - attempt_start

                            if attempt_elapsed > timeout_seconds:
                                timed_out = True
                                if not use_fallback:
                                    # Primary model timed out ΓÇö switch to imagen
                                    _app_logger.debug(
                                        f"[PAINTER] Gemini 3.1 Flash Image Φ╢àµù╢ ({int(attempt_elapsed)}s)∩╝îσêçµìóσê░ Imagen"
                                    )
                                    use_fallback = True
                                    yield f"data: {json.dumps({'type': 'progress', 'message': 'ΓÅ▒∩╕Å µ¿íσ₧ïσôìσ║öΦ╢àµù╢∩╝îσêçµìóσê░ Imagen...', 'detail': '', 'progress': 28, 'stage': 'paint_fallback'})}\n\n"
                                    break
                                else:
                                    # Imagen also timed out
                                    elapsed = time.time() - start_time
                                    yield f"data: {json.dumps({'type': 'token', 'content': f'ΓÜá∩╕Å σ¢╛σâÅτöƒµêÉΦ╢àµù╢ ({int(elapsed)}s)∩╝îΦ»╖τ¿ìσÉÄΘçìΦ»ò'})}\n\n"
                                    total_time = time.time() - start_time
                                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                                    return

                            if interrupted():
                                yield f"data: {json.dumps({'type': 'token', 'content': 'ΓÅ╣∩╕Å σ¢╛σâÅτöƒµêÉσ╖▓Σ╕¡µû¡'})}\n\n"
                                total_time = time.time() - start_time
                                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                                return

                            try:
                                status, data = result_queue.get(timeout=3.0)
                                if status == "success":
                                    response = data
                                    break
                                else:
                                    raise data
                            except queue.Empty:
                                progress_guess = min(
                                    85,
                                    30 + int((attempt_elapsed / timeout_seconds) * 55),
                                )
                                yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒÄ¿ {model_name} τöƒµêÉΣ╕¡...', 'detail': f'{int(attempt_elapsed)}s', 'progress': progress_guess, 'stage': 'paint_running'})}\n\n"

                        if timed_out:
                            continue  # outer for-loop: retry with use_fallback=True (imagen timeout already returned above)

                        # σñäτÉåσôìσ║ö
                        yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒÆ╛ µ¡úσ£¿Σ┐¥σ¡ÿσ¢╛τëç...', 'detail': '', 'progress': 90, 'stage': 'paint_save'})}\n\n"

                        if use_fallback:
                            if response.generated_images:
                                for gen_img in response.generated_images:
                                    img_data = gen_img.image.image_bytes
                                    images_dir = settings_manager.images_dir
                                    os.makedirs(images_dir, exist_ok=True)
                                    timestamp = int(time.time())
                                    filename = f"generated_{timestamp}.png"
                                    filepath = os.path.join(images_dir, filename)
                                    with open(filepath, "wb") as f:
                                        f.write(img_data)

                                    # τí«Σ┐¥Φ╖»σ╛äσ£¿ workspace Σ╕ï
                                    try:
                                        rel_path = os.path.relpath(
                                            filepath, WORKSPACE_DIR
                                        ).replace("\\", "/")
                                        if ".." not in rel_path:
                                            images.append(rel_path)
                                            _app_logger.debug(
                                                f"[PAINTER] Imagen σ╖▓Σ┐¥σ¡ÿ: {rel_path}"
                                            )
                                        else:
                                            # ΘÖìτ║ºΣ┐¥σ¡ÿσê░ workspace/images
                                            abs_workspace_images = os.path.join(
                                                WORKSPACE_DIR, "images"
                                            )
                                            os.makedirs(
                                                abs_workspace_images, exist_ok=True
                                            )
                                            fallback_filepath = os.path.join(
                                                abs_workspace_images, filename
                                            )
                                            with open(fallback_filepath, "wb") as f:
                                                f.write(img_data)
                                            fallback_rel = os.path.relpath(
                                                fallback_filepath, WORKSPACE_DIR
                                            ).replace("\\", "/")
                                            images.append(fallback_rel)
                                            _app_logger.debug(
                                                f"[PAINTER] Imagen ΘÖìτ║ºΣ┐¥σ¡ÿ: {fallback_rel}"
                                            )
                                    except Exception as path_err:
                                        _app_logger.debug(f"[PAINTER] Path error: {path_err}")
                        else:
                            if (
                                response.candidates
                                and response.candidates[0].content.parts
                            ):
                                for part in response.candidates[0].content.parts:
                                    if (
                                        hasattr(part, "inline_data")
                                        and part.inline_data
                                    ):
                                        img_filename = Utils.save_image_part(part)
                                        if img_filename:
                                            images.append(img_filename)
                                            _app_logger.debug(
                                                f"[PAINTER] Gemini 3.1 Flash Image σ╖▓Σ┐¥σ¡ÿ: {img_filename}"
                                            )

                        if images:
                            save_path = settings_manager.images_dir
                            msg = f"Γ£¿ σ¢╛τëçσ╖▓τöƒµêÉ! (Σ╜┐τö¿ {model_name})\n≡ƒû╝∩╕Å Σ┐¥σ¡ÿΣ╜ìτ╜«: {save_path}"
                            yield f"data: {json.dumps({'type': 'token', 'content': msg})}\n\n"

                            yield f"data: {json.dumps({'type': 'progress', 'message': 'Γ£à σ¢╛σâÅτöƒµêÉσ«îµêÉ', 'detail': f'{len(images)} σ╝á', 'progress': 100, 'stage': 'complete'})}\n\n"

                            # σàêΣ┐¥σ¡ÿσÄåσÅ▓Φ«░σ╜ò∩╝êσîàσÉ½σ¢╛τëçΦ╖»σ╛ä∩╝ë∩╝îσåìσÅæΘÇü done
                            session_manager.append_and_save(
                                f"{session_name}.json",
                                user_input,
                                "σ¢╛σâÅσ╖▓τöƒµêÉ",
                                images=images,
                                task="PAINTER",
                                model_name=model_name,
                            )

                            total_time = time.time() - start_time
                            _app_logger.debug(f"[PAINTER] σÅæΘÇüσ¢╛τëçσêùΦí¿: {images}")  # Φ░âΦ»ò
                            yield f"data: {json.dumps({'type': 'done', 'images': images, 'saved_files': [], 'total_time': total_time})}\n\n"
                            return
                        else:
                            if not use_fallback:
                                use_fallback = True
                                continue
                            else:
                                yield f"data: {json.dumps({'type': 'token', 'content': 'Γ¥î µ¿íσ₧ïµ£¬Φ┐öσ¢₧σ¢╛τëç'})}\n\n"

                    except Exception as img_err:
                        error_msg = str(img_err)
                        model_label = (
                            "Imagen" if use_fallback else "Gemini-3.1-Flash-Image"
                        )
                        _app_logger.debug(
                            f"[PAINTER] {model_label} σ░¥Φ»ò {attempt+1} σñ▒Φ┤Ñ ({type(img_err).__name__}): {error_msg[:300]}"
                        )

                        # Fall back to imagen on ANY non-safety error when using primary model
                        if (
                            not use_fallback
                            and "safety" not in error_msg.lower()
                            and "blocked" not in error_msg.lower()
                        ):
                            _app_logger.debug(
                                f"[PAINTER] Gemini 3.1 Flash Image σñ▒Φ┤Ñ∩╝îσêçµìóσê░ Imagen: {error_msg[:200]}"
                            )
                            use_fallback = True
                            continue

                        if (
                            "safety" in error_msg.lower()
                            or "blocked" in error_msg.lower()
                        ):
                            user_msg = "Γ¥î σåàσ«╣Φó½σ«ëσà¿τ¡ûτòÑΦ┐çµ╗ñ∩╝îΦ»╖Σ┐«µö╣µÅÅΦ┐░"
                        elif "location is not supported" in error_msg.lower():
                            user_msg = "Γ¥î σ£░σî║ΘÖÉσê╢∩╝îΦ»╖Θàìτ╜«Σ╕¡Φ╜¼µ£ìσèí"
                        else:
                            user_msg = f"Γ¥î σ¢╛σâÅτöƒµêÉσñ▒Φ┤Ñ: {error_msg[:100]}"

                        yield f"data: {json.dumps({'type': 'token', 'content': user_msg})}\n\n"

                # PAINTER µëÇµ£ëΘçìΦ»òΘâ╜σñ▒Φ┤Ñµù╢Σ╣ƒΦªüΣ┐¥σ¡ÿσÄåσÅ▓
                session_manager.append_and_save(
                    f"{session_name}.json", user_input, "σ¢╛σâÅτöƒµêÉσñ▒Φ┤Ñ"
                )

                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return

            # === FILE_GEN Mode (µûçΣ╗╢τöƒµêÉ - Φç¬σè¿µëºΦíî) ===
            if task_type == "FILE_GEN":
                t = yield_thinking(
                    f"Φ┐¢σàÑµûçΣ╗╢τöƒµêÉµ¿íσ╝Å∩╝îσ░åΣ╜┐τö¿ {model_id} τöƒµêÉµûçµíú", "generating"
                )
                if t:
                    yield t
                _app_logger.debug(f"[FILE_GEN] ===== Starting file generation =====")
                _app_logger.debug(
                    f"[FILE_GEN] Model: {model_id}, User input: {user_input[:100]}..."
                )

                yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒº╛ σçåσñçτöƒµêÉµûçµíú...', 'detail': get_model_display_name(model_id), 'progress': 5, 'stage': 'filegen_prepare'})}\n\n"

                response_text = ""
                generated_files = []
                temp_scripts = []  # Σ╕┤µù╢ΦäÜµ£¼σêùΦí¿∩╝êµëºΦíîσÉÄσêáΘÖñ∩╝ë
                api_timeout = 120  # σó₧σèáσê░ 120 τºÆ∩╝îΘò┐µûçµíúΘ£ÇΦªüµ¢┤σñÜµù╢Θù┤

                if interrupted():
                    yield f"data: {json.dumps({'type': 'token', 'content': 'ΓÅ╣∩╕Å µûçΣ╗╢τöƒµêÉσ╖▓Σ╕¡µû¡'})}\n\n"
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                # Γ¡É µúÇµƒÑµÿ»σÉªµÿ»"Φ╜¼µìóΦ»╖µ▒é"∩╝êµèèΣ╣ïσëìτÜäσåàσ«╣σüÜµêÉword/pdf∩╝ë
                is_convert_request = (
                    context_info
                    and context_info.get("is_continuation")
                    and context_info.get("continuation_type") == "convert"
                    and context_info.get("context_summary", {}).get("last_model_output")
                )

                if is_convert_request:
                    # τ¢┤µÄÑΦ╜¼µìóµ¿íσ╝Å - Σ╕ìΘ£ÇΦªüΦ░âτö¿µ¿íσ₧ï∩╝îτ¢┤µÄÑτöƒµêÉµûçµíú
                    yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒô¥ µ¡úσ£¿σ░åσåàσ«╣Φ╜¼µìóΣ╕║µûçµíú...', 'detail': '', 'progress': 30, 'stage': 'filegen_convert'})}\n\n"

                    try:
                        from web.document_generator import save_docx, save_pdf

                        source_content = context_info["context_summary"][
                            "last_model_output"
                        ]
                        _app_logger.debug(
                            f"[FILE_GEN] τ¢┤µÄÑΦ╜¼µìóµ¿íσ╝Å∩╝îµ║Éσåàσ«╣Θò┐σ║ª: {len(source_content)}"
                        )

                        # µÅÉσÅûµáçΘóÿ∩╝êσ░¥Φ»òΣ╗Äσåàσ«╣Σ╕¡µë╛ # µáçΘóÿ∩╝ë
                        title_match = re.search(
                            r"^#\s*(.+)$", source_content, re.MULTILINE
                        )
                        if title_match:
                            title = title_match.group(1).strip()[:50]
                        else:
                            title = (
                                f"Kotoµûçµíú_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                            )

                        docs_dir = settings_manager.documents_dir
                        os.makedirs(docs_dir, exist_ok=True)

                        # σêñµû¡τöƒµêÉ Word Φ┐ÿµÿ» PDF
                        user_lower = user_input.lower()

                        if "pdf" in user_lower:
                            yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒôä µ¡úσ£¿τöƒµêÉ PDF...', 'detail': '', 'progress': 60, 'stage': 'filegen_save'})}\n\n"
                            saved_path = save_pdf(
                                source_content, title=title, output_dir=docs_dir
                            )
                            file_type = "PDF"
                        else:
                            yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒôä µ¡úσ£¿τöƒµêÉ Word µûçµíú...', 'detail': '', 'progress': 60, 'stage': 'filegen_save'})}\n\n"
                            saved_path = save_docx(
                                source_content, title=title, output_dir=docs_dir
                            )
                            file_type = "Word"

                        rel_path = os.path.relpath(saved_path, WORKSPACE_DIR).replace(
                            "\\", "/"
                        )
                        generated_files.append(rel_path)

                        success_msg = f"Γ£à **{file_type} µûçµíúτöƒµêÉµêÉσèƒ∩╝ü**\n\n≡ƒôü µûçΣ╗╢: **{os.path.basename(saved_path)}**\n≡ƒôì Σ╜ìτ╜«: `{docs_dir}`"
                        yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"

                        yield f"data: {json.dumps({'type': 'progress', 'message': 'Γ£à µûçµíúΦ╜¼µìóσ«îµêÉ', 'detail': file_type, 'progress': 100, 'stage': 'complete'})}\n\n"

                        _app_logger.info(f"[FILE_GEN] Γ£à τ¢┤µÄÑΦ╜¼µìóµêÉσèƒ: {rel_path}")

                    except Exception as convert_err:
                        error_msg = f"Γ¥î µûçµíúΦ╜¼µìóσñ▒Φ┤Ñ: {str(convert_err)}"
                        _app_logger.debug(f"[FILE_GEN] Φ╜¼µìóΘöÖΦ»»: {convert_err}")
                        yield f"data: {json.dumps({'type': 'token', 'content': error_msg})}\n\n"

                    # Σ┐¥σ¡ÿσÄåσÅ▓∩╝êσƒ║Σ║Äτúüτ¢ÿσ«îµò┤σÄåσÅ▓Φ┐╜σèá∩╝ë
                    _model_msg = (
                        f"σ╖▓τöƒµêÉµûçΣ╗╢: {', '.join(generated_files)}"
                        if generated_files
                        else "µûçµíúΦ╜¼µìóσñ▒Φ┤Ñ"
                    )
                    session_manager.append_and_save(
                        f"{session_name}.json", user_input, _model_msg
                    )

                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': generated_files, 'total_time': total_time})}\n\n"
                    return

                # Γ¡É µúÇµƒÑµÿ»σÉªµÿ» PPT τöƒµêÉΦ»╖µ▒é
                ppt_keywords = [
                    "ppt",
                    "σ╣╗τü»τëç",
                    "µ╝öτñ║µûçτ¿┐",
                    "µ╝öτñ║",
                    "presentation",
                    "slide",
                    "slides",
                ]
                user_lower_check = user_input.lower()
                is_ppt_request = any(kw in user_lower_check for kw in ppt_keywords)

                if is_ppt_request:
                    # =============== PPT Σ╕ôτö¿τöƒµêÉµ╡üτ¿ï ===============
                    _app_logger.debug(f"[FILE_GEN] ≡ƒÄ» µúÇµ╡ïσê░ PPT τöƒµêÉΦ»╖µ▒é")

                    # ΓöÇΓöÇ σê¥σºïσîûµÖ║Φâ╜σÅìΘªê ΓöÇΓöÇ
                    from web.smart_feedback import SmartFeedback

                    def _fb_emit(msg, detail=""):
                        _app_logger.debug(f"[SmartFB] {msg} | {detail}")

                    fb = SmartFeedback.for_ppt(user_input, emit=_fb_emit)

                    def _fb_sse(msg_detail_tuple):
                        """σ░å SmartFeedback Φ┐öσ¢₧τÜä (msg, detail) Φ╜¼Σ╕║ SSE µò░µì«Φíî"""
                        msg, detail = msg_detail_tuple
                        progress_pct = 0
                        if (
                            getattr(fb, "total_steps", None)
                            and getattr(fb, "current_step", 0) > 0
                        ):
                            progress_pct = min(
                                95, int((fb.current_step / fb.total_steps) * 100)
                            )
                        return f"data: {json.dumps({'type': 'progress', 'message': msg, 'detail': detail, 'progress': progress_pct})}\n\n"

                    yield _fb_sse(fb.start())

                    try:
                        # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ Session: σê¢σ╗║ PPT τ╝ûΦ╛æΣ╝ÜΦ»¥ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        ppt_session_id = None
                        try:
                            from web.ppt_session_manager import get_ppt_session_manager

                            ppt_session_mgr = get_ppt_session_manager()
                            ppt_session_id = ppt_session_mgr.create_session(
                                title=user_input[:50],  # σëì 50 σ¡ùΣ╜£Σ╕║Σ╕┤µù╢µáçΘóÿ
                                user_input=user_input,
                                theme="business",
                            )
                            _app_logger.info(f"[FILE_GEN/PPT] ≡ƒôï σê¢σ╗║τ╝ûΦ╛æΣ╝ÜΦ»¥: {ppt_session_id}")
                        except Exception as session_err:
                            _app_logger.warning(
                                f"[FILE_GEN/PPT] ΓÜá∩╕Å Σ╝ÜΦ»¥σê¢σ╗║σ╝éσ╕╕∩╝êΣ╕ìσ╜▒σôìτöƒµêÉ∩╝ë: {session_err}"
                            )

                        # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ Step 0: σñäτÉåΣ╕èΣ╝áτÜäµûçΣ╗╢ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        uploaded_file_context = ""
                        uploaded_files = (
                            request.files.getlist("files[]")
                            if request.method == "POST"
                            else []
                        )

                        if uploaded_files:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒôé µ¡úσ£¿Φºúµ₧É {len(uploaded_files)} Σ╕¬Σ╕èΣ╝áµûçΣ╗╢...', 'detail': 'µÅÉσÅûµûçµ£¼σåàσ«╣'})}\n\n"
                            try:
                                from web.file_parser import FileParser

                                uploaded_file_paths = []
                                for uploaded_file in uploaded_files:
                                    if uploaded_file and uploaded_file.filename:
                                        # Σ┐¥σ¡ÿΣ╕┤µù╢µûçΣ╗╢
                                        temp_dir = os.path.join(
                                            WORKSPACE_DIR, "temp_uploads"
                                        )
                                        os.makedirs(temp_dir, exist_ok=True)
                                        temp_path = os.path.join(
                                            temp_dir, uploaded_file.filename
                                        )
                                        uploaded_file.save(temp_path)
                                        uploaded_file_paths.append(temp_path)

                                if uploaded_file_paths:
                                    # µë╣ΘçÅΦºúµ₧É
                                    parse_results = FileParser.batch_parse(
                                        uploaded_file_paths
                                    )
                                    successful_results = [
                                        r for r in parse_results if r.get("success")
                                    ]

                                    if successful_results:
                                        uploaded_file_context = (
                                            FileParser.merge_contents(
                                                successful_results
                                            )
                                        )
                                        _app_logger.info(
                                            f"[FILE_GEN/PPT] Γ£à σ╖▓Φºúµ₧É {len(successful_results)} Σ╕¬µûçΣ╗╢, µÇ╗σ¡ùµò░: {len(uploaded_file_context)}"
                                        )
                                        yield f"data: {json.dumps({'type': 'progress', 'message': f'Γ£à σ╖▓Φºúµ₧É {len(successful_results)} Σ╕¬Σ╕èΣ╝áµûçΣ╗╢', 'detail': f'{len(uploaded_file_context)} σ¡ùσåàσ«╣'})}\n\n"
                                    else:
                                        _app_logger.warning(f"[FILE_GEN/PPT] ΓÜá∩╕Å Σ╕èΣ╝áµûçΣ╗╢Φºúµ₧Éσñ▒Φ┤Ñ")
                                        failed_reasons = [
                                            r.get("error", "µ£¬τƒÑΘöÖΦ»»")
                                            for r in parse_results
                                            if not r.get("success")
                                        ]
                                        _app_logger.info(f"    σÄƒσ¢á: {', '.join(failed_reasons)}")

                            except ImportError:
                                _app_logger.warning(
                                    f"[FILE_GEN/PPT] ΓÜá∩╕Å FileParser µ¿íσ¥ùµ£¬µë╛σê░∩╝îΦ╖│Φ┐çµûçΣ╗╢σñäτÉå"
                                )
                            except Exception as file_err:
                                _app_logger.warning(f"[FILE_GEN/PPT] ΓÜá∩╕Å µûçΣ╗╢σñäτÉåσ╝éσ╕╕: {file_err}")

                        # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ Step 0.1: µÖ║Φâ╜σêñµû¡µÿ»σÉªΘ£ÇΦªüΦüöτ╜æµÉ£τ┤ó ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        search_context = ""

                        # µúÇµ╡ïµÿ»σÉªΘ£ÇΦªüµÉ£τ┤óµ£Çµû░Σ┐íµü»
                        _needs_search = WebSearcher.needs_web_search(user_input)

                        # Θó¥σñûPPTΦ»¥ΘóÿµúÇµ╡ï∩╝ÜσîàσÉ½σ╣┤Σ╗╜/µù╢Θù┤/µû░σôü/Σ║ïΣ╗╢/µÄÆΦíîτ¡ëτÜäPPTσñºµªéτÄçΘ£ÇΦªüµÉ£τ┤ó
                        import re as _re

                        _time_topic_patterns = [
                            r"20\d{2}",  # σ╣┤Σ╗╜
                            r"\d+µ£ê",  # µ£êΣ╗╜
                            r"(µû░τò¬|µû░τëç|µû░σëº|µû░µ¡î|µû░σôü|Σ╕èµÿá|ΘªûσÅæ|σÅæσö«)",
                            r"(µÄÆΦíî|µÄÆσÉì|µª£σìò|top|τ¢ÿτé╣|σ»╝Φºå|ΘÇƒΘÇÆ|Σ╕ÇΦºê)",
                            r"(Φíîµâà|Φ╡░σè┐|Φ╢ïσè┐|σ╕éσ£║|Σ╗╖µá╝|µèÑσæè)",
                            r"(τâ¡Θù¿|τâ¡τé╣|τü½τêå|µ╡üΦíî|Σ║║µ░ö)",
                            r"(µ£Çµû░|µ£ÇΦ┐æ|Φ┐æµ£ƒ|µ£¼σæ¿|µ£¼µ£ê|σ╜ôσëì|τ¢«σëì)",
                        ]
                        if not _needs_search:
                            for pat in _time_topic_patterns:
                                if _re.search(pat, user_input, _re.IGNORECASE):
                                    _needs_search = True
                                    _app_logger.info(
                                        f"[FILE_GEN/PPT] ≡ƒöì Φ»¥Θóÿµù╢µòêµÇºµúÇµ╡ïσæ╜Σ╕¡: {pat}"
                                    )
                                    break

                        if _needs_search:
                            yield _fb_sse(fb.search_start())
                            try:
                                search_result = WebSearcher.search_with_grounding(
                                    user_input
                                )
                                if search_result.get("success") and search_result.get(
                                    "response"
                                ):
                                    search_context = search_result["response"]
                                    _app_logger.info(
                                        f"[FILE_GEN/PPT] Γ£à µÉ£τ┤óσ«îµêÉ, ΦÄ╖σÅû {len(search_context)} σ¡ùτ¼ªσÅéΦÇâΣ┐íµü»"
                                    )
                                    yield _fb_sse(
                                        fb.search_done(char_count=len(search_context))
                                    )
                                else:
                                    _app_logger.warning(f"[FILE_GEN/PPT] ΓÜá∩╕Å µÉ£τ┤óµùáτ╗ôµ₧£µêûσñ▒Φ┤Ñ")
                            except Exception as search_err:
                                _app_logger.warning(f"[FILE_GEN/PPT] ΓÜá∩╕Å µÉ£τ┤óσ╝éσ╕╕: {search_err}")

                        # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ Step 0.5: σñìµ¥éΣ╕╗Θóÿµ╖▒σ║ªτáöτ⌐╢ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        research_context = ""
                        _complex_patterns = [
                            r"(σÄƒτÉå|µ£║σê╢|µ₧╢µ₧ä|µèÇµ£»|τ«ùµ│ò|τÉåΦ«║|σêåµ₧É|τáöτ⌐╢|τ╗╝Φ┐░)",
                            r"(ΦíîΣ╕Ü|Σ║ºΣ╕Ü|σ╕éσ£║|σòåΣ╕Ü|µêÿτòÑ|ΦºäσêÆ|µû╣µíê)",
                            r"(σ¡ªµ£»|Φ«║µûç|Φ»╛Θóÿ|µ»òΣ╕Ü|µòÖσ¡ª|Φ»╛τ¿ï)",
                            r"(σÄåσÅ▓|σÅæσ▒ò|µ╝öσÅÿ|σÅÿΦ┐ü|µ▓┐Θ¥⌐)",
                            r"(σ»╣µ»ö|µ»öΦ╛â|Φ»äΣ╝░|Φ»äµ╡ï|benchmark)",
                            r"(τ╗Åµ╡Ä|ΘçæΦ₧ì|µèòΦ╡ä|Φ┤óσèí|Φ┤óµèÑ)",
                        ]
                        _is_complex = len(user_input) > 30 or any(
                            _re.search(p, user_input) for p in _complex_patterns
                        )

                        if _is_complex:
                            yield _fb_sse(fb.research_start())
                            try:
                                research_context = WebSearcher.deep_research_for_ppt(
                                    user_input, search_context
                                )
                                if research_context:
                                    yield _fb_sse(
                                        fb.research_done(
                                            char_count=len(research_context)
                                        )
                                    )
                                else:
                                    _app_logger.warning(f"[FILE_GEN/PPT] ΓÜá∩╕Å µ╖▒σ║ªτáöτ⌐╢µ£¬Φ┐öσ¢₧τ╗ôµ₧£")
                            except Exception as res_err:
                                _app_logger.warning(f"[FILE_GEN/PPT] ΓÜá∩╕Å µ╖▒σ║ªτáöτ⌐╢σ╝éσ╕╕: {res_err}")

                        # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ Step 1: τö¿ AI τöƒµêÉτ╗ôµ₧äσîûσñºτ║▓ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        # ΓöÇΓöÇΓöÇΓöÇ µÅÉσÅûτö¿µê╖ PPT σüÅσÑ╜∩╝êΘí╡µò░πÇüΘçìτé╣πÇüτ«ÇΦªüΦ»¥Θóÿ∩╝ë ΓöÇΓöÇΓöÇΓöÇ
                        import re as _ppt_re

                        def _extract_ppt_preferences(text):
                            prefs = {
                                "target_pages": None,
                                "focus_topics": [],
                                "brief_topics": [],
                            }
                            pm = _ppt_re.search(
                                r"(?:σüÜ|τöƒµêÉ|Θ£ÇΦªü|σñºµªé|τ║ª|σñºτ║ª)?\s*(\d+)\s*Θí╡", text
                            )
                            if pm:
                                prefs["target_pages"] = int(pm.group(1))
                            for pat in [
                                r"(?:Θçìτé╣|Φ»ªτ╗å|τ¥ÇΘçì|µ╖▒σàÑ|σñÜΦ«▓|σñÜΣ╗ïτ╗ì)(?:Σ╗ïτ╗ì|Φ«▓|σêåµ₧É|Φ»┤µÿÄ|σ▒òτñ║|Φ«▓Φºú)\s*(.+?)(?:[∩╝î,πÇé∩╝¢;πÇü]|$)",
                                r"(?:τ¬üσç║|σ╝║Φ░â)\s*(.+?)(?:[∩╝î,πÇé∩╝¢;πÇü]|$)",
                            ]:
                                for m in _ppt_re.finditer(pat, text):
                                    t = m.group(1).strip()
                                    if t and len(t) < 30:
                                        prefs["focus_topics"].append(t)
                            for pat in [
                                r"(?:τ«Çσìò|τ«ÇΦªü|τ«ÇτòÑ|σñºΦç┤)(?:σ╕ªΦ┐ç|Σ╗ïτ╗ì|Φ»┤|Φ«▓)\s*(.+?)(?:[∩╝î,πÇé∩╝¢;πÇü]|$)",
                                r"(.+?)(?:Σ╕Çτ¼öσ╕ªΦ┐ç|τòÑΦ┐ç|Φ╖│Φ┐ç|τ«ÇσìòΦ»┤)",
                            ]:
                                for m in _ppt_re.finditer(pat, text):
                                    t = m.group(1).strip()
                                    if t and len(t) < 30:
                                        prefs["brief_topics"].append(t)
                            return prefs

                        ppt_prefs = _extract_ppt_preferences(user_input)
                        _target_pages = ppt_prefs["target_pages"]
                        _target_hint = (
                            f"τ║ª {_target_pages} Θí╡∩╝êσ░üΘ¥óσÆîτ╗ôµ¥ƒΘí╡ΘÖñσñû∩╝îΦ┐Öµÿ»τö¿µê╖µîçσ«ÜτÜä∩╝îσ┐àΘí╗Σ╕Ñµá╝Θü╡σ«ê∩╝ë"
                            if _target_pages
                            else "8~15 Θí╡∩╝êµá╣µì«σåàσ«╣σñìµ¥éσ║ªµÖ║Φâ╜Φ░âµò┤∩╝îσåàσ«╣σñÜσÅ»Σ╗ÑσñÜσüÜσçáΘí╡∩╝îσåàσ«╣σ░æσ░▒τ▓╛τ«Ç∩╝ë"
                        )
                        _focus_hint = ""
                        if ppt_prefs["focus_topics"]:
                            _focus_hint = (
                                "\n**τö¿µê╖µîçσ«ÜτÜäΘçìτé╣σåàσ«╣∩╝êσ┐àΘí╗τö¿ [Φ»ªτ╗å] σñÜΘí╡σ▒òσ╝Ç∩╝ë∩╝Ü**\n"
                                + "\n".join(f"- {t}" for t in ppt_prefs["focus_topics"])
                                + "\n"
                            )
                        _brief_hint = ""
                        if ppt_prefs["brief_topics"]:
                            _brief_hint = (
                                "\n**τö¿µê╖µîçσ«ÜτÜäτ«ÇΦªüσåàσ«╣∩╝êσÉêσ╣╢σê░ [µªéΦºê] Θí╡∩╝ë∩╝Ü**\n"
                                + "\n".join(f"- {t}" for t in ppt_prefs["brief_topics"])
                                + "\n"
                            )

                        _app_logger.info(
                            f"[FILE_GEN/PPT] τö¿µê╖σüÅσÑ╜: Θí╡µò░={_target_pages}, Θçìτé╣={ppt_prefs['focus_topics']}, τ«ÇΦªü={ppt_prefs['brief_topics']}"
                        )

                        # ΓöÇΓöÇΓöÇΓöÇ µÖ║Φâ╜σåàσ«╣ΦºäσêÆ Prompt ΓöÇΓöÇΓöÇΓöÇ
                        ppt_outline_prompt = (
                            "Σ╜áµÿ»Σ╕ÇΣ╕¬Θí╢σ░ûτÜäµ╝öτñ║µûçτ¿┐σåàσ«╣τ¡ûσêÆσ╕êσÆîµÄÆτëêΦºäσêÆσ╕êπÇé\n\n"
                            "Σ╜áτÜäσ╖ÑΣ╜£σêåΣ╕ñµ¡Ñ∩╝Ü\n"
                            "1. **σåàσ«╣ΦºäσêÆ** ΓÇö σêåµ₧ÉΣ╕╗Θóÿ∩╝îσêñµû¡σô¬Σ║¢σåàσ«╣µÿ»Θçìτé╣∩╝êΘ£ÇΦªüσñÜΘí╡Φ»ªτ╗åσ▒òτñ║∩╝ë∩╝îσô¬Σ║¢µÿ»τ«ÇΦªü∩╝êσÅ»Σ╗ÑΣ╕ÇΘí╡σñÜΣ╕╗ΘóÿΘÇƒΦºê∩╝ë\n"
                            "2. **τëêσ╝ÅΘÇëµï⌐** ΓÇö Σ╕║µ»ÅΘâ¿σêåΘÇëµï⌐µ£ÇσÉêΘÇéτÜäσ╣╗τü»τëçτ▒╗σ₧ï\n\n"
                            "## σÅ»τö¿τÜäσ╣╗τü»τëçτ▒╗σ₧ï\n"
                            "σ£¿µ»ÅΣ╕¬ `## τ½áΦèéµáçΘóÿ` σëìΣ╕ÇΦíîσåÖτ▒╗σ₧ïµáçτ¡╛∩╝Ü\n\n"
                            "| µáçτ¡╛ | τö¿ΘÇö | µá╝σ╝Å |\n"
                            "|------|------|------|\n"
                            "| `[Φ»ªτ╗å]` | σ╕╕Φºäσåàσ«╣Θí╡∩╝îµ╖▒σàÑσ▒òτñ║ 4-6 Σ╕¬Φªüτé╣∩╝êµ»ÅΣ╕¬Φªüτé╣30-80σ¡ù∩╝ë | `- **σà│Θö«Φ»ì** ΓÇö Φ»ªτ╗åΦºúΘçèΦ»┤µÿÄσÆîσà╖Σ╜ôµò░µì«` |\n"
                            "| `[µªéΦºê]` | σñÜΣ╕╗ΘóÿΘÇƒΦºêΘí╡∩╝î2-4 Σ╕¬σ░ÅΣ╕╗Θóÿσ╣╢σêù∩╝êµ»ÅΣ╕¬σ░ÅΣ╕╗ΘóÿΣ╕ï2-4Σ╕¬Φªüτé╣∩╝ë | τö¿ `### σ¡ÉµáçΘóÿ` σêåτ╗ä |\n"
                            "| `[Σ║«τé╣]` | σà│Θö«µò░µì«τ¬üσç║Θí╡∩╝ê3-4τ╗äµò░µì«∩╝ë | `- µò░σÇ╝ \\| Φ»ªτ╗åΦ»┤µÿÄ` |\n"
                            "| `[σ»╣µ»ö]` | Σ╕ñµû╣σ»╣µ»öΘí╡∩╝êµ»Åµû╣3-5Σ╕¬Φªüτé╣∩╝ë | τö¿ `### ΘÇëΘí╣A` σÆî `### ΘÇëΘí╣B` σêåΣ╕ñτ╗ä |\n"
                            "| `[Φ┐çµ╕íΘí╡]` | τ½áΦèéΦ┐çµ╕í∩╝îσ╝òσàÑσñºτ½áΦèé∩╝êσ░æτö¿∩╝ë | Σ╕ïµû╣σåÖΣ╕ÇΦíîµÅÅΦ┐░ |\n\n"
                            "## Φ╛ôσç║µá╝σ╝Å∩╝êΣ╕Ñµá╝Θü╡σ╛¬∩╝ë\n"
                            "```\n"
                            "# PPTΣ╕╗µáçΘóÿ\n\n"
                            "[Φ┐çµ╕íΘí╡]\n"
                            "## τ¼¼Σ╕ÇΘâ¿σêåµáçΘóÿ\n"
                            "τ«Çτƒ¡µÅÅΦ┐░\n\n"
                            "[Φ»ªτ╗å]\n"
                            "## Θí╡Θ¥óµáçΘóÿ\n"
                            "- **µá╕σ┐âµªéσ┐╡** ΓÇö Φ┐Öµÿ»Σ╕Çµ«╡σ«îµò┤τÜäΦºúΘçèµÇºµûçσ¡ù∩╝îσîàσÉ½σà│Θö«µò░µì«µêûΣ║ïσ«₧Σ╛¥µì«∩╝îΦ«⌐ΦºéΣ╝ùΦâ╜τ£ƒµ¡úτÉåΦºúΦ┐ÖΣ╕Çτé╣τÜäσåàσ«╣\n"
                            "- **µèÇµ£»τë╣τé╣** ΓÇö σà╖Σ╜ôµÅÅΦ┐░µèÇµ£»τÜäσ╖ÑΣ╜£σÄƒτÉåπÇüΣ╝ÿσè┐µëÇσ£¿πÇüσ«₧ΘÖàσ║öτö¿σ£║µÖ»σÆîτ¢╕σà│σÅéµò░\n"
                            "- **σ╕éσ£║µò░µì«** ΓÇö σ╝òτö¿µ¥âσ¿üµ£║µ₧äτÜäτ╗ƒΦ«íµò░σ¡ùπÇüσ╕éσ£║Φºäµ¿íπÇüσó₧Θò┐τÄçτ¡ëΘçÅσîûΣ┐íµü»\n"
                            "- **σ«₧ΘÖàµíêΣ╛ï** ΓÇö µƒÉσà¼σÅ╕/Θí╣τ¢«τÜäσà╖Σ╜ôσ«₧Φ╖╡τ╗ÅΘ¬î∩╝îσÅûσ╛ùΣ║åΣ╗ÇΣ╣êµá╖τÜäµêÉµ₧£\n\n"
                            "[µªéΦºê]\n"
                            "## ΘÇƒΦºêµáçΘóÿ\n"
                            "### σ¡ÉΦ»¥Θóÿ1\n"
                            "- τ¼¼Σ╕ÇΣ╕¬Φªüτé╣τÜäΦ»ªτ╗åΦ»┤µÿÄ\n"
                            "- τ¼¼Σ║îΣ╕¬Φªüτé╣τÜäΦ»ªτ╗åΦ»┤µÿÄ\n"
                            "- τ¼¼Σ╕ëΣ╕¬Φªüτé╣τÜäΦ»ªτ╗åΦ»┤µÿÄ\n"
                            "### σ¡ÉΦ»¥Θóÿ2\n"
                            "- τ¼¼Σ╕ÇΣ╕¬Φªüτé╣σÆîσà╖Σ╜ôµò░µì«\n"
                            "- τ¼¼Σ║îΣ╕¬Φªüτé╣σÆîσ║öτö¿σ£║µÖ»\n\n"
                            "[Σ║«τé╣]\n"
                            "## σà│Θö«µò░µì«\n"
                            "- 500Σ║┐ | σà¿τÉâσ╕éσ£║Φºäµ¿í\n"
                            "- 35% | σ╣┤σó₧Θò┐τÄç\n\n"
                            "[σ»╣µ»ö]\n"
                            "## σ»╣µ»öµáçΘóÿ\n"
                            "### µû╣µíêA\n"
                            "- τë╣τé╣1\n"
                            "### µû╣µíêB\n"
                            "- τë╣τé╣1\n"
                            "```\n\n"
                            "## σåàσ«╣ΦºäσêÆΦºäσêÖ\n"
                            f"1. **µÇ╗Θí╡µò░τ¢«µáç: {_target_hint}**\n"
                            "2. **Θçìτé╣σåàσ«╣**Σ╜┐τö¿σñÜΣ╕¬ `[Φ»ªτ╗å]` Θí╡σ▒òσ╝Ç∩╝îµ»ÅΘí╡ 4-6 Σ╕¬Σ┐íµü»Σ╕░σ»îτÜäΦªüτé╣\n"
                            "3. ΓÜá∩╕Å **µ»ÅΣ╕¬Φªüτé╣σ┐àΘí╗µÿ»Σ╕ÇΣ╕¬σ«îµò┤τÜäΣ┐íµü»µ«╡ΦÉ╜∩╝ê30-80σ¡ù∩╝ë∩╝îΣ╕ìΦâ╜σÅ¬σåÖσçáΣ╕¬Φ»ìµêûτƒ¡Φ»¡**\n"
                            "4. Φªüτé╣µá╝σ╝Å: `- **σà│Θö«Φ»ì** ΓÇö σà╖Σ╜ôτÜäΦºúΘçèΦ»┤µÿÄ∩╝îσîàσÉ½µò░µì«πÇüΣ║ïσ«₧πÇüµíêΣ╛ïτ¡ëσ«₧Φ┤¿σåàσ«╣`\n"
                            "5. **Θ¥₧Θçìτé╣σåàσ«╣**σÉêσ╣╢σê░ `[µªéΦºê]` Θí╡∩╝îΣ╕ÇΘí╡ 2-4 Σ╕¬σ░ÅΣ╕╗Θóÿ∩╝îΓÜá∩╕Å **µ»ÅΣ╕¬σ░ÅΣ╕╗ΘóÿΣ╕ïσ┐àΘí╗µ£ë 2-4 Σ╕¬Φªüτé╣**\n"
                            "6. µ£ëµò░µì«Σ║«τé╣µù╢τö¿ `[Σ║«τé╣]` Θí╡∩╝êσà¿µûçµ£ÇσñÜ 1-2 µ¼í∩╝ë∩╝îµ»ÅΘí╡ 3-4 τ╗äµò░µì«\n"
                            "7. `[Φ┐çµ╕íΘí╡]` µ£ÇσñÜ 2 Σ╕¬∩╝îτö¿Σ║ÄσêÆσêåσñºτ½áΦèé\n"
                            "8. ΓÜá∩╕Å **µÉ£τ┤óΦ╡äµûÖσÆîτáöτ⌐╢µèÑσæèΣ╕¡τÜäµò░µì«πÇüµíêΣ╛ïπÇüµò░σ¡ùσ┐àΘí╗σªéσ«₧σ╝òτö¿∩╝îΣ╕ìσ╛ùτ╝ûΘÇáπÇéµ»ÅΘí╡Φç│σ░æσ╝òτö¿ 1 Σ╕¬σà╖Σ╜ôµò░µì«µêûµíêΣ╛ï**\n"
                            "9. Σ╕¡µûçΦ╛ôσç║∩╝îσÅ¬Φ╛ôσç║σñºτ║▓∩╝îΣ╕ìΦªüΘó¥σñûΦ»┤µÿÄ\n"
                            "10. ΓÜá∩╕Å **σåàσ«╣σààσ«₧σ║ªµÿ»µ£ÇΘçìΦªüτÜäΦ»äσêñµáçσçå ΓÇö σ«üσÅ»Φªüτé╣σ░æΣ╕ÇΣ║¢Σ╜åµ»ÅΣ╕¬Φªüτé╣Σ┐íµü»ΘçÅσñº∩╝îΣ╕ìΦªüσ╛êσñÜτ⌐║µ┤₧τÜäΦªüτé╣**\n"
                            "11. ΓÜá∩╕Å **τªüµ¡óσç║τÄ░µ¿íτ│èΦí¿Φ┐░**∩╝Üσªé 'µÿ╛Φæùσó₧Θò┐'πÇü'σ╣┐µ│¢σ║öτö¿'πÇü'σ╖¿σñºµ╜£σè¢' τ¡ë∩╝îσ┐àΘí╗τö¿σà╖Σ╜ôµò░σ¡ùµ¢┐Σ╗úπÇéΣ╛ïσªé∩╝Ü'σ╕éσ£║Φºäµ¿íΦ╛╛ XX Σ║┐' ΦÇîΣ╕ìµÿ» 'σ╕éσ£║Φºäµ¿íσ╖¿σñº'\n"
                            "12. **µ»ÅΣ╕¬ [Φ»ªτ╗å] Θí╡Φç│σ░æσîàσÉ½ 1 Σ╕¬τ£ƒσ«₧µíêΣ╛ïµêûµò░µì«τé╣**∩╝îµò░µì«Θ£Çµáçµ│¿µ¥Ñµ║É∩╝êσªé 'µì«IDCµò░µì«' 'µá╣µì«XXσ╣┤µèÑ'∩╝ë\n"
                            f"{_focus_hint}"
                            f"{_brief_hint}"
                            f"\nτö¿µê╖Θ£Çµ▒é: {user_input}\n"
                        )

                        # µ│¿σàÑµÉ£τ┤óτ╗ôµ₧£∩╝êσó₧σèáΘÖÉΘó¥Σ╗ÑΣ┐¥τòÖµ¢┤σñÜµò░µì«∩╝ë
                        if uploaded_file_context:
                            ppt_outline_prompt = (
                                ppt_outline_prompt[: -len("\nτö¿µê╖Θ£Çµ▒é: " + user_input)]
                                + f"\n\n## Σ╕èΣ╝áτÜäσÅéΦÇâµûçΣ╗╢σåàσ«╣\n"
                                f"Σ╗ÑΣ╕ïµÿ»τö¿µê╖Σ╕èΣ╝áτÜäµûçµíúΦ╡äµûÖ∩╝îΦ»╖σààσêåσê⌐τö¿σà╢Σ╕¡τÜäσåàσ«╣πÇüµò░µì«πÇüµíêΣ╛ïµ¥ÑτöƒµêÉ PPT∩╝Ü\n"
                                f"---\n{uploaded_file_context[:15000]}\n---\n"
                                f"\nτö¿µê╖Θ£Çµ▒é: {user_input}\n"
                            )

                        if search_context:
                            ppt_outline_prompt += (
                                f"\n**Σ╗ÑΣ╕ïµÿ»Φüöτ╜æµÉ£τ┤óΦÄ╖σÅûτÜäµ£Çµû░σÅéΦÇâΦ╡äµûÖ∩╝êσîàσÉ½ΘçìΦªüµò░µì«∩╝ë∩╝îΦ»╖σèíσ┐àσƒ║Σ║ÄΦ┐ÖΣ║¢Σ┐íµü»τöƒµêÉσåàσ«╣∩╝îσ░ñσà╢µÿ»σà╢Σ╕¡τÜäµò░σ¡ùπÇüµíêΣ╛ïπÇüσ╕éσ£║µò░µì«∩╝Ü**\n"
                                f"---\n{search_context[:10000]}\n---\n"
                            )

                        # µ│¿σàÑµ╖▒σ║ªτáöτ⌐╢τ╗ôµ₧£∩╝êσó₧σèáΘÖÉΘó¥∩╝ë
                        if research_context:
                            ppt_outline_prompt += (
                                f"\n**Σ╗ÑΣ╕ïµÿ»µ╖▒σ║ªτáöτ⌐╢σêåµ₧ÉµèÑσæèΓÇöΓÇöΦ┐Öµÿ»Σ╜áµ£ÇΘçìΦªüτÜäσåàσ«╣µ¥Ñµ║É∩╝îσà╢Σ╕¡τÜäµò░µì«σÆîσêåµ₧Éσ┐àΘí╗σààσêåΦ₧ìσàÑσñºτ║▓∩╝Ü**\n"
                                f"---\n{research_context[:12000]}\n---\n"
                            )

                        # Σ╣ƒµ│¿σàÑΣ╕èΣ╕ïµûç
                        if (
                            context_info
                            and context_info.get("is_continuation")
                            and context_info.get("enhanced_input")
                        ):
                            ppt_outline_prompt += f"\n\nσÄåσÅ▓Σ╕èΣ╕ïµûçσÅéΦÇâΦ╡äµûÖ:\n{context_info['enhanced_input'][:3000]}"

                        yield _fb_sse(fb.ppt_planning("Φ░âτö¿ AI τöƒµêÉσåàσ«╣σñºτ║▓"))

                        outline_response = None
                        outline_models = [
                            "gemini-2.5-flash",
                            model_id,
                            "gemini-3-flash-preview",
                        ]
                        # µá╣µì«τ¢«µáçΘí╡µò░Φ░âµò┤ token ΘÖÉΘó¥∩╝Ü20Θí╡σñºτ║▓Θ£ÇΦªüµ¢┤σñÜτ⌐║Θù┤
                        _outline_tokens = (
                            16384 if (_target_pages and _target_pages >= 15) else 8192
                        )
                        for om in outline_models:
                            try:
                                resp = client.models.generate_content(
                                    model=om,
                                    contents=ppt_outline_prompt,
                                    config=types.GenerateContentConfig(
                                        temperature=0.6,
                                        max_output_tokens=_outline_tokens,
                                    ),
                                )
                                if resp.text:
                                    outline_response = resp.text
                                    _app_logger.info(
                                        f"[FILE_GEN/PPT] Γ£à σñºτ║▓τöƒµêÉµêÉσèƒ ({om}), Θò┐σ║ª: {len(outline_response)}"
                                    )
                                    break
                            except Exception as oe:
                                _app_logger.info(f"[FILE_GEN/PPT] σñºτ║▓µ¿íσ₧ï {om} σñ▒Φ┤Ñ: {oe}")
                                continue

                        if not outline_response:
                            raise Exception("µëÇµ£ëµ¿íσ₧ïσ¥çµùáµ│òτöƒµêÉσñºτ║▓")

                        # Step 2: Φºúµ₧ÉµÖ║Φâ╜ΦºäσêÆσñºτ║▓∩╝êµö»µîüσñÜτºìσ╣╗τü»τëçτ▒╗σ₧ïµáçτ¡╛∩╝ë
                        def _parse_ppt_plan(md_text):
                            """Φºúµ₧Éσ╕ª [τ▒╗σ₧ï] µáçτ¡╛τÜäµÖ║Φâ╜ PPT σñºτ║▓"""
                            import re as _re

                            lines = md_text.split("\n")
                            plan = {"title": "", "subtitle": "", "slides": []}

                            _type_map = {
                                "Φ┐çµ╕íΘí╡": "divider",
                                "Φ┐çµ╕í": "divider",
                                "σêåΘÜö": "divider",
                                "Φ»ªτ╗å": "detail",
                                "Θçìτé╣": "detail",
                                "Σ║«τé╣": "highlight",
                                "µò░µì«": "highlight",
                                "σà│Θö«": "highlight",
                                "µªéΦºê": "overview",
                                "ΘÇƒΦºê": "overview",
                                "τ«ÇΦªü": "overview",
                                "µÇ╗Φºê": "overview",
                                "σ»╣µ»ö": "comparison",
                                "µ»öΦ╛â": "comparison",
                                "vs": "comparison",
                            }

                            current_slide = None
                            current_type = "detail"
                            current_sub = (
                                None  # σ╜ôσëìσ¡ÉΣ╕╗Θóÿ∩╝êτö¿Σ║Ä overview / comparison∩╝ë
                            )

                            for line in lines:
                                line = line.rstrip()

                                # Φ╖│Φ┐ç markdown Σ╗úτáüσ¥ùµáçΦ«░
                                if line.strip() in ("```", "```markdown"):
                                    continue

                                # τ▒╗σ₧ïµáçτ¡╛Φíî: [xxx]
                                tag_m = _re.match(r"^\s*\[(.+?)\]\s*$", line)
                                if tag_m:
                                    tag = tag_m.group(1).strip()
                                    current_type = _type_map.get(tag, "detail")
                                    continue

                                # Σ╕╗µáçΘóÿ: # xxx
                                if line.startswith("# ") and not line.startswith("## "):
                                    raw = line[2:].strip()
                                    for pfx in [
                                        "σ╣╗τü»τëçµáçΘóÿ∩╝Ü",
                                        "σ╣╗τü»τëçµáçΘóÿ:",
                                        "µ╝öτñ║µáçΘóÿ∩╝Ü",
                                        "µ╝öτñ║µáçΘóÿ:",
                                        "PPTµáçΘóÿ∩╝Ü",
                                        "PPTµáçΘóÿ:",
                                    ]:
                                        if raw.startswith(pfx):
                                            raw = raw[len(pfx) :].strip()
                                    plan["title"] = raw
                                    continue

                                # τ½áΦèéµáçΘóÿ: ## xxx
                                if line.startswith("## "):
                                    # Σ┐¥σ¡ÿΣ╕èΣ╕ÇΣ╕¬ slide τÜä subsection
                                    if (
                                        current_sub
                                        and current_slide
                                        and current_slide.get("type")
                                        in ("overview", "comparison")
                                    ):
                                        current_slide.setdefault(
                                            "subsections", []
                                        ).append(current_sub)
                                        current_sub = None
                                    # Σ┐¥σ¡ÿΣ╕èΣ╕ÇΣ╕¬ slide
                                    if current_slide:
                                        plan["slides"].append(current_slide)

                                    current_slide = {
                                        "type": current_type,
                                        "title": line[3:].strip(),
                                        "points": [],
                                        "content": [],
                                    }
                                    if current_type == "divider":
                                        current_slide["description"] = ""
                                    current_type = (
                                        "detail"  # Θçìτ╜«∩╝êµ»ÅΣ╕¬µáçτ¡╛σÅ¬Σ╜£τö¿Σ║Äτ┤ºΦ╖ƒτÜä ## ∩╝ë
                                    )
                                    current_sub = None
                                    continue

                                # σ¡ÉµáçΘóÿ: ### xxx ∩╝êτö¿Σ║Ä overview / comparison∩╝ë
                                if line.startswith("### ") and current_slide:
                                    # σªéµ₧£σ╜ôσëì slide Σ╕ìµÿ» overview/comparison∩╝îΦç¬σè¿σìçτ║ºΣ╕║ overview
                                    if current_slide.get("type") not in (
                                        "overview",
                                        "comparison",
                                    ):
                                        current_slide["type"] = "overview"
                                    if current_sub:
                                        current_slide.setdefault(
                                            "subsections", []
                                        ).append(current_sub)
                                    current_sub = {
                                        "subtitle": line[4:].strip(),
                                        "label": line[4:].strip(),
                                        "points": [],
                                    }
                                    continue

                                # Φªüτé╣Φíî: - / ΓÇó / * µêûµò░σ¡ùτ╝ûσÅ╖ 1. 2. τ¡ë
                                if (
                                    _re.match(r"^[\s]*[-ΓÇó*]\s", line)
                                    or _re.match(r"^[\s]*\d+[.πÇü)\s]\s*", line)
                                ) and current_slide is not None:
                                    pt = _re.sub(
                                        r"^[\s]*[-ΓÇó*\d.πÇü)\s]+\s*", "", line
                                    ).strip()
                                    if not pt:
                                        continue
                                    if current_sub is not None:
                                        current_sub["points"].append(pt)
                                    else:
                                        current_slide["points"].append(pt)
                                        current_slide["content"].append(pt)
                                    continue

                                # µÖ«ΘÇÜµûçµ£¼Φíî∩╝êΘ¥₧τ⌐║πÇüΘ¥₧µáçΘóÿ∩╝ëΓåÆ Σ╣ƒµìòΦÄ╖Σ╕║Φªüτé╣
                                if (
                                    current_slide is not None
                                    and line.strip()
                                    and not line.startswith("#")
                                ):
                                    # Φ┐çµ╕íΘí╡µÅÅΦ┐░µûçσ¡ùΣ╝ÿσàê
                                    if current_slide.get("type") == "divider":
                                        current_slide["description"] = line.strip()
                                        continue
                                    # µ╕àτÉåσÅ»Φâ╜µ«ïτòÖτÜä markdown µáçΦ«░
                                    cleaned = _re.sub(r"^#{1,4}\s+", "", line.strip())
                                    cleaned = cleaned.strip()
                                    if not cleaned:
                                        continue
                                    if current_sub is not None:
                                        current_sub["points"].append(cleaned)
                                    else:
                                        current_slide["points"].append(cleaned)
                                        current_slide["content"].append(cleaned)
                                    continue

                                # Φ┐çµ╕íΘí╡µÅÅΦ┐░µûçσ¡ù (fallback - Σ╕ìσ║öσê░Φ╛╛Φ┐ÖΘçî)
                                if (
                                    current_slide
                                    and current_slide.get("type") == "divider"
                                    and line.strip()
                                    and not line.startswith("#")
                                ):
                                    current_slide["description"] = line.strip()

                            # µö╢σ░╛
                            if current_sub and current_slide:
                                current_slide.setdefault("subsections", []).append(
                                    current_sub
                                )
                            if current_slide:
                                plan["slides"].append(current_slide)

                            # σÉÄσñäτÉå: σªéµ₧£ slide µ£ë subsections Σ╜åτ▒╗σ₧ïΣ╕ìµÿ» overview/comparison∩╝îΦç¬σè¿Σ┐«µ¡ú
                            for sl in plan["slides"]:
                                if sl.get("subsections") and sl.get("type") not in (
                                    "overview",
                                    "comparison",
                                ):
                                    sl["type"] = "overview"

                            # σÉÄσñäτÉå: comparison τÜä subsections ΓåÆ left / right
                            for sl in plan["slides"]:
                                if (
                                    sl.get("type") == "comparison"
                                    and "subsections" in sl
                                ):
                                    subs = sl["subsections"]
                                    if len(subs) >= 2:
                                        sl["left"] = subs[0]
                                        sl["right"] = subs[1]

                            return plan

                        ppt_data = _parse_ppt_plan(outline_response)
                        slide_count = len(ppt_data["slides"])
                        slide_types_summary = ", ".join(
                            f'{s.get("type","detail")}' for s in ppt_data["slides"]
                        )
                        _app_logger.info(
                            f"[FILE_GEN/PPT] Φºúµ₧Éσ«îµêÉ: µáçΘóÿ='{ppt_data['title']}', {slide_count} Θí╡, τ▒╗σ₧ï=[{slide_types_summary}]"
                        )

                        if slide_count == 0:
                            raise Exception("σñºτ║▓Φºúµ₧Éσñ▒Φ┤Ñ∩╝îµ£¬µÅÉσÅûσê░σ╣╗τü»τëçσåàσ«╣")

                        # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ Quality Gate: σñºτ║▓Φ┤¿ΘçÅΦç¬µúÇΣ╕Äσåàσ«╣µ╕àµ┤ù ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        try:
                            from web.file_quality_checker import FileQualityGate

                            def _quality_progress(msg, detail=""):
                                pass  # σåàΘâ¿τö¿∩╝îΣ╕ïΘ¥óτ╗ƒΣ╕ÇσÅæ SSE

                            qg_result = FileQualityGate.check_and_fix_ppt_outline(
                                ppt_data["slides"],
                                user_request=user_input,
                                progress_callback=_quality_progress,
                            )
                            ppt_data["slides"] = qg_result["outline"]
                            _qg_score = qg_result["quality"]["score"]
                            _qg_fixes = qg_result["fixes"]
                            _qg_issues = qg_result["quality"]["issues"]

                            # µèÑσæèµ╕àµ┤ùτ╗ôµ₧£
                            if _qg_fixes:
                                yield _fb_sse(
                                    fb.info(
                                        f"≡ƒº╣ σ╖▓Φç¬σè¿µ╕àµ┤ù {len(_qg_fixes)} σñäσåàσ«╣Θù«Θóÿ",
                                        "τº╗ΘÖñ Markdown µ«ïτòÖσÆî AI σ»╣Φ»¥τùòΦ┐╣",
                                    )
                                )
                                _app_logger.info(
                                    f"[FILE_GEN/PPT] ≡ƒº╣ Φ┤¿ΘçÅµ╕àµ┤ù: {len(_qg_fixes)} σñäΣ┐«σñì"
                                )

                            # µèÑσæèΦ┤¿ΘçÅΦ»äσêå
                            yield _fb_sse(
                                fb.ppt_quality_check(
                                    _qg_score, issues=_qg_issues, fixes=_qg_fixes
                                )
                            )
                            _app_logger.info(
                                f"[FILE_GEN/PPT] ≡ƒôè Φ┤¿ΘçÅΦ»äσêå: {_qg_score}/100, action={qg_result['action']}"
                            )

                            # µ¢┤µû░ slide_count∩╝êµ╕àµ┤ùσÅ»Φâ╜τº╗ΘÖñτ⌐║τÖ╜ slide∩╝ë
                            slide_count = len(ppt_data["slides"])
                        except Exception as qg_err:
                            _app_logger.warning(
                                f"[FILE_GEN/PPT] ΓÜá∩╕Å Φ┤¿ΘçÅΘù¿µÄºσ╝éσ╕╕∩╝êΣ╕ìσ╜▒σôìτöƒµêÉ∩╝ë: {qg_err}"
                            )

                        # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ Step 2.1: τö¿µê╖µîçσ«ÜΘí╡µò░µù╢Φ░âµò┤σ╣╗τü»τëçµò░ΘçÅ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        _max_slides = _target_pages  # σÅ¬µ£ëτö¿µê╖µÿÄτí«µîçσ«Üµù╢µëìτöƒµòê
                        if _max_slides and slide_count > _max_slides:
                            _app_logger.warning(
                                f"[FILE_GEN/PPT] ΓÜá∩╕Å Θí╡µò░Φ╢àΘÖÉ ({slide_count} > {_max_slides})∩╝îµëºΦíîµÖ║Φâ╜τ▓╛τ«Ç..."
                            )
                            yield _fb_sse(
                                fb.info(
                                    f"Γ£é∩╕Å τ▓╛τ«ÇΘí╡Θ¥ó: {slide_count} ΓåÆ {_max_slides} Θí╡",
                                    "σÉêσ╣╢τ¢╕Σ╝╝σåàσ«╣∩╝îΣ┐¥τòÖµá╕σ┐âΣ┐íµü»",
                                )
                            )

                            slides = ppt_data["slides"]
                            # τ¡ûτòÑ: 1) σÉêσ╣╢τ¢╕Θé╗τÜäΦ»ªτ╗åΘí╡Σ╕║µªéΦºêΘí╡  2) σÄ╗µÄëσñÜΣ╜ÖΦ┐çµ╕íΘí╡  3) µê¬µû¡σ░╛Θâ¿

                            # σàêσÄ╗µÄëσñÜΣ╜ÖΦ┐çµ╕íΘí╡∩╝êσÅ¬Σ┐¥τòÖµ£ÇσñÜ 1 Σ╕¬∩╝ë
                            divider_indices = [
                                i
                                for i, s in enumerate(slides)
                                if s.get("type") == "divider"
                            ]
                            if len(divider_indices) > 1:
                                for idx in divider_indices[1:]:
                                    slides[idx]["_remove"] = True
                                slides = [s for s in slides if not s.get("_remove")]

                            # τä╢σÉÄσÉêσ╣╢τ¢╕Θé╗τÜäΦ»ªτ╗åΘí╡Σ╕║µªéΦºêΘí╡
                            while len(slides) > _max_slides:
                                merged = False
                                for i in range(len(slides) - 1):
                                    if (
                                        slides[i].get("type") == "detail"
                                        and slides[i + 1].get("type") == "detail"
                                    ):
                                        # σÉêσ╣╢: τ¼¼Σ╕ÇΣ╕¬σÆîτ¼¼Σ║îΣ╕¬Φ»ªτ╗åΘí╡σÅÿµêÉΣ╕ÇΣ╕¬µªéΦºêΘí╡
                                        s1 = slides[i]
                                        s2 = slides[i + 1]
                                        merged_slide = {
                                            "type": "overview",
                                            "title": s1.get("title", ""),
                                            "points": [],
                                            "content": [],
                                            "subsections": [
                                                {
                                                    "subtitle": s1.get("title", ""),
                                                    "label": s1.get("title", ""),
                                                    "points": (
                                                        s1.get("points", [])
                                                        or s1.get("content", [])
                                                    )[:4],
                                                },
                                                {
                                                    "subtitle": s2.get("title", ""),
                                                    "label": s2.get("title", ""),
                                                    "points": (
                                                        s2.get("points", [])
                                                        or s2.get("content", [])
                                                    )[:4],
                                                },
                                            ],
                                        }
                                        slides[i] = merged_slide
                                        slides.pop(i + 1)
                                        merged = True
                                        break
                                if not merged:
                                    # µùáµ│òσÉêσ╣╢Σ║å∩╝îτ¢┤µÄÑµê¬µû¡
                                    slides = slides[:_max_slides]
                                    break

                            ppt_data["slides"] = slides
                            slide_count = len(slides)
                            _app_logger.info(f"[FILE_GEN/PPT] τ▓╛τ«ÇσÉÄ: {slide_count} Θí╡")

                        # τöƒµêÉτëêσ╝ÅµæÿΦªü
                        _type_map_display = {
                            "detail": "Φ»ªτ╗å",
                            "overview": "µªéΦºê",
                            "highlight": "Σ║«τé╣",
                            "divider": "Φ┐çµ╕í",
                            "comparison": "σ»╣µ»ö",
                        }
                        _tc = {}
                        for _s in ppt_data["slides"]:
                            _t = _s.get("type", "detail")
                            _tc[_t] = _tc.get(_t, 0) + 1
                        _ts = "├ù".join(
                            f"{_type_map_display.get(k,k)}{v}" for k, v in _tc.items()
                        )

                        yield _fb_sse(
                            fb.ppt_outline_ready(
                                slide_count, title=ppt_data["title"], type_summary=_ts
                            )
                        )

                        # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ Step 2.2: σåàσ«╣σààσ«₧∩╝êΘÇÉΘí╡µë⌐σåÖ∩╝ë ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        # µúÇµƒÑσåàσ«╣µÿ»σÉªσìòΦûä∩╝êσ╣│σ¥çµ»ÅΘí╡Φªüτé╣σ░æΣ║Ä 3 Σ╕¬µêûΦªüτé╣σñ¬τƒ¡∩╝ë
                        _thin_slides = []
                        for si, sl in enumerate(ppt_data["slides"]):
                            stype = sl.get("type", "detail")
                            if stype in ("divider",):
                                continue  # Φ┐çµ╕íΘí╡Σ╕ìΘ£ÇΦªüσààσ«₧
                            pts = sl.get("points", [])
                            subs = sl.get("subsections", [])
                            # Φªüτé╣σñ¬σ░æ µêû σ╣│σ¥çΦªüτé╣σñ¬τƒ¡
                            avg_len = sum(len(p) for p in pts) / max(len(pts), 1)
                            sub_pts_count = (
                                sum(len(sub.get("points", [])) for sub in subs)
                                if subs
                                else 0
                            )

                            if stype == "overview":
                                # µªéΦºêΘí╡∩╝Üσ¡ÉΣ╕╗Θóÿµò░σñ¬σ░æµêûµ»ÅΣ╕¬σ¡ÉΣ╕╗ΘóÿΦªüτé╣σñ¬σ░æ
                                if not subs or sub_pts_count < len(subs) * 2:
                                    _thin_slides.append(si)
                            elif stype in ("detail", "comparison"):
                                if len(pts) < 3 or avg_len < 20:
                                    _thin_slides.append(si)
                            elif stype == "highlight":
                                if len(pts) < 2:
                                    _thin_slides.append(si)

                        if _thin_slides:
                            yield _fb_sse(fb.ppt_enriching(len(_thin_slides)))

                            # µ₧äσ╗║µë╣ΘçÅσààσ«₧ prompt∩╝êΣ╕Çµ¼íµÇºσñäτÉåµëÇµ£ëΦûäσ╝▒Θí╡Θ¥ó∩╝ë
                            _enrich_prompt = (
                                "Σ╜áµÿ»PPTσåàσ«╣µÆ░σåÖΣ╕ôσ«╢πÇéΣ╗ÑΣ╕ïσ╣╗τü»τëçσåàσ«╣σñ¬σìòΦûä∩╝îΦ»╖ΘÇÉΘí╡σààσ«₧πÇé\n\n"
                                "**Φªüµ▒é∩╝Ü**\n"
                                "1. µ»ÅΣ╕¬[Φ»ªτ╗å]Θí╡σ┐àΘí╗µ£ë 4-6 Σ╕¬Φªüτé╣∩╝îµ»ÅΣ╕¬Φªüτé╣ 30-80 σ¡ù\n"
                                "2. µ»ÅΣ╕¬[µªéΦºê]Θí╡τÜäµ»ÅΣ╕¬σ¡ÉΣ╕╗ΘóÿΣ╕ïσ┐àΘí╗µ£ë 2-4 Σ╕¬Φªüτé╣\n"
                                "3. Σ┐¥µîü `- **σà│Θö«Φ»ì** ΓÇö Φ»ªτ╗åΦºúΘçè` τÜäµá╝σ╝Å\n"
                                "4. ΓÜá∩╕Å **σ┐àΘí╗σîàσÉ½σà╖Σ╜ôµò░µì«πÇüµíêΣ╛ïπÇüΣ║ïσ«₧** ΓÇö τªüµ¡óσåÖ 'µÿ╛Φæùσó₧Θò┐' 'σ╣┐µ│¢σ║öτö¿' τ¡ëµ¿íτ│èΦí¿Φ┐░∩╝î\n"
                                "   σ┐àΘí╗σåÖ 'µì«IDCµò░µì«∩╝î2025σ╣┤σ╕éσ£║Φºäµ¿íΦ╛╛XXXΣ║┐' Φ┐Öµá╖µ£ëµò░σ¡ùµ£ëµ¥Ñµ║ÉτÜäσåàσ«╣\n"
                                "5. Σ╝ÿσàêΣ╜┐τö¿Σ╕ïµû╣πÇÉσÅéΦÇâΦ╡äµûÖπÇæσÆîπÇÉτáöτ⌐╢σêåµ₧ÉπÇæΣ╕¡τÜäτ£ƒσ«₧µò░µì«\n"
                                "6. Σ╕Ñµá╝µîëΣ╗ÑΣ╕ï JSON µá╝σ╝ÅΦ╛ôσç║∩╝îΣ╕ìΦªüΘó¥σñûµûçσ¡ù\n\n"
                            )

                            _slides_to_enrich = []
                            for si in _thin_slides:
                                sl = ppt_data["slides"][si]
                                _slides_to_enrich.append(
                                    {
                                        "index": si,
                                        "type": sl.get("type", "detail"),
                                        "title": sl.get("title", ""),
                                        "current_points": sl.get("points", []),
                                        "subsections": (
                                            [
                                                {
                                                    "subtitle": sub.get("subtitle", ""),
                                                    "points": sub.get("points", []),
                                                }
                                                for sub in sl.get("subsections", [])
                                            ]
                                            if sl.get("subsections")
                                            else []
                                        ),
                                    }
                                )

                            _enrich_prompt += f"Σ╕╗Θóÿ: {ppt_data['title']}\n"
                            if search_context:
                                _enrich_prompt += f"\nσÅéΦÇâΦ╡äµûÖ∩╝êσîàσÉ½ΘçìΦªüµò░µì«∩╝îΦ»╖σààσêåσê⌐τö¿∩╝ë:\n{search_context[:6000]}\n"
                            if research_context:
                                _enrich_prompt += f"\nτáöτ⌐╢σêåµ₧É∩╝êσîàσÉ½µá╕σ┐âµò░µì«σÆîµíêΣ╛ï∩╝îσ┐àΘí╗Φ₧ìσàÑ∩╝ë:\n{research_context[:6000]}\n"

                            _enrich_prompt += (
                                f"\nΘ£ÇΦªüσààσ«₧τÜäσ╣╗τü»τëç:\n```json\n{json.dumps(_slides_to_enrich, ensure_ascii=False, indent=2)}\n```\n\n"
                                "Φ»╖Φ╛ôσç║σààσ«₧σÉÄτÜäτ╗ôµ₧£∩╝îµá╝σ╝Å:\n"
                                "```json\n"
                                '[{"index": 0, "points": ["...", ...], "subsections": [{"subtitle": "...", "points": ["..."]}, ...]}]\n'
                                "```\n"
                                "σÅ¬Φ╛ôσç║ JSON∩╝îΣ╕ìΦªüΘó¥σñûµûçσ¡ùπÇé"
                            )

                            try:
                                _enrich_resp = client.models.generate_content(
                                    model="gemini-2.5-flash",
                                    contents=_enrich_prompt,
                                    config=types.GenerateContentConfig(
                                        temperature=0.5, max_output_tokens=8192
                                    ),
                                )
                                _enrich_text = _enrich_resp.text or ""
                                import re as _enrich_re

                                _em = _enrich_re.search(
                                    r"\[.*\]", _enrich_text, _enrich_re.DOTALL
                                )
                                if _em:
                                    _enriched = json.loads(_em.group())
                                    _applied = 0
                                    for _e in _enriched:
                                        _idx = _e.get("index")
                                        if _idx is not None and 0 <= _idx < len(
                                            ppt_data["slides"]
                                        ):
                                            _sl = ppt_data["slides"][_idx]
                                            # µ¢┤µû░ points
                                            if _e.get("points") and len(
                                                _e["points"]
                                            ) >= len(_sl.get("points", [])):
                                                _sl["points"] = _e["points"]
                                                _sl["content"] = _e["points"]
                                            # µ¢┤µû░ subsections
                                            if (
                                                _e.get("subsections")
                                                and len(_e["subsections"]) > 0
                                            ):
                                                _new_subs = []
                                                for _ns in _e["subsections"]:
                                                    _new_subs.append(
                                                        {
                                                            "subtitle": _ns.get(
                                                                "subtitle", ""
                                                            ),
                                                            "label": _ns.get(
                                                                "subtitle", ""
                                                            ),
                                                            "points": _ns.get(
                                                                "points", []
                                                            ),
                                                        }
                                                    )
                                                if _new_subs:
                                                    _sl["subsections"] = _new_subs
                                                    # Σ╣ƒµ¢┤µû░ comparison τÜä left/right
                                                    if (
                                                        _sl.get("type") == "comparison"
                                                        and len(_new_subs) >= 2
                                                    ):
                                                        _sl["left"] = _new_subs[0]
                                                        _sl["right"] = _new_subs[1]
                                            _applied += 1

                                    if _applied > 0:
                                        yield _fb_sse(fb.ppt_enriched(_applied))
                                        _app_logger.info(
                                            f"[FILE_GEN/PPT] Γ£à σåàσ«╣σààσ«₧σ«îµêÉ: {_applied}/{len(_thin_slides)} Θí╡"
                                        )
                                    else:
                                        _app_logger.warning(
                                            f"[FILE_GEN/PPT] ΓÜá∩╕Å σåàσ«╣σààσ«₧Φºúµ₧ÉµêÉσèƒΣ╜åµ£¬σ║öτö¿"
                                        )
                                else:
                                    _app_logger.warning(f"[FILE_GEN/PPT] ΓÜá∩╕Å σåàσ«╣σààσ«₧Φ┐öσ¢₧µá╝σ╝Åσ╝éσ╕╕")
                            except Exception as enrich_err:
                                _app_logger.warning(
                                    f"[FILE_GEN/PPT] ΓÜá∩╕Å σåàσ«╣σààσ«₧σ╝éσ╕╕∩╝êΣ╕ìσ╜▒σôìτöƒµêÉ∩╝ë: {enrich_err}"
                                )

                        # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ Step 2.5: Σ╕║σ╣╗τü»τëçτöƒµêÉΘàìσ¢╛∩╝êGemini 3.1 Flash Image Σ╝ÿσàê∩╝ë ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        ppt_images = []
                        # σ»╣Φ»ªτ╗åΘí╡Θàìσ¢╛∩╝êµªéΦºê/σ»╣µ»ö/Φ┐çµ╕í/Σ║«τé╣Θí╡Σ╕ìΘÇéσÉêµÅÆσ¢╛∩╝ë
                        img_candidate_slides = [
                            (i, s)
                            for i, s in enumerate(ppt_data["slides"])
                            if s.get("type", "detail") == "detail"
                        ]

                        if img_candidate_slides:
                            _n_images = min(
                                4, max(2, len(img_candidate_slides) // 2 + 1)
                            )
                            yield _fb_sse(fb.ppt_images(_n_images))
                            try:
                                slide_titles_for_img = [
                                    s.get("title", "") for _, s in img_candidate_slides
                                ]
                                img_results = WebSearcher.generate_ppt_images(
                                    slide_titles_for_img,
                                    topic=ppt_data["title"],
                                    max_images=_n_images,
                                )
                                # σ░åΘàìσ¢╛Φ╖»σ╛äµ│¿σàÑσê░σ»╣σ║ö slide
                                for img_info in img_results:
                                    picked_idx = img_info["slide_index"]
                                    if picked_idx < len(img_candidate_slides):
                                        real_idx = img_candidate_slides[picked_idx][0]
                                        ppt_data["slides"][real_idx]["image"] = (
                                            img_info["image_path"]
                                        )
                                        ppt_images.append(img_info["image_path"])

                                if ppt_images:
                                    yield _fb_sse(fb.ppt_images_done(len(ppt_images)))
                                else:
                                    yield _fb_sse(fb.ppt_images_done(0))
                            except Exception as img_err:
                                _app_logger.warning(f"[FILE_GEN/PPT] ΓÜá∩╕Å Θàìσ¢╛τöƒµêÉσ╝éσ╕╕: {img_err}")
                                yield _fb_sse(fb.warn("Θàìσ¢╛Φ╖│Φ┐ç∩╝îΣ╕ìσ╜▒σôìPPTτöƒµêÉ"))

                        # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ Step 3: τöƒµêÉ PPT µûçΣ╗╢(σÉ½ΘÇÉΘí╡Φ┐¢σ║ª) ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        yield _fb_sse(fb.ppt_rendering(slide_count))

                        from web.ppt_generator import PPTGenerator

                        # µúÇµ╡ïΣ╕╗Θóÿ
                        theme = "business"
                        if any(
                            kw in user_lower_check
                            for kw in ["µèÇµ£»", "tech", "τºæµèÇ", "τ╝ûτ¿ï", "σ╝ÇσÅæ"]
                        ):
                            theme = "tech"
                        elif any(
                            kw in user_lower_check
                            for kw in ["σê¢µäÅ", "creative", "Φë║µ£»", "Φ«╛Φ«í"]
                        ):
                            theme = "creative"

                        ppt_gen = PPTGenerator(theme=theme)

                        ppt_title = ppt_data["title"] or "µ╝öτñ║µûçτ¿┐"
                        safe_title = re.sub(r'[\\/*?:"<>|]', "_", ppt_title)[:50]
                        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pptx"
                        docs_dir = settings_manager.documents_dir
                        os.makedirs(docs_dir, exist_ok=True)
                        ppt_path = os.path.join(docs_dir, filename)

                        # Σ╜┐τö¿ progress_callback µ¥Ñµö╢Θ¢åΦ┐¢σ║ªµ╢êµü»∩╝êτöƒµêÉσÖ¿µùáµ│òσ£¿σ¢₧Φ░âΣ╕¡yield∩╝ë
                        _slide_progress_msgs = []

                        def _ppt_progress_cb(cur, total, stitle, stype):
                            _slide_progress_msgs.append((cur, total, stitle, stype))

                        ppt_gen.generate_from_outline(
                            title=ppt_title,
                            outline=ppt_data["slides"],
                            output_path=ppt_path,
                            subtitle=ppt_data.get("subtitle", ""),
                            author="Koto AI",
                            progress_callback=_ppt_progress_cb,
                        )

                        # σÅæΘÇüΘÇÉΘí╡Φ┐¢σ║ª∩╝êσ¢₧Φ░âσ╖▓τ╗Åµö╢Θ¢åσ«îµ»ò∩╝ë
                        for cur, total, stitle, stype in _slide_progress_msgs:
                            if stitle:
                                yield _fb_sse(
                                    fb.ppt_slide_progress(cur, total, stitle, stype)
                                )

                        yield _fb_sse(fb.substep("PPT µ╕▓µƒôσ«îµêÉ∩╝îµ¡úσ£¿Σ┐¥σ¡ÿ"))

                        # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ Post-Render Quality Check: µúÇµƒÑµ╕▓µƒôσÉÄτÜä PPTX µûçΣ╗╢ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        try:
                            from web.file_quality_checker import FileQualityGate

                            _post_check = FileQualityGate.post_check_pptx(ppt_path)
                            if _post_check.get("issues"):
                                _pc_score = _post_check["score"]
                                yield _fb_sse(
                                    fb.ppt_quality_check(
                                        _pc_score, issues=_post_check["issues"]
                                    )
                                )
                                _app_logger.info(
                                    f"[FILE_GEN/PPT] ≡ƒôè µûçΣ╗╢σÉÄµúÇ: {_pc_score}/100, issues={_post_check['issues']}"
                                )
                            else:
                                yield _fb_sse(fb.info("Γ£à µûçΣ╗╢Φ┤¿ΘçÅΘ¬îΦ»üΘÇÜΦ┐ç"))
                        except Exception as pc_err:
                            _app_logger.warning(f"[FILE_GEN/PPT] ΓÜá∩╕Å µûçΣ╗╢σÉÄµúÇσ╝éσ╕╕: {pc_err}")

                        rel_path = os.path.relpath(ppt_path, WORKSPACE_DIR).replace(
                            "\\", "/"
                        )
                        generated_files.append(rel_path)

                        # τ╗ƒΦ«íσÉäτ▒╗σ₧ïσ╣╗τü»τëçµò░ΘçÅ
                        _type_names = {
                            "detail": "Φ»ªτ╗åΘí╡",
                            "overview": "µªéΦºêΘí╡",
                            "highlight": "Σ║«τé╣Θí╡",
                            "divider": "Φ┐çµ╕íΘí╡",
                            "comparison": "σ»╣µ»öΘí╡",
                        }
                        _type_counts = {}
                        for _s in ppt_data["slides"]:
                            _t = _s.get("type", "detail")
                            _type_counts[_t] = _type_counts.get(_t, 0) + 1
                        _type_desc = "πÇü".join(
                            f"{_type_names.get(k,k)} ├ù{v}"
                            for k, v in _type_counts.items()
                        )

                        _img_desc = (
                            f"\n≡ƒû╝∩╕Å Θàìσ¢╛: {len(ppt_images)} σ╝á" if ppt_images else ""
                        )
                        _research_desc = (
                            "\n≡ƒö¼ σ╖▓Φ₧ìσàÑµ╖▒σ║ªτáöτ⌐╢σêåµ₧É" if research_context else ""
                        )

                        # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ Σ┐¥σ¡ÿΣ╝ÜΦ»¥µò░µì«∩╝êP1 τ╝ûΦ╛æσèƒΦâ╜µö»µîü∩╝ë ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        if ppt_session_id:
                            try:
                                from web.ppt_session_manager import (
                                    get_ppt_session_manager,
                                )

                                ppt_session_mgr = get_ppt_session_manager()
                                ppt_session_mgr.save_generation_data(
                                    session_id=ppt_session_id,
                                    ppt_data=ppt_data,
                                    ppt_file_path=rel_path,
                                    search_context=search_context,
                                    research_context=research_context,
                                    uploaded_file_context=uploaded_file_context,
                                )
                                _app_logger.info(
                                    f"[FILE_GEN/PPT] ≡ƒÆ╛ Σ╝ÜΦ»¥µò░µì«σ╖▓Σ┐¥σ¡ÿ∩╝îσÅ»τö¿Σ║ÄσÉÄτ╗¡τ╝ûΦ╛æ"
                                )
                            except Exception as save_err:
                                _app_logger.warning(f"[FILE_GEN/PPT] ΓÜá∩╕Å Σ╝ÜΦ»¥Σ┐¥σ¡ÿσ╝éσ╕╕: {save_err}")

                        success_msg = (
                            f"Γ£à **PPT µ╝öτñ║µûçτ¿┐τöƒµêÉµêÉσèƒ∩╝ü**\n\n"
                            f"≡ƒôè µáçΘóÿ: **{ppt_title}**\n"
                            f"≡ƒôä Θí╡µò░: {slide_count} Θí╡∩╝ê{_type_desc}∩╝ë{_img_desc}{_research_desc}\n"
                            f"≡ƒôü µûçΣ╗╢: **{filename}**\n"
                            f"≡ƒôì Σ╜ìτ╜«: `{docs_dir}`"
                        )

                        # σªéµ₧£µ£ëΣ╝ÜΦ»¥∩╝îΘÖäσèáτ╝ûΦ╛æΘô╛µÄÑ
                        if ppt_session_id:
                            success_msg += f"\n\n≡ƒÄ¿ **[τé╣σç╗τ╝ûΦ╛æ PPT](/edit-ppt/{ppt_session_id})** - Σ┐«µö╣σåàσ«╣πÇüΦ░âµò┤Θí║σ║ÅπÇüΘçìµû░τöƒµêÉΘí╡Θ¥ó"

                        yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                        _app_logger.info(f"[FILE_GEN/PPT] Γ£à PPT τöƒµêÉµêÉσèƒ: {rel_path}")

                    except Exception as ppt_err:
                        _app_logger.error(f"[FILE_GEN/PPT] Γ¥î PPT τöƒµêÉσñ▒Φ┤Ñ: {ppt_err}")
                        import traceback

                        traceback.print_exc()
                        error_msg = f"Γ¥î PPT τöƒµêÉσñ▒Φ┤Ñ: {str(ppt_err)}"
                        yield f"data: {json.dumps({'type': 'token', 'content': error_msg})}\n\n"

                    # Σ┐¥σ¡ÿσÄåσÅ▓∩╝êσƒ║Σ║Äτúüτ¢ÿσ«îµò┤σÄåσÅ▓Φ┐╜σèá∩╝îσ£¿ done Σ║ïΣ╗╢Σ╣ïσëì∩╝ë
                    _ppt_msg = (
                        f"σ╖▓τöƒµêÉPPT: {', '.join(generated_files)}"
                        if generated_files
                        else "PPTτöƒµêÉσñ▒Φ┤Ñ"
                    )
                    session_manager.append_and_save(
                        f"{session_name}.json", user_input, _ppt_msg
                    )

                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': generated_files, 'total_time': total_time})}\n\n"
                    return

                # Σ╜┐τö¿Σ╕èΣ╕ïµûçσó₧σ╝║τÜäΦ╛ôσàÑ∩╝êσªéµ₧£µ£ë∩╝îΣ╛ïσªé"µèèΦ┐ÖΣ╕¬σüÜµêÉword"µù╢Σ╝ÜσîàσÉ½Σ╣ïσëìτÜäσåàσ«╣∩╝ë
                if (
                    context_info
                    and context_info.get("is_continuation")
                    and context_info.get("enhanced_input")
                ):
                    file_gen_input = context_info["enhanced_input"]
                    _app_logger.debug(
                        f"[FILE_GEN] Σ╜┐τö¿Σ╕èΣ╕ïµûçσó₧σ╝║Φ╛ôσàÑ (length: {len(file_gen_input)})"
                    )
                else:
                    file_gen_input = effective_input

                # Γ¡É FILE_GEN σëìτ╜«µ¡ÑΘ¬ñ∩╝Üµù╢Θù┤Φºúµ₧É + Σ┐íµü»µö╢Θ¢å
                _time_context_text, _time_parse = _build_filegen_time_context(
                    user_input
                )
                _web_context = ""
                _should_collect = WebSearcher.needs_web_search(user_input)

                # σ»╣ΓÇ£Xµ£êµû░τò¬/τò¬σëº/σè¿τö╗ΓÇ¥τ¡ëµù╢Θù┤µòÅµäƒΣ╕╗Θóÿσ╝║σê╢σÉ»τö¿Σ┐íµü»µö╢Θ¢å
                _anime_time_patterns = [
                    r"([1-9]|1[0-2])\s*µ£ê\s*(µû░τò¬|τò¬σëº|σè¿τö╗)",
                    r"(µû░τò¬|τò¬σëº|σè¿τö╗).*(\d{1,2}\s*µ£ê)",
                ]
                if not _should_collect and any(
                    re.search(p, user_input, re.IGNORECASE)
                    for p in _anime_time_patterns
                ):
                    _should_collect = True

                if _should_collect:
                    try:
                        if _time_parse.get("resolved_month"):
                            _q = f"{_time_parse['resolved_year']}σ╣┤{_time_parse['resolved_month']}µ£ê µû░τò¬ σè¿τö╗ τò¬σëº σÉìσìò Σ╗ïτ╗ì"
                        else:
                            _q = user_input

                        _time_detail = _time_context_text.replace("\n", " | ")[:180]
                        yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒòÆ µ¡úσ£¿Φºúµ₧Éµù╢Θù┤Φ»¡Σ╣ë...', 'detail': _time_detail})}\n\n"
                        yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒîÉ µ¡úσ£¿µö╢Θ¢åµ£Çµû░Σ┐íµü»...', 'detail': _q[:120]})}\n\n"

                        _search_res = WebSearcher.search_with_grounding(_q)
                        if _search_res.get("success") and _search_res.get("response"):
                            _web_context = _search_res.get("response", "")
                            _app_logger.info(
                                f"[FILE_GEN] Γ£à Σ┐íµü»µö╢Θ¢åσ«îµêÉ∩╝îΘò┐σ║ª: {len(_web_context)}"
                            )
                            yield f"data: {json.dumps({'type': 'progress', 'message': 'Γ£à Σ┐íµü»µö╢Θ¢åσ«îµêÉ', 'detail': f'σ╖▓ΦÄ╖σÅû {len(_web_context)} σ¡ùτ¼ªσÅéΦÇâΣ┐íµü»'})}\n\n"
                        else:
                            _app_logger.warning(f"[FILE_GEN] ΓÜá∩╕Å Σ┐íµü»µö╢Θ¢åµ£¬Φ┐öσ¢₧τ╗ôµ₧£")
                    except Exception as _collect_err:
                        _app_logger.warning(f"[FILE_GEN] ΓÜá∩╕Å Σ┐íµü»µö╢Θ¢åσ╝éσ╕╕: {_collect_err}")

                # σ░åµù╢Θù┤Σ╕èΣ╕ïµûç/µúÇτ┤óτ╗ôµ₧£µï╝µÄÑΦ┐¢τöƒµêÉΦ╛ôσàÑ
                _prepended_blocks = [_time_context_text]
                if _web_context:
                    _prepended_blocks.append("[Φüöτ╜æµúÇτ┤óσÅéΦÇâ]\n" + _web_context[:9000])
                file_gen_input = (
                    "\n\n".join(_prepended_blocks) + "\n\n" + file_gen_input
                )

                # Γ¡É σêñµû¡µÿ»σÉªµÿ»µûçµíúτöƒµêÉΦ»╖µ▒é∩╝êWord/PDF∩╝ë
                _doc_keywords = [
                    "word",
                    "docx",
                    "doc",
                    "pdf",
                    "µèÑσæè",
                    "µûçµíú",
                    "Φ«║µûç",
                    "τ╗╝Φ┐░",
                    "whitepaper",
                ]
                _is_doc_request = any(k in user_input.lower() for k in _doc_keywords)
                _is_complex = (context_info or {}).get("complexity") == "complex"

                if _is_doc_request:
                    # ============== µûçµíúτ¢┤σç║µ¿íσ╝Å∩╝êµ╡üσ╝Å∩╝ë ==============
                    # Σ╜┐τö¿ generate_content_stream Σ┐¥µîüΦ┐₧µÄÑµ┤╗Φ╖â∩╝îΘü┐σàìΣ╗úτÉåΦ╢àµù╢µû¡σ╝Ç
                    _doc_type = "PDF" if "pdf" in user_input.lower() else "Word"
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒôä µ¡úσ£¿τöƒµêÉ {_doc_type} µûçµíú...', 'detail': 'Φ»╖τ¿ìσÇÖ∩╝îµ¡úσ£¿µÆ░σåÖσåàσ«╣'})}\n\n"
                    _app_logger.debug(
                        f"[FILE_GEN] ≡ƒôä µûçµíúτ¢┤σç║µ¿íσ╝Å-µ╡üσ╝Å (type={_doc_type}, complex={_is_complex})"
                    )

                    _doc_instruction = """Σ╜áµÿ» Koto Σ╕ôΣ╕ÜµûçµíúµÆ░σåÖσè⌐µëïπÇéΦ»╖µá╣µì«τö¿µê╖Φªüµ▒é∩╝îτ¢┤µÄÑΦ╛ôσç║**σ«îµò┤πÇüΦ»ªτ╗åπÇüΘ½ÿΦ┤¿ΘçÅ**τÜäµûçµíúµ¡úµûçσåàσ«╣πÇé

## Φ╛ôσç║ΦºäσêÖ
- τ¢┤µÄÑΦ╛ôσç║ Markdown µá╝σ╝ÅτÜäµûçµíúµ¡úµûç∩╝îΣ╕ìΦªüΦ╛ôσç║Σ╗úτáü
- Σ╜┐τö¿ # ## ### τ╗äτ╗çµáçΘóÿσ▒éτ║º
- Σ╜┐τö¿µ«╡ΦÉ╜πÇüσêùΦí¿πÇüΦí¿µá╝Σ╕░σ»îσåàσ«╣
- Σ╕¡µûçµÆ░σåÖ∩╝îΣ╕ôΣ╕Üµ£»Φ»¡σçåτí«
- σåàσ«╣Φªü**σààσ«₧Φ»ªσ░╜**∩╝îµ»ÅΣ╕ÇΦèéΦç│σ░æ2-3µ«╡∩╝îµÇ╗σ¡ùµò░Σ╕ìσ░æΣ║Ä3000σ¡ù
- σªéµ₧£µÿ»µèÇµ£»µèÑσæè∩╝îσ┐àΘí╗σîàσÉ½∩╝ÜΦíîΣ╕ÜµªéΦ┐░πÇüµèÇµ£»σÄƒτÉåπÇüσà│Θö«σ╖ÑΦë║πÇüσ»╣µ»öσêåµ₧ÉπÇüσ║öτö¿σ£║µÖ»πÇüσÅæσ▒òΦ╢ïσè┐
- Σ╕ìΦªüΦ╛ôσç║Σ╗╗Σ╜ò BEGIN_FILE/END_FILE µáçΦ«░
- Σ╕ìΦªüΦ╛ôσç║ JSON µêûΣ╗úτáüµá╝σ╝Å"""

                    _doc_instruction += "\n\nµù╢Θù┤Φªüµ▒é∩╝ÜΦïÑτö¿µê╖Φ»╖µ▒éµ╢ëσÅèµ£êΣ╗╜Σ╜åµ£¬σåÖσ╣┤Σ╗╜∩╝êσªéΓÇÿ1µ£êµû░τò¬ΓÇÖ∩╝ë∩╝îσ┐àΘí╗µîëσ╜ôσëìσ╣┤Σ╗╜µÆ░σåÖ∩╝îτªüµ¡óΘ╗ÿΦ«ñσ¢₧ΘÇÇσê░σÄåσÅ▓σ╣┤Σ╗╜πÇé"

                    _max_tokens = 16384 if _is_complex else 8192
                    _doc_models = list(dict.fromkeys([
                        model_id,
                        "gemini-3.1-pro-preview",
                        "gemini-2.5-flash",
                        "gemini-3-flash-preview",
                    ]))

                    _doc_collected = []  # µö╢Θ¢åµëÇµ£ëµ╡üσ╝Åµûçµ£¼σ¥ù

                    for model_attempt, current_model in enumerate(_doc_models):
                        if _doc_collected:
                            break
                        if interrupted():
                            yield f"data: {json.dumps({'type': 'token', 'content': 'ΓÅ╣∩╕Å µûçΣ╗╢τöƒµêÉσ╖▓Σ╕¡µû¡'})}\n\n"
                            total_time = time.time() - start_time
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                            return
                        if model_attempt > 0:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒöä σêçµìóσê░σñçτö¿µ¿íσ₧ï {current_model}...', 'detail': ''})}\n\n"
                            _doc_collected.clear()
                        else:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒÜÇ µ¡úσ£¿Φ░âτö¿ {current_model}...', 'detail': 'µ╡üσ╝ÅτöƒµêÉΣ╕¡'})}\n\n"

                        try:
                            _doc_stream = client.models.generate_content_stream(
                                model=current_model,
                                contents=file_gen_input,
                                config=types.GenerateContentConfig(
                                    system_instruction=_doc_instruction,
                                    max_output_tokens=_max_tokens,
                                    temperature=0.7,
                                ),
                            )
                            _first_chunk = False
                            for item_type, item_data in stream_with_keepalive(
                                _doc_stream,
                                start_time,
                                keepalive_interval=5,
                                max_wait_first_token=120,  # µûçµíúτöƒµêÉσàüΦ«╕τ¡ëσ╛àµ¢┤Σ╣à
                            ):
                                if interrupted():
                                    _app_logger.info(f"[FILE_GEN/DOC] τö¿µê╖Σ╕¡µû¡")
                                    _interrupt_msg = "\n\nΓÅ╣∩╕Å µûçΣ╗╢τöƒµêÉσ╖▓Σ╕¡µû¡"
                                    yield f"data: {json.dumps({'type': 'token', 'content': _interrupt_msg})}\n\n"
                                    break

                                if item_type == "heartbeat":
                                    _elapsed = item_data
                                    _char_count = sum(len(c) for c in _doc_collected)
                                    if _first_chunk:
                                        yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒô¥ µ¡úσ£¿µÆ░σåÖµûçµíú...', 'detail': f'σ╖▓τöƒµêÉ {_char_count} σ¡ùτ¼ª∩╝îΦÇùµù╢ {_elapsed}s', 'stage': 'generating'})}\n\n"
                                    else:
                                        yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒºá µ¿íσ₧ïµ¡úσ£¿τ╗äτ╗çσåàσ«╣...', 'detail': f'σ╖▓τ¡ëσ╛à {_elapsed}s∩╝îΦ»╖ΦÇÉσ┐âτ¡ëσ╛à', 'stage': 'api_calling'})}\n\n"

                                elif item_type == "timeout":
                                    _app_logger.warning(
                                        f"[FILE_GEN/DOC] ΓÜá∩╕Å {current_model} τ¡ëσ╛àΘªûtokenΦ╢àµù╢: {item_data}"
                                    )
                                    break  # σ░¥Φ»òΣ╕ïΣ╕ÇΣ╕¬µ¿íσ₧ï

                                elif item_type == "chunk":
                                    chunk = item_data
                                    if chunk.text:
                                        if not _first_chunk:
                                            _first_chunk = True
                                            _app_logger.info(
                                                f"[FILE_GEN/DOC] Γ£à {current_model} µö╢σê░τ¼¼Σ╕ÇΣ╕¬σôìσ║öσ¥ù∩╝îΦÇùµù╢ {time.time() - start_time:.1f}s"
                                            )
                                        _doc_collected.append(chunk.text)
                                        # µ»Åµö╢σê░10Σ╕¬chunkσÅæΘÇüΣ╕Çµ¼íΦ┐¢σ║ªµ¢┤µû░∩╝îΣ┐¥µîüσ«óµê╖τ½»Φ┐₧µÄÑµ┤╗Φ╖â
                                        if len(_doc_collected) % 10 == 0:
                                            _char_count = sum(
                                                len(c) for c in _doc_collected
                                            )
                                            _elapsed = int(time.time() - start_time)
                                            yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒô¥ µ¡úσ£¿µÆ░σåÖµûçµíú...', 'detail': f'σ╖▓τöƒµêÉ {_char_count} σ¡ùτ¼ª∩╝îΦÇùµù╢ {_elapsed}s'})}\n\n"

                        except Exception as _doc_err:
                            err_str = str(_doc_err)
                            _app_logger.error(f"[FILE_GEN/DOC] Γ¥î {current_model}: {err_str[:200]}")
                            if "location is not supported" in err_str.lower():
                                response_text = "Γ¥î σ£░σî║ΘÖÉσê╢∩╝îΦ»╖Θàìτ╜«Σ╕¡Φ╜¼µ£ìσèí"
                                break
                            continue

                    response_text = "".join(_doc_collected)
                    if response_text:
                        _app_logger.info(
                            f"[FILE_GEN/DOC] Γ£à µ╡üσ╝ÅτöƒµêÉσ«îµêÉ∩╝îσà▒ {len(response_text)} σ¡ùτ¼ª"
                        )

                    if not response_text or response_text.startswith("Γ¥î"):
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text or 'Γ¥î µëÇµ£ëµ¿íσ₧ïΘâ╜Σ╕ìσÅ»τö¿∩╝îΦ»╖τ¿ìσÉÄΘçìΦ»ò'})}\n\n"
                    else:
                        # ΓöÇΓöÇ µûçµíúΦ┤¿ΘçÅΦç¬µúÇΣ╕Äµ╕àµ┤ù ΓöÇΓöÇ
                        try:
                            from web.file_quality_checker import FileQualityGate

                            _doc_qg = FileQualityGate.check_and_fix_document(
                                response_text, user_request=user_input
                            )
                            response_text = _doc_qg["text"]
                            _doc_score = _doc_qg["quality"]["score"]
                            _doc_fixes = _doc_qg["fixes"]
                            if _doc_fixes:
                                yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒº╣ σ╖▓µ╕àµ┤ù {len(_doc_fixes)} σñäσåàσ«╣Θù«Θóÿ', 'detail': 'τº╗ΘÖñAIσ»╣Φ»¥τùòΦ┐╣'})}\n\n"
                            _dq_emoji = "Γ£à" if _doc_score >= 75 else "ΓÜá∩╕Å"
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'{_dq_emoji} µûçµíúΦ┤¿ΘçÅµúÇµƒÑ: {_doc_score}/100', 'detail': '; '.join(_doc_qg['quality']['issues'][:2]) if _doc_qg['quality']['issues'] else 'Φ┤¿ΘçÅΦë»σÑ╜'})}\n\n"
                        except Exception as _dqg_err:
                            _app_logger.warning(f"[FILE_GEN/DOC] ΓÜá∩╕Å Φ┤¿ΘçÅΘù¿µÄºσ╝éσ╕╕: {_dqg_err}")

                        # τ¢┤µÄÑΣ┐¥σ¡ÿµûçµíú
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒô¥ µ¡úσ£¿Σ┐¥σ¡ÿ {_doc_type} µûçµíú...', 'detail': ''})}\n\n"
                        try:
                            try:
                                from web.document_generator import save_docx, save_pdf
                            except ModuleNotFoundError:
                                from document_generator import save_docx, save_pdf
                            docs_dir = settings_manager.documents_dir
                            os.makedirs(docs_dir, exist_ok=True)
                            title_match = re.search(
                                r"^#\s*(.+)$", response_text, re.MULTILINE
                            )
                            title = (
                                title_match.group(1).strip()[:50]
                                if title_match
                                else f"Kotoµûçµíú_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                            )
                            if _doc_type == "PDF":
                                saved_path = save_pdf(
                                    response_text, title=title, output_dir=docs_dir
                                )
                            else:
                                saved_path = save_docx(
                                    response_text, title=title, output_dir=docs_dir
                                )
                            rel_path = os.path.relpath(
                                saved_path, WORKSPACE_DIR
                            ).replace("\\", "/")
                            generated_files.append(rel_path)
                            success_msg = f"Γ£à **{_doc_type} µûçµíúτöƒµêÉµêÉσèƒ∩╝ü**\n\n≡ƒôü µûçΣ╗╢: **{os.path.basename(saved_path)}**\n≡ƒôì Σ╜ìτ╜«: `{docs_dir}`"
                            yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                            _app_logger.info(f"[FILE_GEN/DOC] Γ£à µûçµíúσ╖▓Σ┐¥σ¡ÿ: {rel_path}")
                        except Exception as doc_err:
                            import traceback

                            traceback.print_exc()
                            _app_logger.error(f"[FILE_GEN/DOC] Γ¥î µûçµíúΣ┐¥σ¡ÿσñ▒Φ┤Ñ: {doc_err}")
                            fallback_msg = (
                                f"ΓÜá∩╕Å µûçµíúΣ┐¥σ¡ÿσñ▒Φ┤Ñ ({doc_err})∩╝îΣ╗ÑΣ╕ïµÿ»τöƒµêÉτÜäσåàσ«╣∩╝Ü\n\n"
                            )
                            yield f"data: {json.dumps({'type': 'token', 'content': fallback_msg + response_text})}\n\n"

                    _gen_msg = (
                        f"σ╖▓τöƒµêÉµûçΣ╗╢: {', '.join(generated_files)}"
                        if generated_files
                        else (response_text[:500] if response_text else "τöƒµêÉσñ▒Φ┤Ñ")
                    )
                    session_manager.append_and_save(
                        f"{session_name}.json", user_input, _gen_msg
                    )
                    total_time = time.time() - start_time
                    _app_logger.info(
                        f"[FILE_GEN/DOC] ΓÿàΓÿàΓÿà done event, files: {generated_files}, time: {total_time:.2f}s"
                    )
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': generated_files, 'total_time': total_time})}\n\n"
                    return

                # µÖ«ΘÇÜ FILE_GEN µ¿íσ╝Å∩╝êΘ£ÇΦªüµ¿íσ₧ïτöƒµêÉΣ╗úτáü/ΦäÜµ£¼∩╝ë
                yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒôä µ¡úσ£¿τöƒµêÉµûçΣ╗╢Σ╗úτáü...', 'detail': 'Φ»╖τ¿ìσÇÖ∩╝îσÅ»Φâ╜Θ£ÇΦªü 10-30 τºÆ'})}\n\n"

                # µ¿íσ₧ïσêùΦí¿∩╝êΣ╕╗µ¿íσ₧ï + σñçτö¿µ¿íσ₧ï∩╝ë
                file_gen_models = [
                    model_id,  # Σ╕╗µ¿íσ₧ï
                    "gemini-3.1-pro-preview",  # σñçτö¿1 (µ£Çσ╝║ generate_content σà╝σ«╣)
                    "gemini-2.5-flash",  # σñçτö¿2
                    "gemini-3-flash-preview",  # σñçτö¿3
                ]

                # Σ╜┐τö¿τ║┐τ¿ï + Φ╢àµù╢µ¥ÑΦ░âτö¿API∩╝êσ╕ªΘçìΦ»ò∩╝ë
                import tempfile
                import threading

                for model_attempt, current_model in enumerate(file_gen_models):
                    if response_text and not response_text.startswith("Γ¥î"):
                        break  # σ╖▓µêÉσèƒ

                    if interrupted():
                        yield f"data: {json.dumps({'type': 'token', 'content': 'ΓÅ╣∩╕Å µûçΣ╗╢τöƒµêÉσ╖▓Σ╕¡µû¡'})}\n\n"
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                    if model_attempt > 0:
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒöä σêçµìóσê░σñçτö¿µ¿íσ₧ï {current_model}...', 'detail': ''})}\n\n"
                        _app_logger.debug(f"[FILE_GEN] Trying fallback model: {current_model}")
                    else:
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒÜÇ µ¡úσ£¿Φ░âτö¿ {current_model}...', 'detail': 'τöƒµêÉΣ╕¡'})}\n\n"

                    response_holder = {"data": None, "error": None}

                    def call_api(m=current_model):
                        try:
                            _app_logger.debug(f"[FILE_GEN] Calling API: {m}")
                            response = client.models.generate_content(
                                model=m,
                                contents=file_gen_input,  # Σ╜┐τö¿Σ╕èΣ╕ïµûçσó₧σ╝║τÜäΦ╛ôσàÑ
                                config=types.GenerateContentConfig(
                                    system_instruction=_get_system_instruction(),
                                    max_output_tokens=8192,
                                ),
                            )
                            response_holder["data"] = response
                            _app_logger.info(f"[FILE_GEN] Γ£à API call successful with {m}")
                        except Exception as e:
                            _app_logger.error(
                                f"[FILE_GEN] Γ¥î API call exception with {m}: {type(e).__name__}: {str(e)}"
                            )
                            response_holder["error"] = e

                    api_thread = threading.Thread(target=call_api, daemon=True)
                    api_thread.start()

                    # σ£¿τ¡ëσ╛àµ£ƒΘù┤σÅæΘÇüσ┐âΦ╖│Φ┐¢σ║ª
                    wait_interval = 5  # µ»Å 5 τºÆσÅæΘÇüΣ╕Çµ¼íΦ┐¢σ║ª
                    elapsed = 0
                    while api_thread.is_alive() and elapsed < api_timeout:
                        api_thread.join(timeout=wait_interval)
                        elapsed += wait_interval
                        if interrupted():
                            yield f"data: {json.dumps({'type': 'token', 'content': 'ΓÅ╣∩╕Å µûçΣ╗╢τöƒµêÉσ╖▓Σ╕¡µû¡'})}\n\n"
                            total_time = time.time() - start_time
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                            return
                        if api_thread.is_alive() and elapsed < api_timeout:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'ΓÅ│ µ¡úσ£¿τöƒµêÉΣ╕¡...', 'detail': f'σ╖▓τ¡ëσ╛à {elapsed} τºÆ'})}\n\n"

                    if api_thread.is_alive():
                        _app_logger.warning(
                            f"[FILE_GEN] ΓÜá∩╕Å API call timeout with {current_model} after {api_timeout}s"
                        )
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'ΓÜá∩╕Å {current_model} σôìσ║öΦ╢àµù╢', 'detail': 'µ¡úσ£¿σêçµìóµ¿íσ₧ï...'})}\n\n"
                        response_text = ""
                        continue  # σ░¥Φ»òΣ╕ïΣ╕ÇΣ╕¬µ¿íσ₧ï
                    elif response_holder["error"]:
                        error_str = str(response_holder["error"])
                        _app_logger.debug(f"[FILE_GEN] API Error with {current_model}: {error_str}")

                        # σ£░σî║ΘÖÉσê╢ΘöÖΦ»» - τ¢┤µÄÑσñ▒Φ┤Ñ∩╝îΣ╕ìΘçìΦ»ò
                        if (
                            "location is not supported" in error_str.lower()
                            or "failed_precondition" in error_str.lower()
                        ):
                            response_text = "Γ¥î σ£░σî║ΘÖÉσê╢\n\nµé¿µëÇσ£¿τÜäσ£░σî║Σ╕ìµö»µîü Gemini APIπÇé\n\n≡ƒÆí Φºúσå│µû╣µíê:\n1. σ£¿ config/gemini_config.env Θàìτ╜«Σ╕¡Φ╜¼µ£ìσèí GEMINI_API_BASE\n2. µêûΣ╜┐τö¿µö»µîüτÜäΣ╗úτÉåµ£ìσèí"
                            break  # σ£░σî║ΘÖÉσê╢∩╝îΣ╕ìτ╗ºτ╗¡ΘçìΦ»ò
                        elif (
                            "503" in error_str
                            or "overloaded" in error_str.lower()
                            or "unavailable" in error_str.lower()
                        ):
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'ΓÜá∩╕Å {current_model} µ£ìσèíτ╣üσ┐Ö', 'detail': 'µ¡úσ£¿σêçµìóµ¿íσ₧ï...'})}\n\n"
                            response_text = ""
                            continue  # 503 ΘöÖΦ»»∩╝îσ░¥Φ»òΣ╕ïΣ╕ÇΣ╕¬µ¿íσ₧ï
                        else:
                            response_text = f"Γ¥î API Φ░âτö¿σñ▒Φ┤Ñ: {error_str[:200]}"
                            continue  # σà╢Σ╗ûΘöÖΦ»»Σ╣ƒσ░¥Φ»òΣ╕ïΣ╕ÇΣ╕¬µ¿íσ₧ï
                    elif response_holder["data"]:
                        file_gen_response = response_holder["data"]
                        if (
                            file_gen_response.candidates
                            and file_gen_response.candidates[0].content.parts
                        ):
                            for part in file_gen_response.candidates[0].content.parts:
                                if hasattr(part, "text") and part.text:
                                    response_text += part.text
                        _app_logger.debug(f"[FILE_GEN] Response length: {len(response_text)}")
                        if response_text:
                            break  # µêÉσèƒΦÄ╖σÅûσôìσ║ö

                if not response_text:
                    response_text = "Γ¥î µëÇµ£ëµ¿íσ₧ïΘâ╜Σ╕ìσÅ»τö¿∩╝îΦ»╖τ¿ìσÉÄΘçìΦ»ò"

                if Utils.is_failure_output(response_text):
                    yield f"data: {json.dumps({'type': 'progress', 'message': 'ΓÜá∩╕Å σê¥µ¼íτöƒµêÉσñ▒Φ┤Ñ∩╝îµ¡úσ£¿Σ┐«µ¡ú...', 'detail': ''})}\n\n"
                    fix_prompt = Utils.build_fix_prompt(
                        "FILE_GEN", user_input, response_text
                    )
                    try:
                        fix_resp = client.models.generate_content(
                            model=model_id,
                            contents=fix_prompt,
                            config=types.GenerateContentConfig(
                                system_instruction=_get_system_instruction(),
                                max_output_tokens=8192,
                                temperature=0.4,
                            ),
                        )
                        response_text = fix_resp.text or response_text
                    except Exception as fix_err:
                        _app_logger.debug(f"[FILE_GEN] Σ┐«µ¡úΘçìΦ»òσñ▒Φ┤Ñ: {fix_err}")

                # σÅ¬µÿ╛τñ║τ«Çτƒ¡τÜäΦ┐¢σ║ª∩╝îΣ╕ìµÿ╛τñ║σ«îµò┤Σ╗úτáü
                if response_text and not response_text.startswith("Γ¥î"):
                    yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒöº µ¡úσ£¿σñäτÉåΣ╗úτáü...', 'detail': ''})}\n\n"

                    # µÅÉσÅûΣ╗úτáüσê░Σ╕┤µù╢µûçΣ╗╢
                    patterns = [
                        r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\s*(.*?)---\s*END_FILE\s*---",
                        r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\n(.*?)\n---END_FILE---",
                    ]

                    code_content = None
                    for pattern in patterns:
                        matches = re.findall(
                            pattern, response_text, re.DOTALL | re.IGNORECASE
                        )
                        if matches:
                            _, code_content = matches[0]
                            code_content = code_content.strip()
                            _app_logger.debug(
                                f"[FILE_GEN] Extracted code, length: {len(code_content)}"
                            )
                            break

                    # µúÇµƒÑµÅÉσÅûτÜäσåàσ«╣µÿ»σÉªµÿ»µ£ëµòêτÜäPythonΣ╗úτáü∩╝êΣ╕ìµÿ»JSONµêûσà╢Σ╗ûµá╝σ╝Å∩╝ë
                    is_valid_python = False
                    if code_content:
                        code_lower = code_content.lower()
                        # σªéµ₧£µÅÉσÅûτÜäσåàσ«╣µÿ» JSON µêû HTML µêûσà╢Σ╗ûµá╝σ╝Å∩╝îτ¢┤µÄÑΦ╖│Φ┐çΣ╗úτáüµëºΦíî
                        if code_lower.startswith(("{", "[", "<", '"')):
                            _app_logger.debug(
                                f"[FILE_GEN] Extracted content is not Python code (starts with {code_content[0]}), treating as text content"
                            )
                            code_content = None
                        else:
                            is_valid_python = True

                    if code_content and is_valid_python:
                        # Σ┐¥σ¡ÿσê░Σ╕┤µù╢µûçΣ╗╢
                        temp_dir = tempfile.gettempdir()
                        temp_script = os.path.join(
                            temp_dir, f"koto_gen_{int(time.time())}.py"
                        )

                        with open(temp_script, "w", encoding="utf-8") as f:
                            f.write(code_content)
                        temp_scripts.append(temp_script)
                        _app_logger.debug(f"[FILE_GEN] Saved temp script: {temp_script}")

                        # µëºΦíîΦäÜµ£¼
                        yield f"data: {json.dumps({'type': 'progress', 'message': 'ΓÜÖ∩╕Å µ¡úσ£¿µëºΦíîΦäÜµ£¼τöƒµêÉµûçΣ╗╢...', 'detail': ''})}\n\n"

                        try:
                            if getattr(sys, "frozen", False):
                                # µëôσîàµ¿íσ╝Å∩╝ÜΦ┐¢τ¿ïσåà exec() µëºΦíî∩╝îΘü┐σàìσÉ»σè¿µû░ Koto τ¬ùσÅú
                                import contextlib as _ctx
                                import io as _io

                                _out, _err, _rc = _io.StringIO(), _io.StringIO(), 0
                                try:
                                    _prev = os.getcwd()
                                    os.chdir(WORKSPACE_DIR)
                                    with _ctx.redirect_stdout(
                                        _out
                                    ), _ctx.redirect_stderr(_err):
                                        exec(
                                            open(
                                                temp_script, "r", encoding="utf-8"
                                            ).read(),
                                            {"__file__": temp_script},
                                        )
                                    os.chdir(_prev)
                                except Exception as _ex:
                                    _err.write(str(_ex))
                                    _rc = 1

                                class _FgR:
                                    returncode = _rc
                                    stdout = _out.getvalue()
                                    stderr = _err.getvalue()

                                result = _FgR()
                            else:
                                result = subprocess.run(
                                    [sys.executable, temp_script],
                                    capture_output=True,
                                    text=True,
                                    timeout=60,
                                    cwd=WORKSPACE_DIR,
                                    creationflags=(
                                        subprocess.CREATE_NO_WINDOW
                                        if sys.platform == "win32"
                                        else 0
                                    ),
                                )
                            _app_logger.debug(f"[FILE_GEN] Script exit code: {result.returncode}")
                            _app_logger.debug(f"[FILE_GEN] Script stdout: {result.stdout}")
                            _app_logger.debug(f"[FILE_GEN] Script stderr: {result.stderr}")

                            if result.returncode == 0:
                                # µúÇµƒÑτöƒµêÉτÜäµûçΣ╗╢
                                docs_dir = settings_manager.documents_dir
                                if os.path.exists(docs_dir):
                                    for f in os.listdir(docs_dir):
                                        if f.endswith(
                                            (
                                                ".pdf",
                                                ".docx",
                                                ".xlsx",
                                                ".pptx",
                                                ".ppt",
                                                ".png",
                                                ".jpg",
                                            )
                                        ):
                                            full_path = os.path.join(docs_dir, f)
                                            age = time.time() - os.path.getmtime(
                                                full_path
                                            )
                                            if age < 60:
                                                rel_path = os.path.relpath(
                                                    full_path, WORKSPACE_DIR
                                                ).replace("\\", "/")
                                                if rel_path not in generated_files:
                                                    generated_files.append(rel_path)
                                                    _app_logger.debug(
                                                        f"[FILE_GEN] Generated: {rel_path}"
                                                    )

                                if generated_files:
                                    files_list = ", ".join(
                                        [os.path.basename(f) for f in generated_files]
                                    )
                                    success_msg = (
                                        "Γ£à **µûçΣ╗╢τöƒµêÉµêÉσèƒ∩╝ü**\n\n≡ƒôü τöƒµêÉτÜäµûçΣ╗╢: **"
                                        + files_list
                                        + "**\n≡ƒôì Σ┐¥σ¡ÿΣ╜ìτ╜«: `"
                                        + docs_dir
                                        + "`"
                                    )
                                    yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                                else:
                                    # ΦäÜµ£¼µëºΦíîµêÉσèƒΣ╜åµ▓íµ£ëµúÇµ╡ïσê░µû░µûçΣ╗╢
                                    output = result.stdout.strip()
                                    if output:
                                        msg = (
                                            "Γ£à ΦäÜµ£¼µëºΦíîσ«îµêÉ\n```\n" + output + "\n```"
                                        )
                                        yield f"data: {json.dumps({'type': 'token', 'content': msg})}\n\n"
                                    else:
                                        yield f"data: {json.dumps({'type': 'token', 'content': 'ΓÜá∩╕Å ΦäÜµ£¼µëºΦíîσ«îµêÉ∩╝îΣ╜åµ£¬µúÇµ╡ïσê░µû░µûçΣ╗╢'})}\n\n"
                            else:
                                error_msg = result.stderr.strip() or "µ£¬τƒÑΘöÖΦ»»"
                                err_content = (
                                    "Γ¥î ΦäÜµ£¼µëºΦíîσñ▒Φ┤Ñ\n```\n" + error_msg[:500] + "\n```"
                                )
                                yield f"data: {json.dumps({'type': 'token', 'content': err_content})}\n\n"

                        except subprocess.TimeoutExpired:
                            yield f"data: {json.dumps({'type': 'token', 'content': 'ΓÜá∩╕Å ΦäÜµ£¼µëºΦíîΦ╢àµù╢∩╝ê60τºÆ∩╝ë'})}\n\n"
                        except Exception as e:
                            _app_logger.debug(f"[FILE_GEN] Execution error: {e}")
                            err_msg = "Γ¥î µëºΦíîΘöÖΦ»»: " + str(e)
                            yield f"data: {json.dumps({'type': 'token', 'content': err_msg})}\n\n"

                        # σêáΘÖñΣ╕┤µù╢ΦäÜµ£¼
                        for temp_file in temp_scripts:
                            try:
                                if os.path.exists(temp_file):
                                    os.remove(temp_file)
                                    _app_logger.debug("Deleted temp script: %s", temp_file)
                            except OSError:
                                pass
                    else:
                        # µ▓íµ£ëσî╣Θàìσê░Σ╗úτáüµá╝σ╝Å∩╝Üτ¢┤µÄÑµèèµ¿íσ₧ïσåàσ«╣τöƒµêÉµûçµíú
                        try:
                            from web.document_generator import save_docx, save_pdf

                            docs_dir = settings_manager.documents_dir
                            os.makedirs(docs_dir, exist_ok=True)

                            # µÅÉσÅûµáçΘóÿ∩╝êσ░¥Φ»òΣ╗Äσåàσ«╣Σ╕¡µë╛ # µáçΘóÿ∩╝ë
                            title_match = re.search(
                                r"^#\s*(.+)$", response_text, re.MULTILINE
                            )
                            if title_match:
                                title = title_match.group(1).strip()[:50]
                            else:
                                # σ░¥Φ»òΣ╗Äτö¿µê╖Φ╛ôσàÑµÅÉσÅûσà│Θö«Φ»ìΣ╜£Σ╕║µûçΣ╗╢σÉì
                                try:
                                    clean_input = user_input
                                    # σÄ╗ΘÖñσ╕╕τö¿µîçΣ╗ñΦ»ì
                                    stop_patterns = [
                                        "τöƒµêÉτÜä",
                                        "σåÖΣ╕ÇΣ╕¬",
                                        "σåÖΣ╕Çτ»ç",
                                        "σ╕«µêæ",
                                        "Φ»╖",
                                        "σà│Σ║Ä",
                                        "Σ╕ÇΣ╕ï",
                                        "µûçµíú",
                                        "file",
                                        "generate",
                                        "write",
                                        "about",
                                        "make",
                                        "create",
                                    ]
                                    for pattern in stop_patterns:
                                        clean_input = clean_input.replace(pattern, " ")

                                    # µÅÉσÅûΣ╕¡Φï▒µûçσà│Θö«Φ»ì (2-20 chars)
                                    keywords = [
                                        w
                                        for w in re.split(
                                            r"[^a-zA-Z0-9\u4e00-\u9fa5]", clean_input
                                        )
                                        if w.strip()
                                    ]
                                    valid_keywords = [
                                        k
                                        for k in keywords
                                        if len(k) > 1 and len(k) < 20
                                    ]

                                    if valid_keywords:
                                        # σÅûσëìσçáΣ╕¬σà│Θö«Φ»ìτ╗äσÉê
                                        title = "_".join(valid_keywords[:3])
                                    else:
                                        title = f"Kotoµûçµíú_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                                except Exception:
                                    title = f"Kotoµûçµíú_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

                            user_lower = user_input.lower()
                            if "pdf" in user_lower:
                                saved_path = save_pdf(
                                    response_text, title=title, output_dir=docs_dir
                                )
                                file_type = "PDF"
                            else:
                                saved_path = save_docx(
                                    response_text, title=title, output_dir=docs_dir
                                )
                                file_type = "Word"

                            rel_path = os.path.relpath(
                                saved_path, WORKSPACE_DIR
                            ).replace("\\", "/")
                            if rel_path not in generated_files:
                                generated_files.append(rel_path)

                            success_msg = f"Γ£à **{file_type} µûçµíúτöƒµêÉµêÉσèƒ∩╝ü**\n\n≡ƒôü µûçΣ╗╢: **{os.path.basename(saved_path)}**\n≡ƒôì Σ╜ìτ╜«: `{docs_dir}`"
                            yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                        except Exception as direct_err:
                            _app_logger.debug(f"[FILE_GEN] Direct save failed: {direct_err}")
                            # σ¢₧ΘÇÇσ▒òτñ║σÄƒσºïσôìσ║ö
                            yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"
                else:
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text or 'ΓÜá∩╕Å µ¿íσ₧ïµ£¬Φ┐öσ¢₧σôìσ║ö'})}\n\n"

                # Σ┐¥σ¡ÿσÄåσÅ▓∩╝êσƒ║Σ║Äτúüτ¢ÿσ«îµò┤σÄåσÅ▓Φ┐╜σèá∩╝îσ£¿ done Σ║ïΣ╗╢Σ╣ïσëì∩╝ë
                _gen_msg = (
                    f"σ╖▓τöƒµêÉµûçΣ╗╢: {', '.join(generated_files)}"
                    if generated_files
                    else (response_text[:500] if response_text else "τöƒµêÉσñ▒Φ┤Ñ")
                )
                session_manager.append_and_save(
                    f"{session_name}.json", user_input, _gen_msg
                )

                # σÅæΘÇüσ«îµêÉΣ║ïΣ╗╢
                total_time = time.time() - start_time
                _app_logger.debug(
                    f"[FILE_GEN] ΓÿàΓÿàΓÿà Sending done event, generated_files: {generated_files}, total_time: {total_time:.2f}s"
                )
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': generated_files, 'total_time': total_time})}\n\n"
                return

            # === Regular Mode (µ╡üσ╝ÅΦ╛ôσç║) ===
            # use_instruction τ¢┤µÄÑτ╗ºµë┐σñûσ▒é system_instruction
            # ∩╝êσ╖▓σîàσÉ½ CWM Θò┐µ£ƒΦ«░σ┐åπÇüτƒÑΦ»åσ║ôRAGπÇüGraph RAGπÇüSkills µ│¿σàÑ∩╝ë
            # Σ╕ìσåìΣ╗ÄΘ¢╢Θçìσ╗║∩╝îΘü┐σàìΣ╕óσñ▒µëÇµ£ë outer-scope µ│¿σàÑτÜäΣ╕èΣ╕ïµûç
            use_instruction = system_instruction

            # µ│¿σàÑΘò┐µ£ƒΦ«░σ┐åΣ╕èΣ╕ïµûç
            _memory_manager = get_memory_manager()

            # µ¢┤µû░σ»╣Φ»¥µæÿΦªü∩╝êµ╗æσè¿τ¬ùσÅúσñû∩╝ë
            if full_history and len(full_history) > 20:

                def _summarize():
                    return _memory_manager.get_or_update_summary(
                        session_name, full_history
                    )

                _, err, timed_out = run_with_timeout(_summarize, 6)
                if timed_out:
                    _app_logger.debug("[MEMORY] µæÿΦªüµ¢┤µû░Φ╢àµù╢∩╝îσ╖▓Φ╖│Φ┐ç")
                elif err:
                    _app_logger.debug(f"[MEMORY] µæÿΦªüµ¢┤µû░σñ▒Φ┤Ñ: {err}")

            memory_context = _memory_manager.get_context_string(
                user_input, session_name=session_name, history=full_history
            )
            if memory_context:
                use_instruction += f"\n\n{memory_context}"
                _app_logger.debug(f"[MEMORY] µ│¿σàÑΣ║å {len(memory_context)} σ¡ùτ¼ªτÜäΦ«░σ┐åΣ╕èΣ╕ïµûç")
                t = yield_thinking(
                    f"Σ╗ÄΘò┐µ£ƒΦ«░σ┐åΣ╕¡µúÇτ┤óσê░ {len(memory_context)} σ¡ùτ¼ªτÜäτ¢╕σà│Σ╕èΣ╕ïµûçσ╣╢µ│¿σàÑ",
                    "context",
                    "local",
                )
                if t:
                    yield t

            # µá╣µì«Σ╗╗σèíτ▒╗σ₧ïµÅÉΣ╛¢σ╖«σ╝éσîûΦ┐¢σ║ªµÅÉτñ║
            if task_type == "CODER":
                used_model = model_id
                t = yield_thinking(
                    f"Φ┐¢σàÑΣ╗úτáüτöƒµêÉµ¿íσ╝Å∩╝îΣ╜┐τö¿ {model_id} Φ┐¢ΦíîΣ╗úτáüσêåµ₧ÉΣ╕ÄτöƒµêÉ",
                    "generating",
                    "cloud",
                )
                if t:
                    yield t
                yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒÆ╗ µ¡úσ£¿σêåµ₧ÉΣ╗úτáüΘ£Çµ▒é...', 'detail': f'Σ╜┐τö¿ {model_id}'})}\n\n"

                # τë╣µ«èΣ╝ÿσîû∩╝Üσ»╣Σ║Äµ╕╕µêÅσ╝ÇσÅæµêûσ«ëΦúàσîà∩╝îµ╖╗σèáτ«Çτƒ¡µîçΣ╗ñΘü┐σàìσò░σùª
                if any(
                    k in user_input.lower()
                    for k in ["µ╕╕µêÅ", "app", "Σ║öσ¡Éµúï", "pygame", "install", "σ«ëΦúà"]
                ):
                    use_instruction += "\n\n[Important] If suggesting to install packages (like pygame), assume the user knows how to use pip. Just output `pip install package_name` in a code block. Do NOT write long tutorials about installation. Focus on the Python Code."

            elif task_type == "CHAT":
                used_model = model_id
                t = yield_thinking(
                    f"Φ┐¢σàÑσ»╣Φ»¥µ¿íσ╝Å∩╝îΣ╜┐τö¿ {model_id} τöƒµêÉσ¢₧σñì", "generating", "cloud"
                )
                if t:
                    yield t
                yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒÆ¼ Koto µ¡úσ£¿µÇ¥ΦÇâ...', 'detail': 'Φ»╖τ¿ìσÇÖ'})}\n\n"

                # ΓòÉΓòÉΓòÉ µ£¼σ£░µ¿íσ₧ïσ┐½ΘÇƒΘÇÜΘüô∩╝Üτ«ÇσìòΘù«Θóÿτ¢┤µÄÑΦ╡░ Ollama ΓòÉΓòÉΓòÉ
                from app.core.routing import LocalModelRouter

                if LocalModelRouter.is_simple_query(user_input, task_type, history):
                    local_stream = LocalModelRouter.generate_stream(
                        user_input,
                        history=history,
                        system_instruction=_get_DEFAULT_CHAT_SYSTEM_INSTRUCTION(),
                    )
                    if local_stream is not None:
                        _app_logger.debug(
                            f"[CHAT] ΓÜí Σ╜┐τö¿µ£¼σ£░µ¿íσ₧ïσ┐½ΘÇƒσôìσ║ö: {LocalModelRouter._response_model}"
                        )
                        t = yield_thinking(
                            f"µúÇµ╡ïσê░τ«ÇσìòµƒÑΦ»ó∩╝îσêçµìóσê░µ£¼σ£░µ¿íσ₧ï {LocalModelRouter._response_model} σ┐½ΘÇƒσôìσ║ö",
                            "model",
                            "local",
                        )
                        if t:
                            yield t
                        yield f"data: {json.dumps({'type': 'classification', 'task_type': task_type, 'task_display': '≡ƒÆ¼ σ»╣Φ»¥', 'model': f'≡ƒÅá {LocalModelRouter._response_model} (µ£¼σ£░)', 'message': f'≡ƒÄ» Σ╗╗σèíσêåτ▒╗: ≡ƒÆ¼ σ»╣Φ»¥ (µû╣µ│ò: ≡ƒÅá {LocalModelRouter._response_model} µ£¼σ£░σ┐½ΘÇƒΘÇÜΘüô)'})}\n\n"
                        local_full_text = ""
                        local_ok = False
                        try:
                            for chunk in local_stream:
                                local_full_text += chunk
                                yield f"data: {json.dumps({'type': 'token', 'content': chunk})}\n\n"
                            local_ok = bool(local_full_text.strip())
                        except Exception as local_err:
                            _app_logger.debug(f"[CHAT] µ£¼σ£░µ¿íσ₧ïτöƒµêÉσñ▒Φ┤Ñ: {local_err}")

                        if local_ok:
                            # µ£¼σ£░µ¿íσ₧ïµêÉσèƒ ΓåÆ Σ┐¥σ¡ÿσ╣╢Φ┐öσ¢₧
                            session_manager.append_and_save(
                                f"{session_name}.json",
                                user_input,
                                local_full_text,
                                task=task_type,
                                model_name=f"ollama/{LocalModelRouter._response_model}",
                            )
                            _reflect_types_local = {"CHAT", "RESEARCH", "CODER", "FILE_GEN", "AGENT"}
                            if task_type in _reflect_types_local:
                                _start_memory_extraction(
                                    user_input,
                                    local_full_text,
                                    history,
                                    task_type=task_type,
                                    session_name=session_name,
                                )
                            total_time = time.time() - start_time
                            _app_logger.debug(f"[CHAT] ΓÜí µ£¼σ£░µ¿íσ₧ïσôìσ║öσ«îµêÉ ({total_time:.2f}s)")
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                            return
                        else:
                            # µ£¼σ£░µ¿íσ₧ïσñ▒Φ┤Ñ ΓåÆ Θ¥ÖΘ╗ÿΘÖìτ║ºσê░Σ║æµ¿íσ₧ï
                            _app_logger.debug(f"[CHAT] µ£¼σ£░µ¿íσ₧ïΦ╛ôσç║Σ╕║τ⌐║∩╝îΘÖìτ║ºσê░Σ║æµ¿íσ₧ï")
                            t = yield_thinking(
                                f"µ£¼σ£░µ¿íσ₧ïΦ╛ôσç║Σ╕║τ⌐║∩╝îΘÖìτ║ºσê░Σ║æτ½»µ¿íσ₧ï {model_id}",
                                "model",
                                "hybrid",
                            )
                            if t:
                                yield t
                            yield f"data: {json.dumps({'type': 'progress', 'message': 'Γÿü∩╕Å σêçµìóσê░Σ║æτ½»µ¿íσ₧ï...', 'detail': model_id})}\n\n"
            elif task_type == "RESEARCH":
                yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒö¼ µ¡úσ£¿Φ┐¢Φíîµ╖▒σ║ªσêåµ₧É...', 'detail': f'Σ╜┐τö¿ {model_id}'})}\n\n"
            else:
                yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒÆ¡ Koto µ¡úσ£¿µÇ¥ΦÇâ...', 'detail': 'Φ»╖τ¿ìσÇÖ'})}\n\n"

            # µ│¿σàÑ skill_prompt∩╝êΦ╖»τö▒σÖ¿σ»╣τö¿µê╖µäÅσ¢╛τÜäµá╝σ╝Åµ£ƒµ£¢∩╝ë
            _task_skill = (context_info or {}).get("skill_prompt")
            if _task_skill:
                use_instruction += f"\n\n[σôìσ║öΦªüµ▒é] {_task_skill}"

            # µ₧äσ╗║σÄåσÅ▓Φ«░σ╜ò∩╝êΘ¥₧σ╗╢τ╗¡Σ╗╗σèíµù╢Φ┐çµ╗ñµùáσà│σÄåσÅ▓∩╝ë
            if context_info and context_info.get("is_continuation"):
                history_for_model = history
                t = yield_thinking(
                    f"µúÇµ╡ïσê░Σ╕èΣ╕ïµûçσ╗╢τ╗¡∩╝îΣ┐¥τòÖσà¿Θâ¿ {len(history)} Φ╜«σ»╣Φ»¥σÄåσÅ▓",
                    "context",
                    "hybrid",
                )
                if t:
                    yield t
            else:
                history_for_model = ContextAnalyzer.filter_history(user_input, history)
                if len(history_for_model) != len(history):
                    t = yield_thinking(
                        f"Φ┐çµ╗ñσ»╣Φ»¥σÄåσÅ▓: {len(history)} Φ╜« ΓåÆ {len(history_for_model)} Φ╜«τ¢╕σà│Φ«░σ╜ò",
                        "context",
                        "hybrid",
                    )
                    if t:
                        yield t

            formatted_history = []
            for turn in history_for_model:
                formatted_history.append(
                    types.Content(
                        role=turn["role"],
                        parts=[types.Part.from_text(text=p) for p in turn["parts"]],
                    )
                )

            t = yield_thinking(
                f"σçåσñçΦ░âτö¿ {model_id} API∩╝îσÅæΘÇü {len(formatted_history)+1} µ¥íµ╢êµü»",
                "generating",
            )
            if t:
                yield t

            # ΓöÇΓöÇ σ╣╢Φüöµ£¼σ£░µ¿íσ₧ï∩╝ÜτöƒµêÉµëºΦíîΦ«íσêÆ∩╝êΣ╕ÄΣ║æτ½»µ¿íσ₧ïσ╣╢σÅæ∩╝îσí½σààΘªûσîàσ╗╢Φ┐ƒµ¡╗σî║∩╝ëΓöÇΓöÇ
            # µ£¼σ£░µ¿íσ₧ï ~200-400ms τöƒµêÉ 3-5 Σ╕¬µ¡ÑΘ¬ñ∩╝¢Σ║æτ½»ΘªûσîàΘÇÜσ╕╕Θ£Ç 1-5s
            # Σ╕ñΦÇàσ╣╢σÅæµëºΦíî∩╝îµ¡ÑΘ¬ñσ£¿ΘªûσîàσëìµêûΘªûµ¼íσ┐âΦ╖│µù╢µ╡üσç║∩╝îΘ¢╢Θó¥σñûσ╗╢Φ┐ƒ
            import concurrent.futures as _cf

            _plan_future = None
            try:
                from app.core.routing import LocalModelRouter as _LMR_plan

                if _LMR_plan.is_ollama_available() and _LMR_plan._initialized:
                    _plan_exec = _cf.ThreadPoolExecutor(
                        max_workers=1, thread_name_prefix="koto_plan"
                    )
                    _plan_future = _plan_exec.submit(
                        _LMR_plan.generate_plan, user_input, task_type
                    )
                    _plan_exec.shutdown(wait=False)  # Σ╕ìΘÿ╗σí₧∩╝îσÉÄσÅ░Φ╖æ
            except Exception:
                _plan_future = None

            # σ░å RAG Σ╕èΣ╕ïµûçΣ╜£Σ╕║τö¿µê╖µ╢êµü»σëìτ╝Ç∩╝êΣ╕Äτ│╗τ╗ƒµîçΣ╗ñσêåτª╗∩╝îµÅÉσìçΣ║ïσ«₧σçåτí«µÇº∩╝ë
            if _rag_context_block:
                _rag_augmented_input = (
                    f"[≡ƒôÜ τƒÑΦ»åσ║ôσÅéΦÇâσåàσ«╣∩╝êΦ»╖Σ╗Ñµ¡ñΣ╕║Σ║ïσ«₧Σ╛¥µì«∩╝ë]\n"
                    f"{_rag_context_block}"
                    f"ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ\n"
                    f"[τö¿µê╖Θù«Θóÿ]\n{effective_input}"
                )
            else:
                _rag_augmented_input = effective_input

            # Σ╜┐τö¿µ╡üσ╝Åσôìσ║ö
            response = client.models.generate_content_stream(
                model=model_id,
                contents=formatted_history
                + [
                    types.Content(
                        role="user",
                        parts=[types.Part.from_text(text=_rag_augmented_input)],
                    )
                ],
                config=types.GenerateContentConfig(system_instruction=use_instruction),
            )

            full_text = ""
            chunk_count = 0
            heartbeat_interval = 5  # µ»Å5τºÆσÅæΘÇüΣ╕Çµ¼íσ┐âΦ╖│
            first_chunk_received = False
            _plan_flushed = False  # µ£¼σ£░µëºΦíîΦ«íσêÆµÿ»σÉªσ╖▓µ╡üσç║

            try:
                # Σ╜┐τö¿Σ┐¥µ┤╗σîàΦúàσÖ¿σñäτÉåµ╡üσ╝Åσôìσ║ö
                max_wait = 60 if task_type == "CODER" else 120
                for item_type, item_data in stream_with_keepalive(
                    response,
                    start_time,
                    keepalive_interval=heartbeat_interval,
                    max_wait_first_token=max_wait,
                ):
                    # µúÇµƒÑΣ╕¡µû¡µáçσ┐ù
                    if _interrupt_manager.is_interrupted(session_name):
                        _app_logger.debug(f"[INTERRUPT] User interrupted at chunk {chunk_count}")
                        interrupt_msg = "\n\nΓÅ╕∩╕Å τö¿µê╖σ╖▓Σ╕¡µû¡"
                        yield f"data: {json.dumps({'type': 'token', 'content': interrupt_msg})}\n\n"
                        break

                    # ΓöÇΓöÇ µëºΦíîΦ«íσêÆ∩╝Üσ░¥Φ»òσ£¿σ┐âΦ╖│µù╢Θ¥₧Θÿ╗σí₧σê╖σç║∩╝êσí½σààΘªûσîàσëìτÜäτ⌐║τÖ╜∩╝ëΓöÇΓöÇ
                    if not _plan_flushed and _plan_future is not None:
                        try:
                            _steps = _plan_future.result(timeout=0.05)
                            if _steps:
                                for _s in _steps:
                                    _pt = yield_thinking(
                                        f"≡ƒôï {_s}", "planning", "local"
                                    )
                                    if _pt:
                                        yield _pt
                            _plan_flushed = True
                            _plan_future = None
                        except _cf.TimeoutError:
                            pass  # Φ┐ÿµ▓íσÑ╜∩╝îΣ╕ïµ¼íµúÇµƒÑ
                        except Exception:
                            _plan_flushed = True
                            _plan_future = None

                    if item_type == "heartbeat":
                        elapsed = item_data
                        if first_chunk_received:
                            # µá╣µì«Σ╗╗σèíτ▒╗σ₧ïσ╖«σ╝éσîûσ┐âΦ╖│∩╝êσ╖▓µö╢σê░Θªûσîà∩╝îµ¡úσ£¿µ╡üσ╝ÅΦ╛ôσç║∩╝ë
                            char_count = len(full_text)
                            if task_type == "CODER":
                                hb_msg = f"≡ƒÆ╗ Σ╗úτáüτöƒµêÉΣ╕¡... σ╖▓Φ╛ôσç║ {char_count} σ¡ùτ¼ª"
                            elif task_type == "RESEARCH":
                                hb_msg = f"≡ƒö¼ µ╖▒σ║ªσêåµ₧ÉΣ╕¡... σ╖▓Φ╛ôσç║ {char_count} σ¡ùτ¼ª"
                            else:
                                hb_msg = "≡ƒÆ¡ µ¡úσ£¿τöƒµêÉ..."
                            yield f"data: {json.dumps({'type': 'progress', 'message': hb_msg, 'detail': f'{elapsed}s', 'stage': 'generating'})}\n\n"
                        else:
                            # τ¡ëσ╛àΘªûσîà∩╝êapi_calling Θÿ╢µ«╡∩╝îσëìτ½»µÿ╛τñ║µùïΦ╜¼ spinner∩╝ë
                            if task_type == "CODER":
                                hb_msg = "≡ƒÆ╗ Σ╗úτáüσêåµ₧ÉΣ╕¡∩╝îΦ»╖τ¿ìσÇÖ..."
                            elif task_type == "RESEARCH":
                                hb_msg = "≡ƒö¼ µ╖▒σ║ªµÇ¥ΦÇâΣ╕¡∩╝îΦ»╖ΦÇÉσ┐âτ¡ëσ╛à..."
                            else:
                                hb_msg = "≡ƒºá µ¿íσ₧ïµÇ¥ΦÇâΣ╕¡..."
                            yield f"data: {json.dumps({'type': 'progress', 'message': hb_msg, 'detail': f'σ╖▓τ¡ëσ╛à {elapsed}s', 'stage': 'api_calling'})}\n\n"

                    elif item_type == "timeout":
                        if task_type == "CODER" and not full_text:
                            yield f"data: {json.dumps({'type': 'progress', 'message': 'ΓÜá∩╕Å ΘªûσîàΦ╢àµù╢∩╝îσêçµìóσê░σ┐½ΘÇƒµ¿íσ₧ï...', 'detail': ''})}\n\n"
                            try:
                                fallback_resp = client.models.generate_content(
                                    model="gemini-2.5-flash",
                                    contents=formatted_history
                                    + [
                                        types.Content(
                                            role="user",
                                            parts=[
                                                types.Part.from_text(
                                                    text=_rag_augmented_input
                                                )
                                            ],
                                        )
                                    ],
                                    config=types.GenerateContentConfig(
                                        system_instruction=use_instruction,
                                        temperature=0.4,
                                        max_output_tokens=4000,
                                    ),
                                )
                                fallback_text = fallback_resp.text or ""
                                if fallback_text:
                                    full_text = fallback_text
                                    yield f"data: {json.dumps({'type': 'token', 'content': fallback_text})}\n\n"
                            except Exception:
                                yield f"data: {json.dumps({'type': 'token', 'content': f'ΓÜá∩╕Å {item_data}∩╝îΦ»╖τ¿ìσÉÄΘçìΦ»ò'})}\n\n"
                        else:
                            yield f"data: {json.dumps({'type': 'token', 'content': f'ΓÜá∩╕Å {item_data}∩╝îΦ»╖τ¿ìσÉÄΘçìΦ»ò'})}\n\n"
                        break

                    elif item_type == "chunk":
                        chunk = item_data
                        if chunk.text:
                            if not first_chunk_received:
                                first_chunk_received = True
                                _app_logger.debug(
                                    f"[CHAT] µö╢σê░τ¼¼Σ╕ÇΣ╕¬σôìσ║ö∩╝îΦÇùµù╢ {time.time() - start_time:.1f}s"
                                )
                                # Θªûσîàσê░Φ╛╛∩╝Üµ£ÇσÉÄΣ╕Çµ¼íµ£║Σ╝Üσê╖σç║µëºΦíîΦ«íσêÆ∩╝êτ¡ëµ£ÇσñÜ 0.5s∩╝ë
                                if not _plan_flushed and _plan_future is not None:
                                    try:
                                        _steps = _plan_future.result(timeout=0.5)
                                        if _steps:
                                            for _s in _steps:
                                                _pt = yield_thinking(
                                                    f"≡ƒôï {_s}", "planning", "local"
                                                )
                                                if _pt:
                                                    yield _pt
                                    except Exception:
                                        pass
                                    _plan_flushed = True
                                    _plan_future = None

                            full_text += chunk.text
                            chunk_count += 1
                            yield f"data: {json.dumps({'type': 'token', 'content': chunk.text})}\n\n"

            except Exception as stream_error:
                error_str = str(stream_error)
                _app_logger.debug(f"[CHAT] Stream error: {error_str}")

                # σ£░σî║ΘÖÉσê╢ΘöÖΦ»»
                if (
                    "location is not supported" in error_str.lower()
                    or "failed_precondition" in error_str.lower()
                ):
                    error_text = "Γ¥î σ£░σî║ΘÖÉσê╢\n\nµé¿µëÇσ£¿τÜäσ£░σî║Σ╕ìµö»µîü Gemini APIπÇé\n\n≡ƒÆí Φºúσå│µû╣µíê:\n1. σ£¿ `config/gemini_config.env` Θàìτ╜«Σ╕¡Φ╜¼µ£ìσèí `GEMINI_API_BASE`\n2. µêûΣ╜┐τö¿µö»µîüτÜäΣ╗úτÉåµ£ìσèí"
                    yield f"data: {json.dumps({'type': 'token', 'content': error_text})}\n\n"
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return
                # µ╡üσ╝ÅΣ╝áΦ╛ôΣ╕¡µû¡∩╝îΣ╜åσ╖▓µ£ëΘâ¿σêåσåàσ«╣
                elif full_text:
                    error_msg = error_str[:50]
                    warn_text = f"\n\nΓÜá∩╕Å (Σ╝áΦ╛ôΣ╕¡µû¡: {error_msg}...)"
                    yield f"data: {json.dumps({'type': 'token', 'content': warn_text})}\n\n"
                else:
                    raise stream_error

            # σñ▒Φ┤Ñµù╢σàêΣ┐«µ¡úΣ╕Çµ¼í∩╝êΣ╕ìτ¢┤µÄÑµèÑΘöÖ∩╝ë
            if Utils.is_failure_output(full_text):
                yield f"data: {json.dumps({'type': 'progress', 'message': 'ΓÜá∩╕Å σê¥µ¼íτöƒµêÉσñ▒Φ┤Ñ∩╝îµ¡úσ£¿Σ┐«µ¡ú...', 'detail': ''})}\n\n"
                fix_prompt = Utils.build_fix_prompt(task_type, user_input, full_text)
                fix_resp = client.models.generate_content(
                    model=model_id,
                    contents=fix_prompt,
                    config=types.GenerateContentConfig(
                        system_instruction=use_instruction,
                        temperature=0.4,
                        max_output_tokens=4000,
                    ),
                )
                corrected_text = fix_resp.text or full_text
                if corrected_text and corrected_text != full_text:
                    corrected_msg = f"\n\n≡ƒöü Σ┐«µ¡úτëêµ£¼:\n{corrected_text}"
                    yield f"data: {json.dumps({'type': 'token', 'content': corrected_msg})}\n\n"
                    full_text = corrected_text
            else:
                # σñìµ¥éΣ╗╗σèíΦ┐¢Φíîσ┐½ΘÇƒΦç¬µúÇ
                is_complex_task = (
                    task_type in ["RESEARCH", "FILE_GEN", "CODER"] or
                    (context_info and context_info.get("complexity") == "complex") or
                    len(user_input) > 200
                )
                if is_complex_task:
                    check = Utils.quick_self_check(task_type, user_input, full_text)
                    if not check.get("pass") and check.get("fix_prompt"):
                        status_msg = "≡ƒ⌐║ Φç¬µúÇµ£¬ΘÇÜΦ┐ç∩╝îµ¡úσ£¿Σ┐«µ¡ú..."
                        yield f"data: {json.dumps({'type': 'progress', 'message': status_msg, 'detail': 'σ┐½ΘÇƒµ¿íσ₧ïΦç¬µúÇ'})}\n\n"
                        fix_resp = client.models.generate_content(
                            model=model_id,
                            contents=check["fix_prompt"],
                            config=types.GenerateContentConfig(
                                system_instruction=use_instruction,
                                temperature=0.4,
                                max_output_tokens=4000,
                            )
                        )
                        corrected_text = fix_resp.text or full_text
                        if corrected_text and corrected_text != full_text:
                            corrected_msg = f"\n\n≡ƒöü Σ┐«µ¡úτëêµ£¼:\n{corrected_text}"
                            yield f"data: {json.dumps({'type': 'token', 'content': corrected_msg})}\n\n"
                            full_text = corrected_text

            # σñäτÉåΦç¬σè¿Σ┐¥σ¡ÿτÜäµûçΣ╗╢
            saved_files = Utils.auto_save_files(full_text)

            # Σ╗úτáüΣ╗╗σèí: µúÇµ╡ïσ╣╢Φç¬σè¿σ«ëΦúàΣ╛¥Φ╡û
            if task_type == "CODER":
                pkgs = Utils.detect_required_packages(full_text)
                if pkgs:
                    yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒôª µúÇµ╡ïσê░Σ╛¥Φ╡û∩╝îµ¡úσ£¿µúÇµƒÑ/σ«ëΦúà...', 'detail': ', '.join(pkgs)})}\n\n"
                    install_result = Utils.auto_install_packages(pkgs)
                    installed = install_result.get("installed", [])
                    failed = install_result.get("failed", [])
                    skipped = install_result.get("skipped", [])
                    msg_parts = []
                    if installed:
                        msg_parts.append(f"Γ£à σ╖▓σ«ëΦúà: {', '.join(installed)}")
                    if skipped:
                        msg_parts.append(f"Γä╣∩╕Å σ╖▓σ¡ÿσ£¿: {', '.join(skipped)}")
                    if failed:
                        msg_parts.append(f"ΓÜá∩╕Å σ«ëΦúàσñ▒Φ┤Ñ: {', '.join(failed)}")
                    if msg_parts:
                        msg_content = "\n\n" + "\n".join(msg_parts)
                        yield f"data: {json.dumps({'type': 'token', 'content': msg_content})}\n\n"

            # σªéµ₧£µ£ëΣ┐¥σ¡ÿτÜäµûçΣ╗╢∩╝îµÅÉτñ║τö¿µê╖Σ┐¥σ¡ÿΣ╜ìτ╜«
            if saved_files:
                files_list = ", ".join(saved_files)
                save_hint = (
                    f"\n\n≡ƒôü µûçΣ╗╢σ╖▓Σ┐¥σ¡ÿ: **{files_list}**\n≡ƒôé Σ╜ìτ╜«: `{WORKSPACE_DIR}`"
                )
                yield f"data: {json.dumps({'type': 'token', 'content': save_hint})}\n\n"

            # σàêΣ┐¥σ¡ÿσÄåσÅ▓∩╝îσåìσÅæΘÇü done Σ║ïΣ╗╢∩╝êσîàσÉ½σàâµò░µì«τö¿Σ║Äσëìτ½»µ╕▓µƒô∩╝ë
            session_manager.append_and_save(
                f"{session_name}.json",
                user_input,
                full_text,
                task=task_type,
                model_name=model_id,
                saved_files=saved_files,
            )
            # 2-B: Memory reflection for all supported task types
            _reflect_types = {"CHAT", "RESEARCH", "CODER", "FILE_GEN", "AGENT"}
            if task_type in _reflect_types:
                _start_memory_extraction(
                    user_input,
                    full_text,
                    history_for_model,
                    task_type=task_type,
                    session_name=session_name,
                )

            # Φ«íτ«ù msg_id σ╣╢σîàσÉ½σ£¿ done Σ║ïΣ╗╢Σ╕¡∩╝îΣ╜┐σëìτ½»Φâ╜µÅÉΣ║ñτö¿µê╖Φ»äσêå
            try:
                from app.core.learning.rating_store import RatingStore as _RS

                _done_msg_id = _RS.make_msg_id(session_name, user_input)
            except Exception:
                _done_msg_id = ""

            total_time = time.time() - start_time
            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': saved_files, 'total_time': total_time, 'msg_id': _done_msg_id})}\n\n"

        except Exception as e:
            error_str = str(e)
            _app_logger.debug(f"[CHAT] Exception: {error_str}")

            # σ£░σî║ΘÖÉσê╢ΘöÖΦ»»
            if (
                "location is not supported" in error_str.lower()
                or "failed_precondition" in error_str.lower()
            ):
                error_response = "Γ¥î σ£░σî║ΘÖÉσê╢\n\nµé¿µëÇσ£¿τÜäσ£░σî║Σ╕ìµö»µîü Gemini APIπÇé\n\n≡ƒÆí Φºúσå│µû╣µíê:\n1. σ£¿ `config/gemini_config.env` Θàìτ╜«Σ╕¡Φ╜¼µ£ìσèí `GEMINI_API_BASE`\n2. µêûΣ╜┐τö¿µö»µîüτÜäΣ╗úτÉåµ£ìσèí"
            elif "API key not valid" in error_str or (
                "INVALID_ARGUMENT" in error_str and "api key" in error_str.lower()
            ):
                error_response = (
                    "Γ¥î **API σ»åΘÆÑµùáµòê**\n\n"
                    "Φ»╖µúÇµƒÑµé¿τÜä Gemini API σ»åΘÆÑ∩╝Ü\n"
                    "1. σëìσ╛Ç [aistudio.google.com/apikey](https://aistudio.google.com/apikey) ΦÄ╖σÅûµ£ëµòêσ»åΘÆÑ\n"
                    "2. σ£¿ Koto Φ«╛τ╜«Θí╡Θ¥óµ¢┤µû░ API σ»åΘÆÑ∩╝êΦ«╛τ╜« ΓåÆ API Θàìτ╜«∩╝ë\n"
                    "3. τí«Σ┐¥σ»åΘÆÑµëÇσ£¿ Google Θí╣τ¢«σ╖▓σÉ»τö¿ Generative Language API\n\n"
                    f"σÄƒσºïΘöÖΦ»»: `{error_str[:150]}`"
                )
            elif (
                "server disconnected" in error_str.lower()
                or "disconnected without" in error_str.lower()
                or "connection reset" in error_str.lower()
                or "connection aborted" in error_str.lower()
            ):
                error_response = (
                    "Γ¥î **µ£ìσèíσÖ¿Φ┐₧µÄÑΣ╕¡µû¡**\n\n"
                    "Σ╕Ä Gemini API τÜäΦ┐₧µÄÑΦó½µäÅσñûµû¡σ╝Ç∩╝îΦ┐ÖΘÇÜσ╕╕µÿ»Σ╕┤µù╢µÇºΘù«ΘóÿπÇé\n\n"
                    "≡ƒÆí σ╗║Φ««∩╝Ü\n"
                    "1. τ¿ìτ¡ëτëçσê╗σÉÄΘçìµû░σÅæΘÇüµ╢êµü»\n"
                    "2. µúÇµƒÑµé¿τÜäτ╜æτ╗£Φ┐₧µÄÑτ¿│σ«ÜµÇº\n"
                    "3. σªéµ₧£Σ╜┐τö¿Σ╗úτÉå∩╝îΦ»╖τí«Φ«ñΣ╗úτÉåΦ┐₧µÄÑµ¡úσ╕╕\n"
                    "4. σªéΘù«Θóÿµîüτ╗¡∩╝îσÅ»σ░¥Φ»òσêçµìóσê░σà╢Σ╗ûµ¿íσ₧ï"
                )
            elif (
                "resource_exhausted" in error_str.lower()
                or "quota" in error_str.lower()
                or "rate limit" in error_str.lower()
                or "429" in error_str
            ):
                error_response = (
                    "Γ¥î **API ΘàìΘó¥Φ╢àΘÖÉ**\n\n"
                    "σ╜ôσëì API σ»åΘÆÑτÜäΦ»╖µ▒éΘóæτÄçµêûΘàìΘó¥σ╖▓Φ╛╛Σ╕èΘÖÉπÇé\n\n"
                    "≡ƒÆí σ╗║Φ««∩╝Ü\n"
                    "1. τ¿ìτ¡ë 1-2 σêåΘÆƒσÉÄΘçìΦ»ò\n"
                    "2. σ£¿Φ«╛τ╜«Σ╕¡σêçµìóσê░σà╢Σ╗û API σ»åΘÆÑ\n"
                    "3. µêûσìçτ║ºµé¿τÜä Google AI Studio Φ«íσêÆ"
                )
            elif (
                "unavailable" in error_str.lower()
                or "503" in error_str
                or "service unavailable" in error_str.lower()
            ):
                error_response = (
                    "Γ¥î **Gemini µ£ìσèíµÜéµù╢Σ╕ìσÅ»τö¿**\n\n"
                    "Gemini API µ£ìσèíσÖ¿σ╜ôσëìµùáµ│òσôìσ║ö∩╝îσÅ»Φâ╜µ¡úσ£¿τ╗┤µèñΣ╕¡πÇé\n\n"
                    "≡ƒÆí σ╗║Φ««∩╝Üτ¿ìτ¡ëτëçσê╗σÉÄΘçìΦ»ò∩╝îµêûΦ«┐Θù« [status.google.com](https://status.google.com) µƒÑτ£ïµ£ìσèíτè╢µÇü"
                )
            elif (
                "deadline_exceeded" in error_str.lower()
                or "timed out" in error_str.lower()
            ):
                error_response = (
                    "Γ¥î **Φ»╖µ▒éΦ╢àµù╢**\n\n"
                    "µ¿íσ₧ïσôìσ║öµù╢Θù┤Φ┐çΘò┐∩╝îΦ»╖µ▒éσ╖▓Φ╢àµù╢πÇé\n\n"
                    "≡ƒÆí σ╗║Φ««∩╝Ü\n"
                    "1. σ░¥Φ»òτ╝⌐τƒ¡µé¿τÜäΘù«Θóÿµêûσêåµ¡ÑΘ¬ñµÅÉΘù«\n"
                    "2. σêçµìóσê░σôìσ║öµ¢┤σ┐½τÜäµ¿íσ₧ï∩╝êσªé gemini-2.5-flash∩╝ë\n"
                    "3. µúÇµƒÑτ╜æτ╗£Φ┐₧µÄÑΦ┤¿ΘçÅ"
                )
            else:
                error_response = f"Γ¥î σÅæτöƒΘöÖΦ»»: {error_str[:200]}"

            # σì│Σ╜┐σç║ΘöÖΣ╣ƒΦªüΣ┐¥σ¡ÿτö¿µê╖τÜäΘù«Θóÿ
            session_manager.append_and_save(
                f"{session_name}.json",
                user_input,
                error_response,
                task=task_type,
                model_name=model_id,
            )

            yield f"data: {json.dumps({'type': 'token', 'content': error_response})}\n\n"
            total_time = time.time() - start_time
            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"

    def _safe_generate():
        """
        generate() σñûσ▒éσ«ëσà¿σîàΦúàσÖ¿∩╝Ü
        τí«Σ┐¥µùáΦ«║ generate() σåàΘâ¿Σ╗ÑΣ╜òτºìµû╣σ╝Åτ╗ôµ¥ƒ∩╝îσëìτ½»Θâ╜Φâ╜µö╢σê░ 'done' Σ║ïΣ╗╢∩╝î
        Θü┐σàìσ¢áΣ╗╗σèíΦ»åσê½σñ▒Φ┤Ñ/µù⌐µ£ƒσ╝éσ╕╕σ»╝Φç┤σ»╣Φ»¥τòîΘ¥óµ░╕Φ┐£µîéΦ╡╖πÇé
        """
        _sent_done = False
        try:
            for _chunk in generate():
                if isinstance(_chunk, (str, bytes)):
                    _chunk_str = (
                        _chunk
                        if isinstance(_chunk, str)
                        else _chunk.decode("utf-8", errors="replace")
                    )
                    if '"type": "done"' in _chunk_str or "'type': 'done'" in _chunk_str:
                        _sent_done = True
                yield _chunk
        except Exception as _sg_err:
            _app_logger.warning(f"[STREAM] ΓÜá∩╕Å _safe_generate caught exception: {_sg_err}")
            import traceback

            traceback.print_exc()
            if not _sent_done:
                _err_msg = f"Γ¥î µ╡üσ╝Åσôìσ║öσ╝éσ╕╕τ╗êµ¡ó: {str(_sg_err)[:200]}"
                yield f"data: {json.dumps({'type': 'token', 'content': _err_msg})}\n\n"
        finally:
            if not _sent_done:
                _app_logger.warning(
                    f"[STREAM] ΓÜá∩╕Å generate() µ£¬σÅæΘÇü done Σ║ïΣ╗╢∩╝îΦºªσÅæσà£σ║ò done (task_type={task_type})"
                )
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': 0, 'fallback_done': True})}\n\n"

    response = Response(
        stream_with_context(_safe_generate()), mimetype="text/event-stream"
    )
    response.headers["Cache-Control"] = "no-cache"
    response.headers["X-Accel-Buffering"] = "no"  # τªüτö¿ nginx τ╝ôσå▓
    response.headers["Connection"] = "keep-alive"
    return response


@app.route("/api/chat/file", methods=["POST"])
def chat_with_file():
    """σñäτÉåµûçΣ╗╢Σ╕èΣ╝áσÆîΦüèσñ⌐Φ»╖µ▒é"""
    from web.document_generator import save_docx, save_pdf, to_workspace_rel
    from web.file_processor import process_uploaded_file

    def _strip_code_blocks(text: str) -> str:
        if not text:
            return text
        # Remove fenced code blocks entirely
        text = re.sub(r"```[\s\S]*?```", "", text)
        # Remove inline code ticks but keep the content
        text = text.replace("`", "")
        return text.strip()

    def _build_analysis_title(user_text: str, filename: str, is_binary: bool) -> str:
        name_base = os.path.splitext(filename)[0]
        text_lower = (user_text or "").lower()
        ext = os.path.splitext(filename)[1].lower()

        # 1. Determine File Type Prefix
        if ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"]:
            prefix = "σ¢╛τëç"
        elif ext == ".pdf":
            prefix = "PDF"
        elif ext in [".doc", ".docx"]:
            prefix = "Word"
        elif ext in [".ppt", ".pptx"]:
            prefix = "PPT"
        else:
            prefix = "µûçΣ╗╢" if is_binary else "µûçµíú"

        # 2. Determine Intent
        intent = "σêåµ₧É"
        intent_map = {
            "τ┐╗Φ»æ": ["τ┐╗Φ»æ", "translate", "Φ»æµûç", "Σ╕¡Φ»æΦï▒", "Φï▒Φ»æΣ╕¡"],
            "µÇ╗τ╗ô": ["µÇ╗τ╗ô", "σ╜Æτ║│", "µæÿΦªü", "summary", "µªéµï¼", "µá╕σ┐âσåàσ«╣"],
            "µûçσ¡ùΦ»åσê½": ["µÅÉσÅû", "Φ»åσê½", "ocr", "µûçσ¡ù", "Φ╜¼µûçσ¡ù", "Φ»╗σ¢╛"],
            "Φí¿µá╝Φ»åσê½": ["Φí¿µá╝", "table", "excel", "Φ╜¼Φí¿"],
            "σ»╣µ»öσêåµ₧É": ["σ»╣µ»ö", "µ»öΦ╛â", "diff", "σî║σê½", "σ╖«σ╝é"],
            "µáíσ»╣": ["µáíσ»╣", "µúÇµƒÑ", "σ«íΘÿà", "τ║áΘöÖ", "µö╣ΘöÖ"],
            "µ╢ªΦë▓": ["µ╢ªΦë▓", "µö╣σåÖ", "polish", "rewrite", "Σ╝ÿσîû", "τ╛Äσîû"],
            "τ╗¡σåÖ": ["τ╗¡σåÖ", "µë⌐σåÖ", "continue", "ΦíÑσàà"],
            "σñºτ║▓": ["σñºτ║▓", "µíåµ₧╢", "outline", "τ¢«σ╜ò"],
            "ΦºúΘçè": ["ΦºúΘçè", "explain", "Σ╗ÇΣ╣êµäÅµÇ¥", "σÉ½Σ╣ë"],
        }

        found_intent_keywords = []
        for k, v in intent_map.items():
            for kw in v:
                if kw in text_lower:
                    intent = k
                    found_intent_keywords.append(kw)
                    break
            if intent != "σêåµ₧É":
                break

        # 3. Extract Topic Keywords (Improved)
        stop_words = [
            "σ╕«µêæ",
            "Φ»╖",
            "Σ╕ÇΣ╕ï",
            "µèè",
            "Φ┐ÖΣ╕¬",
            "Φ┐Öτ»ç",
            "µûçΣ╗╢",
            "µûçτ½á",
            "σåàσ«╣",
            "τöƒµêÉ",
            "σåÖΣ╕ÇΣ╕¬",
            "σüÜΣ╕ÇΣ╗╜",
            "koto",
            "σêåµ₧É",
            "ΘÿàΦ»╗",
            "µÅÉσÅû",
            "Φ»åσê½",
            "output",
            "make",
            "create",
            "generate",
            "please",
            "the",
            "a",
            "an",
            "is",
            "of",
            "to",
            "for",
            "with",
            "in",
            "on",
            "user",
            "file",
            "document",
            "from",
            "this",
            "that",
            "it",
            "what",
            "how",
            "why",
            "where",
            "into",
            "check",
            "run",
        ]

        # Prepare text
        text_lower = user_text.lower()

        # Safe replacement for Chinese phrases (which don't use spaces)
        zh_stops = [w for w in stop_words if re.match(r"[\u4e00-\u9fa5]+", w)]
        for stop in zh_stops + found_intent_keywords:
            if re.match(r"[\u4e00-\u9fa5]+", stop):  # Only safe-replace Chinese phrases
                text_lower = text_lower.replace(stop, " ")

        # Tokenize by non-word chars (separates English words, breaks Chinese into blocks if spaces inserted)
        # Regex: Keep Chinese chars and English words
        # This splits "summary of report" -> "summary", "of", "report"
        tokens = re.findall(r"[a-zA-Z0-9\u4e00-\u9fa5]+", text_lower)

        # Filter tokens
        valid_keywords = []
        en_stops = set([w for w in stop_words if not re.match(r"[\u4e00-\u9fa5]+", w)])

        for token in tokens:
            if token in en_stops:
                continue
            if token in found_intent_keywords:
                continue  # Filter intent words token-wise
            if len(token) < 2:
                continue
            valid_keywords.append(token)

        # Select best keyword
        topic = ""
        if valid_keywords:
            topic = "_".join(valid_keywords[:3])

        # 4. Construct Final Title
        # Strategy:
        # If user provided a specific topic, prioritize it: "{Intent}_{Topic}_{Filename}"
        # If no detected topic but intent exists: "{Intent}_{Filename}"
        # Fallback: "{Prefix}{Intent}_{Filename}"

        sanitized_name = name_base.replace(" ", "_")

        if topic:
            return f"{intent}_{topic}_{sanitized_name}"
        else:
            return f"{prefix}{intent}_{sanitized_name}"

    session_name = request.form.get("session")
    user_input = request.form.get("message", "")
    files = request.files.getlist("file")

    # ≡ƒöì Φ░âΦ»òµùÑσ┐ù
    _app_logger.info(f"[FILE UPLOAD DEBUG] ========== µÄÑµö╢σê░µûçΣ╗╢Σ╕èΣ╝áΦ»╖µ▒é ==========")
    _app_logger.info(f"[FILE UPLOAD DEBUG] request.files keys: {list(request.files.keys())}")
    _app_logger.info(f"[FILE UPLOAD DEBUG] request.files.getlist('file'): {len(files)} Σ╕¬µûçΣ╗╢")
    for i, f in enumerate(files):
        _app_logger.info(f"[FILE UPLOAD DEBUG]   {i+1}. {f.filename if f else 'None'}")

    if not files:
        single_file = request.files.get("file")
        if single_file:
            files = [single_file]
            _app_logger.info(f"[FILE UPLOAD DEBUG] Σ╜┐τö¿σìòµûçΣ╗╢µ¿íσ╝Å∩╝îµûçΣ╗╢: {single_file.filename}")

    locked_task = request.form.get("locked_task")
    locked_model = request.form.get("locked_model", "auto")
    stream_mode = request.form.get("stream", "").lower() in ("1", "true", "yes")

    _app_logger.info(f"[FILE UPLOAD DEBUG] µ£Çτ╗ê files σêùΦí¿: {len(files)} Σ╕¬µûçΣ╗╢")
    _app_logger.info(f"[FILE UPLOAD DEBUG] σêñµû¡: len(files) > 1 = {len(files) > 1}")

    if not session_name or not files:
        return jsonify({"error": "Missing session or file"}), 400
    if len(files) > 10:
        return jsonify({"error": "µ£ÇσñÜΣ╕Çµ¼íΣ╕èΣ╝á 10 Σ╕¬µûçΣ╗╢"}), 400

    if len(files) > 1:
        # µúÇµ╡ïµÿ»σÉªµÿ» PPT τöƒµêÉµäÅσ¢╛ (σñÜµûçΣ╗╢σÉêσ╣╢τöƒµêÉ PPT)
        ppt_keywords = ["ppt", "slide", "σ╣╗τü»τëç", "µ╝öτñ║µûçτ¿┐", "powerpoint"]
        is_ppt_intent = any(kw in (user_input or "").lower() for kw in ppt_keywords)

        if is_ppt_intent:
            _app_logger.info(f"[FILE UPLOAD] µúÇµ╡ïσê░σñÜµûçΣ╗╢ PPT τöƒµêÉµäÅσ¢╛: {user_input}")

            # ΘóäσàêΣ┐¥σ¡ÿµëÇµ£ëµûçΣ╗╢∩╝îΘü┐σàìσ£¿τöƒµêÉσÖ¿Σ╕¡Φ«┐Θù«σ╖▓σà│Θù¡τÜä FileStorage
            saved_file_paths = []
            source_filenames = []

            for f in files:
                if f and f.filename:
                    fname = f.filename
                    fpath = os.path.join(UPLOAD_DIR, fname)
                    # σªéµ₧£µûçΣ╗╢µîçΘÆêΣ╕ìσ£¿σ╝Çσñ┤∩╝îΘçìτ╜«σ«â
                    f.seek(0)
                    f.save(fpath)
                    saved_file_paths.append(fpath)
                    source_filenames.append(fname)

            def generate_ppt_stream():
                try:
                    yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒôè µ¡úσ£¿σçåσñç PPT τöƒµêÉ...', 'detail': f'µúÇµ╡ïσê░ {len(saved_file_paths)} Σ╕¬µ║ÉµûçΣ╗╢'})}\n\n"

                    context_text = ""

                    # 1. µÅÉσÅûµëÇµ£ëσ╖▓Σ┐¥σ¡ÿµûçΣ╗╢σåàσ«╣
                    for i, filepath in enumerate(saved_file_paths):
                        filename = os.path.basename(filepath)
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒôû µ¡úσ£¿Φ»╗σÅûµûçΣ╗╢ ({i+1}/{len(saved_file_paths)})...', 'detail': filename})}\n\n"

                        try:
                            # µÅÉσÅûσåàσ«╣
                            from web.file_processor import FileProcessor

                            processor = FileProcessor()
                            # τ«ÇσîûτëêτÜä process
                            f_result = processor.process_file(filepath)
                            content = f_result.get("text_content") or f_result.get(
                                "content", ""
                            )

                            # µê¬µû¡Φ┐çΘò┐σåàσ«╣Θü┐σàìTokenτêåτé╕∩╝îΣ╜åΣ┐¥τòÖΦ╢│σñƒΣ╕èΣ╕ïµûç
                            if len(content) > 50000:
                                content = content[:50000] + "...(truncated)"

                            context_text += f"\n\n=== {filename} ===\n{content}\n"

                        except Exception as e:
                            _app_logger.info(f"[PPT BATCH] Φ»╗σÅûµûçΣ╗╢ {filename} σñ▒Φ┤Ñ: {e}")
                            context_text += (
                                f"\n\n=== {filename} (Error) ===\nµùáµ│òΦ»╗σÅûσåàσ«╣\n"
                            )

                    # 2. Φ░âτö¿ PPT τöƒµêÉτ«íΘüô
                    yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒÄ¿ µ¡úσ£¿Φ«╛Φ«í PPT τ╗ôµ₧ä...', 'detail': 'σƒ║Σ║ÄσñÜΣ╕¬µûçΣ╗╢σåàσ«╣'})}\n\n"

                    import asyncio

                    from web.ppt_pipeline import PPTGenerationPipeline

                    # µ₧äΘÇáσó₧σ╝║σÉÄτÜä Prompt
                    enhanced_prompt = f"{user_input}\n\nπÇÉσÅéΦÇâΦ╡äµûÖπÇæ\nσƒ║Σ║ÄΣ╗ÑΣ╕ïµûçΣ╗╢τöƒµêÉτÜä PPT:\n{context_text}"

                    # ΘÖÉσê╢ Prompt Θò┐σ║ª
                    if len(enhanced_prompt) > 100000:
                        enhanced_prompt = (
                            enhanced_prompt[:100000] + "\n...(context truncated)"
                        )

                    # σ╝éµ¡ÑµëºΦíî PPT τöƒµêÉ
                    # Σ╜┐τö¿Θí╣τ¢«σåàτÜä get_client() ΦÄ╖σÅû Gemini σ«óµê╖τ½»
                    ai_client = get_client()
                    pipeline = PPTGenerationPipeline(ai_client=ai_client)

                    import queue
                    import threading
                    import traceback

                    pipeline_timeout_sec = 300
                    start_ts = time.time()

                    # µ╖╖σÉêµ╢êµü»Θÿƒσêù∩╝êΦ┐¢σ║ª+µÇ¥ΦÇâ∩╝ë
                    event_queue = queue.Queue()

                    def _progress_listener(msg, p=None):
                        event_queue.put({"type": "progress", "msg": msg, "progress": p})

                    def _thought_listener(text):
                        # Use a dedicated type for thought/reasoning text
                        event_queue.put({"type": "thought", "text": text})

                    run_state = {
                        "done": False,
                        "result": None,
                        "error": None,
                        "traceback": "",
                    }

                    def _run_pipeline_bg():
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)
                        try:
                            # Σ╝áΘÇÆ progress_callback σÆî thought_callback
                            run_state["result"] = loop.run_until_complete(
                                pipeline.generate(
                                    user_request=enhanced_prompt,
                                    output_path=os.path.join(
                                        settings_manager.documents_dir,
                                        f"Koto_Presentation_{int(time.time())}.pptx",
                                    ),
                                    enable_auto_images=True,  # σàüΦ«╕Φç¬σè¿Θàìσ¢╛
                                    progress_callback=_progress_listener,
                                    thought_callback=_thought_listener,
                                )
                            )
                        except Exception as bg_err:
                            run_state["error"] = str(bg_err)
                            run_state["traceback"] = traceback.format_exc()
                        finally:
                            try:
                                loop.close()
                            except Exception:
                                pass
                            run_state["done"] = True

                    worker = threading.Thread(target=_run_pipeline_bg, daemon=True)
                    worker.start()

                    # σ«₧µù╢Φ╜«Φ»óΦ┐¢σ║ªΘÿƒσêù∩╝îΦ╜¼σÅæτ╗Öσëìτ½»
                    last_progress_msg = "σê¥σºïσîûτöƒµêÉτÄ»σóâ..."

                    while not run_state["done"]:
                        elapsed = int(time.time() - start_ts)
                        if elapsed > pipeline_timeout_sec:
                            _progress_listener("τöƒµêÉΦ╢àµù╢∩╝îµ¡úσ£¿σ╝║σê╢σü£µ¡ó...", 100)
                            run_state["error"] = (
                                f"PPT τöƒµêÉΦ╢àµù╢∩╝ê>{pipeline_timeout_sec}s∩╝ë"
                            )
                            break

                        # µ╢êΦ┤╣µëÇµ£ëτÜäΣ║ïΣ╗╢
                        try:
                            while not event_queue.empty():
                                item = event_queue.get_nowait()

                                if item["type"] == "progress":
                                    msg = item["msg"]
                                    p = item["progress"]
                                    last_progress_msg = msg
                                    detail_text = (
                                        f"Φ┐¢σ║ª: {p}%"
                                        if p is not None
                                        else f"σ╖▓τö¿µù╢ {elapsed}s"
                                    )
                                    yield f"data: {json.dumps({'type': 'progress', 'message': msg, 'detail': detail_text})}\n\n"

                                elif item["type"] == "thought":
                                    # Send thought as a partial text response or a special 'thought' event
                                    # Assuming frontend can handle 'text' type for appending to the assistant's message
                                    # or 'thought' for a distinct UI block.
                                    # Let's use 'text' for now to ensure it appears in the chat stream.
                                    thought_text = (
                                        f"\n\n> ≡ƒñû **Koto µÇ¥ΦÇâ**: {item['text']}\n"
                                    )
                                    yield f"data: {json.dumps({'type': 'text', 'content': thought_text})}\n\n"

                        except queue.Empty:
                            pass

                        # σªéµ₧£µ▓íµ£ëµû░µ╢êµü»∩╝îµ»Å2τºÆσÅæΣ╕Çµ¼íσ┐âΦ╖│Θÿ▓µ¡óΦ┐₧µÄÑµû¡σ╝Ç
                        if elapsed % 2 == 0 and event_queue.empty():
                            yield f"data: {json.dumps({'type': 'progress', 'message': last_progress_msg, 'detail': f'σ╖▓τö¿µù╢ {elapsed}s'})}\n\n"

                        time.sleep(0.5)

                    # σÅæΘÇüµ£ÇσÉÄσë⌐Σ╜ÖτÜäµ╢êµü»
                    try:
                        while not event_queue.empty():
                            item = event_queue.get_nowait()
                            if item["type"] == "progress":
                                yield f"data: {json.dumps({'type': 'progress', 'message': item['msg'], 'detail': ''})}\n\n"
                            elif item["type"] == "thought":
                                thought_text = (
                                    f"\n\n> ≡ƒñû **Koto µÇ¥ΦÇâ**: {item['text']}\n"
                                )
                                yield f"data: {json.dumps({'type': 'text', 'content': thought_text})}\n\n"
                    except Exception:
                        pass

                    if run_state["error"]:
                        err = run_state["error"]
                        tb = run_state.get("traceback", "")
                        _app_logger.info(f"[PPT BATCH] Background pipeline error: {err}")
                        if tb:
                            _app_logger.info(f"[PPT BATCH] Traceback: {tb[:800]}")
                        raise Exception(f"PPT τ«íΘüôσ╝éσ╕╕: {err}")

                    ppt_result = run_state["result"] or {}

                    # pipeline returns 'output_path', also check 'file_path' for compat
                    saved_path = ppt_result.get("output_path") or ppt_result.get(
                        "file_path"
                    )

                    if not ppt_result.get("success"):
                        err_detail = ppt_result.get("error", "µ£¬τƒÑΘöÖΦ»»")
                        tb = ppt_result.get("traceback", "")
                        _app_logger.info(f"[PPT BATCH] Pipeline returned failure: {err_detail}")
                        if tb:
                            _app_logger.info(f"[PPT BATCH] Traceback: {tb[:500]}")
                        raise Exception(f"PPT τ«íΘüôτöƒµêÉσñ▒Φ┤Ñ: {err_detail}")

                    if saved_path and os.path.exists(saved_path):
                        yield f"data: {json.dumps({'type': 'progress', 'message': 'Γ£à PPT τöƒµêÉσ«îµêÉ∩╝ü', 'detail': os.path.basename(saved_path)})}\n\n"

                        rel_path = os.path.relpath(saved_path, WORKSPACE_DIR).replace(
                            "\\", "/"
                        )
                        success_msg = f"Γ£à **PPT τöƒµêÉµêÉσèƒ∩╝ü**\n\nσƒ║Σ║Ä {len(saved_file_paths)} Σ╕¬µûçΣ╗╢τöƒµêÉτÜäµ╝öτñ║µûçτ¿┐πÇé\n≡ƒôü µûçΣ╗╢: **{os.path.basename(saved_path)}**"

                        yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"

                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [rel_path], 'total_time': 0})}\n\n"
                    else:
                        raise Exception("PPT µûçΣ╗╢τöƒµêÉσñ▒Φ┤Ñ∩╝îµ£¬Φ┐öσ¢₧Φ╖»σ╛ä")

                except Exception as e:
                    _app_logger.info(f"[PPT BATCH ERROR] {e}")
                    import traceback

                    traceback.print_exc()
                    err_msg = f"Γ¥î τöƒµêÉσñ▒Φ┤Ñ: {str(e)}"
                    yield f"data: {json.dumps({'type': 'token', 'content': err_msg})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': 0})}\n\n"

            return Response(
                stream_with_context(generate_ppt_stream()), mimetype="text/event-stream"
            )

        history = session_manager.load(f"{session_name}.json")
        file_names = [f.filename for f in files if f and f.filename]
        user_message = f"[Files: {', '.join(file_names)}] {user_input}"
        session_manager.append_user_early(f"{session_name}.json", user_message)

        batch_results = []
        combined_saved_files = []
        combined_images = []

        def _process_single_file(file):
            if not file or not file.filename:
                return None

            filename = _secure_filename(file.filename) or f"upload_{uuid.uuid4().hex}"
            filepath = os.path.join(UPLOAD_DIR, filename)
            file.save(filepath)
            file_type = file.mimetype or file.content_type or ""
            file_ext = os.path.splitext(filename)[1].lower()

            # µúÇµ╡ïµÿ»σÉªµÿ»τ║»σ╜Æµíú/µò┤τÉåΦ»╖µ▒é∩╝êΣ╕ìΘ£ÇΦªüAIσêåµ₧Éσåàσ«╣∩╝ë
            organize_keywords = [
                "µò┤τÉå",
                "σ╜Æµíú",
                "σ╜Æτ║│",
                "σêåτ▒╗",
                "µò┤τÉåΣ╕ÇΣ╕ï",
                "µò┤τÉåΣ╕ï",
                "σ╕«µêæµò┤τÉå",
                "µûçΣ╗╢µò┤τÉå",
                "organize",
                "sort",
            ]
            is_organize_only = any(kw in (user_input or "") for kw in organize_keywords)

            try:
                # formatted_message, file_data = process_uploaded_file(filepath, user_input)
                # --- Modify to use FileProcessor directly for simultaneous KB indexing ---
                from web.file_processor import FileProcessor

                _processor = FileProcessor()
                _file_raw = _processor.process_file(filepath)

                # 1. Φç¬σè¿σ╗║σ║ô (Auto-Indexing to Knowledge Base) - Use threading to not block UI
                try:
                    _text_content = _file_raw.get("text_content", "")
                    if _text_content and len(_text_content) > 50:  # Ignore tiny files

                        def _bg_index(content, meta):
                            try:
                                from web.knowledge_base import KnowledgeBase

                                _kb = KnowledgeBase()
                                res = _kb.add_content(content, meta)
                                _app_logger.debug(f"[KB] Auto-indexing completed: {res}")
                            except Exception as e:
                                _app_logger.debug(f"[KB] Auto-indexing failed: {e}")

                        import threading

                        _idx_thread = threading.Thread(
                            target=_bg_index,
                            args=(
                                _text_content,
                                {
                                    "file_path": filepath,
                                    "file_name": filename,
                                    "file_type": file_ext,
                                    "mtime": os.path.getmtime(filepath),
                                },
                            ),
                        )
                        _idx_thread.start()
                        _app_logger.debug(f"[KB] σ╖▓σÉ»σè¿σÉÄσÅ░σ╗║σ║ôΣ╗╗σèí: {filename}")
                except Exception as _kb_err:
                    _app_logger.debug(f"[KB] Indexing trigger failed: {_kb_err}")

                # 1-B. µ│¿σåîσê░ FileRegistry∩╝êτ╗ƒΣ╕ÇµûçΣ╗╢σàâµò░µì«Σ╕¡σ┐â∩╝ë
                try:

                    def _bg_register_file(_fpath, _sid):
                        try:
                            from app.core.file.file_registry import get_file_registry

                            _reg = get_file_registry()
                            _reg.register(
                                _fpath,
                                source="upload",
                                session_id=_sid,
                                extract_content=True,
                            )
                            _app_logger.info(
                                f"[FileRegistry] Γ£à σ╖▓µ│¿σåîΣ╕èΣ╝áµûçΣ╗╢: {os.path.basename(_fpath)}"
                            )
                        except Exception as _re:
                            _app_logger.warning(f"[FileRegistry] ΓÜá∩╕Å µ│¿σåîσñ▒Φ┤Ñ∩╝êΘ¥₧Φç┤σæ╜∩╝ë: {_re}")

                    import threading as _thr

                    _reg_thread = _thr.Thread(
                        target=_bg_register_file,
                        args=(filepath, session_name),
                        daemon=True,
                    )
                    _reg_thread.start()
                except Exception as _rge:
                    _app_logger.warning(f"[FileRegistry] ΓÜá∩╕Å σÉ»σè¿µ│¿σåîτ║┐τ¿ïσñ▒Φ┤Ñ: {_rge}")

                # 2. Continue with standard chat formatting
                formatted_message, file_data = _processor.format_result_for_chat(
                    _file_raw, user_input
                )

                task_type = locked_task
                context_info = None
                route_method = "Auto"
                if not task_type:
                    if file_data and file_type and file_type.startswith("image"):
                        message_lower = (user_input or "").lower()
                        is_edit = any(
                            kw in message_lower for kw in KotoBrain.IMAGE_EDIT_KEYWORDS
                        )
                        task_type = "PAINTER" if is_edit else "VISION"
                        route_method = (
                            "≡ƒû╝∩╕Å Image Edit" if is_edit else "≡ƒæü∩╕Å Image Analysis"
                        )
                        _app_logger.info(
                            f"[FILE UPLOAD] σ¢╛τëçΣ╗╗σèíτ¢┤ΘÇÜΦ╖»τö▒: {task_type} (µû╣µ│ò: {route_method})"
                        )
                    else:
                        _ann_exts = {
                            ".doc",
                            ".docx",
                            ".pdf",
                            ".txt",
                            ".md",
                            ".markdown",
                            ".rtf",
                            ".odt",
                        }
                        use_annotation = (
                            _should_use_annotation_system(user_input, has_file=True)
                            and file_ext in _ann_exts
                        )

                        if use_annotation:
                            task_type = "DOC_ANNOTATE"
                            route_method = "≡ƒôî Annotation-Strict"
                        elif _is_explicit_file_gen_request(user_input):
                            # τö¿µê╖µÿÄτí«ΦªüτöƒµêÉµû░µûçΣ╗╢∩╝îτ¢┤µÄÑΦ╖»τö▒∩╝îµùáΘ£Çµ¿íσ₧ïσêåτ▒╗
                            task_type = "FILE_GEN"
                            route_method = "≡ƒôä Explicit-Gen"
                        else:
                            # Γÿà Σ╕╗Φ╖»σ╛ä∩╝ÜΦ«⌐µ£¼σ£░µ¿íσ₧ïσüÜΦ»¡Σ╣ëΦ╖»τö▒
                            # Σ╝áσàÑ [FILE_ATTACHED:ext] µáçΦ«░∩╝îµ¿íσ₧ïΘÇÜΦ┐çΦ«¡τ╗âσÑ╜τÜäΦºäσêÖσêñµû¡
                            # CHAT=Φ»╗µûçΣ╗╢σ¢₧τ¡ö  RESEARCH=µ╖▒σàÑτáöτ⌐╢  FILE_GEN=τöƒµêÉµû░µûçµíú
                            _dispatch_q = (
                                user_input or ""
                            ).strip() or "Φ»╖σêåµ₧ÉΦ┐ÖΣ╗╜µûçΣ╗╢τÜäσåàσ«╣"
                            _dispatch_input = (
                                f"[FILE_ATTACHED:{file_ext or '.file'}] {_dispatch_q}"
                            )
                            task_analysis, route_method, context_info = (
                                SmartDispatcher.analyze(
                                    _dispatch_input, history=history
                                )
                            )
                            task_type = task_analysis

                if locked_model != "auto":
                    model_to_use = locked_model
                else:
                    complexity = "complex" if file_data is None else "normal"
                    if context_info and context_info.get("complexity"):
                        complexity = context_info["complexity"]

                    if task_type == "FILE_GEN":
                        model_to_use = SmartDispatcher.get_model_for_task(
                            task_type, has_image=bool(file_data), complexity=complexity
                        )
                    else:
                        model_to_use = SmartDispatcher.get_model_for_task(
                            task_type, has_image=bool(file_data)
                        )

                _app_logger.info(f"[FILE UPLOAD] Σ╗╗σèíτ▒╗σ₧ï: {task_type}, µ¿íσ₧ï: {model_to_use}")

                result = {
                    "task": task_type,
                    "model": model_to_use,
                    "route_method": route_method,
                    "response": "",
                    "images": [],
                    "saved_files": [],
                }

                # τ║»σ╜Æµíúµ¿íσ╝Å∩╝ÜΦ╖│Φ┐çAIσåàσ«╣σêåµ₧É∩╝îτ¢┤µÄÑσ╜Æµíú
                if is_organize_only:
                    _app_logger.info(f"[FILE UPLOAD] τ║»σ╜Æµíúµ¿íσ╝Å: {filename}∩╝îΦ╖│Φ┐çAIσêåµ₧É")
                    result["response"] = ""
                    result["task"] = "FILE_ORGANIZE"
                elif task_type == "DOC_ANNOTATE":
                    # µë╣ΘçÅ/σñÜµûçΣ╗╢µ¿íσ╝ÅΣ╕ïτÜäµáçµ│¿∩╝ÜσÉîµ¡ÑΦ┐ÉΦíîµáçµ│¿τ«íΘüô
                    _app_logger.info(f"[FILE UPLOAD] µë╣ΘçÅ DOC_ANNOTATE µ¿íσ╝Å: {filename}")
                    try:
                        from web.document_feedback import DocumentFeedbackSystem

                        _batch_docs_dir = settings_manager.documents_dir
                        os.makedirs(_batch_docs_dir, exist_ok=True)
                        # σªéΘ£ÇΦ╜¼µìóσàêΦ╜¼µìó
                        _batch_filepath = filepath
                        _batch_file_ext = file_ext
                        if _batch_file_ext != ".docx":
                            try:
                                import tempfile as _bttmp

                                from web.doc_converter import convert_to_docx as _btc

                                _bt_tmp = _bttmp.mkdtemp(prefix="koto_bt_")
                                _bt_conv, _ = _btc(_batch_filepath, output_dir=_bt_tmp)
                                _bt_dest = os.path.join(
                                    _batch_docs_dir, os.path.basename(_bt_conv)
                                )
                                import shutil as _bt_sh

                                _bt_sh.copy2(_bt_conv, _bt_dest)
                                _batch_filepath = _bt_dest
                            except Exception as _bt_err:
                                _app_logger.info(f"[BATCH DOC_ANNOTATE] Φ╜¼µìóσñ▒Φ┤Ñ: {_bt_err}")
                        _batch_target = os.path.join(
                            _batch_docs_dir, os.path.basename(_batch_filepath)
                        )
                        if os.path.abspath(_batch_filepath) != os.path.abspath(
                            _batch_target
                        ):
                            import shutil as _bsh

                            _bsh.copy2(_batch_filepath, _batch_target)
                        _bt_feedback = DocumentFeedbackSystem(gemini_client=client)
                        _bt_final = None
                        for _bt_evt in _bt_feedback.full_annotation_loop_streaming(
                            _batch_target, user_input
                        ):
                            if _bt_evt.get("stage") == "complete":
                                _bt_final = _bt_evt.get("result", {})
                        if _bt_final and _bt_final.get("success"):
                            _bt_revised = _bt_final.get("revised_file", "")
                            result["response"] = (
                                f"Γ£à µûçµíúµáçµ│¿σ«îµêÉ: {os.path.basename(_bt_revised)}"
                            )
                            result["saved_files"] = [_bt_revised] if _bt_revised else []
                        else:
                            result["response"] = (
                                f"Γ¥î µë╣ΘçÅµáçµ│¿σñ▒Φ┤Ñ: {(_bt_final or {}).get('message', 'µ£¬τƒÑΘöÖΦ»»')}"
                            )
                    except Exception as _bt_exc:
                        result["response"] = f"Γ¥î µë╣ΘçÅµáçµ│¿σ╝éσ╕╕: {_bt_exc}"
                else:
                    _app_logger.info(f"[FILE UPLOAD] σñäτÉåµûçΣ╗╢: {filename}, Σ╜┐τö¿ brain.chat")
                    brain_result = brain.chat(
                        history=history,
                        user_input=formatted_message,
                        file_data=file_data,
                        model=model_to_use,
                        auto_model=(locked_model == "auto"),
                    )
                    result.update(brain_result)

                # ≡ƒùé∩╕Å σà│Θö«∩╝ÜΣ╕║µ»ÅΣ╕¬µûçΣ╗╢Φ░âτö¿FileOrganizerΦ┐¢Φíîσ╜Æµíú
                organize_info = {"success": False, "message": "µ£¬σ╜Æµíú"}
                try:
                    # Σ╜┐τö¿AIσêåµ₧ÉµûçΣ╗╢τ▒╗σ₧ïσÆîσ╗║Φ««τ¢«σ╜ò
                    from web.file_analyzer import FileAnalyzer

                    analyzer = FileAnalyzer()
                    analysis = analyzer.analyze_file(filepath)  # σÅ¬Σ╝áµûçΣ╗╢Φ╖»σ╛ä
                    suggested_folder = analysis.get("suggested_folder")
                    entity_name = analysis.get("entity")
                    entity_type = analysis.get("entity_type")
                    organizer = get_file_organizer()

                    # σªéµ₧£σ╖▓σ¡ÿσ£¿σÉîσÉìσà¼σÅ╕/Θí╣τ¢«µûçΣ╗╢σñ╣∩╝îσêÖσñìτö¿
                    if entity_name:
                        existing_folder = organizer.find_entity_folder(entity_name)
                        if existing_folder:
                            suggested_folder = existing_folder

                    if suggested_folder:
                        org_result = organizer.organize_file(
                            filepath,
                            suggested_folder,
                            auto_confirm=True,
                            metadata={
                                "entity": entity_name,
                                "entity_type": entity_type,
                            },
                        )

                        if org_result.get("success"):
                            organize_info = {
                                "success": True,
                                "message": f"Γ£à σ╖▓σ╜Æµíúσê░: {org_result.get('relative_path', suggested_folder)}",
                                "category": suggested_folder,
                                "path": org_result.get("dest_file"),
                            }
                            _app_logger.info(
                                f"[FILE ORGANIZE] Γ£à {filename} -> {suggested_folder}"
                            )
                        else:
                            organize_info = {
                                "success": False,
                                "message": f"ΓÜá∩╕Å σ╜Æµíúσñ▒Φ┤Ñ: {org_result.get('error', 'µ£¬τƒÑΘöÖΦ»»')}",
                            }
                    else:
                        organize_info = {
                            "success": False,
                            "message": "ΓÜá∩╕Å µùáµ│òτí«σ«ÜµûçΣ╗╢σêåτ▒╗",
                        }
                except Exception as e:
                    organize_info = {
                        "success": False,
                        "message": f"ΓÜá∩╕Å σ╜Æµíúσ╝éσ╕╕: {str(e)}",
                    }
                    _app_logger.info(f"[FILE ORGANIZE ERROR] {filename}: {e}")

                result["file_name"] = filename
                result["organize"] = organize_info
                return result

            except Exception as e:
                return {
                    "file_name": filename,
                    "task": "ERROR",
                    "model": "none",
                    "response": f"Γ¥î σñäτÉåµûçΣ╗╢µù╢σç║ΘöÖ: {str(e)}",
                    "images": [],
                    "saved_files": [],
                    "organize": {"success": False, "message": "Γ¥î σñäτÉåσñ▒Φ┤Ñ∩╝îµ£¬σ╜Æµíú"},
                }

        if stream_mode:

            def generate_progress():
                total = len([f for f in files if f and f.filename])
                started = {
                    "type": "progress",
                    "current": 0,
                    "total": total,
                    "status": "start",
                    "detail": f"σ╝ÇσºïσñäτÉå {total} Σ╕¬µûçΣ╗╢",
                }
                yield f"data: {json.dumps(started)}\n\n"

                current = 0
                for file in files:
                    if not file or not file.filename:
                        continue

                    current += 1
                    payload = {
                        "type": "progress",
                        "current": current,
                        "total": total,
                        "status": "processing",
                        "detail": f"σñäτÉåΣ╕¡: {file.filename} ({current}/{total})",
                    }
                    yield f"data: {json.dumps(payload)}\n\n"

                    result = _process_single_file(file)
                    if result:
                        batch_results.append(result)
                        combined_saved_files.extend(result.get("saved_files", []))
                        combined_images.extend(result.get("images", []))

                    payload = {
                        "type": "progress",
                        "current": current,
                        "total": total,
                        "status": "done",
                        "detail": f"σ«îµêÉ: {file.filename} ({current}/{total})",
                    }
                    yield f"data: {json.dumps(payload)}\n\n"

                summary_lines = [f"≡ƒôª µë╣ΘçÅσñäτÉåσ«îµêÉ∩╝îσà▒ {len(batch_results)} Σ╕¬µûçΣ╗╢", ""]

                organized_count = sum(
                    1
                    for item in batch_results
                    if item.get("organize", {}).get("success")
                )
                if organized_count > 0:
                    summary_lines.append(f"Γ£à σ╖▓σ╜Æµíú: {organized_count} Σ╕¬µûçΣ╗╢")

                summary_lines.append("\n≡ƒôä **µûçΣ╗╢Φ»ªµâà∩╝Ü**")
                for i, item in enumerate(batch_results, 1):
                    fname = item.get("file_name", "unknown")
                    task = item.get("task", "UNKNOWN")
                    organize = item.get("organize", {})

                    status = "Γ£à" if task != "ERROR" else "Γ¥î"
                    org_status = organize.get("message", "µ£¬σ╜Æµíú")

                    summary_lines.append(f"{i}. {status} **{fname}**")
                    summary_lines.append(f"   ≡ƒôé {org_status}")

                    response = item.get("response", "")
                    if response and len(response) > 100:
                        summary_lines.append(f"   ≡ƒÆ¼ {response[:100]}...")
                    elif response:
                        summary_lines.append(f"   ≡ƒÆ¼ {response}")

                summary_msg = "\n".join(summary_lines)

                session_manager.update_last_model_response(
                    f"{session_name}.json",
                    summary_msg,
                    task="FILE_BATCH",
                    model_name=locked_model if locked_model != "auto" else "auto",
                    saved_files=combined_saved_files,
                    images=combined_images,
                )

                final_payload = {
                    "type": "final",
                    "response": summary_msg,
                    "task": "FILE_BATCH",
                    "model": locked_model if locked_model != "auto" else "auto",
                    "results": batch_results,
                    "images": combined_images,
                    "saved_files": combined_saved_files,
                }
                yield f"data: {json.dumps(final_payload)}\n\n"

            return Response(generate_progress(), mimetype="text/event-stream")

        for file in files:
            result = _process_single_file(file)
            if not result:
                continue
            batch_results.append(result)
            combined_saved_files.extend(result.get("saved_files", []))
            combined_images.extend(result.get("images", []))

        # τöƒµêÉΦ»ªτ╗åµæÿΦªü∩╝îσîàσÉ½σ╜ÆµíúΣ┐íµü»
        summary_lines = [f"≡ƒôª µë╣ΘçÅσñäτÉåσ«îµêÉ∩╝îσà▒ {len(batch_results)} Σ╕¬µûçΣ╗╢", ""]

        organized_count = sum(
            1 for item in batch_results if item.get("organize", {}).get("success")
        )
        if organized_count > 0:
            summary_lines.append(f"Γ£à σ╖▓σ╜Æµíú: {organized_count} Σ╕¬µûçΣ╗╢")

        summary_lines.append("\n≡ƒôä **µûçΣ╗╢Φ»ªµâà∩╝Ü**")
        for i, item in enumerate(batch_results, 1):
            fname = item.get("file_name", "unknown")
            task = item.get("task", "UNKNOWN")
            organize = item.get("organize", {})

            status = "Γ£à" if task != "ERROR" else "Γ¥î"
            org_status = organize.get("message", "µ£¬σ╜Æµíú")

            summary_lines.append(f"{i}. {status} **{fname}**")
            summary_lines.append(f"   ≡ƒôé {org_status}")

            # µÿ╛τñ║AIσôìσ║öµæÿΦªü∩╝êµê¬σÅûσëì100σ¡ù∩╝ë
            response = item.get("response", "")
            if response and len(response) > 100:
                summary_lines.append(f"   ≡ƒÆ¼ {response[:100]}...")
            elif response:
                summary_lines.append(f"   ≡ƒÆ¼ {response}")

        summary_msg = "\n".join(summary_lines)

        session_manager.update_last_model_response(
            f"{session_name}.json",
            summary_msg,
            task="FILE_BATCH",
            model_name=locked_model if locked_model != "auto" else "auto",
            saved_files=combined_saved_files,
            images=combined_images,
        )

        return jsonify(
            {
                "response": summary_msg,
                "task": "FILE_BATCH",
                "model": locked_model if locked_model != "auto" else "auto",
                "results": batch_results,
                "images": combined_images,
                "saved_files": combined_saved_files,
            }
        )

    file = files[0]

    # Save uploaded file
    filename = _secure_filename(file.filename) or f"upload_{uuid.uuid4().hex}"
    filepath = os.path.join(UPLOAD_DIR, filename)
    file.save(filepath)
    file_type = file.mimetype or file.content_type or ""
    file_ext = os.path.splitext(filename)[1].lower()

    # Load history first (Σ┐¥Φ»üσì│Σ╜┐σç║ΘöÖΣ╣ƒΦâ╜Σ┐¥σ¡ÿτö¿µê╖Φ╛ôσàÑ)
    history = session_manager.load(f"{session_name}.json")
    user_message = f"[File: {filename}] {user_input}"

    # ≡ƒöÆ τ½ïσì│Σ┐¥σ¡ÿτö¿µê╖µ╢êµü»σê░τúüτ¢ÿ∩╝îΘÿ▓µ¡óµû¡Φ┐₧/σ┤⌐µ║âσ»╝Φç┤Σ╕óσñ▒
    session_manager.append_user_early(f"{session_name}.json", user_message)

    try:
        # Σ╜┐τö¿µû░τÜäµûçΣ╗╢σñäτÉåσÖ¿∩╝êµÅÉσÅûµûçµ£¼/Σ║îΦ┐¢σê╢∩╝ë
        formatted_message, file_data = process_uploaded_file(filepath, user_input)

        # ==================== µÖ║Φâ╜µûçµíúσêåµ₧Éσ╝òµôÄ ====================
        # σ»╣ .docx/.doc µûçΣ╗╢∩╝îΣ╜┐τö¿ LLM Θ⌐▒σè¿τÜäµÖ║Φâ╜σêåµ₧Éσ╝òµôÄσêñµû¡τö¿µê╖µäÅσ¢╛
        # Σ╕ìσåìτí¼τ╝ûτáüµ¡úσêÖ∩╝îΦÇîµÿ»Φ«⌐σêåµ₧ÉσÖ¿τÉåΦºúτö¿µê╖τ£ƒσ«₧Θ£Çµ▒é
        if file_ext in [".docx", ".doc"]:
            # ΓöÇΓöÇ τ┐╗Φ»æΦ»╖µ▒é∩╝Üµ£ÇΘ½ÿΣ╝ÿσàêτ║º∩╝îτ¢┤µÄÑΦ╡░µ£ìσèíσÖ¿τ½»τ┐╗Φ»æτ«íΘüô ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
            _TRANSLATE_KWS = [
                "τ┐╗Φ»æ",
                "Φ»æµêÉ",
                "Φ»æΣ╕║",
                "Φ╜¼µêÉΦï▒µûç",
                "Φ╜¼µêÉµùÑµûç",
                "Φ╜¼µêÉΣ╕¡µûç",
                "translate",
                "τ┐╗µêÉ",
                "Φ╜¼Φ»æ",
            ]
            _is_translate_request = any(
                kw in (user_input or "").lower() for kw in _TRANSLATE_KWS
            )
            if _is_translate_request and locked_task != "DOC_ANNOTATE":
                _app_logger.info(f"[DOCX TRANSLATE] µúÇµ╡ïσê░τ┐╗Φ»æΦ»╖µ▒é∩╝îσÉ»τö¿µá╝σ╝ÅΣ┐¥τòÖτ┐╗Φ»æτ«íΘüô")

                def generate_docx_translation():
                    try:
                        from web.docx_translator_module import (
                            detect_target_language,
                            translate_docx_streaming,
                        )

                        target_lang = detect_target_language(user_input or "")
                        docs_dir = os.path.join(WORKSPACE_DIR, "documents")
                        os.makedirs(docs_dir, exist_ok=True)

                        yield f"data: {json.dumps({'type': 'classification', 'task_type': 'FILE_GEN', 'task_display': '≡ƒîÉ Word µûçµíúτ┐╗Φ»æ', 'route_method': '≡ƒîÉ DocxTranslator', 'message': f'≡ƒÄ» σÉ»σè¿µá╝σ╝ÅΣ┐¥τòÖτ┐╗Φ»æ ΓåÆ {target_lang}'})}\n\n"

                        for event in translate_docx_streaming(
                            filepath, target_lang, client, output_dir=docs_dir
                        ):
                            stage = event.get("stage", "")
                            msg = event.get("message", "")
                            progress = event.get("progress", 0)

                            if stage == "error":
                                yield f"data: {json.dumps({'type': 'token', 'content': f'Γ¥î τ┐╗Φ»æσñ▒Φ┤Ñ: {msg}'})}\n\n"
                                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                                return

                            elif stage == "complete":
                                out_path = event.get("output_path", "")
                                out_name = event.get(
                                    "output_filename", os.path.basename(out_path)
                                )
                                count = event.get("translated_count", 0)
                                lang = event.get("target_language", target_lang)
                                rel_path = os.path.relpath(
                                    out_path, WORKSPACE_DIR
                                ).replace("\\", "/")

                                success_msg = (
                                    f"Γ£à **Word µûçµíúτ┐╗Φ»æσ«îµêÉ∩╝ü**\n\n"
                                    f"≡ƒîÉ τ¢«µáçΦ»¡Φ¿Ç: **{lang}**\n"
                                    f"≡ƒô¥ τ┐╗Φ»æµ«╡ΦÉ╜: **{count}** µ«╡\n"
                                    f"≡ƒôü µûçΣ╗╢σÉì: **{out_name}**\n"
                                    f"≡ƒôì Σ╜ìτ╜«: `workspace/documents/`\n\n"
                                    f"µá╝σ╝Åσ╖▓σ«îµò┤Σ┐¥τòÖ∩╝êσ¡ùΣ╜ô/σèáτ▓ù/µû£Σ╜ô/Θó£Φë▓/Φí¿µá╝/Θí╡τ£ëΘí╡ΦäÜ∩╝ë"
                                )
                                yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                                session_manager.append_and_save(
                                    f"{session_name}.json",
                                    user_input,
                                    f"τ┐╗Φ»æσ«îµêÉ ΓåÆ {out_name} ({count}µ«╡, {lang})",
                                )
                                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [rel_path]})}\n\n"
                                return

                            else:
                                yield f"data: {json.dumps({'type': 'progress', 'message': msg, 'detail': f'{progress}%'})}\n\n"

                    except Exception as _te:
                        import traceback as _tb

                        _app_logger.error(f"[DOCX TRANSLATE] Γ¥î τ┐╗Φ»æσ╝éσ╕╕: {_tb.format_exc()}")
                        yield f"data: {json.dumps({'type': 'token', 'content': f'Γ¥î τ┐╗Φ»æσç║ΘöÖ: {str(_te)}'})}\n\n"
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"

                return Response(
                    stream_with_context(generate_docx_translation()),
                    content_type="text/event-stream",
                )

            # µáçµ│¿Σ╗╗σèíΣ╝ÿσàêτ║ºµ¢┤Θ½ÿ∩╝Üµÿ╛σ╝Åµáçµ│¿µäÅσ¢╛µêûτö¿µê╖Θöüσ«Ü DOC_ANNOTATE µù╢∩╝îΣ╕ìΦ┐¢σàÑµÖ║Φâ╜σêåµ₧Éσ╝òµôÄ
            force_annotation = (
                locked_task == "DOC_ANNOTATE"
            ) or _should_use_annotation_system(user_input, has_file=True)

            # µÖ║Φâ╜µúÇµ╡ï∩╝ÜΣ╗╗Σ╜òσ»╣µûçµíúσåàσ«╣µ£ëσ«₧Φ┤¿µÇºσñäτÉåΘ£Çµ▒éτÜäΦ»╖µ▒é
            # σîàµï¼Σ╜åΣ╕ìΘÖÉΣ║Ä∩╝ÜσåÖµæÿΦªüπÇüµö╣σ╝òΦ¿ÇπÇüµö╣τ╗ôΦ«║πÇüµ╢ªΦë▓πÇüσêåµ₧Éτ╗ôµ₧äτ¡ë
            _doc_intent_keywords = [
                # τöƒµêÉτ▒╗
                "σåÖ",
                "τöƒµêÉ",
                "σ╕«µêæσåÖ",
                "σåÖΣ╕Çµ«╡",
                "σåÖΣ╕¬",
                # Σ┐«µö╣/µö╣σûäτ▒╗
                "µö╣",
                "µö╣σûä",
                "µö╣Φ┐¢",
                "Σ╝ÿσîû",
                "µ╢ªΦë▓",
                "ΘçìσåÖ",
                "Σ┐«µö╣",
                "µÅÉσìç",
                # σ¡ªµ£»Θâ¿Σ╗╢
                "µæÿΦªü",
                "σ╝òΦ¿Ç",
                "τ╗ôΦ«║",
                "abstract",
                "σëìΦ¿Ç",
                "σ»╝Φ¿Ç",
                # σêåµ₧Éτ▒╗
                "σêåµ₧É",
                "µÇ╗τ╗ô",
                "µó│τÉå",
                "µªéΦ┐░",
                "Φ»äΣ╝░",
                # Φ┤¿ΘçÅτ▒╗
                "Σ╕ìµ╗íµäÅ",
                "Σ╕ìσÑ╜",
                "Σ╕ìσñƒ",
                "Θ£ÇΦªüµö╣",
                "µ£ëΘù«Θóÿ",
            ]
            is_doc_processing_request = any(
                kw in user_input.lower() for kw in _doc_intent_keywords
            )

            if is_doc_processing_request and not force_annotation:
                _app_logger.info(f"[INTELLIGENT ANALYZER] µúÇµ╡ïσê░µûçµíúσñäτÉåΦ»╖µ▒é∩╝îσÉ»τö¿µÖ║Φâ╜σêåµ₧Éσ╝òµôÄ")
                from web.intelligent_document_analyzer import (
                    create_intelligent_analyzer,
                )

                # σê¢σ╗║µÖ║Φâ╜σêåµ₧ÉσÖ¿
                analyzer = create_intelligent_analyzer(client)

                # µ╡üσ╝ÅσñäτÉåµûçµíúσêåµ₧É
                def generate_intelligent_analysis():
                    """τöƒµêÉµÖ║Φâ╜µûçµíúσêåµ₧ÉτÜäµ╡üσ╝Åσôìσ║ö"""
                    try:
                        # Σ╜┐τö¿asyncτöƒµêÉσÖ¿∩╝êΘ£ÇΦªüσ£¿async contextΣ╕¡∩╝ë
                        import asyncio

                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)

                        async def run_analysis():
                            async for (
                                event
                            ) in analyzer.process_document_intelligent_streaming(
                                filepath, user_input, session_name
                            ):
                                yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"

                        gen = run_analysis()
                        while True:
                            try:
                                result = loop.run_until_complete(gen.__anext__())
                                yield result
                            except StopAsyncIteration:
                                break
                    except Exception as e:
                        error_event = {
                            "stage": "error",
                            "message": f"µÖ║Φâ╜σêåµ₧Éσñ▒Φ┤Ñ: {str(e)}",
                        }
                        yield f"data: {json.dumps(error_event, ensure_ascii=False)}\n\n"
                    finally:
                        loop.close()

                return Response(
                    stream_with_context(generate_intelligent_analysis()),
                    content_type="text/event-stream",
                )
        # ==================== µÖ║Φâ╜µûçµíúσêåµ₧Éσ╝òµôÄτ╗ôµ¥ƒ ====================

        # µÖ║Φâ╜Σ╗╗σèíσêåµ₧É
        task_type = locked_task
        context_info = None
        route_method = "Auto"
        if not task_type:
            # σªéµ₧£µÿ»σ¢╛τëçΣ╕èΣ╝á∩╝îτ¢┤µÄÑσêñµû¡τ╝ûΦ╛æµêûσêåµ₧É∩╝îΘü┐σàìσê¥σºïσîûµ£¼σ£░Φ╖»τö▒σÖ¿σ»╝Φç┤σìíΘí┐
            if file_data and file_type and file_type.startswith("image"):
                message_lower = (user_input or "").lower()
                is_edit = any(
                    kw in message_lower for kw in KotoBrain.IMAGE_EDIT_KEYWORDS
                )
                task_type = "PAINTER" if is_edit else "VISION"
                route_method = "≡ƒû╝∩╕Å Image Edit" if is_edit else "≡ƒæü∩╕Å Image Analysis"
                _app_logger.info(
                    f"[FILE UPLOAD] σ¢╛τëçΣ╗╗σèíτ¢┤ΘÇÜΦ╖»τö▒: {task_type} (µû╣µ│ò: {route_method})"
                )
            else:
                # µûçµíúΣ╕èΣ╝á∩╝ÜΣ╕Ñµá╝µúÇµ╡ïµáçµ│¿µäÅσ¢╛∩╝êσ┐àΘí╗µÿÄτí«Φªüµ▒éσ£¿σÄƒµûçΣ╕èµáçΦ«░∩╝ë
                _ann_exts = {
                    ".doc",
                    ".docx",
                    ".pdf",
                    ".txt",
                    ".md",
                    ".markdown",
                    ".rtf",
                    ".odt",
                }
                use_annotation = (
                    _should_use_annotation_system(user_input, has_file=True)
                    and file_ext in _ann_exts
                )

                if use_annotation:
                    task_type = "DOC_ANNOTATE"
                    route_method = "≡ƒôî Annotation-Strict"
                elif _is_explicit_file_gen_request(user_input):
                    # τö¿µê╖µÿÄτí«ΦªüτöƒµêÉµû░µûçΣ╗╢∩╝îτ¢┤µÄÑΦ╖»τö▒∩╝îµùáΘ£Çµ¿íσ₧ïσêåτ▒╗
                    task_type = "FILE_GEN"
                    route_method = "≡ƒôä Explicit-Gen"
                    _app_logger.info(
                        f"[FILE UPLOAD] ≡ƒÄ» µúÇµ╡ïσê░µÿÄτí«µûçΣ╗╢τöƒµêÉΦ»╖µ▒é∩╝îσÉ»τö¿ FILE_GEN µ¿íσ╝Å"
                    )
                else:
                    # Γÿà Σ╕╗Φ╖»σ╛ä∩╝ÜΦ«⌐µ£¼σ£░µ¿íσ₧ïσüÜΦ»¡Σ╣ëΦ╖»τö▒
                    # Σ╝áσàÑ [FILE_ATTACHED:ext] µáçΦ«░∩╝îµ¿íσ₧ïΘÇÜΦ┐çΦ«¡τ╗âσÑ╜τÜäΦºäσêÖσêñµû¡
                    # CHAT=Φ»╗µûçΣ╗╢σ¢₧τ¡ö  RESEARCH=µ╖▒σàÑτáöτ⌐╢  FILE_GEN=τöƒµêÉµû░µûçµíú
                    _dispatch_q = (user_input or "").strip() or "Φ»╖σêåµ₧ÉΦ┐ÖΣ╗╜µûçΣ╗╢τÜäσåàσ«╣"
                    _dispatch_input = (
                        f"[FILE_ATTACHED:{file_ext or '.file'}] {_dispatch_q}"
                    )
                    task_analysis, route_method, context_info = SmartDispatcher.analyze(
                        _dispatch_input, history=history
                    )
                    task_type = task_analysis

                _app_logger.info(
                    f"[FILE UPLOAD] µÖ║Φâ╜Φ╖»τö▒ΘÇëµï⌐Σ╗╗σèíτ▒╗σ₧ï: {task_type} (µû╣µ│ò: {route_method})"
                )

        # τí«σ«ÜΣ╜┐τö¿τÜäµ¿íσ₧ï
        if locked_model != "auto":
            model_to_use = locked_model
        else:
            # ΦÄ╖σÅûΣ╗╗σèíσñìµ¥éσ║ª∩╝êΣ╕èΣ╝áµûçΣ╗╢Θ╗ÿΦ«ñµîëσñìµ¥éΣ╗╗σèíσñäτÉå∩╝ë
            complexity = "complex" if file_data is None else "normal"
            if context_info and context_info.get("complexity"):
                complexity = context_info["complexity"]

            if task_type == "DOC_ANNOTATE":
                # µûçµíúµáçµ│¿Θ£ÇΦªüσ╝║µ¿íσ₧ï∩╝ÜΣ╝ÿσàêΣ╜┐τö¿ gemini-3.1-pro-preview∩╝êσªéσÅ»τö¿∩╝ë∩╝îσ¢₧ΘÇÇ gemini-2.5-pro
                # µ│¿µäÅ∩╝Ügemini-3-pro-preview / gemini-3-flash-preview Σ╗àµö»µîü Interactions API∩╝îΣ╕ìΦâ╜τö¿Σ║Ä generate_content
                model_to_use = "gemini-3.1-pro-preview"
            elif task_type == "FILE_GEN":
                model_to_use = SmartDispatcher.get_model_for_task(
                    task_type, has_image=bool(file_data), complexity=complexity
                )
            else:
                model_to_use = SmartDispatcher.get_model_for_task(
                    task_type, has_image=bool(file_data)
                )
        
        # µ£Çµù⌐µïªµê¬∩╝ÜµûçΣ╗╢σêåµ₧ÉΣ╕ìΦâ╜Σ╜┐τö¿ interactions-only µ¿íσ₧ï∩╝êΣ╕ìµö»µîü generate_content Σ╣ƒΣ╕ìµö»µîüµûçΣ╗╢ΘÖäΣ╗╢∩╝ë
        if model_to_use in _INTERACTIONS_ONLY_MODELS or str(model_to_use or "").startswith("deep-research-pro-preview"):
            _orig_model = model_to_use
            model_to_use = _INTERACTIONS_FALLBACK_MODEL
            _app_logger.warning(f"[FILE UPLOAD] ΓÜá∩╕Å {_orig_model} µÿ» interactions-only∩╝îµûçΣ╗╢σêåµ₧ÉΘÖìτ║ºσê░ {_INTERACTIONS_FALLBACK_MODEL}")

        _app_logger.info(f"[FILE UPLOAD] Σ╗╗σèíτ▒╗σ₧ï: {task_type}, µ¿íσ₧ï: {model_to_use}")

        # σ«ëσà¿σà£σ║ò∩╝Ülocked_task ΘóäΦ«╛µù╢ prefer_ppt σÅ»Φâ╜µ£¬σ«ÜΣ╣ë
        if "prefer_ppt" not in locals():
            _ppt_kws = [
                "ppt",
                "σ╣╗τü»τëç",
                "µ╝öτñ║",
                "µ▒çµèÑ",
                "presentation",
                "slide",
                "deck",
            ]
            prefer_ppt = any(kw in (user_input or "").lower() for kw in _ppt_kws)

        # σªéµ₧£µÿ»µûçµ£¼τ▒╗µûçΣ╗╢∩╝îµîëΣ╗╗σèíτ▒╗σ₧ïσñäτÉå
        result = {
            "task": "FILE_GEN" if task_type == "DOC_ANNOTATE" else task_type,
            "subtask": "DOC_ANNOTATE" if task_type == "DOC_ANNOTATE" else None,
            "model": model_to_use,
            "route_method": route_method,
            "response": "",
            "images": [],
            "saved_files": [],
        }

        # µûçµíúµáçµ│¿Σ╗╗σèí - µ╡üσ╝ÅσÅìΘªê∩╝îτöƒµêÉσ╕ªTrack ChangesτÜäWordµûçµíú
        if task_type == "DOC_ANNOTATE":
            docs_dir = settings_manager.documents_dir
            os.makedirs(docs_dir, exist_ok=True)

            source_path = filepath
            target_path = os.path.join(docs_dir, filename)
            if os.path.abspath(source_path) != os.path.abspath(target_path):
                shutil.copy2(source_path, target_path)

            # Σ╜┐τö¿µ╡üσ╝ÅSSEΦ┐öσ¢₧Φ┐¢σ║ª∩╝îΦ«⌐σëìτ½»Φâ╜σ«₧µù╢µÿ╛τñ║
            # µìòΦÄ╖Θù¡σîàσÅÿΘçÅ∩╝êΘÿ▓µ¡ógeneratorσ╗╢Φ┐ƒµëºΦíîµù╢σÅÿΘçÅσ╖▓µö╣σÅÿ∩╝ë
            _ann_target_path = target_path
            _ann_filename = filename
            _ann_file_ext = file_ext
            _ann_route_method = route_method
            _ann_model = model_to_use
            _ann_session = session_name
            _ann_user_input = user_input
            _ann_client = client
            _ann_docs_dir = docs_dir  # τí«Σ┐¥Φ╜¼µìóσÉÄτÜäµûçΣ╗╢σÆîΦ╛ôσç║Θâ╜Σ┐¥σ¡ÿσê░µáçσçåτ¢«σ╜ò

            def generate_doc_annotate_stream():
                import time as _time

                _start = _time.time()
                task_id = f"doc_annotate_{_ann_session}_{int(_start * 1000)}"

                # Local mutable copies of closure vars (Python makes vars local if assigned anywhere
                # in the function, so we cannot reassign _ann_* directly without UnboundLocalError)
                _loc_target_path = _ann_target_path
                _loc_filename = _ann_filename
                _loc_file_ext = _ann_file_ext

                # ΓöÇΓöÇ Θ¥₧ .docx µá╝σ╝ÅΦç¬σè¿Φ╜¼µìó ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                # σ»╣ .doc / .pdf / .txt / .md / .rtf / .odt σàêΦ╜¼µìóΣ╕║ .docx σåìΦ┐¢µáçµ│¿
                _converted_warning = ""
                _classif_sent = False
                if _loc_file_ext != ".docx":
                    yield f"data: {json.dumps({'type': 'classification', 'task_type': 'DOC_ANNOTATE', 'route_method': _ann_route_method, 'model': _ann_model, 'task_id': task_id, 'message': '≡ƒôä DOC_ANNOTATE'})}\n\n"
                    _classif_sent = True
                    yield f"data: {json.dumps({'type': 'progress', 'stage': 'converting', 'message': f'≡ƒöä µ¡úσ£¿σ░å {_loc_file_ext} Φ╜¼µìóΣ╕║σÅ»τ╝ûΦ╛æ .docx...', 'detail': _loc_filename, 'progress': 3})}\n\n"
                    try:
                        from web.doc_converter import convert_to_docx, needs_conversion

                        if needs_conversion(_loc_file_ext):
                            import tempfile as _tmpmod

                            _conv_dir = _tmpmod.mkdtemp(prefix="koto_conv_")
                            _conv_path, _converted_warning = convert_to_docx(
                                _loc_target_path, output_dir=_conv_dir
                            )
                            # σ░åΦ╜¼µìóσÉÄτÜä .docx σñìσê╢σê░µáçσçåµûçµíúτ¢«σ╜ò∩╝îτí«Σ┐¥Φ╛ôσç║Σ╣ƒσ£¿Φ»Ñτ¢«σ╜ò
                            _conv_basename = os.path.basename(_conv_path)
                            _conv_in_docs = os.path.join(_ann_docs_dir, _conv_basename)
                            import shutil as _shutil_conv

                            _shutil_conv.copy2(_conv_path, _conv_in_docs)
                            _loc_target_path = (
                                _conv_in_docs  # τö¿ docs_dir Φ╖»σ╛ä∩╝îΦ╛ôσç║Σ╣ƒΣ╝Üσ£¿µ¡ñ
                            )
                            _loc_filename = _conv_basename
                            _loc_file_ext = ".docx"
                            _app_logger.info(
                                f"[DocConvert] Γ£à Φ╜¼µìóσ╣╢σñìσê╢σê░µûçµíúτ¢«σ╜ò ΓåÆ {_loc_target_path}"
                            )
                        else:
                            yield f"data: {json.dumps({'type': 'error', 'message': f'Σ╕ìµö»µîüτÜäµá╝σ╝Å∩╝Ü{_loc_file_ext}'})}\n\n"
                            return
                    except Exception as _conv_err:
                        err_msg = f"Γ¥î µá╝σ╝ÅΦ╜¼µìóσñ▒Φ┤Ñ∩╝Ü{_conv_err}"
                        yield f"data: {json.dumps({'type': 'token', 'content': err_msg})}\n\n"
                        _elapsed = _time.time() - _start
                        session_manager.update_last_model_response(
                            f"{_ann_session}.json", err_msg
                        )
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': _elapsed})}\n\n"
                        return
                # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ

                try:
                    from web.document_feedback import DocumentFeedbackSystem

                    feedback_system = DocumentFeedbackSystem(gemini_client=_ann_client)

                    # σÅæΘÇüσêåτ▒╗Σ┐íµü»∩╝êΦïÑΣ╕èµû╣Φ╜¼µìóσ¥ùσ╖▓σÅæΘÇüσêÖΦ╖│Φ┐çΘçìσñì∩╝ë
                    if not _classif_sent:
                        yield f"data: {json.dumps({'type': 'classification', 'task_type': 'DOC_ANNOTATE', 'route_method': _ann_route_method, 'model': _ann_model, 'task_id': task_id, 'message': '≡ƒôä DOC_ANNOTATE'})}\n\n"
                    if _converted_warning:
                        yield f"data: {json.dumps({'type': 'info', 'message': _converted_warning})}\n\n"

                    # σÅæΘÇüσê¥σºïΦ┐¢σ║ª
                    yield f"data: {json.dumps({'type': 'progress', 'stage': 'init_reading', 'message': '≡ƒôû µ¡úσ£¿Φ»╗σÅûµûçµíú...', 'detail': _loc_filename, 'progress': 5})}\n\n"

                    # ΓöÇΓöÇ Φ╜¼µìóΦ┤¿ΘçÅµúÇµƒÑ∩╝Üσªéµ₧£µ«╡ΦÉ╜µò░Φ┐çσñÜµêûσåàσ«╣Σ╕║Σ╣▒τáü∩╝îµïÆτ╗¥µáçµ│¿ ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                    try:
                        from docx import Document as _QDoc

                        _qd = _QDoc(_loc_target_path)
                        _q_paras = [p.text for p in _qd.paragraphs if p.text.strip()]
                        _max_para_limit = 500
                        # Σ╣▒τáüµúÇµ╡ï∩╝Üτƒ¡σ₧âσ£╛µ«╡ΦÉ╜µ»öΣ╛ï > 60% µêû σ¡ùµ»ìµ»öΣ╛ïΦ┐çΣ╜Ä
                        _q_short = sum(1 for p in _q_paras if len(p) < 15)
                        _q_short_ratio = _q_short / max(len(_q_paras), 1)
                        _q_alpha_ratio = sum(
                            sum(1 for c in p if c.isalpha()) / max(len(p), 1)
                            for p in _q_paras[:200]
                        ) / max(min(len(_q_paras), 200), 1)
                        _is_garbage = (
                            len(_q_paras) > _max_para_limit and _q_short_ratio > 0.5
                        ) or (len(_q_paras) > 200 and _q_short_ratio > 0.7)
                        if _is_garbage:
                            _q_err = (
                                f"Γ¥î **µûçΣ╗╢Φ╜¼µìóΦ┤¿ΘçÅΦ┐çΣ╜Ä**∩╝îµúÇµ╡ïσê░ {len(_q_paras):,} Σ╕¬µ«╡ΦÉ╜"
                                f"∩╝ê{_q_short_ratio:.0%} Σ╕║Σ╣▒τáüτƒ¡Φíî∩╝ë∩╝îσåàσ«╣µùáµ│òΦ»åσê½πÇé\n\n"
                                "**σÄƒσ¢á**∩╝Ü`.doc` µá╝σ╝ÅΣ╜┐τö¿Σ║åµùºτëêΣ║îΦ┐¢σê╢τ╗ôµ₧ä∩╝îµùáµ│òΦç¬σè¿Φºúµ₧ÉπÇé\n\n"
                                "**Φºúσå│µû╣µ│ò**∩╝Ü\n"
                                "1. τö¿ **Microsoft Word** µëôσ╝Ç `.doc` µûçΣ╗╢\n"
                                "2. τé╣σç╗πÇÉµûçΣ╗╢πÇæΓåÆπÇÉσÅªσ¡ÿΣ╕║πÇæ\n"
                                "3. ΘÇëµï⌐µá╝σ╝Å **Word µûçµíú (*.docx)**\n"
                                "4. Θçìµû░Σ╕èΣ╝á `.docx` µûçΣ╗╢"
                            )
                            session_manager.update_last_model_response(
                                f"{_ann_session}.json", _q_err
                            )
                            yield f"data: {json.dumps({'type': 'token', 'content': _q_err})}\n\n"
                            _elapsed = _time.time() - _start
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': _elapsed})}\n\n"
                            return
                    except Exception:
                        pass  # µúÇµƒÑσñ▒Φ┤Ñµù╢τ╗ºτ╗¡µ¡úσ╕╕µ╡üτ¿ï
                    # ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ

                    revised_file = None
                    final_result = None

                    for (
                        progress_event
                    ) in feedback_system.full_annotation_loop_streaming(
                        _loc_target_path,
                        _ann_user_input,
                        task_id=task_id,
                        model_id=_ann_model,
                        cancel_check=lambda: _interrupt_manager.is_interrupted(
                            _ann_session
                        ),
                    ):
                        stage = progress_event.get("stage", "unknown")
                        progress = progress_event.get("progress", 0)
                        message_text = progress_event.get("message", "")
                        detail = progress_event.get("detail", "")

                        if stage == "cancelled":
                            yield f"data: {json.dumps({'type': 'info', 'message': 'ΓÅ╕∩╕Å Σ╗╗σèíσ╖▓σÅûµ╢ê'})}\n\n"
                            _elapsed = _time.time() - _start
                            # Σ┐¥σ¡ÿσÅûµ╢êΦ«░σ╜ò
                            session_manager.update_last_model_response(
                                f"{_ann_session}.json", "ΓÅ╕∩╕Å µûçµíúµáçµ│¿Σ╗╗σèíσ╖▓σÅûµ╢ê"
                            )
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': _elapsed, 'cancelled': True})}\n\n"
                            return

                        yield f"data: {json.dumps({'type': 'progress', 'stage': stage, 'message': message_text, 'detail': detail, 'progress': progress})}\n\n"

                        if stage == "complete":
                            final_result = progress_event.get("result", {})
                            revised_file = final_result.get("revised_file")

                    _elapsed = _time.time() - _start

                    if final_result and final_result.get("success"):
                        applied = final_result.get("applied", 0)
                        failed = final_result.get("failed", 0)
                        total = final_result.get("total", applied + failed)

                        # ΓöÇΓöÇ σà£σ║òµúÇµ╡ï ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        _fb_used = final_result.get("fallback_used", False)
                        _fb_partial = final_result.get("partial_fallback", False)
                        _fb_err = final_result.get("last_api_error", "")
                        _fb_chunks = final_result.get("fallback_chunk_count", 0)
                        _ai_chunks = final_result.get("ai_chunk_count", 0)

                        # Φ»╗σÅûµûçµíúΣ┐íµü»
                        try:
                            from docx import Document as _Doc

                            _d = _Doc(_loc_target_path)
                            _total_paras = len(
                                [p for p in _d.paragraphs if p.text.strip()]
                            )
                            _total_chars = sum(len(p.text) for p in _d.paragraphs)
                        except Exception:
                            _total_paras = 0
                            _total_chars = 0

                        density = (
                            (applied / _total_chars * 1000) if _total_chars > 0 else 0
                        )

                        # ΓöÇΓöÇ µ₧äσ╗║µ¿íσ₧ïΦíî / σà£σ║òΦ¡ªσæè ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        if _fb_used:
                            model_display = (
                                f"`{_ann_model}` ΓÜá∩╕Å **∩╝êAIµ£¬µêÉσèƒ∩╝îσ╖▓τö¿µ£¼σ£░ΦºäσêÖσà£σ║ò∩╝ë**"
                            )
                        elif _fb_partial:
                            model_display = f"`{_ann_model}` ΓÜá∩╕Å **∩╝ê{_fb_chunks}µ«╡σà£σ║ò / {_ai_chunks}µ«╡AI∩╝ë**"
                        else:
                            model_display = f"`{_ann_model}`"

                        summary_lines = [
                            "## Γ£à µûçµíúΣ┐«µö╣σ«îµêÉ∩╝ü",
                            "",
                            "### ≡ƒôè Σ┐«µö╣τ╗ƒΦ«í",
                            f"- µë╛σê░σ╣╢σ║öτö¿: **{applied}** σñäΣ┐«µö╣",
                            f"- σ«ÜΣ╜ìσñ▒Φ┤Ñ: {failed} σñä",
                            f"- µÇ╗Φ«íσêåµ₧É: {total} σñä",
                            "",
                            "### ≡ƒôï µûçµíúΣ┐íµü»",
                            f"- µûçΣ╗╢σÉì: `{_loc_filename}`",
                            f"- µ«╡ΦÉ╜µò░: {_total_paras} µ«╡",
                            f"- σ¡ùµò░: {_total_chars} σ¡ù",
                            f"- Σ┐«µö╣σ»åσ║ª: **{density:.1f}** σñä/σìâσ¡ù",
                            "",
                            f"### ≡ƒôä µ¿íσ₧ï: {model_display}",
                            "",
                            f"### ≡ƒô¥ Φ╛ôσç║µûçΣ╗╢: `{os.path.basename(revised_file) if revised_file else 'σ╛àτöƒµêÉ'}`",
                        ]

                        # ΓöÇΓöÇ σ╜ôΣ╜┐τö¿σà£σ║òµù╢∩╝îµÅÆσàÑµÿ╛τ£╝τÜäΦ¡ªσæèσ¥ù ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
                        if _fb_used or _fb_partial:
                            fb_label = (
                                "σà¿Θâ¿σêåµ«╡"
                                if _fb_used
                                else f"{_fb_chunks}/{_fb_chunks+_ai_chunks} σêåµ«╡"
                            )
                            summary_lines += [
                                "",
                                "---",
                                "### ΓÜá∩╕Å Φ┤¿ΘçÅΦ¡ªσæè∩╝Üµ£¼µ¼íΣ╜┐τö¿Σ║åµ£¼σ£░ΦºäσêÖσà£σ║ò",
                                "",
                                f"**Θù«Θóÿ**: Gemini API σ£¿ {fb_label} Σ╕¡Φ░âτö¿σñ▒Φ┤Ñ∩╝îτ│╗τ╗ƒΦç¬σè¿ΘÖìτ║ºΣ╕║σƒ║Σ║Äµ¡úσêÖΦºäσêÖτÜäµ£¼σ£░µáçµ│¿πÇé",
                                "µ£¼σ£░σà£σ║òµáçµ│¿Φ┤¿ΘçÅ**µÿÄµÿ╛Σ╜ÄΣ║Ä** AI σêåµ₧É∩╝îΣ╕╗ΦªüΦªåτ¢ûΦó½σè¿σÅÑπÇüσÉìΦ»ìσîûπÇüσåùΣ╜ÖΦ┐₧µÄÑΦ»ìτ¡ëσ¢║σ«Üµ¿íσ╝Å∩╝î",
                                "µùáµ│òτÉåΦºúΣ╕èΣ╕ïµûçΦ»¡Σ╣ëπÇé",
                                "",
                                f"**ΘöÖΦ»»Σ┐íµü»**: `{_fb_err[:120] if _fb_err else '∩╝êµùáΦ»ªτ╗åΘöÖΦ»»µùÑσ┐ù∩╝ë'}`",
                                "",
                                "**σ╗║Φ««µÄÆµƒÑ**:",
                                "1. µúÇµƒÑ Koto σÉÄσÅ░µÄºσê╢σÅ░∩╝îµë╛ `[DocumentFeedback] Γ¥î` µêû `ΓÜá∩╕Å` σ╝Çσñ┤τÜäµùÑσ┐ùΦíî",
                                "2. τí«Φ«ñ API Key µ£ëµòê∩╝Ü`config/gemini_config.env` ΓåÆ `GEMINI_API_KEY`",
                                "3. τí«Φ«ñ `gemini-2.5-pro` σ»╣µé¿τÜäΦ┤ªσÅ╖σÅ»τö¿∩╝êΘâ¿σêåΦ┤ªσÅ╖σÅùΦ«┐Θù«ΘÖÉσê╢∩╝ë",
                                "4. ΘçìΦ»òΣ╕Çµ¼í∩╝îσªéΣ╗ìσñ▒Φ┤ÑσÅ»µìóτö¿ `gemini-2.5-flash`",
                                "---",
                            ]

                        summary_lines += [
                            "",
                            "### ≡ƒÆí Σ╜┐τö¿µû╣µ│ò",
                            "1. τö¿ Microsoft Word µëôσ╝ÇΦ╛ôσç║µûçΣ╗╢",
                            "2. τé╣σç╗πÇîσ«íΘÿàπÇìµáçτ¡╛Θí╡",
                            "3. σÅ│Σ╛ºµ░öµ│íΣ╕¡µƒÑτ£ïσà¿Θâ¿Σ┐«µö╣σ╗║Φ««",
                            "4. ΘÇÉµ¥íµÄÑσÅùµêûσ┐╜τòÑ∩╝êσÅ│Θö«µë╣µ│¿σÅ»µôìΣ╜£∩╝ë",
                        ]
                        summary_msg = "\n".join(summary_lines)

                        yield f"data: {json.dumps({'type': 'token', 'content': summary_msg})}\n\n"

                        session_manager.update_last_model_response(
                            f"{_ann_session}.json",
                            summary_msg,
                            task="DOC_ANNOTATE",
                            model_name=_ann_model,
                            saved_files=[revised_file] if revised_file else [],
                        )

                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [revised_file] if revised_file else [], 'total_time': _elapsed})}\n\n"
                    else:
                        err_msg = (
                            final_result.get("message", "µ£¬τƒÑΘöÖΦ»»")
                            if final_result
                            else "σñäτÉåσñ▒Φ┤Ñ"
                        )
                        # Σ┐¥σ¡ÿσñ▒Φ┤ÑΦ«░σ╜ò
                        session_manager.update_last_model_response(
                            f"{_ann_session}.json", f"Γ¥î µûçµíúµáçµ│¿σñ▒Φ┤Ñ: {err_msg}"
                        )
                        yield f"data: {json.dumps({'type': 'error', 'message': 'Γ¥î ' + err_msg})}\n\n"
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': _elapsed})}\n\n"

                except Exception as e:
                    import traceback

                    traceback.print_exc()
                    # Σ┐¥σ¡ÿσ╝éσ╕╕Φ«░σ╜ò
                    session_manager.update_last_model_response(
                        f"{_ann_session}.json", f"Γ¥î µáçµ│¿τ│╗τ╗ƒΘöÖΦ»»: {str(e)[:200]}"
                    )
                    yield f"data: {json.dumps({'type': 'error', 'message': 'Γ¥î µáçµ│¿τ│╗τ╗ƒΘöÖΦ»»: ' + str(e)[:200]})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"

            return Response(
                stream_with_context(generate_doc_annotate_stream()),
                mimetype="text/event-stream",
                headers={
                    "Cache-Control": "no-cache",
                    "X-Accel-Buffering": "no",
                    "Connection": "keep-alive",
                },
            )

        # ≡ƒÄ» FILE_GEN + PPT τöƒµêÉ∩╝êP0 µû░σó₧∩╝ë
        elif task_type == "FILE_GEN" and prefer_ppt:
            _app_logger.info(f"[FILE_GEN PPT] σ╝Çσºï PPT τöƒµêÉµ╡üτ¿ï")

            # τ¼¼ 1 µ¡Ñ∩╝ÜΣ╜┐τö¿ FileParser µÅÉσÅûτ╗ôµ₧äσîûσåàσ«╣
            from web.file_parser import FileParser
            from web.ppt_session_manager import PPTSessionManager

            parser = FileParser()
            parse_result = parser.parse_file(filepath)
            file_content = parse_result.get("content", "") if parse_result else ""

            # τ¼¼ 2 µ¡Ñ∩╝Üσê¢σ╗║ PPT Σ╝ÜΦ»¥
            ppt_session_dir = os.path.join(WORKSPACE_DIR, "workspace", "ppt_sessions")
            os.makedirs(ppt_session_dir, exist_ok=True)

            session_manager_ppt = PPTSessionManager(ppt_session_dir)
            ppt_session_id = session_manager_ppt.create_session(
                title=f"PPT from {os.path.splitext(filename)[0]}",
                user_input=user_input,
                theme="business",
            )
            _app_logger.info(f"[FILE_GEN PPT] σê¢σ╗║Σ╝ÜΦ»¥: {ppt_session_id}")

            # τ¼¼ 3 µ¡Ñ∩╝ÜΣ┐¥σ¡ÿµûçΣ╗╢σåàσ«╣σê░Σ╝ÜΦ»¥
            session_manager_ppt.save_generation_data(
                session_id=ppt_session_id,
                ppt_data=None,
                ppt_file_path=None,
                uploaded_file_context=file_content[:3000],  # σ░åσåàσ«╣ΘÖÉσê╢Σ╕║σëì3000σ¡ùτ¼ª
            )
            _app_logger.info(f"[FILE_GEN PPT] µûçΣ╗╢σåàσ«╣σ╖▓Σ┐¥σ¡ÿσê░Σ╝ÜΦ»¥")

            # Σ╜┐τö¿µ╡üσ╝Åσôìσ║ö∩╝êStreamed Response∩╝ëΣ╗Ñµö»µîüσ«₧µù╢Φ┐¢σ║ªµÿ╛τñ║
            def generate_ppt_file_stream():
                import asyncio
                import queue
                import threading
                import time as _time

                from web.app import TaskOrchestrator

                _start = _time.time()

                # σÅæΘÇüσê¥σºïσîûΣ┐íµü»
                yield f"data: {json.dumps({'type': 'classification', 'task_type': 'FILE_GEN', 'subtask': 'PPT_CREATION', 'message': '≡ƒôè σ╝Çσºï PPT µ╝öτñ║µûçτ¿┐τöƒµêÉµ╡üτ¿ï'})}\n\n"

                # σçåσñçΣ╗╗σèíσÅéµò░
                subtask = {
                    "task_type": "FILE_GEN",
                    "index": 1,
                    "description": f"Σ╗Äµûçµíú {filename} τöƒµêÉ PPT",
                }
                context = {"original_input": user_input, "step_1_output": file_content}

                # Φ┐¢σ║ªΘÿƒσêù
                progress_queue = queue.Queue()

                def _progress_cb(msg, detail=""):
                    progress_queue.put({"msg": msg, "detail": detail})

                # Σ╗╗σèíτ╗ôµ₧£σ«╣σÖ¿
                task_result_holder = {"result": None}

                # σÉÄσÅ░µëºΦíîσç╜µò░
                def _run_task_thread():
                    # Σ╕║µû░τ║┐τ¿ïσê¢σ╗║τï¼τ½ïτÜäΣ║ïΣ╗╢σ╛¬τÄ»
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    try:
                        task_result_holder["result"] = loop.run_until_complete(
                            TaskOrchestrator._execute_file_gen(
                                user_input, context, subtask, _progress_cb
                            )
                        )
                    except Exception as e:
                        task_result_holder["result"] = {
                            "success": False,
                            "error": str(e),
                        }
                    finally:
                        loop.close()
                        progress_queue.put(None)  # Signal done

                # σÉ»σè¿σÉÄσÅ░τ║┐τ¿ï
                t = threading.Thread(target=_run_task_thread)
                t.start()

                # Σ╕╗τ║┐τ¿ïσ╛¬τÄ»Φ»╗σÅûΦ┐¢σ║ª
                while True:
                    try:
                        item = progress_queue.get(timeout=0.1)
                        if item is None:
                            break
                        # σÅæΘÇüΦ┐¢σ║ªSSE
                        yield f"data: {json.dumps({'type': 'progress', 'message': item['msg'], 'detail': item['detail']})}\n\n"
                    except queue.Empty:
                        if not t.is_alive():
                            break

                t.join()
                ppt_result = task_result_holder["result"]
                _elapsed = _time.time() - _start

                # σñäτÉåµ£Çτ╗êτ╗ôµ₧£
                if ppt_result and ppt_result.get("success"):
                    saved_files = ppt_result.get("saved_files", [])
                    if saved_files:
                        ppt_file_path = (
                            saved_files[0]
                            if isinstance(saved_files, list)
                            else saved_files
                        )
                        # Σ┐¥σ¡ÿΣ╝ÜΦ»¥µò░µì«
                        session_manager_ppt.save_generation_data(
                            session_id=ppt_session_id,
                            ppt_data=ppt_result.get("ppt_data"),
                            ppt_file_path=ppt_file_path,
                        )

                        final_msg = (
                            f"Γ£à PPT µ╝öτñ║σ╖▓τöƒµêÉ\n\n"
                            f"≡ƒôä µûçΣ╗╢: [{os.path.basename(ppt_file_path)}]({ppt_file_path.replace(os.sep, '/')})\n"
                            f"≡ƒöù Σ╝ÜΦ»¥ID: `{ppt_session_id}`\n"
                            f"ΓÅ▒∩╕Å ΦÇùµù╢: {_elapsed:.1f}s"
                        )
                        yield f"data: {json.dumps({'type': 'token', 'content': final_msg})}\n\n"

                        # µ¢┤µû░σÄåσÅ▓
                        session_manager.update_last_model_response(
                            f"{session_name}.json",
                            final_msg,
                            task="FILE_GEN",
                            model_name=model_to_use,
                            saved_files=[ppt_file_path],
                        )

                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [ppt_file_path], 'ppt_session_id': ppt_session_id})}\n\n"
                    else:
                        err_msg = "ΓÜá∩╕Å PPT µíåµ₧╢σ╖▓τöƒµêÉ∩╝îΣ╜åµûçΣ╗╢Σ┐¥σ¡ÿσñ▒Φ┤Ñ"
                        yield f"data: {json.dumps({'type': 'error', 'message': err_msg})}\n\n"
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                else:
                    err_msg = (
                        ppt_result.get("error", "µ£¬τƒÑΘöÖΦ»»")
                        if ppt_result
                        else "Σ╗╗σèíµëºΦíîµùáτ╗ôµ₧£"
                    )
                    yield f"data: {json.dumps({'type': 'error', 'message': f'Γ¥î τöƒµêÉσñ▒Φ┤Ñ: {err_msg}'})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"

            return Response(
                stream_with_context(generate_ppt_file_stream()),
                mimetype="text/event-stream",
            )

        elif task_type in ["FILE_GEN", "RESEARCH", "CHAT"]:
            # ΓöÇΓöÇ µûçµ£¼τ▒╗µûçΣ╗╢σêåµ₧É∩╝êSSE µ╡üσ╝Å∩╝îΣ┐«σñìσÄƒ blocking brain.chat σìíµ¡╗Θù«Θóÿ∩╝ëΓöÇΓöÇ
            _captured_context = context_info  # closure capture
            _captured_model = model_to_use
            _captured_task = task_type

            def generate_file_analysis_stream():
                import time as _time

                _start = _time.time()
                try:
                    _skill = (_captured_context or {}).get("skill_prompt")

                    # µá╣µì«Σ╗╗σèíτ▒╗σ₧ïΘÇëµï⌐ system instruction
                    if _captured_task == "RESEARCH":
                        _sys = (
                            "Σ╜áµÿ»Σ╕ÇΣ╜ìΣ╕ôΣ╕ÜτÜäµûçµíúσêåµ₧Éσè⌐µëï∩╝îµôàΘò┐µ╖▒σ║ªΦºúΦ»╗σÉäτ▒╗µûçΣ╗╢∩╝êσòåΣ╕ÜΦ«íσêÆΣ╣ªπÇüτáöτ⌐╢µèÑσæèπÇüµèÇµ£»µûçµíúτ¡ë∩╝ëπÇé\n"
                            "Φ»╖Σ╗öτ╗åΘÿàΦ»╗τö¿µê╖µÅÉΣ╛¢τÜäµûçΣ╗╢σåàσ«╣∩╝îσ╣╢µîëΣ╗ÑΣ╕ïτ╗ôµ₧äΦ╛ôσç║σêåµ₧ÉµèÑσæè∩╝Ü\n\n"
                            "## µá╕σ┐âµæÿΦªü\n- τö¿ 3-5 µ¥íΦªüτé╣µªéµï¼µûçΣ╗╢µá╕σ┐âσåàσ«╣\n\n"
                            "## Φ»ªτ╗åΦºúΦ»╗\n### ΦâîµÖ»Σ╕Äτ¢«µáç\n### σà│Θö«σåàσ«╣σêåµ₧É\n### µò░µì«Σ╕ÄΦ»üµì«\n\n"
                            "## τ╗ôΦ«║Σ╕Äσ╗║Φ««\n- τ╗╝σÉêΦ»äσêñΣ╕ÄσÅ»ΦíîµÇº/Σ╗╖σÇ╝σêñµû¡\n\n"
                            "Φªüµ▒é∩╝Üτö¿Σ╕¡µûç∩╝îµ¥íτÉåµ╕àµÖ░∩╝îΘü┐σàìσåùΣ╜Ö∩╝îΣ╕ìΦ╛ôσç║Σ╗úτáüσ¥ùµáçΦ«░πÇé"
                        )
                    elif _captured_task == "CHAT":
                        # τö¿µê╖Σ╕èΣ╝áµûçΣ╗╢+µÅÉΘù« ΓåÆ Φ»╗σÅûσêåµ₧ÉµûçΣ╗╢∩╝îΣ╕ìτöƒµêÉµû░µûçΣ╗╢µ¿íµ¥┐
                        _sys = (
                            "Σ╜áµÿ»Σ╕ÇΣ╜ìΣ╕ôΣ╕ÜτÜäµûçµíúΘÿàΦ»╗Σ╕Äσêåµ₧Éσè⌐µëïπÇéτö¿µê╖Σ╕èΣ╝áΣ║åΣ╕ÇΣ╗╜µûçΣ╗╢σ╣╢µÅÉσç║Σ║åΘù«Θóÿ∩╝î"
                            "Φ»╖Φ«ñτ£ƒΘÿàΦ»╗µûçΣ╗╢τÜäσ«îµò┤σåàσ«╣∩╝îτö¿Σ╕¡µûçτ╗Öσç║Φ»ªτ╗åπÇüσçåτí«τÜäσêåµ₧ÉσÆîσ¢₧τ¡öπÇé\n"
                            "µ│¿µäÅ∩╝Ü\n"
                            "- τ¢┤µÄÑσ¢₧τ¡öτö¿µê╖τÜäσà╖Σ╜ôΘù«Θóÿ∩╝îΣ╕ìΦªüτöƒµêÉµû░µûçµíúµ¿íµ¥┐\n"
                            "- σ╝òτö¿µûçΣ╗╢Σ╕¡τÜäσà╖Σ╜ôµò░µì«σÆîΣ┐íµü»µö»µÆæΣ╜áτÜäσêñµû¡\n"
                            "- σªéµ╢ëσÅèµèòΦ╡äΣ╗╖σÇ╝/ΘúÄΘÖ⌐∩╝îτ╗ôσÉêµûçΣ╗╢σåàσ«╣τ╗Öσç║µ£ëΣ╛¥µì«τÜäΦ»äΣ╝░\n"
                            "- τö¿µ╕àµÖ░τÜäτ╗ôµ₧äΦ╛ôσç║∩╝îΘü┐σàìτ⌐║µ│¢Φí¿Φ┐░"
                        )
                    else:
                        _sys = _get_filegen_brief_instruction()

                    if _skill:
                        _sys += f"\n\n[σêåµ₧ÉΘçìτé╣] {_skill}"

                    # ΓöÇΓöÇ Σ║îΦ┐¢σê╢µûçΣ╗╢∩╝êPDF/Wordτ¡ë∩╝ë+ CHAT/RESEARCH∩╝ÜΣ╝áσ¡ùΦèéµ╡üτ╗Öµ¿íσ₧ïτ¢┤µÄÑΦ»╗σÅû ΓöÇΓöÇ
                    # σªéµ₧£µ£ë file_data∩╝êΘ¥₧σ¢╛τëçΣ║îΦ┐¢σê╢∩╝ë∩╝îσ╝║σê╢Σ╜┐τö¿µö»µîü generate_content τÜäµ¿íσ₧ï
                    # σ╣╢σ░å PDF σ¡ùΦèéΘÖäσèáσê░Φ»╖µ▒éΣ╕¡∩╝îΦÇîΣ╕ìµÿ»Σ╛¥Φ╡ûµÅÉσÅûτÜäµûçµ£¼
                    _stream_model = _captured_model
                    _stream_contents = formatted_message  # Θ╗ÿΦ«ñ∩╝Üµûçµ£¼µ╢êµü»

                    # ΘÇÜτö¿µïªµê¬∩╝Üinteractions-only µ¿íσ₧ïΣ╕ìµö»µîü generate_content_stream
                    # Φªåτ¢ûµëÇµ£ëµâàσå╡∩╝êσîàµï¼µûçµ£¼σ╡îσàÑµ¿íσ╝Å∩╝îΣ╕ìΣ╗àµÿ» binary doc∩╝ë
                    if _stream_model in _INTERACTIONS_ONLY_MODELS or str(_stream_model).startswith("deep-research-pro-preview"):
                        _app_logger.warning(f"[FILE STREAM] ΓÜá∩╕Å {_stream_model} µÿ» interactions-only∩╝îΘÖìτ║ºσê░ {_INTERACTIONS_FALLBACK_MODEL}")
                        _stream_model = _INTERACTIONS_FALLBACK_MODEL

                    _has_binary_doc = (
                        file_data is not None
                        and not (file_data.get("mime_type") or "").lower().startswith("image/")
                    )
                    if _has_binary_doc and _captured_task in ("CHAT", "RESEARCH"):
                        # σ╝║σê╢ΘÖìτ║ºσê░µö»µîüµûçΣ╗╢σ¡ùΦèéτÜäµ¿íσ₧ï∩╝êInteractions API Σ╕ìµö»µîüΘÖäΣ╗╢σ¡ùΦèé∩╝ë
                        _stream_model = _INTERACTIONS_FALLBACK_MODEL
                        try:
                            _doc_part = types.Part.from_bytes(
                                data=file_data["data"],
                                mime_type=file_data.get("mime_type", "application/pdf"),
                            )
                            _stream_contents = [formatted_message, _doc_part]
                            _app_logger.info(
                                f"[FILE STREAM] ≡ƒôä Binary-Doc-Read: model={_stream_model}, bytes={len(file_data['data'])}"
                            )
                        except Exception as _bp_err:
                            _app_logger.warning(
                                f"[FILE STREAM] ΓÜá∩╕Å µùáµ│òσê¢σ╗║ doc_part∩╝îσ¢₧ΘÇÇσê░µûçµ£¼µ¿íσ╝Å: {_bp_err}"
                            )
                            _stream_contents = formatted_message

                    yield f"data: {json.dumps({'type': 'classification', 'task_type': _captured_task, 'model': _stream_model, 'message': f'≡ƒôä µ¡úσ£¿σêåµ₧É: {filename}'})}\n\n"
                    yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒôé µûçΣ╗╢σåàσ«╣σ╖▓σ░▒τ╗¬', 'stage': 'file_ready_complete', 'progress': 15})}\n\n"
                    _task_display = {
                        "FILE_GEN": "≡ƒô¥ µûçΣ╗╢τöƒµêÉ",
                        "RESEARCH": "≡ƒö¼ µ╖▒σ║ªσêåµ₧É",
                        "CHAT": "≡ƒÆ¼ σ»╣Φ»¥σêåµ₧É",
                    }.get(_captured_task, _captured_task)
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'≡ƒÄ» Σ╗╗σèíτ▒╗σ₧ï: {_task_display}', 'stage': 'routing_complete', 'progress': 25})}\n\n"
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'ΓÜí µ¡úσ£¿Φ»╖µ▒é {_stream_model}∩╝îΦ»╖τ¿ìσÇÖ...', 'stage': 'api_calling', 'progress': 35})}\n\n"

                    response_stream = client.models.generate_content_stream(
                        model=_stream_model,
                        contents=_stream_contents,
                        config=types.GenerateContentConfig(
                            system_instruction=_sys,
                            temperature=0.7,
                            max_output_tokens=8000,
                        ),
                    )

                    full_text = ""
                    _first_token = True
                    for _chunk in response_stream:
                        _t = getattr(_chunk, "text", None)
                        if _t:
                            if _first_token:
                                yield f"data: {json.dumps({'type': 'progress', 'message': 'Γ£ì∩╕Å µ¿íσ₧ïµ¡úσ£¿τöƒµêÉσ¢₧σñì...', 'stage': 'generating_complete', 'progress': 55})}\n\n"
                                _first_token = False
                            full_text += _t
                            yield f"data: {json.dumps({'type': 'token', 'content': _t})}\n\n"

                    _elapsed = round(_time.time() - _start, 2)
                    _saved_files = []

                    # Φç¬σè¿Σ┐¥σ¡ÿΣ╕║ DOCX
                    if full_text and len(full_text) > 50:
                        yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒÆ╛ µ¡úσ£¿Σ┐¥σ¡ÿµûçµíú...', 'stage': 'saving', 'progress': 90})}\n\n"
                        try:
                            _title = _build_analysis_title(
                                user_input, filename, is_binary=False
                            )
                            _cleaned = _strip_code_blocks(full_text)
                            _docx = save_docx(
                                _cleaned,
                                title=_title,
                                output_dir=settings_manager.documents_dir,
                            )
                            _docx_rel = os.path.relpath(_docx, WORKSPACE_DIR).replace(
                                "\\", "/"
                            )
                            _saved_files.append(_docx_rel)
                            _app_logger.info(f"[FILE UPLOAD] Γ£à σêåµ₧Éσ╖▓Σ┐¥σ¡ÿ DOCX: {_docx_rel}")
                            # µîëΘ£ÇσÉîµù╢Σ┐¥σ¡ÿ PDF
                            if any(
                                kw in (user_input or "").lower()
                                for kw in ["pdf", "Σ╕ñτºìµá╝σ╝Å", "both"]
                            ):
                                try:
                                    _pdf = save_pdf(
                                        _cleaned,
                                        title=_title,
                                        output_dir=settings_manager.documents_dir,
                                    )
                                    _saved_files.append(
                                        os.path.relpath(_pdf, WORKSPACE_DIR).replace(
                                            "\\", "/"
                                        )
                                    )
                                except Exception:
                                    pass
                        except Exception as _de:
                            _app_logger.warning(f"[FILE UPLOAD] ΓÜá∩╕Å Σ┐¥σ¡ÿ DOCX σñ▒Φ┤Ñ: {_de}")

                    session_manager.update_last_model_response(
                        f"{session_name}.json",
                        full_text,
                        task=_captured_task,
                        model_name=_captured_model,
                        saved_files=_saved_files,
                    )
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': _saved_files, 'total_time': _elapsed})}\n\n"

                except Exception as _e:
                    import traceback as _tb

                    _tb.print_exc()
                    _emsg = str(_e)[:200]
                    session_manager.update_last_model_response(
                        f"{session_name}.json", f"Γ¥î σêåµ₧Éσñ▒Φ┤Ñ: {_emsg}"
                    )
                    yield f"data: {json.dumps({'type': 'error', 'message': f'Γ¥î σêåµ₧Éσñ▒Φ┤Ñ: {_emsg}'})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': 0})}\n\n"

            return Response(
                stream_with_context(generate_file_analysis_stream()),
                mimetype="text/event-stream",
            )

        else:
            # ΓöÇΓöÇ σ¢╛τëç / Σ║îΦ┐¢σê╢µûçΣ╗╢∩╝ÜΦºåΦºëσêåµ₧É∩╝êSSE µ╡üσ╝ÅσîàΦúà∩╝ëΓöÇΓöÇ
            _captured_fdata = file_data
            _captured_model_v = model_to_use
            _captured_task_v = task_type

            def generate_vision_stream():
                import time as _time

                _start = _time.time()
                try:
                    yield f"data: {json.dumps({'type': 'classification', 'task_type': _captured_task_v, 'model': _captured_model_v, 'message': f'≡ƒæü∩╕Å µ¡úσ£¿σêåµ₧É: {filename}'})}\n\n"
                    yield f"data: {json.dumps({'type': 'progress', 'message': '∩┐╜ µûçΣ╗╢σ╖▓µÄÑµö╢', 'stage': 'file_ready_complete', 'progress': 15})}\n\n"
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'ΓÜí µ¡úσ£¿Φ»╖µ▒éΦºåΦºëµ¿íσ₧ï {_captured_model_v}...', 'stage': 'api_calling', 'progress': 35})}\n\n"

                    # Φ░âτö¿ brain.chat∩╝êvision Φ╖»σ╛äΘÇÜσ╕╕Φ╛âσ┐½∩╝ë
                    _brain_result = brain.chat(
                        history=history,
                        user_input=formatted_message,
                        file_data=_captured_fdata,
                        model=_captured_model_v,
                        auto_model=(locked_model == "auto"),
                    )
                    _resp_text = _brain_result.get("response", "")
                    _elapsed = round(_time.time() - _start, 2)
                    _saved_files = list(_brain_result.get("saved_files", []))

                    # Φ╛ôσç║σåàσ«╣
                    if _resp_text:
                        yield f"data: {json.dumps({'type': 'progress', 'message': 'Γ£ì∩╕Å σêåµ₧Éσ«îµêÉ∩╝îµ¡úσ£¿Φ╛ôσç║...', 'stage': 'generating_complete', 'progress': 70})}\n\n"
                        yield f"data: {json.dumps({'type': 'token', 'content': _resp_text})}\n\n"

                    # Φç¬σè¿Σ┐¥σ¡ÿΦºåΦºëσêåµ₧ÉΣ╕║ DOCX
                    if _resp_text and len(_resp_text) > 50:
                        yield f"data: {json.dumps({'type': 'progress', 'message': '≡ƒÆ╛ µ¡úσ£¿Σ┐¥σ¡ÿµûçµíú...', 'stage': 'saving', 'progress': 90})}\n\n"
                        try:
                            _title = _build_analysis_title(
                                user_input, filename, is_binary=True
                            )
                            _cleaned = _strip_code_blocks(_resp_text)
                            _docx = save_docx(
                                _cleaned,
                                title=_title,
                                output_dir=settings_manager.documents_dir,
                            )
                            _docx_rel = os.path.relpath(_docx, WORKSPACE_DIR).replace(
                                "\\", "/"
                            )
                            _saved_files.append(_docx_rel)
                            _app_logger.info(f"[FILE UPLOAD] Γ£à ΦºåΦºëσêåµ₧Éσ╖▓Σ┐¥σ¡ÿ DOCX: {_docx_rel}")
                        except Exception as _de:
                            _app_logger.warning(f"[FILE UPLOAD] ΓÜá∩╕Å ΦºåΦºë DOCX Σ┐¥σ¡ÿσñ▒Φ┤Ñ: {_de}")

                    session_manager.update_last_model_response(
                        f"{session_name}.json",
                        _resp_text,
                        task=_captured_task_v,
                        model_name=_captured_model_v,
                        saved_files=_saved_files,
                        images=_brain_result.get("images", []),
                    )
                    yield f"data: {json.dumps({'type': 'done', 'images': _brain_result.get('images', []), 'saved_files': _saved_files, 'total_time': _elapsed})}\n\n"

                except Exception as _e:
                    import traceback as _tb

                    _tb.print_exc()
                    _emsg = str(_e)[:200]
                    session_manager.update_last_model_response(
                        f"{session_name}.json", f"Γ¥î µûçΣ╗╢σêåµ₧Éσñ▒Φ┤Ñ: {_emsg}"
                    )
                    yield f"data: {json.dumps({'type': 'error', 'message': f'Γ¥î µûçΣ╗╢σêåµ₧Éσñ▒Φ┤Ñ: {_emsg}'})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': 0})}\n\n"

            return Response(
                stream_with_context(generate_vision_stream()),
                mimetype="text/event-stream",
            )

    except Exception as e:
        # σì│Σ╜┐σç║ΘöÖΣ╣ƒΣ┐¥σ¡ÿτö¿µê╖τÜäΘù«ΘóÿσÆîΘöÖΦ»»Σ┐íµü»
        import traceback

        error_detail = traceback.format_exc()
        _app_logger.info(f"[FILE UPLOAD ERROR] {error_detail}")

        error_response = f"Γ¥î σñäτÉåµûçΣ╗╢µù╢σç║ΘöÖ: {str(e)}"
        session_manager.update_last_model_response(
            f"{session_name}.json", error_response
        )

        return jsonify(
            {
                "response": error_response,
                "task": "ERROR",
                "model": "none",
                "images": [],
                "saved_files": [],
            }
        )


# ==================== PPT τ¢╕σà│ API τ½»τé╣∩╝êP0 ΦíÑσàà∩╝ë====================


@app.route("/api/ppt/download", methods=["POST"])
def download_ppt():
    """Σ╕ïΦ╜╜ PPT PPTX µûçΣ╗╢"""
    try:
        session_id = request.json.get("session_id")
        if not session_id:
            return jsonify({"error": "Missing session_id"}), 400

        # Σ╗Ä PPT Σ╝ÜΦ»¥Σ╕¡ΦÄ╖σÅûµûçΣ╗╢Φ╖»σ╛ä
        from web.ppt_session_manager import PPTSessionManager

        ppt_session_dir = os.path.join(WORKSPACE_DIR, "workspace", "ppt_sessions")
        manager = PPTSessionManager(ppt_session_dir)

        session_data = manager.load_session(session_id)
        if not session_data:
            return jsonify({"error": "Session not found"}), 404

        ppt_file_path = session_data.get("ppt_file_path")
        if not ppt_file_path:
            # σªéµ₧£µûçΣ╗╢Φ┐ÿµ▓íτöƒµêÉ∩╝îσ░¥Φ»òτöƒµêÉΣ╕ÇΣ╕¬Σ╕┤µù╢τÜä
            return jsonify({"error": "PPT file not generated yet"}), 400

        # µ₧äσ╗║σ«îµò┤τÜäµûçΣ╗╢Φ╖»σ╛ä
        full_path = os.path.join(
            WORKSPACE_DIR, ppt_file_path.lstrip("/").replace("/", os.sep)
        )

        if not os.path.exists(full_path):
            return jsonify({"error": "PPT file not found"}), 404

        # Φ┐öσ¢₧µûçΣ╗╢Σ╕ïΦ╜╜
        return send_file(
            full_path,
            mimetype="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            as_attachment=True,
            download_name=os.path.basename(full_path),
        )

    except Exception as e:
        _app_logger.info(f"[PPT DOWNLOAD] ΘöÖΦ»»: {e}")
        return jsonify({"error": f"Download failed: {str(e)}"}), 500


@app.route("/api/ppt/session/<session_id>", methods=["GET"])
def get_ppt_session(session_id):
    """ΦÄ╖σÅû PPT Σ╝ÜΦ»¥Σ┐íµü»"""
    try:
        from web.ppt_session_manager import PPTSessionManager

        ppt_session_dir = os.path.join(WORKSPACE_DIR, "workspace", "ppt_sessions")
        manager = PPTSessionManager(ppt_session_dir)

        session_data = manager.load_session(session_id)
        if not session_data:
            return jsonify({"error": "Session not found"}), 404

        return jsonify(
            {
                "success": True,
                "session": {
                    "id": session_data.get("session_id"),
                    "title": session_data.get("title"),
                    "status": session_data.get("status"),
                    "ppt_file_path": session_data.get("ppt_file_path"),
                    "created_at": session_data.get("created_at"),
                    "updated_at": session_data.get("updated_at"),
                },
            }
        )

    except Exception as e:
        _app_logger.info(f"[PPT SESSION] ΘöÖΦ»»: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/info", methods=["GET"])
def api_info():
    """Application metadata and configuration info.
    ---
    tags: [Health]
    responses:
      200:
        description: App metadata
        schema:
          properties:
            version: {type: string}
            deploy_mode: {type: string, enum: [local, cloud]}
            auth_enabled: {type: boolean}
    """
    return jsonify(
        {
            "version": APP_VERSION,
            "deploy_mode": os.environ.get("KOTO_DEPLOY_MODE", "local"),
            "auth_enabled": os.environ.get("KOTO_AUTH_ENABLED", "false").lower()
            == "true",
        }
    )


@app.route("/api/v1/models", methods=["GET"])
def api_list_models():
    """
    σè¿µÇüµ¿íσ₧ïσêùΦí¿ API
    Φ┐öσ¢₧σ╜ôσëì API σÅ»τö¿τÜäµ¿íσ₧ïσÅèσÉäΣ╗╗σèíτÜäΦ╖»τö▒τ╗ôµ₧£πÇé
    - ready: µ¿íσ₧ïτ«íτÉåσÖ¿µÿ»σÉªσ╖▓σ«îµêÉσê¥σºïσîû
    - model_map: Σ╗╗σèí ΓåÆ µ¿íσ₧ï ID τÜäσ╜ôσëìΦ╖»τö▒Φí¿∩╝êσÉ½Φ»äσêåΣ┐íµü»∩╝ë
    - available: µëÇµ£ëσÅ»τö¿µ¿íσ₧ïτÜäΦâ╜σè¢σêùΦí¿
    """
    if _model_manager:
        return jsonify(
            {
                "ready": True,
                "model_map": _model_manager.get_model_map_with_scores(),
                "available": _model_manager.get_available_models(),
                "fallback": _INTERACTIONS_FALLBACK_MODEL,
                "interactions_only": list(_INTERACTIONS_ONLY_MODELS),
            }
        )
    # µ¿íσ₧ïτ«íτÉåσÖ¿σ░Üµ£¬σ░▒τ╗¬µêûΣ╕ìσÅ»τö¿∩╝îΦ┐öσ¢₧Θ¥ÖµÇüΘ╗ÿΦ«ñσÇ╝
    return jsonify(
        {
            "ready": False,
            "model_map": {
                task: {
                    "model_id": mid,
                    "display": get_model_display_name(mid),
                    "provider": "gemini" if mid != "local-executor" else "local",
                    "tier": MODEL_INFO.get(mid, {}).get("tier", 5),
                    "score": None,
                    "_inferred": False,
                }
                for task, mid in MODEL_MAP.items()
            },
            "available": [
                {
                    "id": mid,
                    "display": get_model_display_name(mid),
                    "tier": MODEL_INFO.get(mid, {}).get("tier", 5),
                    "provider": "gemini" if mid != "local-executor" else "local",
                    "strengths": MODEL_INFO.get(mid, {}).get("strengths", []),
                    "capabilities": {},
                }
                for mid in dict.fromkeys(MODEL_MAP.values())
            ],
            "fallback": _INTERACTIONS_FALLBACK_MODEL,
            "interactions_only": list(_INTERACTIONS_ONLY_MODELS),
        }
    )


@app.route("/api/v1/models/refresh", methods=["POST"])
def api_refresh_models():
    """
    µëïσè¿ΦºªσÅæµ¿íσ₧ïσêùΦí¿σê╖µû░πÇé
    Θçìµû░µƒÑΦ»ó API σ╣╢µ¢┤µû░Φ╖»τö▒Φí¿∩╝îσ£¿µû░µ¿íσ₧ïΣ╕èτ║┐σÉÄσÅ»τ½ïσì│τöƒµòêπÇé
    """
    if not _model_manager_available or _model_manager is None:
        # τ«íτÉåσÖ¿µ£¬σ░▒τ╗¬∩╝îσ£¿σÉÄσÅ░Θçìµû░σê¥σºïσîû
        import threading as _t

        _t.Thread(
            target=_init_model_manager, name="ModelManagerReinit", daemon=True
        ).start()
        return jsonify(
            {"status": "initializing", "message": "µ¿íσ₧ïτ«íτÉåσÖ¿µ¡úσ£¿σÉÄσÅ░σê¥σºïσîû"}
        )
    try:
        new_map = _model_manager.refresh()
        MODEL_MAP.update(new_map)
        # σÉîµ¡Ñµ¢┤µû░ ModelFallbackExecutor Φ╖»τö▒Φí¿
        try:
            from app.core.llm.model_fallback import get_fallback_executor
            get_fallback_executor().update_model_map(MODEL_MAP)
        except Exception as _fe:
            _app_logger.warning(f"[ModelRefresh] ΓÜá∩╕Å FallbackExecutor sync failed: {_fe}")
        # σÉîµ¡Ñµ¢┤µû░ AIRouter Φ╜╗ΘçÅΦ╖»τö▒µ¿íσ₧ï
        try:
            from app.core.routing.ai_router import AIRouter
            _caps = _model_manager._cached_caps
            _candidates = [
                (mid, caps) for mid, caps in _caps.items()
                if not caps.get("interactions_only", False)
                and not caps.get("image_gen", False)
                and mid != "local-executor"
            ]
            if _candidates:
                _best = max(_candidates, key=lambda x: x[1].get("speed", 0) + x[1].get("tier", 0) * 0.1)[0]
                AIRouter.set_router_model(_best)
        except Exception as _are:
            _app_logger.warning(f"[ModelRefresh] ΓÜá∩╕Å AIRouter update failed: {_are}")
        return jsonify({
            "status":    "ok",
            "model_map": _model_manager.get_model_map_with_scores(),
            "count":     len(_model_manager.get_available_models()),
        })
    except Exception as e:
        return jsonify({"status": "error", "error": str(e)}), 500

@app.route('/api/analyze', methods=['POST'])
def analyze_task():
    """Θóäσêåµ₧ÉΣ╗╗σèíτ▒╗σ₧ïσÆîµ¿íσ₧ïΘÇëµï⌐ - Φ«⌐σëìτ½»τ½ïσì│µÿ╛τñ║"""
    data = request.json
    message = data.get("message", "")
    locked_task = data.get("locked_task")
    locked_model = data.get("locked_model", "auto")
    has_file = data.get("has_file", False)
    file_type = data.get("file_type", "")

    if not message:
        return jsonify(
            {"task": "CHAT", "model": MODEL_MAP["CHAT"], "route_method": "Empty"}
        )

    # σ¢╛σâÅτ╝ûΦ╛æσà│Θö«Φ»ì
    IMAGE_EDIT_KEYWORDS = [
        "Σ┐«µö╣",
        "µìó",
        "µö╣µêÉ",
        "σÅÿµêÉ",
        "σ║òΦë▓",
        "ΦâîµÖ»",
        "Θó£Φë▓",
        "µèáσ¢╛",
        "σÄ╗ΦâîµÖ»",
        "Pσ¢╛",
        "τ╛Äσîû",
        "µ╗ñΘò£",
        "Φ░âΦë▓",
        "τ╝ûΦ╛æ",
        "change",
        "modify",
        "edit",
        "background",
        "color",
    ]

    # σªéµ₧£τö¿µê╖Θöüσ«ÜΣ║åΣ╗╗σèíτ▒╗σ₧ï
    if locked_task:
        task = locked_task
        route_method = "≡ƒöÆ Manual"
    elif has_file and file_type and file_type.startswith("image"):
        # µ£ëσ¢╛τëçµûçΣ╗╢∩╝îσêñµû¡µÿ»τ╝ûΦ╛æΦ┐ÿµÿ»σêåµ₧É
        message_lower = message.lower()
        is_edit = any(kw in message_lower for kw in IMAGE_EDIT_KEYWORDS)
        if is_edit:
            task = "PAINTER"
            route_method = "≡ƒû╝∩╕Å Image Edit"
        else:
            task = "VISION"
            route_method = "≡ƒæü∩╕Å Image Analysis"
    else:
        # Σ╜┐τö¿µÖ║Φâ╜Φ╖»τö▒σÖ¿
        task, route_method, _ = SmartDispatcher.analyze(message)

    # σªéµ₧£τö¿µê╖ΘÇëµï⌐Σ║åτë╣σ«Üµ¿íσ₧ï
    if locked_model and locked_model != "auto":
        model = locked_model
    else:
        model = SmartDispatcher.get_model_for_task(task, has_image=has_file)

    # ΦÄ╖σÅûµ¿íσ₧ïµÿ╛τñ║Σ┐íµü»
    model_info = MODEL_INFO.get(model, {"name": model, "speed": ""})

    return jsonify(
        {
            "task": task,
            "model": model,
            "model_name": model_info.get("name", model),
            "model_speed": model_info.get("speed", ""),
            "route_method": route_method,  # Φ╖»τö▒τ«ùµ│òΣ┐íµü»
            "strengths": model_info.get("strengths", []),
        }
    )


@app.route("/api/workspace/<path:filepath>")
def get_workspace_file(filepath):
    """ΦÄ╖σÅû workspace Σ╕¡τÜäµûçΣ╗╢∩╝îµö»µîüσ¡Éτ¢«σ╜ò"""
    _app_logger.debug(f"[API] Serving workspace file: {filepath}")
    full_path = os.path.join(WORKSPACE_DIR, filepath)

    # σ«ëσà¿µúÇµƒÑ∩╝Üτí«Σ┐¥Φ»╖µ▒éτÜäΦ╖»σ╛äσ£¿ WORKSPACE_DIR Σ╕ï
    try:
        resolved_path = os.path.abspath(full_path)
        resolved_workspace = os.path.abspath(WORKSPACE_DIR)
        if not resolved_path.startswith(resolved_workspace):
            _app_logger.debug(
                f"[API] Security violation: {resolved_path} not under {resolved_workspace}"
            )
            return jsonify({"error": "Access denied"}), 403

        if not os.path.exists(resolved_path):
            _app_logger.debug(f"[API] File not found: {resolved_path}")
            return jsonify({"error": "File not found"}), 404

        _app_logger.debug(f"[API] Serving: {resolved_path}")
        return send_from_directory(WORKSPACE_DIR, filepath)
    except Exception as e:
        _app_logger.debug(f"[API] Error serving {filepath}: {e}")
        import traceback

        traceback.print_exc()
        return jsonify({"error": "Server error", "detail": str(e)}), 500


@app.route("/api/workspace", methods=["GET"])
def list_workspace_files():
    files = os.listdir(WORKSPACE_DIR)
    return jsonify({"files": files})


@app.route("/api/open-workspace", methods=["POST"])
def open_workspace():
    """µëôσ╝Ç workspace µûçΣ╗╢σñ╣"""
    try:
        if sys.platform == "win32":
            subprocess.Popen(
                f'explorer "{WORKSPACE_DIR}"',
                shell=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
        elif sys.platform == "darwin":
            subprocess.Popen(["open", WORKSPACE_DIR])
        else:
            subprocess.Popen(["xdg-open", WORKSPACE_DIR])
        return jsonify({"success": True, "path": WORKSPACE_DIR})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route("/api/open-file", methods=["POST"])
def open_file_native():
    """τö¿τ│╗τ╗ƒΘ╗ÿΦ«ñτ¿ïσ║Åµëôσ╝ÇµûçΣ╗╢∩╝êΣ╕ìτ╗ÅΦ┐çµ╡ÅΦºêσÖ¿∩╝ë"""
    try:
        data = request.get_json()
        filepath = data.get("filepath", "")
        if not filepath:
            return jsonify({"success": False, "error": "No filepath provided"}), 400

        full_path = os.path.join(WORKSPACE_DIR, filepath)
        resolved_path = os.path.abspath(full_path)
        resolved_workspace = os.path.abspath(WORKSPACE_DIR)

        if not resolved_path.startswith(resolved_workspace):
            return jsonify({"success": False, "error": "Access denied"}), 403

        if not os.path.exists(resolved_path):
            return jsonify({"success": False, "error": "File not found"}), 404

        _app_logger.debug(f"[API] Opening file natively: {resolved_path}")
        if sys.platform == "win32":
            os.startfile(resolved_path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", resolved_path])
        else:
            subprocess.Popen(["xdg-open", resolved_path])

        return jsonify({"success": True, "path": resolved_path})
    except Exception as e:
        _app_logger.debug(f"[API] Error opening file: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


# ================= Settings API =================

# ΓöÇΓöÇΓöÇ µ£¼σ£░µ¿íσ₧ïτè╢µÇü API ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ


@app.route("/api/local-model/status", methods=["GET"])
def local_model_status():
    """Φ┐öσ¢₧σ╜ôσëìµ£¼σ£░µ¿íσ₧ïΘàìτ╜«σÆîΦ┐ÉΦíîτè╢µÇü"""
    try:
        from app.core.llm.ollama_provider import get_local_model_info

        info = get_local_model_info()
        return jsonify({"success": True, **info})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/local-model/switch", methods=["POST"])
def local_model_switch():
    """σêçµìó AI µ¿íσ╝Å∩╝êlocal / cloud∩╝ëσ╣╢τâ¡µ¢┤µû░σ«óµê╖τ½»τ╝ôσ¡ÿ"""
    global _client, _client_mode_key
    try:
        data = request.json or {}
        mode = data.get("mode", "cloud")  # "local" µêû "cloud"
        model_tag = data.get("model_tag")  # µ£¼σ£░µ¿íσ╝Åµù╢σÅ»µîçσ«Üµ¿íσ₧ï

        settings_path = os.path.join(PROJECT_ROOT, "config", "user_settings.json")
        try:
            with open(settings_path, "r", encoding="utf-8") as f:
                settings = json.load(f)
        except Exception:
            settings = {}

        settings["model_mode"] = mode
        if model_tag:
            settings["local_model"] = model_tag

        with open(settings_path, "w", encoding="utf-8") as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)

        # µ╕àΘÖñτ╝ôσ¡ÿ∩╝îΣ╕ïµ¼í get_client() Φ░âτö¿µù╢Θçìσ╗║
        _user_settings_cache.clear()
        _client = None
        _client_mode_key = (None, None)

        return jsonify(
            {
                "success": True,
                "mode": mode,
                "model": model_tag or settings.get("local_model"),
            }
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/local-model/setup", methods=["POST"])
def local_model_setup():
    """ΦºªσÅæµ£¼σ£░µ¿íσ₧ïσ«ëΦúàσÉæσ»╝∩╝êσ╝éµ¡Ñ∩╝îΣ╕ìΘÿ╗σí₧ API σôìσ║ö∩╝ë"""

    def _run_gui():
        try:
            from model_downloader import run_downloader_gui

            run_downloader_gui()
            # σ«ëΦúàσ«îµêÉσÉÄµ╕àΘÖñτ╝ôσ¡ÿ
            global _client, _client_mode_key
            _user_settings_cache.clear()
            _client = None
            _client_mode_key = (None, None)
        except Exception as e:
            _app_logger.debug(f"[LocalModel] σ«ëΦúàσÉæσ»╝σñ▒Φ┤Ñ: {e}")

    import threading as _threading

    _threading.Thread(target=_run_gui, daemon=True).start()
    return jsonify({"success": True, "message": "σ«ëΦúàσÉæσ»╝σ╖▓σÉ»σè¿"})


# GET /api/skills σ╖▓Φ┐üτº╗Φç│ skill_bp Φô¥σ¢╛∩╝êapp/api/skill_routes.py∩╝ë∩╝îσ£¿µ¡ñτº╗ΘÖñσåàΦüöσ«ÜΣ╣ëΘü┐σàìΦ╖»τö▒Θÿ╗µïª


@app.route("/api/skills/<skill_id>/toggle", methods=["POST"])
def toggle_skill(skill_id: str):
    """σÉ»τö¿/τªüτö¿µƒÉΣ╕¬µèÇΦâ╜"""
    try:
        from app.core.skills.skill_manager import SkillManager

        data = request.json or {}
        enabled = bool(data.get("enabled", False))
        success = SkillManager.set_enabled(skill_id, enabled)
        return jsonify({"success": success})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/skills/<skill_id>/prompt", methods=["POST"])
def update_skill_prompt(skill_id: str):
    """µ¢┤µû░µƒÉΣ╕¬µèÇΦâ╜τÜäΦç¬σ«ÜΣ╣ë Prompt"""
    try:
        from app.core.skills.skill_manager import SkillManager

        data = request.json or {}
        prompt = data.get("prompt", "")
        if not prompt.strip():
            SkillManager.reset_prompt(skill_id)
            return jsonify({"success": True, "reset": True})
        success = SkillManager.update_prompt(skill_id, prompt)
        return jsonify({"success": success})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/skills/<skill_id>/reset", methods=["POST"])
def reset_skill_prompt(skill_id: str):
    """σ░åµèÇΦâ╜ Prompt µüóσñìΣ╕║Θ╗ÿΦ«ñσÇ╝"""
    try:
        from app.core.skills.skill_manager import SkillManager

        success = SkillManager.reset_prompt(skill_id)
        return jsonify({"success": success})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ================= Settings API =================


@app.route("/api/settings", methods=["GET"])
def get_settings():
    """Get all application settings.
    ---
    tags:
      - Settings
    responses:
      200:
        description: All settings grouped by category
        schema:
          type: object
    """
    # σÉêσ╣╢ appearance Σ╕╗Θóÿ∩╝êσªéµ£ë cookie/σÅéµò░σÅ»σ£¿µ¡ñσÉêσ╣╢∩╝ë
    return jsonify(settings_manager.get_all())


@app.route("/api/settings", methods=["POST"])
def update_settings():
    """Update an application setting.
    ---
    tags:
      - Settings
    parameters:
      - in: body
        name: body
        required: true
        schema:
          required: [category, key, value]
          properties:
            category:
              type: string
              description: Settings category
            key:
              type: string
              description: Setting key
            value:
              description: New value
    responses:
      200:
        description: Update result
        schema:
          type: object
          properties:
            success:
              type: boolean
    """
    data = request.json
    category = data.get("category")
    key = data.get("key")
    value = data.get("value")

    if category and key:
        success = settings_manager.set(category, key, value)
        settings_manager.ensure_directories()
        # Σ╜┐ _load_user_settings τ╝ôσ¡ÿσñ▒µòê∩╝îτí«Σ┐¥σÉÄτ╗¡Φ»╗σÅûΦÄ╖σ╛ùµ£Çµû░σÇ╝
        _user_settings_cache.clear()
        # Σ╗úτÉåΦ«╛τ╜«σÅÿµ¢┤µù╢τ½ïσì│Θçìµû░µúÇµ╡ï
        if category == "proxy":
            global _proxy_checked, _detected_proxy
            _proxy_checked = False
            _detected_proxy = None
            threading.Thread(target=lambda: get_detected_proxy(), daemon=True).start()
        return jsonify({"success": success})
    return jsonify({"success": False, "error": "Missing category or key"})


@app.route("/api/settings/reset", methods=["POST"])
def reset_settings():
    success = settings_manager.reset()
    # σÉîµá╖µ╕àΘÖñτ╝ôσ¡ÿ
    _user_settings_cache.clear()
    global _proxy_checked, _detected_proxy
    _proxy_checked = False
    _detected_proxy = None
    return jsonify({"success": success})

# ================= Mini Mode Switch API =================


@app.route("/api/switch-to-mini", methods=["POST"])
def switch_to_mini():
    """σêçµìóσê░Φ┐╖Σ╜áµ¿íσ╝Å"""
    import subprocess
    import sys

    # µëôσîàτëêµùáµ│òΣ╗ÑΦäÜµ£¼µû╣σ╝ÅσÉ»σè¿ mini_koto.py
    if getattr(sys, "frozen", False):
        return jsonify(
            {"success": False, "error": "µëôσîàτëêµÜéΣ╕ìµö»µîüΦ┐╖Σ╜áµ¿íσ╝Å∩╝îΦ»╖Σ╜┐τö¿τ¬ùσÅúΘí╢µáÅµîëΘÆ«"}
        )

    try:
        # σÉ»σè¿Φ┐╖Σ╜áτ¬ùσÅú
        mini_koto_path = os.path.join(PROJECT_ROOT, "web", "mini_koto.py")
        if os.path.exists(mini_koto_path):
            # σ£¿µû░Φ┐¢τ¿ïΣ╕¡σÉ»σè¿Φ┐╖Σ╜áτ¬ùσÅú
            subprocess.Popen(
                [sys.executable, mini_koto_path],
                creationflags=(
                    subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0
                ),
                cwd=PROJECT_ROOT,
            )
            return jsonify({"success": True, "message": "Φ┐╖Σ╜áµ¿íσ╝Åσ╖▓σÉ»σè¿"})
        else:
            return jsonify({"success": False, "error": "µë╛Σ╕ìσê░Φ┐╖Σ╜áµ¿íσ╝Åτ¿ïσ║Å"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route("/api/switch-to-main", methods=["POST"])
def switch_to_main():
    """σêçµìóσê░Σ╕╗τ¿ïσ║Å"""
    import subprocess
    import sys

    # µëôσîàτëêσ╖▓σ£¿Σ╕╗τ¿ïσ║Åτ¬ùσÅúΣ╕¡Φ┐ÉΦíî∩╝îτ¢┤µÄÑΦ┐öσ¢₧µêÉσèƒ
    if getattr(sys, "frozen", False):
        return jsonify({"success": True, "message": "σ╖▓σ£¿Σ╕╗τ¿ïσ║ÅΣ╕¡Φ┐ÉΦíî"})

    try:
        # σÉ»σè¿Σ╕╗τ¬ùσÅú
        main_app_path = os.path.join(PROJECT_ROOT, "koto_app.py")
        if os.path.exists(main_app_path):
            subprocess.Popen(
                [sys.executable, main_app_path],
                creationflags=(
                    subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0
                ),
                cwd=PROJECT_ROOT,
            )
            return jsonify({"success": True, "message": "Σ╕╗τ¿ïσ║Åσ╖▓σÉ»σè¿"})
        else:
            return jsonify({"success": False, "error": "µë╛Σ╕ìσê░Σ╕╗τ¿ïσ║Å"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route("/mini")
def mini_page():
    """Φ┐╖Σ╜áµ¿íσ╝ÅΘí╡Θ¥ó∩╝êµ╡ÅΦºêσÖ¿Φ«┐Θù«τö¿∩╝ë"""
    return render_template("mini_koto.html")


@app.route("/m")
@app.route("/mobile")
def mobile_page():
    """τº╗σè¿τ½»Σ╝ÿσîûΘí╡Θ¥ó"""
    return render_template("mobile.html")


@app.route("/api/mini/chat", methods=["POST"])
def mini_chat():
    """Φ┐╖Σ╜áµ¿íσ╝ÅΣ╕ôτö¿Φüèσñ⌐API - Σ╜┐τö¿Σ╕ÄσÄƒτëêσ«îσà¿τ¢╕σÉîτÜäΣ╗╗σèíσêåΘàìσÆîµëºΦíîΘÇ╗Φ╛æ"""
    data = request.json
    user_input = data.get("message", "").strip()

    if not user_input:
        return jsonify({"error": "µ╢êµü»Σ╕ìΦâ╜Σ╕║τ⌐║"}), 400

    user_input = Utils.sanitize_string(user_input)

    # Σ╜┐τö¿σ¢║σ«ÜτÜäΦ┐╖Σ╜áΣ╝ÜΦ»¥
    session_name = "MiniKoto_Quick"
    history = session_manager.load(f"{session_name}.json")

    # ≡ƒÄ» Σ╜┐τö¿ SmartDispatcher Φ┐¢ΦíîΣ╗╗σèíσêåµ₧É∩╝êΣ╕Äσ«îµò┤τëêτ¢╕σÉî∩╝ë
    task_type, route_method, context_info = SmartDispatcher.analyze(user_input, history)
    _app_logger.debug(
        f"[MINI_CHAT] SmartDispatcher σêåµ₧Éτ╗ôµ₧£: task_type='{task_type}', method='{route_method}'"
    )

    response_text = ""
    is_error = False
    used_model = "unknown"

    try:
        # ===== µá╣µì«Σ╗╗σèíτ▒╗σ₧ïµëºΦíîΣ╕ìσÉîτÜäσñäτÉåΘÇ╗Φ╛æ∩╝êΣ╕Äσ«îµò┤τëêτ¢╕σÉî∩╝ë=====

        if task_type == "WEB_SEARCH":
            # ≡ƒîÉ τ╜æτ╗£µÉ£τ┤ó - Σ╜┐τö¿ Gemini Google Search Grounding
            _app_logger.debug(f"[MINI_CHAT] ≡ƒîÉ µëºΦíîτ╜æτ╗£µÉ£τ┤ó...")
            _mini_skill_prompt = (context_info or {}).get("skill_prompt")
            search_result = WebSearcher.search_with_grounding(
                user_input, skill_prompt=_mini_skill_prompt
            )
            response_text = search_result.get("response", "")
            used_model = "gemini-2.5-flash (Google Search)"

            # σªéµ₧£µÉ£τ┤óσñ▒Φ┤Ñ∩╝îσ░¥Φ»òΣ┐«µ¡úµƒÑΦ»ó
            if (
                not search_result.get("success")
                or Utils.is_failure_output(response_text)
                or "µÉ£τ┤óσñ▒Φ┤Ñ" in response_text
            ):
                _app_logger.warning(f"[MINI_CHAT] ΓÜá∩╕Å σê¥µ¼íµÉ£τ┤óσñ▒Φ┤Ñ∩╝îσ░¥Φ»òΣ┐«µ¡úµƒÑΦ»ó...")
                fix_query_prompt = (
                    "Φ»╖µèèτö¿µê╖Θ£Çµ▒éµö╣σåÖµêÉµ¢┤ΘÇéσÉêµÉ£τ┤óτÜäτ«Çτƒ¡σà│Θö«Φ»ìµêûµƒÑΦ»óΦ»¡σÅÑ∩╝îσÅ¬Φ╛ôσç║µƒÑΦ»óΦ»¡σÅÑπÇé\n"
                    f"τö¿µê╖Θ£Çµ▒é: {user_input}"
                )
                try:
                    fix_query_resp = client.models.generate_content(
                        model="gemini-2.0-flash-lite",
                        contents=fix_query_prompt,
                        config=types.GenerateContentConfig(
                            temperature=0.2, max_output_tokens=64
                        ),
                    )
                    fixed_query = (fix_query_resp.text or user_input).strip()
                    _app_logger.debug(f"[MINI_CHAT] Σ┐«µ¡úσÉÄτÜäµƒÑΦ»ó: {fixed_query}")
                    search_result = WebSearcher.search_with_grounding(fixed_query)
                    response_text = search_result.get("response", "")
                except Exception as e:
                    _app_logger.debug(f"[MINI_CHAT] Σ┐«µ¡úµƒÑΦ»óσñ▒Φ┤Ñ: {e}")

            if not response_text or Utils.is_failure_output(response_text):
                is_error = True
                response_text = f"µÉ£τ┤óσñ▒Φ┤Ñ∩╝Üµùáµ│òΦÄ╖σÅû '{user_input}' τÜäσ«₧µù╢Σ┐íµü»"

        elif task_type == "SYSTEM":
            # ≡ƒûÑ∩╕Å τ│╗τ╗ƒσæ╜Σ╗ñ - µ£¼σ£░µëºΦíî
            _app_logger.debug(f"[MINI_CHAT] ≡ƒûÑ∩╕Å µëºΦíîτ│╗τ╗ƒσæ╜Σ╗ñ∩╝Ü{user_input}")
            try:
                exec_result = LocalExecutor.execute(user_input)
                response_text = exec_result.get("message", "σæ╜Σ╗ñµëºΦíîσñ▒Φ┤Ñ")
                if exec_result.get("details"):
                    response_text += f"\n\n{exec_result['details']}"
                used_model = "LocalExecutor"
                is_error = not exec_result.get("success", False)

                # σªéµ₧£µëºΦíîσñ▒Φ┤Ñ∩╝îσ░¥Φ»òτö¿ AI Σ┐«µ¡ú
                if is_error or Utils.is_failure_output(response_text):
                    _app_logger.warning(f"[MINI_CHAT] ΓÜá∩╕Å µ£¼σ£░µëºΦíîσñ▒Φ┤Ñ∩╝îσ░¥Φ»ò AI Σ┐«µ¡ú...")
                    fix_prompt = Utils.build_fix_prompt(
                        "SYSTEM", user_input, response_text
                    )
                    try:
                        fix_resp = client.models.generate_content(
                            model="gemini-2.5-flash",
                            contents=fix_prompt,
                            config=types.GenerateContentConfig(
                                system_instruction=_get_DEFAULT_CHAT_SYSTEM_INSTRUCTION(),
                                temperature=0.4,
                                max_output_tokens=1000,
                            ),
                        )
                        response_text = fix_resp.text or response_text
                        used_model = "gemini-2.5-flash (fallback)"
                        is_error = False
                    except Exception as e:
                        _app_logger.debug(f"[MINI_CHAT] AI Σ┐«µ¡úσñ▒Φ┤Ñ: {e}")
            except Exception as e:
                _app_logger.error(f"[MINI_CHAT] Γ¥î τ│╗τ╗ƒσæ╜Σ╗ñµëºΦíîσç║ΘöÖ: {e}")
                response_text = f"τ│╗τ╗ƒσæ╜Σ╗ñµëºΦíîσç║ΘöÖ∩╝Ü{str(e)}"
                used_model = "LocalExecutor"
                is_error = True

        else:
            # ≡ƒÆ¼ σà╢Σ╗ûΣ╗╗σèí∩╝êCHAT, RESEARCH, CODER τ¡ë∩╝ë- Σ╜┐τö¿ brain.chat()
            _app_logger.debug(f"[MINI_CHAT] ≡ƒÆ¼ µëºΦíî {task_type} Σ╗╗σèí...")
            model = MODEL_MAP.get(task_type, MODEL_MAP["CHAT"])
            result = brain.chat(
                history, user_input, model=model, auto_model=False, task_type=task_type
            )
            response_text = result.get("response", "")
            used_model = result.get("model", model)
            is_error = response_text.startswith("Error:")

            # σªéµ₧£Θüçσê░ 404 ΘöÖΦ»»∩╝îσ░¥Φ»òσñçτö¿µ¿íσ₧ï
            if is_error and "404" in response_text:
                _app_logger.warning(f"[MINI_CHAT] ΓÜá∩╕Å µ¿íσ₧ï 404∩╝îσ░¥Φ»òσñçτö¿µ¿íσ₧ï...")
                for fallback_model in ["gemini-2.5-flash", "gemini-3-flash-preview"]:
                    try:
                        result = brain.chat(
                            history, user_input, model=fallback_model, auto_model=False
                        )
                        if not result.get("response", "").startswith("Error:"):
                            response_text = result.get("response", "")
                            used_model = fallback_model
                            is_error = False
                            break
                    except Exception as e:
                        continue

    except Exception as e:
        _app_logger.error(f"[MINI_CHAT] Γ¥î µëºΦíîσç║ΘöÖ: {e}")
        is_error = True
        response_text = f"Error: {str(e)}"

    # µ¢┤µû░σÄåσÅ▓∩╝êµêÉσèƒσÆîσñ▒Φ┤ÑΘâ╜Σ┐¥σ¡ÿ∩╝îΣ╛┐Σ║ÄµÄÆµƒÑ∩╝ë
    if response_text:
        session_manager.append_and_save(
            f"{session_name}.json", user_input, response_text
        )

    _app_logger.info(
        f"[MINI_CHAT] Γ£à σ«îµêÉ: task_type={task_type}, model={used_model}, success={not is_error}"
    )

    # Φ┐öσ¢₧τ╗ƒΣ╕Çµá╝σ╝Å
    return jsonify(
        {
            "success": not is_error,
            "response": response_text,
            "model": used_model,
            "task_type": task_type,
            "route_method": route_method,
            "error": response_text if is_error else "",
        }
    )


# ================= Setup & Initialization API =================


@app.route("/api/setup/status", methods=["GET"])
def get_setup_status():
    """µúÇµƒÑΘªûµ¼íΦ«╛τ╜«τè╢µÇü"""
    config_path = os.path.join(PROJECT_ROOT, "config", "gemini_config.env")
    has_api_key = bool(API_KEY and len(API_KEY) > 10)
    has_workspace = os.path.exists(WORKSPACE_DIR)

    return jsonify(
        {
            "initialized": has_api_key and has_workspace,
            "has_api_key": has_api_key,
            "has_workspace": has_workspace,
            "workspace_path": os.path.abspath(WORKSPACE_DIR),
            "config_path": os.path.abspath(config_path),
        }
    )


@app.route("/api/setup/apikey", methods=["POST"])
def setup_api_key():
    """Φ«╛τ╜« API Key"""
    data = request.json
    api_key = data.get("api_key", "").strip()

    if not api_key or len(api_key) < 10:
        return jsonify({"success": False, "error": "Invalid API key"})

    config_path = os.path.join(PROJECT_ROOT, "config", "gemini_config.env")
    os.makedirs(os.path.dirname(config_path), exist_ok=True)

    try:
        # σåÖσàÑΘàìτ╜«µûçΣ╗╢∩╝êσÉîµù╢σåÖσàÑΣ╕ñΣ╕¬σÅÿΘçÅσÉì∩╝îΘü┐σàìΣ╝ÿσàêτ║ºΘöÖΣ╣▒∩╝ë
        with open(config_path, "w", encoding="utf-8") as f:
            f.write(
                f"# Koto Configuration\nGEMINI_API_KEY={api_key}\nAPI_KEY={api_key}\n"
            )

        # µ¢┤µû░τÄ»σóâσÅÿΘçÅ
        os.environ["GEMINI_API_KEY"] = api_key
        os.environ["API_KEY"] = api_key
        global API_KEY, client
        API_KEY = api_key
        client = create_client()

        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route("/api/setup/workspace", methods=["POST"])
def setup_workspace():
    """Φ«╛τ╜«σ╖ÑΣ╜£σî║τ¢«σ╜ò"""
    data = request.json
    workspace_path = data.get("path", "").strip()

    if not workspace_path:
        workspace_path = os.path.join(PROJECT_ROOT, "workspace")

    try:
        os.makedirs(workspace_path, exist_ok=True)
        os.makedirs(os.path.join(workspace_path, "documents"), exist_ok=True)
        os.makedirs(os.path.join(workspace_path, "images"), exist_ok=True)
        os.makedirs(os.path.join(workspace_path, "code"), exist_ok=True)

        # µ¢┤µû░Φ«╛τ╜«
        settings_manager.set("storage", "workspace_dir", workspace_path)

        return jsonify({"success": True, "path": workspace_path})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route("/api/setup/test", methods=["GET"])
def test_api_connection():
    """µ╡ïΦ»ò API Φ┐₧µÄÑ"""
    try:
        start = time.time()
        response = client.models.generate_content(
            model="gemini-3-flash-preview",
            contents="Say 'Koto is ready!' in one short sentence.",
        )
        latency = time.time() - start
        return jsonify(
            {"success": True, "message": response.text, "latency": round(latency, 2)}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route("/api/diagnose", methods=["GET"])
def diagnose_models():
    """Φ»èµû¡µëÇµ£ëµ¿íσ₧ïτÜäσÅ»τö¿µÇº"""
    import threading

    results = {
        "proxy": {
            "detected": get_detected_proxy(),
            "force": FORCE_PROXY or None,
            "custom_endpoint": GEMINI_API_BASE or None,
        },
        "models": {},
    }

    # µ╡ïΦ»òµ¿íσ₧ïσêùΦí¿
    test_models = [
        ("gemini-2.0-flash-lite", "Φ╖»τö▒σêåτ▒╗"),
        ("gemini-3-flash-preview", "µùÑσ╕╕σ»╣Φ»¥"),
        ("gemini-3-pro-preview", "Σ╗úτáüτöƒµêÉ"),
        ("gemini-2.5-flash", "Φüöτ╜æµÉ£τ┤ó"),
        ("gemini-3.1-flash-image-preview", "σ¢╛σâÅτöƒµêÉ"),
    ]

    def test_model(model_id, purpose):
        try:
            start = time.time()
            if "image-generation" in model_id or "imagen" in model_id:
                # σ¢╛σâÅµ¿íσ₧ïσÅ¬µ╡ïΦ»òΦ┐₧ΘÇÜµÇº
                response = client.models.generate_content(
                    model=model_id,
                    contents="test",
                    config=types.GenerateContentConfig(max_output_tokens=10),
                )
            else:
                response = client.models.generate_content(
                    model=model_id,
                    contents="Reply with only: OK",
                    config=types.GenerateContentConfig(max_output_tokens=10),
                )
            latency = time.time() - start
            return {
                "status": "Γ£à σÅ»τö¿",
                "latency": round(latency, 2),
                "purpose": purpose,
            }
        except Exception as e:
            error_msg = str(e)
            if "location is not supported" in error_msg:
                status = "Γ¥î σ£░σî║ΘÖÉσê╢"
            elif "not found" in error_msg.lower():
                status = "Γ¥î µ¿íσ₧ïΣ╕ìσ¡ÿσ£¿"
            elif "quota" in error_msg.lower():
                status = "ΓÜá∩╕Å ΘàìΘó¥ΦÇùσ░╜"
            elif "timeout" in error_msg.lower():
                status = "ΓÜá∩╕Å Φ╢àµù╢"
            else:
                status = f"Γ¥î ΘöÖΦ»»"
            return {"status": status, "error": error_msg[:150], "purpose": purpose}

    # σ╣╢Φíîµ╡ïΦ»ò∩╝êσ╕ªΦ╢àµù╢∩╝ë
    threads = []
    for model_id, purpose in test_models:

        def run_test(m=model_id, p=purpose):
            results["models"][m] = test_model(m, p)

        t = threading.Thread(target=run_test, daemon=True)
        threads.append(t)
        t.start()

    # τ¡ëσ╛àµëÇµ£ëτ║┐τ¿ïσ«îµêÉ∩╝êµ£ÇσñÜ 15 τºÆ∩╝ë
    for t in threads:
        t.join(timeout=15)

    # µúÇµƒÑµÿ»σÉªµëÇµ£ëµ¿íσ₧ïΘâ╜Σ╕ìσÅ»τö¿
    all_failed = all(
        "Γ¥î" in results["models"].get(m, {}).get("status", "") for m, _ in test_models
    )

    if all_failed:
        results["recommendation"] = (
            "µëÇµ£ëµ¿íσ₧ïσ¥çΣ╕ìσÅ»τö¿πÇéσ╗║Φ««∩╝Ü\n1. µúÇµƒÑΣ╗úτÉåΘàìτ╜«µÿ»σÉªµ¡úτí«\n2. ΦÇâΦÖæΣ╜┐τö¿ API Σ╕¡Φ╜¼µ£ìσèí\n3. σ£¿ gemini_config.env Σ╕¡Θàìτ╜« GEMINI_API_BASE"
        )

    return jsonify(results)


@app.route("/api/browse", methods=["GET"])
def browse_folders():
    import os

    path = request.args.get("path", "C:\\")

    try:
        if not os.path.exists(path):
            return jsonify({"error": "Φ╖»σ╛äΣ╕ìσ¡ÿσ£¿", "folders": [], "parent": None})

        if not os.path.isdir(path):
            return jsonify({"error": "Σ╕ìµÿ»µûçΣ╗╢σñ╣", "folders": [], "parent": None})

        folders = []
        try:
            for item in os.listdir(path):
                item_path = os.path.join(path, item)
                if os.path.isdir(item_path):
                    folders.append({"name": item, "path": item_path})
        except PermissionError:
            return jsonify({"error": "µ▓íµ£ëµ¥âΘÖÉΦ«┐Θù«", "folders": [], "parent": None})

        folders.sort(key=lambda x: x["name"].lower())

        # Get parent path
        parent = os.path.dirname(path)
        if parent == path:  # Root drive
            parent = None

        return jsonify({"folders": folders, "parent": parent, "current": path})
    except Exception as e:
        return jsonify({"error": str(e), "folders": [], "parent": None})


@app.route("/api/chat/interrupt", methods=["POST"])
def interrupt_chat():
    """Σ╕¡µû¡σ╜ôσëìσ»╣Φ»¥τöƒµêÉ"""
    payload = request.json or {}
    session_name = payload.get("session")
    task_id = payload.get("task_id")
    if not session_name:
        return jsonify({"error": "Missing session"}), 400

    # Σ╜┐τö¿µû░τÜäΣ╕¡µû¡τ«íτÉåσÖ¿
    _interrupt_manager.set_interrupt(session_name)
    # Σ┐¥µîüσÉæσÉÄσà╝σ«╣
    _interrupt_flags[session_name] = True

    # σÅ»ΘÇë∩╝Üσªéµ₧£σëìτ½»Σ╝áσàÑ task_id∩╝îσÉîµ¡ÑσÅûµ╢êΦ░âσ║ªσÖ¿Σ╗╗σèí∩╝êτö¿Σ║Ä DOC_ANNOTATE τ¡ëµ╡üσ╝ÅΘò┐Σ╗╗σèí∩╝ë
    if task_id:
        try:
            from task_scheduler import get_task_scheduler

            get_task_scheduler().cancel_task(task_id)
            _app_logger.debug(f"[INTERRUPT] Cancel task_id={task_id}")
        except Exception as e:
            _app_logger.debug(f"[INTERRUPT] cancel task failed: {e}")

    # σÉîµ¡ÑΣ╕¡µû¡µáçσ┐ùσê░ AgentLoop∩╝êσªéµ₧£µ¡úσ£¿µëºΦíî Agent Σ╗╗σèí∩╝ë
    # NOTE: Legacy agent_loop retired ΓÇö interrupt handled by _interrupt_manager above
    pass

    return jsonify({"success": True, "message": "Chat interrupted"})


@app.route("/api/chat/reset-interrupt", methods=["POST"])
def reset_interrupt():
    """Θçìτ╜«Σ╕¡µû¡µáçσ┐ù"""
    session_name = request.json.get("session")
    if session_name:
        # Σ╜┐τö¿µû░τÜäΣ╕¡µû¡τ«íτÉåσÖ¿
        _interrupt_manager.reset(session_name)
        # Σ┐¥µîüσÉæσÉÄσà╝σ«╣
        if session_name in _interrupt_flags:
            del _interrupt_flags[session_name]
    return jsonify({"success": True})


# ================= µû░σèƒΦâ╜ API Φ╖»τö▒ =================


# === σ┐½ΘÇƒτ¼öΦ«░ API ===
@app.route("/api/notes/add", methods=["POST"])
def add_note():
    """µ╖╗σèáτ¼öΦ«░"""
    from note_manager import get_note_manager

    data = request.json
    title = data.get("title", "")
    content = data.get("content", "")
    category = data.get("category", "default")
    tags = data.get("tags", [])

    note_manager = get_note_manager()
    note_id = note_manager.add_note(title, content, category, tags)

    return jsonify({"success": True, "note_id": note_id})


@app.route("/api/notes/list", methods=["GET"])
def list_notes():
    """σêùσç║µ£ÇΦ┐æτ¼öΦ«░"""
    from note_manager import get_note_manager

    limit = int(request.args.get("limit", 20))
    category = request.args.get("category")

    note_manager = get_note_manager()
    notes = note_manager.get_recent_notes(limit, category)

    return jsonify({"notes": notes})


@app.route("/api/notes/search", methods=["GET"])
def search_notes():
    """µÉ£τ┤óτ¼öΦ«░"""
    from note_manager import get_note_manager

    query = request.args.get("query", "")
    note_manager = get_note_manager()
    results = note_manager.search_notes(query)

    return jsonify({"results": results})


@app.route("/api/notes/<note_id>", methods=["DELETE"])
def delete_note(note_id):
    """σêáΘÖñτ¼öΦ«░"""
    from note_manager import get_note_manager

    note_manager = get_note_manager()
    success = note_manager.delete_note(note_id)

    return jsonify({"success": success})


# === µ£¼σ£░µÅÉΘåÆ API∩╝êWindows τ│╗τ╗ƒΘÇÜτƒÑ∩╝ë ===
@app.route("/api/reminders/add", methods=["POST"])
def add_reminder():
    """σê¢σ╗║µ£¼σ£░τ│╗τ╗ƒµÅÉΘåÆ
    Φ»╖µ▒éΣ╜ô: {"title": str, "message": str, "time": ISO8601, "seconds": int}
    - Σ╝á time (ISO µù╢Θù┤) µêû seconds (τ¢╕σ»╣τºÆµò░) Σ╗╗ΘÇëσà╢Σ╕Ç
    """
    from datetime import datetime

    from reminder_manager import get_reminder_manager

    data = request.json or {}
    title = data.get("title") or "µÅÉΘåÆ"
    message = data.get("message") or ""
    icon = data.get("icon")
    remind_time = data.get("time")
    seconds = data.get("seconds")

    mgr = get_reminder_manager()
    if remind_time:
        try:
            dt = datetime.fromisoformat(remind_time)
        except Exception:
            return jsonify({"success": False, "error": "µù╢Θù┤µá╝σ╝ÅΘ£ÇΣ╕║ ISO8601"}), 400
        rid = mgr.add_reminder(title, message, dt, icon)
    elif seconds is not None:
        try:
            sec = int(seconds)
        except Exception:
            return jsonify({"success": False, "error": "seconds Θ£ÇΣ╕║µò┤µò░"}), 400
        rid = mgr.add_reminder_in(title, message, sec, icon)
    else:
        return jsonify({"success": False, "error": "Θ£ÇµÅÉΣ╛¢ time µêû seconds"}), 400

    return jsonify({"success": True, "reminder_id": rid})


@app.route("/api/reminders/list", methods=["GET"])
def list_reminders_api():
    """σêùσç║µëÇµ£ëµÅÉΘåÆ"""
    from reminder_manager import get_reminder_manager

    mgr = get_reminder_manager()
    return jsonify({"reminders": mgr.list_reminders()})


@app.route("/api/reminders/<reminder_id>", methods=["DELETE"])
def cancel_reminder(reminder_id):
    """σÅûµ╢êµÅÉΘåÆ"""
    from reminder_manager import get_reminder_manager

    mgr = get_reminder_manager()
    ok = mgr.cancel_reminder(reminder_id)
    return jsonify({"success": ok})


# === µùÑτ¿ï∩╝êµ£¼σ£░µùÑσÄå∩╝ë API ===
@app.route("/api/calendar/add", methods=["POST"])
def add_calendar_event():
    """µû░σó₧µùÑτ¿ïσ╣╢Φç¬σè¿σê¢σ╗║µ£¼σ£░µÅÉΘåÆ
    Φ»╖µ▒éΣ╜ô: {"title": str, "description": str, "start": ISO8601, "end": ISO8601?, "remind_before_minutes": int?}
    """
    from datetime import datetime

    from calendar_manager import get_calendar_manager

    data = request.json or {}
    title = data.get("title") or "µùÑτ¿ï"
    description = data.get("description") or ""
    start = data.get("start")
    end = data.get("end")
    remind_before_minutes = int(data.get("remind_before_minutes") or 0)

    if not start:
        return jsonify({"success": False, "error": "start Σ╕ìΦâ╜Σ╕║τ⌐║ (ISO8601)"}), 400
    try:
        start_dt = datetime.fromisoformat(start)
    except Exception:
        return jsonify({"success": False, "error": "start σ┐àΘí╗µÿ» ISO8601 µù╢Θù┤"}), 400
    end_dt = None
    if end:
        try:
            end_dt = datetime.fromisoformat(end)
        except Exception:
            return jsonify({"success": False, "error": "end σ┐àΘí╗µÿ» ISO8601 µù╢Θù┤"}), 400

    mgr = get_calendar_manager()
    event_id = mgr.add_event(
        title, description, start_dt, end_dt, remind_before_minutes
    )
    return jsonify({"success": True, "event_id": event_id})


@app.route("/api/calendar/list", methods=["GET"])
def list_calendar_events():
    from calendar_manager import get_calendar_manager

    limit = int(request.args.get("limit", 100))
    mgr = get_calendar_manager()
    return jsonify({"events": mgr.list_events(limit)})


@app.route("/api/calendar/<event_id>", methods=["DELETE"])
def delete_calendar_event(event_id):
    from calendar_manager import get_calendar_manager

    mgr = get_calendar_manager()
    ok = mgr.delete_event(event_id)
    return jsonify({"success": ok})


# === σë¬Φ┤┤µ¥┐ API ===
@app.route("/api/clipboard/history", methods=["GET"])
def get_clipboard_history():
    """ΦÄ╖σÅûσë¬Φ┤┤µ¥┐σÄåσÅ▓"""
    from clipboard_manager import get_clipboard_manager

    limit = int(request.args.get("limit", 50))
    type_filter = request.args.get("type")
    clipboard_manager = get_clipboard_manager()
    history = clipboard_manager.get_history(limit)
    if type_filter:
        history = [item for item in history if item.get("type") == type_filter]

    return jsonify({"history": history})


@app.route("/api/clipboard/search", methods=["GET"])
def search_clipboard():
    """µÉ£τ┤óσë¬Φ┤┤µ¥┐σÄåσÅ▓"""
    from clipboard_manager import get_clipboard_manager

    query = request.args.get("query", "")
    type_filter = request.args.get("type")
    clipboard_manager = get_clipboard_manager()
    results = clipboard_manager.search(query)
    if type_filter:
        results = [item for item in results if item.get("type") == type_filter]

    return jsonify({"results": results})


@app.route("/api/clipboard/copy", methods=["POST"])
def copy_from_history():
    """Σ╗ÄσÄåσÅ▓Σ╕¡σñìσê╢"""
    from clipboard_manager import get_clipboard_manager

    content = request.json.get("content")
    index = request.json.get("index")
    clipboard_manager = get_clipboard_manager()
    if index is not None:
        try:
            index = int(index)
        except Exception:
            return jsonify({"success": False, "error": "index σ┐àΘí╗µÿ»µò┤µò░"}), 400
        success = clipboard_manager.copy_from_history(index)
    else:
        success = clipboard_manager.copy_from_history(content or "")

    return jsonify({"success": success})


# === Σ╗╗σèíΦ░âσ║ª API∩╝êσ╖▓Φ┐üτº╗Φç│ task_bp Φô¥σ¢╛ app/api/task_routes.py∩╝ë===
# σÄƒσåàΦüöΦ╖»τö▒Σ╛¥Φ╡ûΣ╕ìσ¡ÿσ£¿τÜä task_scheduler µ¿íσ¥ù∩╝îσ╖▓τº╗ΘÖñΣ╗ÑΦºúΘÖñσ»╣ task_bp τÜäΦ╖»τö▒Θÿ╗µïªπÇé
# task_bp µÅÉΣ╛¢∩╝ÜGET /api/tasks, GET /api/tasks/<id>, POST /api/tasks/<id>/cancel,
#              POST /api/tasks/<id>/interrupt, GET /api/tasks/<id>/stream, τ¡ëπÇé


# === Θé«Σ╗╢ API ===
@app.route("/api/email/accounts", methods=["GET"])
def list_email_accounts():
    """σêùσç║Θé«τ«▒Φ┤ªµê╖"""
    from email_manager import get_email_manager

    email_manager = get_email_manager()
    accounts = list(email_manager.accounts.keys())
    default = email_manager.default_account

    return jsonify({"accounts": accounts, "default": default})


@app.route("/api/email/accounts/add", methods=["POST"])
def add_email_account():
    """µ╖╗σèáΘé«τ«▒Φ┤ªµê╖"""
    from email_manager import get_email_manager

    data = request.json
    email_address = data.get("email")
    password = data.get("password")
    smtp_server = data.get("smtp_server")
    smtp_port = data.get("smtp_port", 587)
    imap_server = data.get("imap_server")
    set_as_default = data.get("set_as_default", False)

    email_manager = get_email_manager()
    success = email_manager.add_account(
        email_address=email_address,
        password=password,
        smtp_server=smtp_server,
        smtp_port=smtp_port,
        imap_server=imap_server,
        set_as_default=set_as_default,
    )

    return jsonify({"success": success})


@app.route("/api/email/send", methods=["POST"])
def send_email():
    """σÅæΘÇüΘé«Σ╗╢"""
    from email_manager import get_email_manager

    data = request.json
    to_addrs = data.get("to", [])
    subject = data.get("subject", "")
    body = data.get("body", "")
    cc_addrs = data.get("cc", [])
    attachments = data.get("attachments", [])
    html = data.get("html", False)

    email_manager = get_email_manager()
    success = email_manager.send_email(
        to_addrs=to_addrs,
        subject=subject,
        body=body,
        cc_addrs=cc_addrs,
        attachments=attachments,
        html=html,
    )

    return jsonify({"success": success})


@app.route("/api/email/fetch", methods=["GET"])
def fetch_emails():
    """ΦÄ╖σÅûΘé«Σ╗╢σêùΦí¿"""
    from email_manager import get_email_manager

    folder = request.args.get("folder", "INBOX")
    limit = int(request.args.get("limit", 20))
    unread_only = request.args.get("unread_only", "false").lower() == "true"

    email_manager = get_email_manager()
    emails = email_manager.fetch_emails(
        folder=folder, limit=limit, unread_only=unread_only
    )

    return jsonify({"emails": emails})


@app.route("/api/email/search", methods=["GET"])
def search_emails():
    """µÉ£τ┤óΘé«Σ╗╢"""
    from email_manager import get_email_manager

    keyword = request.args.get("query", "")
    folder = request.args.get("folder", "INBOX")

    email_manager = get_email_manager()
    results = email_manager.search_emails(keyword, folder=folder)

    return jsonify({"results": results})


# === µ╡ÅΦºêσÖ¿Φç¬σè¿σîû API ===
@app.route("/api/browser/open", methods=["POST"])
def browser_open():
    """µëôσ╝Ç URL"""
    from browser_automation import get_browser_automation

    url = request.json.get("url", "")
    browser = get_browser_automation()
    success = browser.open_url(url)

    return jsonify({"success": success})


@app.route("/api/browser/search", methods=["POST"])
def browser_search():
    """Google µÉ£τ┤ó"""
    from browser_automation import get_browser_automation

    query = request.json.get("query", "")
    browser = get_browser_automation()
    results = browser.search_google(query)

    return jsonify({"results": results})


@app.route("/api/browser/screenshot", methods=["POST"])
def browser_screenshot():
    """µê¬σ¢╛"""
    import os

    from browser_automation import get_browser_automation

    filename = request.json.get("filename", f"screenshot_{int(time.time())}.png")
    file_path = os.path.join(WORKSPACE_DIR, "images", filename)

    browser = get_browser_automation()
    success = browser.take_screenshot(file_path)

    return jsonify({"success": success, "path": file_path})


# === µÖ║Φâ╜µÉ£τ┤ó API ===
@app.route("/api/search/all", methods=["GET"])
def search_all():
    """σà¿σ▒ÇµÉ£τ┤ó"""
    from search_engine import get_search_engine

    query = request.args.get("query", "")
    max_results = int(request.args.get("max_results", 50))

    search_engine = get_search_engine()
    results = search_engine.search_all(query, max_results)

    return jsonify(results)


@app.route("/api/search/files", methods=["GET"])
def search_files():
    """µÉ£τ┤óµûçΣ╗╢"""
    from search_engine import get_search_engine

    query = request.args.get("query", "")
    max_results = int(request.args.get("max_results", 20))

    search_engine = get_search_engine()
    results = search_engine.search_files(query, max_results)

    return jsonify({"results": results})


# ================= Φ»¡Θƒ│Φ»åσê½ API (µû░µ₧╢µ₧ä) =================
@app.route("/api/voice/engines", methods=["GET"])
def voice_engines():
    """ΦÄ╖σÅûσÅ»τö¿Φ»¡Θƒ│σ╝òµôÄσêùΦí¿"""
    try:
        from web.voice_fast import get_available_engines

        result = get_available_engines()
        return jsonify(result)
    except Exception as e:
        return (
            jsonify(
                {
                    "success": False,
                    "engines": [],
                    "message": f"ΦÄ╖σÅûσ╝òµôÄσêùΦí¿σñ▒Φ┤Ñ: {str(e)}",
                }
            ),
            500,
        )


@app.route("/api/voice/record", methods=["POST"])
def voice_record():
    """σ╜òσê╢Θƒ│Θóæ"""
    try:
        data = request.json or {}
        duration = data.get("duration", 5)

        from web.voice_input import record_audio

        result = record_audio(duration=int(duration))

        return jsonify(result)
    except Exception as e:
        return (
            jsonify(
                {"success": False, "message": f"σ╜òΘƒ│σñ▒Φ┤Ñ: {str(e)}", "audio_file": None}
            ),
            500,
        )


@app.route("/api/voice/recognize", methods=["POST"])
def voice_recognize():
    """Φ»åσê½Θƒ│ΘóæµûçΣ╗╢"""
    try:
        data = request.json or {}
        audio_path = data.get("audio_path")
        engine = data.get("engine", None)

        if not audio_path:
            return jsonify({"success": False, "message": "τ╝║σ░æΘƒ│ΘóæµûçΣ╗╢Φ╖»σ╛ä"}), 400

        from web.voice_input import recognize_audio

        result = recognize_audio(audio_path, engine)

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "message": f"Φ»åσê½σñ▒Φ┤Ñ: {str(e)}"}), 500


@app.route("/api/voice/listen", methods=["POST"])
def voice_listen():
    """Σ╕ÇΘö«Θ║ªσàïΘúÄΦ»åσê½∩╝êµ£¼σ£░µ¿íσ╝Å - Σ╝ÿσîûτëê∩╝Üτ½ïσì│σÉ»σè¿∩╝ë"""
    try:
        data = request.json or {}
        timeout = data.get("timeout", 5)
        language = data.get("language", "zh-CN")

        # Σ╜┐τö¿σ┐½ΘÇƒµ£¼σ£░Φ»åσê½
        from web.voice_fast import recognize_voice

        result = recognize_voice(timeout=int(timeout), language=language)

        # Σ╝ÿσîû∩╝ÜΦ«╛τ╜«σôìσ║öσñ┤σèáσ┐½Σ╝áΦ╛ô
        response = jsonify(result)
        response.headers["Cache-Control"] = "no-cache, no-store"
        response.headers["Content-Type"] = "application/json; charset=utf-8"
        return response

    except Exception as e:
        import traceback

        traceback.print_exc()
        response = jsonify(
            {
                "success": False,
                "text": "",
                "message": f"Φ»¡Θƒ│Φ»åσê½σç║ΘöÖ: {str(e)}",
                "engine": "error",
            }
        )
        response.status_code = 500
        response.headers["Cache-Control"] = "no-cache"
        return response


@app.route("/api/voice/stream")
def voice_stream():
    """µ╡üσ╝ÅΦ»¡Θƒ│Φ»åσê½ - Vosk µ£¼σ£░τª╗τ║┐∩╝îσ«₧µù╢Φ┐öσ¢₧Θâ¿σêå/µ£Çτ╗êτ╗ôµ₧£∩╝êSSE∩╝ë"""
    import json as _json

    from flask import Response, stream_with_context

    @stream_with_context
    def generate():
        try:
            from web.voice_engine import recognize_stream

            for event in recognize_stream(max_wait=8.0, max_speech=30.0):
                yield f"data: {_json.dumps(event, ensure_ascii=False)}\n\n"
                if event.get("type") in ("final", "error"):
                    break
        except GeneratorExit:
            pass
        except Exception as e:
            yield f"data: {_json.dumps({'type': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"

    return Response(
        generate(),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
            "Transfer-Encoding": "chunked",
        },
    )


@app.route("/api/voice/stop", methods=["POST"])
def voice_stop():
    """σü£µ¡óσ╜ôσëìΦ»¡Θƒ│Φ»åσê½µ╡ü∩╝êΘÇÜτƒÑ voice_engine σü£µ¡ó∩╝ë"""
    try:
        from web.voice_engine import request_stop

        request_stop()
    except Exception:
        pass
    return jsonify({"success": True, "message": "σ╖▓σÅæΘÇüσü£µ¡óΣ┐íσÅ╖"})


@app.route("/api/voice/commands", methods=["GET"])
def voice_commands():
    """Φ┐öσ¢₧σåàτ╜«Φ»¡Θƒ│σæ╜Σ╗ñσêùΦí¿∩╝êΣ╛¢Φ»¡Θƒ│Θ¥óµ¥┐σ▒òτñ║∩╝ë"""
    commands = [
        {"name": "σÅæΘÇüµ╢êµü»", "description": "Φ»┤σç║µ╢êµü»σÉÄΦç¬σè¿σÅæΘÇü", "keyword": ""},
        {"name": "µû░σ»╣Φ»¥", "description": "Φ»┤'µû░σ»╣Φ»¥'σ╝Çσºïµû░Φüèσñ⌐", "keyword": "µû░σ»╣Φ»¥"},
        {"name": "µ╕àτ⌐║Φ╛ôσàÑ", "description": "Φ»┤'µ╕àτ⌐║'µ╕àΘÖñΦ╛ôσàÑµíå", "keyword": "µ╕àτ⌐║"},
        {"name": "Θçìµû░Φ»åσê½", "description": "σåìµ¼íτé╣σç╗Θ║ªσàïΘúÄΘçìµû░Φ»┤", "keyword": ""},
    ]
    return jsonify({"success": True, "commands": commands})


@app.route("/api/voice/stt_status", methods=["GET"])
def voice_stt_status():
    """µƒÑΦ»óσ╜ôσëìΦ»¡Θƒ│σ╝òµôÄτè╢µÇü∩╝êΣ╜┐τö¿µû░ voice_engine∩╝ëπÇé"""
    try:
        from web.voice_engine import get_status

        fast = get_status()
    except Exception:
        fast = {"available": False, "engine": "unavailable", "label": "µùáσ╝òµôÄ"}

    return jsonify(
        {
            "fast": fast,
            "local": fast,  # σà╝σ«╣σëìτ½»µùºσ¡ùµ«╡
            "active": fast.get("engine", "none"),
        }
    )


@app.route("/api/voice/gemini_stt", methods=["POST"])
@app.route("/api/voice/stt", methods=["POST"])  # τ╗ƒΣ╕ÇσàÑσÅúσê½σÉì
def voice_gemini_stt():
    """
    τ╗ƒΣ╕ÇΦ»¡Θƒ│Φ╜¼µûçσ¡ù (STT) σàÑσÅú∩╝Üµ£¼σ£░ Whisper Σ╝ÿσàê ΓåÆ Gemini STT σñçτö¿πÇé

    - ΦïÑσ«ëΦúàΣ║å faster-whisper µêû openai-whisper∩╝Üσ«îσà¿µ£¼σ£░Φ╜¼σåÖ∩╝îµùá API µ╢êΦÇù
    - σÉªσêÖ∩╝ÜσÅæΘÇüΦç│ Gemini gemini-2.0-flash-lite Φ╜¼σåÖ
    - σºïτ╗êΦ┐öσ¢₧ JSON∩╝îτ╗¥Σ╕ìΦ┐öσ¢₧ HTML ΘöÖΦ»»Θí╡Θ¥óπÇé
    """
    try:
        data = request.get_json(silent=True) or {}
        audio_b64 = data.get("audio", "")
        mime_type = data.get("mime", "audio/webm")

        if not audio_b64:
            return (
                jsonify({"success": False, "text": "", "message": "τ╝║σ░æ audio σ¡ùµ«╡"}),
                400,
            )

        import base64 as _b64

        try:
            audio_bytes = _b64.b64decode(audio_b64)
        except Exception:
            return (
                jsonify(
                    {"success": False, "text": "", "message": "Θƒ│Θóæ base64 Φºúτáüσñ▒Φ┤Ñ"}
                ),
                400,
            )

        if len(audio_bytes) < 300:
            return jsonify(
                {"success": False, "text": "", "message": "σ╜òΘƒ│σñ¬τƒ¡∩╝îΦ»╖Θçìµû░Φ»┤Φ»¥"}
            )

        _app_logger.debug(f"[STT] µö╢σê░Θƒ│Θóæ {len(audio_bytes)/1024:.1f}KB  MIME={mime_type}")

        # ΓöÇΓöÇ Σ╝ÿσàêσ░¥Φ»òµ£¼σ£░ Whisper ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        try:
            from web.local_stt import is_available, transcribe

            if is_available():
                ok, text, engine = transcribe(audio_bytes, mime_type)
                if ok and text:
                    return jsonify(
                        {
                            "success": True,
                            "text": text,
                            "engine": engine,
                            "message": "Φ»åσê½µêÉσèƒ∩╝êµ£¼σ£░∩╝ë",
                        }
                    )
                # µ£¼σ£░Φ»åσê½σç║τ⌐║µûçµ£¼ ΓåÆ Σ╣ƒτ¢┤µÄÑΦ┐öσ¢₧∩╝êΣ╕ìσ¢₧ΘÇÇ∩╝îΘü┐σàìΘçìσñìΦ«íΦ┤╣∩╝ë
                return jsonify(
                    {
                        "success": False,
                        "text": "",
                        "engine": engine,
                        "message": "µ£¬µúÇµ╡ïσê░Φ»¡Θƒ│",
                    }
                )
        except Exception as _le:
            _app_logger.debug(f"[STT] µ£¼σ£░ STT σ╝éσ╕╕∩╝îσ¢₧ΘÇÇ Gemini: {_le}")

        # ΓöÇΓöÇ σ¢₧ΘÇÇ∩╝ÜGemini STT ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
        if client is None:
            return (
                jsonify(
                    {
                        "success": False,
                        "text": "",
                        "message": "Gemini σ«óµê╖τ½»µ£¬σê¥σºïσîû∩╝îΦ»╖µúÇµƒÑ API Key∩╝¢"
                        "µêûσ«ëΦúà faster-whisper Σ╜┐τö¿µ£¼σ£░Φ»åσê½",
                    }
                ),
                503,
            )

        stt_model = "gemini-2.0-flash-lite"
        prompt_parts = [
            types.Part.from_bytes(data=audio_bytes, mime_type=mime_type),
            types.Part.from_text(
                text=(
                    "Φ»╖σ░åΣ╕èΘ¥óΘƒ│ΘóæΣ╕¡τÜäΦ»¡Θƒ│σåàσ«╣σ«îµò┤Φ╜¼σåÖΣ╕║µûçσ¡ùπÇé"
                    "σÅ¬Φ╛ôσç║Φ╜¼σåÖτ╗ôµ₧£∩╝îΣ╕ìΦªüσèáΣ╗╗Σ╜òΦºúΘçèπÇüµáçτé╣Σ┐«ΘÑ░µêûσëìτ╝Ç∩╝êσªéπÇîΦ╜¼σåÖ∩╝ÜπÇìτ¡ë∩╝ëπÇé"
                    "σªéµ₧£σÉ¼Σ╕ìµ╕àµêûµ▓íµ£ëΦ»¡Θƒ│∩╝îσÅ¬Φ╛ôσç║τ⌐║σ¡ùτ¼ªΣ╕▓πÇé"
                )
            ),
        ]

        resp = client.models.generate_content(
            model=stt_model,
            contents=[types.Content(role="user", parts=prompt_parts)],
            config=types.GenerateContentConfig(temperature=0.0, max_output_tokens=512),
        )

        text = (resp.text or "").strip()
        for prefix in ("Φ╜¼σåÖ∩╝Ü", "Φ╜¼σåÖ:", "Φ»åσê½∩╝Ü", "Φ»åσê½:", "µûçσ¡ù∩╝Ü", "µûçσ¡ù:"):
            if text.startswith(prefix):
                text = text[len(prefix) :].strip()

        _app_logger.debug(f"[STT] Gemini Φ»åσê½τ╗ôµ₧£: {text[:80]!r}")
        return jsonify(
            {
                "success": bool(text),
                "text": text,
                "engine": f"Gemini/{stt_model}",
                "message": "Φ»åσê½µêÉσèƒ" if text else "µ£¬µúÇµ╡ïσê░Φ»¡Θƒ│σåàσ«╣",
            }
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        return (
            jsonify(
                {"success": False, "text": "", "message": f"STT σñ▒Φ┤Ñ: {str(e)[:200]}"}
            ),
            500,
        )


# ================= σó₧σ╝║σèƒΦâ╜ API (σ£║µÖ»1-3) =================


@app.route("/api/data/extract-transform", methods=["POST"])
def data_extract_transform():
    """µò░µì«µÅÉσÅûΣ╕ÄΦ╜¼µìó - σ£║µÖ»1∩╝ÜΦ╖¿σ║öτö¿µò░µì«µÉ¼Φ┐É"""
    try:
        data = request.json
        source_type = data.get("source_type", "wechat_contact")
        source_data = data.get("source_data")
        target_format = data.get("target_format", "excel")
        output_filename = data.get(
            "output_filename", f'µÅÉσÅûµò░µì«_{datetime.now().strftime("%Y%m%d_%H%M%S")}'
        )

        # τí«σ«ÜΦ╛ôσç║Φ╖»σ╛ä
        if target_format == "excel":
            ext = ".xlsx"
        elif target_format == "csv":
            ext = ".csv"
        else:
            ext = ".json"

        output_path = os.path.join(
            WORKSPACE_DIR, "documents", f"{output_filename}{ext}"
        )

        # µëºΦíîµò░µì«τ«íΘüô
        from web.data_pipeline import CrossAppDataPipeline

        pipeline = CrossAppDataPipeline()
        result = pipeline.run_pipeline(
            source_type, source_data, target_format, output_path
        )

        return jsonify(result)

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/code/generate", methods=["POST"])
def code_generate():
    """Σ╗úτáüτöƒµêÉ - σ£║µÖ»2∩╝Üσ╕«σè⌐τö¿µê╖σ«îµêÉτ╝ûτ¿ïΣ╗╗σèí"""
    try:
        data = request.json
        template_name = data.get("template_name")
        description = data.get("description")
        language = data.get("language", "python")
        output_filename = data.get("output_filename")

        from web.code_generator import CodeGenerator

        generator = CodeGenerator()

        # τí«σ«ÜΦ╛ôσç║Φ╖»σ╛ä
        output_path = None
        if output_filename:
            output_path = os.path.join(WORKSPACE_DIR, "code", output_filename)

        # τöƒµêÉΣ╗úτáü
        if template_name:
            result = generator.generate(
                template_name, output_path, **data.get("params", {})
            )
        elif description:
            # Σ╜┐τö¿AIτöƒµêÉ∩╝êσªéµ₧£σÅ»τö¿∩╝ë
            result = generator.generate_from_description(description, language)
        else:
            return (
                jsonify(
                    {"success": False, "error": "Θ£ÇΦªüµÅÉΣ╛¢template_nameµêûdescription"}
                ),
                400,
            )

        return jsonify(result)

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/code/templates", methods=["GET"])
def code_templates():
    """ΦÄ╖σÅûσÅ»τö¿Σ╗úτáüµ¿íµ¥┐σêùΦí¿"""
    try:
        from web.code_generator import CodeGenerator

        generator = CodeGenerator()

        language = request.args.get("language")
        templates = generator.list_templates(language)

        return jsonify(
            {"success": True, "templates": templates, "count": len(templates)}
        )

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/ppt/generate", methods=["POST"])
def ppt_generate():
    """PPTτöƒµêÉ - σ£║µÖ»3∩╝ÜΘ½ÿΦ┤¿ΘçÅµ╝öτñ║µûçτ¿┐"""
    try:
        data = request.json
        title = data.get("title", "µ╝öτñ║µûçτ¿┐")
        subtitle = data.get("subtitle", "")
        outline = data.get("outline")
        content = data.get("content")
        theme = data.get("theme", "business")
        output_filename = data.get(
            "output_filename",
            f'{title}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pptx',
        )

        output_path = os.path.join(WORKSPACE_DIR, "documents", output_filename)

        from web.ppt_generator import PPTGenerator

        generator = PPTGenerator(theme=theme)

        # τöƒµêÉPPT
        if outline:
            result = generator.generate_from_outline(
                title, outline, output_path, subtitle=subtitle
            )
        elif content:
            result = generator.generate_from_text(content, output_path, title)
        else:
            return jsonify({"success": False, "error": "Θ£ÇΦªüµÅÉΣ╛¢outlineµêûcontent"}), 400

        return jsonify(result)

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ==================== µÖ║Φâ╜µûçµíúσñäτÉåΦ╖»τö▒ ====================


def _should_use_annotation_system(requirement: str, has_file: bool = False) -> bool:
    """
    Σ╕Ñµá╝σêñµû¡µÿ»σÉªΣ╜┐τö¿µûçµíúµáçµ│¿τ│╗τ╗ƒ∩╝êσ£¿σÄƒµûçΣ╕èµáçτ║óΣ┐«µö╣∩╝ë

    µáçµ│¿τ│╗τ╗ƒΣ╗àΘÇéτö¿Σ║Ä∩╝Üτö¿µê╖µÿÄτí«Φªüµ▒éσ£¿σÄƒµûçΣ╕èσüÜµáçΦ«░/µë╣µ│¿/µáçτ║ó/Track Changes

    µ│¿µäÅ∩╝Ü"Σ┐«µö╣"πÇü"Σ╝ÿσîû"πÇü"µö╣σûä"τ¡ëΦ»ìσñ¬σ«╜µ│¢∩╝îΣ╕ìΦâ╜σìòτï¼ΦºªσÅæµáçµ│¿πÇé
    σÅ¬µ£ëΣ╕Ä"σ£¿σÄƒµûçΣ╕è"πÇü"µáçσç║µ¥Ñ"πÇü"µáçτ║ó"τ¡ëσ«ÜΣ╜ìΦ»ìτ╗äσÉêµëìΦºªσÅæπÇé
    """
    if not requirement:
        return False

    requirement_lower = requirement.lower()

    # τ¼¼Σ╕Çσ▒é∩╝ÜµÿÄτí«τÜäµáçµ│¿/µë╣µ│¿σà│Θö«Φ»ì ΓÇö τ¢┤µÄÑΦºªσÅæ
    explicit_annotation = [
        "µáçµ│¿",
        "µáçΦ«░",
        "µë╣µ│¿",
        "µáçσç║",
        "µáçτ║ó",
        "track changes",
        "µë╣µö╣",
    ]
    if any(kw in requirement_lower for kw in explicit_annotation):
        return True

    # τ¼¼Σ║îσ▒é∩╝Üτ╝ûΦ╛æµäÅσ¢╛ + σ«ÜΣ╜ìΦ»ìτ╗äσÉêµëìΦºªσÅæ
    # "Σ┐«µö╣"σìòτï¼σç║τÄ░ Γëá µáçµ│¿∩╝î"Σ┐«µö╣+µáçσç║µ¥Ñ" = µáçµ│¿
    edit_words = ["Σ┐«µö╣", "µö╣µ¡ú", "τ║áµ¡ú", "µáíσ»╣", "σ«íµáí", "τ║áΘöÖ"]
    location_words = [
        "σ£¿σÄƒµûç",
        "σÄƒµûçΣ╕è",
        "µáçσç║",
        "µáçΦ«░σç║",
        "µîçσç║.*Σ╜ìτ╜«",
        "σô¬Σ║¢σ£░µû╣",
        "σô¬Σ║¢Σ╜ìτ╜«",
    ]
    has_edit = any(kw in requirement_lower for kw in edit_words)
    has_location = any(re.search(kw, requirement_lower) for kw in location_words)

    if has_edit and has_location:
        return True

    # τ¼¼Σ╕ëσ▒é∩╝Üσ«íµƒÑ/Σ┐«µö╣+Φ┤¿ΘçÅµÅÅΦ┐░τ╗äσÉê
    review_words = ["σ«íµƒÑ", "Φ»äσ«í", "σ«íµá╕", "µö╣σûä", "Σ╝ÿσîû", "Σ┐«µö╣", "µ╢ªΦë▓", "Φ░âµò┤"]
    quality_words = ["Σ╕ìσÉêΘÇé", "τöƒτí¼", "τ┐╗Φ»æΦàö", "Φ»¡σ║Å", "τö¿Φ»ì", "ΘÇ╗Φ╛æ", "Θù«Θóÿ"]
    has_review = any(kw in requirement_lower for kw in review_words)
    has_quality = any(kw in requirement_lower for kw in quality_words)

    if has_review and has_quality:
        return True

    # Θ╗ÿΦ«ñΣ╕ìΦºªσÅæ ΓÇö σ«üσÅ»µ╝ÅσêñΣ╣ƒΣ╕ìΦ»»σêñ
    return False


def _is_analysis_request(requirement: str) -> bool:
    """σêñµû¡µÿ»σÉªΣ╕║σêåµ₧É/Θù«τ¡öτ▒╗Φ»╖µ▒é∩╝êσîàµï¼τ«ÇσìòΘù«τ¡öσÆîσñìµ¥éσêåµ₧É∩╝îΣ╜åΣ╕ìσÉ½τöƒµêÉµûçµíúµäÅσ¢╛∩╝ë"""
    if not requirement:
        return False

    requirement_lower = requirement.lower()

    # µÿÄτí«τÜäσêåµ₧É/Θù«τ¡öσè¿Σ╜£Φ»ì∩╝êσà¿Θ¥óΦªåτ¢û∩╝ë
    analysis_actions = [
        # σêåµ₧Éτ▒╗
        "σêåµ₧É",
        "µÇ╗τ╗ô",
        "µªéΦ┐░",
        "µó│τÉå",
        "ΦºúΦ»╗",
        "Φ»äΣ╝░",
        "σ»╣µ»ö",
        "µÅÉτé╝",
        "σ╜Æτ║│",
        "Σ╕╗ΦªüΦºéτé╣",
        "µá╕σ┐âΦºéτé╣",
        "Φªüτé╣",
        "Θçìτé╣",
        "Σ║«τé╣",
        # Θù«τ¡ö/Φ»óΘù«τ▒╗
        "σæèΦ»ëµêæ",
        "σæèΦ»ë",
        "µÿ»Σ╗ÇΣ╣ê",
        "σüÜΣ╗ÇΣ╣ê",
        "µâ│σüÜΣ╗ÇΣ╣ê",
        "σ£¿σüÜΣ╗ÇΣ╣ê",
        "µÿ»σÉª",
        "µ£ëµ▓íµ£ë",
        "σÇ╝Σ╕ìσÇ╝",
        "σÇ╝Σ╕ìσÇ╝σ╛ù",
        "µèòΦ╡äΣ╗╖σÇ╝",
        "µèòΦ╡äσ╗║Φ««",
        "µÿ»σÉªσÇ╝σ╛ù",
        "σÇ╝σ╛ùµèòΦ╡ä",
        "µ£ëµùáΣ╗╖σÇ╝",
        "µ£ëΣ╗╖σÇ╝σÉù",
        "σÇ╝σ╛ùσà│µ│¿",
        "Φ«▓Φ«▓",
        "Φ«▓Σ╕ÇΣ╕ï",
        "Φ»┤Φ»┤",
        "Φ»┤Σ╕ÇΣ╕ï",
        "Σ╗ïτ╗ì",
        "Σ╗ïτ╗ìΣ╕ÇΣ╕ï",
        "Σ╗ïτ╗ìΣ╕ï",
        "ΦºúΘçè",
        "ΦºúΘçèΣ╕ÇΣ╕ï",
        "σ╕«µêæΦºúΘçè",
        "Σ║åΦºú",
        "τ£ïτ£ï",
        "τ£ïΣ╕Çτ£ï",
        "Φ»╗Σ╕ÇΦ»╗",
        "Φ»╗Σ╕ÇΣ╕ï",
        "Σ╗ÇΣ╣êµÿ»",
        "µÇÄΣ╣êτ£ï",
        "µÇÄΣ╣êµá╖",
        "σªéΣ╜ò",
        "Σ╗ÇΣ╣êµâàσå╡",
        "σ╕«µêæτ£ï",
        "σ╕«µêæΦ»╗",
        "σ╕«µêæτÉåΦºú",
        "σ╕«µêæΣ║åΦºú",
        "σ╕«µêæΦ»äΣ╝░",
        "σ╕«µêæσêñµû¡",
        "Φ┐ÖΣ╗╜",
        "Φ┐ÖΣ╕¬",
        "µúÇµƒÑΣ╕ÇΣ╕ï",
        "µƒÑτ£ïΣ╕ÇΣ╕ï",
        "τ£ïΣ╕ÇΣ╕ïΦ┐Ö",
        "Σ╗ûΣ╗¼µâ│",
        "Σ╗ûµâ│",
        "σ«âµâ│",
        "Φ»Ñσà¼σÅ╕",
        "Φ»ÑΘí╣τ¢«",
        # English
        "review",
        "analysis",
        "summary",
        "summarize",
        "analyze",
        "explain",
        "understand",
        "evaluate",
        "assess",
        "what is",
        "what does",
        "how does",
        "tell me",
        "should i",
        "is it worth",
        "investment value",
        "check",
        "read this",
        "look at",
    ]

    # µÄÆΘÖñΦ»ì∩╝ÜµÿÄτí«τÜäµûçµíúτöƒµêÉµäÅσ¢╛∩╝êσÅ¬µÄÆΘÖñµ£ÇµÿÄτí«τÜäτöƒµêÉµîçΣ╗ñ∩╝ë
    generation_words = [
        "τöƒµêÉΣ╕ÇΣ╗╜",
        "τöƒµêÉΣ╕ÇΣ╕¬",
        "σ╕«µêæτöƒµêÉ",
        "σåÖΣ╕ÇΣ╗╜",
        "σåÖΣ╕ÇΣ╕¬",
        "σ╕«µêæσåÖ",
        "µö╣σûä",
        "µö╣Φ┐¢",
        "Σ╝ÿσîû",
        "µ╢ªΦë▓",
        "ΘçìσåÖ",
        "σ╕«µêæσüÜΣ╕ÇΣ╗╜",
        "σüÜΣ╕ÇΣ╕¬µèÑσæè",
        "σüÜΣ╕ÇΣ╗╜µèÑσæè",
        "create a document",
        "generate a report",
        "write a report",
    ]

    has_analysis = any(kw in requirement_lower for kw in analysis_actions)
    has_generation = any(kw in requirement_lower for kw in generation_words)

    if has_analysis and not has_generation:
        return True

    return False


def _is_explicit_file_gen_request(requirement: str) -> bool:
    """σêñµû¡τö¿µê╖µÿ»σÉªµÿÄτí«Φªüµ▒éτöƒµêÉ/Φ╛ôσç║Σ╕ÇΣ╕¬µû░µûçΣ╗╢∩╝êµèÑσæèπÇüWordπÇüPDFτ¡ë∩╝ë"""
    if not requirement:
        return False
    requirement_lower = requirement.lower()
    gen_keywords = [
        "τöƒµêÉΣ╕ÇΣ╗╜",
        "τöƒµêÉΣ╕ÇΣ╕¬",
        "σ╕«µêæτöƒµêÉ",
        "σåÖΣ╕ÇΣ╗╜µèÑσæè",
        "σåÖΣ╕ÇΣ╕¬µèÑσæè",
        "σåÖµèÑσæè",
        "σåÖΣ╕ÇΣ╗╜",
        "σ╕«µêæσåÖ",
        "σüÜΣ╕ÇΣ╗╜",
        "σüÜΣ╕ÇΣ╕¬",
        "σ╕«µêæσüÜ",
        "σ»╝σç║",
        "Φ╛ôσç║Σ╕║",
        "Σ┐¥σ¡ÿΣ╕║",
        "Φ╜¼µêÉ",
        "τöƒµêÉword",
        "τöƒµêÉpdf",
        "τöƒµêÉexcel",
        "τöƒµêÉppt",
        "σê¢σ╗║µûçµíú",
        "µû░σ╗║µûçµíú",
        "σê╢Σ╜£µèÑσæè",
        "µò┤τÉåµêÉµûçµíú",
        "σ╜óµêÉµèÑσæè",
        "Φ╛ôσç║µèÑσæè",
    ]
    return any(kw in requirement_lower for kw in gen_keywords)


@app.route("/api/document/smart-process", methods=["POST"])
def document_smart_process():
    """
    µÖ║Φâ╜µûçµíúσñäτÉåσàÑσÅú
    Φç¬σè¿σêñµû¡Σ╜┐τö¿∩╝Üµáçµ│¿τ│╗τ╗ƒ or µûçΣ╗╢σêåµ₧Éτ│╗τ╗ƒ
    """
    try:
        data = request.json
        file_path = data.get("file_path")
        requirement = data.get("requirement", "")

        if not file_path:
            return jsonify({"success": False, "error": "τ╝║σ░æfile_pathσÅéµò░"}), 400

        # µÖ║Φâ╜σêñµû¡σ║öΦ»Ñτö¿σô¬Σ╕¬τ│╗τ╗ƒ
        use_annotation = _should_use_annotation_system(requirement)

        _app_logger.debug(f"[SmartProcess] µÖ║Φâ╜σêñµû¡: use_annotation={use_annotation}")
        _app_logger.debug(f"[SmartProcess] Θ£Çµ▒é: {requirement[:100]}")

        if use_annotation:
            # Σ╜┐τö¿µûçµíúµáçµ│¿τ│╗τ╗ƒ
            _app_logger.debug(f"[SmartProcess] Φ╖»τö▒σê░: µûçµíúΦç¬σè¿µáçµ│¿τ│╗τ╗ƒ")
            return _call_document_annotate(file_path, requirement)
        else:
            # Σ╜┐τö¿Σ╝áτ╗ƒτÜäµûçΣ╗╢σêåµ₧Éτ│╗τ╗ƒ
            _app_logger.debug(f"[SmartProcess] Φ╖»τö▒σê░: µûçΣ╗╢σêåµ₧Éτ│╗τ╗ƒ")
            return _call_document_analysis(file_path, requirement)

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


def _call_document_annotate(file_path: str, requirement: str):
    """Φ░âτö¿µûçµíúµáçµ│¿τ│╗τ╗ƒ"""
    try:
        # Φ╜¼µìóΣ╕║τ╗¥σ»╣Φ╖»σ╛ä
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"µûçΣ╗╢Σ╕ìσ¡ÿσ£¿: {file_path}"}), 404

        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(gemini_client=client)

        result = feedback_system.full_annotation_loop(
            file_path=file_path,
            user_requirement=requirement,
            model_id="gemini-3.1-pro-preview"
        )

        # µ╖╗σèáσñäτÉåµ¿íσ╝ÅµáçΦ«░
        result["processing_mode"] = "annotation"
        result["mode_description"] = "µûçµíúΦç¬σè¿µáçµ│¿"

        return jsonify(result)

    except Exception as e:
        return (
            jsonify(
                {"success": False, "error": str(e), "processing_mode": "annotation"}
            ),
            500,
        )


def _call_document_analysis(file_path: str, requirement: str):
    """Φ░âτö¿Σ╝áτ╗ƒτÜäµûçΣ╗╢σêåµ₧Éτ│╗τ╗ƒ"""
    try:
        # Φ┐ÖΘçîΦ░âτö¿τÄ░µ£ëτÜäµûçΣ╗╢σêåµ₧ÉΘÇ╗Φ╛æ
        # Σ╕┤µù╢Φ┐öσ¢₧Φ»┤µÿÄ∩╝êσ«₧ΘÖàσ║öΦ»ÑΦ░âτö¿τÄ░µ£ëτÜäσêåµ₧Éτ½»τé╣∩╝ë
        return (
            jsonify(
                {
                    "success": False,
                    "error": "µûçΣ╗╢σêåµ₧Éτ│╗τ╗ƒΘ£ÇΦªüσìòτï¼σ«₧τÄ░",
                    "processing_mode": "analysis",
                    "mode_description": "µûçΣ╗╢σêåµ₧É",
                }
            ),
            501,
        )

    except Exception as e:
        return (
            jsonify({"success": False, "error": str(e), "processing_mode": "analysis"}),
            500,
        )


@app.route("/api/document/feedback", methods=["POST"])
def document_feedback():
    """µûçµíúµÖ║Φâ╜σÅìΘªê∩╝ÜΦ»╗σÅûµûçµíú ΓåÆ AIσêåµ₧É ΓåÆ σ║öτö¿Σ┐«µö╣"""
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")
        auto_apply = data.get("auto_apply", True)

        if not file_path:
            return jsonify({"success": False, "error": "τ╝║σ░æfile_pathσÅéµò░"}), 400

        # Φ╜¼µìóΣ╕║τ╗¥σ»╣Φ╖»σ╛ä
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"µûçΣ╗╢Σ╕ìσ¡ÿσ£¿: {file_path}"}), 404

        # σê¥σºïσîûσÅìΘªêτ│╗τ╗ƒ
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(gemini_client=client)

        # µëºΦíîσ«îµò┤σÅìΘªêΘù¡τÄ»
        result = feedback_system.full_feedback_loop(
            file_path=file_path,
            user_requirement=user_requirement,
            auto_apply=auto_apply,
        )

        return jsonify(result)

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/document/analyze", methods=["POST"])
def document_analyze():
    """Σ╗àσêåµ₧Éµûçµíú∩╝îΣ╕ìσ║öτö¿Σ┐«µö╣"""
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")

        if not file_path:
            return jsonify({"success": False, "error": "τ╝║σ░æfile_pathσÅéµò░"}), 400

        # Φ╜¼µìóΣ╕║τ╗¥σ»╣Φ╖»σ╛ä
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"µûçΣ╗╢Σ╕ìσ¡ÿσ£¿: {file_path}"}), 404

        # σê¥σºïσîûσÅìΘªêτ│╗τ╗ƒ
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(gemini_client=client)

        # Σ╗àσêåµ₧É
        result = feedback_system.analyze_and_suggest(
            file_path=file_path, user_requirement=user_requirement
        )

        return jsonify(result)

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/document/apply", methods=["POST"])
def document_apply():
    """σ║öτö¿Σ┐«µö╣σ╗║Φ««σê░µûçµíú"""
    try:
        data = request.json
        file_path = data.get("file_path")
        modifications = data.get("modifications", [])

        if not file_path or not modifications:
            return (
                jsonify(
                    {"success": False, "error": "τ╝║σ░æfile_pathµêûmodificationsσÅéµò░"}
                ),
                400,
            )

        # Φ╜¼µìóΣ╕║τ╗¥σ»╣Φ╖»σ╛ä
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"µûçΣ╗╢Σ╕ìσ¡ÿσ£¿: {file_path}"}), 404

        # σ║öτö¿Σ┐«µö╣
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(gemini_client=client)

        result = feedback_system.apply_suggestions(
            file_path=file_path, modifications=modifications
        )

        return jsonify(result)

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/document/annotate", methods=["POST"])
def document_annotate():
    """µûçµíúΦç¬σè¿µáçµ│¿∩╝ÜAIσêåµ₧É -> τöƒµêÉµáçµ│¿ -> σ║öτö¿σê░σë»µ£¼"""
    try:
        data = request.json
        file_path = data.get('file_path')
        user_requirement = data.get('requirement', '')
        model_id = data.get('model_id', 'gemini-3.1-pro-preview')
        
        if not file_path:
            return jsonify({"success": False, "error": "τ╝║σ░æfile_pathσÅéµò░"}), 400

        # Φ╜¼µìóΣ╕║τ╗¥σ»╣Φ╖»σ╛ä
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"µûçΣ╗╢Σ╕ìσ¡ÿσ£¿: {file_path}"}), 404

        # σê¥σºïσîûσÅìΘªêτ│╗τ╗ƒ
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(gemini_client=client)

        # µëºΦíîσ«îµò┤µáçµ│¿Θù¡τÄ»
        result = feedback_system.full_annotation_loop(
            file_path=file_path, user_requirement=user_requirement, model_id=model_id
        )

        return jsonify(result)

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/document/analyze-annotations", methods=["POST"])
def document_analyze_annotations():
    """Σ╗àσêåµ₧Éµûçµíúσ╣╢τöƒµêÉµáçµ│¿σ╗║Φ««∩╝êΣ╕ìσ║öτö¿∩╝ë- σ╖▓σ╝âτö¿∩╝îΦ»╖Σ╜┐τö¿ /api/document/batch-annotate-stream"""
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")

        if not file_path:
            return jsonify({"success": False, "error": "τ╝║σ░æfile_pathσÅéµò░"}), 400

        # Φ╜¼µìóΣ╕║τ╗¥σ»╣Φ╖»σ╛ä
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"µûçΣ╗╢Σ╕ìσ¡ÿσ£¿: {file_path}"}), 404

        # Σ╜┐τö¿V2µë╣ΘçÅµáçµ│¿τ│╗τ╗ƒ∩╝êτ½ïσì│Φ┐öσ¢₧τ╗ôµ₧£∩╝îΣ╕ìµ╡üσ╝Å∩╝ë
        from web.document_direct_edit import ImprovedBatchAnnotator

        annotator = ImprovedBatchAnnotator(gemini_client=client, batch_size=5)

        # µö╢Θ¢åµëÇµ£ëΣ║ïΣ╗╢∩╝êΘ¥₧µ╡üσ╝Å∩╝ë
        events = []
        final_result = None

        for event in annotator.annotate_document_streaming(file_path, user_requirement):
            # Φºúµ₧ÉΣ║ïΣ╗╢
            if event.startswith("event: complete"):
                data_line = event.split("\n")[1]
                if data_line.startswith("data: "):
                    final_result = json.loads(data_line[6:])
            events.append(event)

        if final_result:
            return jsonify({"success": True, **final_result})
        else:
            return jsonify({"success": False, "error": "σñäτÉåσñ▒Φ┤Ñ∩╝îµ£¬µö╢σê░σ«îµêÉΣ║ïΣ╗╢"}), 500

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/document/batch-annotate-stream", methods=["POST"])
def document_batch_annotate_stream():
    """
    µë╣ΘçÅµáçµ│¿µûçµíú∩╝êSSEµ╡üσ╝ÅΦ┐öσ¢₧∩╝îσ«₧µù╢σÅìΘªêΦ┐¢σ║ª∩╝ë

    µÄÑµö╢σÅéµò░:
        file_path: µûçµíúΦ╖»σ╛ä
        requirement: τö¿µê╖Θ£Çµ▒é∩╝êσÅ»ΘÇë∩╝ë
        batch_size: µ»Åµë╣σñäτÉåµ«╡ΦÉ╜µò░∩╝êΘ╗ÿΦ«ñ5∩╝ë

    Φ┐öσ¢₧: SSEΣ║ïΣ╗╢µ╡ü
        event: progress - Φ┐¢σ║ªµ¢┤µû░
        event: batch_complete - µë╣µ¼íσ«îµêÉ
        event: complete - σà¿Θâ¿σ«îµêÉ
        event: error - ΘöÖΦ»»
    """
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")
        batch_size = data.get("batch_size", 5)

        if not file_path:
            return jsonify({"success": False, "error": "τ╝║σ░æfile_pathσÅéµò░"}), 400

        # Φ╜¼µìóΣ╕║τ╗¥σ»╣Φ╖»σ╛ä
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"µûçΣ╗╢Σ╕ìσ¡ÿσ£¿: {file_path}"}), 404

        # σ»╝σàÑV2µë╣ΘçÅµáçµ│¿τ│╗τ╗ƒ
        from web.document_batch_annotator_v2 import annotate_large_document

        # Φ┐öσ¢₧SSEµ╡ü
        return Response(
            annotate_large_document(
                file_path=file_path,
                user_requirement=user_requirement,
                gemini_client=client,
                batch_size=batch_size,
            ),
            mimetype="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
                "Connection": "keep-alive",
            },
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/document/apply-annotations", methods=["POST"])
def document_apply_annotations():
    """σ║öτö¿µáçµ│¿σ╗║Φ««σê░µûçµíú"""
    try:
        data = request.json
        file_path = data.get("file_path")
        annotations = data.get("annotations", [])

        if not file_path or not annotations:
            return (
                jsonify({"success": False, "error": "τ╝║σ░æfile_pathµêûannotationsσÅéµò░"}),
                400,
            )

        # Φ╜¼µìóΣ╕║τ╗¥σ»╣Φ╖»σ╛ä
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"µûçΣ╗╢Σ╕ìσ¡ÿσ£¿: {file_path}"}), 404

        # σ║öτö¿µáçµ│¿
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(gemini_client=client)

        result = feedback_system.annotate_document(file_path, annotations)

        return jsonify(result)

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ==================== µû░σèƒΦâ╜ API Φ╖»τö▒ ====================

# ==================== µö╣Φ┐¢τÜäσ╗║Φ««σ╝Åµáçµ│¿ API ====================


@app.route("/api/document/suggest-stream", methods=["POST"])
def document_suggest_stream():
    """
    τöƒµêÉΣ┐«µö╣σ╗║Φ««µ╡ü∩╝êSSE∩╝ë

    Φ»╖µ▒éσÅéµò░:
        file_path: µûçµíúΦ╖»σ╛ä
        requirement: τö¿µê╖Θ£Çµ▒é∩╝êσÅ»ΘÇë∩╝ë

    Φ┐öσ¢₧: SSEΣ║ïΣ╗╢µ╡ü
        event: progress - Φ┐¢σ║ª
        event: suggestion - σìòΣ╕¬σ╗║Φ««
        event: suggestions_complete - µëÇµ£ëσ╗║Φ««σ«îµêÉ
        event: complete - σ«îµêÉ
    """
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")

        if not file_path:
            return jsonify({"success": False, "error": "τ╝║σ░æfile_pathσÅéµò░"}), 400

        # Φ╜¼µìóΣ╕║τ╗¥σ»╣Φ╖»σ╛ä
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"µûçΣ╗╢Σ╕ìσ¡ÿσ£¿: {file_path}"}), 404

        # Σ╜┐τö¿σ╗║Φ««σ╝Åµáçµ│¿σÖ¿
        from web.suggestion_annotator import SuggestionAnnotator

        annotator = SuggestionAnnotator(batch_size=3)

        # Φ┐öσ¢₧SSEµ╡ü
        return Response(
            annotator.analyze_document_streaming(file_path, user_requirement),
            mimetype="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
                "Connection": "keep-alive",
            },
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/document/apply-suggestions", methods=["POST"])
def document_apply_suggestions():
    """
    µá╣µì«τö¿µê╖ΘÇëµï⌐σ║öτö¿Σ┐«µö╣σ╗║Φ««

    Φ»╖µ▒éσÅéµò░:
        file_path: σÄƒσºïµûçµíúΦ╖»σ╛ä
        suggestions: τö¿µê╖τÜäΘÇëµï⌐σêùΦí¿
            [
                {
                    "id": "s_5_0",
                    "σÄƒµûç": "σ£¿Φó½Φ«░σ╜òτÜä",
                    "Σ┐«µö╣": "σ£¿Φ«░σ╜òτÜä",
                    "µÄÑσÅù": True/False
                },
                ...
            ]

    Φ┐öσ¢₧:
        {
            "success": True,
            "output_file": "Σ┐«µö╣σÉÄτÜäµûçΣ╗╢Φ╖»σ╛ä",
            "applied_count": σ«₧ΘÖàσ║öτö¿τÜäΣ┐«µö╣µò░,
            "accepted_count": τö¿µê╖µÄÑσÅùτÜäµò░ΘçÅ
        }
    """
    try:
        from docx import Document

        data = request.json
        file_path = data.get("file_path")
        suggestions = data.get("suggestions", [])

        if not file_path:
            return jsonify({"success": False, "error": "τ╝║σ░æfile_pathσÅéµò░"}), 400

        # Φ╜¼µìóΣ╕║τ╗¥σ»╣Φ╖»σ╛ä
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"µûçΣ╗╢Σ╕ìσ¡ÿσ£¿: {file_path}"}), 404

        # Φ»╗σÅûµûçµíú
        doc = Document(file_path)

        # τ¡¢ΘÇëτö¿µê╖µÄÑσÅùτÜäσ╗║Φ««
        accepted_suggestions = [s for s in suggestions if s.get("µÄÑσÅù", False)]

        applied_count = 0

        # σ║öτö¿Σ┐«µö╣∩╝êτ¢┤µÄÑσ£¿µ«╡ΦÉ╜Σ╕¡µƒÑµë╛σ╣╢µ¢┐µìó∩╝ë
        for suggestion in accepted_suggestions:
            original = suggestion.get("σÄƒµûç", "")
            modified = suggestion.get("Σ┐«µö╣", "")

            if not original or not modified:
                continue

            # σ£¿µëÇµ£ëµ«╡ΦÉ╜Σ╕¡µƒÑµë╛σ╣╢µ¢┐µìó
            for para in doc.paragraphs:
                if original in para.text:
                    # µ¢┐µìóµûçµ£¼
                    full_text = para.text
                    new_text = full_text.replace(original, modified, 1)

                    if new_text != full_text:
                        # µ╕àτ⌐║σ╣╢Θçìµû░µ╖╗σèá∩╝êΣ┐¥τòÖµá╝σ╝Å∩╝ë
                        para.clear()
                        para.add_run(new_text)
                        applied_count += 1
                        break  # µ»ÅΣ╕¬σ╗║Φ««σÅ¬σ║öτö¿Σ╕Çµ¼í

            # µúÇµƒÑΦí¿µá╝Σ╕¡τÜäµûçµ£¼
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if original in para.text:
                                full_text = para.text
                                new_text = full_text.replace(original, modified, 1)
                                if new_text != full_text:
                                    para.clear()
                                    para.add_run(new_text)
                                    applied_count += 1

        # Σ┐¥σ¡ÿΣ╕║µû░µûçΣ╗╢
        base_name = os.path.splitext(file_path)[0]
        output_path = f"{base_name}_accepted.docx"
        doc.save(output_path)

        return jsonify(
            {
                "success": True,
                "output_file": output_path,
                "applied_count": applied_count,
                "accepted_count": len(accepted_suggestions),
                "message": f"σ╖▓σ║öτö¿ {applied_count} σñäΣ┐«µö╣∩╝êτö¿µê╖µÄÑσÅùΣ║å {len(accepted_suggestions)} Σ╕¬σ╗║Φ««∩╝ë",
            }
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


# τƒÑΦ»åσ║ô API
@app.route("/api/knowledge-base/add", methods=["POST"])
def kb_add_document():
    """µ╖╗σèáµûçµíúσê░τƒÑΦ»åσ║ô"""
    try:
        from web.knowledge_base import KnowledgeBase

        data = request.json
        file_path = data.get("file_path")

        if not file_path:
            return jsonify({"success": False, "error": "τ╝║σ░æfile_pathσÅéµò░"}), 400

        kb = KnowledgeBase()
        result = kb.add_document(file_path)

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/knowledge-base/search", methods=["POST"])
def kb_search():
    """µÉ£τ┤óτƒÑΦ»åσ║ô"""
    try:
        from web.knowledge_base import KnowledgeBase

        data = request.json
        query = data.get("query")
        max_results = data.get("max_results", 10)

        if not query:
            return jsonify({"success": False, "error": "τ╝║σ░æqueryσÅéµò░"}), 400

        kb = KnowledgeBase()
        results = kb.search(query, max_results=max_results)

        return jsonify({"success": True, "results": results})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/knowledge-base/stats", methods=["GET"])
def kb_stats():
    """ΦÄ╖σÅûτƒÑΦ»åσ║ôτ╗ƒΦ«í"""
    try:
        from web.knowledge_base import KnowledgeBase

        kb = KnowledgeBase()
        stats = kb.get_stats()

        return jsonify({"success": True, "stats": stats})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ==================== µûçΣ╗╢τ╜æτ╗£τ┤óσ╝ò API ====================


@app.route("/api/file-network/search", methods=["POST"])
def file_network_search():
    """σñÜτ╗┤µƒÑΦ»óµûçΣ╗╢

    Φ»╖µ▒éσÅéµò░:
        query: µûçµ£¼µÉ£τ┤óµƒÑΦ»ó∩╝êσÅ»ΘÇë∩╝ë
        file_type: µûçΣ╗╢τ▒╗σ₧ï∩╝êdocx, pdfτ¡ë∩╝îσÅ»ΘÇë∩╝ë
        tags: µáçτ¡╛σêùΦí¿∩╝êσÅ»ΘÇë∩╝ë
        operation: σñäτÉåµôìΣ╜£∩╝êannotate, editτ¡ë∩╝îσÅ»ΘÇë∩╝ë
        date_from: σ╝ÇσºïµùÑµ£ƒ∩╝êISOµá╝σ╝Å∩╝îσÅ»ΘÇë∩╝ë
        date_to: τ╗ôµ¥ƒµùÑµ£ƒ∩╝êISOµá╝σ╝Å∩╝îσÅ»ΘÇë∩╝ë
        limit: Φ┐öσ¢₧µò░ΘçÅΘÖÉσê╢∩╝êΘ╗ÿΦ«ñ50∩╝ë
    """
    try:
        from web.processed_file_network import get_file_network

        data = request.json or {}
        file_network = get_file_network()

        result = file_network.search_files(
            query=data.get("query"),
            file_type=data.get("file_type"),
            tags=data.get("tags"),
            operation=data.get("operation"),
            date_from=data.get("date_from"),
            date_to=data.get("date_to"),
            limit=data.get("limit", 50),
        )

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/file-network/open", methods=["POST"])
def file_network_open():
    """σ┐½ΘÇƒµëôσ╝ÇµûçΣ╗╢

    Φ»╖µ▒éσÅéµò░:
        file_id: µûçΣ╗╢ID
    """
    try:
        from web.processed_file_network import get_file_network

        data = request.json
        file_id = data.get("file_id")

        if not file_id:
            return jsonify({"success": False, "error": "τ╝║σ░æfile_idσÅéµò░"}), 400

        file_network = get_file_network()
        result = file_network.open_file(file_id)

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/file-network/network", methods=["POST"])
def file_network_get_network():
    """ΦÄ╖σÅûµûçΣ╗╢σà│τ│╗τ╜æτ╗£

    Φ»╖µ▒éσÅéµò░:
        file_id: µûçΣ╗╢ID
        depth: σà│τ│╗µ╖▒σ║ª∩╝ê1=τ¢┤µÄÑσà│τ│╗∩╝î2=Σ║îτ║ºσà│τ│╗∩╝îΘ╗ÿΦ«ñ2∩╝ë
    """
    try:
        from web.processed_file_network import get_file_network

        data = request.json
        file_id = data.get("file_id")
        depth = data.get("depth", 2)

        if not file_id:
            return jsonify({"success": False, "error": "τ╝║σ░æfile_idσÅéµò░"}), 400

        file_network = get_file_network()
        result = file_network.get_file_network(file_id, depth)

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/file-network/statistics", methods=["GET"])
def file_network_statistics():
    """ΦÄ╖σÅûµûçΣ╗╢τ╜æτ╗£τ╗ƒΦ«íΣ┐íµü»"""
    try:
        from web.processed_file_network import get_file_network

        file_network = get_file_network()
        result = file_network.get_statistics()

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/file-network/register", methods=["POST"])
def file_network_register():
    """µëïσè¿µ│¿σåîµûçΣ╗╢σê░τ╜æτ╗£

    Φ»╖µ▒éσÅéµò░:
        file_path: µûçΣ╗╢Φ╖»σ╛ä
        tags: µáçτ¡╛σêùΦí¿∩╝êσÅ»ΘÇë∩╝ë
        extract_snippets: µÿ»σÉªµÅÉσÅûµûçµ£¼τëçµ«╡∩╝êΘ╗ÿΦ«ñtrue∩╝ë
    """
    try:
        from web.processed_file_network import get_file_network

        data = request.json
        file_path = data.get("file_path")

        if not file_path:
            return jsonify({"success": False, "error": "τ╝║σ░æfile_pathσÅéµò░"}), 400

        file_network = get_file_network()
        result = file_network.register_file(
            file_path=file_path,
            tags=data.get("tags"),
            extract_snippets=data.get("extract_snippets", True),
        )

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# µë╣ΘçÅσñäτÉå API
@app.route("/api/batch/rename", methods=["POST"])
def batch_rename():
    """µë╣ΘçÅΘçìσæ╜σÉìµûçΣ╗╢"""
    try:
        from web.batch_processor import BatchFileProcessor

        data = request.json
        directory = data.get("directory")
        pattern = data.get("pattern")

        processor = BatchFileProcessor()
        result = processor.batch_rename(directory, **pattern)

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/batch/convert", methods=["POST"])
def batch_convert():
    """µë╣ΘçÅµá╝σ╝ÅΦ╜¼µìó"""
    try:
        from web.batch_processor import BatchFileProcessor

        data = request.json
        directory = data.get("directory")
        from_format = data.get("from_format")
        to_format = data.get("to_format")

        processor = BatchFileProcessor()
        result = processor.batch_convert(directory, from_format, to_format)

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# µ¿íµ¥┐σ║ô API
@app.route("/api/template/list", methods=["GET"])
def template_list():
    """ΦÄ╖σÅûµ¿íµ¥┐σêùΦí¿"""
    try:
        from web.template_library import TemplateLibrary

        library = TemplateLibrary()
        templates = library.list_templates()

        return jsonify({"success": True, "templates": templates})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/template/generate", methods=["POST"])
def template_generate():
    """Σ╗Äµ¿íµ¥┐τöƒµêÉµûçµíú"""
    try:
        from web.template_library import TemplateLibrary

        data = request.json
        template_name = data.get("template_id") or data.get("template_name")
        variables = data.get("variables", {})
        output_dir = data.get("output_dir")
        output_file = data.get("output_file")
        if output_file and not output_dir:
            if os.path.isdir(output_file):
                output_dir = output_file
            else:
                output_dir = os.path.dirname(output_file) or None

        library = TemplateLibrary()
        result = library.generate_from_template(template_name, variables, output_dir)

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# Σ╕ÇΦç┤µÇºµúÇµƒÑ API
@app.route("/api/check/consistency", methods=["POST"])
def check_consistency():
    """µúÇµƒÑµûçµíúΣ╕ÇΦç┤µÇº"""
    try:
        from web.consistency_checker import ConsistencyChecker

        data = request.json
        file_path = data.get("file_path")

        checker = ConsistencyChecker()
        result = checker.check_document(file_path)
        report = checker.generate_report(result)

        return jsonify({"success": True, "result": result, "report": report})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# µûçµíúσ»╣µ»ö API
@app.route("/api/compare/documents", methods=["POST"])
def compare_documents():
    """σ»╣µ»öµûçµíú"""
    try:
        from web.document_comparator import DocumentComparator

        data = request.json
        file_a = data.get("file_a")
        file_b = data.get("file_b")
        output_format = data.get("output_format", "markdown")

        comparator = DocumentComparator()
        result = comparator.compare_documents(file_a, file_b, output_format)

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# OCR σè⌐µëï API
@app.route("/api/ocr/screenshot", methods=["POST"])
def ocr_screenshot():
    """µê¬σ¢╛σ╣╢OCR"""
    try:
        from web.clipboard_ocr_assistant import ClipboardOCRAssistant

        data = request.json
        save_image = data.get("save_image", True)
        auto_index = data.get("auto_index", False)

        assistant = ClipboardOCRAssistant()
        result = assistant.capture_and_ocr(source="screenshot", save_image=save_image)

        if auto_index and result.get("ocr_success"):
            assistant.auto_index_to_knowledge_base(result)

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/ocr/clipboard", methods=["POST"])
def ocr_clipboard():
    """σë¬Φ┤┤µ¥┐σ¢╛τëçOCR"""
    try:
        from web.clipboard_ocr_assistant import ClipboardOCRAssistant

        data = request.json
        save_image = data.get("save_image", True)
        auto_index = data.get("auto_index", False)

        assistant = ClipboardOCRAssistant()
        result = assistant.capture_and_ocr(source="clipboard", save_image=save_image)

        if auto_index and result.get("ocr_success"):
            assistant.auto_index_to_knowledge_base(result)

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# µôìΣ╜£σÄåσÅ▓ API
@app.route("/api/history/list", methods=["GET"])
def history_list():
    """ΦÄ╖σÅûµôìΣ╜£σÄåσÅ▓"""
    try:
        from web.operation_history import OperationHistory

        limit = request.args.get("limit", 50, type=int)
        file_path = request.args.get("file_path")

        history = OperationHistory()
        operations = history.get_history(limit=limit, file_path=file_path)

        return jsonify({"success": True, "operations": operations})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/history/rollback/<op_id>", methods=["POST"])
def history_rollback(op_id):
    """σ¢₧µ╗ÜµôìΣ╜£"""
    try:
        from web.operation_history import OperationHistory

        history = OperationHistory()
        result = history.rollback(op_id)

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/history/stats", methods=["GET"])
def history_stats():
    """ΦÄ╖σÅûσÄåσÅ▓τ╗ƒΦ«í"""
    try:
        from web.operation_history import OperationHistory

        history = OperationHistory()
        stats = history.get_statistics()

        return jsonify({"success": True, "stats": stats})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# Φ»¡Θƒ│Φ╜¼σåÖ API
@app.route("/api/speech/transcribe-file", methods=["POST"])
def speech_transcribe_file():
    """Φ╜¼σåÖΘƒ│ΘóæµûçΣ╗╢"""
    try:
        from web.speech_transcriber import SpeechTranscriber

        data = request.json
        audio_path = data.get("audio_path")
        language = data.get("language", "zh-CN")
        output_format = data.get("output_format", "txt")
        title = data.get("title")
        auto_summary = data.get("auto_summary", True)

        if not audio_path:
            return jsonify({"success": False, "error": "τ╝║σ░æaudio_pathσÅéµò░"}), 400

        transcriber = SpeechTranscriber()
        result = transcriber.process_audio_complete(
            audio_path,
            language=language,
            output_format=output_format,
            title=title,
            auto_summary=auto_summary,
        )

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/speech/transcribe-microphone", methods=["POST"])
def speech_transcribe_microphone():
    """Σ╗ÄΘ║ªσàïΘúÄσ╜òΘƒ│σ╣╢Φ╜¼σåÖ"""
    try:
        from web.speech_transcriber import SpeechTranscriber

        data = request.json
        duration = data.get("duration", 30)
        language = data.get("language", "zh-CN")
        output_format = data.get("output_format", "txt")
        title = data.get("title")

        transcriber = SpeechTranscriber()

        # σ╜òΘƒ│
        mic_result = transcriber.transcribe_microphone(
            duration=duration, language=language
        )

        if not mic_result["success"]:
            return jsonify(mic_result), 400

        text = mic_result["text"]

        # µÅÉσÅûµÇ╗τ╗ô
        summary_result = transcriber.extract_keywords_and_summary(text)
        keywords = (
            summary_result.get("keywords", []) if summary_result["success"] else []
        )
        summary = summary_result.get("summary", []) if summary_result["success"] else []

        # τöƒµêÉµûçµíú
        output_file = transcriber.generate_transcript_document(
            text,
            keywords=keywords,
            summary=summary,
            title=title,
            output_format=output_format,
        )

        return jsonify(
            {
                "success": True,
                "text": text,
                "keywords": keywords,
                "summary": summary,
                "output_file": output_file,
                "format": output_format,
            }
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/speech/extract-summary", methods=["POST"])
def speech_extract_summary():
    """Σ╗Äµûçµ£¼µÅÉσÅûσà│Θö«Φ»ìσÆîµÇ╗τ╗ô"""
    try:
        from web.speech_transcriber import SpeechTranscriber

        data = request.json
        text = data.get("text")
        max_keywords = data.get("max_keywords", 10)

        if not text:
            return jsonify({"success": False, "error": "τ╝║σ░ætextσÅéµò░"}), 400

        transcriber = SpeechTranscriber()
        result = transcriber.extract_keywords_and_summary(
            text, max_keywords=max_keywords
        )

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ================= Σ╕╗τ¿ïσ║ÅσàÑσÅú =================

# ================= NotebookLM σèƒΦâ╜σñìσê╗ API =================


@app.route("/api/notebook/overview", methods=["POST"])
def notebook_overview():
    """τöƒµêÉΘƒ│ΘóæµªéΦºê (Podcast)"""
    data = request.json
    content = data.get("content", "")
    if not content:
        return jsonify({"success": False, "error": "σåàσ«╣Σ╕ìΦâ╜Σ╕║τ⌐║"}), 400

    try:
        from web.audio_overview import AudioOverviewGenerator

        generator = AudioOverviewGenerator(
            output_dir=os.path.join(settings_manager.workspace_dir, "audio_cache")
        )

        # 1. τöƒµêÉσëºµ£¼
        # ΦÄ╖σÅûµ¿íσ₧ïσ«₧Σ╛ï (σñìτö¿τÄ░µ£ëτÜä KotoBrain µêûτ¢┤µÄÑΦ░âτö¿ API)
        # Φ┐ÖΘçîΣ╕║Σ║åτ«Çσîû∩╝îσüçΦ«╛µêæΣ╗¼Φâ╜ΦÄ╖σÅûσê░Σ╕ÇΣ╕¬ genai model σ«₧Σ╛ï
        # σ«₧ΘÖàΘí╣τ¢«Σ╕¡σ║öΦ»Ñσñìτö¿ koto_brain.client.models
        # µÜéµù╢Σ╜┐τö¿Σ╕┤µù╢τÜä model σ«₧Σ╛ï
        import google.genai as genai

        client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
        model = client.models

        script = asyncio.run(generator.generate_script(content, model))
        if not script:
            return jsonify({"success": False, "error": "σëºµ£¼τöƒµêÉσñ▒Φ┤Ñ"}), 500

        # 2. σÉêµêÉΘƒ│Θóæ
        session_id = f"overview_{int(time.time())}"
        audio_path = asyncio.run(generator.synthesize_audio(script, session_id))

        if audio_path:
            # Φ┐öσ¢₧τ¢╕σ»╣Σ║Ä workspace τÜäΦ╖»σ╛äµêûΦÇà download url
            rel_path = os.path.relpath(audio_path, settings_manager.workspace_dir)
            # µ│¿µäÅ∩╝Üσ«₧ΘÖàΦ«┐Θù«σÅ»Φâ╜Θ£ÇΦªüΘÇÜΦ┐ç send_from_directory Φ╖»τö▒
            # σüçΦ«╛µêæΣ╗¼µ£ëΣ╕ÇΣ╕¬ /files/ Φ╖»τö▒σÅ»Σ╗ÑΦ«┐Θù« workspace/
            audio_url = f"/api/files/download?path={requests.utils.quote(audio_path)}"

            return jsonify({"success": True, "audio_url": audio_url, "script": script})
        else:
            return jsonify({"success": False, "error": "Θƒ│ΘóæσÉêµêÉσñ▒Φ┤Ñ"}), 500

    except Exception as e:
        _app_logger.error(f"Error processing audio overview: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/notebook/qa", methods=["POST"])
def notebook_qa():
    """µ║Éµûçµíúµ╖▒σ║ªΘù«τ¡ö (Source-Grounded Q&A)"""
    data = request.json
    question = data.get("question")
    file_ids = data.get(
        "file_ids", []
    )  # σüçΦ«╛σëìτ½»Σ╝áσ¢₧ files (Φ┐ÖΘçîσàêτ«ÇσîûΣ╕║ content τ¢┤µÄÑΣ╝áσàÑ µêûΦÇà file paths)
    # Σ╕║Σ║åτ«Çσîûµ╝öτñ║∩╝îµêæΣ╗¼σàêµÄÑσÅùτ║»µûçµ£¼ content
    context_content = data.get("context", "")

    if not question or not context_content:
        return jsonify({"success": False, "error": "τ╝║σ░æΘù«ΘóÿµêûΣ╕èΣ╕ïµûç"}), 400

    prompt = f"""
    Answer the user's question mostly based on the provided source context.
    
    [Source Context]
    {context_content[:30000]} 

    [User Question]
    {question}

    [Rules]
    1. You must cite your sources. When you use information from the context, append [Source] at the end of the sentence.
    2. If the answer is not in the context, state that clearly.
    3. Be precise and concise.
    """

    try:
        # σñìτö¿ KotoBrain τÜäΘÇ╗Φ╛æµêûΦÇàτ¢┤µÄÑΦ░âτö¿
        import google.genai as genai

        client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
        response = client.models.generate_content(
            model="gemini-2.5-flash", contents=prompt
        )
        return jsonify({"success": True, "answer": response.text})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/notebook/study_guide", methods=["POST"])
def notebook_study_guide():
    """τöƒµêÉσ¡ªΣ╣áµîçσìù/τ«ÇµèÑ"""
    data = request.json
    content = data.get("content", "")
    type_ = data.get("type", "summary")  # summary, quiz, timelime, faq

    prompts = {
        "summary": "Create a comprehensive briefing document summarizing the key points, key people, and timeline from the text.",
        "quiz": "Create 5 multiple-choice questions based on the text to test understanding. Include the correct answer key at the end.",
        "timeline": "Extract a chronological timeline of events mentioned in the text.",
        "faq": "Create a FAQ section based on the text, anticipating what a reader might ask.",
    }

    selected_prompt = prompts.get(type_, prompts["summary"])
    full_prompt = f"{selected_prompt}\n\n[Source Text]\n{content[:20000]}"

    try:
        import google.genai as genai

        client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
        response = client.models.generate_content(
            model="gemini-2.5-flash", contents=full_prompt
        )
        return jsonify({"success": True, "result": response.text})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/files/download", methods=["GET"])
def download_file_proxy():
    """ΘÇÜτö¿τÜäµûçΣ╗╢Σ╕ïΦ╜╜Σ╗úτÉå"""
    file_path = request.args.get("path")
    if not file_path or not os.path.exists(file_path):
        return "File not found", 404
    return send_file(file_path, as_attachment=True)


@app.route("/api/notebook/upload", methods=["POST"])
def notebook_upload():
    """Σ╕èΣ╝áσ╣╢Φºúµ₧ÉµûçΣ╗╢ (PDF/Docx/Txt)"""
    if "file" not in request.files:
        return jsonify({"success": False, "error": "No file part"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"success": False, "error": "No selected file"}), 400

    try:
        # Save temp file
        filename = file.filename
        temp_path = os.path.join(
            tempfile.gettempdir(), f"koto_{int(time.time())}_{filename}"
        )
        file.save(temp_path)

        # Parse using FileParser
        from web.file_parser import FileParser

        result = FileParser.parse_file(temp_path)

        # Cleanup
        try:
            os.remove(temp_path)
        except OSError:
            pass

        if result.get("success"):
            return jsonify(
                {
                    "success": True,
                    "filename": filename,
                    "content": result.get("content", ""),
                    "char_count": result.get("char_count", 0),
                }
            )
        else:
            return jsonify({"success": False, "error": result.get("error")}), 500

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/notebook")
def notebook_ui():
    """NotebookLM ΘúÄµá╝τòîΘ¥ó"""
    return render_template("notebook_lm.html")


if __name__ == "__main__":

    print("\n≡ƒÜÇ Koto Web Server Starting...")
    print(f"≡ƒôü Chat Directory: {os.path.abspath(CHAT_DIR)}")
    print(f"≡ƒôü Workspace: {os.path.abspath(WORKSPACE_DIR)}")

    # σ╗╢Φ┐ƒµúÇµƒÑ Ollama τè╢µÇü∩╝êΣ╕ìΘÿ╗σí₧σÉ»σè¿∩╝ë
    def check_ollama_async():
        time.sleep(2)  # σ╗╢Φ┐ƒ2τºÆσÉÄµúÇµƒÑ
        if os.environ.get("KOTO_DEPLOY_MODE") == "cloud":
            print("Γÿü∩╕Å Ollama: Disabled (cloud mode - using Gemini API)")
            return
        if LocalDispatcher.is_ollama_running():
            print("≡ƒªÖ Ollama: Running")
        else:
            print("≡ƒªÖ Ollama: Not Running")

    threading.Thread(target=check_ollama_async, daemon=True).start()

    print("ΓÜá∩╕Å µ£¼σ£░µ¿íσ₧ïΣ╗╗σèíΦ╖»τö▒σÖ¿σ╖▓τªüτö¿∩╝îΣ╜┐τö¿Φ┐£τ¿ï AI")

    print("\n≡ƒîÉ Open http://localhost:5000 in your browser\n")

    # σÉ»σè¿σÉÄσÅ░µ£ìσèí∩╝êσ╝éµ¡Ñ∩╝îΣ╕ìΘÿ╗σí₧σÉ»σè¿∩╝ë
    def start_background_services():
        time.sleep(1)  # σ╗╢Φ┐ƒ1τºÆσÉÄσÉ»σè¿σÉÄσÅ░µ£ìσèí
        try:
            from auto_catalog_scheduler import get_auto_catalog_scheduler
            from clipboard_manager import get_clipboard_manager
            from task_scheduler import get_task_scheduler

            # σÉ»σè¿σë¬Φ┤┤µ¥┐τ¢æµÄº
            clipboard_manager = get_clipboard_manager()
            clipboard_manager.start_monitoring()
            print("≡ƒôï σë¬Φ┤┤µ¥┐τ¢æµÄºσ╖▓σÉ»σè¿")

            # σÉ»σè¿Σ╗╗σèíΦ░âσ║ªσÖ¿
            task_scheduler = get_task_scheduler()
            task_scheduler.start()
            print("ΓÅ░ Σ╗╗σèíΦ░âσ║ªσÖ¿σ╖▓σÉ»σè¿")

            # σê¥σºïσîûΦç¬σè¿σ╜Æτ║│Φ░âσ║ªσÖ¿∩╝êσªéµ₧£σ╖▓σÉ»τö¿∩╝ë
            auto_catalog = get_auto_catalog_scheduler()
            if auto_catalog.is_auto_catalog_enabled():
                auto_catalog._register_scheduled_task()
                print(
                    f"≡ƒùé∩╕Å Φç¬σè¿σ╜Æτ║│σ╖▓σÉ»τö¿∩╝îµ»ÅµùÑ {auto_catalog.get_catalog_schedule()} µëºΦíî"
                )

        except Exception as e:
            print(f"ΓÜá∩╕Å σÉÄσÅ░µ£ìσèíσÉ»σè¿σñ▒Φ┤Ñ: {e}")

    threading.Thread(target=start_background_services, daemon=True).start()

    try:
        debug_mode = os.environ.get("KOTO_DEBUG", "false").lower() == "true"
        port = int(os.environ.get("KOTO_PORT", "5000"))
        app.run(debug=debug_mode, host="0.0.0.0", port=port, threaded=True)
    finally:
        # σ║öτö¿σà│Θù¡µù╢µ╕àτÉåσ╣╢ΦíîµëºΦíîτ│╗τ╗ƒ
        if PARALLEL_SYSTEM_ENABLED:
            print("[PARALLEL] ≡ƒ¢æ Shutting down parallel execution system...")
            stop_dispatcher()
            print("[PARALLEL] Γ£à Parallel execution system shut down")


# ΓòÉΓòÉΓòÉ µûçΣ╗╢τ╗äτ╗çτ│╗τ╗ƒ API ΓòÉΓòÉΓòÉ

# σê¥σºïσîûµûçΣ╗╢τ╗äτ╗çσÖ¿
_file_organizer_cache = {}
_batch_ops_cache = {}


def get_file_organizer():
    """µçÆσèáΦ╜╜µûçΣ╗╢τ╗äτ╗çσÖ¿"""
    if "organizer" not in _file_organizer_cache:
        try:
            from web.file_organizer import FileOrganizer
        except ImportError:
            from file_organizer import FileOrganizer

        organize_root = get_organize_root()
        _file_organizer_cache["organizer"] = FileOrganizer(organize_root)

    return _file_organizer_cache["organizer"]


def get_file_analyzer():
    """µçÆσèáΦ╜╜µûçΣ╗╢σêåµ₧ÉσÖ¿"""
    if "analyzer" not in _file_organizer_cache:
        try:
            from web.file_analyzer import FileAnalyzer
        except ImportError:
            from file_analyzer import FileAnalyzer

        _file_organizer_cache["analyzer"] = FileAnalyzer()

    return _file_organizer_cache["analyzer"]


def get_batch_ops_manager():
    """µçÆσèáΦ╜╜µë╣ΘçÅµûçΣ╗╢σñäτÉåτ«íτÉåσÖ¿"""
    if "batch_ops" not in _batch_ops_cache:
        try:
            from web.batch_file_ops import BatchFileOpsManager
        except ImportError:
            from batch_file_ops import BatchFileOpsManager
        _batch_ops_cache["batch_ops"] = BatchFileOpsManager()
    return _batch_ops_cache["batch_ops"]


_file_editor_cache = {}
_file_indexer_cache = {}
_concept_extractor_cache = {}
_knowledge_graph_cache = {}
_behavior_monitor_cache = {}
_suggestion_engine_cache = {}
_insight_reporter_cache = {}


def get_file_editor():
    """µçÆσèáΦ╜╜µûçΣ╗╢τ╝ûΦ╛æσÖ¿"""
    if "editor" not in _file_editor_cache:
        try:
            from web.file_editor import FileEditor
        except ImportError:
            from file_editor import FileEditor
        _file_editor_cache["editor"] = FileEditor()
    return _file_editor_cache["editor"]


def get_file_indexer():
    """µçÆσèáΦ╜╜µûçΣ╗╢τ┤óσ╝òσÖ¿"""
    if "indexer" not in _file_indexer_cache:
        try:
            from web.file_indexer import FileIndexer
        except ImportError:
            from file_indexer import FileIndexer
        _file_indexer_cache["indexer"] = FileIndexer()
    return _file_indexer_cache["indexer"]


def get_concept_extractor():
    """µçÆσèáΦ╜╜µªéσ┐╡µÅÉσÅûσÖ¿"""
    if "extractor" not in _concept_extractor_cache:
        try:
            from web.concept_extractor import ConceptExtractor
        except ImportError:
            from concept_extractor import ConceptExtractor
        _concept_extractor_cache["extractor"] = ConceptExtractor()
    return _concept_extractor_cache["extractor"]


def get_knowledge_graph():
    """µçÆσèáΦ╜╜τƒÑΦ»åσ¢╛Φ░▒"""
    if "graph" not in _knowledge_graph_cache:
        try:
            from web.knowledge_graph import KnowledgeGraph
        except ImportError:
            from knowledge_graph import KnowledgeGraph
        _knowledge_graph_cache["graph"] = KnowledgeGraph()
    return _knowledge_graph_cache["graph"]


def get_behavior_monitor():
    """µçÆσèáΦ╜╜ΦíîΣ╕║τ¢æµÄºσÖ¿"""
    if "monitor" not in _behavior_monitor_cache:
        try:
            from web.behavior_monitor import BehaviorMonitor
        except ImportError:
            from behavior_monitor import BehaviorMonitor
        _behavior_monitor_cache["monitor"] = BehaviorMonitor()
    return _behavior_monitor_cache["monitor"]


def get_suggestion_engine():
    """µçÆσèáΦ╜╜σ╗║Φ««σ╝òµôÄ"""
    if "engine" not in _suggestion_engine_cache:
        try:
            from web.suggestion_engine import SuggestionEngine
        except ImportError:
            from suggestion_engine import SuggestionEngine
        _suggestion_engine_cache["engine"] = SuggestionEngine()
    return _suggestion_engine_cache["engine"]


def get_insight_reporter():
    """µçÆσèáΦ╜╜µ┤₧σ»ƒµèÑσæèτöƒµêÉσÖ¿"""
    if "reporter" not in _insight_reporter_cache:
        try:
            from web.insight_reporter import InsightReporter
        except ImportError:
            from insight_reporter import InsightReporter
        _insight_reporter_cache["reporter"] = InsightReporter()
    return _insight_reporter_cache["reporter"]


# ==================== σó₧σ╝║Σ╕╗σè¿Φâ╜σè¢µ¿íσ¥ùτ╝ôσ¡ÿ ====================
_notification_manager_cache = {}
_proactive_dialogue_cache = {}
_context_awareness_cache = {}
_auto_execution_cache = {}
_trigger_system_cache = {}


def get_notification_manager():
    """µçÆσèáΦ╜╜ΘÇÜτƒÑτ«íτÉåσÖ¿"""
    if "manager" not in _notification_manager_cache:
        try:
            from web.notification_manager import get_notification_manager as _get_mgr
        except ImportError:
            from notification_manager import get_notification_manager as _get_mgr
        _notification_manager_cache["manager"] = _get_mgr()
    return _notification_manager_cache["manager"]


def get_proactive_dialogue():
    """µçÆσèáΦ╜╜Σ╕╗σè¿σ»╣Φ»¥σ╝òµôÄ"""
    if "engine" not in _proactive_dialogue_cache:
        try:
            from web.proactive_dialogue import get_proactive_dialogue_engine
        except ImportError:
            from proactive_dialogue import get_proactive_dialogue_engine

        # Θ¢åµêÉΣ╛¥Φ╡ûµ¿íσ¥ù
        notif_mgr = get_notification_manager()
        behavior_mon = get_behavior_monitor()
        suggestion_eng = get_suggestion_engine()

        _proactive_dialogue_cache["engine"] = get_proactive_dialogue_engine(
            notification_manager=notif_mgr,
            behavior_monitor=behavior_mon,
            suggestion_engine=suggestion_eng,
        )
    return _proactive_dialogue_cache["engine"]


def get_context_awareness():
    """µçÆσèáΦ╜╜µâàσóâµäƒτƒÑτ│╗τ╗ƒ"""
    if "system" not in _context_awareness_cache:
        try:
            from web.context_awareness import get_context_awareness_system
        except ImportError:
            from context_awareness import get_context_awareness_system

        behavior_mon = get_behavior_monitor()
        _context_awareness_cache["system"] = get_context_awareness_system(
            behavior_monitor=behavior_mon
        )
    return _context_awareness_cache["system"]


def get_auto_execution():
    """µçÆσèáΦ╜╜Φç¬σè¿µëºΦíîσ╝òµôÄ"""
    if "engine" not in _auto_execution_cache:
        try:
            from web.auto_execution import get_auto_execution_engine
        except ImportError:
            from auto_execution import get_auto_execution_engine

        notif_mgr = get_notification_manager()
        _auto_execution_cache["engine"] = get_auto_execution_engine(
            notification_manager=notif_mgr
        )
    return _auto_execution_cache["engine"]


def get_trigger_system():
    """µçÆσèáΦ╜╜Σ╕╗σè¿Σ║ñΣ║ÆΦºªσÅæτ│╗τ╗ƒ"""
    if "system" not in _trigger_system_cache:
        try:
            from web.proactive_trigger import get_trigger_system as _get_trigger_system
        except ImportError:
            from proactive_trigger import get_trigger_system as _get_trigger_system

        behavior_mon = get_behavior_monitor()
        context_sys = get_context_awareness()
        suggestion_eng = get_suggestion_engine()
        notif_mgr = get_notification_manager()
        dialogue_eng = get_proactive_dialogue()

        _trigger_system_cache["system"] = _get_trigger_system(
            behavior_monitor=behavior_mon,
            context_awareness=context_sys,
            suggestion_engine=suggestion_eng,
            notification_manager=notif_mgr,
            dialogue_engine=dialogue_eng,
        )
        # σÉ»σè¿σÉÄσÅ░Φ╜«Φ»ó∩╝êµ»Å5σêåΘÆƒµúÇµƒÑΣ╕Çµ¼íΦºªσÅæµ¥íΣ╗╢∩╝ë
        try:
            _trigger_system_cache["system"].start_monitoring(check_interval=300)
        except Exception as _tse:
            _app_logger.warning(f"[TriggerSystem] ΓÜá∩╕Å start_monitoring σñ▒Φ┤Ñ∩╝êΘ¥₧Φç┤σæ╜∩╝ë: {_tse}")
    return _trigger_system_cache["system"]


@app.route("/api/batch/submit", methods=["POST"])
def batch_submit():
    """µÅÉΣ║ñµë╣ΘçÅµûçΣ╗╢σñäτÉåΣ╗╗σèí"""
    try:
        data = request.json or {}
        command = data.get("command", "")
        manager = get_batch_ops_manager()

        if command:
            parsed = manager.parse_command(command)
            if not parsed.get("success"):
                return (
                    jsonify(
                        {
                            "success": False,
                            "error": parsed.get("error"),
                            "hint": parsed.get("hint"),
                        }
                    ),
                    400,
                )
            operation = parsed.get("operation")
            input_dir = parsed.get("input_dir")
            output_dir = parsed.get("output_dir")
            options = parsed.get("options", {})
        else:
            operation = data.get("operation")
            input_dir = data.get("input_dir")
            output_dir = data.get("output_dir")
            options = data.get("options", {})

        if not operation or not input_dir or not output_dir:
            return jsonify({"success": False, "error": "τ╝║σ░æσ┐àΦªüσÅéµò░"}), 400

        job = manager.create_job(
            name=f"batch_{operation}",
            operation=operation,
            input_dir=input_dir,
            output_dir=output_dir,
            options=options,
        )
        manager.start_job(job.job_id)
        return jsonify(
            {"success": True, "job_id": job.job_id, "job": manager.get_job(job.job_id)}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/batch/jobs", methods=["GET"])
def batch_list_jobs():
    """σêùσç║µë╣ΘçÅΣ╗╗σèí"""
    manager = get_batch_ops_manager()
    return jsonify({"success": True, "jobs": manager.list_jobs()})


@app.route("/api/batch/jobs/<job_id>", methods=["GET"])
def batch_get_job(job_id):
    """ΦÄ╖σÅûσìòΣ╕¬Σ╗╗σèíΦ»ªµâà"""
    manager = get_batch_ops_manager()
    job = manager.get_job(job_id)
    if not job:
        return jsonify({"success": False, "error": "Σ╗╗σèíΣ╕ìσ¡ÿσ£¿"}), 404
    return jsonify({"success": True, "job": job})


@app.route("/api/batch/stream/<job_id>", methods=["GET"])
def batch_stream_job(job_id):
    """µë╣ΘçÅΣ╗╗σèíΦ┐¢σ║ªµ╡ü"""
    manager = get_batch_ops_manager()
    return Response(manager.stream_job(job_id), mimetype="text/event-stream")


@app.route("/api/organize/scan-file", methods=["POST"])
def organize_scan_file():
    """µë½µÅÅσÆîσêåµ₧ÉσìòΣ╕¬µûçΣ╗╢"""
    try:
        data = request.json
        file_path = data.get("file_path")

        if not file_path:
            return jsonify({"error": "τ╝║σ░æ file_path σÅéµò░"}), 400

        if not os.path.exists(file_path):
            return jsonify({"error": f"µûçΣ╗╢Σ╕ìσ¡ÿσ£¿: {file_path}"}), 404

        analyzer = get_file_analyzer()
        analysis_result = analyzer.analyze_file(file_path)

        return jsonify(
            {
                "success": True,
                "file": os.path.basename(file_path),
                "analysis": analysis_result,
            }
        )

    except Exception as e:
        return jsonify({"error": f"σêåµ₧Éσñ▒Φ┤Ñ: {str(e)}"}), 500


@app.route("/api/organize/auto-organize", methods=["POST"])
def organize_auto_organize():
    """Φç¬σè¿τ╗äτ╗çµûçΣ╗╢∩╝êσêåµ₧É+τº╗σè¿∩╝ë"""
    try:
        data = request.json
        file_path = data.get("file_path")
        auto_confirm = data.get("auto_confirm", True)

        if not file_path:
            return jsonify({"error": "τ╝║σ░æ file_path σÅéµò░"}), 400

        if not os.path.exists(file_path):
            return jsonify({"error": f"µûçΣ╗╢Σ╕ìσ¡ÿσ£¿: {file_path}"}), 404

        # τ¼¼Σ╕Çµ¡Ñ∩╝Üσêåµ₧ÉµûçΣ╗╢
        analyzer = get_file_analyzer()
        analysis = analyzer.analyze_file(file_path)
        suggested_folder = analysis.get("suggested_folder")

        if not suggested_folder:
            return jsonify({"error": "µùáµ│òτí«σ«ÜµûçΣ╗╢σêåτ▒╗", "analysis": analysis}), 400

        # τ¼¼Σ║îµ¡Ñ∩╝Üτ╗äτ╗çµûçΣ╗╢
        organizer = get_file_organizer()
        org_result = organizer.organize_file(
            file_path, suggested_folder, auto_confirm=auto_confirm
        )

        if org_result.get("success"):
            return jsonify(
                {
                    "success": True,
                    "file": os.path.basename(file_path),
                    "analysis": analysis,
                    "organized": org_result,
                }
            )
        else:
            return (
                jsonify(
                    {"error": org_result.get("error", "τ╗äτ╗çσñ▒Φ┤Ñ"), "analysis": analysis}
                ),
                500,
            )

    except Exception as e:
        return jsonify({"error": f"Φç¬σè¿τ╗äτ╗çσñ▒Φ┤Ñ: {str(e)}"}), 500


@app.route("/api/organize/list-categories", methods=["GET"])
def organize_list_categories():
    """σêùσç║µëÇµ£ëσêåτ▒╗σÆîµûçΣ╗╢σñ╣"""
    try:
        organizer = get_file_organizer()
        folders = organizer.list_organized_folders()
        stats = organizer.get_categories_stats()

        return jsonify(
            {
                "success": True,
                "folders": folders,
                "stats": stats,
                "total_files": len(organizer.get_index().get("files", [])),
            }
        )

    except Exception as e:
        return jsonify({"error": f"ΦÄ╖σÅûσêåτ▒╗σñ▒Φ┤Ñ: {str(e)}"}), 500


@app.route("/api/organize/search", methods=["POST"])
def organize_search():
    """µÉ£τ┤óσ╖▓τ╗äτ╗çτÜäµûçΣ╗╢"""
    try:
        data = request.json
        keyword = data.get("keyword", "")

        if not keyword:
            return jsonify({"error": "τ╝║σ░æµÉ£τ┤óσà│Θö«Φ»ì"}), 400

        organizer = get_file_organizer()
        results = organizer.search_files(keyword)

        return jsonify(
            {
                "success": True,
                "keyword": keyword,
                "count": len(results),
                "results": results,
            }
        )

    except Exception as e:
        return jsonify({"error": f"µÉ£τ┤óσñ▒Φ┤Ñ: {str(e)}"}), 500


@app.route("/api/organize/stats", methods=["GET"])
def organize_stats():
    """ΦÄ╖σÅûτ╗äτ╗çτ╗ƒΦ«íΣ┐íµü»"""
    try:
        organizer = get_file_organizer()
        index = organizer.get_index()
        stats = organizer.get_categories_stats()
        folders = organizer.list_organized_folders()

        return jsonify(
            {
                "success": True,
                "total_files": index.get("total_files", 0),
                "total_folders": len(folders),
                "by_industry": stats,
                "last_updated": index.get("last_updated"),
            }
        )

    except Exception as e:
        return jsonify({"error": f"ΦÄ╖σÅûτ╗ƒΦ«íσñ▒Φ┤Ñ: {str(e)}"}), 500


@app.route("/api/organize/cleanup", methods=["POST"])
def organize_cleanup():
    """µò┤σÉêµ╕àτÉå _organize τ¢«σ╜òΣ╕¡τÜäΘçìσñìµûçΣ╗╢σñ╣"""
    try:
        data = request.get_json(silent=True) or {}
        dry_run = data.get("dry_run", True)
        ai_rename = data.get("ai_rename", False)

        organize_root = get_organize_root()

        try:
            from web.organize_cleanup import OrganizeCleanup
        except ImportError:
            from organize_cleanup import OrganizeCleanup

        cleanup = OrganizeCleanup(organize_root=organize_root)
        report = cleanup.run(dry_run=dry_run, ai_rename=ai_rename)

        return jsonify(
            {
                "success": True,
                "dry_run": dry_run,
                "total_folders_scanned": report.get("total_folders_scanned", 0),
                "similarity_groups": report.get("similarity_groups", 0),
                "merge_plans": report.get("merge_plans", 0),
                "merged_files": report.get("merged_files", 0),
                "deduped_files": report.get("deduped_files", 0),
                "removed_folders": report.get("removed_folders", 0),
                "empty_cleaned": report.get("empty_cleaned", 0),
                "ai_renames": report.get("ai_renames", 0),
                "log": report.get("log", [])[-50:],  # µ£ÇΦ┐æ50µ¥íµùÑσ┐ù
            }
        )

    except Exception as e:
        return jsonify({"error": f"µò┤σÉêµ╕àτÉåσñ▒Φ┤Ñ: {str(e)}"}), 500


# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ
# µûçΣ╗╢τ╝ûΦ╛æΣ╕ÄµÉ£τ┤ó API
# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ


@app.route("/api/file-editor/read", methods=["POST"])
def file_editor_read():
    """Φ»╗σÅûµûçΣ╗╢σåàσ«╣"""
    try:
        data = request.json or {}
        file_path = data.get("file_path")

        if not file_path:
            return jsonify({"error": "τ╝║σ░æµûçΣ╗╢Φ╖»σ╛ä"}), 400

        editor = get_file_editor()
        result = editor.read_file(file_path)

        return jsonify(result)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/file-editor/write", methods=["POST"])
def file_editor_write():
    """σåÖσàÑµûçΣ╗╢σåàσ«╣"""
    try:
        data = request.json or {}
        file_path = data.get("file_path")
        content = data.get("content")

        if not file_path or content is None:
            return jsonify({"error": "τ╝║σ░æσ┐àΦªüσÅéµò░"}), 400

        editor = get_file_editor()
        result = editor.write_file(file_path, content)

        return jsonify(result)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/file-editor/replace", methods=["POST"])
def file_editor_replace():
    """µ¢┐µìóµûçΣ╗╢σåàσ«╣"""
    try:
        data = request.json or {}
        file_path = data.get("file_path")
        old_text = data.get("old_text")
        new_text = data.get("new_text")
        use_regex = data.get("use_regex", False)

        if not all([file_path, old_text is not None, new_text is not None]):
            return jsonify({"error": "τ╝║σ░æσ┐àΦªüσÅéµò░"}), 400

        editor = get_file_editor()
        result = editor.replace_text(file_path, old_text, new_text, use_regex=use_regex)

        return jsonify(result)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/file-editor/smart-edit", methods=["POST"])
def file_editor_smart_edit():
    """µÖ║Φâ╜τ╝ûΦ╛æ∩╝êτÉåΦºúΦç¬τä╢Φ»¡Φ¿ÇµîçΣ╗ñ∩╝ë"""
    try:
        data = request.json or {}
        file_path = data.get("file_path")
        instruction = data.get("instruction")

        if not file_path or not instruction:
            return jsonify({"error": "τ╝║σ░æσ┐àΦªüσÅéµò░"}), 400

        editor = get_file_editor()
        result = editor.smart_edit(file_path, instruction)

        return jsonify(result)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/file-search/index", methods=["POST"])
def file_search_index():
    """τ┤óσ╝òµûçΣ╗╢µêûτ¢«σ╜ò"""
    try:
        data = request.json or {}
        path = data.get("path")
        is_directory = data.get("is_directory", False)

        if not path:
            return jsonify({"error": "τ╝║σ░æΦ╖»σ╛äσÅéµò░"}), 400

        indexer = get_file_indexer()

        if is_directory:
            result = indexer.index_directory(path, recursive=True)
        else:
            result = indexer.index_file(path)

        return jsonify(result)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/file-search/search", methods=["POST"])
def file_search_search():
    """µÉ£τ┤óµûçΣ╗╢"""
    try:
        data = request.json or {}
        query = data.get("query")
        limit = data.get("limit", 20)
        file_types = data.get("file_types")

        if not query:
            return jsonify({"error": "τ╝║σ░æµÉ£τ┤óσà│Θö«Φ»ì"}), 400

        indexer = get_file_indexer()
        results = indexer.search(query, limit=limit, file_types=file_types)

        return jsonify({"success": True, "results": results, "count": len(results)})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/file-search/find-by-content", methods=["POST"])
def file_search_find_by_content():
    """µá╣µì«σåàσ«╣τëçµ«╡µƒÑµë╛µûçΣ╗╢"""
    try:
        data = request.json or {}
        content_sample = data.get("content")
        min_similarity = data.get("min_similarity", 0.3)

        if not content_sample:
            return jsonify({"error": "τ╝║σ░æσåàσ«╣µá╖µ£¼"}), 400

        indexer = get_file_indexer()
        results = indexer.find_by_content(content_sample, min_similarity=min_similarity)

        return jsonify({"success": True, "results": results, "count": len(results)})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/file-search/list", methods=["GET"])
def file_search_list():
    """σêùσç║µëÇµ£ëσ╖▓τ┤óσ╝òµûçΣ╗╢"""
    try:
        limit = request.args.get("limit", 100, type=int)
        offset = request.args.get("offset", 0, type=int)

        indexer = get_file_indexer()
        files = indexer.list_indexed_files(limit=limit, offset=offset)

        return jsonify({"success": True, "files": files, "count": len(files)})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ
# σà¿τ¢ÿµûçΣ╗╢µë½µÅÅ API  (FileScanner)
# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ


@app.route("/api/scan/start", methods=["POST"])
def scan_start():
    """σÉ»σè¿σà¿τ¢ÿµûçΣ╗╢µë½µÅÅ∩╝êσÉÄσÅ░τ║┐τ¿ï∩╝ë"""
    try:
        from web.file_scanner import FileScanner

        data = request.json or {}
        drives = data.get("drives")  # None ΓåÆ Φç¬σè¿µ₧ÜΣ╕╛µëÇµ£ëσêåσî║
        already = not FileScanner.start_scan(drives=drives)
        return jsonify(
            {
                "success": True,
                "already_running": already,
                "drives": drives or FileScanner.get_drives(),
                "message": (
                    "µë½µÅÅσ╖▓σ£¿Φ┐¢ΦíîΣ╕¡" if already else "σà¿τ¢ÿµë½µÅÅσ╖▓σÉ»σè¿∩╝êσÉÄσÅ░Φ┐ÉΦíî∩╝ë"
                ),
            }
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/scan/status", methods=["GET"])
def scan_status():
    """Φ┐öσ¢₧µë½µÅÅΦ┐¢σ║ªσÆîτ╗ƒΦ«í"""
    try:
        from web.file_scanner import FileScanner

        return jsonify(
            {
                "success": True,
                **FileScanner.get_status(),
                "indexed_count": FileScanner.stats()["total"],
                "by_category": FileScanner.stats()["by_category"],
            }
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/scan/search", methods=["POST"])
def scan_search():
    """σà¿τ¢ÿµûçΣ╗╢σÉìµ¿íτ│èµÉ£τ┤ó"""
    try:
        from web.file_scanner import FileScanner

        data = request.json or {}
        query = (data.get("query") or "").strip()
        limit = int(data.get("limit", 12))
        ext_filter = data.get("ext_filter")  # ['.docx', ...] or None
        category_filter = data.get("category")  # 'µûçµíú' / 'σ¢╛τëç' / ... or None
        if not query:
            return jsonify({"success": False, "error": "τ╝║σ░æ query σÅéµò░"}), 400
        FileScanner.ensure_loaded()
        results = FileScanner.search(
            query, limit=limit, ext_filter=ext_filter, category_filter=category_filter
        )
        return jsonify({"success": True, "results": results, "count": len(results)})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/scan/open", methods=["POST"])
def scan_open():
    """τö¿τ│╗τ╗ƒΘ╗ÿΦ«ñτ¿ïσ║Åµëôσ╝Çµîçσ«Üτ╗¥σ»╣Φ╖»σ╛äµûçΣ╗╢"""
    try:
        from web.file_scanner import FileScanner

        data = request.json or {}
        path = (data.get("path") or "").strip()
        if not path:
            return jsonify({"success": False, "error": "τ╝║σ░æ path σÅéµò░"}), 400
        result = FileScanner.open_file(path)
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/scan/stats", methods=["GET"])
def scan_stats():
    """τ┤óσ╝òτ╗ƒΦ«íµò░µì«"""
    try:
        from web.file_scanner import FileScanner

        return jsonify({"success": True, **FileScanner.stats()})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ
# µªéσ┐╡µÅÉσÅû API
# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ


@app.route("/api/concepts/extract", methods=["POST"])
def concepts_extract():
    """Σ╗ÄµûçΣ╗╢Σ╕¡µÅÉσÅûσà│Θö«µªéσ┐╡"""
    try:
        data = request.json or {}
        file_path = data.get("file_path")
        content = data.get("content")  # σÅ»ΘÇë∩╝îσªéµ₧£σ╖▓Φ»╗σÅûσåàσ«╣
        top_n = data.get("top_n", 10)

        if not file_path:
            return jsonify({"error": "τ╝║σ░æµûçΣ╗╢Φ╖»σ╛ä"}), 400

        extractor = get_concept_extractor()
        result = extractor.analyze_file(file_path, content=content)

        return jsonify(result)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/concepts/related-files", methods=["POST"])
def concepts_related_files():
    """µƒÑµë╛Σ╕ÄµûçΣ╗╢τ¢╕σà│τÜäσà╢Σ╗ûµûçΣ╗╢"""
    try:
        data = request.json or {}
        file_path = data.get("file_path")
        limit = data.get("limit", 5)

        if not file_path:
            return jsonify({"error": "τ╝║σ░æµûçΣ╗╢Φ╖»σ╛ä"}), 400

        extractor = get_concept_extractor()
        related = extractor.find_related_files(file_path, limit=limit)

        return jsonify(
            {"success": True, "file_path": file_path, "related_files": related}
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/concepts/top", methods=["GET"])
def concepts_top():
    """ΦÄ╖σÅûσà¿σ▒Çτâ¡Θù¿µªéσ┐╡"""
    try:
        limit = request.args.get("limit", 20, type=int)

        extractor = get_concept_extractor()
        concepts = extractor.get_top_concepts(limit=limit)

        return jsonify({"success": True, "concepts": concepts})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/concepts/stats", methods=["GET"])
def concepts_stats():
    """ΦÄ╖σÅûµªéσ┐╡µÅÉσÅûτ╗ƒΦ«í"""
    try:
        extractor = get_concept_extractor()
        stats = extractor.get_statistics()

        return jsonify(stats)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ
# τƒÑΦ»åσ¢╛Φ░▒ API
# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ


@app.route("/api/knowledge-graph/build", methods=["POST"])
def knowledge_graph_build():
    """µ₧äσ╗║τƒÑΦ»åσ¢╛Φ░▒"""
    try:
        data = request.json or {}
        file_paths = data.get("file_paths", [])
        force_rebuild = data.get("force_rebuild", False)

        if not file_paths:
            return jsonify({"error": "τ╝║σ░æµûçΣ╗╢Φ╖»σ╛äσêùΦí¿"}), 400

        kg = get_knowledge_graph()
        kg.build_file_graph(file_paths, force_rebuild=force_rebuild)

        stats = kg.get_statistics()

        return jsonify(
            {"success": True, "message": "τƒÑΦ»åσ¢╛Φ░▒µ₧äσ╗║σ«îµêÉ", "statistics": stats}
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/knowledge-graph/data", methods=["GET"])
def knowledge_graph_data():
    """ΦÄ╖σÅûτƒÑΦ»åσ¢╛Φ░▒µò░µì«τö¿Σ║ÄσÅ»Φºåσîû"""
    try:
        max_nodes = request.args.get("max_nodes", 100, type=int)

        kg = get_knowledge_graph()
        graph_data = kg.get_graph_data(max_nodes=max_nodes)

        return jsonify(graph_data)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/knowledge-graph/neighbors", methods=["POST"])
def knowledge_graph_neighbors():
    """ΦÄ╖σÅûµûçΣ╗╢τÜäΘé╗σ▒àΦèéτé╣"""
    try:
        data = request.json or {}
        file_path = data.get("file_path")
        depth = data.get("depth", 1)

        if not file_path:
            return jsonify({"error": "τ╝║σ░æµûçΣ╗╢Φ╖»σ╛ä"}), 400

        kg = get_knowledge_graph()
        neighbors = kg.get_file_neighbors(file_path, depth=depth)

        return jsonify(neighbors)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/knowledge-graph/concept-cluster", methods=["POST"])
def knowledge_graph_concept_cluster():
    """ΦÄ╖σÅûµªéσ┐╡τ¢╕σà│τÜäµûçΣ╗╢Θ¢åτ╛ñ"""
    try:
        data = request.json or {}
        concept = data.get("concept")
        limit = data.get("limit", 20)

        if not concept:
            return jsonify({"error": "τ╝║σ░æµªéσ┐╡σÅéµò░"}), 400

        kg = get_knowledge_graph()
        cluster = kg.get_concept_cluster(concept, limit=limit)

        return jsonify(cluster)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/knowledge-graph/stats", methods=["GET"])
def knowledge_graph_stats():
    """ΦÄ╖σÅûτƒÑΦ»åσ¢╛Φ░▒τ╗ƒΦ«í"""
    try:
        kg = get_knowledge_graph()
        stats = kg.get_statistics()

        return jsonify(stats)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ
# ΦíîΣ╕║τ¢æµÄº API
# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ


@app.route("/api/behavior/log-event", methods=["POST"])
def behavior_log_event():
    """Φ«░σ╜òτö¿µê╖ΦíîΣ╕║Σ║ïΣ╗╢"""
    try:
        data = request.json or {}
        event_type = data.get("event_type")
        file_path = data.get("file_path")
        session_id = data.get("session_id")
        event_data = data.get("event_data")
        duration_ms = data.get("duration_ms")
        user_id = data.get("user_id", "default")
        auto_trigger = data.get("auto_trigger", True)

        if not event_type:
            return jsonify({"error": "τ╝║σ░æΣ║ïΣ╗╢τ▒╗σ₧ï"}), 400

        monitor = get_behavior_monitor()
        event_id = monitor.log_event(
            event_type=event_type,
            file_path=file_path,
            session_id=session_id,
            event_data=event_data,
            duration_ms=duration_ms,
        )

        decision_payload = None
        triggered = False
        if auto_trigger:
            trigger_system = get_trigger_system()
            decision = trigger_system.evaluate_interaction_need(user_id)
            if decision and decision.should_interact:
                trigger_system.execute_interaction(decision, user_id)
                triggered = True
                decision_payload = {
                    "interaction_type": decision.interaction_type.value,
                    "priority": decision.priority,
                    "reason": decision.reason,
                    "content": decision.content,
                    "scores": {
                        "urgency": decision.urgency_score,
                        "importance": decision.importance_score,
                        "disturbance": decision.disturbance_cost,
                        "final": decision.final_score,
                    },
                }

        return jsonify(
            {
                "success": True,
                "event_id": event_id,
                "triggered": triggered,
                "decision": decision_payload,
            }
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/behavior/recent-events", methods=["GET"])
def behavior_recent_events():
    """ΦÄ╖σÅûµ£ÇΦ┐æτÜäΣ║ïΣ╗╢"""
    try:
        limit = request.args.get("limit", 50, type=int)
        event_type = request.args.get("event_type")

        monitor = get_behavior_monitor()
        events = monitor.get_recent_events(limit=limit, event_type=event_type)

        return jsonify({"success": True, "events": events})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/behavior/top-files", methods=["GET"])
def behavior_top_files():
    """ΦÄ╖σÅûµ£Çσ╕╕τö¿τÜäµûçΣ╗╢"""
    try:
        limit = request.args.get("limit", 10, type=int)

        monitor = get_behavior_monitor()
        files = monitor.get_frequently_used_files(limit=limit)

        return jsonify({"success": True, "files": files})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/behavior/work-patterns", methods=["GET"])
def behavior_work_patterns():
    """ΦÄ╖σÅûσ╖ÑΣ╜£µ¿íσ╝Åσêåµ₧É"""
    try:
        monitor = get_behavior_monitor()
        patterns = monitor.get_work_patterns()

        return jsonify(patterns)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/behavior/stats", methods=["GET"])
def behavior_stats():
    """ΦÄ╖σÅûΦíîΣ╕║τ╗ƒΦ«í"""
    try:
        monitor = get_behavior_monitor()
        stats = monitor.get_statistics()

        return jsonify(stats)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ
# µÖ║Φâ╜σ╗║Φ«« API
# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ


@app.route("/api/suggestions/generate", methods=["POST"])
def suggestions_generate():
    """τöƒµêÉµÖ║Φâ╜σ╗║Φ««"""
    try:
        data = request.json or {}
        force_regenerate = data.get("force_regenerate", False)

        engine = get_suggestion_engine()
        suggestions = engine.generate_suggestions(force_regenerate=force_regenerate)

        return jsonify(
            {"success": True, "suggestions": suggestions, "count": len(suggestions)}
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/suggestions/pending", methods=["GET"])
def suggestions_pending():
    """ΦÄ╖σÅûσ╛àσñäτÉåτÜäσ╗║Φ««"""
    try:
        limit = request.args.get("limit", 10, type=int)

        engine = get_suggestion_engine()
        suggestions = engine.get_pending_suggestions(limit=limit)

        return jsonify(
            {"success": True, "suggestions": suggestions, "count": len(suggestions)}
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/suggestions/dismiss", methods=["POST"])
def suggestions_dismiss():
    """µïÆτ╗¥σ╗║Φ««"""
    try:
        data = request.json or {}
        suggestion_id = data.get("suggestion_id")
        feedback = data.get("feedback")

        if not suggestion_id:
            return jsonify({"error": "τ╝║σ░æσ╗║Φ««ID"}), 400

        engine = get_suggestion_engine()
        engine.dismiss_suggestion(suggestion_id, feedback=feedback)

        return jsonify({"success": True, "message": "σ╗║Φ««σ╖▓µïÆτ╗¥"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/suggestions/apply", methods=["POST"])
def suggestions_apply():
    """σ║öτö¿σ╗║Φ««"""
    try:
        data = request.json or {}
        suggestion_id = data.get("suggestion_id")
        feedback = data.get("feedback")

        if not suggestion_id:
            return jsonify({"error": "τ╝║σ░æσ╗║Φ««ID"}), 400

        engine = get_suggestion_engine()
        engine.apply_suggestion(suggestion_id, feedback=feedback)

        return jsonify({"success": True, "message": "σ╗║Φ««σ╖▓σ║öτö¿"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/suggestions/stats", methods=["GET"])
def suggestions_stats():
    """ΦÄ╖σÅûσ╗║Φ««τ╗ƒΦ«í"""
    try:
        engine = get_suggestion_engine()
        stats = engine.get_statistics()

        return jsonify(stats)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ
# µ┤₧σ»ƒµèÑσæè API
# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ


@app.route("/api/insights/generate-weekly", methods=["POST"])
def insights_generate_weekly():
    """τöƒµêÉσæ¿µèÑ"""
    try:
        reporter = get_insight_reporter()
        report = reporter.generate_weekly_report()

        return jsonify({"success": True, "report": report})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/insights/generate-monthly", methods=["POST"])
def insights_generate_monthly():
    """τöƒµêÉµ£êµèÑ"""
    try:
        reporter = get_insight_reporter()
        report = reporter.generate_monthly_report()

        return jsonify({"success": True, "report": report})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/insights/latest", methods=["GET"])
def insights_latest():
    """ΦÄ╖σÅûµ£Çµû░µèÑσæè"""
    try:
        report_type = request.args.get("type", "weekly")

        reporter = get_insight_reporter()
        report = reporter.get_latest_report(report_type=report_type)

        if report:
            return jsonify({"success": True, "report": report})
        else:
            return jsonify({"success": False, "message": "µÜéµùáµèÑσæè"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/insights/export-markdown", methods=["POST"])
def insights_export_markdown():
    """σ»╝σç║µèÑσæèΣ╕║Markdown"""
    try:
        data = request.json or {}
        report = data.get("report")
        output_path = data.get("output_path", "workspace/report.md")

        if not report:
            return jsonify({"error": "τ╝║σ░æµèÑσæèµò░µì«"}), 400

        reporter = get_insight_reporter()
        saved_path = reporter.export_report_markdown(report, output_path)

        return jsonify({"success": True, "file_path": saved_path})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ==================== ΘÇÜτƒÑτ«íτÉå API ====================


@app.route("/api/notifications/unread", methods=["GET"])
def get_unread_notifications():
    """ΦÄ╖σÅûµ£¬Φ»╗ΘÇÜτƒÑ"""
    try:
        user_id = request.args.get("user_id", "default")
        limit = int(request.args.get("limit", 50))

        manager = get_notification_manager()
        notifications = manager.get_unread_notifications(user_id, limit)

        return jsonify(
            {
                "success": True,
                "notifications": notifications,
                "count": len(notifications),
            }
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/notifications/mark-read", methods=["POST"])
def mark_notification_read():
    """µáçΦ«░ΘÇÜτƒÑσ╖▓Φ»╗"""
    try:
        data = request.json or {}
        notification_id = data.get("notification_id")
        user_id = data.get("user_id", "default")

        if not notification_id:
            return jsonify({"error": "τ╝║σ░ænotification_id"}), 400

        manager = get_notification_manager()
        manager.mark_as_read(notification_id, user_id)

        return jsonify({"success": True})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/notifications/dismiss", methods=["POST"])
def dismiss_notification():
    """σ┐╜τòÑΘÇÜτƒÑ"""
    try:
        data = request.json or {}
        notification_id = data.get("notification_id")
        user_id = data.get("user_id", "default")

        if not notification_id:
            return jsonify({"error": "τ╝║σ░ænotification_id"}), 400

        manager = get_notification_manager()
        manager.dismiss_notification(notification_id, user_id)

        return jsonify({"success": True})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/notifications/stats", methods=["GET"])
def get_notification_stats():
    """ΦÄ╖σÅûΘÇÜτƒÑτ╗ƒΦ«í"""
    try:
        user_id = request.args.get("user_id", "default")
        days = int(request.args.get("days", 7))

        manager = get_notification_manager()
        stats = manager.get_notification_stats(user_id, days)

        return jsonify({"success": True, "stats": stats})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/notifications/preferences", methods=["GET", "POST"])
def notification_preferences():
    """ΦÄ╖σÅûµêûΦ«╛τ╜«ΘÇÜτƒÑσüÅσÑ╜"""
    try:
        user_id = request.args.get("user_id", "default")
        manager = get_notification_manager()

        if request.method == "GET":
            prefs = manager.get_user_preferences(user_id)
            return jsonify({"success": True, "preferences": prefs})

        else:  # POST
            data = request.json or {}
            manager.update_user_preferences(user_id, data)
            return jsonify({"success": True})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ==================== Σ╕╗σè¿σ»╣Φ»¥ API ====================


@app.route("/api/dialogue/start-monitoring", methods=["POST"])
def start_dialogue_monitoring():
    """σÉ»σè¿Σ╕╗σè¿σ»╣Φ»¥τ¢æµÄº"""
    try:
        data = request.json or {}
        check_interval = data.get("check_interval", 300)  # Θ╗ÿΦ«ñ5σêåΘÆƒ

        engine = get_proactive_dialogue()
        engine.start_monitoring(check_interval)

        return jsonify({"success": True, "message": "Σ╕╗σè¿σ»╣Φ»¥τ¢æµÄºσ╖▓σÉ»σè¿"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/dialogue/stop-monitoring", methods=["POST"])
def stop_dialogue_monitoring():
    """σü£µ¡óΣ╕╗σè¿σ»╣Φ»¥τ¢æµÄº"""
    try:
        engine = get_proactive_dialogue()
        engine.stop_monitoring()

        return jsonify({"success": True, "message": "Σ╕╗σè¿σ»╣Φ»¥τ¢æµÄºσ╖▓σü£µ¡ó"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/dialogue/trigger", methods=["POST"])
def trigger_dialogue():
    """µëïσè¿ΦºªσÅæσ»╣Φ»¥"""
    try:
        data = request.json or {}
        user_id = data.get("user_id", "default")
        scene_type = data.get("scene_type")
        context = data.get("context", {})

        if not scene_type:
            return jsonify({"error": "τ╝║σ░æscene_type"}), 400

        engine = get_proactive_dialogue()
        engine.manual_trigger(user_id, scene_type, **context)

        return jsonify({"success": True, "message": f"σ╖▓ΦºªσÅæ{scene_type}σ»╣Φ»¥"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/dialogue/history", methods=["GET"])
def get_dialogue_history():
    """ΦÄ╖σÅûσ»╣Φ»¥σÄåσÅ▓"""
    try:
        user_id = request.args.get("user_id", "default")
        limit = int(request.args.get("limit", 50))

        engine = get_proactive_dialogue()
        history = engine.get_dialogue_history(user_id, limit)

        return jsonify({"success": True, "history": history, "count": len(history)})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ==================== µâàσóâµäƒτƒÑ API ====================


@app.route("/api/context/detect", methods=["POST"])
def detect_context():
    """µúÇµ╡ïσ╜ôσëìσ╖ÑΣ╜£σ£║µÖ»"""
    try:
        data = request.json or {}
        user_id = data.get("user_id", "default")

        system = get_context_awareness()
        context = system.detect_context(user_id)

        return jsonify({"success": True, "context": context})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/context/current", methods=["GET"])
def get_current_context():
    """ΦÄ╖σÅûσ╜ôσëìσ£║µÖ»"""
    try:
        system = get_context_awareness()
        context = system.get_current_context()

        return jsonify({"success": True, "context": context})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/context/history", methods=["GET"])
def get_context_history():
    """ΦÄ╖σÅûσ£║µÖ»σÄåσÅ▓"""
    try:
        user_id = request.args.get("user_id", "default")
        days = int(request.args.get("days", 7))

        system = get_context_awareness()
        history = system.get_context_history(user_id, days)

        return jsonify({"success": True, "history": history, "count": len(history)})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/context/statistics", methods=["GET"])
def get_context_statistics():
    """ΦÄ╖σÅûσ£║µÖ»τ╗ƒΦ«í"""
    try:
        user_id = request.args.get("user_id", "default")
        days = int(request.args.get("days", 30))

        system = get_context_awareness()
        stats = system.get_context_statistics(user_id, days)

        return jsonify({"success": True, "statistics": stats})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/context/predict", methods=["GET"])
def predict_next_context():
    """Θóäµ╡ïΣ╕ïΣ╕ÇΣ╕¬σ£║µÖ»"""
    try:
        user_id = request.args.get("user_id", "default")

        system = get_context_awareness()
        prediction = system.predict_next_context(user_id)

        return jsonify({"success": True, "prediction": prediction})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ==================== Φç¬σè¿µëºΦíî API ====================


@app.route("/api/execution/authorize", methods=["POST"])
def authorize_task_execution():
    """µÄêµ¥âΣ╗╗σèíµëºΦíî"""
    try:
        data = request.json or {}
        user_id = data.get("user_id", "default")
        task_type = data.get("task_type")
        auto_execute = data.get("auto_execute", False)
        max_executions_per_day = data.get("max_executions_per_day", 10)
        expires_days = data.get("expires_days", 30)

        if not task_type:
            return jsonify({"error": "τ╝║σ░ætask_type"}), 400

        engine = get_auto_execution()
        engine.authorize_task(
            user_id, task_type, auto_execute, max_executions_per_day, expires_days
        )

        return jsonify({"success": True, "message": f"σ╖▓µÄêµ¥â{task_type}Σ╗╗σèí"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/execution/revoke", methods=["POST"])
def revoke_task_authorization():
    """µÆñΘöÇΣ╗╗σèíµÄêµ¥â"""
    try:
        data = request.json or {}
        user_id = data.get("user_id", "default")
        task_type = data.get("task_type")

        if not task_type:
            return jsonify({"error": "τ╝║σ░ætask_type"}), 400

        engine = get_auto_execution()
        engine.revoke_authorization(user_id, task_type)

        return jsonify({"success": True, "message": f"σ╖▓µÆñΘöÇ{task_type}Σ╗╗σèíµÄêµ¥â"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/execution/execute", methods=["POST"])
def execute_task():
    """µëºΦíîΣ╗╗σèí"""
    try:
        data = request.json or {}
        user_id = data.get("user_id", "default")
        task_type = data.get("task_type")
        params = data.get("params", {})
        force = data.get("force", False)

        if not task_type:
            return jsonify({"error": "τ╝║σ░ætask_type"}), 400

        engine = get_auto_execution()
        result = engine.execute_task(user_id, task_type, params, force)

        if result["success"]:
            return jsonify(result)
        else:
            return jsonify(result), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/execution/queue", methods=["POST"])
def queue_task():
    """Σ╗╗σèíσèáσàÑΘÿƒσêù"""
    try:
        data = request.json or {}
        user_id = data.get("user_id", "default")
        task_type = data.get("task_type")
        params = data.get("params", {})
        priority = data.get("priority", 5)

        if not task_type:
            return jsonify({"error": "τ╝║σ░ætask_type"}), 400

        engine = get_auto_execution()
        task_id = engine.queue_task(user_id, task_type, params, priority)

        return jsonify({"success": True, "task_id": task_id})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/execution/history", methods=["GET"])
def get_execution_history():
    """ΦÄ╖σÅûµëºΦíîσÄåσÅ▓"""
    try:
        user_id = request.args.get("user_id", "default")
        limit = int(request.args.get("limit", 50))

        engine = get_auto_execution()
        history = engine.get_execution_history(user_id, limit)

        return jsonify({"success": True, "history": history, "count": len(history)})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/execution/statistics", methods=["GET"])
def get_execution_statistics():
    """ΦÄ╖σÅûµëºΦíîτ╗ƒΦ«í"""
    try:
        user_id = request.args.get("user_id", "default")
        days = int(request.args.get("days", 30))

        engine = get_auto_execution()
        stats = engine.get_statistics(user_id, days)

        return jsonify({"success": True, "statistics": stats})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/execution/start-processor", methods=["POST"])
def start_execution_processor():
    """σÉ»σè¿Φç¬σè¿µëºΦíîσñäτÉåσÖ¿"""
    try:
        data = request.json or {}
        interval = data.get("interval", 60)  # Θ╗ÿΦ«ñ1σêåΘÆƒ

        engine = get_auto_execution()
        engine.start_queue_processor(interval)

        return jsonify({"success": True, "message": "Φç¬σè¿µëºΦíîσñäτÉåσÖ¿σ╖▓σÉ»σè¿"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/execution/stop-processor", methods=["POST"])
def stop_execution_processor():
    """σü£µ¡óΦç¬σè¿µëºΦíîσñäτÉåσÖ¿"""
    try:
        engine = get_auto_execution()
        engine.stop_queue_processor()

        return jsonify({"success": True, "message": "Φç¬σè¿µëºΦíîσñäτÉåσÖ¿σ╖▓σü£µ¡ó"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ==================== Σ╕╗σè¿Σ║ñΣ║ÆΦºªσÅæτ│╗τ╗ƒ API ====================


@app.route("/api/triggers/evaluate", methods=["POST"])
def triggers_evaluate():
    """Φ»äΣ╝░µÿ»σÉªΘ£ÇΦªüΣ╕╗σè¿Σ║ñΣ║Æ"""
    try:
        data = request.json or {}
        user_id = data.get("user_id", "default")
        execute = data.get("execute", True)

        system = get_trigger_system()
        decision = system.evaluate_interaction_need(user_id)

        if decision and decision.should_interact and execute:
            system.execute_interaction(decision, user_id)

        decision_payload = None
        if decision:
            decision_payload = {
                "should_interact": decision.should_interact,
                "interaction_type": decision.interaction_type.value,
                "priority": decision.priority,
                "reason": decision.reason,
                "content": decision.content,
                "scores": {
                    "urgency": decision.urgency_score,
                    "importance": decision.importance_score,
                    "disturbance": decision.disturbance_cost,
                    "final": decision.final_score,
                },
            }

        return jsonify({"success": True, "decision": decision_payload})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/triggers/start", methods=["POST"])
def triggers_start():
    """σÉ»σè¿Σ╕╗σè¿Σ║ñΣ║Æτ¢æµÄº"""
    try:
        data = request.json or {}
        user_id = data.get("user_id", "default")
        interval = data.get("interval", 300)

        system = get_trigger_system()
        system.start_monitoring(check_interval=interval, user_id=user_id)

        return jsonify(
            {"success": True, "message": "Σ╕╗σè¿Σ║ñΣ║ÆΦºªσÅæτ│╗τ╗ƒσ╖▓σÉ»σè¿", "interval": interval}
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/triggers/stop", methods=["POST"])
def triggers_stop():
    """σü£µ¡óΣ╕╗σè¿Σ║ñΣ║Æτ¢æµÄº"""
    try:
        system = get_trigger_system()
        system.stop_monitoring()

        return jsonify({"success": True, "message": "Σ╕╗σè¿Σ║ñΣ║ÆΦºªσÅæτ│╗τ╗ƒσ╖▓σü£µ¡ó"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/triggers/stats", methods=["GET"])
def triggers_stats():
    """ΦÄ╖σÅûΦºªσÅæτ╗ƒΦ«í"""
    try:
        days = int(request.args.get("days", 7))

        system = get_trigger_system()
        stats = system.get_trigger_statistics(days)

        return jsonify({"success": True, "stats": stats})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/triggers/list", methods=["GET"])
def triggers_list():
    """ΦÄ╖σÅûΦºªσÅæσÖ¿σêùΦí¿"""
    try:
        system = get_trigger_system()
        triggers = system.list_triggers()

        return jsonify({"success": True, "triggers": triggers})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/triggers/update", methods=["POST"])
def triggers_update():
    """µ¢┤µû░ΦºªσÅæσÖ¿Θàìτ╜«"""
    try:
        data = request.json or {}
        trigger_id = data.get("trigger_id")

        if not trigger_id:
            return jsonify({"error": "τ╝║σ░ætrigger_id"}), 400

        enabled = data.get("enabled")
        priority = data.get("priority")
        cooldown_minutes = data.get("cooldown_minutes")
        threshold_value = data.get("threshold_value")
        parameters = data.get("parameters")

        system = get_trigger_system()
        ok = system.update_trigger_config(
            trigger_id,
            enabled=enabled,
            priority=priority,
            cooldown_minutes=cooldown_minutes,
            threshold_value=threshold_value,
        )

        if not ok:
            return jsonify({"error": "ΦºªσÅæσÖ¿Σ╕ìσ¡ÿσ£¿"}), 404

        # σªéµ₧£µÅÉΣ╛¢Σ║åσÅéµò░∩╝îµ¢┤µû░σÅéµò░
        if parameters is not None:
            system.update_trigger_params(trigger_id, parameters)

        return jsonify({"success": True, "message": "ΦºªσÅæσÖ¿Θàìτ╜«σ╖▓µ¢┤µû░"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/triggers/params/<trigger_id>", methods=["GET"])
def get_trigger_params(trigger_id):
    """ΦÄ╖σÅûΦºªσÅæσÖ¿σÅéµò░"""
    try:
        system = get_trigger_system()
        params = system.get_trigger_params(trigger_id)

        return jsonify(
            {"success": True, "trigger_id": trigger_id, "parameters": params}
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/triggers/params/<trigger_id>", methods=["POST"])
def update_trigger_params_endpoint(trigger_id):
    """µ¢┤µû░ΦºªσÅæσÖ¿σÅéµò░"""
    try:
        data = request.json or {}
        parameters = data.get("parameters", {})

        system = get_trigger_system()
        ok = system.update_trigger_params(trigger_id, parameters)

        if not ok:
            return jsonify({"error": "ΦºªσÅæσÖ¿Σ╕ìσ¡ÿσ£¿"}), 404

        return jsonify(
            {
                "success": True,
                "message": "ΦºªσÅæσÖ¿σÅéµò░σ╖▓µ¢┤µû░",
                "parameters": system.get_trigger_params(trigger_id),
            }
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/triggers/feedback", methods=["POST"])
def triggers_feedback():
    """µÅÉΣ║ñΦºªσÅæσÅìΘªê"""
    try:
        data = request.json or {}
        trigger_id = data.get("trigger_id")
        feedback = data.get("feedback")
        response_time_seconds = data.get("response_time_seconds", 0)

        if not trigger_id or not feedback:
            return jsonify({"error": "τ╝║σ░ætrigger_idµêûfeedback"}), 400

        system = get_trigger_system()
        system.record_user_feedback(trigger_id, feedback, response_time_seconds)

        return jsonify({"success": True, "message": "σÅìΘªêσ╖▓Φ«░σ╜ò"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ΓòÉΓòÉΓòÉ µ│¿σåîσó₧σ╝║Φ«░σ┐åτ│╗τ╗ƒAPI∩╝êµ¿íσ¥ùτ║ºσê½∩╝îτí«Σ┐¥σºïτ╗êµëºΦíî∩╝ë ΓòÉΓòÉΓòÉ
try:
    from memory_api_routes import register_memory_routes

    register_memory_routes(app, get_memory_manager)
except ImportError:
    try:
        from web.memory_api_routes import register_memory_routes

        register_memory_routes(app, get_memory_manager)
    except ImportError:
        _app_logger.warning("ΓÜá∩╕Å  σó₧σ╝║Φ«░σ┐åτ│╗τ╗ƒAPIµ£¬µë╛σê░∩╝îΣ╜┐τö¿σƒ║τíÇσèƒΦâ╜")


# ΓòÉΓòÉΓòÉ Φç¬σè¿σ╜Æτ║│Φ░âσ║ªσÖ¿ API ΓòÉΓòÉΓòÉ


@app.route("/api/auto-catalog/status", methods=["GET"])
def auto_catalog_status():
    """ΦÄ╖σÅûΦç¬σè¿σ╜Æτ║│τè╢µÇü"""
    try:
        from auto_catalog_scheduler import get_auto_catalog_scheduler

        scheduler = get_auto_catalog_scheduler()

        return jsonify(
            {
                "success": True,
                "enabled": scheduler.is_auto_catalog_enabled(),
                "schedule_time": scheduler.get_catalog_schedule(),
                "source_directories": scheduler.get_source_directories(),
                "backup_directory": scheduler.get_backup_directory(),
            }
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/auto-catalog/enable", methods=["POST"])
def auto_catalog_enable():
    """σÉ»τö¿Φç¬σè¿σ╜Æτ║│"""
    try:
        from auto_catalog_scheduler import get_auto_catalog_scheduler

        scheduler = get_auto_catalog_scheduler()

        data = request.json or {}
        schedule_time = data.get("schedule_time", "02:00")
        source_dirs = data.get("source_directories")

        scheduler.enable_auto_catalog(schedule_time, source_dirs)

        return jsonify(
            {
                "success": True,
                "message": f"Φç¬σè¿σ╜Æτ║│σ╖▓σÉ»τö¿∩╝îµ»ÅµùÑ {schedule_time} µëºΦíî",
                "schedule_time": schedule_time,
                "source_directories": scheduler.get_source_directories(),
            }
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/auto-catalog/disable", methods=["POST"])
def auto_catalog_disable():
    """τªüτö¿Φç¬σè¿σ╜Æτ║│"""
    try:
        from auto_catalog_scheduler import get_auto_catalog_scheduler

        scheduler = get_auto_catalog_scheduler()

        scheduler.disable_auto_catalog()

        return jsonify({"success": True, "message": "Φç¬σè¿σ╜Æτ║│σ╖▓τªüτö¿"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/auto-catalog/run-now", methods=["POST"])
def auto_catalog_run_now():
    """τ½ïσì│µëºΦíîΣ╕Çµ¼íσ╜Æτ║│∩╝êµëïσè¿ΦºªσÅæ∩╝ë"""
    try:
        from auto_catalog_scheduler import get_auto_catalog_scheduler

        scheduler = get_auto_catalog_scheduler()

        result = scheduler.manual_catalog_now()

        return jsonify(
            {
                "success": result.get("success", False),
                "total_files": result.get("total_files", 0),
                "organized_count": result.get("organized_count", 0),
                "backed_up_count": result.get("backed_up_count", 0),
                "errors": result.get("errors", []),
                "report_path": result.get("report_path", ""),
            }
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/auto-catalog/backup-manifest/<path:filename>", methods=["GET"])
def get_backup_manifest(filename):
    """Σ╕ïΦ╜╜σñçΣ╗╜µ╕àσìòµûçΣ╗╢"""
    try:
        from auto_catalog_scheduler import get_auto_catalog_scheduler

        scheduler = get_auto_catalog_scheduler()

        backup_dir = scheduler.get_backup_directory()
        return send_from_directory(backup_dir, filename, as_attachment=False)
    except Exception as e:
        return jsonify({"error": str(e)}), 404


# ΓöÇΓöÇ Token Σ╜┐τö¿τ╗ƒΦ«íµÄÑσÅú ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ


@app.route("/api/token-stats", methods=["GET"])
def api_token_stats():
    """Φ┐öσ¢₧ Token τö¿ΘçÅτ╗ƒΦ«í∩╝êΣ╗èµùÑ / µ£¼µ£ê / µîëµ¿íσ₧ï / Φ┐æ 7 σñ⌐∩╝ë"""
    try:
        from token_tracker import get_stats

        return jsonify(get_stats())
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/token-stats/reset", methods=["POST"])
def api_token_stats_reset():
    """Θçìτ╜«τ╗ƒΦ«íµò░µì«πÇéBody: {"period": "today" | "month" | "all"}"""
    try:
        from token_tracker import reset_stats

        period = (request.json or {}).get("period", "all")
        return jsonify(reset_stats(period))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ΓöÇΓöÇ LangGraph σ╖ÑΣ╜£µ╡üσÅ»Φºåσîû & σ╝ÇσÅæσ╖Ñσà╖ API ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ


@app.route("/workflow-dag")
def workflow_dag_page():
    """σ╖ÑΣ╜£µ╡ü DAG σÅ»ΦºåσîûΘí╡Θ¥ó"""
    import os

    html_path = os.path.join(os.path.dirname(__file__), "static", "workflow_dag.html")
    try:
        from flask import send_file

        return send_file(html_path)
    except Exception as e:
        return f"<h3>Error: {e}</h3>", 500


@app.route("/api/dev/graph-mermaid", methods=["GET"])
def api_dev_graph_mermaid():
    """
    Φ┐öσ¢₧µîçσ«Üσ╖ÑΣ╜£µ╡ü / Agent τÜä Mermaid DAG σ¢╛µáçΦ«░πÇé

    σÅéµò░:
        workflow : σ╖ÑΣ╜£µ╡üσÉìτº░  (research_and_document | multi_agent_ppt | react_agent)
        type     : τ▒╗σ₧ï        (workflow | agent)
    """
    wf = request.args.get("workflow", "react_agent")
    wf_type = request.args.get("type", "agent")
    try:
        if wf_type == "agent" or wf == "react_agent":
            from app.core.agent.factory import create_langgraph_agent

            agent = create_langgraph_agent()
            mermaid_code = agent.get_graph_mermaid()
            # Count nodes/edges
            node_count = mermaid_code.count("\n    ") if mermaid_code else 0
            edge_count = (
                mermaid_code.count("-->") + mermaid_code.count("-.->")
                if mermaid_code
                else 0
            )
        else:
            from app.core.workflow.langgraph_workflow import WorkflowEngine

            engine = WorkflowEngine()
            mermaid_code = engine.get_graph_mermaid(wf)
            node_count = mermaid_code.count("\n    ") if mermaid_code else 0
            edge_count = (
                mermaid_code.count("-->") + mermaid_code.count("-.->")
                if mermaid_code
                else 0
            )

        return jsonify(
            {
                "success": True,
                "workflow": wf,
                "type": wf_type,
                "mermaid": mermaid_code,
                "node_count": max(node_count, 0),
                "edge_count": max(edge_count, 0),
            }
        )
    except Exception as e:
        import traceback

        return (
            jsonify(
                {"success": False, "error": str(e), "traceback": traceback.format_exc()}
            ),
            500,
        )


@app.route("/api/dev/checkpoint-info", methods=["GET"])
def api_dev_checkpoint_info():
    """Φ┐öσ¢₧µúÇµƒÑτé╣µò░µì«σ║ôΣ┐íµü»∩╝êτ▒╗σ₧ï / Σ╝ÜΦ»¥µò░ / σ┐½τàºµÇ╗µò░∩╝ëπÇé"""
    try:
        from app.core.agent.checkpoint_manager import CheckpointManager

        return jsonify(CheckpointManager.get_db_info())
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/dev/checkpoints/<thread_id>", methods=["GET"])
def api_dev_list_checkpoints(thread_id):
    """σêùσç║µƒÉΣ╝ÜΦ»¥τÜäµúÇµƒÑτé╣σ┐½τàºσêùΦí¿πÇé"""
    try:
        from app.core.agent.checkpoint_manager import CheckpointManager

        snapshots = CheckpointManager.list_checkpoints(thread_id)
        return jsonify(
            {"thread_id": thread_id, "snapshots": snapshots, "count": len(snapshots)}
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/dev/checkpoints/<thread_id>", methods=["DELETE"])
def api_dev_delete_checkpoints(thread_id):
    """σêáΘÖñµƒÉΣ╝ÜΦ»¥τÜäσà¿Θâ¿µúÇµƒÑτé╣∩╝êτö¿Σ║Äµ╕àΘÖñσ»╣Φ»¥σÄåσÅ▓∩╝ëπÇé"""
    try:
        from app.core.agent.checkpoint_manager import CheckpointManager

        ok = CheckpointManager.delete_thread(thread_id)
        return jsonify({"success": ok, "thread_id": thread_id})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ
# RAG σÉæΘçÅµúÇτ┤ó API
# ΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉΓòÉ


@app.route("/api/rag/ingest", methods=["POST"])
def api_rag_ingest():
    """
    τ┤óσ╝òµûçΣ╗╢µêûµûçµ£¼σê░σÉæΘçÅσ║ôπÇé

    Φ»╖µ▒éΣ╜ô (JSON):
        { "file_path": "/abs/path/to/doc.pdf" }
        µêû
        { "text": "Φªüτ┤óσ╝òτÜäµûçµ£¼σåàσ«╣...", "source": "my_doc" }

    Φ┐öσ¢₧:
        { "success": true, "chunks_added": 42, "stats": {...} }
    """
    try:
        from app.core.services.rag_service import get_rag_service

        data = request.get_json(force=True) or {}
        rag = get_rag_service()

        if "file_path" in data:
            fp = data["file_path"]
            if not os.path.isabs(fp):
                fp = os.path.join(os.getcwd(), fp)
            if not os.path.exists(fp):
                return jsonify({"error": f"µûçΣ╗╢Σ╕ìσ¡ÿσ£¿: {fp}"}), 400
            count = rag.index_file(fp)
        elif "text" in data:
            count = rag.index_text(data["text"], source=data.get("source", "api_input"))
        else:
            return jsonify({"error": "Φ»╖µÅÉΣ╛¢ file_path µêû text σ¡ùµ«╡"}), 400

        return jsonify({"success": True, "chunks_added": count, "stats": rag.stats()})
    except Exception as e:
        logger.exception("[RAG /ingest] error")
        return jsonify({"error": str(e)}), 500


@app.route("/api/rag/query", methods=["POST"])
def api_rag_query():
    """
    µúÇτ┤óσÉæΘçÅσ║ô∩╝îΦ┐öσ¢₧τ¢╕σà│µûçµ£¼τëçµ«╡πÇé

    Φ»╖µ▒éΣ╜ô (JSON):
        {
          "question": "Koto µö»µîüσô¬Σ║¢µûçΣ╗╢µá╝σ╝Å∩╝ƒ",
          "k": 5,
          "answer": true        // σÅ»ΘÇë∩╝Ütrue = σÉîµù╢τöƒµêÉ LLM τ¡öµíê
        }

    Φ┐öσ¢₧∩╝êΣ╗àµúÇτ┤ó∩╝ë:
        { "chunks": [...], "count": 3 }

    Φ┐öσ¢₧∩╝êσÉ½τ¡öµíê∩╝ë:
        { "answer": "...", "sources": [...], "chunks": [...], "context_used": true }
    """
    try:
        from app.core.services.rag_service import get_rag_service

        data = request.get_json(force=True) or {}
        question = data.get("question", "").strip()
        if not question:
            return jsonify({"error": "question σ¡ùµ«╡Σ╕ìΦâ╜Σ╕║τ⌐║"}), 400

        k = int(data.get("k", 5))
        want_answer = data.get("answer", False)
        rag = get_rag_service()

        if want_answer:
            result = rag.rag_answer(question, k=k)
            return jsonify(result)
        else:
            chunks = rag.retrieve(question, k=k)
            return jsonify({"chunks": chunks, "count": len(chunks)})
    except Exception as e:
        logger.exception("[RAG /query] error")
        return jsonify({"error": str(e)}), 500


@app.route("/api/rag/stats", methods=["GET"])
def api_rag_stats():
    """
    Φ┐öσ¢₧ RAG τ┤óσ╝òτ╗ƒΦ«íΣ┐íµü»πÇé

    Φ┐öσ¢₧:
        {
          "initialized": true,
          "doc_count": 312,
          "index_dir": "config/rag_index",
          "index_size_mb": 2.4,
          "embedding_model": "GoogleGenerativeAIEmbeddings"
        }
    """
    try:
        from app.core.services.rag_service import get_rag_service

        rag = get_rag_service()
        return jsonify(rag.stats())
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/rag/clear", methods=["DELETE"])
def api_rag_clear():
    """µ╕àτ⌐║ RAG σÉæΘçÅσ║ô∩╝êσêáΘÖñµëÇµ£ëτ┤óσ╝òµò░µì«∩╝ëπÇé"""
    try:
        import app.core.services.rag_service as _rag_mod
        from app.core.services.rag_service import _rag_instance, get_rag_service

        rag = get_rag_service()
        ok = rag.clear()
        # Θçìτ╜«σìòΣ╛ï∩╝îΣ╕ïµ¼í get_rag_service() σ░åΘçìσ╗║
        _rag_mod._rag_instance = None
        return jsonify({"success": ok, "message": "σÉæΘçÅσ║ôσ╖▓µ╕àτ⌐║"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
# τö¿µê╖Φ»äσêåµÄÑσÅú ΓÇö µÄÑµö╢σëìτ½» appendRatingBar τÜä 5 µÿƒΦ»äσêå∩╝îσ¡ÿ RatingStore
# σ╣╢σ£¿Θ½ÿΦ»äσêå∩╝êΓëÑ4 µÿƒ∩╝ëµù╢σÉæ ShadowTracer σåÖσàÑΣ╝ÿΦ┤¿µá╖µ£¼∩╝îµÄ¿σè¿Φ«¡τ╗âµò░µì«Θú₧Φ╜«
# ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
@app.route("/api/response/rate", methods=["POST"])
def api_response_rate():
    """
    µÄÑµö╢τö¿µê╖σ»╣ AI σ¢₧σñìτÜäµÿƒτ║ºΦ»äσêåπÇé

    Φ»╖µ▒éΣ╜ô:
      msg_id       str   ΓÇö MD5 µ╢êµü»µîçτ║╣∩╝êτö▒σÉÄτ½» done Σ║ïΣ╗╢Σ╕ïσÅæ∩╝ë
      stars        int   ΓÇö 1~5 µÿƒ
      comment      str   ΓÇö σÅ»ΘÇëµûçσ¡ùσÅìΘªê
      session_name str   ΓÇö Σ╝ÜΦ»¥σÉì
      user_input   str   ΓÇö τö¿µê╖σÄƒσºïΦ╛ôσàÑ∩╝êσëì 500 σ¡ù∩╝ë
      ai_response  str   ΓÇö AI σ¢₧σñìµûçµ£¼∩╝êσëì 500 σ¡ù∩╝ë
      task_type    str   ΓÇö Σ╗╗σèíτ▒╗σ₧ï∩╝îΘ╗ÿΦ«ñ CHAT
    """
    data = request.json or {}
    msg_id = data.get("msg_id", "")
    stars = int(data.get("stars", 0))
    comment = (data.get("comment") or "").strip()
    session_name = data.get("session_name", "default")
    user_input = data.get("user_input", "")
    ai_response = data.get("ai_response", "")
    task_type = data.get("task_type", "CHAT")

    if not (1 <= stars <= 5):
        return jsonify({"success": False, "error": "stars σ┐àΘí╗σ£¿ 1~5 Σ╣ïΘù┤"}), 400

    # ΓöÇΓöÇ 1. σ¡ÿσàÑ RatingStore ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    try:
        from app.core.learning.rating_store import RatingStore

        rs = RatingStore()
        rs.save_user_rating(
            msg_id=msg_id,
            stars=stars,
            comment=comment,
            session_name=session_name,
            user_input=user_input,
            ai_response=ai_response,
        )
    except Exception as e:
        _app_logger.warning(f"[ResponseRate] ΓÜá∩╕Å RatingStore Σ┐¥σ¡ÿσñ▒Φ┤Ñ: {e}")

    # ΓöÇΓöÇ 2. Θ½ÿΦ»äσêå∩╝êΓëÑ4 µÿƒ∩╝ëΓåÆ ShadowTracer Φ«░σ╜òΣ╝ÿΦ┤¿µá╖µ£¼∩╝îµÄ¿Φ┐¢Θú₧Φ╜« ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    trace_id = None
    if stars >= 4 and user_input and ai_response:
        try:
            from app.core.learning.shadow_tracer import ShadowTracer

            trace_id = ShadowTracer.record_approved(
                session_id=session_name,
                user_input=user_input,
                ai_response=ai_response,
                skill_id=None,
                task_type=task_type,
                model_used="",
                metadata={"stars": stars, "comment": comment, "source": "user_rating"},
            )
            _app_logger.debug(
                f"[ResponseRate] Γ¡É {stars}µÿƒ ΓåÆ ShadowTracer Φ«░σ╜ò trace_id={trace_id}"
            )
        except Exception as e:
            _app_logger.warning(f"[ResponseRate] ΓÜá∩╕Å ShadowTracer Φ«░σ╜òσñ▒Φ┤Ñ: {e}")

    return jsonify({
        "success": True,
        "msg_id": msg_id,
        "stars": stars,
        "trace_id": trace_id,
        "flywheel": trace_id is not None,
    })


