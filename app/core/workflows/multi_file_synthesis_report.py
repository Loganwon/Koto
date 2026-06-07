# ══════════════════════════════════════════════════════════════
# multi_file_synthesis_report.py — 多文档综合研究报告
#
# 用户场景：
#   上传 2-10 份文档（研报、论文、报告），
#   AI 逐份提取核心发现 → 交叉分析 → 生成结构化 DOCX 报告。
# ══════════════════════════════════════════════════════════════

from __future__ import annotations

import json
import logging
from typing import Any

from app.core.workflow_engine import (
    WorkflowExecutor,
    sse_error,
    sse_output,
    sse_progress,
    sse_step_done,
    sse_step_start,
)

logger = logging.getLogger(__name__)

_MAX_DOC_CHARS = 6000

_EXTRACT_SYSTEM = """你是一名研究分析师。
请从以下文档中提取 {n_findings} 个最重要的核心发现/观点。

输出 JSON 数组，每个元素：
{{
  "finding": "核心发现（一句话）",
  "detail": "支撑细节（2-3句）",
  "source_quote": "原文关键引用（如有）"
}}

只输出 JSON 数组，不要任何说明。"""

_SYNTHESIS_SYSTEM = """你是一名高级研究分析师。
请基于多份文档的核心发现进行交叉分析。

分析维度：
1. 共同主题：多份文档共同指向的趋势或结论
2. 矛盾点：不同文档之间的分歧或冲突
3. 综合结论：综合所有来源得出的整体判断
{focus_hint}

输出 JSON 对象：
{{
  "common_themes": ["主题1", "主题2", ...],
  "contradictions": ["矛盾点1", ...],
  "synthesis": "综合结论（200字以内）",
  "recommendations": ["建议1", "建议2", ...]
}}

只输出 JSON，不要任何说明。"""


class MultiFileSynthesisReport(WorkflowExecutor):
    """
    多文档综合研究报告工作流。

    params 期望字段:
        source_files:  List[str] — 源文件路径
        report_title:  str       — 报告标题（可选）
        focus:         str       — 分析重点（可选）
        model_mode:    str       — "auto" | "local"
    """

    WORKFLOW_ID = "multi_file_synthesis_report"
    WORKFLOW_NAME = "多文档综合报告"

    def execute(self, params: dict, yield_event) -> Any:
        files: list[str] = params.get("source_files") or []
        title: str = params.get("report_title") or "综合研究报告"
        focus: str = params.get("focus") or ""
        model_mode: str = params.get("model_mode") or "auto"

        if len(files) < 2:
            yield sse_error("请至少上传 2 份文档（source_files）")
            return

        is_local = model_mode == "local"
        max_files = 4 if is_local else 10
        files = files[:max_files]
        n_findings = 2 if is_local else 4

        # ── Step 1: 逐文件解析 + 提取核心发现 ───────────────────────
        yield sse_step_start("extract", f"📚 分析 {len(files)} 份文档…")
        all_findings: list[dict] = []  # [{source, findings: [...]}]

        system = _EXTRACT_SYSTEM.format(n_findings=n_findings)

        for idx, fpath in enumerate(files):
            fname = fpath.split("\\")[-1].split("/")[-1]
            yield sse_progress(idx + 1, len(files), fname)

            text = self.parse_file(fpath)
            if not text.strip():
                logger.warning("[SynthesisReport] 文件无内容: %s", fpath)
                all_findings.append({"source": fname, "findings": []})
                continue

            findings = self._extract_findings(text[:_MAX_DOC_CHARS], system, model_mode)
            all_findings.append({"source": fname, "findings": findings})

        yield sse_step_done("extract", f"📚 已提取 {sum(len(f['findings']) for f in all_findings)} 个发现")

        # ── Step 2: 交叉分析 ────────────────────────────────────────
        synthesis = {}
        if not is_local:
            yield sse_step_start("synthesis", "🔬 交叉分析…")
            synthesis = self._cross_analyze(all_findings, focus, model_mode)
            yield sse_step_done("synthesis", "🔬 分析完成")

        # ── Step 3: 生成执行摘要 ────────────────────────────────────
        yield sse_step_start("summary", "📝 生成执行摘要…")
        exec_summary = synthesis.get("synthesis", "")
        if not exec_summary:
            # 简单合并各来源的第一个发现
            parts = []
            for f in all_findings:
                if f["findings"]:
                    parts.append(f["findings"][0].get("finding", ""))
            exec_summary = "；".join(parts[:5])
        yield sse_step_done("summary", "📝 摘要完成")

        # ── Step 4: 构建 DOCX ───────────────────────────────────────
        yield sse_step_start("docx", "📄 生成报告文档…")
        output_path = self.save_output_file(".docx")
        self._build_docx(title, exec_summary, all_findings, synthesis, str(output_path))
        yield sse_step_done("docx", "📄 文档生成完成")

        # ── 输出 ─────────────────────────────────────────────────────
        preview = self._build_markdown(title, exec_summary, all_findings, synthesis)
        yield sse_output(
            "docx_file",
            {"path": str(output_path), "filename": f"{title}.docx"},
            "综合报告",
        )
        yield sse_output("markdown", preview, "报告预览")

    # ── 辅助方法 ──────────────────────────────────────────────────────

    def _extract_findings(self, text: str, system: str, model_mode: str) -> list[dict]:
        prompt = f"请分析以下文档：\n\n{text}"
        try:
            result = self.llm_json(prompt, system=system, model_mode=model_mode)
            if isinstance(result, list):
                return result
        except Exception as e:
            logger.warning("[SynthesisReport] 提取失败: %s", e)
        return []

    def _cross_analyze(self, all_findings: list[dict], focus: str, model_mode: str) -> dict:
        """跨文档交叉分析。"""
        findings_text = json.dumps(
            [{
                "source": f["source"],
                "findings": [fi.get("finding", "") for fi in f["findings"]]
            } for f in all_findings],
            ensure_ascii=False, indent=2
        )
        focus_hint = f"\n特别关注: {focus}" if focus else ""
        system = _SYNTHESIS_SYSTEM.format(focus_hint=focus_hint)
        prompt = f"以下是来自 {len(all_findings)} 份文档的核心发现：\n\n{findings_text}"

        try:
            result = self.llm_json(prompt, system=system, model_mode=model_mode)
            if isinstance(result, dict):
                return result
        except Exception as e:
            logger.warning("[SynthesisReport] 交叉分析失败: %s", e)
        return {}

    def _build_docx(self, title: str, exec_summary: str, all_findings: list[dict],
                    synthesis: dict, output_path: str) -> None:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 标题
        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 执行摘要
        doc.add_heading("执行摘要", level=2)
        p = doc.add_paragraph(exec_summary)
        for run in p.runs:
            run.font.size = Pt(11)

        # 各来源发现
        doc.add_heading("各文档核心发现", level=2)
        for item in all_findings:
            doc.add_heading(item["source"], level=3)
            for fi in item["findings"]:
                finding = fi.get("finding", "")
                detail = fi.get("detail", "")
                p = doc.add_paragraph()
                run = p.add_run(f"• {finding}")
                run.bold = True
                run.font.size = Pt(10)
                if detail:
                    p.add_run(f"\n  {detail}").font.size = Pt(10)

        # 交叉分析
        if synthesis:
            themes = synthesis.get("common_themes", [])
            contradictions = synthesis.get("contradictions", [])
            recommendations = synthesis.get("recommendations", [])

            if themes:
                doc.add_heading("共同主题", level=2)
                for t in themes:
                    doc.add_paragraph(str(t), style="List Bullet")

            if contradictions:
                doc.add_heading("分歧与矛盾", level=2)
                for c in contradictions:
                    doc.add_paragraph(str(c), style="List Bullet")

            if recommendations:
                doc.add_heading("建议", level=2)
                for r in recommendations:
                    doc.add_paragraph(str(r), style="List Number")

        doc.save(output_path)

    def _build_markdown(self, title: str, exec_summary: str, all_findings: list[dict],
                        synthesis: dict) -> str:
        lines = [f"# {title}\n", f"## 执行摘要\n\n{exec_summary}\n"]

        lines.append("## 各文档发现\n")
        for item in all_findings:
            lines.append(f"\n### {item['source']}\n")
            for fi in item["findings"]:
                lines.append(f"- **{fi.get('finding', '')}**: {fi.get('detail', '')}")

        if synthesis:
            themes = synthesis.get("common_themes", [])
            if themes:
                lines.append("\n## 共同主题\n")
                for t in themes:
                    lines.append(f"- {t}")

        return "\n".join(lines)
