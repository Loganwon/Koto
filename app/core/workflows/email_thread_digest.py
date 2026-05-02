# ══════════════════════════════════════════════════════════════
# email_thread_digest.py — 邮件/消息链摘要
#
# 用户场景：
#   导出的邮件往来或群聊记录文件，
#   AI 提取参与者、时间线、决策、待办事项，
#   生成结构化 DOCX 纪要文档。
# ══════════════════════════════════════════════════════════════

from __future__ import annotations

import json
import logging
from typing import Any

from app.core.workflow_engine import (
    WorkflowExecutor,
    sse_error,
    sse_output,
    sse_step_done,
    sse_step_start,
)

logger = logging.getLogger(__name__)

_EXTRACT_SYSTEM = """你是一名高效的会议/沟通纪要助手。
请从以下邮件/消息记录中提取结构化信息。

输出语言: {lang}

输出 JSON 对象，格式：
{{
  "title": "主题摘要（一句话）",
  "participants": ["参与者1", "参与者2", ...],
  "date_range": "时间范围描述",
  "timeline": [
    {{"time": "时间", "speaker": "发言人", "summary": "内容摘要"}}
  ],
  "decisions": ["决策1", "决策2", ...],
  "open_questions": ["未解决问题1", ...],
  "action_items": [
    {{"owner": "负责人", "task": "任务描述", "deadline": "截止日期（如有）"}}
  ]
}}

规则：
1. timeline 最多 15 条（合并低价值的寒暄/确认）
2. action_items 重点提取，不要遗漏
3. 如果找不到明确信息，对应字段设为空数组
4. 只输出 JSON，不要任何说明"""

_MAX_TEXT = 8000


class EmailThreadDigest(WorkflowExecutor):
    """
    邮件/消息链摘要工作流。

    params 期望字段:
        thread_file: str — 邮件/聊天记录文件路径
        output_lang: str — "zh" | "en"
        model_mode:  str — "auto" | "local"
    """

    WORKFLOW_ID = "email_thread_digest"
    WORKFLOW_NAME = "邮件消息摘要"

    def execute(self, params: dict, yield_event) -> Any:
        thread_file: str = params.get("thread_file") or ""
        output_lang: str = params.get("output_lang") or "zh"
        model_mode: str = params.get("model_mode") or "auto"

        if not thread_file:
            yield sse_error("请提供邮件/消息文件（thread_file）")
            return

        # ── Step 1: 解析文件 ─────────────────────────────────────────
        yield sse_step_start("parse", "📧 解析邮件/消息记录…")
        text = self._parse_thread(thread_file)
        if not text.strip():
            yield sse_error("文件内容为空或无法解析")
            return

        if model_mode == "local":
            text = text[:4000]
        else:
            text = text[:_MAX_TEXT]
        yield sse_step_done("parse", f"📧 已读取（{len(text)} 字）")

        # ── Step 2: LLM 结构化提取 ──────────────────────────────────
        yield sse_step_start("extract", "🤖 AI 结构化分析…")
        lang_label = "中文" if output_lang == "zh" else "English"
        system = _EXTRACT_SYSTEM.format(lang=lang_label)

        if model_mode == "local":
            # 本地模型简化：只提取 decisions + action_items
            system = system.replace(
                '"timeline": [',
                '"timeline": [],  // 本地模式跳过 timeline\n    // '
            )

        structured = self._extract_structure(text, system, model_mode)
        if not structured:
            yield sse_error("AI 提取失败，请重试")
            return
        yield sse_step_done("extract", "🤖 结构化分析完成")

        # ── Step 3: 生成 DOCX ───────────────────────────────────────
        yield sse_step_start("docx", "📄 生成纪要文档…")
        output_path = self.save_output_file(".docx")
        self._build_docx(structured, str(output_path), output_lang)
        yield sse_step_done("docx", "📄 文档生成完成")

        # ── Step 4: Markdown 预览 ───────────────────────────────────
        preview = self._build_markdown(structured)

        # ── 输出 ─────────────────────────────────────────────────────
        title = structured.get("title", "纪要")
        yield sse_output(
            "docx_file",
            {"path": str(output_path), "filename": f"{title}.docx"},
            "纪要文档",
        )
        yield sse_output("markdown", preview, "快速预览")

    # ── 辅助方法 ──────────────────────────────────────────────────────

    def _parse_thread(self, file_path: str) -> str:
        """解析邮件/聊天文件。.eml 额外处理 headers。"""
        from pathlib import Path
        ext = Path(file_path).suffix.lower()

        if ext == ".eml":
            return self._parse_eml(file_path)
        return self.parse_file(file_path)

    def _parse_eml(self, file_path: str) -> str:
        """解析 .eml 邮件文件。"""
        import email
        from email.header import decode_header

        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                msg = email.message_from_file(f)

            parts = []
            # Headers
            for h in ("From", "To", "Cc", "Subject", "Date"):
                val = msg.get(h, "")
                if val:
                    decoded = decode_header(val)
                    text_parts = []
                    for part, enc in decoded:
                        if isinstance(part, bytes):
                            text_parts.append(part.decode(enc or "utf-8", errors="replace"))
                        else:
                            text_parts.append(part)
                    parts.append(f"{h}: {''.join(text_parts)}")

            parts.append("")  # separator

            # Body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload:
                            charset = part.get_content_charset() or "utf-8"
                            parts.append(payload.decode(charset, errors="replace"))
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    charset = msg.get_content_charset() or "utf-8"
                    parts.append(payload.decode(charset, errors="replace"))

            return "\n".join(parts)
        except Exception as e:
            logger.warning("[EmailDigest] EML 解析失败: %s", e)
            return self.parse_file(file_path)

    def _extract_structure(self, text: str, system: str, model_mode: str) -> dict | None:
        """调用 LLM 提取结构化数据。"""
        prompt = f"请分析以下邮件/消息记录：\n\n{text}"
        try:
            result = self.llm_json(prompt, system=system, model_mode=model_mode)
            if isinstance(result, dict):
                return result
        except Exception as e:
            logger.warning("[EmailDigest] LLM 提取失败: %s", e)
        return None

    def _build_docx(self, data: dict, output_path: str, lang: str) -> None:
        """用 python-docx 构建结构化纪要文档。"""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 标题
        title = data.get("title", "沟通纪要")
        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 基本信息
        participants = data.get("participants", [])
        date_range = data.get("date_range", "")
        if participants or date_range:
            p = doc.add_paragraph()
            if participants:
                run = p.add_run(f"参与者: {', '.join(participants)}")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            if date_range:
                if participants:
                    p.add_run("  |  ").font.size = Pt(10)
                run = p.add_run(f"时间: {date_range}")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # 时间线
        timeline = data.get("timeline", [])
        if timeline:
            doc.add_heading("时间线", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "时间"
            hdr[1].text = "发言人"
            hdr[2].text = "内容"
            for item in timeline:
                row = table.add_row().cells
                row[0].text = str(item.get("time", ""))
                row[1].text = str(item.get("speaker", ""))
                row[2].text = str(item.get("summary", ""))

        # 决策
        decisions = data.get("decisions", [])
        if decisions:
            doc.add_heading("决策事项", level=2)
            for d in decisions:
                p = doc.add_paragraph(str(d), style="List Bullet")

        # 未解决问题
        questions = data.get("open_questions", [])
        if questions:
            doc.add_heading("待讨论事项", level=2)
            for q in questions:
                doc.add_paragraph(str(q), style="List Bullet")

        # 待办事项
        action_items = data.get("action_items", [])
        if action_items:
            doc.add_heading("待办事项", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "负责人"
            hdr[1].text = "任务"
            hdr[2].text = "截止日期"
            for item in action_items:
                row = table.add_row().cells
                row[0].text = str(item.get("owner", ""))
                row[1].text = str(item.get("task", ""))
                row[2].text = str(item.get("deadline", ""))

        doc.save(output_path)

    def _build_markdown(self, data: dict) -> str:
        """生成 Markdown 快速预览。"""
        lines = [f"# {data.get('title', '纪要')}\n"]

        participants = data.get("participants", [])
        if participants:
            lines.append(f"**参与者**: {', '.join(participants)}\n")

        decisions = data.get("decisions", [])
        if decisions:
            lines.append("## 决策\n")
            for d in decisions:
                lines.append(f"- {d}")

        action_items = data.get("action_items", [])
        if action_items:
            lines.append("\n## 待办事项\n")
            for item in action_items:
                owner = item.get("owner", "?")
                task = item.get("task", "")
                deadline = item.get("deadline", "")
                line = f"- [ ] **{owner}**: {task}"
                if deadline:
                    line += f" (截止: {deadline})"
                lines.append(line)

        return "\n".join(lines)
