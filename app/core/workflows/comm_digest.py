# ══════════════════════════════════════════════════════════════
# comm_digest.py — 沟通纪要生成
#
# 合并原 action_item_extractor + email_thread_digest
#
# 用户场景：
#   上传邮件、群聊记录、会议纪要等沟通文件（或直接粘贴文本），
#   AI 提取参与者、时间线、决策、待办事项。
#   输出可选：
#     - "docx":     结构化 DOCX 纪要（参与者表、时间线表、待办表）
#     - "excel":    彩色待办事项 Excel（按优先级高亮）
#     - "markdown": Markdown 进度报告
#     - "auto":     有文件输入则 DOCX，纯文本输入则 Excel
# ══════════════════════════════════════════════════════════════

from __future__ import annotations

import json
import logging
from typing import Any

from app.core.workflow_engine import (
    WorkflowExecutor,
    sse_error,
    sse_output,
    sse_step_done,
    sse_step_start,
)

logger = logging.getLogger(__name__)

# ── LLM 指令 ─────────────────────────────────────────────────────────────────

_FULL_EXTRACT_SYSTEM = """你是一名高效的会议/沟通纪要助手。
请从以下邮件/消息/聊天记录中提取结构化信息。

输出语言: {lang}

输出 JSON 对象，格式：
{{
  "title": "主题摘要（一句话）",
  "participants": ["参与者1", "参与者2", ...],
  "date_range": "时间范围描述",
  "timeline": [
    {{"time": "时间", "speaker": "发言人", "summary": "内容摘要"}}
  ],
  "decisions": ["决策1", "决策2", ...],
  "open_questions": ["未解决问题1", ...],
  "action_items": [
    {{"owner": "负责人", "task": "任务描述", "deadline": "截止日期（如有）", "priority": "high|medium|low", "status": "pending|in_progress|done"}}
  ]
}}

规则：
1. timeline 最多 15 条（合并低价值的寒暄/确认）
2. action_items 重点提取，不要遗漏
3. 为每条 action_item 推断 priority（high/medium/low）和 status
4. 如果找不到明确信息，对应字段设为空数组
5. 只输出 JSON，不要任何说明"""

_REPORT_SYSTEM = """你是一个项目管理助手。
将以下待办事项列表整理为一份清晰的中文周报/进度报告（Markdown 格式）。

报告结构：
# 待办事项汇总
## ⚠️ 高优先级（需立即处理）
- ...
## 📅 正常推进
- ...
## ✅ 已完成
- ...
## ❓ 待确认责任人
- ...

要求：
- 按优先级分组
- 每条待办格式：`- [责任人] 任务描述（截止日期）`
- 语言简洁专业
"""

_MAX_TEXT = 12000


class CommDigest(WorkflowExecutor):
    """
    沟通纪要生成工作流（合并 action_item_extractor + email_thread_digest）。

    params 期望字段:
        files:       List[str] — 文件路径列表（txt/eml/docx/pdf/md）
        texts:       List[str] — 直接输入的文本片段列表
        output_mode: str       — "auto" | "docx" | "excel" | "markdown"
        output_lang: str       — "zh" | "en"
        model_mode:  str       — "auto" | "local"
    """

    WORKFLOW_ID = "comm_digest"
    WORKFLOW_NAME = "沟通纪要生成"

    def execute(self, params: dict, yield_event) -> Any:
        raw_texts: list[str] = params.get("texts") or []
        files: list[str] = params.get("files") or []
        output_mode: str = params.get("output_mode") or "auto"
        output_lang: str = params.get("output_lang") or "zh"
        model_mode: str = params.get("model_mode") or "auto"

        all_texts: list[str] = list(raw_texts)

        # ── Step 1: 解析文件 ──────────────────────────────────────────
        if files:
            yield sse_step_start("parse", f"📧 解析 {len(files)} 个文件…")
            for fp in files:
                text = self._parse_thread(fp)
                if text.strip():
                    fname = fp.split("\\")[-1].split("/")[-1]
                    all_texts.append(f"[文件: {fname}]\n{text}")
            yield sse_step_done("parse", f"📧 解析完成，累计 {len(all_texts)} 段文本")

        if not all_texts:
            yield sse_error("请提供至少一段文本（texts）或一个文件（files）")
            return

        # 决定输出格式
        if output_mode == "auto":
            output_mode = "docx" if files else "excel"

        # ── Step 2: LLM 结构化提取 ──────────────────────────────────
        yield sse_step_start("extract", "🤖 AI 结构化分析…")

        combined = "\n\n---\n\n".join(all_texts)
        if model_mode == "local":
            combined = combined[:4000]
        else:
            combined = combined[:_MAX_TEXT]

        structured = self._extract_full(combined, output_lang, model_mode)
        if not structured:
            yield sse_error("AI 提取失败，请重试")
            return

        action_items = structured.get("action_items", [])
        yield sse_step_done("extract", f"🤖 提取到 {len(action_items)} 条待办事项")

        # ── Step 3: 根据输出模式生成结果 ──────────────────────────────
        if output_mode == "docx":
            yield from self._output_docx(structured, output_lang)
        elif output_mode == "markdown":
            yield from self._output_markdown(structured, action_items, model_mode)
        else:  # excel
            yield from self._output_excel(structured, action_items, model_mode)

    # ── 输出模式 ──────────────────────────────────────────────────────────────

    def _output_docx(self, structured, output_lang):
        """生成结构化 DOCX 纪要文档。"""
        yield sse_step_start("docx", "📄 生成纪要文档…")
        output_path = self.save_output_file(".docx")
        self._build_docx(structured, str(output_path), output_lang)
        yield sse_step_done("docx", "📄 文档生成完成")

        preview = self._build_preview_markdown(structured)
        title = structured.get("title", "纪要")
        yield sse_output(
            "docx_file",
            {"path": str(output_path), "filename": f"{title}.docx"},
            "纪要文档",
        )
        yield sse_output("markdown", preview, "快速预览")

    def _output_excel(self, structured, action_items, model_mode):
        """生成彩色待办事项 Excel 表格。"""
        if not action_items:
            yield sse_output(
                "markdown", "# 沟通纪要\n\n未发现明确的待办事项。", "提取结果"
            )
            return

        yield sse_step_start("build_output", "📊 生成任务表格…")
        workbook = self._build_workbook(action_items)
        yield sse_step_done("build_output", "📊 表格生成完成")
        yield sse_output("xlsx_data", workbook, f"待办事项（{len(action_items)} 条）")

        preview = self._build_preview_markdown(structured)
        yield sse_output("markdown", preview, "纪要概览")

    def _output_markdown(self, structured, action_items, model_mode):
        """生成 Markdown 进度报告。"""
        if not action_items:
            preview = self._build_preview_markdown(structured)
            yield sse_output("markdown", preview, "沟通纪要")
            return

        yield sse_step_start("gen_report", "📝 生成进度报告…")
        items_json = json.dumps(action_items, ensure_ascii=False, indent=2)
        prompt = f"以下是提取的待办事项列表（JSON 格式）：\n\n{items_json}\n\n请生成一份中文进度报告。"
        try:
            report = self.llm(prompt, system=_REPORT_SYSTEM, model_mode=model_mode)
        except Exception:
            report = self._build_preview_markdown(structured)
        yield sse_step_done("gen_report", "📝 报告生成完成")
        yield sse_output("markdown", report, f"待办报告（{len(action_items)} 条）")

    # ── 文件解析 ──────────────────────────────────────────────────────────────

    def _parse_thread(self, file_path: str) -> str:
        """解析邮件/聊天文件。.eml 额外处理 headers。"""
        from pathlib import Path

        ext = Path(file_path).suffix.lower()

        if ext == ".eml":
            return self._parse_eml(file_path)
        return self.parse_file(file_path)

    def _parse_eml(self, file_path: str) -> str:
        """解析 .eml 邮件文件。"""
        import email
        from email.header import decode_header

        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                msg = email.message_from_file(f)

            parts = []
            for h in ("From", "To", "Cc", "Subject", "Date"):
                val = msg.get(h, "")
                if val:
                    decoded = decode_header(val)
                    text_parts = []
                    for part, enc in decoded:
                        if isinstance(part, bytes):
                            text_parts.append(
                                part.decode(enc or "utf-8", errors="replace")
                            )
                        else:
                            text_parts.append(part)
                    parts.append(f"{h}: {''.join(text_parts)}")

            parts.append("")

            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload:
                            charset = part.get_content_charset() or "utf-8"
                            parts.append(payload.decode(charset, errors="replace"))
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    charset = msg.get_content_charset() or "utf-8"
                    parts.append(payload.decode(charset, errors="replace"))

            return "\n".join(parts)
        except Exception as e:
            logger.warning("[CommDigest] EML 解析失败: %s", e)
            return self.parse_file(file_path)

    # ── LLM 提取 ─────────────────────────────────────────────────────────────

    def _extract_full(self, text: str, lang: str, model_mode: str) -> dict | None:
        """调用 LLM 提取完整结构化数据（含时间线、决策、待办）。"""
        lang_label = "中文" if lang == "zh" else "English"
        system = _FULL_EXTRACT_SYSTEM.format(lang=lang_label)
        prompt = f"请分析以下沟通记录：\n\n{text}"
        try:
            result = self.llm_json(prompt, system=system, model_mode=model_mode)
            if isinstance(result, dict):
                return result
        except Exception as e:
            logger.warning("[CommDigest] LLM 提取失败: %s", e)
        return None

    # ── DOCX 构建 ────────────────────────────────────────────────────────────

    def _build_docx(self, data: dict, output_path: str, lang: str) -> None:
        """用 python-docx 构建结构化纪要文档。"""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        doc = Document()

        # 标题
        title = data.get("title", "沟通纪要")
        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 基本信息
        participants = data.get("participants", [])
        date_range = data.get("date_range", "")
        if participants or date_range:
            p = doc.add_paragraph()
            if participants:
                run = p.add_run(f"参与者: {', '.join(participants)}")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            if date_range:
                if participants:
                    p.add_run("  |  ").font.size = Pt(10)
                run = p.add_run(f"时间: {date_range}")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # 时间线
        timeline = data.get("timeline", [])
        if timeline:
            doc.add_heading("时间线", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "时间"
            hdr[1].text = "发言人"
            hdr[2].text = "内容"
            for item in timeline:
                row = table.add_row().cells
                row[0].text = str(item.get("time", ""))
                row[1].text = str(item.get("speaker", ""))
                row[2].text = str(item.get("summary", ""))

        # 决策
        decisions = data.get("decisions", [])
        if decisions:
            doc.add_heading("决策事项", level=2)
            for d in decisions:
                doc.add_paragraph(str(d), style="List Bullet")

        # 未解决问题
        questions = data.get("open_questions", [])
        if questions:
            doc.add_heading("待讨论事项", level=2)
            for q in questions:
                doc.add_paragraph(str(q), style="List Bullet")

        # 待办事项
        action_items = data.get("action_items", [])
        if action_items:
            doc.add_heading("待办事项", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "负责人"
            hdr[1].text = "任务"
            hdr[2].text = "截止日期"
            hdr[3].text = "优先级"
            for item in action_items:
                row = table.add_row().cells
                row[0].text = str(item.get("owner", ""))
                row[1].text = str(item.get("task", ""))
                row[2].text = str(item.get("deadline", ""))
                row[3].text = str(item.get("priority", ""))

        doc.save(output_path)

    # ── Markdown 预览 ────────────────────────────────────────────────────────

    def _build_preview_markdown(self, data: dict) -> str:
        """生成 Markdown 快速预览。"""
        lines = [f"# {data.get('title', '纪要')}\n"]

        participants = data.get("participants", [])
        if participants:
            lines.append(f"**参与者**: {', '.join(participants)}\n")

        decisions = data.get("decisions", [])
        if decisions:
            lines.append("## 决策\n")
            for d in decisions:
                lines.append(f"- {d}")

        action_items = data.get("action_items", [])
        if action_items:
            lines.append("\n## 待办事项\n")
            for item in action_items:
                owner = item.get("owner", "?")
                task = item.get("task", "")
                deadline = item.get("deadline", "")
                priority = item.get("priority", "")
                line = f"- [ ] **{owner}**: {task}"
                if deadline:
                    line += f" (截止: {deadline})"
                if priority == "high":
                    line += " ⚠️"
                lines.append(line)

        return "\n".join(lines)

    # ── Excel 输出 ───────────────────────────────────────────────────────────

    def _build_workbook(self, items: list[dict]) -> dict:
        """将待办事项列表转为 Univer IWorkbookData 格式。"""
        import uuid as _uuid

        wb_id = str(_uuid.uuid4())[:8]
        sheet_id = "action_items"

        columns = ["task", "owner", "deadline", "status", "priority"]
        col_names_cn = {
            "task": "任务描述",
            "owner": "负责人",
            "deadline": "截止日期",
            "status": "状态",
            "priority": "优先级",
        }
        status_colors = {
            "overdue": "#ffcccc",
            "in_progress": "#d4edda",
            "pending": "#fff3cd",
            "done": "#e2e3e5",
        }
        priority_colors = {
            "high": "#ffcccc",
            "medium": "#fff3cd",
            "low": "#e2e3e5",
        }
        header_style = {"bl": 1, "bg": {"rgb": "#1a73e8"}, "cl": {"rgb": "#ffffff"}}

        cell_data: dict = {}
        row0: dict = {}
        for c, col in enumerate(columns):
            row0[str(c)] = {"v": col_names_cn[col], "t": 1, "s": header_style}
        cell_data["0"] = row0

        for r, item in enumerate(items, start=1):
            row_cells: dict = {}
            for c, col in enumerate(columns):
                v = item.get(col) or ""
                cell: dict = {"v": str(v), "t": 1}
                if col == "status" and str(v) in status_colors:
                    cell["s"] = {"bg": {"rgb": status_colors[str(v)]}}
                elif col == "priority" and str(v) in priority_colors:
                    cell["s"] = {"bg": {"rgb": priority_colors[str(v)]}}
                row_cells[str(c)] = cell
            cell_data[str(r)] = row_cells

        return {
            "id": wb_id,
            "name": "待办事项",
            "appVersion": "0.5.0",
            "sheetOrder": [sheet_id],
            "sheets": {
                sheet_id: {
                    "id": sheet_id,
                    "name": "待办事项",
                    "rowCount": len(items) + 1,
                    "columnCount": len(columns),
                    "cellData": cell_data,
                    "mergeData": [],
                }
            },
            "styles": {},
        }
