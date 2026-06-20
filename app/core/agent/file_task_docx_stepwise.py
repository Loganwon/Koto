# -*- coding: utf-8 -*-
# Copyright (C) 2024-2026 Koto AI. All rights reserved.
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from app.core.agent._file_task_stepwise_helpers import (
    file_task_suffix,
    stepwise_docx_polish_step_index,
    stepwise_docx_polish_window_paragraphs,
)
from app.core.agent.file_task_contract import FileTaskFile, FileTaskRequest
from app.core.agent.file_task_workflow_state import workflow_resume_control
from app.core.agent.tool_design_protocol import extract_first_json_value


def stepwise_docx_polish_target_path(
    request: FileTaskRequest, files: list[FileTaskFile]
) -> str:
    candidates: list[str] = []
    if request.target_path:
        candidates.append(str(request.target_path))
    resume_control = workflow_resume_control(request)
    for value in (
        resume_control.get("target_path"),
        resume_control.get("source_path"),
    ):
        if value:
            candidates.append(str(value))
    for file_info in files:
        if file_task_suffix(file_info) in {"doc", "docx"} and file_info.target:
            candidates.append(str(file_info.path or ""))
    for file_info in files:
        if file_task_suffix(file_info) in {"doc", "docx"}:
            candidates.append(str(file_info.path or ""))
    for candidate in candidates:
        clean = candidate.strip()
        if clean and clean.lower().endswith((".doc", ".docx")):
            return clean
    return ""


def read_docx_paragraph_window(
    request: FileTaskRequest, path: str
) -> dict[str, Any]:
    from docx import Document  # type: ignore

    doc = Document(path)
    visible_indices = [
        index
        for index, paragraph in enumerate(doc.paragraphs)
        if str(paragraph.text or "").strip()
    ]
    window_size = stepwise_docx_polish_window_paragraphs(request)
    step_index = stepwise_docx_polish_step_index(request)
    start_visible = step_index * window_size
    end_visible = min(start_visible + window_size, len(visible_indices))
    selected_indices = visible_indices[start_visible:end_visible]
    paragraphs = [doc.paragraphs[index].text for index in selected_indices]
    return {
        "source": Path(path).name,
        "path": path,
        "paragraph_indices": selected_indices,
        "paragraphs": paragraphs,
        "start_visible_index": start_visible,
        "end_visible_index": end_visible,
        "total_visible_paragraphs": len(visible_indices),
        "window_paragraphs": window_size,
        "step_index": step_index,
        "has_next": end_visible < len(visible_indices),
    }


def docx_polish_window_prompt(
    request: FileTaskRequest, paragraphs: list[str]
) -> str:
    numbered = "\n".join(
        f"{index}. {text}" for index, text in enumerate(paragraphs, start=1)
    )
    return (
        "请润色下面 DOCX 当前段落窗口。要求：\n"
        "1. 保持段落数量完全一致；\n"
        "2. 只改善语病、重复、口语化和不顺畅表达；\n"
        "3. 不改变事实、术语、数字和专名；\n"
        '4. 只返回 JSON 字符串数组，例如 ["润色后第1段", "润色后第2段"]。\n'
        f"用户任务：{request.task}\n\n"
        f"段落窗口：\n{numbered}"
    )


def parse_polished_docx_paragraphs(
    content: str, *, expected_count: int
) -> list[str]:
    text = str(content or "").strip()
    if not text:
        return []
    parsed = extract_first_json_value(text)
    if isinstance(parsed, dict):
        for key in ("paragraphs", "items", "result", "texts"):
            if isinstance(parsed.get(key), list):
                parsed = parsed.get(key)
                break
    if not isinstance(parsed, list):
        try:
            parsed = json.loads(text)
        except Exception:
            parsed = []
    if not isinstance(parsed, list):
        return []
    cleaned = [str(item or "").strip() for item in parsed[:expected_count]]
    cleaned = [item for item in cleaned if item]
    return cleaned if len(cleaned) == expected_count else []


def simple_polish_docx_paragraph(text: str) -> str:
    polished = re.sub(r"[ \t]+", " ", str(text or "")).strip()
    polished = re.sub(r"\s+([，。！？；：、])", r"\1", polished)
    polished = re.sub(r"([（【])\s+", r"\1", polished)
    polished = re.sub(r"\s+([）】])", r"\1", polished)
    return polished or str(text or "")


def rewrite_docx_paragraph_window(
    path: str, paragraph_indices: list[int], paragraphs: list[str]
) -> int:
    from docx import Document  # type: ignore

    doc = Document(path)
    changed = 0
    for paragraph_index, new_text in zip(paragraph_indices, paragraphs):
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            continue
        paragraph = doc.paragraphs[paragraph_index]
        if paragraph.text == new_text:
            continue
        for run in list(paragraph.runs):
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
        changed += 1
    doc.save(path)
    return changed


def docx_polish_wait_artifact(
    request: FileTaskRequest, target_path: str, window: dict[str, Any]
) -> dict[str, Any]:
    next_step_index = int(window.get("step_index") or 0) + 1
    window_paragraphs = int(window.get("window_paragraphs") or 8)
    next_start = int(window.get("end_visible_index") or 0) + 1
    next_end = min(
        next_start + window_paragraphs - 1,
        int(window.get("total_visible_paragraphs") or next_start),
    )
    label = (
        f"继续第 {next_start}-{next_end} 段"
        if bool(window.get("has_next"))
        else "已无下一段"
    )
    return {
        "artifact_type": "koto_stepwise_resume_v1",
        "category": "stepwise_confirmation",
        "route": "long_docx_stepwise_polish_writeback",
        "status": "awaiting_confirmation",
        "summary": "上一段落窗口已写回 DOCX。可以继续处理下一段。",
        "suggested_next_step": label,
        "actions": [
            {
                "type": "file_task_resume",
                "label": label,
                "enabled": bool(window.get("has_next")),
                "request": {
                    "task": f"继续分步润色 {Path(target_path).name}",
                    "target_path": target_path,
                    "files": [
                        {
                            "path": target_path,
                            "name": Path(target_path).name,
                            "type": "docx",
                            "target": True,
                        }
                    ],
                    "options": {
                        "workflow_checkpoint": {
                            "policy": "confirm_each_step",
                            "step_index": next_step_index,
                            "window_paragraphs": window_paragraphs,
                            "target_path": target_path,
                            "source_path": target_path,
                            "original_task": request.task,
                            "route": "long_docx_stepwise_polish_writeback",
                        }
                    },
                },
            }
        ],
        "stepwise": {
            "current_step_index": int(window.get("step_index") or 0),
            "next_step_index": next_step_index,
            "window_paragraphs": window_paragraphs,
            "has_next": bool(window.get("has_next")),
            "paragraph_start": int(window.get("start_visible_index") or 0) + 1,
            "paragraph_end": int(window.get("end_visible_index") or 0),
            "total_visible_paragraphs": int(
                window.get("total_visible_paragraphs") or 0
            ),
        },
    }
