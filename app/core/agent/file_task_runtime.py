from __future__ import annotations

import json
import logging
import os
import re
import tempfile
import threading
import time
import uuid
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Optional

from app.core.agent.file_task_contract import (
    FileTaskClassification,
    FileTaskExecutionContext,
    FileTaskExecutionBrief,
    FileTaskEvent,
    FileTaskFile,
    FileTaskIntentPlan,
    FileTaskLedger,
    FileTaskRequirementSet,
    FileTaskRequest,
    FileTaskToolStreamChunk,
    FileTaskToolStreamResult,
)
from app.core.agent.file_task_intent_planner import FileTaskIntentPlanner
from app.core.agent.file_task_capability import (
    build_request_capability_profiles,
    matched_native_capability_names,
    native_tool_gap_for_request,
)
from app.core.agent.file_task_model import FileTaskModelClient
from app.core.agent.file_task_review_intent import (
    DOCX_REVIEW_INTENT_MARKERS,
    has_explicit_docx_review_intent,
    request_has_file_type,
)
from app.core.agent.file_task_recipes import (
    FileTaskRecipeMatch,
    recipe_matches,
    request_file_types,
    request_target_file_type,
    select_task_recipe,
    semantic_markers,
)
from app.core.agent.file_task_validation import (
    build_file_task_requirements,
    validate_file_task_plan,
)
from app.core.agent.file_task_whitebox import (
    WhiteboxExecutionPlan,
    build_decision_audit,
    build_recipe_skeleton,
    extract_whitebox_execution_plan,
    validate_whitebox_plan,
    whitebox_execution_plan_schema,
)
from app.core.agent.file_task_tool_catalog import (
    extract_koto_paths,
    extract_sandbox_artifacts,
    file_states_for_changes,
    is_file_task_tool,
    is_write_tool,
    parse_file_change,
    supported_file_workflows,
    tool_result_preview,
    stringify_result,
    write_target_for_tool,
)
from app.core.agent.file_task_tool_gateway import (
    FileTaskToolContext,
    FileTaskToolGateway,
    FileTaskToolProvider,
    ToolExecutor,
)
from app.core.agent.tool_design_protocol import (
    TOOL_DESIGN_PROTOCOL,
    build_next_action_artifact,
    extract_first_json_value,
    extract_tool_gap_from_response,
    merge_tool_gaps,
    tool_design_prompt_text,
)
from app.core.shared.tool_parser import parse_task_tool_calls

logger = logging.getLogger(__name__)

ModelCaller = Callable[..., Dict[str, Any]]

_CANCELLED_RUNS: Dict[str, float] = {}
_CANCEL_LOCK = threading.Lock()
_CANCEL_TTL_SECONDS = 60 * 60


def _clean_run_id(run_id: Any) -> str:
    return str(run_id or "").strip()


def _safe_float(value: Any, default: float = 0.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _prune_cancelled_runs(now: Optional[float] = None) -> None:
    current = time.time() if now is None else now
    expired = [
        run_id
        for run_id, ts in _CANCELLED_RUNS.items()
        if current - float(ts or 0) > _CANCEL_TTL_SECONDS
    ]
    for run_id in expired:
        _CANCELLED_RUNS.pop(run_id, None)


def request_cancel(run_id: str) -> bool:
    clean = _clean_run_id(run_id)
    if not clean:
        return False
    with _CANCEL_LOCK:
        _prune_cancelled_runs()
        already_requested = clean in _CANCELLED_RUNS
        _CANCELLED_RUNS[clean] = time.time()
    return not already_requested


def is_cancel_requested(run_id: str) -> bool:
    clean = _clean_run_id(run_id)
    if not clean:
        return False
    with _CANCEL_LOCK:
        _prune_cancelled_runs()
        return clean in _CANCELLED_RUNS


_READ_LIMIT = 12_000
_WRITE_INTENT_WORDS = (
    "修改",
    "写入",
    "生成",
    "创建",
    "替换",
    "插入",
    "更新",
    "保存",
    "导出",
    "写回",
    "加入",
    "添加",
    "追加",
    "附加",
    "导入",
    "合并",
    "填入",
    "填充",
    "批注",
    "标注",
    "审校",
    "校对",
    "润色",
    "改写",
    "重写",
    "美化",
    "排版",
    "套用主题",
    "应用主题",
    "设计主题",
    "设计风格",
    "fill",
    "write",
    "create",
    "insert",
    "update",
    "replace",
    "export",
    "add",
    "append",
    "import",
    "merge",
    "theme",
    "layout",
    "template",
    "style",
    "annotate",
    "comment",
    "review",
    "proofread",
    "rewrite",
    "polish",
)
_WRITE_INTENT_PATTERNS = (
    re.compile(r"放(?:到|进|入).{0,24}(?:页|页里|幻灯片|slide|slides)", re.IGNORECASE),
    re.compile(r"(?:新增|添加|生成|新建).{0,12}(?:页|幻灯片|slide|slides)", re.IGNORECASE),
    re.compile(
        r"(?:总结|概括).{0,20}(?:放(?:到|进|入)|生成).{0,20}(?:页|幻灯片|slide|slides)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:pptx?|slides?|幻灯片|演示文稿).{0,24}(?:补充|充实|扩写|完善).{0,18}(?:每一页|每页|逐页|各页|内容|文字|文本|页|slide|slides)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:每一页|每页|逐页|各页).{0,24}(?:补充|充实|扩写|完善).{0,18}(?:内容|文字|文本|页|slide|slides)?",
        re.IGNORECASE,
    ),
)
_EXPLICIT_WRITE_INTENT_WORDS = (
    "写入",
    "写回",
    "保存",
    "导出",
    "插入",
    "替换",
    "更新到",
    "应用到",
    "应用进",
    "同步到",
    "填入",
    "填充",
    "附加",
    "追加",
    "导入",
    "合并",
    "创建文件",
    "新建文件",
    "批注",
    "标注",
    "审校",
    "校对",
    "write back",
    "save",
    "export",
    "insert",
    "replace",
    "append",
)
_SOFT_WRITE_ACTION_WORDS = (
    "修改",
    "更新",
    "添加",
    "生成",
    "创建",
    "润色",
    "改写",
    "重写",
    "补充",
    "充实",
    "完善",
    "美化",
    "排版",
    "换",
)
_WRITE_TARGET_HINT_WORDS = (
    "文件",
    "文档",
    "word",
    "docx",
    "ppt",
    "pptx",
    "幻灯片",
    "slide",
    "slides",
    "页面",
    "页",
    "excel",
    "xlsx",
    "工作表",
    "sheet",
    "表格",
    "当前",
    "目标",
    "译稿",
    "原文",
    "文本",
    "段落",
)
_ANALYSIS_ADVICE_PATTERNS = (
    re.compile(
        r"(?:看看|看下|分析|评估|审查|review|review一下).{0,32}(?:哪些|哪里|什么地方|哪部分).{0,20}(?:需要|可以)?(?:修改|改进|优化|调整)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:哪些|哪里|什么地方|哪部分).{0,16}(?:需要|可以)?(?:修改|改进|优化|调整)(?:的地方|之处)?",
        re.IGNORECASE,
    ),
    re.compile(r"(?:修改建议|改进建议|优化建议|调整建议)", re.IGNORECASE),
    re.compile(r"(?:从大方向上|整体上|方向上).{0,12}(?:修改|改进|优化)", re.IGNORECASE),
)
_ANALYSIS_CUE_WORDS = (
    "分析",
    "看看",
    "看下",
    "评估",
    "审查",
    "review",
    "指出",
    "列出",
    "说明",
    "找出",
    "发现",
)
_ADVICE_CUE_WORDS = (
    "修改",
    "改进",
    "优化",
    "调整",
    "建议",
    "问题",
    "风险",
    "方向",
)
_DIAGNOSTIC_REQUEST_PATTERNS = (
    re.compile(
        r"^\s*(?:为什么|为啥|为何|怎么会|怎么没有|为什么没有|为什么没|失败原因|原因是什么|怎么回事|哪里出了问题|请解释|解释一下|说明一下|帮我解释|帮我说明)",
        re.IGNORECASE,
    ),
    re.compile(
        r"^\s*(?:这个任务|这次任务|这个结果|这次结果|上一轮|上次|这轮|这个流程|这次审校).{0,18}(?:为什么|为啥|为何|失败|出错|不对|有问题)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:为什么|为啥|为何).{0,20}(?:任务|结果|审校|修订|写回|批注|修改|删除|失败|报错|权限|permission denied)",
        re.IGNORECASE,
    ),
)
_DIAGNOSTIC_NEW_TASK_PATTERNS = (
    re.compile(
        r"(?:并|然后|顺便|同时).{0,8}(?:请|帮我|直接)?(?:修改|删除|写入|应用|批注|润色|重写|继续处理|重新执行)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:请|帮我|麻烦).{0,6}(?:直接|顺便)?(?:修改|删除|写入|应用|批注|润色|重写|继续处理|重新执行)",
        re.IGNORECASE,
    ),
)
_READONLY_WRITE_NEGATION_PATTERNS = (
    re.compile(
        r"(?:不要|不用|无需|不需要|不必|别|不).{0,10}(?:修改|改动|编辑|写入|写回|更新|保存|插入|删除|替换|应用|落盘|生成文件)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:do not|don't|dont|no need to|without).{0,24}(?:modify|edit|write|update|save|insert|replace|apply)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:read[ -]?only|only analyze|analysis only|answer only)", re.IGNORECASE
    ),
)
_RUN_PYTHON_STRONG_WRITE_PATTERNS = (
    re.compile(r"\bKOTO_MODIFIED\b"),
    re.compile(r"\.save\s*\(", re.IGNORECASE),
    re.compile(r"\.write_text\s*\(", re.IGNORECASE),
    re.compile(r"\.write_bytes\s*\(", re.IGNORECASE),
    re.compile(
        r"\bopen\s*\([^\n]{0,220},\s*['\"][^'\"]*[wax+][^'\"]*['\"]", re.IGNORECASE
    ),
    re.compile(r"\bto_(?:excel|csv|json|parquet|markdown)\s*\(", re.IGNORECASE),
    re.compile(r"\b(?:shutil\.)?(?:copy|copy2|move)\s*\(", re.IGNORECASE),
    re.compile(r"\b(?:os\.)?(?:remove|unlink|rename|replace)\s*\(", re.IGNORECASE),
)
_RUN_PYTHON_ARTIFACT_WRITE_PATTERNS = (
    re.compile(r"\bKOTO_CREATED\b"),
    re.compile(r"\bsavefig\s*\(", re.IGNORECASE),
    re.compile(
        r"\.save\s*\([^\n]{0,160}\.(?:png|jpg|jpeg|webp|svg)['\"]", re.IGNORECASE
    ),
)
_IMPERATIVE_WRITE_PATTERNS = (
    re.compile(
        r"^(?:请|帮我|麻烦)?(?:把|将)?(?:这个|当前|这份|该)?(?:文件|文档|word|ppt|表格|内容|文本|段落|译稿|稿件).{0,12}(?:修改|更新|润色|改写|重写|补充|完善)",
        re.IGNORECASE,
    ),
    re.compile(
        r"^(?:请|帮我|麻烦)?(?:直接|立刻)?(?:修改|更新|润色|改写|重写|补充|完善).{0,16}(?:文件|文档|word|ppt|表格|内容|文本|段落|译稿|稿件)",
        re.IGNORECASE,
    ),
)
_DOCX_ANNOTATE_INTENT_WORDS = DOCX_REVIEW_INTENT_MARKERS
_MAX_MODEL_ROUNDS = 6
_MAX_VERIFY_REPAIR_ATTEMPTS = 2
_MAX_WRITE_OPS_PER_FILE = 1
_KOTO_CREATED_MARKER = "__koto_created__:"
_KOTO_MODIFIED_MARKER = "__koto_modified__:"


def _preview(value: Any, limit: int = 700) -> str:
    text = str(value or "").strip()
    if len(text) <= limit:
        return text
    return text[: limit - 3] + "..."


def _file_task_suffix(file_info: FileTaskFile) -> str:
    explicit = str(getattr(file_info, "type", "") or "").strip().lower().lstrip(".")
    if explicit:
        return explicit
    candidate = str(
        getattr(file_info, "path", "") or getattr(file_info, "name", "") or ""
    )
    return Path(candidate).suffix.lower().lstrip(".")


def _looks_like_windowed_pdf_task(
    request: FileTaskRequest, recipe_skeleton: Dict[str, Any]
) -> bool:
    recipe_id = str((recipe_skeleton or {}).get("recipe_id") or "").strip()
    if recipe_id == "long_pdf_stepwise_docx_summary":
        return True
    options = request.options if isinstance(request.options, dict) else {}
    batch_control = (
        options.get("batch_control")
        if isinstance(options.get("batch_control"), dict)
        else {}
    )
    text = "\n".join(
        part
        for part in (
            str(getattr(request, "task", "") or ""),
            str(batch_control.get("original_task") or ""),
        )
        if part
    )
    batch_source_path = str(batch_control.get("source_path") or "").strip().lower()
    if str(
        batch_control.get("policy") or ""
    ).strip().lower() == "confirm_each_step" and (
        "pdf" in request_file_types(request.files) or batch_source_path.endswith(".pdf")
    ):
        return True
    return bool(
        re.search(
            r"(?:分步|一步一步|每一步|继续|下一段|下一页|按页|分页|stepwise|chunk)",
            text,
            re.IGNORECASE,
        )
        and re.search(r"(?:pdf|长文|很长|大量内容)", text, re.IGNORECASE)
    )


def _looks_like_stepwise_docx_polish_task(request: FileTaskRequest) -> bool:
    options = request.options if isinstance(request.options, dict) else {}
    batch_control = (
        options.get("batch_control")
        if isinstance(options.get("batch_control"), dict)
        else {}
    )
    text = "\n".join(
        part
        for part in (
            str(getattr(request, "task", "") or ""),
            str(batch_control.get("original_task") or ""),
        )
        if part
    )
    lowered = text.lower()
    has_docx = any(
        _file_task_suffix(file_info) in {"doc", "docx"}
        for file_info in (request.files or [])
    ) or str(request.target_path or "").lower().endswith((".doc", ".docx"))
    if not has_docx:
        return False
    has_polish = bool(
        re.search(
            r"(?:润色|改写|重写|优化表达|polish|rewrite|humanise|humanize)", text, re.IGNORECASE
        )
    )
    has_stepwise = (
        bool(
            re.search(
                r"(?:每完成一步|每一步|分步|一步一步|等待(?:我|用户)?确认|等我(?:来说)?继续|继续下一段|stepwise|each step|wait for)",
                text,
                re.IGNORECASE,
            )
        )
        or str(batch_control.get("policy") or "").strip().lower() == "confirm_each_step"
    )
    has_long = bool(
        re.search(r"(?:非常长|很长|大量内容|整篇|全文|长文|long|large)", text, re.IGNORECASE)
    ) or bool(batch_control)
    return (
        has_polish
        and has_stepwise
        and (has_long or "docx" in lowered or "word" in lowered or "文档" in text)
    )


def _stepwise_docx_polish_window_paragraphs(request: FileTaskRequest) -> int:
    options = request.options if isinstance(request.options, dict) else {}
    batch_control = (
        options.get("batch_control")
        if isinstance(options.get("batch_control"), dict)
        else {}
    )
    raw_value = (
        batch_control.get("window_paragraphs") or options.get("window_paragraphs") or 8
    )
    try:
        return max(1, min(int(raw_value), 24))
    except Exception:
        return 8


def _stepwise_docx_polish_step_index(request: FileTaskRequest) -> int:
    options = request.options if isinstance(request.options, dict) else {}
    batch_control = (
        options.get("batch_control")
        if isinstance(options.get("batch_control"), dict)
        else {}
    )
    try:
        return max(0, int(batch_control.get("step_index") or 0))
    except Exception:
        return 0


def _should_force_pdf_tool_read(
    request: FileTaskRequest,
    file_info: FileTaskFile,
    recipe_skeleton: Dict[str, Any],
) -> bool:
    if _file_task_suffix(file_info) != "pdf":
        return False
    if not getattr(file_info, "path", ""):
        return False
    if _looks_like_windowed_pdf_task(request, recipe_skeleton):
        return True
    return len(str(getattr(file_info, "content", "") or "")) > 8000


def _pdf_context_read_args(
    request: FileTaskRequest,
    file_info: FileTaskFile,
    recipe_skeleton: Dict[str, Any],
) -> Dict[str, Any]:
    args: Dict[str, Any] = {"path": file_info.path, "max_chars": _READ_LIMIT}
    if _looks_like_windowed_pdf_task(request, recipe_skeleton):
        window_pages = _stepwise_pdf_window_pages(request)
        step_index = _stepwise_pdf_step_index(request)
        start_page = 1 + step_index * window_pages
        args.update(
            {
                "start_page": start_page,
                "end_page": start_page + window_pages - 1,
                "max_chars": min(_READ_LIMIT, 9000),
            }
        )
    return args


def _stepwise_pdf_window_pages(request: FileTaskRequest) -> int:
    options = request.options if isinstance(request.options, dict) else {}
    batch_control = (
        options.get("batch_control")
        if isinstance(options.get("batch_control"), dict)
        else {}
    )
    raw_value = batch_control.get("window_pages") or options.get("window_pages") or 3
    try:
        return max(1, min(int(raw_value), 10))
    except Exception:
        return 3


def _stepwise_pdf_step_index(request: FileTaskRequest) -> int:
    options = request.options if isinstance(request.options, dict) else {}
    batch_control = (
        options.get("batch_control")
        if isinstance(options.get("batch_control"), dict)
        else {}
    )
    try:
        return max(0, int(batch_control.get("step_index") or 0))
    except Exception:
        return 0


def _normalized_pdf_body(value: Any) -> str:
    text = re.sub(r"\[Page\s+\d+\]", " ", str(value or ""), flags=re.IGNORECASE)
    text = re.sub(r"\s+", "", text)
    return text.strip()


def _pdf_text_quality(value: Any) -> Dict[str, Any]:
    body = _normalized_pdf_body(value)
    if not body:
        return {
            "usable": False,
            "reason": "empty_pdf_text",
            "char_count": 0,
            "unique_chars": 0,
        }
    unique_chars = len(set(body))
    alpha_num = len(re.findall(r"[A-Za-z0-9\u4e00-\u9fff]", body))
    cjk_chars = len(re.findall(r"[\u4e00-\u9fff]", body))
    repeated_watermark = bool(
        re.fullmatch(r"(?:考参通海泰国供仅|仅供国泰海通参考|用使点原禾元供仅荐推苇一|-)+", body)
    )
    low_density = alpha_num < 80 or unique_chars < 18
    mostly_single_repeats = (
        cjk_chars >= 20
        and cjk_chars / max(alpha_num, 1) > 0.5
        and (unique_chars / max(len(body), 1)) < 0.08
    )
    usable = not repeated_watermark and not low_density and not mostly_single_repeats
    reason = ""
    if not usable:
        if repeated_watermark:
            reason = "watermark_only_pdf_text"
        elif low_density:
            reason = "low_density_pdf_text"
        else:
            reason = "repetitive_pdf_text"
    return {
        "usable": usable,
        "reason": reason,
        "char_count": len(body),
        "unique_chars": unique_chars,
        "alpha_num_chars": alpha_num,
        "cjk_chars": cjk_chars,
    }


def _compact_line(value: Any, limit: int = 260) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    if len(text) <= limit:
        return text
    return text[: limit - 3] + "..."


def _json_payload(value: Any) -> Dict[str, Any]:
    if isinstance(value, dict):
        return value
    if not isinstance(value, str):
        return {}
    try:
        parsed = json.loads(value)
    except Exception:
        return {}
    return parsed if isinstance(parsed, dict) else {}


def _is_error_result(value: Any) -> bool:
    payload = _json_payload(value)
    if payload.get("error"):
        return True
    text = str(value or "").strip()
    return (
        text.startswith(("Error:", "Sandbox error:", "[error]")) or "\n[error]" in text
    )


def _sanitize_followup_file_changes(value: Any) -> List[Dict[str, Any]]:
    if isinstance(value, str):
        try:
            value = json.loads(value)
        except Exception:
            return []
    if not isinstance(value, list):
        return []

    cleaned: List[Dict[str, Any]] = []
    for item in value[:8]:
        if not isinstance(item, dict):
            continue
        entry: Dict[str, Any] = {}
        for key in (
            "path",
            "file_type",
            "operation",
            "summary",
            "source_path",
            "sheet",
            "requested_sheet",
            "table_title",
            "change_type",
            "original_target_path",
        ):
            text = str(item.get(key) or "").strip()
            if text:
                entry[key] = _preview(text, 400)
        for key in ("rows_written", "columns_written"):
            raw_value = item.get(key)
            if raw_value in (None, ""):
                continue
            try:
                entry[key] = int(raw_value)
            except Exception:
                continue
        if bool(item.get("fallback_copy")):
            entry["fallback_copy"] = True
        if entry:
            cleaned.append(entry)
    return cleaned


def _followup_has_prior_excel_docx_insert(followup_context: Dict[str, Any]) -> bool:
    for change in _sanitize_followup_file_changes(
        followup_context.get("previous_task_file_changes")
    ):
        if str(change.get("operation") or "").strip() == "insert_excel_as_docx_table":
            return True
    return False


class FileTaskRuntime:
    """First Koto-native whitebox runtime for file-assistant complex tasks.

    Typed file-task runtime with an allowlisted model -> tool -> checker loop.

    The model can plan and call tools freely, but only through the Koto-native
    tool catalog. The runtime owns event logging, duplicate-write guards, file
    change detection, and final verification.
    """

    def __init__(
        self,
        *,
        tool_executor: Optional[ToolExecutor] = None,
        tool_provider: Optional[FileTaskToolProvider] = None,
        tool_gateway: Optional[FileTaskToolGateway] = None,
        model_client: Optional[FileTaskModelClient | ModelCaller] = None,
        intent_planner: Optional[FileTaskIntentPlanner] = None,
        gemini_client: Any = None,
        workspace_root: str = "",
        max_rounds: int = _MAX_MODEL_ROUNDS,
    ):
        self._tool_executor = tool_executor
        self._tool_provider = tool_provider
        self._tool_gateway = tool_gateway
        self._model_client = model_client or FileTaskModelClient()
        self._intent_planner = intent_planner or FileTaskIntentPlanner()
        self._gemini_client = gemini_client
        self._workspace_root = workspace_root
        self._max_rounds = max(1, int(max_rounds or _MAX_MODEL_ROUNDS))

    def run(self, request: FileTaskRequest) -> Iterable[FileTaskEvent]:
        ledger = FileTaskLedger(request.run_id)
        if self._is_cancelled(request):
            yield self._cancelled_event(ledger, request)
            return

        context_files = self._context_files(request)
        from app.core.agent.file_task_financial_report_runner import (
            FileTaskFinancialReportRunner,
        )

        financial_report_runner = FileTaskFinancialReportRunner(self)
        if financial_report_runner.should_route(request, context_files):
            yield from financial_report_runner.stream(request, context_files)
            return

        base_classification = self._classify_request(request, context_files)
        intent_adjudication = self._adjudicate_intent_if_needed(
            request, context_files, base_classification
        )
        classification = self._apply_intent_adjudication(
            request, context_files, base_classification, intent_adjudication
        )
        execution_context = self._build_execution_context(
            request,
            context_files,
            classification=classification,
            intent_adjudication=intent_adjudication,
            quick_action_mode=self._quick_action_mode(request),
        )
        known_tool_gap = execution_context.known_tool_gap
        classification = execution_context.classification
        intent_plan = execution_context.intent_plan
        requirements = execution_context.requirements
        plan_check = execution_context.plan_check
        quick_action_mode = execution_context.quick_action_mode
        simple_quick_action = execution_context.simple_quick_action
        write_intent = execution_context.write_intent
        if classification.execution_mode != "doc_annotate_bridge":
            from app.core.agent import file_task_doc_annotate_boundary

            if file_task_doc_annotate_boundary.should_use_bridge_execution(request):
                classification.execution_mode = "doc_annotate_bridge"
                classification.task_family = "annotate"
                classification.operation_kind = "annotate"
                classification.output_mode = "write"
                classification.write_intent = True
                classification.docx_annotation_request = True
                if "annotate_file" not in classification.matched_capabilities:
                    classification.matched_capabilities.append("annotate_file")
                if "read_docx_content" not in classification.matched_capabilities:
                    classification.matched_capabilities.append("read_docx_content")
                if not classification.selected_recipe:
                    classification.selected_recipe = (
                        file_task_doc_annotate_boundary.bridge_recipe_id(request)
                    )
                classification.reason_codes.append(
                    "doc_annotate_bridge_execution_fallback"
                )
                write_intent = True
        bridge_execution_mode = classification.execution_mode == "doc_annotate_bridge"
        if bridge_execution_mode:
            tool_defs = []
            executor = None
        else:
            gateway = self._build_tool_gateway(request, context_files)
            tool_defs = self._tool_defs_for_classification(
                gateway.definitions(),
                classification,
            )
            executor = gateway.execute
        recipe_skeleton = build_recipe_skeleton(
            request,
            context_files,
            classification,
            intent_plan,
            tool_defs,
        )
        constraint_audit = self._constraint_audit(
            request,
            context_files,
            classification,
            intent_plan,
            requirements,
            recipe_skeleton,
        )

        classification_payload = classification.public_dict()
        intent_plan_payload = intent_plan.public_dict()
        requirements_payload = requirements.public_dict()
        plan_check_payload = plan_check.public_dict()
        plan_runtime = self._build_runtime_metadata(
            terminal_status="plan_checked",
            readonly_fallback_used=False,
            model_failed=False,
            planner_payload={
                "backend": execution_context.effective_planner_backend or "native",
                "source": "native",
                "policy": execution_context.effective_planner_policy or "native_only",
                "transport": "native",
                "reason": execution_context.effective_planner_reason
                or "file_task_native_only",
                "round": 1,
            },
            planner_fallback_payload={},
        )

        yield ledger.event(
            "run.started",
            {
                "task": request.task,
                "mode": "whitebox_v1",
                "file_count": len(context_files),
                "target_path": request.target_path,
                "model_mode": request.model_mode,
                "model_id": request.model_id,
                "quick_action_mode": quick_action_mode,
                "workflow_version": recipe_skeleton.get("version"),
                "recipe_skeleton": recipe_skeleton,
                "constraint_audit": constraint_audit,
                **classification_payload,
                "intent_plan": intent_plan_payload,
            },
        )

        if self._is_cancelled(request):
            yield self._cancelled_event(ledger, request)
            return

        if not simple_quick_action:
            yield ledger.event(
                "task.classified",
                execution_context.public_dict(),
                step_id="plan",
            )
        yield ledger.event(
            "plan.checked",
            {
                **plan_check_payload,
                "requirements": requirements_payload,
                "constraint_audit": constraint_audit,
                **(
                    {
                        "quick_action_bypass": True,
                    }
                    if simple_quick_action
                    else {}
                ),
            },
            step_id="plan",
        )

        if not plan_check.passed:
            yield ledger.event(
                "step.result",
                self._build_step_result_payload(
                    title="规划检查",
                    summary=plan_check.summary,
                    status="failed",
                    runtime=plan_runtime,
                    passed=False,
                ),
                step_id="plan",
            )
            yield ledger.event(
                "run.finished",
                {
                    "task": request.task,
                    "mode": "whitebox_v1",
                    "summary": plan_check.summary,
                    "completed_task": False,
                    "context": [],
                    "file_changes": [],
                    "runtime": plan_runtime,
                    "quick_action_mode": quick_action_mode,
                    "intent_plan": intent_plan_payload,
                    "requirements": requirements_payload,
                    "plan_check": plan_check_payload,
                    "recipe_skeleton": recipe_skeleton,
                    "constraint_audit": constraint_audit,
                    **classification_payload,
                },
            )
            return

        plan_steps = intent_plan.dynamic_steps or self._build_plan(
            request,
            context_files,
            write_intent,
            classification.output_mode,
            known_tool_gap,
        )
        if not simple_quick_action:
            yield ledger.event(
                "plan.created",
                {
                    "summary": self._plan_summary(request, context_files, write_intent),
                    "steps": plan_steps,
                    "success_criteria": self._success_criteria(
                        request, write_intent, classification.output_mode
                    ),
                    "tool_families": supported_file_workflows(),
                    "intent_plan": intent_plan_payload,
                    "recipe_skeleton": recipe_skeleton,
                    "constraint_audit": constraint_audit,
                },
            )

        if bridge_execution_mode:
            yield from self._stream_doc_annotate_bridge_execution(
                ledger,
                request,
                classification_payload=classification_payload,
                intent_plan_payload=intent_plan_payload,
                requirements_payload=requirements_payload,
                plan_check_payload=plan_check_payload,
                recipe_skeleton=recipe_skeleton,
                constraint_audit=constraint_audit,
                quick_action_mode=quick_action_mode,
            )
            return

        if _looks_like_stepwise_docx_polish_task(request):
            yield from self._stream_long_docx_stepwise_polish_writeback(
                ledger,
                request,
                context_files,
                classification,
                intent_plan,
                requirements_payload,
                plan_check_payload,
                recipe_skeleton,
                constraint_audit,
                quick_action_mode,
                classification_payload,
                intent_plan_payload,
            )
            return

        context_step_id = "context"
        if self._is_cancelled(request):
            yield self._cancelled_event(ledger, request)
            return
        yield ledger.event(
            "step.started",
            {
                "title": "读取显式上下文",
                "detail": "只使用用户附加、选中或明确指向的文件。",
            },
            step_id=context_step_id,
        )

        snippets: List[Dict[str, Any]] = []
        if request.selection:
            snippets.append(
                {
                    "source": request.selection_source or "selection",
                    "preview": _preview(request.selection, 500),
                    "chars": len(request.selection),
                }
            )
            yield ledger.event(
                "tool.finished",
                {
                    "tool_name": "selection_context",
                    "success": True,
                    "result_preview": _preview(request.selection, 500),
                },
                step_id=context_step_id,
            )

        for file_info in context_files:
            if self._is_cancelled(request):
                yield self._cancelled_event(ledger, request)
                return
            if (
                _looks_like_windowed_pdf_task(request, recipe_skeleton)
                and file_info.target
                and _file_task_suffix(file_info) in {"doc", "docx"}
            ):
                continue
            force_pdf_tool_read = _should_force_pdf_tool_read(
                request, file_info, recipe_skeleton
            )
            if file_info.content and not force_pdf_tool_read:
                snippets.append(
                    {
                        "source": file_info.name or file_info.path,
                        "path": file_info.path,
                        "preview": _preview(file_info.content, 500),
                        "chars": len(file_info.content),
                    }
                )
                yield ledger.event(
                    "tool.finished",
                    {
                        "tool_name": "provided_file_context",
                        "success": True,
                        "path": file_info.path,
                        "result_preview": _preview(file_info.content, 500),
                    },
                    step_id=context_step_id,
                )
                continue

            if not file_info.path:
                continue
            args = (
                _pdf_context_read_args(request, file_info, recipe_skeleton)
                if force_pdf_tool_read
                else {"path": file_info.path, "max_chars": _READ_LIMIT}
            )
            yield ledger.event(
                "tool.started",
                {
                    "tool_name": "parse_file_to_text",
                    "tool_args": args,
                },
                step_id=context_step_id,
            )
            try:
                result = executor("parse_file_to_text", args)
                success = not _is_error_result(result)
                if (
                    success
                    and force_pdf_tool_read
                    and args.get("start_page")
                    and not _pdf_text_quality(result).get("usable")
                ):
                    window_pages = max(
                        1,
                        int(args.get("end_page") or args.get("start_page") or 1)
                        - int(args.get("start_page") or 1)
                        + 1,
                    )
                    for _retry_index in range(3):
                        retry_args = dict(args)
                        retry_start = int(retry_args.get("start_page") or 1) + (
                            window_pages * (_retry_index + 1)
                        )
                        retry_args["start_page"] = retry_start
                        retry_args["end_page"] = retry_start + window_pages - 1
                        retry_result = executor("parse_file_to_text", retry_args)
                        if _is_error_result(retry_result):
                            continue
                        if _pdf_text_quality(retry_result).get("usable"):
                            args = retry_args
                            result = retry_result
                            success = True
                            break
            except Exception as exc:
                result = str(exc)
                success = False
                logger.warning("[FileTaskRuntime] parse_file_to_text failed: %s", exc)
            yield ledger.event(
                "tool.finished",
                {
                    "tool_name": "parse_file_to_text",
                    "success": success,
                    "result_preview": _preview(result),
                },
                step_id=context_step_id,
            )
            if success:
                snippet = {
                    "source": file_info.name or file_info.path,
                    "path": file_info.path,
                    "preview": _preview(result, 500),
                    "chars": len(str(result or "")),
                }
                if str(Path(str(file_info.path or "")).suffix).lower() == ".pdf":
                    if args.get("start_page"):
                        snippet["start_page"] = int(args.get("start_page") or 1)
                    if args.get("end_page"):
                        snippet["end_page"] = int(args.get("end_page") or 0)
                    snippet["_raw_text"] = str(result or "")
                snippets.append(snippet)

        context_summary = (
            f"已整理 {len(snippets)} 份上下文片段。" if snippets else "没有显式文件或选区可读取。"
        )
        yield ledger.event(
            "step.finished",
            {
                "summary": context_summary,
            },
            step_id=context_step_id,
        )
        yield ledger.event(
            "step.result",
            self._build_step_result_payload(
                title="读取显式上下文",
                summary=context_summary,
                status="completed" if snippets else "needs_attention",
                snippet_count=len(snippets),
                snippets=snippets,
            ),
            step_id=context_step_id,
        )

        execute_step_id = "execute"
        yield ledger.event(
            "step.started",
            {
                "title": "模型规划并调用工具",
                "detail": "模型只能调用 Koto 文件工具目录中的 allowlist 工具。",
                "max_rounds": self._max_rounds,
            },
            step_id=execute_step_id,
        )

        messages = self._build_messages(
            request,
            snippets,
            context_files,
            known_tool_gap,
            classification,
            intent_plan,
            execution_context=execution_context,
            recipe_skeleton=recipe_skeleton,
        )
        system = self._build_system_prompt(
            request,
            context_files,
            known_tool_gap,
            classification,
            intent_plan,
            execution_context=execution_context,
            recipe_skeleton=recipe_skeleton,
        )
        model_request = self._initial_model_request(request)
        completed_write_ops: Dict[str, int] = {}
        file_changes: List[Dict[str, Any]] = []
        final_summary = ""
        completed_task = False
        model_failed = False
        readonly_fallback_used = False
        last_tool_batch_signature = ""
        planner_runtime_payload: Dict[str, Any] = {}
        planner_fallback_runtime_payload: Dict[str, Any] = {}
        last_tool_gap_signature = ""
        plan_confirmed_emitted = False
        last_execution_brief_signature = ""
        write_guard_injected = False
        readonly_answer_guard_injected = False
        duplicate_supervisor_guard_injected = False
        readonly_tool_outputs: List[Dict[str, Any]] = []
        repair_attempts = 0
        last_check_payload: Optional[Dict[str, Any]] = None
        tool_gap: Optional[Dict[str, Any]] = None
        next_action_artifact: Optional[Dict[str, Any]] = None
        tool_runtime_outcome: Optional[Dict[str, Any]] = None
        active_execution_plan: Optional[WhiteboxExecutionPlan] = None
        last_execution_plan_signature = ""

        for round_index in range(1, self._max_rounds + 1):
            if self._is_cancelled(request):
                yield self._cancelled_event(ledger, request)
                return
            planner_fallback_runtime_payload = {}
            try:
                response = self._call_model(
                    request=model_request,
                    messages=messages,
                    system=system,
                    tools=tool_defs,
                )
            except Exception as exc:
                logger.warning("[FileTaskRuntime] model call failed: %s", exc)
                deterministic_change = yield from self._write_stepwise_pdf_docx_native(
                    ledger,
                    request,
                    executor,
                    snippets,
                    context_files,
                    recipe_skeleton,
                    execute_step_id,
                    reason="model_unavailable",
                    fallback=True,
                    model_unavailable=True,
                )
                if deterministic_change:
                    file_changes.append(deterministic_change)
                    completed_task = True
                    model_failed = True
                    final_summary = str(
                        deterministic_change.get("summary")
                        or "模型不可用，已使用 Koto 原生流程写入当前分步结果。"
                    )
                    yield ledger.event(
                        "file.changed", deterministic_change, step_id=execute_step_id
                    )
                    yield ledger.event(
                        "step.result",
                        self._build_step_result_payload(
                            title="模型不可用兜底写入",
                            summary=final_summary,
                            status="completed",
                            round_index=round_index,
                            file_changes=file_changes,
                        ),
                        step_id=execute_step_id,
                    )
                    break
                fallback_summary = (
                    ""
                    if write_intent
                    else self._fallback_readonly_summary(
                        request,
                        snippets,
                        context_files,
                        exc,
                    )
                )
                if fallback_summary:
                    readonly_fallback_used = True
                    completed_task = True
                    final_summary = fallback_summary
                    yield ledger.event(
                        "tool.finished",
                        {
                            "tool_name": "model_message",
                            "success": True,
                            "fallback": True,
                            "model_unavailable": True,
                            "result_preview": fallback_summary,
                        },
                        step_id=execute_step_id,
                    )
                    yield ledger.event(
                        "step.result",
                        self._build_step_result_payload(
                            title="模型规划并调用工具",
                            summary=fallback_summary,
                            status="completed",
                            round_index=round_index,
                        ),
                        step_id=execute_step_id,
                    )
                else:
                    model_failed = True
                    error_text = f"模型调用失败：{exc}"
                    yield ledger.event(
                        "run.error",
                        {
                            "text": error_text,
                            "recoverable": not write_intent,
                        },
                        step_id=execute_step_id,
                    )
                    yield ledger.event(
                        "step.result",
                        self._build_step_result_payload(
                            title="模型规划并调用工具",
                            summary=error_text,
                            status="failed",
                            round_index=round_index,
                        ),
                        step_id=execute_step_id,
                    )
                break

            if self._is_cancelled(request):
                yield self._cancelled_event(ledger, request)
                return

            planner_runtime_payload = {
                "backend": execution_context.effective_planner_backend or "native",
                "source": "native",
                "policy": execution_context.effective_planner_policy or "native_only",
                "transport": "native",
                "reason": execution_context.effective_planner_reason
                or "file_task_native_only",
                "round": round_index,
            }
            planner_meta = dict(planner_runtime_payload)

            tool_gap = extract_tool_gap_from_response(response)
            if tool_gap and known_tool_gap:
                tool_gap = merge_tool_gaps(tool_gap, known_tool_gap)
            content_text, tool_calls = self._normalize_model_response(
                response, tool_defs
            )
            execution_brief, content_text = self._extract_execution_brief(
                response, content_text
            )
            tool_execution_brief, tool_calls = self._extract_execution_brief_tool_call(
                tool_calls
            )
            if tool_execution_brief and not execution_brief:
                execution_brief = tool_execution_brief
            execution_plan = extract_whitebox_execution_plan(response, content_text)
            if execution_plan:
                plan_payload = execution_plan.public_dict()
                plan_signature = json.dumps(
                    plan_payload, ensure_ascii=False, sort_keys=True, default=str
                )
                if plan_signature != last_execution_plan_signature:
                    active_execution_plan = execution_plan
                    last_execution_plan_signature = plan_signature
                    yield ledger.event(
                        "plan.proposed", plan_payload, step_id=execute_step_id
                    )
                    gate_payload = validate_whitebox_plan(
                        execution_plan, recipe_skeleton
                    )
                    yield ledger.event(
                        "plan.gated", gate_payload, step_id=execute_step_id
                    )
                    if not gate_payload.get("passed"):
                        if round_index < self._max_rounds:
                            repair_message = self._whitebox_plan_repair_message(
                                gate_payload, recipe_skeleton
                            )
                            yield ledger.event(
                                "tool.finished",
                                {
                                    "tool_name": "plan_gate",
                                    "success": False,
                                    "result_preview": repair_message,
                                },
                                step_id=execute_step_id,
                            )
                            messages.append({"role": "user", "content": repair_message})
                            continue
                        final_summary = str(gate_payload.get("summary") or "白盒计划审查未通过。")
                        completed_task = False
                        yield ledger.event(
                            "step.result",
                            self._build_step_result_payload(
                                title="白盒计划审查",
                                summary=final_summary,
                                status="failed",
                                round_index=round_index,
                            ),
                            step_id=execute_step_id,
                        )
                        break
            external_planner_request = False
            if (
                not tool_gap
                and known_tool_gap
                and not tool_calls
                and not external_planner_request
                and not bool(
                    (model_request.options or {}).get(
                        "planner_runtime_fallback_attempted"
                    )
                )
            ):
                tool_gap = known_tool_gap

            if execution_brief:
                brief_payload = execution_brief.public_dict()
                brief_signature = json.dumps(
                    brief_payload, ensure_ascii=False, sort_keys=True, default=str
                )
                if brief_signature != last_execution_brief_signature:
                    last_execution_brief_signature = brief_signature
                    yield ledger.event(
                        "plan.briefed", brief_payload, step_id=execute_step_id
                    )

            if tool_gap:
                gap_runtime = self._build_runtime_metadata(
                    terminal_status="tool_gap",
                    readonly_fallback_used=readonly_fallback_used,
                    model_failed=model_failed,
                    planner_payload=planner_runtime_payload,
                    planner_fallback_payload=planner_fallback_runtime_payload,
                )
                next_action_artifact = self._with_runtime_context(
                    build_next_action_artifact(request, tool_gap),
                    gap_runtime,
                )
                gap_payload = {
                    "summary": str(tool_gap.get("summary") or ""),
                    "missing_capability": str(tool_gap.get("missing_capability") or ""),
                    "why_missing": str(tool_gap.get("why_missing") or ""),
                    "suggested_next_step": str(
                        tool_gap.get("suggested_next_step") or ""
                    ),
                    "proposed_tool": (
                        tool_gap.get("proposed_tool")
                        if isinstance(tool_gap.get("proposed_tool"), dict)
                        else None
                    ),
                    "next_action_artifact": next_action_artifact,
                    "runtime": gap_runtime,
                    "round": round_index,
                }
                gap_signature = json.dumps(
                    gap_payload, ensure_ascii=False, sort_keys=True, default=str
                )
                if gap_signature != last_tool_gap_signature:
                    last_tool_gap_signature = gap_signature
                    yield ledger.event(
                        "tool.missing", gap_payload, step_id=execute_step_id
                    )

            if tool_calls and not plan_confirmed_emitted:
                tool_gate_payload = validate_whitebox_plan(
                    active_execution_plan,
                    recipe_skeleton,
                    tool_calls=tool_calls,
                )
                yield ledger.event(
                    "plan.gated", tool_gate_payload, step_id=execute_step_id
                )
                if not tool_gate_payload.get("passed"):
                    if round_index < self._max_rounds:
                        repair_message = self._whitebox_plan_repair_message(
                            tool_gate_payload, recipe_skeleton
                        )
                        yield ledger.event(
                            "supervisor.intervention",
                            {
                                "reason": "plan_gate_failed",
                                "summary": repair_message,
                                "gate": tool_gate_payload,
                            },
                            step_id=execute_step_id,
                        )
                        messages.append({"role": "user", "content": repair_message})
                        continue
                    final_summary = str(
                        tool_gate_payload.get("summary") or "工具计划未通过白盒审查。"
                    )
                    completed_task = False
                    yield ledger.event(
                        "step.result",
                        self._build_step_result_payload(
                            title="白盒计划审查",
                            summary=final_summary,
                            status="failed",
                            round_index=round_index,
                        ),
                        step_id=execute_step_id,
                    )
                    break
                plan_confirmed_emitted = True
                yield ledger.event(
                    "plan.confirmed",
                    self._build_confirmed_plan(
                        request,
                        context_files,
                        tool_calls,
                        write_intent,
                        content_text,
                    ),
                    step_id=execute_step_id,
                )

            if content_text and (not tool_calls or len(content_text) <= 220):
                yield ledger.event(
                    "tool.finished",
                    {
                        "tool_name": "model_message",
                        "success": True,
                        "result_preview": _preview(content_text, 600),
                    },
                    step_id=execute_step_id,
                )

            model_turn: Dict[str, Any] = {
                "role": "model",
                "content": content_text or "",
            }
            if isinstance(response, dict) and response.get("reasoning_content"):
                model_turn["reasoning_content"] = str(
                    response.get("reasoning_content") or ""
                )
            if tool_calls:
                for tool_call in tool_calls:
                    tool_call.setdefault("id", uuid.uuid4().hex[:8])
                model_turn["tool_calls"] = tool_calls
            raw_parts = (
                response.get("_raw_parts") if isinstance(response, dict) else None
            )
            if raw_parts:
                model_turn["parts"] = raw_parts
            if tool_gap:
                model_turn["tool_gap"] = tool_gap
            messages.append(model_turn)

            if not tool_calls:
                if tool_gap:
                    final_summary = content_text or str(
                        tool_gap.get("summary") or "当前任务缺少对应的 Koto 原生工具。"
                    )
                    completed_task = False
                    yield ledger.event(
                        "step.result",
                        self._build_step_result_payload(
                            title="模型规划并调用工具",
                            summary=final_summary,
                            status="failed",
                            round_index=round_index,
                            file_changes=file_changes,
                            next_action_artifact=next_action_artifact,
                        ),
                        step_id=execute_step_id,
                    )
                    break
                if execution_brief and round_index < self._max_rounds:
                    model_request = self._request_after_execution_brief(
                        request, model_request, execution_brief
                    )
                    reminder = self._execution_brief_continue_message(
                        request, execution_brief
                    )
                    final_summary = (
                        execution_brief.summary or content_text or "已完成任务分析，准备继续执行。"
                    )
                    yield ledger.event(
                        "step.result",
                        self._build_step_result_payload(
                            title="模型规划并调用工具",
                            summary=final_summary,
                            status="pending",
                            round_index=round_index,
                            file_changes=file_changes,
                        ),
                        step_id=execute_step_id,
                    )
                    messages.append({"role": "user", "content": reminder})
                    continue
                if execution_plan and round_index < self._max_rounds:
                    reminder = self._execution_plan_continue_message(
                        request, execution_plan, recipe_skeleton
                    )
                    final_summary = (
                        execution_plan.plan_summary
                        or execution_plan.goal
                        or content_text
                        or "已完成白盒计划，准备继续执行。"
                    )
                    yield ledger.event(
                        "step.result",
                        self._build_step_result_payload(
                            title="白盒执行计划",
                            summary=final_summary,
                            status="pending",
                            round_index=round_index,
                            file_changes=file_changes,
                        ),
                        step_id=execute_step_id,
                    )
                    messages.append({"role": "user", "content": reminder})
                    continue
                runtime_status = self._tool_runtime_status(tool_runtime_outcome)
                awaiting_confirmation = runtime_status == "awaiting_confirmation"
                terminal_write_blocked = runtime_status in {"blocked", "write_blocked"}
                if (
                    write_intent
                    and not file_changes
                    and not awaiting_confirmation
                    and not terminal_write_blocked
                    and not write_guard_injected
                    and round_index < self._max_rounds
                ):
                    write_guard_injected = True
                    reminder = self._write_retry_message(request, context_files)
                    yield ledger.event(
                        "tool.finished",
                        {
                            "tool_name": "write_guard",
                            "success": False,
                            "result_preview": reminder,
                        },
                        step_id=execute_step_id,
                    )
                    messages.append({"role": "user", "content": reminder})
                    final_summary = content_text or "模型未再请求工具调用。"
                    yield ledger.event(
                        "step.result",
                        self._build_step_result_payload(
                            title="模型规划并调用工具",
                            summary=reminder,
                            status="needs_attention",
                            round_index=round_index,
                        ),
                        step_id=execute_step_id,
                    )
                    continue
                if write_intent:
                    last_check_payload = self._verify_task(
                        request,
                        executor,
                        file_changes,
                        write_intent,
                        classification.output_mode,
                        model_failed,
                        readonly_fallback_used=readonly_fallback_used,
                        tool_runtime_outcome=tool_runtime_outcome,
                        tool_gap=tool_gap,
                        next_action_artifact=next_action_artifact,
                    )
                    if self._should_attempt_repair(
                        last_check_payload,
                        round_index=round_index,
                        repair_attempts=repair_attempts,
                    ):
                        repair_attempts += 1
                        repair_runtime = self._build_runtime_metadata(
                            terminal_status=str(
                                last_check_payload.get("status") or ""
                            ).strip(),
                            readonly_fallback_used=readonly_fallback_used,
                            model_failed=model_failed,
                            planner_payload=planner_runtime_payload,
                            planner_fallback_payload=planner_fallback_runtime_payload,
                        )
                        repair_check_payload = dict(last_check_payload)
                        repair_check_payload["runtime"] = repair_runtime
                        repair_check_payload["repair_attempt"] = repair_attempts
                        yield ledger.event(
                            "check.started",
                            {
                                "title": "检查执行状态",
                                "criteria": self._success_criteria(
                                    request, write_intent, classification.output_mode
                                ),
                                "repair_attempt": repair_attempts,
                            },
                            step_id="check",
                        )
                        yield ledger.event(
                            "check.finished", repair_check_payload, step_id="check"
                        )
                        yield ledger.event(
                            "step.result",
                            self._build_step_result_payload(
                                title="检查执行状态",
                                summary=str(
                                    repair_check_payload.get("summary") or "检查未通过。"
                                ),
                                status=(
                                    "completed"
                                    if repair_check_payload.get("passed")
                                    else "needs_attention"
                                ),
                                runtime=repair_runtime,
                                passed=repair_check_payload.get("passed"),
                                file_changes=file_changes,
                                next_action_artifact=repair_check_payload.get(
                                    "next_action_artifact"
                                ),
                            ),
                            step_id="check",
                        )
                        repair_message = self._repair_retry_message(
                            request, last_check_payload, file_changes
                        )
                        yield ledger.event(
                            "tool.finished",
                            {
                                "tool_name": "repair_guard",
                                "success": False,
                                "result_preview": repair_message,
                            },
                            step_id=execute_step_id,
                        )
                        messages.append({"role": "user", "content": repair_message})
                        completed_write_ops.clear()
                        last_check_payload = None
                        final_summary = (
                            repair_check_payload.get("summary")
                            or content_text
                            or "核验未通过，准备修复。"
                        )
                        continue
                    check_status = (
                        str(last_check_payload.get("status") or "").strip().lower()
                    )
                    final_summary = content_text or str(
                        last_check_payload.get("summary") or "模型未再请求工具调用。"
                    )
                    completed_task = bool(last_check_payload.get("passed"))
                    yield ledger.event(
                        "step.result",
                        self._build_step_result_payload(
                            title="模型规划并调用工具",
                            summary=self._execute_step_summary(
                                round_index=round_index,
                                final_summary=final_summary,
                                model_failed=model_failed,
                                tool_gap=tool_gap,
                                file_changes=file_changes,
                                tool_runtime_outcome=tool_runtime_outcome,
                            ),
                            status=self._execute_step_result_status(
                                completed=completed_task,
                                tool_gap=tool_gap,
                                tool_runtime_outcome=tool_runtime_outcome,
                                model_failed=model_failed,
                            ),
                            round_index=round_index,
                            file_changes=file_changes,
                            next_action_artifact=next_action_artifact,
                        ),
                        step_id=execute_step_id,
                    )
                    break
                if (
                    not content_text
                    and (snippets or readonly_tool_outputs)
                    and not readonly_answer_guard_injected
                    and round_index < self._max_rounds
                ):
                    readonly_answer_guard_injected = True
                    reminder = self._readonly_answer_required_message(
                        request, snippets, readonly_tool_outputs
                    )
                    yield ledger.event(
                        "tool.finished",
                        {
                            "tool_name": "readonly_answer_guard",
                            "success": False,
                            "result_preview": reminder,
                        },
                        step_id=execute_step_id,
                    )
                    messages.append({"role": "user", "content": reminder})
                    final_summary = "已读取内容，正在生成可见分析结果。"
                    yield ledger.event(
                        "step.result",
                        self._build_step_result_payload(
                            title="模型规划并调用工具",
                            summary=final_summary,
                            status="pending",
                            round_index=round_index,
                        ),
                        step_id=execute_step_id,
                    )
                    continue
                final_summary = (
                    content_text
                    or self._readonly_context_summary(
                        request, snippets, readonly_tool_outputs
                    )
                    or "已读取上下文，但模型未生成可见分析结果。"
                )
                completed_task = not write_intent or bool(file_changes)
                yield ledger.event(
                    "step.result",
                    self._build_step_result_payload(
                        title="模型规划并调用工具",
                        summary=self._execute_step_summary(
                            round_index=round_index,
                            final_summary=final_summary,
                            model_failed=model_failed,
                            tool_gap=tool_gap,
                            file_changes=file_changes,
                            tool_runtime_outcome=tool_runtime_outcome,
                        ),
                        status=self._execute_step_result_status(
                            completed=completed_task,
                            tool_gap=tool_gap,
                            tool_runtime_outcome=tool_runtime_outcome,
                            model_failed=model_failed,
                        ),
                        round_index=round_index,
                        file_changes=file_changes,
                        next_action_artifact=next_action_artifact,
                    ),
                    step_id=execute_step_id,
                )
                break

            batch_signature = self._tool_batch_signature(tool_calls)
            if batch_signature and batch_signature == last_tool_batch_signature:
                if (
                    write_intent
                    and not file_changes
                    and not duplicate_supervisor_guard_injected
                    and round_index < self._max_rounds
                ):
                    duplicate_supervisor_guard_injected = True
                    final_summary = "检测到重复读取/重复工具调用，监管层已要求模型回到计划主线继续执行。"
                    reminder = self._duplicate_supervisor_retry_message(
                        request,
                        context_files,
                        classification,
                        intent_plan,
                        tool_calls,
                    )
                    yield ledger.event(
                        "tool.finished",
                        {
                            "tool_name": "supervisor_guard",
                            "success": False,
                            "skipped": True,
                            "result_preview": reminder,
                        },
                        step_id=execute_step_id,
                    )
                    yield ledger.event(
                        "step.result",
                        self._build_step_result_payload(
                            title="监管纠偏",
                            summary=final_summary,
                            status="needs_attention",
                            round_index=round_index,
                            file_changes=file_changes,
                        ),
                        step_id=execute_step_id,
                    )
                    messages.append({"role": "user", "content": reminder})
                    continue
                final_summary = "检测到重复工具调用，已自动停止以避免重复写入。"
                yield ledger.event(
                    "tool.finished",
                    {
                        "tool_name": "duplicate_guard",
                        "success": True,
                        "skipped": True,
                        "result_preview": final_summary,
                    },
                    step_id=execute_step_id,
                )
                yield ledger.event(
                    "step.result",
                    self._build_step_result_payload(
                        title="模型规划并调用工具",
                        summary=final_summary,
                        status="needs_attention",
                        round_index=round_index,
                        file_changes=file_changes,
                    ),
                    step_id=execute_step_id,
                )
                break
            last_tool_batch_signature = batch_signature

            for tool_index, tool_call in enumerate(tool_calls, start=1):
                if self._is_cancelled(request):
                    yield self._cancelled_event(ledger, request)
                    return
                tool_name = str(tool_call.get("name") or "").strip()
                tool_args = dict(tool_call.get("args") or {})
                tool_args = self._repair_tool_args_for_context(
                    tool_name, tool_args, request, context_files
                )
                tool_call_id = str(tool_call.get("id") or uuid.uuid4().hex[:8])
                current_step_id = f"tool_{round_index}_{tool_index}"
                yield ledger.event(
                    "decision.made",
                    build_decision_audit(
                        request=request,
                        skeleton=recipe_skeleton,
                        tool_name=tool_name,
                        tool_args=tool_args,
                        round_index=round_index,
                        tool_index=tool_index,
                        execution_plan=active_execution_plan,
                    ),
                    step_id=current_step_id,
                )

                if not is_file_task_tool(tool_name):
                    error_text = (
                        f"工具 {tool_name or '<empty>'} 不在 Koto 文件任务 allowlist 中。"
                    )
                    yield ledger.event(
                        "tool.finished",
                        {
                            "tool_name": tool_name,
                            "success": False,
                            "result_preview": error_text,
                        },
                        step_id=current_step_id,
                    )
                    messages.append(
                        {
                            "role": "function",
                            "name": tool_name or "invalid_tool",
                            "tool_call_id": tool_call_id,
                            "content": self._tool_feedback_for_model(
                                tool_name or "invalid_tool",
                                tool_args,
                                {"error": error_text},
                                success=False,
                                invalid=True,
                            ),
                        }
                    )
                    continue

                exposed_tool_names = {
                    str(definition.get("name") or "").strip()
                    for definition in tool_defs
                    if str(definition.get("name") or "").strip()
                }
                if exposed_tool_names and tool_name not in exposed_tool_names:
                    error_text = self._recipe_tool_block_message(
                        tool_name,
                        classification,
                        exposed_tool_names,
                    )
                    yield ledger.event(
                        "tool.finished",
                        {
                            "tool_name": tool_name,
                            "success": False,
                            "blocked": True,
                            "result_preview": error_text,
                        },
                        step_id=current_step_id,
                    )
                    messages.append(
                        {
                            "role": "function",
                            "name": tool_name,
                            "tool_call_id": tool_call_id,
                            "content": self._tool_feedback_for_model(
                                tool_name,
                                tool_args,
                                {"error": error_text},
                                success=False,
                                blocked=True,
                            ),
                        }
                    )
                    continue

                if (
                    is_write_tool(tool_name)
                    and tool_name != "run_python_code"
                    and (not write_intent or classification.output_mode != "write")
                ):
                    block_text = self._readonly_write_tool_block_message(
                        tool_name,
                        request,
                        classification.output_mode,
                    )
                    yield ledger.event(
                        "tool.finished",
                        {
                            "tool_name": tool_name,
                            "success": False,
                            "blocked": True,
                            "result_preview": block_text,
                        },
                        step_id=current_step_id,
                    )
                    messages.append(
                        {
                            "role": "function",
                            "name": tool_name,
                            "tool_call_id": tool_call_id,
                            "content": self._tool_feedback_for_model(
                                tool_name,
                                tool_args,
                                {"error": block_text},
                                success=False,
                                blocked=True,
                            ),
                        }
                    )
                    continue

                if tool_name == "run_python_code" and (
                    not write_intent or classification.output_mode != "write"
                ):
                    block_text = self._readonly_run_python_write_block_message(
                        tool_args,
                        request,
                        classification.output_mode,
                    )
                    if block_text:
                        yield ledger.event(
                            "tool.finished",
                            {
                                "tool_name": tool_name,
                                "success": False,
                                "blocked": True,
                                "result_preview": block_text,
                            },
                            step_id=current_step_id,
                        )
                        messages.append(
                            {
                                "role": "function",
                                "name": tool_name,
                                "tool_call_id": tool_call_id,
                                "content": self._tool_feedback_for_model(
                                    tool_name,
                                    tool_args,
                                    {"error": block_text},
                                    success=False,
                                    blocked=True,
                                ),
                            }
                        )
                        continue

                if is_write_tool(tool_name) and tool_name != "run_python_code":
                    target = write_target_for_tool(tool_name, tool_args)
                    write_key = f"{tool_name}::{target}"
                    if completed_write_ops.get(write_key, 0) >= _MAX_WRITE_OPS_PER_FILE:
                        skip_text = (
                            f"{tool_name} 已成功写入过 {target or '同一目标'}，本次跳过以避免重复覆盖。"
                        )
                        yield ledger.event(
                            "tool.finished",
                            {
                                "tool_name": tool_name,
                                "success": True,
                                "skipped": True,
                                "result_preview": skip_text,
                            },
                            step_id=current_step_id,
                        )
                        messages.append(
                            {
                                "role": "function",
                                "name": tool_name,
                                "tool_call_id": tool_call_id,
                                "content": self._tool_feedback_for_model(
                                    tool_name,
                                    tool_args,
                                    {"summary": skip_text},
                                    success=True,
                                    skipped=True,
                                ),
                            }
                        )
                        continue

                stepwise_write_block = self._stepwise_docx_write_block_message(
                    request,
                    snippets,
                    recipe_skeleton,
                    tool_name,
                    tool_args,
                )
                if stepwise_write_block:
                    yield ledger.event(
                        "tool.finished",
                        {
                            "tool_name": "supervisor_guard",
                            "success": False,
                            "blocked": True,
                            "result_preview": stepwise_write_block,
                        },
                        step_id=current_step_id,
                    )
                    messages.append(
                        {
                            "role": "function",
                            "name": tool_name,
                            "tool_call_id": tool_call_id,
                            "content": self._tool_feedback_for_model(
                                tool_name,
                                tool_args,
                                {"error": stepwise_write_block},
                                success=False,
                                blocked=True,
                            ),
                        }
                    )
                    continue

                yield ledger.event(
                    "tool.started",
                    {
                        "tool_name": tool_name,
                        "tool_args": tool_args,
                        "round": round_index,
                    },
                    step_id=current_step_id,
                )

                blocked_message = self._blocked_run_python_message(
                    tool_name, tool_args, request, context_files
                )
                if blocked_message:
                    yield ledger.event(
                        "tool.finished",
                        {
                            "tool_name": tool_name,
                            "success": False,
                            "blocked": True,
                            "result_preview": blocked_message,
                        },
                        step_id=current_step_id,
                    )
                    messages.append(
                        {
                            "role": "function",
                            "name": tool_name,
                            "tool_call_id": tool_call_id,
                            "content": self._tool_feedback_for_model(
                                tool_name,
                                tool_args,
                                {"error": blocked_message},
                                success=False,
                                blocked=True,
                            ),
                        }
                    )
                    continue

                if tool_name == "run_python_code":
                    yield ledger.event(
                        "code.started",
                        {
                            "code": str(tool_args.get("code") or ""),
                        },
                        step_id=current_step_id,
                    )

                try:
                    result = executor(tool_name, tool_args)
                    if isinstance(result, FileTaskToolStreamResult):
                        result = yield from self._consume_streaming_tool_result(
                            ledger,
                            step_id=current_step_id,
                            stream_result=result,
                        )
                    success = not _is_error_result(result)
                except Exception as exc:
                    result = f"Error: {exc}"
                    success = False
                    logger.warning(
                        "[FileTaskRuntime] tool %s failed: %s", tool_name, exc
                    )

                if self._is_cancelled(request):
                    yield self._cancelled_event(ledger, request)
                    return

                model_result = self._tool_result_for_model(tool_name, result)
                current_tool_runtime_outcome = self._extract_tool_runtime_outcome(
                    result
                )
                if current_tool_runtime_outcome:
                    tool_runtime_outcome = current_tool_runtime_outcome
                    artifact = current_tool_runtime_outcome.get("next_action_artifact")
                    if isinstance(artifact, dict):
                        next_action_artifact = artifact
                runtime_status = self._tool_runtime_status(current_tool_runtime_outcome)
                runtime_blocked = runtime_status in {"blocked", "write_blocked"}
                result_text = stringify_result(model_result)
                if success and not is_write_tool(tool_name):
                    readonly_tool_outputs.append(
                        {
                            "tool_name": tool_name,
                            "args": dict(tool_args),
                            "result": model_result,
                            "preview": tool_result_preview(
                                tool_name, model_result, 1200
                            ),
                        }
                    )
                artifacts = self._tool_artifacts(tool_name, result)
                if tool_name == "run_python_code":
                    yield ledger.event(
                        "code.output",
                        {
                            "text": self._code_output_preview(
                                tool_name, result, result_text
                            ),
                            "stream": "stdout" if success else "stderr",
                        },
                        step_id=current_step_id,
                    )
                    yield ledger.event(
                        "code.finished",
                        {
                            "success": success,
                        },
                        step_id=current_step_id,
                    )

                tool_finished_payload = {
                    "tool_name": tool_name,
                    "success": success,
                    "result_preview": tool_result_preview(
                        tool_name, model_result, 1200
                    ),
                }
                if runtime_blocked:
                    tool_finished_payload["blocked"] = True
                if artifacts:
                    tool_finished_payload["artifacts"] = artifacts
                yield ledger.event(
                    "tool.finished", tool_finished_payload, step_id=current_step_id
                )

                messages.append(
                    {
                        "role": "function",
                        "name": tool_name,
                        "tool_call_id": tool_call_id,
                        "content": self._tool_feedback_for_model(
                            tool_name,
                            tool_args,
                            model_result,
                            success=success,
                            blocked=runtime_blocked,
                        ),
                    }
                )

                extracted_changes = self._extract_file_changes(
                    tool_name, tool_args, result
                )
                if (
                    success
                    and is_write_tool(tool_name)
                    and tool_name != "run_python_code"
                ):
                    target = write_target_for_tool(tool_name, tool_args)
                    write_key = f"{tool_name}::{target}"
                    completed_write_ops[write_key] = (
                        completed_write_ops.get(write_key, 0) + 1
                    )

                if extracted_changes:
                    repair_attempts = 0
                for change in extracted_changes:
                    file_changes.append(change)
                    yield ledger.event("file.changed", change, step_id=current_step_id)

            execute_round_summary = self._execute_step_summary(
                round_index=round_index,
                final_summary=final_summary,
                model_failed=model_failed,
                tool_gap=tool_gap,
                file_changes=file_changes,
                tool_runtime_outcome=tool_runtime_outcome,
            )
            yield ledger.event(
                "step.finished",
                {
                    "title": "模型工具执行完成",
                    "summary": execute_round_summary,
                },
                step_id=execute_step_id,
            )
            yield ledger.event(
                "step.result",
                self._build_step_result_payload(
                    title="模型工具执行完成",
                    summary=execute_round_summary,
                    status=self._execute_step_result_status(
                        completed=not model_failed and not tool_gap,
                        tool_gap=tool_gap,
                        tool_runtime_outcome=tool_runtime_outcome,
                        model_failed=model_failed,
                    ),
                    round_index=round_index,
                    file_changes=file_changes,
                    next_action_artifact=next_action_artifact,
                ),
                step_id=execute_step_id,
            )
            if self._tool_runtime_status(tool_runtime_outcome) in {
                "blocked",
                "write_blocked",
            }:
                final_summary = execute_round_summary
                completed_task = False
                break
            if (
                write_intent
                and not file_changes
                and not write_guard_injected
                and round_index < self._max_rounds
                and self._should_prompt_for_write_after_tool_round(
                    request, context_files, tool_calls, round_index
                )
            ):
                write_guard_injected = True
                reminder = self._write_retry_message(request, context_files)
                yield ledger.event(
                    "tool.finished",
                    {
                        "tool_name": "write_guard",
                        "success": False,
                        "result_preview": reminder,
                    },
                    step_id=execute_step_id,
                )
                messages.append({"role": "user", "content": reminder})

        if write_intent and not file_changes:
            deterministic_change = yield from self._write_stepwise_pdf_docx_native(
                ledger,
                request,
                executor,
                snippets,
                context_files,
                recipe_skeleton,
                execute_step_id,
                reason="model_finished_without_write",
                fallback=True,
                model_unavailable=False,
            )
            if deterministic_change:
                file_changes.append(deterministic_change)
                completed_task = True
                last_check_payload = None
                final_summary = str(
                    deterministic_change.get("summary")
                    or "模型未完成写入，已使用 Koto 原生分步流程写入当前结果。"
                )
                yield ledger.event(
                    "file.changed", deterministic_change, step_id=execute_step_id
                )
                yield ledger.event(
                    "step.result",
                    self._build_step_result_payload(
                        title="原生分步兜底写入",
                        summary=final_summary,
                        status="completed",
                        file_changes=file_changes,
                    ),
                    step_id=execute_step_id,
                )

        if not write_intent and not str(final_summary or "").strip():
            final_summary = self._readonly_context_summary(
                request, snippets, readonly_tool_outputs
            )
            if final_summary:
                completed_task = True

        check_step_id = "check"
        if self._is_cancelled(request):
            yield self._cancelled_event(ledger, request)
            return
        yield ledger.event(
            "check.started",
            {
                "title": "检查执行状态",
                "criteria": self._success_criteria(
                    request, write_intent, classification.output_mode
                ),
            },
            step_id=check_step_id,
        )

        check_payload = (
            dict(last_check_payload)
            if isinstance(last_check_payload, dict)
            else self._verify_task(
                request,
                executor,
                file_changes,
                write_intent,
                classification.output_mode,
                model_failed,
                readonly_fallback_used,
                tool_runtime_outcome,
                tool_gap,
                next_action_artifact,
            )
        )
        stepwise_artifact = self._stepwise_docx_wait_artifact(
            request,
            context_files,
            snippets,
            file_changes,
            recipe_skeleton,
        )
        if stepwise_artifact and bool(check_payload.get("passed")):
            next_action_artifact = stepwise_artifact
            check_payload = dict(check_payload)
            check_payload["passed"] = False
            check_payload["status"] = "awaiting_confirmation"
            check_payload["summary"] = "当前步骤已写入 DOCX，等待用户说“继续”后处理下一段。"
            check_payload["remaining"] = ["用户说“继续”后处理下一页窗口，并继续追加 DOCX。"]
            check_payload["next_action_artifact"] = stepwise_artifact
        terminal_runtime = self._build_runtime_metadata(
            terminal_status=str(check_payload.get("status") or "").strip(),
            readonly_fallback_used=readonly_fallback_used,
            model_failed=model_failed,
            planner_payload=planner_runtime_payload,
            planner_fallback_payload=planner_fallback_runtime_payload,
        )
        check_payload["runtime"] = terminal_runtime

        yield ledger.event("check.finished", check_payload, step_id=check_step_id)
        yield ledger.event(
            "step.result",
            self._build_step_result_payload(
                title="检查执行状态",
                summary=str(check_payload.get("summary") or "检查完成。"),
                status=self._check_step_result_status(check_payload),
                runtime=terminal_runtime,
                passed=check_payload.get("passed"),
                file_changes=file_changes,
                next_action_artifact=check_payload.get("next_action_artifact")
                or next_action_artifact,
            ),
            step_id=check_step_id,
        )
        run_summary = check_payload.get("summary") or final_summary or "任务执行结束。"
        if not write_intent and final_summary and not tool_gap:
            run_summary = final_summary
        if classification_payload.get("selected_recipe") == "docx_contract_compare_review":
            contract_risks = None
            for change in file_changes:
                risks = change.get("contract_risk_summary") if isinstance(change, dict) else None
                if isinstance(risks, list) and risks:
                    contract_risks = risks
                    break
            if (
                isinstance(contract_risks, list)
                and contract_risks
                and "风险关注点" not in str(run_summary)
            ):
                risk_lines = "\n".join(f"- {item}" for item in contract_risks[:5])
                run_summary = f"{run_summary}\n风险关注点：\n{risk_lines}"
        run_payload = {
            "task": request.task,
            "mode": "whitebox_v1",
            "summary": run_summary,
            "completed_task": bool(check_payload.get("passed"))
            and (completed_task or not write_intent or bool(file_changes)),
            "context": self._public_context_snippets(snippets[:8]),
            "file_changes": file_changes,
            "runtime": terminal_runtime,
            "quick_action_mode": quick_action_mode,
            "workflow_version": recipe_skeleton.get("version"),
            "recipe_skeleton": recipe_skeleton,
            **classification_payload,
        }
        if tool_gap:
            run_payload["tool_gap"] = tool_gap
        if next_action_artifact:
            run_payload["next_action_artifact"] = next_action_artifact
        yield ledger.event("run.finished", run_payload)

    def _is_cancelled(self, request: FileTaskRequest) -> bool:
        return is_cancel_requested(str(request.run_id or ""))

    def _cancelled_event(
        self, ledger: FileTaskLedger, request: FileTaskRequest
    ) -> FileTaskEvent:
        return ledger.event(
            "run.cancelled",
            {
                "task": request.task,
                "mode": "whitebox_v1",
                "summary": "任务已被用户取消。",
                "completed_task": False,
                "runtime": {
                    "terminal_status": "cancelled",
                    "execution_path": "cancelled",
                    "model_failed": False,
                    "readonly_fallback_used": False,
                },
            },
            step_id="run",
        )

    def _build_runtime_metadata(
        self,
        *,
        terminal_status: str,
        readonly_fallback_used: bool,
        model_failed: bool,
        planner_payload: Optional[Dict[str, Any]] = None,
        planner_fallback_payload: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        planner_payload = planner_payload if isinstance(planner_payload, dict) else {}
        planner_fallback_payload = (
            planner_fallback_payload
            if isinstance(planner_fallback_payload, dict)
            else {}
        )

        backend = str(planner_payload.get("backend") or "")
        source = str(planner_payload.get("source") or "")
        policy = str(planner_payload.get("policy") or "")
        transport = str(planner_payload.get("transport") or "")
        reason = str(planner_payload.get("reason") or "")
        fallback_from = str(planner_fallback_payload.get("from") or "")

        execution_path = "native"
        if readonly_fallback_used:
            execution_path = "readonly_fallback"
        elif fallback_from:
            execution_path = "planner_fallback"
        elif backend and backend != "native":
            execution_path = "planner"
        elif source and source != "native":
            execution_path = "planner"

        planner_runtime = {
            "backend": backend,
            "source": source,
            "policy": policy,
            "transport": transport,
            "reason": reason,
        }
        round_index = planner_payload.get("round")
        if round_index:
            planner_runtime["round"] = round_index
        if fallback_from:
            planner_runtime["fallback_from"] = fallback_from

        return {
            "execution_path": execution_path,
            "terminal_status": str(terminal_status or ""),
            "model_unavailable": bool(model_failed or readonly_fallback_used),
            "readonly_fallback_used": bool(readonly_fallback_used),
            "planner": planner_runtime,
        }

    def _with_runtime_context(
        self,
        artifact: Optional[Dict[str, Any]],
        runtime_metadata: Dict[str, Any],
    ) -> Optional[Dict[str, Any]]:
        if not isinstance(artifact, dict) or not artifact:
            return artifact
        enriched = dict(artifact)
        enriched["runtime_context"] = dict(runtime_metadata)
        return enriched

    def _step_result_file_changes(
        self, file_changes: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        items: List[Dict[str, Any]] = []
        for change in file_changes[:8]:
            if not isinstance(change, dict):
                continue
            item: Dict[str, Any] = {}
            for key in (
                "path",
                "file_path",
                "file_type",
                "operation",
                "summary",
                "warning",
                "annotations_added",
                "rows_written",
                "columns_written",
                "requested_sheet",
                "sheet",
                "source_path",
                "slides_designed",
                "theme_name",
                "table_title",
            ):
                value = change.get(key)
                if value in (None, "", [], {}):
                    continue
                item[key] = value
            if item:
                items.append(item)
        return items

    def _build_step_result_payload(
        self,
        *,
        title: str,
        summary: str,
        status: str = "completed",
        round_index: int = 0,
        snippet_count: int = 0,
        snippets: Optional[List[Dict[str, Any]]] = None,
        file_changes: Optional[List[Dict[str, Any]]] = None,
        runtime: Optional[Dict[str, Any]] = None,
        passed: Optional[bool] = None,
        next_action_artifact: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        payload: Dict[str, Any] = {
            "title": str(title or "").strip() or "步骤结果",
            "summary": str(summary or "").strip()
            or str(title or "步骤结果").strip()
            or "步骤结果",
            "status": str(status or "completed").strip().lower() or "completed",
        }
        if round_index > 0:
            payload["round"] = int(round_index)
        if snippet_count > 0:
            payload["snippet_count"] = int(snippet_count)
        if snippets:
            payload["snippets"] = self._public_context_snippets(snippets[:4])
        safe_changes = self._step_result_file_changes(file_changes or [])
        if safe_changes:
            payload["file_change_count"] = len(file_changes or [])
            payload["file_changes"] = safe_changes
        if isinstance(runtime, dict) and runtime:
            payload["runtime"] = dict(runtime)
        if passed is not None:
            payload["passed"] = bool(passed)
        if isinstance(next_action_artifact, dict) and next_action_artifact:
            payload["next_action_artifact"] = next_action_artifact
        return payload

    def _public_context_snippets(
        self, snippets: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        public_items: List[Dict[str, Any]] = []
        for item in snippets or []:
            if not isinstance(item, dict):
                continue
            public_items.append(
                {
                    key: value
                    for key, value in item.items()
                    if not str(key).startswith("_")
                }
            )
        return public_items

    def _execute_step_summary(
        self,
        *,
        round_index: int,
        final_summary: str,
        model_failed: bool,
        tool_gap: Optional[Dict[str, Any]],
        file_changes: List[Dict[str, Any]],
        tool_runtime_outcome: Optional[Dict[str, Any]],
    ) -> str:
        summary = str(final_summary or "").strip()
        if summary:
            return summary
        runtime_status = self._tool_runtime_status(tool_runtime_outcome)
        if runtime_status == "awaiting_confirmation":
            return "已生成下一步执行方案，等待用户确认继续。"
        if runtime_status in {"blocked", "write_blocked"}:
            return str(
                (tool_runtime_outcome or {}).get("summary") or "目标文件当前不可写，已停止继续重试。"
            )
        if model_failed:
            return "模型调用失败，已停止工具执行。"
        if isinstance(tool_gap, dict) and tool_gap:
            return str(tool_gap.get("summary") or "当前任务缺少对应的 Koto 原生工具。")
        if file_changes:
            return f"已完成第 {round_index} 轮工具执行，累计记录 {len(file_changes)} 次文件变更。"
        if round_index > 0:
            return f"已完成第 {round_index} 轮工具执行。"
        return "模型未再请求工具调用。"

    def _execute_step_result_status(
        self,
        *,
        completed: bool,
        tool_gap: Optional[Dict[str, Any]],
        tool_runtime_outcome: Optional[Dict[str, Any]],
        model_failed: bool,
    ) -> str:
        runtime_status = self._tool_runtime_status(tool_runtime_outcome)
        if runtime_status == "awaiting_confirmation":
            return "needs_attention"
        if runtime_status in {"blocked", "write_blocked"}:
            return "failed"
        if model_failed or (isinstance(tool_gap, dict) and tool_gap):
            return "failed"
        return "completed" if completed else "needs_attention"

    def _check_step_result_status(self, check_payload: Dict[str, Any]) -> str:
        status = str(check_payload.get("status") or "").strip().lower()
        if status == "awaiting_confirmation":
            return "needs_attention"
        return "completed" if check_payload.get("passed") else "failed"

    def _context_files(self, request: FileTaskRequest) -> List[FileTaskFile]:
        seen: set[str] = set()
        result: List[FileTaskFile] = []
        candidates: List[Optional[FileTaskFile]] = [
            *request.files,
            request.current_file,
        ]
        options = request.options if isinstance(request.options, dict) else {}
        batch_control = (
            options.get("batch_control")
            if isinstance(options.get("batch_control"), dict)
            else {}
        )

        def _append_path_candidate(path_value: Any, *, target: bool = False) -> None:
            path_text = str(path_value or "").strip()
            if not path_text:
                return
            suffix = Path(path_text).suffix.lower().lstrip(".")
            if not suffix:
                return
            candidates.append(
                FileTaskFile(
                    path=path_text,
                    name=Path(path_text).name,
                    type=suffix,
                    target=target,
                )
            )

        if (
            str(batch_control.get("policy") or "").strip().lower()
            == "confirm_each_step"
        ):
            _append_path_candidate(batch_control.get("source_path"), target=False)
            _append_path_candidate(
                batch_control.get("target_path") or request.target_path, target=True
            )
        elif request.target_path:
            _append_path_candidate(request.target_path, target=True)

        for file_info in candidates:
            if not file_info:
                continue
            key = (
                (file_info.path or file_info.name or file_info.content[:80])
                .strip()
                .lower()
            )
            if not key:
                continue
            if key in seen:
                if file_info.target:
                    for existing in result:
                        existing_key = (
                            (existing.path or existing.name or existing.content[:80])
                            .strip()
                            .lower()
                        )
                        if existing_key == key:
                            existing.target = True
                            break
                continue
            seen.add(key)
            result.append(file_info)
        return result

    def _build_tool_gateway(
        self, request: FileTaskRequest, context_files: List[FileTaskFile]
    ) -> FileTaskToolGateway:
        if self._tool_gateway is not None:
            return self._tool_gateway
        providers = [self._tool_provider] if self._tool_provider is not None else None
        return FileTaskToolGateway(
            context=FileTaskToolContext(
                task_files=[file_info.public_dict() for file_info in context_files],
                workspace_root=self._workspace_root,
                gemini_client=self._gemini_client,
                request_context={
                    "task": request.task,
                    "target_path": request.target_path,
                    "options": (
                        dict(request.options)
                        if isinstance(request.options, dict)
                        else {}
                    ),
                    "model_mode": request.model_mode,
                    "model_id": request.model_id,
                },
            ),
            providers=providers,
            tool_executor=self._tool_executor,
        )

    def _tool_defs_for_classification(
        self,
        tool_defs: List[Dict[str, Any]],
        classification: FileTaskClassification,
    ) -> List[Dict[str, Any]]:
        selected_recipe = str(classification.selected_recipe or "").strip()
        if selected_recipe not in {
            "docx_compare_annotation",
            "docx_contract_compare_review",
        }:
            return tool_defs
        allowed = {
            "parse_file_to_text",
            "verify_task_completion",
            *[
                str(name or "").strip()
                for name in classification.matched_capabilities
                if str(name or "").strip()
            ],
        }
        forbidden = {"annotate_file"}
        return [
            definition
            for definition in tool_defs
            if str(definition.get("name") or "").strip() in allowed
            and str(definition.get("name") or "").strip() not in forbidden
        ]

    def _recipe_tool_block_message(
        self,
        tool_name: str,
        classification: FileTaskClassification,
        exposed_tool_names: set[str],
    ) -> str:
        selected_recipe = str(classification.selected_recipe or "").strip()
        allowed_text = ", ".join(sorted(exposed_tool_names)) or "当前路线工具集"
        if selected_recipe in {
            "docx_compare_annotation",
            "docx_contract_compare_review",
        }:
            return (
                f"工具 {tool_name} 不属于当前 DOCX 对比批注路线。"
                "这是两份 DOCX 的差异比较任务，不是单文档审校；"
                "请使用 plan_docx_compare_annotations 定位差异，"
                "再使用 write_docx_comments 写入目标 DOCX 原文批注。"
                f" 当前允许工具：{allowed_text}。"
            )
        return (
            f"工具 {tool_name} 不属于当前任务路线 {selected_recipe or '未命名路线'}。"
            f" 当前允许工具：{allowed_text}。"
        )

    def _stream_doc_annotate_bridge_execution(
        self,
        ledger: FileTaskLedger,
        request: FileTaskRequest,
        *,
        classification_payload: Dict[str, Any],
        intent_plan_payload: Dict[str, Any],
        requirements_payload: Dict[str, Any],
        plan_check_payload: Dict[str, Any],
        recipe_skeleton: Dict[str, Any],
        constraint_audit: Dict[str, Any],
        quick_action_mode: str,
    ) -> Iterable[FileTaskEvent]:
        from app.core.agent.file_task_doc_annotate_runner import (
            FileTaskDocAnnotateRunner,
        )

        yield from FileTaskDocAnnotateRunner(self).stream_bridge_execution(
            ledger,
            request,
            classification_payload=classification_payload,
            intent_plan_payload=intent_plan_payload,
            requirements_payload=requirements_payload,
            plan_check_payload=plan_check_payload,
            recipe_skeleton=recipe_skeleton,
            constraint_audit=constraint_audit,
            quick_action_mode=quick_action_mode,
        )

    def _stream_long_docx_stepwise_polish_writeback(
        self,
        ledger: FileTaskLedger,
        request: FileTaskRequest,
        context_files: List[FileTaskFile],
        classification: FileTaskClassification,
        intent_plan: FileTaskIntentPlan,
        requirements_payload: Dict[str, Any],
        plan_check_payload: Dict[str, Any],
        recipe_skeleton: Dict[str, Any],
        constraint_audit: Dict[str, Any],
        quick_action_mode: str,
        classification_payload: Dict[str, Any],
        intent_plan_payload: Dict[str, Any],
    ) -> Iterable[FileTaskEvent]:
        del classification, intent_plan

        target_path = self._stepwise_docx_polish_target_path(request, context_files)
        context_step_id = "context"
        execute_step_id = "execute"
        check_step_id = "check"
        file_changes: List[Dict[str, Any]] = []

        yield ledger.event(
            "step.started",
            {
                "title": "读取当前 DOCX 段落窗口",
                "detail": "按段落窗口读取 Word 当前步骤内容，不一次性润色全文。",
            },
            step_id=context_step_id,
        )

        if not target_path or not Path(target_path).exists():
            summary = "未找到可写回的 DOCX 文件，无法执行分步润色。"
            runtime = self._build_runtime_metadata(
                terminal_status="failed",
                readonly_fallback_used=False,
                model_failed=False,
            )
            yield ledger.event(
                "run.finished",
                {
                    "task": request.task,
                    "mode": "whitebox_v1",
                    "summary": summary,
                    "completed_task": False,
                    "context": [],
                    "file_changes": [],
                    "runtime": runtime,
                    "quick_action_mode": quick_action_mode,
                    "intent_plan": intent_plan_payload,
                    "requirements": requirements_payload,
                    "plan_check": plan_check_payload,
                    "recipe_skeleton": recipe_skeleton,
                    "constraint_audit": constraint_audit,
                    **classification_payload,
                },
            )
            return

        try:
            window = self._read_docx_paragraph_window(request, target_path)
        except Exception as exc:
            summary = f"读取 DOCX 段落失败：{exc}"
            runtime = self._build_runtime_metadata(
                terminal_status="failed",
                readonly_fallback_used=False,
                model_failed=False,
            )
            yield ledger.event(
                "tool.finished",
                {
                    "tool_name": "read_docx_content",
                    "success": False,
                    "path": target_path,
                    "result_preview": summary,
                },
                step_id=context_step_id,
            )
            yield ledger.event(
                "run.finished",
                {
                    "task": request.task,
                    "mode": "whitebox_v1",
                    "summary": summary,
                    "completed_task": False,
                    "context": [],
                    "file_changes": [],
                    "runtime": runtime,
                    "quick_action_mode": quick_action_mode,
                    "intent_plan": intent_plan_payload,
                    "requirements": requirements_payload,
                    "plan_check": plan_check_payload,
                    "recipe_skeleton": recipe_skeleton,
                    "constraint_audit": constraint_audit,
                    **classification_payload,
                },
            )
            return

        if not window["paragraphs"]:
            summary = "当前 DOCX 没有可润色的剩余段落。"
            runtime = self._build_runtime_metadata(
                terminal_status="verified",
                readonly_fallback_used=False,
                model_failed=False,
            )
            yield ledger.event(
                "run.finished",
                {
                    "task": request.task,
                    "mode": "whitebox_v1",
                    "summary": summary,
                    "completed_task": True,
                    "context": [window],
                    "file_changes": [],
                    "runtime": runtime,
                    "quick_action_mode": quick_action_mode,
                    "intent_plan": intent_plan_payload,
                    "requirements": requirements_payload,
                    "plan_check": plan_check_payload,
                    "recipe_skeleton": recipe_skeleton,
                    "constraint_audit": constraint_audit,
                    **classification_payload,
                },
            )
            return

        yield ledger.event(
            "tool.finished",
            {
                "tool_name": "read_docx_content",
                "success": True,
                "path": target_path,
                "result_preview": _preview("\n".join(window["paragraphs"]), 900),
                "paragraph_start": window["start_visible_index"] + 1,
                "paragraph_end": window["end_visible_index"],
            },
            step_id=context_step_id,
        )
        yield ledger.event(
            "step.result",
            self._build_step_result_payload(
                title="读取当前 DOCX 段落窗口",
                summary=(
                    f"已读取第 {window['start_visible_index'] + 1}-"
                    f"{window['end_visible_index']} 个非空段落。"
                ),
                status="completed",
                snippet_count=1,
                snippets=[
                    {
                        "source": Path(target_path).name,
                        "path": target_path,
                        "preview": _preview("\n".join(window["paragraphs"]), 500),
                        "paragraph_start": window["start_visible_index"] + 1,
                        "paragraph_end": window["end_visible_index"],
                    }
                ],
            ),
            step_id=context_step_id,
        )

        yield ledger.event(
            "step.started",
            {
                "title": "润色并写回当前段落",
                "detail": "只处理当前段落窗口，保留文档其他内容。",
            },
            step_id=execute_step_id,
        )

        model_failed = False
        polished: List[str] = []
        try:
            response = self._call_model(
                request=request,
                messages=[
                    {
                        "role": "user",
                        "content": self._docx_polish_window_prompt(
                            request, window["paragraphs"]
                        ),
                    }
                ],
                system=("你是严谨的中文文档润色助手。只润色用户给出的段落窗口，" "保持原意、术语和段落数量；不要扩写成总结，不要添加解释。"),
                tools=[],
            )
            content, _tool_calls = self._normalize_model_response(response, [])
            polished = self._parse_polished_docx_paragraphs(
                content, expected_count=len(window["paragraphs"])
            )
        except Exception as exc:
            model_failed = True
            logger.warning(
                "[FileTaskRuntime] stepwise DOCX polish model failed: %s", exc
            )

        if not polished:
            polished = [
                self._simple_polish_docx_paragraph(text)
                for text in window["paragraphs"]
            ]

        changed_count = self._rewrite_docx_paragraph_window(
            target_path,
            window["paragraph_indices"],
            polished,
        )
        change = {
            "path": target_path,
            "file_type": "docx",
            "operation": "rewrite_docx_paragraph_window",
            "summary": (
                f"已润色并写回第 {window['start_visible_index'] + 1}-"
                f"{window['end_visible_index']} 个非空段落。"
            ),
            "paragraphs_rewritten": changed_count,
            "paragraph_start": window["start_visible_index"] + 1,
            "paragraph_end": window["end_visible_index"],
            "change_type": "modify",
            "focus": True,
        }
        file_changes.append(change)
        yield ledger.event(
            "tool.finished",
            {
                "tool_name": "rewrite_docx_paragraph_window",
                "success": changed_count > 0,
                "path": target_path,
                "result_preview": change["summary"],
                "paragraphs_rewritten": changed_count,
            },
            step_id=execute_step_id,
        )
        yield ledger.event("file.changed", change, step_id=execute_step_id)

        next_artifact = self._docx_polish_wait_artifact(request, target_path, window)
        runtime = self._build_runtime_metadata(
            terminal_status="awaiting_confirmation",
            readonly_fallback_used=False,
            model_failed=model_failed,
        )
        check_payload = self._evaluate_task_quality_gate(
            request,
            file_changes,
            write_intent=True,
            output_mode="write",
        )
        check_payload.update(
            {
                "status": "awaiting_confirmation",
                "summary": "当前段落窗口已写回 DOCX，等待用户说“继续”后处理下一段。",
                "next_action_artifact": next_artifact,
                "runtime": runtime,
            }
        )

        yield ledger.event(
            "step.result",
            self._build_step_result_payload(
                title="润色并写回当前段落",
                summary=change["summary"],
                status="completed",
                file_changes=file_changes,
                runtime=runtime,
                next_action_artifact=next_artifact,
            ),
            step_id=execute_step_id,
        )
        yield ledger.event("check.completed", check_payload, step_id=check_step_id)
        yield ledger.event(
            "step.result",
            self._build_step_result_payload(
                title="核验结果",
                summary="当前步骤已写入 DOCX，等待用户说“继续”后处理下一段。",
                status="awaiting_confirmation",
                file_changes=file_changes,
                runtime=runtime,
                passed=bool(check_payload.get("passed")),
                next_action_artifact=next_artifact,
            ),
            step_id=check_step_id,
        )
        yield ledger.event(
            "run.finished",
            {
                "task": request.task,
                "mode": "whitebox_v1",
                "summary": "当前步骤已写入 DOCX，等待用户说“继续”后处理下一段。",
                "completed_task": True,
                "context": [window],
                "file_changes": file_changes,
                "runtime": runtime,
                "quick_action_mode": quick_action_mode,
                "intent_plan": intent_plan_payload,
                "requirements": requirements_payload,
                "plan_check": plan_check_payload,
                "recipe_skeleton": recipe_skeleton,
                "constraint_audit": constraint_audit,
                "next_action_artifact": next_artifact,
                **classification_payload,
            },
        )

    def _stepwise_docx_polish_target_path(
        self, request: FileTaskRequest, files: List[FileTaskFile]
    ) -> str:
        candidates: List[str] = []
        if request.target_path:
            candidates.append(str(request.target_path))
        options = request.options if isinstance(request.options, dict) else {}
        batch_control = (
            options.get("batch_control")
            if isinstance(options.get("batch_control"), dict)
            else {}
        )
        for value in (
            batch_control.get("target_path"),
            batch_control.get("source_path"),
        ):
            if value:
                candidates.append(str(value))
        for file_info in files:
            if _file_task_suffix(file_info) in {"doc", "docx"} and file_info.target:
                candidates.append(str(file_info.path or ""))
        for file_info in files:
            if _file_task_suffix(file_info) in {"doc", "docx"}:
                candidates.append(str(file_info.path or ""))
        for candidate in candidates:
            clean = candidate.strip()
            if clean and clean.lower().endswith((".doc", ".docx")):
                return clean
        return ""

    def _read_docx_paragraph_window(
        self, request: FileTaskRequest, path: str
    ) -> Dict[str, Any]:
        from docx import Document  # type: ignore

        doc = Document(path)
        visible_indices = [
            index
            for index, paragraph in enumerate(doc.paragraphs)
            if str(paragraph.text or "").strip()
        ]
        window_size = _stepwise_docx_polish_window_paragraphs(request)
        step_index = _stepwise_docx_polish_step_index(request)
        start_visible = step_index * window_size
        end_visible = min(start_visible + window_size, len(visible_indices))
        selected_indices = visible_indices[start_visible:end_visible]
        paragraphs = [doc.paragraphs[index].text for index in selected_indices]
        return {
            "source": Path(path).name,
            "path": path,
            "paragraph_indices": selected_indices,
            "paragraphs": paragraphs,
            "start_visible_index": start_visible,
            "end_visible_index": end_visible,
            "total_visible_paragraphs": len(visible_indices),
            "window_paragraphs": window_size,
            "step_index": step_index,
            "has_next": end_visible < len(visible_indices),
        }

    def _docx_polish_window_prompt(
        self, request: FileTaskRequest, paragraphs: List[str]
    ) -> str:
        numbered = "\n".join(
            f"{index}. {text}" for index, text in enumerate(paragraphs, start=1)
        )
        return (
            "请润色下面 DOCX 当前段落窗口。要求：\n"
            "1. 保持段落数量完全一致；\n"
            "2. 只改善语病、重复、口语化和不顺畅表达；\n"
            "3. 不改变事实、术语、数字和专名；\n"
            '4. 只返回 JSON 字符串数组，例如 ["润色后第1段", "润色后第2段"]。\n'
            f"用户任务：{request.task}\n\n"
            f"段落窗口：\n{numbered}"
        )

    def _parse_polished_docx_paragraphs(
        self, content: str, *, expected_count: int
    ) -> List[str]:
        text = str(content or "").strip()
        if not text:
            return []
        parsed = extract_first_json_value(text)
        if isinstance(parsed, dict):
            for key in ("paragraphs", "items", "result", "texts"):
                if isinstance(parsed.get(key), list):
                    parsed = parsed.get(key)
                    break
        if not isinstance(parsed, list):
            try:
                parsed = json.loads(text)
            except Exception:
                parsed = []
        if not isinstance(parsed, list):
            return []
        cleaned = [str(item or "").strip() for item in parsed[:expected_count]]
        cleaned = [item for item in cleaned if item]
        return cleaned if len(cleaned) == expected_count else []

    def _simple_polish_docx_paragraph(self, text: str) -> str:
        polished = re.sub(r"[ \t]+", " ", str(text or "")).strip()
        polished = re.sub(r"\s+([，。！？；：、])", r"\1", polished)
        polished = re.sub(r"([（【])\s+", r"\1", polished)
        polished = re.sub(r"\s+([）】])", r"\1", polished)
        return polished or str(text or "")

    def _rewrite_docx_paragraph_window(
        self, path: str, paragraph_indices: List[int], paragraphs: List[str]
    ) -> int:
        from docx import Document  # type: ignore

        doc = Document(path)
        changed = 0
        for paragraph_index, new_text in zip(paragraph_indices, paragraphs):
            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                continue
            paragraph = doc.paragraphs[paragraph_index]
            if paragraph.text == new_text:
                continue
            for run in list(paragraph.runs):
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)
            changed += 1
        doc.save(path)
        return changed

    def _docx_polish_wait_artifact(
        self, request: FileTaskRequest, target_path: str, window: Dict[str, Any]
    ) -> Dict[str, Any]:
        next_step_index = int(window.get("step_index") or 0) + 1
        window_paragraphs = int(window.get("window_paragraphs") or 8)
        next_start = int(window.get("end_visible_index") or 0) + 1
        next_end = min(
            next_start + window_paragraphs - 1,
            int(window.get("total_visible_paragraphs") or next_start),
        )
        label = (
            f"继续第 {next_start}-{next_end} 段"
            if bool(window.get("has_next"))
            else "已无下一段"
        )
        return {
            "artifact_type": "koto_stepwise_resume_v1",
            "category": "stepwise_confirmation",
            "route": "long_docx_stepwise_polish_writeback",
            "status": "awaiting_confirmation",
            "summary": "上一段落窗口已写回 DOCX。可以继续处理下一段。",
            "suggested_next_step": label,
            "actions": [
                {
                    "type": "file_task_resume",
                    "label": label,
                    "enabled": bool(window.get("has_next")),
                    "request": {
                        "task": f"继续分步润色 {Path(target_path).name}",
                        "target_path": target_path,
                        "files": [
                            {
                                "path": target_path,
                                "name": Path(target_path).name,
                                "type": "docx",
                                "target": True,
                            }
                        ],
                        "options": {
                            "batch_control": {
                                "policy": "confirm_each_step",
                                "step_index": next_step_index,
                                "window_paragraphs": window_paragraphs,
                                "target_path": target_path,
                                "source_path": target_path,
                                "original_task": request.task,
                                "route": "long_docx_stepwise_polish_writeback",
                            }
                        },
                    },
                }
            ],
            "stepwise": {
                "current_step_index": int(window.get("step_index") or 0),
                "next_step_index": next_step_index,
                "window_paragraphs": window_paragraphs,
                "has_next": bool(window.get("has_next")),
                "paragraph_start": int(window.get("start_visible_index") or 0) + 1,
                "paragraph_end": int(window.get("end_visible_index") or 0),
                "total_visible_paragraphs": int(
                    window.get("total_visible_paragraphs") or 0
                ),
            },
        }

    def _should_route_financial_xlsx_docx_report(
        self, request: FileTaskRequest, files: List[FileTaskFile]
    ) -> bool:
        from app.core.agent.file_task_financial_report_runner import (
            FileTaskFinancialReportRunner,
        )

        return FileTaskFinancialReportRunner(self).should_route(request, files)

    def _stream_financial_xlsx_docx_report(
        self,
        request: FileTaskRequest,
        context_files: List[FileTaskFile],
    ) -> Iterable[FileTaskEvent]:
        from app.core.agent.file_task_financial_report_runner import (
            FileTaskFinancialReportRunner,
        )

        yield from FileTaskFinancialReportRunner(self).stream(request, context_files)
    def _run_builtin_tool(
        self,
        ledger: FileTaskLedger,
        executor: ToolExecutor,
        *,
        step_id: str,
        tool_name: str,
        tool_args: Dict[str, Any],
        file_changes: List[Dict[str, Any]],
    ) -> tuple[Dict[str, Any], List[FileTaskEvent]]:
        events: List[FileTaskEvent] = []
        events.append(
            ledger.event(
                "tool.started",
                {
                    "tool_name": tool_name,
                    "tool_args": dict(tool_args or {}),
                },
                step_id=step_id,
            )
        )
        try:
            result = executor(tool_name, dict(tool_args or {}))
            success = not _is_error_result(result)
        except Exception as exc:
            result = {"error": str(exc)}
            success = False
            logger.warning(
                "[FileTaskRuntime] deterministic tool %s failed: %s", tool_name, exc
            )

        payload: Dict[str, Any]
        try:
            parsed = json.loads(stringify_result(result))
            payload = (
                parsed
                if isinstance(parsed, dict)
                else {"summary": stringify_result(result)}
            )
        except Exception:
            payload = {"summary": stringify_result(result)}

        events.append(
            ledger.event(
                "tool.finished",
                {
                    "tool_name": tool_name,
                    "success": success,
                    "result_preview": tool_result_preview(tool_name, result, 1200),
                },
                step_id=step_id,
            )
        )
        if success:
            for change in self._extract_file_changes(tool_name, tool_args, result):
                file_changes.append(change)
                events.append(ledger.event("file.changed", change, step_id=step_id))
        return payload, events

    def _request_has_file_type(self, request: FileTaskRequest, file_type: str) -> bool:
        return request_has_file_type(request, file_type)

    def _is_docx_annotation_request(self, request: FileTaskRequest) -> bool:
        if not self._request_has_file_type(request, "docx"):
            return False
        options = request.options if isinstance(request.options, dict) else {}
        if bool(options.get("skip_doc_annotate_bridge")):
            return False
        from app.core.agent import file_task_doc_annotate_boundary

        if file_task_doc_annotate_boundary.looks_like_docx_review_clear_request(
            request.task
        ):
            return False
        if file_task_doc_annotate_boundary.looks_like_direct_docx_rewrite_request(
            request.task
        ):
            return False
        if file_task_doc_annotate_boundary.looks_like_multi_file_compare_request(
            request
        ):
            return False
        task_lower = str(request.task or "").strip().lower()
        if not task_lower:
            return False
        if (
            self._request_has_file_type(request, "pdf")
            and any(marker in task_lower for marker in ("翻译", "translation", "译稿"))
            and any(marker in task_lower for marker in ("原文", "source", "pdf"))
            and any(marker in task_lower for marker in ("处理", "分段", "拆成", "batch"))
        ):
            return True
        return has_explicit_docx_review_intent(task_lower)

    def _is_docx_clear_review_request(self, request: FileTaskRequest) -> bool:
        if not self._request_has_file_type(request, "docx"):
            return False
        from app.core.agent import file_task_doc_annotate_boundary

        return file_task_doc_annotate_boundary.looks_like_docx_review_clear_request(
            request.task
        )

    def _consume_streaming_tool_result(
        self,
        ledger: FileTaskLedger,
        *,
        step_id: str,
        stream_result: FileTaskToolStreamResult,
    ) -> Iterable[FileTaskEvent | Any]:
        final_result: Any = None
        for chunk in stream_result.chunks:
            if not isinstance(chunk, FileTaskToolStreamChunk):
                continue
            if str(chunk.kind or "").strip().lower() == "event":
                event_type = str(chunk.event_type or "").strip()
                if not event_type:
                    continue
                payload = dict(chunk.payload) if isinstance(chunk.payload, dict) else {}
                yield ledger.event(event_type, payload, step_id=step_id)
                continue
            if str(chunk.kind or "").strip().lower() == "result":
                final_result = chunk.payload
        return final_result

    def _has_write_intent(self, task: str) -> bool:
        if self._is_diagnostic_request(task):
            return False
        if self._has_readonly_write_negation(task):
            return False
        strong_write_intent = self._has_strong_write_intent(task)
        explicit_write_intent = self._has_explicit_write_intent(task)
        if self._is_advisory_analysis_request(task) and not strong_write_intent:
            return False
        return explicit_write_intent or strong_write_intent

    def _has_strong_write_intent(self, task: str) -> bool:
        if self._has_readonly_write_negation(task):
            return False
        lowered = (task or "").lower()
        task_text = task or ""
        if any(word in lowered for word in _EXPLICIT_WRITE_INTENT_WORDS):
            return True
        if any(pattern.search(task_text) for pattern in _WRITE_INTENT_PATTERNS):
            return True
        if any(pattern.search(task_text) for pattern in _IMPERATIVE_WRITE_PATTERNS):
            return True
        markers = semantic_markers(task_text)
        if (
            markers.get("docx_write_phrase")
            or markers.get("docx_create_phrase")
            or markers.get("ppt_slide_write_request")
            or markers.get("ppt_design_request")
        ):
            return True
        return bool(
            re.search(
                r"(?:加入|添加|插入|放入|写入).{0,18}(?:docx|word|文档|pptx?|幻灯片|slides?)",
                task_text,
                re.IGNORECASE,
            )
        )

    def _has_explicit_write_intent(self, task: str) -> bool:
        if self._has_readonly_write_negation(task):
            return False
        lowered = (task or "").lower()
        task_text = task or ""
        if any(word in lowered for word in _EXPLICIT_WRITE_INTENT_WORDS):
            return True
        if any(pattern.search(task_text) for pattern in _WRITE_INTENT_PATTERNS):
            return True
        if any(pattern.search(task_text) for pattern in _IMPERATIVE_WRITE_PATTERNS):
            return True
        has_soft_action = any(word in lowered for word in _SOFT_WRITE_ACTION_WORDS)
        has_target_hint = any(word in lowered for word in _WRITE_TARGET_HINT_WORDS)
        if has_soft_action and has_target_hint:
            return True
        markers = semantic_markers(task_text)
        if (
            markers.get("docx_write_phrase")
            or markers.get("docx_create_phrase")
            or markers.get("ppt_design_request")
        ):
            return True
        return any(word in lowered for word in _WRITE_INTENT_WORDS)

    def _has_readonly_write_negation(self, task: str) -> bool:
        task_text = str(task or "").strip()
        if not task_text:
            return False
        return any(
            pattern.search(task_text) for pattern in _READONLY_WRITE_NEGATION_PATTERNS
        )

    def _is_advisory_analysis_request(self, task: str) -> bool:
        task_text = str(task or "").strip()
        if not task_text:
            return False
        lowered = task_text.lower()
        if any(pattern.search(task_text) for pattern in _ANALYSIS_ADVICE_PATTERNS):
            return True
        has_analysis_cue = any(word in lowered for word in _ANALYSIS_CUE_WORDS)
        has_advice_cue = any(word in lowered for word in _ADVICE_CUE_WORDS)
        return (
            has_analysis_cue
            and has_advice_cue
            and not self._has_explicit_write_intent(task_text)
        )

    def _is_diagnostic_request(self, task: str) -> bool:
        task_text = str(task or "").strip()
        if not task_text:
            return False
        if any(pattern.search(task_text) for pattern in _DIAGNOSTIC_NEW_TASK_PATTERNS):
            return False
        return any(
            pattern.search(task_text) for pattern in _DIAGNOSTIC_REQUEST_PATTERNS
        )

    def _explicit_output_mode(self, request: FileTaskRequest) -> str:
        options = request.options if isinstance(request.options, dict) else {}
        normalized = str(options.get("output_mode") or "").strip().lower()
        if normalized in {"answer", "write", "hybrid"}:
            return normalized
        return ""

    def _has_target_context(
        self, request: FileTaskRequest, files: List[FileTaskFile]
    ) -> bool:
        if str(request.target_path or "").strip():
            return True
        if request.current_file is not None:
            return True
        return any(bool(file_info and file_info.target) for file_info in files)

    def _infer_output_mode(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        *,
        write_intent: bool,
        diagnostic_request: bool,
        docx_annotation_request: bool,
        advisory_analysis_request: bool,
    ) -> str:
        explicit_mode = self._explicit_output_mode(request)
        if explicit_mode:
            if explicit_mode == "answer" and not diagnostic_request and write_intent:
                return "write"
            return explicit_mode
        if diagnostic_request:
            return "answer"
        if write_intent or docx_annotation_request:
            return "write"
        if advisory_analysis_request and self._has_target_context(request, files):
            return "hybrid"
        return "answer"

    def _quick_action_mode(self, request: FileTaskRequest) -> str:
        options = request.options if isinstance(request.options, dict) else {}
        return str(options.get("quick_action_mode") or "").strip().lower()

    def _classify_request(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        known_tool_gap: Optional[Dict[str, Any]] = None,
    ) -> FileTaskClassification:
        options = request.options if isinstance(request.options, dict) else {}
        followup_context = self._followup_context(request)
        batch_control = (
            options.get("batch_control")
            if isinstance(options.get("batch_control"), dict)
            else {}
        )
        classification_task = self._classification_task_text(request, batch_control)
        classification_request = self._request_with_task(request, classification_task)
        matched_capabilities = matched_native_capability_names(classification_request)
        advisory_analysis_request = self._is_advisory_analysis_request(
            classification_task
        )
        readonly_write_negation = self._has_readonly_write_negation(classification_task)
        raw_write_intent = self._has_explicit_write_intent(classification_task)
        write_intent = self._has_write_intent(classification_task)
        raw_docx_annotation_request = self._is_docx_annotation_request(
            classification_request
        )
        docx_annotation_request = raw_docx_annotation_request
        clear_docx_review_request = self._is_docx_clear_review_request(
            classification_request
        )
        docx_compare_annotate_request = (
            "compare_docx_and_annotate" in matched_capabilities
        )
        if docx_compare_annotate_request:
            if "annotate_file" in matched_capabilities:
                matched_capabilities = [
                    name for name in matched_capabilities if name != "annotate_file"
                ]
            docx_annotation_request = False
            raw_docx_annotation_request = False
        semantic = semantic_markers(
            classification_task,
            file_types=request_file_types(files),
            target_file_type=request_target_file_type(request, files),
        )
        chart_request = semantic.get("chart_request", False)
        table_request = semantic.get("table_request", False)
        summary_request = semantic.get("summary_request", False)
        translation_request = semantic.get("translation_request", False)
        polish_request = semantic.get("polish_request", False)
        financial_request = semantic.get("financial_request", False)
        ppt_slide_write_request = semantic.get("ppt_slide_write_request", False)
        ppt_design_request = semantic.get("ppt_design_request", False)
        docx_report_request = semantic.get("docx_report_request", False)
        if clear_docx_review_request and "annotate_file" in matched_capabilities:
            matched_capabilities = [
                name for name in matched_capabilities if name != "annotate_file"
            ]
        planner_policy, planner_reason, planner_backend = self._planner_classification(
            request
        )
        batch_adapter = str(batch_control.get("adapter") or "").strip().lower()
        diagnostic_request = self._is_diagnostic_request(classification_task)

        request_kind = "new_task"
        execution_mode = "generic_tool_loop"
        reason_codes: List[str] = []
        stepwise_pdf_docx_resume = False
        followup_action = (
            str(followup_context.get("followup_action") or "").strip().lower()
        )
        previous_task_family = (
            str(followup_context.get("previous_task_family") or "").strip().lower()
        )
        previous_task_execution_mode = (
            str(followup_context.get("previous_task_execution_mode") or "")
            .strip()
            .lower()
        )
        previous_task_output_mode = (
            str(followup_context.get("previous_task_output_mode") or "").strip().lower()
        )
        previous_task_intent_can_apply = (
            str(followup_context.get("previous_task_intent_can_apply") or "")
            .strip()
            .lower()
        )
        if (
            semantic.get("pdf_source", False)
            and semantic.get("summary_request", False)
            and semantic.get("stepwise_confirmation_request", False)
            and semantic.get("docx_target", False)
        ):
            summary_request = True
            docx_report_request = True
            write_intent = True
            raw_write_intent = True
            reason_codes.append("long_pdf_stepwise_docx_forced_write_intent")
        if batch_control:
            request_kind = "resume"
            execution_mode = "awaiting_confirmation_resume"
            reason_codes.append("batch_control_resume")
            if batch_adapter:
                reason_codes.append(f"batch_adapter:{batch_adapter}")
            if (
                str(batch_control.get("policy") or "").strip().lower()
                == "confirm_each_step"
                and "pdf" in request_file_types(files)
                and request_target_file_type(request, files) in {"docx", "doc"}
            ):
                stepwise_pdf_docx_resume = True
                summary_request = True
                docx_report_request = True
                write_intent = True
                raw_write_intent = True
                reason_codes.append("stepwise_resume_forced_write_intent")
        elif followup_context:
            request_kind = "followup"
            execution_mode = "followup_contextual"
            if followup_action:
                reason_codes.append(f"followup_action:{followup_action}")
            else:
                reason_codes.append("followup_context")

        if request_kind == "followup" and followup_action == "question":
            diagnostic_request = True
            reason_codes.append("followup_question")

        if clear_docx_review_request:
            reason_codes.append("docx_clear_review_request")
            if not write_intent:
                write_intent = True
                reason_codes.append("docx_clear_review_forced_write_intent")

        if docx_compare_annotate_request:
            reason_codes.append("docx_compare_annotate_request")
            if not write_intent:
                write_intent = True
                reason_codes.append("docx_compare_annotate_forced_write_intent")

        if diagnostic_request:
            reason_codes.append("diagnostic_request")
            if write_intent or raw_write_intent:
                write_intent = False
                reason_codes.append("diagnostic_overrode_write_intent")
            if docx_annotation_request or raw_docx_annotation_request:
                docx_annotation_request = False
                reason_codes.append("diagnostic_overrode_docx_annotation")

        if readonly_write_negation:
            reason_codes.append("readonly_write_negation")
            if write_intent or raw_write_intent:
                write_intent = False
                reason_codes.append("readonly_overrode_write_intent")
            if docx_annotation_request or raw_docx_annotation_request:
                docx_annotation_request = False
                reason_codes.append("readonly_overrode_docx_annotation")

        if self._explicit_output_mode(request) == "answer" and not diagnostic_request:
            if (write_intent or raw_write_intent) and not self._has_strong_write_intent(
                classification_task
            ):
                write_intent = False
                reason_codes.append("answer_mode_overrode_write_intent")
            if docx_annotation_request or raw_docx_annotation_request:
                docx_annotation_request = False
                reason_codes.append("answer_mode_overrode_docx_annotation")

        if batch_adapter == "doc_annotate_bridge":
            docx_annotation_request = True
        if request_kind == "followup" and followup_action == "improve":
            if previous_task_family == "annotate":
                reason_codes.append("followup_previous_task_family:annotate")
                if self._request_has_file_type(classification_request, "docx"):
                    docx_annotation_request = True
            if previous_task_execution_mode in {
                "annotate_tool_loop",
                "awaiting_confirmation_resume",
            }:
                reason_codes.append(
                    f"followup_previous_execution_mode:{previous_task_execution_mode}"
                )
                if self._request_has_file_type(classification_request, "docx"):
                    docx_annotation_request = True
        if request_kind == "followup" and followup_action == "apply":
            if (
                previous_task_output_mode in {"hybrid", "write"}
                or previous_task_intent_can_apply == "true"
            ):
                write_intent = True
                reason_codes.append("followup_apply_write_intent")

        if docx_annotation_request:
            if request_kind == "new_task":
                execution_mode = "annotate_tool_loop"
            reason_codes.append("docx_annotation_request")
            if not write_intent:
                write_intent = True
                reason_codes.append("docx_annotation_forced_write_intent")

        if write_intent:
            reason_codes.append("write_intent")
            if (
                str(options.get("output_mode") or "").strip().lower() == "answer"
                and not diagnostic_request
            ):
                output_mode = "write"
                reason_codes.append("answer_mode_overridden_by_write_intent")

        recipe_match_request = classification_request
        if stepwise_pdf_docx_resume:
            recipe_match_request = FileTaskRequest(
                task=(f"{classification_task}\n" "分步 长PDF DOCX 总结 每一步写入并等待确认"),
                run_id=classification_request.run_id,
                session_id=classification_request.session_id,
                files=classification_request.files,
                current_file=classification_request.current_file,
                selection=classification_request.selection,
                selection_source=classification_request.selection_source,
                target_path=classification_request.target_path,
                model_mode=classification_request.model_mode,
                model_id=classification_request.model_id,
                history=classification_request.history,
                options=classification_request.options,
            )
        recipe_candidates = recipe_matches(
            recipe_match_request, files, write_intent=write_intent
        )
        selected_recipe_match = recipe_candidates[0] if recipe_candidates else None
        if selected_recipe_match:
            reason_codes.extend(selected_recipe_match.reason_codes)
            for capability in selected_recipe_match.recipe.matched_capabilities:
                if capability not in matched_capabilities:
                    matched_capabilities.append(capability)
            if selected_recipe_match.recipe.execution_mode != "generic_tool_loop":
                execution_mode = selected_recipe_match.recipe.execution_mode
            if len(recipe_candidates) > 1:
                reason_codes.extend(
                    f"recipe_candidate:{item.recipe.id}"
                    for item in recipe_candidates[1:4]
                )

        if planner_policy:
            reason_codes.append(f"planner_policy:{planner_policy}")
        elif planner_reason == "deferred_to_execution_brief":
            reason_codes.append("planner_deferred:model_first")
        if planner_backend:
            reason_codes.append(f"planner_backend:{planner_backend}")

        known_gap_name = ""
        if isinstance(known_tool_gap, dict):
            known_gap_name = str(known_tool_gap.get("missing_capability") or "").strip()
            if known_gap_name:
                reason_codes.append(f"native_tool_gap:{known_gap_name}")

        reason_codes.extend(f"capability:{name}" for name in matched_capabilities[:4])
        semantic_reason_markers = {
            "chart_request": chart_request,
            "table_request": table_request,
            "summary_request": summary_request,
            "translation_request": translation_request,
            "polish_request": polish_request,
            "financial_request": financial_request,
            "ppt_slide_write_request": ppt_slide_write_request,
            "ppt_design_request": ppt_design_request,
            "docx_report_request": docx_report_request,
        }
        reason_codes.extend(
            name for name, enabled in semantic_reason_markers.items() if enabled
        )

        task_family = "analyze"
        operation_kind = "read"
        if diagnostic_request:
            task_family = "analyze"
            operation_kind = "read"
        elif clear_docx_review_request:
            task_family = "transform"
            operation_kind = "write"
            docx_annotation_request = False
            if "annotate_file" in matched_capabilities:
                matched_capabilities = [
                    name for name in matched_capabilities if name != "annotate_file"
                ]
        elif (
            selected_recipe_match
            and selected_recipe_match.recipe.execution_mode == "doc_annotate_bridge"
        ):
            task_family = selected_recipe_match.recipe.task_family
            operation_kind = selected_recipe_match.recipe.write_operation_kind
            docx_annotation_request = True
        elif (
            selected_recipe_match
            and selected_recipe_match.recipe.id == "docx_contract_compare_review"
        ):
            task_family = selected_recipe_match.recipe.task_family
            operation_kind = selected_recipe_match.recipe.write_operation_kind
        elif (
            selected_recipe_match
            and selected_recipe_match.recipe.id == "docx_compare_annotation"
        ):
            task_family = selected_recipe_match.recipe.task_family
            operation_kind = selected_recipe_match.recipe.write_operation_kind
            if "annotate_file" in matched_capabilities:
                matched_capabilities = [
                    name for name in matched_capabilities if name != "annotate_file"
                ]
        elif docx_compare_annotate_request:
            task_family = "compare"
            operation_kind = "compare_annotate"
        elif docx_annotation_request or "annotate_file" in matched_capabilities:
            task_family = "annotate"
            operation_kind = "annotate"
        elif selected_recipe_match:
            task_family = selected_recipe_match.recipe.task_family
            operation_kind = (
                selected_recipe_match.recipe.write_operation_kind
                if write_intent
                else selected_recipe_match.recipe.read_operation_kind
            )
            for capability in selected_recipe_match.recipe.matched_capabilities:
                if capability not in matched_capabilities:
                    matched_capabilities.append(capability)
            if selected_recipe_match.recipe.execution_mode != "generic_tool_loop":
                execution_mode = selected_recipe_match.recipe.execution_mode
        elif financial_request and chart_request and docx_report_request:
            task_family = "financial_report"
            operation_kind = (
                "analyze_visualize_write" if write_intent else "analyze_visualize"
            )
        elif "compare_files" in matched_capabilities:
            task_family = "compare"
            operation_kind = "compare"
        elif ppt_slide_write_request:
            task_family = "presentation"
            operation_kind = "write_slides" if write_intent else "read"
        elif translation_request:
            task_family = "translate"
            operation_kind = "write" if write_intent else "read"
        elif polish_request:
            task_family = "polish"
            operation_kind = "write" if write_intent else "read"
        elif self._looks_like_problem_analysis_request(classification_task):
            task_family = "analyze"
            operation_kind = "write" if write_intent else "read"
        elif summary_request:
            task_family = "summarize"
            operation_kind = "write" if write_intent else "read"
        elif chart_request:
            task_family = "visualize"
            operation_kind = "visualize_write" if write_intent else "visualize"
        elif table_request and write_intent:
            task_family = "table_transfer"
            operation_kind = "write_table"
        elif "run_python_code" in matched_capabilities:
            task_family = "automation"
            operation_kind = "compute"
        elif write_intent:
            task_family = "transform"
            operation_kind = "write"

        file_types = sorted(
            {
                str(profile.get("format") or "").strip().lower()
                for profile in build_request_capability_profiles(request)
                if str(profile.get("format") or "").strip()
            }
        )
        target_file_type = (
            Path(str(request.target_path or "")).suffix.lstrip(".").lower()
        )
        if not target_file_type:
            for file_info in files:
                if not file_info.target:
                    continue
                target_file_type = (
                    file_info.type
                    or Path(file_info.path or file_info.name).suffix.lstrip(".")
                ).lower()
                if target_file_type:
                    break

        output_mode = self._infer_output_mode(
            request,
            files,
            write_intent=write_intent,
            diagnostic_request=diagnostic_request,
            docx_annotation_request=docx_annotation_request,
            advisory_analysis_request=advisory_analysis_request,
        )

        confidence = 1.0
        if diagnostic_request:
            confidence = (
                0.7 if (raw_write_intent or raw_docx_annotation_request) else 0.9
            )

        return FileTaskClassification(
            request_kind=request_kind,
            task_family=task_family,
            operation_kind=operation_kind,
            execution_mode=execution_mode,
            output_mode=output_mode,
            write_intent=write_intent,
            diagnostic_request=diagnostic_request,
            docx_annotation_request=docx_annotation_request,
            planner_policy=planner_policy,
            planner_reason=planner_reason,
            planner_backend=planner_backend,
            target_file_type=target_file_type,
            known_native_tool_gap=known_gap_name,
            file_types=file_types,
            matched_capabilities=matched_capabilities,
            reason_codes=reason_codes,
            selected_recipe=(
                selected_recipe_match.recipe.id if selected_recipe_match else ""
            ),
            recipe_candidates=[item.public_dict() for item in recipe_candidates[:5]],
            confidence=confidence,
        )

    def _effective_planner_classification(
        self, request: FileTaskRequest
    ) -> tuple[str, str, str]:
        return self._planner_classification(request)

    def _classification_task_text(
        self, request: FileTaskRequest, batch_control: Dict[str, Any]
    ) -> str:
        task_text = str(request.task or "").strip()
        original_task = ""
        if isinstance(batch_control, dict):
            original_task = str(batch_control.get("original_task") or "").strip()
        if original_task and original_task not in task_text:
            return f"{task_text}\n原始分步任务：{original_task}".strip()
        return task_text

    def _request_with_task(
        self, request: FileTaskRequest, task_text: str
    ) -> FileTaskRequest:
        if str(task_text or "") == str(request.task or ""):
            return request
        return FileTaskRequest(
            task=task_text,
            run_id=request.run_id,
            session_id=request.session_id,
            files=list(request.files),
            current_file=request.current_file,
            selection=request.selection,
            selection_source=request.selection_source,
            target_path=request.target_path,
            model_mode=request.model_mode,
            model_id=request.model_id,
            history=list(request.history),
            options=dict(request.options),
        )

    def _should_adjudicate_intent(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        classification: FileTaskClassification,
    ) -> bool:
        options = request.options if isinstance(request.options, dict) else {}
        if bool(options.get("disable_ai_intent_adjudicator")):
            return False
        if any(
            key in options
            for key in ("planner_backend", "planner_policy", "planner_command")
        ):
            return False
        if str(options.get("quick_action_mode") or "").strip().lower() == "simple":
            return False
        if bool(options.get("enable_ai_intent_adjudicator")):
            return True
        task_text = str(request.task or "").strip()
        if not task_text:
            return False
        if classification.request_kind == "resume":
            return False
        if classification.selected_recipe in {
            "long_pdf_stepwise_docx_summary",
            "financial_xlsx_docx_report",
            "docx_clear_review_marks",
        }:
            return False
        if classification.diagnostic_request or self._has_readonly_write_negation(
            task_text
        ):
            return False
        explicit_mode = self._explicit_output_mode(request)
        if classification.selected_recipe and not explicit_mode:
            return False
        if explicit_mode == "answer" and classification.write_intent:
            return True
        if explicit_mode == "hybrid" and classification.write_intent:
            return True
        has_target = self._has_target_context(request, files)
        if not has_target:
            return False
        lowered = task_text.lower()
        ambiguity_markers = (
            "看看",
            "看下",
            "帮我看",
            "建议",
            "怎么改",
            "如何改",
            "优化",
            "风格",
            "主题",
            "配色",
            "好看",
            "美化",
            "调整",
            "改进",
            "review",
            "suggest",
            "style",
            "theme",
        )
        if any(marker in lowered for marker in ambiguity_markers):
            return True
        return False

    def _intent_adjudicator_system_prompt(self) -> str:
        return (
            "你是 Koto 文件助手的任务意图裁判。你不执行任务，只判断用户希望产生什么结果。\n"
            "请严格区分：\n"
            "1. answer_only：只回答，不改文件。\n"
            "2. analyze_then_confirm：先分析建议，再等用户确认是否应用到文件。\n"
            "3. edit_file：直接修改当前/目标文件。\n"
            "4. create_file：创建新文件。\n"
            "5. resume_stepwise：继续上一步分步任务。\n"
            "6. diagnose_failure：解释任务为什么失败或上一轮哪里不对。\n"
            "判断规则：\n"
            "- “改、换、应用、写入、创建、美化、更新、删除、插入、套用、换成”通常是写入。\n"
            "- “看看、分析、建议、为什么、哪里有问题”通常是只读、先分析后确认或诊断。\n"
            "- “继续”要结合上一轮任务状态；没有上一轮状态时不要臆造。\n"
            "- 明确的“不写入、不修改、只分析、只给答案”必须覆盖其他写入词。\n"
            "- 如果入口模式和用户正文冲突，优先判断用户正文真正要求的产物。\n"
            "只输出严格 JSON，不要输出 Markdown 或解释文本。"
        )

    def _intent_adjudicator_messages(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        classification: FileTaskClassification,
    ) -> List[Dict[str, Any]]:
        file_payload = [file_info.public_dict() for file_info in files[:8]]
        payload = {
            "task": request.task,
            "target_path": request.target_path,
            "files": file_payload,
            "entry_options": (
                dict(request.options) if isinstance(request.options, dict) else {}
            ),
            "rule_classification": classification.public_dict(),
            "required_json_schema": {
                "intent": "answer_only | analyze_then_confirm | edit_file | create_file | resume_stepwise | diagnose_failure",
                "confidence": "0.0-1.0",
                "should_write": "boolean",
                "needs_clarification": "boolean",
                "target_file_type": "string",
                "operation": "short operation name",
                "reason": "brief reason",
            },
        }
        return [
            {
                "role": "user",
                "content": json.dumps(payload, ensure_ascii=False),
            }
        ]

    def _adjudicate_intent_if_needed(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        classification: FileTaskClassification,
    ) -> Dict[str, Any]:
        if not self._should_adjudicate_intent(request, files, classification):
            return {}
        try:
            response = self._call_model(
                request=request,
                messages=self._intent_adjudicator_messages(
                    request, files, classification
                ),
                system=self._intent_adjudicator_system_prompt(),
                tools=[],
            )
        except Exception as exc:
            logger.warning("[FileTaskRuntime] intent adjudicator unavailable: %s", exc)
            return {
                "source": "ai_intent_adjudicator",
                "status": "unavailable",
                "error": _preview(str(exc), 240),
            }

        content = (
            str(response.get("content") or response.get("text") or "").strip()
            if isinstance(response, dict)
            else str(response or "").strip()
        )
        candidate: Any = None
        if isinstance(response, dict):
            for key in ("intent_adjudication", "intent", "classification"):
                if isinstance(response.get(key), dict):
                    candidate = response.get(key)
                    break
        if candidate is None:
            candidate = extract_first_json_value(content)
        if not isinstance(candidate, dict):
            return {
                "source": "ai_intent_adjudicator",
                "status": "invalid",
                "raw_preview": _preview(content, 360),
            }
        intent = str(candidate.get("intent") or "").strip().lower()
        confidence = _safe_float(candidate.get("confidence"), 0.0)
        return {
            "source": "ai_intent_adjudicator",
            "status": "ok" if intent else "invalid",
            "intent": intent,
            "confidence": max(0.0, min(1.0, confidence)),
            "should_write": bool(candidate.get("should_write")),
            "needs_clarification": bool(candidate.get("needs_clarification")),
            "target_file_type": str(candidate.get("target_file_type") or "")
            .strip()
            .lower()
            .lstrip("."),
            "operation": str(candidate.get("operation") or "").strip()[:120],
            "reason": str(candidate.get("reason") or "").strip()[:500],
        }

    def _apply_intent_adjudication(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        classification: FileTaskClassification,
        adjudication: Dict[str, Any],
    ) -> FileTaskClassification:
        if not isinstance(adjudication, dict) or adjudication.get("status") != "ok":
            if isinstance(adjudication, dict) and adjudication.get("status"):
                classification.reason_codes.append(
                    f"ai_intent_adjudicator:{adjudication.get('status')}"
                )
            return classification
        intent = str(adjudication.get("intent") or "").strip().lower()
        confidence = float(adjudication.get("confidence") or 0.0)
        should_write = bool(adjudication.get("should_write"))
        if confidence < 0.55:
            classification.reason_codes.append("ai_intent_adjudicator_low_confidence")
            return classification
        if self._has_readonly_write_negation(request.task):
            classification.reason_codes.append("ai_intent_adjudicator_readonly_guard")
            return classification
        if classification.diagnostic_request:
            classification.reason_codes.append("ai_intent_adjudicator_diagnostic_guard")
            return classification

        output_override = ""
        write_override: Optional[bool] = None
        if intent in {"edit_file", "create_file", "resume_stepwise"} or should_write:
            output_override = "write"
            write_override = True
        elif intent == "analyze_then_confirm":
            output_override = "hybrid"
            write_override = False
        elif intent in {"answer_only", "diagnose_failure"}:
            if not self._has_strong_write_intent(request.task):
                output_override = "answer"
                write_override = False

        if not output_override:
            classification.reason_codes.append("ai_intent_adjudicator_no_override")
            return classification

        original_output = str(classification.output_mode or "").strip().lower()
        original_write = bool(classification.write_intent)
        classification.output_mode = output_override
        if write_override is not None:
            classification.write_intent = bool(write_override)
        classification.confidence = max(
            float(classification.confidence or 0.0), confidence
        )
        classification.reason_codes.append(f"ai_intent_adjudicator:{intent}")
        if (
            original_output != classification.output_mode
            or original_write != classification.write_intent
        ):
            classification.reason_codes.append("ai_intent_adjudicator_override")

        recipe_candidates = recipe_matches(
            request, files, write_intent=classification.write_intent
        )
        selected_recipe_match = recipe_candidates[0] if recipe_candidates else None
        if selected_recipe_match:
            classification.selected_recipe = selected_recipe_match.recipe.id
            classification.recipe_candidates = [
                item.public_dict() for item in recipe_candidates[:5]
            ]
            for capability in selected_recipe_match.recipe.matched_capabilities:
                if capability not in classification.matched_capabilities:
                    classification.matched_capabilities.append(capability)
            classification.task_family = selected_recipe_match.recipe.task_family
            classification.operation_kind = (
                selected_recipe_match.recipe.write_operation_kind
                if classification.write_intent
                else selected_recipe_match.recipe.read_operation_kind
            )
            for code in selected_recipe_match.reason_codes:
                if code not in classification.reason_codes:
                    classification.reason_codes.append(code)
        return classification

    def _build_execution_context(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        *,
        known_tool_gap: Optional[Dict[str, Any]] = None,
        classification: Optional[FileTaskClassification] = None,
        intent_plan: Optional[FileTaskIntentPlan] = None,
        intent_adjudication: Optional[Dict[str, Any]] = None,
        quick_action_mode: str = "",
    ) -> FileTaskExecutionContext:
        resolved_known_tool_gap = (
            known_tool_gap
            if isinstance(known_tool_gap, dict)
            else native_tool_gap_for_request(request)
        )
        resolved_classification = classification or self._classify_request(
            request, files, resolved_known_tool_gap
        )
        resolved_intent_plan = self._resolve_intent_plan(
            request,
            files,
            known_tool_gap=resolved_known_tool_gap,
            classification=resolved_classification,
            intent_plan=intent_plan,
        )
        requirements = build_file_task_requirements(request, resolved_classification)
        plan_check = validate_file_task_plan(
            requirements, resolved_classification, resolved_intent_plan
        )
        (
            effective_planner_policy,
            effective_planner_reason,
            effective_planner_backend,
        ) = self._effective_planner_classification(request)
        resolved_quick_action_mode = (
            str(quick_action_mode or self._quick_action_mode(request)).strip().lower()
        )
        return FileTaskExecutionContext(
            classification=resolved_classification,
            intent_plan=resolved_intent_plan,
            requirements=requirements,
            plan_check=plan_check,
            known_tool_gap=resolved_known_tool_gap,
            intent_adjudication=dict(intent_adjudication or {}),
            effective_planner_policy=effective_planner_policy,
            effective_planner_reason=effective_planner_reason,
            effective_planner_backend=effective_planner_backend,
            quick_action_mode=resolved_quick_action_mode,
            simple_quick_action=resolved_quick_action_mode == "simple",
        )

    def _constraint_audit(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        classification: FileTaskClassification,
        intent_plan: FileTaskIntentPlan,
        requirements: FileTaskRequirementSet,
        recipe_skeleton: Dict[str, Any],
    ) -> Dict[str, Any]:
        hard: List[str] = ["allowlist_tools_only"]
        soft: List[str] = []
        ignored: List[str] = []
        conflicts: List[str] = []

        options = request.options if isinstance(request.options, dict) else {}
        for key in (
            "deterministic_financial_xlsx_docx_report",
            "force_model_financial_xlsx_docx_report",
        ):
            if key in options:
                ignored.append(f"legacy_option_ignored:{key}")

        recipe_id = str(recipe_skeleton.get("recipe_id") or "").strip()
        if recipe_id and recipe_id != "generic_file_task":
            hard.append(f"recipe:{recipe_id}")
        if bool(requirements.write_required):
            hard.append("write_requires_file_changed")
        if bool(recipe_skeleton.get("quality_gates")):
            hard.append("quality_gates_enforced")
        if recipe_id == "financial_xlsx_docx_report":
            hard.append("native_financial_workflow")

        target_type = (
            str(requirements.target_file_type or classification.target_file_type or "")
            .strip()
            .lower()
        )
        if target_type in {"docx", "doc", "pptx", "ppt", "xlsx", "xlsm"}:
            hard.append("explicit_or_unambiguous_target_required")

        if classification.output_mode == "hybrid":
            soft.append("hybrid_mode_default_no_write_without_apply")
        if intent_plan.requires_confirmation:
            soft.append("confirmation_required_before_apply")
        if recipe_id == "generic_file_task":
            soft.append("model_guided_generic_loop")

        if requirements.write_required and classification.output_mode != "write":
            conflicts.append("write_required_output_mode_mismatch")
        if requirements.write_required and intent_plan.write_intent is False:
            conflicts.append("write_required_intent_plan_mismatch")
        if not requirements.write_required and classification.output_mode == "write":
            conflicts.append("readonly_request_escalated_to_write")

        same_type_files = (
            self._context_files_by_type(files, {target_type}) if target_type else []
        )
        if (
            target_type
            and len(same_type_files) > 1
            and not str(request.target_path or "").strip()
        ):
            conflicts.append(f"ambiguous_target:{target_type}")

        return {
            "version": "file_task_constraint_audit_v1",
            "hard_constraints": sorted(set(hard)),
            "soft_constraints": sorted(set(soft)),
            "ignored_legacy_options": sorted(set(ignored)),
            "conflicts": sorted(set(conflicts)),
            "status": "conflict" if conflicts else "clear",
        }

    def _financial_constraint_audit(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        *,
        ambiguous_docx_target: bool = False,
    ) -> Dict[str, Any]:
        options = request.options if isinstance(request.options, dict) else {}
        ignored = [
            f"legacy_option_ignored:{key}"
            for key in (
                "deterministic_financial_xlsx_docx_report",
                "force_model_financial_xlsx_docx_report",
            )
            if key in options
        ]
        conflicts = ["ambiguous_target:docx"] if ambiguous_docx_target else []
        return {
            "version": "file_task_constraint_audit_v1",
            "hard_constraints": [
                "allowlist_tools_only",
                "explicit_or_unambiguous_target_required",
                "native_financial_workflow",
                "quality_gates_enforced",
                "recipe:financial_xlsx_docx_report",
                "write_requires_file_changed",
            ],
            "soft_constraints": [],
            "ignored_legacy_options": ignored,
            "conflicts": conflicts,
            "status": "conflict" if conflicts else "clear",
        }

    def _planner_classification(self, request: FileTaskRequest) -> tuple[str, str, str]:
        return "native_only", "file_task_native_only", "native"

    def _has_explicit_planner_override(self, request: FileTaskRequest) -> bool:
        return False

    def _sanitize_planner_options(self, options: Dict[str, Any]) -> Dict[str, Any]:
        return {
            str(key): value
            for key, value in dict(options or {}).items()
            if "planner" not in str(key)
        }

    def _clone_request_with_options(
        self, request: FileTaskRequest, options: Dict[str, Any]
    ) -> FileTaskRequest:
        return FileTaskRequest(
            task=request.task,
            run_id=request.run_id,
            session_id=request.session_id,
            files=list(request.files),
            current_file=request.current_file,
            selection=request.selection,
            selection_source=request.selection_source,
            target_path=request.target_path,
            model_mode=request.model_mode,
            model_id=request.model_id,
            history=list(request.history),
            options=dict(options),
        )

    def _initial_model_request(self, request: FileTaskRequest) -> FileTaskRequest:
        options = self._sanitize_planner_options(dict(request.options or {}))
        options["planner_policy"] = "native_only"
        options["planner_runtime_reason"] = "file_task_native_only"
        return self._clone_request_with_options(request, options)

    def _request_after_execution_brief(
        self,
        original_request: FileTaskRequest,
        current_request: FileTaskRequest,
        brief: FileTaskExecutionBrief,
    ) -> FileTaskRequest:
        options = self._sanitize_planner_options(dict(current_request.options or {}))
        options["planner_policy"] = "native_only"
        options["planner_runtime_reason"] = "file_task_native_only"
        return self._clone_request_with_options(original_request, options)

    def _build_plan(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        write_intent: bool,
        output_mode: str,
        known_tool_gap: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        recipe_match = select_task_recipe(request, files, write_intent=write_intent)
        if recipe_match and recipe_match.recipe.plan_steps:
            return [dict(step) for step in recipe_match.recipe.plan_steps]
        context_parts = []
        if files:
            context_parts.append(f"{len(files)} 个文件")
        if request.selection:
            context_parts.append("1 段选区")
        context_detail = "和".join(context_parts)
        steps = [
            {
                "id": "context",
                "title": "读取显式上下文",
                "description": (
                    f"读取 {context_detail}，并保留来源引用。"
                    if context_detail
                    else "检查是否有选区、附件或明确当前文件。"
                ),
            },
            {
                "id": "execute",
                "title": "执行任务",
                "description": self._execute_plan_description(
                    write_intent, output_mode, known_tool_gap
                ),
            },
            {
                "id": "check",
                "title": "核验结果",
                "description": "输出检查结论和剩余动作，避免静默失败。",
            },
        ]
        return steps

    def _resolve_intent_plan(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        *,
        known_tool_gap: Optional[Dict[str, Any]] = None,
        classification: Optional[FileTaskClassification] = None,
        intent_plan: Optional[FileTaskIntentPlan] = None,
    ) -> FileTaskIntentPlan:
        resolved_classification = classification or self._classify_request(
            request, files, known_tool_gap
        )
        if isinstance(intent_plan, FileTaskIntentPlan):
            planned = intent_plan
        else:
            try:
                planned = self._intent_planner.plan(
                    request,
                    files,
                    resolved_classification,
                    known_tool_gap=known_tool_gap,
                )
            except Exception as exc:
                logger.warning("[FileTaskRuntime] intent planner failed: %s", exc)
                planned = self._fallback_intent_plan(
                    request, files, resolved_classification, known_tool_gap
                )
            if not isinstance(planned, FileTaskIntentPlan):
                planned = self._fallback_intent_plan(
                    request, files, resolved_classification, known_tool_gap
                )

        planned.intent_type = (
            str(
                planned.intent_type or resolved_classification.task_family or "analyze"
            ).strip()
            or "analyze"
        )
        planned.output_mode = (
            str(resolved_classification.output_mode or planned.output_mode or "answer")
            .strip()
            .lower()
            or "answer"
        )
        planned.confidence = float(
            resolved_classification.confidence
            if resolved_classification.confidence is not None
            else planned.confidence or 0.0
        )
        planned.write_intent = bool(resolved_classification.write_intent)
        if not str(planned.goal_statement or "").strip():
            planned.goal_statement = self._fallback_intent_goal_statement(
                request, resolved_classification, known_tool_gap
            )
        if not planned.dynamic_steps:
            planned.dynamic_steps = self._build_plan(
                request,
                files,
                resolved_classification.write_intent,
                planned.output_mode,
                known_tool_gap,
            )
        if not planned.reason_codes:
            planned.reason_codes = [
                item for item in resolved_classification.reason_codes if item
            ]
        return planned

    def _fallback_intent_plan(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        classification: FileTaskClassification,
        known_tool_gap: Optional[Dict[str, Any]] = None,
    ) -> FileTaskIntentPlan:
        output_mode = (
            str(classification.output_mode or "answer").strip().lower() or "answer"
        )
        recommended_strategy = self._fallback_intent_strategy(
            classification, output_mode, known_tool_gap
        )
        can_apply = output_mode in {
            "write",
            "hybrid",
        } and self._fallback_intent_has_apply_target(request, files)
        requires_confirmation = output_mode == "hybrid"
        reason_codes = [item for item in classification.reason_codes if item]
        reason_codes.extend(
            [
                "intent_plan:fallback",
                f"intent_type:{classification.task_family or 'analyze'}",
                f"strategy:{recommended_strategy}",
            ]
        )
        if can_apply:
            reason_codes.append("can_apply")
        if requires_confirmation:
            reason_codes.append("requires_confirmation")
        return FileTaskIntentPlan(
            intent_type=classification.task_family or "analyze",
            goal_statement=self._fallback_intent_goal_statement(
                request, classification, known_tool_gap
            ),
            output_mode=output_mode,
            confidence=float(classification.confidence or 0.0),
            write_intent=bool(classification.write_intent),
            can_apply=can_apply,
            requires_confirmation=requires_confirmation,
            recommended_strategy=recommended_strategy,
            dynamic_steps=self._build_plan(
                request, files, classification.write_intent, output_mode, known_tool_gap
            ),
            reason_codes=reason_codes,
        )

    def _fallback_intent_goal_statement(
        self,
        request: FileTaskRequest,
        classification: FileTaskClassification,
        known_tool_gap: Optional[Dict[str, Any]] = None,
    ) -> str:
        task_text = _preview(request.task, 180) or "当前文件任务"
        output_mode = (
            str(classification.output_mode or "answer").strip().lower() or "answer"
        )
        if known_tool_gap:
            return f"识别缺失原生能力并输出可落地工具设计：{task_text}"
        if classification.request_kind == "resume":
            return f"延续上一轮待确认的文件任务：{task_text}"
        if output_mode == "write":
            return f"完成真实文件修改并交付结果：{task_text}"
        if output_mode == "hybrid":
            return f"先分析并整理可应用建议，再等待确认：{task_text}"
        return f"基于显式上下文给出结论或答复：{task_text}"

    def _fallback_intent_strategy(
        self,
        classification: FileTaskClassification,
        output_mode: str,
        known_tool_gap: Optional[Dict[str, Any]] = None,
    ) -> str:
        if known_tool_gap:
            return "design_new_tool"
        if classification.request_kind == "resume":
            return "resume_previous_plan"
        if output_mode == "write":
            return "write_through"
        if output_mode == "hybrid":
            return "analyze_then_confirm"
        return "answer_only"

    def _fallback_intent_has_apply_target(
        self, request: FileTaskRequest, files: List[FileTaskFile]
    ) -> bool:
        if str(request.target_path or "").strip():
            return True
        if request.selection:
            return True
        return any(
            file_info.target or file_info.path or file_info.name for file_info in files
        )

    def _execute_plan_description(
        self,
        write_intent: bool,
        output_mode: str,
        known_tool_gap: Optional[Dict[str, Any]],
    ) -> str:
        if known_tool_gap:
            capability = str(known_tool_gap.get("missing_capability") or "缺失能力").strip()
            return f"当前任务触发 Koto 原生能力缺口：{capability}；模型需要产出 {TOOL_DESIGN_PROTOCOL} 工具规格，不调用未注册工具。"
        if write_intent:
            return "模型在 Koto allowlist 工具目录内规划并执行，写入后产生 file.changed 事件。"
        if output_mode == "hybrid":
            return "模型先读取文件并给出可应用的分析建议；当前轮不默认直接写入原文件。"
        return "模型可读取文件、调用分析工具并生成可审计答复。"

    def _output_mode_label(self, output_mode: str) -> str:
        normalized = str(output_mode or "").strip().lower()
        if normalized == "write":
            return "写入文件"
        if normalized == "hybrid":
            return "先分析后决定"
        return "只给答案"

    def _output_mode_guidance(self, classification: FileTaskClassification) -> str:
        output_mode = str(classification.output_mode or "answer").strip().lower()
        label = self._output_mode_label(output_mode)
        if output_mode == "write":
            return (
                f"当前任务反馈模式：{label}。\n"
                "本轮目标是完成真实文件修改；除非进入等待确认状态，否则不要只给建议或总结就结束。\n"
                "如果没有产生真实 file.changed，就不能把任务说成已完成。\n"
            )
        if output_mode == "hybrid":
            return (
                f"当前任务反馈模式：{label}。\n"
                "本轮先基于显式上下文给出分析、问题清单、修改方向或可应用方案。\n"
                "除非用户这轮已经明确要求直接应用到文件，否则不要直接调用写入工具，也不要声称文件已经更新。\n"
                "如果需要后续落盘，应先把建议说清楚，再等待用户继续要求应用。\n"
            )
        return (
            f"当前任务反馈模式：{label}。\n"
            "本轮默认只返回分析、总结、解释或结论。\n"
            "不要调用写入工具，不要伪造 file.changed，也不要把结果描述成已经写入文件。\n"
            "只有当用户明确要求把结果写入文件时，才改走写回路径。\n"
        )

    def _intent_plan_guidance(self, intent_plan: FileTaskIntentPlan) -> str:
        lines = ["高阶意图规划："]
        goal_statement = str(intent_plan.goal_statement or "").strip()
        if goal_statement:
            lines.append(f"- 目标：{goal_statement}")
        lines.append(
            f"- 策略：{str(intent_plan.recommended_strategy or 'answer_only').strip() or 'answer_only'}"
        )
        lines.append(f"- 可应用：{'是' if intent_plan.can_apply else '否'}")
        lines.append(f"- 写回前需要确认：{'是' if intent_plan.requires_confirmation else '否'}")
        if intent_plan.dynamic_steps:
            lines.append("- 计划步骤：")
            for index, step in enumerate(intent_plan.dynamic_steps[:8], start=1):
                if not isinstance(step, dict):
                    continue
                title = str(
                    step.get("title") or step.get("id") or f"步骤 {index}"
                ).strip()
                description = str(step.get("description") or "").strip()
                lines.append(
                    f"  {index}. {title}" + (f"：{description}" if description else "")
                )
        if intent_plan.write_intent:
            lines.append("- 监管约束：写入型任务必须产生真实 file.changed；分步确认任务必须先完成本步骤写入，再进入等待用户继续。")
        return "\n".join(lines) + "\n"

    def _execution_brief_schema(self) -> Dict[str, Any]:
        return {
            "execution_brief": {
                "summary": "一句中文概述当前准备怎么处理",
                "objective": "本轮要完成的目标",
                "steps": [{"title": "步骤标题", "description": "准备做什么"}],
                "planned_tools": ["tool_name"],
                "read_targets": ["会读取的文件或对象"],
                "write_targets": ["会写入的文件或对象"],
                "verification": "准备如何验证结果",
            }
        }

    def _normalize_execution_brief(
        self, value: Any
    ) -> Optional[FileTaskExecutionBrief]:
        candidate = value
        if isinstance(candidate, dict) and isinstance(
            candidate.get("execution_brief"), dict
        ):
            candidate = candidate.get("execution_brief")
        if not isinstance(candidate, dict):
            return None
        brief = FileTaskExecutionBrief.from_mapping(candidate)
        if not any(
            (
                brief.summary,
                brief.objective,
                brief.steps,
                brief.planned_tools,
                brief.read_targets,
                brief.write_targets,
                brief.verification,
            )
        ):
            return None
        return brief

    def _looks_like_brief_only_content(self, content_text: str) -> bool:
        text = str(content_text or "").strip()
        if not text:
            return False
        if text.startswith(("{", "[")):
            return True
        if text.startswith("```") and extract_first_json_value(text) is not None:
            return True
        return False

    def _extract_execution_brief(
        self,
        response: Any,
        content_text: str,
    ) -> tuple[Optional[FileTaskExecutionBrief], str]:
        candidate: Any = None
        if isinstance(response, dict):
            if isinstance(response.get("execution_brief"), dict):
                candidate = response.get("execution_brief")
            elif content_text:
                candidate = extract_first_json_value(content_text)
        elif content_text:
            candidate = extract_first_json_value(content_text)

        brief = self._normalize_execution_brief(candidate)
        if not brief:
            return None, content_text

        cleaned_content = content_text.strip()
        if self._looks_like_brief_only_content(cleaned_content):
            cleaned_content = brief.summary or brief.objective or ""
        return brief, cleaned_content

    def _extract_execution_brief_tool_call(
        self,
        tool_calls: List[Dict[str, Any]],
    ) -> tuple[Optional[FileTaskExecutionBrief], List[Dict[str, Any]]]:
        if not tool_calls:
            return None, []

        brief: Optional[FileTaskExecutionBrief] = None
        remaining: List[Dict[str, Any]] = []
        for tool_call in tool_calls:
            tool_name = str((tool_call or {}).get("name") or "").strip()
            if tool_name != "execution_brief":
                remaining.append(tool_call)
                continue

            candidate = (tool_call or {}).get("args") or {}
            parsed = self._normalize_execution_brief(candidate)
            if parsed and brief is None:
                brief = parsed

        return brief, remaining

    def _execution_brief_continue_message(
        self,
        request: FileTaskRequest,
        brief: FileTaskExecutionBrief,
    ) -> str:
        summary = brief.summary or brief.objective or "已完成任务分析。"
        lines = [
            f"已收到 execution_brief：{summary}",
            "下一轮请在白盒任务骨架内直接调用需要的 Koto 工具继续执行，不要重复输出同一份 brief。",
        ]
        if request.target_path:
            lines.append(f"当前目标文件是：{request.target_path}。")
        return " ".join(lines)

    def _execution_plan_continue_message(
        self,
        request: FileTaskRequest,
        execution_plan: WhiteboxExecutionPlan,
        recipe_skeleton: Dict[str, Any],
    ) -> str:
        summary = execution_plan.plan_summary or execution_plan.goal or "已完成白盒执行计划。"
        lines = [
            f"已收到 execution_plan：{summary}",
            "现在请按该计划继续调用 Koto allowlist 工具执行；不要重复输出计划，也不要跳过必需写入/核验步骤。",
        ]
        required_operations = (
            recipe_skeleton.get("completion_check", {}).get("required_operations")
            if isinstance(recipe_skeleton.get("completion_check"), dict)
            else []
        )
        if required_operations:
            lines.append(
                "完成检查要求包含："
                + "、".join(
                    str(item) for item in required_operations if str(item or "").strip()
                )
            )
        if request.target_path:
            lines.append(f"目标文件：{request.target_path}。")
        return " ".join(lines)

    def _whitebox_plan_repair_message(
        self,
        gate_payload: Dict[str, Any],
        recipe_skeleton: Dict[str, Any],
    ) -> str:
        lines = [
            "白盒计划审查未通过或不完整，请修复计划后继续执行。",
            "必须遵守 recipe_skeleton 的 required_steps、allowed_tools、success_criteria 和 completion_check。",
        ]
        violations = (
            gate_payload.get("violations")
            if isinstance(gate_payload.get("violations"), list)
            else []
        )
        warnings = (
            gate_payload.get("warnings")
            if isinstance(gate_payload.get("warnings"), list)
            else []
        )
        if violations:
            lines.append("阻断问题：" + "；".join(str(item) for item in violations[:6]))
        if warnings:
            lines.append("需要补强：" + "；".join(str(item) for item in warnings[:6]))
        allowed_tools = (
            recipe_skeleton.get("allowed_tools")
            if isinstance(recipe_skeleton.get("allowed_tools"), list)
            else []
        )
        if allowed_tools:
            lines.append(
                "只能调用这些工具："
                + "、".join(
                    str(item) for item in allowed_tools[:30] if str(item or "").strip()
                )
            )
        return "\n".join(lines)

    def _build_confirmed_plan(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        tool_calls: List[Dict[str, Any]],
        write_intent: bool,
        content_text: str,
    ) -> Dict[str, Any]:
        steps: List[Dict[str, Any]] = []
        seen: set[str] = set()
        has_write_step = False

        for idx, tool_call in enumerate(tool_calls, start=1):
            tool_name = str(tool_call.get("name") or "").strip()
            tool_args = dict(tool_call.get("args") or {})
            if not tool_name:
                continue
            signature = json.dumps(
                {"name": tool_name, "args": tool_args},
                ensure_ascii=False,
                sort_keys=True,
                default=str,
            )
            if signature in seen:
                continue
            seen.add(signature)
            has_write_step = has_write_step or (
                is_write_tool(tool_name) and tool_name != "run_python_code"
            )
            steps.append(
                {
                    "id": f"model_step_{idx}",
                    "tool_name": tool_name,
                    "title": self._tool_plan_title(tool_name),
                    "description": self._tool_plan_description(
                        tool_name, tool_args, files, request
                    ),
                }
            )

        if write_intent and not has_write_step:
            steps.append(self._inferred_write_plan_step(request, files))

        steps.append(
            {
                "id": "verify",
                "title": "核验结果",
                "description": "检查目标文件是否真的更新，并给出最终结论。",
            }
        )

        clean_summary = _preview(content_text, 180) if content_text else "AI 已确认执行方案。"
        return {
            "summary": clean_summary,
            "steps": steps,
            "estimated": True,
            "note": "实际步骤会根据读取结果和工具返回自动微调。",
        }

    def _tool_plan_title(self, tool_name: str) -> str:
        labels = {
            "read_sheet_data": "读取 Excel 表格",
            "inspect_workbook_structure": "检查 Excel 结构",
            "audit_financial_workbook": "审计财务模型",
            "read_docx_content": "读取 Word 内容",
            "parse_file_to_text": "解析文件文本",
            "clear_docx_review_marks": "清除 Word 审阅标记",
            "insert_image_into_docx": "插入 Word 图片",
            "insert_excel_as_docx_table": "写入 Word 表格",
            "write_docx_content": "写入 Word 内容",
            "write_sheet_data": "写入 Excel 单元格",
            "design_pptx_theme_layout": "设计 PPT 主题版式",
            "write_pptx_slides": "更新 PPT 页面",
            "add_pptx_slides": "新增 PPT 页面",
            "create_file": "创建文件",
            "copy_file": "复制文件",
            "read_file_range": "读取文本片段",
            "replace_file_selection": "替换文本选区",
            "compare_files": "对比文件",
            "compare_docx_and_annotate": "对比并标注 Word 差异",
            "extract_to_file": "提取到文件",
            "annotate_file": "添加批注",
            "run_python_code": "运行代码处理",
        }
        return labels.get(tool_name, f"调用工具 {tool_name}")

    def _tool_plan_description(
        self,
        tool_name: str,
        tool_args: Dict[str, Any],
        files: List[FileTaskFile],
        request: FileTaskRequest,
    ) -> str:
        if tool_name == "read_sheet_data":
            source = (
                self._display_path(tool_args.get("path"))
                or self._first_file_name(files, {"xlsx", "xlsm", "csv"})
                or "表格文件"
            )
            sheet = str(tool_args.get("sheet_name") or "").strip()
            rows = str(tool_args.get("max_rows") or "").strip()
            suffix = f"，工作表：{sheet}" if sheet else ""
            rows_text = f"，最多 {rows} 行" if rows else ""
            return f"读取 {source} 的表格数据{suffix}{rows_text}。"
        if tool_name == "inspect_workbook_structure":
            source = (
                self._display_path(tool_args.get("path"))
                or self._first_file_name(files, {"xlsx", "xlsm"})
                or "Excel 文件"
            )
            return f"检查 {source} 的工作表结构、公式分布和外部链接依赖。"
        if tool_name == "audit_financial_workbook":
            source = (
                self._display_path(tool_args.get("path"))
                or self._first_file_name(files, {"xlsx", "xlsm"})
                or "财务模型"
            )
            return f"审计 {source} 的三表完整性、外部依赖和关键年份序列红旗。"
        if tool_name == "insert_excel_as_docx_table":
            source = (
                self._display_path(tool_args.get("source_path"))
                or self._first_file_name(files, {"xlsx", "xlsm", "csv"})
                or "表格文件"
            )
            target = (
                self._display_path(tool_args.get("target_path"))
                or request.target_path
                or self._first_file_name(files, {"docx"}, target=True)
                or "Word 文档"
            )
            table_title = str(tool_args.get("table_title") or "").strip()
            title_text = f"，表题：{table_title}" if table_title else ""
            return f"把 {source} 的数据作为真实 Word 表格插入 {self._display_path(target) or target}{title_text}。"
        if tool_name == "insert_image_into_docx":
            target = (
                self._display_path(tool_args.get("path"))
                or request.target_path
                or self._first_file_name(files, {"docx"}, target=True)
                or "Word 文档"
            )
            image_path = (
                self._display_path(tool_args.get("image_path"))
                or str(tool_args.get("image_path") or "图片文件").strip()
                or "图片文件"
            )
            title = str(tool_args.get("title") or "").strip()
            title_text = f"，图题：{title}" if title else ""
            return f"把 {image_path} 作为真实图片插入 {self._display_path(target) or target}{title_text}。"
        if tool_name == "write_docx_content":
            target = (
                self._display_path(tool_args.get("path"))
                or request.target_path
                or self._first_file_name(files, {"docx"}, target=True)
                or "Word 文档"
            )
            return f"把生成后的段落写入 {self._display_path(target) or target}。"
        if tool_name == "clear_docx_review_marks":
            target = (
                self._display_path(tool_args.get("path"))
                or request.target_path
                or self._first_file_name(files, {"docx"}, target=True)
                or "Word 文档"
            )
            scope = (
                str(tool_args.get("scope") or "comments").strip().lower() or "comments"
            )
            if scope == "all":
                return f"清除 {self._display_path(target) or target} 中的批注并接受修订。"
            if scope == "revisions":
                return f"接受并清除 {self._display_path(target) or target} 中的修订标记。"
            return f"清除 {self._display_path(target) or target} 中的全部批注。"
        if tool_name == "write_sheet_data":
            target = (
                self._display_path(tool_args.get("path"))
                or request.target_path
                or self._first_file_name(files, {"xlsx", "xlsm"}, target=True)
                or "Excel 文件"
            )
            sheet = str(tool_args.get("sheet_name") or "").strip()
            sheet_text = f"，工作表：{sheet}" if sheet else ""
            return f"把结构化更新写入 {self._display_path(target) or target}{sheet_text}。"
        if tool_name == "annotate_file":
            target = (
                self._display_path(tool_args.get("path"))
                or request.target_path
                or self._first_file_name(
                    files, {"docx", "pdf", "txt", "md"}, target=True
                )
                or "目标文件"
            )
            requirement = str(tool_args.get("requirement") or "").strip()
            if requirement:
                return f"按要求为 {self._display_path(target) or target} 生成并写回批注：{_compact_line(requirement, 90)}。"
            return f"把结构化批注写入 {self._display_path(target) or target}。"
        if tool_name in {
            "design_pptx_theme_layout",
            "write_pptx_slides",
            "add_pptx_slides",
        }:
            target = (
                self._display_path(tool_args.get("path"))
                or request.target_path
                or self._first_file_name(files, {"pptx"}, target=True)
                or "PPT 文件"
            )
            if tool_name == "design_pptx_theme_layout":
                style_brief = str(tool_args.get("style_brief") or "").strip()
                style_text = f"，风格要求：{style_brief}" if style_brief else ""
                return f"为 {self._display_path(target) or target} 套用统一主题、字体、配色和安全版式{style_text}。"
            action = "新增" if tool_name == "add_pptx_slides" else "更新"
            return f"在 {self._display_path(target) or target} 中{action}幻灯片内容。"
        if tool_name == "parse_file_to_text":
            source = (
                self._display_path(tool_args.get("path"))
                or self._first_file_name(files, set())
                or "文件"
            )
            return f"解析 {source} 的文本内容，供后续分析使用。"
        if tool_name == "read_file_range":
            source = (
                self._display_path(tool_args.get("path"))
                or self._first_file_name(
                    files, {"txt", "md", "csv", "json", "py", "js", "html", "css"}
                )
                or "文本文件"
            )
            start = str(tool_args.get("start_line") or "1").strip()
            end = str(tool_args.get("end_line") or "").strip()
            window = f"第 {start} 到 {end} 行" if end else f"从第 {start} 行开始"
            return f"读取 {source} 的{window}，供后续分析使用。"
        if tool_name == "replace_file_selection":
            target = (
                self._display_path(tool_args.get("path"))
                or request.target_path
                or self._first_file_name(
                    files,
                    {"txt", "md", "csv", "json", "py", "js", "html", "css"},
                    target=True,
                )
                or "文本文件"
            )
            return f"把改写后的选区内容写回 {self._display_path(target) or target}。"
        if tool_name == "compare_files":
            raw_paths = str(tool_args.get("file_paths") or "").strip()
            aspect = str(tool_args.get("aspect") or "content").strip()
            return f"对比文件{f'：{raw_paths}' if raw_paths else ''}，比较维度：{aspect}。"
        if tool_name == "run_python_code":
            return "在沙盒中运行代码处理数据，必要时生成图表或中间文件。"
        target = self._display_path(
            tool_args.get("path")
            or tool_args.get("target_path")
            or tool_args.get("destination")
        )
        return f"执行 {tool_name}{f'，目标：{target}' if target else ''}。"

    def _inferred_write_plan_step(
        self, request: FileTaskRequest, files: List[FileTaskFile]
    ) -> Dict[str, Any]:
        source = self._first_file_name(files, {"xlsx", "xlsm", "csv"})
        docx_target = (
            self._typed_target_display_path(request, {"docx", "doc"})
            or self._first_file_name(files, {"docx"}, target=True)
            or self._first_file_name(files, {"docx"})
        )
        pptx_target = (
            self._typed_target_display_path(request, {"pptx", "ppt"})
            or self._first_file_name(files, {"pptx"}, target=True)
            or self._first_file_name(files, {"pptx"})
        )
        text_target = self._typed_target_display_path(
            request, {"txt", "md", "csv", "json", "py", "js", "html", "css"}
        ) or self._first_file_name(
            files, {"txt", "md", "csv", "json", "py", "js", "html", "css"}, target=True
        )
        task_lower = (request.task or "").lower()
        if (
            source
            and docx_target
            and self._looks_like_financial_xlsx_docx_chart_report_task(request, files)
        ):
            return {
                "id": "inferred_write",
                "title": "写入问题和图表",
                "description": f"读取完成后，先生成真实财务图表图片并整理问题清单，再写入 {docx_target}。",
            }
        if source and docx_target:
            return {
                "id": "inferred_write",
                "title": "写入 Word 表格",
                "description": f"读取完成后，把 {source} 的表格数据写入 {docx_target}。",
            }
        if text_target:
            return {
                "id": "inferred_write",
                "title": "写回文本文件",
                "description": f"读取完成后，把处理结果写回 {text_target}。",
            }
        if pptx_target or "ppt" in task_lower or "幻灯片" in task_lower:
            if any(
                word in task_lower
                for word in (
                    "风格",
                    "主题",
                    "版式",
                    "美化",
                    "排版",
                    "配色",
                    "视觉",
                    "theme",
                    "layout",
                    "design",
                )
            ):
                return {
                    "id": "inferred_write",
                    "title": "设计 PPT 主题版式",
                    "description": f"读取完成后，为 {pptx_target or '目标 PPT'} 套用统一主题、字体、配色和安全版式。",
                }
            return {
                "id": "inferred_write",
                "title": "写入 PPT 内容",
                "description": f"读取完成后，把整理结果写入 {pptx_target or '目标 PPT'}。",
            }
        target = self._display_path(request.target_path) or next(
            (
                self._display_path(file_info.path)
                for file_info in files
                if file_info.target and file_info.path
            ),
            "目标文件",
        )
        return {
            "id": "inferred_write",
            "title": "写入目标文件",
            "description": f"读取完成后，把处理结果写入 {target}。",
        }

    def _typed_target_display_path(
        self, request: FileTaskRequest, types: set[str]
    ) -> str:
        raw = str(request.target_path or "").strip()
        suffix = Path(raw).suffix.lstrip(".").lower()
        if raw and suffix in types:
            return self._display_path(raw)
        return ""

    def _first_file_name(
        self, files: List[FileTaskFile], types: set[str], *, target: bool = False
    ) -> str:
        for file_info in files:
            file_type = (
                file_info.type
                or Path(file_info.path or file_info.name).suffix.lstrip(".")
            ).lower()
            if target and not file_info.target:
                continue
            if types and file_type not in types:
                continue
            return file_info.name or self._display_path(file_info.path)
        return ""

    def _first_context_file(
        self, files: List[FileTaskFile], types: set[str], *, target: bool = False
    ) -> Optional[FileTaskFile]:
        for file_info in files:
            file_type = (
                file_info.type
                or Path(file_info.path or file_info.name).suffix.lstrip(".")
            ).lower()
            if target and not file_info.target:
                continue
            if types and file_type not in types:
                continue
            return file_info
        return None

    def _single_context_file(
        self, files: List[FileTaskFile], types: set[str]
    ) -> Optional[FileTaskFile]:
        matches: List[FileTaskFile] = []
        for file_info in files:
            file_type = (
                file_info.type
                or Path(file_info.path or file_info.name).suffix.lstrip(".")
            ).lower()
            if types and file_type not in types:
                continue
            matches.append(file_info)
        return matches[0] if len(matches) == 1 else None

    def _context_files_by_type(
        self, files: List[FileTaskFile], types: set[str]
    ) -> List[FileTaskFile]:
        matches: List[FileTaskFile] = []
        for file_info in files:
            file_type = (
                file_info.type
                or Path(file_info.path or file_info.name).suffix.lstrip(".")
            ).lower()
            if types and file_type not in types:
                continue
            matches.append(file_info)
        return matches

    def _display_path(self, value: Any) -> str:
        text = str(value or "").strip()
        if not text:
            return ""
        return re.split(r"[\\/]+", text)[-1] or text

    def _repair_tool_args_for_context(
        self,
        tool_name: str,
        tool_args: Dict[str, Any],
        request: FileTaskRequest,
        files: List[FileTaskFile],
    ) -> Dict[str, Any]:
        args = dict(tool_args or {})
        if (
            tool_name
            in {
                "write_docx_content",
                "insert_image_into_docx",
                "clear_docx_review_marks",
            }
            and not str(args.get("path") or "").strip()
        ):
            target = self._single_target_path_for_types(request, files, {"docx", "doc"})
            if target:
                args["path"] = target
        if (
            tool_name == "design_pptx_theme_layout"
            and not str(args.get("path") or "").strip()
        ):
            target = self._single_target_path_for_types(request, files, {"pptx", "ppt"})
            if target:
                args["path"] = target
        if tool_name == "write_sheet_data" and not str(args.get("path") or "").strip():
            target = self._single_target_path_for_types(
                request, files, {"xlsx", "xlsm"}
            )
            if target:
                args["path"] = target
        if tool_name == "insert_excel_as_docx_table":
            if not str(args.get("target_path") or "").strip():
                target = self._single_target_path_for_types(
                    request, files, {"docx", "doc"}
                )
                if target:
                    args["target_path"] = target
            if not str(args.get("source_path") or "").strip():
                source = self._single_source_path_for_types(
                    request, files, {"xlsx", "xlsm"}
                )
                if source:
                    args["source_path"] = source
        if tool_name == "write_docx_content" and "paragraphs" not in args:
            for key in ("content", "text", "markdown", "body"):
                value = args.get(key)
                if str(value or "").strip():
                    args["paragraphs"] = value
                    break
        return args

    def _single_target_path_for_types(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        file_types: set[str],
    ) -> str:
        target_path = str(request.target_path or "").strip()
        if target_path:
            suffix = Path(target_path).suffix.lower().lstrip(".")
            if not suffix or suffix in file_types:
                return target_path
        candidates: List[str] = []
        for file_info in files or []:
            if not getattr(file_info, "target", False):
                continue
            path = str(file_info.path or file_info.name or "").strip()
            suffix = str(file_info.type or "").strip().lower().lstrip(".") or Path(
                path
            ).suffix.lower().lstrip(".")
            if path and suffix in file_types:
                candidates.append(path)
        unique = list(dict.fromkeys(candidates))
        return unique[0] if len(unique) == 1 else ""

    def _single_source_path_for_types(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        file_types: set[str],
    ) -> str:
        candidates: List[str] = []
        for file_info in files or []:
            if getattr(file_info, "target", False):
                continue
            path = str(file_info.path or file_info.name or "").strip()
            suffix = str(file_info.type or "").strip().lower().lstrip(".") or Path(
                path
            ).suffix.lower().lstrip(".")
            if path and suffix in file_types:
                candidates.append(path)
        unique = list(dict.fromkeys(candidates))
        return unique[0] if len(unique) == 1 else ""

    def _resolve_task_file_path(self, path: Any) -> str:
        raw = str(path or "").strip()
        if not raw:
            return ""
        if os.path.isabs(raw) and os.path.exists(raw):
            return os.path.normpath(raw)
        try:
            from app.core.agent import task_tools

            resolved = task_tools._resolve_path(raw)  # type: ignore[attr-defined]
            if resolved and os.path.exists(resolved):
                return os.path.normpath(resolved)
        except Exception:
            pass
        return ""

    def _plan_summary(
        self, request: FileTaskRequest, files: List[FileTaskFile], write_intent: bool
    ) -> str:
        target = request.target_path or next((f.path for f in files if f.target), "")
        has_selection = bool(request.selection)
        if write_intent and target:
            suffix = "，并引用 1 段选区" if has_selection else ""
            return f"准备更新 {Path(target).name}{suffix}。"
        if files and has_selection:
            return f"准备处理 {len(files)} 个文件和 1 段选区。"
        if files:
            return f"准备处理 {len(files)} 个文件。"
        if has_selection:
            return "准备处理 1 段选区。"
        return "准备处理当前任务。"

    def _write_stepwise_pdf_docx_native(
        self,
        ledger: FileTaskLedger,
        request: FileTaskRequest,
        executor: ToolExecutor,
        snippets: List[Dict[str, Any]],
        files: List[FileTaskFile],
        recipe_skeleton: Dict[str, Any],
        step_id: str,
        *,
        reason: str,
        fallback: bool,
        model_unavailable: bool,
    ):
        if not _looks_like_windowed_pdf_task(request, recipe_skeleton):
            return None
        pdf_snippet = next(
            (
                item
                for item in snippets
                if str(item.get("source") or item.get("path") or "")
                .lower()
                .endswith(".pdf")
                or str(item.get("path") or "").lower().endswith(".pdf")
            ),
            None,
        )
        if not pdf_snippet:
            return None
        text_quality = _pdf_text_quality(
            pdf_snippet.get("_raw_text") or pdf_snippet.get("preview") or ""
        )
        if not text_quality.get("usable"):
            reason_text = str(text_quality.get("reason") or "low_quality_pdf_text")
            yield ledger.event(
                "tool.finished",
                {
                    "tool_name": "supervisor_guard",
                    "success": False,
                    "blocked": True,
                    "native_stepwise": True,
                    "result_preview": f"监管层阻止写入：当前 PDF 页窗文本质量不足（{reason_text}），不能据此生成分步 DOCX 摘要。",
                },
                step_id=step_id,
            )
            return None
        target_path = self._stepwise_docx_target_path(request, files)
        if not target_path:
            return None

        paragraphs = self._stepwise_pdf_fallback_paragraphs(
            request, pdf_snippet, RuntimeError(reason)
        )
        tool_args = {
            "path": target_path,
            "paragraphs": json.dumps(paragraphs, ensure_ascii=False),
        }
        started_payload = {
            "tool_name": "write_docx_content",
            "tool_args": tool_args,
            "native_stepwise": True,
            "reason": reason,
        }
        if fallback:
            started_payload["fallback"] = True
        yield ledger.event("tool.started", started_payload, step_id=step_id)
        try:
            result = executor("write_docx_content", tool_args)
            success = not _is_error_result(result)
        except Exception as write_exc:
            result = f"Error: {write_exc}"
            success = False
            logger.warning(
                "[FileTaskRuntime] stepwise PDF native write failed: %s", write_exc
            )

        finished_payload = {
            "tool_name": "write_docx_content",
            "success": success,
            "native_stepwise": True,
            "model_unavailable": bool(model_unavailable),
            "result_preview": tool_result_preview("write_docx_content", result, 1200),
        }
        if fallback:
            finished_payload["fallback"] = True
        yield ledger.event("tool.finished", finished_payload, step_id=step_id)
        if not success:
            return None
        changes = self._extract_file_changes("write_docx_content", tool_args, result)
        return changes[0] if changes else None

    def _stepwise_docx_target_path(
        self, request: FileTaskRequest, files: List[FileTaskFile]
    ) -> str:
        raw_target = str(request.target_path or "").strip()
        if raw_target:
            return (
                raw_target
                if os.path.isabs(raw_target)
                else str(Path(raw_target).resolve())
            )
        docx_file = self._first_context_file(
            files, {"docx"}, target=True
        ) or self._first_context_file(files, {"docx"})
        if docx_file and docx_file.path:
            return docx_file.path
        pdf_file = self._first_context_file(files, {"pdf"})
        if pdf_file and pdf_file.path:
            pdf_path = Path(pdf_file.path)
            return str(pdf_path.with_name(f"{pdf_path.stem}_分步总结.docx"))
        return ""

    def _latest_pdf_snippet_quality(
        self, snippets: List[Dict[str, Any]]
    ) -> Dict[str, Any]:
        pdf_snippets = [
            item
            for item in snippets
            if isinstance(item, dict)
            and (
                str(item.get("source") or item.get("path") or "")
                .lower()
                .endswith(".pdf")
                or str(item.get("path") or "").lower().endswith(".pdf")
            )
        ]
        if not pdf_snippets:
            return {
                "usable": False,
                "reason": "missing_pdf_context",
                "char_count": 0,
                "unique_chars": 0,
            }
        text = str(
            pdf_snippets[-1].get("_raw_text") or pdf_snippets[-1].get("preview") or ""
        )
        return _pdf_text_quality(text)

    def _tool_args_docx_paragraph_text(self, tool_args: Dict[str, Any]) -> str:
        raw_paragraphs = tool_args.get("paragraphs")
        items: Any = []
        if isinstance(raw_paragraphs, str) and raw_paragraphs.strip():
            try:
                items = json.loads(raw_paragraphs)
            except Exception:
                items = []
        elif isinstance(raw_paragraphs, list):
            items = raw_paragraphs
        texts: List[str] = []
        if isinstance(items, list):
            for item in items:
                if isinstance(item, dict):
                    texts.append(str(item.get("text") or ""))
                elif item is not None:
                    texts.append(str(item))
        content = str(tool_args.get("content") or "").strip()
        if content:
            texts.append(content)
        return "\n".join(text for text in texts if text.strip())

    def _stepwise_docx_write_block_message(
        self,
        request: FileTaskRequest,
        snippets: List[Dict[str, Any]],
        recipe_skeleton: Dict[str, Any],
        tool_name: str,
        tool_args: Dict[str, Any],
    ) -> str:
        if not _looks_like_windowed_pdf_task(request, recipe_skeleton):
            return ""
        if tool_name == "create_file":
            target = str(tool_args.get("path") or "").strip().lower()
            if not target.endswith(".docx"):
                return ""
        elif tool_name != "write_docx_content":
            return ""
        quality = self._latest_pdf_snippet_quality(snippets)
        if not quality.get("usable"):
            return (
                "监管层阻止写入：当前 PDF 页窗的可提取文本质量不足，不能把水印、乱码或空内容写成总结。"
                f" 质量原因：{quality.get('reason') or 'unknown'}；"
                f"可用字符数：{quality.get('alpha_num_chars') or quality.get('char_count') or 0}；"
                f"唯一字符数：{quality.get('unique_chars') or 0}。"
                " 下一轮请改用新的 start_page/end_page 读取后续页窗；如果连续页窗仍不可读，应停止写入并提示需要 OCR/视觉解析。"
            )

        text = self._tool_args_docx_paragraph_text(tool_args)
        if tool_name != "create_file" and re.search(
            r"^\s*#{1,6}\s+", text, re.MULTILINE
        ):
            return (
                "监管层阻止写入：write_docx_content 的 paragraphs 不能包含 Markdown 标题符号 #。"
                " 请使用 paragraph.style='Heading 1' 这类 Word 段落样式。"
            )
        progress_patterns = (
            r"^\s*步骤\s*\d+\s*[：:]",
            r"当前进度\s*[：:]",
            r"下一步(?:计划|继续|处理)",
            r"等待(?:用户|我|确认|继续)",
            r"请回复\s*[\"“]?继续",
            r"当前步骤已(?:成功)?(?:完成|写入)",
            r"状态\s*[：:]",
            r"file\.changed",
            r"目标\s*DOCX\s*文件已成功更新",
        )
        if any(
            re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            for pattern in progress_patterns
        ):
            return (
                "监管层阻止写入：DOCX 正文不能包含任务进度、等待确认、下一步计划或 file.changed 这类前端提示。"
                " 请重写 paragraphs：只写当前页窗的实质内容摘要、关键发现、证据/主题；"
                "页码只可作为“来源页码：第 x-y 页”这样的简短来源说明，不要写等待继续或下一步计划。"
            )
        quality_block = self._stepwise_docx_content_quality_block_message(
            snippets, text
        )
        if quality_block:
            return quality_block
        return ""

    def _stepwise_docx_content_quality_block_message(
        self, snippets: List[Dict[str, Any]], text: str
    ) -> str:
        latest_pdf_snippet = next(
            (
                item
                for item in reversed(snippets or [])
                if isinstance(item, dict)
                and (
                    str(item.get("source") or item.get("path") or "")
                    .lower()
                    .endswith(".pdf")
                    or str(item.get("path") or "").lower().endswith(".pdf")
                )
            ),
            {},
        )
        expected_start = int(latest_pdf_snippet.get("start_page") or 0)
        expected_end = int(latest_pdf_snippet.get("end_page") or 0)
        expected_range = (
            (expected_start, expected_end) if expected_start and expected_end else None
        )
        body = str(text or "").strip()
        if not body:
            return "监管层阻止写入：当前分步 DOCX 正文为空。请写入当前页窗的摘要、关键发现、结构/内容线索和来源页码。"
        if re.search(r"(^|\s)(?:\*\*[^*\n]+\*\*|__[^_\n]+__|[-*_]{3,})(\s|$)", body):
            return (
                "监管层阻止写入：分步 DOCX 正文不能包含 Markdown 标记（如 **加粗**、---）。"
                " 请用 Word 段落样式和纯文本标签写入。"
            )

        combined_label_patterns = (
            r"文档识别\s*/\s*核心要点",
            r"段落主题\s*/\s*关键发现",
            r"内容线索\s*/\s*案例线索",
        )
        if any(re.search(pattern, body) for pattern in combined_label_patterns):
            return (
                "监管层阻止写入：分步 DOCX 正文不能使用“文档识别/核心要点”这类合并标签。"
                " 请改用固定独立标签：文档识别、段落主题、结构线索、内容线索、来源页码。"
            )

        section_ranges: List[tuple[int, int]] = []
        declared_page_ranges: List[tuple[int, int]] = []
        for line in body.splitlines():
            match = re.match(
                r"^\s*【?\s*第\s*(\d+)\s*[-－—~至]\s*(\d+)\s*页[^。\n]{0,40}(?:】|[:：])",
                line,
            )
            if match:
                section_ranges.append((int(match.group(1)), int(match.group(2))))
            if re.search(r"(?:当前页窗摘要|来源页码)", line):
                for range_match in re.finditer(
                    r"第\s*(\d+)\s*[-－—~至]\s*(\d+)\s*页",
                    line,
                ):
                    declared_page_ranges.append(
                        (int(range_match.group(1)), int(range_match.group(2)))
                    )
        unique_section_ranges = list(dict.fromkeys(section_ranges))
        unique_declared_ranges = list(dict.fromkeys(declared_page_ranges))
        if len(unique_section_ranges) > 1:
            ranges = "、".join(
                f"第 {start}-{end} 页" for start, end in unique_section_ranges[:4]
            )
            return (
                "监管层阻止写入：单个分步窗口的 DOCX 正文不能同时覆盖多个页窗标题。"
                f" 检测到：{ranges}。请只写当前页窗内容，上一页窗内容不要重复写入。"
            )
        if section_ranges and section_ranges.count(section_ranges[0]) > 1:
            start, end = section_ranges[0]
            return f"监管层阻止写入：第 {start}-{end} 页在本次写入中出现了重复小节标题。" " 请合并为一个小节，删除重复标题和重复要点。"
        if (
            expected_range
            and unique_section_ranges
            and unique_section_ranges[0] != expected_range
        ):
            expected_label = f"第 {expected_start}-{expected_end} 页"
            actual_start, actual_end = unique_section_ranges[0]
            return (
                "监管层阻止写入：DOCX 小节页码与当前读取窗口不一致。"
                f" 当前窗口应为 {expected_label}，但正文标题写成第 {actual_start}-{actual_end} 页。"
            )
        if expected_range and unique_declared_ranges:
            mismatched = [
                item for item in unique_declared_ranges if item != expected_range
            ]
            if mismatched:
                expected_label = f"第 {expected_start}-{expected_end} 页"
                actual_start, actual_end = mismatched[0]
                return (
                    "监管层阻止写入：DOCX 页窗标签与当前读取窗口不一致。"
                    f" 当前窗口应为 {expected_label}，但正文写成第 {actual_start}-{actual_end} 页。"
                )

        normalized_blocks: List[str] = []
        for line in body.splitlines():
            cleaned = re.sub(r"\s+", "", line)
            cleaned = re.sub(r"^[\-*•\d.、（）()]+", "", cleaned)
            if len(cleaned) >= 16:
                normalized_blocks.append(cleaned.lower())
        seen_blocks: set[str] = set()
        repeated_blocks: List[str] = []
        for block in normalized_blocks:
            if block in seen_blocks:
                repeated_blocks.append(block)
            seen_blocks.add(block)
        if repeated_blocks:
            return "监管层阻止写入：当前分步 DOCX 正文存在重复段落。请去重后再写入。"

        label_hits = sum(
            1 for label in ("文档识别", "段落主题", "结构线索", "内容线索", "来源页码") if label in body
        )
        if label_hits < 4:
            return (
                "监管层阻止写入：当前分步摘要缺少稳定结构。"
                " 请按“当前页窗摘要 / 文档识别 / 段落主题 / 结构线索 / 内容线索 / 来源页码”重写。"
            )
        return ""

    def _stepwise_docx_wait_artifact(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        snippets: List[Dict[str, Any]],
        file_changes: List[Dict[str, Any]],
        recipe_skeleton: Dict[str, Any],
    ) -> Optional[Dict[str, Any]]:
        if not _looks_like_windowed_pdf_task(request, recipe_skeleton):
            return None
        docx_change = next(
            (
                change
                for change in file_changes
                if str(change.get("operation") or "") == "write_docx_content"
                and str(change.get("path") or change.get("file_path") or "")
                .lower()
                .endswith(".docx")
            ),
            None,
        )
        if not docx_change:
            return None
        target_path = str(
            docx_change.get("path")
            or docx_change.get("file_path")
            or request.target_path
            or self._stepwise_docx_target_path(request, files)
            or ""
        ).strip()
        pdf_file = self._first_context_file(files, {"pdf"})
        latest_pdf_snippet = next(
            (
                item
                for item in reversed(snippets or [])
                if isinstance(item, dict)
                and (
                    str(item.get("source") or item.get("path") or "")
                    .lower()
                    .endswith(".pdf")
                    or str(item.get("path") or "").lower().endswith(".pdf")
                )
            ),
            {},
        )
        window_pages = _stepwise_pdf_window_pages(request)
        current_step_index = _stepwise_pdf_step_index(request)
        current_start = int(
            latest_pdf_snippet.get("start_page")
            or (1 + current_step_index * window_pages)
        )
        current_end = int(
            latest_pdf_snippet.get("end_page") or (current_start + window_pages - 1)
        )
        next_step_index = current_step_index + 1
        next_start = current_end + 1
        next_end = next_start + window_pages - 1

        resume_files = [
            file_info.public_dict()
            for file_info in files
            if file_info and (file_info.path or file_info.name)
        ]
        for item in resume_files:
            if target_path and str(item.get("path") or "") == target_path:
                item["target"] = True
        original_task = (
            str(
                ((request.options or {}).get("batch_control") or {}).get(
                    "original_task"
                )
                if isinstance((request.options or {}).get("batch_control"), dict)
                else ""
            ).strip()
            or str(request.task or "").strip()
        )
        resume_request = {
            "task": f"继续当前分步文件任务的下一步：处理 PDF 第 {next_start}-{next_end} 页，并把本段实质分析追加到同一个 DOCX。",
            "session_id": request.session_id,
            "model_mode": request.model_mode,
            "model_id": request.model_id,
            "target_path": target_path,
            "files": resume_files,
            "options": {
                "batch_control": {
                    "adapter": "generic_tool_loop",
                    "policy": "confirm_each_step",
                    "step_index": next_step_index,
                    "window_pages": window_pages,
                    "original_task": original_task,
                    "source_path": pdf_file.path if pdf_file and pdf_file.path else "",
                    "target_path": target_path,
                },
                "followup_context": {
                    "kind": "stepwise_task_resume",
                    "source": "koto_stepwise_resume_artifact",
                    "followup_action": "resume",
                    "stepwise": {
                        "policy": "confirm_each_step",
                        "completed_page_range": f"{current_start}-{current_end}",
                        "next_page_range": f"{next_start}-{next_end}",
                        "next_step_index": next_step_index,
                        "original_task": original_task,
                    },
                },
            },
        }
        artifact: Dict[str, Any] = {
            "artifact_type": "koto_stepwise_resume_v1",
            "category": "stepwise_confirmation",
            "title": f"继续处理第 {next_start}-{next_end} 页",
            "summary": f"上一段（第 {current_start}-{current_end} 页）已写入 DOCX。可以继续处理第 {next_start}-{next_end} 页。",
            "suggested_next_step": f"点击继续处理第 {next_start}-{next_end} 页",
            "action_label": f"继续第 {next_start}-{next_end} 页",
            "route": "long_pdf_stepwise_docx_summary",
            "current_step_status": "written",
            "completed_page_range": f"{current_start}-{current_end}",
            "next_page_range": f"{next_start}-{next_end}",
            "next_start_page": next_start,
            "next_end_page": next_end,
            "next_step_index": next_step_index,
            "resume_request": resume_request,
        }
        if target_path:
            artifact["target_path"] = target_path
        if pdf_file and pdf_file.path:
            artifact["source_path"] = pdf_file.path
        return artifact

    def _stepwise_pdf_fallback_paragraphs(
        self,
        request: FileTaskRequest,
        pdf_snippet: Dict[str, Any],
        exc: Exception,
    ) -> List[Dict[str, str]]:
        preview = str(
            pdf_snippet.get("_raw_text") or pdf_snippet.get("preview") or ""
        ).strip()
        pages = [
            int(match.group(1)) for match in re.finditer(r"\[Page\s+(\d+)\]", preview)
        ]
        start_page = int(pdf_snippet.get("start_page") or 0)
        end_page = int(pdf_snippet.get("end_page") or 0)
        if start_page and end_page:
            page_range = (
                f"第 {start_page}-{end_page} 页"
                if start_page != end_page
                else f"第 {start_page} 页"
            )
        elif pages:
            page_range = (
                f"第 {min(pages)}-{max(pages)} 页"
                if min(pages) != max(pages)
                else f"第 {pages[0]} 页"
            )
        else:
            page_range = "当前页窗口"
        insights = self._stepwise_pdf_fallback_insights(preview)
        cleaned_preview = re.sub(r"\[Page\s+\d+\]", " ", preview, flags=re.IGNORECASE)
        cleaned_preview = re.sub(r"\s+", " ", cleaned_preview).strip(" ；;，,")

        def _field(label: str) -> str:
            prefix = f"{label}："
            for item in insights:
                text = str(item or "").strip()
                if text.startswith(prefix):
                    return text[len(prefix) :].strip()
            return ""

        source_name = _compact_line(
            Path(
                str(pdf_snippet.get("source") or pdf_snippet.get("path") or "PDF 文档")
            ).stem,
            160,
        )
        document_value = _field("文档识别") or f"当前页窗来自“{source_name}”。"
        topic_value = _field("段落主题")
        structure_value = _field("结构线索")
        content_value = _field("内容线索")

        supplemental = [
            str(item or "").strip()
            for item in insights
            if item
            and not str(item).startswith(("文档识别：", "段落主题：", "结构线索：", "内容线索：", "来源页码："))
        ]
        if not topic_value:
            topic_seed = (
                content_value
                or (supplemental[0] if supplemental else "")
                or cleaned_preview
            )
            topic_value = _compact_line(topic_seed, 180) or "当前页窗文本较短，主题需结合后续页窗继续确认。"
        if not structure_value:
            structure_seed = "；".join(supplemental[:2])
            structure_value = (
                _compact_line(structure_seed, 220)
                if structure_seed
                else "当前页窗作为本步骤材料，记录可提取的结构与上下文线索，供后续页窗衔接。"
            )
        if not content_value:
            content_seed = cleaned_preview or "当前页窗未提取到足够正文，暂不能形成可靠内容摘要。"
            content_value = _compact_line(content_seed, 260)
        paragraphs = [
            {"text": f"当前页窗摘要（{page_range}）", "style": "Heading 1"},
            {"text": f"文档识别：{document_value}"},
            {"text": f"段落主题：{topic_value}"},
            {"text": f"结构线索：{structure_value}"},
            {"text": f"内容线索：{content_value}"},
            {"text": f"来源页码：{page_range}"},
        ]
        return paragraphs

    def _stepwise_pdf_fallback_insights(self, preview: str) -> List[str]:
        cleaned = re.sub(
            r"\[Page\s+\d+\]", "\n", str(preview or ""), flags=re.IGNORECASE
        )
        raw_lines = [
            re.sub(r"\s+", " ", line).strip(" \t|-") for line in cleaned.splitlines()
        ]

        def _is_running_header(line: str) -> bool:
            compact = re.sub(r"\s+", "", line)
            return bool(
                re.search(
                    r"Annual Report on Digital Technology Application|Case Study in Chinese Museums",
                    line,
                    re.IGNORECASE,
                )
                or "中国博物馆数字技术应用及案例研究年度报告" in compact
                or re.fullmatch(r"(?:SUMMAR|ARTICLE|综|述|篇)", line, flags=re.IGNORECASE)
            )

        def _is_noise_line(line: str) -> bool:
            if not line or len(line) < 2:
                return True
            if line.isdigit() or re.fullmatch(r"\d+\s+\d+", line):
                return True
            if _is_running_header(line):
                return True
            return False

        def _join_pdf_body_lines(source_lines: List[str]) -> List[str]:
            blocks: List[str] = []
            buffer = ""

            def _flush() -> None:
                nonlocal buffer
                text = buffer.strip(" ；;，,")
                if len(text) >= 12:
                    blocks.append(text)
                buffer = ""

            for line in source_lines:
                if _is_noise_line(line):
                    _flush()
                    continue
                if re.fullmatch(r"\d+\.\s+.+", line) and len(line) > 90:
                    _flush()
                    continue
                starts_new = bool(
                    re.match(r"^[一二三四五六七八九十]+、", line)
                    or re.match(r"^\d+[.、]\s*", line)
                    or re.match(r"^《.+》", line)
                    or re.match(r"^表\d+", line)
                )
                if starts_new:
                    _flush()
                    buffer = line
                elif buffer and len(buffer) + len(line) <= 260:
                    separator = "" if re.search(r"[\u4e00-\u9fff]$", buffer) else " "
                    buffer = f"{buffer}{separator}{line}"
                else:
                    _flush()
                    buffer = line
                if re.search(r"[。！？!?]$", line) or len(buffer) >= 220:
                    _flush()
            _flush()

            deduped: List[str] = []
            seen_blocks: set[str] = set()
            for block in blocks:
                key = re.sub(r"\s+", "", block).lower()
                if key in seen_blocks:
                    continue
                seen_blocks.add(key)
                deduped.append(block)
            return deduped

        lines: List[str] = []
        seen: set[str] = set()
        for line in raw_lines:
            if _is_noise_line(line) or len(line) < 4:
                continue
            key = line.lower()
            if key in seen:
                continue
            seen.add(key)
            lines.append(line)
            if len(lines) >= 80:
                break

        annual_title = next(
            (
                line
                for line in raw_lines
                if line
                and (
                    "Annual Report" in line
                    or "年度报告" in line
                    or "中国博物馆数字技术应用及案例研究年度报告" in re.sub(r"\s+", "", line)
                )
            ),
            "",
        )
        section_title = next(
            (
                line
                for line in lines
                if re.match(r"^[一二三四五六七八九十]+、", line)
                or line.startswith("关于")
                or re.match(r"^《.+》", line)
            ),
            "",
        )
        organizer = next(
            (
                line
                for line in lines
                if len(line) <= 120
                and not re.match(r"^\d+[.、]", line)
                and (
                    re.search(
                        r"^(?:主\s*编|执行主编|副\s*主\s*编|专家顾问|编\s*辑|英文翻译|美术编辑)",
                        line,
                    )
                    or "委员会" in line
                    or "中国博物馆协会资助项目" in line
                    or re.search(r"中国博物馆协会.*(?:编|委员会)", line)
                )
            ),
            "",
        )
        editor_lines = [
            line
            for line in lines
            if re.search(r"(?:主\s*编|执行主编|副\s*主\s*编|专家顾问|编辑)", line)
        ][:3]
        toc_lines = [
            line
            for line in lines
            if len(line) <= 50 and re.search(r"(?:目录|引言|综述篇|案例篇|参考文献|作者简介)", line)
        ][:4]
        body_blocks = _join_pdf_body_lines(raw_lines)
        content_lines = [
            line
            for line in body_blocks
            if len(line) >= 18
            and line != annual_title
            and line != section_title
            and not (section_title and line.startswith(section_title))
            and line != organizer
            and line not in editor_lines
            and line not in toc_lines
            and not _is_running_header(line)
            and not re.search(
                r"^(?:主\s*编|执行主编|副\s*主\s*编|专家顾问|编\s*辑|英文翻译|美术编辑)",
                line,
            )
        ][:4]

        insights: List[str] = []
        if annual_title:
            insights.append(f"文档识别：当前页窗来自“{_compact_line(annual_title, 180)}”。")
        if section_title:
            insights.append(f"段落主题：{_compact_line(section_title, 180)}。")
        if organizer and organizer not in {annual_title, section_title}:
            insights.append(f"组织信息：{_compact_line(organizer, 180)}。")
        if editor_lines:
            insights.append(
                "编写线索："
                + "；".join(_compact_line(line, 120) for line in editor_lines)
                + "。"
            )
        if toc_lines:
            insights.append(
                "结构线索：" + "；".join(_compact_line(line, 120) for line in toc_lines) + "。"
            )
        if content_lines:
            insights.append(
                "内容线索："
                + "；".join(_compact_line(line, 180) for line in content_lines)
                + "。"
            )
        if not insights:
            excerpt_lines = [_compact_line(line, 140) for line in lines[:4]]
            if excerpt_lines:
                insights.append("当前页窗可读内容集中在：" + "；".join(excerpt_lines) + "。")
        if not insights:
            insights.append("当前页窗未提取到足够正文，暂不能形成可靠内容摘要。")
        return insights

    def _fallback_readonly_summary(
        self,
        request: FileTaskRequest,
        snippets: List[Dict[str, Any]],
        files: List[FileTaskFile],
        exc: Exception,
    ) -> str:
        if not snippets:
            return ""

        lines = [
            "模型暂不可用，Koto 已先基于显式上下文整理可见内容（非模型推理）：",
        ]
        used_sources: set[str] = set()
        for index, snippet in enumerate(snippets[:5], start=1):
            source = str(
                snippet.get("source") or snippet.get("path") or f"上下文 {index}"
            ).strip()
            if not source and index <= len(files):
                source = files[index - 1].name or files[index - 1].path
            source_label = self._display_path(source) or f"上下文 {index}"
            preview = _compact_line(snippet.get("preview"), 320)
            if not preview:
                continue
            dedupe_key = f"{source_label}:{preview}"
            if dedupe_key in used_sources:
                continue
            used_sources.add(dedupe_key)
            lines.append(f"{index}. {source_label}：{preview}")

        if len(lines) == 1:
            return ""

        lines.append("恢复模型后可以继续生成更完整的总结、改写或写入文件。")
        lines.append(f"模型错误：{_compact_line(exc, 160)}")
        return "\n".join(lines)

    def _readonly_answer_required_message(
        self,
        request: FileTaskRequest,
        snippets: List[Dict[str, Any]],
        readonly_tool_outputs: List[Dict[str, Any]],
    ) -> str:
        lines = [
            "你已经完成了只读文件读取，但还没有给用户可见答案。本轮必须直接输出分析结果，不要空回复。",
            f"用户任务：{request.task}",
            "要求：基于已读取内容给出结构化结论；如果信息不足，也要明确说明已读取到什么、缺什么、下一步怎么做。",
        ]
        source_lines = self._readonly_context_source_lines(
            snippets, readonly_tool_outputs, limit=5
        )
        if source_lines:
            lines.append("已读取内容摘录：")
            lines.extend(source_lines)
        return "\n".join(lines)

    def _readonly_context_source_lines(
        self,
        snippets: List[Dict[str, Any]],
        readonly_tool_outputs: List[Dict[str, Any]],
        *,
        limit: int = 5,
    ) -> List[str]:
        lines: List[str] = []
        seen: set[str] = set()
        for item in readonly_tool_outputs:
            if not isinstance(item, dict):
                continue
            source = self._readonly_tool_source_label(item)
            for point in self._readonly_tool_points(item):
                text = _compact_line(point, 260)
                if not text:
                    continue
                key = f"{source}:{text}"
                if key in seen:
                    continue
                seen.add(key)
                lines.append(f"- {source}：{text}")
                if len(lines) >= limit:
                    return lines
        for index, snippet in enumerate(snippets, start=1):
            if not isinstance(snippet, dict):
                continue
            source = str(
                snippet.get("source") or snippet.get("path") or f"上下文 {index}"
            ).strip()
            text = _compact_line(snippet.get("preview"), 260)
            if not text:
                continue
            key = f"{source}:{text}"
            if key in seen:
                continue
            seen.add(key)
            lines.append(f"- {self._display_path(source) or source}：{text}")
            if len(lines) >= limit:
                break
        return lines

    def _readonly_context_summary(
        self,
        request: FileTaskRequest,
        snippets: List[Dict[str, Any]],
        readonly_tool_outputs: List[Dict[str, Any]],
    ) -> str:
        source_lines = self._readonly_context_source_lines(
            snippets, readonly_tool_outputs, limit=7
        )
        if not source_lines:
            return ""
        lines = [
            "已完成文件读取，但模型没有返回进一步自然语言分析。以下是 Koto 基于已读取内容整理的可见结果：",
            f"任务：{request.task}",
            "已读取内容：",
            *source_lines,
            "结论：本轮为只读分析，没有写入或修改文件。可以继续追问，让模型基于上述内容做更深入的总结、风险识别或访谈提纲整理。",
        ]
        return "\n".join(lines)

    def _readonly_tool_source_label(self, item: Dict[str, Any]) -> str:
        args = item.get("args") if isinstance(item.get("args"), dict) else {}
        raw_path = str(
            args.get("path") or args.get("file_path") or item.get("path") or ""
        ).strip()
        if raw_path:
            return self._display_path(raw_path) or raw_path
        tool_name = str(item.get("tool_name") or "").strip()
        return tool_name or "读取结果"

    def _readonly_tool_points(self, item: Dict[str, Any]) -> List[str]:
        result = item.get("result")
        payload = result if isinstance(result, dict) else _json_payload(result)
        points: List[str] = []
        if isinstance(payload, dict):
            paragraphs = (
                payload.get("paragraphs")
                if isinstance(payload.get("paragraphs"), list)
                else []
            )
            tables = (
                payload.get("tables") if isinstance(payload.get("tables"), list) else []
            )
            total_paragraphs = payload.get("total_paragraphs")
            total_tables = payload.get("total_tables")
            if total_paragraphs is not None or total_tables is not None:
                points.append(
                    f"Word 内容包含 {int(total_paragraphs or len(paragraphs) or 0)} 段文本、{int(total_tables or len(tables) or 0)} 个表格。"
                )
            for paragraph in paragraphs:
                if not isinstance(paragraph, dict):
                    continue
                text = str(paragraph.get("text") or "").strip()
                if text:
                    points.append(text)
                if len(points) >= 6:
                    break
            if not points and payload.get("text"):
                points.append(str(payload.get("text") or ""))
        if not points:
            preview = str(item.get("preview") or "").strip()
            if preview:
                points.append(preview)
        if not points and result is not None:
            points.append(stringify_result(result))
        return points

    def _success_criteria(
        self, request: FileTaskRequest, write_intent: bool, output_mode: str
    ) -> List[str]:
        criteria = [
            "每个步骤都产生 typed event，可被前端时间线渲染",
            "所有上下文来源都来自显式输入",
        ]
        recipe_match = select_task_recipe(
            request, request.files or [], write_intent=write_intent
        )
        if recipe_match and recipe_match.recipe.success_criteria:
            criteria.extend(
                str(item)
                for item in recipe_match.recipe.success_criteria
                if str(item or "").strip()
            )
            return criteria
        if write_intent:
            criteria.extend(
                [
                    "写入工具必须产生 file.changed 事件",
                    "最终 checker 必须确认目标文件已更新",
                ]
            )
        elif output_mode == "hybrid":
            criteria.append("最终摘要必须给出明确建议，且当前轮不默认直接写入原文件")
        else:
            criteria.append("最终摘要说明已使用的上下文和未完成项")
        return criteria

    def _file_types(self, files: List[FileTaskFile]) -> set[str]:
        return request_file_types(files)

    def _looks_like_chart_request(self, task: str) -> bool:
        return semantic_markers(task).get("chart_request", False)

    def _looks_like_problem_analysis_request(self, task: str) -> bool:
        return semantic_markers(task).get("problem_analysis_request", False)

    def _looks_like_financial_request(self, task: str) -> bool:
        return semantic_markers(task).get("financial_request", False)

    def _looks_like_table_request(self, task: str) -> bool:
        return semantic_markers(task).get("table_request", False)

    def _looks_like_summary_request(self, task: str) -> bool:
        return semantic_markers(task).get("summary_request", False)

    def _looks_like_translation_request(self, task: str) -> bool:
        return semantic_markers(task).get("translation_request", False)

    def _looks_like_polish_request(self, task: str) -> bool:
        return semantic_markers(task).get("polish_request", False)

    def _looks_like_ppt_request(self, task: str, files: List[FileTaskFile]) -> bool:
        return semantic_markers(task, file_types=self._file_types(files)).get(
            "ppt_request", False
        )

    def _looks_like_ppt_slide_write_request(
        self, request: FileTaskRequest, files: List[FileTaskFile]
    ) -> bool:
        return semantic_markers(request.task, file_types=self._file_types(files)).get(
            "ppt_slide_write_request", False
        )

    def _looks_like_docx_report_request(
        self, request: FileTaskRequest, files: List[FileTaskFile]
    ) -> bool:
        return semantic_markers(
            request.task,
            file_types=self._file_types(files),
            target_file_type=request_target_file_type(request, files),
        ).get("docx_report_request", False)

    def _looks_like_financial_xlsx_docx_chart_report_task(
        self, request: FileTaskRequest, files: List[FileTaskFile]
    ) -> bool:
        return semantic_markers(
            request.task,
            file_types=self._file_types(files),
            target_file_type=request_target_file_type(request, files),
        ).get("financial_xlsx_docx_chart_report", False)

    def _looks_like_pdf_python_text_read(self, code: Any) -> bool:
        text = str(code or "").lower()
        if not text.strip():
            return False

        pdf_markers = (
            "pypdf2",
            "from pypdf import",
            "pdfreader",
            "pdfplumber",
            "pymupdf",
            "fitz",
            ".pdf",
            "pdf_path",
        )
        read_markers = (
            "extract_text(",
            "get_text(",
            "reader.pages",
            "page.get_text",
            "page.extract_text",
            "pdf.pages",
        )
        return any(marker in text for marker in pdf_markers) and any(
            marker in text for marker in read_markers
        )

    def _blocked_run_python_message(
        self,
        tool_name: str,
        tool_args: Dict[str, Any],
        request: FileTaskRequest,
        files: List[FileTaskFile],
    ) -> str:
        if tool_name != "run_python_code":
            return ""

        file_types = self._file_types(files)
        code = tool_args.get("code")
        if "pdf" not in file_types and ".pdf" not in str(code or "").lower():
            return ""
        if not self._looks_like_pdf_python_text_read(code):
            return ""

        message = (
            "不要用 run_python_code 直接读取 PDF 文本。"
            "请改用 parse_file_to_text(path, max_chars, start_page, end_page)；"
            "长文或原文对照任务必须按页窗口分段读取。"
        )
        if "docx" in file_types:
            message += " 读取完 PDF 分段后，再用 read_docx_content 读取 DOCX 译稿。"
        if request.target_path:
            message += f" 当前目标文件是：{request.target_path}。"
        return message

    def _should_prompt_for_write_after_tool_round(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        tool_calls: List[Dict[str, Any]],
        round_index: int,
    ) -> bool:
        if not tool_calls:
            return False
        if any(
            is_write_tool(str(call.get("name") or ""))
            and str(call.get("name") or "") != "run_python_code"
            for call in tool_calls
        ):
            return False
        if self._looks_like_financial_xlsx_docx_chart_report_task(request, files):
            return True
        return round_index >= 2

    def _write_retry_message(
        self, request: FileTaskRequest, files: List[FileTaskFile]
    ) -> str:
        target = request.target_path or next(
            (
                file_info.path
                for file_info in files
                if file_info.target and file_info.path
            ),
            "",
        )
        file_types = self._file_types(files)
        task_text = str(request.task or "")
        hint = "你还没有完成真实文件写入。不要只总结或结束，下一轮必须调用会修改文件的工具。"
        if self._looks_like_financial_xlsx_docx_chart_report_task(request, files):
            hint += (
                " 当前是 Excel 财务预测图表+问题写入 DOCX 任务：不要只插入 Excel 原表，也不要只输出 Python stdout。"
                " 先用 run_python_code 生成真实 PNG/JPG 图表并输出 KOTO_CREATED 路径；"
                "再调用 write_docx_content 写入问题清单/分析结论；"
                "最后调用 insert_image_into_docx 把生成的真实图片插入目标 DOCX。"
                " 解析 P&L 时不要依赖 pandas 默认列名；如果列名是 Unnamed，应扫描每一行找到 2025E/2026E/2027E/2028E 等年份头，再按这些列抽取收入、毛利、费用、净利润等指标。"
            )
        elif "xlsx" in file_types and "docx" in file_types:
            hint += " 对于把 Excel 加入 Word，优先调用 insert_excel_as_docx_table；如果已经读到真实工作表名，就用真实 sheet 写入目标 docx。"
        if "docx" in file_types and re.search(
            r"(?:图表|可视化|绘图|画图|画.{0,4}图|图片|chart|plot|graph|image)",
            task_text,
            re.IGNORECASE,
        ):
            hint += " 如果用户要求把图表或图片加入 DOCX，先用 run_python_code 生成真实 PNG/JPG 文件，再调用 insert_image_into_docx；不要用 write_docx_content 把图片描述文字写进文档代替真实插图。"
        if {"txt", "md", "csv", "json", "py", "js", "html", "css"}.intersection(
            file_types
        ):
            hint += " 对于 TXT/MD/CSV/JSON 或代码文本文件，如果用户提供了选区并要求润色/改写/替换后写回，优先调用 replace_file_selection，用 original_selection=用户选区原文、new_content=改写结果；不要为了单个选区改写去 run_python_code 整文件覆写。没有选区时，先用 read_file_range 或 parse_file_to_text 读取必要内容，再选择 replace_file_selection 或 run_python_code 写回；如果只是批注/审校可用 annotate_file。不要只输出润色后的文本而不落盘。"
        if "pdf" in file_types:
            hint += " 读取 PDF 原文必须调用 parse_file_to_text；长文必须用 start_page/end_page 分段读取，不要用 run_python_code、PyPDF2、pdfplumber 或 fitz 直接解析 PDF。"
        if "pdf" in file_types and "docx" in file_types:
            hint += " 对于 PDF 原文和 DOCX 译稿对照任务，先分页读取 PDF，再读取 DOCX；不要试图一次性抽取整本 PDF。"
        if "pptx" in file_types:
            hint += " 对于 PPT，读取内容优先用 parse_file_to_text；如果要整体风格、主题、版式、美化或配色，调用 design_pptx_theme_layout；如果要新增总结页，调用 add_pptx_slides；如果是改现有页文本，用 write_pptx_slides。不要对 PPTX 调用 read_docx_content。"
        if target:
            hint += f" 当前目标文件是：{target}。"
        return hint

    def _duplicate_supervisor_retry_message(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        classification: FileTaskClassification,
        intent_plan: FileTaskIntentPlan,
        tool_calls: List[Dict[str, Any]],
    ) -> str:
        repeated_tools = (
            ", ".join(
                str(call.get("name") or "").strip()
                for call in tool_calls
                if str(call.get("name") or "").strip()
            )
            or "上一轮工具"
        )
        lines = [
            "监管层检测到你正在重复上一轮相同工具调用，但当前任务仍未产生任何 file.changed。",
            f"重复工具：{repeated_tools}",
            "不要继续重复读取同一批内容；下一轮必须回到计划主线，改变工具参数或推进到写入/生成/插入步骤。",
        ]
        selected_recipe = str(classification.selected_recipe or "").strip()
        if selected_recipe:
            lines.append(f"当前任务路线：{selected_recipe}")
        if intent_plan.dynamic_steps:
            lines.append("计划账本：")
            for index, step in enumerate(intent_plan.dynamic_steps[:8], start=1):
                if not isinstance(step, dict):
                    continue
                title = str(
                    step.get("title") or step.get("id") or f"步骤 {index}"
                ).strip()
                description = str(step.get("description") or "").strip()
                if title:
                    lines.append(
                        f"{index}. {title}" + (f"：{description}" if description else "")
                    )
        file_types = self._file_types(files)
        if "pdf" in file_types:
            lines.append(
                "PDF 长文任务：如已读取当前页窗，下一轮要么换 start_page/end_page 读取下一段，要么把当前步骤要点写入目标 DOCX；不要再次读取同一页窗。"
            )
        if (
            "docx" in file_types
            or "docx" in str(request.task or "").lower()
            or "word" in str(request.task or "").lower()
        ):
            lines.append(
                "DOCX 输出任务：必须调用 write_docx_content 写入本步骤发现；如果没有明确目标路径，就在源文件同目录创建清晰命名的 DOCX 输出文件。"
            )
        if "xlsx" in file_types:
            lines.append("Excel 任务：如果已完成结构读取，下一轮必须进入真实分析/制图/写回，不要重复打印同一张表。")
        target = request.target_path or next(
            (
                file_info.path
                for file_info in files
                if file_info.target and file_info.path
            ),
            "",
        )
        if target:
            lines.append(f"目标文件：{target}")
        lines.append("只有在本轮已经产生真实文件变更，或任务确实是只读答复时，才允许结束。")
        return "\n".join(lines)

    def _build_system_prompt(
        self,
        request: FileTaskRequest,
        files: List[FileTaskFile],
        known_tool_gap: Optional[Dict[str, Any]] = None,
        classification: Optional[FileTaskClassification] = None,
        intent_plan: Optional[FileTaskIntentPlan] = None,
        execution_context: Optional[FileTaskExecutionContext] = None,
        recipe_skeleton: Optional[Dict[str, Any]] = None,
    ) -> str:
        resolved_context = execution_context or self._build_execution_context(
            request,
            files,
            known_tool_gap=known_tool_gap,
            classification=classification,
            intent_plan=intent_plan,
        )
        resolved_classification = resolved_context.classification
        resolved_intent_plan = resolved_context.intent_plan
        resolved_known_tool_gap = resolved_context.known_tool_gap
        workflows = json.dumps(supported_file_workflows(), ensure_ascii=False, indent=2)
        file_list = (
            ", ".join(
                (file_info.path or file_info.name)
                for file_info in files
                if file_info.path or file_info.name
            )
            or "none"
        )
        capability_profiles = build_request_capability_profiles(request)
        skeleton = recipe_skeleton or build_recipe_skeleton(
            request,
            files,
            resolved_classification,
            resolved_intent_plan,
            [],
        )
        known_gap_text = ""
        if resolved_known_tool_gap:
            known_gap_text = (
                "\n已知原生工具缺口：\n"
                + json.dumps(resolved_known_tool_gap, ensure_ascii=False, indent=2)
                + "\n"
            )
        capability_text = ""
        if capability_profiles:
            capability_text = (
                "文件能力概览：" + json.dumps(capability_profiles, ensure_ascii=False) + "\n"
            )
        followup_context = self._followup_context(request)
        followup_guidance = ""
        financial_chart_docx_guidance = ""
        docx_compare_annotate_guidance = ""
        clear_docx_review_guidance = ""
        single_docx_annotate_guidance = ""
        if self._looks_like_financial_xlsx_docx_chart_report_task(request, files):
            financial_chart_docx_guidance = (
                "Excel 财务预测图表写入 DOCX 任务规则：\n"
                "- 目标不是把原始 Excel 表格塞进 Word，而是生成“问题清单/分析结论 + 真实图表图片”。\n"
                "- 必须先审计 Excel：inspect_workbook_structure 或 audit_financial_workbook；必要时读取 P&L、产品线、Expenses、资本折旧等关键工作表。\n"
                "- 如果 pandas 读出的列名是 Unnamed，不要用 df.columns 找年份列；应扫描每一行，定位包含 2025E/2026E/2027E/2028E 等年份标签的 header row，再按这些列抽取指标。\n"
                "- 优先用 openpyxl/data_only=True 读取公式结果，并通过行标签匹配“收入合计、毛利合计、费用合计、净利润、销量”等指标；不要猜空列名。\n"
                "- 用 run_python_code 生成真实 PNG/JPG 图表，stdout 必须包含 KOTO_CREATED: <图片路径>；仅打印数据或错误栈不算完成。\n"
                "- 随后调用 write_docx_content 写入问题清单/分析结论，再调用 insert_image_into_docx 插入真实图片；没有 file.changed 不能结束。\n"
            )
        if "compare_docx_and_annotate" in resolved_classification.matched_capabilities:
            docx_files = [
                self._display_path(file_info.path or file_info.name)
                for file_info in files
                if str(
                    file_info.type
                    or Path(str(file_info.path or file_info.name)).suffix.lstrip(".")
                ).lower()
                in {"docx", "doc"}
            ]
            docx_compare_annotate_guidance = (
                "DOCX 双文件对比标注任务规则：\n"
                f"- 待对比 DOCX：{', '.join(item for item in docx_files if item) or '已附加的两份 DOCX'}\n"
                "- 这是跨文件差异比较，不是单文档审校；不要调用 annotate_file 批改其中一份文稿，也不要创建独立的对比说明文档。\n"
                "- 目标是修改现有 DOCX：把 Word 原生批注写在 target_path 对应原文条款/段落旁边。\n"
                "- target_path 必须是用户要被标注的那份 DOCX；如果用户说“原文/原文件/当前文档/第一份文档上标注”，必须指向该文件。\n"
                "- 推荐流程：先调用 plan_docx_compare_annotations(original_path, revised_path, target_path) 定位目标文档里的可批注差异锚点；再根据候选差异生成 comments_json 数组，调用 write_docx_comments(path=target_path, comments_json=[...]) 写回原 DOCX。\n"
                "- comments_json 必须直接传数组对象，不要把数组转成需要转义的长字符串；每项必须使用候选中的原文片段/anchor 作为锚点；批注内容应简洁说明“另一版为：... 本版为：...”。\n"
                "- 合同任务的批注可补充“风险：...”和“建议：...”，但这些内容也必须作为 Word 批注写在目标合同原文旁边。\n"
                "- 仅当需要兜底时才使用 compare_docx_and_annotate 一步完成；优先让 AI 基于候选差异撰写批注内容后调用 write_docx_comments。\n"
                "- 完成后必须产生 file.changed，且 annotations_added > 0 才能声称已标注差异；对话框总结只概括批注数量和高风险类别。\n"
            )
        if str(followup_context.get("kind") or "").strip() == "review_last_task":
            followup_action = (
                str(followup_context.get("followup_action") or "").strip().lower()
            )
            if followup_action == "apply":
                followup_guidance = (
                    "当前输入是用户要求把上一轮文件任务中的建议直接应用到文件。这不是一个无关的新任务，而是同一任务的写回续跑。\n"
                    "优先沿用上一轮的目标文件、分析建议、文件变更和约束；必要时只补充最少量上下文后直接执行写入。\n"
                    "如果上一轮已经给出可应用建议，这一轮应进入真实写回路径并产生 file.changed；不要只重复建议文本。\n"
                )
            elif followup_action == "improve":
                followup_guidance = (
                    "当前输入是用户要求围绕上一轮文件任务结果继续优化。这不是一个无关的新任务，而是同一任务的后续回合。\n"
                    "先解释上一轮结果的不足和这次准备如何改进；如果确实需要，可以继续调用工具修正目标文件。"
                    "优先沿用上一轮的目标、目标文件、失败点和约束，不要把上下文重置成新的独立任务。\n"
                )
                if _followup_has_prior_excel_docx_insert(followup_context):
                    followup_guidance += (
                        "如果上一轮已经通过 insert_excel_as_docx_table 把 Excel 表格写入目标 DOCX，"
                        "这轮继续优化时不要再次插入同一张表。"
                        "优先补写摘要、说明、结论或修正已有文字；"
                        "只有用户明确要求重插、替换或追加另一张表时，才再次插表。\n"
                    )
            else:
                followup_guidance = (
                    "当前输入是用户对上一轮文件任务结果的反馈或质问，不要默认把它当作全新的执行任务。\n"
                    "先解释上一轮结果、指出可能的问题，并回答用户的追问。"
                    "只有当用户明确要求重新修改文件、继续执行或调用工具时，才进入新的工具执行。"
                    "如果当前只是反馈上一轮结果，不要调用写入工具，也不要伪造新的完成状态。\n"
                )
        if resolved_classification.docx_annotation_request:
            target_docx = (
                self._display_path(request.target_path)
                or self._first_file_name(files, {"docx"}, target=True)
                or self._first_file_name(files, {"docx"})
                or "当前 DOCX"
            )
            single_docx_annotate_guidance = (
                "DOCX 审校/批注任务规则：\n"
                f"- 目标 DOCX：{target_docx}\n"
                "- 直接调用 annotate_file。对于 AI 生成批注的场景，传 path=<目标 DOCX>、requirement=<用户要求>，annotations 保持空数组即可。\n"
                "- 如果当前任务还附带 PDF 原文、分批继续执行信息或上一轮审校 follow-up，上述 annotate_file 会自动复用这些上下文；不要再绕开白盒工具循环。\n"
                "- 不要自己编造 annotations 的 range_start/range_end 去模拟 Word 定位；annotate_file 会负责分析、定位并把批注写回原文。\n"
                "- 如果目标是把意见直接写回 DOCX，不能只输出批注清单文本后结束。\n"
            )
        elif self._is_docx_clear_review_request(request):
            target_docx = (
                self._display_path(request.target_path)
                or self._first_file_name(files, {"docx"}, target=True)
                or self._first_file_name(files, {"docx"})
                or "当前 DOCX"
            )
            clear_docx_review_guidance = (
                "DOCX 批注/修订清理任务规则：\n"
                f"- 目标 DOCX：{target_docx}\n"
                "- 调用 clear_docx_review_marks。若用户只要求删除批注，scope 用 comments；若明确要求去掉修订或全部审阅标记，scope 用 revisions 或 all。\n"
                "- 不要调用 annotate_file 去重新生成批注，也不要走 doc_annotate_bridge。\n"
                "- 这是一个真实写回任务，完成后必须产生 file.changed。\n"
            )
        return (
            "你是 Koto 文件助手的后端执行 agent。你可以自主规划并调用工具，"
            "但只能调用系统提供的 Koto 文件工具。不要编造工具、文件路径或已经完成的写入。\n\n"
            f"{self._output_mode_guidance(resolved_classification)}"
            f"{self._intent_plan_guidance(resolved_intent_plan)}"
            f"{followup_guidance}"
            f"{financial_chart_docx_guidance}"
            f"{docx_compare_annotate_guidance}"
            f"{clear_docx_review_guidance}"
            f"{single_docx_annotate_guidance}"
            "首轮协议：你可以直接调用工具；如果任务较复杂、需要先拆解执行方案，"
            "也可以先返回 execution_plan。"
            f"execution_plan 格式：{json.dumps(whitebox_execution_plan_schema(), ensure_ascii=False)}\n"
            f"execution_brief 格式：{json.dumps(self._execution_brief_schema(), ensure_ascii=False)}\n"
            "返回 execution_plan 或 execution_brief 后，下一轮必须在白盒任务骨架内继续调用 Koto 工具，不要重复同一份计划。\n"
            "白盒任务骨架：\n"
            f"{json.dumps(skeleton, ensure_ascii=False, indent=2)}\n"
            "执行原则：\n"
            "1. 优先使用显式提供的当前文件、附件、选区和目标路径。\n"
            "2. Office 文件必须使用格式感知工具；DOCX/XLSX/PPTX 优先用专用工具，PDF 默认只读提取。\n"
            "3. 读取 PDF 文本时只能使用 parse_file_to_text；长文必须使用 start_page/end_page 按页窗口分段读取。不要用 run_python_code 调用 PyPDF2/pypdf/pdfplumber/fitz/PyMuPDF 读取 PDF。\n"
            "4. 用户要求分步、每步汇报、等他说继续时，必须把它当作 confirm_each_step 任务：每一步只处理一个小窗口；如果任务要求创建/更新 DOCX，必须先用 write_docx_content 写入当前页窗的实质摘要、关键发现和来源页码，再进入等待确认。分步 DOCX 正文必须使用稳定模板：Heading 1 写“当前页窗摘要（第 x-y 页）”，随后用独立纯文本段落依次写“文档识别：...”“段落主题：...”“结构线索：...”“内容线索：...”“来源页码：第 x-y 页”；这里的 x-y 必须严格等于 context_snippets 当前 PDF 片段的 start_page/end_page，不要使用 PDF 印刷页码、目录页码、章节页码或模型推断页码。不要写“文档识别/核心要点”这类合并标签，不要写 Markdown 的 #、**、---。内容必须由模型基于 context_snippets 中当前页窗文本综合提炼：解释这一页窗在全文结构中的作用，区分目录/标题/正文/案例信息，合并重复页眉页脚，保留关键概念、章节名、案例名和论证线索；不要把页码、目录条目、作者名单或原文碎片机械拼接成摘要。每段应是可读的分析句或紧凑要点，而不是关键词串。一轮只写当前页窗，不要重复前面页窗，不要把同一页窗拆成多个重复标题，不要只堆目录或原文列表。DOCX 正文不能包含“等待继续、下一步计划、当前步骤已完成、当前进度、file.changed、状态”等前端进度提示；这些只放在助手消息/运行事件里。未产生 file.changed 不允许声称“当前步骤完成”。\n"
            "5. 当用户要求创建 DOCX/Word 但没有明确目标路径时，在源文件同目录创建清晰命名的输出文件，例如“源文件名_分步总结.docx”；不要因为缺少目标路径而只输出文字。\n"
            "6. PDF 原文 + DOCX 译稿/润色/审校任务，先分段读取 PDF，再读取 DOCX；不要一次性抽取整本 PDF，也不要用 Python 临时脚本拼接全文。\n"
            "7. Excel 工作表名未知时不要猜 Sheet1；省略 sheet_name，或先读取表格让工具返回真实 sheet 名。若请求的工作表不存在，继续根据 available_sheets 和已读取结果完成分析，并明确说明缺失的报表。\n"
            "8. 遇到财务模型、预算、预测、报表审阅类任务时，先调用 inspect_workbook_structure 或 audit_financial_workbook，先确认工作表完整性、外部链接、年份列和公式缺口，再用 read_sheet_data 深入关键工作表。区分“结构性缺陷/可复算性问题”和“经营假设偏激进”，不要混为一谈。遇到 P&L 第一行不是表头、列名为 Unnamed 的工作簿时，必须扫描行内容定位年份头，不要用空列名或 df.columns 直接取数。\n"
            "9. 读取 PPTX 内容优先用 parse_file_to_text；read_docx_content 只用于 DOCX。\n"
            "10. 需要整体设计 PPTX 的风格、主题、版式、美化或配色时调用 design_pptx_theme_layout；需要新增 PPT 总结页时优先用 add_pptx_slides；修改现有页内容时用 write_pptx_slides。\n"
            "11. 对于 TXT/MD/CSV/JSON/代码等文本文件的直接改写：如果用户有选区，优先用 replace_file_selection 精准替换选区，original_selection=用户选区原文，new_content=改写结果；不要为了单个选区改写去 run_python_code 整文件覆写。没有选区时先用 read_file_range 或 parse_file_to_text 读取必要片段，再用 replace_file_selection 或 run_python_code 写回。不要只返回改写后的文本。\n"
            "12. 需要计算、制图、批量转换或复杂文件处理时使用 run_python_code，并在输出中保留 KOTO_CREATED/KOTO_MODIFIED 标记；但 PDF 文本读取不属于这一类。\n"
            "13. 如果任务要求把图表/图片加入 DOCX，先用 run_python_code 生成真实图片文件，再调用 insert_image_into_docx 把图片写回目标 DOCX；不要把图片描述文字写进文档代替真实插图。\n"
            "14. 生成中文图表时，优先配置 matplotlib 中文字体候选（Microsoft YaHei、SimHei、Noto Sans CJK SC、WenQuanYi Micro Hei、DejaVu Sans）并设置 axes.unicode_minus=False；保存图表时使用 dpi>=220 和 bbox_inches='tight'。\n"
            "15. Excel -> DOCX 任务默认要保留真实表格；优先用 insert_excel_as_docx_table 落盘。但如果用户明确要求整理、总结、分析、说明、结论或要点，先用 write_docx_content 把真实摘要写入目标 DOCX，再按需插入一次支撑表格；不要只插原表就结束。\n"
            "16. 完成写入后直接给出简短结果，不要重复写入同一目标文件。\n"
            "17. 如果任务要求的编辑能力当前工具不支持，必须遵循下面的工具设计协议；不要只说做不了，也不要把任务判定为已完成。\n"
            f"{tool_design_prompt_text()}\n\n"
            f"显式文件：{file_list}\n"
            f"目标路径：{request.target_path or 'none'}\n"
            f"{capability_text}"
            f"工具设计协议：{TOOL_DESIGN_PROTOCOL}\n"
            f"{known_gap_text}"
            f"支持的主流办公文件工作流：\n{workflows}\n\n"
            "如果 provider 原生 tool calling 不可用，也可以在文本中输出 JSON 工具调用，格式为 "
            '{"name": "tool_name", "args": {...}} 或由这些对象组成的数组。'
        )

    def _build_messages(
        self,
        request: FileTaskRequest,
        snippets: List[Dict[str, Any]],
        files: List[FileTaskFile],
        known_tool_gap: Optional[Dict[str, Any]] = None,
        classification: Optional[FileTaskClassification] = None,
        intent_plan: Optional[FileTaskIntentPlan] = None,
        execution_context: Optional[FileTaskExecutionContext] = None,
        recipe_skeleton: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        resolved_context = execution_context or self._build_execution_context(
            request,
            files,
            known_tool_gap=known_tool_gap,
            classification=classification,
            intent_plan=intent_plan,
        )
        resolved_classification = resolved_context.classification
        resolved_intent_plan = resolved_context.intent_plan
        resolved_known_tool_gap = resolved_context.known_tool_gap
        capability_profiles = build_request_capability_profiles(request)
        followup_context = self._followup_context(request)
        skeleton = recipe_skeleton or build_recipe_skeleton(
            request,
            files,
            resolved_classification,
            resolved_intent_plan,
            [],
        )
        context = {
            "task": request.task,
            "target_path": request.target_path,
            "selection_source": request.selection_source,
            "task_feedback_mode": {
                "output_mode": resolved_classification.output_mode,
                "label": self._output_mode_label(resolved_classification.output_mode),
                "write_intent": bool(resolved_classification.write_intent),
                "should_write_this_round": str(
                    resolved_classification.output_mode or ""
                )
                .strip()
                .lower()
                == "write",
            },
            "intent_plan": resolved_intent_plan.public_dict(),
            "files": [file_info.public_dict() for file_info in files],
            "file_capability_profiles": capability_profiles,
            "context_snippets": snippets[:10],
            "recipe_skeleton": skeleton,
            "execution_plan_schema": whitebox_execution_plan_schema(),
            "execution_brief_schema": self._execution_brief_schema(),
            "tool_design_protocol": TOOL_DESIGN_PROTOCOL,
        }
        if isinstance(request.options, dict):
            memory_context = str(request.options.get("memory_context") or "").strip()
            if memory_context:
                context["memory_context"] = _preview(memory_context, 6000)
        if resolved_known_tool_gap:
            context["known_native_tool_gap"] = resolved_known_tool_gap
        if followup_context:
            context["followup_context"] = followup_context
        messages: List[Dict[str, Any]] = []
        for item in request.history[-6:]:
            if not isinstance(item, dict):
                continue
            role = str(item.get("role") or "user").strip().lower()
            content = str(item.get("content") or item.get("text") or "").strip()
            if content and role in {"user", "assistant", "model"}:
                messages.append(
                    {
                        "role": "model" if role == "assistant" else role,
                        "content": _preview(content, 1500),
                    }
                )
        prompt_prefix = "请完成这个文件任务。"
        if str(followup_context.get("kind") or "").strip() == "review_last_task":
            constFollowupAction = (
                str(followup_context.get("followup_action") or "").strip().lower()
            )
            if constFollowupAction == "apply":
                prompt_prefix = (
                    "用户要求把上一轮文件任务中已经给出的建议直接应用到目标文件。"
                    "请把它视为同一任务的写回续跑，优先沿用上一轮建议、目标文件和已知约束，不要重新从头分析。"
                )
            elif constFollowupAction == "improve":
                prompt_prefix = (
                    "用户要求在上一轮文件任务结果基础上继续优化。" "请把它视为同一任务的后续处理回合，先说明你准备如何改进，再继续处理。"
                )
                if _followup_has_prior_excel_docx_insert(followup_context):
                    prompt_prefix += " 上一轮已经有实际 file.changed 记录表明目标 DOCX 插入过 Excel 表格；请先基于这些已写入结果判断缺口，不要重复同一插表。"
            else:
                prompt_prefix = (
                    "用户正在对上一轮文件任务结果提出反馈。"
                    "请先回答上一轮结果为什么会这样、哪里可能有问题，以及是否需要重做。"
                    "除非用户已经明确提出新的文件修改要求，否则不要把这条消息当成新的文件执行任务。"
                )
        messages.append(
            {
                "role": "user",
                "content": prompt_prefix
                + "上下文如下：\n"
                + json.dumps(context, ensure_ascii=False, indent=2),
            }
        )
        return messages

    def _followup_context(self, request: FileTaskRequest) -> Dict[str, Any]:
        if not isinstance(request.options, dict):
            return {}
        value = request.options.get("followup_context")
        if not isinstance(value, dict):
            return {}

        cleaned: Dict[str, Any] = {}
        for key in (
            "kind",
            "source",
            "followup_action",
            "user_feedback",
            "previous_run_id",
            "previous_task_summary",
            "previous_task_status",
            "previous_task_timestamp",
            "previous_user_request",
            "previous_task_request",
            "previous_task_mode",
            "previous_task_request_kind",
            "previous_task_family",
            "previous_task_operation_kind",
            "previous_task_execution_mode",
            "previous_task_selected_recipe",
            "previous_task_output_mode",
            "previous_task_intent_strategy",
            "previous_task_intent_can_apply",
            "previous_task_intent_requires_confirmation",
            "previous_task_target_file_type",
            "previous_completed_task",
        ):
            text = str(value.get(key) or "").strip()
            if text:
                cleaned[key] = _preview(text, 2000)
        previous_task_file_changes = _sanitize_followup_file_changes(
            value.get("previous_task_file_changes")
        )
        if previous_task_file_changes:
            cleaned["previous_task_file_changes"] = previous_task_file_changes
        return cleaned

    def _call_model(
        self,
        *,
        request: FileTaskRequest,
        messages: List[Dict[str, Any]],
        system: str,
        tools: List[Dict[str, Any]],
    ) -> Dict[str, Any]:
        if hasattr(self._model_client, "call"):
            return self._model_client.call(request=request, messages=messages, system=system, tools=tools)  # type: ignore[union-attr]
        return self._model_client(request=request, messages=messages, system=system, tools=tools)  # type: ignore[misc]

    def _normalize_model_response(
        self,
        response: Dict[str, Any],
        tool_defs: List[Dict[str, Any]],
    ) -> tuple[str, List[Dict[str, Any]]]:
        if not isinstance(response, dict):
            return str(response or ""), []
        content_text = str(response.get("content") or "")
        tool_calls = response.get("tool_calls") or []
        normalized = self._coerce_tool_calls(tool_calls)
        if not normalized and content_text:
            allowed = {str(definition.get("name") or "") for definition in tool_defs}
            content_text, normalized = parse_task_tool_calls(content_text, allowed)
        return content_text.strip(), normalized

    def _coerce_tool_calls(self, raw_tool_calls: Any) -> List[Dict[str, Any]]:
        items = raw_tool_calls if isinstance(raw_tool_calls, list) else [raw_tool_calls]
        normalized: List[Dict[str, Any]] = []
        for item in items:
            if not isinstance(item, dict):
                continue
            function_payload = (
                item.get("function") if isinstance(item.get("function"), dict) else {}
            )
            tool_name = str(
                item.get("name")
                or item.get("tool_name")
                or function_payload.get("name")
                or ""
            ).strip()
            if not tool_name:
                continue
            tool_args = item.get("args")
            if tool_args is None:
                tool_args = item.get("arguments")
            if tool_args is None and function_payload:
                tool_args = function_payload.get("arguments")
            if isinstance(tool_args, str):
                try:
                    tool_args = json.loads(tool_args)
                except Exception:
                    tool_args = {}
            if not isinstance(tool_args, dict):
                tool_args = {}
            normalized.append(
                {
                    "id": str(item.get("id") or uuid.uuid4().hex[:8]),
                    "name": tool_name,
                    "args": tool_args,
                }
            )
        return normalized

    def _tool_batch_signature(self, tool_calls: List[Dict[str, Any]]) -> str:
        if not tool_calls:
            return ""
        safe_calls = [
            {"name": item.get("name"), "args": item.get("args") or {}}
            for item in tool_calls
        ]
        try:
            return json.dumps(
                safe_calls, ensure_ascii=False, sort_keys=True, default=str
            )
        except Exception:
            return str(safe_calls)

    def _extract_file_changes(
        self, tool_name: str, tool_args: Dict[str, Any], result: Any
    ) -> List[Dict[str, Any]]:
        changes: List[Dict[str, Any]] = []
        structured = parse_file_change(tool_name, tool_args, result)
        if structured:
            changes.append(structured)
        if tool_name == "run_python_code":
            for path in extract_koto_paths(result, _KOTO_CREATED_MARKER):
                changes.append(
                    {
                        "path": path,
                        "file_type": Path(path).suffix.lstrip(".").lower(),
                        "operation": "run_python_code",
                        "summary": f"Python 代码创建了 {Path(path).name}",
                        "preview": "",
                        "change_type": "create",
                        "focus": True,
                    }
                )
            for path in extract_koto_paths(result, _KOTO_MODIFIED_MARKER):
                changes.append(
                    {
                        "path": path,
                        "file_type": Path(path).suffix.lstrip(".").lower(),
                        "operation": "run_python_code",
                        "summary": f"Python 代码更新了 {Path(path).name}",
                        "preview": "",
                        "change_type": "modify",
                        "focus": True,
                    }
                )
        return changes

    def _tool_result_for_model(self, tool_name: str, result: Any) -> Any:
        if tool_name != "run_python_code" or not isinstance(result, dict):
            return result

        sanitized: Dict[str, Any] = {}
        summary = str(result.get("summary") or "").strip()
        stdout = str(result.get("stdout") or "").strip()
        stderr = str(result.get("stderr") or "").strip()
        error = str(result.get("error") or "").strip()
        if summary:
            sanitized["summary"] = summary
        if stdout:
            sanitized["stdout"] = stdout
        if stderr:
            sanitized["stderr"] = stderr
        if error:
            sanitized["error"] = error

        created = extract_koto_paths(result, _KOTO_CREATED_MARKER)
        modified = extract_koto_paths(result, _KOTO_MODIFIED_MARKER)
        if created:
            sanitized["created_paths"] = created
        if modified:
            sanitized["modified_paths"] = modified

        artifacts = extract_sandbox_artifacts(result)
        if artifacts:
            sanitized["generated_files"] = [
                artifact.get("name") for artifact in artifacts
            ]
            sanitized["generated_file_count"] = len(artifacts)
        return sanitized or {"summary": "(no output)"}

    def _tool_feedback_for_model(
        self,
        tool_name: str,
        tool_args: Dict[str, Any],
        model_result: Any,
        *,
        success: bool,
        blocked: bool = False,
        skipped: bool = False,
        invalid: bool = False,
    ) -> str:
        if success and not blocked and not skipped and not invalid:
            return _preview(stringify_result(model_result), 6000)

        if invalid:
            failure_reason = "invalid_tool"
            next_action = (
                "这个工具当前不在 Koto 文件任务 allowlist 中。"
                "请改用现有 allowlist 工具，或在确实缺原生能力时返回 tool_gap；"
                "不要重复调用同一个无效工具。"
            )
        elif blocked:
            failure_reason = "blocked"
            next_action = (
                "这次调用被运行时拦截。请根据 error 或 summary 改用允许的原生工具或修改方案；" "不要重复完全相同的调用。"
            )
        elif skipped:
            failure_reason = "skipped"
            next_action = "这次调用被运行时跳过。请先理解跳过原因，再修改目标或方案；" "不要原样重复同一个调用。"
        else:
            failure_reason = "execution_failed"
            next_action = (
                "上一个工具调用执行失败。请先根据 error、stderr、stdout 和 summary 判断错在哪；"
                "只有在参数、代码或方案已经改变时才允许再次调用，不要重复完全相同的调用。"
            )

        payload = {
            "tool_name": tool_name,
            "tool_args": self._truncate_tool_feedback_value(tool_args),
            "success": bool(success),
            "blocked": bool(blocked),
            "skipped": bool(skipped),
            "invalid_tool": bool(invalid),
            "failure_reason": failure_reason,
            "retry_same_call_allowed": False,
            "result": self._truncate_tool_feedback_value(model_result),
            "next_action": next_action,
        }
        return json.dumps(payload, ensure_ascii=False, default=str)

    def _readonly_write_tool_block_message(
        self,
        tool_name: str,
        request: FileTaskRequest,
        output_mode: str,
    ) -> str:
        mode_label = self._output_mode_label(output_mode)
        return (
            f"当前任务模式是“{mode_label}”，用户没有授权写入文件；已拦截写入工具 {tool_name}。"
            f"请不要再调用写入工具，直接基于已读取内容回答用户任务：{request.task}"
        )

    def _readonly_run_python_write_block_message(
        self,
        tool_args: Dict[str, Any],
        request: FileTaskRequest,
        output_mode: str,
    ) -> str:
        code = str((tool_args or {}).get("code") or "")
        if not code.strip():
            return ""
        has_strong_write = any(
            pattern.search(code) for pattern in _RUN_PYTHON_STRONG_WRITE_PATTERNS
        )
        has_artifact_write = any(
            pattern.search(code) for pattern in _RUN_PYTHON_ARTIFACT_WRITE_PATTERNS
        )
        explicit_readonly = self._has_readonly_write_negation(request.task)
        if not has_strong_write and not (explicit_readonly and has_artifact_write):
            return ""
        mode_label = self._output_mode_label(output_mode)
        return (
            f"当前任务模式是“{mode_label}”，用户没有授权写入文件；已拦截 run_python_code 中的文件写入/保存代码。"
            f"请只用 Python 读取、计算和汇总，或直接输出分析结论；不要创建、保存、覆盖或移动文件。用户任务：{request.task}"
        )

    def _extract_tool_runtime_outcome(self, result: Any) -> Optional[Dict[str, Any]]:
        payload = result if isinstance(result, dict) else _json_payload(result)
        if not isinstance(payload, dict):
            return None

        raw_status = str(payload.get("status") or "").strip().lower()
        awaiting_confirmation = (
            bool(payload.get("awaiting_confirmation"))
            or raw_status == "awaiting_confirmation"
        )
        artifact = (
            payload.get("next_action_artifact")
            if isinstance(payload.get("next_action_artifact"), dict)
            else None
        )
        summary = str(payload.get("summary") or payload.get("error") or "").strip()
        suggested_next_step = str(payload.get("suggested_next_step") or "").strip()
        status = "awaiting_confirmation" if awaiting_confirmation else raw_status
        if not status and artifact is None:
            return None

        outcome: Dict[str, Any] = {
            "status": status or "needs_attention",
            "summary": summary,
        }
        if suggested_next_step:
            outcome["suggested_next_step"] = suggested_next_step
        if artifact is not None:
            outcome["next_action_artifact"] = artifact
        return outcome

    def _tool_runtime_status(
        self, tool_runtime_outcome: Optional[Dict[str, Any]]
    ) -> str:
        if not isinstance(tool_runtime_outcome, dict):
            return ""
        return str(tool_runtime_outcome.get("status") or "").strip().lower()

    def _truncate_tool_feedback_value(self, value: Any, *, depth: int = 0) -> Any:
        if isinstance(value, str):
            return _preview(value, 2400 if depth == 0 else 1600)
        if value is None or isinstance(value, (int, float, bool)):
            return value
        if depth >= 2:
            try:
                return _preview(
                    json.dumps(value, ensure_ascii=False, default=str), 1600
                )
            except Exception:
                return _preview(str(value), 1600)
        if isinstance(value, dict):
            trimmed: Dict[str, Any] = {}
            for index, (key, item) in enumerate(value.items()):
                if index >= 20:
                    trimmed["__truncated__"] = True
                    break
                trimmed[str(key)] = self._truncate_tool_feedback_value(
                    item, depth=depth + 1
                )
            return trimmed
        if isinstance(value, (list, tuple)):
            items = [
                self._truncate_tool_feedback_value(item, depth=depth + 1)
                for item in list(value)[:20]
            ]
            if len(value) > 20:
                items.append("__truncated__")
            return items
        return _preview(str(value), 1600)

    def _tool_artifacts(self, tool_name: str, result: Any) -> List[Dict[str, Any]]:
        if tool_name != "run_python_code":
            return []
        return extract_sandbox_artifacts(result)

    def _should_attempt_repair(
        self,
        check_payload: Optional[Dict[str, Any]],
        *,
        round_index: int,
        repair_attempts: int,
    ) -> bool:
        if not isinstance(check_payload, dict):
            return False
        if repair_attempts >= _MAX_VERIFY_REPAIR_ATTEMPTS:
            return False
        if round_index >= self._max_rounds:
            return False
        if bool(check_payload.get("passed")):
            return False
        status = str(check_payload.get("status") or "").strip().lower()
        return status in {"needs_attention", "no_file_change", "quality_gate_failed"}

    def _change_operations(self, file_changes: List[Dict[str, Any]]) -> set[str]:
        return {
            str(change.get("operation") or "").strip()
            for change in file_changes
            if str(change.get("operation") or "").strip()
        }

    def _change_sum_int(self, file_changes: List[Dict[str, Any]], key: str) -> int:
        total = 0
        for change in file_changes:
            try:
                total += int(change.get(key) or 0)
            except Exception:
                continue
        return total

    def _target_or_request_type(
        self, request: FileTaskRequest, file_changes: List[Dict[str, Any]]
    ) -> str:
        target_type = Path(str(request.target_path or "")).suffix.lstrip(".").lower()
        if target_type:
            return target_type
        for change in file_changes:
            candidate = (
                str(
                    change.get("file_type")
                    or Path(str(change.get("path") or "")).suffix.lstrip(".")
                )
                .lower()
                .strip()
            )
            if candidate:
                return candidate
        for file_info in request.files or []:
            if file_info.target:
                candidate = (
                    str(
                        file_info.type
                        or Path(str(file_info.path or file_info.name)).suffix.lstrip(
                            "."
                        )
                    )
                    .lower()
                    .strip()
                )
                if candidate:
                    return candidate
        return ""

    def _quality_gate_result(
        self,
        *,
        criterion: str,
        passed: bool,
        detail: str,
        priority: str = "high",
    ) -> Dict[str, Any]:
        return {
            "criterion": criterion,
            "passed": bool(passed),
            "detail": detail,
            "priority": priority,
        }

    def _evaluate_task_quality_gate(
        self,
        request: FileTaskRequest,
        file_changes: List[Dict[str, Any]],
        *,
        write_intent: bool,
        output_mode: str,
    ) -> Dict[str, Any]:
        if not write_intent:
            return {"passed": True, "criteria_results": [], "remaining": []}

        operations = self._change_operations(file_changes)
        target_type = self._target_or_request_type(request, file_changes)
        paragraphs_written = self._change_sum_int(file_changes, "paragraphs_written")
        images_inserted = self._change_sum_int(file_changes, "images_inserted")
        rows_written = self._change_sum_int(file_changes, "rows_written")
        cells_written = self._change_sum_int(file_changes, "cells_written")
        slides_updated = self._change_sum_int(file_changes, "slides_updated")
        slides_added = self._change_sum_int(file_changes, "slides_added")
        slides_designed = self._change_sum_int(file_changes, "slides_designed")
        text_shapes_styled = self._change_sum_int(file_changes, "text_shapes_styled")
        annotations_added = self._change_sum_int(file_changes, "annotations_added")
        differences_detected = self._change_sum_int(
            file_changes, "differences_detected"
        )
        comments_removed = self._change_sum_int(file_changes, "comments_removed")
        revisions_accepted = self._change_sum_int(file_changes, "revisions_accepted")
        paragraphs_rewritten = self._change_sum_int(
            file_changes, "paragraphs_rewritten"
        )
        task_text = str(request.task or "")

        criteria: List[Dict[str, Any]] = []
        metric_values = {
            "paragraphs_written": paragraphs_written,
            "images_inserted": images_inserted,
            "rows_written": rows_written,
            "slides_updated": slides_updated,
            "slides_added": slides_added,
            "slides_designed": slides_designed,
            "text_shapes_styled": text_shapes_styled,
            "annotations_added": annotations_added,
            "differences_detected": differences_detected,
            "comments_removed": comments_removed,
            "revisions_accepted": revisions_accepted,
            "paragraphs_rewritten": paragraphs_rewritten,
            "cells_written": cells_written,
        }
        recipe_match = select_task_recipe(
            request, request.files or [], write_intent=write_intent
        )
        seen_recipe_criteria: set[str] = set()
        if recipe_match:
            for gate in recipe_match.recipe.quality_gates:
                criterion = str(gate.get("criterion") or "").strip()
                if not criterion or criterion in seen_recipe_criteria:
                    continue
                seen_recipe_criteria.add(criterion)
                operation = str(gate.get("operation") or "").strip()
                any_operation = {
                    str(item).strip()
                    for item in gate.get("any_operation") or []
                    if str(item).strip()
                }
                metric_name = str(gate.get("metric") or "").strip()
                actual = int(metric_values.get(metric_name, 0) or 0)
                minimum = int(gate.get("minimum") or 0)
                if any_operation:
                    passed = bool(operations.intersection(any_operation)) and actual >= minimum
                    detail = str(gate.get("detail") or "").format(
                        operations=", ".join(sorted(operations)) or "无", actual=actual
                    )
                else:
                    passed = (
                        not operation or operation in operations
                    ) and actual >= minimum
                    detail = str(gate.get("detail") or "").format(
                        actual=actual, minimum=minimum
                    )
                criteria.append(
                    self._quality_gate_result(
                        criterion=criterion,
                        passed=passed,
                        detail=detail or criterion,
                        priority=str(gate.get("priority") or "high"),
                    )
                )

        if criteria:
            failed = [item for item in criteria if not item.get("passed")]
            return {
                "passed": not failed,
                "criteria_results": criteria,
                "remaining": [
                    str(item.get("detail") or item.get("criterion")) for item in failed
                ],
            }

        if self._looks_like_financial_xlsx_docx_chart_report_task(
            request, request.files or []
        ):
            criteria.extend(
                [
                    self._quality_gate_result(
                        criterion="financial_report_has_narrative",
                        passed="write_docx_content" in operations
                        and paragraphs_written >= 8,
                        detail=f"财务图表报告应写入结构化分析段落；当前段落写入数：{paragraphs_written}。",
                        priority="critical",
                    ),
                    self._quality_gate_result(
                        criterion="financial_report_has_real_chart_image",
                        passed="insert_image_into_docx" in operations
                        and images_inserted >= 1,
                        detail=f"财务图表报告必须插入真实图表图片；当前图片写入数：{images_inserted}。",
                        priority="critical",
                    ),
                ]
            )
        elif target_type in {"docx", "doc"} and self._looks_like_chart_request(
            task_text
        ):
            criteria.append(
                self._quality_gate_result(
                    criterion="docx_chart_request_has_image",
                    passed="insert_image_into_docx" in operations
                    and images_inserted >= 1,
                    detail=f"用户要求图表/图片进入 Word；当前图片写入数：{images_inserted}。",
                    priority="critical",
                )
            )

        if target_type in {"docx", "doc"} and self._looks_like_docx_report_request(
            request, request.files or []
        ):
            narrative_minimum = 2 if self._looks_like_table_request(task_text) else 3
            criteria.append(
                self._quality_gate_result(
                    criterion="docx_report_has_narrative",
                    passed=(
                        "write_docx_content" in operations
                        and paragraphs_written >= narrative_minimum
                    )
                    or paragraphs_written >= narrative_minimum,
                    detail=(
                        "DOCX 报告/分析任务应写入可读文本结构；"
                        f"当前段落写入数：{paragraphs_written}，"
                        f"最低要求：{narrative_minimum}。"
                    ),
                    priority="high",
                )
            )

        if (
            target_type in {"docx", "doc"}
            and self._looks_like_table_request(task_text)
            and not self._looks_like_problem_analysis_request(task_text)
        ):
            criteria.append(
                self._quality_gate_result(
                    criterion="docx_table_request_has_table",
                    passed="insert_excel_as_docx_table" in operations
                    and rows_written > 0,
                    detail=f"用户要求表格数据进入 Word；当前表格写入行数：{rows_written}。",
                    priority="high",
                )
            )

        if target_type in {"docx", "doc"} and operations.intersection(
            {"compare_docx_and_annotate", "write_docx_comments"}
        ):
            criteria.append(
                self._quality_gate_result(
                    criterion="docx_compare_has_difference_annotations",
                    passed=annotations_added > 0,
                    detail=f"DOCX 对比标注任务必须写入真实差异批注；当前批注数：{annotations_added}。",
                    priority="critical",
                )
            )

        if self._looks_like_ppt_slide_write_request(request, request.files or []):
            criteria.append(
                self._quality_gate_result(
                    criterion="ppt_request_has_slide_write",
                    passed=bool(
                        operations.intersection(
                            {
                                "add_pptx_slides",
                                "write_pptx_slides",
                                "design_pptx_theme_layout",
                            }
                        )
                    ),
                    detail=f"PPT 任务应产生幻灯片写入/更新操作；当前操作：{', '.join(sorted(operations)) or '无'}。",
                    priority="critical",
                )
            )

        if target_type in {"docx", "doc"} and not criteria:
            docx_write_ops = {
                "write_docx_content",
                "insert_excel_as_docx_table",
                "insert_image_into_docx",
                "annotate_file",
                "compare_docx_and_annotate",
                "clear_docx_review_marks",
                "rewrite_docx_paragraph_window",
            }
            docx_metric_total = (
                paragraphs_written
                + images_inserted
                + rows_written
                + annotations_added
                + differences_detected
                + comments_removed
                + revisions_accepted
                + paragraphs_rewritten
            )
            run_python_docx_writeback = (
                "run_python_code" in operations
                and bool(file_changes)
                and (
                    self._looks_like_polish_request(task_text)
                    or self._looks_like_translation_request(task_text)
                )
            )
            criteria.append(
                self._quality_gate_result(
                    criterion="generic_docx_has_native_write",
                    passed=bool(operations.intersection(docx_write_ops))
                    and docx_metric_total > 0
                    or run_python_docx_writeback,
                    detail=(
                        "DOCX 写入任务必须产生可核验的 Word 原生写入指标；"
                        f"当前操作：{', '.join(sorted(operations)) or '无'}，"
                        f"段落/图片/表格/批注/修订指标合计：{docx_metric_total}。"
                    ),
                    priority="high",
                )
            )

        if target_type in {"pptx", "ppt"} and not criteria:
            pptx_write_ops = {
                "add_pptx_slides",
                "write_pptx_slides",
                "design_pptx_theme_layout",
            }
            pptx_metric_total = (
                slides_updated + slides_added + slides_designed + text_shapes_styled
            )
            criteria.append(
                self._quality_gate_result(
                    criterion="generic_pptx_has_native_write",
                    passed=bool(operations.intersection(pptx_write_ops))
                    and pptx_metric_total > 0,
                    detail=(
                        "PPTX 写入任务必须产生可核验的幻灯片写入、更新或设计指标；"
                        f"当前操作：{', '.join(sorted(operations)) or '无'}，"
                        f"幻灯片/文本样式指标合计：{pptx_metric_total}。"
                    ),
                    priority="high",
                )
            )

        if target_type in {"xlsx", "xlsm", "csv"} and not criteria:
            spreadsheet_metric_total = rows_written + cells_written
            criteria.append(
                self._quality_gate_result(
                    criterion="generic_spreadsheet_has_native_write",
                    passed="write_sheet_data" in operations
                    and spreadsheet_metric_total > 0,
                    detail=(
                        "表格写入任务必须产生可核验的单元格/行写入指标；"
                        f"当前操作：{', '.join(sorted(operations)) or '无'}，"
                        f"行/单元格指标合计：{spreadsheet_metric_total}。"
                    ),
                    priority="high",
                )
            )

        failed = [item for item in criteria if not item.get("passed")]
        return {
            "passed": not failed,
            "criteria_results": criteria,
            "remaining": [
                str(item.get("detail") or item.get("criterion")) for item in failed
            ],
        }

    def _repair_retry_message(
        self,
        request: FileTaskRequest,
        check_payload: Dict[str, Any],
        file_changes: List[Dict[str, Any]],
    ) -> str:
        lines = [
            "核验未通过，当前任务还不能结束。下一轮必须修复目标文件，而不是重复上一轮完全相同的调用。",
        ]
        status = str(check_payload.get("status") or "").strip()
        summary = str(check_payload.get("summary") or "").strip()
        if status:
            lines.append(f"当前核验状态：{status}")
        if summary:
            lines.append(f"核验摘要：{summary}")
        if request.target_path:
            lines.append(f"目标文件：{request.target_path}")

        request_files = getattr(request, "files", []) or []
        recipe_match = select_task_recipe(request, request_files, write_intent=True)
        if recipe_match:
            lines.append(f"当前任务路线：{recipe_match.recipe.id}")
            if recipe_match.recipe.success_criteria:
                lines.append("本路线验收标准：")
                for criterion in recipe_match.recipe.success_criteria[:5]:
                    text = str(criterion or "").strip()
                    if text:
                        lines.append(f"- {text}")
        if self._looks_like_financial_xlsx_docx_chart_report_task(
            request, request_files
        ):
            lines.append(
                "财务预测图表写入修复要求：本任务不能只完成 Python 计算或打印 stdout。"
                "必须产生写入工具事件：write_docx_content 写入问题清单/分析结论，insert_image_into_docx 插入真实 PNG/JPG 图表。"
            )
            lines.append(
                "Excel 解析要求：如果 pandas 读出的列名是 Unnamed，不要用 df.columns 判断年份列；"
                "应扫描表格行，找到包含 2025E/2026E/2027E/2028E 等年份标签的 header row，"
                "再根据这些列抽取“收入合计、毛利合计、费用合计、净利润、销量”等指标。"
            )

        remaining = (
            check_payload.get("remaining")
            if isinstance(check_payload.get("remaining"), list)
            else []
        )
        if remaining:
            lines.append("仍需满足：")
            for index, item in enumerate(remaining[:5], start=1):
                text = str(item or "").strip()
                if text:
                    lines.append(f"{index}. {text}")

        if file_changes:
            lines.append("已观察到的文件变更：")
            for change in file_changes[-3:]:
                if not isinstance(change, dict):
                    continue
                change_summary = str(change.get("summary") or "").strip()
                path_text = str(
                    change.get("path") or change.get("file_path") or ""
                ).strip()
                if change_summary and path_text:
                    lines.append(f"- {path_text}: {change_summary}")
                elif change_summary:
                    lines.append(f"- {change_summary}")
                elif path_text:
                    lines.append(f"- {path_text}")

        lines.append("要求：先理解核验失败原因；只有当参数、代码、工具选择或写入位置已经改变时，才允许再次调用工具；修复后再结束。")
        return "\n".join(lines)

    def _code_output_preview(
        self, tool_name: str, result: Any, result_text: str
    ) -> str:
        if tool_name != "run_python_code" or not isinstance(result, dict):
            return _preview(result_text, 2000)

        parts: List[str] = []
        stdout = str(result.get("stdout") or "").strip()
        stderr = str(result.get("stderr") or "").strip()
        summary = str(result.get("summary") or "").strip()
        if stdout:
            parts.append(stdout)
        if stderr:
            parts.append(f"[stderr] {stderr}")
        if not parts and summary:
            parts.append(summary)
        return _preview("\n".join(parts) if parts else result_text, 2000)

    def _verify_task(
        self,
        request: FileTaskRequest,
        executor: ToolExecutor,
        file_changes: List[Dict[str, Any]],
        write_intent: bool,
        output_mode: str,
        model_failed: bool,
        readonly_fallback_used: bool = False,
        tool_runtime_outcome: Optional[Dict[str, Any]] = None,
        tool_gap: Optional[Dict[str, Any]] = None,
        next_action_artifact: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        if tool_gap:
            remaining = []
            if tool_gap.get("suggested_next_step"):
                remaining.append(str(tool_gap.get("suggested_next_step")))
            if isinstance(tool_gap.get("proposed_tool"), dict) and tool_gap[
                "proposed_tool"
            ].get("name"):
                remaining.append(
                    f"按 {TOOL_DESIGN_PROTOCOL} 评估并实现新工具：{tool_gap['proposed_tool']['name']}"
                )
            if not remaining:
                remaining = ["根据缺口说明补充 Koto 原生工具或调整任务范围"]
            return {
                "passed": False,
                "status": "tool_gap",
                "remaining": remaining,
                "tool_gap": tool_gap,
                "next_action_artifact": next_action_artifact,
            }
        runtime_status = self._tool_runtime_status(tool_runtime_outcome)
        if (
            runtime_status == "awaiting_confirmation"
            and write_intent
            and not file_changes
            and self._requires_file_change_before_pause(request)
        ):
            return {
                "passed": False,
                "status": "no_file_change",
                "summary": "任务请求分步写入并等待确认，但本步骤尚未产生任何文件变更。",
                "remaining": ["先调用真实写入工具更新目标文件，再进入等待确认状态"],
                "next_action_artifact": next_action_artifact,
                "criteria_results": [
                    {
                        "criterion": "write_before_stepwise_pause",
                        "passed": False,
                        "detail": "分步写入任务必须先产生 file.changed，再等待用户继续。",
                        "priority": "critical",
                    }
                ],
            }
        if runtime_status == "awaiting_confirmation":
            artifact = (
                tool_runtime_outcome.get("next_action_artifact")
                if isinstance(tool_runtime_outcome.get("next_action_artifact"), dict)
                else next_action_artifact
            )
            remaining: List[str] = []
            if isinstance(artifact, dict):
                suggested = str(
                    artifact.get("suggested_next_step") or artifact.get("summary") or ""
                ).strip()
                if suggested:
                    remaining.append(suggested)
            if not remaining:
                remaining = ["等待用户确认后继续下一步。"]
            return {
                "passed": False,
                "status": "awaiting_confirmation",
                "summary": str(
                    tool_runtime_outcome.get("summary") or "任务已暂停，等待用户确认继续。"
                ),
                "remaining": remaining,
                "next_action_artifact": artifact,
            }
        if runtime_status in {"blocked", "write_blocked"}:
            suggested = str(
                (tool_runtime_outcome or {}).get("suggested_next_step") or ""
            ).strip()
            remaining = [suggested] if suggested else ["关闭占用目标文件的程序或页签后重试。"]
            return {
                "passed": False,
                "status": runtime_status,
                "summary": str(
                    (tool_runtime_outcome or {}).get("summary") or "目标文件当前不可写。"
                ),
                "remaining": remaining,
                "next_action_artifact": next_action_artifact,
            }
        if write_intent and not file_changes:
            return {
                "status": "no_file_change",
                "summary": "任务包含写入意图，但没有任何工具报告文件变更。",
                "passed": False,
                "remaining": ["调用真实写入工具并确保产出 file.changed 事件"],
                "criteria_results": [
                    {
                        "criterion": "file_change_emitted",
                        "passed": False,
                        "detail": "任务包含写入意图，但没有任何工具报告文件变更。",
                        "priority": "critical",
                    }
                ],
            }
        if not write_intent and model_failed:
            return {
                "passed": False,
                "status": "model_unavailable",
                "summary": "模型不可用，已完成显式上下文读取但未生成 AI 分析。",
                "remaining": ["检查云端 API Key 或启动本地 Ollama 后重试"],
            }
        if not write_intent and readonly_fallback_used:
            return {
                "passed": True,
                "status": "context_summary_fallback",
                "summary": "模型不可用，已基于显式上下文生成可见内容摘要。",
                "remaining": ["恢复模型后可生成更完整的 AI 分析"],
            }

        if file_changes:
            verify_target_path = str(request.target_path or "").strip()
            verify_args = {
                "task_description": request.task,
                "file_states": json.dumps(
                    file_states_for_changes(file_changes), ensure_ascii=False
                ),
                "file_changes": json.dumps(file_changes, ensure_ascii=False),
                "target_path": verify_target_path,
                "model_mode": request.model_mode,
            }
            try:
                result = executor("verify_task_completion", verify_args)
                payload = _json_payload(result)
            except Exception as exc:
                logger.warning(
                    "[FileTaskRuntime] verify_task_completion failed: %s", exc
                )
                payload = {
                    "completed": False,
                    "summary": f"文件已变更，但 AI 核验工具不可用：{exc}",
                }

            if payload.get("error"):
                return {
                    "passed": False,
                    "status": "verify_error",
                    "summary": f"文件已变更，但核验工具返回错误：{payload.get('error')}",
                    "remaining": ["修复模型/核验工具配置后重新核验"],
                    "criteria_results": [
                        {
                            "criterion": "verification_tool_available",
                            "passed": False,
                            "detail": f"核验工具返回错误：{payload.get('error')}",
                            "priority": "critical",
                        }
                    ],
                }

            completed = payload.get("completed")
            passed = bool(completed) if completed is not None else True
            quality_gate = self._evaluate_task_quality_gate(
                request,
                file_changes,
                write_intent=write_intent,
                output_mode=output_mode,
            )
            verification_criteria = payload.get("criteria_results") or []
            combined_criteria = [
                *verification_criteria,
                *quality_gate.get("criteria_results", []),
            ]
            if not quality_gate.get("passed", True):
                remaining = list(quality_gate.get("remaining") or [])
                if payload.get("remaining_steps"):
                    remaining.extend(
                        str(item)
                        for item in payload.get("remaining_steps") or []
                        if str(item or "").strip()
                    )
                return {
                    "passed": False,
                    "status": "quality_gate_failed",
                    "summary": "文件已有变更，但未满足本任务的关键质量门禁。",
                    "confidence": payload.get("confidence"),
                    "remaining": remaining or ["补齐任务要求的关键产物后重新核验"],
                    "criteria_results": combined_criteria,
                }
            return {
                "passed": passed,
                "status": "verified" if passed else "needs_attention",
                "summary": str(
                    payload.get("summary") or ("文件变更已记录。" if passed else "核验未通过。")
                ),
                "confidence": payload.get("confidence"),
                "remaining": payload.get("remaining_steps")
                or ([] if passed else ["根据核验结果继续修复"]),
                "criteria_results": combined_criteria,
            }

        return {
            "passed": True,
            "status": "completed" if not model_failed else "context_only",
            "summary": (
                "已完成分析建议，当前未直接写入文件。" if output_mode == "hybrid" else "已完成只读任务，没有产生文件写入。"
            ),
            "remaining": [],
        }

    def _requires_file_change_before_pause(self, request: FileTaskRequest) -> bool:
        request_files = getattr(request, "files", []) or []
        recipe_match = select_task_recipe(request, request_files, write_intent=True)
        if not recipe_match:
            return False
        if recipe_match.recipe.quality_gates:
            return any(
                str(gate.get("operation") or "").strip()
                for gate in recipe_match.recipe.quality_gates
            )
        return any(
            "file.changed" in str(item or "")
            for item in recipe_match.recipe.success_criteria
        )
