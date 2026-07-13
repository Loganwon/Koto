# -*- coding: utf-8 -*-
"""Pure DOCX comparison helpers kept separate from the task-tools registry."""

from __future__ import annotations

import os
import re
from typing import Dict, List


def _docx_nonempty_paragraph_texts(path: str) -> List[str]:
    from docx import Document

    doc = Document(path)
    paragraphs: List[str] = []
    for paragraph in doc.paragraphs:
        text = re.sub(r"\s+", " ", str(paragraph.text or "")).strip()
        if text:
            paragraphs.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [
                re.sub(r"\s+", " ", str(cell.text or "")).strip() for cell in row.cells
            ]
            text = " | ".join(cell for cell in cells if cell)
            if text:
                paragraphs.append(text)
    return paragraphs


def _docx_diff_key(text: str) -> str:
    return re.sub(r"\s+", "", str(text or "")).strip().lower()


def _short_docx_diff_text(text: str, limit: int = 220) -> str:
    cleaned = re.sub(r"\s+", " ", str(text or "")).strip()
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[: limit - 1].rstrip() + "..."


def _docx_diff_comment(kind: str, other_text: str = "", target_text: str = "") -> str:
    other = _short_docx_diff_text(other_text, 140)
    target = _short_docx_diff_text(target_text, 140)
    if kind == "replace":
        if other and target:
            return f"另一份为：{other}\n本文件为：{target}"
        return f"内容修改：{target or other}"
    if kind == "insert":
        return f"本文件多出：{target}"
    if kind == "delete":
        return f"另一份有，本文件缺少：{other}"
    return target or other


_CONTRACT_RISK_RULES: tuple[tuple[str, tuple[str, ...], str], ...] = (
    (
        "付款/费用",
        ("付款", "支付", "价款", "费用", "发票", "逾期", "payment", "invoice", "fee"),
        "付款或费用条款发生变化，需确认金额、期限、开票和逾期责任是否可接受。",
    ),
    (
        "违约责任",
        ("违约", "违约金", "赔偿", "损害", "breach", "default", "penalty", "damages"),
        "违约或赔偿安排发生变化，可能影响责任承担和救济成本。",
    ),
    (
        "终止/解除",
        (
            "终止",
            "解除",
            "到期",
            "续约",
            "termination",
            "terminate",
            "renewal",
            "expire",
        ),
        "终止、解除或续约条款发生变化，需关注退出条件和通知期限。",
    ),
    (
        "责任限制",
        (
            "责任限制",
            "责任上限",
            "间接损失",
            "liability",
            "limitation",
            "cap",
            "indirect",
        ),
        "责任限制或损失范围发生变化，可能扩大或缩小一方承担的风险。",
    ),
    (
        "保密/数据",
        ("保密", "数据", "隐私", "confidential", "privacy", "data"),
        "保密、数据或隐私义务发生变化，需核对披露范围和保护责任。",
    ),
    (
        "知识产权",
        (
            "知识产权",
            "著作权",
            "许可",
            "授权",
            "ip",
            "intellectual property",
            "license",
        ),
        "知识产权或许可安排发生变化，需确认权利归属和使用范围。",
    ),
    (
        "争议解决",
        (
            "管辖",
            "仲裁",
            "适用法律",
            "诉讼",
            "jurisdiction",
            "arbitration",
            "governing law",
        ),
        "争议解决或适用法律发生变化，可能影响维权地点、成本和程序。",
    ),
    (
        "交付/验收",
        ("交付", "验收", "服务水平", "sla", "delivery", "acceptance", "service level"),
        "交付、验收或服务水平条款发生变化，需关注履约标准和验收责任。",
    ),
)


def _contract_risk_summary_from_annotations(
    annotations: List[Dict[str, str]], *, max_items: int = 5
) -> List[str]:
    matched: List[str] = []
    seen: set[str] = set()
    combined_items = [
        " ".join(
            str(annotation.get(key) or "")
            for key in ("原文片段", "批注内容", "修改原因")
        ).lower()
        for annotation in annotations
    ]
    combined = "\n".join(combined_items)
    for label, keywords, summary in _CONTRACT_RISK_RULES:
        if label in seen:
            continue
        if any(keyword.lower() in combined for keyword in keywords):
            seen.add(label)
            matched.append(f"{label}：{summary}")
        if len(matched) >= max_items:
            break
    return matched


def _unique_docx_anchor(text: str, seen: set[str]) -> str:
    anchor = _short_docx_diff_text(text, 180)
    if not anchor:
        return ""
    candidate = anchor
    while candidate and candidate in seen and len(candidate) > 18:
        candidate = candidate[:-8].rstrip()
    if candidate in seen:
        return ""
    seen.add(candidate)
    return candidate


def _changed_target_anchor(counterpart_text: str, target_text: str) -> str:
    import difflib

    target = str(target_text or "")
    counterpart = str(counterpart_text or "")
    if not target:
        return ""

    matcher = difflib.SequenceMatcher(None, counterpart, target, autojunk=False)
    spans: List[tuple[int, int]] = []
    for tag, _i1, _i2, j1, j2 in matcher.get_opcodes():
        if tag != "equal" and j1 != j2:
            spans.append((j1, j2))
    if not spans:
        return _short_docx_diff_text(target, 180)

    j1, j2 = max(spans, key=lambda item: item[1] - item[0])
    while j1 > 0 and target[j1 - 1].isspace():
        j1 -= 1
    while j2 < len(target) and target[j2 : j2 + 1].isspace():
        j2 += 1

    anchor = target[j1:j2].strip()
    if len(anchor) < 2:
        anchor = target[max(0, j1 - 4) : min(len(target), j2 + 4)].strip()
    if len(anchor) < 2:
        return _short_docx_diff_text(target, 180)
    if len(anchor) < 6:
        expanded = target[max(0, j1 - 2) : min(len(target), j2 + 2)].strip()
        if 2 <= len(expanded) <= 24:
            anchor = expanded
    return _short_docx_diff_text(anchor, 80)


def _build_docx_compare_annotations(
    counterpart_paragraphs: List[str],
    target_paragraphs: List[str],
    *,
    max_differences: int,
    target_label: str = "当前标注文档",
    counterpart_label: str = "另一份文件",
) -> tuple[List[Dict[str, str]], int]:
    import difflib

    counterpart_keys = [_docx_diff_key(text) for text in counterpart_paragraphs]
    target_keys = [_docx_diff_key(text) for text in target_paragraphs]
    matcher = difflib.SequenceMatcher(
        None, counterpart_keys, target_keys, autojunk=False
    )
    annotations: List[Dict[str, str]] = []
    seen_anchors: set[str] = set()
    differences_detected = 0

    def add_annotation(anchor_text: str, comment: str, reason: str) -> None:
        if len(annotations) >= max_differences:
            return
        anchor = _unique_docx_anchor(anchor_text, seen_anchors)
        if anchor:
            annotations.append(
                {
                    "原文片段": anchor,
                    "批注内容": comment,
                    "批注标签": "差异：",
                    "修改原因": reason,
                }
            )

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        differences_detected += max(i2 - i1, j2 - j1, 1)
        if tag == "replace":
            counterpart_slice = counterpart_paragraphs[i1:i2]
            target_slice = target_paragraphs[j1:j2]
            for offset in range(max(len(counterpart_slice), len(target_slice))):
                other_text = (
                    counterpart_slice[offset] if offset < len(counterpart_slice) else ""
                )
                target_text = target_slice[offset] if offset < len(target_slice) else ""
                if target_text:
                    add_annotation(
                        _changed_target_anchor(other_text, target_text) or target_text,
                        _docx_diff_comment("replace", other_text, target_text),
                        "",
                    )
        elif tag == "insert":
            for target_text in target_paragraphs[j1:j2]:
                add_annotation(
                    _short_docx_diff_text(target_text, 80),
                    _docx_diff_comment("insert", target_text=target_text),
                    "",
                )
        elif tag == "delete":
            missing_text = "；".join(
                _short_docx_diff_text(text, 120)
                for text in counterpart_paragraphs[i1:i2]
            )
            anchor_text = (
                target_paragraphs[j1]
                if 0 <= j1 < len(target_paragraphs)
                else (
                    target_paragraphs[j1 - 1]
                    if 0 <= j1 - 1 < len(target_paragraphs)
                    else ""
                )
            )
            if anchor_text:
                add_annotation(
                    anchor_text,
                    _docx_diff_comment("delete", other_text=missing_text),
                    "",
                )
        if len(annotations) >= max_differences:
            break
    return annotations, differences_detected


def _same_resolved_path(left: str, right: str) -> bool:
    try:
        return os.path.normcase(os.path.abspath(left)) == os.path.normcase(
            os.path.abspath(right)
        )
    except Exception:
        return os.path.normcase(str(left or "")) == os.path.normcase(str(right or ""))


def _docx_compare_annotation_candidates(
    original_resolved: str,
    revised_resolved: str,
    target_resolved: str,
    *,
    max_differences: int,
) -> tuple[List[Dict[str, str]], int]:
    original_paragraphs = _docx_nonempty_paragraph_texts(original_resolved)
    revised_paragraphs = _docx_nonempty_paragraph_texts(revised_resolved)
    target_paragraphs = _docx_nonempty_paragraph_texts(target_resolved)
    if not original_paragraphs and not revised_paragraphs:
        return [], 0
    counterpart_paragraphs = (
        revised_paragraphs
        if _same_resolved_path(target_resolved, original_resolved)
        else original_paragraphs
    )
    return _build_docx_compare_annotations(
        counterpart_paragraphs,
        target_paragraphs,
        max_differences=max_differences,
        target_label=(
            "被标注原文"
            if _same_resolved_path(target_resolved, original_resolved)
            else "被标注文档"
        ),
    )
