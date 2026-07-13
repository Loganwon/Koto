# -*- coding: utf-8 -*-
# Copyright (C) 2024-2026 Koto AI. All rights reserved.
# SPDX-License-Identifier: LicenseRef-Koto-Proprietary
from __future__ import annotations

import base64
import html
import io
import logging
import os
import re
import zipfile
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


def _css_color_to_hex(value: str) -> str | None:
    """Convert a CSS color string like '#aabbcc' or 'rgb(r,g,b)' to a 6-char hex string."""
    if not value:
        return None
    value = value.strip()
    if value.startswith("#"):
        v = value.lstrip("#")
        if len(v) == 3:
            v = "".join(c * 2 for c in v)
        return v.upper() if len(v) == 6 else None
    if value.startswith("rgb"):
        nums = re.findall(r"\d+", value)
        if len(nums) >= 3:
            return "{:02X}{:02X}{:02X}".format(int(nums[0]), int(nums[1]), int(nums[2]))
    return None


def _parse_css_inline(style_str: str) -> dict[str, str]:
    """Parse a CSS inline style string into a {property: value} dict."""
    result: dict[str, str] = {}
    for part in style_str.split(";"):
        part = part.strip()
        if ":" not in part:
            continue
        prop, _, val = part.partition(":")
        result[prop.strip().lower()] = val.strip()
    return result


def _css_length_to_pt(value: str) -> float | None:
    """Convert a CSS length string (pt, px, em, cm, mm, in) to points.

    Returns None if the value cannot be parsed.
    """
    if not value:
        return None
    value = value.strip()
    try:
        if value.endswith("pt"):
            return float(value[:-2])
        if value.endswith("px"):
            return float(value[:-2]) * 0.75
        if value.endswith("em"):
            return float(value[:-2]) * 12  # 1em ≈ 12pt default
        if value.endswith("cm"):
            return float(value[:-2]) * 28.3465
        if value.endswith("mm"):
            return float(value[:-2]) * 2.83465
        if value.endswith("in"):
            return float(value[:-2]) * 72
    except (ValueError, TypeError):
        return None
    return None


def _set_run_east_asian_font(run: Any, font_name: str) -> None:
    """Set East Asian font on a run via direct XML manipulation."""
    try:
        from docx.oxml.ns import qn as _qn

        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(_qn("w:rFonts"))
        if rFonts is None:
            from lxml import etree

            rFonts = etree.SubElement(rPr, _qn("w:rFonts"))
        rFonts.set(_qn("w:eastAsia"), font_name)
    except Exception:
        pass


def _apply_run_inline(run: Any, css: dict[str, str]) -> None:
    """Apply inline CSS to a python-docx Run (font, bold, italic, color, size, etc)."""
    from docx.shared import Pt, RGBColor

    # ── Font family ──────────────────────────────────────────────────────
    ff = css.get("font-family", "")
    if ff:
        # Parse font-family: may be "'Calibri','SimSun'" or "SimSun" etc.
        fonts = [f.strip().strip("'\"") for f in ff.split(",") if f.strip()]
        if fonts:
            run.font.name = fonts[0]
            # If there's a second font, treat it as East Asian font
            if len(fonts) >= 2:
                _set_run_east_asian_font(run, fonts[1])
            else:
                # Check if the single font is a CJK font — also set as eastAsia
                _CJK_FONTS = {
                    "SimHei",
                    "SimSun",
                    "KaiTi",
                    "FangSong",
                    "Microsoft YaHei",
                    "STZhongsong",
                    "STSong",
                    "STHeiti",
                    "STKaiti",
                    "STFangsong",
                    "FZShuSong-Z01",
                    "FZHei-B01",
                    "NSimSun",
                    "DengXian",
                    "YouYuan",
                    "LiSu",
                }
                if fonts[0] in _CJK_FONTS:
                    _set_run_east_asian_font(run, fonts[0])

    # ── Bold / Italic / Underline ────────────────────────────────────────
    if css.get("font-weight") in ("bold", "700", "800", "900"):
        run.bold = True
    if css.get("font-style") == "italic":
        run.italic = True
    td = css.get("text-decoration", "")
    if "underline" in td:
        run.underline = True
    if "line-through" in td:
        run.font.strike = True

    # ── Superscript / Subscript ──────────────────────────────────────────
    va = css.get("vertical-align", "").lower()
    if va == "super":
        run.font.superscript = True
    elif va == "sub":
        run.font.subscript = True

    # ── Text color ───────────────────────────────────────────────────────
    color_hex = _css_color_to_hex(css.get("color", ""))
    if color_hex:
        r, g, b = (
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16),
        )
        run.font.color.rgb = RGBColor(r, g, b)

    # ── Highlight / background color ─────────────────────────────────────
    bg = css.get("background-color", "") or css.get("background", "")
    bg_hex = _css_color_to_hex(bg)
    if bg_hex:
        try:
            from docx.oxml.ns import qn as _qn
            from lxml import etree

            rPr = run._element.get_or_add_rPr()
            shd = rPr.find(_qn("w:shd"))
            if shd is None:
                shd = etree.SubElement(rPr, _qn("w:shd"))
            shd.set(_qn("w:val"), "clear")
            shd.set(_qn("w:color"), "auto")
            shd.set(_qn("w:fill"), bg_hex.upper())
        except Exception:
            pass

    # ── Font size ────────────────────────────────────────────────────────
    fs_pt = _css_length_to_pt(css.get("font-size", ""))
    if fs_pt is not None and fs_pt > 0:
        run.font.size = Pt(fs_pt)

    # ── Letter spacing ───────────────────────────────────────────────────
    ls = css.get("letter-spacing", "")
    ls_pt = _css_length_to_pt(ls)
    if ls_pt is not None:
        try:
            from docx.oxml.ns import qn as _qn

            rPr = run._element.get_or_add_rPr()
            spacing = rPr.find(_qn("w:spacing"))
            if spacing is None:
                from lxml import etree

                spacing = etree.SubElement(rPr, _qn("w:spacing"))
            # OOXML spacing is in half-points (1/144 inch)
            spacing.set(_qn("w:val"), str(int(ls_pt * 20)))
        except Exception:
            pass


def _apply_para_format(para: Any, pcss: dict[str, str]) -> None:
    """Apply paragraph-level CSS properties to a python-docx Paragraph."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Twips

    # Text alignment
    align = pcss.get("text-align", "")
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "left":
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pf = para.paragraph_format

    # Space before (margin-top)
    mt_pt = _css_length_to_pt(pcss.get("margin-top", ""))
    if mt_pt is not None and mt_pt >= 0:
        pf.space_before = Pt(mt_pt)

    # Space after (margin-bottom)
    mb_pt = _css_length_to_pt(pcss.get("margin-bottom", ""))
    if mb_pt is not None and mb_pt >= 0:
        pf.space_after = Pt(mb_pt)

    # Line spacing
    lh = pcss.get("line-height", "")
    if lh:
        try:
            # Only use _css_length_to_pt if the value has a unit suffix
            has_unit = any(
                lh.rstrip().endswith(u) for u in ("pt", "px", "em", "cm", "mm", "in")
            )
            if has_unit:
                lh_pt = _css_length_to_pt(lh)
                if lh_pt is not None and lh_pt > 0:
                    from docx.enum.text import WD_LINE_SPACING

                    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    pf.line_spacing = Pt(lh_pt)
            else:
                # Unitless multiplier (e.g. "1.15", "1.5")
                pf.line_spacing = float(lh)
        except (ValueError, TypeError):
            pass

    # First-line indent (text-indent)
    ti_pt = _css_length_to_pt(pcss.get("text-indent", ""))
    if ti_pt is not None and ti_pt > 0:
        pf.first_line_indent = Pt(ti_pt)

    # Left indent (padding-left or margin-left)
    pl = pcss.get("padding-left", "") or pcss.get("margin-left", "")
    pl_pt = _css_length_to_pt(pl)
    if pl_pt is not None and pl_pt > 0:
        pf.left_indent = Pt(pl_pt)

    # Right indent (padding-right or margin-right)
    pr = pcss.get("padding-right", "") or pcss.get("margin-right", "")
    pr_pt = _css_length_to_pt(pr)
    if pr_pt is not None and pr_pt > 0:
        pf.right_indent = Pt(pr_pt)


def _add_paragraph_from_tag(doc: Any, tag: Any) -> None:
    """Convert a single <p>/<h1>-<h6>/<li> BS4 tag into a python-docx paragraph."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    tag_name = tag.name.lower()
    if tag_name in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(tag_name[1])
        para = doc.add_heading("", level=level)
    else:
        para = doc.add_paragraph()

    # Paragraph-level formatting from style attribute
    style_str = tag.get("style", "")
    pcss: dict[str, str] = {}
    if style_str:
        pcss = _parse_css_inline(style_str)
        _apply_para_format(para, pcss)

    # Walk inline children (text nodes + <strong>, <em>, <span>, <a>, <u>)
    for child in tag.children:
        if hasattr(child, "name") and child.name is None:
            # NavigableString
            text = str(child)
            if text:
                run = para.add_run(text)
                # Apply paragraph-level font info to bare text nodes
                _apply_run_inline(run, pcss)
        elif hasattr(child, "name"):
            name = child.name.lower() if child.name else ""
            # Inline image — void element with no text; insert as inline picture
            if name == "img":
                src = child.get("src", "")
                img_bytes = _resolve_img_src(src)
                if img_bytes:
                    try:
                        from docx.shared import Inches, Pt

                        run = para.add_run()

                        img_css = _parse_css_inline(child.get("style", ""))
                        w_str = img_css.get("width", child.get("width", ""))
                        w_val = None
                        if w_str:
                            try:
                                if w_str.endswith("px"):
                                    w_val = Pt(float(w_str.replace("px", "")) * 0.75)
                                elif w_str.endswith("in"):
                                    w_val = Inches(float(w_str.replace("in", "")))
                                elif w_str.isdigit():
                                    w_val = Pt(float(w_str) * 0.75)
                            except Exception:
                                import logging

                                logging.getLogger(__name__).warning(
                                    "Silenced exception caught", exc_info=True
                                )
                                pass

                        if w_val:
                            run.add_picture(io.BytesIO(img_bytes), width=w_val)
                        else:
                            run.add_picture(io.BytesIO(img_bytes), width=Inches(3))
                    except Exception:
                        pass
                continue
            text = child.get_text()
            if not text:
                continue
            run = para.add_run(text)
            child_css: dict[str, str] = _parse_css_inline(child.get("style", ""))
            # Merge parent paragraph CSS as fallback for font properties
            merged_css = dict(pcss)
            merged_css.update(child_css)
            if name in ("strong", "b"):
                run.bold = True
            if name in ("em", "i"):
                run.italic = True
            if name == "u":
                run.underline = True
            if name == "s":
                run.font.strike = True
            _apply_run_inline(run, merged_css)


def _set_cell_shading(cell: Any, fill_hex: str) -> None:
    """Set the background shading of a python-docx table cell via direct XML."""
    from lxml import etree

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing shd element
    for shd in tcPr.findall(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd"
    ):
        tcPr.remove(shd)
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    shd = etree.SubElement(tcPr, f"{{{WNS}}}shd")
    shd.set(f"{{{WNS}}}val", "clear")
    shd.set(f"{{{WNS}}}color", "auto")
    shd.set(f"{{{WNS}}}fill", fill_hex.upper())


def _apply_border_xml(tcBorders: Any, side: str, val: str, sz: str, color: str) -> None:
    """Helper to add/update an individual border to tcBorders xml."""
    from lxml import etree

    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    el = tcBorders.find(f"{{{WNS}}}{side}")
    if el is None:
        el = etree.SubElement(tcBorders, f"{{{WNS}}}{side}")
    if val == "none":
        el.set(f"{{{WNS}}}val", "none")
    else:
        el.set(f"{{{WNS}}}val", "single")
        el.set(f"{{{WNS}}}sz", sz)
        if color:
            el.set(f"{{{WNS}}}color", color.upper())


def _set_cell_borders(cell: Any, css: dict[str, str]) -> None:
    """Apply borders to a cell by parsing border/border-* CSS."""
    import re

    from lxml import etree

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    tcBorders = tcPr.find(f"{{{WNS}}}tcBorders")
    if tcBorders is None:
        tcBorders = etree.SubElement(tcPr, f"{{{WNS}}}tcBorders")

    # Helper to parse a shorthand border value like "1px solid #000000"
    def parse_border(val_str: str) -> tuple[str, str, str]:
        if not val_str or val_str == "none":
            return "none", "0", ""
        # defaults
        sz, color = "4", "auto"
        parts = val_str.split()
        for p in parts:
            if "px" in p:
                try:  # 1px is approx 4 eighths-of-a-point
                    sz = str(int(float(p.replace("px", "")) * 4))
                except Exception:
                    pass
            elif "pt" in p:
                try:  # 1pt is 8 eighths-of-a-point
                    sz = str(int(float(p.replace("pt", "")) * 8))
                except Exception:
                    pass
            elif p.startswith("#") or p.startswith("rgb"):
                h = _css_color_to_hex(p)
                if h:
                    color = h.upper()
        return "single", sz, color

    # Apply general border
    if "border" in css:
        val, sz, color = parse_border(css["border"])
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            _apply_border_xml(tcBorders, side, val, sz, color)

    # Override sides
    # We map CSS border sides to Word sides
    for css_prop, w_side in [
        ("border-top", "top"),
        ("border-bottom", "bottom"),
        ("border-left", "left"),
        ("border-right", "right"),
    ]:
        if css_prop in css:
            val, sz, color = parse_border(css[css_prop])
            _apply_border_xml(tcBorders, w_side, val, sz, color)


def _resolve_img_src(src: str) -> bytes | None:
    """Decode a base64 data URI or read a /api/v1/workspace/tmp_image/ temp file.

    Handles:
    - ``data:image/png;base64,<b64>`` — decodes the base64 payload
    - ``/api/v1/workspace/tmp_image/<session_id>/<filename>`` — reads the file
      from the session-scoped tmp directory on disk (written by upload_image)

    Returns the raw image bytes, or *None* on any failure.
    """
    if not src:
        return None
    # Case 1: inline base64 data URI
    if src.startswith("data:image/"):
        try:
            _, _, b64_str = src.partition(",")
            return base64.b64decode(b64_str)
        except Exception:
            return None
    # Case 2: server-side temp image URL (from /api/v1/workspace/upload_image)
    if src.startswith("/api/v1/workspace/tmp_image/"):
        parts = src.split("/")
        # expected: ['', 'api', 'v1', 'workspace', 'tmp_image', '{sid}', '{fname}']
        if len(parts) == 7:
            session_id, filename = parts[5], parts[6]
            if (
                len(session_id) == 32
                and all(c in "0123456789abcdef" for c in session_id)
                and "/" not in filename
                and "\\" not in filename
                and ".." not in filename
            ):
                img_path = Path("workspace") / "tmp" / session_id / "images" / filename
                if img_path.is_file():
                    try:
                        return img_path.read_bytes()
                    except Exception:
                        pass
    return None


def _insert_block_image(doc: Any, tag: Any) -> None:
    """Insert an image into the document from an ``<img>`` or ``<figure>`` tag.

    Finds the ``<img>`` element, resolves its ``src`` attribute via
    :func:`_resolve_img_src`, and calls ``doc.add_picture()`` to embed it at
    5 inches wide (constrained by the page width with aspect ratio preserved).
    Silently skips unresolvable images.
    """
    from docx.shared import Inches

    img_el = (
        tag
        if (hasattr(tag, "name") and tag.name and tag.name.lower() == "img")
        else (tag.find("img") if hasattr(tag, "find") else None)
    )
    if not img_el:
        return
    src = img_el.get("src", "")
    img_bytes = _resolve_img_src(src)
    if not img_bytes:
        return

    img_css = _parse_css_inline(img_el.get("style", ""))
    w_str = img_css.get("width", img_el.get("width", ""))
    from docx.shared import Pt

    w_val = None
    if w_str:
        try:
            if w_str.endswith("px"):
                w_val = Pt(float(w_str.replace("px", "")) * 0.75)
            elif w_str.endswith("in"):
                w_val = Inches(float(w_str.replace("in", "")))
            elif w_str.isdigit():
                w_val = Pt(float(w_str) * 0.75)
        except Exception:
            import logging

            logging.getLogger(__name__).warning(
                "Silenced exception caught", exc_info=True
            )
            pass

    try:
        if w_val:
            doc.add_picture(io.BytesIO(img_bytes), width=w_val)
        else:
            doc.add_picture(io.BytesIO(img_bytes), width=Inches(5))
    except Exception as exc:
        logger.debug("[_insert_block_image] 图片插入跳过: %s", exc)


def _setup_blank_doc_defaults(doc: Any) -> None:
    """Configure default styles for a blank python-docx Document."""
    from docx.oxml.ns import qn as _qn
    from docx.shared import Pt

    doc_style = doc.styles["Normal"]
    doc_style.font.name = "Calibri"
    doc_style.font.size = Pt(10.5)  # 五号 — standard Chinese document size
    try:
        rPr = doc_style.element.get_or_add_rPr()
        rFonts = rPr.find(_qn("w:rFonts"))
        if rFonts is None:
            from lxml import etree

            rFonts = etree.SubElement(rPr, _qn("w:rFonts"))
        rFonts.set(_qn("w:eastAsia"), "DengXian")
    except Exception:
        pass
    # Remove default empty paragraph
    for para in list(doc.paragraphs):
        p = para._element
        p.getparent().remove(p)


def _extract_docx_save_parts(docx_input: Any) -> tuple[str, dict[str, Any]]:
    if isinstance(docx_input, dict):
        html_content = docx_input.get("html")
        if html_content is None:
            html_content = docx_input.get("body_html", "")
        payload = {
            "header_html": str(docx_input.get("header_html") or ""),
            "footer_html": str(docx_input.get("footer_html") or ""),
            "sections": (
                docx_input.get("sections")
                if isinstance(docx_input.get("sections"), list)
                else []
            ),
            "comments": (
                docx_input.get("comments")
                if isinstance(docx_input.get("comments"), list)
                else []
            ),
            "footnotes": (
                docx_input.get("footnotes")
                if isinstance(docx_input.get("footnotes"), list)
                else []
            ),
        }
        return str(html_content or ""), payload
    return str(docx_input or ""), {
        "header_html": "",
        "footer_html": "",
        "sections": [],
        "comments": [],
        "footnotes": [],
    }


def _xml_local_name(tag: Any) -> str:
    text = str(tag or "")
    return text.split("}", 1)[-1] if "}" in text else text


def _xml_attr_by_local_name(el: Any, attr_name: str) -> str:
    if el is None:
        return ""
    target = str(attr_name or "").strip()
    if not target:
        return ""
    for key, value in getattr(el, "attrib", {}).items():
        if _xml_local_name(key) == target and value not in (None, ""):
            return str(value)
    return ""


def _coerce_docx_comment_flag(value: Any) -> bool | None:
    if value in (None, ""):
        return None
    normalized = str(value).strip().lower()
    if normalized in {"1", "true", "on", "yes"}:
        return True
    if normalized in {"0", "false", "off", "no"}:
        return False
    return None


def _normalize_docx_export_comments(raw_comments: Any) -> list[dict[str, Any]]:
    normalized: list[dict[str, Any]] = []
    for item in raw_comments if isinstance(raw_comments, list) else []:
        if not isinstance(item, dict):
            continue
        anchor_text = str(
            item.get("anchor_text")
            or item.get("anchorText")
            or item.get("target_text")
            or item.get("selected_text")
            or ""
        ).strip()
        text = str(
            item.get("text")
            or item.get("content")
            or item.get("body")
            or item.get("comment")
            or ""
        ).strip()
        if not anchor_text or not text:
            continue
        normalized_item: dict[str, Any] = {
            "anchor_text": anchor_text,
            "text": text,
            "author": str(item.get("author") or "").strip(),
            "initials": str(item.get("initials") or "").strip(),
            "date": str(item.get("date") or item.get("created_at") or "").strip(),
        }

        anchor_start_offset = item.get("anchor_start_offset")
        if anchor_start_offset in (None, ""):
            anchor_start_offset = item.get("anchorStartOffset")
        if anchor_start_offset in (None, ""):
            anchor_start_offset = item.get("range_start")

        anchor_end_offset = item.get("anchor_end_offset")
        if anchor_end_offset in (None, ""):
            anchor_end_offset = item.get("anchorEndOffset")
        if anchor_end_offset in (None, ""):
            anchor_end_offset = item.get("range_end")

        anchor_occurrence = item.get("anchor_occurrence")
        if anchor_occurrence in (None, ""):
            anchor_occurrence = item.get("anchorOccurrence")

        try:
            if anchor_start_offset not in (None, ""):
                normalized_item["anchor_start_offset"] = int(anchor_start_offset)
        except (TypeError, ValueError):
            pass
        try:
            if anchor_end_offset not in (None, ""):
                normalized_item["anchor_end_offset"] = int(anchor_end_offset)
        except (TypeError, ValueError):
            pass
        try:
            if anchor_occurrence not in (None, ""):
                normalized_item["anchor_occurrence"] = max(0, int(anchor_occurrence))
        except (TypeError, ValueError):
            pass

        anchor_context_before = item.get("anchor_context_before")
        if anchor_context_before in (None, ""):
            anchor_context_before = item.get("anchorContextBefore")
        anchor_context_after = item.get("anchor_context_after")
        if anchor_context_after in (None, ""):
            anchor_context_after = item.get("anchorContextAfter")
        if isinstance(anchor_context_before, str) and anchor_context_before:
            normalized_item["anchor_context_before"] = anchor_context_before
        if isinstance(anchor_context_after, str) and anchor_context_after:
            normalized_item["anchor_context_after"] = anchor_context_after

        for key, alt_key in (
            ("para_id", "paraId"),
            ("parent_para_id", "parentParaId"),
            ("parent_id", "parentId"),
            ("durable_id", "durableId"),
        ):
            raw_value = item.get(key)
            if raw_value in (None, ""):
                raw_value = item.get(alt_key)
            text_value = str(raw_value or "").strip()
            if text_value:
                normalized_item[key] = text_value

        for key in ("done", "resolved"):
            flag_state = _coerce_docx_comment_flag(item.get(key))
            if flag_state is not None:
                normalized_item[key] = flag_state

        normalized.append(normalized_item)
    return normalized


def _apply_docx_export_comments(doc: Any, raw_comments: Any) -> tuple[int, Any | None]:
    comments = _normalize_docx_export_comments(raw_comments)
    if not comments:
        return 0, None
    try:
        from docx.oxml.ns import qn
        from lxml import etree

        from web.track_changes_editor import TrackChangesEditor
    except ImportError as exc:
        logger.warning("[export_docx] comment support unavailable: %s", exc)
        return 0, None

    comments_el = etree.fromstring(
        b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        b' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
    )
    editor = TrackChangesEditor(author="Koto Review")
    applied = 0

    for index, comment in enumerate(comments, start=1):
        comment_author = comment.get("author") or editor.author or "Koto Review"
        editor.author = comment_author
        before_count = len(comments_el)
        if not editor._apply_single_comment(
            doc,
            comments_el,
            comment.get("anchor_text", ""),
            comment.get("text", ""),
            "",
            {
                "anchor_start_offset": comment.get("anchor_start_offset"),
                "anchor_end_offset": comment.get("anchor_end_offset"),
                "anchor_occurrence": comment.get("anchor_occurrence"),
                "anchor_context_before": comment.get("anchor_context_before"),
                "anchor_context_after": comment.get("anchor_context_after"),
            },
        ):
            logger.info(
                "[export_docx] skipped comment %d because anchor was not found: %.80s",
                index,
                comment.get("anchor_text", ""),
            )
            continue

        applied += 1
        if len(comments_el) > before_count:
            latest = comments_el[-1]
            latest.set(qn("w:author"), comment_author)
            comment_initials = str(comment.get("initials") or "").strip()
            if comment_initials:
                latest.set(qn("w:initials"), comment_initials)
            comment_date = comment.get("date") or ""
            if comment_date:
                latest.set(qn("w:date"), comment_date)

    if applied <= 0:
        return 0, None
    return applied, comments_el


def _inject_docx_comments_part(docx_bytes: bytes, comments_el: Any) -> bytes:
    import zipfile

    from lxml import etree

    from web.track_changes_editor import TrackChangesEditor

    helper = TrackChangesEditor(author="Koto Review")
    comments_xml = etree.tostring(
        comments_el, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    src = io.BytesIO(docx_bytes)
    out = io.BytesIO()
    wrote_comments = False

    with zipfile.ZipFile(src, "r") as zin:
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = helper._add_comments_content_type(data)
                elif item.filename == "word/_rels/document.xml.rels":
                    data = helper._add_comments_relationship(data)
                elif item.filename == "word/comments.xml":
                    data = comments_xml
                    wrote_comments = True
                zout.writestr(item, data)

            if not wrote_comments:
                zout.writestr("word/comments.xml", comments_xml)

    return out.getvalue()


def _clear_block_container(container: Any) -> None:
    element = getattr(container, "_element", None)
    if element is None:
        return
    for child in list(element):
        element.remove(child)


def _apply_run_marks(run: Any, marks: dict[str, bool] | None) -> None:
    marks = marks or {}
    if marks.get("bold"):
        run.bold = True
    if marks.get("italic"):
        run.italic = True
    if marks.get("underline"):
        run.underline = True
    if marks.get("strike"):
        run.font.strike = True
    if marks.get("superscript"):
        run.font.superscript = True
    if marks.get("subscript"):
        run.font.subscript = True


def _add_page_number_field(
    para: Any, css: dict[str, str] | None = None, marks: dict[str, bool] | None = None
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _styled_run(text: str = ""):
        run = para.add_run(text)
        _apply_run_marks(run, marks)
        _apply_run_inline(run, css or {})
        return run

    begin_run = _styled_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(fld_begin)

    instr_run = _styled_run()
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)

    sep_run = _styled_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(fld_sep)

    _styled_run("1")

    end_run = _styled_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


def _append_inline_html_to_paragraph(
    para: Any,
    node: Any,
    inherited_css: dict[str, str] | None = None,
    marks: dict[str, bool] | None = None,
) -> None:
    from bs4.element import NavigableString, Tag
    from docx.shared import Inches, Pt

    inherited_css = dict(inherited_css or {})
    marks = dict(marks or {})

    for child in list(getattr(node, "children", [])):
        if isinstance(child, NavigableString):
            text = str(child)
            if not text:
                continue
            run = para.add_run(text)
            _apply_run_marks(run, marks)
            _apply_run_inline(run, inherited_css)
            continue

        if not isinstance(child, Tag):
            continue

        name = child.name.lower() if child.name else ""
        child_css = dict(inherited_css)
        child_css.update(_parse_css_inline(child.get("style", "")))
        child_marks = dict(marks)

        if name in ("strong", "b"):
            child_marks["bold"] = True
        if name in ("em", "i"):
            child_marks["italic"] = True
        if name == "u":
            child_marks["underline"] = True
        if name == "s":
            child_marks["strike"] = True
        if name == "sup":
            child_marks["superscript"] = True
        if name == "sub":
            child_marks["subscript"] = True

        if name == "br":
            para.add_run().add_break()
            continue

        if "koto-hdr-page-num" in (child.get("class") or []):
            _add_page_number_field(para, child_css, child_marks)
            continue

        if name == "img":
            src = child.get("src", "")
            img_bytes = _resolve_img_src(src)
            if img_bytes:
                try:
                    run = para.add_run()
                    _apply_run_marks(run, child_marks)
                    img_css = _parse_css_inline(child.get("style", ""))
                    w_str = img_css.get("width", child.get("width", ""))
                    w_val = None
                    if w_str:
                        try:
                            if w_str.endswith("px"):
                                w_val = Pt(float(w_str.replace("px", "")) * 0.75)
                            elif w_str.endswith("in"):
                                w_val = Inches(float(w_str.replace("in", "")))
                            elif w_str.isdigit():
                                w_val = Pt(float(w_str) * 0.75)
                        except Exception:
                            pass
                    if w_val:
                        run.add_picture(io.BytesIO(img_bytes), width=w_val)
                    else:
                        run.add_picture(io.BytesIO(img_bytes), width=Inches(2.5))
                except Exception:
                    pass
            continue

        _append_inline_html_to_paragraph(para, child, child_css, child_marks)


def _configure_header_footer_tabs(para: Any, section: Any) -> None:
    try:
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
        from docx.shared import Emu

        usable_width = (
            int(section.page_width)
            - int(section.left_margin)
            - int(section.right_margin)
        )
        if usable_width <= 0:
            return
        tab_stops = para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(
            Emu(usable_width // 2), WD_TAB_ALIGNMENT.CENTER, WD_TAB_LEADER.SPACES
        )
        tab_stops.add_tab_stop(
            Emu(usable_width), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
        )
    except Exception:
        pass


def _add_header_footer_paragraph(container: Any, tag: Any, section: Any) -> None:
    para = container.add_paragraph()
    pcss = _parse_css_inline(tag.get("style", ""))
    if pcss:
        _apply_para_format(para, pcss)

    direct_cols = [
        child
        for child in list(getattr(tag, "children", []))
        if getattr(child, "name", None) and "koto-hdr-col" in (child.get("class") or [])
    ]
    if direct_cols:
        _configure_header_footer_tabs(para, section)
        for idx, col in enumerate(direct_cols):
            col_css = dict(pcss)
            col_css.update(_parse_css_inline(col.get("style", "")))
            _append_inline_html_to_paragraph(para, col, col_css, {})
            if idx < len(direct_cols) - 1:
                tab_run = para.add_run("\t")
                _apply_run_inline(tab_run, col_css)
        return

    _append_inline_html_to_paragraph(para, tag, pcss, {})


def _write_docx_header_footer_html(
    container: Any, html_content: str, section: Any
) -> None:
    from bs4 import BeautifulSoup

    _clear_block_container(container)
    soup = BeautifulSoup(html_content or "", "html.parser")
    body = soup.find("body") or soup
    wrote_any = False

    for top in body.children:
        if not hasattr(top, "name") or top.name is None:
            text = str(top).strip()
            if text:
                para = container.add_paragraph()
                run = para.add_run(text)
                _apply_run_inline(run, {})
                wrote_any = True
            continue

        tag_name = top.name.lower()
        if tag_name in ("p", "li", "h1", "h2", "h3", "h4", "h5", "h6"):
            _add_header_footer_paragraph(container, top, section)
            wrote_any = True
            continue
        if tag_name in ("ul", "ol"):
            for li in top.find_all("li", recursive=False):
                _add_header_footer_paragraph(container, li, section)
                wrote_any = True
            continue

        text = top.get_text(separator=" ").strip()
        if text:
            para = container.add_paragraph()
            run = para.add_run(text)
            _apply_run_inline(run, {})
            wrote_any = True

    if not wrote_any:
        container.add_paragraph("")


def _payload_value(
    section_payload: dict[str, Any] | None, key: str, fallback: str = ""
) -> str:
    if isinstance(section_payload, dict) and key in section_payload:
        return str(section_payload.get(key) or "")
    return fallback


def _export_docx_python(docx_input: Any, original_path: str | None = None) -> bytes:
    """
    Build a .docx from editor HTML using python-docx directly.

    If *original_path* points to a valid .docx file, it is used as a
    template so that the original styles, theme, fonts, headers/footers,
    and section properties are preserved.  The body content is replaced
    with what was edited in the browser.
    """

    try:
        from bs4 import BeautifulSoup
        from docx import Document
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError(f"python-docx 或 beautifulsoup4 未安装: {exc}") from exc

    html_content, docx_payload = _extract_docx_save_parts(docx_input)
    default_header_html = docx_payload.get("header_html", "")
    default_footer_html = docx_payload.get("footer_html", "")
    sections_payload = (
        docx_payload.get("sections")
        if isinstance(docx_payload.get("sections"), list)
        else []
    )
    comments_payload = (
        docx_payload.get("comments")
        if isinstance(docx_payload.get("comments"), list)
        else []
    )

    # ── Open original as template, or create blank ────────────────────────
    if original_path and os.path.isfile(original_path):
        try:
            doc = Document(original_path)
            # Clear body content but keep section properties
            body = doc.element.body
            sect_pr = body.findall(qn("w:sectPr"))
            for child in list(body):
                if child.tag != qn("w:sectPr"):
                    body.remove(child)
            logger.info(
                "[export_docx] using original DOCX as template: %s", original_path
            )
        except Exception as exc:
            logger.warning(
                "[export_docx] failed to open original (%s), creating blank", exc
            )
            doc = Document()
            _setup_blank_doc_defaults(doc)
    else:
        doc = Document()
        _setup_blank_doc_defaults(doc)

    soup = BeautifulSoup(html_content, "html.parser")
    body = soup.find("body") or soup

    for top in body.children:
        if not hasattr(top, "name") or top.name is None:
            # bare text node
            text = str(top).strip()
            if text:
                doc.add_paragraph(text)
            continue

        tag_name = top.name.lower()

        # ── Images (standalone <img> or <figure>) ────────────────────────
        if tag_name in ("img", "figure"):
            _insert_block_image(doc, top)
            continue

        # ── Block-level text elements ──────────────────────────────────────
        if tag_name in ("h1", "h2", "h3", "h4", "h5", "h6", "li"):
            _add_paragraph_from_tag(doc, top)
            continue

        if tag_name == "p":
            # <p><img .../></p> — image-only paragraph → embed as picture block
            if not top.get_text().strip() and top.find("img"):
                _insert_block_image(doc, top.find("img"))
            else:
                _add_paragraph_from_tag(doc, top)
            continue

        if tag_name in ("ul", "ol"):
            for li in top.find_all("li", recursive=False):
                _add_paragraph_from_tag(doc, li)
            continue

        # ── Tables ─────────────────────────────────────────────────────────
        if tag_name == "table":
            rows_tags = top.find_all("tr")
            if not rows_tags:
                continue

            # Determine table dimensions
            col_count = 0
            for r in rows_tags:
                c = 0
                for cell in r.find_all(["td", "th"], recursive=False):
                    c += int(cell.get("colspan", 1))
                col_count = max(col_count, c)
            if col_count == 0:
                continue

            tbl = doc.add_table(rows=len(rows_tags), cols=col_count)
            tbl.style = "Table Grid"

            # Track which (row, col) positions are already consumed by a
            # rowspan merge so we skip writing to them again.
            occupied: dict[tuple[int, int], bool] = {}

            for ri, row_tag in enumerate(rows_tags):
                cell_tags = row_tag.find_all(["td", "th"], recursive=False)
                ci = 0  # logical col cursor
                for cell_tag in cell_tags:
                    # Advance past cells already filled by a previous rowspan
                    while occupied.get((ri, ci)):
                        ci += 1

                    colspan = int(cell_tag.get("colspan", 1))
                    rowspan = int(cell_tag.get("rowspan", 1))

                    # Guard against out-of-range
                    if ci >= col_count:
                        break

                    # Mark all cells this span covers as occupied
                    for dr in range(rowspan):
                        for dc in range(colspan):
                            if dr == 0 and dc == 0:
                                continue
                            occupied[(ri + dr, ci + dc)] = True

                    tbl_cell = tbl.cell(ri, ci)

                    # ── Merge ────────────────────────────────────────────
                    end_r = min(ri + rowspan - 1, len(rows_tags) - 1)
                    end_c = min(ci + colspan - 1, col_count - 1)
                    if end_r > ri or end_c > ci:
                        tbl_cell.merge(tbl.cell(end_r, end_c))

                    # ── Cell text ────────────────────────────────────────
                    # Clear the default empty paragraph python-docx adds
                    for p_el in list(tbl_cell.paragraphs):
                        p_el._element.getparent().remove(p_el._element)

                    cell_css = _parse_css_inline(cell_tag.get("style", ""))

                    # Vertical alignment
                    v_align = cell_css.get("vertical-align", "").lower()
                    if v_align == "middle":
                        tbl_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    elif v_align == "bottom":
                        tbl_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

                    # Bold if <th>
                    is_header = cell_tag.name.lower() == "th"

                    # Check if cell contains <p> sub-elements (TipTap style)
                    p_tags = cell_tag.find_all(
                        ["p", "h1", "h2", "h3", "h4", "h5", "h6"], recursive=False
                    )
                    if p_tags:
                        # Multi-paragraph cell: each <p> becomes a paragraph
                        for pi, p_tag in enumerate(p_tags):
                            cell_para = tbl_cell.add_paragraph()
                            p_css = _parse_css_inline(p_tag.get("style", ""))
                            # Merge cell_css as fallback
                            merged_p = dict(cell_css)
                            merged_p.update(p_css)
                            _apply_para_format(cell_para, merged_p)
                            for child in p_tag.children:
                                if not hasattr(child, "name") or child.name is None:
                                    text = str(child)
                                    if text:
                                        run = cell_para.add_run(text)
                                        if is_header:
                                            run.bold = True
                                        _apply_run_inline(run, merged_p)
                                else:
                                    child_name = (
                                        child.name.lower() if child.name else ""
                                    )
                                    text = child.get_text()
                                    if not text:
                                        continue
                                    run = cell_para.add_run(text)
                                    if is_header:
                                        run.bold = True
                                    child_css_inner = _parse_css_inline(
                                        child.get("style", "")
                                    )
                                    merged = dict(merged_p)
                                    merged.update(child_css_inner)
                                    if child_name in ("strong", "b"):
                                        run.bold = True
                                    if child_name in ("em", "i"):
                                        run.italic = True
                                    if child_name == "u":
                                        run.underline = True
                                    if child_name == "s":
                                        run.font.strike = True
                                    _apply_run_inline(run, merged)
                    else:
                        # Flat cell content (no <p> wrappers)
                        cell_para = tbl_cell.add_paragraph()

                        # Text alignment from cell style
                        align = cell_css.get("text-align", "")
                        if align == "center":
                            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif align == "right":
                            cell_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                        for child in cell_tag.children:
                            if not hasattr(child, "name") or child.name is None:
                                text = str(child).strip()
                                if text:
                                    run = cell_para.add_run(text)
                                    if is_header:
                                        run.bold = True
                                    _apply_run_inline(run, cell_css)
                            else:
                                child_name = child.name.lower() if child.name else ""
                                text = child.get_text()
                                if not text:
                                    continue
                                run = cell_para.add_run(text)
                                if is_header:
                                    run.bold = True
                                child_css_inner = _parse_css_inline(
                                    child.get("style", "")
                                )
                                merged = dict(cell_css)
                                merged.update(child_css_inner)
                                if child_name in ("strong", "b"):
                                    run.bold = True
                                if child_name in ("em", "i"):
                                    run.italic = True
                                if child_name == "u":
                                    run.underline = True
                                if child_name == "s":
                                    run.font.strike = True
                                _apply_run_inline(run, merged)

                    # ── Cell background colour ───────────────────────────
                    bg = cell_css.get("background-color", "") or cell_css.get(
                        "background", ""
                    )
                    bg_hex = _css_color_to_hex(bg)
                    if not bg_hex and is_header:
                        bg_hex = "EEF1F8"  # default header shading
                    if bg_hex:
                        _set_cell_shading(tbl_cell, bg_hex)

                    # Default borders for all tables (often lost in HTML conversion)
                    if not any(k.startswith("border") for k in cell_css):
                        _set_cell_borders(tbl_cell, {"border": "1px solid #000000"})
                    else:
                        _set_cell_borders(tbl_cell, cell_css)

                    # ── Border colour override (legacy) ──────────────────
                    border_raw = cell_css.get("border-color", "")
                    border_hex = _css_color_to_hex(border_raw)
                    if border_hex:
                        _set_cell_borders(tbl_cell, border_hex)

                    ci += colspan

            # ── Column widths from <col> or first row <td style="width:..."> ──
            col_tags = top.find_all("col")
            if not col_tags and rows_tags:
                col_tags = rows_tags[0].find_all(["td", "th"], recursive=False)

            if col_tags:
                # Force table to not auto-fit so widths are respected
                try:
                    tbl.allow_autofit = False
                    from docx.oxml.ns import qn

                    tblPr = tbl._element.xpath("w:tblPr")
                    if tblPr:
                        tblW = tblPr[0].xpath("w:tblW")
                        if tblW:
                            tblW[0].set(qn("w:type"), "pct")
                            tblW[0].set(qn("w:w"), "5000")  # 100%
                except Exception:
                    pass

                for idx, col_tag in enumerate(col_tags[:col_count]):
                    w_css = _parse_css_inline(col_tag.get("style", ""))
                    width_str = w_css.get("width", col_tag.get("width", ""))
                    if width_str:
                        try:
                            if width_str.endswith("px"):
                                emu = int(
                                    float(width_str[:-2]) * 9144
                                )  # 1px ≈ 9144 EMU
                                for r in range(len(rows_tags)):
                                    tbl.cell(r, idx).width = emu
                            elif width_str.endswith("%"):
                                pct_val = float(width_str[:-1])
                                w_val = int(pct_val * 50)  # 50 = 1%
                                for r in range(len(rows_tags)):
                                    try:
                                        tc = tbl.cell(r, idx)._tc
                                        tcW = tc.get_or_add_tcPr().get_or_add_tcW()
                                        tcW.type = "pct"
                                        tcW.w = w_val
                                    except Exception:
                                        pass
                        except (ValueError, IndexError):
                            pass

            continue  # done with this table

        # ── Any other block: just extract its text ─────────────────────────
        text = top.get_text(separator=" ").strip()
        if text:
            doc.add_paragraph(text)

    try:
        has_even = any(
            _payload_value(section_payload, "even_header_html")
            or _payload_value(section_payload, "even_footer_html")
            for section_payload in sections_payload
            if isinstance(section_payload, dict)
        )
        if hasattr(doc.settings, "odd_and_even_pages_header_footer"):
            doc.settings.odd_and_even_pages_header_footer = bool(has_even)
    except Exception:
        has_even = False

    for idx, section in enumerate(doc.sections):
        section_payload = (
            sections_payload[idx]
            if idx < len(sections_payload) and isinstance(sections_payload[idx], dict)
            else None
        )
        section_header_html = _payload_value(
            section_payload, "header_html", default_header_html
        )
        section_footer_html = _payload_value(
            section_payload, "footer_html", default_footer_html
        )
        first_header_html = _payload_value(section_payload, "first_header_html", "")
        first_footer_html = _payload_value(section_payload, "first_footer_html", "")
        even_header_html = _payload_value(section_payload, "even_header_html", "")
        even_footer_html = _payload_value(section_payload, "even_footer_html", "")

        try:
            section.header.is_linked_to_previous = False
        except Exception:
            pass
        try:
            section.footer.is_linked_to_previous = False
        except Exception:
            pass
        _write_docx_header_footer_html(section.header, section_header_html, section)
        _write_docx_header_footer_html(section.footer, section_footer_html, section)

        try:
            section.different_first_page_header_footer = bool(
                first_header_html or first_footer_html
            )
        except Exception:
            pass
        try:
            section.first_page_header.is_linked_to_previous = False
        except Exception:
            pass
        try:
            section.first_page_footer.is_linked_to_previous = False
        except Exception:
            pass
        _write_docx_header_footer_html(
            section.first_page_header, first_header_html, section
        )
        _write_docx_header_footer_html(
            section.first_page_footer, first_footer_html, section
        )

        if has_even:
            try:
                section.even_page_header.is_linked_to_previous = False
            except Exception:
                pass
            try:
                section.even_page_footer.is_linked_to_previous = False
            except Exception:
                pass
            _write_docx_header_footer_html(
                section.even_page_header, even_header_html, section
            )
            _write_docx_header_footer_html(
                section.even_page_footer, even_footer_html, section
            )

    applied_comment_count = 0
    comments_el = None
    if comments_payload:
        applied_comment_count, comments_el = _apply_docx_export_comments(
            doc, comments_payload
        )
        logger.info(
            "[export_docx] comment payload=%d applied=%d",
            len(comments_payload),
            applied_comment_count,
        )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    data = buf.read()

    if comments_el is not None and applied_comment_count > 0:
        try:
            data = _inject_docx_comments_part(data, comments_el)
        except Exception as exc:
            logger.warning("[export_docx] failed to inject comments.xml: %s", exc)

    return data


def export_docx(docx_input: Any, original_path: str | None = None) -> bytes:
    """
    将编辑器产出的 HTML 转换为 .docx 字节流。

    如果提供 original_path（原始 .docx 文件路径），则以该文件为模板，
    保留样式、主题、字体、页眉页脚等原始格式，仅替换正文内容。

    主路径：_export_docx_python — 基于 python-docx 直接构建，完整支持：
      - 合并单元格 (colspan/rowspan)
      - 单元格背景色、边框色
      - 列宽、对齐方式
      - 行内格式（加粗、斜体、颜色、字号）

    备用路径：html2docx（当 python-docx/bs4 导入失败时）

    Returns:
        bytes — .docx 文件内容
    """
    html_content, _ = _extract_docx_save_parts(docx_input)
    logger.info(
        "[export_docx] html_content length=%d original=%s preview=%.200s",
        len(html_content or ""),
        original_path or "(none)",
        (html_content or "")[:200],
    )

    # ── Primary: rich python-docx builder ──────────────────────────────────
    try:
        return _export_docx_python(docx_input, original_path=original_path)
    except Exception as exc:
        logger.warning(
            "[export_docx] python-docx builder failed (%s), falling back to html2docx",
            exc,
        )

    # ── Fallback: html2docx (handles edge cases the builder misses) ─────────
    try:
        from html2docx import html2docx

        wrapped_html = (
            '<html><head><meta charset="utf-8">'
            '<style>body{font-family:"Microsoft YaHei","PingFang SC","SimHei",Arial,sans-serif;'
            "font-size:11pt;line-height:1.6;}"
            "h1{font-size:20pt}h2{font-size:16pt}h3{font-size:14pt}"
            "</style></head><body>" + html_content + "</body></html>"
        )
        buf = html2docx(wrapped_html, title="Koto 导出文档")
        buf.seek(0)
        data = buf.read()
        if data:
            return data
    except ImportError:
        pass

    raise RuntimeError(
        "export_docx: 所有路径均失败，请确认 python-docx 和 beautifulsoup4 已安装"
    )


__all__ = ["export_docx"]
