#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Quick integration test for docx_translator_module"""
import sys, os, json, re
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

from web.docx_translator_module import translate_docx_streaming, detect_target_language

# ── Mock LLM ──────────────────────────────────────────────────────────────────
class MockLLM:
    class models:
        @staticmethod
        def generate_content(model, contents, config):
            lines = [l for l in contents.split('\n') if re.match(r'^\d+\.', l.strip())]
            translations = {
                '张三的简历': 'Zhang San Resume',
                '个人信息': 'Personal Information',
                '姓名：张三': 'Name: Zhang San',
                '职位：软件工程师': 'Position: Software Engineer',
                '邮箱：zhangsan@example.com': 'Email: zhangsan@example.com',
                '工作经历': 'Work Experience',
                '2020-2024 ABC科技公司 - 高级工程师': '2020-2024 ABC Technology - Senior Engineer',
                '负责核心系统开发，主导微服务架构升级项目。': 'Led core system development and microservice architecture upgrade.',
                '技能': 'Skills', '级别': 'Level', '年限': 'Years',
                '精通': 'Expert', '5年': '5 years', 'Python': 'Python',
            }
            results = []
            for l in lines:
                txt = re.sub(r'^\d+\.\s+', '', l.strip())
                results.append(translations.get(txt, f'[EN] {txt}'))
            class Resp:
                text = json.dumps(results)
            return Resp()

# ── Language detection tests ───────────────────────────────────────────────────
print("=== Language Detection ===")
cases = [
    ("翻译成英文", "English"),
    ("translate to Japanese", "Japanese"),
    ("翻译成韩文", "Korean"),
    ("给我翻译成法语", "French"),
    ("翻译所有内容", "English"),  # default
]
for q, expected in cases:
    got = detect_target_language(q)
    ok = "✅" if got == expected else "❌"
    print(f"  {ok}  '{q}' → {got} (expected {expected})")

# ── Translation pipeline test ──────────────────────────────────────────────────
print("\n=== Translation Pipeline ===")
test_file = os.path.join(os.path.dirname(__file__), '..', 'workspace', 'documents', 'test_resume.docx')
if not os.path.exists(test_file):
    print("⚠️  test_resume.docx not found, creating it...")
    from docx import Document
    doc = Document()
    doc.add_heading('张三的简历', 0)
    doc.add_heading('个人信息', 1)
    doc.add_paragraph('姓名：张三')
    doc.add_paragraph('职位：软件工程师')
    os.makedirs(os.path.dirname(test_file), exist_ok=True)
    doc.save(test_file)

out_dir = os.path.join(os.path.dirname(__file__), '..', 'workspace', 'documents')
for event in translate_docx_streaming(test_file, 'English', MockLLM(), output_dir=out_dir, batch_size=5):
    stage = event['stage']
    msg = event['message']
    print(f"  [{stage:12s}] {msg}")
    if stage == 'complete':
        out = event.get('output_path', '')
        count = event.get('translated_count', 0)
        print(f"\n✅ Output: {event.get('output_filename')}")
        print(f"✅ Translated: {count} paragraphs")
        if os.path.exists(out):
            # Verify content
            from docx import Document as D
            d2 = D(out)
            print(f"✅ Output has {len(d2.paragraphs)} paragraphs, {len(d2.tables)} tables")
            print(f"   First heading: '{d2.paragraphs[0].text}'")
        else:
            print(f"❌ Output file not found: {out}")
    elif stage == 'error':
        print(f"\n❌ ERROR: {msg}")
        sys.exit(1)
