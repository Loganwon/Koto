#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""End-to-end test: translate the real resume and read back first 10 paragraphs."""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from dotenv import load_dotenv
load_dotenv('config/gemini_config.env')

from google import genai as genai2
api_key = os.environ.get('GEMINI_API_KEY') or os.environ.get('GOOGLE_API_KEY')
print(f'API key found: {bool(api_key)}')

client = genai2.Client(api_key=api_key)

from web.docx_translator_module import translate_docx_streaming

output_path = None
for event in translate_docx_streaming(
    input_path='web/uploads/张梓煜秋招中文简历_fv(1).docx',
    target_language='English',
    llm_client=client,
    output_dir='workspace/documents',
    batch_size=8,
):
    stage = event.get('stage')
    msg = event.get('message', '')
    pct = event.get('progress', 0)
    print(f'[{pct:3d}%] {msg}')
    if stage == 'complete':
        output_path = event.get('output_path')
        print(f'Output: {output_path}')
    elif stage == 'error':
        print('ERROR:', msg)
        sys.exit(1)

if output_path and os.path.exists(output_path):
    print('\n--- First 10 paragraphs of translated doc ---')
    from docx import Document
    doc = Document(output_path)
    count = 0
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            print(f'  P{count}: {t[:80]}')
            count += 1
            if count >= 10:
                break
    # Also check tables
    print('\n--- First 5 table cells ---')
    cell_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    print(f'  C{cell_count}: {t[:80]}')
                    cell_count += 1
                    if cell_count >= 5:
                        break
            if cell_count >= 5:
                break
        if cell_count >= 5:
            break
