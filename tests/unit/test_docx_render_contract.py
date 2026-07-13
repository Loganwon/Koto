"""Regression coverage for the parser-to-editor DOCX render contract."""

from __future__ import annotations

from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[2]


def test_mammoth_fallback_keeps_page_layout_and_navigation_contract(
    tmp_path, monkeypatch
):
    docx = pytest.importorskip("docx", reason="python-docx is required")
    pytest.importorskip("mammoth", reason="mammoth is required")
    from docx.shared import Inches
    from app.core.file.parsers import docx_parser

    document = docx.Document()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.75)
    section.header.paragraphs[0].text = "兼容模式页眉"
    document.add_heading("兼容模式标题", level=1)
    document.add_paragraph("用于验证 Mammoth 回退契约的正文。")
    fixture = tmp_path / "mammoth-fallback-contract.docx"
    document.save(fixture)

    def _force_rich_renderer_failure(*_args, **_kwargs):
        raise RuntimeError("forced rich renderer failure")

    monkeypatch.setattr(docx_parser, "_docx_to_rich_html", _force_rich_renderer_failure)
    result = docx_parser.parse_docx(str(fixture), progressive_preview=True)

    assert result["render_contract_version"] == 1
    assert result["render_source"] == "mammoth_fallback"
    assert result["render_degraded"] is True
    assert result["progressive"] == {"pending": False, "target_pages": 3}
    assert result["page_width_px"] > 700
    assert result["page_height_px"] > 1000
    assert result["margin_top_px"] > 0
    assert result["margin_bottom_px"] > 0
    assert result["margin_left_px"] > 0
    assert result["margin_right_px"] > 0
    assert result["header_html"] == "<p>兼容模式页眉</p>"
    assert result["sections"]
    assert result["sections"][0]["page_width_px"] == result["page_width_px"]
    assert result["headings"] == [
        {"id": "koto-fallback-heading-1", "level": 1, "text": "兼容模式标题"}
    ]
    assert 'data-koto-role="structural_heading"' in result["html"]


def test_workspace_maps_parser_docx_layout_keys_once_at_the_editor_boundary():
    source = (ROOT / "web" / "src" / "workspace" / "file-open.ts").read_text(
        encoding="utf-8"
    )

    assert "function _toDocxRenderOptions(data: any)" in source
    assert "_docxNumber(source, 'pageWidthPx', 'page_width_px')" in source
    assert "_docxNumber(source, 'marginTopPx', 'margin_top_px')" in source
    assert "state.activeEditor!.render(html, _toDocxRenderOptions(data));" in source
