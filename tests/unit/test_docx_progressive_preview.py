from __future__ import annotations

import io
import os
import sys

import pytest


ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
if ROOT not in sys.path:
    sys.path.insert(0, ROOT)


def _make_table_heavy_docx() -> bytes:
    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    table = doc.add_table(rows=48, cols=3)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.text = f"第 {row_idx + 1} 行 / 第 {col_idx + 1} 列 / 渐进加载测试内容"
    for idx in range(80):
        doc.add_paragraph(f"正文补充段落 {idx + 1}。" + "这是用于触发 DOCX 后台补全的测试内容。" * 2)
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture()
def wa_client():
    from flask import Flask
    from web.blueprints.workspace_assistant import workspace_assistant_bp

    app = Flask(__name__)
    app.secret_key = "test-secret-key"
    app.register_blueprint(workspace_assistant_bp)
    app.config["TESTING"] = True

    import tempfile
    import web.blueprints.workspace_assistant as _wa_mod

    tmp_root = tempfile.mkdtemp()
    orig_tmp_root = _wa_mod._TMP_ROOT
    _wa_mod._TMP_ROOT = type(_wa_mod._TMP_ROOT)(tmp_root)
    with app.test_client() as client:
        yield client
    _wa_mod._TMP_ROOT = orig_tmp_root


def test_parse_docx_progressive_preview_truncates_large_table_doc(tmp_path):
    pytest.importorskip("docx", reason="python-docx 未安装")
    docx_path = tmp_path / "table-heavy.docx"
    docx_path.write_bytes(_make_table_heavy_docx())

    from app.core.file.file_parser import parse_docx

    preview = parse_docx(str(docx_path), progressive_preview=True)
    full = parse_docx(str(docx_path))

    assert preview.get("progressive", {}).get("pending") is True
    assert preview.get("progressive", {}).get("target_pages") == 3
    assert len(preview.get("html", "")) < len(full.get("html", ""))
    assert "正在后台加载" in preview.get("html", "")


def test_workspace_docx_open_file_returns_full_payload_by_default(wa_client):
    pytest.importorskip("docx", reason="python-docx 未安装")

    resp = wa_client.post(
        "/api/v1/workspace/open_file",
        data={"file": (io.BytesIO(_make_table_heavy_docx()), "progressive.docx")},
        content_type="multipart/form-data",
    )
    assert resp.status_code == 200, resp.get_data(as_text=True)
    open_json = resp.get_json()
    assert open_json["file_type"] == "docx"
    assert open_json["data"].get("progressive") in (None, {}) or not open_json["data"].get("progressive", {}).get("pending")
    assert "正在后台加载" not in open_json["data"].get("html", "")

    full_resp = wa_client.post(
        "/api/v1/workspace/docx_full",
        json={"file_id": open_json["file_id"]},
    )
    assert full_resp.status_code == 200, full_resp.get_data(as_text=True)
    full_json = full_resp.get_json()
    assert full_json["file_type"] == "docx"
    assert len(full_json["data"].get("html", "")) == len(open_json["data"].get("html", ""))
    assert full_json["data"].get("progressive") in (None, {}) or not full_json["data"].get("progressive", {}).get("pending")