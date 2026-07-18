import base64
import json
import os
import stat
from pathlib import Path


def test_read_sheet_data_accepts_string_max_rows(tmp_path):
    import openpyxl

    from app.core.agent.task_tools import read_sheet_data

    workbook_path = tmp_path / "sales.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.append(["客户名称", "产品名称", "数量"])
    sheet.append(["杭州新汇鑫光电有限公司", "LASER", 1])
    sheet.append(["山东镭鸟激光设备有限公司", "LASER", 2])
    workbook.save(workbook_path)

    payload = json.loads(read_sheet_data(str(workbook_path), max_rows="1"))

    assert payload["headers"] == ["客户名称", "产品名称", "数量"]
    assert payload["row_count"] == 1
    assert payload["rows"][0][0] == "杭州新汇鑫光电有限公司"


def test_read_sheet_data_falls_back_from_generic_sheet1_for_single_sheet_workbook(
    tmp_path,
):
    import openpyxl

    from app.core.agent.task_tools import read_sheet_data

    workbook_path = tmp_path / "sales.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "汇总表"
    sheet.append(["客户名称", "产品名称", "数量"])
    sheet.append(["杭州新汇鑫光电有限公司", "LASER", 1])
    workbook.save(workbook_path)

    payload = json.loads(
        read_sheet_data(str(workbook_path), sheet_name="Sheet1", max_rows="1")
    )

    assert payload["sheet"] == "汇总表"
    assert payload["requested_sheet"] == "Sheet1"
    assert "used '汇总表'" in payload["warning"]
    assert payload["row_count"] == 1


def test_read_sheet_data_reports_available_sheets_when_requested_sheet_is_missing(
    tmp_path,
):
    import openpyxl

    from app.core.agent.task_tools import read_sheet_data

    workbook_path = tmp_path / "financial-model.xlsx"
    workbook = openpyxl.Workbook()
    first_sheet = workbook.active
    first_sheet.title = "P&L"
    first_sheet.append(["科目", "2025E"])
    first_sheet.append(["营业收入", 34807])

    expenses_sheet = workbook.create_sheet("Expenses")
    expenses_sheet.append(["科目", "2025E"])
    expenses_sheet.append(["销售费用", 1200])

    capex_sheet = workbook.create_sheet("资本折旧与投入")
    capex_sheet.append(["科目", "2025E"])
    capex_sheet.append(["资本开支", 800])
    workbook.save(workbook_path)

    payload = json.loads(
        read_sheet_data(str(workbook_path), sheet_name="Balance Sheet", max_rows="1")
    )

    assert payload["sheet"] == ""
    assert payload["headers"] == []
    assert payload["rows"] == []
    assert payload["row_count"] == 0
    assert payload["requested_sheet"] == "Balance Sheet"
    assert payload["available_sheets"] == ["P&L", "Expenses", "资本折旧与投入"]
    assert payload["missing_requested_sheet"] is True
    assert "Available:" in payload["warning"]
    assert "继续分析" in payload["summary"]
    assert "error" not in payload


def test_insert_excel_as_docx_table_accepts_string_max_rows(tmp_path):
    import openpyxl
    from docx import Document

    from app.core.agent.task_tools import insert_excel_as_docx_table

    workbook_path = tmp_path / "sales.xlsx"
    target_path = tmp_path / "target.docx"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "销售台账"
    sheet.append(["客户名称", "产品名称", "数量"])
    sheet.append(["杭州新汇鑫光电有限公司", "LASER", 1])
    sheet.append(["山东镭鸟激光设备有限公司", "LASER", 2])
    workbook.save(workbook_path)

    document = Document()
    document.add_paragraph("雷鸟访谈问题")
    document.save(target_path)

    payload = json.loads(
        insert_excel_as_docx_table(
            str(workbook_path),
            str(target_path),
            sheet_name="销售台账",
            table_title="销售台账数据",
            max_rows="1",
        )
    )

    assert payload["success"] is True
    assert payload["rows_written"] == 1
    assert payload["columns_written"] == 3

    saved = Document(str(target_path))
    assert len(saved.tables) == 1
    assert saved.tables[0].cell(1, 0).text == "杭州新汇鑫光电有限公司"


def test_insert_excel_as_docx_table_compacts_financial_model_for_report(tmp_path):
    import openpyxl
    from docx import Document

    from app.core.agent.task_tools import insert_excel_as_docx_table

    workbook_path = tmp_path / "financial.xlsx"
    target_path = tmp_path / "report.docx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "P&L"
    sheet.append([None] * 12)
    sheet.append([None, "利润表", *([None] * 10)])
    sheet.append([None, "单位：人民币万元", *([None] * 10)])
    sheet.append([None, None, "2025E", "2026E", "2027E", "2028E", *([None] * 6)])
    sheet.append(
        [None, "收入合计", 22327.432, 54884.056, 113756.697, 181242.989, *([None] * 6)]
    )
    sheet.append([None, "增速%", None, 1.4581445, 1.0726729, 0.5932512, *([None] * 6)])
    sheet.append([None, "净利润", -11738, -2041, 952, 13362, *([None] * 6)])
    sheet.append([None, "备注行", "skip", "skip", "skip", "skip", *([None] * 6)])
    workbook.save(workbook_path)

    payload = json.loads(
        insert_excel_as_docx_table(
            str(workbook_path),
            str(target_path),
            sheet_name="P&L",
            table_title="关键财务预测数据",
            max_rows=18,
            financial_compact=True,
        )
    )
    table = Document(target_path).tables[0]

    assert payload["financial_compact"] is True
    assert payload["columns_written"] == 5
    assert payload["rows_written"] == 3
    assert [cell.text for cell in table.rows[0].cells] == [
        "指标",
        "2025E",
        "2026E",
        "2027E",
        "2028E",
    ]
    assert [cell.text for cell in table.rows[2].cells] == [
        "增速%",
        "—",
        "145.8%",
        "107.3%",
        "59.3%",
    ]


def test_insert_excel_as_docx_table_sorts_and_selects_columns_for_top_n(tmp_path):
    import openpyxl
    from docx import Document

    from app.core.agent.task_tools import insert_excel_as_docx_table

    workbook_path = tmp_path / "sales.xlsx"
    target_path = tmp_path / "target.docx"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Sales"
    sheet.append(["Customer", "Region", "Revenue", "Margin", "Risk"])
    sheet.append(["Northwind Labs", "NA", 128000, 0.34, "Security review"])
    sheet.append(["Aurora Retail", "EU", 96000, 0.28, "Payment terms"])
    sheet.append(["Blue Harbor", "APAC", 142000, 0.31, "Capacity"])
    sheet.append(["Delta Foods", "EU", 118000, 0.37, "Upsell"])
    workbook.save(workbook_path)

    payload = json.loads(
        insert_excel_as_docx_table(
            str(workbook_path),
            str(target_path),
            sheet_name="Sales",
            table_title="Top 3 Customers by Revenue",
            max_rows="3",
            sort_by="Revenue",
            sort_order="desc",
            columns='["Customer", "Region", "Revenue", "Margin"]',
        )
    )

    assert payload["success"] is True
    assert payload["rows_written"] == 3
    assert payload["columns_written"] == 4
    assert payload["sort_by"] == "Revenue"
    assert payload["selected_columns"] == ["Customer", "Region", "Revenue", "Margin"]

    saved = Document(str(target_path))
    table = saved.tables[0]
    assert [table.cell(row, 0).text for row in range(1, 4)] == [
        "Blue Harbor",
        "Northwind Labs",
        "Delta Foods",
    ]
    assert len(table.columns) == 4


def test_insert_excel_as_docx_table_falls_back_from_generic_sheet1_for_single_sheet_workbook(
    tmp_path,
):
    import openpyxl
    from docx import Document

    from app.core.agent.task_tools import insert_excel_as_docx_table

    workbook_path = tmp_path / "sales.xlsx"
    target_path = tmp_path / "target.docx"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "汇总表"
    sheet.append(["客户名称", "产品名称", "数量"])
    sheet.append(["杭州新汇鑫光电有限公司", "LASER", 1])
    workbook.save(workbook_path)

    document = Document()
    document.add_paragraph("雷鸟访谈问题")
    document.save(target_path)

    payload = json.loads(
        insert_excel_as_docx_table(
            str(workbook_path),
            str(target_path),
            sheet_name="Sheet1",
            table_title="销售台账数据",
            max_rows="100",
        )
    )

    assert payload["success"] is True
    assert payload["sheet"] == "汇总表"
    assert payload["requested_sheet"] == "Sheet1"
    assert "used '汇总表'" in payload["warning"]
    assert payload["rows_written"] == 1


def test_insert_excel_as_docx_table_continues_when_backup_creation_is_denied(
    tmp_path, monkeypatch
):
    import openpyxl
    from docx import Document

    from app.core.agent import task_tools
    from app.core.agent.task_tools import insert_excel_as_docx_table

    workbook_path = tmp_path / "financial-model.xlsx"
    target_path = tmp_path / "target.docx"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "P&L"
    sheet.append(["科目", "2025E"])
    sheet.append(["收入合计", 34807])
    workbook.save(workbook_path)

    document = Document()
    document.add_paragraph("雷鸟访谈问题")
    document.save(target_path)

    real_copy2 = task_tools.shutil.copy2

    def deny_backup_copy(src, dst, *args, **kwargs):
        if str(dst).endswith(".bak"):
            raise PermissionError("backup file is locked")
        return real_copy2(src, dst, *args, **kwargs)

    monkeypatch.setattr(task_tools.shutil, "copy2", deny_backup_copy)

    payload = json.loads(
        insert_excel_as_docx_table(
            str(workbook_path),
            str(target_path),
            sheet_name="P&L",
            table_title="财务预测 - 利润表",
            max_rows="10",
        )
    )

    assert payload["success"] is True
    assert payload["sheet"] == "P&L"
    assert "无法创建备份" in payload["warning"]

    saved = Document(str(target_path))
    assert len(saved.tables) == 1
    assert saved.tables[0].cell(1, 0).text == "收入合计"


def test_insert_excel_as_docx_table_falls_back_to_unique_backup_when_primary_backup_path_is_locked(
    tmp_path, monkeypatch
):
    import openpyxl
    from docx import Document

    from app.core.agent import task_tools
    from app.core.agent.task_tools import insert_excel_as_docx_table

    workbook_path = tmp_path / "financial-model.xlsx"
    target_path = tmp_path / "target.docx"
    canonical_backup_path = tmp_path / "target.docx.bak"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "P&L"
    sheet.append(["科目", "2025E"])
    sheet.append(["收入合计", 34807])
    workbook.save(workbook_path)

    document = Document()
    document.add_paragraph("雷鸟访谈问题")
    document.save(target_path)

    real_copy2 = task_tools.shutil.copy2

    def deny_only_canonical_backup(src, dst, *args, **kwargs):
        if str(dst) == str(canonical_backup_path):
            raise PermissionError("backup file is locked")
        return real_copy2(src, dst, *args, **kwargs)

    monkeypatch.setattr(task_tools.shutil, "copy2", deny_only_canonical_backup)

    payload = json.loads(
        insert_excel_as_docx_table(
            str(workbook_path),
            str(target_path),
            sheet_name="P&L",
            table_title="财务预测 - 利润表",
            max_rows="10",
        )
    )

    assert payload["success"] is True
    assert "warning" not in payload or "无法创建备份" not in str(
        payload.get("warning") or ""
    )

    backup_candidates = [
        path
        for path in tmp_path.glob("target.docx*.bak")
        if path.name != canonical_backup_path.name
    ]
    assert backup_candidates

    saved = Document(str(target_path))
    assert len(saved.tables) == 1
    assert saved.tables[0].cell(1, 0).text == "收入合计"


def test_insert_excel_as_docx_table_clears_readonly_existing_backup_before_copy(
    tmp_path, monkeypatch
):
    import openpyxl
    from docx import Document

    from app.core.agent import task_tools
    from app.core.agent.task_tools import insert_excel_as_docx_table

    workbook_path = tmp_path / "financial-model.xlsx"
    target_path = tmp_path / "target.docx"
    backup_path = tmp_path / "target.docx.bak"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "P&L"
    sheet.append(["科目", "2025E"])
    sheet.append(["收入合计", 34807])
    workbook.save(workbook_path)

    document = Document()
    document.add_paragraph("雷鸟访谈问题")
    document.save(target_path)
    backup_path.write_bytes(b"stale backup")
    backup_path.chmod(stat.S_IREAD)

    real_copy2 = task_tools.shutil.copy2

    def guarded_copy2(src, dst, *args, **kwargs):
        if str(dst) == str(backup_path):
            assert os.stat(dst).st_mode & stat.S_IWRITE
        return real_copy2(src, dst, *args, **kwargs)

    monkeypatch.setattr(task_tools.shutil, "copy2", guarded_copy2)

    payload = json.loads(
        insert_excel_as_docx_table(
            str(workbook_path),
            str(target_path),
            sheet_name="P&L",
            table_title="财务预测 - 利润表",
            max_rows="10",
        )
    )

    assert payload["success"] is True
    assert str(payload["path"]).endswith("target.docx")


def test_insert_excel_as_docx_table_clears_readonly_target_before_replace(tmp_path):
    import openpyxl
    from docx import Document

    from app.core.agent.task_tools import insert_excel_as_docx_table

    workbook_path = tmp_path / "financial-model.xlsx"
    target_path = tmp_path / "target.docx"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "P&L"
    sheet.append(["科目", "2025E"])
    sheet.append(["收入合计", 34807])
    workbook.save(workbook_path)

    document = Document()
    document.add_paragraph("雷鸟访谈问题")
    document.save(target_path)
    target_path.chmod(stat.S_IREAD)

    payload = json.loads(
        insert_excel_as_docx_table(
            str(workbook_path),
            str(target_path),
            sheet_name="P&L",
            table_title="财务预测 - 利润表",
            max_rows="10",
        )
    )

    assert payload["success"] is True
    assert str(payload["path"]).endswith("target.docx")

    saved = Document(str(target_path))
    assert len(saved.tables) == 1
    assert saved.tables[0].cell(1, 0).text == "收入合计"


def test_write_sheet_data_clears_readonly_target_before_save(tmp_path):
    import openpyxl

    from app.core.agent.task_tools import write_sheet_data

    workbook_path = tmp_path / "sales.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "汇总表"
    sheet.append(["客户名称", "数量"])
    sheet.append(["杭州新汇鑫光电有限公司", 1])
    workbook.save(workbook_path)
    workbook_path.chmod(stat.S_IREAD)

    payload = json.loads(
        write_sheet_data(
            str(workbook_path),
            sheet_name="汇总表",
            updates=[{"row": 2, "col": 2, "value": 3}],
        )
    )

    assert payload["success"] is True
    assert payload["cells_written"] == 1
    assert "自动移除只读属性" in payload["warning"]

    saved = openpyxl.load_workbook(workbook_path, data_only=True)
    assert saved["汇总表"].cell(row=2, column=2).value == 3
    saved.close()


def test_write_docx_content_clears_readonly_target_before_save(tmp_path):
    from docx import Document

    from app.core.agent.task_tools import write_docx_content

    target_path = tmp_path / "notes.docx"
    document = Document()
    document.add_paragraph("原始段落")
    document.save(target_path)
    target_path.chmod(stat.S_IREAD)

    payload = json.loads(
        write_docx_content(
            str(target_path),
            paragraphs=[{"text": "新增段落"}],
        )
    )

    assert payload["success"] is True
    assert payload["paragraphs_written"] == 1
    assert "自动移除只读属性" in payload["warning"]

    saved = Document(str(target_path))
    texts = [paragraph.text for paragraph in saved.paragraphs]
    assert "原始段落" in texts
    assert "新增段落" in texts


def test_insert_image_into_docx_appends_picture_and_caption(tmp_path):
    from docx import Document

    from app.core.agent.task_tools import insert_image_into_docx

    target_path = tmp_path / "report.docx"
    image_path = tmp_path / "chart.png"

    document = Document()
    document.add_paragraph("原始段落")
    document.save(target_path)
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+tm0YAAAAASUVORK5CYII="
        )
    )

    payload = json.loads(
        insert_image_into_docx(
            str(target_path),
            str(image_path),
            title="财务预测图表",
            caption="收入与利润趋势",
            width_inches="5.5",
        )
    )

    assert payload["success"] is True
    assert payload["image_name"] == "chart.png"
    assert payload["images_inserted"] == 1

    saved = Document(str(target_path))
    texts = [paragraph.text for paragraph in saved.paragraphs]
    assert "原始段落" in texts
    assert "财务预测图表" in texts
    assert "收入与利润趋势" in texts
    assert len(saved.inline_shapes) == 1


def test_run_python_code_materializes_sandbox_image_artifacts():
    from app.core.agent.task_tools import run_python_in_sandbox

    png_b64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+tm0YAAAAASUVORK5CYII="
    result = run_python_in_sandbox(
        "import base64, pathlib\n"
        f"pathlib.Path('chart_expense_structure.png').write_bytes(base64.b64decode('{png_b64}'))\n"
        "print('chart ready')\n"
    )

    generated_paths = result["generated_file_paths"]
    assert "chart_expense_structure.png" in generated_paths
    assert Path(generated_paths["chart_expense_structure.png"]).exists()
    assert (
        result["generated_files"][0]["path"]
        == generated_paths["chart_expense_structure.png"]
    )
    assert "chart_expense_structure.png" in result["files"]


def test_run_python_code_syncs_relative_workspace_outputs_without_target_path(
    tmp_path, monkeypatch
):
    from app.core.agent import task_tools
    from app.core.agent.task_tools import run_python_in_sandbox

    monkeypatch.setattr(task_tools, "_WORKSPACE_ROOT", str(tmp_path))

    result = run_python_in_sandbox(
        "from pathlib import Path\n"
        "import openpyxl\n"
        "output = Path('workspace/_codex_frontend_task_tests/real_frontend_ai_chain_validation.xlsx')\n"
        "output.parent.mkdir(parents=True, exist_ok=True)\n"
        "workbook = openpyxl.Workbook()\n"
        "sheet = workbook.active\n"
        "sheet.append(['品类', '数值'])\n"
        "sheet.append(['黄金', 680])\n"
        "sheet.append(['白银', 8])\n"
        "sheet.append(['铂金', 230])\n"
        "workbook.save(output)\n"
        "print('saved ' + str(output))\n",
        timeout=30,
    )

    expected = (
        tmp_path
        / "_codex_frontend_task_tests"
        / "real_frontend_ai_chain_validation.xlsx"
    )
    assert result.get("error") == ""
    assert expected.exists()
    assert str(expected) in result["__koto_created__"]


def test_run_python_code_strips_workspace_prefix_for_explicit_target_path(
    tmp_path, monkeypatch
):
    from app.core.agent import task_tools
    from app.core.agent.task_tools import run_python_in_sandbox

    monkeypatch.setattr(task_tools, "_WORKSPACE_ROOT", str(tmp_path))

    result = run_python_in_sandbox(
        "from pathlib import Path\n"
        "import openpyxl\n"
        "output = Path('workspace/reports/explicit_target.xlsx')\n"
        "output.parent.mkdir(parents=True, exist_ok=True)\n"
        "workbook = openpyxl.Workbook()\n"
        "workbook.active.append(['ok'])\n"
        "workbook.save(output)\n",
        timeout=30,
        target_path="workspace/reports/explicit_target.xlsx",
    )

    expected = tmp_path / "reports" / "explicit_target.xlsx"
    assert result.get("error") == ""
    assert expected.exists()
    assert str(expected) in result["__koto_created__"]


def test_insert_excel_as_docx_table_writes_fallback_copy_when_target_is_locked(
    tmp_path, monkeypatch
):
    import openpyxl
    from docx import Document

    from app.core.agent import task_tools
    from app.core.agent.task_tools import insert_excel_as_docx_table

    monkeypatch.setattr(task_tools, "_WORKSPACE_ROOT", str(tmp_path))

    workbook_path = tmp_path / "financial-model.xlsx"
    target_path = tmp_path / "target.docx"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "P&L"
    sheet.append(["科目", "2025E"])
    sheet.append(["收入合计", 34807])
    workbook.save(workbook_path)

    document = Document()
    document.add_paragraph("雷鸟访谈问题")
    document.save(target_path)

    real_replace = task_tools.os.replace

    def deny_target_replace(src, dst):
        if str(dst) == str(target_path):
            raise PermissionError(32, "locked", dst)
        return real_replace(src, dst)

    monkeypatch.setattr(task_tools.os, "replace", deny_target_replace)

    payload = json.loads(
        insert_excel_as_docx_table(
            str(workbook_path),
            str(target_path),
            sheet_name="P&L",
            table_title="财务预测 - 利润表",
            max_rows="10",
        )
    )

    assert payload["success"] is False
    assert payload["status"] == "write_blocked"
    assert payload["fallback_copy"] is True
    assert payload["blocked_target"] is True
    assert payload["original_target_path"] == "target.docx"
    assert payload["path"].endswith("target.koto-copy.docx")
    assert "尚未写回原文件" in payload["summary"]
    assert "当前不可写" in payload["summary"]

    saved = Document(str(tmp_path / "target.koto-copy.docx"))
    assert len(saved.tables) == 1
    assert saved.tables[0].cell(1, 0).text == "收入合计"


def test_save_docx_via_temp_file_reports_locked_target_with_actionable_message(
    tmp_path, monkeypatch
):
    import pytest

    from app.core.agent import task_tools

    target_path = tmp_path / "target.docx"
    target_path.write_bytes(b"original")

    class FakeDocument:
        def save(self, path):
            with open(path, "wb") as handle:
                handle.write(b"updated")

    def deny_replace(src, dst):
        raise PermissionError(32, "locked", dst)

    monkeypatch.setattr(task_tools.os, "replace", deny_replace)

    with pytest.raises(PermissionError, match="目标文件 target.docx 当前不可写"):
        task_tools._save_docx_via_temp_file(FakeDocument(), str(target_path))


def test_inspect_workbook_structure_reports_formula_headers_and_samples(tmp_path):
    import openpyxl

    from app.core.agent.task_tools import inspect_workbook_structure

    workbook_path = tmp_path / "financial-model.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "P&L"
    sheet.append(["科目", "2025E", "2026E", "2027E", "2028E"])
    sheet.append(["收入合计", 100, 120, 150, 180])
    sheet.append(["增速%", None, "=C2/B2-1", "=D2/C2-1", "=E2/D2-1"])
    expenses = workbook.create_sheet("Expenses")
    expenses.append(["科目", "2025E", "2026E"])
    expenses.append(["销售费用", 20, 24])
    workbook.save(workbook_path)

    payload = json.loads(
        inspect_workbook_structure(
            str(workbook_path),
            sample_rows_per_sheet=2,
            max_formula_examples_per_sheet=3,
        )
    )

    assert payload["sheet_count"] == 2
    assert payload["sheet_names"] == ["P&L", "Expenses"]
    first_sheet = payload["sheets"][0]
    assert first_sheet["name"] == "P&L"
    assert first_sheet["formula_count"] == 3
    assert first_sheet["year_header"]["row"] == 1
    assert [item["header"] for item in first_sheet["year_header"]["columns"]] == [
        "2025E",
        "2026E",
        "2027E",
        "2028E",
    ]
    assert first_sheet["sample_rows"][0]["values"] == [
        "科目",
        "2025E",
        "2026E",
        "2027E",
        "2028E",
    ]
    assert first_sheet["formula_examples"][0]["cell"] == "C3"


def test_audit_financial_workbook_flags_missing_statements_external_refs_and_series_gaps(
    tmp_path,
):
    import openpyxl

    from app.core.agent.task_tools import audit_financial_workbook

    workbook_path = tmp_path / "forecast.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "P&L"
    sheet.append(["科目", "2025E", "2026E", "2027E", "2028E"])
    sheet.append(["收入合计", 100, 120, 150, 180])
    sheet.append(["所得税费用", 0, 0, None, "='[底稿.xlsx]税表'!$B$2"])
    expenses = workbook.create_sheet("Expenses")
    expenses.append(["科目", "2025E", "2026E"])
    expenses.append(["销售费用", 20, 24])
    workbook.save(workbook_path)

    payload = json.loads(audit_financial_workbook(str(workbook_path), max_findings=10))
    finding_types = {item["type"] for item in payload["findings"]}

    assert payload["statement_presence"]["profit_and_loss"]["present"] is True
    assert payload["statement_presence"]["balance_sheet"]["present"] is False
    assert payload["statement_presence"]["cash_flow"]["present"] is False
    assert {"missing_statement", "external_dependency", "year_series_gap"}.issubset(
        finding_types
    )
    gap = next(
        item for item in payload["findings"] if item["type"] == "year_series_gap"
    )
    assert gap["sheet"] == "P&L"
    assert gap["label"] == "所得税费用"
    assert "2027E" in gap["message"]
