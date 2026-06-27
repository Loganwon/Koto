import json
from pathlib import Path

from app.core.agent.file_task_contract import (
    FileTaskFile,
    FileTaskRequest,
    FileTaskToolStreamChunk,
    FileTaskToolStreamResult,
)
from app.core.agent.file_task_runtime import FileTaskRuntime
from app.core.agent.file_task_workflow_state import build_workflow_state


def test_workflow_state_uses_unified_windows_for_large_file_batches():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="继续复杂文件任务的下一步",
        run_id="unified_window_demo",
        target_path="deck.pptx",
        files=[
            FileTaskFile(path="draft.docx", name="draft.docx", type="docx"),
            FileTaskFile(path="model.xlsx", name="model.xlsx", type="xlsx"),
            FileTaskFile(
                path="deck.pptx", name="deck.pptx", type="pptx", target=True
            ),
        ],
        options={
            "workflow_checkpoint": {
                "adapter": "generic_tool_loop",
                "policy": "confirm_each_step",
                "step_index": 2,
                "window_paragraphs": 6,
                "window_slides": 4,
                "source_path": "draft.docx",
                "target_path": "deck.pptx",
            }
        },
    )

    classification = runtime._normalize_mainline_contract(
        request,
        request.files,
        runtime._classify_request(request, request.files),
    )
    state = build_workflow_state(
        request,
        request.files,
        classification,
        {"recipe_id": "generic_file_task"},
    )
    windows_by_unit = {item["unit"]: item for item in state["large_file_windows"]}

    assert state["checkpoint"]["step_index"] == 2
    assert state["checkpoint"]["target_path"] == "deck.pptx"
    assert windows_by_unit["paragraph"]["current"] == {"start": 13, "end": 18}
    assert windows_by_unit["sheet"]["current"] == {"sheet_index": 2}
    assert windows_by_unit["slide"]["current"] == {"start": 9, "end": 12}
    assert state["task_plan"]["version"] == "task_plan_v1"
    assert state["task_plan"]["mainline_locked"] is True
    task_step_ids = [step["id"] for step in state["task_plan"]["steps"]]
    assert task_step_ids[:1] == ["read_context"]
    assert "context" not in task_step_ids
    assert "execute" not in task_step_ids
    assert "check" not in task_step_ids
    assert any(step["id"] == "verify_outputs" for step in state["task_plan"]["steps"])
    assert "unified_large_file_window:v1" in state["reason_codes"]


def test_runtime_context_reads_large_file_windows_from_workflow_state():
    parse_calls = []

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            parse_calls.append(dict(args or {}))
            return f"windowed:{args.get('path')}"
        if tool_name == "verify_task_completion":
            return json.dumps({"passed": True, "summary": "ok"}, ensure_ascii=False)
        return ""

    request = FileTaskRequest(
        task="继续复杂文件任务的下一步",
        run_id="runtime_window_context_demo",
        target_path="deck.pptx",
        files=[
            FileTaskFile(path="draft.docx", name="draft.docx", type="docx"),
            FileTaskFile(path="model.xlsx", name="model.xlsx", type="xlsx"),
            FileTaskFile(
                path="deck.pptx", name="deck.pptx", type="pptx", target=True
            ),
        ],
        options={
            "workflow_checkpoint": {
                "adapter": "generic_tool_loop",
                "policy": "confirm_each_step",
                "step_index": 1,
                "window_paragraphs": 5,
                "window_slides": 3,
                "source_path": "draft.docx",
                "target_path": "deck.pptx",
            }
        },
    )

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor,
            model_client=lambda **kwargs: {"content": "先读取窗口。", "tool_calls": []},
            max_rounds=1,
        ).run(request)
    )
    snippets = next(event for event in events if event.type == "run.finished").payload[
        "context"
    ]
    args_by_path = {item["path"]: item for item in parse_calls}

    assert args_by_path["draft.docx"]["window_unit"] == "paragraph"
    assert args_by_path["draft.docx"]["start"] == 6
    assert args_by_path["draft.docx"]["end"] == 10
    assert args_by_path["model.xlsx"]["window_unit"] == "sheet"
    assert args_by_path["model.xlsx"]["sheet_index"] == 1
    assert args_by_path["deck.pptx"]["window_unit"] == "slide"
    assert args_by_path["deck.pptx"]["start"] == 4
    assert args_by_path["deck.pptx"]["end"] == 6
    assert any(item.get("window_unit") == "paragraph" for item in snippets)
    assert any(item.get("sheet_index") == 1 for item in snippets)


def test_file_task_runtime_treats_awaiting_confirmation_tool_result_as_paused_state():
    artifact = {
        "artifact_type": "koto_large_task_resume_v1",
        "category": "batch_confirmation",
        "title": "继续执行第 1/3 批",
        "summary": "文件较大，等待确认开始第 1/3 批。",
        "suggested_next_step": "确认后继续执行第 1/3 批审校。",
    }

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            return "参考内容"
        if tool_name == "annotate_file":
            return FileTaskToolStreamResult(
                chunks=[
                    FileTaskToolStreamChunk(
                        kind="event",
                        event_type="plan.confirmed",
                        payload={
                            "summary": "按 3 批执行",
                            "steps": [{"id": "batch_1", "title": "第 1 批"}],
                        },
                    ),
                    FileTaskToolStreamChunk(
                        kind="result",
                        payload={
                            "success": True,
                            "summary": "文件较大，已生成 3 批执行计划，等待确认开始第 1/3 批。",
                            "awaiting_confirmation": True,
                            "target_path": "translation.docx",
                            "source_path": "source.pdf",
                            "batch_index": 0,
                            "total_batches": 3,
                            "next_action_artifact": artifact,
                        },
                    ),
                ]
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    def fake_model(**kwargs):
        if any(
            message.get("role") == "function" and message.get("name") == "annotate_file"
            for message in kwargs["messages"]
        ):
            return {"content": "等待确认后继续", "tool_calls": []}
        return {
            "content": "开始分批审校",
            "tool_calls": [
                {
                    "id": "annotate_batch_demo",
                    "name": "annotate_file",
                    "args": {
                        "path": "translation.docx",
                        "requirement": "PDF是原文，docx文件是现有翻译稿。文件较大，请拆成多个分段来处理。",
                        "annotations": "[]",
                    },
                }
            ],
        }

    request = FileTaskRequest(
        task="PDF是原文，docx文件是现有翻译稿。文件较大，请拆成多个分段来处理。",
        run_id="annotate_wait_confirm",
        target_path="translation.docx",
        files=[
            FileTaskFile(path="source.pdf", name="source.pdf", type="pdf"),
            FileTaskFile(
                path="translation.docx",
                name="translation.docx",
                type="docx",
                target=True,
            ),
        ],
    )

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor, model_client=fake_model, max_rounds=2
        ).run(request)
    )

    check_finished = next(event for event in events if event.type == "check.finished")
    run_finished = events[-1]

    assert any(event.type == "plan.confirmed" for event in events)
    assert not any(
        event.type == "tool.finished"
        and event.payload.get("tool_name") == "write_guard"
        for event in events
    )
    assert check_finished.payload["status"] == "awaiting_confirmation"
    assert check_finished.payload["next_action_artifact"] == artifact
    assert run_finished.payload["completed_task"] is False
    assert run_finished.payload["next_action_artifact"] == artifact


def test_file_task_runtime_forces_windowed_pdf_read_for_stepwise_docx_summary():
    tool_calls = []

    def fake_model(**kwargs):
        return {
            "content": "",
            "tool_calls": [
                {
                    "name": "write_docx_content",
                    "args": {
                        "path": "museum-summary.docx",
                        "paragraphs": json.dumps(
                            [
                                {
                                    "text": "当前页窗摘要（第 1-3 页）",
                                    "style": "Heading 1",
                                },
                                {
                                    "text": "文档识别：当前页窗来自中国博物馆数字技术应用年度报告，呈现数智化建设背景、编写组织和研究对象。"
                                },
                                {"text": "段落主题：本页窗用于建立报告开篇背景和目录框架，说明数字技术如何进入博物馆业务。"},
                                {"text": "结构线索：报告先交代数字化建设背景，再通过综述篇和案例篇展开理论方向与实践项目。"},
                                {
                                    "text": "内容线索：模型从当前页窗提炼出藏品管理、观众服务、展览展示、数字敦煌、知识图谱和沉浸式展览等关键词。"
                                },
                                {"text": "来源页码：第 1-3 页"},
                            ],
                            ensure_ascii=False,
                        ),
                    },
                }
            ],
        }

    def fake_executor(tool_name, args):
        tool_calls.append((tool_name, dict(args or {})))
        if tool_name == "parse_file_to_text":
            return (
                "[Page 1] 中国博物馆数字技术应用年度报告介绍了数智化建设背景、编写组织和研究对象。"
                "报告强调数字技术正在推动博物馆藏品管理、观众服务和展览展示方式持续变化。\n"
                "[Page 2] 目录列出综述篇与案例篇，案例覆盖数字敦煌、知识图谱、藏品档案管理系统和沉浸式展览。"
                "综述部分讨论以观众为中心的可持续发展、数据驱动的观众分析、策展逻辑变化和数据要素驱动的建设局面。"
                "这些内容为后续分章案例提供了行业背景，也说明报告关注的是技术、数据、场景和治理之间的关系。"
            )
        if tool_name == "write_docx_content":
            return json.dumps(
                {
                    "success": True,
                    "path": args["path"],
                    "operation": "write_docx_content",
                    "summary": "已写入第 1 步要点。",
                    "paragraphs_written": 6,
                },
                ensure_ascii=False,
            )
        if tool_name == "verify_task_completion":
            return json.dumps(
                {"passed": True, "summary": "已检测到 DOCX 写入。"}, ensure_ascii=False
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    request = FileTaskRequest(
        task="这是一篇非常长的pdf，里面有大量内容。请分步总结，创建一个docx文件，每完成一步等我继续。",
        run_id="long_pdf_window_demo",
        target_path="museum-summary.docx",
        files=[
            FileTaskFile(
                path="museum-report.pdf",
                name="museum-report.pdf",
                type="pdf",
                content="[Page 1] " + ("原始 PDF 预览 " * 2000),
            )
        ],
    )

    events = list(
        FileTaskRuntime(tool_executor=fake_executor, model_client=fake_model).run(
            request
        )
    )
    parse_call = next(
        (args for name, args in tool_calls if name == "parse_file_to_text"), None
    )
    run_started = next(event for event in events if event.type == "run.started")

    assert parse_call is not None
    assert parse_call["path"] == "museum-report.pdf"
    assert parse_call["start_page"] == 1
    assert parse_call["end_page"] == 3
    workflow_state = run_started.payload["workflow_state"]
    assert workflow_state["version"] == "file_task_workflow_state_v1"
    assert workflow_state["mainline"]["selected_recipe"] == "long_pdf_stepwise_docx_summary"
    assert workflow_state["large_file_windows"][0]["unit"] == "page"
    assert workflow_state["large_file_windows"][0]["current"] == {"start": 1, "end": 3}
    assert not any(
        event.type == "tool.finished"
        and event.payload.get("tool_name") == "provided_file_context"
        for event in events
    )
    assert any(event.type == "file.changed" for event in events)
    check_finished = next(event for event in events if event.type == "check.finished")
    run_finished = next(event for event in events if event.type == "run.finished")
    assert check_finished.payload["status"] == "awaiting_confirmation"
    assert (
        check_finished.payload["next_action_artifact"]["artifact_type"]
        == "koto_stepwise_resume_v1"
    )
    assert (
        check_finished.payload["next_action_artifact"]["workflow_checkpoint"][
            "target_path"
        ]
        == "museum-summary.docx"
    )
    assert (
        check_finished.payload["next_action_artifact"]["workflow_checkpoint"][
            "step_index"
        ]
        == 1
    )
    assert (
        check_finished.payload["next_action_artifact"]["workflow_checkpoint"][
            "source"
        ]
        == "workflow_checkpoint"
    )
    assert (
        check_finished.payload["next_action_artifact"]["large_file_windows"][0][
            "next"
        ]
        == {"start": 4, "end": 6}
    )
    assert check_finished.payload["next_action_artifact"]["next_page_range"] == "4-6"
    resume_options = check_finished.payload["next_action_artifact"][
        "resume_request"
    ]["options"]
    assert "batch_control" not in resume_options
    assert (
        resume_options["workflow_checkpoint"]["step_index"]
        == 1
    )
    assert run_finished.payload["completed_task"] is False
    assert run_finished.payload["runtime"]["terminal_status"] == "awaiting_confirmation"


def test_workflow_resume_control_ignores_retired_batch_control():
    from app.core.agent.file_task_workflow_state import workflow_resume_control

    request = FileTaskRequest(
        task="继续",
        options={
            "batch_control": {"step_index": 9, "source": "retired_batch_control"},
        },
    )

    assert workflow_resume_control(request) == {}


def test_request_with_workflow_checkpoint_does_not_recreate_batch_control():
    from app.core.agent.file_task_workflow_state import request_with_workflow_checkpoint

    request = FileTaskRequest(
        task="继续",
        options={
            "workflow_checkpoint": {
                "step_index": 3,
                "target_path": "summary.docx",
            },
            "batch_control": {"step_index": 0},
        },
    )

    normalized = request_with_workflow_checkpoint(request)

    assert "batch_control" not in normalized.options
    assert normalized.options["workflow_checkpoint"]["step_index"] == 3
    assert normalized.options["workflow_checkpoint_normalized"] is True
    assert normalized.target_path == "summary.docx"


def test_request_with_workflow_checkpoint_strips_retired_batch_control_without_resume():
    from app.core.agent.file_task_workflow_state import request_with_workflow_checkpoint

    request = FileTaskRequest(
        task="继续",
        options={"batch_control": {"step_index": 3}},
    )

    normalized = request_with_workflow_checkpoint(request)

    assert "batch_control" not in normalized.options
    assert "workflow_checkpoint" not in normalized.options


def test_file_task_runtime_resumes_from_workflow_checkpoint_without_compat_batch_control():
    tool_calls = []

    def fake_model(**kwargs):
        return {"content": "继续处理 checkpoint 指定窗口。", "tool_calls": []}

    def fake_executor(tool_name, args):
        tool_calls.append((tool_name, dict(args or {})))
        if tool_name == "parse_file_to_text":
            return "[Page 4] checkpoint window content"
        if tool_name == "verify_task_completion":
            return json.dumps({"passed": True, "summary": "ok"}, ensure_ascii=False)
        return ""

    request = FileTaskRequest(
        task="继续",
        run_id="workflow_checkpoint_resume_demo",
        files=[
            FileTaskFile(path="museum-report.pdf", name="museum-report.pdf", type="pdf"),
            FileTaskFile(
                path="museum-summary.docx",
                name="museum-summary.docx",
                type="docx",
                target=True,
            ),
        ],
        options={
            "workflow_checkpoint": {
                "adapter": "generic_tool_loop",
                "policy": "confirm_each_step",
                "step_index": 1,
                "window_pages": 3,
                "source_path": "museum-report.pdf",
                "target_path": "museum-summary.docx",
                "original_task": "长 PDF 分步总结并写入 DOCX，每步等待确认。",
            }
        },
    )

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor,
            model_client=fake_model,
            max_rounds=1,
        ).run(request)
    )
    parse_call = next(
        args
        for name, args in tool_calls
        if name == "parse_file_to_text" and args.get("path") == "museum-report.pdf"
    )
    run_started = next(event for event in events if event.type == "run.started")

    assert parse_call["start_page"] == 4
    assert parse_call["end_page"] == 6
    assert run_started.payload["target_path"] == "museum-summary.docx"
    assert "workflow_checkpoint_resume" in run_started.payload["reason_codes"]
    assert run_started.payload["workflow_state"]["checkpoint"]["step_index"] == 1


def test_file_task_runtime_uses_model_docx_write_for_stepwise_pdf_summary(
    tmp_path, monkeypatch
):
    import app.core.agent.task_tools as task_tools

    monkeypatch.setattr(task_tools, "_WORKSPACE_ROOT", str(tmp_path))
    target_path = tmp_path / "global-rules-summary.docx"
    tool_calls = []
    model_calls = []

    def fake_model(**kwargs):
        model_calls.append(kwargs)
        return {
            "content": "",
            "tool_calls": [
                {
                    "name": "write_docx_content",
                    "args": {
                        "path": str(target_path),
                        "paragraphs": json.dumps(
                            [
                                {
                                    "text": "当前页窗摘要（第 1-3 页）",
                                    "style": "Heading 1",
                                },
                                {
                                    "text": "文档识别：当前页窗来自 The Global Rules of Art，主要呈现书籍出版信息和目录框架。"
                                },
                                {"text": "段落主题：本页窗用于定位全书结构，说明 Part I 关注当代视觉艺术全球场域的形成。"},
                                {
                                    "text": "结构线索：目录从全球艺术场域的理论入口展开，随后进入生成机制、分化结构和文化世界经济中的位置分析。"
                                },
                                {
                                    "text": "内容线索：模型综合当前页窗后识别出作者、出版社、章节序列和核心研究对象，而不是简单复制目录文本。"
                                },
                                {"text": "来源页码：第 1-3 页"},
                            ],
                            ensure_ascii=False,
                        ),
                    },
                }
            ],
        }

    def fake_executor(tool_name, args):
        tool_calls.append((tool_name, dict(args or {})))
        if tool_name == "parse_file_to_text":
            return (
                "[Page 4] The Global Rules of Art. Larissa Buchholz. Princeton University Press 2022.\n"
                "[Page 5] Contents list Part I: The Emergence of a Global Field in the Contemporary Visual Arts.\n"
                "[Page 6] Chapter 1 introduces a global field approach to art and culture.\n"
                "[Page 7] Chapter 2 discusses the genesis of a global art field.\n"
                "[Page 8] Later parts cover divisions, valuation, and positions in a cultural world economy."
            )
        if tool_name == "write_docx_content":
            return task_tools.write_docx_content(
                args["path"], args.get("paragraphs", "[]")
            )
        if tool_name == "verify_task_completion":
            return json.dumps(
                {"completed": True, "summary": "已检测到 DOCX 写入。"},
                ensure_ascii=False,
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    request = FileTaskRequest(
        task=(
            "这是一篇非常长的pdf，里面有大量内容。我需要你做的是一个分步任务，将任务拆分成很多个小任务，"
            "一步一步完成，每完成一步和我汇报一下我来说继续。你将总结整篇文章的核心内容，你创建一个docx文件，"
            "记录你每一步发现的要点，然后每一步完成后更新docx"
        ),
        run_id="long_pdf_create_file_docx_contract",
        target_path=str(target_path),
        files=[
            FileTaskFile(
                path="global-rules.pdf",
                name="global-rules.pdf",
                type="pdf",
                content="[Page 4] " + ("book metadata and table of contents " * 400),
            )
        ],
    )

    events = list(
        FileTaskRuntime(tool_executor=fake_executor, model_client=fake_model).run(
            request
        )
    )

    assert len(model_calls) >= 1
    assert not any(name == "create_file" for name, _args in tool_calls)
    assert any(name == "write_docx_content" for name, _args in tool_calls)
    write_args = next(args for name, args in tool_calls if name == "write_docx_content")
    written_text = "\n".join(
        item["text"] for item in json.loads(write_args["paragraphs"])
    )
    for label in ("文档识别：", "段落主题：", "结构线索：", "内容线索：", "来源页码："):
        assert label in written_text
    assert "模型综合当前页窗后识别出作者" in written_text
    assert "当前进度" not in written_text
    assert "下一步计划" not in written_text
    file_changed = next(event for event in events if event.type == "file.changed")
    assert file_changed.payload["operation"] == "write_docx_content"
    assert file_changed.payload["paragraphs_written"] >= 6
    assert target_path.exists()

    check_finished = next(event for event in events if event.type == "check.finished")
    run_finished = next(event for event in events if event.type == "run.finished")
    assert check_finished.payload["status"] == "awaiting_confirmation"
    assert (
        check_finished.payload["next_action_artifact"]["route"]
        == "long_pdf_stepwise_docx_summary"
    )
    assert run_finished.payload["runtime"]["terminal_status"] == "awaiting_confirmation"


def test_file_task_runtime_stepwise_docx_write_falls_back_when_model_fails():
    tool_calls = []
    model_calls = []

    def fake_model(**kwargs):
        model_calls.append(kwargs)
        raise RuntimeError(
            'Ollama HTTP 500: {"error":"XML syntax error on line 4: element <function> closed by </parameter>"}'
        )

    def fake_executor(tool_name, args):
        tool_calls.append((tool_name, dict(args or {})))
        if tool_name == "parse_file_to_text":
            return (
                "[Page 1] 中国博物馆数字技术应用年度报告介绍了数智化建设背景、编写组织和研究对象。"
                "报告强调数字技术正在推动博物馆藏品管理、观众服务和展览展示方式持续变化。\n"
                "[Page 2] 目录列出综述篇与案例篇，案例覆盖数字敦煌、知识图谱、藏品档案管理系统和沉浸式展览。"
                "综述部分讨论以观众为中心的可持续发展、数据驱动的观众分析、策展逻辑变化和数据要素驱动的建设局面。"
                "这些内容为后续分章案例提供了行业背景，也说明报告关注的是技术、数据、场景和治理之间的关系。"
            )
        if tool_name == "write_docx_content":
            paragraphs = json.loads(args["paragraphs"])
            return json.dumps(
                {
                    "success": True,
                    "path": args["path"],
                    "operation": "write_docx_content",
                    "summary": f"已写入 {len(paragraphs)} 个段落到 Word 文档",
                    "file_type": "docx",
                    "change_type": "modify",
                    "paragraphs_written": len(paragraphs),
                },
                ensure_ascii=False,
            )
        if tool_name == "verify_task_completion":
            return json.dumps(
                {"completed": True, "summary": "已检测到 DOCX 写入。"},
                ensure_ascii=False,
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    request = FileTaskRequest(
        task="这是一篇非常长的pdf，里面有大量内容。我需要你做的是一个分步任务，将任务拆分成很多个小任务，一步一步完成，每完成一步和我汇报一下我来说继续。你将总结整篇文章的核心内容，你创建一个docx文件，记录你每一步发现的要点，然后每一步完成后更新docx",
        run_id="long_pdf_local_xml_fallback",
        target_path="museum-summary.docx",
        model_mode="local",
        files=[
            FileTaskFile(
                path="museum-report.pdf",
                name="museum-report.pdf",
                type="pdf",
                content="[Page 1] " + ("原始 PDF 预览 " * 2000),
            )
        ],
    )

    events = list(
        FileTaskRuntime(tool_executor=fake_executor, model_client=fake_model).run(
            request
        )
    )
    tool_names = [name for name, _args in tool_calls]
    write_args = next(args for name, args in tool_calls if name == "write_docx_content")
    run_finished = next(event for event in events if event.type == "run.finished")
    check_finished = next(event for event in events if event.type == "check.finished")

    assert "parse_file_to_text" in tool_names
    assert "write_docx_content" in tool_names
    assert Path(write_args["path"]).name == "museum-summary.docx"
    written_paragraphs = [item["text"] for item in json.loads(write_args["paragraphs"])]
    assert any("来源页码：第 1-3 页" == text for text in written_paragraphs)
    assert any(text.startswith("文档识别：") for text in written_paragraphs)
    assert any(text.startswith("段落主题：") for text in written_paragraphs)
    assert any(text.startswith("结构线索：") for text in written_paragraphs)
    assert any(text.startswith("内容线索：") for text in written_paragraphs)
    assert not any(
        "下一步计划" in text or "等待用户" in text or text.startswith("状态：")
        for text in written_paragraphs
    )
    assert "模型调用失败" not in run_finished.payload["summary"]
    assert check_finished.payload["status"] == "awaiting_confirmation"
    assert (
        check_finished.payload["next_action_artifact"]["artifact_type"]
        == "koto_stepwise_resume_v1"
    )
    assert check_finished.payload["next_action_artifact"]["next_page_range"] == "4-6"
    assert run_finished.payload["completed_task"] is False
    assert (
        run_finished.payload["next_action_artifact"]["route"]
        == "long_pdf_stepwise_docx_summary"
    )
    assert len(model_calls) == 1
    assert run_finished.payload["runtime"]["model_unavailable"] is True
    assert run_finished.payload["runtime"]["terminal_status"] == "awaiting_confirmation"
    assert any(
        event.type == "file.changed"
        and event.payload.get("operation") == "write_docx_content"
        for event in events
    )


def test_file_task_runtime_stepwise_resume_reads_next_pdf_window():
    tool_calls = []
    resume_task = "\u7ee7\u7eed\u5f53\u524d\u5206\u6b65\u6587\u4ef6\u4efb\u52a1\u7684\u4e0b\u4e00\u6b65\u3002"
    original_task = "\u8fd9\u662f\u4e00\u7bc7\u975e\u5e38\u957f\u7684pdf\uff0c\u5206\u6b65\u603b\u7ed3\u6574\u7bc7\u6587\u7ae0\uff0c\u521b\u5efa\u4e00\u4e2adocx\u8bb0\u5f55\u6bcf\u4e00\u6b65\u53d1\u73b0\uff0c\u6bcf\u5b8c\u6210\u4e00\u6b65\u7b49\u6211\u7ee7\u7eed\u3002"

    def fake_model(**kwargs):
        return {
            "content": "",
            "tool_calls": [
                {
                    "name": "write_docx_content",
                    "args": {
                        "path": "museum-summary.docx",
                        "paragraphs": json.dumps(
                            [
                                {
                                    "text": "当前页窗摘要（第 4-6 页）",
                                    "style": "Heading 1",
                                },
                                {"text": "文档识别：当前页窗继续处理中国博物馆数字技术应用年度报告，覆盖目录收束和引言开端。"},
                                {"text": "段落主题：本段说明报告如何从目录框架进入文化遗产数字化保护的发展背景。"},
                                {
                                    "text": "结构线索：页窗先列出引言、综述篇、案例篇等组成部分，再转入上世纪八十年代以来的行业演进。"
                                },
                                {
                                    "text": "内容线索：模型综合当前页窗后识别出数字敦煌、知识图谱、藏品档案管理系统、数字展览和智慧博物馆建设等主线。"
                                },
                                {"text": "来源页码：第 4-6 页"},
                            ],
                            ensure_ascii=False,
                        ),
                    },
                }
            ],
        }

    def fake_executor(tool_name, args):
        tool_calls.append((tool_name, dict(args or {})))
        if tool_name == "parse_file_to_text":
            return (
                "[Page 4] 目录列出引言、综述篇、案例篇、参考文献和作者简介，显示报告结构较完整。"
                "[Page 5] 案例篇覆盖数字敦煌、知识图谱、江苏省博物馆藏品档案管理系统和数字展览。"
                "[Page 6] 引言回顾上世纪八十年代以来文化遗产数字化保护的发展，强调信息技术、通信技术和数字技术在博物馆中的广泛应用。"
            )
        if tool_name == "write_docx_content":
            return json.dumps(
                {
                    "success": True,
                    "path": args["path"],
                    "operation": "write_docx_content",
                    "summary": "已写入第 2 步要点。",
                    "paragraphs_written": 6,
                },
                ensure_ascii=False,
            )
        if tool_name == "verify_task_completion":
            return json.dumps(
                {"completed": True, "summary": "文件已写入。"}, ensure_ascii=False
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    events = list(
        FileTaskRuntime(tool_executor=fake_executor, model_client=fake_model).run(
            FileTaskRequest(
                task=resume_task,
                run_id="long_pdf_window_resume_demo",
                target_path="museum-summary.docx",
                files=[
                    FileTaskFile(
                        path="museum-report.pdf", name="museum-report.pdf", type="pdf"
                    ),
                    FileTaskFile(
                        path="museum-summary.docx",
                        name="museum-summary.docx",
                        type="docx",
                        target=True,
                    ),
                ],
                options={
                    "workflow_checkpoint": {
                        "adapter": "generic_tool_loop",
                        "policy": "confirm_each_step",
                        "step_index": 1,
                        "window_pages": 3,
                        "original_task": original_task,
                    }
                },
            )
        )
    )

    parse_call = next(
        args
        for name, args in tool_calls
        if name == "parse_file_to_text" and args.get("path") == "museum-report.pdf"
    )
    check_finished = next(event for event in events if event.type == "check.finished")

    assert parse_call["start_page"] == 4
    assert parse_call["end_page"] == 6
    assert check_finished.payload["status"] == "awaiting_confirmation"
    assert (
        check_finished.payload["next_action_artifact"]["completed_page_range"] == "4-6"
    )
    assert check_finished.payload["next_action_artifact"]["next_page_range"] == "7-9"
    resume_options = check_finished.payload["next_action_artifact"][
        "resume_request"
    ]["options"]
    assert "batch_control" not in resume_options
    assert (
        resume_options["workflow_checkpoint"]["step_index"]
        == 2
    )


def test_file_task_runtime_stepwise_resume_rehydrates_files_and_falls_back_when_model_deviates():
    tool_calls = []

    def fake_model(**kwargs):
        return {
            "content": "我先看看目录。",
            "tool_calls": [],
        }

    def fake_executor(tool_name, args):
        tool_calls.append((tool_name, dict(args or {})))
        if tool_name == "parse_file_to_text":
            assert args["path"] == "museum-report.pdf"
            return (
                "[Page 7] Theory section explains museum digital applications for audience learning, "
                "visitor experience, interpretation, evaluation, curatorial design, accessibility, ethics, "
                "heritage protection, cultural communication, and sustainable public services. "
                "[Page 8] Data-driven visitor analysis builds a chain from data collection and association "
                "analysis to visitor profiles, behavioral prediction, emotional feedback, operational "
                "diagnosis, service optimization, privacy governance, and measurable education outcomes. "
                "[Page 9] Case studies cover collection data construction, collection data use, venue data "
                "integration, digital exhibitions, knowledge graph exploration, immersive display, "
                "cross-institution collaboration, open platforms, and public cultural service scenarios."
            )
        if tool_name == "verify_task_completion":
            return json.dumps(
                {"completed": True, "summary": "文件已写入。"}, ensure_ascii=False
            )
        if tool_name == "write_docx_content":
            return json.dumps(
                {
                    "success": True,
                    "path": args["path"],
                    "file_type": "docx",
                    "operation": "write_docx_content",
                    "change_type": "modify",
                    "summary": "已写入 6 个段落到 Word 文档",
                    "paragraphs_written": 6,
                },
                ensure_ascii=False,
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor, model_client=fake_model, max_rounds=1
        ).run(
            FileTaskRequest(
                task="继续当前分步文件任务的下一步：处理 PDF 第 7-9 页，并把本段实质分析追加到同一个 DOCX。",
                run_id="stepwise_rehydrate_demo",
                files=[],
                options={
                    "workflow_checkpoint": {
                        "adapter": "generic_tool_loop",
                        "policy": "confirm_each_step",
                        "step_index": 2,
                        "window_pages": 3,
                        "original_task": "长 PDF 分步总结并写入 DOCX，每步等待确认。",
                        "source_path": "museum-report.pdf",
                        "target_path": "museum-summary.docx",
                    }
                },
            )
        )
    )

    parse_call = next(args for name, args in tool_calls if name == "parse_file_to_text")
    check_finished = next(event for event in events if event.type == "check.finished")

    assert parse_call["start_page"] == 7
    assert parse_call["end_page"] == 9
    assert any(name == "write_docx_content" for name, _args in tool_calls)
    assert any(event.type == "file.changed" for event in events)
    assert check_finished.payload["status"] == "awaiting_confirmation"


def test_file_task_runtime_blocks_stepwise_docx_write_when_pdf_text_is_watermark_only():
    def fake_model(**kwargs):
        return {
            "content": "写入当前分段。",
            "tool_calls": [
                {
                    "name": "write_docx_content",
                    "args": {
                        "path": "watermark-summary.docx",
                        "paragraphs": json.dumps(
                            [
                                {"text": "水印内容摘要", "style": "Heading 1"},
                                {"text": "当前页窗只有考参通海泰国供仅。"},
                                {"text": "来源页码：第 1-3 页"},
                            ],
                            ensure_ascii=False,
                        ),
                    },
                }
            ],
        }

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            return (
                "[Page 1]\n考\n参\n通\n海\n泰\n国\n供\n仅\n\n[Page 2]\n考\n参\n通\n海\n泰\n国\n供\n仅"
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor, model_client=fake_model, max_rounds=1
        ).run(
            FileTaskRequest(
                task="这是一篇非常长的pdf，分步总结，创建docx，每一步等我继续。",
                run_id="watermark_pdf_stepwise_demo",
                target_path="watermark-summary.docx",
                files=[
                    FileTaskFile(path="watermark.pdf", name="watermark.pdf", type="pdf")
                ],
            )
        )
    )

    guard = next(
        event
        for event in events
        if event.payload.get("tool_name") == "supervisor_guard"
    )
    check_finished = next(event for event in events if event.type == "check.finished")

    assert "文本质量不足" in guard.payload["result_preview"]
    assert not any(event.type == "file.changed" for event in events)
    assert check_finished.payload["status"] == "no_file_change"


def test_file_task_runtime_native_stepwise_docx_write_bypasses_frontend_progress_model_output():
    model_calls = []

    def fake_model(**kwargs):
        model_calls.append(kwargs)
        raise RuntimeError("model unavailable for fallback test")

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            return (
                "[Page 1] 中国博物馆数字技术应用年度报告介绍了数智化建设背景、编写组织和研究对象。"
                "报告强调数字技术正在推动博物馆藏品管理、观众服务和展览展示方式持续变化。"
                "[Page 2] 目录列出综述篇与案例篇，案例覆盖数字敦煌、知识图谱、藏品档案管理系统和沉浸式展览。"
                "综述部分讨论以观众为中心的可持续发展、数据驱动的观众分析、策展逻辑变化和数据要素驱动的建设局面。"
                "这些内容为后续分章案例提供了行业背景，也说明报告关注的是技术、数据、场景和治理之间的关系。"
            )
        if tool_name == "write_docx_content":
            paragraphs = json.loads(args["paragraphs"])
            assert not any(
                "下一步计划" in item.get("text", "") or "当前进度" in item.get("text", "")
                for item in paragraphs
            )
            return json.dumps(
                {
                    "success": True,
                    "path": args["path"],
                    "operation": "write_docx_content",
                    "summary": f"已写入 {len(paragraphs)} 个段落到 Word 文档",
                    "paragraphs_written": len(paragraphs),
                },
                ensure_ascii=False,
            )
        if tool_name == "verify_task_completion":
            return json.dumps(
                {"completed": True, "summary": "已写入原生分步摘要。"},
                ensure_ascii=False,
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor, model_client=fake_model, max_rounds=1
        ).run(
            FileTaskRequest(
                task="这是一篇非常长的pdf，分步总结，创建docx，每一步等我继续。",
                run_id="progress_text_stepwise_demo",
                target_path="museum-summary.docx",
                files=[FileTaskFile(path="museum.pdf", name="museum.pdf", type="pdf")],
            )
        )
    )

    check_finished = next(event for event in events if event.type == "check.finished")

    assert len(model_calls) == 1
    assert any(
        event.type == "file.changed"
        and event.payload.get("operation") == "write_docx_content"
        for event in events
    )
    assert check_finished.payload["status"] == "awaiting_confirmation"


def test_file_task_runtime_native_stepwise_docx_write_avoids_markdown_progress_and_combined_labels():
    model_calls = []

    def fake_model(**kwargs):
        model_calls.append(kwargs)
        raise RuntimeError("model unavailable for fallback test")

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            return (
                "[Page 4] 中国博物馆数字技术应用及案例研究年度报告目录。"
                "[Page 5] 综述篇列出以观众为中心、数据驱动的观众分析、数字策展逻辑和数据要素驱动等章节。"
                "[Page 6] 案例篇列出数字敦煌、江苏省博物馆藏品档案管理系统、VR 大空间沉浸式展览等案例。"
            )
        if tool_name == "write_docx_content":
            paragraphs = json.loads(args["paragraphs"])
            text = "\n".join(item.get("text", "") for item in paragraphs)
            assert "#" not in text
            assert "**" not in text
            assert "---" not in text
            assert "文档识别/核心要点" not in text
            assert "段落主题/关键发现" not in text
            return json.dumps(
                {
                    "success": True,
                    "path": args["path"],
                    "operation": "write_docx_content",
                    "summary": f"已写入 {len(paragraphs)} 个段落到 Word 文档",
                    "paragraphs_written": len(paragraphs),
                },
                ensure_ascii=False,
            )
        if tool_name == "verify_task_completion":
            return json.dumps(
                {"completed": True, "summary": "已写入原生分步摘要。"},
                ensure_ascii=False,
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor, model_client=fake_model, max_rounds=1
        ).run(
            FileTaskRequest(
                task="这是一篇非常长的pdf，分步总结，创建docx，每一步等我继续。",
                run_id="markdown_progress_stepwise_demo",
                target_path="museum-summary.docx",
                files=[FileTaskFile(path="museum.pdf", name="museum.pdf", type="pdf")],
            )
        )
    )

    check_finished = next(event for event in events if event.type == "check.finished")

    assert len(model_calls) == 1
    assert any(
        event.type == "file.changed"
        and event.payload.get("operation") == "write_docx_content"
        for event in events
    )
    assert check_finished.payload["status"] == "awaiting_confirmation"


def test_file_task_runtime_blocks_stepwise_docx_write_with_combined_labels_only():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    snippets = [
        {
            "source": "museum.pdf",
            "path": "museum.pdf",
            "start_page": 4,
            "end_page": 6,
            "_raw_text": "有效文本" * 80,
        }
    ]

    block = runtime._stepwise_docx_content_quality_block_message(
        snippets,
        "\n".join(
            [
                "当前页窗摘要（第 4-6 页）",
                "文档识别/核心要点：年报目录部分。",
                "段落主题/关键发现：综述篇与案例篇结构。",
                "结构线索：引言、综述、案例三段式。",
                "内容线索：数字敦煌、知识图谱、VR 展览等案例。",
                "来源页码：第 4-6 页",
            ]
        ),
    )

    assert "合并标签" in block


def test_file_task_runtime_blocks_stepwise_docx_write_when_declared_page_range_mismatches_window():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    snippets = [
        {
            "source": "museum.pdf",
            "path": "museum.pdf",
            "start_page": 1,
            "end_page": 3,
            "_raw_text": "有效文本" * 80,
        }
    ]

    block = runtime._stepwise_docx_content_quality_block_message(
        snippets,
        "\n".join(
            [
                "当前页窗摘要（第 1-10 页）",
                "文档识别：当前页窗来自中国博物馆数字技术应用报告。",
                "段落主题：目录与开篇信息。",
                "结构线索：报告从目录进入引言和案例框架。",
                "内容线索：数字技术、观众中心和案例研究。",
                "来源页码：第 1-10 页",
            ]
        ),
    )

    assert "页窗标签与当前读取窗口不一致" in block


def test_file_task_runtime_native_stepwise_docx_write_avoids_duplicate_page_sections():
    model_calls = []

    def fake_model(**kwargs):
        model_calls.append(kwargs)
        raise RuntimeError("model unavailable for fallback test")

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            return (
                "[Page 10] 藏品数据利用部分讨论展览展示中的一物一展和主题式展览，说明数字技术如何围绕单件重要藏品或主题资源展开阐释。"
                "[Page 11] 场馆数据利用部分讨论南京博物院、三星堆博物馆的多源数据联动、数字资产整合和开放枢纽定位。"
                "[Page 12] 上海博物馆以 BIM 和数字孪生构建透明展厅，实现管理、展柜与展品的一体化监测和主动预警。"
            )
        if tool_name == "write_docx_content":
            paragraphs = json.loads(args["paragraphs"])
            assert not any(
                "【第 10-12 页要点】" in item.get("text", "") for item in paragraphs
            )
            return json.dumps(
                {
                    "success": True,
                    "path": args["path"],
                    "operation": "write_docx_content",
                    "summary": f"已写入 {len(paragraphs)} 个段落到 Word 文档",
                    "paragraphs_written": len(paragraphs),
                },
                ensure_ascii=False,
            )
        if tool_name == "verify_task_completion":
            return json.dumps(
                {"completed": True, "summary": "已写入兜底摘要。"}, ensure_ascii=False
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor, model_client=fake_model, max_rounds=2
        ).run(
            FileTaskRequest(
                task="这是一篇非常长的pdf，分步总结整篇文章，创建一个docx记录每一步发现，每完成一步等我继续。",
                run_id="stepwise_duplicate_page_section_demo",
                target_path="museum-summary.docx",
                files=[FileTaskFile(path="museum.pdf", name="museum.pdf", type="pdf")],
                options={
                    "workflow_checkpoint": {
                        "policy": "confirm_each_step",
                        "step_index": 3,
                        "window_pages": 3,
                        "source_path": "museum.pdf",
                        "original_task": "这是一篇非常长的pdf，分步总结整篇文章，创建一个docx记录每一步发现，每完成一步等我继续。",
                    }
                },
            )
        )
    )

    check_finished = next(event for event in events if event.type == "check.finished")

    assert len(model_calls) == 1
    assert any(
        event.type == "file.changed"
        and event.payload.get("operation") == "write_docx_content"
        for event in events
    )
    assert check_finished.payload["status"] == "awaiting_confirmation"


def test_file_task_runtime_allows_stepwise_docx_write_with_probe_style_structure():
    def fake_model(**kwargs):
        return {
            "content": "写入当前页窗。",
            "tool_calls": [
                {
                    "name": "write_docx_content",
                    "args": {
                        "path": "museum-summary.docx",
                        "paragraphs": json.dumps(
                            [
                                {
                                    "text": "当前页窗摘要（第 10-12 页）",
                                    "style": "Heading 1",
                                },
                                {"text": "文档识别：当前页窗来自《中国博物馆数字技术应用及案例研究年度报告》。"},
                                {"text": "段落主题：藏品数据与场馆数据在博物馆数字化转型中的利用方式。"},
                                {
                                    "text": "结构线索：先讨论展览展示中的“一物一展”和“主题式展览”，再转向场馆运营、数字资产和数字孪生管理。"
                                },
                                {
                                    "text": "内容线索：南京博物院、三星堆博物馆和上海博物馆分别体现数据联动、资产开放和透明展厅管理。"
                                },
                                {"text": "来源页码：第 10-12 页"},
                            ],
                            ensure_ascii=False,
                        ),
                    },
                }
            ],
        }

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            return (
                "[Page 10] 藏品数据利用部分讨论展览展示中的一物一展和主题式展览，说明数字技术如何围绕单件重要藏品或主题资源展开阐释。"
                "[Page 11] 场馆数据利用部分讨论南京博物院、三星堆博物馆的多源数据联动、数字资产整合和开放枢纽定位。"
                "[Page 12] 上海博物馆以 BIM 和数字孪生构建透明展厅，实现管理、展柜与展品的一体化监测和主动预警。"
            )
        if tool_name == "write_docx_content":
            return json.dumps(
                {
                    "success": True,
                    "path": args["path"],
                    "operation": "write_docx_content",
                    "summary": "已写入 6 个段落到 Word 文档",
                    "paragraphs_written": 6,
                },
                ensure_ascii=False,
            )
        if tool_name == "verify_task_completion":
            return json.dumps(
                {"completed": True, "summary": "已写入。"}, ensure_ascii=False
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    events = list(
        FileTaskRuntime(tool_executor=fake_executor, model_client=fake_model).run(
            FileTaskRequest(
                task="这是一篇非常长的pdf，分步总结整篇文章，创建一个docx记录每一步发现，每完成一步等我继续。",
                run_id="stepwise_probe_style_structure_demo",
                target_path="museum-summary.docx",
                files=[FileTaskFile(path="museum.pdf", name="museum.pdf", type="pdf")],
                options={
                    "workflow_checkpoint": {
                        "policy": "confirm_each_step",
                        "step_index": 3,
                        "window_pages": 3,
                        "source_path": "museum.pdf",
                        "original_task": "这是一篇非常长的pdf，分步总结整篇文章，创建一个docx记录每一步发现，每完成一步等我继续。",
                    }
                },
            )
        )
    )

    assert any(
        event.type == "file.changed"
        and event.payload.get("operation") == "write_docx_content"
        for event in events
    )
    check_finished = next(event for event in events if event.type == "check.finished")
    assert check_finished.payload["status"] == "awaiting_confirmation"


def test_file_task_runtime_classifies_resume_requests_before_plan_creation():
    def fake_model(**kwargs):
        return {"content": "等待继续执行", "tool_calls": []}

    request = FileTaskRequest(
        task="继续第 1/3 批",
        run_id="resume_classification_demo",
        target_path="translation.docx",
        files=[
            FileTaskFile(path="source.pdf", name="source.pdf", type="pdf"),
            FileTaskFile(
                path="translation.docx",
                name="translation.docx",
                type="docx",
                target=True,
                content="现有译稿",
            ),
        ],
        options={
            "workflow_checkpoint": {
                "adapter": "doc_annotate_bridge",
                "policy": "confirm_each_batch",
                "batch_index": 0,
                "total_batches": 3,
            }
        },
    )

    events = list(
        FileTaskRuntime(
            tool_executor=lambda name, args: "", model_client=fake_model
        ).run(request)
    )
    event_types = [event.type for event in events]
    classified_index = event_types.index("task.classified")
    plan_checked_index = event_types.index("plan.checked")
    run_started = events[0]
    classified = events[classified_index]

    assert classified_index > 0
    assert plan_checked_index > 0
    assert classified_index < plan_checked_index
    assert run_started.payload["request_kind"] == "resume"
    assert run_started.payload["task_family"] == "annotate"
    assert run_started.payload["execution_mode"] == "awaiting_confirmation_resume"
    assert run_started.payload["docx_annotation_request"] is True
    assert "workflow_checkpoint_resume" in run_started.payload["reason_codes"]
    assert classified.payload["classification"]["request_kind"] == "resume"
    assert classified.payload["classification"]["task_family"] == "annotate"
    assert (
        classified.payload["classification"]["execution_mode"]
        == "awaiting_confirmation_resume"
    )
    assert classified.payload["intent_plan"]["intent_type"] == "annotate"


def test_file_task_runtime_workflow_checkpoint_preserves_original_write_intent():
    request = FileTaskRequest(
        task="继续下一步",
        run_id="generic_stepwise_write_resume",
        target_path="summary.docx",
        files=[FileTaskFile(path="source.pdf", name="source.pdf", type="pdf")],
        options={
            "workflow_checkpoint": {
                "adapter": "generic_tool_loop",
                "policy": "confirm_each_step",
                "step_index": 2,
                "original_task": "这是一篇非常长的pdf，分步总结整篇文章，创建一个docx记录每一步发现，每完成一步等我继续。",
            }
        },
    )

    classification = FileTaskRuntime(
        tool_executor=lambda name, args: ""
    )._classify_request(request, request.files)

    assert classification.request_kind == "resume"
    assert classification.write_intent is True
    assert classification.output_mode == "write"
    assert classification.selected_recipe == "long_pdf_stepwise_docx_summary"
    assert classification.execution_mode == "awaiting_confirmation_resume"


def test_meta_keyword_mentions_do_not_trigger_stepwise_docx_polish_recipe():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task=(
            "请分步处理 report.docx，但这一步只追加一句到目标 DOCX 末尾。"
            "任务描述里故意包含总结、检查、润色、继续下一步这些词，"
            "但不要触发快捷动作关键词路由。"
        ),
        target_path="report.docx",
        files=[FileTaskFile(path="report.docx", name="report.docx", type="docx", target=True)],
    )

    classification = runtime._normalize_mainline_contract(
        request, request.files, runtime._classify_request(request, request.files)
    )

    assert classification.selected_recipe != "long_docx_stepwise_polish_writeback"
    assert classification.execution_mode != "long_docx_stepwise_polish_writeback"


def test_file_task_runtime_does_not_execute_stepwise_polish_by_helper_bypass():
    runtime_source = Path("app/core/agent/file_task_runtime.py").read_text(
        encoding="utf-8"
    )
    helper_source = Path("app/core/agent/_file_task_stepwise_helpers.py").read_text(
        encoding="utf-8"
    )

    assert "looks_like_stepwise_docx_polish_task" not in helper_source
    assert "_looks_like_stepwise_docx_polish_task" not in runtime_source
    assert 'classification.selected_recipe or "").strip()' in runtime_source
    assert '"long_docx_stepwise_polish_writeback"' in runtime_source


def test_workflow_state_does_not_create_docx_windows_from_meta_polish_keywords():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task=(
            "请分步处理 report.docx，但这一步只追加一句到目标 DOCX 末尾。"
            "任务描述里故意包含总结、检查、润色、继续下一步这些词，"
            "但不要触发快捷动作关键词路由。"
        ),
        target_path="report.docx",
        files=[
            FileTaskFile(path="notes.docx", name="notes.docx", type="docx"),
            FileTaskFile(
                path="report.docx", name="report.docx", type="docx", target=True
            ),
        ],
    )
    classification = runtime._normalize_mainline_contract(
        request, request.files, runtime._classify_request(request, request.files)
    )
    state = build_workflow_state(
        request,
        request.files,
        classification,
        {"recipe_id": "generic_file_task"},
    )

    assert state["large_file_windows"] == []


def test_file_task_runtime_stepwise_resume_with_target_docx_keeps_stepwise_recipe():
    request = FileTaskRequest(
        task="继续当前分步文件任务的下一步：处理 PDF 第 4-6 页，并把本段实质分析追加到同一个 DOCX。",
        run_id="generic_stepwise_write_resume_with_target_docx",
        target_path="summary.docx",
        files=[
            FileTaskFile(path="source.pdf", name="source.pdf", type="pdf"),
            FileTaskFile(
                path="summary.docx", name="summary.docx", type="docx", target=True
            ),
        ],
        options={
            "workflow_checkpoint": {
                "adapter": "generic_tool_loop",
                "policy": "confirm_each_step",
                "step_index": 1,
                "window_pages": 3,
                "source_path": "source.pdf",
                "target_path": "summary.docx",
            }
        },
    )

    classification = FileTaskRuntime(
        tool_executor=lambda name, args: ""
    )._classify_request(request, request.files)

    assert classification.request_kind == "resume"
    assert classification.write_intent is True
    assert classification.selected_recipe == "long_pdf_stepwise_docx_summary"
    assert classification.operation_kind == "stepwise_write"
    assert "stepwise_resume_forced_write_intent" in classification.reason_codes


def test_file_task_runtime_native_stepwise_writes_before_pause():
    model_calls = []

    def fake_model(**kwargs):
        model_calls.append(kwargs)
        raise RuntimeError("model unavailable for fallback test")

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            return (
                "[Page 1] 当前页窗介绍博物馆数字技术应用的研究背景，包含数字化保护、数据治理、观众服务和展陈创新。"
                "[Page 2] 目录与引言说明报告将按理论综述和案例研究展开，覆盖数字敦煌、知识图谱、沉浸式展览等方向。"
            )
        if tool_name == "write_docx_content":
            paragraphs = json.loads(args["paragraphs"])
            return json.dumps(
                {
                    "success": True,
                    "path": args["path"],
                    "operation": "write_docx_content",
                    "summary": f"已写入 {len(paragraphs)} 个段落到 Word 文档",
                    "paragraphs_written": len(paragraphs),
                },
                ensure_ascii=False,
            )
        if tool_name == "verify_task_completion":
            return json.dumps(
                {"completed": True, "summary": "已写入。"}, ensure_ascii=False
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor,
            model_client=fake_model,
            max_rounds=2,
        ).run(
            FileTaskRequest(
                task="这是一篇非常长的pdf，分步总结，创建一个docx文件记录每一步发现，每完成一步等我继续。",
                run_id="stepwise_pause_without_write_demo",
                files=[FileTaskFile(path="source.pdf", name="source.pdf", type="pdf")],
            )
        )
    )

    assert len(model_calls) == 1
    assert any(
        event.type == "file.changed"
        and event.payload.get("operation") == "write_docx_content"
        for event in events
    )
    check_finished = next(event for event in events if event.type == "check.finished")
    assert check_finished.payload["status"] == "awaiting_confirmation"
    assert check_finished.payload["passed"] is False
    assert (
        check_finished.payload["next_action_artifact"]["route"]
        == "long_pdf_stepwise_docx_summary"
    )


def test_file_task_runtime_routes_long_docx_stepwise_polish_writeback(tmp_path):
    from docx import Document

    docx_path = tmp_path / "draft.docx"
    doc = Document()
    for index in range(1, 7):
        doc.add_paragraph(f"第{index}段  有一些  不够顺畅 的 表达。")
    doc.save(docx_path)

    calls = []

    def fake_model(**kwargs):
        calls.append(kwargs)
        user_content = kwargs["messages"][-1]["content"]
        if "只返回 JSON 字符串数组" not in user_content:
            return {
                "content": json.dumps(
                    {
                        "intent": "edit_file",
                        "confidence": 0.95,
                        "should_write": True,
                        "target_file_type": "docx",
                    },
                    ensure_ascii=False,
                ),
                "tool_calls": [],
            }
        return {
            "content": json.dumps(
                [
                    "第1段的表达已润色得更加顺畅。",
                    "第2段的表达已润色得更加顺畅。",
                ],
                ensure_ascii=False,
            ),
            "tool_calls": [],
        }

    request = FileTaskRequest(
        task="这是一篇非常长的 DOCX，请分步润色，每完成一步写回文档并等待我说继续。",
        run_id="docx_stepwise_polish",
        target_path=str(docx_path),
        files=[
            FileTaskFile(
                path=str(docx_path), name="draft.docx", type="docx", target=True
            )
        ],
        options={
            "workflow_checkpoint": {"policy": "confirm_each_step", "window_paragraphs": 2}
        },
    )
    events = list(FileTaskRuntime(model_client=fake_model, max_rounds=2).run(request))

    final_payload = events[-1].payload
    changed = [
        event.payload
        for event in events
        if event.type == "file.changed"
        and event.payload.get("operation") == "rewrite_docx_paragraph_window"
    ]
    assert changed
    assert changed[0]["paragraphs_rewritten"] == 2
    assert final_payload["runtime"]["terminal_status"] == "awaiting_confirmation"
    assert (
        final_payload["next_action_artifact"]["route"]
        == "long_docx_stepwise_polish_writeback"
    )
    action_options = final_payload["next_action_artifact"]["actions"][0]["request"][
        "options"
    ]
    assert "batch_control" not in action_options
    assert action_options["workflow_checkpoint"]["step_index"] == 1

    updated = Document(docx_path)
    assert updated.paragraphs[0].text == "第1段的表达已润色得更加顺畅。"
    assert updated.paragraphs[1].text == "第2段的表达已润色得更加顺畅。"
    assert updated.paragraphs[2].text.startswith("第3段")
    assert any("只返回 JSON 字符串数组" in call["messages"][-1]["content"] for call in calls)


def test_file_task_runtime_stepwise_pdf_falls_back_when_model_never_writes(tmp_path):
    pdf_path = tmp_path / "source.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n% fake pdf bytes for executor\n")
    docx_path = tmp_path / "summary.docx"

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            return (
                "[Page 1]\n"
                "Section one discusses museum digitization and audience needs. "
                "The report explains how cultural institutions use digital systems "
                "for collection management, public interpretation, exhibition design, "
                "visitor research, and long term heritage preservation. It also "
                "connects technology adoption with institutional strategy, learning "
                "goals, data governance, and service quality improvements."
            )
        if tool_name == "write_docx_content":
            from app.core.agent.task_tools import write_docx_content

            return write_docx_content(**args)
        if tool_name == "verify_task_completion":
            return {"passed": True, "status": "verified", "summary": "ok"}
        raise AssertionError(f"unexpected tool: {tool_name}")

    def fake_model(**kwargs):
        return {
            "content": "I only analyzed the page and forgot to write.",
            "tool_calls": [],
        }

    request = FileTaskRequest(
        task=(
            "This is a very long PDF. Summarize it stepwise into a DOCX, "
            "write each step, and wait for confirmation."
        ),
        run_id="pdf_no_write_fallback",
        target_path=str(docx_path),
        files=[
            FileTaskFile(path=str(pdf_path), name="source.pdf", type="pdf"),
            FileTaskFile(
                path=str(docx_path), name="summary.docx", type="docx", target=True
            ),
        ],
    )

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor, model_client=fake_model, max_rounds=2
        ).run(request)
    )

    changes = [
        event.payload
        for event in events
        if event.type == "file.changed"
        and event.payload.get("operation") == "write_docx_content"
    ]
    assert changes
    assert docx_path.exists()
    assert events[-1].payload["runtime"]["terminal_status"] == "awaiting_confirmation"


def test_file_task_runtime_stepwise_pdf_skips_low_density_front_window(tmp_path):
    pdf_path = tmp_path / "source.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n% fake pdf bytes for executor\n")
    docx_path = tmp_path / "summary.docx"
    parse_calls = []

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            parse_calls.append(dict(args))
            if int(args.get("start_page") or 1) == 1:
                return "[Page 1]\nTitle only"
            return (
                "[Page 4]\n"
                "This usable window contains enough structured content about museum "
                "digital transformation, collection databases, audience research, "
                "exhibition interpretation, data governance, public services, "
                "immersive display, knowledge graphs, and preservation workflows."
            )
        if tool_name == "write_docx_content":
            return json.dumps(
                {
                    "success": True,
                    "path": args["path"],
                    "file_type": "docx",
                    "operation": "write_docx_content",
                    "change_type": "create",
                    "summary": "已写入 6 个段落到 Word 文档",
                    "paragraphs_written": 6,
                },
                ensure_ascii=False,
            )
        if tool_name == "verify_task_completion":
            return {"passed": True, "status": "verified", "summary": "ok"}
        raise AssertionError(f"unexpected tool: {tool_name}")

    def fake_model(**kwargs):
        return {
            "content": "I only analyzed the page and forgot to write.",
            "tool_calls": [],
        }

    request = FileTaskRequest(
        task=(
            "This is a very long PDF. Summarize it stepwise into a DOCX, "
            "write each step, and wait for confirmation."
        ),
        run_id="pdf_skip_front_window",
        target_path=str(docx_path),
        files=[
            FileTaskFile(path=str(pdf_path), name="source.pdf", type="pdf"),
            FileTaskFile(
                path=str(docx_path), name="summary.docx", type="docx", target=True
            ),
        ],
    )

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor, model_client=fake_model, max_rounds=2
        ).run(request)
    )
    context_result = next(
        event.payload
        for event in events
        if event.type == "step.result" and event.step_id == "context"
    )

    assert [call["start_page"] for call in parse_calls[:2]] == [1, 4]
    assert context_result["snippets"][0]["start_page"] == 4
    assert events[-1].payload["runtime"]["terminal_status"] == "awaiting_confirmation"
