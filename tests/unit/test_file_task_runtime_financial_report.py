import base64
import json
from pathlib import Path

from app.core.agent.file_task_contract import FileTaskFile, FileTaskRequest
from app.core.agent.file_task_runtime import FileTaskRuntime


def test_file_task_runtime_plans_financial_chart_docx_report_as_problem_list_and_image():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="将xlsx财务预测数据做成图，并分析存在的问题，将问题和图加入docx",
        target_path="report.docx",
        files=[
            FileTaskFile(
                path="financial.xlsx", name="雷鸟创新-financial model.xlsx", type="xlsx"
            ),
            FileTaskFile(
                path="report.docx", name="雷鸟访谈问题.docx", type="docx", target=True
            ),
        ],
    )

    step = runtime._inferred_write_plan_step(request, request.files)
    retry = runtime._write_retry_message(request, request.files)
    repair = runtime._repair_retry_message(
        request,
        {
            "status": "write_not_performed",
            "summary": "任务包含写入意图，但没有任何工具报告文件变更。",
        },
        [],
    )

    assert step["title"] == "写入问题和图表"
    assert "生成真实财务图表图片并整理问题清单" in step["description"]
    assert "write_docx_content" in retry
    assert "insert_image_into_docx" in retry
    assert "insert_excel_as_docx_table" in retry
    assert "列名是 Unnamed" in retry
    assert "write_docx_content" in repair
    assert "insert_image_into_docx" in repair
    assert "不要用 df.columns 判断年份列" in repair


def test_file_task_runtime_classifies_semantic_task_profile_for_financial_report():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="分析这个xlsx财务数据的问题，并将数据做成图，然后把问题和图加入docx",
        target_path="report.docx",
        files=[
            FileTaskFile(path="financial.xlsx", name="financial.xlsx", type="xlsx"),
            FileTaskFile(
                path="report.docx", name="report.docx", type="docx", target=True
            ),
        ],
    )

    classification = runtime._classify_request(request, request.files)

    assert classification.task_family == "financial_report"
    assert classification.operation_kind == "analyze_visualize_write"
    assert classification.output_mode == "write"
    assert classification.write_intent is True
    assert classification.selected_recipe == "financial_xlsx_docx_report"
    assert "financial_request" in classification.reason_codes
    assert "chart_request" in classification.reason_codes
    assert "docx_report_request" in classification.reason_codes
    assert "recipe:financial_xlsx_docx_report" in classification.reason_codes


def test_file_task_runtime_classifies_new_financial_docx_report_without_existing_docx():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="分析这个财务预测，找出其中的问题并把数据做成图，把数据和图都放入一个新的docx",
        files=[
            FileTaskFile(
                path="financial_model_clean.xlsx",
                name="financial_model_clean.xlsx",
                type="xlsx",
            )
        ],
    )

    classification = runtime._classify_request(request, request.files)

    assert classification.task_family == "financial_report"
    assert classification.output_mode == "write"
    assert classification.selected_recipe == "financial_xlsx_docx_report"


def test_file_task_runtime_classifies_exact_financial_report_request_from_ui():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="分析这个xlsx财务预测，将数据做成图，并找出其中的问题，将问题和图放入一个新建的docx",
        files=[
            FileTaskFile(
                path="financial_model.xlsx",
                name="financial_model.xlsx",
                type="xlsx",
            )
        ],
    )

    classification = runtime._classify_request(request, request.files)

    assert classification.write_intent is True
    assert classification.task_family == "financial_report"
    assert classification.selected_recipe == "financial_xlsx_docx_report"
    table_args = runtime._repair_tool_args_for_context(
        "insert_excel_as_docx_table",
        {
            "source_path": "financial_model.xlsx",
            "target_path": "financial_report.docx",
            "max_rows": 60,
        },
        FileTaskRequest(
            task=request.task,
            target_path="financial_report.docx",
            files=request.files,
        ),
        request.files,
    )
    assert table_args["financial_compact"] is True
    assert table_args["max_rows"] == 18


def test_new_financial_docx_report_gets_target_and_word_tools_before_first_round():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="分析这个财务预测，找出其中的问题并把数据做成图，把数据和图都放入一个新的docx",
        files=[
            FileTaskFile(
                path="tmp/run/financial_model_clean.xlsx",
                name="financial_model_clean.xlsx",
                type="xlsx",
            )
        ],
    )

    normalized = runtime._request_with_inferred_target_path(request)
    gateway = runtime._build_tool_gateway(normalized, normalized.files)
    tool_names = {item["name"] for item in gateway.definitions()}

    assert normalized.target_path.replace("\\", "/") == (
        "tmp/run/financial_model_clean_财务预测分析报告.docx"
    )
    assert {
        "write_docx_content",
        "insert_excel_as_docx_table",
        "insert_image_into_docx",
    }.issubset(tool_names)


def test_new_financial_docx_report_avoids_existing_partial_output(tmp_path):
    source = tmp_path / "financial_model.xlsx"
    source.write_bytes(b"source")
    existing = tmp_path / "financial_model_财务预测分析报告.docx"
    existing.write_bytes(b"partial report from failed task")
    second = tmp_path / "financial_model_财务预测分析报告_2.docx"
    second.write_bytes(b"another prior report")
    runtime = FileTaskRuntime(
        tool_executor=lambda name, args: "",
        workspace_root=str(tmp_path),
    )
    request = FileTaskRequest(
        task="分析这个财务预测，将数据做成图并分析问题，放入一个新的docx",
        files=[
            FileTaskFile(
                path=source.name,
                name=source.name,
                type="xlsx",
            )
        ],
    )

    normalized = runtime._request_with_inferred_target_path(request)

    assert normalized.target_path == "financial_model_财务预测分析报告_3.docx"
    assert existing.read_bytes() == b"partial report from failed task"
    assert second.read_bytes() == b"another prior report"


def test_financial_chart_recovery_is_available_when_chart_code_reports_no_artifact(
    tmp_path,
):
    source = tmp_path / "financial_model_clean.xlsx"
    source.write_bytes(b"placeholder workbook path for recovery selection")
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="分析这个财务预测，找出其中的问题并把数据做成图，把数据和图都放入一个新的docx",
        target_path=str(tmp_path / "financial_report.docx"),
        files=[FileTaskFile(path=str(source), name=source.name, type="xlsx")],
    )

    recovery = runtime._financial_chart_recovery_tool_args(
        request,
        request.files,
        {"code": "import matplotlib.pyplot as plt\nplt.savefig('chart.png')"},
        [],
    )

    assert recovery["timeout"] == 120
    assert json.dumps(str(source), ensure_ascii=False) in recovery["code"]
    assert recovery["code"].count("KOTO_CREATED:") == 2
    assert "financial_pnl_trend.png" in recovery["code"]
    assert "product_sales_revenue_structure.png" in recovery["code"]

    no_recovery = runtime._financial_chart_recovery_tool_args(
        request,
        request.files,
        {"code": "print('inspect only')"},
        [],
    )
    assert no_recovery == {}


def test_financial_report_recovers_through_native_tools_when_model_times_out(tmp_path):
    import openpyxl
    from docx import Document

    workbook_path = tmp_path / "financial_model.xlsx"
    target_path = tmp_path / "financial_model_财务预测分析报告.docx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "P&L"
    sheet.append(["", "", "2025E", "2026E", "2027E"])
    sheet.append(["", "收入合计", 1000, 1800, 3500])
    sheet.append(["", "毛利合计", 400, 900, 2000])
    sheet.append(["", "综合毛利率", 0.4, 0.5, 0.57])
    sheet.append(["", "净利润", -100, 100, 600])
    sheet.append(["", "净利率", -0.1, 0.055, 0.171])
    sheet.append(["", "所得税费用", 0, -10, 100])
    workbook.save(workbook_path)

    def timeout_model(**_kwargs):
        raise TimeoutError("provider timed out after 45 seconds")

    request = FileTaskRequest(
        task="分析这个xlsx财务预测，将数据做成图，并找出其中的问题，将问题和图放入一个新建的docx",
        run_id="financial_native_timeout_recovery",
        target_path=str(target_path),
        files=[
            FileTaskFile(
                path=str(workbook_path),
                name=workbook_path.name,
                type="xlsx",
            )
        ],
    )
    events = list(
        FileTaskRuntime(
            model_client=timeout_model,
            workspace_root=str(tmp_path),
            max_rounds=1,
        ).run(request)
    )

    event_types = [event.type for event in events]
    recovery_finished = next(
        event for event in events if event.type == "recovery.finished"
    )
    check_finished = next(
        event for event in events if event.type == "check.finished"
    )
    run_finished = events[-1]
    document = Document(target_path)

    assert "run.error" not in event_types
    assert recovery_finished.payload["success"] is True
    assert check_finished.payload["passed"] is True
    assert run_finished.type == "run.finished"
    assert run_finished.payload["completed_task"] is True
    assert "failure" not in run_finished.payload
    assert len([paragraph for paragraph in document.paragraphs if paragraph.text.strip()]) >= 12
    assert len(document.tables) >= 1
    assert len(document.inline_shapes) >= 2
    assert "发现的问题与风险" in "\n".join(
        paragraph.text for paragraph in document.paragraphs
    )
    assert not list(tmp_path.glob(".*.koto-partial.docx"))
    assert all(
        "koto-partial" not in str(event.payload)
        for event in events
        if event.type == "file.changed"
    )


def test_failed_financial_recovery_removes_partial_docx_and_uncommitted_charts(
    tmp_path,
):
    import openpyxl
    from docx import Document

    from app.core.agent.file_task_contract import FileTaskEvent
    from app.core.agent.file_task_financial_report_recovery import (
        recover_financial_report,
    )

    workbook_path = tmp_path / "financial_model.xlsx"
    target_path = tmp_path / "financial_report.docx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "P&L"
    sheet.append(["", "", "2025E", "2026E"])
    sheet.append(["", "收入合计", 1000, 1800])
    sheet.append(["", "毛利合计", 400, 900])
    sheet.append(["", "净利润", -100, 100])
    workbook.save(workbook_path)

    chart_paths = [
        tmp_path / "financial_model_financial_pnl_trend.png",
        tmp_path / "financial_model_product_sales_revenue_structure.png",
    ]

    class RuntimeStub:
        def _resolve_task_file_path(self, path):
            return str(path) if Path(str(path)).exists() else ""

        def _financial_chart_recovery_tool_args(
            self, _request, _files, _tool_args, _artifacts
        ):
            return {"code": "create charts"}

        def _extract_file_changes(self, _tool_name, _tool_args, result):
            return [dict(result)] if isinstance(result, dict) else []

        def _tool_artifacts(self, tool_name, result):
            if tool_name == "run_python_code" and isinstance(result, dict):
                return list(result.get("artifacts") or [])
            return []

    class LedgerStub:
        def __init__(self):
            self.seq = 0

        def event(self, event_type, payload, *, step_id=""):
            self.seq += 1
            return FileTaskEvent(
                type=event_type,
                run_id="failed_atomic_recovery",
                seq=self.seq,
                payload=dict(payload),
                step_id=step_id,
            )

    def executor(tool_name, args):
        if tool_name == "audit_financial_workbook":
            return {"findings": []}
        if tool_name == "run_python_code":
            for chart_path in chart_paths:
                chart_path.write_bytes(b"partial-chart")
            return {
                "artifacts": [
                    {"kind": "image", "path": str(path)} for path in chart_paths
                ]
            }
        if tool_name == "write_docx_content":
            Document().save(args["path"])
            return {
                "path": args["path"],
                "operation": "write_docx_content",
                "paragraphs_written": 1,
            }
        if tool_name == "insert_excel_as_docx_table":
            return {"error": "table insertion failed"}
        if tool_name == "insert_image_into_docx":
            return {
                "path": args["path"],
                "operation": "insert_image_into_docx",
                "image_path": args["image_path"],
                "images_inserted": 1,
            }
        return {"error": f"unexpected tool: {tool_name}"}

    request = FileTaskRequest(
        task="分析财务预测、生成图表并写入新的 docx",
        run_id="failed_atomic_recovery",
        target_path=str(target_path),
        files=[
            FileTaskFile(
                path=str(workbook_path),
                name=workbook_path.name,
                type="xlsx",
            )
        ],
    )
    stream = recover_financial_report(
        RuntimeStub(),
        LedgerStub(),
        request,
        executor,
        request.files,
        {"recipe_id": "financial_xlsx_docx_report"},
        step_id="execute",
    )
    events = []
    while True:
        try:
            events.append(next(stream))
        except StopIteration as stop:
            result = stop.value
            break

    assert result.attempted is True
    assert result.completed is False
    assert result.file_changes == []
    assert result.artifacts == []
    assert not target_path.exists()
    assert not list(tmp_path.glob(".*.koto-partial.docx"))
    assert all(not path.exists() for path in chart_paths)
    assert all(event.type != "file.changed" for event in events)


def test_file_task_runtime_does_not_treat_complex_as_financial_request():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task=(
            "请基于当前打开的 koto_task_smoke.txt 完成一个复杂文件任务，"
            "并保存为 _codex_frontend_task_tests/koto_complex_task_report_20260617_1345.md"
        ),
        target_path="workspace/_codex_frontend_task_tests/koto_task_smoke.txt",
        files=[
            FileTaskFile(
                path="workspace/_codex_frontend_task_tests/koto_task_smoke.txt",
                name="koto_task_smoke.txt",
                type="txt",
                target=True,
            )
        ],
    )

    classification = runtime._classify_request(request, request.files)

    assert "financial_request" not in classification.reason_codes


def test_file_task_runtime_quality_gate_rejects_financial_report_without_chart_image():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="分析这个xlsx财务数据的问题，并将数据做成图，然后把问题和图加入docx",
        target_path="report.docx",
        files=[
            FileTaskFile(path="financial.xlsx", name="financial.xlsx", type="xlsx"),
            FileTaskFile(
                path="report.docx", name="report.docx", type="docx", target=True
            ),
        ],
    )
    file_changes = [
        {
            "operation": "write_docx_content",
            "path": "report.docx",
            "file_type": "docx",
            "paragraphs_written": 12,
        },
        {
            "operation": "insert_excel_as_docx_table",
            "path": "report.docx",
            "file_type": "docx",
            "rows_written": 6,
        },
    ]

    def fake_executor(tool_name, args):
        assert tool_name == "verify_task_completion"
        return json.dumps(
            {"completed": True, "summary": "文件已更新。"}, ensure_ascii=False
        )

    check = runtime._verify_task(
        request,
        fake_executor,
        file_changes,
        write_intent=True,
        output_mode="write",
        model_failed=False,
    )

    assert check["passed"] is False
    assert check["status"] == "quality_gate_failed"
    assert "真实图表图片" in check["summary"]
    assert any(
        item["criterion"] == "financial_report_has_real_chart_image"
        for item in check["criteria_results"]
    )
    assert any("真实图表图片" in item for item in check["remaining"])


def test_file_task_runtime_financial_report_quality_requires_table_and_two_images():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="分析这个财务预测，找出其中的问题并把数据做成图，把数据和图都放入一个新的docx",
        files=[
            FileTaskFile(
                path="financial_model_clean.xlsx",
                name="financial_model_clean.xlsx",
                type="xlsx",
            )
        ],
    )

    one_chart = runtime._evaluate_task_quality_gate(
        request,
        [
            {"operation": "write_docx_content", "paragraphs_written": 10},
            {
                "operation": "insert_excel_as_docx_table",
                "rows_written": 6,
            },
            {"operation": "insert_image_into_docx", "images_inserted": 1},
        ],
        write_intent=True,
        output_mode="write",
    )
    no_table = runtime._evaluate_task_quality_gate(
        request,
        [
            {"operation": "write_docx_content", "paragraphs_written": 10},
            {"operation": "insert_image_into_docx", "images_inserted": 2},
        ],
        write_intent=True,
        output_mode="write",
    )

    assert one_chart["passed"] is False
    assert any(
        item["criterion"] == "financial_report_has_real_chart_image"
        and item["passed"] is False
        for item in one_chart["criteria_results"]
    )
    assert no_table["passed"] is False
    assert any(
        item["criterion"] == "financial_report_has_key_data_table"
        and item["passed"] is False
        for item in no_table["criteria_results"]
    )


def test_financial_report_quality_rejects_unreadable_raw_wide_table():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="分析财务预测、生成图表并写入新的 docx",
        target_path="report.docx",
        files=[
            FileTaskFile(path="financial.xlsx", name="financial.xlsx", type="xlsx"),
            FileTaskFile(path="report.docx", name="report.docx", type="docx", target=True),
        ],
    )
    quality = runtime._evaluate_task_quality_gate(
        request,
        [
            {"operation": "write_docx_content", "paragraphs_written": 12},
            {
                "operation": "insert_excel_as_docx_table",
                "rows_written": 60,
                "columns_written": 13,
            },
            {"operation": "insert_image_into_docx", "images_inserted": 2},
        ],
        write_intent=True,
        output_mode="write",
    )

    assert quality["passed"] is False
    assert any(
        item["criterion"] == "financial_report_has_readable_key_data_table"
        and item["passed"] is False
        for item in quality["criteria_results"]
    )


def test_financial_report_quality_rejects_duplicate_report_body(tmp_path):
    from docx import Document

    target = tmp_path / "financial_report.docx"
    document = Document()
    for _ in range(2):
        document.add_paragraph("财务预测分析报告", style="Title")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "指标"
        table.cell(0, 1).text = "2025E"
        table.cell(1, 0).text = "收入"
        table.cell(1, 1).text = "1000"
    document.save(target)

    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="分析财务预测、生成图表并写入新的 docx",
        target_path=str(target),
        files=[
            FileTaskFile(path="financial.xlsx", name="financial.xlsx", type="xlsx"),
            FileTaskFile(
                path=str(target),
                name=target.name,
                type="docx",
                target=True,
            ),
        ],
    )

    quality = runtime._evaluate_task_quality_gate(
        request,
        [
            {
                "operation": "write_docx_content",
                "path": str(target),
                "paragraphs_written": 12,
            },
            {
                "operation": "insert_excel_as_docx_table",
                "path": str(target),
                "rows_written": 6,
                "columns_written": 5,
            },
            {
                "operation": "insert_image_into_docx",
                "path": str(target),
                "images_inserted": 2,
            },
        ],
        write_intent=True,
        output_mode="write",
    )

    assert quality["passed"] is False
    assert any(
        item["criterion"] == "financial_report_has_no_duplicate_report_body"
        and item["passed"] is False
        for item in quality["criteria_results"]
    )


def test_file_task_runtime_quality_gate_rejects_generated_image_not_inserted_into_docx():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="将数据做成图，并把图加入docx",
        target_path="report.docx",
        files=[
            FileTaskFile(
                path="report.docx", name="report.docx", type="docx", target=True
            ),
        ],
    )

    result = runtime._evaluate_task_quality_gate(
        request,
        [
            {
                "operation": "run_python_code",
                "path": "chart1_revenue_profit_trend.png",
                "file_type": "png",
                "images_inserted": 1,
            }
        ],
        write_intent=True,
        output_mode="write",
    )

    assert result["passed"] is False
    assert any(
        item["criterion"] == "docx_chart_request_has_image" and item["passed"] is False
        for item in result["criteria_results"]
    )
    assert any("进入 Word" in item for item in result["remaining"])


def test_file_task_runtime_repairs_docx_image_path_from_generated_artifact(tmp_path):
    image_path = tmp_path / "chart_expense_structure.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+tm0YAAAAASUVORK5CYII="
        )
    )
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="分析这个xlsx财务模型，将数据做成图，把数据的问题和做成的图都加入docx",
        target_path="report.docx",
        files=[
            FileTaskFile(path="financial.xlsx", name="financial.xlsx", type="xlsx"),
            FileTaskFile(
                path="report.docx", name="report.docx", type="docx", target=True
            ),
        ],
    )

    args = runtime._repair_tool_args_for_context(
        "insert_image_into_docx",
        {
            "path": "report.docx",
            "image_path": r"C:\Users\12524\AppData\Local\Temp\koto-task-bcprm2qo\chart_expense_structure.png",
        },
        request,
        request.files,
        generated_artifacts=[
            {
                "kind": "image",
                "name": "chart_expense_structure.png",
                "path": str(image_path),
            }
        ],
    )

    assert args["image_path"] == str(image_path)


def test_file_task_runtime_requires_all_generated_chart_images_inserted(tmp_path):
    image_payload = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+tm0YAAAAASUVORK5CYII="
    image1 = tmp_path / "chart1_revenue_profit_trend.png"
    image2 = tmp_path / "chart2_product_mix.png"
    image1.write_bytes(base64.b64decode(image_payload))
    image2.write_bytes(base64.b64decode(image_payload))

    seen_messages = []
    responses = iter(
        [
            {
                "content": "生成图表并先写入报告。",
                "tool_calls": [
                    {
                        "name": "run_python_code",
                        "args": {"code": "create two charts"},
                    },
                    {
                        "name": "write_docx_content",
                        "args": {
                            "path": "report.docx",
                            "paragraphs": json.dumps(
                                [{"text": f"问题分析 {index}"} for index in range(10)],
                                ensure_ascii=False,
                            ),
                        },
                    },
                    {
                        "name": "insert_excel_as_docx_table",
                        "args": {
                            "source_path": "financial.xlsx",
                            "target_path": "report.docx",
                            "sheet_name": "P&L",
                            "table_title": "关键预测数据",
                        },
                    },
                    {
                        "name": "insert_image_into_docx",
                        "args": {
                            "path": "report.docx",
                            "image_path": str(image1),
                            "caption": "收入和利润趋势",
                        },
                    },
                ],
            },
            {
                "content": "继续补齐剩余图表。",
                "tool_calls": [
                    {
                        "name": "insert_image_into_docx",
                        "args": {
                            "path": "report.docx",
                            "image_path": str(image2),
                            "caption": "产品结构",
                        },
                    }
                ],
            },
        ]
    )

    def fake_model(**kwargs):
        seen_messages.append(str(kwargs["messages"][-1]["content"]))
        return next(responses, {"content": "完成", "tool_calls": []})

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            return "财务预测上下文"
        if tool_name == "run_python_code":
            return {
                "summary": "已生成 2 张图表",
                "files": {
                    image1.name: image_payload,
                    image2.name: image_payload,
                },
                "generated_file_paths": {
                    image1.name: str(image1),
                    image2.name: str(image2),
                },
            }
        if tool_name == "write_docx_content":
            return {
                "path": "report.docx",
                "file_type": "docx",
                "operation": "write_docx_content",
                "summary": "已写入分析段落",
                "paragraphs_written": 10,
            }
        if tool_name == "insert_excel_as_docx_table":
            return {
                "path": "report.docx",
                "file_type": "docx",
                "operation": "insert_excel_as_docx_table",
                "summary": "已写入关键预测数据表",
                "rows_written": 6,
                "columns_written": 5,
            }
        if tool_name == "insert_image_into_docx":
            image_path = str(args.get("image_path") or "")
            return {
                "path": "report.docx",
                "file_type": "docx",
                "operation": "insert_image_into_docx",
                "summary": f"已插入 {Path(image_path).name}",
                "image_path": image_path,
                "image_name": Path(image_path).name,
                "images_inserted": 1,
            }
        if tool_name == "verify_task_completion":
            return json.dumps(
                {"completed": True, "summary": "DOCX 已更新。"}, ensure_ascii=False
            )
        return ""

    request = FileTaskRequest(
        task="分析这个xlsx财务数据的问题，并将数据做成图，然后把问题和图加入docx",
        target_path="report.docx",
        files=[
            FileTaskFile(path="financial.xlsx", name="financial.xlsx", type="xlsx"),
            FileTaskFile(
                path="report.docx", name="report.docx", type="docx", target=True
            ),
        ],
    )

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor,
            model_client=fake_model,
            max_rounds=3,
        ).run(request)
    )

    image_guard = next(
        event
        for event in events
        if event.type == "tool.finished"
        and event.payload.get("tool_name") == "image_insert_guard"
    )
    insert_changes = [
        event
        for event in events
        if event.type == "file.changed"
        and event.payload.get("operation") == "insert_image_into_docx"
    ]
    check_finished = [event for event in events if event.type == "check.finished"][-1]
    run_finished = events[-1]

    assert image_guard.payload["pending_image_count"] == 1
    assert "chart2_product_mix.png" in image_guard.payload["result_preview"]
    assert any("chart2_product_mix.png" in message for message in seen_messages)
    assert [item.payload["image_name"] for item in insert_changes] == [
        "chart1_revenue_profit_trend.png",
        "chart2_product_mix.png",
    ]
    assert check_finished.payload["passed"] is True
    assert run_finished.payload["completed_task"] is True


def test_generated_docx_image_guard_resolves_workspace_relative_artifacts(tmp_path):
    image = tmp_path / "chart1_revenue_profit_trend.png"
    image.write_bytes(base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+tm0YAAAAASUVORK5CYII="))
    runtime = FileTaskRuntime()
    runtime._resolve_task_file_path = lambda value: str(image) if value == image.name else None
    request = FileTaskRequest(
        task="分析财务预测并将图表写入新的 docx",
        target_path="report.docx",
        files=[
            FileTaskFile(path="financial.xlsx", name="financial.xlsx", type="xlsx"),
            FileTaskFile(path="report.docx", name="report.docx", type="docx", target=True),
        ],
    )

    pending = runtime._pending_generated_docx_images(
        request,
        request.files,
        [{"kind": "image", "name": image.name, "path": image.name}],
        [],
    )

    assert pending == [{"kind": "image", "name": image.name, "path": str(image)}]


def test_file_task_runtime_completes_generated_charts_when_model_runs_out_of_rounds(
    tmp_path,
):
    image_payload = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+tm0YAAAAASUVORK5CYII="
    image1 = tmp_path / "chart1_revenue_profit_trend.png"
    image2 = tmp_path / "chart2_product_mix.png"
    image1.write_bytes(base64.b64decode(image_payload))
    image2.write_bytes(base64.b64decode(image_payload))

    def fake_model(**kwargs):
        return {
            "content": "生成图表并写入报告。",
            "tool_calls": [
                {"name": "run_python_code", "args": {"code": "create two charts"}},
                {
                    "name": "write_docx_content",
                    "args": {
                        "path": "report.docx",
                        "paragraphs": json.dumps(
                            [{"text": f"问题分析 {index}"} for index in range(10)],
                            ensure_ascii=False,
                        ),
                    },
                },
                {
                    "name": "insert_excel_as_docx_table",
                    "args": {
                        "source_path": "financial.xlsx",
                        "target_path": "report.docx",
                        "sheet_name": "P&L",
                        "table_title": "关键预测数据",
                    },
                },
                {
                    "name": "insert_image_into_docx",
                    "args": {
                        "path": "report.docx",
                        "image_path": str(image1),
                        "caption": "收入和利润趋势",
                    },
                },
            ],
        }

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            return "财务预测上下文"
        if tool_name == "run_python_code":
            return {
                "summary": "已生成 2 张图表",
                "files": {
                    image1.name: image_payload,
                    image2.name: image_payload,
                },
                "generated_file_paths": {
                    image1.name: str(image1),
                    image2.name: str(image2),
                },
            }
        if tool_name == "write_docx_content":
            return {
                "path": "report.docx",
                "file_type": "docx",
                "operation": "write_docx_content",
                "summary": "已写入分析段落",
                "paragraphs_written": 10,
            }
        if tool_name == "insert_excel_as_docx_table":
            return {
                "path": "report.docx",
                "file_type": "docx",
                "operation": "insert_excel_as_docx_table",
                "summary": "已写入关键预测数据表",
                "rows_written": 6,
                "columns_written": 5,
            }
        if tool_name == "insert_image_into_docx":
            image_path = str(args.get("image_path") or "")
            return {
                "path": "report.docx",
                "file_type": "docx",
                "operation": "insert_image_into_docx",
                "summary": f"已插入 {Path(image_path).name}",
                "image_path": image_path,
                "image_name": Path(image_path).name,
                "images_inserted": 1,
            }
        if tool_name == "verify_task_completion":
            return json.dumps(
                {"completed": True, "summary": "DOCX 已更新。"}, ensure_ascii=False
            )
        return ""

    request = FileTaskRequest(
        task="分析这个xlsx财务数据的问题，并将数据做成图，然后把问题和图加入docx",
        target_path="report.docx",
        files=[
            FileTaskFile(path="financial.xlsx", name="financial.xlsx", type="xlsx"),
            FileTaskFile(
                path="report.docx", name="report.docx", type="docx", target=True
            ),
        ],
    )

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor,
            model_client=fake_model,
            max_rounds=1,
        ).run(request)
    )

    native_inserts = [
        event
        for event in events
        if event.type == "tool.finished"
        and event.payload.get("tool_name") == "insert_image_into_docx"
        and event.payload.get("native_chart_completion")
    ]
    check_finished = [event for event in events if event.type == "check.finished"][-1]
    run_finished = events[-1]

    assert len(native_inserts) == 1
    assert check_finished.payload["passed"] is True
    assert run_finished.payload["completed_task"] is True


def test_file_task_runtime_routes_financial_xlsx_chart_report_to_docx_via_native_mainline(
    tmp_path,
):
    import openpyxl
    from docx import Document

    workbook_path = tmp_path / "financial.xlsx"
    target_path = tmp_path / "target.docx"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "P&L"
    sheet.append(["", "利润表", "", "", ""])
    sheet.append(["", "", "2026E", "2027E", "2028E"])
    sheet.append(["", "收入合计", 1000, 2300, 3600])
    sheet.append(["", "毛利", 300, 700, 1200])
    sheet.append(["", "净利润", 80, 180, 260])
    workbook.save(workbook_path)

    document = Document()
    document.add_paragraph("雷鸟访谈问题")
    document.save(target_path)

    request = FileTaskRequest(
        task="分析这个xlsx财务数据，将数据做成图并找出存在的问题，然后将图和问题加入docx",
        run_id="financial_xlsx_docx_report",
        target_path=str(target_path),
        files=[
            FileTaskFile(
                path=str(workbook_path),
                name="雷鸟创新-financial model.xlsx",
                type="xlsx",
            ),
            FileTaskFile(
                path=str(target_path),
                name="雷鸟访谈问题.docx",
                type="docx",
                target=True,
            ),
        ],
    )

    runtime = FileTaskRuntime(
        model_client=lambda **kwargs: {"content": "按计划执行。", "tool_calls": []},
        workspace_root=str(tmp_path),
    )

    classification = runtime._classify_request(request, request.files)
    assert classification.selected_recipe == "financial_xlsx_docx_report"
    assert classification.task_family == "financial_report"
    assert classification.operation_kind == "analyze_visualize_write"
    assert classification.write_intent is True
    assert not hasattr(runtime, "_should_route_financial_xlsx_docx_report")


def test_file_task_runtime_financial_report_links_supplemental_sales_ledger(tmp_path):
    import openpyxl
    from docx import Document

    financial_path = tmp_path / "financial.xlsx"
    sales_path = tmp_path / "sales.xlsx"
    target_path = tmp_path / "target.docx"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "P&L"
    sheet.append(["", "", "2026E", "2027E", "2028E"])
    sheet.append(["", "收入合计", 1000, 2200, 3800])
    sheet.append(["", "毛利合计", 350, 860, 1500])
    sheet.append(["", "净利润", 80, 210, 420])
    workbook.save(financial_path)

    sales_workbook = openpyxl.Workbook()
    sales_sheet = sales_workbook.active
    sales_sheet.title = "汇总表"
    sales_sheet.append(["月份", "产品", "客户", "销售额", "销量"])
    sales_sheet.append(["2026-01", "AI Glasses", "A客户", 120, 3])
    sales_sheet.append(["2026-02", "AI Glasses", "B客户", 180, 4])
    sales_sheet.append(["2026-02", "XR", "A客户", 90, 2])
    sales_sheet.append(["2026-03", "AR", "C客户", 240, 5])
    sales_workbook.save(sales_path)

    document = Document()
    document.add_paragraph("雷鸟访谈问题")
    document.save(target_path)

    request = FileTaskRequest(
        task="将新的销售台账也加入分析，并且做成图，内容也加入docx",
        run_id="financial_sales_followup",
        target_path=str(target_path),
        options={
            "disable_financial_model_synthesis": True,
            "followup_context": {
                "previous_task_selected_recipe": "financial_xlsx_docx_report"
            },
        },
        files=[
            FileTaskFile(
                path=str(target_path),
                name="雷鸟访谈问题.docx",
                type="docx",
                target=True,
            ),
            FileTaskFile(
                path=str(financial_path),
                name="雷鸟创新-financial model.xlsx",
                type="xlsx",
            ),
            FileTaskFile(path=str(sales_path), name="销售台账.xlsx", type="xlsx"),
        ],
    )

    runtime = FileTaskRuntime(
        model_client=lambda **kwargs: {"content": "按计划执行。", "tool_calls": []},
        workspace_root=str(tmp_path),
    )

    classification = runtime._classify_request(request, request.files)
    assert classification.selected_recipe == "financial_xlsx_docx_report"
    assert not hasattr(runtime, "_should_route_financial_xlsx_docx_report")


def test_file_task_runtime_routes_financial_reports_through_whitebox_mainline():
    runtime_source = Path("app/core/agent/file_task_runtime.py").read_text(
        encoding="utf-8"
    )
    preflight_policy_source = Path(
        "app/core/agent/file_task_preflight_policy.py"
    ).read_text(encoding="utf-8")

    assert not Path("app/core/agent/file_task_financial_report_runner.py").exists()
    assert "FileTaskFinancialReportRunner" not in runtime_source
    assert "_stream_financial_xlsx_docx_report" not in runtime_source
    assert "_should_route_financial_xlsx_docx_report" not in runtime_source
    assert "build_preflight_constraint_audit" in runtime_source
    assert "financial_xlsx_docx_report" in preflight_policy_source


def test_file_task_runtime_keeps_financial_report_on_whitebox_mainline(tmp_path):
    import openpyxl
    from docx import Document

    workbook_path = tmp_path / "financial.xlsx"
    target_path = tmp_path / "target.docx"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "P&L"
    sheet.append(["", "", "2026E", "2027E"])
    sheet.append(["", "收入合计", 1000, 2000])
    workbook.save(workbook_path)

    document = Document()
    document.add_paragraph("雷鸟访谈问题")
    document.save(target_path)

    request = FileTaskRequest(
        task="分析这个xlsx财务数据，将数据做成图并找出存在的问题，然后将图和问题加入docx",
        run_id="financial_default_mainline",
        target_path=str(target_path),
        files=[
            FileTaskFile(
                path=str(workbook_path),
                name="雷鸟创新-financial model.xlsx",
                type="xlsx",
            ),
            FileTaskFile(
                path=str(target_path),
                name="雷鸟访谈问题.docx",
                type="docx",
                target=True,
            ),
        ],
    )
    runtime = FileTaskRuntime(
        model_client=lambda **kwargs: {"content": "按主任务链继续。", "tool_calls": []},
        workspace_root=str(tmp_path),
    )

    classification = runtime._classify_request(request, request.files)
    assert classification.selected_recipe == "financial_xlsx_docx_report"


def test_file_task_runtime_financial_report_requires_unambiguous_docx_target(tmp_path):
    import openpyxl
    from docx import Document

    workbook_path = tmp_path / "financial.xlsx"
    first_docx_path = tmp_path / "source_notes.docx"
    second_docx_path = tmp_path / "target_report.docx"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "P&L"
    sheet.append(["", "", "2026E", "2027E"])
    sheet.append(["", "收入合计", 1000, 2000])
    workbook.save(workbook_path)

    for path in (first_docx_path, second_docx_path):
        document = Document()
        document.add_paragraph(path.stem)
        document.save(path)

    request = FileTaskRequest(
        task="分析这个xlsx财务数据，将数据做成图并找出存在的问题，然后将图和问题加入docx",
        run_id="financial_ambiguous_docx_target",
        files=[
            FileTaskFile(path=str(workbook_path), name="financial.xlsx", type="xlsx"),
            FileTaskFile(
                path=str(first_docx_path), name="source_notes.docx", type="docx"
            ),
            FileTaskFile(
                path=str(second_docx_path), name="target_report.docx", type="docx"
            ),
        ],
    )

    runtime = FileTaskRuntime(
        model_client=lambda **kwargs: {"content": "", "tool_calls": []},
        workspace_root=str(tmp_path),
    )

    classification = runtime._classify_request(request, request.files)
    assert classification.selected_recipe == "financial_xlsx_docx_report"
    assert classification.task_family == "financial_report"


def test_file_task_runtime_prompt_guides_chart_into_docx_via_real_image_write():
    runtime = FileTaskRuntime(
        tool_executor=lambda name, args: "", model_client=lambda **kwargs: {}
    )
    request = FileTaskRequest(
        task="把 Excel 数据画成图加入 docx",
        run_id="chart_docx_prompt_demo",
        files=[
            FileTaskFile(path="财务模型.xlsx", name="财务模型.xlsx", type="xlsx"),
            FileTaskFile(path="报告.docx", name="报告.docx", type="docx", target=True),
        ],
        target_path="报告.docx",
    )

    prompt = runtime._build_system_prompt(request, request.files)

    assert "insert_image_into_docx" in prompt
    assert "不要把图片描述文字写进文档代替真实插图" in prompt
    assert "dpi>=220" in prompt
    assert "axes.unicode_minus=False" in prompt


def test_file_task_runtime_system_prompt_guides_financial_workbook_audit_first():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="分析这个财务预测模型，找到问题",
        files=[FileTaskFile(path="forecast.xlsx", name="forecast.xlsx", type="xlsx")],
    )
    prompt = runtime._build_system_prompt(request, request.files)

    assert "audit_financial_workbook" in prompt
    assert "inspect_workbook_structure" in prompt
    assert "结构性缺陷/可复算性问题" in prompt


def test_file_task_runtime_system_prompt_guides_financial_chart_docx_writeback():
    runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
    request = FileTaskRequest(
        task="将xlsx财务预测数据做成图，并分析存在的问题，将问题和图加入docx",
        target_path="report.docx",
        files=[
            FileTaskFile(path="forecast.xlsx", name="forecast.xlsx", type="xlsx"),
            FileTaskFile(
                path="report.docx", name="report.docx", type="docx", target=True
            ),
        ],
    )
    prompt = runtime._build_system_prompt(request, request.files)

    assert "Excel 财务预测图表写入 DOCX 任务规则" in prompt
    assert "列名是 Unnamed" in prompt
    assert "定位包含 2025E/2026E/2027E/2028E" in prompt
    assert "KOTO_CREATED" in prompt
    assert "write_docx_content" in prompt
    assert "insert_image_into_docx" in prompt
    assert "没有 file.changed 不能结束" in prompt
