import json
from pathlib import Path

from app.core.agent.file_task_contract import FileTaskFile, FileTaskRequest
from app.core.agent.file_task_runtime import FileTaskRuntime


def test_file_task_runtime_stops_retrying_when_write_target_is_locked():
    model_calls = {"count": 0}

    def fake_model(**kwargs):
        model_calls["count"] += 1
        return {
            "content": "先把 Excel 表格写入目标 Word。",
            "tool_calls": [
                {
                    "id": "insert_locked_demo",
                    "name": "insert_excel_as_docx_table",
                    "args": {
                        "source_path": "financial-model.xlsx",
                        "target_path": "report.docx",
                        "sheet_name": "P&L",
                    },
                }
            ],
        }

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            return f"内容来自 {args['path']}"
        if tool_name == "insert_excel_as_docx_table":
            return json.dumps(
                {
                    "success": False,
                    "path": args["target_path"],
                    "status": "write_blocked",
                    "summary": "目标文件 report.docx 当前不可写，无法写回原文件。",
                    "error": "目标文件 report.docx 当前不可写，无法写回原文件。",
                    "suggested_next_step": "检查文件权限；如果文件正在被占用，请关闭相关 Koto 页签或其他程序后重试。",
                },
                ensure_ascii=False,
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    request = FileTaskRequest(
        task="将 xlsx 财务预测加入 docx",
        run_id="write_locked_demo",
        target_path="report.docx",
        files=[
            FileTaskFile(
                path="financial-model.xlsx", name="financial-model.xlsx", type="xlsx"
            ),
            FileTaskFile(
                path="report.docx", name="report.docx", type="docx", target=True
            ),
        ],
    )

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor, model_client=fake_model, max_rounds=4
        ).run(request)
    )
    check_finished = next(event for event in events if event.type == "check.finished")
    run_finished = next(event for event in events if event.type == "run.finished")

    assert model_calls["count"] == 1
    assert not any(
        event.type == "tool.finished"
        and event.payload.get("tool_name") == "write_guard"
        for event in events
    )
    assert not any(
        event.type == "tool.finished"
        and event.payload.get("tool_name") == "repair_guard"
        for event in events
    )
    assert check_finished.payload["status"] == "write_blocked"
    assert "当前不可写" in check_finished.payload["summary"]
    assert check_finished.payload["remaining"] == [
        "检查文件权限；如果文件正在被占用，请关闭相关 Koto 页签或其他程序后重试。"
    ]
    assert run_finished.payload["completed_task"] is False
    assert "当前不可写" in run_finished.payload["summary"]


def test_file_task_runtime_keeps_recovery_copy_but_does_not_mark_original_write_complete():
    model_calls = {"count": 0}

    def fake_model(**kwargs):
        model_calls["count"] += 1
        return {
            "content": "先把 Excel 表格写入目标 Word。",
            "tool_calls": [
                {
                    "id": "insert_locked_demo",
                    "name": "insert_excel_as_docx_table",
                    "args": {
                        "source_path": "financial-model.xlsx",
                        "target_path": "report.docx",
                        "sheet_name": "P&L",
                    },
                }
            ],
        }

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            return f"内容来自 {args['path']}"
        if tool_name == "insert_excel_as_docx_table":
            return json.dumps(
                {
                    "success": False,
                    "path": "report.koto-copy.docx",
                    "file_type": "docx",
                    "change_type": "create",
                    "operation": "insert_excel_as_docx_table",
                    "summary": "原目标文件 report.docx 当前不可写，尚未写回原文件；已将工作表“P&L”的 50 行数据写入恢复副本 report.koto-copy.docx",
                    "preview": "收入合计",
                    "status": "write_blocked",
                    "error": "原目标文件 report.docx 当前不可写，尚未写回原文件；已将工作表“P&L”的 50 行数据写入恢复副本 report.koto-copy.docx",
                    "suggested_next_step": "检查 report.docx 的文件权限；如果文件正在被占用，请关闭相关 Koto 页签或其他程序后重新执行写回原文件。",
                    "original_target_path": "report.docx",
                    "blocked_target": True,
                    "blocked_reason": "目标文件 report.docx 当前不可写，无法写回原文件。",
                    "fallback_copy": True,
                    "sheet": "P&L",
                    "rows_written": 50,
                    "columns_written": 13,
                },
                ensure_ascii=False,
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    request = FileTaskRequest(
        task="将 xlsx 财务预测加入 docx",
        run_id="write_locked_recovery_copy_demo",
        target_path="report.docx",
        files=[
            FileTaskFile(
                path="financial-model.xlsx", name="financial-model.xlsx", type="xlsx"
            ),
            FileTaskFile(
                path="report.docx", name="report.docx", type="docx", target=True
            ),
        ],
    )

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor, model_client=fake_model, max_rounds=4
        ).run(request)
    )
    check_finished = next(event for event in events if event.type == "check.finished")
    run_finished = next(event for event in events if event.type == "run.finished")
    file_changed = next(event for event in events if event.type == "file.changed")
    tool_finished = next(
        event
        for event in events
        if event.type == "tool.finished"
        and event.payload.get("tool_name") == "insert_excel_as_docx_table"
    )

    assert model_calls["count"] == 1
    assert file_changed.payload["path"] == "report.koto-copy.docx"
    assert file_changed.payload["fallback_copy"] is True
    assert tool_finished.payload["blocked"] is True
    assert not any(
        event.type == "tool.finished"
        and event.payload.get("tool_name") == "write_guard"
        for event in events
    )
    assert not any(
        event.type == "tool.finished"
        and event.payload.get("tool_name") == "repair_guard"
        for event in events
    )
    assert check_finished.payload["status"] == "write_blocked"
    assert check_finished.payload["remaining"] == [
        "检查 report.docx 的文件权限；如果文件正在被占用，请关闭相关 Koto 页签或其他程序后重新执行写回原文件。"
    ]
    assert run_finished.payload["completed_task"] is False
    assert "尚未写回原文件" in run_finished.payload["summary"]
    assert "当前不可写" in run_finished.payload["summary"]


def test_file_task_runtime_prompts_write_after_generic_chart_tool_round_without_file_change():
    seen_messages = []
    responses = [
        {
            "content": "先检查并整理数据。",
            "tool_calls": [
                {
                    "name": "run_python_code",
                    "args": {"code": "print('only analysis output')"},
                },
            ],
        },
        {"content": "收到提醒，下一步写入。", "tool_calls": []},
    ]

    def fake_model(**kwargs):
        seen_messages.append(str(kwargs["messages"][-1]["content"]))
        return responses.pop(0) if responses else {"content": "done", "tool_calls": []}

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            return "显式上下文"
        if tool_name == "run_python_code":
            return {
                "stdout": "only analysis output",
                "stderr": "",
                "summary": "Python 执行完成",
            }
        return ""

    request = FileTaskRequest(
        task="Create a chart from sales data and write the chart into report.docx",
        run_id="generic_chart_write_guard_after_tools",
        target_path="report.docx",
        files=[
            FileTaskFile(
                path="report.docx", name="雷鸟访谈问题.docx", type="docx", target=True
            ),
        ],
    )

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor, model_client=fake_model, max_rounds=3
        ).run(request)
    )
    write_guard = next(
        event for event in events if event.payload.get("tool_name") == "write_guard"
    )

    assert "write_docx_content" in write_guard.payload["result_preview"]
    assert "insert_image_into_docx" in write_guard.payload["result_preview"]
    assert any(
        "write_docx_content" in message and "insert_image_into_docx" in message
        for message in seen_messages
    )


def test_file_task_runtime_xlsx_to_docx_write_loop_handles_sheet1_and_string_rows(
    tmp_path,
):
    import openpyxl
    from docx import Document

    workbook_path = tmp_path / "sales.xlsx"
    target_path = tmp_path / "target.docx"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "汇总表"
    sheet.append(["客户名称", "产品名称", "数量"])
    sheet.append(["杭州新汇鑫光电有限公司", "LASER", 1])
    sheet.append(["山东镭鸟激光设备有限公司", "LASER", 2])
    workbook.save(workbook_path)

    document = Document()
    document.add_paragraph("雷鸟访谈问题")
    document.save(target_path)

    responses = iter(
        [
            {
                "content": "先读取 Excel。",
                "tool_calls": [
                    {
                        "name": "read_sheet_data",
                        "args": {
                            "path": str(workbook_path),
                            "sheet_name": "Sheet1",
                            "max_rows": "2",
                        },
                    }
                ],
            },
            {
                "content": "把 Excel 表格加入 Word。",
                "tool_calls": [
                    {
                        "name": "insert_excel_as_docx_table",
                        "args": {
                            "source_path": str(workbook_path),
                            "target_path": str(target_path),
                            "sheet_name": "Sheet1",
                            "table_title": "销售台账数据",
                            "max_rows": "2",
                        },
                    }
                ],
            },
            {"content": "已将销售台账数据加入 Word。", "tool_calls": []},
        ]
    )

    def fake_model(**kwargs):
        return next(responses, {"content": "", "tool_calls": []})

    request = FileTaskRequest(
        task="把 xlsx 表格加入 docx",
        run_id="xlsx_docx_loop",
        model_mode="local",
        target_path=str(target_path),
        files=[
            FileTaskFile(
                path=str(workbook_path),
                name="销售台账.xlsx",
                type="xlsx",
                content="销售台账 Excel",
            ),
            FileTaskFile(
                path=str(target_path),
                name="雷鸟访谈问题.docx",
                type="docx",
                content="目标 Word 文档",
                target=True,
            ),
        ],
    )

    events = list(FileTaskRuntime(model_client=fake_model).run(request))
    read_finished = next(
        event
        for event in events
        if event.type == "tool.finished"
        and event.payload.get("tool_name") == "read_sheet_data"
    )
    insert_finished = next(
        event
        for event in events
        if event.type == "tool.finished"
        and event.payload.get("tool_name") == "insert_excel_as_docx_table"
    )
    file_changed = next(event for event in events if event.type == "file.changed")
    check_finished = next(event for event in events if event.type == "check.finished")
    run_finished = events[-1]

    assert read_finished.payload["success"] is True
    assert "汇总表" in read_finished.payload["result_preview"]
    assert "Sheet1" in read_finished.payload["result_preview"]
    assert insert_finished.payload["success"] is True
    assert "汇总表" in insert_finished.payload["result_preview"]
    assert Path(file_changed.payload["path"]).resolve() == target_path.resolve()
    assert file_changed.payload["sheet"] == "汇总表"
    assert file_changed.payload["requested_sheet"] == "Sheet1"
    assert file_changed.payload["rows_written"] == 2
    assert file_changed.payload["columns_written"] == 3
    assert check_finished.payload["passed"] is True
    assert check_finished.payload["status"] == "verified"
    assert run_finished.payload["completed_task"] is True

    saved = Document(str(target_path))
    assert len(saved.tables) == 1
    assert saved.tables[0].cell(1, 0).text == "杭州新汇鑫光电有限公司"
    assert saved.tables[0].cell(2, 0).text == "山东镭鸟激光设备有限公司"


def test_file_task_runtime_xlsx_to_docx_write_loop_fails_without_file_change(tmp_path):
    from docx import Document

    target_path = tmp_path / "target.docx"
    document = Document()
    document.add_paragraph("雷鸟访谈问题")
    document.save(target_path)

    responses = iter(
        [
            {
                "content": "尝试把 Excel 表格加入 Word。",
                "tool_calls": [
                    {
                        "name": "insert_excel_as_docx_table",
                        "args": {
                            "source_path": str(tmp_path / "missing.xlsx"),
                            "target_path": str(target_path),
                            "sheet_name": "Sheet1",
                            "max_rows": "50",
                        },
                    }
                ],
            },
            {"content": "已完成。", "tool_calls": []},
        ]
    )

    def fake_model(**kwargs):
        return next(responses, {"content": "", "tool_calls": []})

    request = FileTaskRequest(
        task="把 xlsx 表格加入 docx",
        run_id="xlsx_docx_no_change",
        target_path=str(target_path),
        files=[
            FileTaskFile(
                path=str(target_path),
                name="雷鸟访谈问题.docx",
                type="docx",
                content="目标 Word 文档",
                target=True,
            ),
            FileTaskFile(
                path=str(tmp_path / "missing.xlsx"),
                name="missing.xlsx",
                type="xlsx",
            ),
        ],
    )

    events = list(FileTaskRuntime(model_client=fake_model).run(request))
    insert_finished = next(
        event
        for event in events
        if event.type == "tool.finished"
        and event.payload.get("tool_name") == "insert_excel_as_docx_table"
    )
    check_events = [event for event in events if event.type == "check.finished"]
    run_finished = events[-1]

    assert insert_finished.payload["success"] is False
    assert "File not found" in insert_finished.payload["result_preview"]
    assert check_events[0].payload["passed"] is False
    assert check_events[0].payload["status"] in {
        "write_not_performed",
        "quality_gate_failed",
    }
    assert run_finished.payload["completed_task"] is False
    assert len(Document(str(target_path)).tables) == 0


def test_file_task_runtime_retries_write_task_after_read_only_model_answer(tmp_path):
    import openpyxl
    from docx import Document

    workbook_path = tmp_path / "sales.xlsx"
    target_path = tmp_path / "target.docx"

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "汇总表"
    sheet.append(["客户名称", "产品名称", "数量"])
    sheet.append(["杭州新汇鑫光电有限公司", "LASER", 1])
    sheet.append(["山东镭鸟激光设备有限公司", "LASER", 2])
    workbook.save(workbook_path)

    document = Document()
    document.add_paragraph("雷鸟访谈问题")
    document.save(target_path)

    responses = iter(
        [
            {
                "content": "先读取 Excel。",
                "tool_calls": [
                    {
                        "name": "read_sheet_data",
                        "args": {
                            "path": str(workbook_path),
                            "sheet_name": "Sheet1",
                            "max_rows": "2",
                        },
                    }
                ],
            },
            {"content": "我已经读取完表格内容。", "tool_calls": []},
            {
                "content": "现在写入 Word。",
                "tool_calls": [
                    {
                        "name": "insert_excel_as_docx_table",
                        "args": {
                            "source_path": str(workbook_path),
                            "target_path": str(target_path),
                            "sheet_name": "汇总表",
                            "table_title": "销售台账数据",
                            "max_rows": "2",
                        },
                    }
                ],
            },
            {"content": "已将销售台账数据加入 Word。", "tool_calls": []},
        ]
    )
    seen_last_messages = []

    def fake_model(**kwargs):
        seen_last_messages.append(kwargs["messages"][-1]["content"])
        return next(responses, {"content": "", "tool_calls": []})

    request = FileTaskRequest(
        task="把 xlsx 表格加入 docx",
        run_id="xlsx_docx_retry_after_read_only",
        model_mode="local",
        target_path=str(target_path),
        files=[
            FileTaskFile(
                path=str(workbook_path),
                name="销售台账.xlsx",
                type="xlsx",
                content="销售台账 Excel",
            ),
            FileTaskFile(
                path=str(target_path),
                name="雷鸟访谈问题.docx",
                type="docx",
                content="目标 Word 文档",
                target=True,
            ),
        ],
    )

    events = list(FileTaskRuntime(model_client=fake_model, max_rounds=4).run(request))
    write_guard = next(
        event
        for event in events
        if event.type == "tool.finished"
        and event.payload.get("tool_name") == "write_guard"
    )
    insert_finished = next(
        event
        for event in events
        if event.type == "tool.finished"
        and event.payload.get("tool_name") == "insert_excel_as_docx_table"
    )
    file_changed = next(event for event in events if event.type == "file.changed")
    check_finished = next(event for event in events if event.type == "check.finished")
    run_finished = events[-1]

    assert "insert_excel_as_docx_table" in write_guard.payload["result_preview"]
    assert any(
        "insert_excel_as_docx_table" in message for message in seen_last_messages
    )
    assert insert_finished.payload["success"] is True
    assert file_changed.payload["sheet"] == "汇总表"
    assert check_finished.payload["status"] == "verified"
    assert run_finished.payload["completed_task"] is True
    assert len(Document(str(target_path)).tables) == 1


def test_file_task_runtime_resets_repair_budget_after_real_file_change(tmp_path):
    workbook_path = tmp_path / "sales.xlsx"
    target_path = tmp_path / "target.docx"
    workbook_path.write_text("xlsx placeholder", encoding="utf-8")
    target_path.write_text("docx placeholder", encoding="utf-8")

    responses = iter(
        [
            {"content": "我先看一下。", "tool_calls": []},
            {
                "content": "先插入表格。",
                "tool_calls": [
                    {
                        "name": "insert_excel_as_docx_table",
                        "args": {
                            "source_path": str(workbook_path),
                            "target_path": str(target_path),
                            "sheet_name": "Sales",
                            "table_title": "Q2 月度销售数据",
                        },
                    }
                ],
            },
            {"content": "表格已经插入。", "tool_calls": []},
            {
                "content": "补充表格说明。",
                "tool_calls": [
                    {
                        "name": "write_docx_content",
                        "args": {
                            "path": str(target_path),
                            "paragraphs": '[{"text":"说明：下表展示 Q2 月度销售核心数据。"}]',
                        },
                    }
                ],
            },
            {"content": "已补充说明并保留表格。", "tool_calls": []},
        ]
    )

    def fake_model(**kwargs):
        return next(responses, {"content": "", "tool_calls": []})

    def fake_executor(tool_name, args):
        if tool_name == "parse_file_to_text":
            return "已读取上下文"
        if tool_name == "insert_excel_as_docx_table":
            return json.dumps(
                {
                    "path": args["target_path"],
                    "operation": "insert_excel_as_docx_table",
                    "summary": "已将工作表“Sales”的 3 行数据写入 Word 表格",
                    "file_type": "docx",
                    "change_type": "modify",
                    "rows_written": 3,
                    "columns_written": 5,
                },
                ensure_ascii=False,
            )
        if tool_name == "write_docx_content":
            return json.dumps(
                {
                    "path": args["path"],
                    "operation": "write_docx_content",
                    "summary": "已写入 1 个段落到 Word 文档",
                    "file_type": "docx",
                    "change_type": "modify",
                    "paragraphs_written": 1,
                },
                ensure_ascii=False,
            )
        if tool_name == "verify_task_completion":
            changes = json.loads(args.get("file_changes") or "[]")
            operations = {str(change.get("operation") or "") for change in changes}
            if "write_docx_content" not in operations:
                return json.dumps(
                    {
                        "completed": False,
                        "confidence": 0.45,
                        "summary": "目标 DOCX 已插入表格，但还缺少表格前说明。",
                        "remaining_steps": ["用 write_docx_content 补充说明"],
                        "criteria_results": [
                            {
                                "criterion": "docx_narrative_write_present",
                                "passed": False,
                                "priority": "critical",
                            }
                        ],
                    },
                    ensure_ascii=False,
                )
            return json.dumps(
                {
                    "completed": True,
                    "confidence": 0.95,
                    "summary": "表格和说明均已写入。",
                },
                ensure_ascii=False,
            )
        raise AssertionError(f"unexpected tool call: {tool_name}")

    request = FileTaskRequest(
        task="读取 Excel 主要销售数据，把前 5 行以真实表格追加到目标 Word 文档末尾，并在表格前写一句简短说明。",
        run_id="xlsx_table_then_narrative_repair",
        target_path=str(target_path),
        files=[
            FileTaskFile(path=str(workbook_path), name="sales.xlsx", type="xlsx"),
            FileTaskFile(
                path=str(target_path), name="target.docx", type="docx", target=True
            ),
        ],
    )

    events = list(
        FileTaskRuntime(
            tool_executor=fake_executor, model_client=fake_model, max_rounds=5
        ).run(request)
    )
    tool_names = [
        event.payload.get("tool_name")
        for event in events
        if event.type == "tool.finished"
    ]
    operations = [
        event.payload.get("operation")
        for event in events
        if event.type == "file.changed"
    ]
    check_finished = [event for event in events if event.type == "check.finished"]

    assert "write_guard" in tool_names
    assert "repair_guard" in tool_names
    assert operations == ["insert_excel_as_docx_table", "write_docx_content"]
    assert any(
        event.payload.get("status") == "quality_gate_failed" for event in check_finished
    )
    assert check_finished[-1].payload["status"] == "verified"
    assert events[-1].payload["completed_task"] is True
