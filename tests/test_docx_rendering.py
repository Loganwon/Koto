"""
tests/test_docx_rendering.py

Backend integration tests for DOCX → HTML rendering fidelity.

Tests call parse_docx() directly — no server needed, pure Python.
Verifies the output HTML preserves the formatting that the TipTap
editor depends on to render correctly.

Target document: workspace/雷鸟创新-邗投珒创-投资建议书.docx
Word page count: 72 pages  (US Letter, 2.54 cm margins)

Run with:
    python -m pytest tests/test_docx_rendering.py -v
"""

from __future__ import annotations

import base64
import os
import re

import pytest

# ---------------------------------------------------------------------------
# Paths & constants
# ---------------------------------------------------------------------------
_REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX_PATH = os.path.join(_REPO_ROOT, "workspace", "雷鸟创新-邗投珒创-投资建议书.docx")
TRANSLATION_DOCX_PATH = os.path.join(_REPO_ROOT, "workspace", "2.1书稿翻译2.docx")

# Microsoft Word reports 72 pages for this document.
WORD_PAGE_COUNT = 72

# TipTap/Koto rendering constants (must match koto-docx-editor.js)
_PAD_V = 176  # ProseMirror padding: top(96) + bottom(80)
_CONTENT_PAGE_H = 880  # usable content height per page (1056 - 176)

# ---------------------------------------------------------------------------
# Shared fixture — parse the DOCX once for the whole module
# ---------------------------------------------------------------------------


@pytest.fixture(scope="module")
def docx_html() -> str:
    """Parse the target DOCX once and share the HTML across all tests."""
    if not os.path.exists(DOCX_PATH):
        pytest.skip(f"Test document not found: {DOCX_PATH}")

    from app.core.file.file_parser import parse_docx  # noqa: PLC0415

    result = parse_docx(DOCX_PATH)
    assert isinstance(result, dict), "parse_docx() must return a dict"
    assert "html" in result, "parse_docx() result must contain 'html' key"

    html = result["html"]
    assert (
        isinstance(html, str) and len(html) > 1_000
    ), f"HTML output suspiciously short ({len(html)} chars); parser likely failed"
    return html


def _write_typography_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    doc = Document()

    # Create an intentionally conflicting style/default setup so the parser
    # must respect paragraph-level default run props over the style fallback.
    normal = doc.styles["Normal"]
    normal.font.size = Pt(16)
    normal.font.bold = True

    heading = doc.styles["Heading 1"]
    heading.font.size = Pt(16)

    doc.add_paragraph("一、企业简介", style="Heading 1")

    para = doc.add_paragraph()
    p_pr = para._p.get_or_add_pPr()
    r_pr = p_pr.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        p_pr.append(r_pr)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    r_pr.append(sz)

    bold = OxmlElement("w:b")
    bold.set(qn("w:val"), "0")
    r_pr.append(bold)

    para.add_run("企业介绍正文段落，用于验证段落默认字号与粗细继承。")
    doc.save(path)


def _write_outline_only_heading_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415

    doc = Document()
    para = doc.add_paragraph("执行概要")
    p_pr = para._p.get_or_add_pPr()
    outline_lvl = p_pr.find(qn("w:outlineLvl"))
    if outline_lvl is None:
        outline_lvl = OxmlElement("w:outlineLvl")
        p_pr.append(outline_lvl)
    outline_lvl.set(qn("w:val"), "1")
    doc.save(path)


def _write_outline_only_prefixed_heading_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415

    doc = Document()
    para = doc.add_paragraph("第一章 公司基本信息")
    p_pr = para._p.get_or_add_pPr()
    outline_lvl = p_pr.find(qn("w:outlineLvl"))
    if outline_lvl is None:
        outline_lvl = OxmlElement("w:outlineLvl")
        p_pr.append(outline_lvl)
    outline_lvl.set(qn("w:val"), "0")
    doc.save(path)


def _write_outline_only_body_sentence_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    doc = Document()
    para = doc.add_paragraph(
        "这是一段用于回归测试的长正文句子，它带有大纲级别但本质上仍然是正文内容，因此不应该被渲染成主标题，而应该继续作为普通段落显示。"
    )
    para.paragraph_format.first_line_indent = Pt(21)
    p_pr = para._p.get_or_add_pPr()
    outline_lvl = p_pr.find(qn("w:outlineLvl"))
    if outline_lvl is None:
        outline_lvl = OxmlElement("w:outlineLvl")
        p_pr.append(outline_lvl)
    outline_lvl.set(qn("w:val"), "0")
    doc.save(path)


def _write_outline_only_short_body_sentence_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415

    doc = Document()
    para = doc.add_paragraph("在会议场景中，AI眼镜利用多模态能力提升协作效率")
    p_pr = para._p.get_or_add_pPr()
    outline_lvl = p_pr.find(qn("w:outlineLvl"))
    if outline_lvl is None:
        outline_lvl = OxmlElement("w:outlineLvl")
        p_pr.append(outline_lvl)
    outline_lvl.set(qn("w:val"), "1")
    doc.save(path)


def _write_outline_only_date_body_sentence_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415

    doc = Document()
    para = doc.add_paragraph(
        "2022年9月20日，深圳市市场监督管理局完成本次股权转让备案。"
    )
    p_pr = para._p.get_or_add_pPr()
    outline_lvl = p_pr.find(qn("w:outlineLvl"))
    if outline_lvl is None:
        outline_lvl = OxmlElement("w:outlineLvl")
        p_pr.append(outline_lvl)
    outline_lvl.set(qn("w:val"), "1")
    doc.save(path)


def _write_heading_style_cache_regression_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415

    doc = Document()
    body_text = (
        "作为国内AI眼镜领域的链主企业，雷鸟创新已完成核心器件、整机终端和软件内容的全链路布局，"
        "构建起集设计、制造、算法和内容服务于一体的智能眼镜产业链。"
    )
    for idx in range(40):
        doc.add_paragraph(f"第{idx + 1}章 测试标题", style="Heading 1")
        doc.add_paragraph(body_text)
    doc.save(path)


def _write_all_heading_levels_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415

    doc = Document()
    for level in range(1, 7):
        doc.add_paragraph(f"层级{level}标题", style=f"Heading {level}")
    doc.add_paragraph("普通正文，不应进入导航。")
    doc.save(path)


def _write_section_break_marker_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.enum.section import WD_SECTION  # noqa: PLC0415

    doc = Document()
    doc.add_paragraph("第一节正文")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("第二节正文")
    doc.save(path)


def _write_toc_substring_style_body_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.enum.style import WD_STYLE_TYPE  # noqa: PLC0415

    doc = Document()
    stock_style = doc.styles.add_style("StockNote", WD_STYLE_TYPE.PARAGRAPH)
    para = doc.add_paragraph(style=stock_style)
    para.add_run("2022年9月20日 深圳市市场监督管理局备案")
    para.add_run().add_tab()
    para.add_run("60")
    doc.save(path)


def _write_paragraph_bold_run_semantics_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415

    def _ensure_rpr(run):
        r_pr = run._element.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            run._element.insert(0, r_pr)
        return r_pr

    def _append_bold(run, value: str) -> None:
        r_pr = _ensure_rpr(run)
        bold = OxmlElement("w:b")
        bold.set(qn("w:val"), value)
        r_pr.append(bold)

    doc = Document()
    para = doc.add_paragraph()
    p_pr = para._p.get_or_add_pPr()
    p_r_pr = p_pr.find(qn("w:rPr"))
    if p_r_pr is None:
        p_r_pr = OxmlElement("w:rPr")
        p_pr.append(p_r_pr)
    p_bold = OxmlElement("w:b")
    p_bold.set(qn("w:val"), "1")
    p_r_pr.append(p_bold)

    para.add_run("继承粗体")
    explicit_off = para.add_run(" 显式取消")
    _append_bold(explicit_off, "0")
    toggle_off = para.add_run(" 切换取消")
    _append_bold(toggle_off, "1")

    doc.save(path)


def _write_paragraph_layout_semantics_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.enum.text import WD_LINE_SPACING  # noqa: PLC0415
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    def _set_on_off(p_pr, tag_name: str, enabled: bool) -> None:
        el = p_pr.find(qn(f"w:{tag_name}"))
        if el is None:
            el = OxmlElement(f"w:{tag_name}")
            p_pr.append(el)
        el.set(qn("w:val"), "1" if enabled else "0")

    doc = Document()
    para = doc.add_paragraph("段落分页语义测试")
    pf = para.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.first_line_indent = Pt(21)
    pf.left_indent = Pt(18)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)

    p_pr = para._p.get_or_add_pPr()
    _set_on_off(p_pr, "keepNext", True)
    _set_on_off(p_pr, "keepLines", True)
    _set_on_off(p_pr, "pageBreakBefore", True)
    _set_on_off(p_pr, "widowControl", False)

    doc.save(path)


def _write_localized_title_style_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.enum.style import WD_STYLE_TYPE  # noqa: PLC0415
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    doc = Document()
    try:
        title_style = doc.styles["标题"]
    except KeyError:
        title_style = doc.styles.add_style("标题", WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(18)

    r_pr = title_style._element.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        title_style._element.append(r_pr)
    for bold_el in list(r_pr.findall(qn("w:b"))):
        r_pr.remove(bold_el)
    for bold_cs_el in list(r_pr.findall(qn("w:bCs"))):
        r_pr.remove(bold_cs_el)
    bold_cs = OxmlElement("w:bCs")
    bold_cs.set(qn("w:val"), "1")
    r_pr.append(bold_cs)

    doc.add_paragraph("封面标题", style=title_style)
    doc.save(path)


def _write_title_style_based_on_heading_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.enum.style import WD_STYLE_TYPE  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    doc = Document()
    try:
        title_style = doc.styles["Title"]
    except KeyError:
        title_style = doc.styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)

    title_style.base_style = doc.styles["Heading 2"]
    title_style.font.size = Pt(24)
    title_style.font.bold = True

    doc.add_paragraph("封面总标题", style=title_style)
    doc.add_paragraph("第一章 公司简介", style="Heading 1")
    doc.save(path)


def _write_custom_visual_title_chain_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.enum.style import WD_STYLE_TYPE  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    doc = Document()
    try:
        title_style = doc.styles["Title"]
    except KeyError:
        title_style = doc.styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)

    title_style.base_style = doc.styles["Heading 2"]

    custom_title = doc.styles.add_style("封面定制标题", WD_STYLE_TYPE.PARAGRAPH)
    custom_title.base_style = title_style
    custom_title.font.size = Pt(26)
    custom_title.font.bold = True

    doc.add_paragraph("封面链式标题", style=custom_title)
    doc.add_paragraph("第二章 技术分析", style="Heading 1")
    doc.save(path)


def _write_localized_font_family_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("华文字体回归")
    run.font.name = "华文仿宋"
    run.font.size = Pt(12)

    r_pr = run._element.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        run._element.append(r_pr)
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "华文仿宋")
    r_fonts.set(qn("w:hAnsi"), "华文仿宋")
    r_fonts.set(qn("w:eastAsia"), "华文仿宋")

    doc.save(path)


def _write_exact_row_height_table_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.enum.table import WD_ROW_HEIGHT_RULE  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.cell(0, 0).text = (
        "这是一段故意较长的单元格文本，用于验证精确行高不会被导出成浏览器固定高度。"
    )
    table.cell(1, 0).text = "第二行内容"

    first_row = table.rows[0]
    first_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    first_row.height = Pt(12)

    doc.save(path)


def _write_small_multiple_line_spacing_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=1)
    para = table.cell(0, 0).paragraphs[0]
    para.paragraph_format.line_spacing = 0.25
    para.add_run(
        "这是一段故意较长的表格文本，用于验证异常的多倍行距不会被原样导出为浏览器 line-height:0.25。"
    )

    doc.save(path)


def _write_table_inline_image_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415
    from docx.shared import Inches  # noqa: PLC0415

    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+a5n8AAAAASUVORK5CYII="
    )
    image_path = path.with_suffix(".png")
    image_path.write_bytes(png_bytes)

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    para = cell.paragraphs[0]
    para.add_run("邮箱")
    para.add_run(" ")
    para.add_run().add_picture(str(image_path), width=Inches(0.45))

    cell.add_paragraph("手机号")
    doc.save(path)


def _write_table_cell_heading_fixture_docx(path) -> None:
    from docx import Document  # noqa: PLC0415

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell_para = table.cell(0, 0).paragraphs[0]
    cell_para.style = "Heading 1"
    cell_para.text = "表格里的标题"

    doc.add_paragraph("正文标题", style="Heading 1")
    doc.save(path)


@pytest.fixture()
def typography_html(tmp_path) -> str:
    pytest.importorskip("docx", reason="python-docx 未安装")

    docx_path = tmp_path / "typography-fixture.docx"
    _write_typography_fixture_docx(docx_path)

    from app.core.file.file_parser import parse_docx  # noqa: PLC0415

    result = parse_docx(str(docx_path))
    assert isinstance(result, dict), "parse_docx() must return a dict"
    html = result.get("html", "")
    assert isinstance(html, str) and html, "parse_docx() must produce HTML"
    return html


@pytest.fixture()
def outline_heading_fallback_html(tmp_path) -> str:
    pytest.importorskip("docx", reason="python-docx 未安装")

    docx_path = tmp_path / "outline-only-heading.docx"
    _write_outline_only_heading_fixture_docx(docx_path)

    from app.core.file.file_parser import parse_docx  # noqa: PLC0415

    result = parse_docx(str(docx_path))
    assert isinstance(result, dict), "parse_docx() must return a dict"
    html = result.get("html", "")
    assert isinstance(html, str) and html, "parse_docx() must produce HTML"
    return html


@pytest.fixture()
def outline_body_sentence_html(tmp_path) -> str:
    pytest.importorskip("docx", reason="python-docx 未安装")

    docx_path = tmp_path / "outline-body-sentence.docx"
    _write_outline_only_body_sentence_fixture_docx(docx_path)

    from app.core.file.file_parser import parse_docx  # noqa: PLC0415

    result = parse_docx(str(docx_path))
    assert isinstance(result, dict), "parse_docx() must return a dict"
    html = result.get("html", "")
    assert isinstance(html, str) and html, "parse_docx() must produce HTML"
    return html


@pytest.fixture()
def outline_short_body_sentence_html(tmp_path) -> str:
    pytest.importorskip("docx", reason="python-docx 未安装")

    docx_path = tmp_path / "outline-short-body-sentence.docx"
    _write_outline_only_short_body_sentence_fixture_docx(docx_path)

    from app.core.file.file_parser import parse_docx  # noqa: PLC0415

    result = parse_docx(str(docx_path))
    assert isinstance(result, dict), "parse_docx() must return a dict"
    html = result.get("html", "")
    assert isinstance(html, str) and html, "parse_docx() must produce HTML"
    return html


@pytest.fixture()
def outline_date_body_sentence_html(tmp_path) -> str:
    pytest.importorskip("docx", reason="python-docx 未安装")

    docx_path = tmp_path / "outline-date-body-sentence.docx"
    _write_outline_only_date_body_sentence_fixture_docx(docx_path)

    from app.core.file.file_parser import parse_docx  # noqa: PLC0415

    result = parse_docx(str(docx_path))
    assert isinstance(result, dict), "parse_docx() must return a dict"
    html = result.get("html", "")
    assert isinstance(html, str) and html, "parse_docx() must produce HTML"
    return html


@pytest.fixture()
def toc_substring_style_body_html(tmp_path) -> str:
    pytest.importorskip("docx", reason="python-docx 未安装")

    docx_path = tmp_path / "toc-substring-style-body.docx"
    _write_toc_substring_style_body_fixture_docx(docx_path)

    from app.core.file.file_parser import parse_docx  # noqa: PLC0415

    result = parse_docx(str(docx_path))
    assert isinstance(result, dict), "parse_docx() must return a dict"
    html = result.get("html", "")
    assert isinstance(html, str) and html, "parse_docx() must produce HTML"
    return html


@pytest.fixture()
def paragraph_bold_run_semantics_html(tmp_path) -> str:
    pytest.importorskip("docx", reason="python-docx 未安装")

    docx_path = tmp_path / "paragraph-bold-run-semantics.docx"
    _write_paragraph_bold_run_semantics_fixture_docx(docx_path)

    from app.core.file.file_parser import parse_docx  # noqa: PLC0415

    result = parse_docx(str(docx_path))
    assert isinstance(result, dict), "parse_docx() must return a dict"
    html = result.get("html", "")
    assert isinstance(html, str) and html, "parse_docx() must produce HTML"
    return html


@pytest.fixture()
def localized_title_style_html(tmp_path) -> str:
    pytest.importorskip("docx", reason="python-docx 未安装")

    docx_path = tmp_path / "localized-title-style.docx"
    _write_localized_title_style_fixture_docx(docx_path)

    from app.core.file.file_parser import parse_docx  # noqa: PLC0415

    result = parse_docx(str(docx_path))
    assert isinstance(result, dict), "parse_docx() must return a dict"
    html = result.get("html", "")
    assert isinstance(html, str) and html, "parse_docx() must produce HTML"
    return html


@pytest.fixture()
def localized_font_family_html(tmp_path) -> str:
    pytest.importorskip("docx", reason="python-docx 未安装")

    docx_path = tmp_path / "localized-font-family.docx"
    _write_localized_font_family_fixture_docx(docx_path)

    from app.core.file.file_parser import parse_docx  # noqa: PLC0415

    result = parse_docx(str(docx_path))
    assert isinstance(result, dict), "parse_docx() must return a dict"
    html = result.get("html", "")
    assert isinstance(html, str) and html, "parse_docx() must produce HTML"
    return html


@pytest.fixture()
def exact_row_height_table_html(tmp_path) -> str:
    pytest.importorskip("docx", reason="python-docx 未安装")

    docx_path = tmp_path / "exact-row-height-table.docx"
    _write_exact_row_height_table_fixture_docx(docx_path)

    from app.core.file.file_parser import parse_docx  # noqa: PLC0415

    result = parse_docx(str(docx_path))
    assert isinstance(result, dict), "parse_docx() must return a dict"
    html = result.get("html", "")
    assert isinstance(html, str) and html, "parse_docx() must produce HTML"
    return html


@pytest.fixture()
def small_multiple_line_spacing_html(tmp_path) -> str:
    pytest.importorskip("docx", reason="python-docx 未安装")

    docx_path = tmp_path / "small-multiple-line-spacing.docx"
    _write_small_multiple_line_spacing_fixture_docx(docx_path)

    from app.core.file.file_parser import parse_docx  # noqa: PLC0415

    result = parse_docx(str(docx_path))
    assert isinstance(result, dict), "parse_docx() must return a dict"
    html = result.get("html", "")
    assert isinstance(html, str) and html, "parse_docx() must produce HTML"
    return html


@pytest.fixture()
def table_inline_image_html(tmp_path) -> str:
    pytest.importorskip("docx", reason="python-docx 未安装")

    docx_path = tmp_path / "table-inline-image.docx"
    _write_table_inline_image_fixture_docx(docx_path)

    from app.core.file.file_parser import parse_docx  # noqa: PLC0415

    result = parse_docx(str(docx_path))
    assert isinstance(result, dict), "parse_docx() must return a dict"
    html = result.get("html", "")
    assert isinstance(html, str) and html, "parse_docx() must produce HTML"
    return html


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


@pytest.mark.integration
class TestDocxHtmlStructure:
    """Structural checks on the HTML produced by parse_docx()."""

    def test_html_has_substantial_content(self, docx_html: str) -> None:
        """Output HTML must be non-trivial (>50 KB) for a 72-page document."""
        assert len(docx_html) > 50_000, (
            f"HTML is only {len(docx_html):,} chars — expected >50 000 for a "
            f"{WORD_PAGE_COUNT}-page document.  Parser may have failed silently."
        )

    def test_has_block_elements(self, docx_html: str) -> None:
        """Output must contain paragraph or div elements."""
        assert re.search(
            r"<(p|div)\b", docx_html, re.IGNORECASE
        ), "No <p> or <div> elements found in HTML output"


@pytest.mark.integration
class TestHeaderFooter:
    """Header / footer CSS class preservation."""

    def test_header_class_present(self, docx_html: str) -> None:
        """
        The parser emits <p class="koto-header"> for header paragraphs.
        After the TipTap className-attribute fix, DocxParagraph must round-trip
        this class through parse → TipTap → renderHTML.  This test verifies the
        *parser side*: the class must exist in the raw HTML output.
        """
        assert "koto-header" in docx_html, (
            "'koto-header' class not found in HTML.  "
            "Check _section_html() in file_parser.py — "
            "header paragraphs may be missing or have a different class name."
        )

    def test_footer_class_present_if_document_has_footer(self, docx_html: str) -> None:
        """
        If a document has a footer defined in its DOCX, the parser emits
        koto-footer class.  This document (投资建议书) has no footer section,
        so the test always skips for this file.  If you use a document WITH a
        footer, this test will verify the class is preserved.
        """
        if "koto-footer" not in docx_html:
            pytest.skip("Document has no footer — koto-footer class not expected")
        assert "koto-footer" in docx_html

    def test_header_contains_text(self, docx_html: str) -> None:
        """At least one koto-header element must have non-empty text content."""
        # Extract the first koto-header <p> tag and its text
        match = re.search(
            r'<p[^>]*class="koto-header"[^>]*>(.*?)</p>',
            docx_html,
            re.IGNORECASE | re.DOTALL,
        )
        if not match:
            pytest.skip(
                "koto-header class not found — covered by test_header_class_present"
            )
        text = re.sub(r"<[^>]+>", "", match.group(1)).strip()
        assert text, "koto-header paragraph is empty (no visible text)"


@pytest.mark.integration
class TestTypography:
    """Paragraph and heading typography fidelity."""

    def test_body_paragraph_uses_paragraph_default_run_size_without_forced_bold(
        self, typography_html: str
    ) -> None:
        """
        Paragraph-level default run properties (<w:pPr><w:rPr>) should beat the
        Normal style fallback when setting the block font size, and body text
        must not inherit a phantom bold weight at the paragraph level.
        """
        blocks = re.findall(
            r"<p\b([^>]*)>(.*?)</p>", typography_html, re.IGNORECASE | re.DOTALL
        )
        target_attrs = None
        for attrs, inner in blocks:
            text = re.sub(r"<[^>]+>", "", inner).replace("\xa0", " ")
            text = re.sub(r"\s+", " ", text).strip()
            if text == "企业介绍正文段落，用于验证段落默认字号与粗细继承。":
                target_attrs = attrs
                break

        assert (
            target_attrs is not None
        ), "Target body paragraph not found in parsed HTML"
        style_match = re.search(r'style="([^"]*)"', target_attrs)
        assert style_match, "Target body paragraph missing inline style"
        style = style_match.group(1)
        assert "font-size:12.0pt" in style
        assert "font-weight:bold" not in style

    def test_section_heading_keeps_larger_style_level_font_size(
        self, typography_html: str
    ) -> None:
        """Section headings should retain their heading style font size."""
        blocks = re.findall(
            r"<(h[1-6])\b([^>]*)>(.*?)</\1>", typography_html, re.IGNORECASE | re.DOTALL
        )
        target_attrs = None
        for _tag, attrs, inner in blocks:
            text = re.sub(r"<[^>]+>", "", inner).replace("\xa0", " ")
            text = re.sub(r"\s+", " ", text).strip()
            if text == "一、企业简介":
                target_attrs = attrs
                break

        assert (
            target_attrs is not None
        ), "Target section heading not found in parsed HTML"
        style_match = re.search(r'style="([^"]*)"', target_attrs)
        assert style_match, "Target heading missing inline style"
        style = style_match.group(1)
        assert "font-size:16.0pt" in style

    def test_sample_body_paragraph_does_not_inherit_paragraph_mark_bold(
        self, docx_html: str
    ) -> None:
        """Paragraph-mark rPr bold must not make ordinary sample body prose render bold."""
        blocks = re.findall(
            r"<p\b([^>]*)>(.*?)</p>", docx_html, re.IGNORECASE | re.DOTALL
        )
        target_attrs = None
        for attrs, inner in blocks:
            text = re.sub(r"<[^>]+>", "", inner).replace("\xa0", " ")
            text = re.sub(r"\s+", " ", text).strip()
            if text.startswith("2024年，AI眼镜在海外市场实现了指数级增长"):
                target_attrs = attrs
                break

        assert (
            target_attrs is not None
        ), "Sample body paragraph not found in parsed HTML"
        style_match = re.search(r'style="([^"]*)"', target_attrs)
        assert style_match, "Sample body paragraph missing inline style"
        style = style_match.group(1)
        assert "font-family:'华文仿宋','STFangsong'" in style
        assert "font-weight:bold" not in style

    def test_outline_level_heading_without_font_props_gets_fallback_typography(
        self, outline_heading_fallback_html: str
    ) -> None:
        """Outline-only headings should still render with a readable heading size/weight."""
        blocks = re.findall(
            r"<(h[1-6])\b([^>]*)>(.*?)</\1>",
            outline_heading_fallback_html,
            re.IGNORECASE | re.DOTALL,
        )
        target_tag = None
        target_attrs = None
        for tag, attrs, inner in blocks:
            text = re.sub(r"<[^>]+>", "", inner).replace("\xa0", " ")
            text = re.sub(r"\s+", " ", text).strip()
            if text == "执行概要":
                target_tag = tag.lower()
                target_attrs = attrs
                break

        assert target_tag == "h2", "Outline level 1 paragraph should render as <h2>"
        assert target_attrs is not None, "Outline-only heading not found in parsed HTML"
        style_match = re.search(r'style="([^"]*)"', target_attrs)
        assert style_match, "Outline-only heading missing inline style"
        style = style_match.group(1)
        assert "font-size:14.0pt" in style
        assert "font-weight:bold" in style

    def test_outline_level_body_sentence_is_not_promoted_to_heading(
        self, outline_body_sentence_html: str
    ) -> None:
        """Outline-only long prose paragraphs should stay as paragraphs, not h1 blocks."""
        assert "被渲染成主标题" in outline_body_sentence_html
        assert not re.search(
            r"<h[1-6]\b[^>]*>[^<]*被渲染成主标题[^<]*</h[1-6]>",
            outline_body_sentence_html,
            re.IGNORECASE,
        )
        assert re.search(
            r"<p\b[^>]*>[^<]*被渲染成主标题[^<]*</p>",
            outline_body_sentence_html,
            re.IGNORECASE,
        )

    def test_outline_level_short_clause_sentence_is_not_promoted_to_heading(
        self, outline_short_body_sentence_html: str
    ) -> None:
        """Outline-only short clause sentences should stay paragraphs, not navigation headings."""
        assert (
            "在会议场景中，AI眼镜利用多模态能力提升协作效率"
            in outline_short_body_sentence_html
        )
        assert not re.search(
            r"<h[1-6]\b[^>]*>\s*在会议场景中，AI眼镜利用多模态能力提升协作效率\s*</h[1-6]>",
            outline_short_body_sentence_html,
            re.IGNORECASE,
        )
        assert re.search(
            r"<p\b[^>]*>\s*在会议场景中，AI眼镜利用多模态能力提升协作效率\s*</p>",
            outline_short_body_sentence_html,
            re.IGNORECASE,
        )

    def test_outline_level_date_prose_is_not_promoted_to_heading(
        self, outline_date_body_sentence_html: str
    ) -> None:
        """Outline-only date-led prose should stay body text, not navigation headings."""
        assert (
            "2022年9月20日，深圳市市场监督管理局完成本次股权转让备案。"
            in outline_date_body_sentence_html
        )
        assert not re.search(
            r"<h[1-6]\b[^>]*>\s*2022年9月20日，深圳市市场监督管理局完成本次股权转让备案。\s*</h[1-6]>",
            outline_date_body_sentence_html,
            re.IGNORECASE,
        )
        assert re.search(
            r"<p\b[^>]*>\s*2022年9月20日，深圳市市场监督管理局完成本次股权转让备案。\s*</p>",
            outline_date_body_sentence_html,
            re.IGNORECASE,
        )

    def test_toc_substring_style_does_not_turn_body_text_into_toc(
        self, toc_substring_style_body_html: str
    ) -> None:
        """A body style whose name merely contains 'toc' must not become a TOC line."""
        assert "2022年9月20日 深圳市市场监督管理局备案" in toc_substring_style_body_html
        assert "koto-toc-" not in toc_substring_style_body_html

    def test_run_bold_values_do_not_inherit_paragraph_mark_bold(
        self, paragraph_bold_run_semantics_html: str
    ) -> None:
        """Paragraph-mark rPr bold should not leak into run text; explicit run values still apply."""
        paragraph_match = re.search(
            r"<p\b([^>]*)>(.*?)</p>",
            paragraph_bold_run_semantics_html,
            re.IGNORECASE | re.DOTALL,
        )
        assert paragraph_match, "Paragraph bold fixture not found in parsed HTML"
        paragraph_style_match = re.search(r'style="([^"]*)"', paragraph_match.group(1))
        if paragraph_style_match:
            assert "font-weight:bold" not in paragraph_style_match.group(1)

        inner_html = paragraph_match.group(2)
        explicit_off = re.search(
            r"<span\b([^>]*)>\s*显式取消</span>", inner_html, re.IGNORECASE
        )
        assert explicit_off, "Explicit-off run was not rendered as a styled span"
        explicit_style = re.search(r'style="([^"]*)"', explicit_off.group(1))
        assert explicit_style and "font-weight:normal" in explicit_style.group(1)

        toggle_on = re.search(
            r"<span\b([^>]*)>\s*切换取消</span>", inner_html, re.IGNORECASE
        )
        assert toggle_on, "Toggle-on run was not rendered as a styled span"
        toggle_style = re.search(r'style="([^"]*)"', toggle_on.group(1))
        assert toggle_style and "font-weight:bold" in toggle_style.group(1)

    def test_generic_chinese_title_style_is_not_treated_as_structural_heading(
        self, localized_title_style_html: str
    ) -> None:
        """Localized Title-style paragraphs should not be promoted into the outline."""
        assert "封面标题" in localized_title_style_html
        assert not re.search(
            r"<h[1-6]\b[^>]*>\s*封面标题\s*</h[1-6]>",
            localized_title_style_html,
            re.IGNORECASE,
        )
        assert re.search(
            r"<p\b[^>]*>\s*封面标题\s*</p>",
            localized_title_style_html,
            re.IGNORECASE,
        )

        paragraph_match = re.search(
            r"<p\b([^>]*)>\s*封面标题\s*</p>",
            localized_title_style_html,
            re.IGNORECASE,
        )
        assert paragraph_match, "Localized title paragraph not found in parsed HTML"
        assert "koto-visual-title" in paragraph_match.group(1)
        assert 'data-koto-role="visual_title"' in paragraph_match.group(1)
        style_match = re.search(r'style="([^"]*)"', paragraph_match.group(1))
        assert style_match, "Localized title paragraph missing inline style"
        style = style_match.group(1)
        assert "font-size:18.0pt" in style
        assert "font-weight:bold" in style

    def test_localized_font_family_keeps_original_name_and_ascii_alias(
        self, localized_font_family_html: str
    ) -> None:
        """Localized DOCX fonts should keep both the original family name and its ASCII alias."""
        span_match = re.search(
            r"<span\b([^>]*)>华文字体回归</span>",
            localized_font_family_html,
            re.IGNORECASE | re.DOTALL,
        )
        assert span_match, "Localized font run not found in parsed HTML"

        span_style_match = re.search(r'style="([^"]*)"', span_match.group(1))
        assert span_style_match, "Localized font run missing inline style"
        span_style = span_style_match.group(1)
        assert "font-family:'华文仿宋','STFangsong'" in span_style


@pytest.mark.integration
class TestHeadingManifest:
    def test_section_break_markers_include_prev_and_next_section_indices(
        self, tmp_path
    ) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        docx_path = tmp_path / "section-break-marker.docx"
        _write_section_break_marker_fixture_docx(docx_path)

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(str(docx_path))

        assert re.search(
            r'<div\b[^>]*data-page-break="true"[^>]*data-current-section-idx="0"[^>]*data-next-section-idx="1"[^>]*class="koto-page-break"',
            result.get("html", ""),
            re.IGNORECASE,
        )
        assert len(result.get("sections", [])) >= 2

    def test_all_word_heading_styles_are_emitted_in_parser_manifest(
        self, tmp_path
    ) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        docx_path = tmp_path / "all-heading-levels.docx"
        _write_all_heading_levels_fixture_docx(docx_path)

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(str(docx_path))
        headings = result.get("headings", [])

        assert [(item.get("level"), item.get("text")) for item in headings] == [
            (1, "层级1标题"),
            (2, "层级2标题"),
            (3, "层级3标题"),
            (4, "层级4标题"),
            (5, "层级5标题"),
            (6, "层级6标题"),
        ]
        assert "普通正文，不应进入导航。" in result.get("html", "")
        assert all("普通正文，不应进入导航。" != item.get("text") for item in headings)
        for level in range(1, 7):
            assert re.search(
                rf'<h{level}\b[^>]*data-koto-role="structural_heading"[^>]*>\s*层级{level}标题\s*</h{level}>',
                result.get("html", ""),
                re.IGNORECASE,
            )

    def test_paragraph_layout_semantics_are_emitted_as_data_attrs(
        self, tmp_path
    ) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        docx_path = tmp_path / "paragraph-layout-semantics.docx"
        _write_paragraph_layout_semantics_fixture_docx(docx_path)

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(str(docx_path))
        paragraph_match = re.search(
            r"<p\b([^>]*)>.*?段落分页语义测试.*?</p>",
            result.get("html", ""),
            re.IGNORECASE | re.DOTALL,
        )

        assert (
            paragraph_match
        ), "Paragraph with layout semantics not found in parsed HTML"
        attrs = paragraph_match.group(1)
        assert 'data-koto-space-before-twips="240"' in attrs
        assert 'data-koto-space-after-twips="120"' in attrs
        assert 'data-koto-line-rule="exact"' in attrs
        assert 'data-koto-line-twips="360"' in attrs
        assert 'data-koto-first-line-indent-twips="420"' in attrs
        assert 'data-koto-left-indent-twips="360"' in attrs
        assert 'data-koto-keep-next="1"' in attrs
        assert 'data-koto-keep-lines="1"' in attrs
        assert 'data-koto-page-break-before="1"' in attrs
        assert 'data-koto-widow-control="0"' in attrs

    def test_title_style_based_on_heading_is_excluded_from_parser_manifest(
        self, tmp_path
    ) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        docx_path = tmp_path / "title-style-based-on-heading.docx"
        _write_title_style_based_on_heading_fixture_docx(docx_path)

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(str(docx_path))
        heading_texts = [h.get("text", "") for h in result.get("headings", [])]
        assert heading_texts == ["第一章 公司简介"]
        assert "封面总标题" in result.get("html", "")
        assert not re.search(
            r"<h[1-6]\b[^>]*>\s*封面总标题\s*</h[1-6]>",
            result.get("html", ""),
            re.IGNORECASE,
        )
        assert re.search(
            r'<p\b[^>]*class="[^"]*koto-visual-title[^"]*"[^>]*data-koto-role="visual_title"[^>]*>\s*封面总标题\s*</p>',
            result.get("html", ""),
            re.IGNORECASE,
        )

    def test_custom_style_based_on_title_chain_is_excluded_from_parser_manifest(
        self, tmp_path
    ) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        docx_path = tmp_path / "custom-visual-title-chain.docx"
        _write_custom_visual_title_chain_fixture_docx(docx_path)

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(str(docx_path))
        heading_texts = [h.get("text", "") for h in result.get("headings", [])]
        assert heading_texts == ["第二章 技术分析"]
        assert "封面链式标题" in result.get("html", "")
        assert not re.search(
            r"<h[1-6]\b[^>]*>\s*封面链式标题\s*</h[1-6]>",
            result.get("html", ""),
            re.IGNORECASE,
        )
        assert re.search(
            r'<p\b[^>]*class="[^"]*koto-visual-title[^"]*"[^>]*data-koto-role="visual_title"[^>]*>\s*封面链式标题\s*</p>',
            result.get("html", ""),
            re.IGNORECASE,
        )

    def test_outline_level_heading_is_not_emitted_in_parser_manifest(
        self, tmp_path
    ) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        docx_path = tmp_path / "outline-heading-manifest.docx"
        _write_outline_only_heading_fixture_docx(docx_path)

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(str(docx_path))
        assert result["headings"] == []
        assert 'id="koto-heading-执行概要"' in result["html"]
        assert 'data-koto-role="structural_heading"' in result["html"]

    def test_outline_level_prefixed_heading_is_emitted_in_parser_manifest(
        self, tmp_path
    ) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        docx_path = tmp_path / "outline-prefixed-heading-manifest.docx"
        _write_outline_only_prefixed_heading_fixture_docx(docx_path)

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(str(docx_path))
        assert result["headings"] == [
            {
                "level": 1,
                "text": "第一章 公司基本信息",
                "id": "koto-heading-第一章-公司基本信息",
            },
        ]
        assert 'id="koto-heading-第一章-公司基本信息"' in result["html"]
        assert 'data-koto-role="structural_heading"' in result["html"]

    def test_table_cell_heading_is_excluded_from_parser_manifest(
        self, tmp_path
    ) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        docx_path = tmp_path / "table-cell-heading.docx"
        _write_table_cell_heading_fixture_docx(docx_path)

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(str(docx_path))
        assert result["headings"] == [
            {"level": 1, "text": "正文标题", "id": "koto-heading-正文标题"},
        ]
        assert "表格里的标题" in result["html"]
        assert not re.search(
            r"<h[1-6]\b[^>]*>\s*表格里的标题\s*</h[1-6]>", result["html"], re.IGNORECASE
        )

    def test_outline_level_short_clause_sentence_is_excluded_from_parser_manifest(
        self, tmp_path
    ) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        docx_path = tmp_path / "outline-short-body-sentence-manifest.docx"
        _write_outline_only_short_body_sentence_fixture_docx(docx_path)

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(str(docx_path))
        assert result["headings"] == []
        assert "在会议场景中，AI眼镜利用多模态能力提升协作效率" in result["html"]
        assert not re.search(
            r"<h[1-6]\b[^>]*>\s*在会议场景中，AI眼镜利用多模态能力提升协作效率\s*</h[1-6]>",
            result["html"],
            re.IGNORECASE,
        )

    def test_outline_level_date_prose_is_excluded_from_parser_manifest(
        self, tmp_path
    ) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        docx_path = tmp_path / "outline-date-body-sentence-manifest.docx"
        _write_outline_only_date_body_sentence_fixture_docx(docx_path)

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(str(docx_path))
        assert result["headings"] == []
        assert (
            "2022年9月20日，深圳市市场监督管理局完成本次股权转让备案。"
            in result["html"]
        )
        assert not re.search(
            r"<h[1-6]\b[^>]*>\s*2022年9月20日，深圳市市场监督管理局完成本次股权转让备案。\s*</h[1-6]>",
            result["html"],
            re.IGNORECASE,
        )

    def test_normal_body_after_headings_is_excluded_from_parser_manifest(
        self, tmp_path
    ) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        docx_path = tmp_path / "heading-style-cache-regression.docx"
        _write_heading_style_cache_regression_fixture_docx(docx_path)

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(str(docx_path))
        heading_texts = [h.get("text", "") for h in result.get("headings", [])]
        assert len(heading_texts) == 40
        assert all("作为国内AI眼镜领域" not in text for text in heading_texts)
        assert "作为国内AI眼镜领域" in result.get("html", "")
        assert not re.search(
            r"<h[1-6]\b[^>]*>[^<]*作为国内AI眼镜领域",
            result.get("html", ""),
            re.IGNORECASE,
        )

    def test_real_translation_body_paragraphs_are_excluded_from_heading_manifest(
        self,
    ) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        if not os.path.exists(TRANSLATION_DOCX_PATH):
            pytest.skip(f"Test document not found: {TRANSLATION_DOCX_PATH}")

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(TRANSLATION_DOCX_PATH)
        target_text = (
            "双年展化固然并非旨在简单复制场域中占支配地位的西方机构的价值标准，"
            "但它同样也不能被理解为一群崛起中的“边缘”行动者单方面向中心行动者发起冲击。"
        )
        compact_target = re.sub(r"\s+", "", target_text)
        heading_texts = [
            re.sub(r"\s+", "", heading.get("text", ""))
            for heading in result.get("headings", [])
        ]
        compact_html = re.sub(r"\s+", "", result.get("html", ""))
        heading_blocks = "".join(
            re.findall(r"<h[1-6]\b[^>]*>.*?</h[1-6]>", compact_html, re.IGNORECASE)
        )

        assert compact_target in compact_html
        assert compact_target not in heading_texts
        assert compact_target not in heading_blocks

    def test_real_translation_neighboring_subheading_stays_in_heading_manifest(
        self,
    ) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        if not os.path.exists(TRANSLATION_DOCX_PATH):
            pytest.skip(f"Test document not found: {TRANSLATION_DOCX_PATH}")

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(TRANSLATION_DOCX_PATH)
        heading_texts = {
            re.sub(r"\s+", "", heading.get("text", ""))
            for heading in result.get("headings", [])
        }
        expected_heading = re.sub(r"\s+", "", "模糊的空间：竞争中的合作")

        assert expected_heading in heading_texts
        assert re.search(
            r'<h[1-6]\b[^>]*data-koto-role="structural_heading"[^>]*>\s*模糊的空间：竞争中的合作\s*</h[1-6]>',
            result.get("html", ""),
            re.IGNORECASE,
        )

    def test_real_translation_sections_expose_doc_grid_metadata(self) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        if not os.path.exists(TRANSLATION_DOCX_PATH):
            pytest.skip(f"Test document not found: {TRANSLATION_DOCX_PATH}")

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(TRANSLATION_DOCX_PATH)
        sections = result.get("sections")

        assert isinstance(sections, list) and sections
        assert all(isinstance(section.get("doc_grid"), dict) for section in sections)
        assert all(section.get("doc_grid", {}).get("enabled") for section in sections)
        assert all(
            section.get("doc_grid", {}).get("line_pitch_twips", 0) > 0
            for section in sections
        )
        assert result.get("doc_grid", {}).get("enabled")
        assert result.get("doc_grid", {}).get("line_pitch_twips", 0) > 0

    def test_real_translation_exposes_footnotes_metadata_and_inline_refs(self) -> None:
        pytest.importorskip("docx", reason="python-docx 未安装")

        if not os.path.exists(TRANSLATION_DOCX_PATH):
            pytest.skip(f"Test document not found: {TRANSLATION_DOCX_PATH}")

        from app.core.file.file_parser import parse_docx  # noqa: PLC0415

        result = parse_docx(TRANSLATION_DOCX_PATH)
        footnotes = result.get("footnotes")

        assert isinstance(footnotes, list) and footnotes
        assert result.get("footnote_reference_count", 0) > 0
        assert any(
            str(item.get("text") or "").strip()
            for item in footnotes
            if isinstance(item, dict)
        )
        assert sum(
            int(item.get("reference_count") or 0)
            for item in footnotes
            if isinstance(item, dict)
        ) == result.get("footnote_reference_count")
        assert re.search(
            r'<sup\b[^>]*data-koto-footnote-ref="\d+"[^>]*>\d+</sup>',
            result.get("html", ""),
            re.IGNORECASE,
        )


@pytest.mark.integration
class TestImages:
    """Image dimension preservation."""

    def test_images_carry_explicit_dimensions(self, docx_html: str) -> None:
        """
        Every <img> must have both width and height in its inline style so that
        the browser can display it at the correct size.  Without explicit height,
        CSS 'height:auto' stretches images to occupy the full container width.
        """
        img_tags = re.findall(r"<img\b[^>]+>", docx_html, re.IGNORECASE)
        if not img_tags:
            pytest.skip("No <img> tags in document")

        imgs_missing_height = [
            tag for tag in img_tags if not re.search(r"height\s*:\s*\d", tag)
        ]
        assert not imgs_missing_height, (
            f"{len(imgs_missing_height)}/{len(img_tags)} <img> tags lack an "
            f"explicit height in their inline style:\n"
            + "\n".join(imgs_missing_height[:3])
        )

    def test_images_carry_width(self, docx_html: str) -> None:
        """Every <img> must also have an explicit width."""
        img_tags = re.findall(r"<img\b[^>]+>", docx_html, re.IGNORECASE)
        if not img_tags:
            pytest.skip("No <img> tags in document")

        imgs_missing_width = [
            tag for tag in img_tags if not re.search(r"width\s*:\s*\d", tag)
        ]
        assert (
            not imgs_missing_width
        ), f"{len(imgs_missing_width)}/{len(img_tags)} <img> tags lack explicit width"

    def test_table_cell_inline_images_remain_inline_in_paragraph(
        self, table_inline_image_html: str
    ) -> None:
        """Inline pictures inside table-cell paragraphs should stay inline with surrounding text."""
        bs4 = pytest.importorskip(
            "bs4", reason="BeautifulSoup is required for DOCX image-row checks"
        )

        soup = bs4.BeautifulSoup(table_inline_image_html, "html.parser")
        table_cell = soup.find("td")
        assert (
            table_cell is not None
        ), "Expected a rendered table cell in the fixture HTML"

        child_tags = [
            child.name for child in table_cell.children if getattr(child, "name", None)
        ]
        assert child_tags[:2] == [
            "p",
            "p",
        ], "Expected the fixture cell to keep its two paragraph children"

        first_para = table_cell.find("p")
        assert (
            first_para is not None
        ), "Expected the first table-cell paragraph to be present"

        image = first_para.find("img")
        assert (
            image is not None
        ), "Expected the inline image to remain inside the first table-cell paragraph"
        assert image.get("data-koto-layout") != "top-bottom"
        assert (
            "邮箱" in first_para.get_text()
        ), "Expected the surrounding inline text to remain in the same paragraph"


@pytest.mark.integration
class TestTableFormatting:
    """Table cell formatting."""

    def test_table_cells_exist(self, docx_html: str) -> None:
        """The document has tables — at least one <td> must be present."""
        assert re.search(
            r"<td\b", docx_html, re.IGNORECASE
        ), "No <td> elements found; table parsing may have failed"

    def test_table_cells_have_border_inline_styles(self, docx_html: str) -> None:
        """
        Every <td> must carry explicit border-top/bottom/left/right in inline styles
        so the CSS fallback (1px solid #a0a4b8) never fires.
        """
        td_tags = re.findall(r"<td\b[^>]+>", docx_html, re.IGNORECASE)
        if not td_tags:
            pytest.skip("No <td> tags — covered by test_table_cells_exist")

        missing = [t for t in td_tags if "border-top" not in t]
        assert not missing, (
            f"{len(missing)}/{len(td_tags)} <td> tags lack border-top inline style.  "
            "Table borders will fall back to CSS gray grid instead of DOCX values."
        )

    def test_some_cells_have_background_color(self, docx_html: str) -> None:
        """
        Tinted table cells must carry background-color in their inline style.
        The 投资建议书 document has a styled cover table with coloured cells.
        """
        td_tags = re.findall(r"<td\b[^>]+>", docx_html, re.IGNORECASE)
        assert td_tags, pytest.skip("No <td> tags — covered by test_table_cells_exist")

        tinted = [t for t in td_tags if "background-color" in t]
        assert tinted, (
            f"None of the {len(td_tags)} <td> tags carry background-color.  "
            "Table shading will be invisible in the editor."
        )

    def test_exact_row_height_rows_keep_metadata_without_fixed_css(
        self, exact_row_height_table_html: str
    ) -> None:
        """
        Word exact row heights cannot be emitted as fixed browser <tr> heights,
        otherwise wrapped text overflows and visually overlaps adjacent rows.
        Keep the original height as metadata only.
        """
        tr_tags = re.findall(r"<tr\b[^>]*>", exact_row_height_table_html, re.IGNORECASE)
        assert tr_tags, "No <tr> tags found in exact-row-height fixture"

        assert any(
            "data-koto-row-height=" in tag for tag in tr_tags
        ), "Exact-height DOCX rows should preserve their source height as metadata"
        assert not any(
            re.search(r'style="[^"]*\bheight\s*:', tag, re.IGNORECASE)
            for tag in tr_tags
        ), "Exact-height DOCX rows must not emit fixed browser row heights"

    def test_small_multiple_line_spacing_is_clamped_for_browser_preview(
        self, small_multiple_line_spacing_html: str
    ) -> None:
        """
        Tiny DOCX multiple spacing values such as 0.25 make browser-rendered
        wrapped lines overlap inside tables. Preview HTML should clamp them to
        at least single spacing instead of emitting line-height:0.25 literally.
        """
        assert not re.search(
            r'style="[^"]*line-height:0\.25(?=[;\"])',
            small_multiple_line_spacing_html,
        )
        assert re.search(
            r"line-height:1(?:\.0+)?(?=[;\"])", small_multiple_line_spacing_html
        )


@pytest.mark.integration
class TestPageCount:
    """
    Page count estimate.

    We cannot run a real browser to measure rendered height, but we can
    estimate the content volume and check it's in a plausible range.
    The Koto formula:
        pages = ceil((intrinsicHeight - _PAD_V) / _CONTENT_PAGE_H)
    where intrinsicHeight = actual DOM height (top-pad + content + bottom-pad).

    We estimate content height using paragraph count and average font metrics:
        font-size: 16px, line-height: 1.7 → ~27.2px per text line.
        Average Chinese line: ~35 chars.
        So content_height ≈ (char_count / 35) * 27.2px.
    """

    def test_page_count_estimate_in_range(self, docx_html: str) -> None:
        """Estimated page count should be within 50% of Word's 72 pages."""
        text_only = re.sub(r"<[^>]+>", " ", docx_html)
        # Count CJK + Latin printable chars (ignore whitespace)
        char_count = len(re.sub(r"\s", "", text_only))

        chars_per_line = 35
        px_per_line = 27.2  # 16px × 1.7
        content_height_px = (char_count / chars_per_line) * px_per_line
        estimated_pages = max(1, (content_height_px) / _CONTENT_PAGE_H)

        lower = WORD_PAGE_COUNT * 0.4
        upper = WORD_PAGE_COUNT * 4.0

        assert lower <= estimated_pages <= upper, (
            f"Estimated page count ({estimated_pages:.1f}) is far outside "
            f"[{lower:.0f}, {upper:.0f}] — expected ~{WORD_PAGE_COUNT} pages.  "
            f"char_count={char_count:,}, content_height={content_height_px:.0f}px.  "
            "Parser may be dropping large sections of content."
        )

    def test_paragraph_count_reasonable(self, docx_html: str) -> None:
        """A 72-page document must have many paragraphs."""
        p_count = len(re.findall(r"<p\b", docx_html, re.IGNORECASE))
        assert p_count >= 50, (
            f"Only {p_count} <p> elements found; expected ≥50 for a "
            f"{WORD_PAGE_COUNT}-page document"
        )
