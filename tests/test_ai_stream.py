#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Tests for the File Assistant AI stream endpoint.

Validates:
1. SSE streaming response format
2. All action types (polish, translate, find_replace, find_reference, etc.)
3. Full-text context injection
4. Chart rerun endpoint
"""

import importlib
import json
import io
import re
import sys
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

# ── Ensure project root is importable ──
ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))


# ── Helper ──


def parse_sse_events(response_data: bytes) -> list:
    """Parse SSE response bytes into list of event dicts."""
    events = []
    for chunk in response_data.decode("utf-8", errors="replace").split("\n\n"):
        chunk = chunk.strip()
        if chunk.startswith("data: "):
            try:
                events.append(json.loads(chunk[6:]))
            except json.JSONDecodeError:
                pass
    return events


# ── Mock LLM fixture ──


class FakeChunk:
    def __init__(self, text):
        self.text = text


def _make_fake_stream(prompt_keyword_responses: dict):
    """Create a fake generate_content_stream that returns deterministic output."""

    def fake_stream(model, contents, config=None):
        for keyword, response in prompt_keyword_responses.items():
            if keyword in str(contents):
                yield FakeChunk(response)
                return
        yield FakeChunk("默认 AI 回复。")

    return fake_stream


@pytest.fixture
def app_client():
    """Create a Flask test client with mocked LLM."""
    # Mock the Gemini client before importing app
    mock_client = MagicMock()
    mock_client.models.generate_content_stream = _make_fake_stream(
        {
            "润色": "这是一段经过精心润色的优雅文本。",
            "翻译": "This is the translated text.",
            "总结": "本文主要讨论了三个核心观点。",
            "替换": '{"replacements": [{"from": "你好", "to": "您好"}, {"from": "世界", "to": "地球"}], "summary": "共替换 2 处"}',
            "引用": "1. 【论文】Smith et al. (2024) — AI辅助写作综述\n   链接：待核实",
            "检查": "1. 【第2行】你好 → 您好（更正式）",
            "改写": "这是用全新措辞表达的内容。",
            "续写": "接下来，我们将探讨更深层次的问题。",
        }
    )

    with patch.dict("sys.modules", {}):
        # We need to patch the client object in web.app
        try:
            import app.core.agent.agent_loop as agent_loop_module
            from app.core.security.output_validator import OutputValidator
            from app.core.agent.lifecycle import (
                evt_error,
                evt_stream_chunk,
                evt_task_complete,
            )
            from web.app import app

            app.config["TESTING"] = True
            # Patch the client and types (to avoid google.genai circular import under test)
            import web.app as web_app_module

            original_client = getattr(web_app_module, "client", None)
            original_api_key = getattr(web_app_module, "API_KEY", None)
            original_types = getattr(web_app_module, "types", None)
            original_llm_judge = getattr(OutputValidator, "_llm_judge")
            original_loop_run = agent_loop_module.KotoAgentLoop.run
            mock_types = MagicMock()
            mock_types.GenerateContentConfig.return_value = MagicMock()

            default_action_results = {
                "polish": "这是一段经过精心润色的优雅文本。",
                "translate": "This is the translated text.",
                "summary": "本文主要讨论了三个核心观点。",
                "find_replace": '{"replacements": [{"from": "你好", "to": "您好"}, {"from": "世界", "to": "地球"}], "summary": "共替换 2 处"}',
                "find_reference": "1. 【论文】Smith et al. (2024) — AI辅助写作综述\n   链接：待核实",
                "check": "1. 【第2行】你好 → 您好（更正式）",
                "rewrite": "这是用全新措辞表达的内容。",
                "continue": "接下来，我们将探讨更深层次的问题。",
                "custom_instruction": "已按要求处理当前内容。",
            }

            def fake_loop_run(self, request):
                model_mode = (getattr(request, "model_mode", "") or "").strip().lower()
                action_name = (
                    (getattr(request, "action_type", "") or "").strip().lower()
                )

                if model_mode == "local":
                    if not agent_loop_module._is_ollama_alive():
                        yield evt_error("Ollama not running")
                        return
                    result_text = "本地Ollama响应"
                elif model_mode == "cloud":
                    result_text = "云端Gemini响应"
                else:
                    result_text = default_action_results.get(
                        action_name, "默认 AI 回复。"
                    )

                yield evt_stream_chunk(result_text)
                yield evt_task_complete(result=result_text)

            web_app_module.client = mock_client
            web_app_module.API_KEY = "test-key-mock"
            web_app_module.types = mock_types
            OutputValidator._llm_judge = classmethod(
                lambda cls, text, original_prompt: None
            )
            agent_loop_module.KotoAgentLoop.run = fake_loop_run
            yield app.test_client()
            # Restore
            web_app_module.client = original_client
            web_app_module.API_KEY = original_api_key
            web_app_module.types = original_types
            OutputValidator._llm_judge = original_llm_judge
            agent_loop_module.KotoAgentLoop.run = original_loop_run
        except ImportError as e:
            pytest.skip(f"Cannot import web.app: {e}")


# ══════════════════════════════════════════════════════════════
# Tests
# ══════════════════════════════════════════════════════════════


class TestEditorAIStream:
    """Tests for POST /api/editor/ai/stream"""

    def test_whitebox_task_stream_executes_xlsx_to_docx_write_loop(
        self, app_client, tmp_path, monkeypatch
    ):
        import openpyxl
        from docx import Document

        from app.core.agent.file_task_runtime import FileTaskRuntime

        workbook_path = tmp_path / "sales.xlsx"
        target_path = tmp_path / "target.docx"

        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "汇总表"
        sheet.append(["客户名称", "产品名称", "数量"])
        sheet.append(["杭州新汇鑫光电有限公司", "LASER", 1])
        sheet.append(["山东镭鸟激光设备有限公司", "LASER", 2])
        workbook.save(workbook_path)

        document = Document()
        document.add_paragraph("雷鸟访谈问题")
        document.save(target_path)

        responses = iter(
            [
                {
                    "content": "先读取 Excel。",
                    "tool_calls": [
                        {
                            "name": "read_sheet_data",
                            "args": {
                                "path": str(workbook_path),
                                "sheet_name": "Sheet1",
                                "max_rows": "2",
                            },
                        }
                    ],
                },
                {
                    "content": "把 Excel 表格加入 Word。",
                    "tool_calls": [
                        {
                            "name": "insert_excel_as_docx_table",
                            "args": {
                                "source_path": str(workbook_path),
                                "target_path": str(target_path),
                                "sheet_name": "Sheet1",
                                "table_title": "销售台账数据",
                                "max_rows": "2",
                            },
                        }
                    ],
                },
                {"content": "已将销售台账数据加入 Word。", "tool_calls": []},
            ]
        )

        def fake_call_model(self, **kwargs):
            return next(responses)

        monkeypatch.setattr(FileTaskRuntime, "_call_model", fake_call_model)

        resp = app_client.post(
            "/api/editor/ai/task-stream",
            json={
                "task": "把 xlsx 表格加入 docx",
                "session_id": "editor_demo",
                "model_mode": "local",
                "target_path": str(target_path),
                "files": [
                    {
                        "path": str(workbook_path),
                        "name": "销售台账.xlsx",
                        "type": "xlsx",
                        "content": "Excel 上下文",
                    },
                    {
                        "path": str(target_path),
                        "name": "雷鸟访谈问题.docx",
                        "type": "docx",
                        "content": "Word 上下文",
                        "target": True,
                    },
                ],
            },
        )

        events = parse_sse_events(resp.get_data())
        read_finished = next(
            event
            for event in events
            if event.get("type") == "tool.finished"
            and event.get("payload", {}).get("tool_name") == "read_sheet_data"
        )
        file_changed = next(
            event for event in events if event.get("type") == "file.changed"
        )
        check_finished = next(
            event for event in events if event.get("type") == "check.finished"
        )
        run_finished = events[-1]

        assert resp.status_code == 200
        assert "汇总表" in str(
            read_finished.get("payload", {}).get("result_preview", "")
        )
        assert file_changed["payload"]["sheet"] == "汇总表"
        assert file_changed["payload"]["requested_sheet"] == "Sheet1"
        assert file_changed["payload"]["rows_written"] == 2
        assert file_changed["payload"]["columns_written"] == 3
        assert "used '汇总表' instead" in file_changed["payload"]["warning"]
        assert check_finished["payload"]["status"] == "verified"
        assert run_finished["payload"]["completed_task"] is True

        saved = Document(str(target_path))
        assert len(saved.tables) == 1
        assert saved.tables[0].cell(1, 0).text == "杭州新汇鑫光电有限公司"
        assert saved.tables[0].cell(2, 0).text == "山东镭鸟激光设备有限公司"

    def test_whitebox_task_stream_uses_model_docx_writer_for_stepwise_pdf_summary(
        self,
        app_client,
        tmp_path,
        monkeypatch,
    ):
        from app.core.agent.file_task_runtime import FileTaskRuntime
        import app.core.agent.task_tools as task_tools

        monkeypatch.setattr(task_tools, "_WORKSPACE_ROOT", str(tmp_path))
        monkeypatch.setattr(
            task_tools,
            "parse_file_to_text",
            lambda **kwargs: (
                "[Page 4] The Global Rules of Art. Larissa Buchholz. Princeton University Press.\n"
                "[Page 5] Contents: The Emergence of a Global Field in the Contemporary Visual Arts.\n"
                "[Page 6] Chapter 1 introduces a global field approach to art and culture.\n"
                "[Page 7] Chapter 2 discusses the genesis of a global art field.\n"
                "[Page 8] Later chapters cover divisions and positions in a cultural world economy."
            ),
        )

        model_calls = []

        def fake_call_model(self, **kwargs):
            model_calls.append(kwargs)
            return {
                "content": "",
                "tool_calls": [
                    {
                        "name": "write_docx_content",
                        "args": {
                            "path": str(target_path),
                            "paragraphs": json.dumps(
                                [
                                    {
                                        "text": "当前页窗摘要（第 1-3 页）",
                                        "style": "Heading 1",
                                    },
                                    {
                                        "text": "文档识别：当前页窗来自 The Global Rules of Art，记录书籍出版信息和目录入口。"
                                    },
                                    {
                                        "text": "段落主题：模型将目录页窗概括为全书问题意识和章节结构的定位材料。"
                                    },
                                    {
                                        "text": "结构线索：Part I 从全球场域方法进入，随后讨论艺术场域生成和文化世界经济中的分化。"
                                    },
                                    {
                                        "text": "内容线索：这是模型生成的综合摘要，用于验证文件助手入口没有直接走固定兜底模板。"
                                    },
                                    {"text": "来源页码：第 1-3 页"},
                                ],
                                ensure_ascii=False,
                            ),
                        },
                    }
                ],
            }

        monkeypatch.setattr(FileTaskRuntime, "_call_model", fake_call_model)
        target_path = tmp_path / "global-rules-summary.docx"

        resp = app_client.post(
            "/api/editor/ai/task-stream",
            json={
                "task": (
                    "这是一篇非常长的pdf，里面有大量内容。我需要你做的是一个分步任务，将任务拆分成很多个小任务，"
                    "一步一步完成，每完成一步和我汇报一下我来说继续。你将总结整篇文章的核心内容，你创建一个docx文件，"
                    "记录你每一步发现的要点，然后每一步完成后更新docx"
                ),
                "session_id": "workspace_demo",
                "target_path": str(target_path),
                "files": [
                    {
                        "path": "global-rules.pdf",
                        "name": "The Global Rules of Art.pdf",
                        "type": "pdf",
                        "content": "[Page 4] "
                        + ("book metadata table of contents " * 300),
                    }
                ],
            },
        )
        events = parse_sse_events(resp.get_data())

        assert resp.status_code == 200
        assert len(model_calls) >= 1
        file_changed = next(
            event for event in events if event.get("type") == "file.changed"
        )
        check_finished = next(
            event for event in events if event.get("type") == "check.finished"
        )
        run_finished = events[-1]

        assert file_changed["payload"]["operation"] == "write_docx_content"
        assert file_changed["payload"]["paragraphs_written"] >= 4
        assert check_finished["payload"]["status"] == "awaiting_confirmation"
        assert (
            check_finished["payload"]["next_action_artifact"]["route"]
            == "long_pdf_stepwise_docx_summary"
        )
        assert (
            run_finished["payload"]["runtime"]["terminal_status"]
            == "awaiting_confirmation"
        )
        assert target_path.exists()
        from docx import Document

        written_text = "\n".join(p.text for p in Document(str(target_path)).paragraphs)
        assert "模型生成的综合摘要" in written_text

    def test_whitebox_task_stream_reports_no_write_when_docx_is_unchanged(
        self, app_client, tmp_path, monkeypatch
    ):
        from docx import Document

        from app.core.agent.file_task_runtime import FileTaskRuntime

        target_path = tmp_path / "target.docx"
        document = Document()
        document.add_paragraph("雷鸟访谈问题")
        document.save(target_path)

        responses = iter(
            [
                {
                    "content": "尝试把 Excel 表格加入 Word。",
                    "tool_calls": [
                        {
                            "name": "insert_excel_as_docx_table",
                            "args": {
                                "source_path": str(tmp_path / "missing.xlsx"),
                                "target_path": str(target_path),
                                "sheet_name": "Sheet1",
                                "max_rows": "50",
                            },
                        }
                    ],
                },
                {"content": "已完成。", "tool_calls": []},
            ]
        )

        def fake_call_model(self, **kwargs):
            return next(responses)

        monkeypatch.setattr(FileTaskRuntime, "_call_model", fake_call_model)

        resp = app_client.post(
            "/api/editor/ai/task-stream",
            json={
                "task": "把 xlsx 表格加入 docx",
                "session_id": "editor_demo",
                "target_path": str(target_path),
                "files": [
                    {
                        "path": str(target_path),
                        "name": "雷鸟访谈问题.docx",
                        "type": "docx",
                        "content": "Word 上下文",
                        "target": True,
                    },
                ],
            },
        )

        events = parse_sse_events(resp.get_data())
        insert_finished = next(
            event
            for event in events
            if event.get("type") == "tool.finished"
            and event.get("payload", {}).get("tool_name")
            == "insert_excel_as_docx_table"
        )
        check_finished = next(
            event for event in events if event.get("type") == "check.finished"
        )
        run_finished = events[-1]

        assert resp.status_code == 200
        assert insert_finished["payload"]["success"] is False
        assert "File not found" in insert_finished["payload"]["result_preview"]
        assert not any(event.get("type") == "file.changed" for event in events)
        assert check_finished["payload"]["status"] == "no_file_change"
        assert run_finished["payload"]["completed_task"] is False
        assert len(Document(str(target_path)).tables) == 0

    def test_polish_returns_sse(self, app_client):
        """润色请求应返回 SSE 流，包含 token 和 done 事件"""
        resp = app_client.post(
            "/api/editor/ai/stream",
            json={
                "action": "polish",
                "selection": "这段文字需要被润色一下。",
            },
        )
        assert resp.status_code == 200
        assert "text/event-stream" in resp.content_type
        events = parse_sse_events(resp.data)
        assert len(events) >= 1
        types = {e["type"] for e in events}
        assert "token" in types or "done" in types

    def test_polish_with_full_text_context(self, app_client):
        """润色请求带 full_text 应正常工作"""
        resp = app_client.post(
            "/api/editor/ai/stream",
            json={
                "action": "polish",
                "selection": "这段文字需要润色。",
                "full_text": "第一段落。这段文字需要润色。第三段落结尾。",
            },
        )
        assert resp.status_code == 200
        events = parse_sse_events(resp.data)
        assert len(events) >= 1

    def test_translate_action(self, app_client):
        """翻译请求应返回翻译结果"""
        resp = app_client.post(
            "/api/editor/ai/stream",
            json={
                "action": "translate",
                "selection": "你好世界",
            },
        )
        assert resp.status_code == 200
        events = parse_sse_events(resp.data)
        assert len(events) >= 1

    def test_find_replace_action(self, app_client):
        """查找替换请求应返回 JSON 格式替换列表"""
        resp = app_client.post(
            "/api/editor/ai/stream",
            json={
                "action": "find_replace",
                "instruction": "把所有你好替换成您好",
                "full_text": "你好世界，你好中国，你好大家。",
            },
        )
        assert resp.status_code == 200
        events = parse_sse_events(resp.data)
        assert len(events) >= 1

    def test_find_reference_action(self, app_client):
        """引用查找请求应返回引用列表"""
        resp = app_client.post(
            "/api/editor/ai/stream",
            json={
                "action": "find_reference",
                "selection": "人工智能在教育中的应用越来越广泛。",
                "full_text": "本文探讨人工智能在教育中的应用。",
            },
        )
        assert resp.status_code == 200
        events = parse_sse_events(resp.data)
        assert len(events) >= 1

    def test_empty_selection_returns_error(self, app_client):
        """空选区应返回错误"""
        resp = app_client.post(
            "/api/editor/ai/stream",
            json={
                "action": "polish",
                "selection": "",
                "instruction": "",
            },
        )
        assert resp.status_code == 200
        events = parse_sse_events(resp.data)
        assert any(e.get("type") == "error" for e in events)

    def test_custom_instruction_with_context(self, app_client):
        """自定义指令应带上选区和全文上下文"""
        resp = app_client.post(
            "/api/editor/ai/stream",
            json={
                "action": "custom_instruction",
                "selection": "AI技术",
                "instruction": "用更学术的方式描述",
                "full_text": "本篇论文探讨AI技术的发展趋势。",
            },
        )
        assert resp.status_code == 200
        events = parse_sse_events(resp.data)
        assert len(events) >= 1

    def test_ai_task_action_is_not_an_editor_stream_action(self, app_client):
        resp = app_client.post(
            "/api/editor/ai/stream",
            json={
                "action": "ai_task",
                "instruction": "整理当前文件",
                "session_id": "editor_demo",
                "file_name": "demo.docx",
                "file_type": "docx",
            },
        )

        assert resp.status_code == 400
        assert resp.get_json()["error"] == "Unsupported editor AI action: ai_task"

    def test_whitebox_task_stream_emits_new_contract(self, app_client, monkeypatch):
        from app.core.agent.file_task_runtime import FileTaskRuntime

        monkeypatch.setattr(
            FileTaskRuntime,
            "_call_model",
            lambda self, **kwargs: {"content": "已总结：hello world", "tool_calls": []},
        )

        resp = app_client.post(
            "/api/editor/ai/task-stream",
            json={
                "task": "总结这段选区",
                "selection": "hello world",
                "selection_source": "notes.txt",
                "session_id": "editor_demo",
            },
        )
        events = parse_sse_events(resp.get_data())
        event_types = [event.get("type") for event in events]

        assert resp.status_code == 200
        assert event_types[:5] == [
            "run.started",
            "task.classified",
            "plan.checked",
            "plan.created",
            "step.started",
        ]
        assert "check.started" in event_types
        assert "check.finished" in event_types
        assert event_types[-1] == "run.finished"
        assert [event.get("seq") for event in events] == list(range(1, len(events) + 1))
        assert events[0]["payload"]["mode"] == "whitebox_v1"

    def test_whitebox_task_stream_preserves_runtime_metadata_and_followup_context(
        self, app_client, monkeypatch, tmp_path
    ):
        from app.core.agent.file_task_runtime import FileTaskRuntime

        monkeypatch.setenv(
            "KOTO_FILE_TASK_FOLLOWUP_PATH", str(tmp_path / "followups.json")
        )
        monkeypatch.setattr(
            FileTaskRuntime,
            "_call_model",
            lambda self, **kwargs: {
                "content": "当前任务需要新的 Koto 工具。",
                "tool_calls": [],
                "tool_gap": {
                    "summary": "当前缺少读取 CAD 文件的 Koto 原生工具。",
                    "missing_capability": "read_cad_file",
                    "why_missing": "allowlist 中没有可读取 dwg 的工具。",
                    "suggested_next_step": "先实现只读 CAD 解析工具，再决定写入方案。",
                    "proposed_tool": {
                        "name": "read_cad_file",
                        "description": "解析 DWG/DXF 为可检索的结构化文本。",
                        "parameters": {"type": "object"},
                    },
                },
            },
        )

        resp = app_client.post(
            "/api/editor/ai/task-stream",
            json={
                "task": "修改 CAD 文件并导出总结",
                "session_id": "editor_demo",
                "target_path": "drawing.dwg",
            },
        )

        events = parse_sse_events(resp.get_data())
        tool_missing = next(
            event for event in events if event.get("type") == "tool.missing"
        )
        check_finished = next(
            event for event in events if event.get("type") == "check.finished"
        )
        run_finished = next(
            event for event in events if event.get("type") == "run.finished"
        )

        assert resp.status_code == 200
        assert tool_missing["payload"]["runtime"]["execution_path"] == "native"
        assert tool_missing["payload"]["runtime"]["planner"]["backend"] == "native"
        assert (
            tool_missing["payload"]["next_action_artifact"]["runtime_context"][
                "execution_path"
            ]
            == "native"
        )
        assert (
            tool_missing["payload"]["next_action_artifact"]["missing_capability"]
            == "read_cad_file"
        )
        assert check_finished["payload"]["runtime"]["terminal_status"] == "tool_gap"
        assert run_finished["payload"]["runtime"]["execution_path"] == "native"
        assert run_finished["payload"]["runtime"]["terminal_status"] == "tool_gap"

    def test_whitebox_task_stream_requires_task(self, app_client):
        resp = app_client.post(
            "/api/editor/ai/task-stream", json={"selection": "hello"}
        )
        assert resp.status_code == 400
        assert resp.get_json()["error"] == "Missing 'task' parameter"

    def test_whitebox_task_stream_routes_pdf_docx_translation_review_to_doc_annotate_bridge(
        self, app_client, monkeypatch, tmp_path
    ):
        from app.core.agent.file_task_contract import FileTaskLedger
        from app.core.agent.file_task_runtime import FileTaskRuntime
        import app.core.agent.file_task_doc_annotate_bridge as bridge

        pdf_path = tmp_path / "source.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        docx_path = tmp_path / "translation.docx"
        docx_path.write_bytes(b"PK\x03\x04")

        captured = {}

        def fake_stream(request, **kwargs):
            captured["request"] = request
            ledger = FileTaskLedger(request.run_id)
            yield ledger.event(
                "run.started",
                {
                    "task": request.task,
                    "mode": "doc_annotate_bridge",
                    "target_path": str(docx_path),
                    "source_path": str(pdf_path),
                },
                step_id="run",
            )
            yield ledger.event(
                "plan.created",
                {
                    "summary": "识别为 PDF 原文 + DOCX 译稿审校任务，改走 Word 修订写回流程。",
                    "steps": [{"id": "reference", "title": "整理 PDF 原文窗口"}],
                },
                step_id="plan",
            )
            yield ledger.event(
                "step_progress",
                {
                    "detail": "已写入 2 条审校修订",
                    "progress": 84,
                    "level": "progress",
                    "file_updated": True,
                    "path": str(docx_path),
                    "file_path": str(docx_path),
                    "supported": True,
                },
                step_id="review",
            )
            yield ledger.event(
                "file.changed",
                {
                    "path": str(docx_path),
                    "file_path": str(docx_path),
                    "file_type": "docx",
                    "operation": "annotate_file",
                    "summary": f"已将 2 条修订写回 {docx_path.name}。",
                    "annotations_added": 2,
                    "source_path": str(pdf_path),
                },
                step_id="write",
            )
            yield ledger.event(
                "run.finished",
                {
                    "summary": f"已更新 {docx_path.name}。",
                    "completed_task": True,
                    "mode": "doc_annotate_bridge",
                },
                step_id="run",
            )

        def unexpected_build_tool_gateway(self, request, context_files):
            raise AssertionError(
                "generic FileTaskRuntime loop should not execute for doc annotate bridge requests"
            )

        def unexpected_call_model(self, **kwargs):
            raise AssertionError(
                "generic model loop should not execute for doc annotate bridge requests"
            )

        monkeypatch.setattr(bridge, "stream_request", fake_stream)
        monkeypatch.setattr(
            FileTaskRuntime, "_build_tool_gateway", unexpected_build_tool_gateway
        )
        monkeypatch.setattr(FileTaskRuntime, "_call_model", unexpected_call_model)

        resp = app_client.post(
            "/api/editor/ai/task-stream",
            json={
                "task": "PDF是原文，docx文件是现有翻译稿。根据原文内容，润色翻译稿的语句，重新考虑中文翻译不当的部分或有复议可能性的用词或者替换成别的词的地方都标注出来。学术化翻译和中文学界常用词不对应的地方也标注出来，谨遵原著，不要有任何删减和添加。由于文件比较大内容比较多，我建议你将整个任务拆分成多个分段来处理，以保证最终结果的质量和任务可执行性",
                "session_id": "workspace_demo",
                "target_path": str(docx_path),
                "files": [
                    {"path": str(pdf_path), "name": "source.pdf", "type": "pdf"},
                    {
                        "path": str(docx_path),
                        "name": "translation.docx",
                        "type": "docx",
                        "target": True,
                    },
                ],
            },
        )

        events = parse_sse_events(resp.get_data())

        assert resp.status_code == 200
        assert captured["request"].target_path == str(docx_path)
        run_started = next(
            event for event in events if event.get("type") == "run.started"
        )
        assert run_started["payload"]["mode"] == "doc_annotate_bridge"
        assert any(event["type"] == "plan.created" for event in events)
        assert any(
            event["type"] == "step_progress" and event["payload"].get("file_updated")
            for event in events
        )
        assert any(
            event["type"] == "file.changed"
            and event["payload"].get("path") == str(docx_path)
            for event in events
        )
        assert events[-1]["payload"]["completed_task"] is True

    def test_whitebox_task_stream_routes_single_docx_annotation_to_doc_annotate_bridge(
        self, app_client, monkeypatch, tmp_path
    ):
        from app.core.agent.file_task_contract import FileTaskLedger
        from app.core.agent.file_task_runtime import FileTaskRuntime
        import app.core.agent.file_task_doc_annotate_bridge as bridge

        docx_path = tmp_path / "interview.docx"
        docx_path.write_bytes(b"PK\x03\x04")

        captured = {}

        def fake_stream(request, **kwargs):
            captured["request"] = request
            ledger = FileTaskLedger(request.run_id)
            yield ledger.event(
                "run.started",
                {
                    "task": request.task,
                    "mode": "doc_annotate_bridge",
                    "target_path": str(docx_path),
                },
                step_id="run",
            )
            yield ledger.event(
                "step_progress",
                {
                    "detail": "已写入 1/2 条修订",
                    "progress": 80,
                    "level": "progress",
                    "file_updated": True,
                    "path": str(docx_path),
                    "file_path": str(docx_path),
                    "supported": True,
                },
                step_id="review",
            )
            yield ledger.event(
                "file.changed",
                {
                    "path": str(docx_path),
                    "file_path": str(docx_path),
                    "file_type": "docx",
                    "operation": "annotate_file",
                    "summary": f"已将 2 条修订写回 {docx_path.name}。",
                    "annotations_added": 2,
                },
                step_id="write",
            )
            yield ledger.event(
                "run.finished",
                {
                    "summary": f"已更新 {docx_path.name}。",
                    "completed_task": True,
                    "mode": "doc_annotate_bridge",
                },
                step_id="run",
            )

        def unexpected_build_tool_gateway(self, request, context_files):
            raise AssertionError(
                "generic FileTaskRuntime loop should not execute for doc annotate bridge requests"
            )

        def unexpected_call_model(self, **kwargs):
            raise AssertionError(
                "generic model loop should not execute for doc annotate bridge requests"
            )

        monkeypatch.setattr(bridge, "stream_request", fake_stream)
        monkeypatch.setattr(
            FileTaskRuntime, "_build_tool_gateway", unexpected_build_tool_gateway
        )
        monkeypatch.setattr(FileTaskRuntime, "_call_model", unexpected_call_model)

        resp = app_client.post(
            "/api/editor/ai/task-stream",
            json={
                "task": "将你觉得写得不好的地方批注出来",
                "session_id": "workspace_demo",
                "target_path": str(docx_path),
                "files": [
                    {
                        "path": str(docx_path),
                        "name": "interview.docx",
                        "type": "docx",
                        "target": True,
                    },
                ],
            },
        )

        events = parse_sse_events(resp.get_data())

        assert resp.status_code == 200
        assert captured["request"].target_path == str(docx_path)
        run_started = next(
            event for event in events if event.get("type") == "run.started"
        )
        assert run_started["payload"]["mode"] == "doc_annotate_bridge"
        assert any(
            event["type"] == "step_progress" and event["payload"].get("file_updated")
            for event in events
        )
        assert any(
            event["type"] == "file.changed"
            and event["payload"].get("path") == str(docx_path)
            for event in events
        )
        assert events[-1]["payload"]["completed_task"] is True

    def test_whitebox_task_stream_routes_pdf_docx_translation_review_to_doc_annotate_bridge_in_local_mode(
        self, app_client, monkeypatch, tmp_path
    ):
        from app.core.agent.file_task_contract import FileTaskLedger
        from app.core.agent.file_task_runtime import FileTaskRuntime
        import app.core.agent.file_task_doc_annotate_bridge as bridge

        pdf_path = tmp_path / "source.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        docx_path = tmp_path / "translation.docx"
        docx_path.write_bytes(b"PK\x03\x04")

        captured = {"request": None}

        def fake_stream(request, **kwargs):
            captured["request"] = request
            ledger = FileTaskLedger(request.run_id)
            yield ledger.event(
                "run.started",
                {
                    "task": request.task,
                    "mode": "doc_annotate_bridge",
                    "target_path": str(docx_path),
                    "source_path": str(pdf_path),
                    "model_mode": request.model_mode,
                },
                step_id="run",
            )
            yield ledger.event(
                "run.finished",
                {
                    "summary": f"已更新 {docx_path.name}。",
                    "completed_task": True,
                    "mode": "doc_annotate_bridge",
                },
                step_id="run",
            )

        def unexpected_build_tool_gateway(self, request, context_files):
            raise AssertionError(
                "generic FileTaskRuntime loop should not execute for doc annotate bridge requests"
            )

        def unexpected_call_model(self, **kwargs):
            raise AssertionError(
                "generic model loop should not execute for doc annotate bridge requests"
            )

        monkeypatch.setattr(bridge, "stream_request", fake_stream)
        monkeypatch.setattr(
            FileTaskRuntime, "_build_tool_gateway", unexpected_build_tool_gateway
        )
        monkeypatch.setattr(FileTaskRuntime, "_call_model", unexpected_call_model)

        resp = app_client.post(
            "/api/editor/ai/task-stream",
            json={
                "task": "PDF是原文，docx文件是现有翻译稿。根据原文内容，润色翻译稿的语句，重新考虑中文翻译不当的部分或有复议可能性的用词或者替换成别的词的地方都标注出来。学术化翻译和中文学界常用词不对应的地方也标注出来，谨遵原著，不要有任何删减和添加。由于文件比较大内容比较多，我建议你将整个任务拆分成多个分段来处理，以保证最终结果的质量和任务可执行性",
                "session_id": "workspace_demo",
                "model_mode": "local",
                "model_id": "local",
                "target_path": str(docx_path),
                "files": [
                    {"path": str(pdf_path), "name": "source.pdf", "type": "pdf"},
                    {
                        "path": str(docx_path),
                        "name": "translation.docx",
                        "type": "docx",
                        "target": True,
                    },
                ],
            },
        )

        events = parse_sse_events(resp.get_data())

        assert resp.status_code == 200
        assert captured["request"].model_mode == "local"
        run_started = next(
            event for event in events if event.get("type") == "run.started"
        )
        assert run_started["payload"]["mode"] == "doc_annotate_bridge"
        assert events[-1]["payload"]["completed_task"] is True

    def test_whitebox_task_stream_followup_feedback_bypasses_doc_annotate_bridge(
        self, app_client, monkeypatch, tmp_path
    ):
        from app.core.agent.file_task_contract import FileTaskLedger
        from app.core.agent.file_task_runtime import FileTaskRuntime
        import app.core.agent.file_task_doc_annotate_bridge as bridge

        pdf_path = tmp_path / "source.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        docx_path = tmp_path / "translation.docx"
        docx_path.write_bytes(b"PK\x03\x04")

        captured = {}

        def fail_stream(request, **kwargs):
            raise AssertionError(
                "doc annotate bridge should not handle review-last-task follow-up feedback"
            )

        def fake_run(self, request):
            captured["request"] = request
            ledger = FileTaskLedger(request.run_id)
            yield ledger.event(
                "run.started",
                {"task": request.task, "mode": "whitebox_v1"},
                step_id="run",
            )
            yield ledger.event(
                "run.finished",
                {
                    "summary": "先反馈上一轮结果。",
                    "completed_task": True,
                    "mode": "whitebox_v1",
                },
                step_id="run",
            )

        monkeypatch.setattr(bridge, "stream_request", fail_stream)
        monkeypatch.setattr(FileTaskRuntime, "run", fake_run)

        resp = app_client.post(
            "/api/editor/ai/task-stream",
            json={
                "task": "为什么你这次审校结果这么处理？",
                "session_id": "workspace_demo",
                "target_path": str(docx_path),
                "files": [
                    {"path": str(pdf_path), "name": "source.pdf", "type": "pdf"},
                    {
                        "path": str(docx_path),
                        "name": "translation.docx",
                        "type": "docx",
                        "target": True,
                    },
                ],
                "options": {
                    "followup_context": {
                        "kind": "review_last_task",
                        "user_feedback": "为什么你这次审校结果这么处理？",
                        "previous_task_summary": "已生成带批注的审校稿 translation_revised.docx",
                    }
                },
            },
        )

        events = parse_sse_events(resp.get_data())

        assert resp.status_code == 200
        assert (
            captured["request"].options["followup_context"]["kind"]
            == "review_last_task"
        )
        run_started = next(
            event for event in events if event.get("type") == "run.started"
        )
        assert run_started["payload"]["mode"] == "whitebox_v1"
        assert events[-1]["payload"]["completed_task"] is True

    def test_whitebox_task_stream_followup_improve_routes_back_to_doc_annotate_bridge(
        self, app_client, monkeypatch, tmp_path
    ):
        from app.core.agent.file_task_contract import FileTaskLedger
        from app.core.agent.file_task_runtime import FileTaskRuntime
        import app.core.agent.file_task_doc_annotate_bridge as bridge

        pdf_path = tmp_path / "source.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        docx_path = tmp_path / "translation.docx"
        docx_path.write_bytes(b"PK\x03\x04")

        captured = {"request": None}

        def fake_stream(request, **kwargs):
            captured["request"] = request
            ledger = FileTaskLedger(request.run_id)
            yield ledger.event(
                "run.started",
                {
                    "task": request.task,
                    "mode": "doc_annotate_bridge",
                    "target_path": str(docx_path),
                    "source_path": str(pdf_path),
                },
                step_id="run",
            )
            yield ledger.event(
                "step_progress",
                {
                    "detail": "已继续优化 1 处批注",
                    "progress": 82,
                    "level": "progress",
                    "file_updated": True,
                    "path": str(docx_path),
                    "file_path": str(docx_path),
                    "supported": True,
                },
                step_id="review",
            )
            yield ledger.event(
                "run.finished",
                {
                    "summary": f"已继续优化 {docx_path.name}。",
                    "completed_task": True,
                    "mode": "doc_annotate_bridge",
                },
                step_id="run",
            )

        def unexpected_build_tool_gateway(self, request, context_files):
            raise AssertionError(
                "generic FileTaskRuntime loop should not execute for doc annotate bridge requests"
            )

        def unexpected_call_model(self, **kwargs):
            raise AssertionError(
                "generic model loop should not execute for doc annotate bridge requests"
            )

        monkeypatch.setattr(bridge, "stream_request", fake_stream)
        monkeypatch.setattr(
            FileTaskRuntime, "_build_tool_gateway", unexpected_build_tool_gateway
        )
        monkeypatch.setattr(FileTaskRuntime, "_call_model", unexpected_call_model)

        resp = app_client.post(
            "/api/editor/ai/task-stream",
            json={
                "task": "请继续优化上一轮审校结果",
                "session_id": "workspace_demo",
                "target_path": str(docx_path),
                "files": [
                    {"path": str(pdf_path), "name": "source.pdf", "type": "pdf"},
                    {
                        "path": str(docx_path),
                        "name": "translation.docx",
                        "type": "docx",
                        "target": True,
                    },
                ],
                "options": {
                    "followup_context": {
                        "kind": "review_last_task",
                        "followup_action": "improve",
                        "user_feedback": "请继续优化上一轮审校结果",
                        "previous_task_request": "根据原文审校这个译稿",
                        "previous_task_mode": "doc_annotate_bridge",
                    }
                },
            },
        )

        events = parse_sse_events(resp.get_data())

        assert resp.status_code == 200
        assert (
            captured["request"].options["followup_context"]["followup_action"]
            == "improve"
        )
        run_started = next(
            event for event in events if event.get("type") == "run.started"
        )
        assert run_started["payload"]["mode"] == "doc_annotate_bridge"
        assert events[-1]["payload"]["completed_task"] is True

    def test_doc_annotate_bridge_forwards_review_and_write_progress(
        self, monkeypatch, tmp_path
    ):
        from app.core.agent.file_task_contract import FileTaskRequest
        import app.core.agent.file_task_doc_annotate_bridge as bridge
        import web.document_feedback as feedback_module

        pdf_path = tmp_path / "source.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        docx_path = tmp_path / "translation.docx"
        docx_path.write_bytes(b"PK\x03\x04")

        class FakeFeedback:
            default_model_id = "gemini-2.5-pro"

            def __init__(self, gemini_client=None):
                self.gemini_client = gemini_client

            def full_annotation_loop_streaming(self, *args, **kwargs):
                yield {
                    "stage": "reading",
                    "progress": 5,
                    "message": "📖 正在读取文档",
                    "detail": "解析Word文件结构",
                }
                yield {
                    "stage": "reading_complete",
                    "progress": 10,
                    "message": "✅ 文档读取完成",
                    "detail": "12 段，3000 字",
                }
                yield {
                    "stage": "analyzing",
                    "progress": 24,
                    "message": "🤖 正在分析文档...",
                    "detail": "第 1/5 段已完成，累计 2 条建议",
                }
                yield {
                    "stage": "warning",
                    "progress": 30,
                    "message": "⚠️ AI 分析未成功（1/5分段使用本地规则兜底）",
                    "detail": "API 错误: timeout",
                }
                yield {
                    "stage": "analysis_complete",
                    "progress": 50,
                    "message": "✅ 分析完成",
                    "detail": "找到 6 处修改",
                }
                yield {
                    "stage": "applying",
                    "progress": 78,
                    "message": "✏️ 正在写回 Word 修订",
                    "detail": "已写入 3/6 条修订",
                    "file_updated": True,
                    "path": str(docx_path),
                    "file_path": str(docx_path),
                    "supported": True,
                    "applied": 3,
                }
                yield {
                    "stage": "complete",
                    "progress": 100,
                    "result": {
                        "success": True,
                        "revised_file": str(docx_path),
                        "applied": 6,
                    },
                }

        monkeypatch.setattr(
            bridge,
            "_build_pdf_reference_windows",
            lambda path: (
                ["[Page 1]\nOriginal source text"],
                {"window_count": 1, "page_count": 1, "window_pages": 4},
            ),
        )
        monkeypatch.setattr(feedback_module, "DocumentFeedbackSystem", FakeFeedback)

        request = FileTaskRequest.from_mapping(
            {
                "task": "PDF是原文，docx文件是现有翻译稿。根据原文内容，润色翻译稿的语句，并标注可能需要复议的用词。",
                "target_path": str(docx_path),
                "files": [
                    {"path": str(pdf_path), "name": "source.pdf", "type": "pdf"},
                    {
                        "path": str(docx_path),
                        "name": "translation.docx",
                        "type": "docx",
                        "target": True,
                    },
                ],
            }
        )

        events = list(bridge.stream_request(request, workspace_root=str(tmp_path)))

        assert any(
            event.type == "step_progress"
            and event.step_id == "review"
            and "第 1/5 段已完成" in str(event.payload.get("detail") or "")
            for event in events
        )
        assert any(
            event.type == "step_progress"
            and event.step_id == "write"
            and "已写入 3/6 条修订" in str(event.payload.get("detail") or "")
            for event in events
        )
        assert any(
            event.type == "step_progress"
            and event.step_id == "write"
            and event.payload.get("file_updated") is True
            and event.payload.get("path") == str(docx_path)
            for event in events
        )
        file_changed = next(event for event in events if event.type == "file.changed")
        assert file_changed.payload["path"] == str(docx_path)
        assert file_changed.payload["file_path"] == str(docx_path)
        assert events[-1].type == "run.finished"
        assert events[-1].payload["completed_task"] is True

    def test_doc_annotate_bridge_handles_single_docx_annotation_requests(
        self, monkeypatch, tmp_path
    ):
        from app.core.agent.file_task_contract import FileTaskRequest
        import app.core.agent.file_task_doc_annotate_bridge as bridge
        import web.document_feedback as feedback_module

        docx_path = tmp_path / "interview.docx"
        docx_path.write_bytes(b"PK\x03\x04")

        class FakeFeedback:
            default_model_id = "gemini-2.5-pro"

            def __init__(self, gemini_client=None):
                self.gemini_client = gemini_client

            def full_annotation_loop_streaming(self, *args, **kwargs):
                yield {
                    "stage": "reading",
                    "progress": 5,
                    "message": "📖 正在读取文档",
                    "detail": "解析Word文件结构",
                }
                yield {
                    "stage": "reading_complete",
                    "progress": 10,
                    "message": "✅ 文档读取完成",
                    "detail": "21 段，553 字",
                }
                yield {
                    "stage": "analyzing",
                    "progress": 40,
                    "message": "🤖 正在分析文档...",
                    "detail": "已整理 10 条批注建议",
                }
                yield {
                    "stage": "analysis_complete",
                    "progress": 65,
                    "message": "✅ 分析完成",
                    "detail": "找到 10 处修改",
                }
                yield {
                    "stage": "applying",
                    "progress": 82,
                    "message": "✏️ 正在写回 Word 修订",
                    "detail": "已写入 6/10 条修订",
                    "file_updated": True,
                    "path": str(docx_path),
                    "file_path": str(docx_path),
                    "supported": True,
                    "applied": 6,
                }
                yield {
                    "stage": "complete",
                    "progress": 100,
                    "result": {
                        "success": True,
                        "revised_file": str(docx_path),
                        "applied": 10,
                    },
                }

        monkeypatch.setattr(feedback_module, "DocumentFeedbackSystem", FakeFeedback)

        request = FileTaskRequest.from_mapping(
            {
                "task": "将你觉得写得不好的地方批注出来",
                "target_path": str(docx_path),
                "files": [
                    {
                        "path": str(docx_path),
                        "name": "interview.docx",
                        "type": "docx",
                        "target": True,
                    },
                ],
            }
        )

        events = list(bridge.stream_request(request, workspace_root=str(tmp_path)))

        assert events[0].payload["mode"] == "doc_annotate_bridge"
        assert any(
            event.type == "tool.finished"
            and event.payload.get("tool_name") == "read_docx_content"
            and event.payload.get("path") == str(docx_path)
            for event in events
        )
        assert any(
            event.type == "step_progress"
            and event.step_id == "write"
            and "已写入 6/10 条修订" in str(event.payload.get("detail") or "")
            for event in events
        )
        file_changed = next(event for event in events if event.type == "file.changed")
        assert file_changed.payload["path"] == str(docx_path)
        assert file_changed.payload["file_path"] == str(docx_path)
        assert events[-1].type == "run.finished"
        assert events[-1].payload["completed_task"] is True

    def test_doc_annotate_bridge_merges_followup_improve_requirement(
        self, monkeypatch, tmp_path
    ):
        from app.core.agent.file_task_contract import FileTaskRequest, FileTaskFile
        import app.core.agent.file_task_doc_annotate_bridge as bridge
        import web.document_feedback as feedback_module

        pdf_path = tmp_path / "source.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        docx_path = tmp_path / "translation.docx"
        docx_path.write_bytes(b"PK\x03\x04")

        captured = {}

        class FakeFeedback:
            default_model_id = "gemini-2.5-pro"

            def __init__(self, gemini_client=None):
                self.gemini_client = gemini_client

            def full_annotation_loop_streaming(self, *args, **kwargs):
                captured["user_requirement"] = kwargs.get("user_requirement")
                yield {
                    "stage": "complete",
                    "result": {
                        "success": True,
                        "revised_file": str(docx_path),
                        "applied": 1,
                    },
                }

        monkeypatch.setattr(feedback_module, "DocumentFeedbackSystem", FakeFeedback)
        monkeypatch.setattr(
            bridge,
            "_build_pdf_reference_windows",
            lambda path: (
                ["[Page 1]\nOriginal source text"],
                {
                    "window_count": 1,
                    "page_count": 1,
                    "pages_with_text": 1,
                    "window_pages": 4,
                },
            ),
        )

        request = FileTaskRequest(
            task="请继续优化上一轮审校结果",
            run_id="doc_bridge_followup_improve",
            target_path=str(docx_path),
            files=[
                FileTaskFile(path=str(pdf_path), name="source.pdf", type="pdf"),
                FileTaskFile(
                    path=str(docx_path),
                    name="translation.docx",
                    type="docx",
                    target=True,
                ),
            ],
            options={
                "followup_context": {
                    "kind": "review_last_task",
                    "followup_action": "improve",
                    "user_feedback": "请继续优化上一轮审校结果",
                    "previous_task_request": "根据原文审校这个译稿，给出学术化批注",
                    "previous_task_mode": "doc_annotate_bridge",
                }
            },
        )

        list(
            bridge.stream_request(
                request, workspace_root=str(tmp_path), gemini_client=object()
            )
        )

        assert (
            "上一轮任务要求：根据原文审校这个译稿，给出学术化批注"
            in captured["user_requirement"]
        )
        assert "当前追加反馈：请继续优化上一轮审校结果" in captured["user_requirement"]

    def test_doc_annotate_bridge_emits_confirmed_batch_plan_for_large_files(
        self, monkeypatch, tmp_path
    ):
        from app.core.agent.file_task_contract import FileTaskRequest
        import app.core.agent.file_task_doc_annotate_bridge as bridge
        import web.document_feedback as feedback_module

        monkeypatch.setenv("KOTO_DOC_REVIEW_CLOUD_BATCH_TARGET_MINUTES", "6")
        monkeypatch.setenv("KOTO_DOC_REVIEW_CLOUD_BATCH_CHARS_PER_MINUTE", "4000")

        pdf_path = tmp_path / "source.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        docx_path = tmp_path / "translation.docx"
        docx_path.write_bytes(b"PK\x03\x04")
        revised_path = tmp_path / "translation_revised.docx"

        class FakeFeedback:
            default_model_id = "gemini-2.5-pro"
            client = object()

            def __init__(self, gemini_client=None):
                self.gemini_client = gemini_client

            def full_annotation_loop_streaming(self, *args, **kwargs):
                yield {"stage": "reading_complete", "detail": "539 段，100099 字"}
                yield {"stage": "analysis_complete", "detail": "找到 42 处修改"}
                yield {
                    "stage": "complete",
                    "result": {
                        "success": True,
                        "revised_file": str(revised_path),
                        "applied": 42,
                    },
                }

            def _split_into_chunks_by_paragraphs(self, content, chunk_size):
                return [f"chunk-{idx}" for idx in range(27)]

        monkeypatch.setattr(
            bridge,
            "_build_pdf_reference_windows",
            lambda path: (
                [f"[Page {idx}]\ntext" for idx in range(1, 101)],
                {
                    "window_count": 100,
                    "page_count": 417,
                    "pages_with_text": 399,
                    "window_pages": 4,
                },
            ),
        )
        monkeypatch.setattr(
            bridge,
            "_inspect_docx_review_workload",
            lambda target_docx, feedback: {
                "paragraph_count": 539,
                "table_count": 2,
                "formatted_chars": 102901,
                "content_chars": 102801,
                "chunk_size": 4000,
                "chunk_count": 27,
            },
        )
        monkeypatch.setattr(feedback_module, "DocumentFeedbackSystem", FakeFeedback)

        request = FileTaskRequest.from_mapping(
            {
                "task": "PDF是原文，docx文件是现有翻译稿。根据原文内容，润色翻译稿的语句，并拆成多个分段来处理。",
                "target_path": str(docx_path),
                "files": [
                    {"path": str(pdf_path), "name": "source.pdf", "type": "pdf"},
                    {
                        "path": str(docx_path),
                        "name": "translation.docx",
                        "type": "docx",
                        "target": True,
                    },
                ],
            }
        )

        events = list(bridge.stream_request(request, workspace_root=str(tmp_path)))
        confirmed = next(event for event in events if event.type == "plan.confirmed")

        assert confirmed.step_id == "review"
        assert "按 5 批执行" in confirmed.payload["summary"]
        assert "约 2.4 万字/批" in confirmed.payload["summary"]
        assert confirmed.payload["steps"][0]["title"] == "第 1 批：分段 1-6"
        assert confirmed.payload["steps"][4]["title"] == "第 5 批：分段 25-27"
        assert confirmed.payload["steps"][-2]["title"] == "汇总并写回修订"
        assert "417 页" in confirmed.payload["note"]
        assert "539 段" in confirmed.payload["note"]
        assert "目标单批约 6 分钟" in confirmed.payload["note"]

    def test_large_translation_review_uses_finer_batches_in_local_mode(
        self, monkeypatch, tmp_path
    ):
        from app.core.agent.file_task_contract import FileTaskRequest
        import app.core.agent.file_task_doc_annotate_bridge as bridge
        import web.document_feedback as feedback_module

        monkeypatch.setenv("KOTO_DOC_REVIEW_LOCAL_BATCH_TARGET_MINUTES", "5")
        monkeypatch.setenv("KOTO_DOC_REVIEW_LOCAL_BATCH_CHARS_PER_MINUTE", "2400")

        pdf_path = tmp_path / "source.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        docx_path = tmp_path / "translation.docx"
        docx_path.write_bytes(b"PK\x03\x04")

        revised_path = tmp_path / "translation_revised.docx"

        class FakeFeedback:
            default_model_id = "gemini-2.5-pro"
            client = object()

            def __init__(self, gemini_client=None):
                self.gemini_client = gemini_client

            def full_annotation_loop_streaming(self, *args, **kwargs):
                yield {"stage": "reading_complete", "detail": "539 段，100099 字"}
                yield {"stage": "analysis_complete", "detail": "找到 42 处修改"}
                yield {
                    "stage": "complete",
                    "result": {
                        "success": True,
                        "revised_file": str(revised_path),
                        "applied": 42,
                    },
                }

            def _split_into_chunks_by_paragraphs(self, content, chunk_size):
                return [f"chunk-{idx}" for idx in range(27)]

        monkeypatch.setattr(
            bridge,
            "_build_pdf_reference_windows",
            lambda path: (
                [f"[Page {idx}]\ntext" for idx in range(1, 101)],
                {
                    "window_count": 100,
                    "page_count": 417,
                    "pages_with_text": 399,
                    "window_pages": 4,
                },
            ),
        )
        monkeypatch.setattr(
            bridge,
            "_inspect_docx_review_workload",
            lambda target_docx, feedback: {
                "paragraph_count": 539,
                "table_count": 2,
                "formatted_chars": 102901,
                "content_chars": 102801,
                "chunk_size": 4000,
                "chunk_count": 27,
            },
        )
        monkeypatch.setattr(feedback_module, "DocumentFeedbackSystem", FakeFeedback)

        request = FileTaskRequest.from_mapping(
            {
                "task": "PDF是原文，docx文件是现有翻译稿。根据原文内容，润色翻译稿的语句，并拆成多个分段来处理。",
                "model_mode": "local",
                "model_id": "qwen3.5:9b",
                "target_path": str(docx_path),
                "files": [
                    {"path": str(pdf_path), "name": "source.pdf", "type": "pdf"},
                    {
                        "path": str(docx_path),
                        "name": "translation.docx",
                        "type": "docx",
                        "target": True,
                    },
                ],
            }
        )

        events = list(bridge.stream_request(request, workspace_root=str(tmp_path)))
        confirmed = next(event for event in events if event.type == "plan.confirmed")

        assert confirmed.step_id == "review"
        assert "按 9 批执行" in confirmed.payload["summary"]
        assert confirmed.payload["steps"][0]["title"] == "第 1 批：分段 1-3"
        assert confirmed.payload["steps"][8]["title"] == "第 9 批：分段 25-27"
        assert "约 1.2 万字/批" in confirmed.payload["summary"]
        assert "目标单批约 5 分钟" in confirmed.payload["note"]

    def test_local_docx_review_bridge_uses_local_model_identity(
        self, monkeypatch, tmp_path
    ):
        from app.core.agent.file_task_contract import FileTaskRequest
        import app.core.agent.file_task_doc_annotate_bridge as bridge
        import web.document_feedback as feedback_module

        docx_path = tmp_path / "humanise!.docx"
        docx_path.write_bytes(b"PK\x03\x04")

        captured = {}

        class FakeFeedback:
            client = object()

            def __init__(self, gemini_client=None, default_model_id="gemini-2.5-pro"):
                captured["default_model_id"] = default_model_id
                self.gemini_client = gemini_client
                self.default_model_id = default_model_id

            def full_annotation_loop_streaming(self, *args, **kwargs):
                captured["stream_model_id"] = kwargs.get("model_id")
                yield {"stage": "reading_complete", "detail": "12 段，6300 字"}
                yield {
                    "stage": "analyzing",
                    "progress": 15,
                    "message": "分析中",
                    "detail": f"使用 AI({kwargs.get('model_id')}) 检查 12 段文本",
                }
                yield {"stage": "analysis_complete", "detail": "找到 8 处修改"}
                yield {
                    "stage": "complete",
                    "result": {
                        "success": True,
                        "revised_file": str(docx_path),
                        "applied": 8,
                    },
                }

        monkeypatch.setattr(feedback_module, "DocumentFeedbackSystem", FakeFeedback)

        request = FileTaskRequest.from_mapping(
            {
                "task": "将文中不通顺和较为AI表达的地方标出来，并提出修改意见",
                "model_mode": "local",
                "model_id": "",
                "target_path": str(docx_path),
                "options": {"local_model": "qwen3.5:9b"},
                "files": [
                    {
                        "path": str(docx_path),
                        "name": docx_path.name,
                        "type": "docx",
                        "target": True,
                    },
                ],
            }
        )

        events = list(bridge.stream_request(request, workspace_root=str(tmp_path)))

        review_progress = next(
            event for event in events if event.type == "step_progress"
        )

        assert captured["default_model_id"] == "qwen3.5:9b"
        assert captured["stream_model_id"] == "qwen3.5:9b"
        assert "qwen3.5:9b" in review_progress.payload["detail"]

    def test_large_translation_review_cloud_mode_batches_by_char_budget(
        self, monkeypatch, tmp_path
    ):
        from app.core.agent.file_task_contract import FileTaskRequest
        import app.core.agent.file_task_doc_annotate_bridge as bridge
        import web.document_feedback as feedback_module

        monkeypatch.setenv("KOTO_DOC_REVIEW_CLOUD_BATCH_TARGET_MINUTES", "6")
        monkeypatch.setenv("KOTO_DOC_REVIEW_CLOUD_BATCH_CHARS_PER_MINUTE", "4000")

        pdf_path = tmp_path / "source.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        docx_path = tmp_path / "translation.docx"
        docx_path.write_bytes(b"PK\x03\x04")
        revised_path = tmp_path / "translation_revised.docx"

        class FakeFeedback:
            default_model_id = "gemini-2.5-pro"
            client = object()

            def __init__(self, gemini_client=None):
                self.gemini_client = gemini_client

            def full_annotation_loop_streaming(self, *args, **kwargs):
                yield {"stage": "reading_complete", "detail": "160 段，104000 字"}
                yield {"stage": "analysis_complete", "detail": "找到 12 处修改"}
                yield {
                    "stage": "complete",
                    "result": {
                        "success": True,
                        "revised_file": str(revised_path),
                        "applied": 12,
                    },
                }

        monkeypatch.setattr(
            bridge,
            "_build_pdf_reference_windows",
            lambda path: (
                [f"[Page {idx}]\ntext" for idx in range(1, 21)],
                {
                    "window_count": 20,
                    "page_count": 90,
                    "pages_with_text": 84,
                    "window_pages": 4,
                },
            ),
        )
        monkeypatch.setattr(
            bridge,
            "_inspect_docx_review_workload",
            lambda target_docx, feedback: {
                "paragraph_count": 160,
                "table_count": 0,
                "formatted_chars": 104120,
                "content_chars": 104000,
                "chunk_size": 4000,
                "chunk_count": 8,
                "chunk_char_counts": [
                    20000,
                    4000,
                    4000,
                    20000,
                    4000,
                    4000,
                    24000,
                    24000,
                ],
            },
        )
        monkeypatch.setattr(feedback_module, "DocumentFeedbackSystem", FakeFeedback)

        request = FileTaskRequest.from_mapping(
            {
                "task": "PDF是原文，docx文件是现有翻译稿。根据原文内容，润色翻译稿的语句，并拆成多个分段来处理。",
                "target_path": str(docx_path),
                "files": [
                    {"path": str(pdf_path), "name": "source.pdf", "type": "pdf"},
                    {
                        "path": str(docx_path),
                        "name": "translation.docx",
                        "type": "docx",
                        "target": True,
                    },
                ],
            }
        )

        events = list(bridge.stream_request(request, workspace_root=str(tmp_path)))
        confirmed = next(event for event in events if event.type == "plan.confirmed")

        assert confirmed.step_id == "review"
        assert "按 5 批执行" in confirmed.payload["summary"]
        assert "约 2.4 万字/批" in confirmed.payload["summary"]
        assert confirmed.payload["steps"][0]["title"] == "第 1 批：分段 1-2"
        assert confirmed.payload["steps"][1]["title"] == "第 2 批：分段 3-4"
        assert confirmed.payload["steps"][2]["title"] == "第 3 批：分段 5-6"
        assert confirmed.payload["steps"][3]["title"] == "第 4 批：分段 7-7"
        assert confirmed.payload["steps"][4]["title"] == "第 5 批：分段 8-8"
        assert "预计用时约 6 分钟" in confirmed.payload["steps"][0]["description"]
        assert "目标单批约 6 分钟" in confirmed.payload["note"]

    def test_build_pdf_reference_windows_keeps_all_windows_by_default(
        self, monkeypatch
    ):
        import app.core.agent.file_task_doc_annotate_bridge as bridge

        fake_pages = [{"page": idx, "text": f"Page {idx} text"} for idx in range(1, 10)]

        with patch(
            "app.core.file.file_parser.parse_pdf",
            return_value={"page_count": 9, "pages": fake_pages, "text": ""},
        ):
            windows, meta = bridge._build_pdf_reference_windows(
                "dummy.pdf", window_pages=4, per_window_chars=1000
            )

        assert len(windows) == 3
        assert meta["page_count"] == 9
        assert meta["pages_with_text"] == 9
        assert meta["window_count"] == 3
        assert "[Page 9]" in windows[-1]

    def test_whitebox_task_stream_normalizes_local_model_config(
        self, app_client, monkeypatch
    ):
        from app.core.agent.file_task_contract import FileTaskEvent
        from app.core.agent.file_task_runtime import FileTaskRuntime

        captured = {}

        def fake_run(self, request):
            captured["request"] = request
            yield FileTaskEvent(
                type="run.finished",
                run_id=request.run_id,
                seq=1,
                payload={"summary": "ok", "completed_task": True},
            )

        monkeypatch.setattr(FileTaskRuntime, "run", fake_run)

        with patch("web.app._get_configured_local_model_id", return_value="qwen3.5:9b"):
            resp = app_client.post(
                "/api/editor/ai/task-stream",
                json={
                    "task": "总结当前文件",
                    "model_mode": "local",
                    "model_id": "gemini-3-flash-preview",
                },
            )

        request_payload = captured["request"]
        assert resp.status_code == 200
        assert request_payload.model_mode == "local"
        assert request_payload.model_id == ""
        assert request_payload.options["local_model"] == "qwen3.5:9b"

    def test_whitebox_task_stream_keeps_finished_run_runtime_only(
        self, app_client, monkeypatch
    ):
        from app.core.agent.file_task_contract import FileTaskEvent
        from app.core.agent.file_task_runtime import FileTaskRuntime
        import web.app as web_app_module

        calls = {"save": 0, "extract": 0}

        def fake_run(self, request):
            yield FileTaskEvent(
                type="file.changed",
                run_id=request.run_id,
                seq=1,
                step_id="execute",
                payload={
                    "path": "report.docx",
                    "operation": "write_docx_content",
                    "summary": "已写入",
                },
            )
            yield FileTaskEvent(
                type="run.finished",
                run_id=request.run_id,
                seq=2,
                payload={"summary": "已完成文件任务", "completed_task": True},
            )

        def fake_append(*args, **kwargs):
            calls["save"] += 1
            return []

        def fake_memory(*args, **kwargs):
            calls["extract"] += 1

        monkeypatch.setattr(FileTaskRuntime, "run", fake_run)
        monkeypatch.setattr(
            web_app_module.session_manager, "append_and_save", fake_append
        )
        monkeypatch.setattr(web_app_module, "_start_memory_extraction", fake_memory)

        resp = app_client.post(
            "/api/editor/ai/task-stream",
            json={"task": "把总结写入当前文件", "session_id": "workspace_demo"},
        )
        events = parse_sse_events(resp.get_data())

        assert resp.status_code == 200
        assert events[-1]["type"] == "run.finished"
        assert not any(event.get("type") == "memory.loaded" for event in events)
        assert calls == {"save": 0, "extract": 0}

    def test_whitebox_task_stream_uses_request_history_only(
        self, app_client, monkeypatch
    ):
        from app.core.agent.file_task_contract import FileTaskEvent
        from app.core.agent.file_task_runtime import FileTaskRuntime
        import web.app as web_app_module

        captured = {"loaded": []}

        def fake_load(filename):
            captured["loaded"].append(filename)
            return [
                {"role": "user", "parts": ["上一次任务"]},
                {"role": "model", "parts": ["上一次结果"]},
            ]

        def fake_run(self, request):
            captured["request"] = request
            yield FileTaskEvent(
                type="run.finished",
                run_id=request.run_id,
                seq=1,
                payload={"summary": "本轮完成", "completed_task": True},
            )

        monkeypatch.setattr(web_app_module.session_manager, "load", fake_load)
        monkeypatch.setattr(FileTaskRuntime, "run", fake_run)

        resp = app_client.post(
            "/api/editor/ai/task-stream",
            json={
                "task": "继续处理当前文件",
                "session_id": "workspace_demo",
                "history": [
                    {"role": "user", "content": "前端短期任务"},
                    {"role": "assistant", "content": "前端短期结果"},
                    {"role": "user", "content": "继续处理当前文件"},
                ],
            },
        )
        resp.get_data()

        request_payload = captured["request"]
        assert resp.status_code == 200
        assert captured["loaded"] == []
        assert request_payload.history == [
            {"role": "user", "content": "前端短期任务"},
            {"role": "assistant", "content": "前端短期结果"},
            {"role": "user", "content": "继续处理当前文件"},
        ]

    def test_whitebox_task_stream_does_not_inject_memory_router_context(
        self, app_client, monkeypatch
    ):
        from app.core.agent.file_task_contract import FileTaskEvent
        from app.core.agent.file_task_runtime import FileTaskRuntime
        import web.app as web_app_module

        captured = {}

        def fake_run(self, request):
            captured["request"] = request
            yield FileTaskEvent(
                type="run.finished",
                run_id=request.run_id,
                seq=1,
                payload={"summary": "完成", "completed_task": True},
            )

        monkeypatch.setattr(FileTaskRuntime, "run", fake_run)
        monkeypatch.setattr(
            web_app_module,
            "get_memory_manager",
            lambda: (_ for _ in ()).throw(
                AssertionError("memory manager should not be used")
            ),
        )

        resp = app_client.post(
            "/api/editor/ai/task-stream",
            json={
                "task": "整理当前 Word 文档",
                "session_id": "workspace_demo",
                "file_name": "report.docx",
            },
        )
        events = parse_sse_events(resp.get_data())

        request_payload = captured["request"]
        assert resp.status_code == 200
        assert not any(event.get("type") == "memory.loaded" for event in events)
        assert "memory_context" not in (request_payload.options or {})

    def test_editor_ai_history_is_empty_for_runtime_only_conversations(
        self, app_client, monkeypatch
    ):
        import web.app as web_app_module

        loaded = []

        def fake_load_full(filename):
            loaded.append(filename)
            return []

        monkeypatch.setattr(web_app_module.session_manager, "load_full", fake_load_full)

        resp = app_client.get(
            "/api/editor/ai/history?session_id=workspace_demo&doc_id=demo"
        )
        data = resp.get_json()

        assert resp.status_code == 200
        assert loaded == []
        assert data == {"history": [], "session_id": ""}

    def test_file_task_runtime_includes_memory_context_in_task_prompt(self):
        from app.core.agent.file_task_contract import FileTaskRequest
        from app.core.agent.file_task_runtime import FileTaskRuntime

        runtime = FileTaskRuntime(tool_executor=lambda name, args: "")
        request_payload = FileTaskRequest(
            task="总结文档",
            options={"memory_context": "用户偏好：保留原文格式。"},
        )

        messages = runtime._build_messages(request_payload, [], [])

        assert "memory_context" in messages[-1]["content"]
        assert "保留原文格式" in messages[-1]["content"]

    def test_workspace_open_file_returns_temp_path(self, app_client):
        resp = app_client.post(
            "/api/v1/workspace/open_file",
            data={"file": (io.BytesIO("hello world".encode("utf-8")), "notes.txt")},
            content_type="multipart/form-data",
        )

        assert resp.status_code == 200
        data = resp.get_json()
        assert data["temp_path"].startswith("tmp/")
        assert data["temp_path"].endswith(".txt")

    def test_main_stream_forwards_plan_and_step_events(self, app_client):
        """主 editor_ai_stream 应转发 KotoAgentLoop 的 plan/step 事件。"""
        from app.core.agent.lifecycle import (
            evt_plan,
            evt_step_done,
            evt_step_progress,
            evt_step_start,
            evt_task_complete,
        )

        def fake_run(self, request):
            yield evt_plan([{"id": "understand", "description": "理解需求"}])
            yield evt_step_start("understand", "理解需求")
            yield evt_step_progress("understand", "正在分析上下文…")
            yield evt_step_done("understand", "理解需求完成")
            yield evt_task_complete(result="处理完成")

        with patch("app.core.agent.agent_loop.KotoAgentLoop.run", fake_run):
            resp = app_client.post(
                "/api/editor/ai/stream",
                json={
                    "action": "polish",
                    "selection": "这段文字需要润色。",
                },
            )
            payload = resp.get_data()

        assert resp.status_code == 200
        events = parse_sse_events(payload)
        assert any(e.get("type") == "plan" for e in events)
        assert any(e.get("type") == "step_start" for e in events)
        assert any(e.get("type") == "step_progress" for e in events)
        assert any(e.get("type") == "step_done" for e in events)
        done_events = [e for e in events if e.get("type") == "done"]
        assert done_events and done_events[0].get("result") == "处理完成"

    def test_main_stream_keeps_editor_actions_runtime_only(
        self, app_client, monkeypatch
    ):
        import web.app as web_app_module
        from app.core.agent.lifecycle import evt_task_complete

        calls = {"save": 0}

        def fake_append(*args, **kwargs):
            calls["save"] += 1
            return []

        def fake_run(self, request):
            yield evt_task_complete(result="润色完成")

        monkeypatch.setattr(
            web_app_module.session_manager, "append_and_save", fake_append
        )

        with patch("app.core.agent.agent_loop.KotoAgentLoop.run", fake_run):
            resp = app_client.post(
                "/api/editor/ai/stream",
                json={
                    "action": "polish",
                    "selection": "这段文字需要润色。",
                    "session_id": "workspace_demo",
                    "history": [
                        {"role": "user", "content": "先润色标题"},
                        {"role": "assistant", "content": "标题已经润色"},
                    ],
                },
            )
            payload = resp.get_data()

        assert resp.status_code == 200
        assert calls == {"save": 0}
        events = parse_sse_events(payload)
        done_events = [e for e in events if e.get("type") == "done"]
        assert done_events and done_events[0].get("result") == "润色完成"

    def test_non_task_stream_passes_preferred_and_local_model_into_agent_request(
        self, app_client
    ):
        """普通 editor SSE 请求应把显式云端模型和配置的本地模型一起传入 AgentLoop。"""
        captured = {}
        from app.core.agent.lifecycle import evt_task_complete

        def fake_run(self, request):
            captured["request"] = request
            yield evt_task_complete(result="ok")

        with patch("app.core.agent.agent_loop.KotoAgentLoop.run", fake_run), patch(
            "web.app._get_configured_local_model_id", return_value="qwen3.5:9b"
        ):
            resp = app_client.post(
                "/api/editor/ai/stream",
                json={
                    "action": "polish",
                    "selection": "这段文字需要润色。",
                    "model_mode": "cloud",
                    "model_id": "gemini-2.5-pro",
                },
            )
            _ = resp.get_data()

        assert resp.status_code == 200
        req = captured["request"]
        assert req.model_mode == "cloud"
        assert req.extra["preferred_model"] == "gemini-2.5-pro"
        assert req.extra["local_model"] == "qwen3.5:9b"


class TestEditorAIAgent:
    """Tests for POST /api/editor/ai/agent structured progress events."""

    def test_agent_route_emits_structured_step_events(self, app_client):
        from app.core.agent.types import AgentAction, AgentStep, AgentStepType

        class FakeAgent:
            def run(self, input_text, session_id=None, system_context=None):
                yield AgentStep(
                    step_type=AgentStepType.THOUGHT, content="先理解文档问题"
                )
                yield AgentStep(
                    step_type=AgentStepType.ACTION,
                    content="执行搜索",
                    action=AgentAction(
                        tool_name="web_search", tool_args={"query": "AI"}
                    ),
                )
                yield AgentStep(
                    step_type=AgentStepType.OBSERVATION,
                    content="找到结果",
                    observation="找到 3 条相关结果",
                )
                yield AgentStep(step_type=AgentStepType.ANSWER, content="最终答案")

        with patch("app.api.agent_routes.get_agent", return_value=FakeAgent()):
            resp = app_client.post(
                "/api/editor/ai/agent",
                json={"query": "帮我分析这份文档", "full_text": "文档内容"},
            )
            payload = resp.get_data()

        assert resp.status_code == 200
        events = parse_sse_events(payload)
        types = [e.get("type") for e in events]
        assert "thought" in types
        assert "step_start" in types
        assert "tool_call" in types
        assert "tool_result" in types
        assert "step_done" in types
        assert "token" in types
        assert "done" in types


class TestChartStream:
    """Tests for POST /api/editor/ai/chart streaming progress."""

    def test_chart_stream_emits_step_events(self, app_client):
        fake_result = {
            "stdout": "",
            "stderr": "",
            "files": {"chart.png": "ZmFrZQ=="},
            "error": "",
        }

        with patch("app.core.sandbox.run_python", return_value=fake_result):
            resp = app_client.post(
                "/api/editor/ai/chart",
                json={
                    "data_context": "类别,值\nA,10",
                    "instruction": "画一个简单图表",
                    "lang": "python",
                },
            )
            payload = resp.get_data()

        assert resp.status_code == 200
        events = parse_sse_events(payload)
        types = [e.get("type") for e in events]
        assert "step_start" in types
        assert "step_done" in types
        assert "code" in types
        assert "image" in types
        assert "done" in types


class TestChartRerun:
    """Tests for POST /api/editor/ai/chart-rerun"""

    def test_empty_code_returns_error(self, app_client):
        """空代码应返回错误"""
        resp = app_client.post(
            "/api/editor/ai/chart-rerun",
            json={
                "code": "",
                "lang": "python",
            },
        )
        assert resp.status_code == 200
        data = resp.get_json()
        assert data["error"]

    def test_simple_python_code(self, app_client):
        """简单 Python 代码应成功执行"""
        resp = app_client.post(
            "/api/editor/ai/chart-rerun",
            json={
                "code": "print('hello')",
                "lang": "python",
            },
        )
        assert resp.status_code == 200
        data = resp.get_json()
        assert "hello" in (data.get("stdout") or "")

    def test_chart_generation_code(self, app_client):
        """图表生成代码应产出图片文件"""
        code = (
            "import matplotlib\n"
            "matplotlib.use('Agg')\n"
            "import matplotlib.pyplot as plt\n"
            "plt.plot([1, 2, 3], [1, 4, 9])\n"
            "plt.savefig('chart.png', dpi=72)\n"
            "plt.close()\n"
        )
        resp = app_client.post(
            "/api/editor/ai/chart-rerun",
            json={
                "code": code,
                "lang": "python",
            },
        )
        assert resp.status_code == 200
        data = resp.get_json()
        assert data.get("error") is None or data["error"] == ""
        assert "chart.png" in (data.get("files") or {})


class TestBuildEditorPrompt:
    """Tests for _build_editor_prompt function"""

    def test_polish_includes_full_text(self):
        """润色 prompt 应包含全文上下文"""
        try:
            from web.app import _build_editor_prompt
        except ImportError:
            pytest.skip("Cannot import _build_editor_prompt")

        prompt = _build_editor_prompt(
            "polish",
            "需要润色的内容",
            "",
            "全文开头。需要润色的内容。全文结尾。",
        )
        assert "全文" in prompt or "文档" in prompt

    def test_find_replace_prompt_structure(self):
        """查找替换 prompt 应包含 JSON 格式要求"""
        try:
            from web.app import _build_editor_prompt
        except ImportError:
            pytest.skip("Cannot import _build_editor_prompt")

        prompt = _build_editor_prompt(
            "find_replace",
            "",
            "把你好替换成您好",
            "你好世界，你好中国。",
        )
        assert "replacements" in prompt
        assert "JSON" in prompt

    def test_find_reference_prompt(self):
        """引用查找 prompt 应包含来源格式要求"""
        try:
            from web.app import _build_editor_prompt
        except ImportError:
            pytest.skip("Cannot import _build_editor_prompt")

        prompt = _build_editor_prompt(
            "find_reference",
            "AI 在教育中的应用",
            "",
            "文档全文内容",
        )
        assert "参考" in prompt or "引用" in prompt or "来源" in prompt

    def test_chart_prompt_uses_high_dpi_and_cjk_fallbacks(self):
        """画图 prompt 应显式要求中文字体回退和更高 dpi。"""
        try:
            from web.app import _build_editor_prompt
        except ImportError:
            pytest.skip("Cannot import _build_editor_prompt")

        prompt = _build_editor_prompt(
            "chart",
            "",
            "",
            "A,B\n1,2\n3,4",
        )

        assert "Microsoft YaHei" in prompt
        assert "Noto Sans CJK SC" in prompt
        assert "axes.unicode_minus" in prompt
        assert "dpi=220" in prompt


class TestPromptContextTruncation:
    """Test that full_text context is properly truncated."""

    def test_long_full_text_truncated(self):
        """超长全文应被截断以控制 token"""
        try:
            from web.app import _build_editor_prompt
        except ImportError:
            pytest.skip("Cannot import _build_editor_prompt")

        long_text = "A" * 20000
        prompt = _build_editor_prompt("polish", "选中内容", "", long_text)
        # Prompt should not contain the entire 20K text
        assert len(prompt) < 15000


class TestLocalModelMode:
    """Tests that model_mode='local' properly routes to Ollama (not cloud) in editor_ai_stream."""

    def test_local_mode_uses_ollama_when_alive(self, app_client):
        """model_mode=local + Ollama alive → response comes from Ollama, not cloud."""
        from unittest.mock import patch, MagicMock

        # Mock Ollama provider to return a known response
        mock_provider = MagicMock()
        mock_provider.generate_content.return_value = iter(
            [
                {"content": "本地", "tool_calls": [], "usage": {}},
                {"content": "Ollama", "tool_calls": [], "usage": {}},
                {"content": "响应", "tool_calls": [], "usage": {}},
            ]
        )

        with patch(
            "app.core.socket_handler._is_ollama_alive", return_value=True
        ), patch(
            "app.core.socket_handler._get_local_provider", return_value=mock_provider
        ), patch(
            "app.core.agent.agent_loop._is_ollama_alive", return_value=True
        ), patch(
            "app.core.agent.agent_loop._get_local_provider", return_value=mock_provider
        ):
            resp = app_client.post(
                "/api/editor/ai/stream",
                json={
                    "action": "polish",
                    "selection": "需要润色的文字",
                    "model_mode": "local",
                },
            )
            # Force stream consumption while patches are active.
            resp_data = resp.get_data()

        assert resp.status_code == 200
        events = parse_sse_events(resp_data)
        token_texts = "".join(
            e.get("text", "") for e in events if e.get("type") == "token"
        )
        assert "Ollama" in token_texts or "本地" in token_texts

    def test_local_mode_ollama_not_running_returns_error(self, app_client):
        """model_mode=local + Ollama not running → returns error event."""
        from unittest.mock import patch

        with patch(
            "app.core.socket_handler._is_ollama_alive", return_value=False
        ), patch("app.core.agent.agent_loop._is_ollama_alive", return_value=False):
            resp = app_client.post(
                "/api/editor/ai/stream",
                json={
                    "action": "polish",
                    "selection": "需要润色的文字",
                    "model_mode": "local",
                },
            )
            # Force stream consumption while patches are active.
            resp_data = resp.get_data()

        assert resp.status_code == 200
        events = parse_sse_events(resp_data)
        error_events = [e for e in events if e.get("type") == "error"]
        assert (
            len(error_events) > 0
        ), "Expected an error event when Ollama is not running"

    def test_explicit_cloud_mode_is_not_overridden_by_legacy_local_only_setting(
        self, app_client
    ):
        """model_mode=cloud must keep using cloud first even if legacy local-only is enabled."""

        class FakeCloudProvider:
            def generate_content(
                self,
                prompt=None,
                model=None,
                system_instruction=None,
                stream=False,
                **kwargs,
            ):
                assert stream is True
                return iter(
                    [
                        {"content": "云端", "tool_calls": [], "usage": {}},
                        {"content": "Gemini", "tool_calls": [], "usage": {}},
                    ]
                )

        class FakeLocalProvider:
            def generate_content(
                self,
                prompt=None,
                model=None,
                system_instruction=None,
                stream=False,
                **kwargs,
            ):
                assert stream is True
                return iter(
                    [
                        {"content": "本地", "tool_calls": [], "usage": {}},
                        {"content": "Ollama", "tool_calls": [], "usage": {}},
                    ]
                )

        with patch("web.settings.SettingsManager.get", return_value=True), patch(
            "app.core.agent.agent_loop._get_provider", return_value=FakeCloudProvider()
        ), patch(
            "app.core.agent.agent_loop._get_local_provider",
            return_value=FakeLocalProvider(),
        ), patch(
            "app.core.agent.agent_loop._is_ollama_alive", return_value=True
        ):
            resp = app_client.post(
                "/api/editor/ai/stream",
                json={
                    "action": "polish",
                    "selection": "需要润色的文字",
                    "model_mode": "cloud",
                },
            )
            resp_data = resp.get_data()

        assert resp.status_code == 200
        events = parse_sse_events(resp_data)
        token_texts = "".join(
            e.get("text", "") for e in events if e.get("type") == "token"
        )
        assert "云端" in token_texts
        assert "Gemini" in token_texts
        assert "Ollama" not in token_texts

    def test_workspace_quick_actions_use_whitebox_only(self):
        """Workspace quick actions should use whitebox routes without editor SSE fallback."""
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        quick_actions = Path("web/static/js/workspace-ai-quick-actions.js").read_text(
            encoding="utf-8"
        )
        transport = Path("web/static/js/workspace-ai-transport.js").read_text(
            encoding="utf-8"
        )
        assert "window.WA.sendQuickAction = (action) => {" in src
        assert "window.WA.createWorkspaceQuickActionRuntime" in quick_actions
        assert "attachDispatcher" in quick_actions
        assert (
            "return Promise.reject(new Error(`快捷动作 ${actionId} 未配置可用的执行路径`));"
            in quick_actions
        )
        assert (
            "if (!attachedDispatcher || typeof attachedDispatcher.dispatchMessage !== 'function') {"
            in quick_actions
        )
        assert "/api/editor/ai/stream" not in quick_actions
        assert "legacyEditorFallback" not in quick_actions
        assert "sendEditorAction" not in quick_actions
        assert "editorAction" not in quick_actions
        assert "streamEventBlocks({" not in quick_actions
        assert "typeof options.body === 'string'" in transport
        assert "async function streamEventBlocks(options)" in transport

    def test_workspace_readonly_quick_actions_can_use_simple_whitebox(self):
        quick_actions = Path("web/static/js/workspace-ai-quick-actions.js").read_text(
            encoding="utf-8"
        )

        assert "whiteboxMode: 'simple'" in quick_actions
        assert "function usesSimpleWhitebox(action) {" in quick_actions
        assert (
            "function sendSimpleWhiteboxAction(payload, providedAction) {"
            in quick_actions
        )
        assert "return attachedDispatcher.dispatchMessage({" in quick_actions
        assert "quick_action_mode: 'simple'" in quick_actions

    def test_workspace_edit_quick_actions_can_use_proposal_whitebox(self):
        quick_actions = Path("web/static/js/workspace-ai-quick-actions.js").read_text(
            encoding="utf-8"
        )

        assert "whiteboxMode: 'proposal'" in quick_actions
        assert "function usesProposalWhitebox(action) {" in quick_actions
        assert "function buildProposalWhiteboxTask(payload, action) {" in quick_actions
        assert (
            "function sendProposalWhiteboxAction(payload, providedAction) {"
            in quick_actions
        )
        assert "quick_action_mode: 'proposal'" in quick_actions
        assert "options.handleProposals({" in quick_actions
        assert "sendEditorAction" not in quick_actions

    def test_workspace_model_state_uses_wa_keys_only(self):
        """Workspace assistant should use wa_* model state only and not write legacy editor_* keys."""
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        toggle_start = src.find("function _setWorkspaceModelMode(mode) {")
        toggle_end = src.find("window.WA.refreshModelCatalog", toggle_start)
        assert toggle_start != -1 and toggle_end != -1
        toggle_section = src[toggle_start:toggle_end]
        assert "function _syncEditorModelPreference(" not in src
        assert "editor_model_mode" not in src
        assert "editor_locked_model" not in src
        assert "localStorage.setItem('wa_locked_model', newModel);" in toggle_section
        assert (
            "localStorage.setItem('wa_model_choice_explicit', '1');" in toggle_section
        )

    def test_workspace_chart_requests_delegate_to_whitebox_dispatcher(self):
        """Python chart requests should route through the whitebox dispatcher instead of the legacy chart SSE helper."""
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        quick_actions = Path("web/static/js/workspace-ai-quick-actions.js").read_text(
            encoding="utf-8"
        )
        dispatcher = Path("web/static/js/workspace-task-dispatcher.js").read_text(
            encoding="utf-8"
        )
        assert "async function _sendViaSSEChart(payload) {" not in src
        assert "_sendViaLegacyChartSSE" not in src
        assert "fetch('/api/editor/ai/chart'" not in src
        assert "_waQuickActionRuntime.sendChartAction(payload);" not in src
        assert "attachedDispatcher.dispatchMessage({" in quick_actions
        assert "url: '/api/editor/ai/chart'" not in quick_actions
        assert (
            "model_mode: typeof options.getModelMode === 'function' ? options.getModelMode() : 'auto',"
            in dispatcher
        )
        assert (
            "model_id: typeof options.getSelectedCloudModelId === 'function' ? options.getSelectedCloudModelId() : '',"
            in dispatcher
        )

    def test_workspace_task_renderer_supports_python_artifacts(self):
        """The whitebox task renderer should render image artifacts emitted by run_python_code."""
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )
        assert "function appendToolArtifacts(row, payload)" in renderer
        assert "payload.artifacts" in renderer
        assert "wa-task-artifact-image" in renderer

    def test_workspace_task_renderer_distinguishes_blocked_python_calls(self):
        """Blocked run_python_code guidance should be labeled as an interception reason, not Python output."""
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )
        assert "payload.blocked ? '查看拦截原因' : '查看执行输出'" in renderer
        assert "const blocked = !!payload.blocked;" in renderer
        assert "const chipText = blocked ? '拦截'" in renderer

    def test_workspace_task_renderer_eagerly_refreshes_changed_files(self):
        """Segmented file tasks should try to refresh the edited document as soon as file.changed arrives."""
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )
        refresh = Path("web/static/js/workspace-ai-task-refresh.js").read_text(
            encoding="utf-8"
        )
        assert "window.WA.createFileTaskRefreshController" in refresh
        assert "card._fileRefreshPromise" in refresh
        assert "void flush(card).catch" in refresh
        assert (
            "if (!((card._pendingFileRefreshes && card._pendingFileRefreshes.size) || card._fileRefreshPromise))"
            in refresh
        )
        assert "function upsertEntry(card, item)" in refresh
        assert "status: supported ? 'pending' : 'unsupported'" in refresh
        assert "status: 'refreshing'" in refresh
        assert "status: 'reloaded'" in refresh
        assert "status: 'failed'" in refresh
        assert "function queueFileRefresh(card, payload, options)" in renderer
        assert "controller.queue(card, payload, options);" in renderer
        assert "triggerQueuedFileRefresh(card, {" in renderer

    def test_workspace_task_renderer_keeps_write_tool_milestones_visible(self):
        """Successful file-writing tool events should remain visible instead of being fully suppressed."""
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )

        assert "return isInternalTool(name) || isReadTool(name);" in renderer
        assert "ALWAYS_SUPPRESS_TOOL_FINISHED_NAMES.has(name)" in renderer
        assert "'repair_guard'" in renderer
        assert "'readonly_answer_guard'" in renderer
        assert "return false;" in renderer
        assert "toolStepTitle(payload.tool_name)" in renderer
        assert (
            "setStepTitle(step, blocked ? `${toolLabel(payload.tool_name)}已拦截`"
            in renderer
        )

    def test_workspace_task_renderer_refreshes_terminal_files_without_final_notice(
        self,
    ):
        """Terminal refresh should reload files without adding a separate completion notice."""
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )
        refresh = Path("web/static/js/workspace-ai-task-refresh.js").read_text(
            encoding="utf-8"
        )

        assert "function finalizeTerminalRefresh(card, payload, options)" in renderer
        assert "queueTerminalFileChanges(card, payload || {});" in renderer
        assert "showRefreshingStatus: false" in renderer
        assert "summaryHtml(card)" not in renderer
        assert "function summaryHtml(card)" not in refresh

    def test_workspace_task_renderer_warns_when_whitebox_stream_seq_is_incomplete(self):
        """Dropped or out-of-order SSE events should be surfaced to users as whitebox visibility warnings."""
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )

        assert "function noteStreamIssue(card, key, text)" in renderer
        assert "state.streamIssueKeys" in renderer
        assert "state.lastEventSeq" in renderer
        assert "检测到部分进度事件未按顺序抵达" in renderer
        assert "检测到任务进度事件重放" in renderer

    def test_workspace_task_renderer_supports_step_result_rollups(self):
        """The renderer should display backend step.result rollups for major whitebox phases."""
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )

        assert "function renderStepResult(payload)" in renderer
        assert "function upsertStepResultRow(step, payload)" in renderer
        assert "type === 'step.result'" in renderer
        assert "payload.file_change_count" in renderer
        assert "涉及 ${esc(changeCount)} 个文件" in renderer
        assert "appendToolArtifacts(row, payload);" in renderer

    def test_workspace_task_renderer_supports_step_progress_updates(self):
        """The whitebox task renderer should keep the current step alive with incremental progress updates."""
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )

        assert "function upsertProgressRow(step, payload)" in renderer
        assert "type === 'step_progress' || type === 'step.progress'" in renderer
        assert "step._progressRow" in renderer
        assert (
            "if ((payload.file_updated || payload.fileUpdated) && (payload.path || payload.file_path || payload.output_path || payload.target_path))"
            in renderer
        )
        assert "queueFileRefresh(card, payload, {" in renderer

    def test_workspace_task_renderer_preserves_structured_run_errors(self):
        """Structured run.error messages should survive transport failures instead of collapsing to generic network errors."""
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )
        dispatcher = Path("web/static/js/workspace-task-dispatcher.js").read_text(
            encoding="utf-8"
        )

        assert "function makeTaskError(message) {" in renderer
        assert (
            "const fatalText = payload.text || payload.error || '任务失败';" in renderer
        )
        assert "fatalText," in renderer
        assert (
            "if (card._fatalErrorText) throw makeTaskError(card._fatalErrorText);"
            in renderer
        )
        assert "error && error.waTaskError ? error.message" in dispatcher

    def test_workspace_task_renderer_surfaces_runtime_metadata(self):
        """The task renderer should expose runtime execution metadata from whitebox SSE terminal events."""
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )

        assert "function runtimeExecutionLabel(runtime) {" in renderer
        assert "function runtimeMetaHtml(payload) {" in renderer
        assert "function finalRunStatusText(payload) {" in renderer
        assert "执行：" in renderer
        assert "结果：" in renderer
        assert "回退自：" in renderer
        assert "return '缺少工具';" in renderer
        assert "return '摘要回退';" in renderer
        assert "setStatus(card, finalRunStatusText(payload));" in renderer

        tool_gap_start = renderer.find("function renderToolGap(evt) {")
        tool_gap_end = renderer.find(
            "function renderFollowupRecord(record) {", tool_gap_start
        )
        run_summary_start = renderer.find("function renderRunSummary(payload, card) {")
        run_summary_end = renderer.find(
            "function renderFileChange(evt) {", run_summary_start
        )
        check_finished_start = renderer.find("if (type === 'check.finished') {")
        check_finished_end = renderer.find(
            "if (type === 'step.finished') {", check_finished_start
        )

        assert tool_gap_start != -1 and tool_gap_end != -1
        assert run_summary_start != -1 and run_summary_end != -1
        assert check_finished_start != -1 and check_finished_end != -1
        assert "runtimeMetaHtml(payload)" in renderer[tool_gap_start:tool_gap_end]
        assert (
            "runtimeExecutionLabel(artifact.runtime_context)"
            in renderer[tool_gap_start:run_summary_end]
        )
        assert "runtimeMetaHtml(payload)" in renderer[run_summary_start:run_summary_end]
        assert (
            "runtimeMetaHtml(payload)"
            in renderer[check_finished_start:check_finished_end]
        )

    def test_workspace_task_renderer_surfaces_task_classification_metadata(self):
        """The whitebox task renderer should persist classification metadata from run lifecycle events."""
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )
        dispatcher = Path("web/static/js/workspace-task-dispatcher.js").read_text(
            encoding="utf-8"
        )

        assert "function normalizedTaskLifecyclePayload(payload) {" in renderer
        assert (
            "const classification = data.classification && typeof data.classification === 'object'"
            in renderer
        )
        assert "if (type === 'task.classified') {" in renderer
        assert "function classificationMetaHtml(card) {" in renderer
        assert "function renderTaskClassification(evt, card) {" in renderer
        assert "const data = normalizedTaskLifecyclePayload(payload);" in renderer
        assert (
            "card.dataset.taskRequestKind = String(data.request_kind || '').trim();"
            in renderer
        )
        assert (
            "card.dataset.taskFamily = String(data.task_family || '').trim();"
            in renderer
        )
        assert (
            "card.dataset.taskOutputMode = String(data.output_mode || '').trim();"
            in renderer
        )
        assert (
            "card.dataset.taskSelectedRecipe = String(data.selected_recipe || '').trim();"
            in renderer
        )
        assert "card.dataset.taskIntentStrategy = intentStrategy;" in renderer
        assert (
            "card.dataset.taskIntentCanApply = intentPlan.can_apply ? 'true' : 'false';"
            in renderer
        )
        assert (
            "card.dataset.taskIntentRequiresConfirmation = intentPlan.requires_confirmation ? 'true' : 'false';"
            in renderer
        )
        assert "if (type === 'run.started') {" in renderer
        assert "if (type === 'run.finished') {" in renderer
        assert "产出：" in renderer
        assert "策略：" in renderer
        assert "路线：" in renderer
        assert "后续：" in renderer
        assert "目标：" in renderer
        assert "chips.push(`请求：" not in renderer
        assert "chips.push(`任务：" not in renderer
        assert "chips.push(`操作：" not in renderer
        assert "chips.push(`分类：" not in renderer
        assert "${classificationHtml}${pendingResumeHtml}${runtimeHtml}" in renderer

        assert (
            "if (dataset.taskRequestKind) metadata.task_request_kind = String(dataset.taskRequestKind || '').trim();"
            in dispatcher
        )
        assert (
            "if (dataset.taskFamily) metadata.task_family = String(dataset.taskFamily || '').trim();"
            in dispatcher
        )
        assert (
            "if (dataset.taskExecutionMode) metadata.task_execution_mode = String(dataset.taskExecutionMode || '').trim();"
            in dispatcher
        )
        assert (
            "if (dataset.taskSelectedRecipe) metadata.task_selected_recipe = String(dataset.taskSelectedRecipe || '').trim();"
            in dispatcher
        )
        assert (
            "if (dataset.taskOutputMode) metadata.task_output_mode = String(dataset.taskOutputMode || '').trim();"
            in dispatcher
        )
        assert (
            "if (dataset.taskIntentStrategy) metadata.task_intent_strategy = String(dataset.taskIntentStrategy || '').trim();"
            in dispatcher
        )
        assert (
            "metadata.task_intent_can_apply = String(dataset.taskIntentCanApply || '').trim().toLowerCase() === 'true';"
            in dispatcher
        )
        assert (
            "metadata.task_intent_requires_confirmation = String(dataset.taskIntentRequiresConfirmation || '').trim().toLowerCase() === 'true';"
            in dispatcher
        )

    def test_workspace_task_renderer_labels_answer_and_hybrid_output_modes(self):
        """Answer-first and hybrid file tasks should show explicit non-write guidance in the task card."""
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )
        dispatcher = Path("web/static/js/workspace-task-dispatcher.js").read_text(
            encoding="utf-8"
        )

        assert "if (normalized === 'answer') return '只给答案';" in renderer
        assert "if (normalized === 'write') return '写入文件';" in renderer
        assert "if (normalized === 'hybrid') return '先分析后决定';" in renderer
        assert "本轮只做分析，未写入文件。" in renderer
        assert "本轮先完成分析；确认后可写入文件。" in renderer
        assert "本轮完成分析，未写入文件。" in renderer
        assert 'data-task-followup-action="apply"' in renderer
        assert "应用建议" in renderer
        assert (
            "if (outputMode === 'hybrid' && canApply) chips.push(`后续：${requiresConfirmation ? '确认后可写入' : '可继续写入'}`);"
            in renderer
        )
        assert (
            "const previousTaskOutputMode = previewText(previousTaskTurn.task_output_mode || '', 120);"
            in dispatcher
        )
        assert (
            "if (previousTaskOutputMode) context.previous_task_output_mode = previousTaskOutputMode;"
            in dispatcher
        )

    def test_workspace_task_renderer_surfaces_intent_plan_hints(self):
        """Intent-plan metadata should add user-facing strategy and applyability hints without exposing raw internals."""
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )
        dispatcher = Path("web/static/js/workspace-task-dispatcher.js").read_text(
            encoding="utf-8"
        )

        assert "function intentStrategyLabel(value, outputMode) {" in renderer
        assert (
            "if (normalized === 'analyze_then_confirm') return '先分析后确认';"
            in renderer
        )
        assert "if (normalized === 'design_new_tool') return '需补工具';" in renderer
        assert (
            "if (normalized === 'answer_only' && normalizedOutput === 'answer') return '';"
            in renderer
        )
        assert (
            "if (normalized === 'write_through' && normalizedOutput === 'write') return '';"
            in renderer
        )
        assert "if (strategyLabel) chips.push(`策略：${strategyLabel}`);" in renderer
        assert "继续细化方案" in renderer
        assert (
            "const previousTaskIntentStrategy = previewText(previousTaskTurn.task_intent_strategy || '', 120);"
            in dispatcher
        )
        assert (
            "context.previous_task_intent_strategy = previousTaskIntentStrategy;"
            in dispatcher
        )
        assert (
            "context.previous_task_intent_can_apply = previousTaskIntentCanApply;"
            in dispatcher
        )
        assert (
            "context.previous_task_intent_requires_confirmation = previousTaskIntentRequiresConfirmation;"
            in dispatcher
        )

    def test_workspace_task_renderer_hides_placeholder_classification_metadata(self):
        """Default whitebox classification placeholders should stay internal unless they add user-facing meaning."""
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )

        assert "function shouldDisplayClassificationLabel(kind, value) {" in renderer
        assert "if (kind === 'request') return normalized !== 'new_task';" in renderer
        assert "if (kind === 'family') return normalized !== 'analyze';" in renderer
        assert "if (kind === 'operation') return normalized !== 'read';" in renderer
        assert (
            "if (kind === 'execution') return normalized !== 'generic_tool_loop';"
            in renderer
        )
        assert "if (selectedRecipe) chips.push(`路线：${selectedRecipe}`);" in renderer
        assert (
            "if (targetFileType && chips.length) chips.push(`目标：${targetFileType.toUpperCase()}`);"
            in renderer
        )
        assert "if (!classificationHtml && !reasonHtml) return '';" in renderer
        assert "chips.push(`置信：${Math.round(confidence * 100)}%`);" not in renderer

    def test_workspace_task_renderer_supports_plan_briefed_updates(self):
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )

        assert "if (type === 'plan.briefed') {" in renderer
        assert "setStatus(card, '已分析任务');" in renderer
        assert (
            "const title = String(payload.title || '执行方案').trim() || '执行方案';"
            in renderer
        )

    def test_workspace_task_renderer_surfaces_whitebox_plan_gate(self):
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )

        assert "function renderWhiteboxExecutionPlan(card, payload) {" in renderer
        assert "if (type === 'plan.proposed') {" in renderer
        assert "if (type === 'plan.gated') {" in renderer
        assert "AI 执行计划" in renderer
        assert "计划监管" in renderer
        assert "renderPlanGateIssue(payload)" in renderer

    def test_workspace_task_renderer_supports_simple_quick_action_mode(self):
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )

        assert "card.dataset.taskQuickActionMode = quickActionMode;" in renderer
        assert "data.quick_action_mode" in renderer
        assert (
            "card.dataset.taskQuickActionMode || '').trim().toLowerCase() === 'simple'"
            in renderer
        )
        assert (
            "card.dataset.taskQuickActionMode || '').trim().toLowerCase() === 'proposal'"
            in renderer
        )

    def test_workspace_model_selector_fetches_dynamic_catalog(self):
        """Workspace assistant should fetch the dynamic model catalog instead of relying on hardcoded options."""
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        assert "function _refreshModelCatalog(force = false)" in src
        assert "fetch('/api/v1/models', { cache: 'no-store' })" in src
        assert "_syncModelStatusUi();" in src

    def test_workspace_model_state_syncs_from_server_status_on_init(self):
        """Workspace assistant should honor the persisted server-side local/cloud mode on startup."""
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        assert "function _syncLockedModelFromServer()" in src
        assert "fetch('/api/local-model/status', { cache: 'no-store' })" in src
        assert "_syncLockedModelFromServer().finally(() => {" in src
        assert "_checkOllamaStatus();" in src

    def test_workspace_stream_handlers_consume_classification_events(self):
        """Workspace assistant streams should surface backend classification/model routing events."""
        assistant = Path("web/static/js/workspace-assistant.js").read_text(
            encoding="utf-8"
        )
        quick_actions = Path("web/static/js/workspace-ai-quick-actions.js").read_text(
            encoding="utf-8"
        )
        assert "applyRouteEvent: _applyRouteEvent," in assistant
        assert (
            "if (parsed.type === 'classification' || parsed.type === 'route') {"
            in quick_actions
        )
        assert (
            "if (evt.type === 'classification' || evt.type === 'route') {"
            in quick_actions
        )
        assert "options.applyRouteEvent(parsed);" in quick_actions
        assert "options.applyRouteEvent(evt);" in quick_actions

    def test_workspace_send_quick_action_routes_text_and_chart_via_unified_paths(self):
        """sendQuickAction should call editor/chart SSE helpers instead of legacy JSON quick-action."""
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        fn_start = src.find("window.WA.sendQuickAction = (action) => {")
        fn_end = src.find("window.WA.sendSelectionToAI = () => {", fn_start)
        assert fn_start != -1 and fn_end != -1
        send_quick = src[fn_start:fn_end]
        assert "_waTaskDispatcher.dispatchQuickAction(action, {" in send_quick
        assert "/api/v1/workspace/quick-action" not in send_quick

    def test_workspace_selection_to_ai_reads_live_textarea_selection(self):
        """Plain-text editor selections should be read from the live textarea before pinning AI context."""
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        assert "function _getActiveTextEditorSelectionForAI()" in src
        assert "textarea.selectionStart" in src
        assert "textarea.selectionEnd" in src

        send_selection_start = src.find("window.WA.sendSelectionToAI = () => {")
        send_selection_end = src.find(
            "// Auto-expand the right AI panel if it's collapsed", send_selection_start
        )
        assert send_selection_start != -1 and send_selection_end != -1
        send_selection = src[send_selection_start:send_selection_end]
        assert "const liveSelection = _getLiveEditorSelectionForAI();" in send_selection

    def test_workspace_task_dispatcher_does_not_keep_current_file_heuristics(self):
        """Current-file heuristics should stay removed from the whitebox dispatcher."""
        dispatcher = Path("web/static/js/workspace-task-dispatcher.js").read_text(
            encoding="utf-8"
        )
        assert "function wantsCurrentFile(text)" not in dispatcher
        assert "function currentFileContext(text)" not in dispatcher
        assert "wantsCurrentFile," not in dispatcher
        assert "currentFileContext," not in dispatcher
        assert "当前(?:打开的)?(?:\\s*[\\w.+#-]+)?\\s*(?:文件|文档)" not in dispatcher

    def test_workspace_quick_action_keyword_list_includes_check(self):
        """The workspace assistant quick-action keyword routing must recognize 检查."""
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        quick_actions = Path("web/static/js/workspace-ai-quick-actions.js").read_text(
            encoding="utf-8"
        )
        quick_start = src.find("window.WA.quickAction = (text) => {")
        quick_end = src.find("window.WA.pptxSync = (ta) => {", quick_start)
        assert quick_start != -1 and quick_end != -1
        quick_section = src[quick_start:quick_end]
        assert "_waTaskDispatcher.matchQuickAction(text)" in quick_section
        assert (
            "attachedDispatcher.registerQuickActionKeyword(keyword, action.action);"
            in quick_actions
        )
        assert "action: '检查'" in quick_actions
        assert "keywords: ['检查']" in quick_actions
        assert "WA.sendQuickAction(matchedAction);" in quick_section

    def test_pdf_ai_annotate_is_explicitly_disabled_during_migration(self):
        """PDF AI annotate should not call the legacy quick-action path while the feature is offline."""
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        fn_start = src.find("async aiAnnotate() {")
        fn_end = src.find("// ─── AI Watermark removal", fn_start)
        assert fn_start != -1 and fn_end != -1
        ai_annotate = src[fn_start:fn_end]
        assert "/api/v1/workspace/quick-action" not in ai_annotate
        assert "ai_annotate" not in ai_annotate
        assert "AI 标注功能正在迁移到新的 AI 流程" in ai_annotate

    def test_pdf_ai_annotate_button_title_marks_temporary_unavailability(self):
        """The PDF toolbar button should advertise that AI annotate is temporarily unavailable."""
        html = Path("web/templates/workspace_assistant.html").read_text(
            encoding="utf-8"
        )
        assert 'title="AI 标注迁移中，暂不可用"' in html

    def test_workspace_templates_use_shared_model_controls_partial(self):
        """Workspace templates should share provider/local model controls without exposing a redundant model picker."""
        standalone_html = Path("web/templates/workspace_assistant.html").read_text(
            encoding="utf-8"
        )
        index_html = Path("web/templates/index.html").read_text(encoding="utf-8")
        partial_html = Path("web/templates/_workspace_model_controls.html").read_text(
            encoding="utf-8"
        )
        assert "{% include '_workspace_model_controls.html' %}" in standalone_html
        assert "{% include '_workspace_model_controls.html' %}" in index_html
        assert 'id="wa-model-mode-toggle"' in partial_html
        assert 'id="wa-model-mode-gemini-btn"' in partial_html
        assert 'id="wa-model-mode-deepseek-btn"' in partial_html
        assert 'id="wa-model-mode-local-btn"' in partial_html
        assert 'id="wa-model-select"' not in partial_html
        assert 'id="wa-model-mode-gemini-btn"' not in standalone_html
        assert 'id="wa-model-mode-deepseek-btn"' not in standalone_html
        assert 'id="wa-model-mode-local-btn"' not in standalone_html
        assert 'id="wa-model-mode-gemini-btn"' not in index_html
        assert 'id="wa-model-mode-deepseek-btn"' not in index_html
        assert 'id="wa-model-mode-local-btn"' not in index_html

    def test_workspace_cloud_selection_maps_request_model_mode_to_cloud(self):
        js = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        quick_actions = Path("web/static/js/workspace-ai-quick-actions.js").read_text(
            encoding="utf-8"
        )
        assert (
            "const _WA_MODEL_MODES = new Set(['cloud', 'gemini', 'deepseek', 'local']);"
            in js
        )
        assert (
            "lockedModel: _normalizeWorkspaceModelMode(localStorage.getItem('wa_locked_model'), 'cloud')"
            in js
        )
        assert (
            "const storedLockedModel = localStorage.getItem('wa_locked_model');" in js
        )
        assert "return _normalizeWorkspaceModelMode(state.lockedModel, 'cloud');" in js
        assert "model_mode: payload.model_mode || getModelMode()," in quick_actions
        assert "model_mode: modelMode," in quick_actions
        assert (
            "window.WA.setLockedModel = (val) => {\n    _setWorkspaceModelMode(val);\n  };"
            in js
        )
        assert "body: JSON.stringify({ mode: newModel })," in js

    def test_workspace_quick_actions_do_not_render_raw_tool_result_previews_as_progress(
        self,
    ):
        quick_actions = Path("web/static/js/workspace-ai-quick-actions.js").read_text(
            encoding="utf-8"
        )
        assert "function toolResultProgressText(parsed) {" in quick_actions
        assert "setProgress(toolResultProgressText(parsed));" in quick_actions
        assert "setProgress(parsed.result_preview" not in quick_actions


class TestMainChatProgressRegression:
    """Regression checks for canonical step-event support in the main chat UI."""

    def test_main_chat_normalizes_canonical_step_events(self):
        src = Path("web/static/js/app.js").read_text(encoding="utf-8")
        assert "evt.type === 'plan'" in src
        assert "evt.type === 'phase'" in src
        assert "evt.type === 'step_start'" in src
        assert "evt.type === 'tool_call'" in src
        assert "const canonicalProgressPercent =" in src


class TestRemovedLegacyTaskRoutes:
    """Regression checks that old file-task compatibility routes are gone."""

    def test_skill_execute_route_is_not_registered(self, app_client):
        resp = app_client.post("/api/editor/ai/skill-execute", json={})
        assert resp.status_code == 404

    def test_task_execute_route_is_not_registered(self, app_client):
        resp = app_client.post(
            "/api/editor/ai/task-execute", json={"task": "整理当前文件"}
        )
        assert resp.status_code == 404


class TestLegacyDocumentCompatRoutes:
    """Legacy document APIs should reuse the current DocumentFeedbackSystem paths instead of old batch annotators."""

    def test_web_app_doc_annotate_paths_use_compat_helper(self):
        src = Path("web/app.py").read_text(encoding="utf-8")

        assert "from web.document_annotation_compat import (" in src
        assert "iter_annotation_progress_events(" in src
        assert "collect_annotation_result(" in src
        assert "DocumentFeedbackSystem" not in src

    @staticmethod
    def _make_document_client(monkeypatch, workspace_root):
        from flask import Flask
        import web.blueprints.document as document_routes

        app = Flask(__name__)
        app.config["TESTING"] = True
        monkeypatch.setattr(document_routes, "_get_client", lambda: object())
        monkeypatch.setattr(
            document_routes, "_get_workspace_dir", lambda: str(workspace_root)
        )
        app.register_blueprint(document_routes.document_bp)
        return app.test_client()

    def test_document_annotate_route_uses_streaming_compat_path(
        self, monkeypatch, tmp_path
    ):
        import web.document_feedback as feedback_module

        docx_path = tmp_path / "legacy-annotate.docx"
        docx_path.write_bytes(b"PK\x03\x04")
        revised_path = tmp_path / "legacy-annotate_revised.docx"
        client = self._make_document_client(monkeypatch, tmp_path)

        class FakeFeedback:
            def __init__(self, gemini_client=None, default_model_id="gemini-2.5-pro"):
                self.gemini_client = gemini_client
                self.default_model_id = default_model_id

            def full_annotation_loop(self, *args, **kwargs):
                raise AssertionError("legacy sync annotation loop should not be used")

            def full_annotation_loop_streaming(self, *args, **kwargs):
                yield {
                    "stage": "complete",
                    "result": {
                        "success": True,
                        "revised_file": str(revised_path),
                        "applied": 3,
                    },
                }

        monkeypatch.setattr(feedback_module, "DocumentFeedbackSystem", FakeFeedback)

        resp = client.post(
            "/api/document/annotate",
            json={
                "file_path": str(docx_path),
                "requirement": "请标注翻译问题",
            },
        )

        body = resp.get_json()
        assert resp.status_code == 200
        assert body["success"] is True
        assert body["revised_file"] == str(revised_path)
        assert body["applied"] == 3

    def test_document_analyze_annotations_route_uses_chunked_feedback_path(
        self, monkeypatch, tmp_path
    ):
        import web.document_feedback as feedback_module

        docx_path = tmp_path / "legacy-analyze.docx"
        docx_path.write_bytes(b"PK\x03\x04")
        client = self._make_document_client(monkeypatch, tmp_path)

        class FakeFeedback:
            def __init__(self, gemini_client=None, default_model_id="gemini-2.5-pro"):
                self.gemini_client = gemini_client
                self.default_model_id = default_model_id

            def analyze_for_annotation_chunked(self, *args, **kwargs):
                return {
                    "success": True,
                    "annotations": [{"原文片段": "foo", "修改建议": "bar"}],
                    "annotation_count": 1,
                }

        monkeypatch.setattr(feedback_module, "DocumentFeedbackSystem", FakeFeedback)

        resp = client.post(
            "/api/document/analyze-annotations",
            json={
                "file_path": str(docx_path),
                "requirement": "请只分析并标出问题",
            },
        )

        body = resp.get_json()
        assert resp.status_code == 200
        assert body["success"] is True
        assert body["annotation_count"] == 1
        assert body["annotations"][0]["原文片段"] == "foo"

    def test_document_batch_annotate_stream_route_uses_feedback_streaming_path(
        self, monkeypatch, tmp_path
    ):
        import web.document_feedback as feedback_module

        docx_path = tmp_path / "legacy-batch.docx"
        docx_path.write_bytes(b"PK\x03\x04")
        revised_path = tmp_path / "legacy-batch_revised.docx"
        client = self._make_document_client(monkeypatch, tmp_path)

        class FakeFeedback:
            def __init__(self, gemini_client=None, default_model_id="gemini-2.5-pro"):
                self.gemini_client = gemini_client
                self.default_model_id = default_model_id

            def full_annotation_loop_streaming(self, *args, **kwargs):
                yield {
                    "stage": "reading",
                    "progress": 5,
                    "message": "读取中",
                    "detail": "第 1 批",
                }
                yield {
                    "stage": "complete",
                    "result": {
                        "success": True,
                        "revised_file": str(revised_path),
                        "applied": 2,
                    },
                }

        monkeypatch.setattr(feedback_module, "DocumentFeedbackSystem", FakeFeedback)

        resp = client.post(
            "/api/document/batch-annotate-stream",
            json={
                "file_path": str(docx_path),
                "requirement": "请流式标注",
            },
        )

        body = resp.get_data(as_text=True)
        assert resp.status_code == 200
        assert "event: progress" in body
        assert "event: complete" in body
        assert revised_path.name in body

    def test_document_feedback_route_resolves_relative_path_before_wrapper(
        self, monkeypatch, tmp_path
    ):
        import web.document_feedback as feedback_module

        documents_dir = tmp_path / "documents"
        documents_dir.mkdir()
        docx_path = documents_dir / "legacy-feedback.docx"
        docx_path.write_bytes(b"PK\x03\x04")
        client = self._make_document_client(monkeypatch, tmp_path)
        captured = {}

        class FakeFeedback:
            def __init__(self, gemini_client=None, default_model_id="gemini-2.5-pro"):
                self.gemini_client = gemini_client
                self.default_model_id = default_model_id

            def full_feedback_loop(
                self, file_path, user_requirement="", auto_apply=True
            ):
                captured["file_path"] = file_path
                captured["user_requirement"] = user_requirement
                captured["auto_apply"] = auto_apply
                return {
                    "success": True,
                    "new_file_path": str(
                        docx_path.with_name("legacy-feedback_revised.docx")
                    ),
                    "applied_count": 1,
                }

        monkeypatch.setattr(feedback_module, "DocumentFeedbackSystem", FakeFeedback)

        resp = client.post(
            "/api/document/feedback",
            json={
                "file_path": "legacy-feedback.docx",
                "requirement": "请润色措辞",
                "auto_apply": False,
            },
        )

        body = resp.get_json()
        assert resp.status_code == 200
        assert body["success"] is True
        assert captured == {
            "file_path": str(docx_path),
            "user_requirement": "请润色措辞",
            "auto_apply": False,
        }


class TestLegacyEditorRemovalRegression:
    """Regression checks for removing the old Univer editor stack."""

    def test_legacy_univer_entrypoints_are_removed(self):
        assert not Path("web/univer-editor/index.html").exists()
        assert not Path("web/static/univer-dist/index.html").exists()

    def test_legacy_univer_source_tree_is_removed(self):
        assert not Path("web/univer-editor/main.js").exists()
        assert not Path("web/univer-editor/src").exists()

    def test_workspace_assistant_still_loads_sheets_runtime(self):
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        assert "/static/univer-dist/assets/sheets-main.css" in src
        assert "/static/univer-dist/assets/sheets-main.js" in src
        assert Path("web/static/univer-dist/assets/sheets-main.css").exists()
        assert Path("web/static/univer-dist/assets/sheets-main.js").exists()

    def test_univer_editor_build_only_targets_sheets_runtime(self):
        pkg = Path("web/univer-editor/package.json").read_text(encoding="utf-8")
        assert '"build": "npm run build:sheets && npm run clean:assets"' in pkg
        assert '"build:editor"' not in pkg


class TestTaskAgentDocumentEdits:
    """Regression checks for real file-edit tool execution in TaskAgent."""

    def test_task_agent_local_prompt_prefers_real_excel_to_word_table(self):
        prompt = (ROOT / "app/core/agent/task_agent.py").read_text(encoding="utf-8")

        assert "insert_excel_as_docx_table" in prompt
        assert "insert_image_into_docx" in prompt
        assert (
            "先用 `write_docx_content` 把基于真实表格数据生成的摘要/结论写入目标文档"
            in prompt
        )
        assert "write_docx_content" in prompt

    def test_task_agent_runs_stage_verification_after_file_change(self, monkeypatch):
        from app.core.agent.task_agent import TaskAgent

        class FakeRegistry:
            def __init__(self):
                self.executions = []

            def get_definitions(self):
                return [
                    {"name": "insert_excel_as_docx_table"},
                    {"name": "verify_task_completion"},
                ]

            def execute(self, tool_name, tool_args):
                self.executions.append((tool_name, dict(tool_args)))
                if tool_name == "insert_excel_as_docx_table":
                    return json.dumps(
                        {
                            "success": True,
                            "summary": "已将工作表“汇总表”的 200 行数据写入 Word 表格",
                            "path": "target.docx",
                            "file_type": "docx",
                            "change_type": "modify",
                            "operation": tool_name,
                            "preview": "表格已写入目标文档",
                        },
                        ensure_ascii=False,
                    )
                if tool_name == "verify_task_completion":
                    assert tool_args["model_mode"] == "local"
                    payload = json.loads(tool_args["file_states"])
                    assert payload and payload[0]["path"] == "target.docx"
                    return json.dumps(
                        {
                            "completed": True,
                            "confidence": 0.96,
                            "summary": "结果符合要求，目标文档已经完成更新",
                            "remaining_steps": [],
                        },
                        ensure_ascii=False,
                    )
                raise AssertionError(f"Unexpected tool call: {tool_name}")

        registry = FakeRegistry()
        llm_call_count = {"count": 0}

        def fake_call_llm(self, provider, messages, system, tool_defs, options=None):
            llm_call_count["count"] += 1
            return {
                "content": "先执行插表，再检查结果。",
                "tool_calls": [
                    {
                        "id": "call_1",
                        "name": "insert_excel_as_docx_table",
                        "args": {
                            "source_path": "sales.xlsx",
                            "target_path": "target.docx",
                            "sheet_name": "汇总表",
                        },
                    }
                ],
            }

        monkeypatch.setattr(
            TaskAgent, "_get_provider", lambda self, options=None: object()
        )
        monkeypatch.setattr(
            TaskAgent, "_build_registry", lambda self, files=None: registry
        )
        monkeypatch.setattr(TaskAgent, "_call_llm", fake_call_llm)

        agent = TaskAgent(model_id="test-model")
        payload = "".join(
            agent.execute(
                task="将 xls 表格加入 docx，并确认结果符合要求",
                files=[],
                options={"model_mode": "local"},
            )
        )

        assert llm_call_count["count"] == 1
        assert [name for name, _ in registry.executions] == [
            "insert_excel_as_docx_table",
            "verify_task_completion",
        ]
        events = parse_sse_events(payload.encode("utf-8"))
        assert any(
            e.get("type") == "step_progress"
            and "检查当前结果是否符合任务要求" in str(e.get("detail", ""))
            for e in events
        )
        assert any(
            e.get("type") == "step_done" and "结果符合要求" in str(e.get("text", ""))
            for e in events
        )
        assert any(
            e.get("type") == "done" and "结果符合要求" in str(e.get("summary", ""))
            for e in events
        )

    def test_task_agent_executes_textual_task_tool_call_fallback(self, monkeypatch):
        from app.core.agent.task_agent import TaskAgent

        class FakeRegistry:
            def __init__(self):
                self.executions = []

            def get_definitions(self):
                return [
                    {"name": "insert_excel_as_docx_table"},
                    {"name": "verify_task_completion"},
                ]

            def execute(self, tool_name, tool_args):
                self.executions.append((tool_name, dict(tool_args)))
                if tool_name == "insert_excel_as_docx_table":
                    return json.dumps(
                        {
                            "success": True,
                            "summary": "已将工作表“销售台账”的 24 行数据写入 Word 表格",
                            "path": "target.docx",
                            "file_type": "docx",
                            "change_type": "modify",
                            "operation": tool_name,
                            "preview": "客户 | 地区 | 金额",
                        },
                        ensure_ascii=False,
                    )
                if tool_name == "verify_task_completion":
                    return json.dumps(
                        {
                            "completed": True,
                            "confidence": 0.93,
                            "summary": "目标文档已追加销售台账表格",
                            "remaining_steps": [],
                        },
                        ensure_ascii=False,
                    )
                raise AssertionError(f"Unexpected tool call: {tool_name}")

        registry = FakeRegistry()
        llm_call_count = {"count": 0}

        def fake_call_llm(self, provider, messages, system, tool_defs, options=None):
            llm_call_count["count"] += 1
            return {
                "content": (
                    "为了将《销售台账.xlsx》中的信息加入到《雷鸟访谈问题.docx》中，我将直接把销售台账的数据作为 Word 表格追加到文档末尾。\n\n"
                    "下面开始执行插入操作。\n\n"
                    "```json\n"
                    '{"name": "insert_excel_as_docx_table", "arguments": {"source_path": "sales.xlsx", "target_path": "target.docx", "table_title": "附录：销售台账数据"}}\n'
                    "```"
                ),
                "tool_calls": [],
            }

        monkeypatch.setattr(
            TaskAgent, "_get_provider", lambda self, options=None: object()
        )
        monkeypatch.setattr(
            TaskAgent, "_build_registry", lambda self, files=None: registry
        )
        monkeypatch.setattr(TaskAgent, "_call_llm", fake_call_llm)

        agent = TaskAgent(model_id="test-model")
        payload = "".join(
            agent.execute(
                task="将 xlsx 信息加入 docx，并确认已经写入目标文件",
                files=[],
                options={"model_mode": "cloud"},
            )
        )

        assert llm_call_count["count"] == 1
        assert [name for name, _ in registry.executions] == [
            "insert_excel_as_docx_table",
            "verify_task_completion",
        ]
        events = parse_sse_events(payload.encode("utf-8"))
        assert any(
            e.get("type") == "tool_call"
            and e.get("tool_name") == "insert_excel_as_docx_table"
            for e in events
        )
        assert any(
            e.get("type") == "done"
            and "目标文档已追加销售台账表格" in str(e.get("summary", ""))
            for e in events
        )
        assert not any(
            e.get("type") == "result"
            and "insert_excel_as_docx_table" in str(e.get("data", ""))
            for e in events
        )

    def test_task_agent_reinjects_stage_verification_feedback_when_incomplete(
        self, monkeypatch
    ):
        from app.core.agent.task_agent import TaskAgent

        class FakeRegistry:
            def get_definitions(self):
                return [
                    {"name": "write_docx_content"},
                    {"name": "verify_task_completion"},
                ]

            def execute(self, tool_name, tool_args):
                if tool_name == "write_docx_content":
                    return json.dumps(
                        {
                            "success": True,
                            "summary": "已写入 2 个段落到 Word 文档",
                            "path": "draft.docx",
                            "file_type": "docx",
                            "change_type": "modify",
                            "operation": tool_name,
                            "preview": "第一段\n第二段",
                        },
                        ensure_ascii=False,
                    )
                if tool_name == "verify_task_completion":
                    return json.dumps(
                        {
                            "completed": False,
                            "confidence": 0.41,
                            "summary": "当前文档还缺少结论段",
                            "remaining_steps": ["补充结论段"],
                        },
                        ensure_ascii=False,
                    )
                raise AssertionError(f"Unexpected tool call: {tool_name}")

        seen_message_batches = []
        responses = iter(
            [
                {
                    "content": "先写入主体内容。",
                    "tool_calls": [
                        {
                            "id": "call_1",
                            "name": "write_docx_content",
                            "args": {"path": "draft.docx", "paragraphs": []},
                        }
                    ],
                },
                {
                    "content": "继续补充结论段。",
                    "tool_calls": [],
                },
            ]
        )

        monkeypatch.setattr(
            TaskAgent, "_get_provider", lambda self, options=None: object()
        )
        monkeypatch.setattr(
            TaskAgent, "_build_registry", lambda self, files=None: FakeRegistry()
        )

        def fake_call_llm(self, provider, messages, system, tool_defs, options=None):
            seen_message_batches.append(messages)
            return next(responses)

        monkeypatch.setattr(TaskAgent, "_call_llm", fake_call_llm)

        agent = TaskAgent(model_id="test-model")
        _ = "".join(agent.execute(task="补全文档并保证结构完整", files=[], options={}))

        assert len(seen_message_batches) == 2
        verify_messages = [
            m
            for m in seen_message_batches[1]
            if m.get("name") == "verify_task_completion"
        ]
        assert verify_messages
        assert "缺少结论段" in verify_messages[-1]["content"]

    def test_task_agent_skips_duplicate_tool_calls_within_single_batch(
        self, monkeypatch
    ):
        from app.core.agent.task_agent import TaskAgent

        class FakeRegistry:
            def __init__(self):
                self.executions = []

            def get_definitions(self):
                return [{"name": "insert_excel_as_docx_table"}]

            def execute(self, tool_name, tool_args):
                self.executions.append((tool_name, dict(tool_args)))
                return json.dumps(
                    {
                        "success": True,
                        "summary": "已将工作表“汇总表”的 200 行数据写入 Word 表格",
                        "path": "target.docx",
                        "file_type": "docx",
                        "change_type": "modify",
                        "operation": tool_name,
                    },
                    ensure_ascii=False,
                )

        registry = FakeRegistry()
        duplicate_args = {
            "source_path": "sales.xlsx",
            "target_path": "target.docx",
            "sheet_name": "汇总表",
            "table_title": "汇总表",
        }
        responses = iter(
            [
                {
                    "content": "先把 Excel 插入 Word 表格。",
                    "tool_calls": [
                        {
                            "id": "call_1",
                            "name": "insert_excel_as_docx_table",
                            "args": duplicate_args,
                        },
                        {
                            "id": "call_2",
                            "name": "insert_excel_as_docx_table",
                            "args": dict(duplicate_args),
                        },
                    ],
                },
                {
                    "content": "目标文档已经更新完成。",
                    "tool_calls": [],
                },
            ]
        )

        monkeypatch.setattr(
            TaskAgent, "_get_provider", lambda self, options=None: object()
        )
        monkeypatch.setattr(
            TaskAgent, "_build_registry", lambda self, files=None: registry
        )
        monkeypatch.setattr(
            TaskAgent,
            "_call_llm",
            lambda self, provider, messages, system, tool_defs, options=None: next(
                responses
            ),
        )

        agent = TaskAgent(model_id="test-model")
        payload = "".join(
            agent.execute(task="将 xls 表格插入 docx", files=[], options={})
        )

        assert registry.executions == [("insert_excel_as_docx_table", duplicate_args)]
        events = parse_sse_events(payload.encode("utf-8"))
        assert any(
            e.get("type") == "thought" and "重复工具调用" in str(e.get("text", ""))
            for e in events
        )
        assert any(
            e.get("type") == "tool_result"
            and "已跳过重复工具调用" in str(e.get("result_preview", ""))
            for e in events
        )

    def test_task_agent_stops_before_repeating_identical_successful_tool_batch(
        self, monkeypatch
    ):
        from app.core.agent.task_agent import TaskAgent

        class FakeRegistry:
            def __init__(self):
                self.executions = []

            def get_definitions(self):
                return [{"name": "parse_file_to_text"}]

            def execute(self, tool_name, tool_args):
                self.executions.append((tool_name, dict(tool_args)))
                return "第一轮读取成功"

        registry = FakeRegistry()
        repeated_call = {
            "id": "call_1",
            "name": "parse_file_to_text",
            "args": {"path": "demo.txt", "max_chars": 12000},
        }
        responses = iter(
            [
                {
                    "content": "先读取当前文件。",
                    "tool_calls": [repeated_call],
                },
                {
                    "content": "继续读取当前文件。",
                    "tool_calls": [
                        {
                            "id": "call_2",
                            "name": "parse_file_to_text",
                            "args": {"path": "demo.txt", "max_chars": 12000},
                        }
                    ],
                },
            ]
        )

        monkeypatch.setattr(
            TaskAgent, "_get_provider", lambda self, options=None: object()
        )
        monkeypatch.setattr(
            TaskAgent, "_build_registry", lambda self, files=None: registry
        )
        monkeypatch.setattr(
            TaskAgent,
            "_call_llm",
            lambda self, provider, messages, system, tool_defs, options=None: next(
                responses
            ),
        )

        agent = TaskAgent(model_id="test-model")
        payload = "".join(agent.execute(task="总结当前文件", files=[], options={}))

        assert registry.executions == [
            ("parse_file_to_text", {"path": "demo.txt", "max_chars": 12000})
        ]
        events = parse_sse_events(payload.encode("utf-8"))
        assert any(
            e.get("type") == "thought"
            and "重复请求同一组工具" in str(e.get("text", ""))
            for e in events
        )
        assert any(
            e.get("type") == "done" and "检测到重复步骤" in str(e.get("summary", ""))
            for e in events
        )

    def test_task_agent_reinjects_failed_tool_feedback_for_sandbox_style_errors(
        self, monkeypatch
    ):
        from app.core.agent.task_agent import TaskAgent

        class FakeRegistry:
            def __init__(self):
                self.executions = []

            def get_definitions(self):
                return [{"name": "run_python_code"}]

            def execute(self, tool_name, tool_args):
                self.executions.append((tool_name, dict(tool_args)))
                return "[error] 执行失败：unsupported operand type(s) for +: 'float' and 'str'"

        registry = FakeRegistry()
        seen_message_batches = []
        responses = iter(
            [
                {
                    "content": "先执行 Python 代码。",
                    "tool_calls": [
                        {
                            "id": "call_1",
                            "name": "run_python_code",
                            "args": {"code": "print(1)"},
                        }
                    ],
                },
                {
                    "content": "我已经知道上一轮失败了。",
                    "tool_calls": [],
                },
            ]
        )

        monkeypatch.setattr(
            TaskAgent, "_get_provider", lambda self, options=None: object()
        )
        monkeypatch.setattr(
            TaskAgent, "_build_registry", lambda self, files=None: registry
        )

        def fake_call_llm(self, provider, messages, system, tool_defs, options=None):
            seen_message_batches.append(messages)
            return next(responses)

        monkeypatch.setattr(TaskAgent, "_call_llm", fake_call_llm)

        agent = TaskAgent(model_id="test-model")
        payload = "".join(
            agent.execute(
                task="把 xlsx 信息加入 docx", files=[], options={"model_mode": "local"}
            )
        )

        events = parse_sse_events(payload.encode("utf-8"))
        assert registry.executions == [("run_python_code", {"code": "print(1)"})]
        assert any(
            e.get("type") == "step_error"
            and "unsupported operand type" in str(e.get("error", ""))
            for e in events
        )
        assert not any(
            e.get("type") == "step_done" and e.get("step_id") == "run_python_code_1_1"
            for e in events
        )

        assert len(seen_message_batches) == 2
        corrective_prompts = [
            msg.get("content", "")
            for msg in seen_message_batches[1]
            if msg.get("role") == "user"
            and "上一轮工具调用失败" in str(msg.get("content", ""))
        ]
        assert corrective_prompts
        assert "unsupported operand type" in corrective_prompts[-1]
        assert "不要重复完全相同的工具调用、参数或代码" in corrective_prompts[-1]

    def test_task_agent_stops_before_repeating_identical_failed_tool_batch(
        self, monkeypatch
    ):
        from app.core.agent.task_agent import TaskAgent

        class FakeRegistry:
            def __init__(self):
                self.executions = []

            def get_definitions(self):
                return [{"name": "run_python_code"}]

            def execute(self, tool_name, tool_args):
                self.executions.append((tool_name, dict(tool_args)))
                return "[error] 执行失败：unsupported operand type(s) for +: 'float' and 'str'"

        registry = FakeRegistry()
        repeated_call = {
            "id": "call_1",
            "name": "run_python_code",
            "args": {"code": "print(1)"},
        }
        responses = iter(
            [
                {
                    "content": "先执行 Python 代码。",
                    "tool_calls": [repeated_call],
                },
                {
                    "content": "再试一次相同代码。",
                    "tool_calls": [
                        {
                            "id": "call_2",
                            "name": "run_python_code",
                            "args": {"code": "print(1)"},
                        }
                    ],
                },
            ]
        )

        monkeypatch.setattr(
            TaskAgent, "_get_provider", lambda self, options=None: object()
        )
        monkeypatch.setattr(
            TaskAgent, "_build_registry", lambda self, files=None: registry
        )
        monkeypatch.setattr(
            TaskAgent,
            "_call_llm",
            lambda self, provider, messages, system, tool_defs, options=None: next(
                responses
            ),
        )

        agent = TaskAgent(model_id="test-model")
        payload = "".join(
            agent.execute(
                task="把 xlsx 信息加入 docx", files=[], options={"model_mode": "local"}
            )
        )

        assert registry.executions == [("run_python_code", {"code": "print(1)"})]
        events = parse_sse_events(payload.encode("utf-8"))
        assert any(
            e.get("type") == "thought"
            and "重复提交上一轮失败的工具调用" in str(e.get("text", ""))
            for e in events
        )
        assert any(
            e.get("type") == "done"
            and "检测到重复失败步骤" in str(e.get("summary", ""))
            for e in events
        )

    def test_resolve_requested_model_id_falls_back_when_unavailable(self, monkeypatch):
        import web.app as webapp

        class _DummyManager:
            _cached_caps = {}

            def get_available_models(self):
                return [{"id": "gemini-2.5-flash"}]

            def get_model_for_task(self, task):
                return "gemini-2.5-flash"

        monkeypatch.setattr(webapp, "_model_manager", _DummyManager())

        assert (
            webapp._resolve_requested_model_id(
                "gemini-2.5-pro",
                fallback_model="gemini-3.1-pro-preview",
            )
            == "gemini-2.5-flash"
        )

    def test_resolve_requested_model_id_rejects_image_model_for_chat(self, monkeypatch):
        import web.app as webapp

        class _DummyManager:
            _cached_caps = {
                "gemini-3.1-flash-image-preview": {
                    "image_gen": True,
                    "multimodal": True,
                    "grounding": False,
                    "function_calling": False,
                    "tier": 7,
                }
            }

            def get_available_models(self):
                return [
                    {"id": "gemini-3.1-flash-image-preview"},
                    {"id": "gemini-2.5-flash"},
                ]

        monkeypatch.setattr(webapp, "_model_manager", _DummyManager())

        assert (
            webapp._resolve_requested_model_id(
                "gemini-3.1-flash-image-preview",
                fallback_model="gemini-2.5-flash",
                task_type="CHAT",
            )
            == "gemini-2.5-flash"
        )

    def test_resolve_requested_model_id_rejects_model_without_required_task_capability(
        self, monkeypatch
    ):
        import web.app as webapp

        class _DummyManager:
            _cached_caps = {
                "gemini-2.5-flash": {
                    "speed": 10,
                    "quality": 8,
                    "reasoning": 8,
                    "context": 8,
                    "multimodal": True,
                    "grounding": False,
                    "function_calling": True,
                    "image_gen": False,
                    "tier": 8,
                }
            }

            def get_available_models(self):
                return [
                    {"id": "gemini-2.5-flash"},
                    {"id": "gemini-2.5-pro"},
                ]

        monkeypatch.setattr(webapp, "_model_manager", _DummyManager())

        assert (
            webapp._resolve_requested_model_id(
                "gemini-2.5-flash",
                fallback_model="gemini-2.5-pro",
                task_type="WEB_SEARCH",
            )
            == "gemini-2.5-pro"
        )

    def test_parse_file_to_text_accepts_larger_custom_windows(self, tmp_path):
        from app.core.agent.task_tools import parse_file_to_text

        source_path = tmp_path / "long_notes.txt"
        source_path.write_text("A" * 20_000, encoding="utf-8")

        parsed = parse_file_to_text(str(source_path), max_chars=18_000)

        assert len(parsed) == 18_000
        assert len(parsed) > 12_000

    def test_task_agent_keeps_long_tool_results_in_followup_context(self, monkeypatch):
        from app.core.agent.task_agent import TaskAgent

        class FakeRegistry:
            def get_definitions(self):
                return [{"name": "parse_file_to_text"}]

            def execute(self, tool_name, tool_args):
                return "B" * 12_000

        seen_message_batches = []
        responses = iter(
            [
                {
                    "content": "先读取文件全文。",
                    "tool_calls": [
                        {
                            "id": "call_1",
                            "name": "parse_file_to_text",
                            "args": {"path": "demo.txt", "max_chars": 12000},
                        }
                    ],
                },
                {
                    "content": "已完成。",
                    "tool_calls": [],
                },
            ]
        )

        monkeypatch.setattr(
            TaskAgent, "_get_provider", lambda self, options=None: object()
        )
        monkeypatch.setattr(
            TaskAgent, "_build_registry", lambda self, files=None: FakeRegistry()
        )

        def fake_call_llm(self, provider, messages, system, tool_defs, options=None):
            seen_message_batches.append(messages)
            return next(responses)

        monkeypatch.setattr(TaskAgent, "_call_llm", fake_call_llm)

        agent = TaskAgent(model_id="test-model")
        _ = "".join(agent.execute(task="读取大文件", files=[], options={}))

        assert len(seen_message_batches) == 2
        function_messages = [
            m for m in seen_message_batches[1] if m.get("role") == "function"
        ]
        assert function_messages
        assert len(function_messages[-1]["content"]) > 4000

    def test_task_agent_run_python_code_can_open_attached_file_by_basename(
        self, tmp_path, monkeypatch
    ):
        openpyxl = pytest.importorskip("openpyxl")

        source_path = tmp_path / "销售台账.xlsx"

        workbook = openpyxl.Workbook()
        worksheet = workbook.active
        worksheet.title = "销售数据"
        worksheet.append(["姓名", "地区", "销售额"])
        worksheet.append(["张三", "华东", 120])
        workbook.save(source_path)
        workbook.close()

        from app.core.agent.task_agent import TaskAgent

        responses = iter(
            [
                {
                    "content": "先用 Python 读取附件中的 Excel。",
                    "tool_calls": [
                        {
                            "id": "call_1",
                            "name": "run_python_code",
                            "args": {
                                "code": (
                                    "import openpyxl\n"
                                    "wb = openpyxl.load_workbook('销售台账.xlsx', read_only=True, data_only=True)\n"
                                    "ws = wb.active\n"
                                    "print(ws['A2'].value)\n"
                                    "wb.close()\n"
                                )
                            },
                        }
                    ],
                },
                {
                    "content": "已成功读取附件文件。",
                    "tool_calls": [],
                },
            ]
        )

        monkeypatch.setattr(
            TaskAgent, "_get_provider", lambda self, options=None: object()
        )
        monkeypatch.setattr(
            TaskAgent,
            "_call_llm",
            lambda self, provider, messages, system, tool_defs, options=None: next(
                responses
            ),
        )

        agent = TaskAgent(model_id="test-model")
        payload = "".join(
            agent.execute(
                task="读取附件里的 Excel 并打印第一行数据",
                files=[
                    {
                        "path": str(source_path),
                        "name": source_path.name,
                        "type": "xlsx",
                    },
                ],
                options={},
            )
        )

        events = parse_sse_events(payload.encode("utf-8"))
        assert any(
            e.get("type") == "tool_result"
            and "张三" in str(e.get("result_preview", ""))
            for e in events
        )
        assert any(e.get("type") == "done" for e in events)

    def test_run_python_in_sandbox_syncs_modified_attached_file_when_cleanup_fails(
        self, tmp_path, monkeypatch
    ):
        from app.core.agent import task_tools

        source_path = tmp_path / "任务说明.txt"
        source_path.write_text("before", encoding="utf-8")

        cleanup_calls = {"count": 0}
        original_rmtree = task_tools.shutil.rmtree

        def flaky_rmtree(path, *args, **kwargs):
            cleanup_calls["count"] += 1
            if cleanup_calls["count"] == 1:
                raise PermissionError(32, "locked", path)
            return original_rmtree(path, *args, **kwargs)

        monkeypatch.setattr(task_tools.shutil, "rmtree", flaky_rmtree)

        result = task_tools.run_python_in_sandbox(
            (
                "from pathlib import Path\n"
                f"p = Path(TASK_SANDBOX_FILE_PATHS[{source_path.name!r}])\n"
                "p.write_text('after', encoding='utf-8')\n"
                f"print('KOTO_MODIFIED:' + TASK_SANDBOX_FILE_PATHS[{source_path.name!r}])\n"
            ),
            timeout=10,
            task_files=[{"path": str(source_path), "name": source_path.name}],
        )

        assert "Sandbox error:" not in result
        assert "__koto_modified__" in result
        assert str(source_path) in result
        assert source_path.read_text(encoding="utf-8") == "after"
        assert cleanup_calls["count"] == 2

    def test_run_python_in_sandbox_accepts_string_timeout(self):
        from app.core.agent import task_tools

        result = task_tools.run_python_in_sandbox("print('ok')", timeout="30")

        assert "unsupported operand type" not in result
        assert result.get("error") == ""
        assert result.get("stdout", "").strip() == "ok"

    def test_task_agent_run_python_code_syncs_modified_attached_file_and_emits_file_change(
        self, tmp_path, monkeypatch
    ):
        source_path = tmp_path / "任务说明.txt"
        source_path.write_text("before", encoding="utf-8")

        from app.core.agent.task_agent import TaskAgent

        responses = iter(
            [
                {
                    "content": "先用 Python 修改附件内容。",
                    "tool_calls": [
                        {
                            "id": "call_1",
                            "name": "run_python_code",
                            "args": {
                                "code": (
                                    "from pathlib import Path\n"
                                    f"p = Path(TASK_SANDBOX_FILE_PATHS[{source_path.name!r}])\n"
                                    "p.write_text('after', encoding='utf-8')\n"
                                    f"print('KOTO_MODIFIED:' + TASK_SANDBOX_FILE_PATHS[{source_path.name!r}])\n"
                                )
                            },
                        }
                    ],
                },
                {
                    "content": "已完成附件修改。",
                    "tool_calls": [],
                },
            ]
        )

        monkeypatch.setattr(
            TaskAgent, "_get_provider", lambda self, options=None: object()
        )
        monkeypatch.setattr(
            TaskAgent,
            "_call_llm",
            lambda self, provider, messages, system, tool_defs, options=None: next(
                responses
            ),
        )

        agent = TaskAgent(model_id="test-model")
        payload = "".join(
            agent.execute(
                task="把附件内容改成 after",
                files=[
                    {"path": str(source_path), "name": source_path.name, "type": "txt"},
                ],
                options={"model_mode": "local"},
            )
        )

        events = parse_sse_events(payload.encode("utf-8"))
        assert source_path.read_text(encoding="utf-8") == "after"
        assert any(
            e.get("type") == "file_change"
            and e.get("path") == str(source_path)
            and e.get("change_type") == "modify"
            for e in events
        )
        assert any(e.get("type") == "done" for e in events)

    def test_task_agent_run_python_code_detects_direct_source_file_modification(
        self, tmp_path, monkeypatch
    ):
        source_path = tmp_path / "任务说明.txt"
        source_path.write_text("before", encoding="utf-8")

        from app.core.agent.task_agent import TaskAgent

        responses = iter(
            [
                {
                    "content": "直接修改原始附件路径。",
                    "tool_calls": [
                        {
                            "id": "call_1",
                            "name": "run_python_code",
                            "args": {
                                "code": (
                                    "from pathlib import Path\n"
                                    f"p = Path(TASK_FILE_PATHS[{source_path.name!r}])\n"
                                    "p.write_text('after-direct', encoding='utf-8')\n"
                                    "print('done')\n"
                                )
                            },
                        }
                    ],
                },
                {
                    "content": "已完成原文件修改。",
                    "tool_calls": [],
                },
            ]
        )

        monkeypatch.setattr(
            TaskAgent, "_get_provider", lambda self, options=None: object()
        )
        monkeypatch.setattr(
            TaskAgent,
            "_call_llm",
            lambda self, provider, messages, system, tool_defs, options=None: next(
                responses
            ),
        )

        agent = TaskAgent(model_id="test-model")
        payload = "".join(
            agent.execute(
                task="把原始附件内容改成 after-direct",
                files=[
                    {"path": str(source_path), "name": source_path.name, "type": "txt"},
                ],
                options={"model_mode": "local"},
            )
        )

        events = parse_sse_events(payload.encode("utf-8"))
        assert source_path.read_text(encoding="utf-8") == "after-direct"
        assert any(
            e.get("type") == "file_change"
            and e.get("path") == str(source_path)
            and e.get("change_type") == "modify"
            for e in events
        )
        assert any(e.get("type") == "done" for e in events)

    def test_task_agent_inserts_excel_table_into_docx_and_emits_file_change(
        self, tmp_path, monkeypatch
    ):
        openpyxl = pytest.importorskip("openpyxl")
        docx_module = pytest.importorskip("docx")

        source_path = tmp_path / "销售台账.xlsx"
        target_path = tmp_path / "雷鸟访问问题.docx"

        workbook = openpyxl.Workbook()
        worksheet = workbook.active
        worksheet.title = "销售数据"
        worksheet.append(["姓名", "地区", "销售额"])
        worksheet.append(["张三", "华东", 120])
        worksheet.append(["李四", "华南", 98])
        workbook.save(source_path)
        workbook.close()

        document = docx_module.Document()
        document.add_paragraph("雷鸟访问问题说明")
        document.save(target_path)

        from app.core.agent.task_agent import TaskAgent

        responses = iter(
            [
                {
                    "content": "先读取 Excel，并把数据写成 Word 表格。",
                    "tool_calls": [
                        {
                            "id": "call_1",
                            "name": "insert_excel_as_docx_table",
                            "args": {
                                "source_path": str(source_path),
                                "target_path": str(target_path),
                                "sheet_name": "销售数据",
                                "table_title": "销售台账",
                            },
                        }
                    ],
                },
                {
                    "content": "已完成 Excel 到 Word 表格写入，并校验目标文档。",
                    "tool_calls": [],
                },
            ]
        )

        monkeypatch.setattr(
            TaskAgent, "_get_provider", lambda self, options=None: object()
        )
        monkeypatch.setattr(
            TaskAgent,
            "_call_llm",
            lambda self, provider, messages, system, tool_defs, options=None: next(
                responses
            ),
        )

        agent = TaskAgent(model_id="test-model")
        payload = "".join(
            agent.execute(
                task="将excel数据加入word，做一个新表格",
                files=[
                    {
                        "path": str(source_path),
                        "name": source_path.name,
                        "type": "xlsx",
                    },
                    {
                        "path": str(target_path),
                        "name": target_path.name,
                        "type": "docx",
                    },
                ],
                options={},
            )
        )

        events = parse_sse_events(payload.encode("utf-8"))
        assert any(e.get("type") == "file_change" for e in events)
        assert any(
            e.get("type") == "file_change"
            and str(e.get("path", "")).endswith("雷鸟访问问题.docx")
            for e in events
        )
        assert any(e.get("type") == "done" for e in events)

        updated_doc = docx_module.Document(target_path)
        assert updated_doc.tables
        first_table = updated_doc.tables[0]
        assert first_table.cell(0, 0).text == "姓名"
        assert first_table.cell(1, 0).text == "张三"
        assert first_table.cell(2, 1).text == "华南"


class TestWorkspaceAssistantTaskRemovalRegression:
    """Source-level regressions for removing file-assistant AI task/chat flows."""

    def test_workspace_assistant_removes_old_task_stream_senders(self):
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        retired_external_name = "Open" + "Claw"
        assert f"async function _waSendTo{retired_external_name}Task" not in src
        assert f"function _waBuild{retired_external_name}TaskMessage" not in src
        assert "async function _waSendToAgent" not in src
        assert "async function _waSendToChat" not in src
        assert "async function _waSendToInline" not in src
        assert "window.WA.applyAIResponse = (mode, btn) =>" not in src
        assert "window.WA.setOutputMode = () =>" not in src
        assert "function _refreshWorkflowChips() {}" not in src
        assert "async function _appendWorkflowChips() {}" not in src
        assert "async function _suggestWorkflows() {}" not in src
        assert "fetch('/api/chat/stream'" not in src
        assert "fetch('/api/agent/chat'" not in src
        assert "action: 'ai_task'" not in src

    def test_workspace_send_message_keeps_open_file_and_uses_whitebox_stream(self):
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        retired_external_name = "Open" + "Claw"
        send_start = src.index("window.WA.sendMessage = () => {")
        send_end = src.index("// ── Auto-save", send_start)
        send_block = src[send_start:send_end]
        assert "appendUserMessageWithLoading" in send_block
        assert "_waTaskDispatcher.dispatchMessage({" in send_block
        assert "/api/editor/ai/task-stream" not in send_block
        assert f"_waSendTo{retired_external_name}Task(" not in send_block
        assert "_waSendToAgent(" not in send_block
        assert "_waSendToChat(" not in send_block

    def test_workspace_send_message_builds_whitebox_payload_with_target_path_history_and_model_state(
        self,
    ):
        assistant = Path("web/static/js/workspace-assistant.js").read_text(
            encoding="utf-8"
        )
        dispatcher = Path("web/static/js/workspace-task-dispatcher.js").read_text(
            encoding="utf-8"
        )

        send_start = assistant.index("window.WA.sendMessage = () => {")
        send_end = assistant.index("// ── Auto-save", send_start)
        send_block = assistant[send_start:send_end]

        assert "pinnedSelText," in send_block
        assert "pinnedSelSource," in send_block
        assert "_waTaskDispatcher.dispatchMessage({" in send_block
        assert (
            "function buildWhiteboxTaskPayload(text, pinnedSelText, pinnedSelSource, overrides) {"
            in dispatcher
        )
        assert "current_file:" not in dispatcher
        assert "getCurrentAIContextPath" not in dispatcher
        assert (
            "target_path: targetFile ? (targetFile.path || targetFile.name || '') : '',"
            in dispatcher
        )
        assert (
            "model_mode: typeof options.getModelMode === 'function' ? options.getModelMode() : 'auto',"
            in dispatcher
        )
        assert (
            "model_id: typeof options.getSelectedCloudModelId === 'function' ? options.getSelectedCloudModelId() : '',"
            in dispatcher
        )
        assert "options: overrideOptions," in dispatcher
        assert (
            "history: typeof options.getConversationHistory === 'function'"
            in dispatcher
        )
        assert (
            "const payload = buildWhiteboxTaskPayload(context.text, context.pinnedSelText, context.pinnedSelSource, context);"
            in dispatcher
        )
        assert "return Promise.resolve(streamWhiteboxTask({" in dispatcher
        assert "payload," in dispatcher
        assert "taskTurnMetadataFromLoadingEl(loadingEl)" in dispatcher

    def test_workspace_dispatcher_marks_short_task_critiques_as_followup_context(self):
        dispatcher = Path("web/static/js/workspace-task-dispatcher.js").read_text(
            encoding="utf-8"
        )

        assert "function latestCompletedFileTaskTurn()" in dispatcher
        assert "function looksLikeDiagnosticLead(text)" in dispatcher
        assert "function looksLikeTaskCritique(text)" in dispatcher
        assert "function buildTaskFollowupContext(text)" in dispatcher
        assert "kind: 'review_last_task'" in dispatcher
        assert "followup_action:" in dispatcher
        assert "overrideOptions.followup_context = followupContext;" in dispatcher
        assert "context.previous_run_id = previousRunId;" in dispatcher
        assert "context.previous_task_mode = previousTaskMode;" in dispatcher
        assert (
            "context.previous_task_file_changes = previousTaskFileChanges;"
            in dispatcher
        )

    def test_workspace_task_cards_offer_run_bound_followup_actions(self):
        assistant = Path("web/static/js/workspace-assistant.js").read_text(
            encoding="utf-8"
        )
        task_renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )

        assert "window.WA.beginTaskResultFollowup = (details) => {" in assistant
        assert "state._pendingTaskFollowupContext = followupContext;" in assistant
        assert (
            "options: pendingTaskFollowupContext ? { followup_context: pendingTaskFollowupContext } : {},"
            in assistant
        )
        assert (
            "请把上一轮已经给出的建议直接应用到目标文件；沿用同一任务上下文继续写回，不要重新从头分析。"
            in assistant
        )
        assert (
            "previous_task_output_mode: String(payload.output_mode || '').trim(),"
            in assistant
        )
        assert (
            "previous_task_intent_strategy: String(payload.intent_strategy || '').trim(),"
            in assistant
        )
        assert "previous_task_intent_can_apply" in assistant
        assert "previous_task_intent_requires_confirmation" in assistant
        assert "function taskResultActionsHtml(card) {" in task_renderer
        assert 'data-task-followup-action="apply"' in task_renderer
        assert 'data-task-followup-action="question"' in task_renderer
        assert 'data-task-followup-action="improve"' in task_renderer
        assert "window.WA.beginTaskResultFollowup({" in task_renderer
        assert "output_mode: card.dataset.taskOutputMode || ''," in task_renderer
        assert (
            "intent_strategy: card.dataset.taskIntentStrategy || ''," in task_renderer
        )
        assert (
            "intent_can_apply: boolAttr(card.dataset.taskIntentCanApply),"
            in task_renderer
        )
        assert (
            "intent_requires_confirmation: boolAttr(card.dataset.taskIntentRequiresConfirmation),"
            in task_renderer
        )
        assert (
            "file_changes: Array.isArray(taskState.fileChanges) ? taskState.fileChanges.slice(-8) : [],"
            in task_renderer
        )
        assert "card.dataset.taskRunId" in task_renderer
        assert "previous_task_file_changes = previousTaskFileChanges" in assistant

    def test_workspace_whitebox_renderer_is_extracted(self):
        assistant = Path("web/static/js/workspace-assistant.js").read_text(
            encoding="utf-8"
        )
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )
        ai_transport = Path("web/static/js/workspace-ai-transport.js").read_text(
            encoding="utf-8"
        )
        ai_results = Path("web/static/js/workspace-ai-results.js").read_text(
            encoding="utf-8"
        )
        quick_actions = Path("web/static/js/workspace-ai-quick-actions.js").read_text(
            encoding="utf-8"
        )
        conversation = Path("web/static/js/workspace-ai-conversation.js").read_text(
            encoding="utf-8"
        )
        dispatcher = Path("web/static/js/workspace-task-dispatcher.js").read_text(
            encoding="utf-8"
        )
        standalone = Path("web/templates/workspace_assistant.html").read_text(
            encoding="utf-8"
        )
        main = Path("web/templates/index.html").read_text(encoding="utf-8")
        asset_scripts = Path("web/templates/_workspace_asset_scripts.html").read_text(
            encoding="utf-8"
        )

        assert "window.WA.streamWhiteboxTask" in renderer
        assert "fetch('/api/editor/ai/task-stream'" in renderer
        assert "window.WA.createWorkspaceAiTransport" in ai_transport
        assert "window.WA.createWorkspaceAiResultsRuntime" in ai_results
        assert "window.WA.createWorkspaceQuickActionRuntime" in quick_actions
        assert "window.WA.createWorkspaceAiConversation" in conversation
        assert "model' || value === 'ai'" in conversation
        assert "window.WA.createTaskDispatcher" in dispatcher
        assert (
            "typeof window.WA.createWorkspaceAiResultsRuntime === 'function'"
            in assistant
        )
        assert (
            "typeof window.WA.createWorkspaceAiConversation === 'function'" in assistant
        )
        assert "window.WA.hydrateAiHistory" in assistant
        assert (
            "typeof window.WA.createWorkspaceQuickActionRuntime === 'function'"
            in assistant
        )
        assert "_waQuickActionRuntime.attachDispatcher(_waTaskDispatcher);" in assistant
        assert "typeof window.WA.createTaskDispatcher === 'function'" in assistant
        assert "fetch('/api/editor/ai/task-stream'" not in assistant
        assert "{% include '_workspace_asset_scripts.html' %}" in standalone
        assert "{% include '_workspace_asset_scripts.html' %}" in main
        assert "workspace-ai-task.js" in asset_scripts
        assert "workspace-ai-transport.js" in asset_scripts
        assert "workspace-ai-results.js" in asset_scripts
        assert "workspace-ai-quick-actions.js" in asset_scripts
        assert "workspace-ai-conversation.js" in asset_scripts
        assert "workspace-task-dispatcher.js" in asset_scripts
        assert "doc-agent-ui.js" not in standalone
        assert "wa-doc-agent-phases" not in standalone
        assert "wa-inline-ai" not in main
        assert 'data-dm="inline"' not in main

    def test_workspace_dispatcher_exposes_extension_registration_points(self):
        assistant = Path("web/static/js/workspace-assistant.js").read_text(
            encoding="utf-8"
        )
        dispatcher = Path("web/static/js/workspace-task-dispatcher.js").read_text(
            encoding="utf-8"
        )
        quick_actions = Path("web/static/js/workspace-ai-quick-actions.js").read_text(
            encoding="utf-8"
        )

        assert "registerMessageRoute" in dispatcher
        assert "registerQuickActionHandler" in dispatcher
        assert "registerQuickActionKeyword" in dispatcher
        assert "registerAction(definition)" in quick_actions
        assert "window.WA.registerTaskQuickAction" in assistant
        assert "window.WA.registerTaskEntryRoute" in assistant
        assert "window.WA.registerTaskActionHandler" in assistant
        assert "window.WA.registerTaskActionKeyword" in assistant

    def test_workspace_dispatcher_records_assistant_turns_for_task_history(self):
        dispatcher = Path("web/static/js/workspace-task-dispatcher.js").read_text(
            encoding="utf-8"
        )

        assert "function appendAssistantConversationTurn(text, metadata)" in dispatcher
        assert "options.appendAssistantTurn(content" in dispatcher
        assert "options.getConversationHistory()" in dispatcher
        assert ".then((streamResult) =>" in dispatcher
        assert (
            "function finalizeWhiteboxTaskTurn(taskTurnId, loadingEl, result, fallbackStatus, skipModelContext)"
            in dispatcher
        )
        assert (
            "appendAssistantConversationTurn(assistantText, turnMetadata);"
            in dispatcher
        )

    def test_workspace_conversation_runtime_uses_in_memory_session_store_and_model_context(
        self,
    ):
        conversation = Path("web/static/js/workspace-ai-conversation.js").read_text(
            encoding="utf-8"
        )

        assert "function normalizeRole(role)" in conversation
        assert (
            "if (value === 'model' || value === 'ai') return 'assistant';"
            in conversation
        )
        assert "const sessionStore = new Map();" in conversation
        assert "function normalizedSessionId(rawSessionId)" in conversation
        assert "async function hydrate(params)" in conversation
        assert "renderHistory(sessionTurns(sessionId));" in conversation
        assert "query.set('session_id', sessionId);" not in conversation
        assert "renderHistory(history.map" not in conversation
        assert "function getHistoryForModel(limit)" in conversation
        assert "turn.status !== 'error'" in conversation

    def test_workspace_assistant_uses_runtime_scoped_ai_session(self):
        assistant = Path("web/static/js/workspace-assistant.js").read_text(
            encoding="utf-8"
        )

        assert "_linkedAiSessionId" not in assistant
        assert "function _waConversationDocId()" not in assistant
        assert "const _WA_RUNTIME_SESSION_ID = (() => {" in assistant
        assert "workspace_runtime_" in assistant
        assert "return _WA_RUNTIME_SESSION_ID;" in assistant
        assert "return 'workspace_' + (state.fileId || 'default');" not in assistant
        assert "window.WA.openAiSessionFromSidebar" not in assistant
        assert "window.WA.renameLinkedAiSession" not in assistant
        assert "window.WA.removeLinkedAiSession" not in assistant
        assert (
            "getConversationHistory: () => _waConversationRuntime && typeof _waConversationRuntime.getHistoryForModel === 'function'"
            in assistant
        )

    def test_workspace_task_cards_are_snapshotted_for_runtime_history_restore(self):
        conversation = Path("web/static/js/workspace-ai-conversation.js").read_text(
            encoding="utf-8"
        )
        dispatcher = Path("web/static/js/workspace-task-dispatcher.js").read_text(
            encoding="utf-8"
        )
        assistant = Path("web/static/js/workspace-assistant.js").read_text(
            encoding="utf-8"
        )
        task_renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )

        assert (
            "turn.task_card_snapshot && window.WA && typeof window.WA.restoreTaskRunCard === 'function'"
            in conversation
        )
        assert "task_card_snapshot:" in conversation
        assert "function beginAssistantTaskTurn(metadata) {" in conversation
        assert "function syncAssistantTaskTurn(turnId, metadata) {" in conversation
        assert "任务处理中…" in conversation
        assert "loadingEl.classList.contains('wa-task-run')" in conversation
        assert "if (loadingEl && loadingEl.isConnected) {" in conversation
        assert (
            "const taskTurn = typeof options.beginAssistantTaskTurn === 'function'"
            in dispatcher
        )
        assert "onTaskCardSnapshot: (card) => {" in dispatcher
        assert "options.syncAssistantTaskTurn(taskTurnId" in dispatcher
        assert (
            "beginAssistantTaskTurn: (metadata) => _waConversationRuntime && typeof _waConversationRuntime.beginAssistantTaskTurn === 'function'"
            in assistant
        )
        assert (
            "syncAssistantTaskTurn: (turnId, metadata) => _waConversationRuntime && typeof _waConversationRuntime.syncAssistantTaskTurn === 'function'"
            in assistant
        )
        assert (
            "window.WA.restoreTaskRunCard = function restoreTaskRunCard(snapshot) {"
            in task_renderer
        )
        assert "function isTaskCardElement(value) {" in task_renderer
        assert (
            "if (!isTaskCardElement(card) || card._waRunCardBehaviorAttached) return card;"
            in task_renderer
        )
        assert "if (!isTaskCardElement(card)) return null;" in task_renderer
        assert (
            "const card = isTaskCardElement(loadingEl) ? loadingEl : document.createElement('div');"
            in task_renderer
        )
        assert "if (typeof opts.onTaskCardSnapshot === 'function') {" in task_renderer
        assert "return attachRunCardBehavior(card);" in task_renderer

    def test_main_chat_filters_workspace_assistant_sessions(self):
        app_js = Path("web/static/js/app.js").read_text(encoding="utf-8")
        sessions_bp = Path("web/blueprints/sessions.py").read_text(encoding="utf-8")

        assert "_isWorkspaceAssistantSession" not in app_js
        assert "_maybeOpenWorkspaceAssistantSession" not in app_js
        assert "_notifyWorkspaceAssistantSessionRenamed" not in app_js
        assert "_notifyWorkspaceAssistantSessionDeleted" not in app_js
        assert (
            "def _is_workspace_assistant_session(filename: str) -> bool:" in sessions_bp
        )
        assert "if not _is_workspace_assistant_session(session)" in sessions_bp

    def test_workspace_task_renderer_drops_loaded_memory_context_step(self):
        renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )

        assert "type === 'memory.loaded'" not in renderer
        assert "已读取相关对话记忆" not in renderer
        assert 'wa-task-chip ok">记忆' not in renderer

    def test_workspace_quick_actions_remove_editor_ai_stream_fallback(self):
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        quick_actions = Path("web/static/js/workspace-ai-quick-actions.js").read_text(
            encoding="utf-8"
        )
        assert "async function _sendViaEditorActionSSE(payload)" not in src
        assert "window.WA.sendQuickAction = (action) => {" in src
        assert (
            "getConversationHistory: () => _waConversationRuntime && typeof _waConversationRuntime.getHistoryForModel === 'function'"
            in src
        )
        assert "/api/editor/ai/stream" not in quick_actions
        assert "action: editorAction" not in quick_actions
        assert (
            "history: typeof options.getConversationHistory === 'function'"
            not in quick_actions
        )
        assert "options.appendAssistantTurn(trimmed" not in quick_actions
        assert "options.getSessionId()" not in quick_actions
        assert "streamEventBlocks({" not in quick_actions
        assert "legacyEditorFallback" not in quick_actions
        assert "sendEditorAction" not in quick_actions
        assert "console.count('[WA legacy-editor-fallback]');" not in quick_actions
        assert "window.WA.quickAction =" in src

    def test_workspace_transport_accepts_pre_serialized_json_bodies(self):
        transport = Path("web/static/js/workspace-ai-transport.js").read_text(
            encoding="utf-8"
        )
        assert "const requestBody = typeof options.body === 'string'" in transport
        assert "body: requestBody," in transport

    def test_workspace_retired_inline_ai_entrypoints_are_removed(self):
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        assert "window.WA.sendInlineMessage = () =>" not in src
        assert "window.WA.inlineQuickAction = (text) =>" not in src
        assert "window.WA.handleInlineInputKeydown = (e) =>" not in src
        assert "window.WA.setAIDisplayMode = (mode) =>" not in src
        assert "wa_ai_display_mode" not in src

    def test_workspace_topic_ai_stub_remains_disabled(self):
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        assert "window.WA.extractTopics = async () =>" in src
        assert (
            "showToast('文件助手 AI 对话任务流已移除；请使用快捷功能键。', 'warn');"
            in src
        )

    def test_workspace_templates_do_not_expose_retired_inline_ai_controls(self):
        embedded_html = Path("web/templates/index.html").read_text(encoding="utf-8")
        standalone_html = Path("web/templates/workspace_assistant.html").read_text(
            encoding="utf-8"
        )

        assert not Path("web/static/js/doc-agent-ui.js").exists()
        assert not Path("web/static/css/doc-agent.css").exists()
        assert "wa-inline-ai" not in embedded_html
        assert "WA.sendInlineMessage()" not in embedded_html
        assert "WA.inlineQuickAction(" not in embedded_html
        assert 'data-dm="inline"' not in embedded_html
        assert "doc-agent.css" not in standalone_html
        assert "doc-agent-ui.js" not in standalone_html
        assert "wa-doc-agent-phases" not in standalone_html

    def test_workspace_input_autopin_requires_live_editor_selection(self):
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        assert "function _getLiveEditorSelectionForAI()" in src
        input_start = src.find("const _waInput = $('wa-user-input');")
        input_end = src.find("// ── Split.js Init", input_start)
        assert input_start != -1 and input_end != -1
        input_section = src[input_start:input_end]
        assert "const liveSelection = _getLiveEditorSelectionForAI();" in input_section
        assert "if (liveSelection) {" in input_section
        assert "_pinSelectionChip(liveSelection);" in input_section

    def test_workspace_scripts_do_not_reference_retired_ai_input_id(self):
        assistant = Path("web/static/js/workspace-assistant.js").read_text(
            encoding="utf-8"
        )
        embedded_html = Path("web/templates/index.html").read_text(encoding="utf-8")

        assert "wa-ai-input')" not in assistant
        assert 'wa-ai-input")' not in assistant
        assert "wa-ai-input')" not in embedded_html
        assert 'wa-ai-input")' not in embedded_html

    def test_workspace_pptx_save_routes_do_not_use_legacy_export_fallback(self):
        bp = Path("web/blueprints/workspace_assistant.py").read_text(encoding="utf-8")

        assert "def _export_workspace_pptx(file_id: str, data) -> bytes:" in bp
        assert "from app.core.file.file_parser import export_docx, export_xlsx" in bp
        assert ("export_" + "pptx") not in bp
        assert ("Legacy simple format " + "fallback") not in bp

    def test_retired_external_planner_names_are_not_in_runtime_code(self):
        runtime_files = [
            Path("app"),
            Path("web"),
        ]
        retired_terms = (
            "Open" + "Claw",
            "open" + "claw",
            "OPEN" + "CLAW",
            "Her" + "mes",
            "her" + "mes",
        )
        offenders = []
        ignored_parts = {"node_modules", "vendor", "__pycache__"}
        for root in runtime_files:
            for path in root.rglob("*"):
                if not path.is_file():
                    continue
                if ignored_parts.intersection(path.parts):
                    continue
                if path.suffix.lower() not in {".py", ".js", ".html", ".md"}:
                    continue
                text = path.read_text(encoding="utf-8", errors="ignore")
                for term in retired_terms:
                    if term in text:
                        offenders.append(f"{path}:{term}")
        assert offenders == []

    def test_workspace_proposal_card_filters_duplicate_rationale_text(self):
        """Proposal cards should hide rationale text when it just repeats original/proposed content."""
        assistant = Path("web/static/js/workspace-assistant.js").read_text(
            encoding="utf-8"
        )
        results = Path("web/static/js/workspace-ai-results.js").read_text(
            encoding="utf-8"
        )
        assert "function _getProposalRationaleText(proposal)" in assistant
        assert "_waAiResultsRuntime.getProposalRationaleText(proposal)" in assistant
        assert "function getProposalRationaleText(proposal)" in results
        assert "rationaleKey === originalKey || rationaleKey === proposedKey" in results

    def test_workspace_ai_review_tool_calls_route_into_docx_review_surface(self):
        assistant = Path("web/static/js/workspace-assistant.js").read_text(
            encoding="utf-8"
        )
        results = Path("web/static/js/workspace-ai-results.js").read_text(
            encoding="utf-8"
        )
        task_renderer = Path("web/static/js/workspace-ai-task.js").read_text(
            encoding="utf-8"
        )
        assert "window.WA.applyStructuredDocToolCall" in assistant
        assert "window.WA.applyStructuredReviewChangePayload" in assistant
        assert "window.WA.applyStructuredDocToolCall(proposal.tool_call" in results
        assert "window.WA.applyStructuredDocToolCall(toolCall" in results
        assert "window.WA.applyStructuredReviewChangePayload(payload" in task_renderer

    def test_workspace_proposal_buttons_stay_single_line_and_equal_width(self):
        """Proposal action buttons should share width and keep labels on one line."""
        css = Path("web/static/css/workspace.css").read_text(encoding="utf-8")
        assert ".wa-proposal-actions .wa-proposal-btn" in css
        assert "flex: 1 1 0;" in css
        assert "white-space: nowrap;" in css
        assert "min-height: 34px;" in css

    def test_agent_loop_sends_sanitized_proposal_summary(self):
        """Structured proposal summary should reuse the sanitized note, not raw clean_text."""
        src = Path("app/core/agent/agent_loop.py").read_text(encoding="utf-8")
        assert 'proposal_summary = proposals[0].get("rationale", "")' in src
        assert "yield evt_proposal(proposals, proposal_summary)" in src

    def test_workspace_assistant_docx_helpers_stay_outside_browser_ctx(self):
        src = Path("web/static/js/workspace-assistant.js").read_text(encoding="utf-8")
        show_ctx_idx = src.index("window.WA._showBrowserCtx")
        for symbol in (
            "function _cloneSerializable(",
            "function _getDocxRenderOpts(",
            "function _cacheDocxTabState(",
            "function _serializeEditorForTab(",
        ):
            symbol_idx = src.index(symbol)
            assert symbol_idx < show_ctx_idx, (
                f"{symbol} must remain top-level before window.WA._showBrowserCtx "
                "so DOCX open/render helpers stay visible outside the browser context menu handler"
            )
